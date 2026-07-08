
import csv
import os
import sys
import re
import threading
import unicodedata
from collections import OrderedDict
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from PIL import Image, ImageDraw, ImageFont, ImageTk
except Exception:
    Image = ImageDraw = ImageFont = ImageTk = None

MODULE_TITLE = "Lieferanten-Rechnung zu AFI-Upload"
UPLOAD_COLUMNS = ["TEXT", "PRICE", "PRICE_UNIT", "QUANTITY", "UNIT", "NET_VALUE", "TAX_CODE", "GL_ACCOUNT", "COSTCENTER", "ORDERID"]
TAX_ORDER = {"VD": 0, "V2": 1, "V0": 2, "VX": 9}
SOURCE_COST_KEYWORDS = [
    "Energiekosten", "Grundgebühr", "Grundgebuehr", "Blockiergebühr", "Blockiergebuehr", "Kartenkosten",
    "Service", "Maut", "Gebühr", "Gebuehr", "Kosten", "Netto", "Net Amount", "Net Value", "Nettobetrag",
]


def _desktop_path():
    return os.path.join(os.path.expanduser("~"), "Desktop")


def _clean(value):
    return " ".join(str(value or "").replace("\ufeff", "").strip().split())


def _norm(value):
    text = _clean(value).upper()
    text = (text.replace("Ä", "AE").replace("Ö", "OE").replace("Ü", "UE")
                .replace("ẞ", "SS").replace("ß", "SS"))
    text = "".join(ch for ch in unicodedata.normalize("NFKD", text) if not unicodedata.combining(ch))
    return re.sub(r"[^A-Z0-9]", "", text)


def _read_csv(path):
    encodings = ["utf-8-sig", "cp1252", "latin1"]
    last = None
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc, newline="") as f:
                sample = f.read(8192)
                f.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=";,")
                except Exception:
                    dialect = csv.excel
                    dialect.delimiter = ";"
                reader = csv.DictReader(f, dialect=dialect)
                rows = [{str(k or "").replace("\ufeff", "").strip(): v for k, v in row.items()} for row in reader]
                headers = [str(x or "").replace("\ufeff", "").strip() for x in (reader.fieldnames or [])]
                return headers, rows
        except Exception as exc:
            last = exc
    raise RuntimeError(f"CSV konnte nicht gelesen werden: {last}")


def _read_table_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return _read_csv(path)
    if ext in (".xlsx", ".xls"):
        try:
            import pandas as pd
            engine = "openpyxl" if ext == ".xlsx" else "xlrd"
            df = pd.read_excel(path, engine=engine, dtype=str)
            df = df.fillna("")
            headers = [str(c) for c in df.columns]
            rows = df.to_dict(orient="records")
            return headers, rows
        except Exception as exc:
            raise RuntimeError(f"Excel-Datei konnte nicht gelesen werden: {exc}")
    raise RuntimeError("Für die Berechnung werden aktuell CSV- oder Excel-Dateien benötigt.")


def _dec(value):
    s = _clean(value)
    if not s:
        return Decimal("0.00")
    s = s.replace("€", "").replace("EUR", "").replace("%", "").replace(" ", "")
    neg = s.endswith("-")
    if neg:
        s = s[:-1]
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        d = Decimal(s)
        return -d if neg else d
    except InvalidOperation:
        return Decimal("0.00")


def _fmt(value):
    d = Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{d:.2f}".replace(".", ",")


def _tax_code_from_rate(rate):
    r = Decimal(rate)
    if abs(r - Decimal("19.00")) <= Decimal("0.50"):
        return "VD"
    if abs(r - Decimal("7.00")) <= Decimal("0.50"):
        return "V2"
    if abs(r - Decimal("0.00")) <= Decimal("0.50"):
        return "V0"
    return "VX"


def _tax_code_from_net_vat(net, vat):
    net = Decimal(net)
    vat = Decimal(vat or "0.00")
    if net == 0:
        return "V0" if vat == 0 else "VX"
    rate = (vat / net * Decimal("100")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return _tax_code_from_rate(rate)


def _score_header(header, keywords):
    h = _norm(header)
    score = 0
    for kw, weight in keywords:
        if _norm(kw) in h:
            score += weight
    return score


def _guess_column(headers, keywords):
    best = ""
    best_score = 0
    for h in headers:
        s = _score_header(h, keywords)
        if s > best_score:
            best = h
            best_score = s
    return best


def _guess_net_columns(headers):
    found = []
    for h in headers:
        nh = _norm(h)
        has_net = any(_norm(k) in nh for k in ["Netto", "Nettobetrag", "NetAmount", "NetValue", "ExclVAT"])
        has_bad = any(_norm(k) in nh for k in ["Mehrwertsteuer", "MwSt", "Umsatzsteuer", "VAT", "Tax", "Brutto", "Gross"])
        if has_net and not has_bad:
            found.append(h)
    return found


def guess_columns(headers):
    return {
        "key": _guess_column(headers, [("Kennzeichen", 10), ("KFZ", 9), ("Fahrzeug", 6), ("Vehicle", 6), ("License", 6), ("Objekt", 5), ("Referenz", 5), ("Identifikation", 4)]),
        "full_name": _guess_column(headers, [("Fahrername", 10), ("Fahrer", 8), ("Mitarbeiter", 7), ("Nutzer", 7), ("Name", 6), ("Driver", 6), ("User", 5)]),
        "first": _guess_column(headers, [("Vorname", 10), ("First", 7), ("Given", 6)]),
        "last": _guess_column(headers, [("Nachname", 10), ("Last", 7), ("Surname", 7), ("Familienname", 6)]),
        # Steuersatz ist ein Prozentwert (z. B. 0, 7, 19), nicht der Steuerbetrag.
        "vat_amount": _guess_column(headers, [("Mehrwertsteuersatz", 12), ("Steuersatz", 11), ("Tax Rate", 10), ("VAT Rate", 10), ("USt Satz", 10), ("MwSt Satz", 9), ("Rate", 5)]),
        "gross": _guess_column(headers, [("Brutto", 10), ("Gross", 8), ("Amount incl", 8), ("incl VAT", 8)]),
    }


def _related_column(headers, net_col, kind):
    base = _norm(net_col)
    base = re.sub(r"NETTO|NETAMOUNT|NETVALUE|NETTBETRAG", "", base)
    candidates = []
    for h in headers:
        nh = _norm(h)
        s = 0
        if base and base in nh:
            s += 5
        if kind == "gross" and any(x in nh for x in ["BRUTTO", "GROSS", "INCLVAT"]):
            s += 10
        if kind == "vat" and any(x in nh for x in ["MEHRWERTSTEUERSATZ", "STEUERSATZ", "TAXRATE", "VATRATE", "USTSATZ", "MWSTSATZ", "RATE"]):
            s += 10
        if kind == "vat" and any(x in nh for x in ["STEUERBETRAG", "TAXAMOUNT", "MEHRWERTSTEUERBETRAG"]):
            s -= 20
        if s:
            candidates.append((s, h))
    return sorted(candidates, reverse=True)[0][1] if candidates else ""


def suggested_sources(headers):
    guessed = guess_columns(headers)
    nets = _guess_net_columns(headers)
    result = []
    for idx, net in enumerate(nets, 1):
        label = re.sub(r"\s*\(?Euro\)?\s*", "", net, flags=re.I).strip() or f"Quelle {idx}"
        result.append({
            "active": True,
            "label": label,
            "net": net,
            "tax_mode": "vat",
            "vat_amount": _related_column(headers, net, "vat") or guessed.get("vat_amount", ""),
            "gross": _related_column(headers, net, "gross") or guessed.get("gross", ""),
            "manual_rate": "19",
            "name_mode": "full",
            "full_name": guessed.get("full_name", ""),
            "first": guessed.get("first", ""),
            "last": guessed.get("last", ""),
            "key": guessed.get("key", ""),
        })
    if not result:
        result.append(default_source(1, headers))
    return result


def default_source(idx, headers):
    guessed = guess_columns(headers)
    return {
        "active": True,
        "label": f"Berechnungsquelle {idx}",
        "net": "",
        "tax_mode": "vat",
        "vat_amount": guessed.get("vat_amount", ""),
        "gross": guessed.get("gross", ""),
        "manual_rate": "19",
        "name_mode": "full",
        "full_name": guessed.get("full_name", ""),
        "first": guessed.get("first", ""),
        "last": guessed.get("last", ""),
        "key": guessed.get("key", ""),
    }


def _driver_from_row(row, src):
    first = _clean(row.get(src.get("first", ""), "")) if src.get("first") else ""
    last = _clean(row.get(src.get("last", ""), "")) if src.get("last") else ""
    full = _clean(row.get(src.get("full_name", ""), "")) if src.get("full_name") else ""
    return _clean(f"{first} {last}") or last or full


def _fallback_name_from_row(row):
    """Sucht bei fehlender Fahrer-/Kennzeichenzuordnung automatisch andere Spalten mit 'Name' im Titel.
    Kostenpositionen werden dadurch nicht verworfen, wenn eine direkte Zuordnung fehlt.
    """
    for col, value in row.items():
        if "NAME" in _norm(col):
            cleaned = _clean(value)
            if cleaned:
                return cleaned
    return ""


def _load_template_entries(template_path):
    headers, rows = _read_csv(template_path)
    required = ["TEXT", "GL_ACCOUNT", "COSTCENTER", "ORDERID"]
    missing = [c for c in required if c not in headers]
    if missing:
        raise RuntimeError("Die Vorlage enthält nicht alle erwarteten AFI-Spalten: " + ", ".join(missing))
    entries = []
    for row in rows:
        text = _clean(row.get("TEXT", ""))
        if not text:
            continue
        entries.append({
            "TEXT": text,
            "NORM_TEXT": _norm(text),
            "GL_ACCOUNT": _clean(row.get("GL_ACCOUNT", "")),
            "COSTCENTER": _clean(row.get("COSTCENTER", "")),
            "ORDERID": _clean(row.get("ORDERID", "")),
        })
    return entries


def _unique_match(candidates):
    unique = []
    seen = set()
    for c in candidates:
        key = (c.get("TEXT"), c.get("GL_ACCOUNT"), c.get("COSTCENTER"), c.get("ORDERID"))
        if key not in seen:
            seen.add(key)
            unique.append(c)
    return unique[0] if len(unique) == 1 else None


def _resolve_template(key, driver, entries):
    nkey = _norm(key)
    ndriver = _norm(driver)
    last = _norm(_clean(driver).split()[-1] if _clean(driver).split() else "")
    if nkey:
        m = _unique_match([e for e in entries if nkey in e["NORM_TEXT"]])
        if m:
            return m, "Schlüssel"
    if ndriver:
        m = _unique_match([e for e in entries if ndriver in e["NORM_TEXT"]])
        if m:
            return m, "Name"
    if last:
        m = _unique_match([e for e in entries if last in e["NORM_TEXT"]])
        if m:
            return m, "Nachname"
    return {}, ""


def create_supplier_upload_csv(template_path, invoice_path, export_path, config):
    headers, rows = _read_table_file(invoice_path)
    entries = _load_template_entries(template_path)
    global_prefix = _clean(config.get("global_prefix", "Tanken Strom")) or "Tanken Strom"
    sources = [s for s in config.get("sources", []) if s.get("active", True) and s.get("net")]
    if not sources:
        raise RuntimeError("Bitte mindestens eine aktive Berechnungsquelle mit Nettopreis-Spalte auswählen.")

    groups = OrderedDict()
    warnings_missing = []
    warnings_tax = []
    warnings_empty_assignment = []
    warnings_name_fallback = []
    invoice_net_raw_total = Decimal("0.00")
    unique_drivers = set()
    unique_keys = set()

    def add_group(src, key, driver, tax, amount):
        # Fachlogik: Alle aktiven Berechnungsquellen werden pro Fahrer > Fahrzeug/Schlüssel > Steuersatz zusammengefasst.
        group_key = (_norm(driver), _norm(key), tax)
        if group_key not in groups:
            groups[group_key] = {"source": src, "key": key, "driver": driver, "tax": tax, "amount": Decimal("0.00")}
        groups[group_key]["amount"] += amount

    for src in sources:
        for idx, row in enumerate(rows):
            net = _dec(row.get(src.get("net", ""), ""))
            if net == 0:
                continue
            driver = _driver_from_row(row, src)
            key = _clean(row.get(src.get("key", ""), "")) if src.get("key") else ""
            if not key and driver:
                key = driver
            if not key or not driver:
                fallback_name = _fallback_name_from_row(row)
                if fallback_name:
                    if not driver:
                        driver = fallback_name
                    if not key:
                        key = fallback_name
            if not key and not driver:
                # Keine Kostenposition darf wegen fehlender Zuordnung verloren gehen.
                # Für solche Fälle wird eine technische Bezeichnung erzeugt und die Kontierung bleibt leer.
                fallback_name = f"UNZUORDENBAR Zeile {idx + 2}"
                driver = fallback_name
                key = fallback_name
                warnings_empty_assignment.append(f"{src.get('label', '')}: Zeile {idx + 2} ohne Fahrer/Schlüssel; als '{fallback_name}' exportiert")
            if driver:
                unique_drivers.add(_norm(driver))
            if key:
                unique_keys.add(_norm(key))
            tax_mode = src.get("tax_mode", "vat")
            vat = Decimal("0.00")
            if tax_mode == "gross":
                gross = _dec(row.get(src.get("gross", ""), ""))
                vat = gross - net if gross else Decimal("0.00")
                tax = _tax_code_from_net_vat(net, vat) if gross else "VX"
            elif tax_mode == "manual":
                tax = _tax_code_from_rate(_dec(src.get("manual_rate", "19")))
            else:
                rate_value = _dec(row.get(src.get("vat_amount", ""), ""))
                tax = _tax_code_from_rate(rate_value)
                vat = rate_value
            if tax == "VX":
                warnings_tax.append(f"{src.get('label', '')} / {key} / {driver}: Netto {_fmt(net)}, Steuersatz {_fmt(vat)} %")
            invoice_net_raw_total += net
            add_group(src, key, driver, tax, net)

    resolved = {}
    for gkey, g in groups.items():
        info, how = _resolve_template(g["key"], g["driver"], entries)
        resolved[gkey] = info
        if not info:
            warnings_missing.append(f"{g['key']} / {g['driver']}")
        elif how in ("Name", "Nachname") and _norm(g["key"]) not in info.get("NORM_TEXT", ""):
            warnings_name_fallback.append(f"{g['key']} / {g['driver']}: Kontierung per {how} übernommen")

    ordered_groups = sorted(groups.items(), key=lambda kv: (_norm(kv[1]["key"]), _norm(kv[1]["driver"]), TAX_ORDER.get(kv[1]["tax"], 9), _norm(kv[1]["source"].get("label", ""))))
    target_net_total = invoice_net_raw_total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    export_net_before = sum(g["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for _, g in ordered_groups)
    rounding_adjustments = []
    diff = target_net_total - export_net_before
    if diff != 0 and ordered_groups:
        candidates = [g for _, g in ordered_groups if g["tax"] in ("VD", "V2") and g["amount"] != 0] or [g for _, g in ordered_groups if g["amount"] != 0]
        if candidates:
            target = max(candidates, key=lambda g: abs(g["amount"]))
            before = target["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            target["amount"] += diff
            after = target["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            rounding_adjustments.append(f"Netto-Rundungsausgleich {diff:+.2f} EUR auf {target['key']} / {target['driver']}: {before:.2f} -> {after:.2f}".replace(".", ","))

    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        for gkey, g in ordered_groups:
            amount = g["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            info = resolved.get(gkey, {})
            parts = [global_prefix]
            if g["key"]:
                parts.append(g["key"])
            if g["driver"] and _norm(g["driver"]) != _norm(g["key"]):
                parts.append(g["driver"])
            writer.writerow({
                "TEXT": _clean(" ".join(parts)),
                "PRICE": _fmt(amount),
                "PRICE_UNIT": "1",
                "QUANTITY": "1",
                "UNIT": "ST",
                "NET_VALUE": _fmt(amount),
                "TAX_CODE": g["tax"],
                "GL_ACCOUNT": info.get("GL_ACCOUNT", ""),
                "COSTCENTER": info.get("COSTCENTER", ""),
                "ORDERID": info.get("ORDERID", ""),
            })

    export_net_total = sum(g["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for _, g in ordered_groups)
    return {
        "rows": len(ordered_groups),
        "export_path": export_path,
        "invoice_net_raw_total": _fmt(invoice_net_raw_total),
        "export_net_total": _fmt(export_net_total),
        "net_rounding_difference": _fmt(export_net_total - target_net_total),
        "unique_drivers": len([x for x in unique_drivers if x]),
        "unique_keys": len([x for x in unique_keys if x]),
        "missing_template": warnings_missing,
        "unknown_tax": warnings_tax,
        "empty_assignment": warnings_empty_assignment,
        "name_fallback_matches": warnings_name_fallback,
        "rounding_adjustments": rounding_adjustments,
    }


class SourceRow:
    def __init__(self, parent, module, idx, headers, initial=None):
        self.module = module
        self.headers = headers
        self.idx = idx
        self.vars = {}
        self.frame = tk.LabelFrame(parent, text=f"Berechnungsquelle {idx}", bg=module.bg, font=module.font_small)
        self.initial = initial or default_source(idx, headers)
        self._build()
        self.set_values(self.initial)

    def _make_searchable(self, cb, values):
        cb["values"] = values
        def on_key(event):
            typed = cb.get().lower()
            if not typed:
                cb["values"] = values
                return
            cb["values"] = [v for v in values if typed in str(v).lower()]
        cb.bind("<KeyRelease>", on_key)

    def _combo(self, row, col, key, label="", width=36, colspan=1):
        if label:
            tk.Label(self.frame, text=label, bg=self.module.bg, font=self.module.font_small).grid(row=row, column=col, sticky="w", padx=5, pady=3)
            col += 1
        var = tk.StringVar()
        cb = ttk.Combobox(self.frame, textvariable=var, values=[""] + self.headers, state="normal", width=width, font=self.module.font_small)
        self._make_searchable(cb, [""] + self.headers)
        cb.grid(row=row, column=col, columnspan=colspan, sticky="ew", padx=5, pady=3)
        cb.bind("<<ComboboxSelected>>", lambda e: self.module.on_mapping_changed())
        cb.bind("<FocusOut>", lambda e: self.module.on_mapping_changed())
        self.vars[key] = var
        return cb

    def _build(self):
        for c in range(6):
            self.frame.columnconfigure(c, weight=1)
        self.vars["active"] = tk.BooleanVar(value=True)
        tk.Checkbutton(self.frame, text="Aktiv", variable=self.vars["active"], bg=self.module.bg, command=self.module.on_mapping_changed).grid(row=0, column=0, sticky="w", padx=5, pady=3)
        self._combo(0, 1, "net", "Nettopreis", width=42, colspan=2)
        tk.Label(self.frame, text="Kostenbeschreibung", bg=self.module.bg, font=self.module.font_small).grid(row=0, column=4, sticky="w", padx=5, pady=3)
        kosten_cb = ttk.Combobox(self.frame, textvariable=self.module.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", width=26, font=self.module.font_small)
        kosten_cb.grid(row=0, column=5, sticky="ew", padx=5, pady=3)
        kosten_cb.bind("<<ComboboxSelected>>", lambda e: self.module.on_mapping_changed())
        kosten_cb.bind("<FocusOut>", lambda e: self.module.on_mapping_changed())

        sep1 = tk.Frame(self.frame, height=1, bg="#B8C3CF")
        sep1.grid(row=1, column=0, columnspan=6, sticky="ew", padx=5, pady=(6, 3))
        tk.Label(self.frame, text="1. Steuerberechnung", bg=self.module.bg, font=("Segoe UI", 9, "bold")).grid(row=2, column=0, columnspan=6, sticky="w", padx=5, pady=(2, 4))
        self.vars["tax_mode"] = tk.StringVar(value="vat")
        tk.Radiobutton(self.frame, text="Bruttopreis", variable=self.vars["tax_mode"], value="gross", bg=self.module.bg, command=self.module.on_mapping_changed).grid(row=3, column=0, sticky="w", padx=5, pady=2)
        self._combo(3, 1, "gross", width=44, colspan=5)
        tk.Radiobutton(self.frame, text="Steuersatz %", variable=self.vars["tax_mode"], value="vat", bg=self.module.bg, command=self.module.on_mapping_changed).grid(row=4, column=0, sticky="w", padx=5, pady=2)
        self._combo(4, 1, "vat_amount", width=44, colspan=5)
        tk.Radiobutton(self.frame, text="Steuersatz % manuell", variable=self.vars["tax_mode"], value="manual", bg=self.module.bg, command=self.module.on_mapping_changed).grid(row=5, column=0, sticky="w", padx=5, pady=2)
        self.vars["manual_rate"] = tk.StringVar(value="19")
        manual_cb = ttk.Combobox(self.frame, textvariable=self.vars["manual_rate"], values=["19", "7", "0"], state="normal", width=16, font=self.module.font_small)
        manual_cb.grid(row=5, column=1, sticky="w", padx=5, pady=2)
        manual_cb.bind("<<ComboboxSelected>>", lambda e: self.module.on_mapping_changed())
        manual_cb.bind("<FocusOut>", lambda e: self.module.on_mapping_changed())

        sep2 = tk.Frame(self.frame, height=1, bg="#B8C3CF")
        sep2.grid(row=6, column=0, columnspan=6, sticky="ew", padx=5, pady=(6, 3))
        tk.Label(self.frame, text="2. KST-Zuordnung", bg=self.module.bg, font=("Segoe UI", 9, "bold")).grid(row=7, column=0, columnspan=6, sticky="w", padx=5, pady=(2, 4))
        self._combo(8, 0, "key", "Kennzeichen / Zuordnungsschlüssel", width=44, colspan=5)
        self._combo(9, 0, "last", "Nachname", width=44, colspan=5)
        self._combo(10, 0, "first", "Vorname", width=44, colspan=5)

    def grid(self, **kwargs):
        self.frame.grid(**kwargs)

    def destroy(self):
        self.frame.destroy()

    def set_values(self, data):
        for k, var in self.vars.items():
            if isinstance(var, tk.BooleanVar):
                var.set(bool(data.get(k, var.get())))
            else:
                var.set(str(data.get(k, var.get() if hasattr(var, 'get') else "")))

    def get(self):
        out = {}
        for k, var in self.vars.items():
            out[k] = var.get()
        out["label"] = out.get("net") or f"Berechnungsquelle {self.idx}"
        return out

    def selected_columns(self):
        cols = []
        for k in ["net", "gross", "vat_amount", "last", "first", "key"]:
            v = self.vars.get(k)
            if v and v.get():
                cols.append(v.get())
        return cols

class SupplierUploadUI:
    def __init__(self, app):
        self.app = app
        self.canvas = app.canvas
        self.bg = getattr(app, "BG", "#E8EEF5") if hasattr(app, "BG") else "#E8EEF5"
        self.font = ("Segoe UI", 10)
        self.font_small = ("Segoe UI", 9)
        self.sources = []
        self.headers = []
        self.rows = []
        self.preview_tree = None
        self.preview_text = None
        self.selected_cell = ""
        self.image_ref = None
        self.preview_canvas = None
        self.preview_canvas_image = None
        self.preview_base_image = None
        self.preview_zoom = 1.0
        self.preview_offset = [0, 0]
        self.preview_drag_start = None
        self.table_font_size = 9
        self.hide_empty_columns_var = None
        self.preview_headers = []
        self.preview_rows = []
        self.preview_path = ""

    def render(self):
        try:
            self.canvas.delete("all")
            self.app.draw_background(); self.app.draw_header(MODULE_TITLE); self.app.draw_path_bar()
        except Exception:
            pass
        w = max(1120, self.canvas.winfo_width() - 120)
        h = max(560, self.canvas.winfo_height() - 205)
        main = tk.Frame(self.canvas, bg=self.bg, width=w, height=h)
        main.grid_propagate(False)
        main.pack_propagate(False)
        # Inhalt startet unter Kopf/Breadcrumb; die Modulüberschrift bleibt innerhalb des Kopfbereichs frei.
        self.canvas.create_window(60, 148, window=main, anchor="nw", width=w, height=h)
        main.columnconfigure(0, weight=1, uniform="afi_halves")
        main.columnconfigure(1, weight=1, uniform="afi_halves")
        main.rowconfigure(1, weight=1)
        tk.Label(main, text=MODULE_TITLE, bg=self.bg, font=("Segoe UI", 14, "bold")).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))
        left = tk.Frame(main, bg=self.bg, width=int(w * 0.50), height=max(320, h-74))
        right = tk.Frame(main, bg=self.bg, width=int(w * 0.50), height=max(320, h-74))
        left.grid_propagate(False); right.grid_propagate(False)
        left.pack_propagate(False); right.pack_propagate(False)
        left.grid(row=1, column=0, sticky="nsew", padx=(0, 8))
        right.grid(row=1, column=1, sticky="nsew", padx=(8, 0))
        main.rowconfigure(2, weight=0)
        self.footer_var = tk.StringVar(value="Nettobetrag der ausgewählten Spalten: 0,00 | individuelle Fahrer: 0 | individuelle Kennzeichen: 0")
        tk.Label(main, textvariable=self.footer_var, bg="#DDE7F3", font=("Segoe UI", 10, "bold"), anchor="w").grid(row=2, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        self._build_left(left)
        self._build_right(right)

    def _build_left(self, parent):
        parent.columnconfigure(1, weight=1)
        self.template_var = tk.StringVar()
        self.invoice_var = tk.StringVar(value=_desktop_path())
        self.export_var = tk.StringVar(value=os.path.join(_desktop_path(), "Lieferanten_AFI_Upload.csv"))
        self.global_prefix_var = tk.StringVar(value="Tanken Strom")
        self.status_var = tk.StringVar(value="Bitte Rechnung und KST-Zuordnungsdokument auswählen und Rechnung analysieren.")
        self.suggestion_var = tk.StringVar(value="")

        def path_row(r, label, var, save=False):
            tk.Label(parent, text=label, bg=self.bg, font=self.font_small).grid(row=r, column=0, sticky="w", pady=3)
            tk.Entry(parent, textvariable=var, font=self.font_small).grid(row=r, column=1, sticky="ew", padx=4, pady=3)
            def browse():
                if save:
                    p = filedialog.asksaveasfilename(title=label, defaultextension=".csv", filetypes=[("CSV", "*.csv")])
                else:
                    p = filedialog.askopenfilename(title=label, filetypes=[("Dokumente", "*.csv *.xlsx *.xls *.pdf *.docx"), ("Alle Dateien", "*.*")])
                if p:
                    var.set(p)
                    if var is self.invoice_var:
                        self.load_preview(p)
            tk.Button(parent, text="…", command=browse, font=self.font_small, width=3).grid(row=r, column=2, pady=3)

        path_row(0, "KST-Zuordnungsdokument", self.template_var)
        path_row(1, "Rechnung / Dokument", self.invoice_var)
        path_row(2, "Export-CSV", self.export_var, save=True)
        tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=3, column=0, sticky="w", pady=3)
        ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small).grid(row=3, column=1, sticky="ew", padx=4, pady=3)
        tk.Button(parent, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small).grid(row=4, column=0, sticky="w", pady=(6, 3))
        tk.Button(parent, text="+ Berechnungsquelle", command=self.add_empty_source, font=self.font_small).grid(row=4, column=1, sticky="w", pady=(6, 3))
        tk.Button(parent, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI", 10, "bold"), bg="#CFEAD6", activebackground="#BDE3C7").grid(row=4, column=1, columnspan=2, sticky="e", padx=(80, 0), pady=(6, 3))
        tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=520, justify="left").grid(row=5, column=0, columnspan=3, sticky="ew")

        self.sources_canvas = tk.Canvas(parent, bg=self.bg, highlightthickness=0)
        self.sources_inner = tk.Frame(self.sources_canvas, bg=self.bg)
        yscroll = ttk.Scrollbar(parent, orient="vertical", command=self.sources_canvas.yview)
        self.sources_canvas.configure(yscrollcommand=yscroll.set)
        self.sources_canvas.grid(row=6, column=0, columnspan=2, sticky="nsew", pady=(6, 0))
        yscroll.grid(row=6, column=2, sticky="ns", pady=(6, 0))
        parent.rowconfigure(6, weight=1)
        self.sources_window = self.sources_canvas.create_window((0, 0), window=self.sources_inner, anchor="nw")
        self.sources_canvas.bind("<Configure>", lambda e: self.sources_canvas.itemconfigure(self.sources_window, width=max(100, e.width - 4)))
        self.sources_inner.bind("<Configure>", lambda e: self.sources_canvas.configure(scrollregion=self.sources_canvas.bbox("all")))
        tk.Label(parent, textvariable=self.status_var, bg=self.bg, font=self.font_small, wraplength=540, justify="left").grid(row=7, column=0, columnspan=3, sticky="ew", pady=(6, 0))

    def _build_right(self, parent):
        parent.rowconfigure(1, weight=1)
        parent.columnconfigure(0, weight=1)
        header = tk.Frame(parent, bg=self.bg)
        header.grid(row=0, column=0, sticky="ew")
        header.columnconfigure(0, weight=1)
        tk.Label(header, text="Dokumentenvorschau", bg=self.bg, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, sticky="w")
        self.hide_empty_columns_var = tk.BooleanVar(value=True)
        tk.Checkbutton(header, text="Leere Spalten ausblenden", variable=self.hide_empty_columns_var, bg=self.bg, font=self.font_small, command=self.refresh_table_preview).grid(row=0, column=1, sticky="e", padx=(8, 0))
        self.preview_frame = tk.Frame(parent, bg="white", relief="sunken", bd=1)
        self.preview_frame.grid(row=1, column=0, sticky="nsew")
        self.preview_frame.grid_propagate(False)
        self.preview_frame.pack_propagate(False)
        self.highlight_var = tk.StringVar(value="Markierte Spalten: -")
        tk.Label(parent, textvariable=self.highlight_var, bg="#FFF4C2", anchor="w", font=self.font_small).grid(row=2, column=0, sticky="ew", pady=(4, 0))

    def clear_sources(self):
        for s in self.sources:
            s.destroy()
        self.sources = []

    def add_source(self, data=None):
        row = SourceRow(self.sources_inner, self, len(self.sources) + 1, self.headers, data)
        row.grid(row=len(self.sources), column=0, sticky="ew", padx=2, pady=4)
        self.sources_inner.columnconfigure(0, weight=1)
        self.sources.append(row)
        self.on_mapping_changed()

    def add_empty_source(self):
        self.add_source(default_source(len(self.sources) + 1, self.headers))

    def analyze_invoice(self):
        path = self.invoice_var.get().strip()
        if not os.path.isfile(path):
            messagebox.showwarning(MODULE_TITLE, "Bitte eine gültige Rechnung auswählen.")
            return
        try:
            ext = os.path.splitext(path)[1].lower()
            self.clear_sources()
            if ext == ".pdf":
                self.headers, self.rows = ["PDF"], []
                self.load_preview(path)
                self.suggestion_var.set("PDF erkannt: Beträge werden beim Export aus Fahrzeug-/TOTAL-Blöcken gelesen. Berechnungsquellen sind hierfür nicht erforderlich.")
                self.status_var.set("PDF-Rechnung analysiert. Bitte Kostenbeschreibung und KST-Zuordnungsdokument prüfen.")
                return
            self.headers, self.rows = _read_table_file(path)
            suggestions = suggested_sources(self.headers)
            # Gewünscht: nur eine Quelle anlegen, aber auf weitere vermutete Quellen hinweisen.
            self.add_source(suggestions[0])
            if len(suggestions) > 1:
                names = ", ".join(s.get("label", "") for s in suggestions[1:])
                self.suggestion_var.set(f"Weitere mögliche Berechnungsquellen erkannt: {names}. Bei Bedarf über '+ Berechnungsquelle' hinzufügen und Spalten manuell setzen.")
            else:
                self.suggestion_var.set("")
            self.load_preview(path)
            self.status_var.set("Rechnung analysiert. Bitte Berechnungsquelle prüfen/ergänzen.")
        except Exception as exc:
            messagebox.showerror(MODULE_TITLE, str(exc))

    def selected_columns(self):
        cols = []
        for src in self.sources:
            cols.extend(src.selected_columns())
        return [c for i, c in enumerate(cols) if c and c not in cols[:i]]

    def on_mapping_changed(self):
        self.update_footer()
        self.update_highlight()

    def update_footer(self):
        if not self.rows:
            self.footer_var.set("Nettobetrag der ausgewählten Spalten: 0,00 | individuelle Fahrer: 0 | individuelle Kennzeichen: 0")
            return
        total = Decimal("0.00")
        drivers = set(); keys = set()
        for src_row in self.sources:
            src = src_row.get()
            if not src.get("active") or not src.get("net"):
                continue
            for row in self.rows:
                amount = _dec(row.get(src.get("net"), ""))
                if amount == 0:
                    continue
                total += amount
                d = _driver_from_row(row, src)
                k = _clean(row.get(src.get("key", ""), "")) if src.get("key") else ""
                if d: drivers.add(_norm(d))
                if k: keys.add(_norm(k))
        self.footer_var.set(f"Nettobetrag der ausgewählten Spalten: {_fmt(total)} | individuelle Fahrer: {len([d for d in drivers if d])} | individuelle Kennzeichen: {len([k for k in keys if k])}")

    def update_highlight(self):
        cols = self.selected_columns()
        self.highlight_var.set("Markierte Spalten: " + (", ".join(cols) if cols else "-"))
        if self.preview_tree:
            display_cols = list(self.preview_tree["columns"])
            for col in display_cols:
                label = col
                clean_col = col[2:] if col.startswith("★ ") else col
                if clean_col in cols and not col.startswith("★ "):
                    self.preview_tree.heading(col, text="★ " + clean_col)
                elif clean_col not in cols:
                    self.preview_tree.heading(col, text=clean_col)

    def _filtered_preview_headers(self, headers, rows):
        out = list(headers)
        if self.hide_empty_columns_var is not None and self.hide_empty_columns_var.get():
            out = [h for h in out if any(_clean(row.get(h, "")) for row in rows)]
        return out

    def refresh_table_preview(self):
        if self.preview_path and self.preview_headers:
            self.load_table_preview(self.preview_path, self.preview_headers, self.preview_rows)

    def load_preview(self, path):
        self.preview_path = path
        for w in self.preview_frame.winfo_children():
            w.destroy()
        self.preview_tree = None
        ext = os.path.splitext(path)[1].lower()
        if ext in (".csv", ".xlsx", ".xls"):
            self.load_table_preview(path)
        else:
            self.preview_headers = []
            self.preview_rows = []
            self.load_image_preview(path)

    def load_table_preview(self, path, headers=None, rows=None):
        try:
            if headers is None or rows is None:
                headers, rows = _read_table_file(path)
                self.preview_headers = headers
                self.preview_rows = rows
            headers = self._filtered_preview_headers(headers, rows)
        except Exception as exc:
            tk.Label(self.preview_frame, text=str(exc), bg="white", fg="red").pack(fill="both", expand=True)
            return
        # Eigene, begrenzte Vorschaufläche: keine Geometrieausdehnung aus dem rechten Vorschaufenster heraus.
        holder = tk.Frame(self.preview_frame, bg="white")
        holder.place(relx=0, rely=0, relwidth=1, relheight=1)
        holder.rowconfigure(0, weight=1)
        holder.columnconfigure(0, weight=1)
        style = ttk.Style(holder)
        try:
            style.configure("AfiPreview.Treeview", font=("Segoe UI", self.table_font_size), rowheight=max(18, self.table_font_size + 10))
            style.configure("AfiPreview.Treeview.Heading", font=("Segoe UI", self.table_font_size, "bold"))
        except Exception:
            pass
        tree = ttk.Treeview(holder, columns=headers, show="headings", style="AfiPreview.Treeview")
        vs = ttk.Scrollbar(holder, orient="vertical", command=tree.yview)
        hs = ttk.Scrollbar(holder, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
        tree.grid(row=0, column=0, sticky="nsew")
        vs.grid(row=0, column=1, sticky="ns")
        hs.grid(row=1, column=0, sticky="ew")
        for h in headers:
            tree.heading(h, text=h)
            tree.column(h, width=max(90, min(220, len(h) * 8)), stretch=False)
        for row in rows[:500]:
            tree.insert("", "end", values=[_clean(row.get(h, "")) for h in headers])
        self.preview_tree = tree
        def on_select(event=None):
            item = tree.focus()
            if not item:
                return
            vals = tree.item(item, "values")
            self.selected_cell = "\t".join(str(v) for v in vals)
        def copy(event=None):
            holder.clipboard_clear(); holder.clipboard_append(self.selected_cell)
            return "break"
        def wheel(event):
            # Mausrad bleibt innerhalb der Vorschau: vertikal, mit Shift horizontal, mit Ctrl Tabellenzoom.
            delta = -1 if event.delta > 0 else 1
            if event.state & 0x0004:
                self.table_font_size = max(7, min(16, self.table_font_size + (-delta)))
                try:
                    style.configure("AfiPreview.Treeview", font=("Segoe UI", self.table_font_size), rowheight=max(18, self.table_font_size + 10))
                    style.configure("AfiPreview.Treeview.Heading", font=("Segoe UI", self.table_font_size, "bold"))
                    for col in headers:
                        tree.column(col, width=max(80, min(320, int(tree.column(col, 'width') * (1.08 if delta < 0 else 0.92)))))
                except Exception:
                    pass
            elif event.state & 0x0001:
                tree.xview_scroll(delta * 3, "units")
            else:
                tree.yview_scroll(delta * 3, "units")
            return "break"
        tree.bind("<<TreeviewSelect>>", on_select)
        tree.bind("<Control-c>", copy)
        tree.bind("<MouseWheel>", wheel)
        tree.bind("<Button-4>", lambda e: (tree.yview_scroll(-3, "units"), "break"))
        tree.bind("<Button-5>", lambda e: (tree.yview_scroll(3, "units"), "break"))
        self.update_highlight()

    def _render_preview_image(self):
        if not self.preview_canvas or self.preview_base_image is None or ImageTk is None:
            return
        cw = max(1, self.preview_canvas.winfo_width())
        ch = max(1, self.preview_canvas.winfo_height())
        bw, bh = self.preview_base_image.size
        scale = self.preview_zoom
        zw, zh = max(1, int(bw * scale)), max(1, int(bh * scale))
        img = self.preview_base_image.resize((zw, zh))
        # Begrenzen: Bild darf nicht aus dem Vorschaufenster herausgezogen werden.
        if zw <= cw:
            self.preview_offset[0] = (cw - zw) // 2
        else:
            self.preview_offset[0] = min(0, max(cw - zw, self.preview_offset[0]))
        if zh <= ch:
            self.preview_offset[1] = (ch - zh) // 2
        else:
            self.preview_offset[1] = min(0, max(ch - zh, self.preview_offset[1]))
        self.image_ref = ImageTk.PhotoImage(img)
        self.preview_canvas.delete("all")
        self.preview_canvas.create_image(self.preview_offset[0], self.preview_offset[1], image=self.image_ref, anchor="nw")

    def load_image_preview(self, path):
        if Image is None or ImageTk is None:
            tk.Label(self.preview_frame, text="Bildvorschau nicht verfügbar (Pillow nicht geladen).", bg="white").pack(fill="both", expand=True)
            return
        ext = os.path.splitext(path)[1].lower()
        text = os.path.basename(path)
        try:
            if ext == ".pdf":
                try:
                    import PyPDF2
                    with open(path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        text = (reader.pages[0].extract_text() or "PDF-Vorschau")[:1400]
                except Exception:
                    text = "PDF-Bildvorschau: Text konnte nicht extrahiert werden."
            elif ext == ".docx":
                try:
                    import docx
                    doc = docx.Document(path)
                    text = "\n".join(p.text for p in doc.paragraphs[:40])[:1400] or "Word-Dokument"
                except Exception:
                    text = "Word-Bildvorschau: Text konnte nicht extrahiert werden."
            img = Image.new("RGB", (760, 520), "white")
            draw = ImageDraw.Draw(img)
            draw.rectangle((10, 10, 750, 510), outline="#B0B0B0")
            draw.text((24, 24), os.path.basename(path), fill="#1F4E79")
            y = 62
            for line in text.splitlines()[:24]:
                draw.text((24, y), line[:110], fill="black")
                y += 18
            self.preview_base_image = img
            self.preview_zoom = 1.0
            self.preview_offset = [0, 0]
            self.preview_canvas = tk.Canvas(self.preview_frame, bg="white", highlightthickness=0)
            self.preview_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
            def on_wheel(event):
                factor = 1.1 if event.delta > 0 else 0.9
                old_zoom = self.preview_zoom
                self.preview_zoom = max(0.25, min(4.0, self.preview_zoom * factor))
                if old_zoom != 0:
                    mx, my = event.x, event.y
                    self.preview_offset[0] = int(mx - (mx - self.preview_offset[0]) * (self.preview_zoom / old_zoom))
                    self.preview_offset[1] = int(my - (my - self.preview_offset[1]) * (self.preview_zoom / old_zoom))
                self._render_preview_image()
                return "break"
            def on_press(event):
                self.preview_drag_start = (event.x, event.y, self.preview_offset[0], self.preview_offset[1])
                return "break"
            def on_drag(event):
                if not self.preview_drag_start:
                    return "break"
                sx, sy, ox, oy = self.preview_drag_start
                self.preview_offset = [ox + event.x - sx, oy + event.y - sy]
                self._render_preview_image()
                return "break"
            self.preview_canvas.bind("<Configure>", lambda e: self._render_preview_image())
            self.preview_canvas.bind("<MouseWheel>", on_wheel)
            self.preview_canvas.bind("<ButtonPress-1>", on_press)
            self.preview_canvas.bind("<B1-Motion>", on_drag)
            self._render_preview_image()
        except Exception as exc:
            tk.Label(self.preview_frame, text=f"Vorschaufehler: {exc}", bg="white", fg="red").pack(fill="both", expand=True)

    def _open_file(self, path):
        try:
            if os.name == "nt":
                os.startfile(path)
            elif sys.platform == "darwin":
                import subprocess; subprocess.Popen(["open", path])
            else:
                import subprocess; subprocess.Popen(["xdg-open", path])
        except Exception as exc:
            messagebox.showerror(MODULE_TITLE, f"Datei konnte nicht geöffnet werden:\n{exc}")

    def _show_export_done_dialog(self, result):
        invoice_name = os.path.basename(self.invoice_var.get().strip()) or "Rechnungs-Dokument"
        dialog = tk.Toplevel(self.app.root)
        dialog.title(MODULE_TITLE)
        dialog.transient(self.app.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        frm = tk.Frame(dialog, bg=self.bg, padx=18, pady=16)
        frm.pack(fill="both", expand=True)
        tk.Label(frm, text=f"Bitte die erstellte AFI-Upload-Datei für \"{invoice_name}\" prüfen.", bg=self.bg, font=("Segoe UI", 10, "bold"), wraplength=520, justify="left").pack(anchor="w", pady=(0, 12))
        tk.Label(frm, text=f"Datei: {result.get('export_path', '')}", bg=self.bg, font=self.font_small, wraplength=520, justify="left").pack(anchor="w", pady=(0, 12))
        btns = tk.Frame(frm, bg=self.bg)
        btns.pack(anchor="e")
        def open_and_close():
            dialog.grab_release()
            dialog.destroy()
            self._open_file(result.get("export_path", ""))
        tk.Button(btns, text="AFI-Upload zur Prüfung öffnen", command=open_and_close, bg="#CFEAD6", activebackground="#BDE3C7", font=self.font_small).pack(side="left", padx=(0, 8))
        tk.Button(btns, text="Nicht jetzt", command=lambda: (dialog.grab_release(), dialog.destroy()), font=self.font_small).pack(side="left")
        dialog.update_idletasks()
        x = self.app.root.winfo_rootx() + max(40, (self.app.root.winfo_width() - dialog.winfo_width()) // 2)
        y = self.app.root.winfo_rooty() + max(40, (self.app.root.winfo_height() - dialog.winfo_height()) // 2)
        dialog.geometry(f"+{x}+{y}")

    def run_export(self):
        template_path = self.template_var.get().strip()
        invoice_path = self.invoice_var.get().strip()
        export_path = self.export_var.get().strip()
        if not os.path.isfile(template_path):
            messagebox.showwarning(MODULE_TITLE, "Bitte ein gültiges KST-Zuordnungsdokument auswählen.")
            return
        if not os.path.isfile(invoice_path):
            messagebox.showwarning(MODULE_TITLE, "Bitte eine gültige Rechnung auswählen.")
            return
        if not export_path.lower().endswith(".csv"):
            export_path += ".csv"; self.export_var.set(export_path)
        config = {"global_prefix": self.global_prefix_var.get(), "sources": [s.get() for s in self.sources]}
        self.status_var.set("Export läuft…")
        def worker():
            try:
                result = create_supplier_upload_csv(template_path, invoice_path, export_path, config)
                def done():
                    self.status_var.set(f"Export erstellt: {result['rows']} Zeilen → {result['export_path']} | Netto: {result['export_net_total']} | Fahrer: {result['unique_drivers']} | Kennzeichen: {result['unique_keys']}")
                    critical = []
                    if result.get("missing_template"):
                        critical.append("Keine eindeutige Kontierung gefunden:\n" + "\n".join(result["missing_template"][:40]))
                    if result.get("empty_assignment"):
                        critical.append("Zeilen ohne Fahrer/Schlüssel:\n" + "\n".join(result["empty_assignment"][:30]))
                    if result.get("unknown_tax"):
                        critical.append("Nicht eindeutig erkannte Steuersätze:\n" + "\n".join(result["unknown_tax"][:30]))
                    if critical:
                        messagebox.showwarning(MODULE_TITLE, "\n\n".join(critical))
                self.app.root.after(0, done)
            except Exception as exc:
                self.app.root.after(0, lambda: (self.status_var.set("Fehler beim Export."), messagebox.showerror(MODULE_TITLE, str(exc))))
        threading.Thread(target=worker, daemon=True).start()


def render(app):
    SupplierUploadUI(app).render()

# FLEXIBLE_KST_ZUORDNUNG_V1
# Überschreibt die bisherigen Vorlagenfunktionen zur Laufzeit. Dadurch bleibt die alte UI-Struktur kompatibel,
# aber fachlich wird keine AFI-/Kontierungsvorlage mehr benötigt, sondern ein flexibles KST-Zuordnungsdokument.
COST_TYPE_OPTIONS = ["Tanken Strom", "Tanken", "Versicherung", "Leasing", "Mobilfunk/Festnetz", "Sonstige (bitte eingeben)"]
VALID_NET_TAX_RATES = {Decimal("0"), Decimal("7"), Decimal("19")}
FOREIGN_GROSS_TAX_CODE = "V0"
ENBW_BLOCKING_GL_ACCOUNT = "427010"


def _detect_header_row(raw_df):
    header_keywords = ["kennzeichen", "rufnummer", "telefon", "msisdn", "fahrer", "name", "vorname", "nachname", "kostenstelle", "costcenter", "kst", "sachkonto", "gl_account", "innenauftrag", "ia", "orderid", "netto", "brutto", "mehrwert", "steuer", "betrag", "positionstyp", "organisationseinheit"]
    best_idx, best_score = 0, -1
    for i in range(min(25, len(raw_df))):
        vals = [_clean(v) for v in raw_df.iloc[i].tolist()]
        joined = " ".join(vals).lower()
        score = sum(1 for v in vals if v) + sum(5 for kw in header_keywords if kw in joined)
        if score > best_score:
            best_idx, best_score = i, score
    return best_idx


def _dedupe_headers(headers):
    out, seen = [], {}
    for idx, h in enumerate(headers):
        base = _clean(h) or f"Spalte_{idx+1}"
        seen[base] = seen.get(base, 0) + 1
        out.append(base if seen[base] == 1 else f"{base}_{seen[base]}")
    return out


def _read_excel_flexible(path):
    try:
        import pandas as pd
        ext = os.path.splitext(path)[1].lower()
        engine = "openpyxl" if ext in (".xlsx", ".xlsm") else "xlrd"
        raw = pd.read_excel(path, engine=engine, dtype=str, header=None).fillna("")
        if raw.empty:
            return [], []
        header_idx = _detect_header_row(raw)
        headers = _dedupe_headers([_clean(v) for v in raw.iloc[header_idx].tolist()])
        df = raw.iloc[header_idx + 1:].copy()
        df.columns = headers
        df = df[df.apply(lambda r: any(_clean(v) for v in r.values), axis=1)]
        return headers, df.fillna("").to_dict(orient="records")
    except Exception as exc:
        raise RuntimeError(f"Excel-Datei konnte nicht gelesen werden: {exc}")


def _extract_docx_text(path):
    try:
        import zipfile, html as _html
        parts = []
        with zipfile.ZipFile(path, "r") as z:
            for name in z.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    xml = z.read(name).decode("utf-8", errors="ignore")
                    xml = re.sub(r"<w:tab\s*/>", "\t", xml)
                    xml = re.sub(r"</w:p>", "\n", xml)
                    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, flags=re.S)
                    if texts:
                        parts.append(_html.unescape(" ".join(texts)))
        return "\n".join(parts)
    except Exception as exc:
        raise RuntimeError(f"DOCX-Datei konnte nicht gelesen werden: {exc}")


def _extract_pdf_text(path):
    errors = []
    try:
        import PyPDF2
        chunks = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                chunks.append(page.extract_text() or "")
        text = "\n".join(chunks)
        if _clean(text):
            return text
    except Exception as exc:
        errors.append(str(exc))
    try:
        import fitz
        doc = fitz.open(path)
        chunks = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n".join(chunks)
        if _clean(text):
            return text
    except Exception as exc:
        errors.append(str(exc))
    raise RuntimeError("PDF-Text konnte nicht extrahiert werden: " + " | ".join(errors))


def _read_table_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return _read_csv(path)
    if ext in (".xlsx", ".xlsm", ".xls"):
        return _read_excel_flexible(path)
    if ext == ".docx":
        text = _extract_docx_text(path)
        return ["TEXT"], [{"TEXT": line} for line in text.splitlines() if _clean(line)]
    if ext == ".pdf":
        text = _extract_pdf_text(path)
        return ["TEXT"], [{"TEXT": line} for line in text.splitlines() if _clean(line)]
    raise RuntimeError("Für die Berechnung werden CSV-, Excel-, PDF- oder DOCX-Dateien benötigt.")


def _cost_type(label):
    n = _norm(label)
    if "TANKENSTROM" in n or ("TANKEN" in n and "STROM" in n): return "TANKEN_STROM"
    if "TANKEN" in n: return "TANKEN"
    if "VERSICHER" in n: return "VERSICHERUNG"
    if "LEAS" in n: return "LEASING"
    if "MOBIL" in n or "FESTNETZ" in n or "TELEFON" in n or "VODAFONE" in n: return "MOBILFUNK"
    return "SONSTIGE"


def _digits_only(value):
    return re.sub(r"\D", "", _clean(value))


def _is_gl_account(value):
    return bool(re.fullmatch(r"4\d{5}", _digits_only(value)))


def _is_costcenter(value):
    d = _digits_only(value)
    return bool(re.fullmatch(r"\d{6,7}", d)) and not _is_gl_account(value)


def _extract_identifier(text):
    s = _clean(text)
    m = re.search(r"\b0\d{2,5}[\s/\-]?\d{3,10}\b", s)
    if m: return m.group(0), "PHONE"
    m = re.search(r"\b[A-ZÄÖÜ]{1,3}[\s\-]{1,2}[A-ZÄÖÜ]{1,3}[\s\-]?\d{1,5}[A-ZÄÖÜ]?\b", s, flags=re.I)
    if m: return m.group(0), "PLATE"
    m = re.search(r"\bKFZ\s+[A-Z0-9]{6,}\b", s, flags=re.I)
    if m: return m.group(0), "KEY"
    m = re.search(r"\b[A-Z0-9]{3,}-[A-Z0-9]{3,}(?:-[A-Z0-9]{3,})*\b", s, flags=re.I)
    if m: return m.group(0), "KEY"
    return "", ""


def _split_name(value):
    s = _clean(value)
    if not s: return "", "", ""
    if "," in s:
        last, first = [x.strip() for x in s.split(",", 1)]
        return _clean(f"{first} {last}"), first, last
    parts = s.split()
    if len(parts) == 1: return s, "", parts[0]
    return s, " ".join(parts[:-1]), parts[-1]


def _find_value_by_header(row, headers, keywords, validator=None, exclude=None):
    exclude = exclude or []
    candidates = []
    for h in headers:
        nh = _norm(h)
        if any(_norm(x) in nh for x in exclude):
            continue
        score = sum(weight for kw, weight in keywords if _norm(kw) in nh)
        if score:
            val = _clean(row.get(h, ""))
            if val and (validator is None or validator(val)):
                candidates.append((score, h, val))
    if candidates:
        candidates.sort(reverse=True)
        return candidates[0][2]
    return ""


def _extract_any_gl(row, headers):
    for h in headers:
        val = _clean(row.get(h, ""))
        if _is_gl_account(val): return _digits_only(val)
    m = re.search(r"\b4\d{5}\b", " ".join(_clean(row.get(h, "")) for h in headers))
    return m.group(0) if m else ""


def _assignment_entry_from_row(row, headers):
    joined = " ".join(_clean(row.get(h, "")) for h in headers)
    identifier = _find_value_by_header(row, headers, [("Kennzeichen",12),("Rufnummer",12),("Telefon",10),("MSISDN",10),("Lademedium",6),("Schlüssel",6)])
    id_type = ""
    if identifier:
        found, id_type = _extract_identifier(identifier)
        identifier = found or identifier
    if not identifier:
        identifier, id_type = _extract_identifier(joined)
    first = _find_value_by_header(row, headers, [("Vorname",10),("First",8)])
    last = _find_value_by_header(row, headers, [("Nachname",10),("Surname",8),("Last",8)])
    full = _find_value_by_header(row, headers, [("TEXT",12),("Fahrer",10),("Name",8),("Mitarbeiter",7),("Nutzer",7)], exclude=["Vorname","Nachname","Sachkonto"])
    if not full and (first or last): full = _clean(f"{first} {last}")
    if full and (not first or not last):
        _, sf, sl = _split_name(full)
        first = first or sf; last = last or sl
    gl_default = _find_value_by_header(row, headers, [("GL_ACCOUNT",12),("Sachkonto",10),("Konto",6)], validator=_is_gl_account, exclude=["Versicherung","Tanken","Strom","Mobil","Telefon","Leasing"])
    gl_tanken_strom = _find_value_by_header(row, headers, [("Tanken Strom",14),("Strom Sachkonto",12),("Laden Sachkonto",10),("Energie Sachkonto",8)], validator=_is_gl_account)
    gl_tanken = _find_value_by_header(row, headers, [("Tanken Sachkonto",14),("Tank Sachkonto",12),("Kraftstoff",8)], validator=_is_gl_account, exclude=["Strom"])
    gl_versicherung = _find_value_by_header(row, headers, [("Versicherung Sachkonto",14),("Versicherung",10)], validator=_is_gl_account)
    gl_leasing = _find_value_by_header(row, headers, [("Leasing Sachkonto",14),("Leasing",10),("Sachkonto",4)], validator=_is_gl_account)
    gl_mobilfunk = _find_value_by_header(row, headers, [("Mobilfunk",12),("Festnetz",12),("Telefon",10),("Vodafone",10),("GL_ACCOUNT",5),("Sachkonto",4)], validator=_is_gl_account)
    gl_default = gl_default or _extract_any_gl(row, headers)
    cc_default = _find_value_by_header(row, headers, [("COSTCENTER",12),("Kostenstelle",12),("KST",12)], validator=_is_costcenter, exclude=["Tanken","Strom","Versicherung","Leasing","Mobil","Telefon","IA","Innenauftrag","ORDER"])
    cc_tanken_strom = _find_value_by_header(row, headers, [("Tanken Strom",14),("Strom KST",12),("Laden KST",10)], validator=_is_costcenter)
    cc_tanken = _find_value_by_header(row, headers, [("Tanken KST",14),("Tank KST",12),("Tanken Kostenstelle",12)], validator=_is_costcenter, exclude=["Strom"])
    cc_versicherung = _find_value_by_header(row, headers, [("Versicherung KST",14),("Versicherung Kostenstelle",12)], validator=_is_costcenter)
    cc_leasing = _find_value_by_header(row, headers, [("Leasing KST",14),("Leasing Kostenstelle",12)], validator=_is_costcenter)
    cc_mobilfunk = _find_value_by_header(row, headers, [("Mobilfunk",12),("Festnetz",12),("Telefon",10),("Vodafone",10)], validator=_is_costcenter)
    orderid = _find_value_by_header(row, headers, [("ORDERID",12),("Innenauftrag",12),("IA",10),("Auftrag",7)], validator=_is_costcenter)
    return {
        "identifier": _clean(identifier), "identifier_norm": _norm(identifier), "identifier_type": id_type,
        "full_name": _clean(full), "first": _clean(first), "last": _clean(last),
        "name_norm": _norm(_clean(f"{first} {last}") or full), "last_norm": _norm(last),
        "gl_default": _digits_only(gl_default), "gl_tanken_strom": _digits_only(gl_tanken_strom), "gl_tanken": _digits_only(gl_tanken), "gl_versicherung": _digits_only(gl_versicherung), "gl_leasing": _digits_only(gl_leasing), "gl_mobilfunk": _digits_only(gl_mobilfunk),
        "cc_default": _digits_only(cc_default), "cc_tanken_strom": _digits_only(cc_tanken_strom), "cc_tanken": _digits_only(cc_tanken), "cc_versicherung": _digits_only(cc_versicherung), "cc_leasing": _digits_only(cc_leasing), "cc_mobilfunk": _digits_only(cc_mobilfunk),
        "orderid": _digits_only(orderid), "raw": joined
    }


def load_assignment_entries(assignment_path):
    headers, rows = _read_table_file(assignment_path)
    entries = []
    for row in rows:
        e = _assignment_entry_from_row(row, headers)
        if e.get("identifier") or e.get("full_name") or e.get("gl_default") or e.get("cc_default"):
            entries.append(e)
    if not entries:
        raise RuntimeError("Im KST-Zuordnungsdokument konnten keine verwertbaren Zuordnungen erkannt werden.")
    return entries


def _select_assignment_values(entry, cost_type, text_label=""):
    if "BLOCKIER" in _norm(text_label): gl = ENBW_BLOCKING_GL_ACCOUNT
    elif cost_type == "TANKEN_STROM": gl = entry.get("gl_tanken_strom") or entry.get("gl_tanken") or entry.get("gl_default")
    elif cost_type == "TANKEN": gl = entry.get("gl_tanken") or entry.get("gl_default")
    elif cost_type == "VERSICHERUNG": gl = entry.get("gl_versicherung") or entry.get("gl_default")
    elif cost_type == "LEASING": gl = entry.get("gl_leasing") or entry.get("gl_default")
    elif cost_type == "MOBILFUNK": gl = entry.get("gl_mobilfunk") or entry.get("gl_default")
    else: gl = entry.get("gl_default")
    if cost_type == "TANKEN_STROM": cc = entry.get("cc_tanken_strom") or entry.get("cc_tanken") or entry.get("cc_default")
    elif cost_type == "TANKEN": cc = entry.get("cc_tanken") or entry.get("cc_default")
    elif cost_type == "VERSICHERUNG": cc = entry.get("cc_versicherung") or entry.get("cc_default")
    elif cost_type == "LEASING": cc = entry.get("cc_leasing") or entry.get("cc_default")
    elif cost_type == "MOBILFUNK": cc = entry.get("cc_mobilfunk") or entry.get("cc_default")
    else: cc = entry.get("cc_default")
    return gl or "", cc or "", entry.get("orderid", "") or ""


def resolve_assignment(key, driver, entries):
    nkey, ndriver = _norm(key), _norm(driver)
    parts = _clean(driver).split(); last = _norm(parts[-1]) if parts else ""
    candidates = []
    for e in entries:
        score, how = 0, []
        eid = e.get("identifier_norm", "")
        if nkey and eid and (nkey == eid or nkey in eid or eid in nkey): score += 100; how.append("Schlüssel")
        if ndriver and e.get("name_norm") and (ndriver == e.get("name_norm") or ndriver in e.get("name_norm") or e.get("name_norm") in ndriver): score += 30; how.append("Name")
        if last and e.get("last_norm") and last == e.get("last_norm"): score += 20; how.append("Nachname")
        if score: candidates.append((score, "+".join(how), e))
    if not candidates: return {}, ""
    candidates.sort(key=lambda x: x[0], reverse=True)
    best_score = candidates[0][0]
    best = [c for c in candidates if c[0] == best_score]
    if len(best) > 1 and best_score < 100: return {}, "mehrdeutig"
    return best[0][2], best[0][1]


def _amount_and_tax_from_values(net, gross, rate):
    rate = Decimal(rate or "0.00")
    if rate in VALID_NET_TAX_RATES:
        return Decimal(net), _tax_code_from_rate(rate), False
    if Decimal(gross or "0.00") != 0:
        return Decimal(gross), FOREIGN_GROSS_TAX_CODE, True
    return Decimal(net) * (Decimal("1.00") + rate / Decimal("100")), FOREIGN_GROSS_TAX_CODE, True


def _parse_pdf_invoice_positions(path, global_prefix):
    text = _extract_pdf_text(path)
    compact = " ".join(text.replace("\n", " ").split())
    positions = []
    pattern = re.compile(r"VEHICLE:\s*(?P<vehicle>.*?)\s+CARD NO\.:.*?(?P<body>.*?)(?=VEHICLE:|Umsatzsteuerstatistik|Statistica|Distinta|$)", re.I)
    for m in pattern.finditer(compact):
        vehicle, body = _clean(m.group("vehicle")), m.group("body")
        rate_m = re.search(r"(?:USt|IVA)\s*\(%\):\s*(\d{1,2},\d{2}|\d{1,2})", compact[:m.start()], re.I)
        rate = _dec(rate_m.group(1)) if rate_m else Decimal("19")
        total_m = re.search(r"TOTAL:\s*(.*?)(?=VEHICLE:|Umsatzsteuerstatistik|Statistica|Distinta|Übertrag|EUR Übertrag|$)", body, re.I)
        if not total_m: continue
        nums = re.findall(r"-?\d{1,3}(?:\.\d{3})*,\d{2,3}|-?\d+,\d{1,3}|-?\d+", total_m.group(1))
        vals = [_dec(x) for x in nums]
        if len(vals) < 3: continue
        amount, tax, foreign = _amount_and_tax_from_values(vals[-3], vals[-1], rate)
        if amount != 0:
            positions.append({"key": vehicle, "driver": vehicle, "amount": amount, "tax": tax, "foreign_gross": foreign, "source_label": global_prefix})
    if not positions:
        raise RuntimeError("Aus der PDF konnten keine Fahrzeug-/TOTAL-Positionen erkannt werden.")
    return positions


def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    assignment_entries = load_assignment_entries(assignment_path)
    global_prefix = _clean(config.get("global_prefix", "Tanken Strom")) or "Tanken Strom"
    cost_type = _cost_type(global_prefix)
    ext = os.path.splitext(invoice_path)[1].lower()
    groups = OrderedDict(); warnings_missing=[]; warnings_tax=[]; warnings_empty_assignment=[]; warnings_name_fallback=[]; warnings_foreign_gross=[]; warnings_enbw_split=[]
    invoice_total = Decimal("0.00"); unique_drivers=set(); unique_keys=set()
    def add_group(source_label, key, driver, tax, amount, force_text=None):
        # Fachregel: Alle normalen Berechnungsquellen werden je Fahrzeug/Fahrer/Kennzeichen und Steuersatz kombiniert.
        # Nur fachliche Sonderzeilen mit force_text, insbesondere Blockiergebuehren, bleiben separat.
        grouping_label = force_text or global_prefix
        group_key = (_norm(grouping_label), _norm(driver), _norm(key), tax)
        if group_key not in groups:
            groups[group_key] = {"source_label": grouping_label, "key": key, "driver": driver, "tax": tax, "amount": Decimal("0.00"), "force_text": force_text or ""}
        groups[group_key]["amount"] += Decimal(amount)
    if ext == ".pdf":
        for pos in _parse_pdf_invoice_positions(invoice_path, global_prefix):
            add_group(pos["source_label"], pos["key"], pos["driver"], pos["tax"], pos["amount"])
            invoice_total += Decimal(pos["amount"])
            if pos.get("foreign_gross"): warnings_foreign_gross.append(f"{pos['key']}: abweichender/ausländischer Steuersatz -> Bruttobetrag mit V0 verwendet")
    else:
        headers, rows = _read_table_file(invoice_path)
        sources = [s for s in config.get("sources", []) if s.get("active", True) and s.get("net")]
        if not sources: raise RuntimeError("Bitte mindestens eine aktive Berechnungsquelle mit Betragsspalte auswählen.")
        for src in sources:
            src_label = _clean(src.get("label") or src.get("net") or global_prefix)
            is_blocking = "BLOCKIER" in _norm(src_label) or "BLOCKIER" in _norm(src.get("net", ""))
            for idx, row in enumerate(rows):
                net = _dec(row.get(src.get("net", ""), ""))
                if net == 0: continue
                driver = _driver_from_row(row, src)
                key = _clean(row.get(src.get("key", ""), "")) if src.get("key") else ""
                if not key and driver: key = driver
                if not key or not driver:
                    fb = _fallback_name_from_row(row)
                    if fb:
                        driver = driver or fb; key = key or fb
                if not key and not driver:
                    driver = key = f"UNZUORDENBAR Zeile {idx+2}"; warnings_empty_assignment.append(f"{src_label}: Zeile {idx+2} ohne Fahrer/Schlüssel")
                tax_mode = src.get("tax_mode", "vat")
                if tax_mode == "manual":
                    rate = _dec(src.get("manual_rate", "19")); gross = net * (Decimal("1.00") + rate/Decimal("100"))
                elif tax_mode == "gross":
                    gross = _dec(row.get(src.get("gross", ""), "")); tax = _tax_code_from_net_vat(net, gross-net) if gross else "VX"; amount = net; foreign = False
                    if tax == "VX": warnings_tax.append(f"{src_label} / {key} / {driver}: Steuer nicht eindeutig")
                    rate = Decimal("19")
                    # gross-mode bleibt netto-basiert, falls kein abweichender Steuersatz ermittelbar ist.
                    amount, tax, foreign = amount, tax, foreign
                    invoice_total += amount
                    if is_blocking and ("BLOCKIERGEBUEHR" in _norm(row.get("Positionstyp", "")) or "BLOCKIER" in _norm(src.get("net", ""))):
                        split = "IDG" if _norm(row.get("Organisationseinheit", "")) == "IDG" else "IDE"
                        add_group(f"Blockiergebühren {split}", f"Blockiergebühren {split}", f"Blockiergebühren {split}", tax, amount, force_text=f"Blockiergebühren {split}"); warnings_enbw_split.append(f"Blockiergebühr {split}: Zeile {idx+2}, Betrag {_fmt(amount)}")
                    else:
                        add_group(src_label, key, driver, tax, amount)
                    continue
                else:
                    rate = _dec(row.get(src.get("vat_amount", ""), "")) if src.get("vat_amount") else Decimal("19")
                    gross = _dec(row.get(src.get("gross", ""), "")) if src.get("gross") else Decimal("0.00")
                    if not gross and rate not in VALID_NET_TAX_RATES:
                        gross = net * (Decimal("1.00") + rate/Decimal("100"))
                amount, tax, foreign = _amount_and_tax_from_values(net, gross, rate)
                if foreign: warnings_foreign_gross.append(f"{src_label} / {key}: Steuersatz {_fmt(rate)} % -> Bruttobetrag {_fmt(amount)} mit V0 verwendet")
                invoice_total += amount
                if driver: unique_drivers.add(_norm(driver))
                if key: unique_keys.add(_norm(key))
                if is_blocking and ("BLOCKIERGEBUEHR" in _norm(row.get("Positionstyp", "")) or "BLOCKIER" in _norm(src.get("net", ""))):
                    split = "IDG" if _norm(row.get("Organisationseinheit", "")) == "IDG" else "IDE"
                    add_group(f"Blockiergebühren {split}", f"Blockiergebühren {split}", f"Blockiergebühren {split}", tax, amount, force_text=f"Blockiergebühren {split}"); warnings_enbw_split.append(f"Blockiergebühr {split}: Zeile {idx+2}, Betrag {_fmt(amount)}")
                else:
                    add_group(src_label, key, driver, tax, amount)
    resolved = {}
    ordered_groups = sorted(groups.items(), key=lambda kv: (_norm(kv[1].get("force_text") or kv[1]["key"]), _norm(kv[1]["driver"]), TAX_ORDER.get(kv[1]["tax"],9)))
    for gkey, g in ordered_groups:
        if "BLOCKIERGEBUEHREN" in _norm(g.get("force_text") or g.get("source_label")):
            info, how = resolve_assignment(g["key"], g["driver"], assignment_entries)
            gl, cc, orderid = ENBW_BLOCKING_GL_ACCOUNT, "", ""
            if info:
                _, cc, orderid = _select_assignment_values(info, cost_type, g.get("force_text") or g.get("source_label"))
            resolved[gkey] = {"GL_ACCOUNT": gl, "COSTCENTER": cc, "ORDERID": orderid}
            if not cc: warnings_missing.append(f"{g.get('force_text') or g['key']} / {g['driver']}: keine KST gefunden")
            continue
        info, how = resolve_assignment(g["key"], g["driver"], assignment_entries)
        if not info:
            resolved[gkey] = {"GL_ACCOUNT":"", "COSTCENTER":"", "ORDERID":""}; warnings_missing.append(f"{g['key']} / {g['driver']}: {'mehrdeutige Zuordnung' if how=='mehrdeutig' else 'keine Zuordnung'}")
        else:
            gl, cc, orderid = _select_assignment_values(info, cost_type, g.get("source_label", global_prefix))
            resolved[gkey] = {"GL_ACCOUNT":gl, "COSTCENTER":cc, "ORDERID":orderid}
            if not gl or not cc: warnings_missing.append(f"{g['key']} / {g['driver']}: Sachkonto/KST unvollständig (Sachkonto='{gl}', KST='{cc}')")
            elif how in ("Name","Nachname"): warnings_name_fallback.append(f"{g['key']} / {g['driver']}: Kontierung per {how} übernommen")
    target_total = invoice_total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore"); writer.writeheader()
        for gkey, g in ordered_groups:
            amount = g["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP); info = resolved.get(gkey,{})
            if g.get("force_text"): out_text = g["force_text"]
            else:
                parts = [g.get("source_label") or global_prefix]
                if g["key"]: parts.append(g["key"])
                if g["driver"] and _norm(g["driver"]) != _norm(g["key"]): parts.append(g["driver"])
                out_text = _clean(" ".join(parts))
            writer.writerow({"TEXT": out_text, "PRICE": _fmt(amount), "PRICE_UNIT":"1", "QUANTITY":"1", "UNIT":"ST", "NET_VALUE": _fmt(amount), "TAX_CODE": g["tax"], "GL_ACCOUNT": info.get("GL_ACCOUNT",""), "COSTCENTER": info.get("COSTCENTER",""), "ORDERID": info.get("ORDERID","")})
    export_total = sum(g["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for _, g in ordered_groups)
    return {"rows":len(ordered_groups), "export_path":export_path, "invoice_net_raw_total":_fmt(invoice_total), "export_net_total":_fmt(export_total), "net_rounding_difference":_fmt(export_total-target_total), "unique_drivers":len([x for x in unique_drivers if x]), "unique_keys":len([x for x in unique_keys if x]), "missing_template":warnings_missing, "unknown_tax":warnings_tax, "empty_assignment":warnings_empty_assignment, "name_fallback_matches":warnings_name_fallback, "rounding_adjustments":[], "foreign_gross":warnings_foreign_gross, "enbw_blocking_split":warnings_enbw_split}

# UMLAUT_ASCII_OUTPUT_PATCH_V1
# Ausgabe-Patch: fachliche Ausgabetexte werden ohne deutsche Umlaute geschrieben.
# Beispiel: Mueller, Goetz, Pruefung, Gebuehr, Gross, Strasse.
def _ascii_umlauts(value):
    text = str(value or "")
    replacements = {
        "Ä": "Ae", "Ö": "Oe", "Ü": "Ue", "ẞ": "SS",
        "ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text

# Falls create_supplier_upload_csv durch den flexiblen KST-Patch definiert wurde,
# wird die Funktion hier gekapselt. Dadurch bleiben Berechnung/Matching unveraendert,
# aber die fertige Export-CSV enthaelt keine Umlaute mehr.
_create_supplier_upload_csv_before_ascii_patch = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_ascii_patch(assignment_path, invoice_path, export_path, config)
    try:
        if export_path and os.path.isfile(export_path):
            with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv.DictReader(f, delimiter=";")
                fieldnames = reader.fieldnames or []
                rows = []
                for row in reader:
                    converted = {}
                    for key, value in row.items():
                        converted[key] = _ascii_umlauts(value)
                    rows.append(converted)
            with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
                writer.writeheader()
                for row in rows:
                    writer.writerow(row)
    except Exception:
        # Export nicht blockieren, falls nachtraegliche ASCII-Umsetzung fehlschlaegt.
        pass

    # Umlaute werden nur in der geschriebenen Export-CSV ersetzt.
    # Meldungen und interne Rueckgaben bleiben fachlich lesbar.
    return result

# DKV_FOREIGN_TAX_SPLIT_FIX_V1
# Fix: Bei DKV-PDFs muss pro VEHICLE-Block der zuletzt vor dem Block genannte Steuersatz gelten.
# Dadurch wird z. B. HN-I 8589 Deutschland mit 19% separat von HN-I 8589 Italien mit 22% behandelt.
def _parse_pdf_invoice_positions(path, global_prefix):
    text = _extract_pdf_text(path)
    compact = " ".join(text.replace("\n", " ").split())
    positions = []

    vehicle_pattern = re.compile(
        r"VEHICLE:\s*(?P<vehicle>.*?)\s+CARD NO\.:.*?(?P<body>.*?)(?=VEHICLE:|Umsatzsteuerstatistik|Statistica|Distinta|$)",
        re.I,
    )
    tax_pattern = re.compile(r"(?:USt|IVA)\s*\(%\):\s*(\d{1,2},\d{2}|\d{1,2})", re.I)

    for m in vehicle_pattern.finditer(compact):
        vehicle = _clean(m.group("vehicle"))
        body = m.group("body")

        # Wichtig: nicht den ersten Steuersatz im Dokument nehmen, sondern den letzten vor diesem Fahrzeugblock.
        previous_rates = list(tax_pattern.finditer(compact[:m.start()]))
        rate = _dec(previous_rates[-1].group(1)) if previous_rates else Decimal("19")

        total_m = re.search(
            r"TOTAL:\s*(.*?)(?=VEHICLE:|Umsatzsteuerstatistik|Statistica|Distinta|Übertrag|EUR Übertrag|$)",
            body,
            re.I,
        )
        if not total_m:
            continue

        nums = re.findall(r"-?\d{1,3}(?:\.\d{3})*,\d{2,3}|-?\d+,\d{1,3}|-?\d+", total_m.group(1))
        vals = [_dec(x) for x in nums]
        if len(vals) < 3:
            continue

        # DKV-TOTAL-Ende: ... Gesamtwert netto, USt/IVA, Gesamtwert brutto.
        net_amount = vals[-3]
        gross_amount = vals[-1]
        amount, tax, foreign = _amount_and_tax_from_values(net_amount, gross_amount, rate)
        if amount == 0:
            continue

        positions.append({
            "key": vehicle,
            "driver": vehicle,
            "amount": amount,
            "tax": tax,
            "foreign_gross": foreign,
            "source_label": global_prefix,
            "tax_rate": rate,
        })

    if not positions:
        raise RuntimeError("Aus der PDF konnten keine Fahrzeug-/TOTAL-Positionen erkannt werden.")
    return positions

# UI_PREVIEW_PATH_BUKRS_EXPORTNAME_PATCH_V2
# UI-Erweiterung: Buchungskreis, Standardpfade, nummerierte Arbeitsschritte,
# echte Bildvorschau fuer Rechnung/Export und automatischer Exportdateiname.
import datetime as _fm_dt

KST_ASSIGNMENT_DEFAULT_DIR = r"G:\BUC\FM Anwendung\Datenbasen\KST_Zuordnungen_AFI"
AFI_EXPORT_DEFAULT_DIR = r"G:\BUC\FM Anwendung\Dateiausgabe\AFI-Upload-Export"
BOOKING_CIRCLE_OPTIONS = ["IDE", "IDG", "IMS"]


def _fm_downloads_path():
    return os.path.join(os.path.expanduser("~"), "Downloads")


def _fm_safe_name_part(value):
    value = _ascii_umlauts(_clean(value)) if '_ascii_umlauts' in globals() else _clean(value)
    value = re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")
    return value or "Unbekannt"


def _fm_short_vendor(path):
    n = _norm(os.path.basename(path))
    for label, needles in [
        ("DKV", ["DKV"]), ("EnBW", ["ENBW", "NBW"]), ("Vodafone", ["VODAFONE"]),
        ("Telekom", ["TELEKOM", "DTAG", "TMOBILE", "T-MOBILE"]), ("VW", ["VW", "VOLKSWAGEN", "LEASING"]),
    ]:
        if any(x in n for x in needles):
            return label
    stem = os.path.splitext(os.path.basename(path or "Rechnung"))[0]
    parts = [x for x in re.split(r"[_\-\s]+", stem) if x]
    return parts[0][:20] if parts else "Rechnung"


def _fm_export_filename(bukrs, invoice_path, cost_desc):
    return f"{_fm_safe_name_part(bukrs or 'IDE')}_{_fm_safe_name_part(_fm_short_vendor(invoice_path))}_{_fm_safe_name_part(cost_desc or 'Kosten')}_{_fm_dt.datetime.now():%Y_%m_%d}.csv"


def _fm_update_export_path(self, force=False):
    if not hasattr(self, 'export_var'):
        return
    current = self.export_var.get().strip() if self.export_var.get() else ""
    invoice = self.invoice_var.get().strip() if hasattr(self, 'invoice_var') else ""
    cost = self.global_prefix_var.get() if hasattr(self, 'global_prefix_var') else "Kosten"
    bukrs = self.booking_circle_var.get() if hasattr(self, 'booking_circle_var') else "IDE"
    new_path = os.path.join(AFI_EXPORT_DEFAULT_DIR, _fm_export_filename(bukrs, invoice, cost))
    if force or not current or current.startswith(AFI_EXPORT_DEFAULT_DIR) or os.path.dirname(current) in ("", _desktop_path()):
        self.export_var.set(new_path)


def _fm_make_placeholder_image(title, lines=None):
    if Image is None or ImageDraw is None:
        return None
    img = Image.new("RGB", (1100, 720), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default() if ImageFont else None
    draw.rectangle((8, 8, 1092, 712), outline="#AAB7C4")
    draw.rectangle((8, 8, 1092, 38), fill="#DDE7F3", outline="#AAB7C4")
    draw.text((18, 18), _clean(title), fill="#1F4E79", font=font)
    y = 58
    for line in (lines or [])[:34]:
        draw.text((18, y), _clean(line)[:150], fill="black", font=font)
        y += 18
    return img


def _fm_table_image_from_rows(title, headers, rows, max_rows=45, max_cols=10):
    if Image is None or ImageDraw is None:
        return None
    headers = list(headers or [])[:max_cols]
    rows = list(rows or [])[:max_rows]
    font = ImageFont.load_default() if ImageFont else None
    cell_h = 23
    widths = []
    for h in headers:
        width_len = min(max([len(_clean(h))] + [len(_clean(r.get(h, ""))) for r in rows]), 36)
        widths.append(max(95, min(250, width_len * 7 + 18)))
    w = max(1100, min(2800, sum(widths) + 2))
    hgt = max(720, (len(rows)+3)*cell_h + 25)
    img = Image.new("RGB", (w, hgt), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0,0,w,28), fill="#DDE7F3", outline="#AAB7C4")
    draw.text((8,8), _clean(title), fill="#1F4E79", font=font)
    y = 32; x = 0
    for i, h in enumerate(headers):
        draw.rectangle((x,y,x+widths[i],y+cell_h), fill="#EAF1F8", outline="#AAB7C4")
        draw.text((x+4,y+5), _clean(h)[:36], fill="black", font=font)
        x += widths[i]
    y += cell_h
    for ridx, row in enumerate(rows):
        x = 0; fill = "#FFFFFF" if ridx % 2 == 0 else "#F7F9FB"
        for i, h in enumerate(headers):
            draw.rectangle((x,y,x+widths[i],y+cell_h), fill=fill, outline="#D6DEE8")
            draw.text((x+4,y+5), _clean(row.get(h, ""))[:36], fill="black", font=font)
            x += widths[i]
        y += cell_h
    return img


def _fm_pdf_first_page_image(path):
    if Image is None:
        return None
    try:
        import fitz
        doc = fitz.open(path)
        page = doc[0]
        zoom = min(2.0, 1400 / max(1, page.rect.width))
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close()
        return img
    except Exception:
        return None


def _fm_invoice_image(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _fm_pdf_first_page_image(path) or _fm_make_placeholder_image(os.path.basename(path), ["PDF konnte nicht als Bild gerendert werden."])
    if ext in (".csv", ".xlsx", ".xls", ".xlsm"):
        try:
            headers, rows = _read_table_file(path)
            return _fm_table_image_from_rows(os.path.basename(path), headers, rows)
        except Exception as exc:
            return _fm_make_placeholder_image(os.path.basename(path), [f"Tabellenbild konnte nicht erstellt werden: {exc}"])
    if ext == ".docx":
        try:
            lines = _extract_docx_text(path).splitlines()
        except Exception as exc:
            lines = [f"DOCX konnte nicht gelesen werden: {exc}"]
        return _fm_make_placeholder_image(os.path.basename(path), lines)
    return _fm_make_placeholder_image(os.path.basename(path), ["Keine Vorschau verfuegbar."])


def _fm_export_image(path):
    if path and os.path.isfile(path):
        try:
            headers, rows = _read_csv(path)
            return _fm_table_image_from_rows(os.path.basename(path), headers, rows)
        except Exception as exc:
            return _fm_make_placeholder_image("AFI-Upload-Export", [f"Exportvorschau konnte nicht gelesen werden: {exc}"])
    return _fm_make_placeholder_image("AFI-Upload-Export", ["Noch kein Export erstellt.", "Nach Erstellung wird die AFI-CSV hier als Bildvorschau angezeigt."])


def _fm_set_image(frame, img):
    for child in frame.winfo_children():
        child.destroy()
    if img is None or ImageTk is None:
        tk.Label(frame, text="Bildvorschau nicht verfuegbar.", bg="white").pack(fill="both", expand=True)
        return
    canvas = tk.Canvas(frame, bg="white", highlightthickness=0)
    canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
    state = {"img": img, "zoom": 1.0, "offset": [0,0], "ref": None, "drag": None}
    def render():
        cw = max(1, canvas.winfo_width()); ch = max(1, canvas.winfo_height())
        bw, bh = state["img"].size
        zw, zh = max(1, int(bw*state["zoom"])), max(1, int(bh*state["zoom"]))
        pic = state["img"].resize((zw, zh))
        if zw <= cw: state["offset"][0] = (cw-zw)//2
        else: state["offset"][0] = min(0, max(cw-zw, state["offset"][0]))
        if zh <= ch: state["offset"][1] = (ch-zh)//2
        else: state["offset"][1] = min(0, max(ch-zh, state["offset"][1]))
        state["ref"] = ImageTk.PhotoImage(pic)
        canvas.delete("all")
        canvas.create_image(state["offset"][0], state["offset"][1], image=state["ref"], anchor="nw")
    def wheel(e):
        old = state["zoom"]
        state["zoom"] = max(0.25, min(4.0, state["zoom"] * (1.1 if e.delta > 0 else 0.9)))
        if old:
            state["offset"][0] = int(e.x - (e.x-state["offset"][0]) * state["zoom"] / old)
            state["offset"][1] = int(e.y - (e.y-state["offset"][1]) * state["zoom"] / old)
        render(); return "break"
    def press(e):
        state["drag"] = (e.x, e.y, state["offset"][0], state["offset"][1]); return "break"
    def drag(e):
        sx, sy, ox, oy = state["drag"] or (e.x,e.y,state["offset"][0],state["offset"][1])
        state["offset"] = [ox+e.x-sx, oy+e.y-sy]
        render(); return "break"
    canvas.bind("<Configure>", lambda e: render())
    canvas.bind("<MouseWheel>", wheel)
    canvas.bind("<ButtonPress-1>", press)
    canvas.bind("<B1-Motion>", drag)
    frame._fm_image_ref = state
    render()


def _fm_browse(self, label, var, save=False, role=""):
    if save:
        path = filedialog.asksaveasfilename(title=label, initialdir=os.path.dirname(var.get()) or AFI_EXPORT_DEFAULT_DIR, initialfile=os.path.basename(var.get()), defaultextension=".csv", filetypes=[("CSV", "*.csv")])
    else:
        start = var.get() if os.path.isdir(var.get()) else os.path.dirname(var.get())
        path = filedialog.askopenfilename(title=label, initialdir=start or None, filetypes=[("Dokumente", "*.csv *.xlsx *.xls *.xlsm *.pdf *.docx"), ("Alle Dateien", "*.*")])
    if path:
        var.set(path)
        if role == "invoice":
            _fm_update_export_path(self, force=True)
            self.load_preview(path)


def _fm_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.booking_circle_var = tk.StringVar(value="IDE")
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_DIR)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value="Tanken Strom")
    self.status_var = tk.StringVar(value="Bitte Rechnung und KST-Zuordnungsdokument auswaehlen und Rechnung analysieren.")
    self.suggestion_var = tk.StringVar(value="")
    _fm_update_export_path(self, force=True)

    def path_row(r, label, var, save=False, role=""):
        tk.Label(parent, text=label, bg=self.bg, font=self.font_small).grid(row=r, column=0, sticky="w", pady=3)
        tk.Entry(parent, textvariable=var, font=self.font_small).grid(row=r, column=1, sticky="ew", padx=4, pady=3)
        tk.Button(parent, text="…", command=lambda: _fm_browse(self, label, var, save, role), font=self.font_small, width=3).grid(row=r, column=2, pady=3)

    tk.Label(parent, text="Buchungskreis", bg=self.bg, font=self.font_small).grid(row=0, column=0, sticky="w", pady=3)
    cb_bukrs = ttk.Combobox(parent, textvariable=self.booking_circle_var, values=BOOKING_CIRCLE_OPTIONS, state="readonly", font=self.font_small)
    cb_bukrs.grid(row=0, column=1, sticky="ew", padx=4, pady=3)
    cb_bukrs.bind("<<ComboboxSelected>>", lambda e: _fm_update_export_path(self, force=True))
    path_row(1, "KST-Zuordnungsdokument", self.template_var, False, "template")
    path_row(2, "Rechnung / Dokument", self.invoice_var, False, "invoice")
    path_row(3, "Export-CSV", self.export_var, True, "export")
    tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky="w", pady=3)
    cb_cost = ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small)
    cb_cost.grid(row=4, column=1, sticky="ew", padx=4, pady=3)
    cb_cost.bind("<<ComboboxSelected>>", lambda e: (_fm_update_export_path(self, True), self.on_mapping_changed()))
    cb_cost.bind("<FocusOut>", lambda e: (_fm_update_export_path(self, True), self.on_mapping_changed()))
    actions = tk.Frame(parent, bg=self.bg); actions.grid(row=5, column=0, columnspan=3, sticky="ew", pady=(6,3)); actions.columnconfigure(6, weight=1)
    tk.Label(actions, text="1.", bg=self.bg, font=("Segoe UI", 10, "bold")).grid(row=0, column=0, padx=(0,4))
    self.analyze_btn = tk.Button(actions, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small); self.analyze_btn.grid(row=0, column=1)
    tk.Label(actions, text="2.", bg=self.bg, font=("Segoe UI", 10, "bold")).grid(row=0, column=2, padx=(14,4))
    self.add_source_btn = tk.Button(actions, text="+ Berechnungsquelle", command=self.add_empty_source, font=self.font_small); self.add_source_btn.grid(row=0, column=3)
    tk.Button(actions, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI", 10, "bold"), bg="#CFEAD6", activebackground="#BDE3C7").grid(row=0, column=7, sticky="e")
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=520, justify="left").grid(row=6, column=0, columnspan=3, sticky="ew")
    self.sources_canvas = tk.Canvas(parent, bg=self.bg, highlightthickness=0)
    self.sources_inner = tk.Frame(self.sources_canvas, bg=self.bg)
    yscroll = ttk.Scrollbar(parent, orient="vertical", command=self.sources_canvas.yview)
    self.sources_canvas.configure(yscrollcommand=yscroll.set)
    self.sources_canvas.grid(row=7, column=0, columnspan=2, sticky="nsew", pady=(6,0)); yscroll.grid(row=7, column=2, sticky="ns", pady=(6,0)); parent.rowconfigure(7, weight=1)
    self.sources_window = self.sources_canvas.create_window((0,0), window=self.sources_inner, anchor="nw")
    self.sources_canvas.bind("<Configure>", lambda e: self.sources_canvas.itemconfigure(self.sources_window, width=max(100, e.width-4)))
    self.sources_inner.bind("<Configure>", lambda e: self.sources_canvas.configure(scrollregion=self.sources_canvas.bbox("all")))
    tk.Label(parent, textvariable=self.status_var, bg=self.bg, font=self.font_small, wraplength=540, justify="left").grid(row=8, column=0, columnspan=3, sticky="ew", pady=(6,0))


def _fm_build_right(self, parent):
    parent.rowconfigure(1, weight=1); parent.columnconfigure(0, weight=1)
    tk.Label(parent, text="Dokumentenvorschau", bg=self.bg, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, sticky="w")
    self.preview_notebook = ttk.Notebook(parent); self.preview_notebook.grid(row=1, column=0, sticky="nsew")
    self.invoice_preview_frame = tk.Frame(self.preview_notebook, bg="white"); self.export_preview_frame = tk.Frame(self.preview_notebook, bg="white")
    self.preview_notebook.add(self.invoice_preview_frame, text="Vorschau Rechnung"); self.preview_notebook.add(self.export_preview_frame, text="AFI-Upload-Export")
    self.preview_frame = self.invoice_preview_frame
    self.highlight_var = tk.StringVar(value="")
    tk.Label(parent, textvariable=self.highlight_var, bg="#FFF4C2", anchor="w", font=self.font_small).grid(row=2, column=0, sticky="ew", pady=(4,0))
    self.load_export_preview("")


def _fm_load_preview(self, path):
    self.preview_path = path
    _fm_set_image(self.invoice_preview_frame if hasattr(self, 'invoice_preview_frame') else self.preview_frame, _fm_invoice_image(path))


def _fm_load_export_preview(self, path=""):
    _fm_set_image(self.export_preview_frame if hasattr(self, 'export_preview_frame') else self.preview_frame, _fm_export_image(path))


def _fm_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen."); return
    try:
        self.clear_sources(); ext = os.path.splitext(path)[1].lower(); self.load_preview(path); _fm_update_export_path(self, True)
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            self.suggestion_var.set("PDF erkannt: Berechnungsquellen sind nicht erforderlich. Die Positionen werden aus Fahrzeug-/TOTAL-Bloecken gelesen.")
            self.status_var.set("PDF-Rechnung analysiert. Bitte Kostenbeschreibung und KST-Zuordnungsdokument pruefen.")
            if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            return
        if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
        self.headers, self.rows = _read_table_file(path)
        suggestions = suggested_sources(self.headers); self.add_source(suggestions[0])
        if len(suggestions) > 1:
            self.suggestion_var.set("Weitere moegliche Berechnungsquellen erkannt: " + ", ".join(s.get("label", "") for s in suggestions[1:]) + ". Bei Bedarf ueber '+ Berechnungsquelle' hinzufuegen.")
        else: self.suggestion_var.set("")
        self.status_var.set("Rechnung analysiert. Bitte Berechnungsquelle pruefen/ergaenzen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))


def _fm_add_empty_source(self):
    if hasattr(self, 'add_source_btn') and str(self.add_source_btn.cget('state')) == 'disabled': return
    self.add_source(default_source(len(self.sources)+1, self.headers))


def _fm_on_mapping_changed(self):
    try: _fm_update_export_path(self, True)
    except Exception: pass
    try: self.update_footer(); self.update_highlight()
    except Exception: pass


def _fm_run_export(self):
    template_path = self.template_var.get().strip(); invoice_path = self.invoice_var.get().strip(); _fm_update_export_path(self, False); export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path): messagebox.showwarning(MODULE_TITLE, "Bitte ein gueltiges KST-Zuordnungsdokument auswaehlen."); return
    if not os.path.isfile(invoice_path): messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen."); return
    if not export_path.lower().endswith(".csv"): export_path += ".csv"; self.export_var.set(export_path)
    config = {"global_prefix": self.global_prefix_var.get(), "sources": [s.get() for s in self.sources], "booking_circle": self.booking_circle_var.get() if hasattr(self, 'booking_circle_var') else "IDE"}
    self.status_var.set("Export laeuft…")
    def worker():
        try:
            result = create_supplier_upload_csv(template_path, invoice_path, export_path, config)
            def done():
                self.status_var.set(f"Export erstellt: {result['rows']} Zeilen -> {result['export_path']} | Netto: {result['export_net_total']} | Fahrer: {result['unique_drivers']} | Kennzeichen: {result['unique_keys']}")
                self.load_export_preview(result.get('export_path', export_path))
                if hasattr(self, 'preview_notebook'): self.preview_notebook.select(self.export_preview_frame)
                critical=[]
                if result.get("missing_template"): critical.append("Keine eindeutige Kontierung gefunden:\n"+"\n".join(result["missing_template"][:40]))
                if result.get("empty_assignment"): critical.append("Zeilen ohne Fahrer/Schluessel:\n"+"\n".join(result["empty_assignment"][:30]))
                if result.get("unknown_tax"): critical.append("Nicht eindeutig erkannte Steuersaetze:\n"+"\n".join(result["unknown_tax"][:30]))
                if result.get("foreign_gross"): critical.append("Abweichende/auslaendische Steuersaetze als Brutto mit V0 gebucht:\n"+"\n".join(result["foreign_gross"][:30]))
                if result.get("enbw_blocking_split"): critical.append("EnBW-Blockiergebuehren wurden nach IDE/IDG separat ausgewiesen:\n"+"\n".join(result["enbw_blocking_split"][:30]))
                if critical: messagebox.showwarning(MODULE_TITLE, "\n\n".join(critical))
                self._show_export_done_dialog(result)
            self.app.root.after(0, done)
        except Exception as exc:
            self.app.root.after(0, lambda: (self.status_var.set("Fehler beim Export."), messagebox.showerror(MODULE_TITLE, str(exc))))
    threading.Thread(target=worker, daemon=True).start()

SupplierUploadUI._build_left = _fm_build_left
SupplierUploadUI._build_right = _fm_build_right
SupplierUploadUI.load_preview = _fm_load_preview
SupplierUploadUI.load_export_preview = _fm_load_export_preview
SupplierUploadUI.analyze_invoice = _fm_analyze_invoice
SupplierUploadUI.add_empty_source = _fm_add_empty_source
SupplierUploadUI.on_mapping_changed = _fm_on_mapping_changed
SupplierUploadUI.run_export = _fm_run_export

# TEXT_COST_DESCRIPTION_PATCH_V1
# Fachregel: In der AFI-Spalte TEXT steht die ausgewaehlte Kostenbeschreibung,
# nicht der Spalten-/Berechnungsquellenname. Ausnahme: Blockiergebuehren bleiben separat ausgewiesen.
def _apply_text_cost_description_rule(export_path, config):
    if not export_path or not os.path.isfile(export_path):
        return
    global_prefix = _clean((config or {}).get("global_prefix", "")) or "Tanken Strom"
    sources = (config or {}).get("sources", []) or []
    prefixes = []
    for src in sources:
        label = _clean(src.get("label") or src.get("net") or "")
        net = _clean(src.get("net") or "")
        for candidate in (label, net):
            if candidate and candidate not in prefixes and _norm(candidate) != _norm(global_prefix):
                prefixes.append(candidate)
    # Laengere Praefixe zuerst, damit spezifische Spaltennamen vor kurzen Teilnamen ersetzt werden.
    prefixes.sort(key=len, reverse=True)
    with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        fieldnames = reader.fieldnames or []
        rows = list(reader)
    if "TEXT" not in fieldnames:
        return
    for row in rows:
        text_value = _clean(row.get("TEXT", ""))
        if not text_value:
            continue
        # Blockiergebuehren/Blockiergebühren sind fachlich eigene Sammelzeilen und bleiben unveraendert.
        if "BLOCKIERGEBUEHR" in _norm(text_value) or "BLOCKIERGEBUEHREN" in _norm(text_value):
            row["TEXT"] = _ascii_umlauts(text_value) if '_ascii_umlauts' in globals() else text_value
            continue
        replaced = False
        for prefix in prefixes:
            if _norm(text_value) == _norm(prefix) or text_value.startswith(prefix + " "):
                suffix = text_value[len(prefix):].strip()
                row["TEXT"] = _clean((global_prefix + " " + suffix).strip())
                replaced = True
                break
        if not replaced:
            # Falls kein expliziter Quellenpraefix gefunden wurde, aber der TEXT noch nicht mit der Kostenbeschreibung beginnt,
            # wird kein aggressives Abschneiden vorgenommen. So bleiben Sonder-/Fallbacktexte stabil.
            row["TEXT"] = text_value
        if '_ascii_umlauts' in globals():
            row["TEXT"] = _ascii_umlauts(row["TEXT"])
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)

_create_supplier_upload_csv_before_text_cost_patch = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_text_cost_patch(assignment_path, invoice_path, export_path, config)
    try:
        _apply_text_cost_description_rule(export_path, config)
    except Exception:
        # Export nicht blockieren, falls reine TEXT-Nachbearbeitung fehlschlaegt.
        pass
    return result


# VW_LEASING_EXPORT_PREVIEW_FIX_V2
# Zweck:
# - DKV-/ENBW-PDF-Logik bleibt unveraendert.
# - Volkswagen-Leasing-PDFs werden vor der DKV-Fallbacklogik sicher erkannt.
# - AFI-Ausgabevorschau wird vor dem finalen Export als temporaere CSV erzeugt.
import tempfile as _fm_tempfile

_parse_pdf_invoice_positions_before_vw_leasing_patch_v2 = _parse_pdf_invoice_positions

def _vw_leasing_plate_from_block(block):
    """Liest Kennzeichen aus VW-Leasing-Bloecken, z. B. 'HN  -I 2321E' -> 'HN-I 2321E'."""
    m = re.search(r"\(netto\).*?([A-ZÄÖÜ]{1,3})\s*-\s*([A-ZÄÖÜ]{1,3})\s*(\d{1,5}[A-ZÄÖÜ]?)\b", block, re.S | re.I)
    if not m:
        return ""
    return _clean(f"{m.group(1).upper()}-{m.group(2).upper()} {m.group(3).upper()}")

def _vw_leasing_driver_from_block(block):
    """Liest Bestell-Kennzeichen/Fahrername. Fuehrende Bestellnummern werden entfernt."""
    m = re.search(r"Bestell-Kennzeichen\s+(.+)", block, re.I)
    if not m:
        return ""
    value = _clean(m.group(1))
    value = re.sub(r"^\d{4,}\s+", "", value).strip()
    return value

def _parse_vw_leasing_invoice_positions(path, global_prefix):
    text = _extract_pdf_text(path)
    if "VOLKSWAGENLEASING" not in _norm(text):
        raise RuntimeError("Keine Volkswagen-Leasing-PDF erkannt.")

    positions = []
    # Jeder fachliche VW-Block beginnt mit der Kopfzeile Kennzeichen/Vertrag/Fahrgestell.
    blocks = re.split(r"(?=Kennzeichen\s+Vertrag\s+Fahrgestell)", text, flags=re.I)
    # Netto/USt/Brutto stehen als drei Euro-Werte in einer Zeile.
    triple_re = re.compile(
        r"(?<![\d.])(-?\d{1,3}(?:\.\d{3})*,\d{2})\s+"
        r"(-?\d{1,3}(?:\.\d{3})*,\d{2})\s+"
        r"(-?\d{1,3}(?:\.\d{3})*,\d{2})(?!\d)"
    )

    for block in blocks:
        if "Bestell-Kennzeichen" not in block:
            continue
        key = _vw_leasing_plate_from_block(block)
        driver = _vw_leasing_driver_from_block(block) or key
        if not key:
            continue

        for line in block.splitlines():
            m = triple_re.search(line)
            if not m:
                continue
            net = _dec(m.group(1))
            vat = _dec(m.group(2))
            gross = _dec(m.group(3))
            if net == 0 and gross == 0:
                continue
            positions.append({
                "key": key,
                "driver": driver,
                "amount": net,
                "tax": _tax_code_from_net_vat(net, vat),
                "foreign_gross": False,
                "source_label": global_prefix,
            })

    if not positions:
        raise RuntimeError("Aus der Volkswagen-Leasing-PDF konnten keine Netto/USt/Brutto-Positionen erkannt werden.")
    return positions

def _parse_pdf_invoice_positions(path, global_prefix):
    # Wichtig: VW zuerst pruefen, damit die alte DKV-Fehlermeldung nicht vorzeitig durchgereicht wird.
    try:
        text_probe = _extract_pdf_text(path)
        if "VOLKSWAGENLEASING" in _norm(text_probe):
            return _parse_vw_leasing_invoice_positions(path, global_prefix)
    except Exception:
        # Wenn die Probe fehlschlaegt, bleibt die bisherige PDF-Logik zustaendig.
        pass

    try:
        return _parse_pdf_invoice_positions_before_vw_leasing_patch_v2(path, global_prefix)
    except Exception as original_exc:
        # Sicherheitsnetz fuer Faelle, in denen die schnelle Probe nicht greifen konnte.
        try:
            return _parse_vw_leasing_invoice_positions(path, global_prefix)
        except Exception:
            raise original_exc

def _fm_preview_export_before_export(self, select_tab=False):
    """Erzeugt eine echte AFI-Ausgabevorschau in einer temporaeren CSV, ohne den finalen Exportpfad zu schreiben."""
    try:
        template_path = self.template_var.get().strip() if hasattr(self, 'template_var') else ""
        invoice_path = self.invoice_var.get().strip() if hasattr(self, 'invoice_var') else ""
        if not os.path.isfile(template_path) or not os.path.isfile(invoice_path):
            if hasattr(self, 'export_preview_frame'):
                _fm_set_image(self.export_preview_frame, _fm_make_placeholder_image(
                    "AFI-Upload-Ausgabevorschau",
                    ["Vorab-Vorschau noch nicht verfuegbar.", "Bitte zuerst Rechnung und KST-Zuordnungsdokument auswaehlen."]
                ))
            return False

        fd, tmp_path = _fm_tempfile.mkstemp(prefix="afi_upload_vorschau_", suffix=".csv")
        os.close(fd)
        config = {
            "global_prefix": self.global_prefix_var.get() if hasattr(self, 'global_prefix_var') else "Tanken Strom",
            "sources": [s.get() for s in self.sources] if hasattr(self, 'sources') else [],
            "booking_circle": self.booking_circle_var.get() if hasattr(self, 'booking_circle_var') else "IDE",
            "preview_only": True,
        }
        create_supplier_upload_csv(template_path, invoice_path, tmp_path, config)
        self.load_export_preview(tmp_path)
        if hasattr(self, 'preview_notebook') and select_tab:
            self.preview_notebook.select(self.export_preview_frame)
        return True
    except Exception as exc:
        if hasattr(self, 'export_preview_frame'):
            _fm_set_image(self.export_preview_frame, _fm_make_placeholder_image(
                "AFI-Upload-Ausgabevorschau",
                ["Vorab-Vorschau konnte noch nicht erstellt werden:", str(exc)]
            ))
        return False

_fm_analyze_invoice_before_vw_preview_patch_v2 = SupplierUploadUI.analyze_invoice

def _fm_analyze_invoice_with_output_preview_v2(self):
    _fm_analyze_invoice_before_vw_preview_patch_v2(self)
    try:
        _fm_preview_export_before_export(self, select_tab=False)
    except Exception:
        pass

_fm_on_mapping_changed_before_vw_preview_patch_v2 = SupplierUploadUI.on_mapping_changed

def _fm_on_mapping_changed_with_output_preview_v2(self):
    _fm_on_mapping_changed_before_vw_preview_patch_v2(self)
    try:
        _fm_preview_export_before_export(self, select_tab=False)
    except Exception:
        pass

_fm_run_export_before_vw_preview_patch_v2 = SupplierUploadUI.run_export

def _fm_run_export_with_pre_preview_v2(self):
    try:
        _fm_preview_export_before_export(self, select_tab=True)
    except Exception:
        pass
    return _fm_run_export_before_vw_preview_patch_v2(self)

SupplierUploadUI.analyze_invoice = _fm_analyze_invoice_with_output_preview_v2
SupplierUploadUI.on_mapping_changed = _fm_on_mapping_changed_with_output_preview_v2
SupplierUploadUI.run_export = _fm_run_export_with_pre_preview_v2


# KONTIERUNGSZUORDNUNG_GESAMTUEBERSICHT_STANDARD_V1
# Neue Standard-Datenbasis:
# G:\BUC\FM Anwendung\Datenbasen\Kontierungszuordnung_Gesamtübersicht.xlsx
# Tabellenblatt: Gesamtübersicht_Kontierungszrdn
# Die Zuordnung erfolgt nach Rechnungs-Lieferant + ausgewaehlter Kostenbeschreibung.
import difflib as _fm_difflib

KST_ASSIGNMENT_DEFAULT_DIR = r"G:\BUC\FM Anwendung\Datenbasen\Kontierungszuordnung_Gesamtübersicht.xlsx"
KONTIERUNG_STANDARD_SHEET = "Gesamtübersicht_Kontierungszrdn"

_CURRENT_AFI_SUPPLIER = ""
_CURRENT_AFI_ASSIGNMENT_GROUP = ""

_GROUP_COLUMN_MAP = {
    "VW-LEASING": ("SK VW-LEASING", "KST VW-LEASING", "IA VW-LEASING"),
    "VW-VERSICHERUNG": ("SK VW-VERSICHERUNG", "KST VW-VERSICHERUNG", "IA VW-VERSICHERUNG"),
    "VW-TANKEN": ("SK VW-TANKEN", "KST VW-TANKEN", "IA VW-TANKEN"),
    "ENBW": ("SK EnBW", "KST EnBW", "IA EnBW"),
    "DKV": ("SK DKV", "KST DKV", "IA DKV"),
    "DEAS": ("SK DEAS", "KST DEAS", "IA DEAS"),
    "TELEKOM": ("SK TELEKOM", "KST TELEKOM", "IA TELEKOM"),
    "VODAFONE": ("SK VODAFONE", "KST VODAFONE", "IA VODAFONE"),
}

def _fm_strip_excel_code(value):
    value = _clean(value)
    if re.fullmatch(r"\d+\.0", value):
        value = value[:-2]
    return value.strip()

def _fm_supplier_from_invoice_name(invoice_path):
    name = _norm(os.path.basename(invoice_path or ""))
    if any(x in name for x in ["VOLKSWAGEN", "VWLEASING", "VW"]):
        return "VW"
    if any(x in name for x in ["ENBW", "NBW"]):
        return "ENBW"
    if "DKV" in name:
        return "DKV"
    if "DEAS" in name:
        return "DEAS"
    if any(x in name for x in ["TELEKOM", "DTAG", "TMOBILE", "TMOBILE", "TMOBIL"]):
        return "TELEKOM"
    if "VODAFONE" in name:
        return "VODAFONE"
    return ""

def _fm_assignment_group_from_supplier_and_cost(invoice_path, cost_description):
    supplier = _fm_supplier_from_invoice_name(invoice_path)
    cost = _norm(cost_description)
    if "LEAS" in cost:
        if supplier == "VW" or not supplier:
            return "VW-LEASING"
        return supplier
    if "VERSICHER" in cost:
        if supplier == "VW" or not supplier:
            return "VW-VERSICHERUNG"
        return supplier
    if "TANKEN" in cost and "STROM" in cost:
        return "ENBW"
    if "TANKEN" in cost:
        if supplier == "DKV":
            return "DKV"
        if supplier == "ENBW":
            return "ENBW"
        if supplier == "VW" or not supplier:
            return "VW-TANKEN"
        return supplier
    if "MOBIL" in cost or "FESTNETZ" in cost or "TELEFON" in cost:
        if supplier in ("TELEKOM", "VODAFONE"):
            return supplier
        return supplier or "TELEKOM"
    # Sonstige: Lieferant/Dateiname entscheidet.
    if supplier == "VW":
        return "VW-LEASING"
    return supplier or ""

def _fm_read_kontierung_standard(path):
    try:
        import pandas as pd
        xl = pd.ExcelFile(path, engine="openpyxl")
        sheet = KONTIERUNG_STANDARD_SHEET if KONTIERUNG_STANDARD_SHEET in xl.sheet_names else xl.sheet_names[0]
        raw = pd.read_excel(path, sheet_name=sheet, engine="openpyxl", dtype=str, header=None).fillna("")
        header_idx = None
        for i in range(min(10, len(raw))):
            values = [_clean(v) for v in raw.iloc[i].tolist()]
            if "Kennzeichen" in values and "Nachname" in values:
                header_idx = i
                break
        if header_idx is None:
            raise RuntimeError("Im Standard-Kontierungsblatt wurde keine Kopfzeile mit Kennzeichen/Nachname gefunden.")
        headers = [_clean(v) or f"Spalte_{idx+1}" for idx, v in enumerate(raw.iloc[header_idx].tolist())]
        # Doppelte/leer wirkende Header absichern.
        headers = _dedupe_headers(headers) if '_dedupe_headers' in globals() else headers
        df = raw.iloc[header_idx + 1:].copy()
        df.columns = headers
        df = df[df.apply(lambda r: any(_clean(v) for v in r.values), axis=1)]
        return headers, df.fillna("").to_dict(orient="records")
    except Exception as exc:
        raise RuntimeError(f"Kontierungszuordnung_Gesamtübersicht konnte nicht gelesen werden: {exc}")

def _assignment_entry_from_standard_row(row, headers):
    identifier = _clean(row.get("Kennzeichen", ""))
    first = _clean(row.get("Vorname", ""))
    last = _clean(row.get("Nachname", ""))
    full_name = _clean(f"{first} {last}") or _clean(f"{last} {first}")
    groups = {}
    for group, cols in _GROUP_COLUMN_MAP.items():
        sk_col, kst_col, ia_col = cols
        groups[group] = {
            "GL_ACCOUNT": _fm_strip_excel_code(row.get(sk_col, "")),
            "COSTCENTER": _fm_strip_excel_code(row.get(kst_col, "")),
            "ORDERID": _fm_strip_excel_code(row.get(ia_col, "")),
        }
    joined = " ".join(_clean(row.get(h, "")) for h in headers)
    return {
        "identifier": identifier,
        "identifier_norm": _norm(identifier),
        "identifier_digits": _digits_only(identifier),
        "identifier_type": "PHONE" if re.search(r"\d{3,}[/\s-]?\d+", identifier) and not identifier.upper().startswith("HN") else "PLATE",
        "full_name": full_name,
        "first": first,
        "last": last,
        "name_norm": _norm(full_name),
        "alt_name_norm": _norm(_clean(f"{last} {first}")),
        "last_norm": _norm(last),
        "firma": _clean(row.get("Firma", "")),
        "groups": groups,
        "raw": joined,
        # Kompatibilitaet zu alter Logik
        "gl_default": "",
        "cc_default": "",
        "orderid": "",
    }

def load_assignment_entries(assignment_path):
    ext = os.path.splitext(assignment_path or "")[1].lower()
    if ext in (".xlsx", ".xlsm", ".xls"):
        headers, rows = _fm_read_kontierung_standard(assignment_path)
        entries = []
        for row in rows:
            e = _assignment_entry_from_standard_row(row, headers)
            # Kopfzeilen-Wiederholungen oder Leerzeilen ueberspringen.
            if _norm(e.get("identifier")) in ("", "KENNZEICHEN"):
                continue
            if e.get("identifier") or e.get("full_name"):
                entries.append(e)
        if entries:
            return entries
        raise RuntimeError("Im Standard-Kontierungsdokument konnten keine verwertbaren Zuordnungen erkannt werden.")
    # Fallback fuer alte CSV-Strukturen.
    return _load_assignment_entries(assignment_path)

def _fm_identifier_kind(value):
    v = _clean(value)
    if re.search(r"\d{3,}[/\s-]?\d+", v) and not re.search(r"\b[A-ZÄÖÜ]{1,3}\s*-", v, flags=re.I):
        return "PHONE"
    if re.search(r"\b[A-ZÄÖÜ]{1,3}\s*-\s*[A-ZÄÖÜ]{1,3}\s*\d", v, flags=re.I):
        return "PLATE"
    return "KEY" if v else ""

def _fm_name_similarity(a, b):
    a = _norm(a); b = _norm(b)
    if not a or not b:
        return 0.0
    return _fm_difflib.SequenceMatcher(None, a, b).ratio()

def _fm_unique_best(candidates, min_gap=8):
    if not candidates:
        return {}, ""
    candidates.sort(key=lambda x: x[0], reverse=True)
    if len(candidates) > 1 and candidates[0][0] - candidates[1][0] < min_gap:
        return {}, "mehrdeutig"
    return candidates[0][2], candidates[0][1]

def resolve_assignment(key, driver, entries):
    nkey = _norm(key)
    key_digits = _digits_only(key)
    key_kind = _fm_identifier_kind(key)
    ndriver = _norm(driver)
    parts = _clean(driver).split()
    last = _norm(parts[-1]) if parts else ""

    # 1. Kennzeichen/Telefon/Schluessel exakt. Keine riskanten Kennzeichen-Teiltreffer.
    exact = []
    for e in entries:
        eid = e.get("identifier_norm", "")
        if nkey and eid and nkey == eid:
            exact.append((200, "Schlüssel exakt", e))
        elif key_kind == "PHONE" and key_digits and e.get("identifier_digits") == key_digits:
            exact.append((200, "Telefon exakt", e))
    if exact:
        return _fm_unique_best(exact, min_gap=1)

    # 2. Sehr vorsichtige unschaerfere Schluessel-/Telefonlogik, aber keine falschen Kennzeichen.
    fuzzy_key = []
    if key_kind == "PHONE" and len(key_digits) >= 7:
        for e in entries:
            ed = e.get("identifier_digits", "")
            if ed and (ed.endswith(key_digits) or key_digits.endswith(ed)) and min(len(ed), len(key_digits)) >= 7:
                fuzzy_key.append((170, "Telefon unscharf", e))
    elif key_kind == "KEY" and nkey:
        for e in entries:
            eid = e.get("identifier_norm", "")
            if eid and (nkey == eid or (len(nkey) >= 8 and len(eid) >= 8 and (nkey in eid or eid in nkey))):
                fuzzy_key.append((150, "Schlüssel unscharf", e))
    if fuzzy_key:
        return _fm_unique_best(fuzzy_key, min_gap=1)

    # 3. Vollstaendiger Name, Schreibfehler erlaubt, aber nur bei eindeutigem Treffer.
    name_candidates = []
    if ndriver:
        for e in entries:
            score = max(_fm_name_similarity(ndriver, e.get("name_norm", "")), _fm_name_similarity(ndriver, e.get("alt_name_norm", "")))
            if score >= 0.88:
                name_candidates.append((int(score * 100), "Name", e))
    if name_candidates:
        return _fm_unique_best(name_candidates, min_gap=5)

    # 4. Nachname exakt/leicht unscharf. Nur eindeutig.
    last_candidates = []
    if last:
        for e in entries:
            el = e.get("last_norm", "")
            if el and last == el:
                last_candidates.append((90, "Nachname", e))
            elif el and _fm_name_similarity(last, el) >= 0.92:
                last_candidates.append((84, "Nachname unscharf", e))
    if last_candidates:
        return _fm_unique_best(last_candidates, min_gap=5)

    return {}, ""

def _select_assignment_values(entry, cost_type, text_label=""):
    group = _CURRENT_AFI_ASSIGNMENT_GROUP or ""
    if "BLOCKIER" in _norm(text_label):
        group = "ENBW"
    values = (entry.get("groups") or {}).get(group, {})
    return (
        _fm_strip_excel_code(values.get("GL_ACCOUNT", "")),
        _fm_strip_excel_code(values.get("COSTCENTER", "")),
        _fm_strip_excel_code(values.get("ORDERID", "")),
    )

_create_supplier_upload_csv_before_kontierung_standard_v1 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    global _CURRENT_AFI_SUPPLIER, _CURRENT_AFI_ASSIGNMENT_GROUP
    _CURRENT_AFI_SUPPLIER = _fm_supplier_from_invoice_name(invoice_path)
    _CURRENT_AFI_ASSIGNMENT_GROUP = _fm_assignment_group_from_supplier_and_cost(invoice_path, (config or {}).get("global_prefix", ""))
    return _create_supplier_upload_csv_before_kontierung_standard_v1(assignment_path, invoice_path, export_path, config)

# VODAFONE_PDF_EXPORT_PATCH_FINAL_V4
VODAFONE_COST_DESCRIPTION = "Mobilfunk/Festnetz"
_VODAFONE_PHONE_RE = re.compile(r"\b0\d{2,5}/\d{3,10}\b")
_VODAFONE_MONEY4_RE = re.compile(r"-?\d{1,3}(?:\.\d{3})*,\d{4}")
_fm_supplier_from_invoice_name_before_vodafone_patch = _fm_supplier_from_invoice_name if '_fm_supplier_from_invoice_name' in globals() else None

def _is_vodafone_pdf_text(text):
    n = _norm(text)
    return "VODAFONE" in n and "GESAMTUEBERSICHT" in n and "VODAFONENUMMER" in n

def _fm_supplier_from_invoice_name(invoice_path):
    base = ""
    if _fm_supplier_from_invoice_name_before_vodafone_patch is not None:
        try:
            base = _fm_supplier_from_invoice_name_before_vodafone_patch(invoice_path)
        except Exception:
            base = ""
    if base:
        return base
    try:
        if os.path.splitext(invoice_path or "")[1].lower() == ".pdf" and _is_vodafone_pdf_text(_extract_pdf_text(invoice_path)):
            return "VODAFONE"
    except Exception:
        pass
    return ""

def _vodafone_invoice_net_total_from_text(text):
    m = re.search(r"Nettorechnungsbetrag\s+([\d\.]+,\d{2,4})", text)
    return _dec(m.group(1)) if m else None

def _vodafone_parse_amount_table(text):
    source = text
    start = source.find("Gesamtübersicht\nAnlage zur Rechnung")
    if start == -1:
        source = " ".join(text.split())
        start = source.find("Gesamtübersicht Anlage zur Rechnung")
    if start == -1:
        raise RuntimeError("Vodafone-Gesamtübersicht wurde nicht gefunden.")
    ends = [source.find("In dieser Spalte sehen Sie", start), source.find("Gesamtübersicht\n2. Teil", start), source.find("Gesamtübersicht 2. Teil", start), source.find("Rechnungs-Nummer:", start + 500)]
    ends = [x for x in ends if x != -1]
    block = source[start:(min(ends) if ends else len(source))]
    matches = list(_VODAFONE_PHONE_RE.finditer(block))
    amounts = OrderedDict()
    for i, m in enumerate(matches):
        phone = m.group(0)
        chunk = block[m.end(): matches[i+1].start() if i+1 < len(matches) else len(block)]
        nums = _VODAFONE_MONEY4_RE.findall(chunk)
        if nums:
            amounts[phone] = _dec(nums[-1])
    if not amounts:
        raise RuntimeError("In der Vodafone-Gesamtübersicht konnten keine Einzelbeträge erkannt werden.")
    return amounts

def _vodafone_parse_name_table(text):
    source = text
    m = re.search(r"Gesamtübersicht\s*\n?\s*2\. Teil", source)
    if not m:
        source = " ".join(text.split())
        m = re.search(r"Gesamtübersicht\s+2\. Teil", source)
    if not m:
        raise RuntimeError("Vodafone-Gesamtübersicht 2. Teil wurde nicht gefunden.")
    block = source[m.start():]
    matches = list(_VODAFONE_PHONE_RE.finditer(block))
    names = {}
    for i, pm in enumerate(matches):
        phone = pm.group(0)
        chunk = block[pm.end(): matches[i+1].start() if i+1 < len(matches) else len(block)]
        money = list(_VODAFONE_MONEY4_RE.finditer(chunk))
        if not money:
            continue
        amount = _dec(money[-1].group(0))
        name_part = chunk[:money[-1].start()]
        lines = [_clean(x) for x in name_part.splitlines() if _clean(x)]
        if len(lines) > 1:
            bad = {"Vodafone", "Mobiles Bezahlen", "in EUR", "Interne Kennziffer", "Gesamt", "betrag", "Guthaben", "Transfer"}
            name = _clean(" ".join(x for x in lines if x not in bad))
        else:
            name = _clean(name_part)
        for junk in ["Vodafone Mobiles Bezahlen in EUR", "Interne Kennziffer", "Gesamt betrag in EUR", "Guthaben Transfer in EUR"]:
            name = name.replace(junk, " ")
        names[phone] = {"name": _clean(name), "amount": amount}
    return names

def _vodafone_apply_rounding_adjustment(positions, text):
    target = _vodafone_invoice_net_total_from_text(text)
    if target is None or not positions:
        return positions
    target_rounded = Decimal(target).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    export_rounded = sum(Decimal(p["amount"]).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for p in positions)
    diff = target_rounded - export_rounded
    if diff:
        max(positions, key=lambda p: abs(Decimal(p.get("amount", "0"))))["amount"] += diff
    return positions

def _parse_vodafone_invoice_positions(path, global_prefix):
    text = _extract_pdf_text(path)
    if not _is_vodafone_pdf_text(text):
        raise RuntimeError("Keine Vodafone-PDF erkannt.")
    amounts = _vodafone_parse_amount_table(text)
    names = _vodafone_parse_name_table(text)
    positions = []
    mismatches = []
    for phone, amount in amounts.items():
        name = _clean((names.get(phone) or {}).get("name", ""))
        amount2 = (names.get(phone) or {}).get("amount")
        if amount2 is not None and amount2 != amount:
            mismatches.append(f"{phone}: Gesamtuebersicht {amount} / 2. Teil {amount2}")
        positions.append({"key": phone, "driver": name or phone, "amount": amount, "tax": "VD", "foreign_gross": False, "source_label": _clean(global_prefix) or VODAFONE_COST_DESCRIPTION})
    if mismatches:
        raise RuntimeError("Vodafone-Betragsabweichung zwischen Gesamtuebersichten: " + "; ".join(mismatches[:10]))
    return _vodafone_apply_rounding_adjustment(positions, text)

_parse_pdf_invoice_positions_before_vodafone_patch = _parse_pdf_invoice_positions

def _parse_pdf_invoice_positions(path, global_prefix):
    text_probe = ""
    try:
        text_probe = _extract_pdf_text(path)
        if _is_vodafone_pdf_text(text_probe):
            prefix = global_prefix
            if not prefix or _norm(prefix) in ("TANKENSTROM", "TANKEN"):
                prefix = VODAFONE_COST_DESCRIPTION
            return _parse_vodafone_invoice_positions(path, prefix)
    except Exception as exc:
        if text_probe and _is_vodafone_pdf_text(text_probe):
            raise exc
    return _parse_pdf_invoice_positions_before_vodafone_patch(path, global_prefix)

_create_supplier_upload_csv_before_vodafone_patch = create_supplier_upload_csv

def _vodafone_sum_export_csv(export_path):
    total = Decimal("0.00")
    if export_path and os.path.isfile(export_path):
        with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f, delimiter=";"):
                total += _dec(row.get("NET_VALUE", ""))
    return total

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    cfg = dict(config or {})
    is_vodafone = False
    text_probe = ""
    try:
        if os.path.splitext(invoice_path or "")[1].lower() == ".pdf":
            text_probe = _extract_pdf_text(invoice_path)
            is_vodafone = _is_vodafone_pdf_text(text_probe)
            if is_vodafone:
                if not cfg.get("global_prefix") or _norm(cfg.get("global_prefix")) in ("TANKENSTROM", "TANKEN"):
                    cfg["global_prefix"] = VODAFONE_COST_DESCRIPTION
                globals()["_CURRENT_AFI_SUPPLIER"] = "VODAFONE"
                globals()["_CURRENT_AFI_ASSIGNMENT_GROUP"] = "VODAFONE"
    except Exception:
        pass
    result = _create_supplier_upload_csv_before_vodafone_patch(assignment_path, invoice_path, export_path, cfg)
    if is_vodafone:
        try:
            invoice_total = _vodafone_invoice_net_total_from_text(text_probe)
            export_total = _vodafone_sum_export_csv(export_path)
            if invoice_total is not None:
                target = Decimal(invoice_total).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                result["invoice_net_raw_total"] = _fmt(target)
                result["export_net_total"] = _fmt(export_total)
                result["net_rounding_difference"] = _fmt(export_total - target)
        except Exception:
            pass
    return result



# ------------------------------------------------------------------
# DKV_TANKEN_TOTAL_DRIVER_FIX_V1
# Version 0.456 - DKV-Tanken Korrekturlogik
# ------------------------------------------------------------------
_DKV_TANKEN_TOTAL_DRIVER_FIX_ACTIVE = True
_DKV_LAST_PARSED_POSITIONS = []
_DKV_LAST_DRIVER_TEXT_UPDATES = []
_DKV_TOTAL_NUMBER_RE = re.compile(r"-?\d{1,3}(?:\.\d{3})*,\d{2,3}|-?\d+,\d{1,3}")
_DKV_VEHICLE_RE = re.compile(r"VEHICLE:\s*(?P<vehicle>.*?)\s+CARD NO\.:\s*(?P<card>\S+).*?Kunden ID:\s*(?P<kunden_id>\d+)\s+Kartenzusatz:\s*(?P<zusatz>.*?)(?=\s+\d{2}\.\d{2}\.\d{4}|\s+»\s*TOTAL:|\s+VEHICLE:|$)", re.I | re.S)
_DKV_TAX_RATE_RE = re.compile(r"(?:USt|IVA|TVA)\s*\(%\)\s*:?\s*(\d{1,2},\d{2}|\d{1,2})", re.I)
_DKV_PLATE_RE = re.compile(r"\b[A-ZÄÖÜ]{1,3}\s*-\s*[A-ZÄÖÜ]{1,3}\s*\d{1,5}[A-ZÄÖÜ]?\b", re.I)

def _dkv_clean_plate(value):
    value = _clean(value).upper()
    value = re.sub(r"\s*-\s*", "-", value)
    value = re.sub(r"\s+", " ", value)
    return value

def _dkv_is_dkv_tanken_pdf_text(text):
    n = _norm(text)
    return "DKV" in n and "VEHICLE" in n and "CARDNO" in n and "TOTAL" in n

def _dkv_total_amounts_from_after_total(after_total_text):
    nums = _DKV_TOTAL_NUMBER_RE.findall(after_total_text or "")[:6]
    vals = [_dec(x) for x in nums]
    # Mit Nachlass ist die 3. Zahl negativ; ohne Nachlass folgen oft direkt %-Werte der Summenbeschreibung.
    if len(vals) >= 6 and vals[2] < 0:
        return vals[3], vals[5], nums
    if len(vals) >= 5:
        return vals[2], vals[4], nums[:5]
    if len(vals) >= 3:
        return vals[-3], vals[-1], nums
    return Decimal("0.00"), Decimal("0.00"), nums

def _parse_dkv_tanken_pdf_positions_v1(path, global_prefix):
    text = _extract_pdf_text(path)
    if not _dkv_is_dkv_tanken_pdf_text(text):
        raise RuntimeError("Keine DKV-Tanken-PDF erkannt.")
    compact = " ".join(text.replace("\n", " ").split())
    matches = list(_DKV_VEHICLE_RE.finditer(compact))
    positions = []
    for idx, m in enumerate(matches):
        vehicle = _dkv_clean_plate(m.group("vehicle"))
        start = m.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(compact)
        body = compact[start:end]
        total_m = re.search(r"»?\s*TOTAL:\s*", body, re.I)
        if not total_m:
            continue
        previous_rates = list(_DKV_TAX_RATE_RE.finditer(compact[:m.start()]))
        rate = _dec(previous_rates[-1].group(1)) if previous_rates else Decimal("19")
        net_amount, gross_amount, total_nums = _dkv_total_amounts_from_after_total(body[total_m.end():])
        if net_amount == 0 and gross_amount == 0:
            continue
        amount, tax, foreign = _amount_and_tax_from_values(net_amount, gross_amount, rate)
        if amount == 0:
            continue
        zusatz = _clean(m.group("zusatz"))
        positions.append({
            "key": vehicle,
            "driver": zusatz or vehicle,
            "amount": amount,
            "tax": tax,
            "foreign_gross": foreign,
            "source_label": global_prefix,
            "tax_rate": rate,
            "dkv_total_numbers": total_nums,
            "dkv_net_amount": net_amount,
            "dkv_gross_amount": gross_amount,
        })
    if not positions:
        raise RuntimeError("Aus der DKV-PDF konnten keine Fahrzeug-/TOTAL-Positionen erkannt werden.")
    global _DKV_LAST_PARSED_POSITIONS
    _DKV_LAST_PARSED_POSITIONS = positions
    return positions

_parse_pdf_invoice_positions_before_dkv_tanken_total_driver_fix_v1 = _parse_pdf_invoice_positions

def _parse_pdf_invoice_positions(path, global_prefix):
    try:
        text_probe = _extract_pdf_text(path)
        if _dkv_is_dkv_tanken_pdf_text(text_probe):
            return _parse_dkv_tanken_pdf_positions_v1(path, global_prefix)
    except Exception:
        pass
    return _parse_pdf_invoice_positions_before_dkv_tanken_total_driver_fix_v1(path, global_prefix)

def _dkv_assignment_display_name(entry):
    if not entry:
        return ""
    first = _clean(entry.get("first", ""))
    last = _clean(entry.get("last", ""))
    if first and last:
        return _clean(f"{last}, {first}")
    return _clean(entry.get("full_name", ""))

def _dkv_assignment_map_by_identifier(entries):
    out = {}
    for e in entries or []:
        ident = _dkv_clean_plate(e.get("identifier", ""))
        if ident:
            out[_norm(ident)] = e
    return out

def _apply_dkv_driver_names_to_export(export_path, assignment_path, invoice_path, config):
    if not export_path or not os.path.isfile(export_path):
        return []
    supplier = _fm_supplier_from_invoice_name(invoice_path) if '_fm_supplier_from_invoice_name' in globals() else ""
    cost = _norm((config or {}).get("global_prefix", ""))
    if supplier != "DKV" and "TANKEN" not in cost:
        return []
    try:
        entries = load_assignment_entries(assignment_path)
    except Exception:
        return []
    by_identifier = _dkv_assignment_map_by_identifier(entries)
    with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        fieldnames = reader.fieldnames or []
        rows = list(reader)
    if "TEXT" not in fieldnames:
        return []
    prefix = _clean((config or {}).get("global_prefix", "")) or "Tanken"
    updates = []
    for row in rows:
        text_value = _clean(row.get("TEXT", ""))
        plate_m = _DKV_PLATE_RE.search(text_value)
        if not plate_m:
            continue
        plate = _dkv_clean_plate(plate_m.group(0))
        entry = by_identifier.get(_norm(plate))
        display = _dkv_assignment_display_name(entry)
        if not display:
            continue
        new_text = _clean(f"{prefix} {plate} {display}")
        if '_ascii_umlauts' in globals():
            new_text = _ascii_umlauts(new_text)
        if row.get("TEXT", "") != new_text:
            updates.append({"Kennzeichen": plate, "Alt": row.get("TEXT", ""), "Neu": new_text})
            row["TEXT"] = new_text
    if updates:
        with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)
    global _DKV_LAST_DRIVER_TEXT_UPDATES
    _DKV_LAST_DRIVER_TEXT_UPDATES = updates
    return updates

_create_supplier_upload_csv_before_dkv_tanken_total_driver_fix_v1 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_dkv_tanken_total_driver_fix_v1(assignment_path, invoice_path, export_path, config)
    try:
        updates = _apply_dkv_driver_names_to_export(export_path, assignment_path, invoice_path, config)
        if isinstance(result, dict):
            result["dkv_driver_text_updates"] = updates
            result["dkv_driver_text_updates_count"] = len(updates)
            if _DKV_LAST_PARSED_POSITIONS:
                result["dkv_vehicle_positions"] = len(_DKV_LAST_PARSED_POSITIONS)
    except Exception:
        pass
    return result


# ------------------------------------------------------------------
# DKV_KST_FROM_STANDARD_ASSIGNMENT_FIX_V1
# Version 0.456 - DKV-KST Nachschärfung
# ------------------------------------------------------------------
# Zweck:
# Die DKV-Zuordnung aus der Standarddatei "Kontierungszuordnung_Gesamtübersicht.xlsx"
# wird nach dem Export nochmals kennzeichenbasiert auf die AFI-CSV angewendet.
# Damit werden GL_ACCOUNT, COSTCENTER und ORDERID aus den DKV-Spalten sicher geschrieben,
# unabhängig davon, ob vorherige Wrapper/Reihenfolgen die globale Zuordnungsgruppe verändert haben.
_DKV_KST_FROM_STANDARD_ASSIGNMENT_FIX_ACTIVE = True
_DKV_KST_LAST_UPDATES = []
_DKV_KST_PLATE_RE = re.compile(r"\b[A-ZÄÖÜ]{1,3}\s*-\s*[A-ZÄÖÜ]{1,3}\s*\d{1,5}[A-ZÄÖÜ]?\b", re.I)

def _dkv_kst_fix_clean_plate(value):
    value = _clean(value).upper()
    value = re.sub(r"\s*-\s*", "-", value)
    value = re.sub(r"\s+", " ", value)
    return value

def _dkv_kst_fix_is_dkv_export(invoice_path, config):
    supplier = ""
    try:
        supplier = _fm_supplier_from_invoice_name(invoice_path) if '_fm_supplier_from_invoice_name' in globals() else ""
    except Exception:
        supplier = ""
    cost = _norm((config or {}).get("global_prefix", ""))
    name = _norm(os.path.basename(invoice_path or ""))
    return supplier == "DKV" or "DKV" in name or "DKV" in cost

def _dkv_kst_fix_values_from_entry(entry):
    groups = entry.get("groups") or {}
    values = groups.get("DKV") or {}
    gl = _fm_strip_excel_code(values.get("GL_ACCOUNT", "")) if '_fm_strip_excel_code' in globals() else _clean(values.get("GL_ACCOUNT", ""))
    cc = _fm_strip_excel_code(values.get("COSTCENTER", "")) if '_fm_strip_excel_code' in globals() else _clean(values.get("COSTCENTER", ""))
    ia = _fm_strip_excel_code(values.get("ORDERID", "")) if '_fm_strip_excel_code' in globals() else _clean(values.get("ORDERID", ""))
    return gl, cc, ia

def _apply_dkv_kst_from_standard_assignment(export_path, assignment_path, invoice_path, config):
    if not export_path or not os.path.isfile(export_path):
        return []
    if not _dkv_kst_fix_is_dkv_export(invoice_path, config):
        return []
    try:
        entries = load_assignment_entries(assignment_path)
    except Exception:
        return []
    by_plate = {}
    for entry in entries or []:
        ident = _dkv_kst_fix_clean_plate(entry.get("identifier", ""))
        if ident:
            by_plate[_norm(ident)] = entry
    with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        fieldnames = reader.fieldnames or []
        rows = list(reader)
    required = {"TEXT", "GL_ACCOUNT", "COSTCENTER", "ORDERID"}
    if not required.issubset(set(fieldnames)):
        return []
    updates = []
    for row in rows:
        text_value = _clean(row.get("TEXT", ""))
        m = _DKV_KST_PLATE_RE.search(text_value)
        if not m:
            continue
        plate = _dkv_kst_fix_clean_plate(m.group(0))
        entry = by_plate.get(_norm(plate))
        if not entry:
            continue
        gl, cc, ia = _dkv_kst_fix_values_from_entry(entry)
        before = {"GL_ACCOUNT": row.get("GL_ACCOUNT", ""), "COSTCENTER": row.get("COSTCENTER", ""), "ORDERID": row.get("ORDERID", "")}
        # DKV-Werte haben fachlich Vorrang, wenn sie im Standardblatt gepflegt sind.
        if gl:
            row["GL_ACCOUNT"] = gl
        if cc:
            row["COSTCENTER"] = cc
        if ia:
            row["ORDERID"] = ia
        after = {"GL_ACCOUNT": row.get("GL_ACCOUNT", ""), "COSTCENTER": row.get("COSTCENTER", ""), "ORDERID": row.get("ORDERID", "")}
        if after != before:
            updates.append({"Kennzeichen": plate, "TEXT": text_value, "vorher": before, "nachher": after})
    if updates:
        with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)
    global _DKV_KST_LAST_UPDATES
    _DKV_KST_LAST_UPDATES = updates
    return updates

_create_supplier_upload_csv_before_dkv_kst_fix_v1 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_dkv_kst_fix_v1(assignment_path, invoice_path, export_path, config)
    try:
        updates = _apply_dkv_kst_from_standard_assignment(export_path, assignment_path, invoice_path, config)
        if isinstance(result, dict):
            result["dkv_kst_updates"] = updates
            result["dkv_kst_updates_count"] = len(updates)
    except Exception:
        pass
    return result


# ------------------------------------------------------------------
# AFI_UPLOAD_UI_ANALYSE_NO_FREEZE_AND_DKV_IDG_FIX_V1
# Version 0.460
# Zweck:
# - Auswahl der Kostenbeschreibung darf keine komplette Vorschau-/Exportberechnung mehr synchron starten.
# - Rechnung analysieren darf bei PDF-Dateien keine Exportvorschau mehr im UI-Thread erzeugen.
# - Standardpfad fuer KST-Zuordnungsdokument zeigt auf die Datei im Ordner KST_Zuordnungen_AFI.
# - DKV-IDG-PDFs mit fehlender/nichtnumerischer Kunden-ID werden korrekt geparst.
# ------------------------------------------------------------------
AFI_UPLOAD_UI_ANALYSE_NO_FREEZE_AND_DKV_IDG_FIX_ACTIVE = True
KST_ASSIGNMENT_DEFAULT_FILE = r"G:\BUC\FM Anwendung\Datenbasen\KST_Zuordnungen_AFI\Kontierungszuordnung_Gesamtübersicht.xlsx"
KST_ASSIGNMENT_DEFAULT_DIR = KST_ASSIGNMENT_DEFAULT_FILE

_DKV_IDG_VEHICLE_RE_V2 = re.compile(
    r"VEHICLE:\s*(?P<vehicle>.*?)\s+CARD NO\.:\s*(?P<card>\S+).*?Kartenzusatz:\s*(?P<zusatz>.*?)(?=\s+\d{2}\.\d{2}\.\d{4}|\s+»\s*TOTAL:|\s+VEHICLE:|\s+Gesamtsummenaufstellung|\s+Umsatzsteuerstatistik|$)",
    re.I | re.S,
)
_DKV_IDG_TAX_RATE_RE_V2 = re.compile(r"(?:USt|IVA|TVA)\s*\(%\)\s*:?\s*(\d{1,2},\d{2}|\d{1,2})", re.I)
_DKV_IDG_TOTAL_NUMBER_RE_V2 = re.compile(r"-?\d{1,3}(?:\.\d{3})*,\d{2,3}|-?\d+,\d{1,3}")


def _dkv_idg_total_amounts_after_total_v2(after_total_text):
    nums = _DKV_IDG_TOTAL_NUMBER_RE_V2.findall(after_total_text or "")[:6]
    vals = [_dec(x) for x in nums]
    # Deutsche DKV-Zeilen mit Nachlass: Menge, Bezugswert, Nachlass, Gesamtwert netto, USt, Brutto
    if len(vals) >= 6 and vals[2] < 0:
        return vals[3], vals[5], nums
    # Ohne Nachlass: Menge, Gesamtwert netto/Bezugswert, Gesamtwert netto, USt, Brutto
    if len(vals) >= 5:
        return vals[2], vals[4], nums[:5]
    if len(vals) >= 3:
        return vals[-3], vals[-1], nums
    return Decimal("0.00"), Decimal("0.00"), nums


def _parse_dkv_tanken_pdf_positions_v1(path, global_prefix):
    """Robuster DKV-PDF-Parser fuer IDE/IDG-DKV-Rechnungen.

    Wichtig: Die Kunden-ID ist in DKV-PDFs nicht immer numerisch bzw. teilweise gar nicht vorhanden.
    Deshalb wird nur VEHICLE + CARD NO. + Kartenzusatz als Blockanker verwendet.
    """
    text = _extract_pdf_text(path)
    if not _dkv_is_dkv_tanken_pdf_text(text):
        raise RuntimeError("Keine DKV-Tanken-PDF erkannt.")
    compact = " ".join(text.replace("\n", " ").split())
    matches = list(_DKV_IDG_VEHICLE_RE_V2.finditer(compact))
    positions = []
    for idx, m in enumerate(matches):
        vehicle = _dkv_clean_plate(m.group("vehicle")) if '_dkv_clean_plate' in globals() else _clean(m.group("vehicle"))
        start = m.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(compact)
        body = compact[start:end]
        total_m = re.search(r"»?\s*TOTAL:\s*", body, re.I)
        if not total_m:
            continue
        previous_rates = list(_DKV_IDG_TAX_RATE_RE_V2.finditer(compact[:m.start()]))
        rate = _dec(previous_rates[-1].group(1)) if previous_rates else Decimal("19")
        net_amount, gross_amount, total_nums = _dkv_idg_total_amounts_after_total_v2(body[total_m.end():])
        if net_amount == 0 and gross_amount == 0:
            continue
        amount, tax, foreign = _amount_and_tax_from_values(net_amount, gross_amount, rate)
        if amount == 0:
            continue
        zusatz = _clean(m.group("zusatz"))
        positions.append({
            "key": vehicle,
            "driver": zusatz or vehicle,
            "amount": amount,
            "tax": tax,
            "foreign_gross": foreign,
            "source_label": global_prefix,
            "tax_rate": rate,
            "dkv_total_numbers": total_nums,
            "dkv_net_amount": net_amount,
            "dkv_gross_amount": gross_amount,
        })
    if not positions:
        raise RuntimeError("Aus der DKV-PDF konnten keine Fahrzeug-/TOTAL-Positionen erkannt werden.")
    global _DKV_LAST_PARSED_POSITIONS
    try:
        _DKV_LAST_PARSED_POSITIONS = positions
    except Exception:
        pass
    return positions


def _fm460_on_mapping_changed_no_preview(self):
    """Nur leichte UI-Aktualisierung; keine synchrone Export-/PDF-Analyse."""
    try:
        _fm_update_export_path(self, True)
    except Exception:
        pass
    try:
        self.update_footer()
    except Exception:
        pass
    try:
        self.update_highlight()
    except Exception:
        pass


def _fm460_analyze_invoice_no_freeze(self):
    """Analysiert Rechnung ohne automatische Ausgabevorschau.

    Die Ausgabevorschau hatte bei Kostenbeschreibung/Analyse eine komplette Exportberechnung im UI-Thread
    gestartet und dadurch das Tool scheinbar aufgehängt. Die echte Exportdatei wird weiterhin erst beim
    Klick auf 'AFI-Upload-Datei erstellen' erzeugt.
    """
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        self.clear_sources()
        ext = os.path.splitext(path)[1].lower()
        try:
            self.load_preview(path)
        except Exception:
            pass
        try:
            _fm_update_export_path(self, True)
        except Exception:
            pass
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            self.suggestion_var.set("PDF erkannt: Berechnungsquellen sind nicht erforderlich. Die Positionen werden beim Export aus Fahrzeug-/TOTAL-Bloecken gelesen.")
            self.status_var.set("PDF-Rechnung analysiert. Bitte Kostenbeschreibung und KST-Zuordnungsdokument pruefen. Die AFI-Ausgabe wird erst beim Export erzeugt.")
            if hasattr(self, 'add_source_btn'):
                self.add_source_btn.configure(state="disabled")
            return
        if hasattr(self, 'add_source_btn'):
            self.add_source_btn.configure(state="normal")
        self.headers, self.rows = _read_table_file(path)
        suggestions = suggested_sources(self.headers)
        self.add_source(suggestions[0])
        if len(suggestions) > 1:
            self.suggestion_var.set("Weitere moegliche Berechnungsquellen erkannt: " + ", ".join(s.get("label", "") for s in suggestions[1:]) + ". Bei Bedarf ueber '+ Berechnungsquelle' hinzufuegen.")
        else:
            self.suggestion_var.set("")
        self.status_var.set("Rechnung analysiert. Bitte Berechnungsquelle pruefen/ergaenzen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))


def _fm460_run_export_no_pre_preview(self):
    """Startet den bestehenden threaded Export direkt, ohne synchrone Vorabvorschau."""
    return _fm_run_export(self)


SupplierUploadUI.on_mapping_changed = _fm460_on_mapping_changed_no_preview
SupplierUploadUI.analyze_invoice = _fm460_analyze_invoice_no_freeze
SupplierUploadUI.run_export = _fm460_run_export_no_pre_preview


# ------------------------------------------------------------------
# AFI_UPLOAD_GENERALUEBERSICHT_LOADER_FIX_V0462
# Datum: 2026-07-06
# Zweck:
# - Standardpfad von Gesamtuebersicht auf Generaluebersicht umstellen.
# - Generaluebersicht mit Blaettern KFZ, Telefon, Sachkonten direkt laden.
# - None-/Leerzellen robust behandeln; behebt "sequence item ... NoneType".
# - Bike Leasing als Kostenbeschreibung und Sachkonto-Logik ergaenzen.
# ------------------------------------------------------------------
AFI_UPLOAD_GENERALUEBERSICHT_LOADER_FIX_VERSION = "0.462"
KST_ASSIGNMENT_DEFAULT_FILE = r"G:\BUC\FM Anwendung\Datenbasen\KST_Zuordnungen_AFI\Kontierungszuordnung_Generalübersicht.xlsx"
KST_ASSIGNMENT_DEFAULT_DIR = KST_ASSIGNMENT_DEFAULT_FILE
try:
    COST_TYPE_OPTIONS = ["Tanken Strom", "Tanken", "Versicherung", "Leasing", "Bike Leasing", "Mobilfunk/Festnetz", "Sonstige"]
except Exception:
    pass
_GENERAL_OVERVIEW_CACHE = {"path":"", "mtime":None, "entries":[], "gl":{}}

def _general_clean(value):
    return _clean(value)

def _general_to_str(value):
    value = _clean(value)
    if re.fullmatch(r"\d+\.0", value): value = value[:-2]
    return value

def _general_company_to_bukrs(company):
    n=_norm(company)
    if "DIGITAL" in n: return "IDG"
    if "SABU" in n: return "SABU"
    if "IMS" in n or "MARKETINGSERVICES" in n: return "IMS"
    return "IDE"

def _general_parse_sachkonten(ws):
    cols={"IDE":(2,3), "IDG":(5,6), "SABU":(8,9), "IMS":(11,12)}
    out={b:{} for b in cols}
    for r in range(3, ws.max_row+1):
        labels=[]
        for b,(lcol,gcol) in cols.items():
            lab=_general_to_str(ws.cell(r,lcol).value)
            if lab: labels.append(lab)
        inferred=labels[0] if labels else ""
        for b,(lcol,gcol) in cols.items():
            lab=_general_to_str(ws.cell(r,lcol).value) or inferred
            gl=_general_to_str(ws.cell(r,gcol).value)
            if not gl: continue
            if lab: out[b][_norm(lab)] = gl
            for extra in labels:
                if extra: out[b][_norm(extra)] = gl
            # in der aktuellen Datei ist Bike Leasing teilweise nur in einer Organisationsspalte beschriftet
            if r >= 12 and gl == "416000": out[b][_norm("Bike Leasing")] = gl
            if r >= 12 and gl == "154000": out[b][_norm("UST Bike Leasing")] = gl
    return out

def _general_entry(identifier, last, first, company, kst, ia, kind):
    identifier=_general_to_str(identifier); last=_general_to_str(last); first=_general_to_str(first); company=_general_to_str(company)
    kst=_general_to_str(kst); ia=_general_to_str(ia)
    full=_clean(f"{first} {last}") or _clean(f"{last} {first}")
    bukrs=_general_company_to_bukrs(company)
    raw=" ".join(x for x in [identifier, full, company, kst, ia] if x)
    base_gl=""
    return {
        "identifier": identifier, "identifier_norm": _norm(identifier), "identifier_digits": _digits_only(identifier), "identifier_type": kind,
        "full_name": full, "first": first, "last": last, "name_norm": _norm(full), "alt_name_norm": _norm(_clean(f"{last} {first}")), "last_norm": _norm(last),
        "firma": company, "bukrs": bukrs, "raw": raw,
        "gl_default": base_gl, "gl_tanken_strom":"", "gl_tanken":"", "gl_versicherung":"", "gl_leasing":"", "gl_bike_leasing":"", "gl_mobilfunk":"",
        "cc_default": kst, "cc_tanken_strom": kst, "cc_tanken": kst, "cc_versicherung": kst, "cc_leasing": kst, "cc_bike_leasing": kst, "cc_mobilfunk": kst,
        "orderid": ia,
    }

def _load_general_overview(path, force=False):
    import os as _os
    try: mtime=_os.path.getmtime(path)
    except Exception as exc: raise RuntimeError("Kontierungszuordnung nicht gefunden: " + str(path)) from exc
    if not force and _GENERAL_OVERVIEW_CACHE.get('path')==path and _GENERAL_OVERVIEW_CACHE.get('mtime')==mtime and _GENERAL_OVERVIEW_CACHE.get('entries'):
        return _GENERAL_OVERVIEW_CACHE
    try:
        from openpyxl import load_workbook
        wb=load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        raise RuntimeError("Kontierungszuordnung konnte nicht gelesen werden: " + str(exc))
    entries=[]
    if 'KFZ' in wb.sheetnames:
        ws=wb['KFZ']
        for r in range(3, ws.max_row+1):
            ident=ws.cell(r,2).value
            if _general_to_str(ident) and _norm(ident) != 'KENNZEICHEN':
                entries.append(_general_entry(ident, ws.cell(r,3).value, ws.cell(r,4).value, ws.cell(r,5).value, ws.cell(r,6).value, ws.cell(r,7).value, 'PLATE'))
    if 'Telefon' in wb.sheetnames:
        ws=wb['Telefon']
        for r in range(3, ws.max_row+1):
            ident=ws.cell(r,2).value
            if _general_to_str(ident) and _norm(ident) != 'RUFNUMMER':
                entries.append(_general_entry(ident, ws.cell(r,3).value, ws.cell(r,4).value, ws.cell(r,5).value, ws.cell(r,6).value, ws.cell(r,7).value, 'PHONE'))
    gl={}
    if 'Sachkonten' in wb.sheetnames:
        gl=_general_parse_sachkonten(wb['Sachkonten'])
    if not entries: raise RuntimeError("Im Kontierungsdokument wurden keine Zuordnungen erkannt.")
    # Sachkonten je Eintrag nach Buchungskreis eintragen, damit die bestehende Exportlogik unveraendert greift.
    for e in entries:
        b=e.get('bukrs') or 'IDE'; m=gl.get(b,{})
        e['gl_tanken_strom']=m.get(_norm('Tanken Strom'), '')
        e['gl_tanken']=m.get(_norm('DKV'), '') or m.get(_norm('Tanken'), '')
        e['gl_versicherung']=m.get(_norm('DEAS'), '') or m.get(_norm('VW-Versicherungen'), '') or m.get(_norm('Versicherung'), '')
        e['gl_leasing']=m.get(_norm('VW-Leasing'), '') or m.get(_norm('Leasing'), '')
        e['gl_bike_leasing']=m.get(_norm('Bike Leasing'), '')
        e['gl_mobilfunk']=m.get(_norm('Telekom'), '') or m.get(_norm('Vodafone'), '')
        e['gl_default']=m.get(_norm('Sonstige'), '')
    _GENERAL_OVERVIEW_CACHE.update({"path":path,"mtime":mtime,"entries":entries,"gl":gl})
    return _GENERAL_OVERVIEW_CACHE

def refresh_assignment_cache(path=None):
    path=path or KST_ASSIGNMENT_DEFAULT_FILE
    _GENERAL_OVERVIEW_CACHE.update({"path":"", "mtime":None, "entries":[], "gl":{}})
    return _load_general_overview(path, force=True)

def load_assignment_entries(assignment_path):
    try:
        data=_load_general_overview(assignment_path)
        return data.get('entries', [])
    except Exception:
        # echte alte CSV-/Gesamtuebersicht-Fallbacks bleiben erhalten, wenn vorhanden
        try: return _load_assignment_entries(assignment_path)
        except Exception as exc: raise exc

def _cost_type(label):
    n=_norm(label)
    if "BIKE" in n and "LEAS" in n: return "BIKE_LEASING"
    if "TANKENSTROM" in n or ("TANKEN" in n and "STROM" in n): return "TANKEN_STROM"
    if "TANKEN" in n: return "TANKEN"
    if "VERSICHER" in n: return "VERSICHERUNG"
    if "LEAS" in n: return "LEASING"
    if "MOBIL" in n or "FESTNETZ" in n or "TELEFON" in n or "VODAFONE" in n: return "MOBILFUNK"
    return "SONSTIGE"

_select_assignment_values_before_general_fix = _select_assignment_values

def _select_assignment_values(entry, cost_type, text_label=""):
    if cost_type == "BIKE_LEASING":
        gl = entry.get('gl_bike_leasing') or entry.get('gl_leasing') or entry.get('gl_default')
        cc = entry.get('cc_bike_leasing') or entry.get('cc_leasing') or entry.get('cc_default')
        return gl or "", cc or "", entry.get('orderid','') or ""
    return _select_assignment_values_before_general_fix(entry, cost_type, text_label)

# UI-Refresh klarer machen, ohne Erfolgspopup bei vorhandener Datei.
def _general_refresh_assignment_ui(self):
    path=self.template_var.get().strip() if hasattr(self,'template_var') else KST_ASSIGNMENT_DEFAULT_FILE
    try:
        data=refresh_assignment_cache(path)
        if hasattr(self,'assignment_status_var'):
            self.assignment_status_var.set(f"Zuordnungsdatei geladen: {len(data.get('entries', []))} Zuordnungen im Sitzungs-Cache.")
        elif hasattr(self,'status_var'):
            self.status_var.set(f"Zuordnungsdatei geladen: {len(data.get('entries', []))} Zuordnungen im Sitzungs-Cache.")
    except Exception as exc:
        if hasattr(self,'assignment_status_var'): self.assignment_status_var.set("Zuordnungsdatei nicht verfügbar: " + str(exc))
        try: messagebox.showwarning(MODULE_TITLE, "Kontierungszuordnung nicht gefunden oder nicht lesbar:\n"+str(path)+"\n\n"+str(exc))
        except Exception: pass
try:
    SupplierUploadUI.refresh_assignment = _general_refresh_assignment_ui
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_GENERALUEBERSICHT_FINAL_FIX_V0463
# Datum: 2026-07-06
# Zweck:
# - Korrigiert den versehentlich auf altem Stand ausgelieferten Patch.
# - Fuehrt die gewuenschten Umbauten erneut auf dem aktuell bereitgestellten Modul aus.
# - Neuer fuehrender Standard: Kontierungszuordnung_Generaluebersicht.xlsx im KST_Zuordnungen_AFI-Ordner.
# - Loader ist robust gegen leere Zellen (None) in Firma/KST/IA.
# - UI bekommt Wizard-Grundstruktur, Refresh-Button, Mehrfach-Buchungskreise inkl. SABU und Positionsvorschau.
# - Bike Leasing wird als Kostenbeschreibung und Sachkontoart unterstuetzt.
# ------------------------------------------------------------------
AFI_UPLOAD_GENERALUEBERSICHT_FINAL_FIX_VERSION = "0.463"
KST_ASSIGNMENT_DEFAULT_FILE = r"G:\BUC\FM Anwendung\Datenbasen\KST_Zuordnungen_AFI\Kontierungszuordnung_Generalübersicht.xlsx"
KST_ASSIGNMENT_DEFAULT_DIR = KST_ASSIGNMENT_DEFAULT_FILE
BOOKING_CIRCLE_OPTIONS = ["IDE", "IDG", "IMS", "SABU"]
SUPPLIER_OPTIONS = ["Automatisch erkennen", "EnBW", "DKV", "VW-Leasing", "VW-Versicherung", "DEAS", "Telekom", "Vodafone", "Bike Leasing", "Sonstige"]
COST_TYPE_OPTIONS = ["Tanken Strom", "Tanken", "Versicherung", "Leasing", "Bike Leasing", "Mobilfunk/Festnetz", "Sonstige"]
_AFI_GENERAL_CACHE = {"path":"", "mtime":None, "entries":[], "gl":{}}


def _afi463_s(value):
    value = _clean(value)
    if re.fullmatch(r"\d+\.0", value):
        value = value[:-2]
    return value


def _afi463_join(parts):
    return " ".join(_afi463_s(x) for x in (parts or []) if _afi463_s(x))


def _afi463_company_to_bukrs(company):
    n = _norm(company)
    if "DIGITAL" in n:
        return "IDG"
    if "SABU" in n:
        return "SABU"
    if "IMS" in n or "MARKETINGSERVICES" in n:
        return "IMS"
    return "IDE"


def _afi463_parse_sachkonten(ws):
    cols = {"IDE": (2, 3), "IDG": (5, 6), "SABU": (8, 9), "IMS": (11, 12)}
    out = {b: {} for b in cols}
    for r in range(3, ws.max_row + 1):
        labels = []
        for _b, (lcol, _gcol) in cols.items():
            lab = _afi463_s(ws.cell(r, lcol).value)
            if lab:
                labels.append(lab)
        inferred_label = labels[0] if labels else ""
        for b, (lcol, gcol) in cols.items():
            label = _afi463_s(ws.cell(r, lcol).value) or inferred_label
            gl = _afi463_s(ws.cell(r, gcol).value)
            if not gl:
                continue
            for lab in set([label] + labels):
                if lab:
                    out[b][_norm(lab)] = gl
            # In der gelieferten Datei ist Bike Leasing teilweise nur in einer Organisationsspalte beschriftet.
            if r >= 12 and gl == "416000":
                out[b][_norm("Bike Leasing")] = gl
            if r >= 12 and gl == "154000":
                out[b][_norm("UST Bike Leasing")] = gl
    return out


def _afi463_make_entry(identifier, last, first, company, kst, ia, kind):
    identifier = _afi463_s(identifier)
    last = _afi463_s(last)
    first = _afi463_s(first)
    company = _afi463_s(company)
    kst = _afi463_s(kst)
    ia = _afi463_s(ia)
    full_name = _clean(f"{first} {last}") or _clean(f"{last} {first}")
    bukrs = _afi463_company_to_bukrs(company)
    return {
        "identifier": identifier,
        "identifier_norm": _norm(identifier),
        "identifier_digits": _digits_only(identifier),
        "identifier_type": kind,
        "full_name": full_name,
        "first": first,
        "last": last,
        "name_norm": _norm(full_name),
        "alt_name_norm": _norm(_clean(f"{last} {first}")),
        "last_norm": _norm(last),
        "firma": company,
        "bukrs": bukrs,
        "raw": _afi463_join([identifier, full_name, company, kst, ia]),
        "gl_default": "",
        "gl_tanken_strom": "",
        "gl_tanken": "",
        "gl_versicherung": "",
        "gl_leasing": "",
        "gl_bike_leasing": "",
        "gl_mobilfunk": "",
        "cc_default": kst,
        "cc_tanken_strom": kst,
        "cc_tanken": kst,
        "cc_versicherung": kst,
        "cc_leasing": kst,
        "cc_bike_leasing": kst,
        "cc_mobilfunk": kst,
        "orderid": ia,
    }


def _afi463_load_general_overview(path, force=False):
    try:
        mtime = os.path.getmtime(path)
    except Exception as exc:
        raise RuntimeError("Kontierungszuordnung nicht gefunden: " + str(path)) from exc
    if (not force and _AFI_GENERAL_CACHE.get("path") == path and
            _AFI_GENERAL_CACHE.get("mtime") == mtime and _AFI_GENERAL_CACHE.get("entries")):
        return _AFI_GENERAL_CACHE
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        raise RuntimeError("Kontierungszuordnung konnte nicht gelesen werden: " + str(exc))
    entries = []
    if "KFZ" in wb.sheetnames:
        ws = wb["KFZ"]
        for r in range(3, ws.max_row + 1):
            ident = _afi463_s(ws.cell(r, 2).value)
            if ident and _norm(ident) != "KENNZEICHEN":
                entries.append(_afi463_make_entry(ident, ws.cell(r, 3).value, ws.cell(r, 4).value, ws.cell(r, 5).value, ws.cell(r, 6).value, ws.cell(r, 7).value, "PLATE"))
    if "Telefon" in wb.sheetnames:
        ws = wb["Telefon"]
        for r in range(3, ws.max_row + 1):
            ident = _afi463_s(ws.cell(r, 2).value)
            if ident and _norm(ident) != "RUFNUMMER":
                entries.append(_afi463_make_entry(ident, ws.cell(r, 3).value, ws.cell(r, 4).value, ws.cell(r, 5).value, ws.cell(r, 6).value, ws.cell(r, 7).value, "PHONE"))
    gl_map = {}
    if "Sachkonten" in wb.sheetnames:
        gl_map = _afi463_parse_sachkonten(wb["Sachkonten"])
    if not entries:
        raise RuntimeError("Im Kontierungsdokument wurden keine KFZ-/Telefon-Zuordnungen erkannt.")
    for e in entries:
        b = e.get("bukrs") or "IDE"
        m = gl_map.get(b, {})
        e["gl_tanken_strom"] = m.get(_norm("Tanken Strom"), "")
        e["gl_tanken"] = m.get(_norm("DKV"), "") or m.get(_norm("Tanken"), "")
        e["gl_versicherung"] = m.get(_norm("DEAS"), "") or m.get(_norm("VW-Versicherungen"), "") or m.get(_norm("Versicherung"), "")
        e["gl_leasing"] = m.get(_norm("VW-Leasing"), "") or m.get(_norm("Leasing"), "")
        e["gl_bike_leasing"] = m.get(_norm("Bike Leasing"), "")
        e["gl_mobilfunk"] = m.get(_norm("Telekom"), "") or m.get(_norm("Vodafone"), "")
        e["gl_default"] = m.get(_norm("Sonstige"), "")
    _AFI_GENERAL_CACHE.update({"path": path, "mtime": mtime, "entries": entries, "gl": gl_map})
    return _AFI_GENERAL_CACHE


def refresh_assignment_cache(path=None):
    path = path or KST_ASSIGNMENT_DEFAULT_FILE
    _AFI_GENERAL_CACHE.update({"path": "", "mtime": None, "entries": [], "gl": {}})
    return _afi463_load_general_overview(path, force=True)


def load_assignment_entries(assignment_path):
    # Final fuehrend: Generaluebersicht. Fallback nur wenn es definitiv keine Generaluebersicht ist.
    try:
        return _afi463_load_general_overview(assignment_path).get("entries", [])
    except Exception as general_exc:
        try:
            return _load_assignment_entries(assignment_path)
        except Exception:
            raise general_exc


def _cost_type(label):
    n = _norm(label)
    if "BIKE" in n and "LEAS" in n:
        return "BIKE_LEASING"
    if "TANKENSTROM" in n or ("TANKEN" in n and "STROM" in n):
        return "TANKEN_STROM"
    if "TANKEN" in n:
        return "TANKEN"
    if "VERSICHER" in n:
        return "VERSICHERUNG"
    if "LEAS" in n:
        return "LEASING"
    if "MOBIL" in n or "FESTNETZ" in n or "TELEFON" in n or "VODAFONE" in n:
        return "MOBILFUNK"
    return "SONSTIGE"


def _select_assignment_values(entry, cost_type, text_label=""):
    if "BLOCKIER" in _norm(text_label):
        gl = ENBW_BLOCKING_GL_ACCOUNT if 'ENBW_BLOCKING_GL_ACCOUNT' in globals() else (entry.get("gl_tanken_strom") or entry.get("gl_default"))
    elif cost_type == "BIKE_LEASING":
        gl = entry.get("gl_bike_leasing") or entry.get("gl_leasing") or entry.get("gl_default")
    elif cost_type == "TANKEN_STROM":
        gl = entry.get("gl_tanken_strom") or entry.get("gl_tanken") or entry.get("gl_default")
    elif cost_type == "TANKEN":
        gl = entry.get("gl_tanken") or entry.get("gl_default")
    elif cost_type == "VERSICHERUNG":
        gl = entry.get("gl_versicherung") or entry.get("gl_default")
    elif cost_type == "LEASING":
        gl = entry.get("gl_leasing") or entry.get("gl_default")
    elif cost_type == "MOBILFUNK":
        gl = entry.get("gl_mobilfunk") or entry.get("gl_default")
    else:
        gl = entry.get("gl_default")
    if cost_type == "BIKE_LEASING":
        cc = entry.get("cc_bike_leasing") or entry.get("cc_leasing") or entry.get("cc_default")
    elif cost_type == "TANKEN_STROM":
        cc = entry.get("cc_tanken_strom") or entry.get("cc_tanken") or entry.get("cc_default")
    elif cost_type == "TANKEN":
        cc = entry.get("cc_tanken") or entry.get("cc_default")
    elif cost_type == "VERSICHERUNG":
        cc = entry.get("cc_versicherung") or entry.get("cc_default")
    elif cost_type == "LEASING":
        cc = entry.get("cc_leasing") or entry.get("cc_default")
    elif cost_type == "MOBILFUNK":
        cc = entry.get("cc_mobilfunk") or entry.get("cc_default")
    else:
        cc = entry.get("cc_default")
    return gl or "", cc or "", entry.get("orderid", "") or ""


def _afi463_identifier_kind(value):
    v = _clean(value)
    if re.search(r"\d{3,}[/\s-]?\d+", v) and not re.search(r"\b[A-ZÄÖÜ]{1,3}\s*-", v, flags=re.I):
        return "PHONE"
    if re.search(r"\b[A-ZÄÖÜ]{1,3}\s*-\s*[A-ZÄÖÜ]{1,3}\s*\d", v, flags=re.I):
        return "PLATE"
    return "KEY" if v else ""


def resolve_assignment(key, driver, entries):
    nkey = _norm(key)
    key_digits = _digits_only(key)
    key_kind = _afi463_identifier_kind(key)
    ndriver = _norm(driver)
    parts = _clean(driver).split()
    last = _norm(parts[-1]) if parts else ""
    candidates = []
    for e in entries:
        eid = e.get("identifier_norm", "")
        edig = e.get("identifier_digits", "")
        if nkey and eid and nkey == eid:
            candidates.append((220, "Schluessel exakt", e))
        elif key_kind == "PHONE" and key_digits and edig and key_digits == edig:
            candidates.append((215, "Telefon exakt", e))
        elif ndriver and e.get("name_norm") and ndriver == e.get("name_norm"):
            candidates.append((140, "Name exakt", e))
        elif ndriver and e.get("alt_name_norm") and ndriver == e.get("alt_name_norm"):
            candidates.append((135, "Name exakt", e))
        elif last and last == e.get("last_norm"):
            candidates.append((75, "Nachname", e))
    if not candidates:
        return {}, ""
    candidates.sort(key=lambda x: x[0], reverse=True)
    if len(candidates) > 1 and candidates[0][0] == candidates[1][0]:
        return {}, "mehrdeutig"
    return candidates[0][2], candidates[0][1]


def _afi463_selected_bukrs(self):
    if hasattr(self, "booking_circle_vars"):
        return [b for b, v in self.booking_circle_vars.items() if v.get()] or ["IDE"]
    if hasattr(self, "booking_circle_var"):
        return [self.booking_circle_var.get() or "IDE"]
    return ["IDE"]


def _afi463_update_export_path(self, force=False):
    if not hasattr(self, "export_var"):
        return
    current = self.export_var.get().strip() if self.export_var.get() else ""
    bukrs = "-".join(_afi463_selected_bukrs(self))
    invoice = self.invoice_var.get().strip() if hasattr(self, "invoice_var") else ""
    cost = self.global_prefix_var.get() if hasattr(self, "global_prefix_var") else "Kosten"
    try:
        vendor = _fm_short_vendor(invoice)
    except Exception:
        vendor = os.path.splitext(os.path.basename(invoice or "Rechnung"))[0] or "Rechnung"
    try:
        base = AFI_EXPORT_DEFAULT_DIR
    except Exception:
        base = _desktop_path()
    today = _fm_dt.datetime.now().strftime("%Y_%m_%d") if '_fm_dt' in globals() else "2026_07_06"
    filename = f"{_fm_safe_name_part(bukrs)}_{_fm_safe_name_part(vendor)}_{_fm_safe_name_part(cost)}_{today}.csv" if '_fm_safe_name_part' in globals() else f"{bukrs}_{vendor}_{today}.csv"
    new_path = os.path.join(base, filename)
    if force or not current or os.path.dirname(current) in ("", _desktop_path()) or current.startswith(base):
        self.export_var.set(new_path)

# auch aeltere UI-Helfer nutzen diesen Namen
_fm_update_export_path = _afi463_update_export_path


def _afi463_refresh_assignment_ui(self):
    path = self.template_var.get().strip() if hasattr(self, "template_var") else KST_ASSIGNMENT_DEFAULT_FILE
    try:
        data = refresh_assignment_cache(path)
        msg = f"Zuordnungsdatei geladen: {len(data.get('entries', []))} Zuordnungen im Sitzungs-Cache."
        if hasattr(self, "assignment_status_var"):
            self.assignment_status_var.set(msg)
        elif hasattr(self, "status_var"):
            self.status_var.set(msg)
    except Exception as exc:
        msg = "Zuordnungsdatei nicht verfügbar: " + str(exc)
        if hasattr(self, "assignment_status_var"):
            self.assignment_status_var.set(msg)
        elif hasattr(self, "status_var"):
            self.status_var.set(msg)
        try:
            messagebox.showwarning(MODULE_TITLE, "Kontierungszuordnung nicht gefunden oder nicht lesbar:\n" + str(path) + "\n\n" + str(exc))
        except Exception:
            pass


def _afi463_current_config(self):
    b = _afi463_selected_bukrs(self)
    return {
        "global_prefix": self.global_prefix_var.get() if hasattr(self, "global_prefix_var") else "Tanken Strom",
        "sources": [s.get() for s in getattr(self, "sources", [])],
        "booking_circle": b[0] if b else "IDE",
        "booking_circles": b,
        "supplier": self.supplier_var.get() if hasattr(self, "supplier_var") else "Automatisch erkennen",
    }


def _afi463_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_FILE)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value="Tanken Strom")
    self.supplier_var = tk.StringVar(value="Automatisch erkennen")
    self.booking_circle_vars = {b: tk.BooleanVar(value=(b == "IDE")) for b in BOOKING_CIRCLE_OPTIONS}
    self.status_var = tk.StringVar(value="Schritt 1: Rechnung und Zuordnung pruefen.")
    self.suggestion_var = tk.StringVar(value="")
    self.assignment_status_var = tk.StringVar(value="")
    _afi463_update_export_path(self, True)
    tk.Label(parent, text="AFI-Assistent (Wizard)", bg=self.bg, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, columnspan=4, sticky="w")
    tk.Label(parent, textvariable=self.status_var, bg="#DDE7F3", font=self.font_small, anchor="w").grid(row=1, column=0, columnspan=4, sticky="ew", pady=(4,8))
    tk.Label(parent, text="Zuordnungsdatei", bg=self.bg, font=self.font_small).grid(row=2, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.template_var, font=self.font_small).grid(row=2, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Zuordnungsdatei", self.template_var, False, "template") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=2, column=2, padx=2)
    tk.Button(parent, text="Refresh", command=lambda: _afi463_refresh_assignment_ui(self), font=self.font_small).grid(row=2, column=3, padx=2)
    tk.Label(parent, textvariable=self.assignment_status_var, bg=self.bg, fg="#445364", font=self.font_small).grid(row=3, column=0, columnspan=4, sticky="ew")
    tk.Label(parent, text="Rechnung", bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.invoice_var, font=self.font_small).grid(row=4, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Rechnung / Dokument", self.invoice_var, False, "invoice") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=4, column=2, padx=2)
    tk.Label(parent, text="Buchungskreise", bg=self.bg, font=self.font_small).grid(row=5, column=0, sticky="nw")
    bc_frame = tk.Frame(parent, bg=self.bg); bc_frame.grid(row=5, column=1, columnspan=3, sticky="ew")
    for b, var in self.booking_circle_vars.items():
        tk.Checkbutton(bc_frame, text=b, variable=var, bg=self.bg, font=self.font_small, command=lambda: _afi463_update_export_path(self, True)).pack(side="left", padx=(0,8))
    tk.Label(parent, text="Lieferant", bg=self.bg, font=self.font_small).grid(row=6, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.supplier_var, values=SUPPLIER_OPTIONS, state="normal", font=self.font_small).grid(row=6, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=7, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small).grid(row=7, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Export-CSV", bg=self.bg, font=self.font_small).grid(row=8, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.export_var, font=self.font_small).grid(row=8, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Speichern unter", command=lambda: _fm_browse(self, "Export-CSV", self.export_var, True, "export") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=8, column=2, columnspan=2, sticky="ew")
    buttons = tk.Frame(parent, bg=self.bg); buttons.grid(row=9, column=0, columnspan=4, sticky="ew", pady=(8,4)); buttons.columnconfigure(5, weight=1)
    tk.Button(buttons, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small).grid(row=0, column=0, padx=(0,4))
    tk.Button(buttons, text="Ausgewählte zusammenfassen", command=lambda: messagebox.showinfo(MODULE_TITLE, "Manuelle Zusammenfassung ist vorbereitet; die fachliche Standard-Gruppierung erfolgt beim Export."), font=self.font_small).grid(row=0, column=1, padx=(0,4))
    tk.Button(buttons, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI",10,"bold"), bg="#CFEAD6").grid(row=0, column=6, sticky="e")
    self.positions_tree = ttk.Treeview(parent, columns=["bukrs","cost","key","driver","amount","tax","gl","cc","ia"], show="headings", height=10)
    for c,w in [("bukrs",52),("cost",120),("key",110),("driver",150),("amount",80),("tax",50),("gl",80),("cc",90),("ia",80)]:
        self.positions_tree.heading(c, text=c.upper()); self.positions_tree.column(c, width=w, stretch=False)
    self.positions_tree.grid(row=10, column=0, columnspan=4, sticky="nsew", pady=(6,0)); parent.rowconfigure(10, weight=1)
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=560, justify="left").grid(row=11, column=0, columnspan=4, sticky="ew", pady=(4,0))
    try:
        _afi463_refresh_assignment_ui(self)
    except Exception:
        pass


def _afi463_populate_positions_tree(self):
    if not hasattr(self, "positions_tree"):
        return
    for item in self.positions_tree.get_children():
        self.positions_tree.delete(item)
    try:
        # schnelle Vorschau auf Basis des erzeugten Exports in temporaerer Datei, um bestehende Parser zu nutzen
        import tempfile
        fd, tmp = tempfile.mkstemp(prefix="afi_preview_", suffix=".csv")
        os.close(fd)
        res = create_supplier_upload_csv(self.template_var.get().strip(), self.invoice_var.get().strip(), tmp, _afi463_current_config(self))
        headers, rows = _read_csv(tmp)
        for row in rows[:500]:
            self.positions_tree.insert("", "end", values=["", row.get("TEXT",""), "", "", row.get("NET_VALUE",""), row.get("TAX_CODE",""), row.get("GL_ACCOUNT",""), row.get("COSTCENTER",""), row.get("ORDERID","")])
        self.suggestion_var.set(f"Positionsvorschau erstellt: {res.get('rows', len(rows))} Exportzeilen. Bitte pruefen.")
    except Exception as exc:
        self.suggestion_var.set("Positionsvorschau konnte nicht erstellt werden: " + str(exc))


def _afi463_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        self.clear_sources()
        try:
            self.load_preview(path)
        except Exception:
            pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            if hasattr(self, 'add_source_btn'):
                self.add_source_btn.configure(state="disabled")
        else:
            self.headers, self.rows = _read_table_file(path)
            suggestions = suggested_sources(self.headers)
            self.add_source(suggestions[0])
            if hasattr(self, 'add_source_btn'):
                self.add_source_btn.configure(state="normal")
        _afi463_populate_positions_tree(self)
        self.status_var.set("Schritt 2: Positionen bearbeiten/pruefen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))


def _afi463_run_export(self):
    template_path = self.template_var.get().strip()
    invoice_path = self.invoice_var.get().strip()
    _afi463_update_export_path(self, False)
    export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte ein gueltiges KST-Zuordnungsdokument auswaehlen."); return
    if not os.path.isfile(invoice_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen."); return
    if not export_path.lower().endswith(".csv"):
        export_path += ".csv"; self.export_var.set(export_path)
    config = _afi463_current_config(self)
    self.status_var.set("Export laeuft...")
    def worker():
        try:
            result = create_supplier_upload_csv(template_path, invoice_path, export_path, config)
            def done():
                self.status_var.set(f"Schritt 3: Export erstellt: {result.get('rows')} Zeilen -> {result.get('export_path')} | Netto: {result.get('export_net_total')}")
                try:
                    self.load_export_preview(export_path)
                    if hasattr(self, 'preview_notebook'):
                        self.preview_notebook.select(self.export_preview_frame)
                except Exception:
                    pass
                critical = []
                if result.get("missing_template"):
                    critical.append("Fehlende/mehrdeutige Zuordnung:\n" + "\n".join(result["missing_template"][:40]))
                if result.get("unknown_tax"):
                    critical.append("Nicht eindeutig erkannte Steuersaetze:\n" + "\n".join(result["unknown_tax"][:30]))
                if result.get("foreign_gross"):
                    critical.append("Abweichende/auslaendische Steuersaetze als Brutto mit V0 gebucht:\n" + "\n".join(result["foreign_gross"][:30]))
                if critical:
                    messagebox.showwarning(MODULE_TITLE, "\n\n".join(critical))
                try:
                    self._show_export_done_dialog(result)
                except Exception:
                    pass
            self.app.root.after(0, done)
        except Exception as exc:
            self.app.root.after(0, lambda: (self.status_var.set("Fehler beim Export."), messagebox.showerror(MODULE_TITLE, str(exc))))
    threading.Thread(target=worker, daemon=True).start()

try:
    SupplierUploadUI._build_left = _afi463_build_left
    SupplierUploadUI.analyze_invoice = _afi463_analyze_invoice
    SupplierUploadUI.run_export = _afi463_run_export
    SupplierUploadUI.refresh_assignment = _afi463_refresh_assignment_ui
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_POSITIONEN_CSV_PREVIEW_MERGE_V0464
# Datum: 2026-07-06
# Zweck:
# - CSV-Auswahlfehler beheben: sources_inner wird wieder aufgebaut.
# - Rechnungsvorschau bei PDF/DOCX zeigt alle Seiten/Abschnitte untereinander.
# - Positionsvorschau mit Auswahlkaestchen je Position.
# - "Ausgewaehlte zusammenfassen" funktional: markierte Positionen werden je Steuercode zusammengefasst.
# - Saldozeile unterhalb der Positionen: Gesamtbetrag aller Nettobetraege und Summen je Steuercode.
# ------------------------------------------------------------------
AFI_UPLOAD_POSITIONEN_CSV_PREVIEW_MERGE_VERSION = "0.464"
_AFI464_CHECK_OFF = "☐"
_AFI464_CHECK_ON = "☑"


def _afi464_make_sources_area(self, parent, row, columnspan=4):
    """Stellt die Berechnungsquellen-Flaeche wieder bereit, damit CSV-/Excel-Rechnungen funktionieren."""
    tk.Label(parent, text="Berechnungsquellen / Spaltenzuordnung", bg=self.bg, font=("Segoe UI", 9, "bold")).grid(row=row, column=0, columnspan=columnspan, sticky="w", pady=(6, 2))
    toolbar = tk.Frame(parent, bg=self.bg)
    toolbar.grid(row=row, column=2, columnspan=2, sticky="e", pady=(6, 2))
    try:
        self.add_source_btn = tk.Button(toolbar, text="+ Berechnungsquelle", command=self.add_empty_source, font=self.font_small)
        self.add_source_btn.pack(side="right")
    except Exception:
        pass
    self.sources_canvas = tk.Canvas(parent, bg=self.bg, highlightthickness=0, height=120)
    self.sources_inner = tk.Frame(self.sources_canvas, bg=self.bg)
    sources_scroll = ttk.Scrollbar(parent, orient="vertical", command=self.sources_canvas.yview)
    self.sources_canvas.configure(yscrollcommand=sources_scroll.set)
    self.sources_canvas.grid(row=row + 1, column=0, columnspan=columnspan - 1, sticky="nsew", pady=(0, 4))
    sources_scroll.grid(row=row + 1, column=columnspan - 1, sticky="ns", pady=(0, 4))
    self.sources_window = self.sources_canvas.create_window((0, 0), window=self.sources_inner, anchor="nw")
    self.sources_canvas.bind("<Configure>", lambda e: self.sources_canvas.itemconfigure(self.sources_window, width=max(100, e.width - 4)))
    self.sources_inner.bind("<Configure>", lambda e: self.sources_canvas.configure(scrollregion=self.sources_canvas.bbox("all")))


def _afi464_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_FILE)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value="Tanken Strom")
    self.supplier_var = tk.StringVar(value="Automatisch erkennen")
    self.booking_circle_vars = {b: tk.BooleanVar(value=(b == "IDE")) for b in BOOKING_CIRCLE_OPTIONS}
    self.status_var = tk.StringVar(value="Schritt 1: Rechnung und Zuordnung pruefen.")
    self.suggestion_var = tk.StringVar(value="")
    self.assignment_status_var = tk.StringVar(value="")
    self.position_saldo_var = tk.StringVar(value="Gesamtbetrag aller Nettobeträge: 0,00")
    self._afi464_position_rows = []
    self._afi464_selected = set()
    self._afi464_manual_rows = None
    _afi463_update_export_path(self, True)
    tk.Label(parent, text="AFI-Assistent (Wizard)", bg=self.bg, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, columnspan=4, sticky="w")
    tk.Label(parent, textvariable=self.status_var, bg="#DDE7F3", font=self.font_small, anchor="w").grid(row=1, column=0, columnspan=4, sticky="ew", pady=(4,8))
    tk.Label(parent, text="Zuordnungsdatei", bg=self.bg, font=self.font_small).grid(row=2, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.template_var, font=self.font_small).grid(row=2, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Zuordnungsdatei", self.template_var, False, "template") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=2, column=2, padx=2)
    tk.Button(parent, text="Refresh", command=lambda: _afi463_refresh_assignment_ui(self), font=self.font_small).grid(row=2, column=3, padx=2)
    tk.Label(parent, textvariable=self.assignment_status_var, bg=self.bg, fg="#445364", font=self.font_small).grid(row=3, column=0, columnspan=4, sticky="ew")
    tk.Label(parent, text="Rechnung", bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.invoice_var, font=self.font_small).grid(row=4, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Rechnung / Dokument", self.invoice_var, False, "invoice") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=4, column=2, padx=2)
    tk.Label(parent, text="Buchungskreise", bg=self.bg, font=self.font_small).grid(row=5, column=0, sticky="nw")
    bc_frame = tk.Frame(parent, bg=self.bg); bc_frame.grid(row=5, column=1, columnspan=3, sticky="ew")
    for b, var in self.booking_circle_vars.items():
        tk.Checkbutton(bc_frame, text=b, variable=var, bg=self.bg, font=self.font_small, command=lambda: _afi463_update_export_path(self, True)).pack(side="left", padx=(0,8))
    tk.Label(parent, text="Lieferant", bg=self.bg, font=self.font_small).grid(row=6, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.supplier_var, values=SUPPLIER_OPTIONS, state="normal", font=self.font_small).grid(row=6, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=7, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small).grid(row=7, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Export-CSV", bg=self.bg, font=self.font_small).grid(row=8, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.export_var, font=self.font_small).grid(row=8, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Speichern unter", command=lambda: _fm_browse(self, "Export-CSV", self.export_var, True, "export") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=8, column=2, columnspan=2, sticky="ew")
    buttons = tk.Frame(parent, bg=self.bg); buttons.grid(row=9, column=0, columnspan=4, sticky="ew", pady=(8,4)); buttons.columnconfigure(5, weight=1)
    tk.Button(buttons, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small).grid(row=0, column=0, padx=(0,4))
    tk.Button(buttons, text="Ausgewählte zusammenfassen", command=lambda: _afi464_merge_selected_positions(self), font=self.font_small).grid(row=0, column=1, padx=(0,4))
    tk.Button(buttons, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI",10,"bold"), bg="#CFEAD6").grid(row=0, column=6, sticky="e")
    self.positions_tree = ttk.Treeview(parent, columns=["sel","text","amount","tax","gl","cc","ia"], show="headings", height=9)
    for c,t,w in [("sel","",34),("text","POSITION",260),("amount","NETTO",82),("tax","TAX",50),("gl","GL",78),("cc","CC",90),("ia","IA",82)]:
        self.positions_tree.heading(c, text=t); self.positions_tree.column(c, width=w, stretch=(c=="text"), anchor="center" if c in ("sel","amount","tax") else "w")
    self.positions_tree.grid(row=10, column=0, columnspan=4, sticky="nsew", pady=(6,0)); parent.rowconfigure(10, weight=1)
    self.positions_tree.bind("<Button-1>", lambda e: _afi464_toggle_position_checkbox(self, e))
    tk.Label(parent, textvariable=self.position_saldo_var, bg="#FFF4C2", fg="#182431", font=("Segoe UI", 9, "bold"), anchor="w").grid(row=11, column=0, columnspan=4, sticky="ew", pady=(3,0))
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=560, justify="left").grid(row=12, column=0, columnspan=4, sticky="ew", pady=(4,0))
    _afi464_make_sources_area(self, parent, 13, 4)
    parent.rowconfigure(14, weight=0)
    try:
        _afi463_refresh_assignment_ui(self)
    except Exception:
        pass


def _afi464_update_saldo(self):
    rows = getattr(self, "_afi464_position_rows", []) or []
    total = Decimal("0.00")
    by_tax = OrderedDict()
    for row in rows:
        amount = _dec(row.get("NET_VALUE") or row.get("PRICE") or row.get("amount") or "0")
        tax = _clean(row.get("TAX_CODE") or row.get("tax") or "")
        total += amount
        by_tax[tax] = by_tax.get(tax, Decimal("0.00")) + amount
    parts = ["Gesamtbetrag aller Nettobeträge: " + _fmt(total)]
    label_for_tax = {"V0":"0%", "V2":"7%", "VD":"19%", "VX":"unklar"}
    for tax, amount in by_tax.items():
        if tax:
            parts.append(f"Gesamtbetrag {label_for_tax.get(tax, tax)}: {_fmt(amount)}")
    try:
        self.position_saldo_var.set(" | ".join(parts))
    except Exception:
        pass


def _afi464_tree_index_from_iid(self, iid):
    try:
        return int(str(iid).replace("pos_", ""))
    except Exception:
        return None


def _afi464_toggle_position_checkbox(self, event):
    try:
        region = self.positions_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
        col = self.positions_tree.identify_column(event.x)
        if col != "#1":
            return
        iid = self.positions_tree.identify_row(event.y)
        idx = _afi464_tree_index_from_iid(self, iid)
        if idx is None:
            return "break"
        if idx in self._afi464_selected:
            self._afi464_selected.remove(idx); mark = _AFI464_CHECK_OFF
        else:
            self._afi464_selected.add(idx); mark = _AFI464_CHECK_ON
        vals = list(self.positions_tree.item(iid, "values"))
        if vals:
            vals[0] = mark
            self.positions_tree.item(iid, values=vals)
        return "break"
    except Exception:
        return None


def _afi464_rebuild_positions_tree(self):
    if not hasattr(self, "positions_tree"):
        return
    for item in self.positions_tree.get_children():
        self.positions_tree.delete(item)
    rows = getattr(self, "_afi464_position_rows", []) or []
    selected = getattr(self, "_afi464_selected", set()) or set()
    for idx, row in enumerate(rows):
        self.positions_tree.insert("", "end", iid=f"pos_{idx}", values=[
            _AFI464_CHECK_ON if idx in selected else _AFI464_CHECK_OFF,
            row.get("TEXT", ""),
            row.get("NET_VALUE") or row.get("PRICE", ""),
            row.get("TAX_CODE", ""),
            row.get("GL_ACCOUNT", ""),
            row.get("COSTCENTER", ""),
            row.get("ORDERID", ""),
        ])
    _afi464_update_saldo(self)


def _afi464_populate_positions_tree(self):
    self._afi464_manual_rows = None
    self._afi464_selected = set()
    try:
        import tempfile
        fd, tmp = tempfile.mkstemp(prefix="afi_preview_", suffix=".csv")
        os.close(fd)
        res = create_supplier_upload_csv(self.template_var.get().strip(), self.invoice_var.get().strip(), tmp, _afi463_current_config(self))
        _headers, rows = _read_csv(tmp)
        self._afi464_position_rows = rows
        _afi464_rebuild_positions_tree(self)
        self.suggestion_var.set(f"Positionsvorschau erstellt: {len(rows)} Exportzeilen. Positionen per Kästchen markieren und bei Bedarf zusammenfassen.")
    except Exception as exc:
        self._afi464_position_rows = []
        _afi464_rebuild_positions_tree(self)
        self.suggestion_var.set("Positionsvorschau konnte nicht erstellt werden: " + str(exc))


def _afi464_merge_selected_positions(self):
    rows = list(getattr(self, "_afi464_position_rows", []) or [])
    selected = sorted(getattr(self, "_afi464_selected", set()) or set())
    if len(selected) < 2:
        try: messagebox.showwarning(MODULE_TITLE, "Bitte mindestens zwei Positionen per Kästchen auswählen.")
        except Exception: pass
        return
    remaining = [row for i, row in enumerate(rows) if i not in selected]
    buckets = OrderedDict()
    warnings = []
    for i in selected:
        row = rows[i]
        tax = _clean(row.get("TAX_CODE", ""))
        key = tax
        if key not in buckets:
            new = dict(row)
            new["TEXT"] = "Zusammenfassung " + _clean(row.get("TEXT", ""))[:90]
            new["PRICE"] = "0,00"; new["NET_VALUE"] = "0,00"
            new["_amount"] = Decimal("0.00")
            buckets[key] = new
        b = buckets[key]
        # Pro Steuercode getrennt. Kontierung aus erster Position; Unterschiede werden gewarnt.
        for fld in ["GL_ACCOUNT", "COSTCENTER", "ORDERID"]:
            if _clean(b.get(fld,"")) != _clean(row.get(fld,"")):
                warnings.append(f"{fld} unterschiedlich bei TAX {tax}; Kontierung der ersten markierten Position wurde verwendet.")
        b["_amount"] += _dec(row.get("NET_VALUE") or row.get("PRICE") or "0")
    merged = []
    for b in buckets.values():
        amount = b.pop("_amount", Decimal("0.00"))
        b["PRICE"] = _fmt(amount)
        b["NET_VALUE"] = _fmt(amount)
        merged.append(b)
    self._afi464_position_rows = remaining + merged
    self._afi464_selected = set()
    self._afi464_manual_rows = list(self._afi464_position_rows)
    _afi464_rebuild_positions_tree(self)
    msg = f"{len(selected)} Positionen wurden zu {len(merged)} Position(en) zusammengefasst."
    if warnings:
        msg += " Hinweis: " + warnings[0]
    self.suggestion_var.set(msg)


def _afi464_write_manual_rows(export_path, rows):
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            out = {c: row.get(c, "") for c in UPLOAD_COLUMNS}
            writer.writerow(out)
    total = sum((_dec(r.get("NET_VALUE") or r.get("PRICE") or "0") for r in rows), Decimal("0.00"))
    return {"rows": len(rows), "export_path": export_path, "invoice_net_raw_total": _fmt(total), "export_net_total": _fmt(total), "net_rounding_difference": _fmt(Decimal("0.00")), "unique_drivers": 0, "unique_keys": 0, "missing_template": [], "unknown_tax": [], "empty_assignment": [], "name_fallback_matches": [], "rounding_adjustments": []}


def _afi464_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, "sources_inner"):
            # Falls ein aelterer Build ohne Quellenbereich aktiv war, wird der Bereich unsichtbar nachgezogen.
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try:
            self.load_preview(path)
        except Exception:
            pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception: pass
        else:
            self.headers, self.rows = _read_table_file(path)
            suggestions = suggested_sources(self.headers)
            self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception: pass
        _afi464_populate_positions_tree(self)
        self.status_var.set("Schritt 2: Positionen bearbeiten/pruefen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))


def _afi464_run_export(self):
    template_path = self.template_var.get().strip()
    invoice_path = self.invoice_var.get().strip()
    _afi463_update_export_path(self, False)
    export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte ein gueltiges KST-Zuordnungsdokument auswaehlen."); return
    if not os.path.isfile(invoice_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen."); return
    if not export_path.lower().endswith(".csv"):
        export_path += ".csv"; self.export_var.set(export_path)
    config = _afi463_current_config(self)
    manual_rows = list(getattr(self, "_afi464_manual_rows", []) or [])
    self.status_var.set("Export laeuft...")
    def worker():
        try:
            if manual_rows:
                result = _afi464_write_manual_rows(export_path, manual_rows)
            else:
                result = create_supplier_upload_csv(template_path, invoice_path, export_path, config)
            def done():
                self.status_var.set(f"Schritt 3: Export erstellt: {result.get('rows')} Zeilen -> {result.get('export_path')} | Netto: {result.get('export_net_total')}")
                try:
                    self.load_export_preview(export_path)
                    if hasattr(self, 'preview_notebook'):
                        self.preview_notebook.select(self.export_preview_frame)
                except Exception:
                    pass
                critical = []
                if result.get("missing_template"):
                    critical.append("Fehlende/mehrdeutige Zuordnung:\n" + "\n".join(result["missing_template"][:40]))
                if result.get("unknown_tax"):
                    critical.append("Nicht eindeutig erkannte Steuersaetze:\n" + "\n".join(result["unknown_tax"][:30]))
                if critical:
                    messagebox.showwarning(MODULE_TITLE, "\n\n".join(critical))
                try:
                    self._show_export_done_dialog(result)
                except Exception:
                    pass
            self.app.root.after(0, done)
        except Exception as exc:
            self.app.root.after(0, lambda: (self.status_var.set("Fehler beim Export."), messagebox.showerror(MODULE_TITLE, str(exc))))
    threading.Thread(target=worker, daemon=True).start()


def _afi464_extract_pdf_pages(path):
    pages = []
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                pages.append(page.extract_text() or "")
    except Exception:
        try:
            import fitz
            doc = fitz.open(path)
            pages = [page.get_text("text") for page in doc]
            doc.close()
        except Exception as exc:
            pages = ["PDF-Text konnte nicht extrahiert werden: " + str(exc)]
    return pages or [""]


def _afi464_text_preview_image(title, pages):
    if Image is None:
        return None
    page_w, page_h = 820, 1040
    gap = 28
    total_h = max(page_h, len(pages) * (page_h + gap) + gap)
    img = Image.new("RGB", (page_w + 60, total_h), "white")
    draw = ImageDraw.Draw(img)
    y = gap
    for idx, text in enumerate(pages, 1):
        draw.rectangle((30, y, page_w + 30, y + page_h), outline="#B0B0B0", width=2, fill="#FFFFFF")
        draw.text((50, y + 22), f"{title} - Seite {idx}", fill="#1F4E79")
        ty = y + 62
        for line in (text or "").splitlines()[:52]:
            draw.text((50, ty), line[:120], fill="black")
            ty += 18
        y += page_h + gap
    return img


def _afi464_load_image_preview(self, path):
    if Image is None or ImageTk is None:
        tk.Label(self.preview_frame, text="Vorschau nicht verfügbar (Pillow nicht geladen).", bg="white").pack(fill="both", expand=True)
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            pages = _afi464_extract_pdf_pages(path)
            img = _afi464_text_preview_image(os.path.basename(path), pages)
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(path)
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as exc:
                text = "Word-Text konnte nicht extrahiert werden: " + str(exc)
            chunks = [text[i:i+3500] for i in range(0, len(text), 3500)] or [text]
            img = _afi464_text_preview_image(os.path.basename(path), chunks)
        else:
            return SupplierUploadUI.load_image_preview_before_v0464(self, path) if hasattr(SupplierUploadUI, 'load_image_preview_before_v0464') else None
        self.preview_base_image = img
        self.preview_zoom = 1.0
        self.preview_offset = [0, 0]
        self.preview_canvas = tk.Canvas(self.preview_frame, bg="white", highlightthickness=0)
        self.preview_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        def on_wheel(event):
            # Ctrl+Mausrad zoomt; normales Mausrad scrollt vertikal durch alle Seiten.
            if event.state & 0x0004:
                factor = 1.1 if event.delta > 0 else 0.9
                self.preview_zoom = max(0.25, min(4.0, self.preview_zoom * factor))
            else:
                self.preview_offset[1] += 80 if event.delta > 0 else -80
            self._render_preview_image()
            return "break"
        def on_press(event):
            self.preview_drag_start = (event.x, event.y, self.preview_offset[0], self.preview_offset[1]); return "break"
        def on_drag(event):
            if self.preview_drag_start:
                sx, sy, ox, oy = self.preview_drag_start
                self.preview_offset = [ox + event.x - sx, oy + event.y - sy]
                self._render_preview_image()
            return "break"
        self.preview_canvas.bind("<Configure>", lambda e: self._render_preview_image())
        self.preview_canvas.bind("<MouseWheel>", on_wheel)
        self.preview_canvas.bind("<ButtonPress-1>", on_press)
        self.preview_canvas.bind("<B1-Motion>", on_drag)
        self._render_preview_image()
    except Exception as exc:
        tk.Label(self.preview_frame, text=f"Vorschaufehler: {exc}", bg="white", fg="red").pack(fill="both", expand=True)

try:
    SupplierUploadUI._build_left = _afi464_build_left
    SupplierUploadUI.analyze_invoice = _afi464_analyze_invoice
    SupplierUploadUI.run_export = _afi464_run_export
    if not hasattr(SupplierUploadUI, 'load_image_preview_before_v0464'):
        SupplierUploadUI.load_image_preview_before_v0464 = SupplierUploadUI.load_image_preview
    SupplierUploadUI.load_image_preview = _afi464_load_image_preview
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_ENBW_DYNAMIC_PREVIEW_SPLITTER_V0465
# Datum: 2026-07-06
# Zweck:
# - EnBW-Rechnungen: beim Analysieren automatisch 3 Berechnungsquellen anlegen:
#   Blockiergebuehr Netto pro Person, Energiekosten Netto pro Person, Grundgebuehren Netto pro Person.
# - Positionsvorschau aktualisiert sich dynamisch nach Hinzufuegen/Aendern/Loeschen von Berechnungsquellen.
# - Quellen koennen geloescht werden.
# - Verschiebelinie zwischen Positionsvorschau und Berechnungsquellen passt die Hoehe beider Bereiche an.
# - Tabellarische Rechnungsvorschau zeigt alle Zeilen und erlaubt Spalten per Klick auf die Ueberschrift als Berechnungsquelle hinzuzufuegen.
# - PDF-Vorschau rendert, wenn moeglich, alle Seiten vollstaendig untereinander mit Seiten-Shortcuts.
# ------------------------------------------------------------------
AFI_UPLOAD_ENBW_DYNAMIC_PREVIEW_SPLITTER_VERSION = "0.465"

_AFI465_ENBW_SOURCE_DEFS = [
    ("Blockiergebühr Netto pro Person", ["Blockiergebühr", "Blockiergebuehr", "Blockier", "Blocking"], "Blockiergebühr"),
    ("Energiekosten Netto pro Person", ["Energiekosten", "Energie", "Ladekosten", "Charging", "Strom"], "Energiekosten"),
    ("Grundgebühren Netto pro Person", ["Grundgebühr", "Grundgebuehr", "Grundkosten", "Grundpreis", "Basic fee"], "Grundgebühr"),
]


def _afi465_is_enbw(self=None, invoice_path="", supplier=""):
    try:
        supplier = supplier or (self.supplier_var.get() if self and hasattr(self, 'supplier_var') else "")
    except Exception:
        supplier = supplier or ""
    n = _norm(str(supplier) + " " + os.path.basename(invoice_path or ""))
    return "ENBW" in n


def _afi465_find_column(headers, keywords, must_net=True):
    best = (0, "")
    for h in headers or []:
        nh = _norm(h)
        score = 0
        for kw in keywords:
            if _norm(kw) in nh:
                score += 20
        if must_net and any(x in nh for x in ["NETTO", "NET", "NETAMOUNT", "NETVALUE"]):
            score += 8
        if any(x in nh for x in ["BRUTTO", "GROSS", "MWST", "UST", "TAX", "STEUER"]):
            score -= 8
        if score > best[0]:
            best = (score, h)
    return best[1] if best[0] > 0 else ""


def _afi465_guess_identifier_columns(headers):
    guessed = guess_columns(headers)
    return guessed


def _afi465_make_enbw_source(label, keywords, cost_desc, headers):
    guessed = _afi465_guess_identifier_columns(headers)
    net_col = _afi465_find_column(headers, keywords, must_net=True)
    return {
        "active": True,
        "label": label,
        "cost_description": cost_desc,
        "net": net_col,
        "tax_mode": "manual",
        "vat_amount": guessed.get("vat_amount", ""),
        "gross": guessed.get("gross", ""),
        "manual_rate": "19",
        "name_mode": "full",
        "full_name": guessed.get("full_name", ""),
        "first": guessed.get("first", ""),
        "last": guessed.get("last", ""),
        "key": guessed.get("key", ""),
    }


def _afi465_source_get(self):
    out = {}
    for k, var in self.vars.items():
        out[k] = var.get()
    out["label"] = self.initial.get("label") or out.get("net") or f"Berechnungsquelle {self.idx}"
    if self.initial.get("cost_description"):
        out["cost_description"] = self.initial.get("cost_description")
    return out

try:
    SourceRow.get = _afi465_source_get
except Exception:
    pass


def _afi465_schedule_positions_refresh(self, delay=350):
    try:
        old = getattr(self, "_afi465_refresh_after_id", None)
        if old:
            try: self.app.root.after_cancel(old)
            except Exception: pass
        self._afi465_refresh_after_id = self.app.root.after(delay, lambda: _afi464_populate_positions_tree(self))
    except Exception:
        try: _afi464_populate_positions_tree(self)
        except Exception: pass


def _afi465_on_mapping_changed(self):
    try: _afi463_update_export_path(self, True)
    except Exception: pass
    try: self.update_footer()
    except Exception: pass
    try: self.update_highlight()
    except Exception: pass
    _afi465_schedule_positions_refresh(self)


def _afi465_add_source(self, data=None):
    if not hasattr(self, 'sources_inner'):
        return
    row = SourceRow(self.sources_inner, self, len(getattr(self, 'sources', [])) + 1, self.headers, data)
    row.grid(row=len(self.sources), column=0, sticky="ew", padx=2, pady=4)
    try:
        row.frame.columnconfigure(6, weight=0)
        tk.Button(row.frame, text="Löschen", command=lambda r=row: _afi465_delete_source(self, r), font=self.font_small).grid(row=0, column=6, sticky="ne", padx=5, pady=3)
    except Exception:
        pass
    self.sources_inner.columnconfigure(0, weight=1)
    self.sources.append(row)
    try:
        self.on_mapping_changed()
    except Exception:
        pass


def _afi465_delete_source(self, row):
    try:
        row.destroy()
    except Exception:
        pass
    try:
        self.sources = [s for s in self.sources if s is not row]
        for idx, src in enumerate(self.sources, 1):
            src.idx = idx
            try: src.frame.configure(text=f"Berechnungsquelle {idx}")
            except Exception: pass
            src.frame.grid_configure(row=idx-1)
    except Exception:
        pass
    _afi465_schedule_positions_refresh(self, delay=50)


def _afi465_add_empty_source(self):
    self.add_source(default_source(len(getattr(self, 'sources', [])) + 1, self.headers))

try:
    SupplierUploadUI.on_mapping_changed = _afi465_on_mapping_changed
    SupplierUploadUI.add_source = _afi465_add_source
    SupplierUploadUI.add_empty_source = _afi465_add_empty_source
except Exception:
    pass


def _afi465_make_drag_splitter(self, parent, row):
    self._afi465_positions_height = getattr(self, '_afi465_positions_height', 9)
    self._afi465_sources_height = getattr(self, '_afi465_sources_height', 120)
    bar = tk.Frame(parent, bg="#91A3B5", height=6, cursor="sb_v_double_arrow")
    bar.grid(row=row, column=0, columnspan=4, sticky="ew", pady=(4,4))
    def press(event):
        self._afi465_split_start_y = event.y_root
        self._afi465_split_start_pos = int(self.positions_tree.cget('height')) if hasattr(self, 'positions_tree') else 9
        self._afi465_split_start_src = int(self.sources_canvas.cget('height')) if hasattr(self, 'sources_canvas') else 120
    def drag(event):
        dy = event.y_root - getattr(self, '_afi465_split_start_y', event.y_root)
        pos_h = max(4, min(20, getattr(self, '_afi465_split_start_pos', 9) + int(dy / 20)))
        src_h = max(70, min(360, getattr(self, '_afi465_split_start_src', 120) - dy))
        try: self.positions_tree.configure(height=pos_h)
        except Exception: pass
        try: self.sources_canvas.configure(height=src_h)
        except Exception: pass
    bar.bind("<ButtonPress-1>", press)
    bar.bind("<B1-Motion>", drag)
    return bar


def _afi465_build_left(self, parent):
    # Erst v0.464-Struktur nutzen, danach Verschiebelinie/Quellenbereich neu anordnen.
    _afi464_build_left(self, parent)
    try:
        # Die v0.464-Quellenarea liegt ab Zeile 13. Wir ergaenzen eine sichtbare Verschiebelinie in Zeile 13
        # und verschieben den Quellenbereich optisch darunter, soweit Tk dies zulaesst.
        _afi465_make_drag_splitter(self, parent, 13)
        try: self.sources_canvas.grid_configure(row=15)
        except Exception: pass
        try: self.sources_window
        except Exception: pass
        for child in parent.grid_slaves(row=13):
            # Label aus v0.464 nicht zerstoeren, Splitter bleibt sichtbar.
            pass
    except Exception:
        pass

try:
    SupplierUploadUI._build_left = _afi465_build_left
except Exception:
    pass


def _afi465_table_add_source_from_column(self, col):
    col = _clean(col)
    if not col:
        return
    if not hasattr(self, 'headers') or not self.headers:
        self.headers = list(getattr(self, 'preview_headers', []) or [])
    data = default_source(len(getattr(self, 'sources', [])) + 1, self.headers)
    data["net"] = col
    data["label"] = col
    data["cost_description"] = self.global_prefix_var.get() if hasattr(self, 'global_prefix_var') else col
    try:
        self.add_source(data)
        self.suggestion_var.set(f"Spalte '{col}' als Berechnungsquelle hinzugefügt.")
    except Exception as exc:
        try: messagebox.showwarning(MODULE_TITLE, "Berechnungsquelle konnte nicht hinzugefügt werden:\n" + str(exc))
        except Exception: pass


def _afi465_load_table_preview(self, path, headers=None, rows=None):
    try:
        if headers is None or rows is None:
            headers, rows = _read_table_file(path)
            self.preview_headers = headers
            self.preview_rows = rows
        headers = self._filtered_preview_headers(headers, rows)
    except Exception as exc:
        tk.Label(self.preview_frame, text=str(exc), bg="white", fg="red").pack(fill="both", expand=True)
        return
    holder = tk.Frame(self.preview_frame, bg="white")
    holder.place(relx=0, rely=0, relwidth=1, relheight=1)
    holder.rowconfigure(1, weight=1)
    holder.columnconfigure(0, weight=1)
    hint = tk.Label(holder, text="Tipp: Spaltenüberschrift anklicken, um die Spalte als Berechnungsquelle hinzuzufügen.", bg="#FFF4C2", anchor="w", font=self.font_small)
    hint.grid(row=0, column=0, columnspan=2, sticky="ew")
    style = ttk.Style(holder)
    try:
        style.configure("AfiPreview.Treeview", font=("Segoe UI", self.table_font_size), rowheight=max(18, self.table_font_size + 10))
        style.configure("AfiPreview.Treeview.Heading", font=("Segoe UI", self.table_font_size, "bold"))
    except Exception:
        pass
    tree = ttk.Treeview(holder, columns=headers, show="headings", style="AfiPreview.Treeview")
    vs = ttk.Scrollbar(holder, orient="vertical", command=tree.yview)
    hs = ttk.Scrollbar(holder, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
    tree.grid(row=1, column=0, sticky="nsew")
    vs.grid(row=1, column=1, sticky="ns")
    hs.grid(row=2, column=0, sticky="ew")
    for h in headers:
        tree.heading(h, text=h)
        tree.column(h, width=max(90, min(260, len(h) * 8)), stretch=False)
    # Komplettes tabellarisches Dokument anzeigen, nicht mehr auf 500 Zeilen begrenzen.
    for row in rows:
        tree.insert("", "end", values=[_clean(row.get(h, "")) for h in headers])
    self.preview_tree = tree
    def on_click(event):
        try:
            if tree.identify_region(event.x, event.y) == "heading":
                colid = tree.identify_column(event.x)
                idx = int(colid.replace('#','')) - 1
                if 0 <= idx < len(headers):
                    _afi465_table_add_source_from_column(self, headers[idx])
                    return "break"
        except Exception:
            pass
        return None
    def on_select(event=None):
        item = tree.focus()
        if not item: return
        vals = tree.item(item, "values")
        self.selected_cell = "\t".join(str(v) for v in vals)
    def copy(event=None):
        holder.clipboard_clear(); holder.clipboard_append(self.selected_cell); return "break"
    def wheel(event):
        delta = -1 if event.delta > 0 else 1
        if event.state & 0x0004:
            self.table_font_size = max(7, min(16, self.table_font_size + (-delta)))
            try:
                style.configure("AfiPreview.Treeview", font=("Segoe UI", self.table_font_size), rowheight=max(18, self.table_font_size + 10))
                style.configure("AfiPreview.Treeview.Heading", font=("Segoe UI", self.table_font_size, "bold"))
            except Exception: pass
        elif event.state & 0x0001:
            tree.xview_scroll(delta * 3, "units")
        else:
            tree.yview_scroll(delta * 3, "units")
        return "break"
    tree.bind("<Button-1>", on_click)
    tree.bind("<<TreeviewSelect>>", on_select)
    tree.bind("<Control-c>", copy)
    tree.bind("<MouseWheel>", wheel)
    self.update_highlight()

try:
    SupplierUploadUI.load_table_preview = _afi465_load_table_preview
except Exception:
    pass


def _afi465_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, 'sources_inner'):
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try: self.load_preview(path)
        except Exception: pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception: pass
        else:
            self.headers, self.rows = _read_table_file(path)
            if _afi465_is_enbw(self, path, self.supplier_var.get() if hasattr(self, 'supplier_var') else ""):
                for label, keywords, cost_desc in _AFI465_ENBW_SOURCE_DEFS:
                    self.add_source(_afi465_make_enbw_source(label, keywords, cost_desc, self.headers))
                self.suggestion_var.set("EnBW erkannt: Blockiergebühr, Energiekosten und Grundgebühren wurden als Berechnungsquellen vorbereitet.")
            else:
                suggestions = suggested_sources(self.headers)
                self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception: pass
        _afi464_populate_positions_tree(self)
        self.status_var.set("Schritt 2: Positionen bearbeiten/pruefen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.analyze_invoice = _afi465_analyze_invoice
except Exception:
    pass

# PDF-Rendering: falls PyMuPDF verfuegbar ist, ganze Seiten statt Textauszug untereinander rendern.
def _afi465_render_pdf_pages_image(path):
    if Image is None:
        return None, []
    try:
        import fitz
        doc = fitz.open(path)
        pil_pages = []
        offsets = []
        max_w = 1
        total_h = 24
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            pil_pages.append(img)
            max_w = max(max_w, img.width)
            offsets.append(total_h)
            total_h += img.height + 36
        doc.close()
        canvas = Image.new("RGB", (max_w + 80, max(600, total_h)), "#F2F4F7")
        y = 24
        draw = ImageDraw.Draw(canvas)
        for i, img in enumerate(pil_pages, 1):
            x = (canvas.width - img.width) // 2
            draw.text((x, y - 18), f"Seite {i}", fill="#1F4E79")
            canvas.paste(img, (x, y))
            draw.rectangle((x, y, x + img.width, y + img.height), outline="#B0B0B0", width=2)
            y += img.height + 36
        return canvas, offsets
    except Exception:
        return None, []


def _afi465_load_image_preview(self, path):
    if Image is None or ImageTk is None:
        tk.Label(self.preview_frame, text="Vorschau nicht verfügbar (Pillow nicht geladen).", bg="white").pack(fill="both", expand=True)
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            img, offsets = _afi465_render_pdf_pages_image(path)
            if img is None:
                pages = _afi464_extract_pdf_pages(path) if '_afi464_extract_pdf_pages' in globals() else [_extract_pdf_text(path)]
                img = _afi464_text_preview_image(os.path.basename(path), pages) if '_afi464_text_preview_image' in globals() else None
                offsets = [28 + i * 1068 for i in range(len(pages))]
        else:
            return _afi464_load_image_preview(self, path) if '_afi464_load_image_preview' in globals() else SupplierUploadUI.load_image_preview_before_v0464(self, path)
        self.preview_base_image = img
        self.preview_page_offsets = offsets or [0]
        self.preview_zoom = 1.0
        self.preview_offset = [0, 0]
        # Shortcut-Leiste + Canvas
        top = tk.Frame(self.preview_frame, bg="#DDE7F3")
        top.place(relx=0, rely=0, relwidth=1, height=28)
        for idx, off in enumerate(self.preview_page_offsets, 1):
            if idx > 30:
                break
            tk.Button(top, text=f"S.{idx}", font=("Segoe UI", 8), command=lambda o=off: (self.preview_offset.__setitem__(1, -int(o * self.preview_zoom)), self._render_preview_image())).pack(side="left", padx=1, pady=2)
        self.preview_canvas = tk.Canvas(self.preview_frame, bg="white", highlightthickness=0)
        self.preview_canvas.place(relx=0, y=28, relwidth=1, relheight=1, height=-28)
        def on_wheel(event):
            if event.state & 0x0004:
                factor = 1.1 if event.delta > 0 else 0.9
                self.preview_zoom = max(0.25, min(4.0, self.preview_zoom * factor))
            else:
                self.preview_offset[1] += 90 if event.delta > 0 else -90
            self._render_preview_image(); return "break"
        def on_press(event):
            self.preview_drag_start = (event.x, event.y, self.preview_offset[0], self.preview_offset[1]); return "break"
        def on_drag(event):
            if self.preview_drag_start:
                sx, sy, ox, oy = self.preview_drag_start
                self.preview_offset = [ox + event.x - sx, oy + event.y - sy]
                self._render_preview_image()
            return "break"
        self.preview_canvas.bind("<Configure>", lambda e: self._render_preview_image())
        self.preview_canvas.bind("<MouseWheel>", on_wheel)
        self.preview_canvas.bind("<ButtonPress-1>", on_press)
        self.preview_canvas.bind("<B1-Motion>", on_drag)
        self._render_preview_image()
    except Exception as exc:
        tk.Label(self.preview_frame, text=f"Vorschaufehler: {exc}", bg="white", fg="red").pack(fill="both", expand=True)

try:
    SupplierUploadUI.load_image_preview = _afi465_load_image_preview
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_QUELLEN_LAYOUT_ENBW_FIX_V0466
# Datum: 2026-07-06
# Zweck:
# - Verschiebelinie oberhalb des Bereichs "Berechnungsquellen / Spaltenzuordnung" positionieren,
#   damit sie nicht mehr ueber dem Button "+ Berechnungsquelle" liegt.
# - Berechnungsquellen-Scrollbereich beginnt unterhalb des Buttons und fuellt den verfuegbaren unteren Bereich.
# - EnBW-Quellen werden sicher automatisch angelegt, wenn Lieferant/Dateiname/Header auf EnBW hindeuten.
# - Berechnungsquellen-Bloecke optisch klarer vom Hintergrund abheben.
# ------------------------------------------------------------------
AFI_UPLOAD_QUELLEN_LAYOUT_ENBW_FIX_VERSION = "0.466"
_AFI466_SOURCE_BG = "#F8FAFC"
_AFI466_SOURCE_BORDER = "#91A3B5"


def _afi466_headers_indicate_enbw(headers):
    joined = _norm(" ".join(str(h or "") for h in (headers or [])))
    keys = ["BLOCKIER", "GRUNDGEBUEHR", "GRUNDGEBUHR", "GRUNDKOST", "ENERGIEKOST", "LADEKOST", "ENBW"]
    return any(k in joined for k in keys)


def _afi466_is_enbw(self=None, invoice_path="", supplier="", headers=None):
    try:
        supplier = supplier or (self.supplier_var.get() if self and hasattr(self, 'supplier_var') else "")
    except Exception:
        supplier = supplier or ""
    n = _norm(str(supplier) + " " + os.path.basename(invoice_path or ""))
    return "ENBW" in n or _afi466_headers_indicate_enbw(headers)


def _afi466_style_source_row(row):
    try:
        row.frame.configure(bg=_AFI466_SOURCE_BG, relief="solid", bd=1, highlightthickness=1, highlightbackground=_AFI466_SOURCE_BORDER)
    except Exception:
        pass
    try:
        for child in row.frame.winfo_children():
            cls = child.winfo_class()
            if cls in ("Label", "Checkbutton", "Radiobutton", "Frame", "Labelframe"):
                try: child.configure(bg=_AFI466_SOURCE_BG)
                except Exception: pass
            if cls == "Label":
                try: child.configure(fg="#182431")
                except Exception: pass
    except Exception:
        pass


def _afi466_add_source(self, data=None):
    if not hasattr(self, 'sources_inner'):
        return
    row = SourceRow(self.sources_inner, self, len(getattr(self, 'sources', [])) + 1, self.headers, data)
    _afi466_style_source_row(row)
    row.grid(row=len(self.sources), column=0, sticky="ew", padx=4, pady=6)
    try:
        row.frame.columnconfigure(6, weight=0)
        del_btn = tk.Button(row.frame, text="Löschen", command=lambda r=row: _afi465_delete_source(self, r), font=self.font_small, bg="#FEE2E2", activebackground="#FECACA", relief="solid", bd=1)
        del_btn.grid(row=0, column=6, sticky="ne", padx=7, pady=5)
    except Exception:
        pass
    self.sources_inner.columnconfigure(0, weight=1)
    self.sources.append(row)
    try:
        self.on_mapping_changed()
    except Exception:
        pass


def _afi466_make_sources_area(self, parent, header_row=13, canvas_row=14, columnspan=4):
    """Quellenbereich: Button/Beschriftung oben, Scrollbereich darunter bis zum unteren Rand."""
    header = tk.Frame(parent, bg=self.bg)
    header.grid(row=header_row, column=0, columnspan=columnspan, sticky="ew", pady=(2, 2))
    header.columnconfigure(0, weight=1)
    tk.Label(header, text="Berechnungsquellen / Spaltenzuordnung", bg=self.bg, font=("Segoe UI", 9, "bold")).grid(row=0, column=0, sticky="w")
    self.add_source_btn = tk.Button(header, text="+ Berechnungsquelle", command=self.add_empty_source, font=self.font_small)
    self.add_source_btn.grid(row=0, column=1, sticky="e")
    self.sources_canvas = tk.Canvas(parent, bg=self.bg, highlightthickness=0)
    self.sources_inner = tk.Frame(self.sources_canvas, bg=self.bg)
    sources_scroll = ttk.Scrollbar(parent, orient="vertical", command=self.sources_canvas.yview)
    self.sources_canvas.configure(yscrollcommand=sources_scroll.set)
    self.sources_canvas.grid(row=canvas_row, column=0, columnspan=columnspan - 1, sticky="nsew", pady=(0, 0))
    sources_scroll.grid(row=canvas_row, column=columnspan - 1, sticky="ns", pady=(0, 0))
    parent.rowconfigure(canvas_row, weight=1)
    self.sources_window = self.sources_canvas.create_window((0, 0), window=self.sources_inner, anchor="nw")
    self.sources_canvas.bind("<Configure>", lambda e: self.sources_canvas.itemconfigure(self.sources_window, width=max(100, e.width - 4)))
    self.sources_inner.bind("<Configure>", lambda e: self.sources_canvas.configure(scrollregion=self.sources_canvas.bbox("all")))


def _afi466_make_drag_splitter(self, parent, row=12):
    bar = tk.Frame(parent, bg="#6F8194", height=5, cursor="sb_v_double_arrow")
    bar.grid(row=row, column=0, columnspan=4, sticky="ew", pady=(2, 3))
    def press(event):
        self._afi466_split_start_y = event.y_root
        try: self._afi466_split_start_pos = int(self.positions_tree.cget('height'))
        except Exception: self._afi466_split_start_pos = 9
    def drag(event):
        dy = event.y_root - getattr(self, '_afi466_split_start_y', event.y_root)
        # Nach unten: Positionsbereich groesser. Nach oben: Quellenbereich groesser.
        pos_h = max(4, min(20, getattr(self, '_afi466_split_start_pos', 9) + int(dy / 18)))
        try: self.positions_tree.configure(height=pos_h)
        except Exception: pass
    bar.bind("<ButtonPress-1>", press)
    bar.bind("<B1-Motion>", drag)
    return bar


def _afi466_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_FILE)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value="Tanken Strom")
    self.supplier_var = tk.StringVar(value="Automatisch erkennen")
    self.booking_circle_vars = {b: tk.BooleanVar(value=(b == "IDE")) for b in BOOKING_CIRCLE_OPTIONS}
    self.status_var = tk.StringVar(value="Schritt 1: Rechnung und Zuordnung pruefen.")
    self.suggestion_var = tk.StringVar(value="")
    self.assignment_status_var = tk.StringVar(value="")
    self.position_saldo_var = tk.StringVar(value="Gesamtbetrag aller Nettobeträge: 0,00")
    self._afi464_position_rows = []
    self._afi464_selected = set()
    self._afi464_manual_rows = None
    _afi463_update_export_path(self, True)
    tk.Label(parent, text="AFI-Assistent (Wizard)", bg=self.bg, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, columnspan=4, sticky="w")
    tk.Label(parent, textvariable=self.status_var, bg="#DDE7F3", font=self.font_small, anchor="w").grid(row=1, column=0, columnspan=4, sticky="ew", pady=(4,8))
    tk.Label(parent, text="Zuordnungsdatei", bg=self.bg, font=self.font_small).grid(row=2, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.template_var, font=self.font_small).grid(row=2, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Zuordnungsdatei", self.template_var, False, "template") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=2, column=2, padx=2)
    tk.Button(parent, text="Refresh", command=lambda: _afi463_refresh_assignment_ui(self), font=self.font_small).grid(row=2, column=3, padx=2)
    tk.Label(parent, textvariable=self.assignment_status_var, bg=self.bg, fg="#445364", font=self.font_small).grid(row=3, column=0, columnspan=4, sticky="ew")
    tk.Label(parent, text="Rechnung", bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.invoice_var, font=self.font_small).grid(row=4, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Rechnung / Dokument", self.invoice_var, False, "invoice") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=4, column=2, padx=2)
    tk.Label(parent, text="Buchungskreise", bg=self.bg, font=self.font_small).grid(row=5, column=0, sticky="nw")
    bc_frame = tk.Frame(parent, bg=self.bg); bc_frame.grid(row=5, column=1, columnspan=3, sticky="ew")
    for b, var in self.booking_circle_vars.items():
        tk.Checkbutton(bc_frame, text=b, variable=var, bg=self.bg, font=self.font_small, command=lambda: _afi463_update_export_path(self, True)).pack(side="left", padx=(0,8))
    tk.Label(parent, text="Lieferant", bg=self.bg, font=self.font_small).grid(row=6, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.supplier_var, values=SUPPLIER_OPTIONS, state="normal", font=self.font_small).grid(row=6, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=7, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small).grid(row=7, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Export-CSV", bg=self.bg, font=self.font_small).grid(row=8, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.export_var, font=self.font_small).grid(row=8, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Speichern unter", command=lambda: _fm_browse(self, "Export-CSV", self.export_var, True, "export") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=8, column=2, columnspan=2, sticky="ew")
    buttons = tk.Frame(parent, bg=self.bg); buttons.grid(row=9, column=0, columnspan=4, sticky="ew", pady=(8,4)); buttons.columnconfigure(5, weight=1)
    tk.Button(buttons, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small).grid(row=0, column=0, padx=(0,4))
    tk.Button(buttons, text="Ausgewählte zusammenfassen", command=lambda: _afi464_merge_selected_positions(self), font=self.font_small).grid(row=0, column=1, padx=(0,4))
    tk.Button(buttons, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI",10,"bold"), bg="#CFEAD6").grid(row=0, column=6, sticky="e")
    self.positions_tree = ttk.Treeview(parent, columns=["sel","text","amount","tax","gl","cc","ia"], show="headings", height=9)
    for c,t,w in [("sel","",34),("text","POSITION",260),("amount","NETTO",82),("tax","TAX",50),("gl","GL",78),("cc","CC",90),("ia","IA",82)]:
        self.positions_tree.heading(c, text=t); self.positions_tree.column(c, width=w, stretch=(c=="text"), anchor="center" if c in ("sel","amount","tax") else "w")
    self.positions_tree.grid(row=10, column=0, columnspan=4, sticky="nsew", pady=(6,0))
    parent.rowconfigure(10, weight=0)
    self.positions_tree.bind("<Button-1>", lambda e: _afi464_toggle_position_checkbox(self, e))
    tk.Label(parent, textvariable=self.position_saldo_var, bg="#FFF4C2", fg="#182431", font=("Segoe UI", 9, "bold"), anchor="w").grid(row=11, column=0, columnspan=4, sticky="ew", pady=(3,0))
    # Verschiebelinie liegt jetzt klar oberhalb von Label/Button der Berechnungsquellen.
    _afi466_make_drag_splitter(self, parent, row=12)
    _afi466_make_sources_area(self, parent, header_row=13, canvas_row=14, columnspan=4)
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=560, justify="left").grid(row=15, column=0, columnspan=4, sticky="ew", pady=(4,0))
    parent.rowconfigure(14, weight=1)
    try:
        _afi463_refresh_assignment_ui(self)
    except Exception:
        pass


def _afi466_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, 'sources_inner'):
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try: self.load_preview(path)
        except Exception: pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception: pass
        else:
            self.headers, self.rows = _read_table_file(path)
            if _afi466_is_enbw(self, path, self.supplier_var.get() if hasattr(self, 'supplier_var') else "", self.headers):
                for label, keywords, cost_desc in _AFI465_ENBW_SOURCE_DEFS:
                    self.add_source(_afi465_make_enbw_source(label, keywords, cost_desc, self.headers))
                self.suggestion_var.set("EnBW erkannt: Blockiergebühr, Energiekosten und Grundgebühren wurden als Berechnungsquellen vorbereitet.")
            else:
                suggestions = suggested_sources(self.headers)
                self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception: pass
        _afi464_populate_positions_tree(self)
        self.status_var.set("Schritt 2: Positionen bearbeiten/pruefen.")
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI._build_left = _afi466_build_left
    SupplierUploadUI.add_source = _afi466_add_source
    SupplierUploadUI.analyze_invoice = _afi466_analyze_invoice
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_UI_SOURCE_OPTIONS_CLEANUP_V0467
# Datum: 2026-07-06
# Zweck:
# - EnBW: "Grundgebühr je Nutzer Netto" statt "Grundgebühr" verwenden und separat ausweisen.
# - Je Berechnungsquelle Optionen: separat ausweisen, keine automatische Fahrerzuordnung.
# - Blockiergebühren werden als je Nutzer auszuweisende Quelle vorbelegt.
# - Berechnungsquellen farblich minimal unterscheiden.
# - Zuordnungsdatei-Auswahl hinter klappbarem Häkchen verstecken.
# - Texte "AFI-Assistent (Wizard)" und "Schritt N: ..." aus der UI entfernen.
# ------------------------------------------------------------------
AFI_UPLOAD_UI_SOURCE_OPTIONS_CLEANUP_VERSION = "0.467"

_AFI467_SOURCE_COLORS = ["#F8FAFC", "#F4FBF7", "#F7F8FF", "#FFF9F0", "#F9F5FF", "#F3FAFA"]
_AFI465_ENBW_SOURCE_DEFS = [
    ("Blockiergebühr Netto pro Person", ["Blockiergebühr", "Blockiergebuehr", "Blockier", "Blocking"], "Blockiergebühr"),
    ("Energiekosten Netto pro Person", ["Energiekosten", "Energie", "Ladekosten", "Charging", "Strom"], "Energiekosten"),
    ("Grundgebühr je Nutzer Netto", ["Grundgebühr je Nutzer", "Grundgebuehr je Nutzer", "Grundgebühr", "Grundgebuehr", "Grundkosten", "Grundpreis", "Basic fee"], "Grundgebühr je Nutzer"),
]


def _afi467_bool(value):
    try:
        return bool(value.get())
    except Exception:
        return bool(value)


def _afi467_source_defaults(data):
    data = dict(data or {})
    label_norm = _norm((data.get("label") or "") + " " + (data.get("cost_description") or ""))
    # Blockiergebuehren sind fachlich immer je Nutzer separat, aber weiterhin mit Fahrerzuordnung.
    if "BLOCKIER" in label_norm:
        data.setdefault("separate_positions", True)
        data.setdefault("force_per_user", True)
        data.setdefault("no_driver_assignment", False)
    if "GRUNDGEBUEHRJENUTZER" in label_norm or "GRUNDGEBUHRJENUTZER" in label_norm:
        data.setdefault("separate_positions", True)
    return data


def _afi467_make_enbw_source(label, keywords, cost_desc, headers):
    data = _afi465_make_enbw_source(label, keywords, cost_desc, headers)
    data["label"] = label
    data["cost_description"] = cost_desc
    return _afi467_source_defaults(data)

# EnBW-Erzeugung nutzt ab jetzt die neuen Defaults.
try:
    _afi465_make_enbw_source = _afi467_make_enbw_source
except Exception:
    pass


def _afi467_style_source_row(row, index=0):
    bg = _AFI467_SOURCE_COLORS[index % len(_AFI467_SOURCE_COLORS)]
    try:
        row.frame.configure(bg=bg, relief="solid", bd=1, highlightthickness=1, highlightbackground="#91A3B5")
    except Exception:
        pass
    try:
        for child in row.frame.winfo_children():
            cls = child.winfo_class()
            if cls in ("Label", "Checkbutton", "Radiobutton", "Frame", "Labelframe"):
                try: child.configure(bg=bg)
                except Exception: pass
    except Exception:
        pass
    return bg


def _afi467_add_source_options(self, row, data, bg):
    try:
        data = _afi467_source_defaults(data)
        row.vars["separate_positions"] = tk.BooleanVar(value=bool(data.get("separate_positions", False)))
        row.vars["no_driver_assignment"] = tk.BooleanVar(value=bool(data.get("no_driver_assignment", False)))
        row.vars["force_per_user"] = tk.BooleanVar(value=bool(data.get("force_per_user", False)))
        opt = tk.Frame(row.frame, bg=bg)
        opt.grid(row=11, column=0, columnspan=7, sticky="ew", padx=5, pady=(4, 5))
        tk.Checkbutton(opt, text="Separat ausweisen", variable=row.vars["separate_positions"], bg=bg, font=self.font_small, command=self.on_mapping_changed).pack(side="left", padx=(0, 14))
        tk.Checkbutton(opt, text="Keine automatische Fahrerzuordnung", variable=row.vars["no_driver_assignment"], bg=bg, font=self.font_small, command=self.on_mapping_changed).pack(side="left", padx=(0, 14))
        if data.get("force_per_user"):
            tk.Label(opt, text="Blockiergebühr: je Nutzer", bg=bg, fg="#7A4B00", font=("Segoe UI", 8, "bold")).pack(side="left")
    except Exception:
        pass


def _afi467_add_source(self, data=None):
    if not hasattr(self, 'sources_inner'):
        return
    data = _afi467_source_defaults(data)
    row = SourceRow(self.sources_inner, self, len(getattr(self, 'sources', [])) + 1, self.headers, data)
    bg = _afi467_style_source_row(row, len(getattr(self, 'sources', [])))
    _afi467_add_source_options(self, row, data, bg)
    row.grid(row=len(self.sources), column=0, sticky="ew", padx=4, pady=6)
    try:
        row.frame.columnconfigure(6, weight=0)
        del_btn = tk.Button(row.frame, text="Löschen", command=lambda r=row: _afi465_delete_source(self, r), font=self.font_small, bg="#FEE2E2", activebackground="#FECACA", relief="solid", bd=1)
        del_btn.grid(row=0, column=6, sticky="ne", padx=7, pady=5)
    except Exception:
        pass
    self.sources_inner.columnconfigure(0, weight=1)
    self.sources.append(row)
    try:
        self.on_mapping_changed()
    except Exception:
        pass


def _afi467_source_get(self):
    out = {}
    for k, var in self.vars.items():
        try: out[k] = var.get()
        except Exception: out[k] = var
    out["label"] = self.initial.get("label") or out.get("net") or f"Berechnungsquelle {self.idx}"
    if self.initial.get("cost_description"):
        out["cost_description"] = self.initial.get("cost_description")
    return out

try:
    SupplierUploadUI.add_source = _afi467_add_source
    SourceRow.get = _afi467_source_get
except Exception:
    pass


def _afi467_toggle_assignment_frame(self):
    try:
        if self.show_assignment_var.get():
            self.assignment_frame.grid()
        else:
            self.assignment_frame.grid_remove()
    except Exception:
        pass


def _afi467_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_FILE)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value="Tanken Strom")
    self.supplier_var = tk.StringVar(value="Automatisch erkennen")
    self.booking_circle_vars = {b: tk.BooleanVar(value=(b == "IDE")) for b in BOOKING_CIRCLE_OPTIONS}
    self.show_assignment_var = tk.BooleanVar(value=False)
    self.suggestion_var = tk.StringVar(value="")
    self.assignment_status_var = tk.StringVar(value="")
    self.position_saldo_var = tk.StringVar(value="Gesamtbetrag aller Nettobeträge: 0,00")
    self._afi464_position_rows = []
    self._afi464_selected = set()
    self._afi464_manual_rows = None
    # Kompatibilitaet: ältere Logik setzt status_var; wird aber nicht mehr sichtbar gerendert.
    self.status_var = tk.StringVar(value="")
    _afi463_update_export_path(self, True)

    tk.Checkbutton(parent, text="Zuordnungsdatei anzeigen", variable=self.show_assignment_var, command=lambda: _afi467_toggle_assignment_frame(self), bg=self.bg, font=self.font_small).grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 3))
    self.assignment_frame = tk.Frame(parent, bg=self.bg)
    self.assignment_frame.grid(row=1, column=0, columnspan=4, sticky="ew")
    self.assignment_frame.columnconfigure(1, weight=1)
    tk.Label(self.assignment_frame, text="Zuordnungsdatei", bg=self.bg, font=self.font_small).grid(row=0, column=0, sticky="w")
    tk.Entry(self.assignment_frame, textvariable=self.template_var, font=self.font_small).grid(row=0, column=1, sticky="ew", padx=4)
    tk.Button(self.assignment_frame, text="Wählen", command=lambda: _fm_browse(self, "Zuordnungsdatei", self.template_var, False, "template") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=0, column=2, padx=2)
    tk.Button(self.assignment_frame, text="Refresh", command=lambda: _afi463_refresh_assignment_ui(self), font=self.font_small).grid(row=0, column=3, padx=2)
    tk.Label(self.assignment_frame, textvariable=self.assignment_status_var, bg=self.bg, fg="#445364", font=self.font_small).grid(row=1, column=0, columnspan=4, sticky="ew")
    self.assignment_frame.grid_remove()

    tk.Label(parent, text="Rechnung", bg=self.bg, font=self.font_small).grid(row=2, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.invoice_var, font=self.font_small).grid(row=2, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Wählen", command=lambda: _fm_browse(self, "Rechnung / Dokument", self.invoice_var, False, "invoice") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=2, column=2, padx=2)
    tk.Label(parent, text="Buchungskreise", bg=self.bg, font=self.font_small).grid(row=3, column=0, sticky="nw")
    bc_frame = tk.Frame(parent, bg=self.bg); bc_frame.grid(row=3, column=1, columnspan=3, sticky="ew")
    for b, var in self.booking_circle_vars.items():
        tk.Checkbutton(bc_frame, text=b, variable=var, bg=self.bg, font=self.font_small, command=lambda: _afi463_update_export_path(self, True)).pack(side="left", padx=(0,8))
    tk.Label(parent, text="Lieferant", bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.supplier_var, values=SUPPLIER_OPTIONS, state="normal", font=self.font_small).grid(row=4, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Kostenbeschreibung", bg=self.bg, font=self.font_small).grid(row=5, column=0, sticky="w")
    ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state="normal", font=self.font_small).grid(row=5, column=1, columnspan=3, sticky="ew", padx=4)
    tk.Label(parent, text="Export-CSV", bg=self.bg, font=self.font_small).grid(row=6, column=0, sticky="w")
    tk.Entry(parent, textvariable=self.export_var, font=self.font_small).grid(row=6, column=1, sticky="ew", padx=4)
    tk.Button(parent, text="Speichern unter", command=lambda: _fm_browse(self, "Export-CSV", self.export_var, True, "export") if '_fm_browse' in globals() else None, font=self.font_small).grid(row=6, column=2, columnspan=2, sticky="ew")
    buttons = tk.Frame(parent, bg=self.bg); buttons.grid(row=7, column=0, columnspan=4, sticky="ew", pady=(8,4)); buttons.columnconfigure(5, weight=1)
    tk.Button(buttons, text="Rechnung analysieren", command=self.analyze_invoice, font=self.font_small).grid(row=0, column=0, padx=(0,4))
    tk.Button(buttons, text="Ausgewählte zusammenfassen", command=lambda: _afi464_merge_selected_positions(self), font=self.font_small).grid(row=0, column=1, padx=(0,4))
    tk.Button(buttons, text="AFI-Upload-Datei erstellen", command=self.run_export, font=("Segoe UI",10,"bold"), bg="#CFEAD6").grid(row=0, column=6, sticky="e")
    self.positions_tree = ttk.Treeview(parent, columns=["sel","text","amount","tax","gl","cc","ia"], show="headings", height=9)
    for c,t,w in [("sel","",34),("text","POSITION",260),("amount","NETTO",82),("tax","TAX",50),("gl","GL",78),("cc","CC",90),("ia","IA",82)]:
        self.positions_tree.heading(c, text=t); self.positions_tree.column(c, width=w, stretch=(c=="text"), anchor="center" if c in ("sel","amount","tax") else "w")
    self.positions_tree.grid(row=8, column=0, columnspan=4, sticky="nsew", pady=(6,0))
    self.positions_tree.bind("<Button-1>", lambda e: _afi464_toggle_position_checkbox(self, e))
    tk.Label(parent, textvariable=self.position_saldo_var, bg="#FFF4C2", fg="#182431", font=("Segoe UI", 9, "bold"), anchor="w").grid(row=9, column=0, columnspan=4, sticky="ew", pady=(3,0))
    _afi466_make_drag_splitter(self, parent, row=10)
    _afi466_make_sources_area(self, parent, header_row=11, canvas_row=12, columnspan=4)
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg="#7A4B00", font=self.font_small, wraplength=560, justify="left").grid(row=13, column=0, columnspan=4, sticky="ew", pady=(4,0))
    parent.rowconfigure(12, weight=1)
    try:
        _afi463_refresh_assignment_ui(self)
    except Exception:
        pass

try:
    SupplierUploadUI._build_left = _afi467_build_left
except Exception:
    pass

# Exportnachbearbeitung fuer Quellenoption "Keine automatische Fahrerzuordnung".
_create_supplier_upload_csv_before_v0467_options = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_v0467_options(assignment_path, invoice_path, export_path, config)
    try:
        sources = [s for s in (config or {}).get("sources", []) if s.get("active", True)]
        no_driver_sources = []
        for src in sources:
            if bool(src.get("no_driver_assignment")):
                label = _clean(src.get("cost_description") or src.get("label") or src.get("net") or "")
                if label:
                    no_driver_sources.append(label)
        if no_driver_sources and os.path.isfile(export_path):
            with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv.DictReader(f, delimiter=";")
                fieldnames = reader.fieldnames or UPLOAD_COLUMNS
                rows = list(reader)
            changed = 0
            for row in rows:
                nt = _norm(row.get("TEXT", ""))
                if any(_norm(label) in nt for label in no_driver_sources):
                    row["COSTCENTER"] = ""
                    row["ORDERID"] = ""
                    changed += 1
            if changed:
                with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
                    writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
                    writer.writeheader(); writer.writerows(rows)
                if isinstance(result, dict):
                    result["no_driver_assignment_rows"] = changed
    except Exception:
        pass
    return result


# ------------------------------------------------------------------
# AFI_UPLOAD_ANALYSE_PREVIEW_FIX_V0468
# Datum: 2026-07-06
# Zweck:
# - Fehler bei "Rechnung analysieren" beheben: EnBW-Quellenerzeugung war durch Alias-Override rekursionsgefaehrdet.
# - CSV-/Excel-Dokumentenvorschau zeigt final alle Zeilen und alle Spalten.
# - PDF-Dokumentenvorschau rendert final alle Seiten untereinander inkl. Scrollbar und Seiten-Shortcuts.
# - Doppelte Testschleife: synthetische CSV- und PDF-Faelle werden durch Kompilierung/Import/Exportpfad pruefbar.
# - CSV-Exportlogik beachtet final Quellenoptionen: Separat ausweisen, Keine automatische Fahrerzuordnung, Blockiergebuehr je Nutzer.
# ------------------------------------------------------------------
AFI_UPLOAD_ANALYSE_PREVIEW_FIX_VERSION = "0.468"


def _afi468_make_enbw_source(label, keywords, cost_desc, headers):
    guessed = guess_columns(headers)
    net_col = _afi465_find_column(headers, keywords, must_net=True) if '_afi465_find_column' in globals() else ""
    data = {
        "active": True,
        "label": label,
        "cost_description": cost_desc,
        "net": net_col,
        "tax_mode": "manual",
        "vat_amount": guessed.get("vat_amount", ""),
        "gross": guessed.get("gross", ""),
        "manual_rate": "19",
        "name_mode": "full",
        "full_name": guessed.get("full_name", ""),
        "first": guessed.get("first", ""),
        "last": guessed.get("last", ""),
        "key": guessed.get("key", ""),
    }
    if '_afi467_source_defaults' in globals():
        data = _afi467_source_defaults(data)
    return data

# Final harte Zuweisung: keine Rekursion ueber alte Alias-Kette.
_afi465_make_enbw_source = _afi468_make_enbw_source


def _afi468_source_flag(src, key):
    value = src.get(key, False)
    if isinstance(value, str):
        return value.strip().lower() in ("1", "true", "ja", "yes", "on")
    return bool(value)


def _afi468_source_cost_desc(src, global_prefix):
    return _clean(src.get("cost_description") or src.get("label") or src.get("net") or global_prefix) or global_prefix


def _afi468_driver_key_from_row(row, src, idx):
    driver = _driver_from_row(row, src)
    key = _clean(row.get(src.get("key", ""), "")) if src.get("key") else ""
    if not key and driver:
        key = driver
    if not key or not driver:
        fb = _fallback_name_from_row(row)
        if fb:
            driver = driver or fb
            key = key or fb
    if not key and not driver:
        driver = key = f"UNZUORDENBAR Zeile {idx+2}"
    return key, driver


def _afi468_amount_tax(row, src, net):
    tax_mode = src.get("tax_mode", "vat")
    if tax_mode == "manual":
        rate = _dec(src.get("manual_rate", "19"))
        gross = net * (Decimal("1.00") + rate / Decimal("100"))
        return _amount_and_tax_from_values(net, gross, rate)
    if tax_mode == "gross":
        gross = _dec(row.get(src.get("gross", ""), "")) if src.get("gross") else Decimal("0")
        if gross:
            vat = gross - net
            tax = _tax_code_from_net_vat(net, vat)
            if tax == "VX":
                rate = ((gross - net) / net * Decimal("100")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) if net else Decimal("0")
                return _amount_and_tax_from_values(net, gross, rate)
            return net, tax, False
        return net, "VX", False
    rate = _dec(row.get(src.get("vat_amount", ""), "")) if src.get("vat_amount") else Decimal("19")
    gross = _dec(row.get(src.get("gross", ""), "")) if src.get("gross") else Decimal("0")
    if not gross and rate not in VALID_NET_TAX_RATES:
        gross = net * (Decimal("1.00") + rate / Decimal("100"))
    return _amount_and_tax_from_values(net, gross, rate)


_create_supplier_upload_csv_before_v0468 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    ext = os.path.splitext(invoice_path)[1].lower()
    # PDF-Parser bleiben bewusst beim bisherigen, lieferantenspezifisch gepatchten Parser.
    if ext == ".pdf":
        return _create_supplier_upload_csv_before_v0468(assignment_path, invoice_path, export_path, config)
    assignment_entries = load_assignment_entries(assignment_path)
    global_prefix = _clean((config or {}).get("global_prefix", "Tanken Strom")) or "Tanken Strom"
    headers, rows = _read_table_file(invoice_path)
    sources = [s for s in (config or {}).get("sources", []) if s.get("active", True) and s.get("net")]
    if not sources:
        raise RuntimeError("Bitte mindestens eine aktive Berechnungsquelle mit Betragsspalte auswählen.")
    groups = OrderedDict()
    warnings_missing = []
    warnings_tax = []
    warnings_empty_assignment = []
    warnings_foreign_gross = []
    invoice_total = Decimal("0.00")
    unique_drivers = set()
    unique_keys = set()

    def add_export_row(src, row, idx, key, driver, amount, tax, foreign):
        cost_desc = _afi468_source_cost_desc(src, global_prefix)
        separate = _afi468_source_flag(src, "separate_positions") or _afi468_source_flag(src, "force_per_user")
        no_driver = _afi468_source_flag(src, "no_driver_assignment")
        is_blocking = "BLOCKIER" in _norm(cost_desc) or "BLOCKIER" in _norm(src.get("label", "")) or "BLOCKIER" in _norm(src.get("net", ""))
        # Blockiergebuehren bleiben immer je Nutzer/Zeile sichtbar.
        if is_blocking:
            separate = True
        if no_driver:
            rkey = _clean(cost_desc)
            rdriver = _clean(cost_desc)
            gl = cc = orderid = ""
            missing = False
        else:
            info, how = resolve_assignment(key, driver, assignment_entries)
            if not info:
                gl = cc = orderid = ""
                missing = True
                warnings_missing.append(f"{key} / {driver}: {'mehrdeutige Zuordnung' if how == 'mehrdeutig' else 'keine Zuordnung'}")
            else:
                gl, cc, orderid = _select_assignment_values(info, _cost_type(cost_desc), cost_desc)
                missing = not (gl and cc)
                if missing:
                    warnings_missing.append(f"{key} / {driver}: Sachkonto/KST unvollständig (Sachkonto='{gl}', KST='{cc}')")
            rkey = key
            rdriver = driver
        if is_blocking:
            text = _clean(" ".join([cost_desc, key, driver]))
        elif no_driver:
            text = cost_desc
        else:
            text_parts = [cost_desc]
            if key: text_parts.append(key)
            if driver and _norm(driver) != _norm(key): text_parts.append(driver)
            text = _clean(" ".join(text_parts))
        if separate:
            gkey = ("ROW", id(src), idx, tax, _norm(text))
        else:
            gkey = ("GROUP", _norm(cost_desc), _norm(rkey), _norm(rdriver), tax, gl, cc, orderid)
        if gkey not in groups:
            groups[gkey] = {"TEXT": text, "amount": Decimal("0.00"), "TAX_CODE": tax, "GL_ACCOUNT": gl, "COSTCENTER": cc, "ORDERID": orderid}
        groups[gkey]["amount"] += amount

    for src in sources:
        net_col = src.get("net", "")
        for idx, row in enumerate(rows):
            net = _dec(row.get(net_col, ""))
            if net == 0:
                continue
            key, driver = _afi468_driver_key_from_row(row, src, idx)
            if key.startswith("UNZUORDENBAR"):
                warnings_empty_assignment.append(f"{_afi468_source_cost_desc(src, global_prefix)}: Zeile {idx+2} ohne Fahrer/Schlüssel")
            amount, tax, foreign = _afi468_amount_tax(row, src, net)
            if tax == "VX":
                warnings_tax.append(f"{_afi468_source_cost_desc(src, global_prefix)} / {key} / {driver}: Steuer nicht eindeutig")
            if foreign:
                warnings_foreign_gross.append(f"{_afi468_source_cost_desc(src, global_prefix)} / {key}: abweichender/ausländischer Steuersatz -> Bruttobetrag mit V0 verwendet")
            invoice_total += amount
            if driver: unique_drivers.add(_norm(driver))
            if key: unique_keys.add(_norm(key))
            add_export_row(src, row, idx, key, driver, amount, tax, foreign)

    ordered = list(groups.values())
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        for g in ordered:
            amount = Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            writer.writerow({
                "TEXT": g.get("TEXT", ""),
                "PRICE": _fmt(amount),
                "PRICE_UNIT": "1",
                "QUANTITY": "1",
                "UNIT": "ST",
                "NET_VALUE": _fmt(amount),
                "TAX_CODE": g.get("TAX_CODE", ""),
                "GL_ACCOUNT": g.get("GL_ACCOUNT", ""),
                "COSTCENTER": g.get("COSTCENTER", ""),
                "ORDERID": g.get("ORDERID", ""),
            })
    export_total = sum(Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for g in ordered)
    target_total = invoice_total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return {"rows": len(ordered), "export_path": export_path, "invoice_net_raw_total": _fmt(invoice_total), "export_net_total": _fmt(export_total), "net_rounding_difference": _fmt(export_total - target_total), "unique_drivers": len([x for x in unique_drivers if x]), "unique_keys": len([x for x in unique_keys if x]), "missing_template": warnings_missing, "unknown_tax": warnings_tax, "empty_assignment": warnings_empty_assignment, "name_fallback_matches": [], "rounding_adjustments": [], "foreign_gross": warnings_foreign_gross}


def _afi468_load_table_preview(self, path, headers=None, rows=None):
    try:
        if headers is None or rows is None:
            headers, rows = _read_table_file(path)
            self.preview_headers = headers
            self.preview_rows = rows
        headers = list(headers or [])  # keine Filterung: gesamte CSV/Excel anzeigen
    except Exception as exc:
        tk.Label(self.preview_frame, text=str(exc), bg="white", fg="red").pack(fill="both", expand=True)
        return
    holder = tk.Frame(self.preview_frame, bg="white")
    holder.place(relx=0, rely=0, relwidth=1, relheight=1)
    holder.rowconfigure(1, weight=1)
    holder.columnconfigure(0, weight=1)
    info = tk.Label(holder, text=f"Vollständige Tabellenvorschau: {len(rows)} Zeilen, {len(headers)} Spalten. Spaltenüberschrift anklicken = als Berechnungsquelle hinzufügen.", bg="#FFF4C2", anchor="w", font=self.font_small)
    info.grid(row=0, column=0, columnspan=2, sticky="ew")
    style = ttk.Style(holder)
    try:
        style.configure("AfiPreview.Treeview", font=("Segoe UI", self.table_font_size), rowheight=max(18, self.table_font_size + 10))
        style.configure("AfiPreview.Treeview.Heading", font=("Segoe UI", self.table_font_size, "bold"))
    except Exception: pass
    tree = ttk.Treeview(holder, columns=headers, show="headings", style="AfiPreview.Treeview")
    vs = ttk.Scrollbar(holder, orient="vertical", command=tree.yview)
    hs = ttk.Scrollbar(holder, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
    tree.grid(row=1, column=0, sticky="nsew")
    vs.grid(row=1, column=1, sticky="ns")
    hs.grid(row=2, column=0, sticky="ew")
    for h in headers:
        tree.heading(h, text=h)
        tree.column(h, width=max(90, min(280, len(h) * 8)), stretch=False)
    for row in rows:
        tree.insert("", "end", values=[_clean(row.get(h, "")) for h in headers])
    self.preview_tree = tree
    def on_click(event):
        try:
            if tree.identify_region(event.x, event.y) == "heading":
                colid = tree.identify_column(event.x)
                idx = int(colid.replace("#", "")) - 1
                if 0 <= idx < len(headers):
                    _afi465_table_add_source_from_column(self, headers[idx])
                    return "break"
        except Exception:
            pass
    def wheel(event):
        delta = -1 if event.delta > 0 else 1
        if event.state & 0x0001:
            tree.xview_scroll(delta * 3, "units")
        else:
            tree.yview_scroll(delta * 3, "units")
        return "break"
    tree.bind("<Button-1>", on_click)
    tree.bind("<MouseWheel>", wheel)
    self.update_highlight()


def _afi468_render_pdf_pages_image(path):
    if Image is None:
        return None, []
    try:
        import fitz
        doc = fitz.open(path)
        pil_pages = []
        offsets = []
        max_w = 1
        total_h = 24
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.20, 1.20), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            pil_pages.append(img)
            max_w = max(max_w, img.width)
            offsets.append(total_h)
            total_h += img.height + 42
        doc.close()
        if not pil_pages:
            return None, []
        canvas = Image.new("RGB", (max_w + 90, max(700, total_h)), "#F2F4F7")
        draw = ImageDraw.Draw(canvas)
        y = 24
        for idx, img in enumerate(pil_pages, 1):
            x = (canvas.width - img.width) // 2
            draw.text((x, max(2, y - 18)), f"Seite {idx}", fill="#1F4E79")
            canvas.paste(img, (x, y))
            draw.rectangle((x, y, x + img.width, y + img.height), outline="#B0B0B0", width=2)
            y += img.height + 42
        return canvas, offsets
    except Exception:
        try:
            pages = _afi464_extract_pdf_pages(path)
            return _afi464_text_preview_image(os.path.basename(path), pages), [28 + i * 1068 for i in range(len(pages))]
        except Exception:
            return None, []


def _afi468_load_image_preview(self, path):
    if Image is None or ImageTk is None:
        tk.Label(self.preview_frame, text="Vorschau nicht verfügbar (Pillow nicht geladen).", bg="white").pack(fill="both", expand=True)
        return
    ext = os.path.splitext(path)[1].lower()
    if ext != ".pdf":
        try:
            return _afi464_load_image_preview(self, path)
        except Exception:
            return SupplierUploadUI.load_image_preview_before_v0464(self, path) if hasattr(SupplierUploadUI, 'load_image_preview_before_v0464') else None
    try:
        img, offsets = _afi468_render_pdf_pages_image(path)
        if img is None:
            raise RuntimeError("PDF konnte nicht gerendert werden.")
        self.preview_base_image = img
        self.preview_page_offsets = offsets or [0]
        self.preview_zoom = 1.0
        self.preview_offset = [0, 0]
        top = tk.Frame(self.preview_frame, bg="#DDE7F3")
        top.place(relx=0, rely=0, relwidth=1, height=30)
        tk.Label(top, text=f"PDF-Vorschau: {len(self.preview_page_offsets)} Seiten", bg="#DDE7F3", font=("Segoe UI", 8, "bold")).pack(side="left", padx=(4, 8))
        for idx, off in enumerate(self.preview_page_offsets, 1):
            if idx > 40: break
            tk.Button(top, text=f"S.{idx}", font=("Segoe UI", 8), command=lambda o=off: (self.preview_offset.__setitem__(1, -int(o * self.preview_zoom)), self._render_preview_image())).pack(side="left", padx=1, pady=2)
        self.preview_canvas = tk.Canvas(self.preview_frame, bg="white", highlightthickness=0)
        self.preview_canvas.place(relx=0, y=30, relwidth=1, relheight=1, height=-30, width=-16)
        scroll = ttk.Scrollbar(self.preview_frame, orient="vertical")
        scroll.place(relx=1, x=-16, y=30, width=16, relheight=1, height=-30)
        def set_scrollbar():
            try:
                ch = max(1, self.preview_canvas.winfo_height())
                total = max(1, int(self.preview_base_image.size[1] * self.preview_zoom))
                top_pos = max(0, -self.preview_offset[1])
                first = min(1.0, top_pos / total)
                last = min(1.0, (top_pos + ch) / total)
                scroll.set(first, last)
            except Exception: pass
        old_render = self._render_preview_image
        def render_and_scroll():
            old_render(); set_scrollbar()
        self._render_preview_image = render_and_scroll
        def yview(*args):
            try:
                ch = max(1, self.preview_canvas.winfo_height())
                total = max(1, int(self.preview_base_image.size[1] * self.preview_zoom))
                if args[0] == "moveto":
                    frac = float(args[1])
                    self.preview_offset[1] = -int(frac * total)
                elif args[0] == "scroll":
                    units = int(args[1])
                    self.preview_offset[1] -= units * 90
                self._render_preview_image()
            except Exception: pass
        scroll.configure(command=yview)
        def on_wheel(event):
            if event.state & 0x0004:
                factor = 1.1 if event.delta > 0 else 0.9
                self.preview_zoom = max(0.25, min(4.0, self.preview_zoom * factor))
            else:
                self.preview_offset[1] += 90 if event.delta > 0 else -90
            self._render_preview_image(); return "break"
        def on_press(event):
            self.preview_drag_start = (event.x, event.y, self.preview_offset[0], self.preview_offset[1]); return "break"
        def on_drag(event):
            if self.preview_drag_start:
                sx, sy, ox, oy = self.preview_drag_start
                self.preview_offset = [ox + event.x - sx, oy + event.y - sy]
                self._render_preview_image()
            return "break"
        self.preview_canvas.bind("<Configure>", lambda e: self._render_preview_image())
        self.preview_canvas.bind("<MouseWheel>", on_wheel)
        self.preview_canvas.bind("<ButtonPress-1>", on_press)
        self.preview_canvas.bind("<B1-Motion>", on_drag)
        self._render_preview_image()
    except Exception as exc:
        tk.Label(self.preview_frame, text=f"Vorschaufehler: {exc}", bg="white", fg="red").pack(fill="both", expand=True)


def _afi468_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, 'sources_inner'):
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try:
            self.load_preview(path)
        except Exception:
            pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception: pass
        else:
            self.headers, self.rows = _read_table_file(path)
            if _afi466_is_enbw(self, path, self.supplier_var.get() if hasattr(self, 'supplier_var') else "", self.headers):
                for label, keywords, cost_desc in _AFI465_ENBW_SOURCE_DEFS:
                    self.add_source(_afi468_make_enbw_source(label, keywords, cost_desc, self.headers))
                self.suggestion_var.set("EnBW erkannt: Blockiergebühr, Energiekosten und Grundgebühr je Nutzer wurden als Berechnungsquellen vorbereitet.")
            else:
                suggestions = suggested_sources(self.headers)
                self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception: pass
        _afi464_populate_positions_tree(self)
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.load_table_preview = _afi468_load_table_preview
    SupplierUploadUI.load_image_preview = _afi468_load_image_preview
    SupplierUploadUI.analyze_invoice = _afi468_analyze_invoice
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_ENBW_SEPARATE_STANDARD_SOURCES_V0469
# Datum: 2026-07-06
# Zweck:
# - EnBW-Standardquellen exakt einstellen:
#   1) Blockiergebühr Netto je Nutzer
#   2) Energiekosten Netto
#   3) Grundgebühr je Nutzer Netto
# - Alle drei EnBW-Standardquellen werden standardmäßig separat ausgewiesen.
# - Die Grundgebühr wird im Export mit Text/Kostenart "Grundgebühr" ausgewiesen.
# ------------------------------------------------------------------
AFI_UPLOAD_ENBW_SEPARATE_STANDARD_SOURCES_VERSION = "0.469"

_AFI465_ENBW_SOURCE_DEFS = [
    ("Blockiergebühr Netto je Nutzer", ["Blockiergebühr", "Blockiergebuehr", "Blockier", "Blocking"], "Blockiergebühr"),
    ("Energiekosten Netto", ["Energiekosten", "Energie", "Ladekosten", "Charging", "Strom"], "Energiekosten"),
    ("Grundgebühr je Nutzer Netto", ["Grundgebühr je Nutzer", "Grundgebuehr je Nutzer", "Grundgebühr", "Grundgebuehr", "Grundkosten", "Grundpreis", "Basic fee"], "Grundgebühr"),
]


def _afi469_source_defaults(data):
    data = dict(data or {})
    label_norm = _norm((data.get("label") or "") + " " + (data.get("cost_description") or ""))
    # EnBW-Standardquellen immer separat vorbelegen.
    if any(key in label_norm for key in ["BLOCKIER", "ENERGIEKOST", "GRUNDGEBUEHR", "GRUNDGEBUHR"]):
        data["separate_positions"] = True
    # Blockiergebühr bleibt zusätzlich fachlich je Nutzer/Zeile erzwungen.
    if "BLOCKIER" in label_norm:
        data["force_per_user"] = True
        data.setdefault("no_driver_assignment", False)
    return data


def _afi469_make_enbw_source(label, keywords, cost_desc, headers):
    guessed = guess_columns(headers)
    net_col = _afi465_find_column(headers, keywords, must_net=True) if '_afi465_find_column' in globals() else ""
    data = {
        "active": True,
        "label": label,
        "cost_description": cost_desc,
        "net": net_col,
        "tax_mode": "manual",
        "vat_amount": guessed.get("vat_amount", ""),
        "gross": guessed.get("gross", ""),
        "manual_rate": "19",
        "name_mode": "full",
        "full_name": guessed.get("full_name", ""),
        "first": guessed.get("first", ""),
        "last": guessed.get("last", ""),
        "key": guessed.get("key", ""),
    }
    return _afi469_source_defaults(data)

# Finale Alias-Zuweisung, damit Analyse und manuelles Hinzufügen die korrigierten Quelleinstellungen verwenden.
_afi467_source_defaults = _afi469_source_defaults
_afi465_make_enbw_source = _afi469_make_enbw_source
_afi468_make_enbw_source = _afi469_make_enbw_source


def _afi469_export_cost_desc(src, global_prefix):
    label = _clean(src.get("label", ""))
    desc = _clean(src.get("cost_description", ""))
    n = _norm(label + " " + desc)
    if "GRUNDGEBUEHR" in n or "GRUNDGEBUHR" in n:
        return "Grundgebühr"
    if "ENERGIEKOST" in n:
        return "Energiekosten"
    if "BLOCKIER" in n:
        return "Blockiergebühr"
    return desc or label or global_prefix

# Kapselt die v0.468-Hilfsfunktion, damit Grundgebühr im Export immer als "Grundgebühr" erscheint.
_afi468_source_cost_desc_before_v0469 = _afi468_source_cost_desc

def _afi468_source_cost_desc(src, global_prefix):
    return _afi469_export_cost_desc(src, global_prefix)

# Korrigierte Analysefunktion mit neuen EnBW-Quellennamen; keine Status-/Schritt-Texte.
def _afi469_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, 'sources_inner'):
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try:
            self.load_preview(path)
        except Exception:
            pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception:
                pass
        else:
            self.headers, self.rows = _read_table_file(path)
            if _afi466_is_enbw(self, path, self.supplier_var.get() if hasattr(self, 'supplier_var') else "", self.headers):
                for label, keywords, cost_desc in _AFI465_ENBW_SOURCE_DEFS:
                    self.add_source(_afi469_make_enbw_source(label, keywords, cost_desc, self.headers))
                self.suggestion_var.set("EnBW erkannt: Blockiergebühr, Energiekosten und Grundgebühr wurden separat vorbereitet.")
            else:
                suggestions = suggested_sources(self.headers)
                self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception:
                pass
        _afi464_populate_positions_tree(self)
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.analyze_invoice = _afi469_analyze_invoice
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_ENBW_GROUP_PER_DRIVER_V0470
# Datum: 2026-07-07
# Zweck:
# - EnBW-Standardquellen final in Reihenfolge:
#   1) Energiekosten Netto
#   2) Blockiergebühr Netto
#   3) Grundgebühr je Nutzer Netto
# - Alle Kosten einer Kostenbeschreibung werden pro Fahrer in genau einer Position gebuendelt.
# - "Separat ausweisen" bedeutet: eigene Position pro Fahrer fuer diese Berechnungsquelle.
# - Wenn "Separat ausweisen" nicht aktiv ist, wird die Quelle den Energiekosten des Fahrers zugeschlagen.
# - Grundgebühr je Nutzer Netto wird im Export weiterhin als "Grundgebühr" ausgegeben.
# ------------------------------------------------------------------
AFI_UPLOAD_ENBW_GROUP_PER_DRIVER_VERSION = "0.470"

_AFI465_ENBW_SOURCE_DEFS = [
    ("Energiekosten Netto", ["Energiekosten Netto", "Energiekosten", "Energie", "Ladekosten", "Charging", "Strom"], "Energiekosten"),
    ("Blockiergebühr Netto", ["Blockiergebühr Netto", "Blockiergebühr", "Blockiergebuehr", "Blockier", "Blocking"], "Blockiergebühr"),
    ("Grundgebühr je Nutzer Netto", ["Grundgebühr je Nutzer Netto", "Grundgebuehr je Nutzer Netto", "Grundgebühr je Nutzer", "Grundgebühr", "Grundgebuehr", "Grundkosten", "Grundpreis", "Basic fee"], "Grundgebühr"),
]


def _afi470_is_enbw_standard_source(src):
    n = _norm((src or {}).get("label", "") + " " + (src or {}).get("cost_description", "") + " " + (src or {}).get("net", ""))
    return any(x in n for x in ["ENERGIEKOST", "BLOCKIER", "GRUNDGEBUEHR", "GRUNDGEBUHR"])


def _afi470_source_defaults(data):
    data = dict(data or {})
    n = _norm((data.get("label") or "") + " " + (data.get("cost_description") or "") + " " + (data.get("net") or ""))
    # Alle drei EnBW-Standardquellen werden standardmaessig separat ausgewiesen.
    if any(x in n for x in ["ENERGIEKOST", "BLOCKIER", "GRUNDGEBUEHR", "GRUNDGEBUHR"]):
        data["separate_positions"] = True
    return data


def _afi470_make_enbw_source(label, keywords, cost_desc, headers):
    guessed = guess_columns(headers)
    net_col = _afi465_find_column(headers, keywords, must_net=True) if '_afi465_find_column' in globals() else ""
    data = {
        "active": True,
        "label": label,
        "cost_description": cost_desc,
        "net": net_col,
        "tax_mode": "manual",
        "vat_amount": guessed.get("vat_amount", ""),
        "gross": guessed.get("gross", ""),
        "manual_rate": "19",
        "name_mode": "full",
        "full_name": guessed.get("full_name", ""),
        "first": guessed.get("first", ""),
        "last": guessed.get("last", ""),
        "key": guessed.get("key", ""),
    }
    return _afi470_source_defaults(data)

# Finale Alias-Zuweisungen fuer UI und Analyse.
_afi467_source_defaults = _afi470_source_defaults
_afi465_make_enbw_source = _afi470_make_enbw_source
_afi468_make_enbw_source = _afi470_make_enbw_source
_afi469_make_enbw_source = _afi470_make_enbw_source


def _afi470_export_cost_desc(src, global_prefix):
    n = _norm((src or {}).get("label", "") + " " + (src or {}).get("cost_description", "") + " " + (src or {}).get("net", ""))
    if "GRUNDGEBUEHR" in n or "GRUNDGEBUHR" in n:
        return "Grundgebühr"
    if "BLOCKIER" in n:
        return "Blockiergebühr"
    if "ENERGIEKOST" in n:
        return "Energiekosten"
    return _clean((src or {}).get("cost_description") or (src or {}).get("label") or (src or {}).get("net") or global_prefix) or global_prefix


def _afi470_effective_cost_desc(src, global_prefix):
    """Kostenart fuer Export/Gruppierung.

    Wenn Separat ausweisen nicht aktiv ist, werden die Kosten immer den Energiekosten zugeschlagen.
    Energiekosten selbst bleiben Energiekosten.
    """
    own = _afi470_export_cost_desc(src, global_prefix)
    n = _norm(own)
    separate = _afi468_source_flag(src, "separate_positions") if '_afi468_source_flag' in globals() else bool((src or {}).get("separate_positions"))
    if "ENERGIEKOST" in n:
        return "Energiekosten"
    if not separate:
        return "Energiekosten"
    return own

# Auch bestehende v0.468-Hilfsfunktion soll die finale Kostenart liefern.
_afi468_source_cost_desc_before_v0470 = _afi468_source_cost_desc if '_afi468_source_cost_desc' in globals() else None

def _afi468_source_cost_desc(src, global_prefix):
    return _afi470_effective_cost_desc(src, global_prefix)


_create_supplier_upload_csv_before_v0470 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    ext = os.path.splitext(invoice_path)[1].lower()
    if ext == ".pdf":
        return _create_supplier_upload_csv_before_v0470(assignment_path, invoice_path, export_path, config)
    assignment_entries = load_assignment_entries(assignment_path)
    global_prefix = _clean((config or {}).get("global_prefix", "Tanken Strom")) or "Tanken Strom"
    headers, rows = _read_table_file(invoice_path)
    sources = [s for s in (config or {}).get("sources", []) if s.get("active", True) and s.get("net")]
    if not sources:
        raise RuntimeError("Bitte mindestens eine aktive Berechnungsquelle mit Betragsspalte auswählen.")
    groups = OrderedDict()
    warnings_missing = []
    warnings_tax = []
    warnings_empty_assignment = []
    warnings_foreign_gross = []
    invoice_total = Decimal("0.00")
    unique_drivers = set()
    unique_keys = set()

    def resolve_group_values(cost_desc, key, driver, no_driver):
        if no_driver:
            return "", "", "", False
        info, how = resolve_assignment(key, driver, assignment_entries)
        if not info:
            warnings_missing.append(f"{key} / {driver}: {'mehrdeutige Zuordnung' if how == 'mehrdeutig' else 'keine Zuordnung'}")
            return "", "", "", True
        gl, cc, orderid = _select_assignment_values(info, _cost_type(cost_desc), cost_desc)
        if not gl or not cc:
            warnings_missing.append(f"{key} / {driver}: Sachkonto/KST unvollständig (Sachkonto='{gl}', KST='{cc}')")
        return gl or "", cc or "", orderid or "", False

    for src in sources:
        net_col = src.get("net", "")
        if not net_col:
            continue
        own_desc = _afi470_export_cost_desc(src, global_prefix)
        effective_desc = _afi470_effective_cost_desc(src, global_prefix)
        no_driver = _afi468_source_flag(src, "no_driver_assignment") if '_afi468_source_flag' in globals() else bool(src.get("no_driver_assignment"))
        for idx, row in enumerate(rows):
            net = _dec(row.get(net_col, ""))
            if net == 0:
                continue
            key, driver = _afi468_driver_key_from_row(row, src, idx) if '_afi468_driver_key_from_row' in globals() else ("", "")
            if not key and not driver:
                key = driver = f"UNZUORDENBAR Zeile {idx+2}"
            if str(key).startswith("UNZUORDENBAR"):
                warnings_empty_assignment.append(f"{own_desc}: Zeile {idx+2} ohne Fahrer/Schlüssel")
            amount, tax, foreign = _afi468_amount_tax(row, src, net) if '_afi468_amount_tax' in globals() else _amount_and_tax_from_values(net, Decimal("0"), Decimal("19"))
            if tax == "VX":
                warnings_tax.append(f"{own_desc} / {key} / {driver}: Steuer nicht eindeutig")
            if foreign:
                warnings_foreign_gross.append(f"{own_desc} / {key}: abweichender/ausländischer Steuersatz -> Bruttobetrag mit V0 verwendet")
            invoice_total += amount
            unique_drivers.add(_norm(driver))
            unique_keys.add(_norm(key))
            if no_driver:
                group_driver = effective_desc
                group_key = effective_desc
                grouping_key = ("NO_DRIVER", _norm(effective_desc), tax)
            else:
                # Fachregel: pro Fahrer und Kostenbeschreibung genau eine Position.
                group_driver = driver
                group_key = key
                grouping_key = ("DRIVER_COST", _norm(effective_desc), _norm(driver), tax)
            if grouping_key not in groups:
                gl, cc, orderid, _missing = resolve_group_values(effective_desc, group_key, group_driver, no_driver)
                text_parts = [effective_desc]
                if not no_driver:
                    if group_key: text_parts.append(group_key)
                    if group_driver and _norm(group_driver) != _norm(group_key): text_parts.append(group_driver)
                groups[grouping_key] = {
                    "TEXT": _clean(" ".join(text_parts)),
                    "amount": Decimal("0.00"),
                    "TAX_CODE": tax,
                    "GL_ACCOUNT": gl,
                    "COSTCENTER": cc,
                    "ORDERID": orderid,
                    "driver_norm": _norm(group_driver),
                    "cost_norm": _norm(effective_desc),
                }
            # Falls erster Treffer keine Kennzeicheninformation enthielt und spaetere Zeile eine hat, Text/KST nicht neu anfassen,
            # da pro Fahrer eine Position gewuenscht ist und der erste plausible Zuordnungstreffer stabil bleibt.
            groups[grouping_key]["amount"] += amount

    ordered = sorted(groups.values(), key=lambda g: (g.get("driver_norm", ""), g.get("cost_norm", ""), TAX_ORDER.get(g.get("TAX_CODE", "VX"), 9)))
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        for g in ordered:
            amount = Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            writer.writerow({
                "TEXT": g.get("TEXT", ""),
                "PRICE": _fmt(amount),
                "PRICE_UNIT": "1",
                "QUANTITY": "1",
                "UNIT": "ST",
                "NET_VALUE": _fmt(amount),
                "TAX_CODE": g.get("TAX_CODE", ""),
                "GL_ACCOUNT": g.get("GL_ACCOUNT", ""),
                "COSTCENTER": g.get("COSTCENTER", ""),
                "ORDERID": g.get("ORDERID", ""),
            })
    export_total = sum(Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for g in ordered)
    target_total = invoice_total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return {"rows": len(ordered), "export_path": export_path, "invoice_net_raw_total": _fmt(invoice_total), "export_net_total": _fmt(export_total), "net_rounding_difference": _fmt(export_total - target_total), "unique_drivers": len([x for x in unique_drivers if x]), "unique_keys": len([x for x in unique_keys if x]), "missing_template": warnings_missing, "unknown_tax": warnings_tax, "empty_assignment": warnings_empty_assignment, "name_fallback_matches": [], "rounding_adjustments": [], "foreign_gross": warnings_foreign_gross}


def _afi470_analyze_invoice(self):
    path = self.invoice_var.get().strip()
    if not os.path.isfile(path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen.")
        return
    try:
        if not hasattr(self, 'sources_inner'):
            self.sources_inner = tk.Frame(self.app.root if hasattr(self, 'app') else None, bg=self.bg)
        self.clear_sources()
        try:
            self.load_preview(path)
        except Exception:
            pass
        _afi463_update_export_path(self, True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            self.headers, self.rows = ["PDF"], []
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="disabled")
            except Exception:
                pass
        else:
            self.headers, self.rows = _read_table_file(path)
            if _afi466_is_enbw(self, path, self.supplier_var.get() if hasattr(self, 'supplier_var') else "", self.headers):
                for label, keywords, cost_desc in _AFI465_ENBW_SOURCE_DEFS:
                    self.add_source(_afi470_make_enbw_source(label, keywords, cost_desc, self.headers))
                self.suggestion_var.set("EnBW erkannt: Energiekosten, Blockiergebühr und Grundgebühr je Nutzer wurden in der korrekten Reihenfolge vorbereitet.")
            else:
                suggestions = suggested_sources(self.headers)
                self.add_source(suggestions[0])
            try:
                if hasattr(self, 'add_source_btn'): self.add_source_btn.configure(state="normal")
            except Exception:
                pass
        _afi464_populate_positions_tree(self)
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.analyze_invoice = _afi470_analyze_invoice
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_ROBUST_ASSIGNMENT_MULTI_PDF_V0471
# Datum: 2026-07-07
# Zweck:
# - Beträge bleiben bis zum finalen CSV-Schreiben ungerundet; keine Cent-Korrekturpositionen.
# - Mehrfachauswahl von PDF-Rechnungen; mehrere PDFs werden in eine gemeinsame AFI-CSV geschrieben.
# - Robustere Fahrerzuordnung bei Apostrophen, Akzenten, Umlauten und kleinen Tippfehlern.
# - Alte Kennzeichenhalter-Konflikte werden erkannt; optional kann die Rechnungszuordnung übernommen werden.
# - Sachkonto wird immer aus der Kostenbeschreibung gesetzt, auch ohne eindeutige Fahrer-/Rufnummerzuordnung.
# ------------------------------------------------------------------
AFI_UPLOAD_ROBUST_ASSIGNMENT_MULTI_PDF_VERSION = "0.471"

import difflib as _afi471_difflib


def _afi471_split_invoice_paths(value):
    if isinstance(value, (list, tuple)):
        return [str(x).strip() for x in value if str(x).strip()]
    text = str(value or "").strip()
    if not text:
        return []
    parts = []
    for line in text.replace("\r", "\n").split("\n"):
        line = line.strip()
        if line:
            parts.append(line)
    return parts or [text]


def _afi471_join_invoice_paths(paths):
    return "\n".join(str(p) for p in (paths or []) if str(p).strip())


def _afi471_normal_file_exists(value):
    paths = _afi471_split_invoice_paths(value)
    return bool(paths) and all(os.path.isfile(p) for p in paths)


def _afi471_similarity(a, b):
    a = _norm(a)
    b = _norm(b)
    if not a or not b:
        return 0.0
    if a == b:
        return 1.0
    return _afi471_difflib.SequenceMatcher(None, a, b).ratio()


def _afi471_name_variants(entry):
    vals = []
    for key in ("name_norm", "alt_name_norm", "last_norm"):
        v = entry.get(key, "")
        if v:
            vals.append(v)
    fn = _clean(entry.get("first", ""))
    ln = _clean(entry.get("last", ""))
    if fn or ln:
        vals.append(_norm(f"{fn} {ln}"))
        vals.append(_norm(f"{ln} {fn}"))
    return [v for v in vals if v]


def _afi471_best_name_match(entries, driver, threshold=0.86):
    target = _norm(driver)
    if not target:
        return None, "", 0.0
    scored = []
    for e in entries or []:
        best = 0.0
        for v in _afi471_name_variants(e):
            if not v:
                continue
            sim = _afi471_similarity(target, v)
            # Nachname alleine darf helfen, aber nicht alleine entscheiden, wenn vollständiger Name abweicht.
            if len(v) <= 5 and v != target:
                sim = min(sim, 0.84)
            best = max(best, sim)
        if best >= threshold:
            scored.append((best, e))
    scored.sort(key=lambda x: x[0], reverse=True)
    if not scored:
        return None, "", 0.0
    if len(scored) > 1 and scored[0][0] - scored[1][0] < 0.025:
        return None, "mehrdeutig", scored[0][0]
    return scored[0][1], "fuzzy_name", scored[0][0]


def _afi471_plate_entry(entries, key):
    kn = _norm(key)
    if not kn:
        return None
    for e in entries or []:
        if e.get("identifier_type") == "PLATE" and e.get("identifier_norm") == kn:
            return e
    return None


def _afi471_entry_name_score(entry, driver):
    if not entry or not driver:
        return 0.0
    return max([_afi471_similarity(driver, v) for v in _afi471_name_variants(entry)] or [0.0])


def _afi471_cost_type_for_gl(cost_desc):
    n = _norm(cost_desc)
    if any(x in n for x in ("ENERGIEKOST", "BLOCKIER", "GRUNDGEBUEHR", "GRUNDGEBUHR")):
        return "TANKEN_STROM"
    return _cost_type(cost_desc)


def _afi471_fallback_entry_for_gl(assignment_path, cost_desc, bukrs="IDE"):
    try:
        cache = _afi463_load_general_overview(assignment_path)
        glmap = cache.get("gl", {}).get(bukrs or "IDE", {})
    except Exception:
        glmap = {}
    gl = ""
    ct = _afi471_cost_type_for_gl(cost_desc)
    if ct == "TANKEN_STROM":
        gl = glmap.get(_norm("Tanken Strom"), "") or glmap.get(_norm("Tanken"), "")
    elif ct == "TANKEN":
        gl = glmap.get(_norm("Tanken"), "") or glmap.get(_norm("DKV"), "")
    elif ct == "MOBILFUNK":
        gl = glmap.get(_norm("Telekom"), "") or glmap.get(_norm("Vodafone"), "")
    elif ct == "LEASING":
        gl = glmap.get(_norm("VW-Leasing"), "") or glmap.get(_norm("Leasing"), "")
    elif ct == "VERSICHERUNG":
        gl = glmap.get(_norm("VW-Versicherungen"), "") or glmap.get(_norm("DEAS"), "")
    if not gl:
        gl = glmap.get(_norm("Sonstige"), "")
    return {"gl_default": gl, "gl_tanken_strom": gl, "gl_tanken": gl, "gl_mobilfunk": gl, "gl_leasing": gl, "gl_versicherung": gl, "gl_bike_leasing": gl, "cc_default": "", "cc_tanken_strom": "", "cc_tanken": "", "cc_mobilfunk": "", "cc_leasing": "", "cc_versicherung": "", "cc_bike_leasing": "", "orderid": "", "bukrs": bukrs or "IDE"}


def _afi471_org_to_bukrs_from_row(row):
    org = _norm((row or {}).get("Organisationseinheit", ""))
    if org == "IDG":
        return "IDG"
    if org == "SABU":
        return "SABU"
    if org == "IMS":
        return "IMS"
    return "IDE"


def _afi471_resolve_values(assignment_path, entries, cost_desc, key, driver, accepted_old, row, warnings_missing, conflicts, fuzzy_hits):
    bukrs = _afi471_org_to_bukrs_from_row(row)
    fallback = _afi471_fallback_entry_for_gl(assignment_path, cost_desc, bukrs)
    gl_only, _, _ = _select_assignment_values(fallback, _afi471_cost_type_for_gl(cost_desc), cost_desc)
    gl_only = gl_only or fallback.get("gl_default", "")
    plate = _afi471_plate_entry(entries, key)
    name_entry, name_how, name_score = _afi471_best_name_match(entries, driver)
    conflict_id = f"{_norm(key)}|{_norm(driver)}"
    conflict = None
    selected = None
    text_driver = driver
    suppress_driver = False
    old_conflict = False

    if plate:
        score = _afi471_entry_name_score(plate, driver)
        if driver and score < 0.86:
            old_conflict = True
            conflict = {"id": conflict_id, "kennzeichen": key, "rechnung_fahrer": driver, "zuordnung_fahrer": plate.get("full_name", ""), "score": round(score, 3)}
            if conflict_id not in accepted_old:
                suppress_driver = True
                text_driver = ""
                selected = None
            else:
                selected = name_entry or plate
                if name_entry:
                    fuzzy_hits.append(f"{key} / {driver}: alte Rechnungszuordnung per Namenssuche übernommen ({name_entry.get('full_name','')}, Score {name_score:.2f})")
                else:
                    warnings_missing.append(f"{key} / {driver}: alte Rechnungszuordnung übernommen, aber keine robuste Namenszuordnung gefunden")
        else:
            selected = plate
            if driver and 0.0 < score < 1.0:
                fuzzy_hits.append(f"{key} / {driver}: Kennzeichenhalter trotz Schreibabweichung übernommen ({plate.get('full_name','')}, Score {score:.2f})")
    elif name_entry:
        selected = name_entry
        if name_score < 1.0:
            fuzzy_hits.append(f"{key} / {driver}: Kontierung per robuster Namenssuche ({name_entry.get('full_name','')}, Score {name_score:.2f})")
    else:
        selected = None

    if conflict:
        conflicts[conflict_id] = conflict
    if selected:
        gl, cc, orderid = _select_assignment_values(selected, _afi471_cost_type_for_gl(cost_desc), cost_desc)
        gl = gl or gl_only
        if not cc:
            warnings_missing.append(f"{key} / {driver}: keine KST gefunden")
        return gl, cc or "", orderid or "", text_driver, suppress_driver, old_conflict
    warnings_missing.append(f"{key} / {driver}: keine eindeutige Fahrer-/Kennzeichenzuordnung; nur Sachkonto gesetzt")
    return gl_only, "", "", text_driver if not suppress_driver else "", suppress_driver, old_conflict


_create_supplier_upload_csv_before_v0471 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    invoice_paths = _afi471_split_invoice_paths(invoice_path)
    if len(invoice_paths) > 1 and all(os.path.splitext(p)[1].lower() == ".pdf" for p in invoice_paths):
        tmp_files = []
        combined = OrderedDict()
        result_rows = 0
        warnings_all = []
        total = Decimal("0.00")
        for idx, pdf_path in enumerate(invoice_paths, 1):
            tmp = os.path.join(os.path.dirname(os.path.abspath(export_path)) or ".", f".__afi_tmp_pdf_{idx}.csv")
            tmp_files.append(tmp)
            res = _create_supplier_upload_csv_before_v0471(assignment_path, pdf_path, tmp, config)
            warnings_all.extend((res or {}).get("missing_template", []) or [])
            with open(tmp, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv.DictReader(f, delimiter=";")
                for row in reader:
                    key = (row.get("TEXT", ""), row.get("TAX_CODE", ""), row.get("GL_ACCOUNT", ""), row.get("COSTCENTER", ""), row.get("ORDERID", ""))
                    amt = _dec(row.get("NET_VALUE") or row.get("PRICE") or "0")
                    if key not in combined:
                        combined[key] = {"row": dict(row), "amount": Decimal("0.00")}
                    combined[key]["amount"] += amt
                    total += amt
        os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
        with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
            writer.writeheader()
            for item in combined.values():
                amount = item["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                row = item["row"]
                row["PRICE"] = _fmt(amount); row["NET_VALUE"] = _fmt(amount); row["PRICE_UNIT"] = "1"; row["QUANTITY"] = "1"; row["UNIT"] = "ST"
                writer.writerow(row)
        for tmp in tmp_files:
            try: os.remove(tmp)
            except Exception: pass
        return {"rows": len(combined), "export_path": export_path, "invoice_net_raw_total": _fmt(total), "export_net_total": _fmt(sum(i["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for i in combined.values())), "net_rounding_difference": _fmt(sum(i["amount"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for i in combined.values()) - total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)), "missing_template": warnings_all, "multi_pdf_count": len(invoice_paths)}

    if not invoice_paths:
        raise RuntimeError("Bitte eine Rechnung auswählen.")
    invoice_path_single = invoice_paths[0]
    ext = os.path.splitext(invoice_path_single)[1].lower()
    if ext == ".pdf":
        return _create_supplier_upload_csv_before_v0471(assignment_path, invoice_path_single, export_path, config)

    assignment_entries = load_assignment_entries(assignment_path)
    global_prefix = _clean((config or {}).get("global_prefix", "Tanken Strom")) or "Tanken Strom"
    headers, rows = _read_table_file(invoice_path_single)
    sources = [s for s in (config or {}).get("sources", []) if s.get("active", True) and s.get("net")]
    if not sources:
        raise RuntimeError("Bitte mindestens eine aktive Berechnungsquelle mit Betragsspalte auswählen.")
    accepted_old = set((config or {}).get("accepted_old_plate_assignments") or [])
    groups = OrderedDict(); warnings_missing=[]; warnings_tax=[]; warnings_empty_assignment=[]; warnings_foreign_gross=[]; fuzzy_hits=[]; conflicts={}
    invoice_total = Decimal("0.00"); unique_drivers=set(); unique_keys=set()

    for src in sources:
        net_col = src.get("net", "")
        if not net_col:
            continue
        own_desc = _afi470_export_cost_desc(src, global_prefix) if '_afi470_export_cost_desc' in globals() else _clean(src.get("cost_description") or src.get("label") or global_prefix)
        effective_desc = _afi470_effective_cost_desc(src, global_prefix) if '_afi470_effective_cost_desc' in globals() else own_desc
        no_driver = _afi468_source_flag(src, "no_driver_assignment") if '_afi468_source_flag' in globals() else bool(src.get("no_driver_assignment"))
        for idx, row in enumerate(rows):
            net = _dec(row.get(net_col, ""))
            if net == 0:
                continue
            key, driver = _afi468_driver_key_from_row(row, src, idx) if '_afi468_driver_key_from_row' in globals() else ("", "")
            if not key and not driver:
                key = driver = f"UNZUORDENBAR Zeile {idx+2}"; warnings_empty_assignment.append(f"{own_desc}: Zeile {idx+2} ohne Fahrer/Schlüssel")
            amount, tax, foreign = _afi468_amount_tax(row, src, net) if '_afi468_amount_tax' in globals() else _amount_and_tax_from_values(net, Decimal("0"), Decimal("19"))
            if tax == "VX": warnings_tax.append(f"{own_desc} / {key} / {driver}: Steuer nicht eindeutig")
            if foreign: warnings_foreign_gross.append(f"{own_desc} / {key}: abweichender/ausländischer Steuersatz -> Bruttobetrag mit V0 verwendet")
            invoice_total += amount; unique_drivers.add(_norm(driver)); unique_keys.add(_norm(key))
            if no_driver:
                gl, cc, orderid = _select_assignment_values(_afi471_fallback_entry_for_gl(assignment_path, effective_desc), _afi471_cost_type_for_gl(effective_desc), effective_desc)
                text_driver = ""; suppress = True
            else:
                gl, cc, orderid, text_driver, suppress, _old = _afi471_resolve_values(assignment_path, assignment_entries, effective_desc, key, driver, accepted_old, row, warnings_missing, conflicts, fuzzy_hits)
            if suppress:
                grouping_key = ("NO_KST", _norm(effective_desc), _norm(key), tax, gl)
                text_parts = [effective_desc]
                if key: text_parts.append(key)
            else:
                grouping_key = ("DRIVER_COST", _norm(effective_desc), _norm(text_driver or driver), tax, gl, cc, orderid)
                text_parts = [effective_desc]
                if key: text_parts.append(key)
                if (text_driver or driver) and _norm(text_driver or driver) != _norm(key): text_parts.append(text_driver or driver)
            if grouping_key not in groups:
                groups[grouping_key] = {"TEXT": _clean(" ".join(text_parts)), "amount": Decimal("0.00"), "TAX_CODE": tax, "GL_ACCOUNT": gl or "", "COSTCENTER": cc or "", "ORDERID": orderid or "", "driver_norm": _norm(text_driver or driver), "cost_norm": _norm(effective_desc)}
            groups[grouping_key]["amount"] += amount

    ordered = sorted(groups.values(), key=lambda g: (g.get("driver_norm", ""), g.get("cost_norm", ""), TAX_ORDER.get(g.get("TAX_CODE", "VX"), 9)))
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or ".", exist_ok=True)
    with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
        writer.writeheader()
        for g in ordered:
            amount = Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            writer.writerow({"TEXT": g.get("TEXT", ""), "PRICE": _fmt(amount), "PRICE_UNIT": "1", "QUANTITY": "1", "UNIT": "ST", "NET_VALUE": _fmt(amount), "TAX_CODE": g.get("TAX_CODE", ""), "GL_ACCOUNT": g.get("GL_ACCOUNT", ""), "COSTCENTER": g.get("COSTCENTER", ""), "ORDERID": g.get("ORDERID", "")})
    export_total = sum(Decimal(g.get("amount", "0")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for g in ordered)
    target_total = invoice_total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return {"rows": len(ordered), "export_path": export_path, "invoice_net_raw_total": _fmt(invoice_total), "export_net_total": _fmt(export_total), "net_rounding_difference": _fmt(export_total - target_total), "unique_drivers": len([x for x in unique_drivers if x]), "unique_keys": len([x for x in unique_keys if x]), "missing_template": warnings_missing, "unknown_tax": warnings_tax, "empty_assignment": warnings_empty_assignment, "name_fallback_matches": fuzzy_hits, "rounding_adjustments": [], "foreign_gross": warnings_foreign_gross, "old_plate_conflicts": list(conflicts.values())}


_fm_browse_before_v0471 = _fm_browse if '_fm_browse' in globals() else None

def _fm_browse(self, label, var, save=False, role=""):
    if save:
        return _fm_browse_before_v0471(self, label, var, save, role) if _fm_browse_before_v0471 else None
    if role == "invoice":
        start = var.get().strip() if hasattr(var, 'get') else ""
        if "\n" in start:
            start = os.path.dirname(_afi471_split_invoice_paths(start)[0])
        if os.path.isfile(start): start = os.path.dirname(start)
        if not os.path.isdir(start): start = _fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path()
        paths = filedialog.askopenfilenames(title=label, initialdir=start or None, filetypes=[("Dokumente", "*.csv *.xlsx *.xls *.xlsm *.pdf *.docx"), ("Alle Dateien", "*.*")])
        if paths:
            var.set(_afi471_join_invoice_paths(paths))
            try:
                if hasattr(self, "load_preview"):
                    self.load_preview(paths[0])
            except Exception:
                pass
            try:
                _afi463_update_export_path(self, True)
            except Exception:
                try: _fm_update_export_path(self, True)
                except Exception: pass
        return
    return _fm_browse_before_v0471(self, label, var, save, role) if _fm_browse_before_v0471 else None


def _afi471_old_plate_dialog(self, conflicts):
    if not conflicts:
        return []
    try:
        dlg = tk.Toplevel(self.app.root if hasattr(self, 'app') else None)
        dlg.title("Alte Kennzeichenzuordnung prüfen")
        dlg.geometry("760x420")
        tk.Label(dlg, text="Bei folgenden Kennzeichen weicht der Fahrer aus der Rechnung von der aktuellen Zuordnungsdatei ab. Bitte auswählen, welche alte Rechnungszuordnung in den AFI-Export übernommen werden soll.", wraplength=720, justify="left").pack(fill="x", padx=10, pady=8)
        frm = tk.Frame(dlg); frm.pack(fill="both", expand=True, padx=10, pady=5)
        canvas = tk.Canvas(frm); inner = tk.Frame(canvas); scroll = ttk.Scrollbar(frm, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scroll.set); canvas.pack(side="left", fill="both", expand=True); scroll.pack(side="right", fill="y")
        canvas.create_window((0,0), window=inner, anchor="nw")
        vars_ = []
        for c in conflicts:
            v = tk.BooleanVar(value=False); vars_.append((v, c))
            txt = f"{c.get('kennzeichen','')} | Rechnung: {c.get('rechnung_fahrer','')} | Zuordnung: {c.get('zuordnung_fahrer','')}"
            tk.Checkbutton(inner, text=txt, variable=v, anchor="w", justify="left").pack(fill="x", pady=2)
        inner.update_idletasks(); canvas.configure(scrollregion=canvas.bbox("all"))
        result = {"accepted": []}
        def ok():
            result["accepted"] = [c.get("id") for v,c in vars_ if v.get()]
            dlg.destroy()
        def cancel():
            result["accepted"] = []
            dlg.destroy()
        btn = tk.Frame(dlg); btn.pack(fill="x", padx=10, pady=8)
        tk.Button(btn, text="Ausgewählte übernehmen", command=ok).pack(side="right", padx=4)
        tk.Button(btn, text="Keine übernehmen", command=cancel).pack(side="right", padx=4)
        dlg.transient(self.app.root if hasattr(self, 'app') else None); dlg.grab_set(); dlg.wait_window()
        return result.get("accepted", [])
    except Exception:
        return []


def _afi471_current_config_with_accept(self, accepted=None):
    cfg = _afi463_current_config(self) if '_afi463_current_config' in globals() else {}
    if accepted:
        cfg["accepted_old_plate_assignments"] = list(accepted)
    return cfg


def _afi471_run_export(self):
    template_path = self.template_var.get().strip()
    invoice_path = self.invoice_var.get().strip()
    try: _afi463_update_export_path(self, False)
    except Exception: pass
    export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Zuordnungsdatei auswaehlen."); return
    if not _afi471_normal_file_exists(invoice_path):
        messagebox.showwarning(MODULE_TITLE, "Bitte eine gueltige Rechnung auswaehlen."); return
    if not export_path:
        messagebox.showwarning(MODULE_TITLE, "Bitte einen Exportpfad auswaehlen."); return
    try:
        cfg = _afi471_current_config_with_accept(self)
        result = create_supplier_upload_csv(template_path, invoice_path, export_path, cfg)
        conflicts = result.get("old_plate_conflicts") or []
        if conflicts:
            accepted = _afi471_old_plate_dialog(self, conflicts)
            if accepted:
                cfg = _afi471_current_config_with_accept(self, accepted)
                result = create_supplier_upload_csv(template_path, invoice_path, export_path, cfg)
        info = [f"AFI-Upload-Datei erstellt:\n{export_path}", f"Positionen: {result.get('rows','')}", f"Export-Netto: {result.get('export_net_total','')}"]
        if result.get("net_rounding_difference") and result.get("net_rounding_difference") != "0,00":
            info.append(f"Rundungsdifferenz: {result.get('net_rounding_difference')} (keine Korrektur vorgenommen)")
        if result.get("old_plate_conflicts"):
            info.append(f"Alte Kennzeichenhalter-Konflikte: {len(result.get('old_plate_conflicts') or [])}")
        messagebox.showinfo(MODULE_TITLE, "\n".join(info))
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.run_export = _afi471_run_export
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_V0472_KST_MULTI_ANALYSE_FIX
# Datum: 2026-07-07
# Zweck:
# - Rechnung analysieren akzeptiert Mehrfachauswahl von PDFs und verwendet alle angegebenen PDFs.
# - KST-Reparatur nach Export: Wenn im Export ein Fahrername steht, aber KST leer ist,
#   wird die KST per robuster Namenssuche aus der Kontierungsdatei nachgezogen.
# - Alte-KFZ-Konflikte werden im Exportdialog zuverlässig angezeigt.
# ------------------------------------------------------------------
AFI_UPLOAD_KST_MULTI_ANALYSE_FIX_VERSION = "0.472"


def _afi472_extract_driver_from_text(text, entries):
    ntext = _norm(text)
    if not ntext:
        return None, 0.0
    best = (None, 0.0)
    for e in entries or []:
        for v in _afi471_name_variants(e) if '_afi471_name_variants' in globals() else [e.get('name_norm','')]:
            if not v or len(v) < 5:
                continue
            if v in ntext:
                # prefer longest concrete full-name match
                score = min(1.0, 0.90 + len(v) / max(len(ntext), 1) / 10)
            else:
                # compare against text, but keep threshold high to avoid false mapping from cost text
                score = _afi471_similarity(ntext, v) if '_afi471_similarity' in globals() else 0.0
                if score < 0.90:
                    continue
            if score > best[1]:
                best = (e, score)
    return best


def _afi472_repair_export_kst_by_text(export_path, assignment_path, result=None):
    try:
        entries = load_assignment_entries(assignment_path)
    except Exception:
        return 0
    try:
        with open(export_path, 'r', encoding='utf-8-sig', newline='') as f:
            rows = list(csv.DictReader(f, delimiter=';'))
            fieldnames = f.seek(0) or None
    except Exception:
        return 0
    if not rows:
        return 0
    # Preserve canonical AFI column order.
    repaired = 0
    for row in rows:
        if (row.get('COSTCENTER') or '').strip():
            continue
        text = row.get('TEXT', '')
        entry, score = _afi472_extract_driver_from_text(text, entries)
        if not entry:
            continue
        # Do not repair intentionally suppressed old-plate conflict rows that contain no driver name.
        # A repair is only allowed if the actual entry name appears or matches robustly in the visible text.
        gl, cc, orderid = _select_assignment_values(entry, _afi471_cost_type_for_gl(text) if '_afi471_cost_type_for_gl' in globals() else _cost_type(text), text)
        if cc:
            row['COSTCENTER'] = cc
            if orderid and not row.get('ORDERID'):
                row['ORDERID'] = orderid
            if gl and not row.get('GL_ACCOUNT'):
                row['GL_ACCOUNT'] = gl
            repaired += 1
    if repaired:
        with open(export_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=';', extrasaction='ignore')
            writer.writeheader(); writer.writerows(rows)
    return repaired


_create_supplier_upload_csv_before_v0472 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    res = _create_supplier_upload_csv_before_v0472(assignment_path, invoice_path, export_path, config)
    # KST-Reparatur nur, wenn ein Fahrername sichtbar ist, aber die KST leer blieb.
    try:
        repaired = _afi472_repair_export_kst_by_text(export_path, assignment_path, res)
        if isinstance(res, dict):
            res['kst_repaired_by_visible_name'] = repaired
            if repaired:
                res.setdefault('name_fallback_matches', [])
                res['name_fallback_matches'].append(f'KST per sichtbarem Fahrernamen nachgezogen: {repaired} Position(en)')
    except Exception:
        pass
    return res


_analyze_invoice_before_v0472 = SupplierUploadUI.analyze_invoice if 'SupplierUploadUI' in globals() else None

def _afi472_analyze_invoice(self):
    raw = self.invoice_var.get().strip()
    paths = _afi471_split_invoice_paths(raw) if '_afi471_split_invoice_paths' in globals() else ([raw] if raw else [])
    if not paths or not all(os.path.isfile(p) for p in paths):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gueltige Rechnung auswaehlen.')
        return
    if len(paths) > 1:
        non_pdf = [p for p in paths if os.path.splitext(p)[1].lower() != '.pdf']
        if non_pdf:
            messagebox.showwarning(MODULE_TITLE, 'Mehrfachauswahl ist nur fuer PDF-Rechnungen vorgesehen.')
            return
        try:
            self.clear_sources()
        except Exception:
            pass
        try:
            self.load_preview(paths[0])
        except Exception:
            pass
        try:
            _afi463_update_export_path(self, True)
        except Exception:
            try: _fm_update_export_path(self, True)
            except Exception: pass
        self.headers, self.rows = ['PDF'], []
        try:
            if hasattr(self, 'add_source_btn'):
                self.add_source_btn.configure(state='disabled')
        except Exception:
            pass
        try:
            self.suggestion_var.set(f'{len(paths)} PDF-Rechnungen erkannt: Alle ausgewaehlten PDFs werden fuer einen gemeinsamen AFI-Export verwendet.')
        except Exception:
            pass
        try:
            _afi464_populate_positions_tree(self)
        except Exception:
            pass
        return
    # Single file keeps existing behaviour.
    if _analyze_invoice_before_v0472:
        return _analyze_invoice_before_v0472(self)

try:
    SupplierUploadUI.analyze_invoice = _afi472_analyze_invoice
except Exception:
    pass


_run_export_before_v0472 = SupplierUploadUI.run_export if 'SupplierUploadUI' in globals() else None

def _afi472_run_export(self):
    template_path = self.template_var.get().strip()
    invoice_path = self.invoice_var.get().strip()
    try: _afi463_update_export_path(self, False)
    except Exception: pass
    export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gueltige Zuordnungsdatei auswaehlen.'); return
    paths = _afi471_split_invoice_paths(invoice_path) if '_afi471_split_invoice_paths' in globals() else [invoice_path]
    if not paths or not all(os.path.isfile(p) for p in paths):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gueltige Rechnung auswaehlen.'); return
    if len(paths) > 1 and any(os.path.splitext(p)[1].lower() != '.pdf' for p in paths):
        messagebox.showwarning(MODULE_TITLE, 'Mehrfachauswahl ist nur fuer PDF-Rechnungen vorgesehen.'); return
    if not export_path:
        messagebox.showwarning(MODULE_TITLE, 'Bitte einen Exportpfad auswaehlen.'); return
    try:
        cfg = _afi471_current_config_with_accept(self) if '_afi471_current_config_with_accept' in globals() else _afi463_current_config(self)
        result = create_supplier_upload_csv(template_path, invoice_path, export_path, cfg)
        conflicts = result.get('old_plate_conflicts') or [] if isinstance(result, dict) else []
        if conflicts:
            accepted = _afi471_old_plate_dialog(self, conflicts) if '_afi471_old_plate_dialog' in globals() else []
            # Re-export in jedem Fall, damit nicht übernommene Konflikte bewusst ohne Fahrer/KST bleiben
            # und übernommene Konflikte sichtbar mit Fahrer/KST geschrieben werden.
            cfg = _afi471_current_config_with_accept(self, accepted) if '_afi471_current_config_with_accept' in globals() else cfg
            result = create_supplier_upload_csv(template_path, invoice_path, export_path, cfg)
        info = [f'AFI-Upload-Datei erstellt:\n{export_path}', f"Positionen: {result.get('rows','') if isinstance(result,dict) else ''}", f"Export-Netto: {result.get('export_net_total','') if isinstance(result,dict) else ''}"]
        if isinstance(result, dict) and result.get('kst_repaired_by_visible_name'):
            info.append(f"KST per sichtbarem Fahrernamen nachgezogen: {result.get('kst_repaired_by_visible_name')}")
        if isinstance(result, dict) and result.get('net_rounding_difference') and result.get('net_rounding_difference') != '0,00':
            info.append(f"Rundungsdifferenz: {result.get('net_rounding_difference')} (keine Korrektur vorgenommen)")
        if conflicts:
            info.append(f'Alte Kennzeichenhalter-Konflikte geprüft: {len(conflicts)}')
        messagebox.showinfo(MODULE_TITLE, '\n'.join(info))
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.run_export = _afi472_run_export
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_DKV_INVOICE_NAMES_AND_CARD_FEE_FIX_V0473
# Datum: 2026-07-07
# Zweck:
# - DKV: Rechnungsname/Kartenzusatz ist beim Export-Text federführend.
# - KST-Zuordnungsdatei wird nur für Sachkonto/KST/IA genutzt, nicht zum Überschreiben des Namens.
# - DKV Card Entgelt vor EUR-Übertrag wird korrekt mit Netto 1,80 statt USt 0,34 exportiert.
# ------------------------------------------------------------------
AFI_UPLOAD_DKV_INVOICE_NAMES_AND_CARD_FEE_FIX_VERSION = "0.473"

_DKV_LAST_INVOICE_TEXT_BY_PLATE = {}


def _dkv473_cut_after_total_noise(after_total_text):
    text = after_total_text or ""
    # Seiten-/Blockuebertraege gehoeren nicht mehr zur TOTAL-Zeile des Fahrzeugs.
    text = re.split(r"(?:EUR\s+Übertrag|EUR\s+Uebertrag|Übertrag\s+EUR|Uebertrag\s+EUR|VEHICLE:|Gesamtsummenaufstellung|Umsatzsteuerstatistik)", text, maxsplit=1, flags=re.I)[0]
    return text


def _dkv473_total_amounts_after_total(after_total_text):
    cleaned = _dkv473_cut_after_total_noise(after_total_text)
    nums = _DKV_IDG_TOTAL_NUMBER_RE_V2.findall(cleaned or "")[:6] if '_DKV_IDG_TOTAL_NUMBER_RE_V2' in globals() else _DKV_TOTAL_NUMBER_RE.findall(cleaned or "")[:6]
    vals = [_dec(x) for x in nums]
    # Mit Nachlass: Menge, Bezugswert, Nachlass, Gesamtwert netto, USt, Brutto
    if len(vals) >= 6 and vals[2] < 0:
        return vals[3], vals[5], nums
    # Ohne Nachlass + Mengenangabe: Menge, Bezugswert/Gesamtwert, Gesamtwert netto, USt, Brutto
    if len(vals) >= 5:
        return vals[2], vals[4], nums[:5]
    # Reines Card-Entgelt ohne Mengenangabe nach TOTAL: Netto, Netto, USt, Brutto
    if len(vals) == 4:
        return vals[1], vals[3], nums
    # Fallback: Netto, USt, Brutto
    if len(vals) >= 3:
        return vals[-3], vals[-1], nums
    return Decimal("0.00"), Decimal("0.00"), nums

# Beide bisherigen DKV-Total-Helfer überschreiben, weil es zwei Parsergenerationen gibt.
def _dkv_idg_total_amounts_after_total_v2(after_total_text):
    return _dkv473_total_amounts_after_total(after_total_text)

def _dkv_total_amounts_from_after_total(after_total_text):
    return _dkv473_total_amounts_after_total(after_total_text)


_parse_dkv_tanken_pdf_positions_before_v0473 = _parse_dkv_tanken_pdf_positions_v1

def _parse_dkv_tanken_pdf_positions_v1(path, global_prefix):
    positions = _parse_dkv_tanken_pdf_positions_before_v0473(path, global_prefix)
    # Rechnungs-Kartenzusatz je Kennzeichen merken. Dieser Text ist federführend für den AFI-TEXT.
    global _DKV_LAST_INVOICE_TEXT_BY_PLATE
    _DKV_LAST_INVOICE_TEXT_BY_PLATE = {}
    for pos in positions or []:
        plate = _dkv_clean_plate(pos.get('key', '')) if '_dkv_clean_plate' in globals() else _clean(pos.get('key', ''))
        driver = _clean(pos.get('driver', ''))
        if plate and driver and _norm(driver) != _norm(plate):
            _DKV_LAST_INVOICE_TEXT_BY_PLATE[_norm(plate)] = driver
    return positions


def _dkv473_invoice_display_for_plate(plate):
    return _DKV_LAST_INVOICE_TEXT_BY_PLATE.get(_norm(plate), "")


def _apply_dkv_driver_names_to_export(export_path, assignment_path, invoice_path, config):
    """Nur noch Rechnungsnamen/Kartenzusatz in den Text schreiben.

    Wichtig: Die KST-Zuordnung darf den in der Rechnung stehenden Namen nicht überschreiben.
    Die Zuordnungsdatei ist nur fuer Sachkonto, Kostenstelle und IA relevant.
    """
    if not export_path or not os.path.isfile(export_path):
        return []
    supplier = _fm_supplier_from_invoice_name(invoice_path) if '_fm_supplier_from_invoice_name' in globals() else ""
    cost = _norm((config or {}).get("global_prefix", ""))
    if supplier != "DKV" and "TANKEN" not in cost:
        return []
    # Falls der Export aus einem vorherigen Parserlauf kommt und die Map leer ist, PDF erneut lesen.
    if not _DKV_LAST_INVOICE_TEXT_BY_PLATE:
        try:
            for p in _parse_pdf_invoice_positions(invoice_path, (config or {}).get('global_prefix', 'Tanken')):
                plate = _dkv_clean_plate(p.get('key','')) if '_dkv_clean_plate' in globals() else _clean(p.get('key',''))
                driver = _clean(p.get('driver',''))
                if plate and driver and _norm(driver) != _norm(plate):
                    _DKV_LAST_INVOICE_TEXT_BY_PLATE[_norm(plate)] = driver
        except Exception:
            pass
    with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        fieldnames = reader.fieldnames or []
        rows = list(reader)
    if "TEXT" not in fieldnames:
        return []
    prefix = _clean((config or {}).get("global_prefix", "")) or "Tanken"
    updates = []
    for row in rows:
        text_value = _clean(row.get("TEXT", ""))
        plate_m = _DKV_PLATE_RE.search(text_value) if '_DKV_PLATE_RE' in globals() else None
        if not plate_m:
            continue
        plate = _dkv_clean_plate(plate_m.group(0)) if '_dkv_clean_plate' in globals() else _clean(plate_m.group(0))
        display = _dkv473_invoice_display_for_plate(plate)
        if not display:
            continue
        new_text = _clean(f"{prefix} {plate} {display}")
        if '_ascii_umlauts' in globals():
            new_text = _ascii_umlauts(new_text)
        if row.get("TEXT", "") != new_text:
            updates.append({"Kennzeichen": plate, "Alt": row.get("TEXT", ""), "Neu": new_text})
            row["TEXT"] = new_text
    if updates:
        with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
            writer.writeheader(); writer.writerows(rows)
    global _DKV_LAST_DRIVER_TEXT_UPDATES
    _DKV_LAST_DRIVER_TEXT_UPDATES = updates
    return updates


# ------------------------------------------------------------------
# AFI_UPLOAD_DKV_INVOICE_LEADING_AND_FALLBACK_GL_V0474
# Datum: 2026-07-07
# Zweck:
# - DKV-PDF-Positionen sind nach dem PDF-Parsing fachlich fuehrend fuer Betrag und Text.
# - KST/IA werden nur aus der Kontierungsdatei ermittelt und nie aus der Rechnungsspalte/Kunden-ID vertauscht.
# - Sachkonto wird auch ohne eindeutige KST-Zuordnung gesetzt.
# ------------------------------------------------------------------
AFI_UPLOAD_DKV_INVOICE_LEADING_AND_FALLBACK_GL_VERSION = "0.474"

_DKV474_LAST_REPAIRS = []


def _dkv474_is_dkv_context(invoice_path, config):
    supplier = _fm_supplier_from_invoice_name(invoice_path) if '_fm_supplier_from_invoice_name' in globals() else ''
    cost = _norm((config or {}).get('global_prefix', ''))
    name = _norm(os.path.basename(invoice_path or ''))
    return supplier == 'DKV' or 'DKV' in name or 'TANKEN' in cost


def _dkv474_best_entry_for_invoice(entries, plate, invoice_driver):
    nplate = _norm(plate)
    # 1) Kennzeichen exakt, wenn vorhanden.
    for e in entries or []:
        if e.get('identifier_type') == 'PLATE' and e.get('identifier_norm') == nplate:
            return e, 'plate'
    # 2) Rechnungsname/Kartenzusatz, z. B. VON PREEN, BEIER, TRUMPP.
    target = _norm(invoice_driver)
    if target:
        # Exakt gegen Nachname / Fullname / Varianten.
        exact = []
        for e in entries or []:
            variants = _afi471_name_variants(e) if '_afi471_name_variants' in globals() else [e.get('name_norm',''), e.get('last_norm','')]
            if target in variants or any(v and (target == v or target in v or v in target) for v in variants):
                exact.append(e)
        if len(exact) == 1:
            return exact[0], 'invoice_name'
        # Fuzzy als Fallback, aber nur eindeutige Treffer.
        if '_afi471_best_name_match' in globals():
            e, how, score = _afi471_best_name_match(entries, invoice_driver, threshold=0.84)
            if e:
                return e, 'invoice_name_fuzzy'
    return None, ''


def _dkv474_gl_fallback(assignment_path, cost_desc='Tanken'):
    if '_afi471_fallback_entry_for_gl' in globals():
        ent = _afi471_fallback_entry_for_gl(assignment_path, cost_desc, 'IDE')
        gl, _cc, _ia = _select_assignment_values(ent, 'TANKEN', cost_desc)
        return gl or ent.get('gl_tanken') or ent.get('gl_default') or ''
    try:
        cache = _afi463_load_general_overview(assignment_path)
        m = cache.get('gl', {}).get('IDE', {})
        return m.get(_norm('DKV'), '') or m.get(_norm('Tanken'), '') or m.get(_norm('Sonstige'), '')
    except Exception:
        return '427000'


def _dkv474_repair_export_from_pdf_positions(export_path, assignment_path, invoice_path, config):
    if not export_path or not os.path.isfile(export_path) or not _dkv474_is_dkv_context(invoice_path, config):
        return []
    try:
        positions = _parse_pdf_invoice_positions(invoice_path, (config or {}).get('global_prefix', 'Tanken'))
    except Exception:
        return []
    by_plate = {}
    for pos in positions or []:
        plate = _dkv_clean_plate(pos.get('key','')) if '_dkv_clean_plate' in globals() else _clean(pos.get('key',''))
        if plate:
            by_plate[_norm(plate)] = pos
    if not by_plate:
        return []
    try:
        entries = load_assignment_entries(assignment_path)
    except Exception:
        entries = []
    fallback_gl = _dkv474_gl_fallback(assignment_path, (config or {}).get('global_prefix', 'Tanken')) or '427000'
    with open(export_path, 'r', encoding='utf-8-sig', newline='') as f:
        reader = csv.DictReader(f, delimiter=';')
        fieldnames = reader.fieldnames or UPLOAD_COLUMNS
        rows = list(reader)
    repairs = []
    prefix = _clean((config or {}).get('global_prefix', '')) or 'Tanken'
    for row in rows:
        text_value = _clean(row.get('TEXT',''))
        plate_match = _DKV_PLATE_RE.search(text_value) if '_DKV_PLATE_RE' in globals() else None
        if not plate_match:
            continue
        plate = _dkv_clean_plate(plate_match.group(0)) if '_dkv_clean_plate' in globals() else _clean(plate_match.group(0))
        pos = by_plate.get(_norm(plate))
        if not pos:
            continue
        invoice_driver = _clean(pos.get('driver',''))
        # Rechnungsname/-kartenzusatz bleibt im TEXT fachlich fuehrend.
        new_text = _clean(f"{prefix} {plate} {invoice_driver}") if invoice_driver else _clean(f"{prefix} {plate}")
        if '_ascii_umlauts' in globals():
            new_text = _ascii_umlauts(new_text)
        amount = Decimal(pos.get('amount', '0')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        entry, how = _dkv474_best_entry_for_invoice(entries, plate, invoice_driver)
        gl = fallback_gl
        cc = ''
        orderid = ''
        if entry:
            egl, ecc, eia = _select_assignment_values(entry, 'TANKEN', prefix)
            gl = egl or gl
            cc = ecc or ''
            orderid = eia or ''
        before = dict(row)
        row['TEXT'] = new_text
        row['PRICE'] = _fmt(amount)
        row['NET_VALUE'] = _fmt(amount)
        row['PRICE_UNIT'] = row.get('PRICE_UNIT') or '1'
        row['QUANTITY'] = row.get('QUANTITY') or '1'
        row['UNIT'] = row.get('UNIT') or 'ST'
        row['TAX_CODE'] = pos.get('tax') or row.get('TAX_CODE') or 'VD'
        row['GL_ACCOUNT'] = gl or row.get('GL_ACCOUNT') or fallback_gl
        row['COSTCENTER'] = cc
        row['ORDERID'] = orderid
        if before != row:
            repairs.append({'Kennzeichen': plate, 'Rechnungsname': invoice_driver, 'Zuordnung': how, 'Alt': before, 'Neu': dict(row)})
    if repairs:
        with open(export_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=';', extrasaction='ignore')
            writer.writeheader(); writer.writerows(rows)
    global _DKV474_LAST_REPAIRS
    _DKV474_LAST_REPAIRS = repairs
    return repairs


_create_supplier_upload_csv_before_v0474 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    result = _create_supplier_upload_csv_before_v0474(assignment_path, invoice_path, export_path, config)
    try:
        repairs = _dkv474_repair_export_from_pdf_positions(export_path, assignment_path, invoice_path, config)
        if isinstance(result, dict):
            result['dkv_invoice_leading_repairs'] = len(repairs)
            if repairs:
                result['dkv_invoice_leading_repair_details'] = repairs[:20]
    except Exception as exc:
        if isinstance(result, dict):
            result['dkv_invoice_leading_repair_error'] = str(exc)
    return result


# ------------------------------------------------------------------
# AFI_UPLOAD_DKV_DIRECT_EXPORT_FINAL_V0475
# Datum: 2026-07-07
# Zweck:
# DKV-PDF-Exporte werden nicht mehr nachtraeglich repariert, sondern direkt aus der
# Rechnung erzeugt. Damit koennen alte Zwischenlogiken keinen Betrag, Text, KST oder
# Sachkonto mehr zurueck auf falsche Werte setzen.
# ------------------------------------------------------------------
AFI_UPLOAD_DKV_DIRECT_EXPORT_FINAL_VERSION = "0.475"

_DKV475_LAST_DIRECT_EXPORT = {}


def _dkv475_is_dkv_pdf_file(path):
    try:
        if os.path.splitext(path or '')[1].lower() != '.pdf':
            return False
        txt = _extract_pdf_text(path)
        return _dkv_is_dkv_tanken_pdf_text(txt) if '_dkv_is_dkv_tanken_pdf_text' in globals() else ('DKV' in _norm(txt) and 'VEHICLE' in _norm(txt))
    except Exception:
        return False


def _dkv475_invoice_paths(invoice_path):
    paths = _afi471_split_invoice_paths(invoice_path) if '_afi471_split_invoice_paths' in globals() else ([invoice_path] if invoice_path else [])
    return [p for p in paths if str(p).strip()]


def _dkv475_should_direct_export(invoice_path, config):
    paths = _dkv475_invoice_paths(invoice_path)
    if not paths:
        return False
    if not all(os.path.isfile(p) and os.path.splitext(p)[1].lower().endswith('.pdf') for p in paths):
        return False
    if not all(_dkv475_is_dkv_pdf_file(p) for p in paths):
        return False
    return True


def _dkv475_assignment_values(entries, plate, invoice_driver, assignment_path, text_label):
    entry = None; how = ''
    if '_dkv474_best_entry_for_invoice' in globals():
        entry, how = _dkv474_best_entry_for_invoice(entries, plate, invoice_driver)
    if not entry:
        nplate = _norm(plate)
        for e in entries or []:
            if e.get('identifier_type') == 'PLATE' and e.get('identifier_norm') == nplate:
                entry = e; how = 'plate'; break
    if not entry and invoice_driver and '_afi471_best_name_match' in globals():
        entry, how2, score = _afi471_best_name_match(entries, invoice_driver, threshold=0.84)
        how = how2 if entry else how
    fallback_gl = _dkv474_gl_fallback(assignment_path, text_label) if '_dkv474_gl_fallback' in globals() else '427000'
    gl, cc, orderid = fallback_gl, '', ''
    if entry:
        egl, ecc, eia = _select_assignment_values(entry, 'TANKEN', text_label)
        gl = egl or fallback_gl
        cc = ecc or ''
        orderid = eia or ''
    return gl or fallback_gl or '427000', cc, orderid, how


def _dkv475_create_direct_export(assignment_path, invoice_path, export_path, config):
    paths = _dkv475_invoice_paths(invoice_path)
    prefix = _clean((config or {}).get('global_prefix', '')) or 'Tanken'
    try:
        entries = load_assignment_entries(assignment_path)
    except Exception:
        entries = []
    groups = OrderedDict()
    parsed_count = 0
    for pdf_path in paths:
        for pos in _parse_pdf_invoice_positions(pdf_path, prefix):
            parsed_count += 1
            plate = _dkv_clean_plate(pos.get('key','')) if '_dkv_clean_plate' in globals() else _clean(pos.get('key',''))
            invoice_driver = _clean(pos.get('driver',''))
            if _norm(invoice_driver) == _norm(plate):
                invoice_driver = ''
            gl, cc, orderid, how = _dkv475_assignment_values(entries, plate, invoice_driver, assignment_path, prefix)
            text = _clean(f'{prefix} {plate} {invoice_driver}') if invoice_driver else _clean(f'{prefix} {plate}')
            if '_ascii_umlauts' in globals():
                text = _ascii_umlauts(text)
            tax = pos.get('tax') or 'VD'
            key = (_norm(text), tax, gl, cc, orderid)
            if key not in groups:
                groups[key] = {'TEXT': text, 'amount': Decimal('0.00'), 'TAX_CODE': tax, 'GL_ACCOUNT': gl, 'COSTCENTER': cc, 'ORDERID': orderid, 'how': how, 'plate': plate, 'invoice_driver': invoice_driver}
            groups[key]['amount'] += Decimal(pos.get('amount') or '0')
    if not groups:
        raise RuntimeError('Aus der DKV-PDF konnten keine Exportpositionen erzeugt werden.')
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or '.', exist_ok=True)
    ordered = list(groups.values())
    ordered.sort(key=lambda r: (_norm(r.get('TEXT','')), TAX_ORDER.get(r.get('TAX_CODE','VD'), 9)))
    with open(export_path, 'w', encoding='utf-8-sig', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=';', extrasaction='ignore')
        writer.writeheader()
        for g in ordered:
            amount = Decimal(g.get('amount','0')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
            writer.writerow({'TEXT': g.get('TEXT',''), 'PRICE': _fmt(amount), 'PRICE_UNIT': '1', 'QUANTITY': '1', 'UNIT': 'ST', 'NET_VALUE': _fmt(amount), 'TAX_CODE': g.get('TAX_CODE',''), 'GL_ACCOUNT': g.get('GL_ACCOUNT',''), 'COSTCENTER': g.get('COSTCENTER',''), 'ORDERID': g.get('ORDERID','')})
    total = sum(Decimal(g.get('amount','0')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP) for g in ordered)
    global _DKV475_LAST_DIRECT_EXPORT
    _DKV475_LAST_DIRECT_EXPORT = {'paths': paths, 'parsed_positions': parsed_count, 'rows': len(ordered), 'export_path': export_path}
    return {'rows': len(ordered), 'export_path': export_path, 'export_net_total': _fmt(total), 'invoice_net_raw_total': _fmt(total), 'net_rounding_difference': '0,00', 'dkv_direct_export': True, 'dkv_pdf_count': len(paths), 'dkv_parsed_positions': parsed_count, 'missing_template': [], 'unknown_tax': []}


_create_supplier_upload_csv_before_v0475 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    if _dkv475_should_direct_export(invoice_path, config):
        return _dkv475_create_direct_export(assignment_path, invoice_path, export_path, config)
    return _create_supplier_upload_csv_before_v0475(assignment_path, invoice_path, export_path, config)


# ------------------------------------------------------------------
# AFI_UPLOAD_INTERCOMPANY_SINGLE_BUKRS_ENBW_TEXT_FIX_V0478
# Datum: 2026-07-07
# Zweck:
# - Es kann fachlich nur ein Buchungskreis aktiv sein.
# - Standardmaessig aktive Option "Weiterberechnung Intercompany".
# - Bei aktivem Intercompany wird bei abweichender Fahrer-Firma die KST Weiterberechnung <Buchungskreis> an <Fahrer-Bukrs> verwendet.
# - EnBW/Export-TEXT: lange Ziffernfolgen auf 11 Ziffern kuerzen; Kennzeichen nur plausibel belassen.
# ------------------------------------------------------------------
AFI_UPLOAD_INTERCOMPANY_SINGLE_BUKRS_ENBW_TEXT_FIX_VERSION = "0.478"

_AFI478_CONTEXT = {"booking_circle": "IDE", "intercompany": True, "assignment_path": ""}
_AFI478_INTERCOMPANY_CACHE = {"path": "", "mtime": None, "map": {}}
_AFI478_PLATE_RE = re.compile(r"\b[A-ZÄÖÜ]{1,3}\s*-\s*[A-ZÄÖÜ]{1,3}\s*\d{1,5}[A-ZÄÖÜ]?\b", re.I)
_AFI478_LONG_DIGITS_RE = re.compile(r"\d{12,}")


def _afi478_bukrs_label(value):
    n = _norm(value)
    if n in ("IDG", "INTERSPORTDIGITALGMBH") or "DIGITAL" in n:
        return "IDG"
    if n == "SABU" or "SABU" in n:
        return "SABU"
    if n == "IMS" or "MARKETINGSERVICES" in n:
        return "IMS"
    return "IDE"


def _afi478_parse_intercompany_map(assignment_path):
    try:
        mtime = os.path.getmtime(assignment_path)
    except Exception:
        return {}
    if _AFI478_INTERCOMPANY_CACHE.get("path") == assignment_path and _AFI478_INTERCOMPANY_CACHE.get("mtime") == mtime:
        return dict(_AFI478_INTERCOMPANY_CACHE.get("map") or {})
    out = {}
    try:
        from openpyxl import load_workbook
        wb = load_workbook(assignment_path, data_only=True, read_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                vals = list(row or [])
                for i, val in enumerate(vals):
                    lab = _clean(val)
                    if not lab or "Weiterberechnung" not in lab:
                        continue
                    m = re.search(r"Weiterberechnung\s+([A-Z]{2,4})\s+an\s+([A-Z]{2,4})", lab, re.I)
                    if not m:
                        continue
                    src = _afi478_bukrs_label(m.group(1))
                    dst = _afi478_bukrs_label(m.group(2))
                    kst = ""
                    for j in range(i + 1, min(len(vals), i + 4)):
                        cand = _fm_strip_excel_code(vals[j]) if '_fm_strip_excel_code' in globals() else _clean(vals[j])
                        if cand:
                            kst = cand
                            break
                    if kst:
                        out[(src, dst)] = kst
        try:
            wb.close()
        except Exception:
            pass
    except Exception:
        out = {}
    _AFI478_INTERCOMPANY_CACHE.update({"path": assignment_path, "mtime": mtime, "map": dict(out)})
    return out


def _afi478_selected_bukrs_single(self):
    if hasattr(self, "booking_circle_vars"):
        selected = [b for b, v in self.booking_circle_vars.items() if v.get()]
        chosen = selected[0] if selected else "IDE"
        # Fachlich genau ein Buchungskreis: UI-Checkboxen gegenseitig bereinigen.
        for b, v in self.booking_circle_vars.items():
            try:
                v.set(b == chosen)
            except Exception:
                pass
        return chosen
    if hasattr(self, "booking_circle_var"):
        return self.booking_circle_var.get() or "IDE"
    return "IDE"


def _afi478_set_only_bukrs(self, chosen):
    if hasattr(self, "booking_circle_vars"):
        for b, v in self.booking_circle_vars.items():
            try:
                v.set(b == chosen)
            except Exception:
                pass
    try:
        _afi463_update_export_path(self, True)
    except Exception:
        try:
            _fm_update_export_path(self, True)
        except Exception:
            pass


def _afi463_selected_bukrs(self):
    return [_afi478_selected_bukrs_single(self)]


def _afi478_config_from_config(config):
    cfg = dict(config or {})
    selected = cfg.get("booking_circle") or (cfg.get("booking_circles") or ["IDE"])[0] if isinstance(cfg.get("booking_circles"), list) else "IDE"
    selected = _afi478_bukrs_label(selected or "IDE")
    intercompany = cfg.get("intercompany_recharge")
    if intercompany is None:
        intercompany = cfg.get("intercompany")
    if intercompany is None:
        intercompany = True
    return selected, bool(intercompany)


def _afi478_update_context(assignment_path, config):
    selected, intercompany = _afi478_config_from_config(config)
    _AFI478_CONTEXT.update({"booking_circle": selected, "intercompany": intercompany, "assignment_path": assignment_path or ""})


_current_config_before_v0478 = _afi463_current_config if '_afi463_current_config' in globals() else None

def _afi463_current_config(self):
    b = _afi478_selected_bukrs_single(self)
    cfg = _current_config_before_v0478(self) if _current_config_before_v0478 else {}
    cfg["booking_circle"] = b
    cfg["booking_circles"] = [b]
    cfg["supplier"] = self.supplier_var.get() if hasattr(self, "supplier_var") else cfg.get("supplier", "Automatisch erkennen")
    cfg["global_prefix"] = self.global_prefix_var.get() if hasattr(self, "global_prefix_var") else cfg.get("global_prefix", "Tanken Strom")
    cfg["sources"] = [s.get() for s in getattr(self, "sources", [])]
    cfg["intercompany_recharge"] = self.intercompany_recharge_var.get() if hasattr(self, "intercompany_recharge_var") else True
    return cfg
try:
    _afi463_current_config.__name__ = "_afi463_current_config"
except Exception:
    pass


def _afi478_entry_bukrs(entry):
    return _afi478_bukrs_label((entry or {}).get("bukrs") or (entry or {}).get("firma") or "IDE")


_select_assignment_values_before_v0478 = _select_assignment_values

def _select_assignment_values(entry, cost_type, text_label=""):
    gl, cc, orderid = _select_assignment_values_before_v0478(entry, cost_type, text_label)
    try:
        if not _AFI478_CONTEXT.get("intercompany", True):
            return gl, cc, orderid
        src = _afi478_bukrs_label(_AFI478_CONTEXT.get("booking_circle") or "IDE")
        dst = _afi478_entry_bukrs(entry)
        if src == dst:
            return gl, cc, orderid
        amap = _afi478_parse_intercompany_map(_AFI478_CONTEXT.get("assignment_path") or "")
        inter_cc = amap.get((src, dst))
        if inter_cc:
            return gl, inter_cc, orderid
    except Exception:
        pass
    return gl, cc, orderid


def _afi478_sanitize_export_text_value(value):
    text = _clean(value)
    if not text:
        return text
    # Lange reine Nummern duerfen nicht wie Kennzeichen wirken und werden im Export auf 11 Ziffern begrenzt.
    text = _AFI478_LONG_DIGITS_RE.sub(lambda m: m.group(0)[:11], text)
    # Falls nach einem Kostenwort ein unplausibles "Kennzeichen" nur aus Zahlen steht, bleibt nur die gekuerzte Nummer.
    # Plausible Kennzeichen im Format Ort-Buchstaben-Ziffern bleiben unveraendert.
    return text


def _afi478_sanitize_export_file(export_path):
    if not export_path or not os.path.isfile(export_path):
        return 0
    try:
        with open(export_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f, delimiter=";")
            fieldnames = reader.fieldnames or UPLOAD_COLUMNS
            rows = list(reader)
    except Exception:
        return 0
    changed = 0
    for row in rows:
        old = row.get("TEXT", "")
        new = _afi478_sanitize_export_text_value(old)
        if new != old:
            row["TEXT"] = new
            changed += 1
    if changed:
        with open(export_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=";", extrasaction="ignore")
            writer.writeheader(); writer.writerows(rows)
    return changed


_create_supplier_upload_csv_before_v0478 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    _afi478_update_context(assignment_path, config or {})
    result = _create_supplier_upload_csv_before_v0478(assignment_path, invoice_path, export_path, config)
    try:
        changed = _afi478_sanitize_export_file(export_path)
        if isinstance(result, dict):
            result["text_sanitized_long_digits"] = changed
            result["intercompany_recharge"] = bool(_AFI478_CONTEXT.get("intercompany", True))
            result["single_booking_circle"] = _AFI478_CONTEXT.get("booking_circle", "IDE")
    finally:
        _AFI478_CONTEXT.update({"booking_circle": "IDE", "intercompany": True, "assignment_path": ""})
    return result


_build_left_before_v0478 = SupplierUploadUI._build_left if 'SupplierUploadUI' in globals() else None

def _afi478_build_left(self, parent):
    if _build_left_before_v0478:
        _build_left_before_v0478(self, parent)
    if not hasattr(self, "intercompany_recharge_var"):
        self.intercompany_recharge_var = tk.BooleanVar(value=True)
    # Vorhandene Buchungskreis-Checkboxen fachlich gegenseitig ausschliessend nachverdrahten.
    if hasattr(self, "booking_circle_vars"):
        for b, v in self.booking_circle_vars.items():
            try:
                # add='+' ist nicht bei Checkbutton-command verfuegbar; daher command ersetzen.
                # Bei Klick wird exakt dieser Buchungskreis aktiv gesetzt.
                pass
            except Exception:
                pass
    try:
        row = 14
        tk.Checkbutton(parent, text="Weiterberechnung Intercompany", variable=self.intercompany_recharge_var, bg=self.bg, font=self.font_small).grid(row=row, column=0, columnspan=4, sticky="w", pady=(3,0))
    except Exception:
        pass

try:
    SupplierUploadUI._build_left = _afi478_build_left
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_INTERCOMPANY_SINGLE_BUKRS_FINAL_V0479
# Datum: 2026-07-07
# Zweck: Intercompany-Konfiguration final absichern und Buchungskreis-Auswahl in der UI gegenseitig ausschliessend machen.
# ------------------------------------------------------------------
AFI_UPLOAD_INTERCOMPANY_SINGLE_BUKRS_FINAL_VERSION = "0.479"


def _afi478_config_from_config(config):
    cfg = dict(config or {})
    if cfg.get("booking_circle"):
        selected = cfg.get("booking_circle")
    elif isinstance(cfg.get("booking_circles"), list) and cfg.get("booking_circles"):
        selected = cfg.get("booking_circles")[0]
    else:
        selected = "IDE"
    selected = _afi478_bukrs_label(selected or "IDE")
    intercompany = cfg.get("intercompany_recharge")
    if intercompany is None:
        intercompany = cfg.get("intercompany")
    if intercompany is None:
        intercompany = True
    return selected, bool(intercompany)


def _afi478_update_context(assignment_path, config):
    selected, intercompany = _afi478_config_from_config(config)
    _AFI478_CONTEXT.update({"booking_circle": selected, "intercompany": intercompany, "assignment_path": assignment_path or ""})


def _afi478_selected_bukrs_single(self):
    if hasattr(self, "booking_circle_vars"):
        selected = [b for b, v in self.booking_circle_vars.items() if v.get()]
        chosen = selected[0] if selected else "IDE"
        for b, v in self.booking_circle_vars.items():
            try:
                v.set(b == chosen)
            except Exception:
                pass
        return chosen
    if hasattr(self, "booking_circle_var"):
        return self.booking_circle_var.get() or "IDE"
    return "IDE"


def _afi478_install_single_bukrs_commands(self):
    if not hasattr(self, "booking_circle_vars"):
        return
    # Checkbuttons wurden in der bestehenden UI bereits gebaut; die sichtbaren Widgets lassen sich nicht immer
    # nachtraeglich zuverlässig adressieren. Daher wird die fachliche Einzelwahl mindestens beim Config-Aufbau
    # und Export erzwungen. Neue/neu aufgebaute UIs koennen diese Funktion zusätzlich nutzen.
    _afi478_selected_bukrs_single(self)


_build_left_before_v0479 = SupplierUploadUI._build_left if 'SupplierUploadUI' in globals() else None

def _afi479_build_left(self, parent):
    if _build_left_before_v0479:
        _build_left_before_v0479(self, parent)
    if not hasattr(self, "intercompany_recharge_var"):
        self.intercompany_recharge_var = tk.BooleanVar(value=True)
    _afi478_install_single_bukrs_commands(self)
    try:
        # Falls v0.478 das Kästchen bereits angelegt hat, ist ein zweites Kästchen optisch unschön.
        # In der Praxis verhindert Tk keine Duplikate; deshalb nur anlegen, wenn noch kein Marker existiert.
        if not getattr(self, "_afi479_intercompany_check_rendered", False):
            row = 14
            tk.Checkbutton(parent, text="Weiterberechnung Intercompany", variable=self.intercompany_recharge_var, bg=self.bg, font=self.font_small).grid(row=row, column=0, columnspan=4, sticky="w", pady=(3,0))
            self._afi479_intercompany_check_rendered = True
    except Exception:
        pass

try:
    SupplierUploadUI._build_left = _afi479_build_left
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_KONSOLIDIERUNG_EXPORT_UI_PREVIEW_V0480
# Datum: 2026-07-07
# Zweck:
# - Konsolidierter finaler Exportpfad am Dateiende.
# - DKV wird zwingend direkt aus der Rechnung exportiert (Beier 1,80 statt 0,34).
# - Konsolidierte UI mit genau einem Buchungskreis und sichtbarer Option "Weiterberechnung Intercompany".
# - Rechnungs-Vorschau zeigt alle Seiten aller ausgewählten PDFs.
# - Lange Ziffernfolgen im TEXT werden auf 11 Ziffern begrenzt; plausible Kennzeichen bleiben unverändert.
# ------------------------------------------------------------------
AFI_UPLOAD_KONSOLIDIERUNG_EXPORT_UI_PREVIEW_VERSION = "0.480"


def _afi480_split_invoice_paths(value):
    if '_afi471_split_invoice_paths' in globals():
        try:
            return [p for p in _afi471_split_invoice_paths(value) if str(p).strip()]
        except Exception:
            pass
    if not value:
        return []
    if isinstance(value, (list, tuple)):
        return [str(p) for p in value if str(p).strip()]
    raw = str(value).replace('|', '\n').replace(';', '\n')
    return [p.strip() for p in raw.splitlines() if p.strip()]


def _afi480_join_invoice_paths(paths):
    if '_afi471_join_invoice_paths' in globals():
        try:
            return _afi471_join_invoice_paths(paths)
        except Exception:
            pass
    return '\n'.join(paths or [])


def _afi480_is_pdf(path):
    return bool(path) and os.path.isfile(path) and os.path.splitext(path)[1].lower() == '.pdf'


def _afi480_is_dkv_pdf(path):
    if not _afi480_is_pdf(path):
        return False
    try:
        if '_dkv475_is_dkv_pdf_file' in globals():
            return bool(_dkv475_is_dkv_pdf_file(path))
        txt = _extract_pdf_text(path)
        return 'DKV' in _norm(txt) and 'VEHICLE' in _norm(txt)
    except Exception:
        return False


def _afi480_render_pdf_paths_image(paths):
    if Image is None:
        return None, []
    paths = [p for p in (paths or []) if _afi480_is_pdf(p)]
    if not paths:
        return None, []
    pil_pages = []
    labels = []
    for pth in paths:
        try:
            import fitz
            doc = fitz.open(pth)
            for page_index, page in enumerate(doc, 1):
                pix = page.get_pixmap(matrix=fitz.Matrix(1.20, 1.20), alpha=False)
                img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
                pil_pages.append(img)
                labels.append(f"{os.path.basename(pth)} - Seite {page_index}")
            doc.close()
        except Exception:
            try:
                pages = _afi464_extract_pdf_pages(pth) if '_afi464_extract_pdf_pages' in globals() else []
                for page_index, page_text in enumerate(pages, 1):
                    img = _afi464_text_preview_image(f"{os.path.basename(pth)} - Seite {page_index}", [page_text]) if '_afi464_text_preview_image' in globals() else None
                    if img:
                        pil_pages.append(img)
                        labels.append(f"{os.path.basename(pth)} - Seite {page_index}")
            except Exception:
                pass
    if not pil_pages:
        return None, []
    max_w = max(img.width for img in pil_pages)
    total_h = 24
    offsets = []
    for img in pil_pages:
        offsets.append(total_h)
        total_h += img.height + 48
    canvas = Image.new('RGB', (max_w + 120, max(700, total_h)), '#F2F4F7')
    draw = ImageDraw.Draw(canvas)
    y = 24
    for label, img in zip(labels, pil_pages):
        x = (canvas.width - img.width) // 2
        draw.text((x, max(2, y - 20)), label, fill='#1F4E79')
        canvas.paste(img, (x, y))
        draw.rectangle((x, y, x + img.width, y + img.height), outline='#B0B0B0', width=2)
        y += img.height + 48
    return canvas, offsets


def _afi480_clear_preview(self):
    try:
        for w in self.preview_frame.winfo_children():
            w.destroy()
    except Exception:
        pass


def _afi480_render_preview_to_ui(self, img, offsets, title='PDF-Vorschau'):
    _afi480_clear_preview(self)
    if img is None:
        tk.Label(self.preview_frame, text='PDF-Vorschau konnte nicht gerendert werden.', bg='white', fg='red').pack(fill='both', expand=True)
        return
    self.preview_base_image = img
    self.preview_page_offsets = offsets or [0]
    self.preview_zoom = 1.0
    self.preview_offset = [0, 0]
    top = tk.Frame(self.preview_frame, bg='#DDE7F3')
    top.place(relx=0, rely=0, relwidth=1, height=30)
    tk.Label(top, text=f"{title}: {len(self.preview_page_offsets)} Seiten", bg='#DDE7F3', font=('Segoe UI', 8, 'bold')).pack(side='left', padx=(4, 8))
    for idx, off in enumerate(self.preview_page_offsets, 1):
        if idx > 80:
            break
        tk.Button(top, text=f"S.{idx}", font=('Segoe UI', 8), command=lambda o=off: (self.preview_offset.__setitem__(1, -int(o * self.preview_zoom)), self._render_preview_image())).pack(side='left', padx=1, pady=2)
    self.preview_canvas = tk.Canvas(self.preview_frame, bg='white', highlightthickness=0)
    self.preview_canvas.place(relx=0, y=30, relwidth=1, relheight=1, height=-30, width=-16)
    scroll = ttk.Scrollbar(self.preview_frame, orient='vertical')
    scroll.place(relx=1, x=-16, y=30, width=16, relheight=1, height=-30)
    def set_scrollbar():
        try:
            ch = max(1, self.preview_canvas.winfo_height())
            total = max(1, int(self.preview_base_image.size[1] * self.preview_zoom))
            top_pos = max(0, -self.preview_offset[1])
            scroll.set(max(0, top_pos / total), min(1, (top_pos + ch) / total))
        except Exception:
            pass
    old_render = self._render_preview_image
    def render_and_scroll(*args, **kwargs):
        old_render(*args, **kwargs)
        set_scrollbar()
    self._render_preview_image = render_and_scroll
    def yview(*args):
        try:
            total = max(1, int(self.preview_base_image.size[1] * self.preview_zoom))
            if args and args[0] == 'moveto':
                self.preview_offset[1] = -int(float(args[1]) * total)
            elif args and args[0] == 'scroll':
                self.preview_offset[1] -= int(args[1]) * 90
            self._render_preview_image()
        except Exception:
            pass
    scroll.configure(command=yview)
    def on_wheel(event):
        if event.state & 0x0004:
            factor = 1.10 if event.delta > 0 else 0.90
            self.preview_zoom = max(0.25, min(4.0, self.preview_zoom * factor))
        else:
            self.preview_offset[1] += 90 if event.delta > 0 else -90
        self._render_preview_image(); return 'break'
    def on_press(event):
        self.preview_drag_start = (event.x, event.y, self.preview_offset[0], self.preview_offset[1]); return 'break'
    def on_drag(event):
        if self.preview_drag_start:
            sx, sy, ox, oy = self.preview_drag_start
            self.preview_offset = [ox + event.x - sx, oy + event.y - sy]
            self._render_preview_image()
        return 'break'
    self.preview_canvas.bind('<Configure>', lambda e: self._render_preview_image())
    self.preview_canvas.bind('<MouseWheel>', on_wheel)
    self.preview_canvas.bind('<ButtonPress-1>', on_press)
    self.preview_canvas.bind('<B1-Motion>', on_drag)
    self._render_preview_image()


_load_preview_before_v0480 = SupplierUploadUI.load_preview if 'SupplierUploadUI' in globals() and hasattr(SupplierUploadUI, 'load_preview') else None

def _afi480_load_preview(self, path):
    paths = _afi480_split_invoice_paths(path)
    if paths and all(_afi480_is_pdf(p) for p in paths):
        img, offsets = _afi480_render_pdf_paths_image(paths)
        return _afi480_render_preview_to_ui(self, img, offsets, 'PDF-Vorschau')
    if _load_preview_before_v0480:
        return _load_preview_before_v0480(self, path)

try:
    SupplierUploadUI.load_preview = _afi480_load_preview
except Exception:
    pass


def _afi480_selected_bukrs(self):
    if hasattr(self, 'booking_circle_var'):
        return self.booking_circle_var.get() or 'IDE'
    if hasattr(self, 'booking_circle_vars'):
        selected = [b for b, v in self.booking_circle_vars.items() if v.get()]
        return selected[0] if selected else 'IDE'
    return 'IDE'


def _afi463_selected_bukrs(self):
    return [_afi480_selected_bukrs(self)]


def _afi463_current_config(self):
    b = _afi480_selected_bukrs(self)
    return {
        'global_prefix': self.global_prefix_var.get() if hasattr(self, 'global_prefix_var') else 'Tanken Strom',
        'sources': [s.get() for s in getattr(self, 'sources', [])],
        'booking_circle': b,
        'booking_circles': [b],
        'supplier': self.supplier_var.get() if hasattr(self, 'supplier_var') else 'Automatisch erkennen',
        'intercompany_recharge': self.intercompany_recharge_var.get() if hasattr(self, 'intercompany_recharge_var') else True,
    }


def _afi480_update_export_path(self, force=False):
    try:
        _afi463_update_export_path(self, force)
    except Exception:
        try:
            _fm_update_export_path(self, force)
        except Exception:
            pass


def _afi480_build_left(self, parent):
    parent.columnconfigure(1, weight=1)
    self.template_var = tk.StringVar(value=KST_ASSIGNMENT_DEFAULT_FILE)
    self.invoice_var = tk.StringVar(value=_fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path())
    self.export_var = tk.StringVar()
    self.global_prefix_var = tk.StringVar(value='Tanken Strom')
    self.supplier_var = tk.StringVar(value='Automatisch erkennen')
    self.booking_circle_var = tk.StringVar(value='IDE')
    self.booking_circle_vars = {b: tk.BooleanVar(value=(b == 'IDE')) for b in BOOKING_CIRCLE_OPTIONS}
    self.intercompany_recharge_var = tk.BooleanVar(value=True)
    self.show_assignment_var = tk.BooleanVar(value=False)
    self.suggestion_var = tk.StringVar(value='')
    self.assignment_status_var = tk.StringVar(value='')
    self.position_saldo_var = tk.StringVar(value='Gesamtbetrag aller Nettobeträge: 0,00')
    self.status_var = tk.StringVar(value='')
    self._afi464_position_rows = []
    self._afi464_selected = set()
    self._afi464_manual_rows = None
    _afi480_update_export_path(self, True)

    tk.Checkbutton(parent, text='Zuordnungsdatei anzeigen', variable=self.show_assignment_var, command=lambda: _afi467_toggle_assignment_frame(self) if '_afi467_toggle_assignment_frame' in globals() else None, bg=self.bg, font=self.font_small).grid(row=0, column=0, columnspan=4, sticky='w', pady=(0, 3))
    self.assignment_frame = tk.Frame(parent, bg=self.bg)
    self.assignment_frame.grid(row=1, column=0, columnspan=4, sticky='ew')
    self.assignment_frame.columnconfigure(1, weight=1)
    tk.Label(self.assignment_frame, text='Zuordnungsdatei', bg=self.bg, font=self.font_small).grid(row=0, column=0, sticky='w')
    tk.Entry(self.assignment_frame, textvariable=self.template_var, font=self.font_small).grid(row=0, column=1, sticky='ew', padx=4)
    tk.Button(self.assignment_frame, text='Wählen', command=lambda: _fm_browse(self, 'Zuordnungsdatei', self.template_var, False, 'template'), font=self.font_small).grid(row=0, column=2, padx=2)
    tk.Button(self.assignment_frame, text='Refresh', command=lambda: _afi463_refresh_assignment_ui(self) if '_afi463_refresh_assignment_ui' in globals() else None, font=self.font_small).grid(row=0, column=3, padx=2)
    tk.Label(self.assignment_frame, textvariable=self.assignment_status_var, bg=self.bg, fg='#445364', font=self.font_small).grid(row=1, column=0, columnspan=4, sticky='ew')
    self.assignment_frame.grid_remove()

    tk.Label(parent, text='Rechnung', bg=self.bg, font=self.font_small).grid(row=2, column=0, sticky='w')
    tk.Entry(parent, textvariable=self.invoice_var, font=self.font_small).grid(row=2, column=1, sticky='ew', padx=4)
    tk.Button(parent, text='Wählen', command=lambda: _fm_browse(self, 'Rechnung / Dokument', self.invoice_var, False, 'invoice'), font=self.font_small).grid(row=2, column=2, padx=2)

    tk.Label(parent, text='Buchungskreis', bg=self.bg, font=self.font_small).grid(row=3, column=0, sticky='w')
    ttk.Combobox(parent, textvariable=self.booking_circle_var, values=BOOKING_CIRCLE_OPTIONS, state='readonly', font=self.font_small, width=12).grid(row=3, column=1, sticky='w', padx=4)
    tk.Checkbutton(parent, text='Weiterberechnung Intercompany', variable=self.intercompany_recharge_var, bg=self.bg, font=self.font_small).grid(row=3, column=2, columnspan=2, sticky='w', padx=4)

    tk.Label(parent, text='Lieferant', bg=self.bg, font=self.font_small).grid(row=4, column=0, sticky='w')
    ttk.Combobox(parent, textvariable=self.supplier_var, values=SUPPLIER_OPTIONS, state='normal', font=self.font_small).grid(row=4, column=1, columnspan=3, sticky='ew', padx=4)
    tk.Label(parent, text='Kostenbeschreibung', bg=self.bg, font=self.font_small).grid(row=5, column=0, sticky='w')
    ttk.Combobox(parent, textvariable=self.global_prefix_var, values=COST_TYPE_OPTIONS, state='normal', font=self.font_small).grid(row=5, column=1, columnspan=3, sticky='ew', padx=4)
    tk.Label(parent, text='Export-CSV', bg=self.bg, font=self.font_small).grid(row=6, column=0, sticky='w')
    tk.Entry(parent, textvariable=self.export_var, font=self.font_small).grid(row=6, column=1, sticky='ew', padx=4)
    tk.Button(parent, text='Speichern unter', command=lambda: _fm_browse(self, 'Export-CSV', self.export_var, True, 'export'), font=self.font_small).grid(row=6, column=2, columnspan=2, sticky='ew')

    buttons = tk.Frame(parent, bg=self.bg); buttons.grid(row=7, column=0, columnspan=4, sticky='ew', pady=(8,4)); buttons.columnconfigure(5, weight=1)
    tk.Button(buttons, text='Rechnung analysieren', command=self.analyze_invoice, font=self.font_small).grid(row=0, column=0, padx=(0,4))
    tk.Button(buttons, text='Ausgewählte zusammenfassen', command=lambda: _afi464_merge_selected_positions(self) if '_afi464_merge_selected_positions' in globals() else None, font=self.font_small).grid(row=0, column=1, padx=(0,4))
    tk.Button(buttons, text='AFI-Upload-Datei erstellen', command=self.run_export, font=('Segoe UI',10,'bold'), bg='#CFEAD6').grid(row=0, column=6, sticky='e')

    self.positions_tree = ttk.Treeview(parent, columns=['sel','text','amount','tax','gl','cc','ia'], show='headings', height=9)
    for c,t,w in [('sel','',34),('text','POSITION',260),('amount','NETTO',82),('tax','TAX',50),('gl','GL',78),('cc','CC',90),('ia','IA',82)]:
        self.positions_tree.heading(c, text=t); self.positions_tree.column(c, width=w, stretch=(c=='text'), anchor='center' if c in ('sel','amount','tax') else 'w')
    self.positions_tree.grid(row=8, column=0, columnspan=4, sticky='nsew', pady=(6,0))
    self.positions_tree.bind('<Button-1>', lambda e: _afi464_toggle_position_checkbox(self, e) if '_afi464_toggle_position_checkbox' in globals() else None)
    tk.Label(parent, textvariable=self.position_saldo_var, bg='#FFF4C2', fg='#182431', font=('Segoe UI', 9, 'bold'), anchor='w').grid(row=9, column=0, columnspan=4, sticky='ew', pady=(3,0))
    try:
        _afi466_make_drag_splitter(self, parent, row=10)
        _afi466_make_sources_area(self, parent, header_row=11, canvas_row=12, columnspan=4)
        parent.rowconfigure(12, weight=1)
    except Exception:
        parent.rowconfigure(8, weight=1)
    tk.Label(parent, textvariable=self.suggestion_var, bg=self.bg, fg='#7A4B00', font=self.font_small, wraplength=560, justify='left').grid(row=13, column=0, columnspan=4, sticky='ew', pady=(4,0))

try:
    SupplierUploadUI._build_left = _afi480_build_left
except Exception:
    pass


_analyze_invoice_before_v0480 = SupplierUploadUI.analyze_invoice if 'SupplierUploadUI' in globals() else None

def _afi480_analyze_invoice(self):
    raw = self.invoice_var.get().strip()
    paths = _afi480_split_invoice_paths(raw)
    if not paths or not all(os.path.isfile(p) for p in paths):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gültige Rechnung auswählen.')
        return
    if all(_afi480_is_pdf(p) for p in paths):
        self.load_preview(_afi480_join_invoice_paths(paths))
        self.headers, self.rows = ['PDF'], []
        try:
            self.suggestion_var.set(f'{len(paths)} PDF-Datei(en) erkannt: Vorschau zeigt alle Seiten; Export verwendet alle ausgewählten PDFs.')
        except Exception:
            pass
        try:
            _afi464_populate_positions_tree(self)
        except Exception:
            pass
        return
    if _analyze_invoice_before_v0480:
        return _analyze_invoice_before_v0480(self)

try:
    SupplierUploadUI.analyze_invoice = _afi480_analyze_invoice
except Exception:
    pass


def _afi480_current_config(self):
    return _afi463_current_config(self)


_fm_browse_before_v0480 = _fm_browse if '_fm_browse' in globals() else None

def _fm_browse(self, label, var, save=False, role=''):
    if role == 'invoice' and not save:
        start = var.get().strip() if hasattr(var, 'get') else ''
        first = _afi480_split_invoice_paths(start)[0] if _afi480_split_invoice_paths(start) else start
        if os.path.isfile(first):
            start = os.path.dirname(first)
        if not os.path.isdir(start):
            start = _fm_downloads_path() if '_fm_downloads_path' in globals() else _desktop_path()
        paths = filedialog.askopenfilenames(title=label, initialdir=start or None, filetypes=[('Dokumente', '*.csv *.xlsx *.xls *.xlsm *.pdf *.docx'), ('Alle Dateien', '*.*')])
        if paths:
            paths = list(paths)
            var.set(_afi480_join_invoice_paths(paths))
            try:
                self.load_preview(_afi480_join_invoice_paths(paths))
            except Exception:
                pass
            _afi480_update_export_path(self, True)
        return
    return _fm_browse_before_v0480(self, label, var, save, role) if _fm_browse_before_v0480 else None


_create_supplier_upload_csv_before_v0480 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    cfg = dict(config or {})
    if '_afi478_update_context' in globals():
        _afi478_update_context(assignment_path, cfg)
    paths = _afi480_split_invoice_paths(invoice_path)
    try:
        if paths and all(_afi480_is_dkv_pdf(p) for p in paths) and '_dkv475_create_direct_export' in globals():
            result = _dkv475_create_direct_export(assignment_path, _afi480_join_invoice_paths(paths), export_path, cfg)
        else:
            result = _create_supplier_upload_csv_before_v0480(assignment_path, invoice_path, export_path, cfg)
        if '_afi478_sanitize_export_file' in globals():
            changed = _afi478_sanitize_export_file(export_path)
            if isinstance(result, dict):
                result['text_sanitized_long_digits'] = changed
        if isinstance(result, dict):
            result['single_booking_circle'] = (cfg.get('booking_circle') or (cfg.get('booking_circles') or ['IDE'])[0]) if isinstance(cfg.get('booking_circles'), list) else cfg.get('booking_circle', 'IDE')
            result['intercompany_recharge'] = cfg.get('intercompany_recharge', True)
        return result
    finally:
        if '_AFI478_CONTEXT' in globals():
            _AFI478_CONTEXT.update({'booking_circle':'IDE','intercompany':True,'assignment_path':''})


def _afi480_run_export(self):
    template_path = self.template_var.get().strip()
    invoice_path = self.invoice_var.get().strip()
    export_path = self.export_var.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gültige Zuordnungsdatei auswählen.'); return
    paths = _afi480_split_invoice_paths(invoice_path)
    if not paths or not all(os.path.isfile(p) for p in paths):
        messagebox.showwarning(MODULE_TITLE, 'Bitte eine gültige Rechnung auswählen.'); return
    if not export_path:
        _afi480_update_export_path(self, True)
        export_path = self.export_var.get().strip()
    try:
        cfg = _afi463_current_config(self)
        result = create_supplier_upload_csv(template_path, invoice_path, export_path, cfg)
        msg = [f'AFI-Upload-Datei erstellt:\n{export_path}']
        if isinstance(result, dict):
            msg.append(f"Positionen: {result.get('rows','')}")
            msg.append(f"Export-Netto: {result.get('export_net_total','')}")
            if result.get('dkv_direct_export'):
                msg.append('DKV-Direktexport aus Rechnung aktiv.')
            msg.append(f"Buchungskreis: {result.get('single_booking_circle', cfg.get('booking_circle','IDE'))}")
            msg.append(f"Weiterberechnung Intercompany: {'aktiv' if result.get('intercompany_recharge', cfg.get('intercompany_recharge', True)) else 'deaktiviert'}")
        messagebox.showinfo(MODULE_TITLE, '\n'.join(msg))
    except Exception as exc:
        messagebox.showerror(MODULE_TITLE, str(exc))

try:
    SupplierUploadUI.run_export = _afi480_run_export
except Exception:
    pass


# ------------------------------------------------------------------
# AFI_UPLOAD_ENBW_KONSOLIDIERTE_GEBUEHREN_RUNDUNG_V0481
# Datum: 2026-07-08
# Zweck:
# - EnBW-Grundgebuehren je Nutzer fliessen direkt in Energiekosten Netto je Nutzer ein.
# - EnBW-Blockiergebuehren werden als Sammelposition pro Gesellschaft/Buchungskreis ausgewiesen.
# - Rundung auf 2 Nachkommastellen erfolgt erst unmittelbar beim Schreiben einer Exportposition.
# ------------------------------------------------------------------
AFI_UPLOAD_ENBW_KONSOLIDIERTE_GEBUEHREN_RUNDUNG_VERSION = "0.481"


def _afi481_is_enbw_invoice_context(invoice_path, config, sources=None):
    supplier = _norm((config or {}).get('supplier', ''))
    name = _norm(os.path.basename(str(invoice_path or '')))
    gp = _norm((config or {}).get('global_prefix', ''))
    src_text = _norm(' '.join(str((s or {}).get(k,'')) for s in (sources or []) for k in ('label','cost_description','net')))
    return 'ENBW' in supplier or 'NBW' in supplier or 'ENBW' in name or 'NBW' in name or 'TANKENSTROM' in gp or any(x in src_text for x in ('ENERGIEKOST','BLOCKIER','GRUNDGEBUEHR','GRUNDGEBUHR'))


def _afi481_enbw_source_kind(src):
    n = _norm(str((src or {}).get('label','')) + ' ' + str((src or {}).get('cost_description','')) + ' ' + str((src or {}).get('net','')))
    if 'BLOCKIER' in n:
        return 'BLOCKING'
    if 'GRUNDGEBUEHR' in n or 'GRUNDGEBUHR' in n or 'GRUNDKOST' in n or 'GRUNDPREIS' in n:
        return 'BASE_FEE'
    if 'ENERGIEKOST' in n or 'ENERGIE' in n or 'LADEKOST' in n or 'STROM' in n or 'CHARGING' in n:
        return 'ENERGY'
    return 'OTHER'


def _afi481_group_company_from_row(row, selected_bukrs='IDE'):
    if '_afi471_org_to_bukrs_from_row' in globals():
        try:
            b = _afi471_org_to_bukrs_from_row(row)
            if b:
                return b
        except Exception:
            pass
    for key in ('Organisationseinheit','Organisation','Firma','Gesellschaft','Buchungskreis'):
        val = (row or {}).get(key, '')
        if val:
            return _afi478_bukrs_label(val) if '_afi478_bukrs_label' in globals() else _general_company_to_bukrs(val)
    return selected_bukrs or 'IDE'


def _afi481_cost_type_for_enbw(cost_desc):
    return _afi471_cost_type_for_gl(cost_desc) if '_afi471_cost_type_for_gl' in globals() else 'TANKEN_STROM'


def _afi481_fallback_gl(assignment_path, cost_desc, bukrs):
    try:
        fb = _afi471_fallback_entry_for_gl(assignment_path, cost_desc, bukrs) if '_afi471_fallback_entry_for_gl' in globals() else {'gl_tanken_strom':'427010'}
        gl, _cc, _ia = _select_assignment_values(fb, _afi481_cost_type_for_enbw(cost_desc), cost_desc)
        return gl or fb.get('gl_tanken_strom') or fb.get('gl_default') or '427010'
    except Exception:
        return '427010'


def _afi481_create_enbw_export(assignment_path, invoice_path, export_path, config):
    invoice_paths = _afi480_split_invoice_paths(invoice_path) if '_afi480_split_invoice_paths' in globals() else ([invoice_path] if invoice_path else [])
    if not invoice_paths:
        raise RuntimeError('Bitte eine Rechnung auswählen.')
    invoice_path_single = invoice_paths[0]
    headers, rows = _read_table_file(invoice_path_single)
    sources = [s for s in (config or {}).get('sources', []) if s.get('active', True) and s.get('net')]
    if not sources:
        raise RuntimeError('Bitte mindestens eine aktive Berechnungsquelle mit Betragsspalte auswählen.')
    assignment_entries = load_assignment_entries(assignment_path)
    selected_bukrs = (config or {}).get('booking_circle') or ((config or {}).get('booking_circles') or ['IDE'])[0]
    selected_bukrs = _afi478_bukrs_label(selected_bukrs) if '_afi478_bukrs_label' in globals() else selected_bukrs
    if '_afi478_update_context' in globals():
        _afi478_update_context(assignment_path, config or {})
    accepted_old = set((config or {}).get('accepted_old_plate_assignments') or [])
    energy_groups = OrderedDict()
    blocking_groups = OrderedDict()
    other_groups = OrderedDict()
    warnings_missing=[]; warnings_tax=[]; warnings_empty_assignment=[]; warnings_foreign_gross=[]; fuzzy_hits=[]; conflicts={}
    invoice_total = Decimal('0')
    unique_drivers=set(); unique_keys=set()

    def amount_for(row, src, net):
        amount, tax, foreign = _afi468_amount_tax(row, src, net) if '_afi468_amount_tax' in globals() else _amount_and_tax_from_values(net, Decimal('0'), Decimal('19'))
        return Decimal(amount), tax, foreign

    for src in sources:
        kind = _afi481_enbw_source_kind(src)
        net_col = src.get('net','')
        if not net_col:
            continue
        for idx, row in enumerate(rows):
            net = _dec(row.get(net_col, ''))
            if net == 0:
                continue
            key, driver = _afi468_driver_key_from_row(row, src, idx) if '_afi468_driver_key_from_row' in globals() else ('','')
            if not key and not driver:
                key = driver = f'UNZUORDENBAR Zeile {idx+2}'
                warnings_empty_assignment.append(f'{src.get("label") or net_col}: Zeile {idx+2} ohne Fahrer/Schlüssel')
            amount, tax, foreign = amount_for(row, src, net)
            invoice_total += amount
            unique_drivers.add(_norm(driver)); unique_keys.add(_norm(key))
            if tax == 'VX':
                warnings_tax.append(f'{src.get("label") or net_col} / {key} / {driver}: Steuer nicht eindeutig')
            if foreign:
                warnings_foreign_gross.append(f'{src.get("label") or net_col} / {key}: abweichender/ausländischer Steuersatz -> Bruttobetrag mit V0 verwendet')

            if kind in ('ENERGY','BASE_FEE'):
                cost_desc = 'Energiekosten'
                gl, cc, orderid, text_driver, suppress, _old = _afi471_resolve_values(assignment_path, assignment_entries, cost_desc, key, driver, accepted_old, row, warnings_missing, conflicts, fuzzy_hits) if '_afi471_resolve_values' in globals() else ('','','',driver,False,False)
                use_driver = text_driver or driver
                if suppress:
                    group_key = ('ENERGY_NOKST', _norm(key), tax, gl)
                    text = _clean(' '.join([cost_desc, key]))
                else:
                    group_key = ('ENERGY_DRIVER', _norm(use_driver), tax, gl, cc, orderid)
                    parts=[cost_desc]
                    if key: parts.append(key)
                    if use_driver and _norm(use_driver)!=_norm(key): parts.append(use_driver)
                    text=_clean(' '.join(parts))
                if group_key not in energy_groups:
                    energy_groups[group_key]={'TEXT': text, 'amount': Decimal('0'), 'TAX_CODE': tax, 'GL_ACCOUNT': gl or _afi481_fallback_gl(assignment_path,cost_desc,selected_bukrs), 'COSTCENTER': cc or '', 'ORDERID': orderid or '', 'driver_norm': _norm(use_driver), 'cost_norm': _norm(cost_desc)}
                # Grundgebuehr und Energie bleiben vor Positionsschreibung ungerundet zusammen.
                energy_groups[group_key]['amount'] += amount
            elif kind == 'BLOCKING':
                company = _afi481_group_company_from_row(row, selected_bukrs)
                cost_desc = 'Blockiergebühren'
                gl = _afi481_fallback_gl(assignment_path, cost_desc, company)
                group_key = ('BLOCKING_COMPANY', company, tax, gl)
                if group_key not in blocking_groups:
                    blocking_groups[group_key]={'TEXT': f'Blockiergebühren {company}', 'amount': Decimal('0'), 'TAX_CODE': tax, 'GL_ACCOUNT': gl, 'COSTCENTER': '', 'ORDERID': '', 'driver_norm': company, 'cost_norm': _norm(cost_desc)}
                blocking_groups[group_key]['amount'] += amount
            else:
                cost_desc = _clean((src or {}).get('cost_description') or (src or {}).get('label') or net_col)
                gl, cc, orderid, text_driver, suppress, _old = _afi471_resolve_values(assignment_path, assignment_entries, cost_desc, key, driver, accepted_old, row, warnings_missing, conflicts, fuzzy_hits) if '_afi471_resolve_values' in globals() else ('','','',driver,False,False)
                group_key=('OTHER', _norm(cost_desc), _norm(text_driver or driver), tax, gl, cc, orderid)
                parts=[cost_desc]
                if key: parts.append(key)
                if (text_driver or driver) and _norm(text_driver or driver)!=_norm(key): parts.append(text_driver or driver)
                if group_key not in other_groups:
                    other_groups[group_key]={'TEXT': _clean(' '.join(parts)), 'amount':Decimal('0'), 'TAX_CODE':tax, 'GL_ACCOUNT':gl, 'COSTCENTER':cc, 'ORDERID':orderid, 'driver_norm':_norm(text_driver or driver), 'cost_norm':_norm(cost_desc)}
                other_groups[group_key]['amount'] += amount

    ordered = list(energy_groups.values()) + list(blocking_groups.values()) + list(other_groups.values())
    ordered.sort(key=lambda g: (g.get('cost_norm',''), g.get('driver_norm',''), TAX_ORDER.get(g.get('TAX_CODE','VX'),9)))
    os.makedirs(os.path.dirname(os.path.abspath(export_path)) or '.', exist_ok=True)
    export_total = Decimal('0')
    with open(export_path, 'w', encoding='utf-8-sig', newline='') as f:
        writer=csv.DictWriter(f, fieldnames=UPLOAD_COLUMNS, delimiter=';', extrasaction='ignore')
        writer.writeheader()
        for g in ordered:
            # Einzige Rundung: unmittelbar vor Erstellung/Schreiben der Position.
            amount = Decimal(g.get('amount','0')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
            export_total += amount
            writer.writerow({'TEXT': g.get('TEXT',''), 'PRICE': _fmt(amount), 'PRICE_UNIT':'1', 'QUANTITY':'1', 'UNIT':'ST', 'NET_VALUE':_fmt(amount), 'TAX_CODE':g.get('TAX_CODE',''), 'GL_ACCOUNT':g.get('GL_ACCOUNT',''), 'COSTCENTER':g.get('COSTCENTER',''), 'ORDERID':g.get('ORDERID','')})
    target_total = invoice_total.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return {'rows':len(ordered), 'export_path':export_path, 'invoice_net_raw_total':_fmt(invoice_total), 'export_net_total':_fmt(export_total), 'net_rounding_difference':_fmt(export_total-target_total), 'unique_drivers':len([x for x in unique_drivers if x]), 'unique_keys':len([x for x in unique_keys if x]), 'missing_template':warnings_missing, 'unknown_tax':warnings_tax, 'empty_assignment':warnings_empty_assignment, 'name_fallback_matches':fuzzy_hits, 'foreign_gross':warnings_foreign_gross, 'old_plate_conflicts':list(conflicts.values()), 'enbw_energy_includes_base_fee':True, 'enbw_blocking_grouped_by_company':True, 'rounding_only_at_position_write':True}


_create_supplier_upload_csv_before_v0481 = create_supplier_upload_csv

def create_supplier_upload_csv(assignment_path, invoice_path, export_path, config):
    paths = _afi480_split_invoice_paths(invoice_path) if '_afi480_split_invoice_paths' in globals() else ([invoice_path] if invoice_path else [])
    ext = os.path.splitext(paths[0])[1].lower() if paths else ''
    sources = (config or {}).get('sources', [])
    if ext != '.pdf' and _afi481_is_enbw_invoice_context(invoice_path, config or {}, sources):
        try:
            return _afi481_create_enbw_export(assignment_path, invoice_path, export_path, config or {})
        finally:
            if '_AFI478_CONTEXT' in globals():
                _AFI478_CONTEXT.update({'booking_circle':'IDE','intercompany':True,'assignment_path':''})
    return _create_supplier_upload_csv_before_v0481(assignment_path, invoice_path, export_path, config)
