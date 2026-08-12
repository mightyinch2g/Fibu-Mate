# -*- coding: utf-8 -*-
"""FiBu Mate Tool: Weiterberechnungen - Rechnungsgenerator.

Erzeugt je Rechnungs-Nr. WB und Gesellschaft eine Weiterberechnungsrechnung als DOCX und PDF,
liest/legt den Excel-Stammdaten-Reiter WB_Stammdaten an und schreibt Datum/Bildbeleg erstellt zurück.
"""

from __future__ import annotations

import os
import re
import sys
import math
import shutil
import subprocess
from dataclasses import dataclass
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
except Exception:
    Document = None

STANDARD_EXCEL_PATH = r"G:\BUC\Sachkonten-Kreditoren\Weiterberechnungen Company\Auflistung Weiterberechnung IDE-Eingangsrechnungen_Dauerarbeitsmappe_Serienbrief.xlsx"
STANDARD_OUTPUT_DIR = r"G:\BUC\Sachkonten-Kreditoren\Weiterberechnungen Company\Rechnungen Bildbeleg\Rechnungen_nicht_gebucht"
STAMMDATEN_SHEET = "WB_Stammdaten"
FIXED_TEMPLATE_PATH = r"C:\python\bin\tools\400326_Rechnung_Vodafone_SABU.docx"
DEFAULT_TEMPLATE = "400326_Rechnung_Vodafone_SABU.docx"
TOOL_TITLE = "Weiterberechnungen - Rechnungsgenerator"

DEFAULT_STAMMDATEN = [
    {"Kürzel": "IDG", "Gesellschaftsname": "INTERSPORT Digital GmbH", "Adresszeile 1": "Wannenäckerstr. 36", "Adresszeile 2": "74078 Heilbronn", "Adresszeile 3": "", "Adresszeile 4": "", "Lieferantennummer": "224711", "Aktiv": "ja"},
    {"Kürzel": "IMS", "Gesellschaftsname": "INTERSPORT Marketing Services GmbH", "Adresszeile 1": "Wannenäckerstr. 50", "Adresszeile 2": "74078 Heilbronn", "Adresszeile 3": "", "Adresszeile 4": "", "Lieferantennummer": "224712", "Aktiv": "ja"},
    {"Kürzel": "SABU", "Gesellschaftsname": "SABU GmbH", "Adresszeile 1": "Wannenäckerstr. 34", "Adresszeile 2": "74078 Heilbronn", "Adresszeile 3": "", "Adresszeile 4": "", "Lieferantennummer": "274743", "Aktiv": "ja"},
]

REQUIRED_COLUMNS = [
    "Rechnungsbetrag", "Nettowert", "MwSt-Betrag", "MwSt-Satz", "Rechnungsnummer",
    "Belegdatum", "Text", "Rechnungs-Nr. WB", "Datum", "Bildbeleg erstellt"
]


def _norm_header(value):
    return re.sub(r"\s+", " ", str(value or "").strip()).casefold()


def _safe_float(value):
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    txt = str(value).strip().replace("€", "").replace(" ", "")
    txt = txt.replace(".", "").replace(",", ".") if "," in txt else txt
    try:
        return float(txt)
    except Exception:
        return 0.0


def _money(value):
    value = Decimal(str(_safe_float(value))).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return s


def _date_value(value):
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if value in (None, ""):
        return None
    txt = str(value).strip()
    for fmt in ("%d.%m.%Y", "%d.%m.%y", "%m/%d/%Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(txt, fmt).date()
        except Exception:
            pass
    return None


def _ddmmyyyy(value):
    d = _date_value(value)
    return d.strftime("%d.%m.%Y") if d else ""


def _ddmmyy(value):
    d = _date_value(value)
    return d.strftime("%d.%m.%y") if d else ""


def _sanitize_filename(value):
    return re.sub(r"[^A-Za-z0-9ÄÖÜäöüß._ -]+", "_", str(value or "")).strip(" ._")[:140]


def _initials(app):
    raw = ""
    for attr in ("current_user_display", "current_user_key"):
        raw = str(getattr(app, attr, "") or "").strip()
        if raw:
            break
    raw = raw.replace(".", " ").replace("_", " ").replace(",", " ")
    parts = [p for p in raw.split() if p]
    if len(parts) >= 2:
        return (parts[0][0] + parts[-1][0]).upper()
    if parts:
        return parts[0][:2].upper()
    return "FM"


def _find_template_near_app(app):
    """Die Word-Vorlage ist fachlich fest vorgegeben und wird nicht mehr aktiv vom Nutzer gewählt."""
    return FIXED_TEMPLATE_PATH


def _detect_headers(ws):
    headers = {}
    row = 1
    for col in range(1, ws.max_column + 1):
        val = ws.cell(row, col).value
        name = str(val or "").strip()
        if not name:
            # Spezialfall Dauerarbeitsmappe: Kürzel steht rechts neben "Berechnung an" ohne eigene Überschrift.
            left = str(ws.cell(row, col - 1).value or "").strip() if col > 1 else ""
            if _norm_header(left) == "berechnung an":
                name = "Berechnung an Kürzel"
            else:
                continue
        headers[_norm_header(name)] = col
    return headers


def _col(headers, *names):
    for name in names:
        key = _norm_header(name)
        if key in headers:
            return headers[key]
    return None


def ensure_stammdaten_sheet(xlsx_path):
    if load_workbook is None:
        raise RuntimeError("openpyxl ist nicht verfügbar.")
    wb = load_workbook(xlsx_path)
    created = False
    if STAMMDATEN_SHEET not in wb.sheetnames:
        ws = wb.create_sheet(STAMMDATEN_SHEET)
        cols = list(DEFAULT_STAMMDATEN[0].keys())
        for c, h in enumerate(cols, 1):
            ws.cell(1, c).value = h
        for r, row in enumerate(DEFAULT_STAMMDATEN, 2):
            for c, h in enumerate(cols, 1):
                ws.cell(r, c).value = row.get(h, "")
        created = True
    else:
        ws = wb[STAMMDATEN_SHEET]
        existing = {str(ws.cell(r, 1).value or "").strip().upper() for r in range(2, ws.max_row + 1)}
        if ws.max_row < 1 or not ws.cell(1, 1).value:
            for c, h in enumerate(DEFAULT_STAMMDATEN[0].keys(), 1):
                ws.cell(1, c).value = h
        for item in DEFAULT_STAMMDATEN:
            if item["Kürzel"].upper() not in existing:
                r = ws.max_row + 1
                for c, h in enumerate(DEFAULT_STAMMDATEN[0].keys(), 1):
                    ws.cell(r, c).value = item.get(h, "")
                created = True
    wb.save(xlsx_path)
    return created


def load_stammdaten(wb):
    if STAMMDATEN_SHEET not in wb.sheetnames:
        raise RuntimeError(f"Stammdaten-Reiter '{STAMMDATEN_SHEET}' fehlt. Bitte zuerst 'Stammdaten anlegen/prüfen' ausführen.")
    ws = wb[STAMMDATEN_SHEET]
    headers = { _norm_header(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) }
    out = {}
    for r in range(2, ws.max_row + 1):
        code = str(ws.cell(r, headers.get(_norm_header("Kürzel"), 1)).value or "").strip().upper()
        if not code:
            continue
        nummer_col = headers.get(_norm_header("Lieferantennummer"))
        name_col = headers.get(_norm_header("Gesellschaftsname"))
        aktiv_col = headers.get(_norm_header("Aktiv"))
        aktiv = str(ws.cell(r, aktiv_col).value or "ja").strip().casefold() if aktiv_col else "ja"
        if aktiv in ("nein", "n", "false", "0"):
            continue
        adr = []
        if name_col:
            adr.append(str(ws.cell(r, name_col).value or "").strip())
        for h in ["Adresszeile 1", "Adresszeile 2", "Adresszeile 3", "Adresszeile 4"]:
            c = headers.get(_norm_header(h))
            val = str(ws.cell(r, c).value or "").strip() if c else ""
            if val:
                adr.append(val)
        out[code] = {
            "code": code,
            "name": str(ws.cell(r, name_col).value or "").strip() if name_col else code,
            "adresse": adr,
            "lieferantennummer": str(ws.cell(r, nummer_col).value or "").strip() if nummer_col else "",
        }
    return out


def read_invoice_groups(xlsx_path):
    wb = load_workbook(xlsx_path)
    ws = wb[wb.sheetnames[0]]
    headers = _detect_headers(ws)
    missing = []
    for col_name in REQUIRED_COLUMNS:
        if _col(headers, col_name, col_name.strip()) is None:
            missing.append(col_name)
    if missing:
        raise RuntimeError("Pflichtspalten fehlen: " + ", ".join(missing))
    stammdaten = load_stammdaten(wb)
    c_wb = _col(headers, "Rechnungs-Nr. WB")
    c_code = _col(headers, "Berechnung an Kürzel", "Berechnung an")
    c_calc = _col(headers, "Berechnung an")
    groups = {}
    for r in range(2, ws.max_row + 1):
        wb_no = ws.cell(r, c_wb).value
        if wb_no in (None, ""):
            continue
        code = str(ws.cell(r, c_code).value or "").strip().upper() if c_code else ""
        if not code or code.replace(".", "", 1).isdigit():
            # Wenn die Spalte "Berechnung an" numerisch ist, steckt das Kürzel in der rechten Nachbarspalte.
            if c_calc and c_calc + 1 <= ws.max_column:
                code = str(ws.cell(r, c_calc + 1).value or "").strip().upper()
        key = (str(wb_no).split(".")[0], code)
        groups.setdefault(key, {"wb": key[0], "code": code, "rows": [], "excel_rows": []})
        row = {"excel_row": r}
        for name in ["Lieferantenname", "Rechnungsbetrag", "Nettowert", "MwSt-Betrag", "MwSt-Satz", "Rechnungsnummer", "Belegdatum", "Text", "Datum", "Bildbeleg erstellt"]:
            c = _col(headers, name, name.strip())
            row[name] = ws.cell(r, c).value if c else None
        if not row.get("Lieferantenname"):
            c_lief = _col(headers, "Lieferantennummer")
            row["Lieferantenname"] = str(ws.cell(r, c_lief).value or "") if c_lief else ""
        groups[key]["rows"].append(row)
        groups[key]["excel_rows"].append(r)
    # Validieren und Stammdaten ergänzen
    for g in groups.values():
        if not g["code"]:
            raise RuntimeError(f"Zu Rechnungs-Nr. WB {g['wb']} fehlt die Gesellschaft/Kürzel-Spalte.")
        if g["code"] not in stammdaten:
            raise RuntimeError(f"Zu Gesellschaft '{g['code']}' fehlen aktive Stammdaten im Reiter {STAMMDATEN_SHEET}.")
        g["stamm"] = stammdaten[g["code"]]
        g["sum_rebetrag"] = sum(_safe_float(x.get("Rechnungsbetrag")) for x in g["rows"])
        g["sum_netto"] = sum(_safe_float(x.get("Nettowert")) for x in g["rows"])
        g["sum_mwst"] = sum(_safe_float(x.get("MwSt-Betrag")) for x in g["rows"])
    return wb, ws, headers, list(groups.values())


def _short_supplier_name(value):
    txt = str(value or "").strip()
    return txt or "Weiterberechnung"


def _make_subject(group):
    first = group["rows"][0]
    supplier = _short_supplier_name(first.get("Lieferantenname"))
    re_nr = str(first.get("Rechnungsnummer") or "").strip()
    bd = _ddmmyyyy(first.get("Belegdatum")) or _ddmmyy(first.get("Belegdatum"))
    return f"Weiterberechnung {supplier} Rechnung Nr. {re_nr} vom {bd}".strip()


def _create_docx(group, output_path, is_credit=False):
    if Document is None:
        raise RuntimeError("python-docx ist nicht verfügbar.")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7); sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.7); sec.bottom_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("INTERSPORT")
    run.bold = True; run.font.size = Pt(26)
    run.font.name = "Arial"

    detail = doc.add_paragraph("INTERSPORT Deutschland eG | Wannenäckerstraße 50 | 74078 Heilbronn")
    detail.runs[0].font.size = Pt(7)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(9.5); table.columns[1].width = Cm(7.0)
    addr = "\n".join(group["stamm"].get("adresse") or [])
    table.cell(0, 0).text = addr
    table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    inv_date = _date_value(group["rows"][0].get("Datum")) or date.today()
    right = table.cell(0, 1)
    right.text = "Unsere Zeichen\nNK/JN\n\nAnsprechpartner\nNicole Knobloch\n\nDurchwahl\n07131 288-4444\n\nTelefax\n07131 288-4236\n\nE-Mail\nNicole.Knobloch@intersport.de\n\nDatum\n" + inv_date.strftime("%d.%m.%Y")

    doc.add_paragraph("")
    doc.add_paragraph("Lieferantennummer")
    doc.add_paragraph(str(group["stamm"].get("lieferantennummer") or ""))
    doc.add_paragraph(_make_subject(group))
    typ = "Gutschrift Nr." if is_credit else "Rechnung Nr."
    p = doc.add_paragraph(f"{typ} {group['wb']}")
    p.runs[0].bold = True
    doc.add_paragraph("Wir berechnen Ihnen wie folgt: " if not is_credit else "Wir schreiben Ihnen wie folgt gut: ")

    pos_table = doc.add_table(rows=1, cols=5)
    pos_table.style = "Table Grid"
    hdr = pos_table.rows[0].cells
    for i, h in enumerate(["Pos.", "Beschreibung", "MwSt", "Nettowert", "MwSt-Betrag"]):
        hdr[i].text = h
    tax = {}
    for idx, row in enumerate(group["rows"], 1):
        cells = pos_table.add_row().cells
        rate = str(row.get("MwSt-Satz") or 0).replace(".0", "")
        cells[0].text = str(idx)
        cells[1].text = str(row.get("Text") or "")
        cells[2].text = f"{rate}%"
        cells[3].text = _money(row.get("Nettowert"))
        cells[4].text = _money(row.get("MwSt-Betrag"))
        tax.setdefault(rate, {"netto": 0.0, "mwst": 0.0})
        tax[rate]["netto"] += _safe_float(row.get("Nettowert"))
        tax[rate]["mwst"] += _safe_float(row.get("MwSt-Betrag"))

    doc.add_paragraph("")
    for rate, vals in sorted(tax.items(), key=lambda kv: _safe_float(kv[0])):
        doc.add_paragraph(f"Nettowert {rate}%\t\t\t€ {_money(vals['netto'])}")
        doc.add_paragraph(f"+ {rate}% Mehrwertsteuer\t\t€ {_money(vals['mwst'])}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Rechnungsbetrag\t\t\t€ {_money(group['sum_rebetrag'])}")
    run.bold = True; run.underline = True
    doc.add_paragraph("Zahlbar sofort rein netto")

    foot = doc.add_paragraph()
    foot.add_run("INTERSPORT Deutschland eG, ").bold = True
    foot.add_run("Wannenäckerstraße 50, 74078 Heilbronn\n")
    foot.add_run("Telefon: ").bold = True; foot.add_run("07131 288-0, ")
    foot.add_run("Telefax: ").bold = True; foot.add_run("07131 21257, ")
    foot.add_run("Mail: ").bold = True; foot.add_run("info@intersport.de, ")
    foot.add_run("Web: ").bold = True; foot.add_run("www.intersport.de\n")
    foot.add_run("Vorstand: ").bold = True; foot.add_run("Dr. Alexander von Preen, Thomas Storck, Henriette Tesch\n")
    foot.add_run("Amtsgericht: ").bold = True; foot.add_run("Stuttgart GNR 100124, USt-Id-Nr.: DE 145787599")
    for run in foot.runs:
        run.font.size = Pt(7)

    doc.save(output_path)


def _convert_docx_to_pdf(docx_path, pdf_path):
    # Primär: Microsoft Word per COM auf Windows.
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return
    except Exception as exc:
        last_exc = exc
    # Fallback: LibreOffice, falls vorhanden.
    try:
        outdir = os.path.dirname(os.path.abspath(pdf_path))
        subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, os.path.abspath(docx_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        generated = os.path.join(outdir, Path(docx_path).with_suffix(".pdf").name)
        if os.path.exists(generated) and os.path.abspath(generated) != os.path.abspath(pdf_path):
            os.replace(generated, pdf_path)
        if os.path.exists(pdf_path):
            return
    except Exception:
        pass
    raise RuntimeError("PDF konnte nicht erzeugt werden. Microsoft Word/COM oder LibreOffice erforderlich. Letzter Fehler: " + str(last_exc))


class WBGeneratorUI:
    def __init__(self, app):
        self.app = app
        self.excel_path = tk.StringVar(value=STANDARD_EXCEL_PATH)
        self.template_path = tk.StringVar(value=FIXED_TEMPLATE_PATH)
        self.output_dir = tk.StringVar(value=STANDARD_OUTPUT_DIR)
        self.options_visible = tk.BooleanVar(value=False)
        self.status = tk.StringVar(value="Bitte Excel-Datei auswählen oder Stammdaten prüfen.")
        self.selection = []
        self.tree = None
        self.options_frame = None

    def render(self):
        app = self.app
        root = app.root
        w = app.canvas.winfo_width(); h = app.canvas.winfo_height()
        frame = tk.Frame(root, bg=getattr(app, "BG", "#F6F8FA"))
        app.widget_items.append(frame)
        app.canvas.create_window(40, 128, window=frame, anchor="nw", width=max(900, w-80), height=max(520, h-190))

        tk.Label(frame, text=TOOL_TITLE, bg=frame["bg"], fg="#1f4e79", font=("Segoe UI", 16, "bold")).pack(anchor="w")
        tk.Label(frame, text="Erzeugt je Rechnungs-Nr. WB und Gesellschaft eine DOCX- und PDF-Rechnung, aktualisiert Datum/Bildbeleg erstellt und nutzt Stammdaten aus WB_Stammdaten.", bg=frame["bg"], fg="#333", font=("Segoe UI", 10)).pack(anchor="w", pady=(4, 10))

        tk.Label(frame, text="Standardpfade sind hinterlegt. Excel-Datei, feste Word-Vorlage und Ausgabeordner liegen unter 'Weitere Optionen'.", bg=frame["bg"], fg="#555", font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 6))

        opt_btn = tk.Button(frame, text="Weitere Optionen anzeigen" if not self.options_visible.get() else "Weitere Optionen ausblenden", command=self._toggle_options, bg="white", relief="solid", bd=1)
        opt_btn.pack(anchor="w", pady=(4, 2))
        self.options_frame = tk.Frame(frame, bg=frame["bg"])
        if self.options_visible.get():
            self.options_frame.pack(fill="x")
            self._row_file(self.options_frame, "Excel-Datei", self.excel_path, self._browse_excel)
            self._row_file(self.options_frame, "Word-Vorlage", self.template_path, self._browse_template, readonly=True)
            self._row_file(self.options_frame, "Ausgabeordner", self.output_dir, self._browse_output, folder=True)

        btns = tk.Frame(frame, bg=frame["bg"]); btns.pack(fill="x", pady=(10, 8))
        tk.Button(btns, text="Stammdaten anlegen/prüfen", command=self.ensure_master, bg="#e8f0fe", relief="solid", bd=1, padx=9).pack(side="left", padx=(0, 8))
        tk.Button(btns, text="WB-Vorschau laden", command=self.load_preview, bg="#e8f0fe", relief="solid", bd=1, padx=9).pack(side="left", padx=(0, 8))
        tk.Button(btns, text="Rechnungen erstellen", command=self.create_selected, bg="#CFEAD6", relief="solid", bd=1, padx=11).pack(side="left")

        cols = ("wb", "gesellschaft", "zeilen", "netto", "mwst", "betrag")
        self.tree = ttk.Treeview(frame, columns=cols, show="headings", height=13, selectmode="extended")
        for col, title, width in [("wb","WB-Nr.",90),("gesellschaft","Gesellschaft",110),("zeilen","Zeilen",55),("netto","Netto",100),("mwst","MwSt",100),("betrag","Betrag",110)]:
            self.tree.heading(col, text=title); self.tree.column(col, width=width, anchor="w")
        self.tree.pack(fill="both", expand=True)
        tk.Label(frame, textvariable=self.status, bg=frame["bg"], fg="#555", font=("Segoe UI", 9)).pack(anchor="w", pady=(8, 0))

    def _row_file(self, parent, label, var, command, folder=False, readonly=False):
        row = tk.Frame(parent, bg=parent["bg"]); row.pack(fill="x", pady=2)
        tk.Label(row, text=label, width=14, anchor="w", bg=parent["bg"], font=("Segoe UI", 9, "bold")).pack(side="left")
        state = "readonly" if readonly else "normal"
        tk.Entry(row, textvariable=var, state=state).pack(side="left", fill="x", expand=True, padx=(0, 6))
        if readonly:
            tk.Label(row, text="fest", bg=parent["bg"], fg="#666").pack(side="left")
        else:
            tk.Button(row, text="Auswählen", command=command, relief="solid", bd=1).pack(side="left")

    def _toggle_options(self):
        self.options_visible.set(not self.options_visible.get())
        self.app.render_page()

    def _browse_excel(self):
        p = filedialog.askopenfilename(title="WB-Excel auswählen", filetypes=[("Excel", "*.xlsx")])
        if p: self.excel_path.set(p)

    def _browse_template(self):
        # Vorlage ist bewusst fest eingebaut. Methode bleibt aus Kompatibilitätsgründen vorhanden.
        self.template_path.set(FIXED_TEMPLATE_PATH)
        try:
            messagebox.showinfo(TOOL_TITLE, "Die Word-Vorlage ist fest hinterlegt:\n\n" + FIXED_TEMPLATE_PATH)
        except Exception:
            pass


    def _browse_output(self):
        p = filedialog.askdirectory(title="Ausgabeordner auswählen")
        if p: self.output_dir.set(p)

    def ensure_master(self):
        try:
            if not self.excel_path.get():
                self.excel_path.set(STANDARD_EXCEL_PATH)
            if not self.excel_path.get(): return
            created = ensure_stammdaten_sheet(self.excel_path.get())
            self.status.set("Stammdaten-Reiter angelegt/ergänzt." if created else "Stammdaten-Reiter ist vorhanden.")
            messagebox.showinfo(TOOL_TITLE, self.status.get())
        except Exception as exc:
            messagebox.showerror(TOOL_TITLE, str(exc))

    def load_preview(self):
        try:
            if not self.excel_path.get():
                self.excel_path.set(STANDARD_EXCEL_PATH)
            if not self.excel_path.get(): return
            wb, ws, headers, groups = read_invoice_groups(self.excel_path.get())
            self.selection = groups
            self.tree.delete(*self.tree.get_children())
            for idx, g in enumerate(groups):
                self.tree.insert("", "end", iid=str(idx), values=(g["wb"], g["code"], len(g["rows"]), _money(g["sum_netto"]), _money(g["sum_mwst"]), _money(g["sum_rebetrag"])))
            self.status.set(f"{len(groups)} WB-Rechnung(en) geladen. Bitte auswählen und erstellen.")
        except Exception as exc:
            messagebox.showerror(TOOL_TITLE, str(exc))

    def create_selected(self):
        try:
            if not self.selection:
                self.load_preview()
            ids = self.tree.selection() if self.tree else []
            groups = [self.selection[int(i)] for i in ids] if ids else list(self.selection)
            if not groups:
                messagebox.showwarning(TOOL_TITLE, "Keine WB-Rechnung ausgewählt."); return
            outdir = self.output_dir.get() or STANDARD_OUTPUT_DIR
            os.makedirs(outdir, exist_ok=True)
            initials = _initials(self.app)
            today = date.today()
            created = []
            for g in groups:
                is_credit = False
                if g["sum_rebetrag"] < 0:
                    ans = messagebox.askyesnocancel(TOOL_TITLE, f"Rechnungs-Nr. WB {g['wb']} hat einen negativen Rechnungsbetrag.\n\nJa = als Gutschrift erstellen\nNein = als Rechnung mit Negativbetrag erstellen\nAbbrechen = überspringen")
                    if ans is None:
                        continue
                    is_credit = bool(ans)
                supplier = _sanitize_filename(_short_supplier_name(g["rows"][0].get("Lieferantenname")))
                kind = "Gutschrift" if is_credit else "Rechnung"
                base = _sanitize_filename(f"WB_{g['wb']}_{g['code']}_{supplier}_{today.strftime('%Y-%m-%d')}_{kind}")
                docx_path = os.path.join(outdir, base + ".docx")
                pdf_path = os.path.join(outdir, base + ".pdf")
                _create_docx(g, docx_path, is_credit=is_credit)
                _convert_docx_to_pdf(docx_path, pdf_path)
                created.append((g, docx_path, pdf_path))
            if created:
                wb = load_workbook(self.excel_path.get())
                ws = wb[wb.sheetnames[0]]
                headers = _detect_headers(ws)
                c_datum = _col(headers, "Datum", "Datum ")
                c_bild = _col(headers, "Bildbeleg erstellt")
                for g, _, _ in created:
                    for r in g["excel_rows"]:
                        if c_datum: ws.cell(r, c_datum).value = today
                        if c_bild: ws.cell(r, c_bild).value = initials
                wb.save(self.excel_path.get())
            self.status.set(f"{len(created)} Rechnung(en) erstellt. Excel aktualisiert mit {today.strftime('%d.%m.%Y')} / {initials}.")
            messagebox.showinfo(TOOL_TITLE, self.status.get())
        except Exception as exc:
            messagebox.showerror(TOOL_TITLE, str(exc))


def render(app):
    ui = WBGeneratorUI(app)
    ui.render()
    try:
        app.draw_bottom_logo()
    except Exception:
        pass
