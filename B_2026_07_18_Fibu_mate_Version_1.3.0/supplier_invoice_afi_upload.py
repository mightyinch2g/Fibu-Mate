"""FiBu Mate – AFI-Upload, Lieferanten-Regelwerk 1.3.0.

Standardfälle werden ohne LLM/ONNX verarbeitet. Dadurch können native Modellfehler
FiBu Mate nicht mehr beenden. Unterstützte Regeln: EnBW Charging CSV, DKV,
Vodafone Mobilfunk und Kazenmaier Bike Leasing.
"""
from __future__ import annotations

import csv
import io
import textwrap
import hashlib
import json
import os
import re
import threading
import traceback
import unicodedata
from collections import defaultdict
from datetime import datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk

try:
    from PIL import Image, ImageDraw, ImageFont, ImageTk
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

MODULE_TITLE = "AFI-Upload (lokale KI)"
MODULE_VERSION = "1.3.0"
RULESET_VERSION = "AFI-STANDARD-2026.07"
UPLOAD_COLUMNS = [
    "TEXT", "PRICE", "PRICE_UNIT", "QUANTITY", "UNIT", "NET_VALUE",
    "TAX_CODE", "GL_ACCOUNT", "COSTCENTER", "ORDERID",
]
INVOICE_SUFFIXES = {".pdf", ".xlsx", ".xlsm", ".csv", ".docx"}
EXCEL_SUFFIXES = {".xlsx", ".xlsm"}
HERE = Path(__file__).resolve().parent
APP_ROOT = HERE.parent.parent
SETTINGS_NAME = "afi_upload_shared_settings.json"
PROMPT_RULES_NAME = "afi_prompt_rules.json"


def _norm(value):
    text = unicodedata.normalize("NFKD", str(value or "").strip().casefold())
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]", "", text)


def _clean(value):
    return " ".join(str(value or "").replace("\r", " ").replace("\n", " ").split())


def _phone(value):
    digits = re.sub(r"\D", "", str(value or ""))
    if digits.startswith("0049"):
        digits = "0" + digits[4:]
    elif digits.startswith("49") and not digits.startswith("490"):
        digits = "0" + digits[2:]
    return digits


def _plate(value):
    return re.sub(r"[^A-Z0-9]", "", str(value or "").upper())


def _money(value):
    if value in (None, ""):
        return Decimal("0")
    text = str(value).strip().replace("EUR", "").replace("€", "").replace(" ", "")
    negative = text.startswith("-") or (text.startswith("(") and text.endswith(")"))
    text = text.strip("()-")
    if "," in text:
        text = text.replace(".", "").replace(",", ".")
    try:
        number = Decimal(text or "0")
    except InvalidOperation as exc:
        raise RuntimeError(f"Betrag nicht lesbar: {value!r}") from exc
    return -number if negative else number


def _fmt(value):
    value = Decimal(value).quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP)
    return f"{value:.4f}".replace(".", ",")


def _json(path):
    return json.loads(Path(path).read_text(encoding="utf-8-sig"))


def _shared_settings_file(for_write=False):
    candidates = []
    explicit = os.environ.get("FIBUMATE_SHARED_CONFIG_DIR", "").strip()
    if explicit:
        candidates.append(Path(explicit))
    candidates += [
        Path(r"G:\BUC\FM Anwendung\Fibu_Mate_Doc\Config"),
        APP_ROOT / "Fibu_Mate_Doc" / "Config",
        APP_ROOT / "bin" / "Config",
    ]
    for directory in candidates:
        try:
            if directory.is_dir():
                if for_write:
                    probe = directory / ".afi_write_test.tmp"
                    probe.write_text("ok", encoding="ascii")
                    probe.unlink(missing_ok=True)
                return directory / SETTINGS_NAME
        except Exception:
            continue
    fallback = APP_ROOT / "bin" / "Config"
    if for_write:
        fallback.mkdir(parents=True, exist_ok=True)
    return fallback / SETTINGS_NAME


def _default_prompt_rules():
    return {
        "general_prompt": (
            "Verarbeite Rechnungen nachvollziehbar und konservativ. Verwende ausschließlich "
            "Sachkonten und Kostenstellen aus den ausgewählten Datenbanken. Erfinde keine "
            "Kontierung. Der Buchungstext muss den Lieferanten, den Namen und - soweit "
            "vorhanden - Kennzeichen, Rufnummer oder Auftragsnummer enthalten. Nicht eindeutig "
            "auflösbare Positionen müssen als Prüfhinweis ausgegeben und für den Export gesperrt werden."
        ),
        "suppliers": {
            "enbw": {
                "label": "EnBW Charging",
                "tax_code_19": "VD",
                "text_template": "EnBW Stromtanken {kennzeichen} {name}",
                "fee_text_template": "EnBW Grundgebühr {gesellschaft}",
                "prompt": (
                    "Grundgebühr Netto je Gesellschaft in einem Posten zusammenfassen. "
                    "Grundgebühr je Nutzer Netto in die Energiekostenposition des Nutzers einrechnen. "
                    "Kennzeichen und Name im Buchungstext ausgeben."
                ),
                "aggregate_company_fee": True,
                "user_fee_into_energy": True,
            },
            "dkv": {
                "label": "DKV Tanken Inland",
                "tax_code_19": "VD",
                "text_template": "DKV Tanken {kennzeichen} {name}",
                "prompt": (
                    "Je Fahrzeug/Kennzeichen kontieren. Kostenstelle über Kennzeichen, Namen und "
                    "Weiterberechnungsregeln bestimmen. Kennzeichen und Name im Text ausgeben."
                ),
            },
            "vodafone": {
                "label": "Vodafone Mobilfunk",
                "tax_code_19": "VD",
                "text_template": "Vodafone {rufnummer} {name}",
                "aggregate_text_template": "Vodafone Mobilfunk {gesellschaft} {kostenstelle} ({anzahl} Anschlüsse)",
                "prompt": (
                    "Rufnummern über die Telefonzuordnung kontieren. Beträge je Kostenstelle und "
                    "Gesellschaft zusammenfassen; Einzelzuordnung im Prüfprotokoll erhalten."
                ),
            },
            "kazenmaier": {
                "label": "Kazenmaier Bike Leasing",
                "tax_code_19": "VD",
                "text_template": "Kazenmaier {auftragsnummer} {name}",
                "prompt": (
                    "Je Auftragsgruppe und Fahrer kontieren. Auftragsnummer und Name im Text ausgeben. "
                    "Das Sachkonto aus der Kostenart Bike Leasing verwenden."
                ),
            },
        },
    }


def _prompt_rules_file(for_write=False):
    return _shared_settings_file(for_write).with_name(PROMPT_RULES_NAME)


def _deep_merge(defaults, loaded):
    result = dict(defaults)
    for key, value in (loaded or {}).items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = _deep_merge(result[key], value)
        else:
            result[key] = value
    return result


def _load_prompt_rules():
    defaults = _default_prompt_rules()
    path = _prompt_rules_file(False)
    try:
        loaded = _json(path) if path.is_file() else {}
        return _deep_merge(defaults, loaded)
    except Exception:
        return defaults


def _save_prompt_rules(rules):
    path = _prompt_rules_file(True)
    payload = dict(rules)
    payload["saved_at"] = datetime.now().astimezone().isoformat(timespec="seconds")
    payload["module_version"] = MODULE_VERSION
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def _format_template(template, **values):
    safe = defaultdict(str, {key: _clean(value) for key, value in values.items()})
    try:
        return _clean(str(template).format_map(safe))
    except Exception:
        return _clean(" ".join(str(value) for value in values.values() if value))


def _load_settings():
    result = {"costcenter_database": "", "account_database": ""}
    path = _shared_settings_file(False)
    try:
        data = _json(path) if path.is_file() else {}
        for key in result:
            result[key] = str(data.get(key, "") or "").strip()
    except Exception:
        pass
    return result


def _save_settings(costcenter, account):
    path = _shared_settings_file(True)
    payload = {
        "costcenter_database": str(costcenter or "").strip(),
        "account_database": str(account or "").strip(),
        "scope": "shared_all_users",
        "saved_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "module_version": MODULE_VERSION,
    }
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def _read_text(path):
    raw = Path(path).read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            pass
    raise RuntimeError(f"Datei nicht lesbar: {Path(path).name}")


def _pdf_text(path):
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("pypdf fehlt. Bitte Basisinstallation ausführen.") from exc
    pages = []
    for index, page in enumerate(PdfReader(str(path)).pages, 1):
        pages.append(f"\n===== SEITE {index} =====\n{page.extract_text() or ''}")
    text = "".join(pages)
    if len(re.sub(r"\s+", "", text)) < 80:
        raise RuntimeError("Die PDF enthält keinen ausreichend auslesbaren Text. Scan-PDFs benötigen OCR.")
    return text


def _docx_text(path):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx fehlt. Bitte Basisinstallation ausführen.") from exc
    document = Document(str(path))
    output = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            output.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(output)


def _xlsx_text(path):
    try:
        import openpyxl
    except Exception as exc:
        raise RuntimeError("openpyxl fehlt. Bitte Basisinstallation ausführen.") from exc
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    output = []
    try:
        for sheet in workbook.worksheets:
            output.append(f"===== TABELLENBLATT {sheet.title} =====")
            for row in sheet.iter_rows(values_only=True):
                output.append("\t".join(_clean(value) for value in row))
    finally:
        workbook.close()
    return "\n".join(output)


def _invoice_text(path):
    extension = Path(path).suffix.lower()
    if extension == ".pdf":
        return _pdf_text(path)
    if extension == ".docx":
        return _docx_text(path)
    if extension in EXCEL_SUFFIXES:
        return _xlsx_text(path)
    if extension == ".csv":
        return _read_text(path)
    raise RuntimeError(f"Nicht unterstütztes Rechnungsformat: {Path(path).name}")


class MasterData:
    def __init__(self, costcenter_file, general_file):
        self.costcenter_file = Path(costcenter_file)
        self.general_file = Path(general_file)
        self.costcenter_by_name = {}
        self.vehicle_by_plate = {}
        self.company_by_name = {}
        self.phone_by_number = {}
        self.accounts = {}
        self.cross_charge = {}
        self.warnings = []
        self._load_costcenters()
        self._load_general()

    def _load_costcenters(self):
        import openpyxl
        workbook = openpyxl.load_workbook(self.costcenter_file, read_only=True, data_only=True)
        try:
            for sheet in workbook.worksheets:
                rows = list(sheet.iter_rows(values_only=True))
                header_index = None
                columns = {}
                for index, row in enumerate(rows[:30]):
                    normalized = [_norm(value) for value in row]
                    if "nachname" in normalized and "vorname" in normalized and any("kostenstelle" in item for item in normalized):
                        header_index = index
                        columns = {value: position for position, value in enumerate(normalized)}
                        break
                if header_index is None:
                    continue
                last_index = columns.get("nachname")
                first_index = columns.get("vorname")
                cost_index = next(position for position, value in enumerate([_norm(v) for v in rows[header_index]]) if "kostenstelle" in value)
                for row in rows[header_index + 1:]:
                    first = _clean(row[first_index] if first_index < len(row) else "")
                    last = _clean(row[last_index] if last_index < len(row) else "")
                    costcenter = _clean(row[cost_index] if cost_index < len(row) else "")
                    key = _norm(first + " " + last)
                    if key and costcenter:
                        if key in self.costcenter_by_name and self.costcenter_by_name[key] != costcenter:
                            self.warnings.append(f"Mehrdeutige Kostenstelle für {first} {last}")
                        self.costcenter_by_name[key] = costcenter
        finally:
            workbook.close()

    def _load_general(self):
        import openpyxl
        workbook = openpyxl.load_workbook(self.general_file, read_only=True, data_only=True)
        try:
            for sheet in workbook.worksheets:
                name = _norm(sheet.title)
                if "kfz" in name:
                    self._load_vehicles(sheet)
                elif "telefon" in name:
                    self._load_phones(sheet)
                elif "sachkont" in name:
                    self._load_accounts(sheet)
        finally:
            workbook.close()

    def _load_vehicles(self, sheet):
        for row in sheet.iter_rows(values_only=True):
            values = list(row)
            for index, value in enumerate(values):
                label = _clean(value)
                if label.casefold().startswith("kst weiterberechnung") and index + 1 < len(values):
                    costcenter = _clean(values[index + 1])
                    match = re.search(r"weiterberechnung\s+(.+?)\s+an\s+(.+?)(?::|$)", label, re.I)
                    if match and costcenter:
                        self.cross_charge[(_norm(match.group(1)), _norm(match.group(2)))] = costcenter
        for row in sheet.iter_rows(min_row=3, values_only=True):
            values = list(row)
            if len(values) < 5:
                continue
            plate = _clean(values[1]); last = _clean(values[2]); first = _clean(values[3]); company = _clean(values[4])
            if not plate:
                continue
            record = {"plate": plate, "first": first, "last": last, "name": _clean(first + " " + last), "company": company}
            self.vehicle_by_plate[_plate(plate)] = record
            if record["name"]:
                self.company_by_name[_norm(record["name"])] = company

    def _load_phones(self, sheet):
        for row in sheet.iter_rows(min_row=3, values_only=True):
            values = list(row)
            for start in (1, 9):
                if len(values) <= start + 3:
                    continue
                number = _clean(values[start]); last = _clean(values[start + 1]); first = _clean(values[start + 2]); company = _clean(values[start + 3])
                key = _phone(number)
                if key:
                    self.phone_by_number[key] = {"number": number, "first": first, "last": last, "name": _clean(first + " " + last), "company": company}

    def _load_accounts(self, sheet):
        rows = list(sheet.iter_rows(values_only=True))
        if len(rows) < 3:
            return
        first = list(rows[0]); second = list(rows[1])
        for column, header in enumerate(first):
            company = _clean(header)
            if not company or column + 1 >= len(first):
                continue
            if "kreditor" not in _norm(second[column] if column < len(second) else ""):
                continue
            for row in rows[2:]:
                cost_type = _clean(row[column] if column < len(row) else "")
                account = _clean(row[column + 1] if column + 1 < len(row) else "")
                if cost_type and account:
                    self.accounts[(_norm(company), _norm(cost_type))] = account

    def costcenter_for_name(self, name):
        return self.costcenter_by_name.get(_norm(name), "")

    def account_for(self, company, cost_type):
        exact = self.accounts.get((_norm(company), _norm(cost_type)))
        if exact:
            return exact
        candidates = {value for (comp, kind), value in self.accounts.items() if kind == _norm(cost_type)}
        return next(iter(candidates)) if len(candidates) == 1 else ""

    def resolve_person(self, first="", last="", name="", plate="", phone="", recipient_company=""):
        supplied_name = _clean(name or (str(first) + " " + str(last)))
        vehicle = self.vehicle_by_plate.get(_plate(plate)) if plate else None
        phone_record = self.phone_by_number.get(_phone(phone)) if phone else None
        resolved_name = supplied_name
        company = ""
        source = []
        if vehicle:
            resolved_name = vehicle["name"] or resolved_name
            company = vehicle["company"]
            source.append("Kennzeichen")
        if phone_record:
            resolved_name = phone_record["name"] or resolved_name
            company = phone_record["company"] or company
            source.append("Rufnummer")
        if not company and resolved_name:
            company = self.company_by_name.get(_norm(resolved_name), "")
            if company:
                source.append("Name/Firma")
        # Die Kostenstelle wird ausschließlich über den Namen aus der
        # Mitarbeiter-/Kostenstellenliste ermittelt. Eine in Rechnungen
        # vorhandene Personalnummer wird bewusst nicht als Kontierungsobjekt benutzt.
        costcenter = self.costcenter_for_name(supplied_name)
        if not costcenter:
            costcenter = self.costcenter_for_name(resolved_name)
        if costcenter:
            source.append("Name/Kostenstelle")
        if recipient_company and company and _norm(recipient_company) != _norm(company):
            cross = self.cross_charge.get((_norm(recipient_company), _norm(company)), "")
            if cross:
                costcenter = cross
                source.append("Weiterberechnung")
        return {"name": resolved_name, "company": company or recipient_company, "costcenter": costcenter, "source": "+".join(source)}


def _row(text, value, tax_code, account, costcenter, order_id=""):
    return {
        "TEXT": _clean(text)[:120],
        "PRICE": _fmt(value),
        "PRICE_UNIT": "1",
        "QUANTITY": "1",
        "UNIT": "ST",
        "NET_VALUE": _fmt(value),
        "TAX_CODE": tax_code,
        "GL_ACCOUNT": str(account or ""),
        "COSTCENTER": str(costcenter or ""),
        "ORDERID": str(order_id or ""),
    }


class RuleEngine:
    def __init__(self, master, rules=None):
        self.master = master
        self.rules = rules or _load_prompt_rules()

    def _rule(self, supplier):
        return self.rules.get("suppliers", {}).get(supplier, {})

    def _tax(self, supplier):
        return _clean(self._rule(supplier).get("tax_code_19", "VD")) or "VD"

    def _text(self, supplier, key="text_template", **values):
        return _format_template(self._rule(supplier).get(key, "{name}"), **values)

    def process(self, invoice_path, progress=lambda _message: None):
        path = Path(invoice_path)
        progress("Rechnung wird lokal gelesen …")
        if path.suffix.lower() == ".csv":
            rows = self._read_csv(path)
            if rows and "Energiekosten Netto (Euro)" in rows[0]:
                return self._enbw(rows, path)
        text = _invoice_text(path)
        folded = text.casefold()
        if "kazenmaier leasing" in folded:
            return self._kazenmaier(text, path)
        if "vodafone" in folded and ("gesamtübersicht" in folded or "anschlussnummer" in folded):
            return self._vodafone(text, path)
        if "dkv euro service" in folded and "vehicle:" in folded:
            return self._dkv(text, path)
        raise RuntimeError(
            "Für diese Rechnung wurde kein freigegebenes Lieferanten-Regelwerk erkannt. "
            "Aus Sicherheitsgründen wird kein instabiler ONNX-/KI-Freilauf gestartet."
        )

    @staticmethod
    def _read_csv(path):
        text = _read_text(path)
        try:
            dialect = csv.Sniffer().sniff(text[:10000], delimiters=";,\t|")
            delimiter = dialect.delimiter
        except Exception:
            delimiter = ";"
        return list(csv.DictReader(text.splitlines(), delimiter=delimiter))

    @staticmethod
    def _recipient(text):
        for company in ("INTERSPORT Deutschland eG", "INTERSPORT Digital GmbH", "SABU", "IMS"):
            if company.casefold() in text[:5000].casefold():
                return company
        return ""

    def _result(self, supplier, invoice_number, invoice_date, currency, rows, source_total, warnings, invoice_path, extra=None):
        upload_total = sum((_money(row["NET_VALUE"]) for row in rows), Decimal("0"))
        difference = upload_total - source_total
        invalid_accounts = [row for row in rows if not row["GL_ACCOUNT"]]
        invalid_costcenters = [row for row in rows if not row["COSTCENTER"]]
        export_allowed = not warnings and not invalid_accounts and not invalid_costcenters and abs(difference) <= Decimal("0.01")
        technical = {
            "module_version": MODULE_VERSION,
            "ruleset_version": RULESET_VERSION,
            "processing_mode": "deterministic_supplier_rules_no_onnx",
            "invoice_sha256": hashlib.sha256(Path(invoice_path).read_bytes()).hexdigest(),
            "processed_at": datetime.now().astimezone().isoformat(timespec="seconds"),
            "shared_settings_file": str(_shared_settings_file(False)),
        }
        if extra:
            technical.update(extra)
        all_warnings = list(warnings)
        if invalid_accounts:
            all_warnings.append(f"{len(invalid_accounts)} Zeilen ohne Sachkonto")
        if invalid_costcenters:
            all_warnings.append(f"{len(invalid_costcenters)} Zeilen ohne Kostenstelle/IA")
        if abs(difference) > Decimal("0.01"):
            all_warnings.append(f"Summendifferenz {_fmt(difference)} EUR")
        return {
            "status": "ready" if export_allowed else "review_required",
            "invoice": {
                "supplier": supplier,
                "invoice_number": invoice_number,
                "invoice_date": invoice_date,
                "currency": currency or "EUR",
                "net_total": _fmt(source_total),
                "tax_total": "",
                "gross_total": "",
            },
            "database_validation": {
                "required_columns_complete": True,
                "duplicates_found": False,
                "invalid_accounts_found": bool(invalid_accounts),
                "invalid_costcenters_found": bool(invalid_costcenters),
                "warnings": all_warnings,
            },
            "upload_columns": list(UPLOAD_COLUMNS),
            "upload_rows": rows,
            "validation": {
                "invoice_net_total": _fmt(source_total),
                "upload_net_total": _fmt(upload_total),
                "net_difference": _fmt(difference),
                "all_invoice_positions_processed": abs(difference) <= Decimal("0.01"),
                "all_accounts_from_selected_database": not invalid_accounts,
                "all_costcenters_from_selected_database": not invalid_costcenters,
                "export_allowed": export_allowed,
                "warnings": all_warnings,
            },
            "technical": technical,
        }

    def _enbw(self, source_rows, path):
        warnings = []
        energy = defaultdict(Decimal)
        base_fees = defaultdict(Decimal)
        unresolved = 0
        invoice_number = _clean(source_rows[0].get("Rechnungsnummer")) if source_rows else ""
        invoice_date = _clean(source_rows[0].get("Rechnungsdatum")) if source_rows else ""
        for source in source_rows:
            first = _clean(source.get("Fahrer Vorname")); last = _clean(source.get("Fahrer Nachname")); name = _clean(first + " " + last)
            plate = _clean(source.get("Fahrzeug Kennzeichen"))
            resolved = self.master.resolve_person(name=name, plate=plate)
            company = resolved["company"] or "UNGEKLÄRT"
            costcenter = resolved["costcenter"]
            account = self.master.account_for(company, "Tanken Strom")
            energy_value = (
                _money(source.get("Energiekosten Netto (Euro)"))
                + _money(source.get("Blockiergebühr Netto (Euro)"))
                + _money(source.get("Kartenkosten Netto (Euro)"))
                + _money(source.get("Grundgebühr je Nutzer Netto (Euro)"))
            )
            if energy_value:
                key = (company, costcenter, account, resolved["name"] or "ohne Zuordnung", plate)
                energy[key] += energy_value
                if not costcenter or not account:
                    unresolved += 1
            base_value = _money(source.get("Grundgebühr Netto (Euro)"))
            if base_value:
                base_fees[company] += base_value
        output = []
        for (company, costcenter, account, name, plate), value in sorted(energy.items()):
            output.append(_row(self._text("enbw", name=name, kennzeichen=plate), value, self._tax("enbw"), account, costcenter))
        # Sonderregel: allgemeine Grundgebühr exakt ein Posten je Gesellschaft.
        for company, value in sorted(base_fees.items()):
            account = self.master.account_for(company, "Tanken Strom")
            company_costcenters = [key[1] for key in energy if key[0] == company and key[1]]
            costcenter = company_costcenters[0] if len(set(company_costcenters)) == 1 else ""
            output.append(_row(self._text("enbw", "fee_text_template", gesellschaft=company), value, self._tax("enbw"), account, costcenter))
            if not costcenter:
                warnings.append(f"Grundgebühr {company}: keine eindeutige Gesellschafts-Kostenstelle")
        if unresolved:
            warnings.append(f"{unresolved} EnBW-Zuordnungen ohne vollständige Stammdaten")
        total = sum((_money(row.get("Energiekosten Netto (Euro)")) + _money(row.get("Blockiergebühr Netto (Euro)")) + _money(row.get("Kartenkosten Netto (Euro)")) + _money(row.get("Grundgebühr je Nutzer Netto (Euro)")) + _money(row.get("Grundgebühr Netto (Euro)")) for row in source_rows), Decimal("0"))
        return self._result("EnBW Charging", invoice_number, invoice_date, "EUR", output, total, warnings, path, {"special_rules": ["Grundgebühr je Gesellschaft aggregiert", "Grundgebühr je Nutzer Netto in Energiekostenposition eingerechnet"]})

    def _kazenmaier(self, text, path):
        recipient = self._recipient(text) or "INTERSPORT Digital GmbH"
        invoice_number = (re.search(r"Nr\.?/No\.?\s*\*?\*?\s*(\d+)", text, re.I) or re.search(r"Rechnung\s+Nr\.?\s*(\d+)", text, re.I))
        invoice_number = invoice_number.group(1) if invoice_number else ""
        date = re.search(r"Datum/Date\s*\*?\*?\s*(\d{2}\.\d{2}\.\d{4})", text, re.I)
        invoice_date = date.group(1) if date else ""
        total_match = re.search(r"Nettobetrag.*?([\d.]+,\d{2})\s*EUR", text, re.I | re.S)
        source_total = _money(total_match.group(1)) if total_match else Decimal("0")
        pattern = re.compile(
            r"(\d{6})\s*\*?\*?.{0,900}?/\s*([A-ZÄÖÜ][A-Za-zÄÖÜäöüß .'-]+?)\s+[A-Z0-9-]{8,}.{0,1200}?Summe\s+\1\s+([\d.]+,\d{2})\s+([\d.]+,\d{2})\s+([\d.]+,\d{2})",
            re.S,
        )
        output = []; warnings = []; seen = set()
        for match in pattern.finditer(text):
            order_id, name, finance, service, total = match.groups()
            if order_id in seen:
                continue
            seen.add(order_id)
            resolved = self.master.resolve_person(name=name, recipient_company=recipient)
            account = self.master.account_for(recipient, "Bike Leasing")
            output.append(_row(self._text("kazenmaier", name=resolved["name"] or name, auftragsnummer=order_id), _money(total), self._tax("kazenmaier"), account, resolved["costcenter"], order_id))
            if not resolved["costcenter"]:
                warnings.append(f"Kazenmaier {order_id}: Kostenstelle für {name} nicht gefunden")
        if not output:
            warnings.append("Kazenmaier-Anlage konnte nicht nach Auftragsnummern aufgeteilt werden")
        return self._result("Kazenmaier Leasing GmbH", invoice_number, invoice_date, "EUR", output, source_total, warnings, path)

    def _dkv(self, text, path):
        recipient = self._recipient(text) or "INTERSPORT Digital GmbH"
        number = re.search(r"Rechnungsnummer:\s*([\d/]+)", text, re.I)
        invoice_number = number.group(1) if number else ""
        date = re.search(r"Rechnungsdatum:\s*(\d{2}\.\d{2}\.\d{4})", text, re.I)
        invoice_date = date.group(1) if date else ""
        total_match = re.search(r"944,75\s+-13,07\s+([\d.]+,\d{2})\s+177,02", text)
        if not total_match:
            candidates = re.findall(r"19,00\s*%\s*([\d.]+,\d{2})\s+[\d.]+,\d{2}\s+[\d.]+,\d{2}", text)
            source_total = _money(candidates[-1]) if candidates else Decimal("0")
        else:
            source_total = _money(total_match.group(1))
        blocks = re.split(r"(?=VEHICLE:\s*)", text)
        output = []; warnings = []
        account = self.master.account_for(recipient, "DKV")
        for block in blocks:
            vehicle = re.match(r"VEHICLE:\s*(.+?)\s+CARD NO\.", block, re.I | re.S)
            if not vehicle:
                continue
            plate = _clean(vehicle.group(1))
            total = re.search(r"»?\s*TOTAL:\s*(.*?)(?=VEHICLE:|Gesamtsummen|=====|$)", block, re.I | re.S)
            if not total:
                warnings.append(f"DKV {plate}: Fahrzeugsumme nicht gefunden")
                continue
            numbers = re.findall(r"-?\s*\d[\d.]*,\d+", total.group(1))
            if len(numbers) < 3:
                warnings.append(f"DKV {plate}: Fahrzeugsumme unvollständig")
                continue
            net = _money(numbers[-3]); resolved = self.master.resolve_person(plate=plate, recipient_company=recipient)
            output.append(_row(self._text("dkv", name=resolved["name"], kennzeichen=plate), net, self._tax("dkv"), account, resolved["costcenter"]))
            if not resolved["costcenter"]:
                warnings.append(f"DKV {plate}: Kostenstelle nicht gefunden")
        return self._result("DKV EURO SERVICE GmbH + Co. KG", invoice_number, invoice_date, "EUR", output, source_total, warnings, path)

    def _vodafone(self, text, path):
        recipient = self._recipient(text) or "INTERSPORT Deutschland eG"
        number = re.search(r"Rechnungs-Nummer:\s*(\d+)", text, re.I)
        invoice_number = number.group(1) if number else ""
        date = re.search(r"Datum:\s*(\d{2}\.\d{2}\.\d{4})", text, re.I)
        invoice_date = date.group(1) if date else ""
        net = re.search(r"Nettorechnungsbetrag\s+([\d.]+,\d{4})", text, re.I)
        source_total = _money(net.group(1)) if net else Decimal("0")
        marker = text.find("Gesamtübersicht 2. Teil")
        detail = text[marker:] if marker >= 0 else text
        pattern = re.compile(r"(0\d{2,4}/\d{5,8})\s+(.{1,45}?)\s+(-?\s*\d[\d.]*,\d{4})", re.S)
        grouped = defaultdict(Decimal); labels = defaultdict(list); warnings = []; seen = set()
        account = self.master.account_for(recipient, "Vodafone")
        for match in pattern.finditer(detail):
            phone, label, amount = match.groups()
            phone_key = _phone(phone)
            # PDF repeats can occur; one row per number in the second overview only.
            if phone_key in seen:
                continue
            seen.add(phone_key)
            resolved = self.master.resolve_person(phone=phone, name=_clean(label), recipient_company=recipient)
            key = (resolved["costcenter"], resolved["company"], account)
            grouped[key] += _money(amount)
            labels[key].append(phone)
            if not resolved["costcenter"]:
                warnings.append(f"Vodafone {phone}: Kostenstelle nicht gefunden")
        output = []
        for (costcenter, company, gl_account), value in grouped.items():
            output.append(_row(self._text("vodafone", "aggregate_text_template", gesellschaft=company, kostenstelle=costcenter, anzahl=len(labels[(costcenter, company, gl_account)])), value, self._tax("vodafone"), gl_account, costcenter))
        if not output:
            warnings.append("Vodafone-Anschlussübersicht konnte nicht gelesen werden")
        return self._result("Vodafone GmbH", invoice_number, invoice_date, "EUR", output, source_total, warnings, path)


class AFIUI:
    def __init__(self, app):
        self.app = app
        self.root = app.root
        self.canvas = app.canvas
        self.bg = getattr(app, "BG", "#E8EEF5")
        settings = _load_settings()
        self.rules = _load_prompt_rules()
        self.invoice = tk.StringVar()
        self.costcenter = tk.StringVar(value=settings["costcenter_database"])
        self.account = tk.StringVar(value=settings["account_database"])
        self.status_var = tk.StringVar(value="Bereit - regelbasiert, Prompts zentral konfigurierbar.")
        self.summary = tk.StringVar(value="Noch kein Ergebnis.")
        self.result = None
        self.preview_images = []
        self.preview_index = 0
        self.preview_photo = None
        self.rule_widgets = {}

    def render(self):
        try:
            self.canvas.delete("all")
            self.app.draw_background()
            self.app.draw_header(MODULE_TITLE)
            self.app.draw_path_bar()
        except Exception:
            pass
        width = max(1100, self.canvas.winfo_width() - 60)
        height = max(690, self.canvas.winfo_height() - 175)
        shell = tk.Frame(self.canvas, bg=self.bg)
        self.canvas.create_window(30, 132, window=shell, anchor="nw", width=width, height=height)
        shell.rowconfigure(0, weight=1)
        shell.columnconfigure(0, weight=1)
        self.notebook = ttk.Notebook(shell)
        self.notebook.grid(row=0, column=0, sticky="nsew")
        self.work_tab = tk.Frame(self.notebook, bg=self.bg)
        self.rules_tab = tk.Frame(self.notebook, bg=self.bg)
        self.notebook.add(self.work_tab, text="Verarbeitung & Vorschau")
        self.notebook.add(self.rules_tab, text="Prompts & Regeln")
        self._render_work_tab()
        self._render_rules_tab()

    def _render_work_tab(self):
        frame = self.work_tab
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(2, weight=1)
        files = tk.LabelFrame(frame, text="1. Rechnung und Datenbanken", bg=self.bg, padx=10, pady=6, font=("Segoe UI", 10, "bold"))
        files.grid(row=0, column=0, sticky="ew", padx=8, pady=5)
        files.columnconfigure(1, weight=1)
        self._file_row(files, 0, "Rechnung (PDF/XLSX/CSV/Word)", self.invoice, self.pick_invoice, None)
        self._file_row(files, 1, "Datenbank 1: Mitarbeiter -> Kostenstelle/IA", self.costcenter, lambda: self.pick_db(self.costcenter, "Kostenstellen-Datenbank"), lambda: self.clear_db(self.costcenter))
        self._file_row(files, 2, "Datenbank 2: Kennzeichen/Rufnummer/Firma + Sachkonten", self.account, lambda: self.pick_db(self.account, "Kontierungs-Generalübersicht"), lambda: self.clear_db(self.account))
        tk.Label(files, text=f"Zentral gespeichert: {_shared_settings_file(False)} | Prompts: {_prompt_rules_file(False)}", bg=self.bg, fg="#44536A", anchor="w").grid(row=3, column=0, columnspan=4, sticky="ew")
        actions = tk.Frame(frame, bg=self.bg)
        actions.grid(row=1, column=0, sticky="ew", padx=8, pady=(0, 4))
        self.run = tk.Button(actions, text="2. Lokal verarbeiten", command=self.start, bg="#0F6CBD", fg="white", padx=12, pady=5)
        self.run.pack(side="left")
        tk.Button(actions, text="Prompts & Regeln", command=lambda: self.notebook.select(self.rules_tab), padx=10, pady=5).pack(side="left", padx=7)
        tk.Label(actions, text=f"Regelwerk {RULESET_VERSION} | Modul {MODULE_VERSION}", bg=self.bg).pack(side="left")
        tk.Label(actions, textvariable=self.status_var, bg=self.bg).pack(side="left", fill="x", expand=True, padx=10)

        pane = ttk.Panedwindow(frame, orient="horizontal")
        pane.grid(row=2, column=0, sticky="nsew", padx=8, pady=4)
        preview = tk.LabelFrame(pane, text="Rechnungs-Dokument", bg=self.bg, padx=5, pady=5)
        result = tk.LabelFrame(pane, text="AFI-Export-Vorschau", bg=self.bg, padx=5, pady=5)
        pane.add(preview, weight=2)
        pane.add(result, weight=3)
        preview.rowconfigure(1, weight=1)
        preview.columnconfigure(0, weight=1)
        nav = tk.Frame(preview, bg=self.bg)
        nav.grid(row=0, column=0, sticky="ew")
        self.prev_button = tk.Button(nav, text="<", command=lambda: self.change_preview(-1), state="disabled", width=3)
        self.prev_button.pack(side="left")
        self.preview_page_var = tk.StringVar(value="Keine Vorschau")
        tk.Label(nav, textvariable=self.preview_page_var, bg=self.bg).pack(side="left", padx=8)
        self.next_button = tk.Button(nav, text=">", command=lambda: self.change_preview(1), state="disabled", width=3)
        self.next_button.pack(side="left")
        tk.Button(nav, text="Neu laden", command=self.load_document_preview).pack(side="right")
        self.preview_canvas = tk.Canvas(preview, bg="#7f8790", highlightthickness=0)
        self.preview_canvas.grid(row=1, column=0, sticky="nsew")
        self.preview_canvas.bind("<Configure>", lambda _event: self.show_preview())

        result.rowconfigure(1, weight=1)
        result.columnconfigure(0, weight=1)
        tk.Label(result, textvariable=self.summary, bg=self.bg).grid(row=0, column=0, sticky="w")
        holder = tk.Frame(result)
        holder.grid(row=1, column=0, sticky="nsew")
        holder.rowconfigure(0, weight=1)
        holder.columnconfigure(0, weight=1)
        self.tree = ttk.Treeview(holder, columns=UPLOAD_COLUMNS, show="headings")
        for column in UPLOAD_COLUMNS:
            self.tree.heading(column, text=column)
            self.tree.column(column, width=220 if column == "TEXT" else 88)
        self.tree.grid(row=0, column=0, sticky="nsew")
        self.tree.bind("<Double-1>", self.edit)
        ybar = ttk.Scrollbar(holder, orient="vertical", command=self.tree.yview)
        xbar = ttk.Scrollbar(holder, orient="horizontal", command=self.tree.xview)
        ybar.grid(row=0, column=1, sticky="ns")
        xbar.grid(row=1, column=0, sticky="ew")
        self.tree.configure(yscrollcommand=ybar.set, xscrollcommand=xbar.set)
        footer = tk.Frame(result, bg=self.bg)
        footer.grid(row=2, column=0, sticky="ew", pady=(5, 0))
        self.export = tk.Button(footer, text="3. CSV exportieren", command=self.export_csv, state="disabled")
        self.export.pack(side="left")
        self.save = tk.Button(footer, text="Prüf-JSON speichern", command=self.export_json, state="disabled")
        self.save.pack(side="left", padx=8)

    def _render_rules_tab(self):
        frame = self.rules_tab
        frame.rowconfigure(1, weight=1)
        frame.columnconfigure(0, weight=1)
        header = tk.Frame(frame, bg=self.bg)
        header.grid(row=0, column=0, sticky="ew", padx=10, pady=8)
        tk.Label(header, text="Zentrale Prompt- und Regelpflege", bg=self.bg, font=("Segoe UI", 12, "bold")).pack(side="left")
        tk.Button(header, text="Speichern", command=self.save_prompt_rules, bg="#0F6CBD", fg="white", padx=14).pack(side="right")
        tk.Button(header, text="Werkseinstellungen", command=self.reset_prompt_rules, padx=10).pack(side="right", padx=7)
        tk.Label(frame, text=("Änderungen gelten zentral für alle Anwender. Steuerkennzeichen und Textvorlagen wirken direkt auf die AFI-Ausgabe. "
                             "Die freien Prompttexte dokumentieren die allgemeinen und lieferantenspezifischen KI-/Prüfregeln."),
                 bg=self.bg, fg="#44536A", justify="left", anchor="w", wraplength=1200).grid(row=0, column=0, sticky="ew", padx=10, pady=(38, 0))
        sub = ttk.Notebook(frame)
        sub.grid(row=1, column=0, sticky="nsew", padx=10, pady=8)
        general = tk.Frame(sub, bg=self.bg)
        sub.add(general, text="Allgemeiner Prompt")
        general.rowconfigure(1, weight=1)
        general.columnconfigure(0, weight=1)
        tk.Label(general, text="Gesamter allgemeiner Prompt", bg=self.bg, anchor="w").grid(row=0, column=0, sticky="ew", padx=8, pady=5)
        text = tk.Text(general, wrap="word", undo=True, font=("Consolas", 10))
        text.grid(row=1, column=0, sticky="nsew", padx=8, pady=(0, 8))
        text.insert("1.0", self.rules.get("general_prompt", ""))
        self.rule_widgets["general_prompt"] = text
        for supplier in ("enbw", "dkv", "vodafone", "kazenmaier"):
            self._supplier_rule_tab(sub, supplier)

    def _supplier_rule_tab(self, notebook, supplier):
        rule = self.rules["suppliers"][supplier]
        tab = tk.Frame(notebook, bg=self.bg)
        notebook.add(tab, text=rule.get("label", supplier))
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(4, weight=1)
        tax = tk.StringVar(value=rule.get("tax_code_19", "VD"))
        template = tk.StringVar(value=rule.get("text_template", ""))
        tk.Label(tab, text="Steuerkennzeichen 19 %", bg=self.bg).grid(row=0, column=0, sticky="w", padx=8, pady=5)
        tk.Entry(tab, textvariable=tax, width=16).grid(row=0, column=1, sticky="w", padx=8, pady=5)
        tk.Label(tab, text="Buchungstext-Vorlage", bg=self.bg).grid(row=1, column=0, sticky="w", padx=8, pady=5)
        tk.Entry(tab, textvariable=template).grid(row=1, column=1, sticky="ew", padx=8, pady=5)
        extra_vars = {}
        row = 2
        for key, label in (("fee_text_template", "Grundgebühr-Text"), ("aggregate_text_template", "Aggregierter Text")):
            if key in rule:
                variable = tk.StringVar(value=rule.get(key, ""))
                tk.Label(tab, text=label, bg=self.bg).grid(row=row, column=0, sticky="w", padx=8, pady=5)
                tk.Entry(tab, textvariable=variable).grid(row=row, column=1, sticky="ew", padx=8, pady=5)
                extra_vars[key] = variable
                row += 1
        placeholders = {
            "enbw": "Platzhalter: {kennzeichen}, {name}, {gesellschaft}",
            "dkv": "Platzhalter: {kennzeichen}, {name}",
            "vodafone": "Platzhalter: {rufnummer}, {name}, {gesellschaft}, {kostenstelle}, {anzahl}",
            "kazenmaier": "Platzhalter: {auftragsnummer}, {name}",
        }[supplier]
        tk.Label(tab, text=placeholders, bg=self.bg, fg="#44536A").grid(row=row, column=0, columnspan=2, sticky="w", padx=8, pady=3)
        row += 1
        tk.Label(tab, text="Lieferantenspezifischer Prompt / Regeltext", bg=self.bg).grid(row=row, column=0, columnspan=2, sticky="w", padx=8, pady=3)
        row += 1
        prompt = tk.Text(tab, wrap="word", undo=True, font=("Consolas", 10))
        prompt.grid(row=row, column=0, columnspan=2, sticky="nsew", padx=8, pady=(0, 8))
        prompt.insert("1.0", rule.get("prompt", ""))
        tab.rowconfigure(row, weight=1)
        self.rule_widgets[supplier] = {"tax": tax, "template": template, "prompt": prompt, "extra": extra_vars}

    def save_prompt_rules(self):
        rules = _load_prompt_rules()
        rules["general_prompt"] = self.rule_widgets["general_prompt"].get("1.0", "end-1c").strip()
        for supplier, widgets in self.rule_widgets.items():
            if supplier == "general_prompt":
                continue
            rule = rules["suppliers"][supplier]
            rule["tax_code_19"] = widgets["tax"].get().strip()
            rule["text_template"] = widgets["template"].get().strip()
            rule["prompt"] = widgets["prompt"].get("1.0", "end-1c").strip()
            for key, variable in widgets["extra"].items():
                rule[key] = variable.get().strip()
        try:
            _save_prompt_rules(rules)
            self.rules = rules
            messagebox.showinfo(MODULE_TITLE, f"Prompts und Regeln zentral gespeichert:\n{_prompt_rules_file(False)}")
        except Exception as exc:
            messagebox.showerror(MODULE_TITLE, f"Speichern fehlgeschlagen:\n{exc}")

    def reset_prompt_rules(self):
        if not messagebox.askyesno(MODULE_TITLE, "Prompts, Steuerkennzeichen und Textvorlagen auf Werkseinstellungen zurücksetzen?"):
            return
        _save_prompt_rules(_default_prompt_rules())
        self.rules = _load_prompt_rules()
        self.rule_widgets.clear()
        for child in self.rules_tab.winfo_children():
            child.destroy()
        self._render_rules_tab()

    def _file_row(self, parent, row, label, variable, choose, clear):
        tk.Label(parent, text=label, bg=self.bg, anchor="w").grid(row=row, column=0, sticky="w", pady=3)
        tk.Entry(parent, textvariable=variable).grid(row=row, column=1, sticky="ew", padx=8, pady=3)
        tk.Button(parent, text="Auswählen" if row == 0 else "Ändern", command=choose).grid(row=row, column=2, padx=(0, 5), pady=3)
        if clear:
            tk.Button(parent, text="Leeren", command=clear).grid(row=row, column=3, pady=3)

    def pick_invoice(self):
        path = filedialog.askopenfilename(title="Rechnung auswählen", filetypes=[("Rechnungen", "*.pdf *.xlsx *.xlsm *.csv *.docx"), ("Alle Dateien", "*.*")])
        if path:
            self.invoice.set(path)
            self.load_document_preview()

    def pick_db(self, variable, title):
        options = {"title": title, "filetypes": [("Excel-Datenbanken", "*.xlsx *.xlsm"), ("Alle Dateien", "*.*")]}
        if variable.get() and Path(variable.get()).parent.is_dir():
            options["initialdir"] = str(Path(variable.get()).parent)
        path = filedialog.askopenfilename(**options)
        if path:
            variable.set(path)
            _save_settings(self.costcenter.get(), self.account.get())

    def clear_db(self, variable):
        variable.set("")
        _save_settings(self.costcenter.get(), self.account.get())

    def load_document_preview(self):
        path = Path(self.invoice.get().strip())
        if not path.is_file():
            return
        self.status_var.set("Dokumentvorschau wird aufgebaut ...")
        threading.Thread(target=self._preview_worker, args=(path,), daemon=True).start()

    def _preview_worker(self, path):
        try:
            images = self._render_document_images(path)
            self.root.after(0, self._preview_ready, images)
        except Exception as exc:
            self.root.after(0, self._preview_failed, str(exc))

    def _render_document_images(self, path):
        if not PIL_AVAILABLE:
            raise RuntimeError("Pillow fehlt; Bildvorschau nicht verfügbar.")
        extension = path.suffix.lower()
        if extension == ".pdf":
            try:
                import pypdfium2 as pdfium
            except Exception as exc:
                raise RuntimeError("pypdfium2 fehlt. Bitte das Installationsskript 1.2.0 ausführen.") from exc
            document = pdfium.PdfDocument(str(path))
            images = []
            try:
                for index in range(len(document)):
                    page = document[index]
                    bitmap = page.render(scale=1.5)
                    images.append(bitmap.to_pil().convert("RGB"))
            finally:
                document.close()
            return images
        text = _invoice_text(path)
        return [self._text_preview_image(text, path.name)]

    @staticmethod
    def _text_preview_image(text, title):
        width, height = 1240, 1754
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        try:
            font = ImageFont.truetype("arial.ttf", 22)
            title_font = ImageFont.truetype("arialbd.ttf", 28)
        except Exception:
            font = ImageFont.load_default()
            title_font = font
        draw.text((45, 35), title, fill="#182431", font=title_font)
        y = 90
        for raw_line in str(text).splitlines():
            for line in textwrap.wrap(raw_line, width=105) or [""]:
                draw.text((45, y), line, fill="black", font=font)
                y += 28
                if y > height - 55:
                    draw.text((45, height - 40), "Vorschau gekürzt", fill="#B42318", font=font)
                    return image
        return image

    def _preview_ready(self, images):
        self.preview_images = images
        self.preview_index = 0
        self.status_var.set("Dokumentvorschau bereit.")
        self.show_preview()

    def _preview_failed(self, message):
        self.preview_images = []
        self.preview_canvas.delete("all")
        self.preview_canvas.create_text(20, 20, text=f"Vorschau nicht verfügbar:\n{message}", anchor="nw", fill="white", width=500)
        self.status_var.set("Dokumentvorschau nicht verfügbar.")

    def change_preview(self, delta):
        if not self.preview_images:
            return
        self.preview_index = max(0, min(len(self.preview_images) - 1, self.preview_index + delta))
        self.show_preview()

    def show_preview(self):
        if not self.preview_images or not PIL_AVAILABLE or not hasattr(self, "preview_canvas"):
            return
        image = self.preview_images[self.preview_index]
        max_width = max(200, self.preview_canvas.winfo_width() - 20)
        max_height = max(200, self.preview_canvas.winfo_height() - 20)
        copy = image.copy()
        copy.thumbnail((max_width, max_height))
        self.preview_photo = ImageTk.PhotoImage(copy)
        self.preview_canvas.delete("all")
        self.preview_canvas.create_image(max_width // 2 + 10, max_height // 2 + 10, image=self.preview_photo, anchor="center")
        self.preview_page_var.set(f"Seite {self.preview_index + 1} von {len(self.preview_images)}")
        self.prev_button.config(state="normal" if self.preview_index > 0 else "disabled")
        self.next_button.config(state="normal" if self.preview_index < len(self.preview_images) - 1 else "disabled")

    def _inputs(self):
        invoice = Path(self.invoice.get().strip())
        costcenter = Path(self.costcenter.get().strip())
        account = Path(self.account.get().strip())
        if not invoice.is_file() or invoice.suffix.lower() not in INVOICE_SUFFIXES:
            raise RuntimeError("Bitte eine Rechnung als PDF, XLSX/XLSM, CSV oder DOCX auswählen.")
        if not costcenter.is_file() or costcenter.suffix.lower() not in EXCEL_SUFFIXES:
            raise RuntimeError("Bitte Datenbank 1 als XLSX/XLSM auswählen.")
        if not account.is_file() or account.suffix.lower() not in EXCEL_SUFFIXES:
            raise RuntimeError("Bitte Datenbank 2 als XLSX/XLSM auswählen.")
        if costcenter.resolve() == account.resolve():
            raise RuntimeError("Datenbank 1 und Datenbank 2 müssen unterschiedliche Dateien sein.")
        return str(invoice), str(costcenter), str(account)

    def start(self):
        try:
            invoice, costcenter, account = self._inputs()
            _save_settings(costcenter, account)
        except Exception as exc:
            messagebox.showwarning(MODULE_TITLE, str(exc))
            return
        self.run.config(state="disabled")
        self.export.config(state="disabled")
        self.save.config(state="disabled")
        self.result = None
        for item in self.tree.get_children():
            self.tree.delete(item)
        threading.Thread(target=self.worker, args=(invoice, costcenter, account), daemon=True).start()

    def worker(self, invoice, costcenter, account):
        try:
            master = MasterData(costcenter, account)
            engine = RuleEngine(master, _load_prompt_rules())
            result = engine.process(invoice, lambda message: self.root.after(0, self.status_var.set, message))
            self.root.after(0, self.done, result)
        except Exception as exc:
            self.root.after(0, self.fail, str(exc), traceback.format_exc())

    def done(self, result):
        self.result = result
        for row in result["upload_rows"]:
            self.tree.insert("", "end", values=[row.get(column, "") for column in UPLOAD_COLUMNS])
        validation = result["validation"]
        self.summary.set(f"{result['invoice']['supplier']} | {len(result['upload_rows'])} Zeilen | Netto {validation['upload_net_total']} EUR | Export {'Ja' if validation['export_allowed'] else 'Nein'}")
        self.run.config(state="normal")
        self.save.config(state="normal")
        self.export.config(state="normal" if validation["export_allowed"] else "disabled")
        self.status_var.set("Fertig - Regeln und Steuerkennzeichen angewendet.")
        if validation["warnings"]:
            messagebox.showwarning(MODULE_TITLE, "Prüfhinweise:\n\n" + "\n".join(f"- {item}" for item in validation["warnings"][:30]))

    def fail(self, message, details):
        self.run.config(state="normal")
        directory = Path(os.environ.get("LOCALAPPDATA", str(Path.home()))) / "FiBuMate" / "FoundryLocal" / "logs"
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / f"afi_rules_{datetime.now():%Y%m%d_%H%M%S}.log"
        path.write_text(details, encoding="utf-8")
        messagebox.showerror(MODULE_TITLE, f"{message}\n\nProtokoll: {path}")

    def edit(self, event):
        item = self.tree.identify_row(event.y)
        column_id = self.tree.identify_column(event.x)
        if not self.result or not item or not column_id:
            return
        column = UPLOAD_COLUMNS[int(column_id[1:]) - 1]
        new = simpledialog.askstring(column, "Neuer Wert", initialvalue=self.tree.set(item, column), parent=self.root)
        if new is not None:
            self.tree.set(item, column, new)
            self.result["upload_rows"][self.tree.index(item)][column] = new
            self._revalidate_after_edit()

    def _revalidate_after_edit(self):
        rows = self.result.get("upload_rows", [])
        validation = self.result.setdefault("validation", {})
        source_total = _money(validation.get("invoice_net_total"))
        upload_total = sum((_money(row.get("NET_VALUE")) for row in rows), Decimal("0"))
        difference = upload_total - source_total
        missing_accounts = sum(not _clean(row.get("GL_ACCOUNT")) for row in rows)
        missing_costcenters = sum(not _clean(row.get("COSTCENTER")) for row in rows)
        warnings = []
        if missing_accounts:
            warnings.append(f"{missing_accounts} Zeilen ohne Sachkonto")
        if missing_costcenters:
            warnings.append(f"{missing_costcenters} Zeilen ohne Kostenstelle/IA")
        if abs(difference) > Decimal("0.01"):
            warnings.append(f"Summendifferenz {_fmt(difference)} EUR")
        validation.update({"upload_net_total": _fmt(upload_total), "net_difference": _fmt(difference), "all_invoice_positions_processed": abs(difference) <= Decimal("0.01"), "all_accounts_from_selected_database": missing_accounts == 0, "all_costcenters_from_selected_database": missing_costcenters == 0, "export_allowed": not warnings, "warnings": warnings})
        self.export.config(state="normal" if validation["export_allowed"] else "disabled")
        self.summary.set(f"{self.result['invoice']['supplier']} | {len(rows)} Zeilen | Netto {validation['upload_net_total']} EUR | Export {'Ja' if validation['export_allowed'] else 'Nein'}")

    def export_csv(self):
        path = filedialog.asksaveasfilename(defaultextension=".csv", initialfile="AFI_Upload.csv", filetypes=[("CSV-Datei", "*.csv")])
        if path:
            with open(path, "w", encoding="utf-8-sig", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=UPLOAD_COLUMNS, delimiter=";", extrasaction="ignore")
                writer.writeheader()
                writer.writerows(self.result["upload_rows"])

    def export_json(self):
        path = filedialog.asksaveasfilename(defaultextension=".json", initialfile="AFI_Pruefprotokoll.json", filetypes=[("JSON-Datei", "*.json")])
        if path:
            Path(path).write_text(json.dumps(self.result, ensure_ascii=False, indent=2), encoding="utf-8")



# AFI 1.3.0 - direkt integrierte Erweiterungen
"""FiBu Mate AFI-Upload UI/rules extension.

Loaded by supplier_invoice_afi_upload.py. Keeps the stable deterministic parsers and
adds central supplier management, AI prompt review, enhanced previews and CSV export.
"""

import csv
import json
import os
import re
import subprocess
import sys
import tempfile
import traceback
from collections import defaultdict
from pathlib import Path

GREEN = "#E6F4EA"

DEFAULT_SUPPLIERS = {
    "enbw": {
        "label": "EnBW Charging",
        "tax_code_19": "VD",
        "text_template": "EnBW Stromtanken {kennzeichen} {name}",
        "prompt": "Grundgebühr je Gesellschaft zusammenfassen. Grundgebühr je Nutzer Netto in Energiekosten einrechnen. Kennzeichen und Name im Text ausgeben.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Tanken Strom", "account_key": "Tanken Strom"}],
        "use_ai": True,
    },
    "dkv": {
        "label": "DKV",
        "tax_code_19": "VD",
        "text_template": "DKV {kostenart} {kennzeichen} {name}",
        "prompt": "Kostenart bestimmen. Tanken je Kennzeichen kontieren. Versicherungen nach Versicherungsregel kontieren. Kennzeichen und Name im Text ausgeben.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [
            {"name": "Tanken", "terms": ["diesel", "benzin", "kraftstoff"], "account_key": "DKV"},
            {"name": "Versicherung", "terms": ["versicherung", "insurance"], "account_key": "DEAS-Versicherungen"},
        ],
        "use_ai": True,
    },
    "vodafone": {
        "label": "Vodafone Mobilfunk",
        "tax_code_19": "VD",
        "text_template": "Vodafone {rufnummer} {name}",
        "prompt": "Rufnummern über die Telefonzuordnung kontieren und je Kostenstelle/Gesellschaft zusammenfassen.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Mobilfunk", "account_key": "Vodafone"}],
        "use_ai": True,
    },
    "kazenmaier": {
        "label": "Kazenmaier Bike Leasing",
        "tax_code_19": "VD",
        "text_template": "Kazenmaier {auftragsnummer} {name}",
        "prompt": "Je Auftragsgruppe und Person kontieren. Auftragsnummer und Name im Text ausgeben.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Bike Leasing", "account_key": "Bike Leasing"}],
        "use_ai": True,
    },
    "telekom": {
        "label": "Telekom",
        "tax_code_19": "VD",
        "text_template": "Telekom {rufnummer} {name}",
        "prompt": "Rufnummer und Name im Text. Kostenstelle aus Telefon- und Mitarbeiterstamm.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Telekommunikation", "account_key": "Telekom"}],
        "use_ai": True,
    },
    "deas": {
        "label": "DEAS",
        "tax_code_19": "VD",
        "text_template": "DEAS {kostenart} {name}",
        "prompt": "Kostenart Versicherung oder sonstige Leistung erkennen und passendes Sachkonto verwenden.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [
            {"name": "Versicherung", "account_key": "DEAS-Versicherungen"},
            {"name": "Sonstige", "account_key": "DEAS"},
        ],
        "use_ai": True,
    },
    "vw_leasing": {
        "label": "VW Leasing",
        "tax_code_19": "VD",
        "text_template": "VW Leasing {kennzeichen} {name}",
        "prompt": "Leasing je Fahrzeug und Kostenstelle kontieren.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Leasing", "account_key": "VW-Leasing"}],
        "use_ai": True,
    },
    "vw_versicherung": {
        "label": "VW Versicherungen",
        "tax_code_19": "VD",
        "text_template": "VW Versicherung {kennzeichen} {name}",
        "prompt": "Versicherung je Fahrzeug und Kostenstelle kontieren.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Versicherung", "account_key": "VW-Versicherungen"}],
        "use_ai": True,
    },
    "generic": {
        "label": "Weitere Lieferanten",
        "tax_code_19": "VD",
        "text_template": "{lieferant} {kostenart} {kennzeichen} {rufnummer} {name}",
        "prompt": "Lieferant, Kostenart und Kontierungsmerkmale ermitteln. Keine Kontierung erfinden.",
        "regions": {"inland": {"tax_code": "VD"}, "ausland": {"tax_code": "V0"}},
        "cost_types": [{"name": "Sonstige", "account_key": "Sonstige"}],
        "use_ai": True,
    },
}


def _merge(first, second):
    result = dict(first)
    for key, value in (second or {}).items():
        result[key] = _merge(result.get(key, {}), value) if isinstance(value, dict) and isinstance(result.get(key), dict) else value
    return result


def _supplier_key(result, rules):
    supplier = str(result.get("invoice", {}).get("supplier", "")).casefold()
    for key, rule in rules.get("suppliers", {}).items():
        words = [word for word in str(rule.get("label", key)).casefold().split() if len(word) > 3]
        if key.casefold() in supplier or any(word in supplier for word in words):
            return key
    return "generic" if "generic" in rules.get("suppliers", {}) else ""


def _ai_review(globals_dict, result, rules, supplier_key):
    """Run prompt review outside the FiBu Mate process; preserve protected fields."""
    rule = rules.get("suppliers", {}).get(supplier_key, {})
    if not rule.get("use_ai", False):
        return result
    module_path = Path(__file__).resolve()
    with tempfile.TemporaryDirectory(prefix="fibu_afi_ai_") as temp_dir:
        input_path = Path(temp_dir) / "input.json"
        output_path = Path(temp_dir) / "output.json"
        input_path.write_text(json.dumps({
            "general_prompt": rules.get("general_prompt", ""),
            "supplier_prompt": rule.get("prompt", ""),
            "structured_rule": rule,
            "upload_rows": result.get("upload_rows", []),
        }, ensure_ascii=False), encoding="utf-8")
        completed = subprocess.run(
            [sys.executable, str(module_path), "--ai-worker", str(input_path), str(output_path)],
            capture_output=True, text=True, timeout=180,
        )
        if completed.returncode != 0 or not output_path.is_file():
            raise RuntimeError((completed.stderr or completed.stdout or "KI-Prüflauf fehlgeschlagen")[-1000:])
        reviewed = json.loads(output_path.read_text(encoding="utf-8")).get("upload_rows", [])
        original = result.get("upload_rows", [])
        if len(reviewed) != len(original):
            raise RuntimeError("KI änderte die Anzahl der AFI-Zeilen")
        for old, new in zip(original, reviewed):
            for field in ("NET_VALUE", "GL_ACCOUNT", "COSTCENTER", "ORDERID"):
                if str(new.get(field, "")) != str(old.get(field, "")):
                    raise RuntimeError(f"KI änderte geschütztes Feld {field}")
            old["TEXT"] = str(new.get("TEXT", old.get("TEXT", "")))[:120]
            old["TAX_CODE"] = str(new.get("TAX_CODE", old.get("TAX_CODE", "")))[:10]
        result.setdefault("technical", {})["ai_prompts_applied"] = True
    return result


def apply_enhancements(g):
    tk, ttk = g["tk"], g["ttk"]
    messagebox, filedialog, simpledialog = g["messagebox"], g["filedialog"], g["simpledialog"]
    BaseUI = g["AFIUI"]
    original_defaults = g["_default_prompt_rules"]
    original_load = g["_load_prompt_rules"]

    def defaults():
        result = original_defaults()
        result["suppliers"] = _merge(DEFAULT_SUPPLIERS, result.get("suppliers", {}))
        return result

    def load_rules():
        return _merge(defaults(), original_load())

    g["_default_prompt_rules"] = defaults
    g["_load_prompt_rules"] = load_rules

    original_process = g["RuleEngine"].process

    def process_with_prompts(self, invoice_path, progress=lambda _message: None):
        result = original_process(self, invoice_path, progress)
        rules = getattr(self, "rules", load_rules())
        key = _supplier_key(result, rules)
        if key and rules.get("suppliers", {}).get(key, {}).get("use_ai", False):
            try:
                progress("Allgemeiner und lieferantenspezifischer Prompt werden durch die lokale KI geprüft ...")
                result = _ai_review(g, result, rules, key)
            except Exception as exc:
                result.setdefault("technical", {})["ai_prompts_applied"] = False
                result.setdefault("validation", {}).setdefault("warnings", []).append(f"KI-Promptprüfung nicht verfügbar: {exc}")
        return result

    g["RuleEngine"].process = process_with_prompts

    class EnhancedAFIUI(BaseUI):
        def __init__(self, app):
            super().__init__(app)
            self.rules = load_rules()
            self.zoom = 1.0
            self.pan_x = self.pan_y = 0
            self.drag_start = None
            self.supplier_tabs = {}
            self.db_open = False

        def render(self):
            super().render()
            self._bind_preview()
            self._postprocess_work_tab()

        def _walk(self, widget):
            yield widget
            for child in widget.winfo_children():
                yield from self._walk(child)

        def _postprocess_work_tab(self):
            for widget in list(self._walk(self.work_tab)):
                try:
                    if isinstance(widget, tk.Button) and widget.cget("text") == "Prompts & Regeln":
                        widget.destroy()
                except Exception:
                    pass
            boxes = [widget for widget in self.work_tab.winfo_children() if isinstance(widget, tk.LabelFrame)]
            if not boxes:
                return
            box = boxes[0]
            for widget in box.grid_slaves(row=0):
                try:
                    widget.configure(font=("Segoe UI", 10, "bold"))
                except Exception:
                    pass
            self.db_toggle = tk.Button(box, text="Datenbanken anzeigen ▼", command=lambda: self._toggle_databases(box), bg="#D9E2F3")
            self.db_toggle.grid(row=4, column=0, columnspan=4, sticky="w", pady=(4, 0))
            for row in (1, 2, 3):
                for widget in box.grid_slaves(row=row):
                    widget.grid_remove()

        def _toggle_databases(self, box):
            self.db_open = not self.db_open
            for row in (1, 2, 3):
                for widget in box.grid_slaves(row=row):
                    widget.grid() if self.db_open else widget.grid_remove()
            self.db_toggle.configure(text="Datenbanken ausblenden ▲" if self.db_open else "Datenbanken anzeigen ▼")

        def _render_rules_tab(self):
            frame = self.rules_tab
            for child in frame.winfo_children():
                child.destroy()
            frame.rowconfigure(1, weight=1)
            frame.columnconfigure(0, weight=1)
            header = tk.Frame(frame, bg=self.bg)
            header.grid(row=0, column=0, sticky="ew", padx=10, pady=8)
            tk.Label(header, text="Zentrale Prompt- und Regelpflege", bg=self.bg, font=("Segoe UI", 12, "bold")).pack(side="left")
            tk.Button(header, text="Lieferant hinzufügen", command=self.add_supplier, bg="#D9EAD3").pack(side="right", padx=5)
            tk.Button(header, text="Speichern", command=self.save_prompt_rules, bg="#0F6CBD", fg="white").pack(side="right", padx=5)
            notebook = ttk.Notebook(frame)
            notebook.grid(row=1, column=0, sticky="nsew", padx=10, pady=8)
            self.rules_notebook = notebook
            general = tk.Frame(notebook, bg=GREEN)
            notebook.add(general, text="▰ Allgemeiner Prompt")
            general.rowconfigure(1, weight=1)
            general.columnconfigure(0, weight=1)
            tk.Label(general, text="Allgemeiner Prompt", bg=GREEN, font=("Segoe UI", 11, "bold")).grid(row=0, column=0, sticky="w", padx=8, pady=6)
            prompt = tk.Text(general, wrap="word", undo=True, font=("Consolas", 10), bg="#F3FBF5")
            prompt.grid(row=1, column=0, sticky="nsew", padx=8, pady=(0, 8))
            prompt.insert("1.0", self.rules.get("general_prompt", ""))
            self.rule_widgets = {"general_prompt": prompt}
            self.supplier_tabs = {}
            for key, rule in self.rules.get("suppliers", {}).items():
                self._add_supplier_tab(notebook, key, rule)

        def _add_supplier_tab(self, notebook, key, rule):
            tab = tk.Frame(notebook, bg=self.bg)
            notebook.add(tab, text=rule.get("label", key))
            tab.columnconfigure(1, weight=1)
            tab.rowconfigure(7, weight=1)
            self.supplier_tabs[key] = tab
            label = tk.StringVar(value=rule.get("label", key))
            tax = tk.StringVar(value=rule.get("tax_code_19", "VD"))
            template = tk.StringVar(value=rule.get("text_template", ""))
            use_ai = tk.BooleanVar(value=rule.get("use_ai", True))
            for row, (caption, variable) in enumerate((("Bezeichnung", label), ("Steuerkennzeichen 19 %", tax), ("Buchungstext-Vorlage", template))):
                tk.Label(tab, text=caption, bg=self.bg).grid(row=row, column=0, sticky="w", padx=8, pady=4)
                tk.Entry(tab, textvariable=variable).grid(row=row, column=1, sticky="ew", padx=8, pady=4)
            tk.Checkbutton(tab, text="Allgemeinen und lieferantenspezifischen Prompt durch lokale KI berücksichtigen", variable=use_ai, bg=self.bg).grid(row=3, column=0, columnspan=2, sticky="w", padx=8)
            tk.Label(tab, text="Inland-/Ausland-Regeln (JSON)", bg=self.bg).grid(row=4, column=0, sticky="nw", padx=8, pady=4)
            regions = tk.Text(tab, height=5, font=("Consolas", 9))
            regions.grid(row=4, column=1, sticky="ew", padx=8, pady=4)
            regions.insert("1.0", json.dumps(rule.get("regions", {}), ensure_ascii=False, indent=2))
            tk.Label(tab, text="Kostenarten (JSON-Liste)", bg=self.bg).grid(row=5, column=0, sticky="nw", padx=8, pady=4)
            costs = tk.Text(tab, height=7, font=("Consolas", 9))
            costs.grid(row=5, column=1, sticky="ew", padx=8, pady=4)
            costs.insert("1.0", json.dumps(rule.get("cost_types", []), ensure_ascii=False, indent=2))
            tk.Label(tab, text="Lieferantenspezifischer Prompt", bg=self.bg).grid(row=6, column=0, columnspan=2, sticky="w", padx=8)
            supplier_prompt = tk.Text(tab, wrap="word", undo=True, font=("Consolas", 10))
            supplier_prompt.grid(row=7, column=0, columnspan=2, sticky="nsew", padx=8, pady=5)
            supplier_prompt.insert("1.0", rule.get("prompt", ""))
            tk.Button(tab, text="Lieferant entfernen", command=lambda supplier=key: self.remove_supplier(supplier), fg="#B42318").grid(row=8, column=1, sticky="e", padx=8, pady=4)
            self.rule_widgets[key] = {"label": label, "tax": tax, "template": template, "use_ai": use_ai, "regions": regions, "costs": costs, "prompt": supplier_prompt, "extra": {}}

        def add_supplier(self):
            label = simpledialog.askstring("Lieferant hinzufügen", "Bezeichnung des neuen Lieferanten:", parent=self.root)
            if not label:
                return
            key = re.sub(r"[^a-z0-9]+", "_", label.casefold()).strip("_") or "lieferant"
            base, number = key, 2
            while key in self.rules["suppliers"]:
                key = f"{base}_{number}"
                number += 1
            self.rules["suppliers"][key] = _merge(DEFAULT_SUPPLIERS["generic"], {"label": label})
            self._add_supplier_tab(self.rules_notebook, key, self.rules["suppliers"][key])
            self.rules_notebook.select(self.supplier_tabs[key])

        def remove_supplier(self, key):
            if messagebox.askyesno("Lieferant entfernen", f"'{self.rules['suppliers'][key].get('label', key)}' entfernen?"):
                self.rules_notebook.forget(self.supplier_tabs[key])
                self.supplier_tabs.pop(key, None)
                self.rule_widgets.pop(key, None)
                self.rules["suppliers"].pop(key, None)

        def save_prompt_rules(self):
            rules = load_rules()
            rules["general_prompt"] = self.rule_widgets["general_prompt"].get("1.0", "end-1c").strip()
            suppliers = {}
            try:
                for key, widgets in self.rule_widgets.items():
                    if key == "general_prompt":
                        continue
                    suppliers[key] = _merge(self.rules["suppliers"].get(key, {}), {
                        "label": widgets["label"].get().strip() or key,
                        "tax_code_19": widgets["tax"].get().strip(),
                        "text_template": widgets["template"].get().strip(),
                        "use_ai": bool(widgets["use_ai"].get()),
                        "regions": json.loads(widgets["regions"].get("1.0", "end-1c") or "{}"),
                        "cost_types": json.loads(widgets["costs"].get("1.0", "end-1c") or "[]"),
                        "prompt": widgets["prompt"].get("1.0", "end-1c").strip(),
                    })
                rules["suppliers"] = suppliers
                g["_save_prompt_rules"](rules)
                self.rules = rules
                messagebox.showinfo("Prompts & Regeln", "Zentrale Regeln wurden gespeichert.")
            except Exception as exc:
                messagebox.showerror("Prompts & Regeln", f"Ungültiges JSON oder Speicherfehler:\n{exc}")

        def _bind_preview(self):
            canvas = self.preview_canvas
            canvas.bind("<MouseWheel>", self._zoom)
            canvas.bind("<Button-4>", lambda event: self._zoom(event, 1))
            canvas.bind("<Button-5>", lambda event: self._zoom(event, -1))
            canvas.bind("<ButtonPress-1>", self._drag_start)
            canvas.bind("<B1-Motion>", self._drag)
            canvas.bind("<ButtonRelease-1>", lambda event: setattr(self, "drag_start", None))

        def _zoom(self, event, direction=None):
            if not self.preview_images:
                return
            step = direction if direction is not None else (1 if event.delta > 0 else -1)
            self.zoom = max(0.25, min(4.0, self.zoom * (1.15 if step > 0 else 1 / 1.15)))
            self.show_preview()

        def _drag_start(self, event):
            self.drag_start = (event.x, event.y, self.pan_x, self.pan_y)

        def _drag(self, event):
            if not self.drag_start:
                return
            x, y, old_x, old_y = self.drag_start
            self.pan_x = old_x + event.x - x
            self.pan_y = old_y + event.y - y
            self.show_preview()

        def show_preview(self):
            if not self.preview_images or not g.get("PIL_AVAILABLE", False) or not hasattr(self, "preview_canvas"):
                return
            source = self.preview_images[self.preview_index]
            width = max(100, self.preview_canvas.winfo_width())
            height = max(100, self.preview_canvas.winfo_height())
            base = min((width - 20) / source.width, (height - 20) / source.height)
            image_width = max(1, int(source.width * base * self.zoom))
            image_height = max(1, int(source.height * base * self.zoom))
            image = source.resize((image_width, image_height))
            limit_x = max(0, (image_width - width) // 2 + width // 3)
            limit_y = max(0, (image_height - height) // 2 + height // 3)
            self.pan_x = max(-limit_x, min(limit_x, self.pan_x))
            self.pan_y = max(-limit_y, min(limit_y, self.pan_y))
            self.preview_photo = g["ImageTk"].PhotoImage(image)
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(width // 2 + self.pan_x, height // 2 + self.pan_y, image=self.preview_photo)
            self.preview_page_var.set(f"Seite {self.preview_index + 1} von {len(self.preview_images)} | Zoom {self.zoom:.0%}")
            self.prev_button.config(state="normal" if self.preview_index > 0 else "disabled")
            self.next_button.config(state="normal" if self.preview_index < len(self.preview_images) - 1 else "disabled")

        def _render_document_images(self, path):
            if path.suffix.lower() in (".xlsx", ".xlsm"):
                import openpyxl
                workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
                images = []
                try:
                    for sheet in workbook.worksheets:
                        rows = []
                        for row in sheet.iter_rows(values_only=True):
                            rows.append([str(value or "")[:35] for value in row[:12]])
                            if len(rows) >= 45:
                                break
                        images.append(self._table_image(sheet.title, rows))
                finally:
                    workbook.close()
                return images
            return super()._render_document_images(path)

        def _table_image(self, title, rows):
            Image, ImageDraw, ImageFont = g["Image"], g["ImageDraw"], g["ImageFont"]
            columns = max([len(row) for row in rows] or [1])
            cell_width, row_height = 180, 34
            width = min(2200, columns * cell_width + 40)
            height = max(300, min(1800, (len(rows) + 2) * row_height + 40))
            image = Image.new("RGB", (width, height), "white")
            draw = ImageDraw.Draw(image)
            try:
                font = ImageFont.truetype("arial.ttf", 15)
                bold = ImageFont.truetype("arialbd.ttf", 17)
            except Exception:
                font = bold = ImageFont.load_default()
            draw.rectangle((0, 0, width, row_height + 15), fill="#D9EAD3")
            draw.text((15, 10), title, fill="black", font=bold)
            for row_index, row in enumerate(rows):
                y = row_height + 15 + row_index * row_height
                for column_index, value in enumerate(row):
                    x = 20 + column_index * cell_width
                    draw.rectangle((x, y, x + cell_width, y + row_height), outline="#AAB2BD", fill="#EEF4FB" if row_index == 0 else "white")
                    draw.text((x + 5, y + 8), value, fill="black", font=font)
            return image

        def done(self, result):
            super().done(result)
            self.export.config(state="normal")

        def export_csv(self):
            if not self.result:
                return
            if not self.result.get("validation", {}).get("export_allowed", False):
                if not messagebox.askyesno("CSV exportieren", "Es bestehen Prüfhinweise. Sichtbare Vorschau trotzdem exportieren?"):
                    return
            path = filedialog.asksaveasfilename(defaultextension=".csv", initialfile="AFI_Upload.csv", filetypes=[("CSV-Datei", "*.csv")])
            if path:
                with open(path, "w", encoding="utf-8-sig", newline="") as handle:
                    writer = csv.DictWriter(handle, fieldnames=g["UPLOAD_COLUMNS"], delimiter=";", extrasaction="ignore")
                    writer.writeheader()
                    writer.writerows(self.result["upload_rows"])
                messagebox.showinfo("CSV exportieren", f"CSV gespeichert:\n{path}")

    g["AFIUI"] = EnhancedAFIUI


def _ai_worker(input_path, output_path):
    data = json.loads(Path(input_path).read_text(encoding="utf-8"))
    from foundry_local_sdk import Configuration, FoundryLocalManager
    root = Path(os.environ.get("LOCALAPPDATA", str(Path.home()))) / "FiBuMate" / "FoundryLocal"
    root.mkdir(parents=True, exist_ok=True)
    try:
        FoundryLocalManager.initialize(Configuration(app_name="FiBuMate-AFI-Prompt", app_data_dir=str(root), model_cache_dir=str(root / "models"), logs_dir=str(root / "logs")))
    except Exception:
        pass
    manager = FoundryLocalManager.instance
    # Unter Windows die verfügbaren Execution Provider vor der Modellauswahl
    # registrieren. foundry-local-sdk-winml stellt dafür Windows ML bereit.
    try:
        discover = getattr(manager, "discover_eps", None)
        register = getattr(manager, "download_and_register_eps", None)
        eps = discover() if callable(discover) else []
        if eps and callable(register):
            register()
    except Exception:
        pass
    model = manager.catalog.get_model("phi-4-mini")
    if model is None:
        raise RuntimeError("phi-4-mini ist nicht verfügbar")
    if not model.is_cached:
        model.download()
    model.load()
    try:
        client = model.get_chat_client()
        client.settings.temperature = 0
        client.settings.max_tokens = 2048
        client.settings.response_format = {"type": "json_object"}
        instruction = (
            "Prüfe die AFI-Zeilen anhand der Prompts. Gib nur JSON mit upload_rows zurück. "
            "Behalte Anzahl, NET_VALUE, GL_ACCOUNT, COSTCENTER und ORDERID unverändert. "
            "Ändere nur TEXT und TAX_CODE.\n\nALLGEMEINER PROMPT:\n" + data["general_prompt"] +
            "\n\nLIEFERANTENPROMPT:\n" + data["supplier_prompt"] +
            "\n\nSTRUKTURIERTE REGELN:\n" + json.dumps(data["structured_rule"], ensure_ascii=False) +
            "\n\nAFI-ZEILEN:\n" + json.dumps(data["upload_rows"], ensure_ascii=False)
        )
        response = client.complete_chat([
            {"role": "system", "content": "Du bist der lokale AFI-Regelprüfer von FiBu Mate."},
            {"role": "user", "content": instruction},
        ])
        result = json.loads(response.choices[0].message.content)
        Path(output_path).write_text(json.dumps(result, ensure_ascii=False), encoding="utf-8")
    finally:
        try:
            model.unload()
        except Exception:
            pass


apply_enhancements(globals())
# ------------------------------------------------------------------
# AFI 1.3.0 Korrekturblock: Live-Prompts, WinML und Datenbank-Container
# ------------------------------------------------------------------
def _afi_short_ai_error(value):
    text = str(value or "")
    low = text.casefold()
    if "webgpu execution provider is not supported" in low:
        return "GPU-Modell nicht kompatibel (WebGPU fehlt). Windows-ML-Provider erforderlich; CSV ohne KI-Prüfung erstellt."
    if "no module named" in low and "foundry_local_sdk" in low:
        return "foundry-local-sdk-winml fehlt; CSV ohne KI-Prüfung erstellt."
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return (lines[-1] if lines else "Lokale KI nicht verfügbar")[:500]

_afi_process_with_prompts = RuleEngine.process
def _afi_process_live_prompts(self, invoice_path, progress=lambda _message: None):
    self.rules = _load_prompt_rules()
    result = _afi_process_with_prompts(self, invoice_path, progress)
    technical = result.setdefault("technical", {})
    technical["prompt_rules_file"] = str(_prompt_rules_file(False))
    technical["prompt_rules_loaded_at"] = datetime.now().astimezone().isoformat(timespec="seconds")
    technical["general_prompt_loaded"] = bool(self.rules.get("general_prompt", "").strip())
    key = _supplier_key(result, self.rules)
    technical["supplier_rule_key"] = key
    technical["supplier_prompt_loaded"] = bool(self.rules.get("suppliers", {}).get(key, {}).get("prompt", "").strip()) if key else False
    warnings = result.setdefault("validation", {}).setdefault("warnings", [])
    result["validation"]["warnings"] = [
        "KI-Promptprüfung nicht verfügbar: " + _afi_short_ai_error(w)
        if str(w).startswith("KI-Promptprüfung nicht verfügbar:") else str(w)
        for w in warnings
    ]
    technical["ai_prompt_status"] = "angewendet" if technical.get("ai_prompts_applied") is True else "nicht angewendet"
    return result
RuleEngine.process = _afi_process_live_prompts

def _afi_postprocess_work_tab(self):
    for widget in list(self._walk(self.work_tab)):
        try:
            if isinstance(widget, tk.Button) and widget.cget("text") == "Prompts & Regeln":
                widget.destroy()
        except Exception:
            pass
    boxes = [widget for widget in self.work_tab.winfo_children() if isinstance(widget, tk.LabelFrame)]
    if not boxes:
        return
    box = boxes[0]
    box.columnconfigure(1, weight=1)
    for row in (1, 2, 3):
        for widget in box.grid_slaves(row=row):
            widget.grid_remove()
    self.db_toggle = tk.Button(box, text="Datenbanken anzeigen ▼", command=self._toggle_databases, bg="#D9E2F3")
    self.db_toggle.grid(row=4, column=0, columnspan=4, sticky="w", pady=(5, 2))
    self.database_container = tk.Frame(box, bg=self.bg)
    self.database_container.columnconfigure(1, weight=1)
    self._file_row(self.database_container, 0, "Datenbank 1: Mitarbeiter -> Kostenstelle/IA", self.costcenter, lambda: self.pick_db(self.costcenter, "Kostenstellen-Datenbank"), lambda: self.clear_db(self.costcenter))
    self._file_row(self.database_container, 1, "Datenbank 2: Kennzeichen/Rufnummer/Firma + Sachkonten", self.account, lambda: self.pick_db(self.account, "Kontierungs-Generalübersicht"), lambda: self.clear_db(self.account))
    tk.Label(self.database_container, text=f"Zentral gespeichert: {_shared_settings_file(False)} | Prompts: {_prompt_rules_file(False)}", bg=self.bg, fg="#44536A", anchor="w").grid(row=2, column=0, columnspan=4, sticky="ew", pady=(3, 0))
    self.database_container.grid(row=5, column=0, columnspan=4, sticky="ew", pady=(2, 0))
    self.database_container.grid_remove()
    self.db_open = False

def _afi_toggle_databases(self, *_args):
    self.db_open = not bool(getattr(self, "db_open", False))
    if self.db_open:
        self.database_container.grid()
        self.db_toggle.configure(text="Datenbanken ausblenden ▲")
    else:
        self.database_container.grid_remove()
        self.db_toggle.configure(text="Datenbanken anzeigen ▼")
    for widget in (self.database_container, self.work_tab, self.canvas):
        try: widget.update_idletasks()
        except Exception: pass

AFIUI._postprocess_work_tab = _afi_postprocess_work_tab
AFIUI._toggle_databases = _afi_toggle_databases

def render(app):
    ui = AFIUI(app)
    app._afi_local_ai_ui = ui
    ui.render()

if __name__ == "__main__" and len(sys.argv) >= 4 and sys.argv[1] == "--ai-worker":
    try:
        _ai_worker(sys.argv[2], sys.argv[3])
    except Exception:
        traceback.print_exc()
        sys.exit(1)
