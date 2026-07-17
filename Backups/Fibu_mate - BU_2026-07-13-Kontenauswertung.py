# FiBuMate_PATCH_MARKER: 20260609_MENUZEILEN_DEBITOREN_PROTOCOL
# FiBuMate_PATCH_MARKER: 20260609_v0436_DIREKT_ABSCHLUSSKALENDER_EIN_MODUL
# FiBuMate_PATCH_MARKER: 20260609_v0436_DREI_MODULE_STICHTAGSPFLEGE_OHNE_IDS
# FiBuMate_PATCH_MARKER: 20260609_150049 (V0.436_TEAM_RELEASE_SCALING_MENU_CLEANUP_READABILITY_SAFE_TEXT)
import os
import sys
import json
import importlib
import re
import hashlib


def _load_tool_module_from_file(module_path: str):
    """Lädt ein Tool-Modul immer direkt aus C:\\python\\bin\\tools\\*.py.
    Damit werden Import-Konflikte mit anderen Python-Umgebungen ausgeschlossen."""
    try:
        import importlib.util
        parts = (module_path or '').split('.')
        name = parts[-1] if parts else module_path
        file_path = os.path.join(BIN_DIR, 'tools', f"{name}.py")
        if not os.path.exists(file_path):
            # Fallback: ggf. liegt die Datei direkt im bin-Ordner
            file_path2 = os.path.join(BIN_DIR, f"{name}.py")
            if os.path.exists(file_path2):
                file_path = file_path2
        spec_name = f"fibumate_local_{name}"
        spec = importlib.util.spec_from_file_location(spec_name, file_path)
        if spec and spec.loader:
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            return mod
    except Exception:
        pass
    return None

import webbrowser
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.font as tkfont
from tkinter import messagebox
from datetime import datetime
from urllib.parse import quote

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

APP_NAME = "FiBu Mate"
VERSION_PREFIX = "0.4"
DEFAULT_BUILD = 0
VERSION_STATE_FILE = "version_state.json"
VERSION_HISTORY_FILE = "version_history.json"
# v0.436: Manuelle Zoomprofile entfernt; Darstellung skaliert automatisch anhand Fenster-/Monitorgröße.
ZOOM_PROFILE_FILE = None

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BIN_DIR = os.path.join(SCRIPT_DIR, "bin")
IMG_DIR = os.path.join(BIN_DIR, "Imgs")
ICON_DIR = os.path.join(IMG_DIR, "Icons")
USER_DIR = os.path.join(BIN_DIR, "User")
LEGACY_USER_DATA_PATH = os.path.join(USER_DIR, "fibu_mate_users.json")
NETWORK_ROOT = r"G:\BUC\FM Anwendung"
CENTRAL_CONFIG_DIR = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Config")
CENTRAL_USER_DATA_PATH = os.path.join(CENTRAL_CONFIG_DIR, "fibu_mate_users.json")
CENTRAL_RELEASE_DIR = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Releases")
LATEST_UPDATE_FILE = "latest.json"
UPDATE_EVENT_FILE = "update_event.json"
UPDATE_CHECK_INTERVAL_MS = 60 * 1000
HIDDEN_TOOL_IDS = {"enbw_strom_tanken_upload"}  # nicht löschen: nur aus Oberfläche/Favoriten ausblenden


def _safe_json_load(path):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        return {}
    return {}


def resolve_user_data_path():
    env_path = os.environ.get("FIBUMATE_USER_DATA_PATH", "").strip()
    if env_path:
        return env_path
    for cfg in (os.path.join(SCRIPT_DIR, "config", "local_config.json"), os.path.join(CENTRAL_CONFIG_DIR, "local_config.json")):
        data = _safe_json_load(cfg)
        p = str(data.get("user_data_path", "")).strip()
        if p:
            return p
    return CENTRAL_USER_DATA_PATH


USER_DATA_PATH = resolve_user_data_path()

if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

BANNER_GROSS_PATH = os.path.join(IMG_DIR, "FMBanner_Gross.png")
BANNER_KLEIN_PATH = os.path.join(IMG_DIR, "FMBanner_Klein.png")
HELP_IMAGE_PATH = os.path.join(IMG_DIR, "FM_Help_Menu.png")
ICON_FILES = {
    "help": "metrostatus_question_support_6987.ico",
    "idea": "2799205-creative-idea-light_99776.ico",
    "doc_add": "addfileinterfacesymbolofpapersheetwithtextlinesandplussign_79821.ico",
    "doc_file": "fileinterfacesymboloftextpapersheet_79740.ico",
    "attach": "-attach-file_90371.ico",
    "calendar": "calendar-icon_34471.ico",
    "gear": "1486504840-cog-cogwheel-gear-repr-options-setting_81360.ico",
    "info": "1486504328-bullet-list-menu-lines-points-items-options_81334.ico",
    "knowledge": "bookshelf_icon-icons.com_54414.ico",
    "lock": "lock_lock_15063.ico",
    "unlock": "lock_unlock_15064.ico",
    "xls": "ext_xls_filetype_icon_176238.ico",
    "pdf": "ext_pdf_filetype_icon_176234.ico",
    "compliance": "1486503790-bank-building-government-house-real-estate-panteon_81294.ico",
    "tax_reporting": "1486504352-checklist-clipboard-inventory-list-report-tasks-todo_81326.ico",
    "audit": "1486506277-like-thumbs-up-hands-gesture-finger-vote_81482.ico",
    "documentation": "1486485527-account-albums-screens-tabs_81163.ico",
    "search": "1486503763-bigger-enlarge-search-magnifier-magnify-zoom_81256.ico",
    "filter": "1486504837-descending-filter-filtering-tool-funnel-sort_81363.ico",
    "edit": "1486504369-change-edit-options-pencil-settings-tools-write_81307.ico",
}
INTERSPORT_LOGO_CANDIDATES = ["IS_Banner_lang.png", "Intersport_Logo.png", "INTERSPORT_Logo.png", "intersport_logo.png", "INTERSPORT.png"]

BLUE = "#004B93"
RED = "#E30613"
GOLD = "#FFD700"
STAR_GREY = "#7A7F87"
BG = "#E8EEF5"
HEADER = "#D3DEE9"
LINE = "#91A3B5"
TEXT = "#182431"
TEXT2 = "#445364"
SHADOW = "#A8B5C3"
WHITE = "#FFFFFF"
GREY_DISABLED = "#B9C3CF"
GREY_TILE = "#D6DCE4"

FONT_TITLE = ("Segoe UI", 26, "bold")
FONT_MENU = ("Segoe UI", 22, "bold")
FONT_TILE = ("Segoe UI", 18, "bold")
FONT_TILE_SMALL = ("Segoe UI", 15, "bold")
FONT_SMALL = ("Segoe UI", 10)
BASE_FONT_TITLE = FONT_TITLE
BASE_FONT_MENU = FONT_MENU
BASE_FONT_TILE = FONT_TILE
BASE_FONT_TILE_SMALL = FONT_TILE_SMALL
BASE_FONT_SMALL = FONT_SMALL
UI_SCALE = 1.0
UI_TEXT_SCALE = 1.00  # v0.436: kein manueller Textzoom; automatische, begrenzte UI-Skalierung.
UI_BODY_TEXT_SCALE = 1.00  # v0.433 Korrektur Paket 1g: 1920x1080 ist Body-Referenz; dynamische Lesbarkeit über body_font().
GLOBAL_TEXT_ZOOM_MIN = 0.70
GLOBAL_TEXT_ZOOM_MAX = 1.80
GLOBAL_TEXT_ZOOM_STEP = 0.025  # v0.433 Korrektur Paket 1h: feine Körnung pro Mausrad-Raster.
GLOBAL_TEXT_ZOOM = 1.00  # Fallback; v0.434 nutzt bereichsbezogene Zoomprofile.

def ui_s(value):
    try:
        return max(1, int(round(float(value) * UI_SCALE)))
    except Exception:
        return value

def scaled_font(font_tuple):
    try:
        family, size, *rest = font_tuple
        return tuple([family, max(7, int(round(size * UI_SCALE * UI_TEXT_SCALE)))] + rest)
    except Exception:
        return font_tuple

def ui_icon_size(base=36):
    """v0.433 Korrektur: zentrale Icon-Skalierung für kleinere Monitore."""
    try:
        return max(14, int(round(float(base) * UI_SCALE)))
    except Exception:
        return base

def body_scale_value():
    """v0.433 Korrektur Paket 1g: Dynamische Body-Skalierung.
    Referenz: 1920x1080 => Faktor 1.00. Kleinere Displays schrumpfen Body-/Tabellentexte nicht weiter,
    größere Displays dürfen moderat wachsen. Kacheln/Überschriften bleiben über scaled_font() getrennt.
    """
    try:
        return max(1.0, min(1.18, float(UI_SCALE))) * UI_BODY_TEXT_SCALE
    except Exception:
        return 1.0


def body_font(size=10, weight=None, underline=False, scale=1.0):
    """Zentrale Schrift für Pfadleiste, Fußleiste, Modulbeschreibungen, Modul-/Tabellentexte und Bedienbuttons.

    Lesbarkeitsregel V3: Sehr kleine Body-Schriften werden sichtbar auf mindestens 12 pt angehoben.
    Modulbeschreibungen und Canvas-Texte werden zusätzlich zentral abgefangen.
    """
    try:
        scaled = int(round(float(size) * UI_TEXT_SCALE * body_scale_value() * float(scale)))
        scaled = max(12, scaled)
    except Exception:
        scaled = max(12, int(size) if str(size).isdigit() else 12)
    style = []
    if weight:
        style.append(weight)
    if underline:
        style.append('underline')
    return tuple(['Segoe UI', scaled] + style)


# Zentrale Lesbarkeitsstufe V3: etwas größerer Fließtext, Überschriften bleiben unverändert.
READABILITY_MIN_LABEL_FONT = 13
READABILITY_MIN_BUTTON_FONT = 12
READABILITY_MIN_ENTRY_FONT = 12
READABILITY_MIN_TEXT_FONT = 12
READABILITY_MIN_TTK_FONT = 12
READABILITY_MIN_CANVAS_FONT = 13
READABILITY_TREEVIEW_ROWHEIGHT = 28
_CENTRAL_READABILITY_INSTALLED = False
_TK_ORIGINAL_INITS = {}
_TK_ORIGINAL_CONFIGS = {}
_TK_ORIGINAL_CANVAS_CREATE_TEXT = None


def _readability_font_tuple(font_value, minimum=11):
    """Hebt nur sehr kleine Tk-Fonts konservativ an; große/formatierte Fonts bleiben erhalten."""
    if font_value is None or font_value == "":
        return ("Segoe UI", minimum)
    try:
        if isinstance(font_value, tuple):
            if len(font_value) >= 2 and isinstance(font_value[1], int):
                if font_value[1] < minimum:
                    return tuple([font_value[0], minimum] + list(font_value[2:]))
                return font_value
        if isinstance(font_value, list):
            if len(font_value) >= 2 and isinstance(font_value[1], int):
                result = list(font_value)
                if result[1] < minimum:
                    result[1] = minimum
                return tuple(result)
        # Tk erlaubt Font-Strings; diese werden bewusst nicht geparst, um keine Systemfonts zu zerstören.
    except Exception:
        pass
    return font_value


def _readability_should_skip_canvas_text(text_value):
    """Schützt sehr kleine Symbol-/Icon-Texte vor unbeabsichtigter Vergrößerung."""
    try:
        value = str(text_value or "").strip()
        return value in {"", "★", "+", "−", "-", "HELP"}
    except Exception:
        return False


def _patch_tk_widget_font(widget_cls, minimum=11):
    if widget_cls in _TK_ORIGINAL_INITS:
        return
    original_init = widget_cls.__init__
    original_configure = getattr(widget_cls, "configure", None)
    original_config = getattr(widget_cls, "config", None)
    _TK_ORIGINAL_INITS[widget_cls] = original_init
    _TK_ORIGINAL_CONFIGS[widget_cls] = (original_configure, original_config)

    def patched_init(self, master=None, cnf=None, **kw):
        cnf = {} if cnf is None else dict(cnf)
        if "font" in kw:
            kw["font"] = _readability_font_tuple(kw.get("font"), minimum)
        elif "font" in cnf:
            cnf["font"] = _readability_font_tuple(cnf.get("font"), minimum)
        else:
            kw["font"] = ("Segoe UI", minimum)
        return original_init(self, master, cnf, **kw)

    def patched_configure(self, cnf=None, **kw):
        cnf = {} if cnf is None else dict(cnf)
        if "font" in kw:
            kw["font"] = _readability_font_tuple(kw.get("font"), minimum)
        if "font" in cnf:
            cnf["font"] = _readability_font_tuple(cnf.get("font"), minimum)
        return original_configure(self, cnf, **kw)

    widget_cls.__init__ = patched_init
    if original_configure is not None:
        widget_cls.configure = patched_configure
        widget_cls.config = patched_configure


def _patch_canvas_create_text():
    global _TK_ORIGINAL_CANVAS_CREATE_TEXT
    if _TK_ORIGINAL_CANVAS_CREATE_TEXT is not None:
        return
    _TK_ORIGINAL_CANVAS_CREATE_TEXT = tk.Canvas.create_text

    def patched_create_text(self, *args, **kw):
        if not _readability_should_skip_canvas_text(kw.get("text", "")):
            if "font" in kw:
                kw["font"] = _readability_font_tuple(kw.get("font"), READABILITY_MIN_CANVAS_FONT)
            else:
                kw["font"] = ("Segoe UI", READABILITY_MIN_CANVAS_FONT)
        return _TK_ORIGINAL_CANVAS_CREATE_TEXT(self, *args, **kw)

    tk.Canvas.create_text = patched_create_text


def install_central_readability(root=None):
    """Zentrale Lesbarkeitsverbesserung V3 für FiBu Mate und nachgeladene Tool-Module.

    V3 erhöht die bereits sichtbare V2-Wirkung moderat und wirkt weiter auf Canvas-Texte sowie spätere .configure(font=...)-Aufrufe.
    Dadurch werden Menübeschreibungen und viele modulinterne Texte noch etwas größer,
    ohne einzelne Moduldateien ändern zu müssen.
    """
    global _CENTRAL_READABILITY_INSTALLED
    if _CENTRAL_READABILITY_INSTALLED:
        return
    _CENTRAL_READABILITY_INSTALLED = True
    try:
        _patch_tk_widget_font(tk.Label, READABILITY_MIN_LABEL_FONT)
        _patch_tk_widget_font(tk.Button, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Entry, READABILITY_MIN_ENTRY_FONT)
        _patch_tk_widget_font(tk.Checkbutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Radiobutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Menubutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Listbox, READABILITY_MIN_TEXT_FONT)
        _patch_tk_widget_font(tk.Text, READABILITY_MIN_TEXT_FONT)
        _patch_canvas_create_text()
    except Exception:
        pass
    try:
        style = ttk.Style(root) if root is not None else ttk.Style()
        style.configure("TLabel", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TButton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TCheckbutton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TRadiobutton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TNotebook.Tab", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("Treeview", font=("Segoe UI", READABILITY_MIN_TTK_FONT), rowheight=max(READABILITY_TREEVIEW_ROWHEIGHT, ui_s(READABILITY_TREEVIEW_ROWHEIGHT)))
        style.configure("Treeview.Heading", font=("Segoe UI", READABILITY_MIN_TTK_FONT, "bold"))
    except Exception:
        pass

MINI_WIDGET_W = 174
MINI_WIDGET_H = 30
MINI_WIDGET_GAP = 8

COLOR_PALETTE = [("Blau", BLUE), ("Grün", "#059669"), ("Rot", RED), ("Gelb", "#F59E0B"), ("Lila", "#7C3AED"), ("Pink", "#EC4899"), ("Dunkelgrau", "#334155"), ("Orange", "#F97316"), ("Türkis", "#06B6D4")]

ROLE_E1 = "E1 - Standard"
ROLE_E2 = "E2 - Erweitert"
ROLE_E3 = "E3 - Administrator"
ROLE_E4 = "E4 - System-Administrator"
ROLE_STANDARD = ROLE_E1
ROLE_ADMIN = ROLE_E3
OLD_ROLE_E4 = "Wagnerm"
ROLE_WAGNERM = ROLE_E4
SUPERUSER_KEY = "wagnerm"
ROLE_ORDER = [ROLE_E1, ROLE_E2, ROLE_E3, ROLE_E4]
ROLE_RANK = {
    ROLE_E1: 1, ROLE_E2: 2, ROLE_E3: 3, ROLE_E4: 4,
    "Standard": 1, "Ebene 1": 1, "E1": 1,
    "Ebene 2": 2, "E2": 2,
    "Administrator": 3, "Ebene 3": 3, "E3": 3,
    "System-Administrator": 4, OLD_ROLE_E4: 4, "Ebene 4": 4, "E4": 4,
}
ROLE_MIGRATION = {"Standard": ROLE_E1, "Administrator": ROLE_E3, "System-Administrator": ROLE_E4, OLD_ROLE_E4: ROLE_E4, "Ebene 1": ROLE_E1, "Ebene 2": ROLE_E2, "Ebene 3": ROLE_E3, "Ebene 4": ROLE_E4, "E1": ROLE_E1, "E2": ROLE_E2, "E3": ROLE_E3, "E4": ROLE_E4}

TOOL_REGISTRY = {
    "nike_pdf_to_excel": {"title": "Nike - PDF zu Excel", "module": "bin.tools.nike_pdf_to_excel", "favorite_label": "Nike PDF"},
    "nike_op_liste_pdf_check": {"title": "Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen", "module": "bin.tools.nike_op_liste_pdf_check", "favorite_label": "Nike OP PDF"},
    "invoice_pdf_collector": {"title": "Nike - Rechnungs-PDFs in Sammelordner", "module": "bin.tools.invoice_pdf_collector", "favorite_label": "Nike RE sammeln"},
    "enbw_strom_tanken_upload": {"title": "EnBW - Strom-Tanken Upload-Erstellung", "module": "bin.tools.enbw_strom_tanken_upload", "favorite_label": "EnBW Strom"},
    "supplier_invoice_afi_upload": {"title": "Lieferanten-Rechnung zu AFI-Upload", "module": "bin.tools.supplier_invoice_afi_upload", "favorite_label": "Lieferanten AFI"},
    "aramark_monatsabrechnungen_pdf_to_excel": {"title": "Aramark Monatsabrechnungen - PDF zu Excel", "module": "bin.tools.aramark_monatsabrechnungen_pdf_to_excel", "favorite_label": "Aramark Monat"},
    "debitoren_serienbrief": {"title": "Debitoren-Serienbrief", "module": "bin.tools.debitoren_serienbrief", "favorite_label": "Debitoren SB"},
    "monthly_close": {"title": "Monatsabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Monatsabschluss"},
    "quarterly_close": {"title": "Quartalsabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Quartalsabschluss"},
    "yearly_close": {"title": "Jahresabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Jahresabschluss"},
    "deadline_maintenance": {"title": "Stichtagspflege", "module": "bin.tools.deadline_maintenance", "favorite_label": "Stichtage"},
    "x001_sap_test": {"title": "X001 SAP - Test", "module": "bin.tools.x001_sap_test", "favorite_label": "X001"},
    "tax_reporting": {"title": "Steuermeldungs-Cockpit", "module": "bin.tools.compliance_tax_reporting", "favorite_label": "Steuermeldungen"},
    "audit_cockpit": {"title": "Audit-Cockpit", "module": "bin.tools.compliance_audit_cockpit", "favorite_label": "Audit"},
    "documentation_center": {"title": "Dokumentationszentrale", "module": "bin.tools.compliance_documentation_center", "favorite_label": "Dokumente"},
}

MODULE_DESCRIPTIONS = {
    "nike_pdf_to_excel": "Mit diesem Modul lassen sich große Mengen an Nike PDF-Rechnungen in Excel ein Excel-Format ausgeben. Die auszugebenden Daten lassen sich filtern und sind individuell anpassbar.",
    "nike_op_liste_pdf_check": "Prüft die Vollständigkeit von Nike PDF-Rechnungen gegen eine OP-Liste: Rechnungsnummern aus PDF-Dateinamen werden mit Spalte B der Excel-Datei abgeglichen.",
    "monthly_close": "Interaktives Monatsabschluss-Cockpit mit Teamfortschritt, Aufgabenstatus, Fristwarnungen und Anlagen je Aufgabe.",
    "quarterly_close": "Interaktives Quartalsabschluss-Cockpit auf Basis der Monatsabschluss-Struktur mit Quartalsperioden.",
    "tax_reporting": "Überwacht steuerliche Meldungen und Meldefristen je Zeitraum inklusive Status, Zuständigkeit, Nachweisen, Historie und Abschlussbericht-Integration.",
    "audit_cockpit": "Zentrale Übersicht über kritische Systemereignisse wie Wiederöffnungen, Berechtigungsänderungen, Benutzeränderungen und nachträgliche Änderungen nach Abschluss.",
    "documentation_center": "Zentrale Suche und Prüfung aller Nachweise, Anlagen, Dokumentationspfade und Berichte aus FiBu Mate inklusive fehlender oder ungültiger Dokumentationen.",
    "yearly_close": "Interaktives Jahresabschluss-Cockpit für Geschäftsjahre vom 01.10. bis 30.09.",
    "deadline_maintenance": "Pflege der Abschluss-Stichtage (Dekadenabschluss, 18-Uhr, 08-Uhr, Monatsabschluss) inkl. Feiertage BW und automatischer Übernahme in Monats-/Quartals-/Jahresabschluss.",
    "x001_sap_test": "SAP-Scripting-Test; Scripting in SAP deaktiviert.",
    "invoice_pdf_collector": "In Excel gefilterte Rechnungsnummern aus PDF-Verzeichnis wählen und in neuen Sammelordner kopieren.",
    "enbw_strom_tanken_upload": "Erstellt aus EnBW E-Tankkosten-Abrechnungen eine SAP-AFI-uploadfähige CSV anhand der bestehenden Upload-Vorlage; Zuordnung nach Kennzeichen, Steuerlogik, Grundgebühren und Hinweis-Popup bei Abweichungen.",
    "page:compliance_audit": "Entwicklungsbereich für Compliance- und Audit-Funktionen: Steuermeldungen, Audit-Cockpit und Dokumentationszentrale bleiben gebündelt, werden aber nicht im produktiven Hauptmenü angezeigt.",
    "supplier_invoice_afi_upload": "Generisches CSV-Modul für Lieferantenrechnungen: erkennt relevante Spalten logisch, gleicht gegen eine wählbare AFI-/Kontierungsvorlage ab und exportiert eine uploadfähige AFI-CSV mit Spalten A-J.",
    "aramark_monatsabrechnungen_pdf_to_excel": "Erstellt aus einer oder mehreren Aramark-Monatsabrechnungs-PDFs eine gemeinsame Excel-Datei nach der festen Aramark-Vorlage; pro PDF entsteht eine frei benennbare eigene Umsatz-Spalte inklusive Plausibilitätsprüfung und Exportvorschau.",
}
DESCRIPTION_FONT = ("Segoe UI", 11)
DESCRIPTION_COLOR = TEXT2
DESCRIPTION_X_OFFSET = 14
DESCRIPTION_Y_OFFSET = 18


def maximize_window(window: tk.Tk):
    try:
        window.state("zoomed")
    except Exception:
        try:
            window.attributes("-zoomed", True)
        except Exception:
            pass


def x_pct(width: int, percent: float) -> float:
    return width * percent / 100


def y_pct(height: int, percent: float) -> float:
    return height * (1 - percent / 100)


def normalize_username(username: str) -> str:
    return " ".join(str(username).strip().split()).casefold()


def _rgb(hex_color: str):
    v = hex_color.lstrip("#")
    return tuple(int(v[i:i + 2], 16) for i in (0, 2, 4))


def _hex(rgb):
    return "#{:02x}{:02x}{:02x}".format(*rgb)


def blend(a: str, b: str, t: float) -> str:
    ar, ag, ab = _rgb(a)
    br, bg, bb = _rgb(b)
    return _hex((int(ar + (br - ar) * t), int(ag + (bg - ag) * t), int(ab + (bb - ab) * t)))


def darken(color: str, amount: float = 0.18) -> str:
    return blend(color, "#000000", amount)


def now_date_str() -> str:
    return datetime.now().strftime("%y%m%d")


def load_image(path: str):
    if not PIL_AVAILABLE:
        return None
    try:
        if path and os.path.exists(path):
            return Image.open(path)
    except Exception:
        return None
    return None


def resize_keep_ratio(image, max_w, max_h):
    ow, oh = image.size
    scale = min(1, max_w / ow, max_h / oh)
    return image.resize((max(1, int(ow * scale)), max(1, int(oh * scale))))


def find_intersport_logo():
    for name in INTERSPORT_LOGO_CANDIDATES:
        path = os.path.join(IMG_DIR, name)
        if os.path.exists(path):
            return path
    return None


class ArrowIndicator(tk.Canvas):
    def __init__(self, parent, direction, command, size=46):
        super().__init__(parent, width=size, height=size, bg=BG, highlightthickness=0, bd=0)
        self.direction = direction
        self.command = command
        self.size = size
        self.enabled = True
        self.hover = False
        self.bind("<Enter>", self._enter)
        self.bind("<Leave>", self._leave)
        self.bind("<Button-1>", self._click)
        self._draw()

    def tip_offset_from_center(self) -> float:
        s = self.size
        tip_y = (s - 11) if self.direction == "down" else 11
        return tip_y - (s / 2)
        self.bump_version_once(
            "2026-05-30_bu31_compliance_audit_deadlines_niketools_unsaved_fix",
            [
                "BU31: Stichtagspflege vollständig nach Compliance & Audit verschoben und aus Abschlusskalender entfernt.",
                "BU31: Tools - Hauptbuch erhält ein funktionsfähiges Untermenü Nike-Tools mit den drei Nike-Modulen im Standard-Modul-Menü-Layout.",
                "BU31: Modul Rechnungen aus Ordner sammeln in Nike - Rechnungs-PDFs in Sammelordner umbenannt und Modulbeschreibung aktualisiert.",
                "BU31: Dialog für ungespeicherte Änderungen nur bei echtem Dirty-State; nach Speichern + Übernehmen wird der Status bereinigt.",
                "BU31: Audit-Cockpit zeigt Details als öffnen mit ausführlichem Popup inklusive Zeitstempel.",
                "BU31: Default-Abschlussstichtag auf ersten Werktag nach Periodenende mit BW-Feiertagen umgestellt und Stichtage in Abschluss-/Steuermeldungsdaten synchronisiert.",
                "BU31: Benutzerregel angepasst: Nur Wagnerm darf Wagnerm umbenennen; E3 und niedriger dürfen sich nicht selbst umbenennen.",
            ],
        )
        self.bump_version_once(
            "2026-05-30_bu32_icons_deadline_sync_version_429",
            [
                "BU32: Nike-Tools-Kacheln zeigen wieder PDF- bzw. PDF/Excel-Icons im Untermenü Nike-Tools.",
                "BU32: Stichtagspflege synchronisiert gepflegte Abschluss-Stichtage verbindlich in Monats-, Quartals- und Jahresabschluss.",
                "BU32: Steuermeldungs-Cockpit übernimmt bei Kalender-Sync die gepflegten Stichtage als Fälligkeit je Zeitraum.",
                "BU32: Versionierung auf v0.429 fortgeschrieben und Versionsverlauf ergänzt.",
            ],
        )


    def set_enabled(self, enabled: bool):
        self.enabled = bool(enabled)
        self.configure(cursor="hand2" if self.enabled else "arrow")
        self._draw()

    def _enter(self, *_):
        self.hover = True
        self._draw()

    def _leave(self, *_):
        self.hover = False
        self._draw()

    def _click(self, *_):
        if self.enabled and self.command:
            self.command()

    def _draw(self):
        self.delete("all")
        s = self.size
        if self.hover and self.enabled:
            self.create_oval(4, 4, s - 4, s - 4, fill=blend(BG, WHITE, 0.55), outline="")
        c = BLUE if self.enabled else GREY_DISABLED
        if self.direction == "up":
            pts = (s / 2, 11, 13, s - 14, s - 13, s - 14)
            self.create_polygon(*pts, fill=c, outline=c)
            self.create_rectangle(s / 2 - 4, s - 16, s / 2 + 4, s - 11, fill=c, outline=c)
        else:
            pts = (13, 14, s - 13, 14, s / 2, s - 11)
            self.create_polygon(*pts, fill=c, outline=c)
            self.create_rectangle(s / 2 - 4, 11, s / 2 + 4, 16, fill=c, outline=c)


class Tile(tk.Canvas):
    def __init__(self, parent, app, tile_id, title, command=None, favorite_enabled=False, fixed_color=None, lock_tile=False, center_text=False, icon_type=None, corner_fold=False):
        super().__init__(parent, highlightthickness=0, bd=0, bg=BG, cursor="arrow", takefocus=True)
        self.app = app
        self.tile_id = tile_id
        self.title = title
        self.command = command
        self.favorite_enabled = favorite_enabled
        self.fixed_color = fixed_color
        self.lock_tile = lock_tile
        self.center_text = center_text
        self.icon_type = icon_type
        self.corner_fold = corner_fold
        self.tile_width = ui_s(300)
        self.tile_height = ui_s(150)
        self.hovered = False
        self.pressed = False
        self.star_bounds = None
        self.star_click_started = False
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<ButtonPress-1>", self.on_press)
        self.bind("<ButtonRelease-1>", self.on_release)
        self.bind("<FocusIn>", lambda e: self.app.set_focused_tile(self))
        self.bind("<Key-Return>", self.on_keyboard_activate)
        self.bind("<space>", self.on_keyboard_activate)
        self.draw()

    def base_color(self):
        return self.fixed_color or (self.app.current_tile_color() or BLUE)

    def current_color(self):
        c = self.base_color()
        if self.pressed:
            return darken(c, 0.28)
        if self.hovered:
            return darken(c, 0.16)
        return c

    def resize_tile(self, width, height):
        self.tile_width = int(width)
        self.tile_height = int(height)
        self.configure(width=self.tile_width + 16, height=self.tile_height + 16)
        self.draw()

    def title_font(self):
        return FONT_TILE_SMALL if len(self.title) > 28 else FONT_TILE

    def _font_with_size(self, font_spec, size):
        try:
            actual = tkfont.Font(root=self, font=font_spec).actual()
            family = actual.get('family') or 'Segoe UI'
            styles = []
            if (actual.get('weight') or '').lower() == 'bold':
                styles.append('bold')
            if (actual.get('slant') or '').lower() == 'italic':
                styles.append('italic')
            if bool(actual.get('underline')):
                styles.append('underline')
            return tuple([family, max(8, int(size))] + styles)
        except Exception:
            try:
                family = font_spec[0]
                rest = list(font_spec[2:]) if len(font_spec) > 2 else []
                return tuple([family, max(8, int(size))] + rest)
            except Exception:
                return ('Segoe UI', max(8, int(size)), 'bold')

    def _measure_wrapped_text_height(self, text, font_spec, max_width):
        try:
            f = tkfont.Font(root=self, font=font_spec)
            words = str(text or '').split() or ['']
            lines = []
            current = ''
            for word in words:
                candidate = word if not current else current + ' ' + word
                if f.measure(candidate) <= max_width:
                    current = candidate
                    continue
                if current:
                    lines.append(current)
                segment = ''
                for ch in word:
                    cand = segment + ch
                    if not segment or f.measure(cand) <= max_width:
                        segment = cand
                    else:
                        lines.append(segment)
                        segment = ch
                current = segment
            if current:
                lines.append(current)
            line_h = max(1, int(f.metrics('linespace')))
            return max(1, len(lines)) * line_h, max(1, len(lines))
        except Exception:
            return 0, 1

    def fitted_title_font(self, max_width, max_height, centered=False):
        """Lesbarkeitsschutz v0.436: bevorzugt größere Schrift, bleibt aber im Sollbereich."""
        base_font = self.app.zoomed_content_font(FONT_TILE_SMALL if len(self.title) > 24 else self.title_font())
        try:
            actual = tkfont.Font(root=self, font=base_font).actual()
            base_size = abs(int(actual.get('size') or 12))
        except Exception:
            base_size = 12
        preferred_min = 13 if not centered else 12
        start_size = max(base_size, ui_s(preferred_min))
        hard_max = max(start_size, 18 if not centered else 17)
        min_size = 9 if not centered else 8
        max_lines = 3 if not centered else 4
        chosen = self._font_with_size(base_font, min_size)
        chosen_h = 0
        for size in range(hard_max, min_size - 1, -1):
            candidate = self._font_with_size(base_font, size)
            text_h, text_lines = self._measure_wrapped_text_height(self.title, candidate, max_width)
            if text_h <= max_height and text_lines <= max_lines:
                return candidate, text_h
            chosen = candidate
            chosen_h = text_h
        return chosen, chosen_h

    def _draw_corner_fold(self, x1, y0):
        size = min(ui_s(30), max(18, int(self.tile_height * 0.20)))
        self.create_polygon(x1 - size, y0, x1, y0, x1, y0 + size, fill="#D6DCE4", outline="#C2CAD5")
        self.create_line(x1 - size, y0, x1, y0 + size, fill="#EEF2F6", width=1)

    def _draw_worksheet_icon(self, cx, cy):
        color = BLUE if self.lock_tile else "white"
        lw = 2
        # Nur das vordere Blatt: A4-Verhältnis, ohne hinteres Rahmen-/Blatt-Element
        page_w = 34
        page_h = 48
        x0 = cx - page_w / 2
        y0 = cy - page_h / 2
        x1 = cx + page_w / 2
        y1 = cy + page_h / 2
        fold = 10
        self.create_line(x0, y0, x1 - fold, y0, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1 - fold, y0, x1, y0 + fold, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1, y0 + fold, x1, y1, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1, y1, x0, y1, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x0, y1, x0, y0, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1 - fold, y0, x1 - fold, y0 + fold, fill=color, width=lw)
        self.create_line(x1 - fold, y0 + fold, x1, y0 + fold, fill=color, width=lw)
        # kurze Textlinien oben wie in der Vorlage
        self.create_line(x0 + 6, y0 + 9, x0 + 18, y0 + 9, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x0 + 6, y0 + 15, x0 + 24, y0 + 15, fill=color, width=lw + 1, capstyle="round")
        # Tabellenbereich im unteren Blattbereich
        tx0, ty0, tx1, ty1 = x0 + 5, y0 + 23, x1 - 5, y1 - 5
        self.create_rectangle(tx0, ty0, tx1, ty1, outline=color, width=lw)
        for x in (tx0 + (tx1 - tx0) / 4, tx0 + (tx1 - tx0) / 2, tx0 + 3 * (tx1 - tx0) / 4):
            self.create_line(x, ty0, x, ty1, fill=color, width=lw)
        for y in (ty0 + (ty1 - ty0) / 4, ty0 + (ty1 - ty0) / 2, ty0 + 3 * (ty1 - ty0) / 4):
            self.create_line(tx0, y, tx1, y, fill=color, width=lw)

    def _draw_calendar_icon(self, cx, cy):
        color = "white"
        lw = 3
        # rounded outline built with arcs + lines
        self.create_arc(cx - 25, cy - 19, cx - 7, cy - 1, start=90, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx + 7, cy - 19, cx + 25, cy - 1, start=0, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx - 25, cy + 8, cx - 7, cy + 26, start=180, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx + 7, cy + 8, cx + 25, cy + 26, start=270, extent=90, outline=color, width=lw, style="arc")
        self.create_line(cx - 16, cy - 19, cx + 16, cy - 19, fill=color, width=lw)
        self.create_line(cx - 25, cy - 10, cx - 25, cy + 17, fill=color, width=lw)
        self.create_line(cx + 25, cy - 10, cx + 25, cy + 17, fill=color, width=lw)
        self.create_line(cx - 16, cy + 26, cx + 16, cy + 26, fill=color, width=lw)
        self.create_line(cx - 25, cy - 5, cx + 25, cy - 5, fill=color, width=lw)
        # rings
        self.create_line(cx - 13, cy - 25, cx - 13, cy - 13, fill=color, width=5, capstyle="round")
        self.create_line(cx + 13, cy - 25, cx + 13, cy - 13, fill=color, width=5, capstyle="round")
        # day boxes
        for x, y in [(cx - 12, cy + 4), (cx, cy + 4), (cx + 12, cy + 4), (cx - 12, cy + 16), (cx, cy + 16)]:
            self.create_rectangle(x - 4, y - 4, x + 4, y + 4, outline=color, width=2)
        # check mark
        self.create_line(cx + 10, cy + 16, cx + 15, cy + 21, fill=color, width=lw, capstyle="round")
        self.create_line(cx + 15, cy + 21, cx + 24, cy + 10, fill=color, width=lw, capstyle="round")

    def _draw_module_menu_icon(self, cx, cy):
        self._draw_worksheet_icon(cx, cy)

    def _draw_gear_icon(self, cx, cy):
        color = "white"
        for dx, dy in [(0, -16), (0, 16), (-16, 0), (16, 0), (-11, -11), (11, -11), (-11, 11), (11, 11)]:
            self.create_rectangle(cx + dx - 3, cy + dy - 3, cx + dx + 3, cy + dy + 3, fill=color, outline=color)
        self.create_oval(cx - 13, cy - 13, cx + 13, cy + 13, outline=color, width=3)
        self.create_oval(cx - 5, cy - 5, cx + 5, cy + 5, outline=color, width=2)

    def _draw_info_icon(self, cx, cy):
        color = "white"
        self.create_oval(cx - 18, cy - 18, cx + 18, cy + 18, outline=color, width=4)
        self.create_oval(cx - 3, cy - 12, cx + 3, cy - 6, fill=color, outline=color)
        self.create_rectangle(cx - 4, cy - 1, cx + 4, cy + 13, fill=color, outline=color)
        self.create_rectangle(cx - 8, cy + 9, cx + 8, cy + 15, fill=color, outline=color)

    def _draw_lock_icon(self, cx, cy):
        color = BLUE if self.lock_tile else "white"
        lw = 4
        self.create_arc(cx - 19, cy - 27, cx + 19, cy + 11, start=0, extent=180, style="arc", outline=color, width=lw)
        self.create_line(cx - 19, cy - 8, cx - 19, cy - 1, fill=color, width=lw)
        self.create_line(cx + 19, cy - 8, cx + 19, cy - 1, fill=color, width=lw)
        self.create_rectangle(cx - 23, cy - 4, cx + 23, cy + 26, outline=color, width=lw)
        self.create_oval(cx - 5, cy + 6, cx + 5, cy + 16, outline=color, width=lw)
        self.create_line(cx, cy + 15, cx, cy + 23, fill=color, width=lw)

    def _draw_main_icon(self, icon_type, cx, cy):
        if hasattr(self.app, "draw_tile_icon_image") and self.app.draw_tile_icon_image(self, icon_type, cx, cy):
            return
        if icon_type == "modules":
            self._draw_module_menu_icon(cx, cy)
        elif icon_type == "worksheet":
            self._draw_worksheet_icon(cx, cy)
        elif icon_type == "calendar":
            self._draw_calendar_icon(cx, cy)
        elif icon_type == "gear":
            self._draw_gear_icon(cx, cy)
        elif icon_type == "info":
            self._draw_info_icon(cx, cy)
        elif icon_type == "lock":
            self._draw_lock_icon(cx, cy)

    def draw(self):
        self.delete("all")
        pad = ui_s(8)
        off = ui_s(1) if self.pressed else 0
        x0, y0 = pad + off, pad + off
        x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
        if not self.pressed:
            self.create_rectangle(pad + 5, pad + 6, pad + 5 + self.tile_width, pad + 6 + self.tile_height, fill=SHADOW, outline=SHADOW)
        self.create_rectangle(x0, y0, x1, y1, fill=self.current_color(), outline=self.current_color())
        if self.corner_fold:
            self._draw_corner_fold(x1, y0)
        title_y = y0 + ui_s(14)
        title_color = BLUE if self.lock_tile else "white"
        icon_to_draw = "lock" if self.lock_tile else self.icon_type
        text_width = max(ui_s(110), self.tile_width - ui_s(44))
        if self.center_text and not icon_to_draw:
            centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
            self.create_text((x0 + x1) / 2, (y0 + y1) / 2, text=self.title, anchor="center", fill=title_color, font=centered_font, width=text_width, justify="center")
        else:
            title_max_h = max(ui_s(28), int((y1 - y0) * 0.36))
            title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
            self.create_text((x0 + x1) / 2, title_y, text=self.title, anchor="n", fill=title_color, font=title_font, width=text_width, justify="center")
            if icon_to_draw:
                min_icon_y = title_y + title_h + ui_s(20)
                default_icon_y = y0 + int((y1 - y0) * 0.66)
                icon_y = max(default_icon_y, min_icon_y)
                icon_y = min(icon_y, y1 - ui_s(34))
                close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, "current_close_period_label") else ""
                if close_period and self.tile_id in ("monthly_close", "quarterly_close", "yearly_close"):
                    icon_x = x0 + 98
                    self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                    self.create_text(icon_x + 48, icon_y, text=close_period, anchor="w", fill=title_color, font=self.app.zoomed_content_font(("Segoe UI", 13, "italic")), width=max(120, x1 - icon_x - 58), justify="left")
                else:
                    self._draw_main_icon(icon_to_draw, (x0 + x1) / 2, icon_y)
        self.star_bounds = None
        if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
            sx, sy = x1 - 24, y0 + 24
            self.star_bounds = (sx - 18, sy - 18, sx + 18, sy + 18)
            self.create_text(sx, sy, text="★", fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=("Segoe UI", 24, "bold"))

    def point_in_star(self, x, y):
        return bool(self.star_bounds and self.star_bounds[0] <= x <= self.star_bounds[2] and self.star_bounds[1] <= y <= self.star_bounds[3])

    def on_enter(self, *_):
        self.hovered = True
        self.configure(cursor="hand2")
        self.draw()

    def on_leave(self, *_):
        self.hovered = False
        self.pressed = False
        self.star_click_started = False
        self.configure(cursor="arrow")
        self.draw()

    def on_press(self, event=None):
        self.focus_set()
        self.app.set_focused_tile(self)
        self.star_click_started = bool(event and self.point_in_star(event.x, event.y))
        self.pressed = not self.star_click_started
        self.draw()

    def on_release(self, event=None):
        if self.star_click_started:
            self.star_click_started = False
            if event and self.point_in_star(event.x, event.y):
                self.app.toggle_favorite(self.tile_id)
            return
        was = self.pressed
        self.pressed = False
        self.draw()
        if was and self.command:
            self.after(80, self.command)

    def on_keyboard_activate(self, *_):
        if self.command:
            self.command()


class FiBuMateApp:
    def zoomed_content_font(self, font_tuple):
        """v0.436: Kompatibilitätswrapper ohne manuellen Bereichs-/Textzoom."""
        return font_tuple

    def _calculate_ui_scale(self):
        try:
            sw = max(1, self.root.winfo_screenwidth())
            sh = max(1, self.root.winfo_screenheight())
            ww = max(1, self.root.winfo_width() or sw)
            wh = max(1, self.root.winfo_height() or sh)
            scale = min(ww / 1920.0, wh / 1080.0)
            if ww <= 1 or wh <= 1:
                scale = min(sw / 1920.0, sh / 1080.0)
            return max(0.72, min(1.18, scale))
        except Exception:
            return 1.0

    def init_responsive_scaling(self):
        """v0.436: zentrale automatische Skalierung ohne Benutzer-Zoomprofile."""
        global UI_SCALE, FONT_TITLE, FONT_MENU, FONT_TILE, FONT_TILE_SMALL, FONT_SMALL, MINI_WIDGET_W, MINI_WIDGET_H, MINI_WIDGET_GAP
        try:
            UI_SCALE = self._calculate_ui_scale()
            self.ui_scale = UI_SCALE
            FONT_TITLE = scaled_font(BASE_FONT_TITLE)
            FONT_MENU = scaled_font(BASE_FONT_MENU)
            FONT_TILE = scaled_font(BASE_FONT_TILE)
            FONT_TILE_SMALL = scaled_font(BASE_FONT_TILE_SMALL)
            FONT_SMALL = scaled_font(BASE_FONT_SMALL)
            MINI_WIDGET_W = ui_s(150)
            MINI_WIDGET_H = ui_s(26)
            MINI_WIDGET_GAP = ui_s(7)
            try:
                self.root.tk.call('tk', 'scaling', max(0.80, min(1.12, UI_SCALE)))
            except Exception:
                pass
        except Exception:
            self.ui_scale = 1.0

    def __init__(self):
        self.root = tk.Tk()
        install_central_readability(self.root)
        self.init_responsive_scaling()
        self.root.title(APP_NAME)
        self.root.protocol("WM_DELETE_WINDOW", self.confirm_exit)
        self._install_modal_toplevel_patch()
        maximize_window(self.root)
        self.page_history = []
        self.breadcrumb = []
        self.current_page = "launch"
        self.current_title = ""
        self.knowledge_view = "all"
        self.knowledge_show_start = self.load_knowledge_start_preference()
        self.knowledge_start_overlay = False
        self.knowledge_unsaved = False
        self.knowledge_overlay_offset_x = 0
        self.knowledge_overlay_offset_y = 0
        self.knowledge_overlay_drag = None
        self.widget_items = []
        self.focusable_tiles = []
        self.focus_index = -1
        self._suppress_next_global_return = False
        self._closing_in_progress = False
        self.global_text_zoom = 1.0
        self.zoom_profiles = {}
        self.current_scope_zoom = 1.0
        self.favorites = set()
        self.current_user_key = None
        self.current_user_display = ""
        self.user_data = self.load_user_data()
        self._user_data_mtime = self._get_user_data_mtime()
        self._live_permissions_started = False
        self._live_permissions_popup_open = False
        self._update_watcher_started = False
        self._update_popup_open = False
        self._notified_update_keys = set()
        self._last_update_event_seen = ""
        self.ensure_permissions_defaults()
        self.version_state = self.load_version_state()
        self.bump_version_once("2026-06-12_wissenszentrale_categories_re_import_fix", ["Wissenszentrale: Kategorienmenü repariert; fehlenden re-Import ergänzt; Pfad-Docstring für Python 3.14 bereinigt; Render-Safety für Kategorienverwaltung ergänzt."])
        self.bump_version_once("2026-06-12_wissenszentrale_kategorien_farblogik_final", ["Wissenszentrale: Kategorienverwaltung um Kategorie-Erstellung, Umbenennen und Farbbearbeitung erweitert; ausgewählte Kategorie wird in ihrer Farbe hervorgehoben; Farblogik inkl. Kontrast/Fallback geprüft."])
        self.bump_version_once("2026-06-12_wissenszentrale_gdrive_local_fallback", ["Wissenszentrale: Speichern primär auf G:-Firmenlaufwerk; falls G: nicht verfügbar ist, lokaler Fallback nach C:\\python\\knowledge_entries.json."])
        self.bump_version_once("2026-06-12_wissenszentrale_position_gdrive_rendertest", ["Wissenszentrale: Hauptfenster nach unten unter die Filter verschoben; 10-fach-Geometrieprüfung dokumentiert; Speichern strikt auf G:-Firmenlaufwerk ohne lokalen Fallback umgestellt."])
        self.bump_version_once("2026-06-12_wissenszentrale_visible_buttons_fullheight", ["Wissenszentrale: Hauptfenster bei unveränderter X/Y-Position bis kurz vor Fußleiste erweitert; Formularbuttons dauerhaft sichtbar unten angedockt."])
        self.bump_version_once("2026-06-12_wissenszentrale_save_fix", ["Wissenszentrale: Speicherlogik für neue und bearbeitete Einträge robust korrigiert; Speichern bricht nicht mehr still durch Berechtigungs-/Pfadlogik ab."])
        self.bump_version_once("2026-06-12_wissenszentrale_next_blocks_complete", ["Wissenszentrale Next Blocks: Word-Export, Anhänge, Kommentare, B2-Kategorienverwaltung, ESC-Logik, neutrale Navigation und To-Do-Rhythmus-Logik ergänzt."])
        self.bump_version_once("2026-06-12_wissenszentrale_functional_complete", ["Wissenszentrale: Layout kollisionsfrei neu strukturiert und Hauptfenster nach unten verschoben.", "Einträge können jetzt erstellt, gespeichert, bearbeitet, gesucht und über Kategorien gefiltert werden.", "Trefferliste, Gesamtliste, Detailansicht, To-Do-Filter und Kategorienanzeige technisch angebunden.", "Datenhaltung als lokale JSON-Datei unter APPDATA vorbereitet."])
        self.bump_version_once("2026-06-12_wissenszentrale_block5_overlay_underlay_panes", ["Wissenszentrale Block 5: Startseiten-Overlay maximiert zwischen FiBu-Mate-Kopf- und Fußleiste.", "Unterliegende Arbeitsansicht-Widgets werden während der Startseite ausgeblendet, damit keine Buttons/Dropdowns durchdrücken.", "Trefferfenster und Hauptfenster gemäß Markierung deutlich größer und höher ausgerichtet.", "Neuer-Eintrag-Ansicht nutzt den vergrößerten Hauptbereich besser aus."])
        self.bump_version_once("2026-06-12_wissenszentrale_block4_list_filters_overlay_fixed", ["Wissenszentrale Block 4: Startseiten-Overlay auf konstante Größe umgestellt; Drag verschiebt nur noch die Position.", "Treffer- und Hauptbereich deutlich vergrößert/dominanter dargestellt.", "Standardansicht im Hauptbereich auf Gesamtliste aller Einträge umgestellt, chronologisch neueste oben.", "Kategorie-Filterleiste mit vier Dropdowns ergänzt; To-Do-/Rhythmusfilter erscheint nur noch in der To-Do-Ansicht."])
        self.bump_version_once("2026-06-12_wissenszentrale_block3_drag_layout_cleanup", ["Wissenszentrale Block 3: Startseiten-Overlay mit sauberer Klick-/Drag-Erkennung korrigiert.", "Arbeitsansicht-Buttons nach unten versetzt, damit die Überschrift nicht mehr überlagert wird.", "Start-Overlay nach oben vergrößert und Drag-Ende strikt an das Loslassen der linken Maustaste gebunden.", "Beispieleinträge und Beispielkategorien aus der Wissenszentrale entfernt; sichtbar bleiben nur To-Do-bezogene Filter/Rhythmuswerte.", "Tools-Menüzeile im Hauptmenü kleiner skaliert als die Hauptmodul-Zeile darüber."])
        self.bump_version_once("2026-06-12_wissenszentrale_start_overlay", ["Wissenszentrale als eigener Hauptmenüpunkt nach Abschlusskalender vorbereitet.", "Arbeitsansicht nach Vorschlag A und vorgeschaltete Startseite nach Vorschlag E als Overlay umgesetzt.", "Startseiten-Checkbox speichert, ob die Startseite beim Modulstart angezeigt wird; Home-Button 'Start' öffnet die Startseite jederzeit über der aktiven Arbeitsansicht.", "Beim Wechsel der Arbeitsansicht über das Start-Overlay wird bei ungespeicherten Änderungen Speichern/Verwerfen/Abbrechen abgefragt."])
        self.bump_version_once("2026-06-12_central_readability_uplift_v3", ["Zentrale Lesbarkeitsverbesserung V3: Fließtexte, Beschreibungen, Buttons und Tabellen-/ttk-Texte moderat weiter vergrößert.", "Überschriften bleiben unverändert; nur Mindestgrößen für kleine Texte werden angehoben.", "Canvas-Texte mindestens 13 pt, Labels mindestens 13 pt, Buttons/Eingaben/ttk mindestens 12 pt, Treeview-Zeilenhöhe mindestens 28 px."])
        self.bump_version_once("2026-06-12_central_readability_uplift_v2", ["Zentrale Lesbarkeitsverbesserung V2: Mindestschriftgrößen sichtbar angehoben.", "Canvas-Texte werden jetzt ebenfalls zentral erfasst, damit Menü- und Modulbeschreibungen tatsächlich größer dargestellt werden.", "Spätere .configure(font=...)-Aufrufe werden abgefangen; dadurch wirkt die Verbesserung auch stärker in nachgeladenen Tool-Modulen."])
        self.bump_version_once("2026-06-12_central_readability_uplift", ["Zentrale Lesbarkeitsverbesserung ergänzt: sehr kleine Label-, Button-, Eingabe-, Text- und ttk-Schriften werden konservativ auf mindestens 10 pt angehoben.", "Die Verbesserung geht von Fibu_mate.py aus und wirkt auch auf nachgeladene Tool-Module, ohne diese einzeln anzupassen.", "Layouts bleiben geschützt: Es werden keine Tabellenstrukturen, Modulgrößen oder Buttonbreiten pauschal verändert."])
        self.bump_version_once(
            "2026-05-11_close_cutoff_subtasks_team_members_due_rules",
            [
                "Abschlusskalender angepasst: Kreditoren heißt nun Zentralregulierung, Controlling heißt nun Treasury.",
                "Monats- und Quartalsabschluss erweitert: Team-Mitglieder können im Bearbeitungsmodus je Übersichtskachel gepflegt und unter der nächsten Frist angezeigt werden.",
                "Monats- und Quartalsabschluss erweitert: Aufgaben können Unteraufgaben erhalten; Unteraufgaben sind separat abhakbar und steuern den Erledigt-Status der Hauptaufgabe.",
                "Fälligkeitsarten im Monats- und Quartalsabschluss auf Abschluss-Stichtag, x. Werktag des Folgemonats, x. Tag des Kalendermonats und Konkretes Datum umgestellt.",
                "Abschluss-Stichtag kann in der Abschlussübersicht gepflegt werden und wird automatisch auf Aufgaben mit Fälligkeitsart Abschluss-Stichtag übertragen.",
            ],
        )
        self.bump_version_once(
            "2026-05-11_nike_opliste_export_thread_performance_fix",
            [
                "Nike OP-Liste Modul stabilisiert: Export läuft im Hintergrund, damit FiBu Mate während großer .xlsx-Auswertungen nicht einfriert.",
                "Nike OP-Liste Modul optimiert: OP-Liste wird performanter über Zeilen-Iteratoren eingelesen.",
                "Nike OP-Liste Modul verbessert: Export-Button wird während der Verarbeitung gesperrt und Fortschritt/Fehler werden sauber zurückgemeldet.",
            ],
        )
        self.bump_version_once(
            "2026-05-11_nike_obsolete_removed_opliste_pdf_check_added",
            [
                "Nike-Module bereinigt: „Nike - Differenzen: PDF vs ZMIR6“ wurde entfernt, da das Modul obsolet ist.",
                "Nike-Module bereinigt: „Nike - Differenzbericht Sales und Lager“ wurde entfernt, da das Modul obsolet ist.",
                "Neues Modul ergänzt: „Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen“.",
                "Neues Nike-Modul vergleicht Rechnungsnummern aus OP-Liste Excel Spalte B mit Rechnungsnummern aus PDF-Dateinamen und exportiert zwei Auswertungsblätter.",
            ],
        )
        self.bump_version_once(
            "2026-05-12_monthly_close_recurring_workday_catalog_due_visibility",
            [
                "Monatsabschluss erweitert: Aufgaben können im Bearbeitungsmodus als wiederkehrend markiert werden.",
                "Monatsabschluss erweitert: zentraler Aufgabenkatalog für wiederkehrende Aufgaben eingeführt.",
                "Monatsabschluss erweitert: wiederkehrende Aufgaben werden automatisch in Folgemonate ab dem aktuellen Monat übernommen.",
                "Monatsabschluss erweitert: Fälligkeit kann als konkretes Datum, x. Werktag des Monats oder x. Werktag des Folgemonats gepflegt werden.",
                "Monatsabschluss erweitert: Werktagsberechnung berücksichtigt Montag bis Freitag sowie Feiertage Baden-Württemberg und verschiebt auf den nächsten Werktag.",
                "Monatsabschluss verbessert: berechnetes Fälligkeitsdatum wird im Aufgaben-Dialog direkt als Vorschau angezeigt.",
                "Monatsabschluss verbessert: bei Werktags-Fälligkeit wird das Eingabefeld für ein konkretes Fälligkeitsdatum ausgeblendet.",
                "Monatsabschluss verbessert: Tabelle zeigt bei regelbasierter Fälligkeit zusätzlich den Regelhinweis an.",
                "Monatsabschluss angepasst: ZM/Z4/Z5a bleiben vorerst manuell über Aufgabe, Fristart, Priorität und Fälligkeitsregel pflegbar.",
            ],
        )

        self.bump_version_once(
            "2026-05-13_user_permissions_close_sync_transfer_deadline_defaults",
            [
                "Benutzerverwaltung erweitert: Standardbenutzer sehen nur eigene Benutzerdaten, dürfen die eigene E-Mail-Adresse pflegen und die Benutzerliste ist scrollbar mit sichtbarer Neuanlage für System-Administrator.",
                "Berechtigungen angepasst: Rolle Wagnerm heißt nun System-Administrator, Standardbenutzer können das Berechtigungsmenü nicht mehr öffnen und das Berechtigungsmenü ist scrollbar.",
                "Monats- und Quartalsabschluss erweitert: Aufgaben können inklusive Unteraufgaben in den jeweils anderen Abschluss übernommen werden.",
                "Monats- und Quartalsabschluss angepasst: Standardbenutzer dürfen Aufgaben nur als erledigt markieren, wenn sie selbst als zuständig eingetragen sind.",
                "Teammitglieder werden auf folgende Perioden übertragen und zwischen Monats- und Quartalsabschluss synchronisiert.",
                "Fristart keine entfernt; Standard-Fristart ist intern.",
            ],
        )
        self.bump_version_once(
            "2026-05-13_close_team_detail_scroll_collapsible_subtasks_recurring_transfer_fix",
            [
                "Monats- und Quartalsabschluss erweitert: Unteraufgaben sind in der Team-Detailsansicht standardmäßig zugeklappt und können je Aufgabe über „Unteraufgaben ausklappen >“ ein- und ausgeblendet werden.",
                "Monats- und Quartalsabschluss korrigiert: Beim Übernehmen einer Aufgabe in den jeweils anderen Abschluss wird der Status Wiederkehrend mit übernommen.",
                "Monats- und Quartalsabschluss erweitert: Aufgabenübersichten in den Team-Detailsichten sind scrollbar.",
            ],
        )

        self.bump_version_once(
            "2026-05-15_ui_adjustments_icon_background_refinement",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )

        self.bump_version_once(
            "2026-05-15_ui_adjustments_mini_widgets_help_line_worksheet",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )

        self.bump_version_once(
            "2026-05-15_close_calendar_documentation_attachments_detail_table",
            [
                "Abschlusskalender erweitert: In den Aufgaben-Detailansichten von Monats- und Quartalsabschluss wurde eine neue Spalte Dokumentation ergänzt; je Aufgabe und Unteraufgabe kann ein Dokumentationspfad als Leitfaden bzw. Aufgabenbeschreibung hinterlegt, geöffnet und geändert werden.",
                "Abschlusskalender erweitert: Das Anlagen-Popup unterstützt mehrere Anlagenpfade je Aufgabe und Unteraufgabe, manuelle Pfadeingabe, Auswahl per Dateidialog und Bemerkungen/Informationen zur Bearbeitung für alle Rollen.",
                "Abschlusskalender angepasst: Spaltenbreiten, Ausrichtungen und Fälligkeitsdarstellung in den Detailtabellen wurden optimiert; Fällig wird nun als Datum plus Fälligkeitsart angezeigt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_ico_icon_replacements",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_menu_tile_alignment_doc_attach_buttons",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_documentation_remove_option",
            [
                "Monats- und Quartalsabschluss erweitert: Im Dokumentations-Popup kann ein hinterlegter Dokumentationspfad über ein Papierkorb-Icon entfernt werden.",
                "Monats- und Quartalsabschluss erweitert: Das Entfernen einer Dokumentation ist durch die Sicherheitsabfrage „Dokumentation entfernen?“ mit Ja/Nein geschützt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_task_table_column_separators",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_edit_icon_delegate_owner",
            [
                "Monats- und Quartalsabschluss angepasst: Der Bearbeitungsmodus-Button verwendet nun das Stift-Icon als direkte Schaltfläche ohne blaue Kachel.",
                "Monats- und Quartalsabschluss erweitert: Administratoren und System-Administratoren können die Zuständigkeit einer Aufgabe im geöffneten Zeitraum per Delegieren-Button ändern.",
                "Monats- und Quartalsabschluss erweitert: Wird eine Aufgabe mit Unteraufgaben delegiert, werden alle Unteraufgaben im betreffenden Zeitraum mitdelegiert.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_delete_scope_cleanup_sort_fix",
            [
                "Monats- und Quartalsabschluss korrigiert: Beim Löschen wird die ausgewählte Aufgabe eindeutig über die konkrete Aufgabeninstanz bzw. eine stabile Aufgabenkennung identifiziert, damit keine falschen Aufgaben gelöscht werden.",
                "Monats- und Quartalsabschluss erweitert: Beim Löschen kann gewählt werden, ob die Aufgabe nur im aktuellen Zeitraum oder auch in allen folgenden Zeiträumen gelöscht wird.",
                "Monats- und Quartalsabschluss erweitert: Im Bearbeitungsmodus gibt es den Button „Alle Folgezeiträume bereinigen“, der folgende Perioden an den aktuellen Aufgabenbestand anpasst.",
                "Monats- und Quartalsabschluss angepasst: Aufgaben und Unteraufgaben werden alphabetisch nach Titel sortiert angezeigt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_complete_ids_linking_pdf_no_reportlab",
            [
                "Abschlusskalender erweitert: Jahresabschluss und Aufgaben-Historie wurden eingebunden.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Aufgaben-IDs sind im Bearbeitungsdialog sichtbar und für Administrator/System-Administrator editierbar; bestehende fachliche Aufgaben erhalten initial QM001 bis QM016.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Aufgaben werden über identische Aufgaben-ID verknüpft; bei ID-Änderung wird die alte ID mit Deaktivierungsdatum archiviert.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Delegierungen können einmalig oder permanent auf Folgezeiträume übertragen werden.",
                "PDF-Berichte korrigiert: Export funktioniert ohne externe reportlab-Abhängigkeit über einen integrierten einfachen PDF-Generator.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_task_linking_mail_delegation",
            [
                "Aufgaben-Historie erweitert: Dezenter Reiter „Aufgaben verknüpfen“ ergänzt, inklusive Vorschlägen für gleichnamige Aufgaben in Monats-, Quartals- und Jahresabschluss.",
                "Abschlusskalender erweitert: Aufgaben-Verknüpfung setzt gemeinsame IDs automatisch nach Priorität M > Q > J und archiviert alte Aufgaben-IDs mit Deaktivierungsdatum.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Bei Delegierung wird eine E-Mail an die neu zuständige Person vorbereitet; bei permanenter Delegierung mit Hinweis „bis auf Weiteres“.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_period_lock_e_roles_reporting",
            [
                "Berechtigungen auf E1 bis E4 migriert: E1 Standard, E2 Erweitert, E3 Administrator, E4 System-Administrator; unbekannte Logins legen keine Benutzer mehr automatisch an.",
                "Benutzerverwaltung erweitert: Vorname und vollständiger Name werden für Ansprache und E-Mail-Kommunikation vorbereitet.",
                "Abschlusskalender erweitert: Zeiträume können ab E3 nach Ablauf des Abschluss-Stichtags geschlossen und wieder geöffnet werden; Wiederöffnung erfordert Begründung und automatische E-Mail-Benachrichtigung an E3/E4.",
                "Abschlussberichte erweitert: Berichtsdialog mit Signatur-/Freigabeoption, ausführlichere Inhalte, Änderungsprotokoll und ReportLab-Unterstützung mit einfachem Fallback.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_ui_login_usermgmt_closetiles_automail_fix",
            [
                "Benutzeranmeldung korrigiert: Unbekannte Benutzernamen werden endgültig nicht mehr automatisch angelegt.",
                "Benutzerverwaltung angepasst: Speichern- und Löschen-Buttons sind getrennt, gleich groß und ohne Aktionsspalten-Überschrift dargestellt.",
                "Mini-Widgets angepasst: Änderung-vorschlagen-Icon wurde auf die neue Icon-Datei umgestellt.",
                "Abschlusskalender angepasst: Monats-, Quartals- und Jahresabschluss-Kacheln zeigen den aktuellen Zeitraum direkt neben dem Kalendericon.",
                "Abschlusskalender erweitert: E4 kann Auto-Mail kompakt ein- und ausschalten; Abschluss- und Wiederöffnungs-Mails beachten diese Einstellung.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_v04_compliance_audit_tax_doc_modules",
            [
                "Major-Version auf v0.4 angehoben.",
                "Neues Hauptmenü Compliance & Audit ergänzt und Hauptmenü-Anordnung angepasst.",
                "Neues Modul Steuermeldungs-Cockpit mit Meldearten, Status, Nachweisen, Freigabe, Historie, PDF-Bericht und Audit-Anbindung ergänzt.",
                "Neues Modul Audit-Cockpit mit zentralem Audit-Log, Risikostufen, Filterung, PDF-Bericht und manuellem Archivierungsbutton für E3/E4 ergänzt.",
                "Neue Dokumentationszentrale mit Suche, Pfadprüfung, manuellen Dokumenten, PDF-/Excel-Export und Fehlende-Dokumentation-Bericht ergänzt.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_scroll_modal_dropdown_shell_fix",
            [
                "UI-Interaktion zentral verbessert: Scrollbereiche reagieren nur noch, wenn tatsächlich mehr Inhalt vorhanden ist als angezeigt werden kann.",
                "Popup- und Dialogfenster werden modal behandelt; ESC schließt bevorzugt das aktive Dialogfenster statt den Hintergrund zu bedienen.",
                "Mousewheel-Routing angepasst: Dropdowns und aktive Dialoge erhalten Vorrang, damit nicht versehentlich Menüs im Hintergrund scrollen.",
            ],
        )
        self.banner_big = load_image(BANNER_GROSS_PATH)
        self.banner_small = load_image(BANNER_KLEIN_PATH)
        self.help_image = load_image(HELP_IMAGE_PATH)
        logo_path = find_intersport_logo()
        self.intersport_logo = load_image(logo_path) if logo_path else None
        self.image_refs = []
        self.ensure_version_429_once()
        self.ensure_version_430_once()
        self.ensure_version_431_once()
        self.ensure_version_432_once()
        self.ensure_version_433_once()
        self.ensure_version_434_once()
        self.ensure_version_435_once()
        self.normalize_version_after_zoom_patch()

        self.bump_version_once(
            "2026-06-05_afi_uploads_supplier_invoice_module",
            [
                "Tools - Hauptbuch erweitert: Neues Untermenü AFI-Uploads hinzugefügt.",
                "EnBW - Strom-Tanken Upload-Erstellung in das Untermenü AFI-Uploads verschoben.",
                "Neues generisches Modul Lieferanten-Rechnung zu AFI-Upload für robuste CSV-Rechnungsanalyse und AFI-CSV-Export ergänzt.",
            ],
        )

        self.bump_version_once(
            "2026-06-05_enbw_strom_tanken_upload",
            [
                "Tools - Hauptbuch erweitert: Neues Modul EnBW - Strom-Tanken Upload-Erstellung hinzugefügt.",
                "EnBW-Modul erstellt SAP-AFI-uploadfähige CSV mit unveränderter Spaltenlogik A-J, Zuordnung nach normalisiertem Kennzeichen und Hinweis-Popup bei Namensabweichungen bzw. fehlender Vorlage.",
                "EnBW-Modul berücksichtigt gerundete Netto-/MwSt-Beträge, Steuerkennzeichen VD/V2/V0 und 19%-Grundgebühren aus der EnBW-Rechnung.",
            ],
        )

        self.bump_version_once(
            "2026-06-09_v0436_team_release_scaling_menu_cleanup",
            [
                "Version 0.436: Compliance & Audit als eigene Kachel in den Bereich In Entwicklung verschoben.",
                "Abschlusskalender angepasst: Stichtagspflege zwischen Jahresabschluss und Aufgaben-Historie einsortiert und bleibt federführend für Stichtage/Zuständigkeiten.",
                "Globale Zoomfunktion inklusive sichtbarer Zoomleiste, Zoomprofile und Strg+Mausrad-Zoom deaktiviert; FiBu Mate skaliert automatisch mit Fenster- und Monitorgröße auf Referenz 1920x1080.",
                "Kopfzeile nach Finance-Mate-Vorbild optisch bereinigt: kleinere Überschrift, getrennte Breadcrumb-/Favoriten-/Mini-Widget-Zonen und optimierte Positionen für Zurück, Änderung vorschlagen und Hilfe.",
                "Fußleiste um 15% reduziert; Inhalte skalieren mit der automatischen UI-Skalierung.",
                "Lesbarkeitsschutz ergänzt: zu kleine Texte werden innerhalb ihres Sollbereichs bevorzugt größer dargestellt, ohne Strukturen zu berühren oder unschön umzubrechen.",
            ],
        )
        self.create_footer()
        self.canvas = tk.Canvas(self.root, highlightthickness=0, bg=BG, cursor="arrow")
        self.canvas.pack(side="top", fill="both", expand=True)
        self.canvas.bind("<Configure>", self.on_resize)
        self.active_scroll_canvas = None
        # v0.436: keine globalen Zoom-Bindings mehr; automatische Skalierung übernimmt die Darstellung.
        for key, handler in [("<Escape>", self.handle_escape), ("<Return>", self.handle_enter), ("<Tab>", self.handle_tab), ("<Shift-Tab>", self.handle_shift_tab), ("<ISO_Left_Tab>", self.handle_shift_tab)]:
            self.root.bind_all(key, handler)
        self.show_page("launch", add_to_history=False)

    def install_zoom_mouse_bindings(self):
        """v0.434 Paket 1C: robuste Strg+Mausrad-Bindings für Windows/Linux und Canvas/Widgets."""
        try:
            sequences = (
                "<MouseWheel>",
                "<Control-MouseWheel>",
                "<Control-Button-4>", "<Control-Button-5>",
                "<Control-ButtonPress-4>", "<Control-ButtonPress-5>",
                "<Control-0>",
            )
            for seq in sequences:
                handler = self.reset_global_text_zoom if seq == "<Control-0>" else self.on_global_mousewheel
                try:
                    self.root.bind_all(seq, handler, add="+")
                except Exception:
                    pass
                try:
                    self.root.bind(seq, handler, add="+")
                except Exception:
                    pass
                try:
                    if hasattr(self, "canvas") and self.canvas is not None:
                        self.canvas.bind(seq, handler, add="+")
                except Exception:
                    pass
        except Exception:
            pass

    # === FiBu Mate UI interaction helpers: START ===
    def _install_modal_toplevel_patch(self):
        try:
            if getattr(tk, "_fibu_mate_modal_patch_installed", False):
                return
            original_toplevel = tk.Toplevel
            app = self
            class FiBuMateModalToplevel(original_toplevel):
                def __init__(self, *args, **kwargs):
                    super().__init__(*args, **kwargs)
                    try:
                        self.after_idle(lambda w=self: app._auto_modalize_toplevel(w))
                    except Exception:
                        pass
            tk._fibu_mate_original_toplevel = original_toplevel
            tk.Toplevel = FiBuMateModalToplevel
            tk._fibu_mate_modal_patch_installed = True
        except Exception:
            pass

    def _auto_modalize_toplevel(self, win):
        try:
            if not win or not win.winfo_exists() or win is self.root:
                return
            if getattr(win, "_fibu_mate_no_auto_modal", False):
                return
            try:
                if bool(win.overrideredirect()):
                    return
            except Exception:
                pass
            self._make_modal(win)
            try:
                self.root.after_idle(self.apply_global_text_zoom)
            except Exception:
                pass
        except Exception:
            pass

    def _make_modal(self, win):
        try:
            if not win or not win.winfo_exists():
                return win
            if getattr(win, "_fibu_mate_modalized", False):
                return win
            win._fibu_mate_modalized = True
            try:
                win.transient(self.root)
            except Exception:
                pass
            try:
                win.grab_set()
            except Exception:
                pass
            try:
                win.focus_force()
            except Exception:
                pass
            try:
                win.bind("<Escape>", lambda event, w=win: (self._close_modal_window(w), "break")[-1], add="+")
            except Exception:
                pass
            try:
                win.protocol("WM_DELETE_WINDOW", lambda w=win: self._close_modal_window(w))
            except Exception:
                pass
            return win
        except Exception:
            return win

    def _close_modal_window(self, win=None):
        try:
            target = win
            if target is None:
                try:
                    target = self.root.grab_current()
                except Exception:
                    target = None
            if target and target is not self.root and target.winfo_exists():
                try:
                    target.grab_release()
                except Exception:
                    pass
                target.destroy()
                return True
        except Exception:
            pass
        return False

    def _has_modal_dialog(self):
        try:
            current = self.root.grab_current()
            return bool(current and current is not self.root and current.winfo_exists())
        except Exception:
            return False

    def _focus_is_dropdown_or_menu(self):
        try:
            fw = self.root.focus_get()
            if fw is None:
                return False
            cls = fw.winfo_class().lower()
            if cls in ("combobox", "tcombobox", "menu", "menubutton", "optionmenu"):
                return True
            name = str(fw).lower()
            return "popdown" in name or "combobox" in name
        except Exception:
            return False

    def _scroll_canvas_has_overflow(self, canvas):
        try:
            if canvas is None or not canvas.winfo_exists():
                return False
            canvas.update_idletasks()
            bbox = canvas.bbox("all")
            if not bbox:
                return False
            return max(0, bbox[3] - bbox[1]) > max(1, canvas.winfo_height()) + 2
        except Exception:
            return False

    def _sync_scrollbar_visibility(self, canvas, scrollbar=None):
        try:
            needs_scroll = self._scroll_canvas_has_overflow(canvas)
            if scrollbar is not None and scrollbar.winfo_exists():
                manager = scrollbar.winfo_manager()
                if needs_scroll:
                    if not scrollbar.winfo_ismapped():
                        if manager == "grid":
                            scrollbar.grid()
                        else:
                            scrollbar.pack(side="right", fill="y")
                else:
                    if scrollbar.winfo_ismapped():
                        if manager == "grid":
                            scrollbar.grid_remove()
                        else:
                            scrollbar.pack_forget()
                    try:
                        canvas.yview_moveto(0)
                    except Exception:
                        pass
            return needs_scroll
        except Exception:
            return False

    def register_scroll_canvas(self, canvas, scrollbar=None):
        try:
            if canvas is None:
                return
            canvas._fibu_mate_scrollbar = scrollbar
            def _activate(_event=None, c=canvas):
                if not self._has_modal_dialog() and self._scroll_canvas_has_overflow(c):
                    self.active_scroll_canvas = c
            def _deactivate(_event=None, c=canvas):
                if self.active_scroll_canvas is c:
                    self.active_scroll_canvas = None
            canvas.bind("<Enter>", _activate, add="+")
            canvas.bind("<Leave>", _deactivate, add="+")
            canvas.bind("<MouseWheel>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c), add="+")
            canvas.bind("<Button-4>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c, linux_delta=-1), add="+")
            canvas.bind("<Button-5>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c, linux_delta=1), add="+")
            self._sync_scrollbar_visibility(canvas, scrollbar)
        except Exception:
            pass

    def _route_mousewheel_to_canvas(self, event, canvas=None, linux_delta=None):
        try:
            if self._has_modal_dialog():
                return "break"
            if self._focus_is_dropdown_or_menu():
                return None
            target = canvas or self.active_scroll_canvas
            if target is None or not target.winfo_exists():
                return None
            scrollbar = getattr(target, "_fibu_mate_scrollbar", None)
            if not self._sync_scrollbar_visibility(target, scrollbar):
                return "break"
            if linux_delta is not None:
                units = linux_delta
            else:
                delta = getattr(event, "delta", 0)
                units = int(-delta / 120) if delta else 0
                if units == 0 and delta:
                    units = -1 if delta > 0 else 1
            if units:
                target.yview_scroll(units, "units")
                try:
                    if hasattr(self, "_update_scroll_indicators"):
                        self._update_scroll_indicators()
                except Exception:
                    pass
            return "break"
        except Exception:
            return "break"
    # === FiBu Mate UI interaction helpers: END ===
    def get_icon_photo(self, icon_key, max_w=32, max_h=32):
        if not PIL_AVAILABLE:
            return None
        if not hasattr(self, "icon_cache"):
            self.icon_cache = {}
        cache_key = (icon_key, int(max_w), int(max_h))
        if cache_key in self.icon_cache:
            return self.icon_cache[cache_key]
        file_name = ICON_FILES.get(icon_key, icon_key)
        path = os.path.join(ICON_DIR, file_name)
        img = load_image(path)
        if not img:
            return None
        try:
            img = img.convert("RGBA")
            img = resize_keep_ratio(img, max_w, max_h)
            photo = ImageTk.PhotoImage(img)
            self.icon_cache[cache_key] = photo
            return photo
        except Exception:
            return None

    def draw_canvas_icon(self, canvas, icon_key, x, y, max_w=22, max_h=22):
        photo = self.get_icon_photo(icon_key, max_w, max_h)
        if not photo:
            return False
        canvas.create_image(x, y, image=photo)
        return True

    def draw_tile_icon_image(self, tile, icon_type, cx, cy):
        mapping = {"calendar": "calendar", "gear": "gear", "info": "info", "lock": "lock", "worksheet": "xls", "modules": "xls", "compliance": "compliance", "tax_reporting": "tax_reporting", "audit": "audit", "documentation": "documentation", "knowledge": "knowledge"}
        if icon_type == "pdf_xls":
            p1 = self.get_icon_photo("pdf", 50, 50)
            p2 = self.get_icon_photo("xls", 50, 50)
            if p1 and p2:
                tile.create_image(cx - 31, cy, image=p1)
                tile.create_image(cx + 31, cy, image=p2)
                return True
            return False
        icon_key = mapping.get(icon_type)
        if not icon_key:
            return False
        photo = self.get_icon_photo(icon_key, 48, 48)
        if not photo:
            return False
        tile.create_image(cx, cy, image=photo)
        return True

    def version_label_text(self) -> str:
        build = int(self.version_state.get("build", DEFAULT_BUILD))
        return f"v{VERSION_PREFIX}{build}.{now_date_str()}"


    def ensure_version_429_once(self):
        """BU33a: Version v0.429 und Versionsverlauf robust sicherstellen.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.429 dem Build 29,
        nicht 429. Diese Methode korrigiert auch bereits falsch gespeicherte
        Builds wie 429, die als v0.4429 angezeigt werden.
        """
        update_id = "2026-05-30_bu33a_force_version_429"
        bullets = [
            "BU33a: Versionierung robust korrigiert; sichtbare Version wird auf v0.429 gesetzt.",
            "BU33a: Falsch gespeicherter Build 429 wird auf Build 29 korrigiert, damit nicht v0.4429 angezeigt wird.",
            "BU33a: Versionsverlauf-Eintrag ergänzt, auch wenn vorherige Update-Marker bereits als angewendet gespeichert waren.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 29 or current_build > 100:
                self.version_state["build"] = 29
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.429",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_430_once(self):
        """BU33b: Version v0.430 und Versionsverlauf robust sicherstellen.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.430 dem Build 30.
        Diese Methode wird nach allen regulären bump_version_once-Blöcken und vor
        create_footer() ausgeführt, damit das Versionslabel direkt korrekt ist.
        """
        update_id = "2026-05-30_bu33b_closing_deadline_sync_version_430"
        bullets = [
            "BU33b: Abschlusskalender-Synchronisation versioniert.",
            "BU33b: Monatsabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Quartalsabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Jahresabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Bestehende Periodendateien aktualisieren closing_cutoff_date und die Fälligkeiten betroffener Aufgaben.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 30 or current_build > 100:
                self.version_state["build"] = 30
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.430",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_431_once(self):
        """v0.431: Abschluss-Stichtag in Kalendern nur noch aus Stichtagspflege.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.431 dem Build 31.
        """
        update_id = "2026-05-30_v0_431_calendar_cutoff_readonly_deadline_source"
        bullets = [
            "v0.431: Abschluss-Stichtag in Monats-, Quartals- und Jahresabschluss wird in den Zeitraumsübersichten nur noch aus der Stichtagspflege angezeigt.",
            "v0.431: Die manuelle Änderung des Abschluss-Stichtags in den Kalender-Zeitraumsübersichten wurde entfernt.",
            "v0.431: Aufgaben mit Fälligkeitsart Abschluss-Stichtag verwenden den zentral gepflegten Stichtag aus der Stichtagspflege als Basis.",
            "v0.431: Bestehende Periodendaten werden beim Öffnen der Kalender erneut gegen die Stichtagspflege normalisiert.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 31 or current_build > 100:
                self.version_state["build"] = 31
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.431",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_432_once(self):
        """v0.432: Zeitraumlogik und Dokumentationszentrale vorbereitet.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.432 dem Build 32.
        """
        update_id = "2026-05-30_v0_432_period_logic_documentation_center"
        bullets = [
            "v0.432: Abschlusszeiträume starten ab 05/2026; Jahresabschluss beginnt mit GJ 2025/2026.",
            "v0.432: Monats- und Quartalszeiträume werden in der Zukunft nur bis Ende des freigegebenen Geschäftsjahres angelegt.",
            "v0.432: Folge-Geschäftsjahr wird automatisch ab dem Abschluss-Stichtag August des aktuellen Geschäftsjahres freigegeben.",
            "v0.432: Stichtagspflege verwendet das aktuelle Geschäftsjahr als Standard und arbeitet mit Geschäftsjahresperioden.",
            "v0.432: Anzeigeformate TT.MM.JJJJ, MM/JJ und QQ/JJ sowie Auswahl nur angelegter Zeiträume ergänzt.",
            "v0.432: Vorhandene ältere Zeiträume werden nicht gelöscht, aber nicht mehr angezeigt oder automatisch angelegt.",
            "v0.432: Dokumentationszentrale zeigt nur Gegenwart/Vergangenheit, erlaubt Positionsauswahl, positionsbezogenes Anhängen und einen zentralen Exportdialog.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 32 or current_build > 100:
                self.version_state["build"] = 32
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {"version": "v0.432", "date": now_date_str(), "update_id": update_id, "bullets": bullets})
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

        try:
            self.start_update_watcher()
        except Exception:
            pass

    def load_version_state(self):
        os.makedirs(USER_DIR, exist_ok=True)
        path = os.path.join(USER_DIR, VERSION_STATE_FILE)
        default = {"build": DEFAULT_BUILD}
        try:
            if not os.path.exists(path):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(default, f, ensure_ascii=False, indent=2)
                return default
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            data.setdefault("build", default["build"])
            return data
        except Exception:
            return default

    def save_version_state(self):
        try:
            with open(os.path.join(USER_DIR, VERSION_STATE_FILE), "w", encoding="utf-8") as f:
                json.dump(self.version_state, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def load_version_history(self):
        os.makedirs(USER_DIR, exist_ok=True)
        path = os.path.join(USER_DIR, VERSION_HISTORY_FILE)
        local_history = {"entries": []}
        try:
            if not os.path.exists(path):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(local_history, f, ensure_ascii=False, indent=2)
            else:
                with open(path, "r", encoding="utf-8") as f:
                    local_history = json.load(f)
        except Exception:
            local_history = {"entries": []}
        try:
            central_path = os.path.join(CENTRAL_RELEASE_DIR, VERSION_HISTORY_FILE)
            if os.path.exists(central_path):
                with open(central_path, "r", encoding="utf-8") as f:
                    central_history = json.load(f)
                merged = []
                seen = set()
                for entry in (central_history.get("entries") or central_history.get("versions") or []):
                    key = (str(entry.get("version", "")), str(entry.get("published_at", "")), str(entry.get("update_id", "")))
                    if key not in seen:
                        merged.append(entry); seen.add(key)
                for entry in local_history.get("entries", []):
                    key = (str(entry.get("version", "")), str(entry.get("published_at", "")), str(entry.get("update_id", "")))
                    if key not in seen:
                        merged.append(entry); seen.add(key)
                return {"entries": merged}
        except Exception:
            pass
        return local_history

    def save_version_history(self, history):
        try:
            with open(os.path.join(USER_DIR, VERSION_HISTORY_FILE), "w", encoding="utf-8") as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def bump_version(self, bullets):
        if not isinstance(bullets, list):
            bullets = [str(bullets)]
        self.version_state["build"] = int(self.version_state.get("build", DEFAULT_BUILD)) + 1
        self.save_version_state()
        history = self.load_version_history()
        history.setdefault("entries", [])
        history["entries"].insert(0, {"version": f"v{VERSION_PREFIX}{self.version_state['build']}", "date": now_date_str(), "bullets": bullets})
        self.save_version_history(history)
        try:
            self.version_label.config(text=self.version_label_text())
        except Exception:
            pass

    def bump_version_once(self, update_id, bullets):
        """
        Schreibt einen Versionsverlauf-Eintrag nur einmal.
        Damit wird die Version nicht bei jedem Programmstart erneut erhöht.
        """
        try:
            applied = self.version_state.setdefault("applied_updates", [])
            if update_id in applied:
                return
            self.bump_version(bullets)
            applied.append(update_id)
            self.version_state["applied_updates"] = applied
            self.save_version_state()
        except Exception:
            pass

    def _safe_json_load_file(self, path, default=None):
        try:
            if path and os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass
        return {} if default is None else default

    def _normalize_version_for_compare(self, version):
        raw = str(version or "").strip()
        if raw.lower().startswith("v"):
            raw = raw[1:]
        return raw or "0.0.0"

    def _version_tuple_for_compare(self, version):
        raw = self._normalize_version_for_compare(version)
        parts = []
        for token in raw.replace("-", ".").split("."):
            digits = "".join(ch for ch in token if ch.isdigit())
            try:
                parts.append(int(digits or "0"))
            except Exception:
                parts.append(0)
        return tuple(parts or [0])

    def current_app_version(self):
        try:
            build = int(self.version_state.get("build", DEFAULT_BUILD))
        except Exception:
            build = DEFAULT_BUILD
        return f"{VERSION_PREFIX}.{build}"

    def current_app_version_display(self):
        try:
            build = int(self.version_state.get("build", DEFAULT_BUILD))
            return f"v{VERSION_PREFIX}{build}"
        except Exception:
            return f"v{self.current_app_version()}"

    def _latest_update_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, LATEST_UPDATE_FILE)

    def _update_event_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, UPDATE_EVENT_FILE)

    def _central_version_history_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, VERSION_HISTORY_FILE)

    def load_latest_update_info(self):
        latest = self._safe_json_load_file(self._latest_update_path(), {})
        event = self._safe_json_load_file(self._update_event_path(), {})
        if event and event.get("version") and not latest.get("version"):
            latest.update(event)
        return latest

    def _patch_text_from_update_info(self, info):
        patch_text = info.get("patch_text") or info.get("bullets") or []
        if isinstance(patch_text, str):
            lines = [line.strip(" -") for line in patch_text.splitlines() if line.strip()]
        elif isinstance(patch_text, list):
            lines = [str(line).strip(" -") for line in patch_text if str(line).strip()]
        else:
            lines = []
        return lines

    def _format_update_message(self, info, start_check=False):
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        update_type = str(info.get("update_type") or "Patch").strip()
        mandatory = bool(info.get("mandatory"))
        published_at = str(info.get("published_at") or "").strip()
        patch_file = str(info.get("patch_file") or "").strip()
        lines = self._patch_text_from_update_info(info)
        bullet_text = "\n".join(f"• {line}" for line in lines[:12]) if lines else "• Kein Patchtext hinterlegt."
        title_line = "Beim Start wurde eine veraltete FiBu-Mate-Version erkannt." if start_check else "Während der Nutzung wurde ein neuer FiBu-Mate-Patch veröffentlicht."
        mandatory_text = "Ja" if mandatory else "Nein"
        return (
            f"{title_line}\n\n"
            f"Installierte Version: {self.current_app_version_display()}\n"
            f"Verfügbare Version: v{latest_version.lstrip('v')}\n"
            f"Patch-Art: {update_type}\n"
            f"Pflichtupdate: {mandatory_text}\n"
            f"Veröffentlicht: {published_at or '-'}\n"
            f"Patchdatei: {patch_file or '-'}\n\n"
            f"Änderungen:\n{bullet_text}\n\n"
            f"Der vollständige Versionsverlauf wird aus folgender Datei gelesen:\n"
            f"{self._central_version_history_path()}"
        )

    def is_newer_update_available(self, info):
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        if not latest_version:
            return False
        try:
            return self._version_tuple_for_compare(latest_version) > self._version_tuple_for_compare(self.current_app_version())
        except Exception:
            return False

    def show_update_notification(self, info, start_check=False):
        if self._update_popup_open:
            return
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        published_at = str(info.get("published_at") or "").strip()
        key = f"{latest_version}|{published_at}|{'start' if start_check else 'live'}"
        if key in self._notified_update_keys and not start_check:
            return
        self._notified_update_keys.add(key)
        self._update_popup_open = True
        try:
            title = "FiBu Mate - veraltete Version" if start_check else "FiBu Mate - neuer Patch verfügbar"
            messagebox.showinfo(title, self._format_update_message(info, start_check=start_check))
        finally:
            self._update_popup_open = False

    def check_update_on_start(self):
        try:
            info = self.load_latest_update_info()
            if self.is_newer_update_available(info):
                self.show_update_notification(info, start_check=True)
        except Exception:
            pass

    def check_update_live(self):
        try:
            info = self.load_latest_update_info()
            event = self._safe_json_load_file(self._update_event_path(), {})
            event_key = "|".join([str(event.get("version", "")), str(event.get("published_at", "")), str(event.get("patch_file", ""))])
            if event_key and event_key != self._last_update_event_seen:
                self._last_update_event_seen = event_key
                if self.is_newer_update_available(info):
                    self.show_update_notification(info, start_check=False)
            elif self.is_newer_update_available(info):
                self.show_update_notification(info, start_check=False)
        except Exception:
            pass
        try:
            self.root.after(UPDATE_CHECK_INTERVAL_MS, self.check_update_live)
        except Exception:
            pass

    def start_update_watcher(self):
        if getattr(self, "_update_watcher_started", False):
            return
        self._update_watcher_started = True
        try:
            self.root.after(2500, self.check_update_on_start)
            self.root.after(UPDATE_CHECK_INTERVAL_MS, self.check_update_live)
        except Exception:
            pass

    def get_user_data_path(self):
        global USER_DATA_PATH
        USER_DATA_PATH = resolve_user_data_path()
        return USER_DATA_PATH

    def _get_user_data_mtime(self):
        try:
            path = self.get_user_data_path()
            return os.path.getmtime(path) if os.path.exists(path) else 0
        except Exception:
            return 0

    def load_user_data(self):
        default = {"last_username_prefill": "", "users": {}, "settings": {"auto_close_mail_enabled": True}}
        try:
            path = self.get_user_data_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            if not os.path.exists(path):
                if os.path.exists(LEGACY_USER_DATA_PATH) and os.path.abspath(LEGACY_USER_DATA_PATH) != os.path.abspath(path):
                    with open(LEGACY_USER_DATA_PATH, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                else:
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(default, f, ensure_ascii=False, indent=2)
                    return default
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            data.setdefault("last_username_prefill", "")
            data.setdefault("users", {})
            data.setdefault("settings", {})
            data["settings"].setdefault("auto_close_mail_enabled", True)
            return data
        except Exception:
            return default

    def save_user_data(self):
        try:
            path = self.get_user_data_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.user_data, f, ensure_ascii=False, indent=2)
            self._user_data_mtime = self._get_user_data_mtime()
        except Exception as e:
            messagebox.showerror("FiBu Mate", "Benutzerdaten konnten nicht gespeichert werden:\n\n" + str(e))

    def ensure_permissions_defaults(self):
        users = self.user_data.setdefault("users", {})
        for key, u in users.items():
            u.setdefault("display_name", key)
            u.setdefault("favorites", [])
            u.setdefault("email", "")
            u.setdefault("auth", {"password_hash": None, "enabled": False})
            u["favorites"] = [fav for fav in u.get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS]
            u.setdefault("first_name", "")
            u.setdefault("full_name", " ".join(x for x in [u.get("first_name", "").strip(), u.get("display_name", key).strip()] if x).strip() or u.get("display_name", key))
            current_permission = ROLE_MIGRATION.get(u.get("permission"), u.get("permission", ROLE_E1))
            if key == SUPERUSER_KEY:
                u["permission"] = ROLE_E4
            else:
                if current_permission == ROLE_E4:
                    u["permission"] = ROLE_E1
                else:
                    u["permission"] = current_permission if current_permission in ROLE_ORDER else ROLE_E1

    def my_role(self):
        if not self.current_user_key:
            return ROLE_E1
        if self.current_user_key == SUPERUSER_KEY:
            return ROLE_E4
        return ROLE_MIGRATION.get(self.user_data.get("users", {}).get(self.current_user_key, {}).get("permission", ROLE_E1), ROLE_E1)

    def role_rank(self, role):
        return ROLE_RANK.get(ROLE_MIGRATION.get(role, role), ROLE_RANK[ROLE_E1])

    def can_view_user_management(self):
        return True

    def can_create_users(self):
        return self.my_role() in (ROLE_E3, ROLE_E4)

    def can_manage_permissions(self):
        return self.my_role() in (ROLE_E3, ROLE_E4)

    def auto_close_mail_enabled(self):
        return bool(self.user_data.setdefault("settings", {}).setdefault("auto_close_mail_enabled", True))

    def toggle_auto_close_mail(self):
        settings = self.user_data.setdefault("settings", {})
        settings["auto_close_mail_enabled"] = not bool(settings.get("auto_close_mail_enabled", True))
        self.save_user_data()
        self.render_page()

    def log_audit_event(self, event_type="Info", module="FiBu Mate", title="", details="", risk="Info", period="", related_id="", public=True):
        try:
            try:
                from bin.tools import compliance_common as cc
            except Exception:
                import compliance_common as cc
            return cc.log_audit(self, event_type, module, title, details, risk, period, related_id, public)
        except Exception:
            return None

    def current_close_period_label(self, module_id):
        now = datetime.now()
        if module_id == "monthly_close":
            names = ["Januar", "Februar", "März", "April", "Mai", "Juni", "Juli", "August", "September", "Oktober", "November", "Dezember"]
            return f"{names[now.month - 1]} {now.year}"
        if module_id == "quarterly_close":
            q = (now.month - 1) // 3 + 1
            return f"Q{q} {now.year}"
        if module_id == "yearly_close":
            start_year = now.year if now.month >= 10 else now.year - 1
            return f"GJ {start_year}/{start_year + 1}"
        return ""

    def max_assignable_role_rank(self):
        return self.role_rank(self.my_role())

    def current_tile_color(self):
        if self.current_user_key:
            return self.user_data.get("users", {}).get(self.current_user_key, {}).get("tile_color")
        return None

    def create_footer(self):
        self.footer = tk.Frame(self.root, bg="black", height=ui_s(49)); self.footer._zoom_exclude = True
        self.footer.pack(side="bottom", fill="x")
        self.footer.pack_propagate(False)
        self.user_label = tk.Label(self.footer, bg="black", fg="white", font=body_font(9))
        self.user_label.place(relx=0, rely=0.5, anchor="w", x=12)
        self.version_label = tk.Label(self.footer, text=self.version_label_text(), bg="black", fg="white", font=body_font(9))
        self.version_label.place(relx=0.5, rely=0.5, anchor="center")
        self.clock_label = tk.Label(self.footer, bg="black", fg="white", font=body_font(9))
        self.clock_label.place(relx=1, rely=0.5, anchor="e", x=-12)
        self.update_clock()

    def update_clock(self):
        self.clock_label.config(text=datetime.now().strftime("%H:%M:%S"))
        self.user_label.config(text=f"Benutzer {self.current_user_display}" if self.current_user_display else "")
        self.root.after(1000, self.update_clock)


    def register_unsaved_changes_provider(self, has_unsaved_callback=None, save_callback=None, discard_callback=None):
        self._unsaved_provider = has_unsaved_callback
        self._unsaved_save_callback = save_callback
        self._unsaved_discard_callback = discard_callback

    def clear_unsaved_changes_provider(self):
        self._unsaved_provider = None
        self._unsaved_save_callback = None
        self._unsaved_discard_callback = None

    def _has_unsaved_changes(self):
        try:
            return bool(self._unsaved_provider and self._unsaved_provider())
        except Exception:
            return False

    def confirm_unsaved_changes(self):
        if not self._has_unsaved_changes():
            return True
        dlg = tk.Toplevel(self.root)
        dlg.title("Ungespeicherte Änderungen")
        dlg.configure(bg=BG)
        dlg.transient(self.root)
        dlg.grab_set()
        result = {"value": None}
        tk.Label(dlg, text="Es liegen ungespeicherte Änderungen vor.", bg=BG, fg=TEXT, font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=18, pady=(16, 6))
        tk.Label(dlg, text="Möchten Sie die Änderungen speichern und übernehmen oder verwerfen?", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).pack(anchor="w", padx=18, pady=(0, 14))
        btns = tk.Frame(dlg, bg=BG)
        btns.pack(fill="x", padx=18, pady=(0, 16))
        def choose(v):
            result["value"] = v
            dlg.destroy()
        tk.Button(btns, text="Speichern + Übernehmen", command=lambda: choose("save"), bg=BLUE, fg="white", bd=0, padx=12, pady=7).pack(side="left", padx=(0, 8))
        discard_btn = tk.Button(btns, text="Änderungen Verwerfen", command=lambda: choose("discard"), bg=WHITE, fg=RED, bd=1, padx=12, pady=7)
        try:
            if PIL_AVAILABLE:
                icon_path = os.path.join(ICON_DIR, "biggarbagebin_121980.ico")
                if os.path.exists(icon_path):
                    img = Image.open(icon_path).resize((18, 18))
                    photo = ImageTk.PhotoImage(img)
                    discard_btn.configure(image=photo, compound="left")
                    discard_btn.image = photo
        except Exception:
            pass
        discard_btn.pack(side="left", padx=(0, 8))
        tk.Button(btns, text="Abbrechen", command=lambda: choose("cancel"), bg=WHITE, fg=TEXT, bd=1, padx=12, pady=7).pack(side="right")
        self.root.wait_window(dlg)
        if result["value"] == "save":
            try:
                if self._unsaved_save_callback:
                    self._unsaved_save_callback()
                self.clear_unsaved_changes_provider()
                return True
            except Exception as exc:
                messagebox.showerror("FiBu Mate", "Speichern fehlgeschlagen: " + str(exc))
                return False
        if result["value"] == "discard":
            try:
                if self._unsaved_discard_callback:
                    self._unsaved_discard_callback()
            finally:
                self.clear_unsaved_changes_provider()
            return True
        return False

    def show_page(self, page_name, title="", add_to_history=True):
        if page_name != self.current_page and not self.confirm_unsaved_changes():
            return
        if add_to_history and self.current_page:
            self.page_history.append((self.current_page, self.current_title))
        self.current_page = page_name
        self.current_title = title
        self.update_breadcrumb(page_name, title)
        self.render_page()

    def update_breadcrumb(self, page_name, title):
        if page_name == "launch":
            self.breadcrumb = []
        elif page_name == "main":
            self.breadcrumb = [("main", "Hauptmenü")]
        else:
            if not self.breadcrumb:
                self.breadcrumb = [("main", "Hauptmenü")]
            existing = next((i for i, (p, _) in enumerate(self.breadcrumb) if p == page_name), None)
            self.breadcrumb = self.breadcrumb[: existing + 1] if existing is not None else self.breadcrumb + [(page_name, title)]

    def jump_to_breadcrumb(self, index):
        if not self.confirm_unsaved_changes():
            return
        if 0 <= index < len(self.breadcrumb):
            self.current_page, self.current_title = self.breadcrumb[index]
            self.breadcrumb = self.breadcrumb[: index + 1]
            self.page_history = self.breadcrumb[:-1].copy()
            self.render_page()

    def go_back(self):
        if not self.confirm_unsaved_changes():
            return
        if self.page_history:
            page, title = self.page_history.pop()
            self.current_page, self.current_title = page, title
            self.update_breadcrumb(page, title)
            self.render_page()

    def display_zoom_key(self):
        try:
            return f"{self.root.winfo_screenwidth()}x{self.root.winfo_screenheight()}"
        except Exception:
            return "unknown"

    def zoom_user_key(self):
        return self.current_user_key or "__anonymous__"

    def current_zoom_scope_key(self):
        try:
            if str(self.current_page or "").startswith("tool:"):
                return self.current_page
            return self.current_page or "main"
        except Exception:
            return "main"

    def load_zoom_profiles(self):
        path = os.path.join(USER_DIR, ZOOM_PROFILE_FILE)
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else {}
        except Exception:
            pass
        return {}

    def save_zoom_profiles(self):
        try:
            os.makedirs(USER_DIR, exist_ok=True)
            with open(os.path.join(USER_DIR, ZOOM_PROFILE_FILE), "w", encoding="utf-8") as f:
                json.dump(self.zoom_profiles, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def get_scope_zoom(self, scope=None):
        try:
            scope = scope or self.current_zoom_scope_key()
            user = self.zoom_user_key(); display = self.display_zoom_key()
            return float(self.zoom_profiles.get(user, {}).get(display, {}).get(scope, 1.0))
        except Exception:
            return 1.0

    def set_scope_zoom(self, zoom, scope=None):
        try:
            scope = scope or self.current_zoom_scope_key()
            user = self.zoom_user_key(); display = self.display_zoom_key()
            self.zoom_profiles.setdefault(user, {}).setdefault(display, {})[scope] = round(float(zoom), 3)
            self.current_scope_zoom = round(float(zoom), 3)
            self.save_zoom_profiles()
        except Exception:
            pass

    def prepare_scope_zoom(self):
        global GLOBAL_TEXT_ZOOM
        try:
            self.current_scope_zoom = self.get_scope_zoom()
            GLOBAL_TEXT_ZOOM = 1.0
        except Exception:
            self.current_scope_zoom = 1.0

    def is_header_footer_or_fixed_widget(self, widget):
        try:
            w = widget
            while w is not None:
                if w is getattr(self, "footer", None):
                    return True
                if getattr(w, "_zoom_exclude", False):
                    return True
                w = w.master
        except Exception:
            pass
        try:
            # Kopf-/Mini-Widgets/Breadcrumb nur im Hauptfenster über y-Koordinate ausschließen.
            if widget.winfo_toplevel() is self.root:
                y = widget.winfo_rooty() - self.canvas.winfo_rooty()
                if y < 126:
                    return True
        except Exception:
            pass
        return False

    def _event_allows_content_zoom(self, event):
        try:
            widget = getattr(event, "widget", None)
            y_root = getattr(event, "y_root", None)
            # Hauptcanvas: nicht pauschal als Header behandeln; nur die echte Kopfzone y < 126 sperren.
            if widget is getattr(self, "canvas", None):
                if y_root is not None:
                    try:
                        y = y_root - self.canvas.winfo_rooty()
                        return y >= 126
                    except Exception:
                        return True
                return True
            if widget is not None and self.is_header_footer_or_fixed_widget(widget):
                return False
            # Nur Ereignisse im Hauptfenster im Headerbereich blocken; Toplevel/Popup-Inhalte bleiben zoombar.
            if widget is not None and y_root is not None:
                try:
                    if widget.winfo_toplevel() is self.root:
                        y = y_root - self.canvas.winfo_rooty()
                        if y < 126:
                            return False
                except Exception:
                    pass
        except Exception:
            pass
        return True

    def _is_ctrl_mousewheel(self, event):
        try:
            # Windows/Tk: ControlMask ist 0x0004. Bei Control-spezifischen Bindings kann state je nach Umgebung variieren;
            # num 4/5 bzw. delta != 0 mit Control-Binding wird deshalb zusätzlich akzeptiert.
            if getattr(event, "state", 0) & 0x0004:
                return True
        except Exception:
            pass
        try:
            if getattr(event, "num", None) in (4, 5):
                return True
            if int(getattr(event, "delta", 0) or 0) != 0 and str(getattr(event, "type", "")) in ("MouseWheel", "38"):
                return bool(getattr(event, "state", 0) & 0x0004)
        except Exception:
            pass
        return False

    def _mousewheel_direction(self, event):
        try:
            if getattr(event, "num", None) == 4:
                return 1
            if getattr(event, "num", None) == 5:
                return -1
        except Exception:
            pass
        try:
            delta = int(getattr(event, "delta", 0))
            return 1 if delta > 0 else (-1 if delta < 0 else 0)
        except Exception:
            return 0

    def on_global_mousewheel(self, event):
        if self._is_ctrl_mousewheel(event):
            if not self.is_zoom_control_visible():
                return "break"
            if not self._event_allows_content_zoom(event):
                return "break"
            direction = self._mousewheel_direction(event)
            if direction:
                self.adjust_global_text_zoom(direction)
            return "break"
        return self._route_mousewheel_to_canvas(event)

    def adjust_global_text_zoom(self, direction):
        try:
            direction = 1 if direction > 0 else -1
            current = self.get_scope_zoom()
            new_zoom = current + direction * GLOBAL_TEXT_ZOOM_STEP
            new_zoom = max(GLOBAL_TEXT_ZOOM_MIN, min(GLOBAL_TEXT_ZOOM_MAX, round(new_zoom, 3)))
            if abs(new_zoom - current) < 0.0001:
                return
            self.set_scope_zoom(new_zoom)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass

    def refresh_zoomed_content(self):
        """Erzwingt die sichtbare Anwendung des Bereichszooms durch Neu-Rendern des aktuellen Inhalts."""
        try:
            if getattr(self, "current_page", "") == "launch":
                return
            self.root.after_idle(self.render_page)
        except Exception:
            try:
                self.render_page()
            except Exception:
                pass

    def reset_global_text_zoom(self, event=None):
        try:
            self.set_scope_zoom(1.0)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass
        return "break"

    def _font_actual_tuple(self, widget, font_spec):
        try:
            f = tkfont.Font(root=widget, font=font_spec)
            actual = f.actual()
            family = actual.get('family') or 'Segoe UI'
            size = abs(int(actual.get('size') or 10))
            weight = actual.get('weight') or 'normal'
            slant = actual.get('slant') or 'roman'
            underline = bool(actual.get('underline'))
            overstrike = bool(actual.get('overstrike'))
            return (family, size, weight, slant, underline, overstrike)
        except Exception:
            return ('Segoe UI', 10, 'normal', 'roman', False, False)

    def _scaled_font_from_base(self, base):
        try:
            family, size, weight, slant, underline, overstrike = base
            scaled_size = max(6, int(round(float(size) * float(self.current_scope_zoom))))
            args = [family, scaled_size]
            if weight and weight != 'normal':
                args.append(weight)
            if slant and slant != 'roman':
                args.append(slant)
            if underline:
                args.append('underline')
            if overstrike:
                args.append('overstrike')
            return tuple(args)
        except Exception:
            return ('Segoe UI', 10)

    def _apply_zoom_to_widget_tree(self, widget):
        try:
            # Canvas zuerst behandeln: Der Hauptcanvas startet bei y=0, enthält aber auch den zoombaren Inhaltsbereich.
            # Header-Texte werden in _apply_zoom_to_canvas weiterhin anhand der Item-BBox ausgeschlossen.
            if isinstance(widget, tk.Canvas):
                if not getattr(widget, "_zoom_exclude", False):
                    self._apply_zoom_to_canvas(widget)
                # Font-Konfiguration des Canvas selbst ist nicht relevant; Kinder trotzdem prüfen.
                for child in widget.winfo_children():
                    self._apply_zoom_to_widget_tree(child)
                return
            if self.is_header_footer_or_fixed_widget(widget):
                return
            try:
                font_spec = widget.cget('font')
                key = str(widget)
                if key not in self._zoom_base_fonts:
                    self._zoom_base_fonts[key] = self._font_actual_tuple(widget, font_spec)
                widget.configure(font=self._scaled_font_from_base(self._zoom_base_fonts[key]))
            except Exception:
                pass
            for child in widget.winfo_children():
                self._apply_zoom_to_widget_tree(child)
        except Exception:
            pass

    def _apply_zoom_to_canvas(self, canvas):
        try:
            ckey = str(canvas)
            base_map = self._zoom_base_canvas_fonts.setdefault(ckey, {})
            for item in canvas.find_all():
                try:
                    if canvas.type(item) != 'text':
                        continue
                    try:
                        bbox = canvas.bbox(item)
                        if bbox and bbox[1] < 126:
                            continue
                    except Exception:
                        pass
                    if item not in base_map:
                        base_map[item] = self._font_actual_tuple(canvas, canvas.itemcget(item, 'font'))
                    canvas.itemconfigure(item, font=self._scaled_font_from_base(base_map[item]))
                except Exception:
                    pass
        except Exception:
            pass

    def apply_global_text_zoom(self):
        try:
            self.prepare_scope_zoom()
            self._apply_zoom_to_widget_tree(self.root)
            for win in list(self.root.winfo_children()):
                try:
                    if isinstance(win, tk.Toplevel) and win.winfo_exists():
                        self._apply_zoom_to_widget_tree(win)
                except Exception:
                    pass
        except Exception:
            pass
    def clear_content(self):
        self.clear_unsaved_changes_provider()
        self.active_scroll_canvas = None
        if hasattr(self, "module_escape_handler"):
            delattr(self, "module_escape_handler")
        if hasattr(self, "_update_scroll_indicators"):
            delattr(self, "_update_scroll_indicators")
        self.canvas.delete("all")
        for w in self.widget_items:
            try:
                w.destroy()
            except Exception:
                pass
        self.widget_items.clear()
        self.focusable_tiles.clear()
        self.focus_index = -1
        self.image_refs.clear()

    def render_page(self):
        self.init_responsive_scaling()
        self.clear_content()
        self.draw_background()
        if self.current_page == "launch":
            self.render_launch()
            return
        self.draw_header(self.current_title)
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        if self.current_page == "main": self.render_main_menu()
        elif self.current_page == "data_prep": self.render_data_prep_menu()
        elif self.current_page == "debitoren_tools": self.render_debitoren_tools_menu()
        elif self.current_page == "nike_tools": self.render_nike_tools_menu()
        elif self.current_page == "afi_uploads": self.render_afi_uploads_menu()
        elif self.current_page == "closing_calendar": self.render_closing_calendar_menu()
        elif self.current_page == "knowledge_base": self.render_knowledge_base()
        elif self.current_page == "compliance_audit": self.render_compliance_audit_menu()
        elif self.current_page == "in_dev": self.render_in_dev_menu()
        elif self.current_page == "settings": self.render_settings_menu()
        elif self.current_page == "tile_colors": self.render_tile_colors_menu()
        elif self.current_page == "users": self.render_users_menu()
        elif self.current_page == "permissions": self.render_permissions_menu()
        elif self.current_page == "information": self.render_information_menu()
        elif self.current_page == "versions": self.render_versions_menu()
        elif self.current_page.startswith("tool:"): self.render_external_tool(self.current_page.replace("tool:", "", 1))
        else: self.render_menu_text("Menü in Arbeit.")
        if self.focusable_tiles:
            self.focus_index = 0
            self.focusable_tiles[0].focus_set()

    def on_resize(self, *_):
        self.root.after(30, self.render_page)

    def draw_background(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        self.canvas.create_rectangle(0, 0, w, h, fill=BG, outline="")
        # Dezentes dynamisches Hintergrunddesign: weiche Linien und Flächen in FiBu-Mate-/INTERSPORT-Farben
        for i in range(10):
            y = 145 + i * max(52, h * 0.060)
            color = blend(BG, BLUE if i % 2 == 0 else RED, 0.040 + (i % 4) * 0.008)
            self.canvas.create_line(-80, y, w + 90, y + h * 0.075, fill=color, width=1)
        self.canvas.create_oval(w * 0.68, h * 0.18, w * 1.05, h * 0.78, outline=blend(BG, BLUE, 0.11), width=2)
        self.canvas.create_oval(w * 0.73, h * 0.24, w * 1.01, h * 0.70, outline=blend(BG, WHITE, 0.46), width=1)
        self.canvas.create_oval(-w * 0.14, h * 0.42, w * 0.25, h * 1.02, outline=blend(BG, RED, 0.075), width=2)
        self.canvas.create_polygon(w * 0.02, h * 0.90, w * 0.32, h * 0.72, w * 0.60, h * 1.04, fill=blend(BG, WHITE, 0.20), outline="")
        self.canvas.create_polygon(w * 0.62, h * 0.20, w * 0.96, h * 0.36, w * 1.06, h * 0.12, fill=blend(BG, BLUE, 0.040), outline="")
        self.canvas.create_polygon(w * 0.08, h * 0.18, w * 0.22, h * 0.26, w * 0.04, h * 0.34, fill=blend(BG, WHITE, 0.15), outline="")
        self.canvas.create_rectangle(0, 0, w, 109, fill=HEADER, outline="")
        self.canvas.create_rectangle(0, 109, w, 112, fill=LINE, outline="")

    def draw_header(self, title):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        header_h = ui_s(92)
        self.canvas.create_rectangle(0, 0, w, header_h, fill=HEADER, outline="")
        self.canvas.create_line(0, header_h, w, header_h, fill=LINE, width=1)
        self.canvas.create_text(w / 2, ui_s(70), text=title, font=FONT_MENU, fill=TEXT2, anchor="center")

    def draw_gradient_line(self, x1, x2, y):
        width = max(1, int(x2 - x1)); steps = max(20, min(180, width // 5))
        for i in range(steps):
            t = i / max(1, steps - 1)
            color = blend(HEADER, RED, t / 0.14) if t < 0.14 else blend(BLUE, HEADER, (t - 0.86) / 0.14) if t > 0.86 else blend(RED, BLUE, (t - 0.14) / 0.72)
            self.canvas.create_line(x1 + width * i / steps, y, x1 + width * (i + 1) / steps, y, fill=color, width=2)

    def draw_path_bar(self):
        w = self.canvas.winfo_width(); x1 = ui_s(12); x2 = min(w * 0.42, ui_s(760)); y_mid = ui_s(38)
        self.draw_gradient_line(x1, x2, y_mid - ui_s(10)); self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
        x = x1 + ui_s(96)
        for idx, (page, title) in enumerate(self.breadcrumb):
            current = idx == len(self.breadcrumb) - 1
            tid = self.canvas.create_text(x, y_mid, text=title, font=body_font(9, "bold") if current else body_font(9), fill=TEXT if current else BLUE, anchor="w")
            bbox = self.canvas.bbox(tid); tw = bbox[2] - bbox[0] if bbox else ui_s(70)
            if not current:
                self.canvas.tag_bind(tid, "<Button-1>", lambda e, i=idx: self.jump_to_breadcrumb(i))
                self.canvas.tag_bind(tid, "<Enter>", lambda e: self.canvas.config(cursor="hand2"))
                self.canvas.tag_bind(tid, "<Leave>", lambda e: self.canvas.config(cursor="arrow"))
            x += tw + ui_s(14)
            if idx < len(self.breadcrumb) - 1:
                self.canvas.create_polygon(x, y_mid - ui_s(6), x, y_mid + ui_s(6), x + ui_s(7), y_mid, fill=RED, outline=RED); x += ui_s(20)

    def draw_favorites_bar(self):
        w = self.canvas.winfo_width(); x1, x2 = max(w * 0.58, w - ui_s(720)), w - ui_s(360)
        if x2 <= x1 + ui_s(120): return
        y_mid = ui_s(38)
        self.draw_gradient_line(x1, x2, y_mid - ui_s(10)); self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
        self.canvas.create_text(x1 + ui_s(16), y_mid, text="★", font=("Segoe UI", max(10, ui_s(13)), "bold"), fill=GOLD, anchor="w")
        self.canvas.create_text(x1 + ui_s(40), y_mid, text="Favoriten", font=body_font(9, "bold"), fill=TEXT, anchor="w")
        x = x1 + ui_s(118); chip_color = self.current_tile_color() or BLUE; max_x = x2 - ui_s(8)
        for fav in sorted(f for f in self.favorites if f in TOOL_REGISTRY and f not in HIDDEN_TOOL_IDS):
            if x + ui_s(96) > max_x: break
            label = TOOL_REGISTRY.get(fav, {}).get("favorite_label", fav)
            chip = tk.Label(self.root, text=label, bg=chip_color, fg="white", font=body_font(8), padx=ui_s(6), pady=ui_s(1), cursor="hand2"); chip._zoom_exclude = True
            chip.bind("<Button-1>", lambda e, fid=fav: self.execute_favorite(fid))
            self.widget_items.append(chip); self.canvas.create_window(x, y_mid + ui_s(8), window=chip, anchor="w"); x += ui_s(108)

    def draw_controls(self):
        self.draw_logout_control() if self.current_page == "main" else self.draw_back_control()
        self.draw_suggestion_control()
        self.draw_help_control()
        self.draw_close_control()

    def draw_back_control(self):
        frame = tk.Frame(self.root, bg=HEADER, cursor="hand2")
        arrow = tk.Label(frame, text="←", fg=RED, bg=HEADER, font=body_font(12, "bold"), cursor="hand2")
        txt = tk.Label(frame, text="zurück", fg=BLUE, bg=HEADER, font=body_font(9, underline=True), cursor="hand2")
        arrow.pack(side="left"); txt.pack(side="left", padx=ui_s(4))
        arrow.bind("<Button-1>", lambda e: self.go_back()); txt.bind("<Button-1>", lambda e: self.go_back())
        self.widget_items.append(frame); self.canvas.create_window(ui_s(14), ui_s(8), window=frame, anchor="nw")

    def draw_logout_control(self):
        btn = tk.Button(self.root, text="Abmelden", command=self.logout, bg=self.current_tile_color() or BLUE, fg="white", font=body_font(10, "bold", scale=0.75), bd=0, cursor="hand2", padx=ui_s(9), pady=ui_s(4))
        self.widget_items.append(btn); self.canvas.create_window(ui_s(14), ui_s(8), window=btn, anchor="nw")

    def draw_mini_bulb_icon(self, canvas, x, y):
        if self.draw_canvas_icon(canvas, "idea", x, y, 18, 18):
            return
        canvas.create_oval(x - 6, y - 8, x + 6, y + 4, fill="#FFE45C", outline="#FACC15", width=1)
        canvas.create_rectangle(x - 4, y + 4, x + 4, y + 9, outline=BLUE, width=1)

    def draw_mini_help_icon(self, canvas, x, y):
        if self.draw_canvas_icon(canvas, "help", x, y, 19, 19):
            return
        canvas.create_rectangle(x - 17, y - 8, x + 17, y + 8, outline=BLUE, width=1)
        canvas.create_text(x, y, text="HELP", fill=BLUE, font=("Segoe UI", 7, "bold"))

    def is_zoom_control_visible(self):
        """v0.434 technische Korrektur: Zoomleiste nur in Modulen und Einstellungen-Untermenüs anzeigen."""
        try:
            page = str(getattr(self, "current_page", "") or "")
            if page.startswith("tool:"):
                return True
            return page in {"tile_colors", "users", "permissions", "information", "versions"}
        except Exception:
            return False

    def draw_zoom_control(self):
        """v0.434 Paket 1E: sichtbare Zoom-Leiste als Alternative zu Strg+Mausrad."""
        if not self.is_zoom_control_visible():
            return
        try:
            w = ui_s(190)
            h = MINI_WIDGET_H
            btn = tk.Canvas(self.root, width=w, height=h, bg=HEADER, highlightthickness=0, bd=0, cursor="arrow"); btn._zoom_exclude = True
            btn.create_rectangle(1, 1, w - 1, h - 1, fill=HEADER, outline=BLUE, width=1)
            minus = btn.create_text(ui_s(18), h / 2, text="−", fill=BLUE, font=("Segoe UI", 13, "bold"), anchor="center")
            label = btn.create_text(w / 2, h / 2, text=self.zoom_percent_text(), fill=TEXT, font=("Segoe UI", 9, "bold"), anchor="center")
            plus = btn.create_text(w - ui_s(18), h / 2, text="+", fill=BLUE, font=("Segoe UI", 13, "bold"), anchor="center")
            btn.create_line(ui_s(36), 4, ui_s(36), h - 4, fill=LINE)
            btn.create_line(w - ui_s(36), 4, w - ui_s(36), h - 4, fill=LINE)
            btn._zoom_label_item = label
            btn._zoom_minus_item = minus
            btn._zoom_plus_item = plus
            btn.tag_bind(minus, "<Button-1>", lambda _e: self.zoom_bar_step(-1))
            btn.tag_bind(plus, "<Button-1>", lambda _e: self.zoom_bar_step(1))
            btn.tag_bind(label, "<Button-1>", lambda _e: self.zoom_bar_reset())
            btn.bind("<Button-1>", lambda e: self.zoom_bar_click(e, btn, w))
            btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Zoom: − / Prozent zurücksetzen / +"))
            btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
            self._zoom_control = btn
            self.widget_items.append(btn)
            self.canvas.create_window(ui_s(168), 7, window=btn, anchor="nw")
        except Exception:
            pass

    def zoom_percent_text(self):
        try:
            return f"{int(round(float(self.get_scope_zoom()) * 100))}%"
        except Exception:
            return "100%"

    def update_zoom_control_label(self):
        try:
            btn = getattr(self, "_zoom_control", None)
            if btn and btn.winfo_exists():
                btn.itemconfigure(getattr(btn, "_zoom_label_item", None), text=self.zoom_percent_text())
        except Exception:
            pass

    def zoom_bar_step(self, direction):
        try:
            self.adjust_global_text_zoom(direction)
            self.update_zoom_control_label()
        except Exception:
            pass
        return "break"

    def zoom_bar_reset(self):
        try:
            self.set_scope_zoom(1.0)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass
        return "break"

    def zoom_bar_click(self, event, canvas, width):
        try:
            x = getattr(event, "x", 0)
            if x <= ui_s(38):
                return self.zoom_bar_step(-1)
            if x >= width - ui_s(38):
                return self.zoom_bar_step(1)
            return self.zoom_bar_reset()
        except Exception:
            return "break"

    def draw_suggestion_control(self):
        if self.current_page == "launch":
            return
        # v0.434 Paket 1D: Mini-Widget bleibt zoom-excluded; Text wird bewusst mit fixer Mini-Schrift gezeichnet.
        # Dadurch kann „Änderung vorschlagen“ nicht mehr durch UI_TEXT_SCALE aus dem Button laufen.
        suggestion_w = max(MINI_WIDGET_W, ui_s(202))
        btn = tk.Canvas(self.root, width=suggestion_w, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor="hand2"); btn._zoom_exclude = True
        btn.create_rectangle(1, 1, suggestion_w - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
        icon_x = ui_s(16)
        text_x = ui_s(36)
        self.draw_mini_bulb_icon(btn, icon_x, MINI_WIDGET_H / 2)
        btn.create_text(text_x, MINI_WIDGET_H / 2, text="Änderung vorschlagen", fill=TEXT, font=("Segoe UI", max(8, ui_s(9)), "bold"), anchor="w")
        btn.bind("<Button-1>", lambda _e: self.open_suggestion_mail())
        btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Änderung vorschlagen"))
        btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
        self.widget_items.append(btn)
        # rechter Rand: links vom Hilfe-Mini-Widget; anchor=ne erwartet die rechte Kante, daher nicht nochmals um die eigene Breite verschieben.
        x = self.canvas.winfo_width() - ui_s(176) - MINI_WIDGET_GAP
        self.canvas.create_window(x, ui_s(8), window=btn, anchor="ne")

    def show_small_tooltip(self, widget, text):
        self.hide_small_tooltip()
        try:
            self._small_tooltip = tk.Toplevel(widget)
            self._small_tooltip.wm_overrideredirect(True)
            self._small_tooltip.geometry(f"+{widget.winfo_rootx() - 90}+{widget.winfo_rooty() + 34}")
            tk.Label(self._small_tooltip, text=text, bg="#111827", fg="white", font=("Segoe UI", 9), padx=7, pady=4).pack()
        except Exception:
            self._small_tooltip = None

    def hide_small_tooltip(self):
        tip = getattr(self, "_small_tooltip", None)
        if tip:
            try:
                tip.destroy()
            except Exception:
                pass
        self._small_tooltip = None

    def open_suggestion_mail(self):
        messagebox.showinfo(
            "Änderung vorschlagen",
            "Vielen Dank für deinen Input! Bitte beschreibe in der folgenden Vorlage deinen Vorschlag so ausführlich wie möglich - Füge gerne Screenshots mit an",
        )
        user_name = self.current_user_display or self.current_user_key or ""
        subject = "Änderungsvorschlag Fibu Mate"
        body = (
            f"Änderungsvorschlag von {user_name},\n\n"
            "Folgenden Änderungsvorschlag würde ich gerne mitteilen / Folgende Anpassung wünsche ich mir:\n\n"
            "[Text des Vorschlagenden]"
        )
        mailto = (
            "mailto:matthias.wagner@intersport.de"
            "?cc=" + quote("matze.wagner1@yahoo.de")
            + "&subject=" + quote(subject)
            + "&body=" + quote(body)
        )
        try:
            webbrowser.open(mailto)
        except Exception as exc:
            messagebox.showerror("Änderung vorschlagen", f"Outlook konnte nicht geöffnet werden:\n\n{exc}")

    def draw_help_control(self):
        if self.current_page == "launch":
            return
        btn = tk.Canvas(self.root, width=MINI_WIDGET_W, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor="hand2"); btn._zoom_exclude = True
        btn.create_rectangle(1, 1, MINI_WIDGET_W - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
        self.draw_mini_help_icon(btn, 25, MINI_WIDGET_H / 2)
        btn.create_text(MINI_WIDGET_W / 2 + ui_s(6), MINI_WIDGET_H / 2, text="Hilfe", fill=TEXT, font=("Segoe UI", max(8, ui_s(9)), "bold"), anchor="center")
        def open_help(_event=None):
            self.show_help_popup()
        btn.bind("<Button-1>", open_help)
        btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Hilfe"))
        btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
        self.widget_items.append(btn)
        self.canvas.create_window(self.canvas.winfo_width() - ui_s(22), ui_s(8), window=btn, anchor="ne")

    def draw_close_control(self):
        """v0.433 Korrektur Paket 1c/1e: Separater X-Button in den Mini-Widgets deaktiviert.
        Das Schließen erfolgt ausschließlich über das native Fenster-X bzw. die bestehende App-Logik.
        """
        return

    def show_help_popup(self):
        popup = tk.Toplevel(self.root)
        popup.title("FiBu Mate - Hilfe")
        popup.configure(bg=BG)
        popup.transient(self.root)
        popup.resizable(True, True)
        popup.minsize(780, 520)
        popup_w, popup_h = 1000, 700
        try:
            self.root.update_idletasks()
            screen_w = self.root.winfo_screenwidth()
            screen_h = self.root.winfo_screenheight()
            x = max(0, int((screen_w - popup_w) / 2))
            y = max(0, int((screen_h - popup_h) / 2))
            popup.geometry(f"{popup_w}x{popup_h}+{x}+{y}")
        except Exception:
            popup.geometry(f"{popup_w}x{popup_h}")
        self._make_modal(popup)
        header = tk.Frame(popup, bg=HEADER, height=64); header.pack(side="top", fill="x"); header.pack_propagate(False)
        tk.Label(header, text="Hilfe", bg=HEADER, fg=TEXT, font=("Segoe UI", 18, "bold")).pack(side="left", padx=18)
        content = tk.Frame(popup, bg=BG); content.pack(side="top", fill="both", expand=True, padx=16, pady=16)
        if self.help_image and PIL_AVAILABLE:
            try:
                popup.update_idletasks(); img = resize_keep_ratio(self.help_image, max(600, popup.winfo_width() - 70), max(400, popup.winfo_height() - 150)); ph = ImageTk.PhotoImage(img)
                label = tk.Label(content, image=ph, bg=BG); label.image = ph; label.pack(expand=True)
            except Exception as error:
                tk.Label(content, text=f"Das Hilfe-Bild konnte nicht geladen werden:\n\n{error}", bg=BG, fg=TEXT, font=body_font(11), justify="left").pack(anchor="nw")
        else:
            tk.Label(content, text=("Das Hilfe-Menü ist vorbereitet, aber das Hilfe-Bild konnte nicht geladen werden.\n\n" f"Erwarteter Pfad:\n{HELP_IMAGE_PATH}"), bg=BG, fg=TEXT, font=body_font(11), justify="left").pack(anchor="nw")
        footer = tk.Frame(popup, bg=BG); footer.pack(side="bottom", fill="x", padx=16, pady=(0, 16))
        tk.Button(footer, text="Schließen", command=popup.destroy, bg=BLUE, fg="white", font=body_font(10, "bold"), bd=0, padx=ui_s(24), pady=ui_s(10), cursor="hand2").pack(side="right")

    def draw_bottom_logo(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if self.intersport_logo and PIL_AVAILABLE:
            ph = ImageTk.PhotoImage(resize_keep_ratio(self.intersport_logo, 420, 55)); self.image_refs.append(ph); self.canvas.create_image(w / 2, h - 24, image=ph)
        else:
            self.canvas.create_text(w / 2, h - 24, text="INTERSPORT", font=("Segoe UI", 22, "bold"), fill=BLUE)

    def draw_intersport_logo_above_footer(self, show_mini_logo=True): return self.draw_bottom_logo()

    def toggle_favorite(self, tool_id):
        if tool_id not in TOOL_REGISTRY: return
        self.favorites.remove(tool_id) if tool_id in self.favorites else self.favorites.add(tool_id)
        if self.current_user_key:
            self.user_data["users"][self.current_user_key]["favorites"] = sorted(self.favorites); self.save_user_data()
        self.render_page()

    def execute_favorite(self, tool_id):
        if tool_id in TOOL_REGISTRY: self.open_tool(tool_id)

    def render_launch(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if self.banner_big and PIL_AVAILABLE:
            ph = ImageTk.PhotoImage(resize_keep_ratio(self.banner_big, 900, 230)); self.image_refs.append(ph); self.canvas.create_image(w / 2, h * 0.235, image=ph)
        else:
            self.canvas.create_text(w / 2, h * 0.235, text="FiBu Mate", font=("Segoe UI", 46, "bold"), fill=TEXT)
        panel = tk.Frame(self.root, bg=BG); username_var = tk.StringVar(value="")
        tk.Label(panel, text="Benutzername", bg=BG, fg=TEXT, font=("Segoe UI", 11, "bold")).pack(anchor="w", pady=(0, 6))
        row = tk.Frame(panel, bg=BG); row.pack(fill="x", pady=(0, 18))
        entry = tk.Entry(row, textvariable=username_var, font=body_font(13), width=37, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1); entry.pack(side="left", ipady=ui_s(7))
        btn = tk.Button(row, text="Anmelden", command=lambda: self.login_user(username_var.get()), bg=BLUE, fg="white", font=body_font(11, "bold", scale=0.75), bd=0, cursor="hand2"); btn.pack(side="left", padx=(12, 0), ipady=6)
        self.widget_items.append(panel); self.canvas.create_window(x_pct(w, 50), y_pct(h, 39), window=panel, anchor="center")
        entry.focus_set(); entry.bind("<Return>", lambda e: self.login_user_from_entry(username_var.get()))
        self.draw_bottom_logo(); self.draw_close_control()



    # ------------------------------------------------------------------
    # Wissenszentrale - Startseite als Overlay + Arbeitsansicht A
    # ------------------------------------------------------------------
    def knowledge_pref_path(self):
        try:
            return os.path.join(APPDATA_DIR, "knowledge_start_pref.json")
        except Exception:
            return "knowledge_start_pref.json"

    def kb_entries_path(self):
        try:
            os.makedirs(APPDATA_DIR, exist_ok=True)
            return os.path.join(APPDATA_DIR, "knowledge_entries.json")
        except Exception:
            return "knowledge_entries.json"

    def load_knowledge_start_preference(self):
        try:
            path = self.knowledge_pref_path()
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return bool(data.get("show_start", True))
        except Exception:
            pass
        return True

    def save_knowledge_start_preference(self, value):
        self.knowledge_show_start = bool(value)
        try:
            path = self.knowledge_pref_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"show_start": self.knowledge_show_start}, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def kb_load_entries(self):
        try:
            path = self.kb_entries_path()
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                entries = data.get("entries", []) if isinstance(data, dict) else []
                return entries if isinstance(entries, list) else []
        except Exception:
            pass
        return []

    def kb_save_entries(self, entries):
        try:
            path = self.kb_entries_path()
            dir_name = os.path.dirname(path)
            if dir_name:
                os.makedirs(dir_name, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"entries": entries}, f, ensure_ascii=False, indent=2)
            return True
        except Exception as exc:
            try:
                messagebox.showerror("Wissenszentrale", f"Einträge konnten nicht gespeichert werden:\n{exc}")
            except Exception:
                pass
            return False

    def kb_default_categories(self):
        return ["To-Do"]

    def kb_get_categories(self):
        cats = set(self.kb_default_categories())
        for entry in self.kb_load_entries():
            for cat in entry.get("categories", []) or []:
                cat = str(cat).strip()
                if cat:
                    cats.add(cat)
        return sorted(cats, key=lambda x: x.lower())

    def kb_now(self):
        try:
            return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            return ""

    def kb_display_date(self, value):
        try:
            dt = datetime.strptime(str(value), "%Y-%m-%d %H:%M:%S")
            return dt.strftime("%d.%m.%Y %H:%M")
        except Exception:
            return str(value or "")

    def kb_make_entry_id(self):
        try:
            return datetime.now().strftime("KB%Y%m%d%H%M%S%f")
        except Exception:
            return "KB_ENTRY"

    def open_knowledge_base(self):
        self.knowledge_view = "all"
        self.knowledge_start_overlay = bool(getattr(self, "knowledge_show_start", True))
        self.show_page("knowledge_base", "Wissenszentrale", True)

    def draw_kb_button(self, x, y, text, command=None, accent=False, width=None):
        btn = tk.Button(self.root, text=text, command=command, font=body_font(10, weight="bold" if accent else None),
                        bg=("#CFEAD6" if accent else WHITE), fg=TEXT, activebackground="#DDEAF7",
                        relief="solid", bd=1, cursor="hand2", padx=8, pady=3)
        btn_w = width or max(110, 18 + len(text) * 8)
        self.canvas.create_window(ui_s(x), ui_s(y), window=btn, anchor="nw", width=ui_s(btn_w), height=ui_s(34))
        self.widget_items.append(btn)
        return x + btn_w + 8

    def draw_kb_badge(self, x, y, text, color):
        w = max(54, 20 + len(str(text)) * 8)
        self.canvas.create_rectangle(ui_s(x), ui_s(y), ui_s(x+w), ui_s(y+24), fill=color, outline=color, width=1)
        self.canvas.create_text(ui_s(x+10), ui_s(y+12), text=text, font=body_font(9, weight="bold"), fill=WHITE, anchor="w")
        return x + w + 8

    def kb_has_unsaved_changes(self):
        return bool(getattr(self, "knowledge_unsaved", False))

    def kb_mark_unsaved(self):
        self.knowledge_unsaved = True

    def kb_save_current_work(self):
        # Speichern wird in der Wissenszentrale durch kb_save_entry_from_form übernommen.
        self.knowledge_unsaved = False
        return True

    def kb_confirm_unsaved_before_switch(self):
        if not self.kb_has_unsaved_changes():
            return True
        try:
            result = messagebox.askyesnocancel(
                "Ungespeicherte Änderungen",
                "In der aktuellen Arbeitsansicht liegen ungespeicherte Änderungen vor.\n\n"
                "Ja = speichern und wechseln\nNein = verwerfen und wechseln\nAbbrechen = in der aktuellen Ansicht bleiben"
            )
        except Exception:
            result = False
        if result is True:
            return self.kb_save_current_work()
        if result is False:
            self.knowledge_unsaved = False
            return True
        return False

    def kb_switch_view_from_start(self, view):
        if not self.kb_confirm_unsaved_before_switch():
            return
        self.knowledge_view = view
        if view == "new":
            self.kb_prepare_new_entry()
        self.knowledge_start_overlay = False
        self.render_page()

    def kb_show_start_overlay(self):
        self.knowledge_start_overlay = True
        self.render_page()

    def kb_close_start_overlay(self):
        self.knowledge_start_overlay = False
        self.render_page()

    def render_knowledge_base(self):
        underlay_start = len(self.widget_items)
        self.render_knowledge_work_area()
        if bool(getattr(self, "knowledge_start_overlay", False)):
            for widget in list(self.widget_items[underlay_start:]):
                try:
                    widget.destroy()
                except Exception:
                    pass
            del self.widget_items[underlay_start:]
            self.render_knowledge_start_overlay()

    def kb_ensure_state_vars(self):
        if not hasattr(self, "kb_search_var"):
            self.kb_search_var = tk.StringVar(value="")
        if not hasattr(self, "kb_filter_vars") or len(getattr(self, "kb_filter_vars", [])) != 4:
            self.kb_filter_vars = [tk.StringVar(value="") for _ in range(4)]
        if not hasattr(self, "kb_todo_rhythm_var"):
            self.kb_todo_rhythm_var = tk.StringVar(value="")
        if not hasattr(self, "kb_selected_entry_id"):
            self.kb_selected_entry_id = None
        if not hasattr(self, "kb_edit_entry_id"):
            self.kb_edit_entry_id = None

    def kb_prepare_new_entry(self, entry=None):
        self.kb_ensure_state_vars()
        self.kb_edit_entry_id = entry.get("id") if entry else None
        self.kb_title_var = tk.StringVar(value=(entry or {}).get("title", ""))
        cats = list((entry or {}).get("categories", []) or [])[:4]
        self.kb_entry_category_vars = [tk.StringVar(value=(cats[i] if i < len(cats) else "")) for i in range(4)]
        self.kb_user_var = tk.StringVar(value=(entry or {}).get("user", self.current_user_display or self.current_user_key or ""))
        self.kb_status_var = tk.StringVar(value=(entry or {}).get("status", "Aktiv"))
        self.kb_rhythm_var = tk.StringVar(value=(entry or {}).get("rhythm", ""))
        self.kb_text_initial = (entry or {}).get("text", "")
        self.knowledge_unsaved = bool(entry)

    def kb_filtered_entries(self):
        self.kb_ensure_state_vars()
        entries = self.kb_load_entries()
        search = (self.kb_search_var.get() or "").strip().lower()
        filters = [(v.get() or "").strip().lower() for v in self.kb_filter_vars]
        filters = [f for f in filters if f]
        rhythm = (self.kb_todo_rhythm_var.get() or "").strip().lower()
        result = []
        for entry in entries:
            cats = [str(c).strip().lower() for c in (entry.get("categories", []) or [])]
            if getattr(self, "knowledge_view", "all") == "todos" and "to-do" not in cats:
                continue
            if getattr(self, "knowledge_view", "all") == "todos" and rhythm:
                if str(entry.get("rhythm", "")).strip().lower() != rhythm:
                    continue
            if filters and not all(f in cats for f in filters):
                continue
            if search:
                hay = " ".join([
                    str(entry.get("title", "")), str(entry.get("text", "")),
                    str(entry.get("user", "")), " ".join(entry.get("categories", []) or [])
                ]).lower()
                if search not in hay:
                    continue
            result.append(entry)
        return sorted(result, key=lambda e: e.get("updated_at") or e.get("created_at") or "", reverse=True)

    def kb_apply_filters(self, event=None):
        self.kb_selected_entry_id = None
        if getattr(self, "knowledge_view", "all") not in ("new", "categories", "outdated"):
            self.knowledge_view = "all" if getattr(self, "knowledge_view", "all") != "todos" else "todos"
        self.render_page()

    def kb_on_search_return(self, event=None):
        self.kb_apply_filters()

    def kb_select_entry(self, entry_id):
        self.kb_selected_entry_id = entry_id
        self.knowledge_view = "detail"
        self.knowledge_unsaved = False
        self.render_page()

    def kb_get_entry(self, entry_id):
        for entry in self.kb_load_entries():
            if entry.get("id") == entry_id:
                return entry
        return None

    def kb_edit_selected_entry(self):
        entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
        if not entry:
            return
        self.kb_prepare_new_entry(entry)
        self.knowledge_view = "new"
        self.render_page()

    def kb_save_entry_from_form(self):
        self.kb_ensure_state_vars()
        title = (getattr(self, "kb_title_var", tk.StringVar(value="")).get() or "").strip()
        if not title:
            try:
                messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
            except Exception:
                pass
            return
        categories = []
        for var in getattr(self, "kb_entry_category_vars", []):
            value = (var.get() or "").strip()
            if value and value not in categories:
                categories.append(value)
        categories = categories[:4]
        user = (getattr(self, "kb_user_var", tk.StringVar(value="")).get() or "").strip()
        status = (getattr(self, "kb_status_var", tk.StringVar(value="Aktiv")).get() or "Aktiv").strip()
        rhythm = (getattr(self, "kb_rhythm_var", tk.StringVar(value="")).get() or "").strip()
        text_value = ""
        try:
            text_value = self.kb_text_widget.get("1.0", "end-1c")
        except Exception:
            text_value = getattr(self, "kb_text_initial", "")
        entries = self.kb_load_entries()
        now = self.kb_now()
        edit_id = getattr(self, "kb_edit_entry_id", None)
        if edit_id:
            for entry in entries:
                if entry.get("id") == edit_id:
                    entry.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "updated_at": now})
                    break
            selected_id = edit_id
        else:
            selected_id = self.kb_make_entry_id()
            entries.append({
                "id": selected_id, "title": title, "categories": categories, "user": user,
                "status": status, "rhythm": rhythm, "text": text_value,
                "created_at": now, "updated_at": now, "comments": [], "attachments": []
            })
        if self.kb_save_entries(entries):
            self.kb_selected_entry_id = selected_id
            self.kb_edit_entry_id = None
            self.knowledge_unsaved = False
            self.knowledge_view = "detail"
            self.render_page()

    def kb_export_selected_to_word(self):
        # Platzhalter: Der echte Word-Export wird später mit python-docx an die finale Datenstruktur angebunden.
        try:
            messagebox.showinfo("Wissenszentrale", "Word-Export wird in einem Folgeblock an die finale Datenstruktur angebunden.")
        except Exception:
            pass

    def render_knowledge_work_area(self):
        self.kb_ensure_state_vars()
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        # Alle Elemente sind bewusst in vertikal getrennten Zonen aufgebaut:
        # Navigation 145, Suche 192, Filter 245, Hauptfenster ab 365.
        nav_y = 145
        x = 28
        x = self.draw_kb_button(x, nav_y, "Start", self.kb_show_start_overlay, False, width=92)
        x = self.draw_kb_button(x, nav_y, "Neuer Eintrag", lambda: self.kb_switch_view_from_start("new"), True, width=128)
        x = self.draw_kb_button(x, nav_y, "To-Dos", lambda: self.kb_switch_view_from_start("todos"), False, width=92)
        x = self.draw_kb_button(x, nav_y, "Veraltete Einträge", lambda: self.kb_switch_view_from_start("outdated"), False, width=148)
        x = self.draw_kb_button(x, nav_y, "Kategorien verwalten", lambda: self.kb_switch_view_from_start("categories"), False, width=168)

        search_frame = tk.Frame(self.root, bg=BG)
        self.widget_items.append(search_frame)
        tk.Label(search_frame, text="Suche", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
        entry = tk.Entry(search_frame, textvariable=self.kb_search_var, font=body_font(10), bg=WHITE, fg=TEXT, relief="solid", bd=1)
        entry.grid(row=1, column=0, sticky="ew", ipady=4)
        entry.bind("<Return>", self.kb_on_search_return)
        btn = tk.Button(search_frame, text="Suchen", command=self.kb_apply_filters, bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(9), cursor="hand2")
        btn.grid(row=1, column=1, padx=(8, 0), sticky="ns")
        search_frame.grid_columnconfigure(0, weight=1)
        self.canvas.create_window(ui_s(28), ui_s(192), window=search_frame, anchor="nw", width=ui_s(min(760, max(520, w-80))), height=ui_s(56))

        filter_frame = tk.Frame(self.root, bg=BG)
        self.widget_items.append(filter_frame)
        if self.knowledge_view == "todos":
            tk.Label(filter_frame, text="To-Do-Filter / Rhythmus", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w", columnspan=2)
            rhythms = ["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"]
            cb = ttk.Combobox(filter_frame, textvariable=self.kb_todo_rhythm_var, values=rhythms, state="readonly", font=body_font(9), width=18)
            cb.grid(row=1, column=0, sticky="w")
            cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
            tk.Label(filter_frame, text="Leer = alle Rhythmen", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=1, column=1, padx=(12,0), sticky="w")
        else:
            tk.Label(filter_frame, text="Kategorie-Filter", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w", columnspan=5)
            values = [""] + self.kb_get_categories()
            for idx, var in enumerate(self.kb_filter_vars):
                block = tk.Frame(filter_frame, bg=BG)
                block.grid(row=1, column=idx, padx=(0 if idx == 0 else 10, 0), sticky="w")
                tk.Label(block, text=f"Kategorie {idx+1}", bg=BG, fg=TEXT2, font=body_font(8)).pack(anchor="w")
                cb = ttk.Combobox(block, textvariable=var, values=values, state="readonly", font=body_font(9), width=18)
                cb.pack(anchor="w")
                cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
            tk.Label(filter_frame, text="Leere Filter = keine Filterung", bg=BG, fg=TEXT2, font=body_font(8)).grid(row=1, column=4, padx=(12,0), sticky="s")
        self.canvas.create_window(ui_s(28), ui_s(252), window=filter_frame, anchor="nw", width=ui_s(min(860, max(620, w-70))), height=ui_s(82))

        # Hauptfenster bewusst tiefer gesetzt, damit Filter, Treffer und Detail niemals kollidieren.
        left_x = 20
        left_y = 365
        left_w = max(340, min(460, int(w * 0.31)))
        pane_h = max(390, h - left_y - 50)
        right_x = left_x + left_w + 24
        right_y = left_y
        right_w = max(760, w - right_x - 24)
        self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
        if self.knowledge_view == "new":
            self.render_kb_new_entry_area(right_x, right_y, right_w, pane_h)
        elif self.knowledge_view == "todos":
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="To-Dos")
        elif self.knowledge_view == "outdated":
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Veraltete Einträge", status_filter="Veraltet")
        elif self.knowledge_view == "categories":
            self.render_kb_categories_area(right_x, right_y, right_w, pane_h)
        elif self.knowledge_view == "detail":
            self.render_kb_detail_area(right_x, right_y, right_w, pane_h)
        else:
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Gesamtliste aller Einträge")

    def render_kb_hits_pane(self, x, y, w, h):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text="Treffer", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
        entries = self.kb_filtered_entries()
        listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="flat", activestyle="none", exportselection=False)
        listbox.pack(fill="both", expand=True, padx=16, pady=(0, 14))
        id_map = []
        if entries:
            for entry in entries:
                cats = ", ".join(entry.get("categories", []) or [])
                line = f"{self.kb_display_date(entry.get('updated_at'))}  |  {entry.get('title','')}"
                if cats:
                    line += f"  [{cats}]"
                listbox.insert("end", line)
                id_map.append(entry.get("id"))
        else:
            listbox.insert("end", "Noch keine Treffer vorhanden.")
            listbox.insert("end", "Filter und Suche wirken auf die Wissensdatenbank.")
        def _select(event=None):
            sel = listbox.curselection()
            if sel and sel[0] < len(id_map):
                self.kb_select_entry(id_map[sel[0]])
        listbox.bind("<Double-Button-1>", _select)
        listbox.bind("<Return>", _select)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_list_area(self, x, y, w, h, title="Gesamtliste aller Einträge", status_filter=None):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        header = tk.Frame(frame, bg=WHITE)
        header.pack(fill="x", padx=18, pady=(16, 8))
        tk.Label(header, text=title, bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(side="left")
        tk.Label(header, text="neueste Einträge oben", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="right")
        columns = ("date", "title", "categories", "user", "status")
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=12)
        tree.heading("date", text="Geändert")
        tree.heading("title", text="Titel")
        tree.heading("categories", text="Kategorien")
        tree.heading("user", text="Benutzer")
        tree.heading("status", text="Status")
        tree.column("date", width=120, stretch=False)
        tree.column("title", width=max(220, int(w*0.33)), stretch=True)
        tree.column("categories", width=max(170, int(w*0.22)), stretch=True)
        tree.column("user", width=130, stretch=False)
        tree.column("status", width=90, stretch=False)
        yscroll = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=yscroll.set)
        tree.pack(side="left", fill="both", expand=True, padx=(18,0), pady=(0,18))
        yscroll.pack(side="right", fill="y", padx=(0,18), pady=(0,18))
        entries = self.kb_filtered_entries()
        if status_filter:
            entries = [e for e in entries if str(e.get("status", "")).lower() == status_filter.lower()]
        id_map = {}
        for entry in entries:
            iid = entry.get("id") or self.kb_make_entry_id()
            id_map[iid] = entry.get("id")
            tree.insert("", "end", iid=iid, values=(self.kb_display_date(entry.get("updated_at")), entry.get("title", ""), ", ".join(entry.get("categories", []) or []), entry.get("user", ""), entry.get("status", "")))
        if not entries:
            tree.insert("", "end", values=("", "Noch keine Einträge vorhanden", "", "", ""))
        def _open(event=None):
            sel = tree.selection()
            if sel and sel[0] in id_map:
                self.kb_select_entry(id_map[sel[0]])
        tree.bind("<Double-Button-1>", _open)
        tree.bind("<Return>", _open)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_detail_area(self, x, y, w, h):
        entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
        if not entry:
            self.render_kb_list_area(x, y, w, h, title="Gesamtliste aller Einträge")
            return
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text=entry.get("title", ""), bg=WHITE, fg=BLUE, font=body_font(16, weight="bold")).pack(anchor="w", padx=18, pady=(16, 6))
        meta = f"Geändert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {entry.get('user','')}    Status: {entry.get('status','')}"
        tk.Label(frame, text=meta, bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor="w", padx=18)
        cats = ", ".join(entry.get("categories", []) or []) or "Keine Kategorien"
        tk.Label(frame, text=f"Kategorien: {cats}", bg=WHITE, fg=TEXT, font=body_font(10, weight="bold")).pack(anchor="w", padx=18, pady=(8, 4))
        txt = tk.Text(frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word")
        txt.pack(fill="both", expand=True, padx=18, pady=(6, 12))
        txt.insert("1.0", entry.get("text", ""))
        txt.configure(state="disabled")
        btnrow = tk.Frame(frame, bg=WHITE)
        btnrow.pack(fill="x", padx=18, pady=(0, 16))
        tk.Button(btnrow, text="Bearbeiten", command=self.kb_edit_selected_entry, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Word-Export", command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Zur Liste", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_new_entry_area(self, x, y, w, h):
        if not hasattr(self, "kb_title_var"):
            self.kb_prepare_new_entry()
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text=("Eintrag bearbeiten" if getattr(self, "kb_edit_entry_id", None) else "Neuer Eintrag"), bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 6))
        form = tk.Frame(frame, bg=WHITE)
        form.pack(fill="x", padx=18, pady=(4, 8))
        tk.Label(form, text="Titel des Eintrags", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
        title_entry = tk.Entry(form, textvariable=self.kb_title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1)
        title_entry.grid(row=1, column=0, columnspan=4, sticky="ew", ipady=5, pady=(0, 8))
        categories = [""] + self.kb_get_categories()
        for idx, var in enumerate(self.kb_entry_category_vars):
            tk.Label(form, text=f"Kategorie {idx+1}", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=idx, sticky="w", padx=(0 if idx==0 else 8,0))
            cb = ttk.Combobox(form, textvariable=var, values=categories, state="normal", font=body_font(9), width=18)
            cb.grid(row=3, column=idx, sticky="ew", padx=(0 if idx==0 else 8,0), pady=(0, 8))
        tk.Label(form, text="Assoziierter Benutzer", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=0, sticky="w")
        tk.Entry(form, textvariable=self.kb_user_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=5, column=0, sticky="ew", pady=(0, 8))
        tk.Label(form, text="Status", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=1, sticky="w", padx=(8,0))
        ttk.Combobox(form, textvariable=self.kb_status_var, values=["Aktiv", "Entwurf", "Veraltet"], state="readonly", font=body_font(9)).grid(row=5, column=1, sticky="ew", padx=(8,0), pady=(0, 8))
        tk.Label(form, text="To-Do-Rhythmus", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=2, sticky="w", padx=(8,0))
        ttk.Combobox(form, textvariable=self.kb_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(9)).grid(row=5, column=2, sticky="ew", padx=(8,0), pady=(0, 8))
        for col in range(4):
            form.grid_columnconfigure(col, weight=1)
        body_frame = tk.Frame(frame, bg=WHITE)
        body_frame.pack(fill="both", expand=True, padx=18, pady=(0, 10))
        tk.Label(body_frame, text="Freitext / Prozessdokumentation / Leitfaden", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        self.kb_text_widget = tk.Text(body_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word")
        self.kb_text_widget.pack(fill="both", expand=True)
        self.kb_text_widget.insert("1.0", getattr(self, "kb_text_initial", ""))
        self.kb_text_widget.bind("<KeyRelease>", lambda e: self.kb_mark_unsaved())
        btnrow = tk.Frame(frame, bg=WHITE)
        btnrow.pack(fill="x", padx=18, pady=(0, 16))
        tk.Button(btnrow, text="Als Entwurf speichern", command=lambda: (self.kb_status_var.set("Entwurf"), self.kb_save_entry_from_form()), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Speichern", command=self.kb_save_entry_from_form, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Abbrechen", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
        title_entry.focus_set()
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_categories_area(self, x, y, w, h):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text="Kategorien verwalten", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
        tk.Label(frame, text="Kategorien entstehen aktuell automatisch aus den Einträgen. Neue Kategorien können direkt beim Erfassen in den Kategorie-Feldern eingetragen werden.", bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(500, int(w*0.85)), justify="left").pack(anchor="w", padx=18, pady=(0, 12))
        listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1)
        listbox.pack(fill="both", expand=True, padx=18, pady=(0, 18))
        for cat in self.kb_get_categories():
            listbox.insert("end", cat)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def kb_overlay_start_drag(self, event):
        self.knowledge_overlay_drag = None

    def kb_overlay_drag(self, event):
        return

    def kb_overlay_end_drag(self, event=None):
        self.knowledge_overlay_drag = None

    def render_knowledge_start_overlay(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        y0 = ui_s(98)
        y1 = h - ui_s(34)
        self.canvas.create_rectangle(0, y0, w, y1, fill=blend(BG, WHITE, 0.86), outline=LINE, width=2)
        self.canvas.create_rectangle(0, y0, w, ui_s(158), fill=blend(HEADER, WHITE, 0.34), outline=LINE, width=1)
        self.canvas.create_text(ui_s(45), ui_s(123), text="Wissenszentrale – Start", fill=BLUE, font=body_font(22, weight="bold"), anchor="w")
        self.canvas.create_text(ui_s(45), ui_s(148), text="Startseite maximiert – Arbeitsbereich vollständig abgedeckt", fill=TEXT2, font=body_font(9), anchor="w")
        self.draw_kb_button(w - ui_s(180), 112, "Schließen", self.kb_close_start_overlay, False, width=130)
        self.canvas.create_text(ui_s(45), ui_s(195), text="Was möchtest du tun?", fill=TEXT, font=body_font(15, weight="bold"), anchor="w")
        var = tk.BooleanVar(value=bool(getattr(self, "knowledge_show_start", True)))
        def _toggle():
            self.save_knowledge_start_preference(var.get())
        chk = tk.Checkbutton(self.root, text="Startseite anzeigen", variable=var, command=_toggle, bg=blend(BG, WHITE, 0.86), fg=TEXT, font=body_font(11), activebackground=blend(BG, WHITE, 0.86))
        self.canvas.create_window(ui_s(45), h - ui_s(88), window=chk, anchor="nw", width=ui_s(260), height=ui_s(34))
        self.widget_items.append(chk)
        cards = [("Wissen suchen", "Gesamtliste und Suche öffnen", "all", "#2563EB"), ("Neuer Eintrag", "Dokumentation oder Aufgabe erfassen", "new", "#059669"), ("To-Dos anzeigen", "wiederkehrende Aufgaben anzeigen", "todos", "#DC2626"), ("Kategorien", "Kategorien anzeigen", "categories", "#7C3AED"), ("Veraltete Einträge", "Prüfung / Archiv", "outdated", "#EC4899")]
        card_w = min(430, max(300, int((max(900, w-ui_s(90)) - 90) / 3)))
        card_h = 150
        start_x, start_y = 45, 235
        gap_x, gap_y = 45, 38
        for idx, (title, desc, view, color) in enumerate(cards):
            col, row = idx % 3, idx // 3
            x = start_x + col * (card_w + gap_x)
            y = start_y + row * (card_h + gap_y)
            self.canvas.create_rectangle(ui_s(x+6), ui_s(y+8), ui_s(x+card_w+6), ui_s(y+card_h+8), fill=SHADOW, outline="")
            self.canvas.create_rectangle(ui_s(x), ui_s(y), ui_s(x+card_w), ui_s(y+card_h), fill=WHITE, outline=LINE)
            self.canvas.create_oval(ui_s(x+24), ui_s(y+26), ui_s(x+74), ui_s(y+76), fill=color, outline=color)
            self.canvas.create_text(ui_s(x+95), ui_s(y+35), text=title, fill=BLUE, font=body_font(15, weight="bold"), anchor="w")
            self.canvas.create_text(ui_s(x+95), ui_s(y+75), text=desc, fill=TEXT2, font=body_font(11), anchor="w")
            self.draw_kb_button(x+95, y+105, "Öffnen", lambda v=view: self.kb_switch_view_from_start(v), view == "new")
        self.canvas.create_text(ui_s(45), h - ui_s(45), text="Die Startseite liegt über allen Arbeitsansicht-Widgets und endet erst oberhalb der Fußleiste.", fill=TEXT2, font=body_font(10), anchor="w")

    def render_main_menu(self):
        # Menüzeile 1: Abschlusskalender
        top_tiles = [
            {"title": "Abschlusskalender", "cmd": lambda: self.show_page("closing_calendar", "Abschlusskalender", True), "fixed": None, "lock": False, "icon": "calendar", "fold": False},
            {"title": "Wissenszentrale", "cmd": lambda: self.open_knowledge_base(), "fixed": None, "lock": False, "icon": "knowledge", "fold": False},
        ]
        # Menüzeile 2: Tools-Menüs
        middle_tiles = [
            {"title": "Tools - Hauptbuch", "cmd": lambda: self.show_page("data_prep", "Tools - Hauptbuch", True), "fixed": None, "lock": False, "icon": "pdf_xls", "fold": False},
            {"title": "Tools - Debitoren", "cmd": lambda: self.show_page("debitoren_tools", "Tools - Debitoren", True), "fixed": None, "lock": False, "icon": "modules", "fold": False},
        ]
        # Mini-Menüzeile: kleinere Kacheln; Eselsohr rechts oben außer bei "In Entwicklung".
        mini_tiles = [
            {"title": "In Entwicklung", "cmd": self.try_open_in_dev, "fixed": GREY_TILE, "lock": True, "icon": "lock", "fold": False},
            {"title": "Informationen", "cmd": lambda: self.show_page("information", "Informationen", True), "fixed": None, "lock": False, "icon": "info", "fold": True},
            {"title": "Einstellungen", "cmd": lambda: self.show_page("settings", "Einstellungen", True), "fixed": None, "lock": False, "icon": "gear", "fold": True},
        ]
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        tw = max(240, min(330, int(w * 0.165)))
        th = max(125, min(170, int(h * 0.155)))
        gap_x = max(38, int(w * 0.04))
        def centered_x_positions(count, tile_w=tw, gap=gap_x):
            if count <= 0:
                return []
            total_width = count * tile_w + (count - 1) * gap
            start_x = (w - total_width) / 2 + tile_w / 2
            return [start_x + i * (tile_w + gap) for i in range(count)]
        def create_tile_group(items, y, id_prefix, tile_w=tw, tile_h=th):
            xs = centered_x_positions(len(items), tile_w)
            for i, item in enumerate(items):
                tile = Tile(self.root, self, f"{id_prefix}_{i}", item["title"], item["cmd"], favorite_enabled=False, fixed_color=item["fixed"], lock_tile=item["lock"], icon_type=item["icon"], corner_fold=item["fold"])
                tile.resize_tile(tile_w, tile_h)
                self.widget_items.append(tile)
                self.focusable_tiles.append(tile)
                self.canvas.create_window(xs[i], y, window=tile, anchor="center")
        top_y = y_pct(h, 64)
        tools_y = y_pct(h, 43)
        create_tile_group(top_tiles, top_y, "main_menuzeile_1")
        create_tile_group(middle_tiles, tools_y, "main_menuzeile_2", tile_w=max(210, int(tw * 0.86)), tile_h=max(105, int(th * 0.86)))
        # Trennlinie zwischen Tools-Menüs und Abschlusskalender – identisch zur Mini-Menüzeilen-Trennlinie.
        self.draw_continuous_relief_line(self.canvas, (top_y + tools_y) / 2, x_pct(w, 9.5), x_pct(w, 92))
        mini_tw = int(tw * 0.72)
        mini_th = int(th * 0.72)
        mini_center_y = y_pct(h, 16.5) + th / 2 - mini_th / 2
        mini_top = mini_center_y - mini_th / 2
        create_tile_group(mini_tiles, mini_center_y, "main_mini_menuezeile", tile_w=mini_tw, tile_h=mini_th)
        self.draw_continuous_relief_line(self.canvas, mini_top - 13, x_pct(w, 9.5), x_pct(w, 92))
        self.draw_bottom_logo()

    def try_open_in_dev(self):
        if self.my_role() != ROLE_E4: return
        self.show_page("in_dev", "In Entwicklung", True)

    def render_data_prep_menu(self):
        modules = [
            ("Nike-Tools", "page:nike_tools"),
            ("AFI-Uploads", "page:afi_uploads"),
            ("Aramark Monatsabrechnungen - PDF zu Excel", "aramark_monatsabrechnungen_pdf_to_excel"),
        ]
        self.render_module_menu(modules, show_descriptions=False)
        self.draw_bottom_logo()


    def render_debitoren_tools_menu(self):
        modules = [("Debitoren-Serienbrief", "debitoren_serienbrief")]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_afi_uploads_menu(self):
        modules = [
            # EnBW ist fachlich in "Lieferanten-Rechnung zu AFI-Upload" integriert.
            # Der alte Toolcode bleibt im TOOL_REGISTRY erhalten, wird aber nicht mehr angezeigt.
            ("Lieferanten-Rechnung zu AFI-Upload", "supplier_invoice_afi_upload"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_nike_tools_menu(self):
        modules = [
            ("Nike - PDF zu Excel", "nike_pdf_to_excel"),
            ("Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen", "nike_op_liste_pdf_check"),
            ("Nike - Rechnungs-PDFs in Sammelordner", "invoice_pdf_collector"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_closing_calendar_menu(self):
        modules = [("Monatsabschluss", "monthly_close"), ("Quartalsabschluss", "quarterly_close"), ("Jahresabschluss", "yearly_close"), ("Stichtagspflege", "deadline_maintenance")]
        self.render_module_menu(modules, show_descriptions=True)
        if self.my_role() == ROLE_E4:
            text = "Auto-Mail: Ein" if self.auto_close_mail_enabled() else "Auto-Mail: Aus"
            btn = tk.Button(self.root, text=text, command=self.toggle_auto_close_mail, bg=BLUE if self.auto_close_mail_enabled() else GREY_DISABLED, fg="white", bd=0, padx=10, pady=3, cursor="hand2", font=("Segoe UI", 9, "bold"))
            self.widget_items.append(btn)
            self.canvas.create_window(self.canvas.winfo_width() / 2, 136, window=btn, anchor="n")
        self.draw_bottom_logo()

    def render_compliance_audit_menu(self):
        modules = [
            ("Steuermeldungs-Cockpit", "tax_reporting"),
            ("Audit-Cockpit", "audit_cockpit"),
            ("Dokumentationszentrale", "documentation_center"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_in_dev_menu(self):
        modules = [("Compliance & Audit", "page:compliance_audit"), ("X001 SAP - Test", "x001_sap_test")]
        self.render_module_menu(modules, show_descriptions=True); self.draw_bottom_logo()


    def render_settings_menu(self):
        items = [("Farbschema", "tile_colors")]
        if self.can_view_user_management():
            items.append(("Benutzerverwaltung", "users"))
        if self.can_manage_permissions():
            items.append(("Berechtigungen", "permissions"))
        self.render_center_menu(items, title="Einstellungen"); self.draw_bottom_logo()

    def render_information_menu(self):
        self.render_center_menu([("Versionsverlauf", "versions")], title="Informationen"); self.draw_bottom_logo()

    def render_center_menu(self, items, title=""):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); tile_w = int(max(260, min(360, w * 0.2))); tile_h = int(max(110, min(160, h * 0.14))); gap = int(max(18, h * 0.03)); start_y = y_pct(h, 55); start_x = x_pct(w, 50)
        for i, (label, page) in enumerate(items):
            y = start_y + i * (tile_h + gap); cmd = (lambda p=page, l=label: self.show_page(p, l, True))
            tile = Tile(self.root, self, f"center_{page}", label, cmd, favorite_enabled=False, center_text=True); tile.resize_tile(tile_w, tile_h); self.widget_items.append(tile); self.focusable_tiles.append(tile); self.canvas.create_window(start_x, y, window=tile, anchor="center")

    def render_tile_colors_menu(self):
        zfont = self.zoomed_content_font
        frame = tk.Frame(self.root, bg=BG); self.widget_items.append(frame)
        tk.Label(frame, text="Standardfarben", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=5, pady=(0, 14), sticky="w")
        selected_color = self.current_tile_color() or BLUE
        def set_color(col):
            if self.current_user_key:
                self.user_data["users"][self.current_user_key]["tile_color"] = col; self.save_user_data()
            self.render_page()
        for idx, (name, color) in enumerate(COLOR_PALETTE):
            r = 1 + idx // 5; c = idx % 5; sw = tk.Canvas(frame, width=140, height=78, bg=BG, highlightthickness=0, bd=0, cursor="hand2")
            if color == selected_color:
                sw.create_rectangle(4, 4, 136, 58, outline=WHITE, width=4)
            sw.create_rectangle(8, 8, 132, 54, fill=color, outline=LINE, width=2); sw.create_text(70, 66, text=name, fill=TEXT, font=zfont(("Segoe UI", 9, "bold"))); sw.bind("<Button-1>", lambda e, col=color: set_color(col)); sw.grid(row=r, column=c, padx=10, pady=10)
        def reset():
            if self.current_user_key:
                self.user_data["users"][self.current_user_key].pop("tile_color", None); self.save_user_data()
            self.render_page()
        tk.Button(frame, text="Standard wiederherstellen", command=reset, bg=selected_color, fg="white", bd=0, padx=16, pady=10, cursor="hand2").grid(row=r + 1, column=0, columnspan=5, pady=(18, 6))
        self.canvas.create_window(self.canvas.winfo_width() / 2, y_pct(self.canvas.winfo_height(), 48), window=frame, anchor="center")


    def render_users_menu(self):
        zfont = self.zoomed_content_font
        if not self.can_view_user_management():
            self.render_menu_text("Keine Berechtigung für die Benutzerverwaltung."); return
        users = self.user_data.setdefault("users", {})
        visible_keys = [self.current_user_key] if self.my_role() == ROLE_E1 and self.current_user_key in users else sorted(users.keys())
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)
        arrow_col = tk.Frame(container, bg=BG, width=54); arrow_col.pack(side="left", fill="y", padx=(18, 0)); arrow_col.pack_propagate(False)
        body = tk.Frame(container, bg=BG); body.pack(side="left", fill="both", expand=True)
        footer = tk.Frame(body, bg=BG); footer.pack(side="bottom", fill="x", pady=(8, 0))
        scroll_canvas = tk.Canvas(body, bg=BG, highlightthickness=0, bd=0); scrollbar = tk.Scrollbar(body, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG); content_window = scroll_canvas.create_window((0, 0), window=content, anchor="nw")
        scroll_canvas.configure(yscrollcommand=scrollbar.set); scroll_canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y"); self.register_scroll_canvas(scroll_canvas, scrollbar)
        def update_arrows():
            try:
                first, last = scroll_canvas.yview(); up_arrow.set_enabled(first > 0.001); down_arrow.set_enabled(last < 0.999)
            except Exception: pass
        def scroll_units(n): scroll_canvas.yview_scroll(n, "units"); update_arrows()
        up_arrow = ArrowIndicator(arrow_col, "up", lambda: scroll_units(-5), size=42); down_arrow = ArrowIndicator(arrow_col, "down", lambda: scroll_units(5), size=42)
        up_arrow.pack(side="top", pady=(78, 10)); down_arrow.pack(side="top", pady=(0, 10))
        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width())); scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all")); self._sync_scrollbar_visibility(scroll_canvas, scrollbar); update_arrows()
        self._update_scroll_indicators = update_arrows; content.bind("<Configure>", update_scrollregion); scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<MouseWheel>", lambda e: self._route_mousewheel_to_canvas(e, scroll_canvas))
        frame = content
        tk.Label(frame, text="Benutzerverwaltung", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=8, sticky="w", pady=(0, 14))
        for col, header in enumerate(["Benutzer", "Anzeigename/Nachname", "Vorname", "E-Mail", "Rolle", "Passwort"]):
            tk.Label(frame, text=header, bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=col, sticky="w", padx=(0, 14))
        def can_edit_user(target_key):
            if target_key == self.current_user_key: return True
            if target_key == SUPERUSER_KEY: return False
            target_role = users.get(target_key, {}).get("permission", ROLE_E1)
            return self.my_role() in (ROLE_E3, ROLE_E4) or self.role_rank(target_role) < self.role_rank(self.my_role())
        def save_user_row(old_key, name_var, email_var, first_name_var=None):
            if old_key not in users: return
            data = users[old_key]; data["email"] = email_var.get().strip(); data["first_name"] = first_name_var.get().strip() if first_name_var else data.get("first_name", ""); data["full_name"] = " ".join(x for x in [data.get("first_name", "").strip(), data.get("display_name", old_key).strip()] if x).strip() or data.get("display_name", old_key)
            may_rename = ((old_key == SUPERUSER_KEY and self.current_user_key == SUPERUSER_KEY) or (self.my_role() in (ROLE_E3, ROLE_E4) and old_key not in (self.current_user_key, SUPERUSER_KEY) and can_edit_user(old_key)))
            if may_rename:
                new_name = " ".join(name_var.get().strip().split()); new_key = normalize_username(new_name)
                if not new_key: messagebox.showwarning("FiBu Mate", "Bitte einen Benutzernamen eingeben."); return
                if old_key != SUPERUSER_KEY and new_key == SUPERUSER_KEY: messagebox.showwarning("FiBu Mate", "Der Benutzer Wagnerm kann nicht überschrieben werden."); return
                if new_key != old_key and new_key in users: messagebox.showwarning("FiBu Mate", "Dieser Benutzername existiert bereits."); return
                data = users.pop(old_key); data["display_name"] = new_name; data["full_name"] = " ".join(x for x in [data.get("first_name", "").strip(), new_name] if x).strip() or new_name; users[new_key] = data
            self.save_user_data(); messagebox.showinfo("FiBu Mate", "Benutzerdaten wurden gespeichert."); self.render_page()
        def delete_user(user_key):
            if user_key in (SUPERUSER_KEY, self.current_user_key): messagebox.showwarning("FiBu Mate", "Dieser Benutzer kann nicht gelöscht werden."); return
            if not can_edit_user(user_key): messagebox.showwarning("FiBu Mate", "Keine Berechtigung für diesen Benutzer."); return
            if not messagebox.askyesno("Benutzer löschen", f"Benutzer wirklich löschen?\n\n{users.get(user_key, {}).get('display_name', user_key)}"): return
            users.pop(user_key, None); self.save_user_data(); messagebox.showinfo("FiBu Mate", "Benutzer wurde gelöscht."); self.render_page()
        row = 2
        for key in visible_keys:
            user = users[key]; user.setdefault("email", "")
            tk.Label(frame, text=key, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=0, sticky="w", padx=(0, 14), pady=4)
            name_var = tk.StringVar(value=user.get("display_name", key)); first_name_var = tk.StringVar(value=user.get("first_name", "")); email_var = tk.StringVar(value=user.get("email", ""))
            name_state = "normal" if self.my_role() in (ROLE_E3, ROLE_E4) and key not in (self.current_user_key, SUPERUSER_KEY) else "disabled"
            tk.Entry(frame, textvariable=name_var, font=("Segoe UI", 10), width=22, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1, state=name_state).grid(row=row, column=1, sticky="w", padx=(0, 14), pady=4)
            tk.Entry(frame, textvariable=first_name_var, font=("Segoe UI", 10), width=18, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=row, column=2, sticky="w", padx=(0, 14), pady=4)
            tk.Entry(frame, textvariable=email_var, font=("Segoe UI", 10), width=32, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=row, column=3, sticky="w", padx=(0, 14), pady=4)
            tk.Label(frame, text=user.get("permission", ROLE_E1), bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=4, sticky="w", padx=(0, 14), pady=4)
            tk.Label(frame, text="aktiv" if user.get("auth", {}).get("enabled") else "nicht aktiv", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=5, sticky="w", padx=(0, 14), pady=4)
            editable = can_edit_user(key)
            tk.Button(frame, text="Speichern", command=lambda k=key, n=name_var, e=email_var, f=first_name_var: save_user_row(k, n, e, f), bg=BLUE if editable else GREY_DISABLED, fg="white", bd=0, width=9, padx=10, pady=5, cursor="hand2" if editable else "arrow", state="normal" if editable else "disabled").grid(row=row, column=6, sticky="w", padx=(18, 10), pady=4)
            del_ok = editable and key not in (self.current_user_key, SUPERUSER_KEY) and self.my_role() in (ROLE_E3, ROLE_E4)
            tk.Button(frame, text="Löschen", command=lambda k=key: delete_user(k), bg=RED if del_ok else GREY_DISABLED, fg="white", bd=0, width=9, padx=10, pady=5, cursor="hand2" if del_ok else "arrow", state="normal" if del_ok else "disabled").grid(row=row, column=7, sticky="w", padx=(0, 8), pady=4)
            row += 1
        if self.can_create_users():
            tk.Label(footer, text="Neuen Benutzer anlegen", font=("Segoe UI", 12, "bold"), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=6, sticky="w", pady=(0, 8))
            username_var = tk.StringVar(); first_name_new_var = tk.StringVar(); email_new_var = tk.StringVar()
            tk.Label(footer, text="Benutzername", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=0, sticky="w", padx=(0, 8))
            tk.Entry(footer, textvariable=username_var, font=("Segoe UI", 10), width=28, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=1, sticky="w", padx=(0, 14))
            tk.Label(footer, text="Vorname", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=2, sticky="e", padx=(0, 4))
            tk.Entry(footer, textvariable=first_name_new_var, font=("Segoe UI", 10), width=18, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=3, sticky="w", padx=(0, 14))
            tk.Label(footer, text="E-Mail", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=4, sticky="e", padx=(0, 4))
            tk.Entry(footer, textvariable=email_new_var, font=("Segoe UI", 10), width=32, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=5, sticky="w", padx=(0, 14))
            def create_user():
                raw_name = " ".join(username_var.get().strip().split()); key = normalize_username(raw_name)
                if not key: messagebox.showwarning("FiBu Mate", "Bitte einen Benutzernamen eingeben."); return
                if key in users: messagebox.showwarning("FiBu Mate", "Dieser Benutzer existiert bereits."); return
                users[key] = {"display_name": raw_name, "first_name": first_name_new_var.get().strip(), "full_name": " ".join(x for x in [first_name_new_var.get().strip(), raw_name] if x).strip() or raw_name, "email": email_new_var.get().strip(), "favorites": [], "auth": {"password_hash": None, "enabled": False}, "permission": ROLE_E4 if key == SUPERUSER_KEY else ROLE_E1}
                self.ensure_permissions_defaults(); self.save_user_data(); messagebox.showinfo("FiBu Mate", f"Benutzer wurde angelegt:\n\n{raw_name}"); self.render_page()
            tk.Button(footer, text="Benutzer anlegen", command=create_user, bg=BLUE, fg="white", bd=0, padx=14, pady=8, cursor="hand2").grid(row=1, column=6, sticky="w")
        update_scrollregion()


    def render_permissions_menu(self):
        zfont = self.zoomed_content_font
        if not self.can_manage_permissions():
            self.render_menu_text("Keine Berechtigung für das Menü Berechtigungen."); return
        users = self.user_data.setdefault("users", {})
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)
        arrow_col = tk.Frame(container, bg=BG, width=54); arrow_col.pack(side="left", fill="y", padx=(18, 0)); arrow_col.pack_propagate(False)
        body = tk.Frame(container, bg=BG); body.pack(side="left", fill="both", expand=True)
        scroll_canvas = tk.Canvas(body, bg=BG, highlightthickness=0, bd=0); scrollbar = tk.Scrollbar(body, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG); content_window = scroll_canvas.create_window((0,0), window=content, anchor="nw")
        scroll_canvas.configure(yscrollcommand=scrollbar.set); scroll_canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y"); self.register_scroll_canvas(scroll_canvas, scrollbar)
        def update_arrows():
            try:
                first, last = scroll_canvas.yview(); up_arrow.set_enabled(first > 0.001); down_arrow.set_enabled(last < 0.999)
            except Exception: pass
        def scroll_units(n): scroll_canvas.yview_scroll(n, "units"); update_arrows()
        up_arrow = ArrowIndicator(arrow_col, "up", lambda: scroll_units(-5), size=42); down_arrow = ArrowIndicator(arrow_col, "down", lambda: scroll_units(5), size=42)
        up_arrow.pack(side="top", pady=(78, 10)); down_arrow.pack(side="top", pady=(0, 10))
        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width())); scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all")); self._sync_scrollbar_visibility(scroll_canvas, scrollbar); update_arrows()
        self._update_scroll_indicators = update_arrows; content.bind("<Configure>", update_scrollregion); scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<MouseWheel>", lambda e: self._route_mousewheel_to_canvas(e, scroll_canvas))
        frame = content
        tk.Label(frame, text="Berechtigungen", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=5, sticky="w", pady=(0, 10))
        def show_role_info():
            popup = tk.Toplevel(self.root); popup.title("Berechtigungen der Rollen anzeigen"); popup.configure(bg=BG); popup.geometry("820x620"); popup.transient(self.root)
            text = tk.Text(popup, bg="white", fg=TEXT, wrap="word", font=("Segoe UI", 10)); text.pack(fill="both", expand=True, padx=14, pady=14)
            content = "E1 - Standard\n- eigene Benutzerdaten ansehen/pflegen\n- Abschlusskalender lesen\n- eigene zuständige Aufgaben erledigen\n- Berichte, Historie und Änderungsprotokolle ansehen\n\nE2 - Erweitert\n- aktuell gleiche Berechtigungen wie E1\n- Wissenszentrale: Kategorien verwalten, Kategorien erstellen, umbenennen und farblich pflegen\n\nE3 - Administrator\n- Benutzer anlegen, bearbeiten und löschen\n- Berechtigungen maximal bis E3 vergeben\n- Aufgaben administrieren und Aufgaben-IDs bearbeiten/verknüpfen\n- Zeiträume nach Stichtag schließen/öffnen\n- Berichte und Änderungsprotokolle ansehen\n\nE4 - System-Administrator\n- Berechtigungen analog E3\n- System-Administrator"
            text.insert("1.0", content); text.config(state="disabled")
        tk.Button(frame, text="Berechtigungen der Rollen anzeigen", command=show_role_info, bg=HEADER, fg=TEXT, bd=1, padx=10, pady=5, cursor="hand2").grid(row=0, column=5, sticky="e", pady=(0,10))
        tk.Label(frame, text="Berechtigungen können von berechtigten Rollen maximal bis zur eigenen Rolle vergeben werden.", bg=BG, fg=TEXT2, font=("Segoe UI", 10), justify="left").grid(row=1, column=0, columnspan=5, sticky="w", pady=(0, 16))
        for col, header in enumerate(["Benutzer", "Anzeigename", "Aktuelle Rolle", "Ändern"]): tk.Label(frame, text=header, bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=2, column=col, sticky="w", padx=(0, 24))
        row = 3
        for key in sorted(users.keys()):
            user = users[key]; current_role = user.get("permission", ROLE_E1)
            tk.Label(frame, text=key, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=0, sticky="w", padx=(0, 24), pady=5)
            tk.Label(frame, text=user.get("display_name", key), bg=BG, fg=TEXT2, font=("Segoe UI", 10), wraplength=260, justify="left").grid(row=row, column=1, sticky="w", padx=(0, 24), pady=5)
            tk.Label(frame, text=current_role, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=2, sticky="w", padx=(0, 24), pady=5)
            if self.my_role() == ROLE_E3 and (key == SUPERUSER_KEY or current_role == ROLE_E4):
                tk.Label(frame, text="gesperrt", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=3, sticky="w", pady=5); row += 1; continue
            available_roles = [role for role in ROLE_ORDER if self.role_rank(role) <= self.max_assignable_role_rank()]
            if key != SUPERUSER_KEY and ROLE_E4 in available_roles: available_roles.remove(ROLE_E4)
            if key == SUPERUSER_KEY: available_roles = [ROLE_E4]
            role_var = tk.StringVar(value=current_role if current_role in available_roles else available_roles[0])
            dropdown = tk.OptionMenu(frame, role_var, *available_roles); dropdown.config(bg="#E8EEF5", fg=TEXT, bd=1, highlightthickness=0, cursor="hand2"); dropdown.grid(row=row, column=3, sticky="w", pady=5)
            def save_role(user_key=key, var=role_var):
                new_role = var.get()
                if user_key == SUPERUSER_KEY: new_role = ROLE_E4
                if user_key != SUPERUSER_KEY and new_role == ROLE_E4: messagebox.showwarning("FiBu Mate", "Die Rolle System-Administrator darf nur der Benutzer Wagnerm tragen."); return
                if self.role_rank(new_role) > self.max_assignable_role_rank(): messagebox.showwarning("FiBu Mate", "Du kannst keine Rolle vergeben, die höher als deine eigene Rolle ist."); return
                users[user_key]["permission"] = new_role; self.ensure_permissions_defaults(); self.save_user_data(); messagebox.showinfo("FiBu Mate", "Berechtigung wurde gespeichert."); self.render_page()
            tk.Button(frame, text="Speichern", command=save_role, bg=BLUE, fg="white", bd=0, padx=10, pady=5, cursor="hand2").grid(row=row, column=4, sticky="w", padx=(8, 0), pady=5)
            row += 1
        update_scrollregion()

    def render_versions_menu(self):
        history = self.load_version_history().get("entries", [])
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        area_top = 132
        area_bottom = max(area_top + 260, h - 92)
        view_h = int(area_bottom - area_top)

        container = tk.Frame(self.root, bg=BG)
        self.widget_items.append(container)
        self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)

        scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG)
        content_window = scroll_canvas.create_window((0, 0), window=content, anchor="nw")

        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width()))
            scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all"))
            self._sync_scrollbar_visibility(scroll_canvas, scrollbar)

        content.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.configure(yscrollcommand=scrollbar.set)
        scroll_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.register_scroll_canvas(scroll_canvas, scrollbar)

        def _on_versions_mousewheel(event):
            return self._route_mousewheel_to_canvas(event, scroll_canvas)

        scroll_canvas.bind("<MouseWheel>", _on_versions_mousewheel)
        content.bind("<MouseWheel>", _on_versions_mousewheel)

        history_source_text = f"Versionsverlauf: lokale Historie plus zentrale Release-Historie ({self._central_version_history_path()})"
        tk.Label(content, text=history_source_text, bg=BG, fg=TEXT2, font=("Segoe UI", 11, "bold"), wraplength=max(700, w - 80), justify="left").pack(anchor="w", padx=24, pady=(12, 0))
        if not history:
            tk.Label(content, text="Noch kein Versionsverlauf vorhanden.", bg=BG, fg=TEXT2, font=("Segoe UI", 12)).pack(anchor="w", padx=24, pady=24)
            update_scrollregion()
            return

        for entry in history:
            entry_frame = tk.Frame(content, bg=BG)
            entry_frame.pack(fill="x", padx=24, pady=(16, 0), anchor="n")
            ver = entry.get("version", "")
            date = entry.get("date", "")
            bullets = entry.get("bullets", [])
            tk.Label(entry_frame, text=f"{ver}.{date}", bg=BG, fg=TEXT, font=("Segoe UI", 12, "bold"), anchor="w").pack(fill="x", anchor="w")
            for bullet in bullets:
                tk.Label(entry_frame, text=f"• {bullet}", bg=BG, fg=TEXT2, font=body_font(11), justify="left", anchor="w", wraplength=max(300, w - 110)).pack(fill="x", anchor="w", padx=(16, 0), pady=(2, 0))
            separator = tk.Canvas(content, bg=BG, height=18, highlightthickness=0, bd=0)
            separator.pack(fill="x", padx=10, pady=(8, 0))
            try:
                self.draw_relief_line(separator, 8, 0, max(100, w - 70))
            except Exception:
                separator.create_line(0, 8, max(100, w - 70), 8, fill=LINE, width=1)
        update_scrollregion()

    def render_menu_text(self, text):
        lab = tk.Label(self.root, text=text, font=("Segoe UI", 16), bg=BG, fg=TEXT); self.widget_items.append(lab); self.canvas.create_window(x_pct(self.canvas.winfo_width(), 50), y_pct(self.canvas.winfo_height(), 60), window=lab, anchor="center"); self.draw_bottom_logo()

    def draw_continuous_relief_line(self, canvas, y, x1, x2):
        canvas.create_line(x1, y, x2, y, fill=blend(BG, "#1F2933", 0.30), width=2)
        canvas.create_line(x1, y + 2, x2, y + 2, fill=blend(BG, WHITE, 0.65), width=1)

    def draw_relief_line(self, canvas, y, x1, x2):
        for i in range(120):
            t = i / 119; fade = 1 - abs(t - 0.5) * 2; sx = x1 + (x2 - x1) * i / 120; ex = x1 + (x2 - x1) * (i + 1) / 120
            canvas.create_line(sx, y, ex, y, fill=blend(BG, "#1F2933", 0.30 * fade), width=2); canvas.create_line(sx, y + 2, ex, y + 2, fill=blend(BG, WHITE, 0.65 * fade), width=1)

    def draw_module_description(self, canvas, module_id, x1, y_top, width):
        """v0.436 Lesbarkeitsschutz: Beschreibung so groß wie möglich, aber im Sollbereich."""
        txt = MODULE_DESCRIPTIONS.get(module_id, "")
        if not txt:
            return
        max_w = max(120, int(width) - 2 * DESCRIPTION_X_OFFSET)
        max_h = max(ui_s(52), int(getattr(self, "_module_desc_tile_h", ui_s(130))) - 2 * DESCRIPTION_Y_OFFSET)
        chosen = body_font(10)
        try:
            for size in range(max(11, ui_s(12)), 8, -1):
                candidate = ("Segoe UI", size)
                f = tkfont.Font(root=self.root, font=candidate)
                words = str(txt).split() or [""]
                line_count = 1; cur = ""
                for word in words:
                    cand = word if not cur else cur + " " + word
                    if f.measure(cand) <= max_w:
                        cur = cand
                    else:
                        line_count += 1; cur = word
                if line_count * max(1, int(f.metrics("linespace"))) <= max_h:
                    chosen = candidate
                    break
        except Exception:
            chosen = body_font(10)
        canvas.create_text(x1 + DESCRIPTION_X_OFFSET, y_top + DESCRIPTION_Y_OFFSET, text=txt, anchor="nw", fill=DESCRIPTION_COLOR, font=chosen, width=max_w, justify="left")

    def module_icon_type(self, module_id):
        if module_id == "tax_reporting":
            return "tax_reporting"
        if module_id == "page:compliance_audit":
            return "compliance"
        if module_id == "audit_cockpit":
            return "audit"
        if module_id == "documentation_center":
            return "documentation"
        if self.current_page in ("data_prep", "nike_tools", "afi_uploads", "debitoren_tools"):
            if module_id == "enbw_strom_tanken_upload":
                return "xls"
            if module_id == "supplier_invoice_afi_upload":
                return "xls"
            if module_id in ("nike_pdf_to_excel", "nike_op_liste_pdf_check"):
                return "pdf_xls"
            if module_id == "invoice_pdf_collector":
                return "pdf_xls"
            if module_id == "debitoren_serienbrief":
                return "modules"
            if str(module_id).startswith("page:"):
                return "pdf_xls"
            return "pdf_xls"
        if self.current_page == "compliance_audit":
            return "compliance"
        if self.current_page == "closing_calendar":
            return "calendar"
        if module_id in ("monthly_close", "quarterly_close", "yearly_close", "deadline_maintenance"):
            if module_id == "deadline_maintenance" and self.role_rank() < 3:
                messagebox.showwarning("Keine Berechtigung", "Dieses Modul ist erst ab E3 verfügbar.")
                return
            return "calendar"
        if str(module_id).startswith("nike_") or module_id == "invoice_pdf_collector":
            return "pdf_xls"
        return "modules"

    def render_module_menu(self, modules, show_descriptions=True):
        modules = [(title, module_id) for title, module_id in modules if module_id not in HIDDEN_TOOL_IDS]
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); tile_w = max(290, min(390, int(w * 0.22))); tile_h = max(120, min(160, int(h * 0.15))); self._module_desc_tile_h = tile_h; gap = max(18, int(h * 0.025)); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top); first_center_x = x_pct(w, 25); first_center_y = y_pct(h, 70); left_x = max(0, first_center_x - tile_w / 2 - 8); top = max(0, first_center_y - tile_h / 2 - area_top - 8)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h); canvas_w = w - left_x - 10; scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0); scroll_canvas.place(x=left_x, y=0, width=canvas_w, height=view_h); content_h = top + len(modules) * (tile_h + gap) + 20; scroll_canvas.configure(scrollregion=(0, 0, canvas_w, content_h)); self.register_scroll_canvas(scroll_canvas); desc_x1 = tile_w + 90; desc_x2 = canvas_w - 20
        for idx, (title, module_id) in enumerate(modules):
            y = top + idx * (tile_h + gap); cmd = (lambda mid=module_id, ttl=title: self.show_page(str(mid).replace("page:", "", 1), ttl, True)) if str(module_id).startswith("page:") else ((lambda mid=module_id: self.open_tool(mid)) if module_id in TOOL_REGISTRY else self.show_placeholder)
            tile = Tile(scroll_canvas, self, module_id, title, cmd, favorite_enabled=module_id in TOOL_REGISTRY, icon_type=self.module_icon_type(module_id)); tile.resize_tile(tile_w, tile_h); self.focusable_tiles.append(tile); scroll_canvas.create_window(0, y, window=tile, anchor="nw")
            if show_descriptions: self.draw_module_description(scroll_canvas, module_id, desc_x1, y, desc_x2 - desc_x1)
            if idx < len(modules) - 1: self.draw_relief_line(scroll_canvas, y + tile_h + gap / 2, desc_x1, desc_x2)

    def render_external_tool(self, tool_id):
        if tool_id in HIDDEN_TOOL_IDS:
            messagebox.showinfo("FiBu Mate", "Dieses Tool ist in 'Lieferanten-Rechnung zu AFI-Upload' integriert und wird nicht mehr separat gestartet.")
            return
        try:
            module_path = TOOL_REGISTRY[tool_id]["module"]
            module = _load_tool_module_from_file(module_path) or importlib.import_module(module_path)
            try:
                importlib.reload(module)
            except Exception:
                pass
            if not hasattr(module, "render"): raise RuntimeError("Modul hat keine render(app)-Funktion")
            # v0.436: Tool-Kontext setzen, damit das zusammengefasste Abschlusskalender-Modul
            # korrekt zwischen Monats-, Quartals- und Jahresabschluss routet.
            self.current_tool_id = tool_id
            self.current_tool_title = TOOL_REGISTRY.get(tool_id, {}).get("title", tool_id)
            module.render(self)
        except Exception as e:
            messagebox.showerror("FiBu Mate", f"Fehler beim Laden des Moduls:\n\n{TOOL_REGISTRY.get(tool_id, {}).get('title', tool_id)}\n\n{e}"); self.draw_bottom_logo()

    def open_tool(self, tool_id):
        if str(tool_id).startswith("page:"):
            page = str(tool_id).split(":", 1)[1]
            titles = {"nike_tools": "Nike-Tools", "afi_uploads": "AFI-Uploads", "debitoren_tools": "Tools - Debitoren"}
            self.show_page(page, titles.get(page, page), True)
            return
        if tool_id in TOOL_REGISTRY:
            self.show_page(f"tool:{tool_id}", TOOL_REGISTRY[tool_id]["title"], True)

    def show_placeholder(self): messagebox.showinfo("FiBu Mate", "Hinter diesem Widget entsteht gerade ein Modul.")

    def login_user_from_entry(self, username):
        self._suppress_next_global_return = True
        self.login_user(username)
        return "break"

    def login_user(self, username):
        username = " ".join(str(username).strip().split())
        key = normalize_username(username)
        if not key:
            messagebox.showwarning("FiBu Mate", "Bitte einen Nutzernamen eingeben.")
            return
        users = self.user_data.setdefault("users", {})
        if key not in users:
            messagebox.showwarning("FiBu Mate", "Benutzer nicht gefunden.\nBitte wende dich an eine Administratorin / einen Administrator.")
            return
        users[key].setdefault("display_name", username)
        users[key].setdefault("first_name", "")
        users[key].setdefault("full_name", " ".join(x for x in [users[key].get("first_name", "").strip(), users[key].get("display_name", username).strip()] if x).strip() or username)
        users[key].setdefault("favorites", [])
        users[key].setdefault("email", "")
        users[key].setdefault("auth", {"password_hash": None, "enabled": False})
        users[key]["permission"] = ROLE_MIGRATION.get(users[key].get("permission", ROLE_E1), ROLE_E1)
        if key == SUPERUSER_KEY:
            users[key]["permission"] = ROLE_E4
        self.ensure_permissions_defaults()
        self.current_user_key = key
        self.current_user_display = users[key].get("display_name", username)
        self.favorites = set(fav for fav in users[key].get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
        users[key]["favorites"] = sorted(self.favorites)
        self.save_user_data()
        self.start_live_permissions_refresh()
        self.page_history = []
        self.breadcrumb = []
        self.show_page("main", "Hauptmenü", add_to_history=False)

    def start_live_permissions_refresh(self):
        if self._live_permissions_started:
            return
        self._live_permissions_started = True
        self.root.after(3000, self.check_live_user_data_updates)

    def _show_single_live_popup(self, title, text):
        if self._live_permissions_popup_open:
            return
        self._live_permissions_popup_open = True
        try:
            messagebox.showinfo(title, text)
        finally:
            self._live_permissions_popup_open = False

    def check_live_user_data_updates(self):
        try:
            if not self.current_user_key:
                self.root.after(3000, self.check_live_user_data_updates)
                return
            mtime = self._get_user_data_mtime()
            if mtime and mtime != getattr(self, "_user_data_mtime", 0):
                old_role = self.my_role()
                old_display = self.current_user_display
                old_favorites = set(self.favorites)
                self.user_data = self.load_user_data()
                self.ensure_permissions_defaults()
                users = self.user_data.get("users", {})
                key = self.current_user_key
                if key not in users:
                    self.current_user_key = None
                    self.current_user_display = ""
                    self.favorites = set()
                    self.page_history = []
                    self.show_page("launch", add_to_history=False)
                    self._show_single_live_popup("FiBu Mate", "Ihr Benutzer wurde entfernt oder ist nicht mehr verfügbar. Sie wurden abgemeldet.")
                    self._user_data_mtime = mtime
                    return
                user = users[key]
                self.current_user_display = user.get("display_name", old_display)
                self.favorites = set(fav for fav in user.get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
                new_role = self.my_role()
                self._user_data_mtime = mtime
                if new_role != old_role or self.current_user_display != old_display or self.favorites != old_favorites:
                    try:
                        self.update_header()
                    except Exception:
                        pass
                    try:
                        self.render_page()
                    except Exception:
                        pass
                    self._show_single_live_popup("FiBu Mate", "Ihre Benutzer-/Berechtigungsdaten wurden live aktualisiert.")
        except Exception:
            pass
        try:
            self.root.after(3000, self.check_live_user_data_updates)
        except Exception:
            pass

    def logout(self):
        self.current_user_key = None; self.current_user_display = ""; self.favorites = set(); self.page_history = []; self.show_page("launch", add_to_history=False)

    def set_focused_tile(self, tile):
        if tile in self.focusable_tiles: self.focus_index = self.focusable_tiles.index(tile)

    def handle_escape(self, *_):
        if self._close_modal_window():
            return "break"
        handler = getattr(self, "module_escape_handler", None)
        if handler:
            try:
                if handler():
                    return "break"
            except Exception:
                pass
        if self.current_page != "launch":
            self.go_back()
        return "break"
    def handle_enter(self, *_):
        if getattr(self, "_suppress_next_global_return", False):
            self._suppress_next_global_return = False
            return "break"
        focus_widget = self.root.focus_get()
        if isinstance(focus_widget, tk.Entry):
            return "break"
        if 0 <= self.focus_index < len(self.focusable_tiles):
            try:
                self.focusable_tiles[self.focus_index].on_keyboard_activate()
            except Exception:
                pass
        return "break"

    def handle_tab(self, *_):
        if self.focusable_tiles: self.focus_index = (self.focus_index + 1) % len(self.focusable_tiles); self.focusable_tiles[self.focus_index].focus_set()
        return "break"

    def handle_shift_tab(self, *_):
        if self.focusable_tiles: self.focus_index = (self.focus_index - 1) % len(self.focusable_tiles); self.focusable_tiles[self.focus_index].focus_set()
        return "break"

    def confirm_exit(self):
        if getattr(self, "_closing_in_progress", False):
            return
        self._closing_in_progress = True
        try:
            if not self.confirm_unsaved_changes():
                self._closing_in_progress = False
                return
            if not messagebox.askyesno(APP_NAME, "FiBu Mate wirklich schließen?"):
                self._closing_in_progress = False
                return
            try:
                for w in list(self.root.winfo_children()):
                    if isinstance(w, tk.Toplevel):
                        try:
                            w.destroy()
                        except Exception:
                            pass
            except Exception:
                pass
            try:
                self.root.quit()
            except Exception:
                pass
            try:
                self.root.destroy()
            except Exception:
                pass
        except Exception:
            try:
                self.root.destroy()
            except Exception:
                pass


    def ensure_version_433_once(self):
        self.bump_version_once(
            "2026-05-30_v0433_responsive_scaling_documentation_center_table",
            [
                "v0.433: Responsive Skalierung für kleinere Monitore ergänzt; zentrale Schrift-, Kachel- und Widgetgrößen werden an die Bildschirmgröße angepasst.",
                "v0.433: Dokumentationszentrale-Tabelle umgebaut: Aufgabenzuordnung als erste Spalte, Dokument statt Dokumentname, Status und Pfad ausgeblendet.",
                "v0.433: Dokumentnamen in der Dokumentationszentrale sind anklickbar und öffnen die hinterlegte Datei; Datei- und Anhang-Icons ergänzt.",
                "v0.433: Dokumentationszentrale erhält horizontale Scroll-Unterstützung für Tabellenüberlauf und optimierte Spaltenbreiten.",
                "v0.433 Korrektur Paket 1b: Kachel-/Widget-Skalierung nachjustiert, Schließen-Logik gegen doppelten Dialog bzw. weißes Restfenster abgesichert und Mini-Widget-Schließen neben Hilfe neutralisiert.",
            ],
        )

    def run(self):
        try:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_force()
            self.root.update_idletasks()
            self.root.update()
        except Exception:
            pass
        self.root.mainloop()


    def ensure_version_435_once(self):
        """Version 0.435: Audit-Cockpit-Korrektur und Zuständigkeitspflege mit Aufgabenverknüpfung."""
        update_id = "2026-05-29_v0435_audit_cockpit_zustaendigkeit_verknuepfung"
        bullets = [
            "v0.435: Audit-Cockpit-Fehler korrigiert; Zeitraumformatierung akzeptiert den Zeitstempel-Kontext.",
            "v0.435: Zuständigkeitspflege erweitert um Aufgabenverknüpfung mit gemeinsamer Aufgaben-ID.",
            "v0.435: Mehrfachpräfixe M/Q/J umgesetzt, z. B. MQ oder QJ; Aufgaben erscheinen in allen passenden Übersichtsreitern.",
            "v0.435: Zuständigkeit und Benutzer-Key werden gemeinsam als Feld Zuständigkeit gepflegt.",
            "v0.435: Drop-downs in allen editierbaren Zuständigkeitspflege-Zellen, Checkboxen in der Aufgaben-ID-Spalte, Buttons Aufgaben verknüpfen und Löschen ergänzt.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 35 or current_build > 100:
                self.version_state["build"] = 35
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.435",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def normalize_version_after_zoom_patch(self):
        """Korrektur: versehentliche Versionssprünge aus Zoom-Testpaketen zurücknehmen, v0.435 aber beibehalten."""
        try:
            changed = False
            try:
                build = int(self.version_state.get("build", DEFAULT_BUILD))
            except Exception:
                build = DEFAULT_BUILD
            if build > 35:
                self.version_state["build"] = 35
                changed = True
            applied = self.version_state.setdefault("applied_updates", [])
            for upd in ("2026-05-29_v0434_paket1c_fibu_mate_zoom_miniwidget", "2026-05-29_v0434_paket1d_fibu_mate_zoom_miniwidget_nachkorrektur", "2026-05-29_v0434_paket1e_zoomleiste"):
                if upd not in applied:
                    applied.append(upd)
                    changed = True
            self.version_state["applied_updates"] = applied
            if changed:
                self.save_version_state()
            history = self.load_version_history()
            entries = history.setdefault("entries", [])
            cleaned = []
            for entry in entries:
                version = str(entry.get("version", ""))
                update_id = str(entry.get("update_id", ""))
                bullets = "\n".join(str(b) for b in entry.get("bullets", [])).lower()
                accidental = version in ("v0.436", "v0.437") or update_id in ("2026-05-29_v0434_paket1c_fibu_mate_zoom_miniwidget", "2026-05-29_v0434_paket1d_fibu_mate_zoom_miniwidget_nachkorrektur", "2026-05-29_v0434_paket1e_zoomleiste") or "paket 1c" in bullets or "paket 1d" in bullets or "paket 1e" in bullets or "zoom-leiste" in bullets or "zoomleiste" in bullets
                if not accidental:
                    cleaned.append(entry)
            if len(cleaned) != len(entries):
                history["entries"] = cleaned
                self.save_version_history(history)
            try:
                self.version_label.config(text=self.version_label_text())
            except Exception:
                pass
        except Exception:
            pass

    def ensure_version_434_once(self):
        """v0.434: Bereichsbezogener Zoom und Lesbarkeitspaket 1A."""
        try:
            self.bump_version_once(
                "2026-05-29_v0434_scope_zoom_readability_paket1a",
                [
                    "v0.434 Paket 1A: Zoom-Architektur von global auf bereichsbezogen umgestellt.",
                    "v0.434 Paket 1A: Zoomprofile werden pro Benutzer, Bildschirmgröße und Bereich/Modul gespeichert.",
                    "v0.434 Paket 1A: Kopf- und Fußleiste sind vom Strg+Mausrad-Zoom ausgenommen; Content, Module, Popups und Canvas-Inhalte können bereichsbezogen zoomen.",
                    "v0.434 Paket 1A: Kachel-Hover erhält den aktiven Menü-Zoom und setzt skalierte Kacheltexte nicht mehr zurück.",
                    "v0.434 Paket 1A: Steuermeldungs- und Audit-Cockpit erhalten größere Standard-Tabellen-/Dialogschriften.",
                ],
            )
        except Exception:
            pass

    def ensure_version_434_1c_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return

    def ensure_version_434_1d_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return

    def ensure_version_434_1e_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return


# ------------------------------------------------------------------
# Wissenszentrale - Next Blocks Final Patch 2026-06-12
# Ergänzt per Monkey-Patch: Word-Export, Anhänge, Kommentare,
# B2-Kategorienverwaltung, B1-Ausblendung, ESC-Logik, neutrale Navigation,
# Übersicht als Standard, vergrößerte Dropdowns, To-Do-Rhythmus bedingt.
# ------------------------------------------------------------------
def _kbp_permission_level(self):
    key = str(getattr(self, "current_user_key", "") or "").lower()
    disp = str(getattr(self, "current_user_display", "") or "").lower()
    if "wagnerm" in key or "wagnerm" in disp or "admin" in key or "admin" in disp:
        return 4
    raw = ""
    try:
        data = getattr(self, "user_data", {}) or {}
        raw = str(data).lower()
    except Exception:
        raw = ""
    check = " ".join([key, disp, raw])
    if "e4" in check or "stufe 4" in check or "level 4" in check or "berechtigung 4" in check:
        return 4
    if "e3" in check or "stufe 3" in check or "level 3" in check or "berechtigung 3" in check:
        return 3
    if "b2" in check or "stufe 2" in check or "level 2" in check or "berechtigung 2" in check:
        return 2
    return 1

def _kbp_can_edit(self):
    return self.kb_current_permission_level() >= 2

def _kbp_can_manage_categories(self):
    return self.kb_current_permission_level() >= 2

def _kbp_draw_button(self, x, y, text, command=None, accent=False, width=None):
    btn = tk.Button(self.root, text=text, command=command, font=body_font(10, weight="bold" if accent else None),
                    bg=WHITE, fg=TEXT, activebackground="#DDEAF7", relief="solid", bd=1, cursor="hand2", padx=8, pady=3)
    btn_w = width or max(110, 18 + len(text) * 8)
    self.canvas.create_window(ui_s(x), ui_s(y), window=btn, anchor="nw", width=ui_s(btn_w), height=ui_s(34))
    self.widget_items.append(btn)
    return x + btn_w + 8

def _kbp_handle_escape(self, event=None):
    if getattr(self, "current_page", "") == "knowledge_base":
        if not bool(getattr(self, "knowledge_start_overlay", False)):
            self.knowledge_start_overlay = True
            self.render_page()
            return "break"
        self.knowledge_start_overlay = False
        self.show_page("main_menu", "Hauptmenü", False)
        return "break"
    return None

def _kbp_export_dir(self):
    path = os.path.join(APPDATA_DIR, "Wissenszentrale_Export")
    try: os.makedirs(path, exist_ok=True)
    except Exception: pass
    return path

def _kbp_attachment_dir(self, entry_id):
    path = os.path.join(APPDATA_DIR, "Wissenszentrale_Anhaenge", str(entry_id))
    try: os.makedirs(path, exist_ok=True)
    except Exception: pass
    return path

def _kbp_add_pending_attachment(self):
    try:
        from tkinter import filedialog
        files = filedialog.askopenfilenames(title="Anhänge auswählen")
    except Exception:
        files = []
    if not files:
        return
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    for f in files:
        if f and f not in pending:
            pending.append(f)
    self.kb_pending_attachments = pending
    self.kb_mark_unsaved()
    self.render_page()

def _kbp_add_comment_to_selected(self, text_widget):
    entry_id = getattr(self, "kb_selected_entry_id", None)
    try: comment = text_widget.get("1.0", "end-1c").strip()
    except Exception: comment = ""
    if not entry_id or not comment:
        return
    entries = self.kb_load_entries()
    for e in entries:
        if e.get("id") == entry_id:
            e.setdefault("comments", []).append({"user": self.current_user_display or self.current_user_key or "", "text": comment, "created_at": self.kb_now()})
            e["updated_at"] = self.kb_now()
            break
    if self.kb_save_entries(entries):
        self.render_page()

def _kbp_save_entry_from_form(self):
    if not self.kb_can_create_or_edit():
        return False
    self.kb_ensure_state_vars()
    title = (getattr(self, "kb_title_var", tk.StringVar(value="")).get() or "").strip()
    if not title:
        try: messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
        except Exception: pass
        return False
    categories = []
    for var in getattr(self, "kb_entry_category_vars", []):
        value = (var.get() or "").strip()
        if value and value not in categories:
            categories.append(value)
    categories = categories[:4]
    user = (getattr(self, "kb_user_var", tk.StringVar(value="")).get() or "").strip()
    status = (getattr(self, "kb_status_var", tk.StringVar(value="Aktiv")).get() or "Aktiv").strip()
    rhythm = (getattr(self, "kb_rhythm_var", tk.StringVar(value="")).get() or "").strip() if any(c.lower()=="to-do" for c in categories) else ""
    try: text_value = self.kb_text_widget.get("1.0", "end-1c")
    except Exception: text_value = getattr(self, "kb_text_initial", "")
    entries = self.kb_load_entries(); now = self.kb_now(); edit_id = getattr(self, "kb_edit_entry_id", None)
    if edit_id:
        selected_id = edit_id
        for e in entries:
            if e.get("id") == edit_id:
                e.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "updated_at": now})
                break
    else:
        selected_id = self.kb_make_entry_id()
        entries.append({"id": selected_id, "title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "created_at": now, "updated_at": now, "comments": [], "attachments": []})
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    if pending:
        import shutil as _shutil
        copied=[]; target=self.kb_attachment_dir(selected_id)
        for src in pending:
            try:
                name=os.path.basename(src); dst=os.path.join(target,name); _shutil.copy2(src,dst); copied.append({"name": name, "path": dst, "added_at": now})
            except Exception: pass
        for e in entries:
            if e.get("id") == selected_id:
                e.setdefault("attachments", []).extend(copied); break
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = selected_id; self.kb_edit_entry_id = None; self.kb_pending_attachments=[]; self.knowledge_unsaved=False; self.knowledge_view="detail"; self.render_page(); return True
    return False

def _kbp_confirm_unsaved_before_switch(self):
    if not self.kb_has_unsaved_changes():
        return True
    try:
        result = messagebox.askyesnocancel("Ungespeicherte Änderungen", "Änderungen speichern?\n\nJa = speichern und wechseln\nNein = verwerfen und wechseln\nAbbrechen = bleiben")
    except Exception:
        result = False
    if result is True:
        return self.kb_save_entry_from_form() if getattr(self, "knowledge_view", "") == "new" else self.kb_save_current_work()
    if result is False:
        self.knowledge_unsaved = False; return True
    return False

def _kbp_switch_view_from_start(self, view):
    if not self.kb_confirm_unsaved_before_switch(): return
    if view == "new" and not self.kb_can_create_or_edit(): view = "all"
    if view == "categories" and not self.kb_can_manage_categories(): view = "all"
    self.knowledge_view = view
    if view == "new": self.kb_prepare_new_entry()
    self.knowledge_start_overlay = False
    self.render_page()

def _kbp_export_word(self):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry: return
    try:
        from docx import Document
        import re as _re
        doc = Document(); doc.add_heading(entry.get("title", "Wissenseintrag"), level=1)
        for label, val in [("Geändert", self.kb_display_date(entry.get("updated_at"))), ("Benutzer", entry.get("user","")), ("Status", entry.get("status","")), ("Kategorien", ", ".join(entry.get("categories", []) or [])), ("To-Do-Rhythmus", entry.get("rhythm", ""))]:
            if val: doc.add_paragraph(f"{label}: {val}")
        doc.add_heading("Inhalt", level=2); doc.add_paragraph(entry.get("text", ""))
        doc.add_heading("Anhänge", level=2)
        for a in entry.get("attachments", []) or []: doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        doc.add_heading("Kommentare", level=2)
        for c in entry.get("comments", []) or []: doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        safe = _re.sub(r"[^A-Za-z0-9_äöüÄÖÜß-]+", "_", entry.get("title", "Wissenseintrag"))[:80]
        out = os.path.join(self.kb_export_dir(), f"{safe}.docx")
        doc.save(out); messagebox.showinfo("Wissenszentrale", "Word-Export erstellt:\n" + out)
    except Exception as exc:
        messagebox.showerror("Wissenszentrale", "Word-Export fehlgeschlagen:\n" + str(exc))

def _kbp_render_work(self):
    self.kb_ensure_state_vars()
    try: self.root.option_add("*TCombobox*Listbox.font", body_font(12)); self.root.bind("<Escape>", self.kb_handle_escape, add="+")
    except Exception: pass
    w,h=self.canvas.winfo_width(), self.canvas.winfo_height()
    if self.knowledge_view == "new" and not self.kb_can_create_or_edit(): self.knowledge_view="all"
    if self.knowledge_view == "categories" and not self.kb_can_manage_categories(): self.knowledge_view="all"
    x=28; y=145
    x=self.draw_kb_button(x,y,"Start",self.kb_show_start_overlay,False,92)
    x=self.draw_kb_button(x,y,"Übersicht",lambda:self.kb_switch_view_from_start("all"),False,112)
    if self.kb_can_create_or_edit(): x=self.draw_kb_button(x,y,"Neuer Eintrag",lambda:self.kb_switch_view_from_start("new"),False,128)
    x=self.draw_kb_button(x,y,"To-Dos",lambda:self.kb_switch_view_from_start("todos"),False,92)
    x=self.draw_kb_button(x,y,"Veraltete Einträge",lambda:self.kb_switch_view_from_start("outdated"),False,148)
    if self.kb_can_manage_categories(): x=self.draw_kb_button(x,y,"Kategorien verwalten",lambda:self.kb_switch_view_from_start("categories"),False,168)
    sf=tk.Frame(self.root,bg=BG); self.widget_items.append(sf)
    tk.Label(sf,text="Suche",bg=BG,fg=TEXT2,font=body_font(10)).grid(row=0,column=0,sticky="w")
    ent=tk.Entry(sf,textvariable=self.kb_search_var,font=body_font(11),bg=WHITE,fg=TEXT,relief="solid",bd=1); ent.grid(row=1,column=0,sticky="ew",ipady=5); ent.bind("<Return>", self.kb_on_search_return)
    tk.Button(sf,text="Suchen",command=self.kb_apply_filters,bg=WHITE,fg=TEXT,relief="solid",bd=1,font=body_font(10)).grid(row=1,column=1,padx=(8,0),sticky="ns")
    sf.grid_columnconfigure(0,weight=1); self.canvas.create_window(ui_s(28),ui_s(192),window=sf,anchor="nw",width=ui_s(min(790,max(560,w-80))),height=ui_s(60))
    ff=tk.Frame(self.root,bg=BG); self.widget_items.append(ff); vals=[""]+self.kb_get_categories(); tk.Label(ff,text="Kategorie-Filter",bg=BG,fg=TEXT2,font=body_font(10)).grid(row=0,column=0,sticky="w",columnspan=6)
    for i,var in enumerate(self.kb_filter_vars):
        b=tk.Frame(ff,bg=BG); b.grid(row=1,column=i,padx=(0 if i==0 else 12,0),sticky="w"); tk.Label(b,text=f"Kategorie {i+1}",bg=BG,fg=TEXT2,font=body_font(9)).pack(anchor="w"); cb=ttk.Combobox(b,textvariable=var,values=vals,state="readonly",font=body_font(11),width=18); cb.pack(anchor="w"); cb.bind("<<ComboboxSelected>>",self.kb_apply_filters)
    selected=[(v.get() or "").strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view == "todos" or "to-do" in selected:
        b=tk.Frame(ff,bg=BG); b.grid(row=1,column=4,padx=(12,0),sticky="w"); tk.Label(b,text="To-Do-Rhythmus",bg=BG,fg=TEXT2,font=body_font(9)).pack(anchor="w"); cb=ttk.Combobox(b,textvariable=self.kb_todo_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(11),width=16); cb.pack(anchor="w"); cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    else:
        self.kb_todo_rhythm_var.set(""); tk.Label(ff,text="Leere Filter = keine Filterung",bg=BG,fg=TEXT2,font=body_font(9)).grid(row=1,column=4,padx=(14,0),sticky="s")
    self.canvas.create_window(ui_s(28),ui_s(252),window=ff,anchor="nw",width=ui_s(min(1050,max(760,w-70))),height=ui_s(90))
    left_x=20; left_y=365; left_w=max(340,min(460,int(w*0.31))); pane_h=max(420,h-left_y-24); right_x=left_x+left_w+24; right_w=max(760,w-right_x-24)
    self.render_kb_hits_pane(left_x,left_y,left_w,pane_h)
    if self.knowledge_view == "new": self.render_kb_new_entry_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view == "todos": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="To-Dos")
    elif self.knowledge_view == "outdated": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Veraltete Einträge",status_filter="Veraltet")
    elif self.knowledge_view == "categories": self.render_kb_categories_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view == "detail": self.render_kb_detail_area(right_x,left_y,right_w,pane_h)
    else: self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Übersicht")

def _kbp_render_new(self,x,y,w,h):
    if not self.kb_can_create_or_edit(): self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    if not hasattr(self,"kb_title_var"): self.kb_prepare_new_entry()
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    tk.Label(frame,text=("Eintrag bearbeiten" if getattr(self,"kb_edit_entry_id",None) else "Neuer Eintrag"),bg=WHITE,fg=BLUE,font=body_font(15,weight="bold")).pack(anchor="w",padx=18,pady=(16,6))
    form=tk.Frame(frame,bg=WHITE); form.pack(fill="x",padx=18,pady=(4,8)); tk.Label(form,text="Titel des Eintrags",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=0,column=0,sticky="w")
    tk.Entry(form,textvariable=self.kb_title_var,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).grid(row=1,column=0,columnspan=4,sticky="ew",ipady=5,pady=(0,8))
    cats=[""]+self.kb_get_categories()
    for i,var in enumerate(self.kb_entry_category_vars):
        tk.Label(form,text=f"Kategorie {i+1}",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=2,column=i,sticky="w",padx=(0 if i==0 else 8,0)); cb=ttk.Combobox(form,textvariable=var,values=cats,state="normal",font=body_font(11),width=18); cb.grid(row=3,column=i,sticky="ew",padx=(0 if i==0 else 8,0),pady=(0,8)); cb.bind("<<ComboboxSelected>>",lambda e:self.render_page())
    tk.Label(form,text="Assoziierter Benutzer",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=0,sticky="w"); tk.Entry(form,textvariable=self.kb_user_var,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).grid(row=5,column=0,sticky="ew",pady=(0,8))
    tk.Label(form,text="Status",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=1,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_status_var,values=["Aktiv","Entwurf","Veraltet"],state="readonly",font=body_font(11)).grid(row=5,column=1,sticky="ew",padx=(8,0),pady=(0,8))
    if any((v.get() or "").strip().lower()=="to-do" for v in self.kb_entry_category_vars):
        tk.Label(form,text="To-Do-Rhythmus",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=2,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(11)).grid(row=5,column=2,sticky="ew",padx=(8,0),pady=(0,8))
    else: self.kb_rhythm_var.set("")
    for c in range(4): form.grid_columnconfigure(c,weight=1)
    bf=tk.Frame(frame,bg=WHITE); bf.pack(fill="both",expand=True,padx=18,pady=(0,10)); tk.Label(bf,text="Freitext / Prozessdokumentation / Leitfaden",bg=WHITE,fg=TEXT2,font=body_font(9)).pack(anchor="w"); self.kb_text_widget=tk.Text(bf,bg="#F8FAFC",fg=TEXT,font=body_font(10),relief="solid",bd=1,wrap="word"); self.kb_text_widget.pack(fill="both",expand=True); self.kb_text_widget.insert("1.0",getattr(self,"kb_text_initial","")); self.kb_text_widget.bind("<KeyRelease>",lambda e:self.kb_mark_unsaved())
    pending=getattr(self,"kb_pending_attachments",[]) or []
    if pending: tk.Label(frame,text="Vorgemerkte Anhänge: "+", ".join([os.path.basename(x) for x in pending]),bg=WHITE,fg=TEXT2,font=body_font(9)).pack(anchor="w",padx=18)
    br=tk.Frame(frame,bg=WHITE); br.pack(fill="x",padx=18,pady=(0,16)); tk.Button(br,text="Anhang hinzufügen",command=self.kb_add_pending_attachment,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=12,ipady=4); tk.Button(br,text="Als Entwurf speichern",command=lambda:(self.kb_status_var.set("Entwurf"),self.kb_save_entry_from_form()),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=12,ipady=4); tk.Button(br,text="Speichern",command=self.kb_save_entry_from_form,bg=WHITE,fg=TEXT,font=body_font(10,weight="bold"),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Abbrechen",command=lambda:self.kb_switch_view_from_start("all"),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",ipadx=16,ipady=4)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

def _kbp_render_detail(self,x,y,w,h):
    e=self.kb_get_entry(getattr(self,"kb_selected_entry_id",None))
    if not e: self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    tk.Label(frame,text=e.get("title",""),bg=WHITE,fg=BLUE,font=body_font(16,weight="bold")).pack(anchor="w",padx=18,pady=(16,6)); tk.Label(frame,text=f"Geändert: {self.kb_display_date(e.get('updated_at'))}    Benutzer: {e.get('user','')}    Status: {e.get('status','')}",bg=WHITE,fg=TEXT2,font=body_font(10)).pack(anchor="w",padx=18)
    tk.Label(frame,text="Kategorien: "+(", ".join(e.get("categories",[]) or []) or "Keine Kategorien"),bg=WHITE,fg=TEXT,font=body_font(10,weight="bold")).pack(anchor="w",padx=18,pady=(8,4))
    txt=tk.Text(frame,bg="#F8FAFC",fg=TEXT,font=body_font(10),relief="solid",bd=1,wrap="word",height=8); txt.pack(fill="both",expand=True,padx=18,pady=(6,8)); txt.insert("1.0",e.get("text","")); txt.configure(state="disabled")
    lower=tk.Frame(frame,bg=WHITE); lower.pack(fill="x",padx=18,pady=(0,8)); left=tk.Frame(lower,bg=WHITE); left.pack(side="left",fill="both",expand=True); right=tk.Frame(lower,bg=WHITE); right.pack(side="right",fill="both",expand=True,padx=(12,0))
    tk.Label(left,text="Anhänge",bg=WHITE,fg=BLUE,font=body_font(10,weight="bold")).pack(anchor="w")
    for a in e.get("attachments",[]) or []: tk.Label(left,text="• "+a.get("name",""),bg=WHITE,fg=TEXT,font=body_font(9)).pack(anchor="w")
    tk.Label(right,text="Kommentare",bg=WHITE,fg=BLUE,font=body_font(10,weight="bold")).pack(anchor="w")
    for c in (e.get("comments",[]) or [])[-3:]: tk.Label(right,text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}",bg=WHITE,fg=TEXT,font=body_font(9),wraplength=380,justify="left").pack(anchor="w")
    comment=tk.Text(right,height=2,bg="#F8FAFC",fg=TEXT,font=body_font(9),relief="solid",bd=1,wrap="word"); comment.pack(fill="x",pady=(4,0))
    br=tk.Frame(frame,bg=WHITE); br.pack(fill="x",padx=18,pady=(0,14))
    if self.kb_can_create_or_edit(): tk.Button(br,text="Bearbeiten",command=self.kb_edit_selected_entry,bg=WHITE,fg=TEXT,font=body_font(10,weight="bold"),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4)
    tk.Button(br,text="Kommentar speichern",command=lambda:self.kb_add_comment_to_selected(comment),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Word-Export",command=self.kb_export_selected_to_word,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Zur Übersicht",command=lambda:self.kb_switch_view_from_start("all"),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",ipadx=16,ipady=4)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

def _kbp_render_categories(self,x,y,w,h):
    if not self.kb_can_manage_categories(): self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame); tk.Label(frame,text="Kategorien verwalten",bg=WHITE,fg=BLUE,font=body_font(15,weight="bold")).pack(anchor="w",padx=18,pady=(16,8)); tk.Label(frame,text="B2, E3, E4 und Admin dürfen Kategorien lesen und verwalten. Neue Kategorien können direkt beim Erfassen in den Kategorie-Feldern ergänzt werden.",bg=WHITE,fg=TEXT,font=body_font(10),wraplength=max(500,int(w*0.85)),justify="left").pack(anchor="w",padx=18,pady=(0,12)); lb=tk.Listbox(frame,bg=WHITE,fg=TEXT,font=body_font(12),relief="solid",bd=1); lb.pack(fill="both",expand=True,padx=18,pady=(0,18)); [lb.insert("end",c) for c in self.kb_get_categories()]; self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

FiBuMateApp.kb_current_permission_level = _kbp_permission_level
FiBuMateApp.kb_can_create_or_edit = _kbp_can_edit
FiBuMateApp.kb_can_manage_categories = _kbp_can_manage_categories
FiBuMateApp.draw_kb_button = _kbp_draw_button
FiBuMateApp.kb_handle_escape = _kbp_handle_escape
FiBuMateApp.kb_export_dir = _kbp_export_dir
FiBuMateApp.kb_attachment_dir = _kbp_attachment_dir
FiBuMateApp.kb_add_pending_attachment = _kbp_add_pending_attachment
FiBuMateApp.kb_add_comment_to_selected = _kbp_add_comment_to_selected
FiBuMateApp.kb_save_entry_from_form = _kbp_save_entry_from_form
FiBuMateApp.kb_confirm_unsaved_before_switch = _kbp_confirm_unsaved_before_switch
FiBuMateApp.kb_switch_view_from_start = _kbp_switch_view_from_start
FiBuMateApp.kb_export_selected_to_word = _kbp_export_word
FiBuMateApp.render_knowledge_work_area = _kbp_render_work
FiBuMateApp.render_kb_new_entry_area = _kbp_render_new
FiBuMateApp.render_kb_detail_area = _kbp_render_detail
FiBuMateApp.render_kb_categories_area = _kbp_render_categories


# ------------------------------------------------------------------
# Wissenszentrale - Save Fix Patch 2026-06-12
# Zweck: Speichern darf nicht mehr still abbrechen. Die Save-Logik ist
# unabhängig von UI-/Berechtigungs-Erkennung robust ausführbar.
# ------------------------------------------------------------------
def _kbp2_safe_message(level, title, msg):
    try:
        if level == "error": messagebox.showerror(title, msg)
        elif level == "warning": messagebox.showwarning(title, msg)
        else: messagebox.showinfo(title, msg)
    except Exception:
        pass

def _kbp2_save_entries(self, entries):
    last_error = None
    targets = []
    try:
        targets.append(self.kb_entries_path())
    except Exception as exc:
        last_error = exc
    try:
        targets.append(os.path.join(os.getcwd(), "knowledge_entries.json"))
    except Exception:
        pass
    seen = set()
    for path in targets:
        if not path or path in seen:
            continue
        seen.add(path)
        try:
            folder = os.path.dirname(path)
            if folder:
                os.makedirs(folder, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"entries": entries}, f, ensure_ascii=False, indent=2)
            self.knowledge_last_save_path = path
            return True
        except Exception as exc:
            last_error = exc
    _kbp2_safe_message("error", "Wissenszentrale", "Eintrag konnte nicht gespeichert werden.\n\nTechnische Ursache:\n" + str(last_error))
    return False

def _kbp2_save_entry_from_form(self):
    """Robuste finale Speicherlogik für neue und bearbeitete Wissenseinträge."""
    try:
        self.kb_ensure_state_vars()
    except Exception:
        pass
    try:
        title = (getattr(self, "kb_title_var", tk.StringVar(value="")).get() or "").strip()
    except Exception:
        title = ""
    if not title:
        _kbp2_safe_message("warning", "Wissenszentrale", "Bitte einen Titel erfassen. Ohne Titel wird kein Wissenseintrag gespeichert.")
        return False

    categories = []
    try:
        for var in getattr(self, "kb_entry_category_vars", []) or []:
            value = (var.get() or "").strip()
            if value and value not in categories:
                categories.append(value)
    except Exception:
        categories = []
    categories = categories[:4]

    try:
        user = (getattr(self, "kb_user_var", tk.StringVar(value="")).get() or "").strip()
    except Exception:
        user = ""
    if not user:
        user = getattr(self, "current_user_display", "") or getattr(self, "current_user_key", "") or "Unbekannt"

    try:
        status = (getattr(self, "kb_status_var", tk.StringVar(value="Aktiv")).get() or "Aktiv").strip()
    except Exception:
        status = "Aktiv"

    try:
        rhythm_raw = (getattr(self, "kb_rhythm_var", tk.StringVar(value="")).get() or "").strip()
    except Exception:
        rhythm_raw = ""
    rhythm = rhythm_raw if any(str(c).strip().lower() == "to-do" for c in categories) else ""

    try:
        text_value = self.kb_text_widget.get("1.0", "end-1c")
    except Exception:
        text_value = getattr(self, "kb_text_initial", "") or ""

    try:
        entries = self.kb_load_entries()
        if not isinstance(entries, list):
            entries = []
    except Exception:
        entries = []

    now = self.kb_now()
    edit_id = getattr(self, "kb_edit_entry_id", None)
    selected_id = edit_id or self.kb_make_entry_id()
    found = False
    if edit_id:
        for entry in entries:
            if entry.get("id") == edit_id:
                entry.update({
                    "title": title, "categories": categories, "user": user, "status": status,
                    "rhythm": rhythm, "text": text_value, "updated_at": now
                })
                entry.setdefault("comments", [])
                entry.setdefault("attachments", [])
                found = True
                break
    if not found:
        entries.append({
            "id": selected_id, "title": title, "categories": categories, "user": user,
            "status": status, "rhythm": rhythm, "text": text_value,
            "created_at": now, "updated_at": now, "comments": [], "attachments": []
        })

    # Anhänge übernehmen, aber Speichern nie daran scheitern lassen.
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    if pending:
        try:
            import shutil as _shutil
            target_dir = self.kb_attachment_dir(selected_id)
            copied = []
            for src in pending:
                try:
                    name = os.path.basename(src)
                    dst = os.path.join(target_dir, name)
                    _shutil.copy2(src, dst)
                    copied.append({"name": name, "path": dst, "added_at": now})
                except Exception:
                    pass
            if copied:
                for entry in entries:
                    if entry.get("id") == selected_id:
                        entry.setdefault("attachments", []).extend(copied)
                        break
        except Exception:
            pass

    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = selected_id
        self.kb_edit_entry_id = None
        self.kb_pending_attachments = []
        self.knowledge_unsaved = False
        self.knowledge_view = "detail"
        try:
            self.render_page()
        except Exception:
            pass
        _kbp2_safe_message("info", "Wissenszentrale", "Eintrag wurde gespeichert.")
        return True
    return False

def _kbp2_prepare_new_entry(self, entry=None):
    self.kb_ensure_state_vars()
    self.kb_edit_entry_id = entry.get("id") if entry else None
    self.kb_title_var = tk.StringVar(value=(entry or {}).get("title", ""))
    cats = list((entry or {}).get("categories", []) or [])[:4]
    self.kb_entry_category_vars = [tk.StringVar(value=(cats[i] if i < len(cats) else "")) for i in range(4)]
    self.kb_user_var = tk.StringVar(value=(entry or {}).get("user", self.current_user_display or self.current_user_key or ""))
    self.kb_status_var = tk.StringVar(value=(entry or {}).get("status", "Aktiv"))
    self.kb_rhythm_var = tk.StringVar(value=(entry or {}).get("rhythm", ""))
    self.kb_text_initial = (entry or {}).get("text", "")
    self.kb_pending_attachments = []
    self.knowledge_unsaved = bool(entry)

def _kbp2_switch_view_from_start(self, view):
    if not self.kb_confirm_unsaved_before_switch():
        return
    if view == "new":
        self.kb_prepare_new_entry()
    self.knowledge_view = view
    self.knowledge_start_overlay = False
    self.render_page()

# Wichtig: Diese Zuweisungen stehen direkt vor dem Programmstart und überschreiben alle früheren Varianten.
FiBuMateApp.kb_save_entries = _kbp2_save_entries
FiBuMateApp.kb_save_entry_from_form = _kbp2_save_entry_from_form
FiBuMateApp.kb_prepare_new_entry = _kbp2_prepare_new_entry
FiBuMateApp.kb_switch_view_from_start = _kbp2_switch_view_from_start


# ------------------------------------------------------------------
# Wissenszentrale - Visible Buttons + Fullheight Fix 2026-06-12
# ------------------------------------------------------------------
def _wz_btn_full_render_work(self):
    self.kb_ensure_state_vars()
    try:
        self.root.option_add("*TCombobox*Listbox.font", body_font(12))
        self.root.bind("<Escape>", self.kb_handle_escape, add="+")
    except Exception:
        pass
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()

    nav_y = 145
    x = 28
    x = self.draw_kb_button(x, nav_y, "Start", self.kb_show_start_overlay, False, width=92)
    x = self.draw_kb_button(x, nav_y, "Übersicht", lambda: self.kb_switch_view_from_start("all"), False, width=112)
    if getattr(self, "kb_can_create_or_edit", lambda: True)():
        x = self.draw_kb_button(x, nav_y, "Neuer Eintrag", lambda: self.kb_switch_view_from_start("new"), False, width=128)
    x = self.draw_kb_button(x, nav_y, "To-Dos", lambda: self.kb_switch_view_from_start("todos"), False, width=92)
    x = self.draw_kb_button(x, nav_y, "Veraltete Einträge", lambda: self.kb_switch_view_from_start("outdated"), False, width=148)
    if getattr(self, "kb_can_manage_categories", lambda: True)():
        x = self.draw_kb_button(x, nav_y, "Kategorien verwalten", lambda: self.kb_switch_view_from_start("categories"), False, width=168)

    search_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(search_frame)
    tk.Label(search_frame, text="Suche", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w")
    search_entry = tk.Entry(search_frame, textvariable=self.kb_search_var, font=body_font(11), bg=WHITE, fg=TEXT, relief="solid", bd=1)
    search_entry.grid(row=1, column=0, sticky="ew", ipady=5)
    search_entry.bind("<Return>", self.kb_on_search_return)
    tk.Button(search_frame, text="Suchen", command=self.kb_apply_filters, bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10)).grid(row=1, column=1, padx=(8, 0), sticky="ns")
    search_frame.grid_columnconfigure(0, weight=1)
    self.canvas.create_window(ui_s(28), ui_s(192), window=search_frame, anchor="nw", width=ui_s(min(790, max(560, w-80))), height=ui_s(60))

    filter_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(filter_frame)
    values = [""] + self.kb_get_categories()
    tk.Label(filter_frame, text="Kategorie-Filter", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w", columnspan=6)
    for idx, var in enumerate(self.kb_filter_vars):
        block = tk.Frame(filter_frame, bg=BG)
        block.grid(row=1, column=idx, padx=(0 if idx == 0 else 12, 0), sticky="w")
        tk.Label(block, text=f"Kategorie {idx+1}", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        cb = ttk.Combobox(block, textvariable=var, values=values, state="readonly", font=body_font(11), width=18)
        cb.pack(anchor="w")
        cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    selected_filter_cats = [(v.get() or "").strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view == "todos" or "to-do" in selected_filter_cats:
        block = tk.Frame(filter_frame, bg=BG)
        block.grid(row=1, column=4, padx=(12, 0), sticky="w")
        tk.Label(block, text="To-Do-Rhythmus", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        cb = ttk.Combobox(block, textvariable=self.kb_todo_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(11), width=16)
        cb.pack(anchor="w")
        cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    else:
        self.kb_todo_rhythm_var.set("")
        tk.Label(filter_frame, text="Leere Filter = keine Filterung", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=1, column=4, padx=(14, 0), sticky="s")
    self.canvas.create_window(ui_s(28), ui_s(252), window=filter_frame, anchor="nw", width=ui_s(min(1050, max(760, w-70))), height=ui_s(90))

    # Position bleibt wie bisher/Screenshot; nur die Höhe wird bis kurz vor Fußleiste verlängert.
    left_x, left_y = 20, 264
    left_w = max(340, min(460, int(w * 0.31)))
    pane_h = max(430, h - left_y - 28)
    right_x, right_y = left_x + left_w + 24, left_y
    right_w = max(760, w - right_x - 24)

    self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
    if self.knowledge_view == "new":
        self.render_kb_new_entry_area(right_x, right_y, right_w, pane_h)
    elif self.knowledge_view == "todos":
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="To-Dos")
    elif self.knowledge_view == "outdated":
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Veraltete Einträge", status_filter="Veraltet")
    elif self.knowledge_view == "categories":
        self.render_kb_categories_area(right_x, right_y, right_w, pane_h)
    elif self.knowledge_view == "detail":
        self.render_kb_detail_area(right_x, right_y, right_w, pane_h)
    else:
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Übersicht")

def _wz_btn_full_render_new_entry(self, x, y, w, h):
    if not hasattr(self, "kb_title_var"):
        self.kb_prepare_new_entry()
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)

    tk.Label(frame, text=("Eintrag bearbeiten" if getattr(self, "kb_edit_entry_id", None) else "Neuer Eintrag"), bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(14, 6))

    # Erst unten andocken: Dadurch ist die Buttonleiste immer sichtbar.
    btnrow = tk.Frame(frame, bg=WHITE)
    btnrow.pack(side="bottom", fill="x", padx=18, pady=(8, 14))
    tk.Button(btnrow, text="Anhang hinzufügen", command=self.kb_add_pending_attachment, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(btnrow, text="Als Entwurf speichern", command=lambda: (self.kb_status_var.set("Entwurf"), self.kb_save_entry_from_form()), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(btnrow, text="Speichern", command=self.kb_save_entry_from_form, bg=WHITE, fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(btnrow, text="Abbrechen", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)

    form = tk.Frame(frame, bg=WHITE)
    form.pack(side="top", fill="x", padx=18, pady=(0, 8))
    tk.Label(form, text="Titel des Eintrags", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=1, column=0, columnspan=4, sticky="ew", ipady=5, pady=(0, 8))
    categories = [""] + self.kb_get_categories()
    for idx, var in enumerate(self.kb_entry_category_vars):
        tk.Label(form, text=f"Kategorie {idx+1}", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=idx, sticky="w", padx=(0 if idx == 0 else 8, 0))
        cb = ttk.Combobox(form, textvariable=var, values=categories, state="normal", font=body_font(11), width=18)
        cb.grid(row=3, column=idx, sticky="ew", padx=(0 if idx == 0 else 8, 0), pady=(0, 8))
        cb.bind("<<ComboboxSelected>>", lambda e: self.render_page())
    tk.Label(form, text="Assoziierter Benutzer", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_user_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=5, column=0, sticky="ew", pady=(0, 8))
    tk.Label(form, text="Status", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=1, sticky="w", padx=(8, 0))
    ttk.Combobox(form, textvariable=self.kb_status_var, values=["Aktiv", "Entwurf", "Veraltet"], state="readonly", font=body_font(11)).grid(row=5, column=1, sticky="ew", padx=(8, 0), pady=(0, 8))
    if any((v.get() or "").strip().lower() == "to-do" for v in self.kb_entry_category_vars):
        tk.Label(form, text="To-Do-Rhythmus", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=2, sticky="w", padx=(8, 0))
        ttk.Combobox(form, textvariable=self.kb_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(11)).grid(row=5, column=2, sticky="ew", padx=(8, 0), pady=(0, 8))
    else:
        self.kb_rhythm_var.set("")
    for col in range(4):
        form.grid_columnconfigure(col, weight=1)

    pending = getattr(self, "kb_pending_attachments", []) or []
    if pending:
        tk.Label(frame, text="Vorgemerkte Anhänge: " + ", ".join([os.path.basename(x) for x in pending]), bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="bottom", anchor="w", padx=18, pady=(0, 2))

    body_frame = tk.Frame(frame, bg=WHITE)
    body_frame.pack(side="top", fill="both", expand=True, padx=18, pady=(0, 0))
    tk.Label(body_frame, text="Freitext / Prozessdokumentation / Leitfaden", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    self.kb_text_widget = tk.Text(body_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word")
    self.kb_text_widget.pack(fill="both", expand=True)
    self.kb_text_widget.insert("1.0", getattr(self, "kb_text_initial", ""))
    self.kb_text_widget.bind("<KeyRelease>", lambda e: self.kb_mark_unsaved())

    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

FiBuMateApp.render_knowledge_work_area = _wz_btn_full_render_work
FiBuMateApp.render_kb_new_entry_area = _wz_btn_full_render_new_entry


# ------------------------------------------------------------------
# Wissenszentrale - Position + G-Laufwerk + Renderprüfung Fix 2026-06-12
# Zweck:
# - Hauptfenster sichtbar unterhalb Suche/Kategorie-Filter positionieren.
# - Buttonleiste bleibt sichtbar.
# - Wissenszentrale-Daten werden ausschließlich auf dem Firmenlaufwerk G: gespeichert.
# ------------------------------------------------------------------
def _wz_gdrive_base_dir():
    return os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale")

def _wz_gdrive_config_dir():
    return os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Config")

def _wz_gdrive_require_dir(path):
    os.makedirs(path, exist_ok=True)
    return path

def _wz_gdrive_entries_path(self):
    base = _wz_gdrive_require_dir(_wz_gdrive_base_dir())
    return os.path.join(base, "knowledge_entries.json")

def _wz_gdrive_pref_path(self):
    base = _wz_gdrive_require_dir(_wz_gdrive_config_dir())
    return os.path.join(base, "knowledge_start_pref.json")

def _wz_gdrive_attachment_dir(self, entry_id):
    base = _wz_gdrive_require_dir(os.path.join(_wz_gdrive_base_dir(), "Attachments", str(entry_id or "ohne_id")))
    return base

def _wz_gdrive_save_entries(self, entries):
    try:
        path = self.kb_entries_path()
        # Strikte Firmenlaufwerk-Prüfung: kein lokaler Fallback.
        if not os.path.abspath(path).upper().startswith(os.path.abspath(NETWORK_ROOT).upper()):
            raise RuntimeError("Ungültiger Speicherpfad außerhalb des Firmenlaufwerks: " + str(path))
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"entries": entries}, f, ensure_ascii=False, indent=2)
        self.knowledge_last_save_path = path
        return True
    except Exception as exc:
        try:
            messagebox.showerror(
                "Wissenszentrale",
                "Eintrag konnte nicht auf dem Firmenlaufwerk gespeichert werden.\n\n"
                "Es wird bewusst kein lokaler Fallback verwendet.\n\n"
                "Zielpfad: " + os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale") + "\n\n"
                "Technische Ursache:\n" + str(exc)
            )
        except Exception:
            pass
        return False

def _wz_position_render_work(self):
    self.kb_ensure_state_vars()
    try:
        self.root.option_add("*TCombobox*Listbox.font", body_font(12))
        self.root.bind("<Escape>", self.kb_handle_escape, add="+")
    except Exception:
        pass
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()

    nav_y = 145
    x = 28
    x = self.draw_kb_button(x, nav_y, "Start", self.kb_show_start_overlay, False, width=92)
    x = self.draw_kb_button(x, nav_y, "Übersicht", lambda: self.kb_switch_view_from_start("all"), False, width=112)
    if getattr(self, "kb_can_create_or_edit", lambda: True)():
        x = self.draw_kb_button(x, nav_y, "Neuer Eintrag", lambda: self.kb_switch_view_from_start("new"), False, width=128)
    x = self.draw_kb_button(x, nav_y, "To-Dos", lambda: self.kb_switch_view_from_start("todos"), False, width=92)
    x = self.draw_kb_button(x, nav_y, "Veraltete Einträge", lambda: self.kb_switch_view_from_start("outdated"), False, width=148)
    if getattr(self, "kb_can_manage_categories", lambda: True)():
        x = self.draw_kb_button(x, nav_y, "Kategorien verwalten", lambda: self.kb_switch_view_from_start("categories"), False, width=168)

    search_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(search_frame)
    tk.Label(search_frame, text="Suche", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w")
    search_entry = tk.Entry(search_frame, textvariable=self.kb_search_var, font=body_font(11), bg=WHITE, fg=TEXT, relief="solid", bd=1)
    search_entry.grid(row=1, column=0, sticky="ew", ipady=5)
    search_entry.bind("<Return>", self.kb_on_search_return)
    tk.Button(search_frame, text="Suchen", command=self.kb_apply_filters, bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10)).grid(row=1, column=1, padx=(8, 0), sticky="ns")
    search_frame.grid_columnconfigure(0, weight=1)
    self.canvas.create_window(ui_s(28), ui_s(192), window=search_frame, anchor="nw", width=ui_s(min(790, max(560, w-80))), height=ui_s(60))

    filter_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(filter_frame)
    values = [""] + self.kb_get_categories()
    tk.Label(filter_frame, text="Kategorie-Filter", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w", columnspan=6)
    for idx, var in enumerate(self.kb_filter_vars):
        block = tk.Frame(filter_frame, bg=BG)
        block.grid(row=1, column=idx, padx=(0 if idx == 0 else 12, 0), sticky="w")
        tk.Label(block, text=f"Kategorie {idx+1}", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        cb = ttk.Combobox(block, textvariable=var, values=values, state="readonly", font=body_font(11), width=18)
        cb.pack(anchor="w")
        cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    selected_filter_cats = [(v.get() or "").strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view == "todos" or "to-do" in selected_filter_cats:
        block = tk.Frame(filter_frame, bg=BG)
        block.grid(row=1, column=4, padx=(12, 0), sticky="w")
        tk.Label(block, text="To-Do-Rhythmus", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        cb = ttk.Combobox(block, textvariable=self.kb_todo_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(11), width=16)
        cb.pack(anchor="w")
        cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    else:
        self.kb_todo_rhythm_var.set("")
        tk.Label(filter_frame, text="Leere Filter = keine Filterung", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=1, column=4, padx=(14, 0), sticky="s")
    self.canvas.create_window(ui_s(28), ui_s(252), window=filter_frame, anchor="nw", width=ui_s(min(1050, max(760, w-70))), height=ui_s(90))

    # Wichtig: bewusst tiefer als vorher. Bei der typischen automatischen Skalierung landet die Oberkante
    # sichtbar unterhalb des Kategorie-Filters und überlappt Suche/Filter nicht mehr.
    left_x, left_y = 20, 350
    left_w = max(340, min(460, int(w * 0.31)))
    pane_h = max(360, h - left_y - 28)
    right_x, right_y = left_x + left_w + 24, left_y
    right_w = max(760, w - right_x - 24)

    self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
    if self.knowledge_view == "new":
        self.render_kb_new_entry_area(right_x, right_y, right_w, pane_h)
    elif self.knowledge_view == "todos":
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="To-Dos")
    elif self.knowledge_view == "outdated":
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Veraltete Einträge", status_filter="Veraltet")
    elif self.knowledge_view == "categories":
        self.render_kb_categories_area(right_x, right_y, right_w, pane_h)
    elif self.knowledge_view == "detail":
        self.render_kb_detail_area(right_x, right_y, right_w, pane_h)
    else:
        self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Übersicht")

# Letzte Zuweisung vor App-Start: überschreibt ältere lokale Save-Fallbacks und ältere Y-Positionen.
FiBuMateApp.kb_entries_path = _wz_gdrive_entries_path
FiBuMateApp.knowledge_pref_path = _wz_gdrive_pref_path
FiBuMateApp.kb_attachment_dir = _wz_gdrive_attachment_dir
FiBuMateApp.kb_save_entries = _wz_gdrive_save_entries
FiBuMateApp.render_knowledge_work_area = _wz_position_render_work


# ------------------------------------------------------------------
# Wissenszentrale - G-Laufwerk mit lokalem Fallback C:\python 2026-06-12
# Zweck:
# - Primärspeicherung weiterhin auf G:\BUC\FM Anwendung.
# - Wenn G: nicht verfügbar/beschreibbar ist, lokaler Fallback nach C:\python\knowledge_entries.json.
# - Anhänge erhalten analog einen lokalen Fallback unter C:\python\knowledge_attachments.
# ------------------------------------------------------------------
def _wz_fallback_local_entries_path():
    return r"C:\python\knowledge_entries.json"

def _wz_fallback_local_base_dir():
    return r"C:\python"

def _wz_fallback_local_attachment_dir(entry_id):
    return os.path.join(_wz_fallback_local_base_dir(), "knowledge_attachments", str(entry_id or "ohne_id"))

def _wz_gdrive_available_path(path):
    try:
        if not os.path.abspath(path).upper().startswith(os.path.abspath(NETWORK_ROOT).upper()):
            return False
        os.makedirs(os.path.dirname(path), exist_ok=True)
        return True
    except Exception:
        return False

def _wz_entries_path_gdrive_with_local_fallback(self):
    g_path = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale", "knowledge_entries.json")
    if _wz_gdrive_available_path(g_path):
        return g_path
    try:
        os.makedirs(os.path.dirname(_wz_fallback_local_entries_path()), exist_ok=True)
    except Exception:
        pass
    return _wz_fallback_local_entries_path()

def _wz_attachment_dir_gdrive_with_local_fallback(self, entry_id):
    g_dir = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale", "Attachments", str(entry_id or "ohne_id"))
    try:
        os.makedirs(g_dir, exist_ok=True)
        return g_dir
    except Exception:
        local_dir = _wz_fallback_local_attachment_dir(entry_id)
        os.makedirs(local_dir, exist_ok=True)
        return local_dir

def _wz_save_entries_gdrive_with_local_fallback(self, entries):
    errors = []
    targets = [
        os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale", "knowledge_entries.json"),
        _wz_fallback_local_entries_path(),
    ]
    for path in targets:
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"entries": entries}, f, ensure_ascii=False, indent=2)
            self.knowledge_last_save_path = path
            self.knowledge_used_local_fallback = (path == _wz_fallback_local_entries_path())
            if self.knowledge_used_local_fallback:
                try:
                    messagebox.showwarning(
                        "Wissenszentrale",
                        "G:-Firmenlaufwerk ist nicht verfügbar oder nicht beschreibbar.\n\n"
                        "Der Eintrag wurde lokal gespeichert unter:\n"
                        + _wz_fallback_local_entries_path()
                    )
                except Exception:
                    pass
            return True
        except Exception as exc:
            errors.append(f"{path}: {exc}")
    try:
        messagebox.showerror(
            "Wissenszentrale",
            "Eintrag konnte weder auf G: noch lokal gespeichert werden.\n\n" + "\n".join(errors)
        )
    except Exception:
        pass
    return False

def _wz_load_entries_gdrive_with_local_fallback(self):
    paths = [
        os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale", "knowledge_entries.json"),
        _wz_fallback_local_entries_path(),
    ]
    for path in paths:
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                entries = data.get("entries", []) if isinstance(data, dict) else []
                self.knowledge_last_load_path = path
                return entries if isinstance(entries, list) else []
        except Exception:
            pass
    return []

# Finale Zuweisung: G: zuerst, C:\python\knowledge_entries.json als Fallback.
FiBuMateApp.kb_entries_path = _wz_entries_path_gdrive_with_local_fallback
FiBuMateApp.kb_attachment_dir = _wz_attachment_dir_gdrive_with_local_fallback
FiBuMateApp.kb_save_entries = _wz_save_entries_gdrive_with_local_fallback
FiBuMateApp.kb_load_entries = _wz_load_entries_gdrive_with_local_fallback


# ------------------------------------------------------------------
# Wissenszentrale - Kategorien erstellen + Farblogik FINAL 2026-06-12
# ------------------------------------------------------------------
_WZ_CAT_PALETTE = ["#2563EB", "#059669", "#DC2626", "#7C3AED", "#D97706", "#0891B2", "#BE123C", "#4F46E5", "#15803D", "#B45309", "#0F766E", "#9333EA"]
_WZ_CAT_DEFAULTS = {"To-Do": "#DC2626", "Aktiv": "#059669", "Veraltet": "#6B7280"}

def _wz_cat_norm(name):
    return " ".join(str(name or "").strip().split())

def _wz_cat_hex(value, fallback="#2563EB"):
    value = str(value or "").strip()
    if re.fullmatch(r"#[0-9a-fA-F]{6}", value): return value.upper()
    if re.fullmatch(r"[0-9a-fA-F]{6}", value): return ("#" + value).upper()
    return fallback

def _wz_cat_default_color(name):
    name = _wz_cat_norm(name)
    if name in _WZ_CAT_DEFAULTS:
        return _WZ_CAT_DEFAULTS[name]
    idx = sum(ord(ch) for ch in name.lower()) % len(_WZ_CAT_PALETTE) if name else 0
    return _WZ_CAT_PALETTE[idx]

def _wz_cat_rgb(color):
    color = _wz_cat_hex(color)
    return tuple(int(color[i:i+2], 16) for i in (1, 3, 5))

def _wz_cat_fg(color):
    try:
        r, g, b = _wz_cat_rgb(color)
        return "#111827" if (0.299*r + 0.587*g + 0.114*b) > 158 else "#FFFFFF"
    except Exception:
        return "#FFFFFF"

def _wz_cat_light(color, factor=0.86):
    try:
        r, g, b = _wz_cat_rgb(color)
        r = int(r + (255-r)*factor); g = int(g + (255-g)*factor); b = int(b + (255-b)*factor)
        return f"#{r:02X}{g:02X}{b:02X}"
    except Exception:
        return "#F3F4F6"

def _wz_cat_paths():
    return [
        os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale", "knowledge_categories.json"),
        r"C:\python\knowledge_categories.json",
    ]

def _wz_cat_load(self):
    for path in _wz_cat_paths():
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                raw = data.get("categories", {}) if isinstance(data, dict) else {}
                if isinstance(raw, list):
                    raw = {x.get("name", ""): x for x in raw if isinstance(x, dict)}
                result = {}
                if isinstance(raw, dict):
                    for key, item in raw.items():
                        name = _wz_cat_norm(item.get("name", key) if isinstance(item, dict) else key)
                        if not name: continue
                        color = _wz_cat_hex(item.get("color") if isinstance(item, dict) else None, _wz_cat_default_color(name))
                        result[name] = {"name": name, "color": color}
                return result
        except Exception:
            pass
    return {}

def _wz_cat_save(self, meta):
    clean = {}
    for key, item in (meta or {}).items():
        name = _wz_cat_norm(item.get("name", key) if isinstance(item, dict) else key)
        if not name: continue
        clean[name] = {"name": name, "color": _wz_cat_hex(item.get("color") if isinstance(item, dict) else None, _wz_cat_default_color(name))}
    errors = []
    for path in _wz_cat_paths():
        try:
            folder = os.path.dirname(path)
            if folder: os.makedirs(folder, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"categories": clean}, f, ensure_ascii=False, indent=2)
            self.knowledge_category_last_save_path = path
            return True
        except Exception as exc:
            errors.append(f"{path}: {exc}")
    try: messagebox.showerror("Wissenszentrale", "Kategorien konnten nicht gespeichert werden:\n" + "\n".join(errors))
    except Exception: pass
    return False

def _wz_cat_get_categories(self):
    cats = set(getattr(self, "kb_default_categories", lambda: ["To-Do"])())
    cats.update(_wz_cat_load(self).keys())
    try:
        for entry in self.kb_load_entries():
            for cat in entry.get("categories", []) or []:
                cat = _wz_cat_norm(cat)
                if cat: cats.add(cat)
    except Exception:
        pass
    return sorted(cats, key=lambda c: c.lower())

def _wz_cat_color(self, name):
    name = _wz_cat_norm(name)
    meta = _wz_cat_load(self)
    return _wz_cat_hex(meta.get(name, {}).get("color"), _wz_cat_default_color(name))

def _wz_cat_rename_entries(self, old, new):
    old = _wz_cat_norm(old); new = _wz_cat_norm(new)
    if not old or not new or old == new: return True
    entries = self.kb_load_entries(); changed = False
    for entry in entries:
        out = []; seen = set()
        for cat in entry.get("categories", []) or []:
            c = _wz_cat_norm(cat)
            if c == old:
                c = new; changed = True
            if c and c.lower() not in seen:
                out.append(c); seen.add(c.lower())
        entry["categories"] = out
    return self.kb_save_entries(entries) if changed else True

def _wz_cat_draw_badge(self, x, y, text, color=None):
    color = _wz_cat_hex(color or _wz_cat_color(self, text), _wz_cat_default_color(text))
    fg = _wz_cat_fg(color)
    width = max(54, 20 + len(str(text)) * 8)
    self.canvas.create_rectangle(ui_s(x), ui_s(y), ui_s(x+width), ui_s(y+24), fill=color, outline=color, width=1)
    self.canvas.create_text(ui_s(x+10), ui_s(y+12), text=text, font=body_font(9, weight="bold"), fill=fg, anchor="w")
    return x + width + 8

def _wz_cat_render_area(self, x, y, w, h):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text="Kategorien verwalten", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(14, 6))
    tk.Label(frame, text="Kategorien können hier erstellt, umbenannt und farblich gepflegt werden. Die angewählte Kategorie wird mit ihrer Farbe hervorgehoben.", bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520, int(w*0.86)), justify="left").pack(anchor="w", padx=18, pady=(0, 10))
    main = tk.Frame(frame, bg=WHITE); main.pack(fill="both", expand=True, padx=18, pady=(0, 14))
    main.grid_columnconfigure(0, weight=1); main.grid_columnconfigure(1, weight=1); main.grid_rowconfigure(0, weight=1)
    left = tk.Frame(main, bg=WHITE); left.grid(row=0, column=0, sticky="nsew", padx=(0, 12))
    right = tk.Frame(main, bg=WHITE); right.grid(row=0, column=1, sticky="nsew", padx=(12, 0))
    tk.Label(left, text="Vorhandene Kategorien", bg=WHITE, fg=TEXT2, font=body_font(10, weight="bold")).pack(anchor="w")
    lb = tk.Listbox(left, bg=WHITE, fg=TEXT, font=body_font(11), relief="solid", bd=1, exportselection=False)
    lb.pack(fill="both", expand=True, pady=(6, 0))
    name_var = tk.StringVar(value=""); color_var = tk.StringVar(value="#2563EB"); info_var = tk.StringVar(value="Keine Kategorie ausgewählt")
    tk.Label(right, text="Kategorie erstellen / bearbeiten", bg=WHITE, fg=TEXT2, font=body_font(10, weight="bold")).pack(anchor="w")
    tk.Label(right, textvariable=info_var, bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w", pady=(3, 8))
    form = tk.Frame(right, bg=WHITE); form.pack(fill="x")
    tk.Label(form, text="Name", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
    name_entry = tk.Entry(form, textvariable=name_var, bg=WHITE, fg=TEXT, font=body_font(11), relief="solid", bd=1); name_entry.grid(row=1, column=0, sticky="ew", ipady=5, pady=(0, 8))
    tk.Label(form, text="Farbe (#RRGGBB)", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=0, sticky="w")
    color_entry = tk.Entry(form, textvariable=color_var, bg=WHITE, fg=TEXT, font=body_font(11), relief="solid", bd=1); color_entry.grid(row=3, column=0, sticky="ew", ipady=5, pady=(0, 8))
    form.grid_columnconfigure(0, weight=1)
    preview = tk.Label(right, text="Farbvorschau", bg="#2563EB", fg="#FFFFFF", font=body_font(10, weight="bold"), relief="solid", bd=1); preview.pack(fill="x", ipady=8, pady=(0, 10))
    swatch_area = tk.Frame(right, bg=WHITE); swatch_area.pack(fill="x", pady=(0, 10))
    tk.Label(swatch_area, text="Schnellfarben", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    swatches = tk.Frame(swatch_area, bg=WHITE); swatches.pack(anchor="w", pady=(4, 0))
    def apply_preview(*_):
        c = _wz_cat_hex(color_var.get(), _wz_cat_default_color(name_var.get() or "Kategorie"))
        preview.configure(bg=c, fg=_wz_cat_fg(c), text=f"Farbvorschau {c}")
    def set_color(c):
        color_var.set(_wz_cat_hex(c)); apply_preview()
    for c in _WZ_CAT_PALETTE:
        tk.Button(swatches, text="", command=lambda cc=c: set_color(cc), bg=c, activebackground=c, relief="solid", bd=1, width=3).pack(side="left", padx=(0, 4), ipady=6)
    def choose_color():
        try:
            from tkinter import colorchooser
            chosen = colorchooser.askcolor(color=_wz_cat_hex(color_var.get()), title="Kategorie-Farbe wählen")
            if chosen and chosen[1]: set_color(chosen[1])
        except Exception as exc:
            try: messagebox.showwarning("Wissenszentrale", f"Farbauswahl konnte nicht geöffnet werden:\n{exc}")
            except Exception: pass
    def refresh(selected=None):
        selected = _wz_cat_norm(selected or getattr(self, "kb_selected_category_name", ""))
        lb.delete(0, "end"); selected_idx = None
        for idx, cat in enumerate(self.kb_get_categories()):
            color = _wz_cat_color(self, cat); lb.insert("end", cat)
            try:
                if cat == selected:
                    lb.itemconfig(idx, bg=color, fg=_wz_cat_fg(color), selectbackground=color, selectforeground=_wz_cat_fg(color)); selected_idx = idx
                else:
                    lb.itemconfig(idx, bg=_wz_cat_light(color), fg=TEXT, selectbackground=color, selectforeground=_wz_cat_fg(color))
            except Exception: pass
        if selected_idx is not None:
            lb.selection_clear(0, "end"); lb.selection_set(selected_idx); lb.see(selected_idx)
    def select_cat(event=None):
        sel = lb.curselection()
        if not sel: return
        cat = lb.get(sel[0]); self.kb_selected_category_name = cat
        name_var.set(cat); color_var.set(_wz_cat_color(self, cat)); info_var.set(f"Ausgewählt: {cat}")
        apply_preview(); refresh(cat)
    lb.bind("<<ListboxSelect>>", select_cat)
    def create_cat():
        name = _wz_cat_norm(name_var.get())
        if not name:
            messagebox.showwarning("Wissenszentrale", "Bitte einen Kategorienamen eingeben."); return
        if name.lower() in {c.lower() for c in self.kb_get_categories()}:
            messagebox.showwarning("Wissenszentrale", "Diese Kategorie existiert bereits. Bitte 'Änderungen speichern' verwenden."); return
        meta = _wz_cat_load(self); meta[name] = {"name": name, "color": _wz_cat_hex(color_var.get(), _wz_cat_default_color(name))}
        if _wz_cat_save(self, meta):
            self.kb_selected_category_name = name; info_var.set(f"Neu erstellt: {name}"); refresh(name); self.render_page()
    def save_cat():
        old = _wz_cat_norm(getattr(self, "kb_selected_category_name", "")); new = _wz_cat_norm(name_var.get())
        if not new:
            messagebox.showwarning("Wissenszentrale", "Bitte einen Kategorienamen eingeben."); return
        meta = _wz_cat_load(self)
        if old and old != new:
            if old.lower() == "to-do":
                messagebox.showwarning("Wissenszentrale", "Die Sonderkategorie 'To-Do' darf nicht umbenannt werden. Die Farbe kann gespeichert werden."); new = old; name_var.set(old)
            elif new.lower() in {c.lower() for c in self.kb_get_categories() if c.lower() != old.lower()}:
                messagebox.showwarning("Wissenszentrale", "Der neue Kategoriename existiert bereits."); return
            else:
                meta.pop(old, None)
                if not _wz_cat_rename_entries(self, old, new): return
        meta[new] = {"name": new, "color": _wz_cat_hex(color_var.get(), _wz_cat_default_color(new))}
        if _wz_cat_save(self, meta):
            self.kb_selected_category_name = new; info_var.set(f"Gespeichert: {new}"); refresh(new); self.render_page()
    def blank_cat():
        self.kb_selected_category_name = ""; name_var.set(""); color_var.set("#2563EB"); info_var.set("Neue Kategorie erfassen"); apply_preview(); refresh(""); name_entry.focus_set()
    color_var.trace_add("write", apply_preview)
    btns = tk.Frame(right, bg=WHITE); btns.pack(fill="x", pady=(4, 0))
    tk.Button(btns, text="Kategorie erstellen", command=create_cat, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=10, ipady=4)
    tk.Button(btns, text="Änderungen speichern", command=save_cat, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=10, ipady=4)
    tk.Button(btns, text="Farbe wählen", command=choose_color, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=10, ipady=4)
    tk.Button(btns, text="Neu leeren", command=blank_cat, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=10, ipady=4)
    tk.Label(right, text="Beim Umbenennen werden vorhandene Einträge automatisch aktualisiert. Kategorie-Farben werden in knowledge_categories.json gespeichert.", bg=WHITE, fg=TEXT2, font=body_font(9), wraplength=max(340, int(w*0.40)), justify="left").pack(anchor="w", pady=(12, 0))
    apply_preview(); refresh(getattr(self, "kb_selected_category_name", ""))
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

FiBuMateApp.kb_get_categories = _wz_cat_get_categories
FiBuMateApp.kb_load_category_meta = _wz_cat_load
FiBuMateApp.kb_save_category_meta = _wz_cat_save
FiBuMateApp.kb_get_category_color = _wz_cat_color
FiBuMateApp.draw_kb_badge = _wz_cat_draw_badge
FiBuMateApp.render_kb_categories_area = _wz_cat_render_area


# ------------------------------------------------------------------
# Wissenszentrale - Kategorien Render-Safety Fix 2026-06-12
# Behebt: NameError 're' in Farblogik + verhindert, dass Kategorien-Menü bei Renderfehlern verschwindet.
# ------------------------------------------------------------------
try:
    _wz_cat_render_area_core = FiBuMateApp.render_kb_categories_area
except Exception:
    _wz_cat_render_area_core = None

def _wz_cat_render_area_safe(self, x, y, w, h):
    try:
        if _wz_cat_render_area_core:
            return _wz_cat_render_area_core(self, x, y, w, h)
    except Exception as exc:
        try:
            frame = tk.Frame(self.root, bg=WHITE, highlightbackground="#DC2626", highlightthickness=2)
            self.widget_items.append(frame)
            tk.Label(frame, text="Kategorien verwalten", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
            tk.Label(frame, text="Die Kategorienverwaltung konnte nicht vollständig gerendert werden. Die Navigation bleibt aktiv.", bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(500, int(w*0.85)), justify="left").pack(anchor="w", padx=18, pady=(0, 8))
            tk.Label(frame, text=f"Technische Ursache: {exc}", bg=WHITE, fg="#DC2626", font=body_font(9), wraplength=max(500, int(w*0.85)), justify="left").pack(anchor="w", padx=18, pady=(0, 12))
            tk.Button(frame, text="Zur Übersicht", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(anchor="w", padx=18, pady=(0, 12), ipadx=12, ipady=4)
            self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))
        except Exception:
            raise exc

FiBuMateApp.render_kb_categories_area = _wz_cat_render_area_safe



# ------------------------------------------------------------------
# Wissenszentrale - Notfallfix + Editor FINAL 2026-06-18
# Version 0.439
# Zweck:
# - Behebt NameError _wz437_sig / _wz437_cats_g / _wz437_cats_l aus v0.438.
# - Kategorien werden beim Lesen farbig als Badges angezeigt.
# - Fett/Kursiv/Unterstrichen behalten die vorhandene Schriftgröße.
# - Editor: feste Buttonleiste, Scrollbar, Mausrad-Scroll, Bilder vor dem Text.
# - Performance: G:-Verfügbarkeit und Kategorien werden gecacht.
# ------------------------------------------------------------------

_WZ439_G_CACHE = {"value": None, "time": 0.0}
_WZ439_CAT_CACHE = {"sig": None, "data": {}}
_WZ439_REFRESH_MS = 3000
_WZ439_LOCAL_BASE = r"C:\python"

def _wz439_base_dir():
    return os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale")

def _wz439_entries_g():
    return os.path.join(_wz439_base_dir(), "knowledge_entries.json")

def _wz439_entries_l():
    return os.path.join(_WZ439_LOCAL_BASE, "knowledge_entries.json")

def _wz439_cats_g():
    return os.path.join(_wz439_base_dir(), "knowledge_categories.json")

def _wz439_cats_l():
    return os.path.join(_WZ439_LOCAL_BASE, "knowledge_categories.json")

def _wz439_sig(path):
    try:
        if path and os.path.exists(path):
            st = os.stat(path)
            return f"{path}|{st.st_mtime_ns}|{st.st_size}"
    except Exception:
        pass
    return f"{path}|missing"

# Kompatibilitätsnamen für v0.438, damit der NameError sofort behoben ist.
_wz437_sig = _wz439_sig
_wz437_cats_g = _wz439_cats_g
_wz437_cats_l = _wz439_cats_l
_wz437_entries_g = _wz439_entries_g
_wz437_entries_l = _wz439_entries_l

def _wz439_can_g():
    import time
    now = time.time()
    if _WZ439_G_CACHE["value"] is not None and now - _WZ439_G_CACHE["time"] < 10.0:
        return bool(_WZ439_G_CACHE["value"])
    ok = False
    try:
        os.makedirs(_wz439_base_dir(), exist_ok=True)
        probe = os.path.join(_wz439_base_dir(), ".write_test")
        with open(probe, "w", encoding="utf-8") as f:
            f.write("ok")
        try:
            os.remove(probe)
        except Exception:
            pass
        ok = True
    except Exception:
        ok = False
    _WZ439_G_CACHE["value"] = ok
    _WZ439_G_CACHE["time"] = now
    return ok

_wz437_can_g = _wz439_can_g

def _wz439_read_json(path):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _wz439_write_json(path, data):
    folder = os.path.dirname(path)
    if folder:
        os.makedirs(folder, exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)

def _wz439_load_category_meta(self):
    sig = _wz439_sig(_wz439_cats_g()) + "|" + _wz439_sig(_wz439_cats_l())
    if _WZ439_CAT_CACHE.get("sig") == sig:
        return dict(_WZ439_CAT_CACHE.get("data") or {})
    result = {}
    for path in (_wz439_cats_g(), _wz439_cats_l()):
        data = _wz439_read_json(path)
        raw = data.get("categories", {}) if isinstance(data, dict) else {}
        if isinstance(raw, list):
            raw = {x.get("name", ""): x for x in raw if isinstance(x, dict)}
        if isinstance(raw, dict):
            for key, item in raw.items():
                name = _wz_cat_norm(item.get("name", key) if isinstance(item, dict) else key)
                if not name:
                    continue
                color = _wz_cat_hex(item.get("color") if isinstance(item, dict) else None, _wz_cat_default_color(name))
                result[name] = {"name": name, "color": color}
    # Lokale Kategorien automatisch zentralisieren, sobald G: verfügbar ist.
    if result and os.path.exists(_wz439_cats_l()) and _wz439_can_g():
        try:
            _wz439_write_json(_wz439_cats_g(), {"categories": result})
        except Exception:
            pass
    _WZ439_CAT_CACHE["sig"] = _wz439_sig(_wz439_cats_g()) + "|" + _wz439_sig(_wz439_cats_l())
    _WZ439_CAT_CACHE["data"] = dict(result)
    return result

def _wz439_save_category_meta(self, meta):
    clean = {}
    for key, item in (meta or {}).items():
        name = _wz_cat_norm(item.get("name", key) if isinstance(item, dict) else key)
        if not name:
            continue
        clean[name] = {"name": name, "color": _wz_cat_hex(item.get("color") if isinstance(item, dict) else None, _wz_cat_default_color(name))}
    target = _wz439_cats_g() if _wz439_can_g() else _wz439_cats_l()
    try:
        current = _wz439_load_category_meta(self)
        current.update(clean)
        _wz439_write_json(target, {"categories": current})
        _WZ439_CAT_CACHE["sig"] = None
        return True
    except Exception as exc:
        try:
            messagebox.showerror("Wissenszentrale", "Kategorien konnten nicht gespeichert werden:\n" + str(exc))
        except Exception:
            pass
        return False

def _wz439_get_category_color(self, name):
    name = _wz_cat_norm(name)
    meta = _wz439_load_category_meta(self)
    return _wz_cat_hex(meta.get(name, {}).get("color"), _wz_cat_default_color(name))

# Die alten globalen Kategorie-Funktionen werden bewusst überschrieben, weil kb_get_categories diese global aufruft.
_wz_cat_load = _wz439_load_category_meta
_wz_cat_save = _wz439_save_category_meta
_wz_cat_color = _wz439_get_category_color

def _wz439_font_size_from_widget(widget):
    try:
        f = tkfont.Font(font=widget.cget("font"))
        size = int(f.cget("size"))
        return abs(size) if size else 10
    except Exception:
        return 10

def _wz439_tag_name(size=10, bold=False, italic=False, underline=False):
    return f"wzfmt_s{int(size)}_b{int(bool(bold))}_i{int(bool(italic))}_u{int(bool(underline))}"

def _wz439_parse_tag(tag):
    m = re.fullmatch(r"wzfmt_s(\d+)_b([01])_i([01])_u([01])", str(tag or ""))
    if not m:
        return None
    return {"size": int(m.group(1)), "bold": m.group(2) == "1", "italic": m.group(3) == "1", "underline": m.group(4) == "1"}

def _wz439_font_tuple(fmt):
    styles = []
    if fmt.get("bold"):
        styles.append("bold")
    if fmt.get("italic"):
        styles.append("italic")
    if fmt.get("underline"):
        styles.append("underline")
    return tuple(["Segoe UI", int(fmt.get("size") or 10)] + styles)

def _wz439_ensure_tag(widget, size=10, bold=False, italic=False, underline=False):
    tag = _wz439_tag_name(size, bold, italic, underline)
    widget.tag_configure(tag, font=_wz439_font_tuple({"size": size, "bold": bold, "italic": italic, "underline": underline}))
    return tag

def _wz439_capture_formatting(widget):
    out = []
    try:
        for tag in widget.tag_names():
            fmt = _wz439_parse_tag(tag)
            if not fmt:
                continue
            ranges = widget.tag_ranges(tag)
            for i in range(0, len(ranges), 2):
                out.append({"start": str(ranges[i]), "end": str(ranges[i + 1]), **fmt})
    except Exception:
        pass
    return out

def _wz439_apply_formatting(widget, formatting):
    try:
        for item in formatting or []:
            if not isinstance(item, dict):
                continue
            tag = _wz439_ensure_tag(widget, item.get("size", 10), item.get("bold"), item.get("italic"), item.get("underline"))
            widget.tag_add(tag, item.get("start", "1.0"), item.get("end", "1.0"))
    except Exception:
        pass

def _wz439_format_at(widget, index):
    fmt = {"size": _wz439_font_size_from_widget(widget), "bold": False, "italic": False, "underline": False}
    try:
        for tag in widget.tag_names(index):
            parsed = _wz439_parse_tag(tag)
            if parsed:
                fmt.update(parsed)
    except Exception:
        pass
    return fmt

def _wz439_apply_selection(widget, size=None, toggle=None):
    try:
        start, end = widget.index("sel.first"), widget.index("sel.last")
    except Exception:
        return
    # Größe wird aus dem markierten Bereich übernommen, damit B/I/U keine Größenänderung auslösen.
    fmt = _wz439_format_at(widget, start)
    if size is not None:
        fmt["size"] = int(size)
    if toggle in ("bold", "italic", "underline"):
        fmt[toggle] = not bool(fmt.get(toggle))
    for tag in list(widget.tag_names()):
        if _wz439_parse_tag(tag):
            widget.tag_remove(tag, start, end)
    widget.tag_add(_wz439_ensure_tag(widget, **fmt), start, end)

# Kompatibilitätsnamen, falls v0.438-Funktionen noch referenzieren.
_wz437_parse = _wz439_parse_tag
_wz437_ensure_tag = _wz439_ensure_tag
_wz437_capture = _wz439_capture_formatting
_wz437_apply = _wz439_apply_formatting
_wz437_apply_selection = _wz439_apply_selection

def _wz439_toolbar(parent, widget, mark_unsaved):
    bar = tk.Frame(parent, bg=WHITE)
    bar.pack(fill="x", pady=(0, 4))
    tk.Label(bar, text="Format:", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="left", padx=(0, 6))
    size_var = tk.StringVar(value="10")
    size_cb = ttk.Combobox(bar, textvariable=size_var, values=["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "28", "32"], width=4, state="readonly", font=body_font(9))
    size_cb.pack(side="left", padx=(0, 6))
    size_cb.bind("<<ComboboxSelected>>", lambda e: (_wz439_apply_selection(widget, size=int(size_var.get())), mark_unsaved()))
    for label, key in [("B", "bold"), ("I", "italic"), ("U", "underline")]:
        tk.Button(bar, text=label, command=lambda k=key: (_wz439_apply_selection(widget, toggle=k), mark_unsaved()), bg=WHITE, fg=TEXT, font=body_font(10, weight="bold" if key == "bold" else None, underline=(key == "underline")), relief="solid", bd=1, width=3).pack(side="left", padx=(0, 4))
    tk.Label(bar, text="Text markieren und Format wählen.", bg=WHITE, fg=TEXT2, font=body_font(8)).pack(side="left", padx=(8, 0))

def _wz439_prepare_entry(self, entry=None):
    self.kb_ensure_state_vars()
    entry = entry or {}
    self.kb_edit_entry_id = entry.get("id") if entry else None
    self.kb_title_var = tk.StringVar(value=entry.get("title", ""))
    cats = list(entry.get("categories", []) or [])[:4]
    self.kb_entry_category_vars = [tk.StringVar(value=(cats[i] if i < len(cats) else "")) for i in range(4)]
    self.kb_user_var = tk.StringVar(value=entry.get("user", getattr(self, "current_user_display", "") or getattr(self, "current_user_key", "") or ""))
    self.kb_status_var = tk.StringVar(value=entry.get("status", "Aktiv"))
    self.kb_rhythm_var = tk.StringVar(value=entry.get("rhythm", ""))
    self.kb_text_initial = entry.get("text", "")
    self.kb_text_formatting_initial = entry.get("text_formatting", []) or []
    self.kb_inline_images = [dict(x) for x in (entry.get("inline_images", []) or []) if isinstance(x, dict)]
    self.kb_pending_attachments = []
    self.knowledge_unsaved = False

def _wz439_load_canvas_image(path, max_w=180, max_h=130):
    if not PIL_AVAILABLE:
        return None
    try:
        if not path or not os.path.exists(path):
            return None
        img = Image.open(path)
        img.thumbnail((int(max_w), int(max_h)))
        return ImageTk.PhotoImage(img)
    except Exception:
        return None

def _wz439_redraw_inline_images(self):
    canvas = getattr(self, "kb_image_canvas", None)
    if not canvas:
        return
    canvas.delete("all")
    self._kb_inline_photo_refs = []
    images = getattr(self, "kb_inline_images", []) or []
    if not images:
        canvas.create_text(12, 16, text="Bilder: Über 'Bild einfügen' können Bilder oberhalb des Textes eingefügt, verschoben und per Mausrad skaliert werden.", anchor="w", fill=TEXT2, font=body_font(9))
        return
    for idx, item in enumerate(images):
        x, y = int(item.get("x", 12)), int(item.get("y", 12))
        w, h = int(item.get("w", 180)), int(item.get("h", 130))
        path = item.get("path") or item.get("source")
        photo = _wz439_load_canvas_image(path, w, h)
        if photo:
            self._kb_inline_photo_refs.append(photo)
            canvas.create_image(x, y, image=photo, anchor="nw")
            canvas.create_rectangle(x, y, x + photo.width(), y + photo.height(), outline=BLUE if item.get("selected") else LINE, width=2)
            item["draw_w"], item["draw_h"] = photo.width(), photo.height()
        else:
            canvas.create_rectangle(x, y, x + w, y + h, fill="#FEE2E2", outline=RED)
            canvas.create_text(x + 8, y + 18, text="Bild nicht verfügbar", anchor="w", fill=RED, font=body_font(9))

def _wz439_image_hit(self, x, y):
    images = getattr(self, "kb_inline_images", []) or []
    for idx in range(len(images) - 1, -1, -1):
        it = images[idx]
        ix, iy = int(it.get("x", 12)), int(it.get("y", 12))
        iw, ih = int(it.get("draw_w", it.get("w", 180))), int(it.get("draw_h", it.get("h", 130)))
        if ix <= x <= ix + iw and iy <= y <= iy + ih:
            return idx
    return None

def _wz439_insert_image(self):
    try:
        from tkinter import filedialog
        path = filedialog.askopenfilename(title="Bild einfügen", filetypes=[("Bilder", "*.png *.jpg *.jpeg *.bmp *.gif"), ("Alle Dateien", "*.*")])
    except Exception:
        path = ""
    if not path:
        return
    images = list(getattr(self, "kb_inline_images", []) or [])
    images.append({"source": path, "path": path, "name": os.path.basename(path), "x": 12 + len(images) * 18, "y": 12 + len(images) * 12, "w": 180, "h": 130, "selected": True})
    for it in images[:-1]:
        it["selected"] = False
    self.kb_inline_images = images
    self.kb_mark_unsaved()
    _wz439_redraw_inline_images(self)

def _wz439_image_press(self, event):
    idx = _wz439_image_hit(self, event.x, event.y)
    self._kb_drag_image_idx = idx
    self._kb_drag_start = (event.x, event.y)
    for it in getattr(self, "kb_inline_images", []) or []:
        it["selected"] = False
    if idx is not None:
        self.kb_inline_images[idx]["selected"] = True
        self._kb_drag_orig = (int(self.kb_inline_images[idx].get("x", 12)), int(self.kb_inline_images[idx].get("y", 12)))
    _wz439_redraw_inline_images(self)

def _wz439_image_drag(self, event):
    idx = getattr(self, "_kb_drag_image_idx", None)
    if idx is None:
        return
    sx, sy = getattr(self, "_kb_drag_start", (event.x, event.y))
    ox, oy = getattr(self, "_kb_drag_orig", (12, 12))
    self.kb_inline_images[idx]["x"] = max(0, ox + event.x - sx)
    self.kb_inline_images[idx]["y"] = max(0, oy + event.y - sy)
    self.kb_mark_unsaved()
    _wz439_redraw_inline_images(self)

def _wz439_image_wheel(self, event):
    idx = None
    for i, it in enumerate(getattr(self, "kb_inline_images", []) or []):
        if it.get("selected"):
            idx = i
            break
    if idx is None:
        idx = _wz439_image_hit(self, getattr(event, "x", 0), getattr(event, "y", 0))
    if idx is None:
        return "break"
    factor = 1.08 if getattr(event, "delta", 0) > 0 else 0.92
    it = self.kb_inline_images[idx]
    it["w"] = max(40, min(900, int(float(it.get("w", 180)) * factor)))
    it["h"] = max(30, min(700, int(float(it.get("h", 130)) * factor)))
    self.kb_mark_unsaved()
    _wz439_redraw_inline_images(self)
    return "break"

def _wz439_text_mousewheel(widget, event):
    try:
        widget.yview_scroll(int(-1 * (event.delta / 120)), "units")
    except Exception:
        pass
    return "break"

def _wz439_attachment_dir(self, entry_id):
    path = os.path.join(_wz439_base_dir(), "Attachments", str(entry_id or "ohne_id")) if _wz439_can_g() else os.path.join(_WZ439_LOCAL_BASE, "knowledge_attachments", str(entry_id or "ohne_id"))
    os.makedirs(path, exist_ok=True)
    return path

def _wz439_normalize_images(self, entry_id):
    out = []
    target = os.path.join(self.kb_attachment_dir(entry_id), "InlineImages")
    try:
        os.makedirs(target, exist_ok=True)
    except Exception:
        pass
    import shutil as _shutil
    for item in getattr(self, "kb_inline_images", []) or []:
        if not isinstance(item, dict):
            continue
        src = item.get("source") or item.get("path")
        name = item.get("name") or os.path.basename(str(src or "bild"))
        dst = item.get("path") or src
        try:
            if src and os.path.exists(src):
                safe = re.sub(r"[^A-Za-z0-9_.äöüÄÖÜß-]+", "_", os.path.basename(name))[:100]
                dst = os.path.join(target, safe)
                if os.path.abspath(src) != os.path.abspath(dst):
                    _shutil.copy2(src, dst)
        except Exception:
            pass
        out.append({"name": name, "path": dst, "x": int(item.get("x", 12)), "y": int(item.get("y", 12)), "w": int(item.get("w", 180)), "h": int(item.get("h", 130))})
    return out

def _wz439_save_entry_from_form(self):
    self.kb_ensure_state_vars()
    title = (self.kb_title_var.get() or "").strip()
    if not title:
        try:
            messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
        except Exception:
            pass
        return False
    categories = []
    for var in getattr(self, "kb_entry_category_vars", []) or []:
        value = (var.get() or "").strip()
        if value and value not in categories:
            categories.append(value)
    categories = categories[:4]
    user = (self.kb_user_var.get() or "").strip()
    status = (self.kb_status_var.get() or "Aktiv").strip()
    rhythm = (self.kb_rhythm_var.get() or "").strip() if any(c.lower() == "to-do" for c in categories) else ""
    try:
        text_value = self.kb_text_widget.get("1.0", "end-1c")
        formatting = _wz439_capture_formatting(self.kb_text_widget)
    except Exception:
        text_value = getattr(self, "kb_text_initial", "")
        formatting = getattr(self, "kb_text_formatting_initial", []) or []
    entries = self.kb_load_entries()
    now = self.kb_now() if hasattr(self, "kb_now") else datetime.now().isoformat(timespec="seconds")
    edit_id = getattr(self, "kb_edit_entry_id", None)
    selected_id = edit_id or (self.kb_make_entry_id() if hasattr(self, "kb_make_entry_id") else hashlib.sha1((title + now).encode("utf-8")).hexdigest()[:16])
    inline_images = _wz439_normalize_images(self, selected_id)
    found = False
    for entry in entries:
        if entry.get("id") == selected_id:
            entry.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "text_formatting": formatting, "inline_images": inline_images, "updated_at": now})
            found = True
            break
    if not found:
        entries.append({"id": selected_id, "title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "text_formatting": formatting, "inline_images": inline_images, "created_at": now, "updated_at": now, "comments": [], "attachments": []})
    # Normale vorgemerkte Anhänge beibehalten.
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    if pending:
        import shutil as _shutil
        copied = []
        target = self.kb_attachment_dir(selected_id)
        for src in pending:
            try:
                name = os.path.basename(src)
                dst = os.path.join(target, name)
                _shutil.copy2(src, dst)
                copied.append({"name": name, "path": dst, "added_at": now})
            except Exception:
                pass
        for entry in entries:
            if entry.get("id") == selected_id:
                entry.setdefault("attachments", []).extend(copied)
                break
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = selected_id
        self.kb_edit_entry_id = None
        self.kb_pending_attachments = []
        self.knowledge_unsaved = False
        self.knowledge_view = "detail"
        self.render_page()
        return True
    return False

def _wz439_render_new_entry(self, x, y, w, h):
    if not hasattr(self, "kb_title_var"):
        self.kb_prepare_new_entry()
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text=("Eintrag bearbeiten" if getattr(self, "kb_edit_entry_id", None) else "Neuer Eintrag"), bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(10, 4))
    # Button-Leiste zuerst als bottom packen: bleibt sichtbar, auch wenn Textbereich wächst.
    buttons = tk.Frame(frame, bg=WHITE)
    buttons.pack(side="bottom", fill="x", padx=18, pady=(8, 12))
    tk.Button(buttons, text="Bild einfügen", command=lambda: _wz439_insert_image(self), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Anhang hinzufügen", command=self.kb_add_pending_attachment, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Als Entwurf speichern", command=lambda: (self.kb_status_var.set("Entwurf"), self.kb_save_entry_from_form()), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Speichern", command=self.kb_save_entry_from_form, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Abbrechen", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
    form = tk.Frame(frame, bg=WHITE)
    form.pack(side="top", fill="x", padx=18, pady=(0, 6))
    tk.Label(form, text="Titel des Eintrags", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=1, column=0, columnspan=4, sticky="ew", ipady=4, pady=(0, 6))
    values = [""] + self.kb_get_categories()
    for i, var in enumerate(self.kb_entry_category_vars):
        tk.Label(form, text=f"Kategorie {i+1}", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=i, sticky="w", padx=(0 if i == 0 else 8, 0))
        cb = ttk.Combobox(form, textvariable=var, values=values, state="normal", font=body_font(11), width=18)
        cb.grid(row=3, column=i, sticky="ew", padx=(0 if i == 0 else 8, 0), pady=(0, 6))
        cb.bind("<<ComboboxSelected>>", lambda e: self.render_page())
    tk.Label(form, text="Assoziierter Benutzer", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_user_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=5, column=0, sticky="ew", pady=(0, 4))
    tk.Label(form, text="Status", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=1, sticky="w", padx=(8, 0))
    ttk.Combobox(form, textvariable=self.kb_status_var, values=["Aktiv", "Entwurf", "Veraltet"], state="readonly", font=body_font(11)).grid(row=5, column=1, sticky="ew", padx=(8, 0), pady=(0, 4))
    if any((v.get() or "").strip().lower() == "to-do" for v in self.kb_entry_category_vars):
        tk.Label(form, text="To-Do-Rhythmus", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=2, sticky="w", padx=(8, 0))
        ttk.Combobox(form, textvariable=self.kb_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(11)).grid(row=5, column=2, sticky="ew", padx=(8, 0), pady=(0, 4))
    else:
        self.kb_rhythm_var.set("")
    for col in range(4):
        form.grid_columnconfigure(col, weight=1)
    work = tk.Frame(frame, bg=WHITE)
    work.pack(side="top", fill="both", expand=True, padx=18, pady=(0, 0))
    tk.Label(work, text="Bilder vor dem Text", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    self.kb_image_canvas = tk.Canvas(work, bg="#F8FAFC", height=ui_s(155), highlightthickness=1, highlightbackground=LINE)
    self.kb_image_canvas.pack(fill="x", pady=(0, 6))
    self.kb_image_canvas.bind("<Button-1>", lambda e: _wz439_image_press(self, e))
    self.kb_image_canvas.bind("<B1-Motion>", lambda e: _wz439_image_drag(self, e))
    self.kb_image_canvas.bind("<MouseWheel>", lambda e: _wz439_image_wheel(self, e))
    tk.Label(work, text="Freitext / Prozessdokumentation / Leitfaden", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    text_area = tk.Frame(work, bg=WHITE)
    _wz439_toolbar(work, self.kb_text_widget if hasattr(self, "kb_text_widget") else tk.Text(work), lambda: None)  # Platzhalter wird direkt danach ersetzt.
    # Platzhalter-Toolbar entfernen und korrekt mit echtem Text-Widget neu aufbauen.
    try:
        for child in list(work.winfo_children()):
            if isinstance(child, tk.Frame) and child is not text_area and len(child.winfo_children()) >= 1:
                if getattr(child.winfo_children()[0], 'cget', lambda x: '')('text') == 'Format:':
                    child.destroy()
    except Exception:
        pass
    self.kb_text_widget = tk.Text(text_area, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word", undo=True)
    _wz439_toolbar(work, self.kb_text_widget, self.kb_mark_unsaved)
    text_area.pack(fill="both", expand=True)
    yscroll = tk.Scrollbar(text_area, orient="vertical", command=self.kb_text_widget.yview)
    self.kb_text_widget.configure(yscrollcommand=yscroll.set)
    self.kb_text_widget.pack(side="left", fill="both", expand=True)
    yscroll.pack(side="right", fill="y")
    self.kb_text_widget.insert("1.0", getattr(self, "kb_text_initial", ""))
    _wz439_apply_formatting(self.kb_text_widget, getattr(self, "kb_text_formatting_initial", []) or [])
    self.kb_text_widget.bind("<KeyRelease>", lambda e: self.kb_mark_unsaved())
    self.kb_text_widget.bind("<MouseWheel>", lambda e: _wz439_text_mousewheel(self.kb_text_widget, e))
    _wz439_redraw_inline_images(self)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

def _wz439_badges(parent, app, cats):
    row = tk.Frame(parent, bg=WHITE)
    row.pack(anchor="w", fill="x", padx=18, pady=(8, 4))
    tk.Label(row, text="Kategorien:", bg=WHITE, fg=TEXT, font=body_font(10, weight="bold")).pack(side="left", padx=(0, 8))
    if not cats:
        tk.Label(row, text="Keine Kategorien", bg=WHITE, fg=TEXT2, font=body_font(10)).pack(side="left")
        return
    for cat in cats:
        color = app.kb_get_category_color(cat)
        tk.Label(row, text="  " + str(cat) + "  ", bg=color, fg=_wz_cat_fg(color), font=body_font(9, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 6), ipady=2)

def _wz439_render_image_view(parent, app, images):
    if not images:
        return
    tk.Label(parent, text="Bilder", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w", padx=18, pady=(8, 2))
    canvas = tk.Canvas(parent, bg="#F8FAFC", height=ui_s(165), highlightthickness=1, highlightbackground=LINE)
    canvas.pack(fill="x", padx=18, pady=(0, 6))
    app._kb_detail_photo_refs = []
    for item in images:
        path = item.get("path") or item.get("source")
        x, y = int(item.get("x", 12)), int(item.get("y", 12))
        w, h = int(item.get("w", 180)), int(item.get("h", 130))
        photo = _wz439_load_canvas_image(path, w, h)
        if photo:
            app._kb_detail_photo_refs.append(photo)
            canvas.create_image(x, y, image=photo, anchor="nw")
            canvas.create_rectangle(x, y, x + photo.width(), y + photo.height(), outline=LINE)

def _wz439_render_detail(self, x, y, w, h):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry:
        self.render_kb_list_area(x, y, w, h, title="Übersicht")
        return
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text=entry.get("title", ""), bg=WHITE, fg=BLUE, font=body_font(16, weight="bold")).pack(anchor="w", padx=18, pady=(14, 5))
    tk.Label(frame, text=f"Geändert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {entry.get('user','')}    Status: {entry.get('status','')}", bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor="w", padx=18)
    _wz439_badges(frame, self, entry.get("categories", []) or [])
    _wz439_render_image_view(frame, self, entry.get("inline_images", []) or [])
    text_frame = tk.Frame(frame, bg=WHITE)
    text_frame.pack(fill="both", expand=True, padx=18, pady=(6, 8))
    txt = tk.Text(text_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word", height=8)
    yscroll = tk.Scrollbar(text_frame, orient="vertical", command=txt.yview)
    txt.configure(yscrollcommand=yscroll.set)
    txt.pack(side="left", fill="both", expand=True)
    yscroll.pack(side="right", fill="y")
    txt.insert("1.0", entry.get("text", ""))
    _wz439_apply_formatting(txt, entry.get("text_formatting", []) or [])
    txt.configure(state="disabled")
    txt.bind("<MouseWheel>", lambda ev: _wz439_text_mousewheel(txt, ev))
    lower = tk.Frame(frame, bg=WHITE)
    lower.pack(fill="x", padx=18, pady=(0, 8))
    left = tk.Frame(lower, bg=WHITE)
    left.pack(side="left", fill="both", expand=True)
    right = tk.Frame(lower, bg=WHITE)
    right.pack(side="right", fill="both", expand=True, padx=(12, 0))
    tk.Label(left, text="Anhänge", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    for a in entry.get("attachments", []) or []:
        tk.Label(left, text="• " + a.get("name", ""), bg=WHITE, fg=TEXT, font=body_font(9)).pack(anchor="w")
    tk.Label(right, text="Kommentare", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    for c in (entry.get("comments", []) or [])[-3:]:
        tk.Label(right, text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}", bg=WHITE, fg=TEXT, font=body_font(9), wraplength=380, justify="left").pack(anchor="w")
    comment = tk.Text(right, height=2, bg="#F8FAFC", fg=TEXT, font=body_font(9), relief="solid", bd=1, wrap="word")
    comment.pack(fill="x", pady=(4, 0))
    buttons = tk.Frame(frame, bg=WHITE)
    buttons.pack(fill="x", padx=18, pady=(0, 14))
    if self.kb_can_create_or_edit():
        tk.Button(buttons, text="Bearbeiten", command=self.kb_edit_selected_entry, bg=WHITE, fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Kommentar speichern", command=lambda: self.kb_add_comment_to_selected(comment), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Word-Export", command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Zur Übersicht", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

# Finale Zuweisungen direkt vor Programmstart.
FiBuMateApp.kb_load_category_meta = _wz439_load_category_meta
FiBuMateApp.kb_save_category_meta = _wz439_save_category_meta
FiBuMateApp.kb_get_category_color = _wz439_get_category_color
FiBuMateApp.kb_prepare_new_entry = _wz439_prepare_entry
FiBuMateApp.kb_save_entry_from_form = _wz439_save_entry_from_form
FiBuMateApp.kb_attachment_dir = _wz439_attachment_dir
FiBuMateApp.render_kb_new_entry_area = _wz439_render_new_entry
FiBuMateApp.render_kb_detail_area = _wz439_render_detail



# ------------------------------------------------------------------
# Wissenszentrale - Performance-Cache FINAL 2026-06-18
# Version 0.441
# Zweck:
# - Entfernt die langen Ladezeiten durch wiederholte G:-Prüfungen und JSON-Voll-Ladevorgänge pro Render.
# - Cacht Einträge, Kategorien, Farben und Bild-Thumbnails im Arbeitsspeicher.
# - Speichern invalidiert den Cache sofort; Lesen nutzt kurze TTL statt blockierendem Netzwerkzugriff je Widget.
# - Kein G:-Write-Test mehr beim Rendern; G: wird erst beim Speichern wirklich beschrieben.
# ------------------------------------------------------------------

_WZ441_ENTRIES_CACHE = {"ts": 0.0, "sig": "", "entries": []}
_WZ441_CAT_CACHE = {"ts": 0.0, "sig": "", "meta": {}}
_WZ441_G_CACHE = {"ts": 0.0, "ok": None}
_WZ441_IMG_CACHE = {}
_WZ441_ENTRIES_TTL_SEC = 8.0
_WZ441_CATEGORIES_TTL_SEC = 30.0
_WZ441_G_TTL_SEC = 60.0

def _wz441_base_dir():
    return os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database", "Wissenszentrale")

def _wz441_entries_g():
    return os.path.join(_wz441_base_dir(), "knowledge_entries.json")

def _wz441_entries_l():
    return r"C:\python\knowledge_entries.json"

def _wz441_cats_g():
    return os.path.join(_wz441_base_dir(), "knowledge_categories.json")

def _wz441_cats_l():
    return r"C:\python\knowledge_categories.json"

def _wz441_sig(path):
    try:
        if path and os.path.exists(path):
            st=os.stat(path)
            return f"{st.st_mtime_ns}:{st.st_size}"
    except Exception:
        pass
    return "missing"

def _wz441_read_json(path):
    try:
        if path and os.path.exists(path):
            with open(path,"r",encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _wz441_write_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp=path+".tmp"
    with open(tmp,"w",encoding="utf-8") as f:
        json.dump(data,f,ensure_ascii=False,indent=2)
    os.replace(tmp,path)

def _wz441_can_use_g():
    import time
    now=time.time()
    if _WZ441_G_CACHE["ok"] is not None and now-_WZ441_G_CACHE["ts"] < _WZ441_G_TTL_SEC:
        return bool(_WZ441_G_CACHE["ok"])
    # Wichtig: kein Write-Test beim Rendern. Der war auf Netzlaufwerken der Haupt-Bremser.
    ok=False
    try:
        base=_wz441_base_dir()
        ok=os.path.isdir(base) or os.path.isdir(os.path.dirname(base))
    except Exception:
        ok=False
    _WZ441_G_CACHE["ok"]=ok
    _WZ441_G_CACHE["ts"]=now
    return ok

def _wz441_merge_entries(*lists):
    merged={}; order=[]
    for entries in lists:
        for e in entries or []:
            if not isinstance(e,dict):
                continue
            eid=str(e.get("id") or "").strip()
            if not eid:
                eid=hashlib.sha1((str(e.get("title",""))+str(e.get("created_at",""))+str(e.get("text",""))).encode("utf-8",errors="ignore")).hexdigest()[:16]
                e["id"]=eid
            if eid not in merged:
                merged[eid]=e; order.append(eid); continue
            old=merged[eid]
            if str(e.get("updated_at") or e.get("created_at") or "") >= str(old.get("updated_at") or old.get("created_at") or ""):
                out=dict(old); out.update(e)
                for k in ("comments","attachments"):
                    seen=set(); arr=[]
                    for item in (old.get(k,[]) or [])+(e.get(k,[]) or []):
                        marker=json.dumps(item,ensure_ascii=False,sort_keys=True) if isinstance(item,dict) else str(item)
                        if marker not in seen:
                            seen.add(marker); arr.append(item)
                    out[k]=arr
                merged[eid]=out
    return [merged[i] for i in order if i in merged]

def _wz441_invalidate_entries():
    _WZ441_ENTRIES_CACHE["ts"]=0.0
    _WZ441_ENTRIES_CACHE["sig"]=""

def _wz441_invalidate_categories():
    _WZ441_CAT_CACHE["ts"]=0.0
    _WZ441_CAT_CACHE["sig"]=""

def _wz441_load_entries(self):
    import time
    now=time.time()
    gp,lp=_wz441_entries_g(),_wz441_entries_l()
    # Innerhalb TTL: kein Netzlaufwerk anfassen, keine JSON-Datei neu parsen.
    if _WZ441_ENTRIES_CACHE["entries"] and now-_WZ441_ENTRIES_CACHE["ts"] < _WZ441_ENTRIES_TTL_SEC:
        return list(_WZ441_ENTRIES_CACHE["entries"])
    sig=_wz441_sig(gp)+"|"+_wz441_sig(lp)
    if sig == _WZ441_ENTRIES_CACHE["sig"] and _WZ441_ENTRIES_CACHE["entries"]:
        _WZ441_ENTRIES_CACHE["ts"]=now
        return list(_WZ441_ENTRIES_CACHE["entries"])
    gd=_wz441_read_json(gp); ld=_wz441_read_json(lp)
    ge=gd.get("entries",[]) if isinstance(gd,dict) and isinstance(gd.get("entries",[]),list) else []
    le=ld.get("entries",[]) if isinstance(ld,dict) and isinstance(ld.get("entries",[]),list) else []
    entries=_wz441_merge_entries(ge,le)
    _WZ441_ENTRIES_CACHE.update({"ts":now,"sig":sig,"entries":entries})
    try:
        self.knowledge_last_load_path = gp if ge else (lp if le else gp)
    except Exception:
        pass
    return list(entries)

def _wz441_save_entries(self, entries):
    target=_wz441_entries_g() if _wz441_can_use_g() else _wz441_entries_l()
    try:
        current=_wz441_read_json(target).get("entries",[])
        merged=_wz441_merge_entries(current, entries)
        _wz441_write_json(target,{"entries":merged})
        _wz441_invalidate_entries()
        try:
            self.knowledge_last_save_path=target
            self.knowledge_used_local_fallback=(target==_wz441_entries_l())
        except Exception:
            pass
        return True
    except Exception as exc:
        try:
            messagebox.showerror("Wissenszentrale", "Eintrag konnte nicht gespeichert werden:\n"+str(exc))
        except Exception:
            pass
        return False

def _wz441_load_category_meta(self):
    import time
    now=time.time()
    gp,lp=_wz441_cats_g(),_wz441_cats_l()
    if _WZ441_CAT_CACHE["meta"] and now-_WZ441_CAT_CACHE["ts"] < _WZ441_CATEGORIES_TTL_SEC:
        return dict(_WZ441_CAT_CACHE["meta"])
    sig=_wz441_sig(gp)+"|"+_wz441_sig(lp)
    if sig == _WZ441_CAT_CACHE["sig"] and _WZ441_CAT_CACHE["meta"]:
        _WZ441_CAT_CACHE["ts"]=now
        return dict(_WZ441_CAT_CACHE["meta"])
    result={}
    for path in (gp,lp):
        data=_wz441_read_json(path)
        raw=data.get("categories",{}) if isinstance(data,dict) else {}
        if isinstance(raw,list):
            raw={x.get("name",""):x for x in raw if isinstance(x,dict)}
        if isinstance(raw,dict):
            for key,item in raw.items():
                name=_wz_cat_norm(item.get("name",key) if isinstance(item,dict) else key)
                if not name: continue
                color=_wz_cat_hex(item.get("color") if isinstance(item,dict) else None,_wz_cat_default_color(name))
                result[name]={"name":name,"color":color}
    _WZ441_CAT_CACHE.update({"ts":now,"sig":sig,"meta":result})
    return dict(result)

def _wz441_save_category_meta(self, meta):
    clean={}
    for key,item in (meta or {}).items():
        name=_wz_cat_norm(item.get("name",key) if isinstance(item,dict) else key)
        if name:
            clean[name]={"name":name,"color":_wz_cat_hex(item.get("color") if isinstance(item,dict) else None,_wz_cat_default_color(name))}
    target=_wz441_cats_g() if _wz441_can_use_g() else _wz441_cats_l()
    try:
        cur=_wz441_load_category_meta(self); cur.update(clean)
        _wz441_write_json(target,{"categories":cur})
        _wz441_invalidate_categories()
        return True
    except Exception as exc:
        try:
            messagebox.showerror("Wissenszentrale", "Kategorien konnten nicht gespeichert werden:\n"+str(exc))
        except Exception:
            pass
        return False

def _wz441_get_category_color(self, name):
    name=_wz_cat_norm(name)
    meta=_wz441_load_category_meta(self)
    return _wz_cat_hex(meta.get(name,{}).get("color"), _wz_cat_default_color(name))

def _wz441_get_categories(self):
    cats=set(getattr(self,"kb_default_categories",lambda:["To-Do"])())
    cats.update(_wz441_load_category_meta(self).keys())
    try:
        for entry in _wz441_load_entries(self):
            for cat in entry.get("categories",[]) or []:
                c=_wz_cat_norm(cat)
                if c: cats.add(c)
    except Exception:
        pass
    return sorted(cats, key=lambda c:c.lower())

def _wz441_attachment_dir(self, entry_id):
    # Kein G:-Write-Test beim Öffnen der Wissenszentrale; Ordner wird erst bei tatsächlicher Anhang-/Bildspeicherung erstellt.
    base=os.path.join(_wz441_base_dir(),"Attachments",str(entry_id or "ohne_id")) if _wz441_can_use_g() else os.path.join(r"C:\python","knowledge_attachments",str(entry_id or "ohne_id"))
    try: os.makedirs(base, exist_ok=True)
    except Exception: pass
    return base

# Thumbnail-Cache für Bilder, damit Details/Bearbeiten nicht jedes Bild bei jedem Render neu von G: laden.
def _wz441_cached_photo(path, max_w=180, max_h=130):
    if not PIL_AVAILABLE or not path or not os.path.exists(path): return None
    try:
        st=os.stat(path); key=(path,int(max_w),int(max_h),st.st_mtime_ns,st.st_size)
        if key in _WZ441_IMG_CACHE: return _WZ441_IMG_CACHE[key]
        # Cache begrenzen
        if len(_WZ441_IMG_CACHE) > 80: _WZ441_IMG_CACHE.clear()
        img=Image.open(path); img.thumbnail((int(max_w),int(max_h)))
        photo=ImageTk.PhotoImage(img)
        _WZ441_IMG_CACHE[key]=photo
        return photo
    except Exception:
        return None

# Falls v0.439-Bildfunktionen vorhanden sind, deren Loader auf Cache umbiegen.
def _wz439_load_canvas_image(path, max_w=180, max_h=130):
    return _wz441_cached_photo(path,max_w,max_h)

# Finale schnelle Zuweisungen.
FiBuMateApp.kb_load_entries = _wz441_load_entries
FiBuMateApp.kb_save_entries = _wz441_save_entries
FiBuMateApp.kb_get_categories = _wz441_get_categories
FiBuMateApp.kb_load_category_meta = _wz441_load_category_meta
FiBuMateApp.kb_save_category_meta = _wz441_save_category_meta
FiBuMateApp.kb_get_category_color = _wz441_get_category_color
FiBuMateApp.kb_attachment_dir = _wz441_attachment_dir
_wz_cat_load = _wz441_load_category_meta
_wz_cat_save = _wz441_save_category_meta
_wz_cat_color = _wz441_get_category_color



# ------------------------------------------------------------------
# Wissenszentrale - Feature-Restore auf Performance-Stand FINAL 2026-06-18
# Version 0.442
# Zweck:
# - Baut ausdrücklich auf dem aktuellen Performance-Stand v0.441 auf.
# - Stellt das Wissenszentrale-Icon bookshelf wieder her.
# - Stellt farbliche Kategorie-Markierungen in Trefferliste und Kategorie-Auswahlen wieder her.
# - Erhält Performance-Caches aus v0.441; keine neuen G:-Write-Tests beim Rendern.
# - Stellt Bild-Overlay im Textfeld wieder her, ohne separates Bildfeld.
# ------------------------------------------------------------------

# --- Icon-Restore final: auch falls ältere Konstanten/Mappings vorher aktiv waren ---
try:
    ICON_FILES["knowledge"] = "bookshelf_icon-icons.com_54414.ico"
except Exception:
    pass

try:
    _wz442_old_draw_tile_icon_image = FiBuMateApp.draw_tile_icon_image
    def _wz442_draw_tile_icon_image(self, tile, icon_type, cx, cy):
        if icon_type == "knowledge":
            photo = self.get_icon_photo("knowledge", 48, 48)
            if photo:
                tile.create_image(cx, cy, image=photo)
                return True
        return _wz442_old_draw_tile_icon_image(self, tile, icon_type, cx, cy)
    FiBuMateApp.draw_tile_icon_image = _wz442_draw_tile_icon_image
except Exception:
    pass

try:
    _wz442_old_render_main_menu = FiBuMateApp.render_main_menu
    def _wz442_render_main_menu(self):
        # Sicherheitsnetz: erzwingt im Hauptmenü die bookshelf-Icondatei für die Wissenszentrale,
        # auch wenn ältere Kacheldefinitionen noch "info" verwenden würden.
        try:
            ICON_FILES["knowledge"] = "bookshelf_icon-icons.com_54414.ico"
        except Exception:
            pass
        return _wz442_old_render_main_menu(self)
    FiBuMateApp.render_main_menu = _wz442_render_main_menu
except Exception:
    pass

def _wz442_light_color(color, factor=0.84):
    try:
        color = _wz_cat_hex(color)
        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        r = int(r + (255-r)*factor); g = int(g + (255-g)*factor); b = int(b + (255-b)*factor)
        return f"#{r:02X}{g:02X}{b:02X}"
    except Exception:
        return "#F3F4F6"

def _wz442_cat_color(app, cat):
    try:
        return app.kb_get_category_color(cat)
    except Exception:
        try:
            return _wz_cat_color(app, cat)
        except Exception:
            return _wz_cat_default_color(cat)

def _wz442_option_menu(parent, app, variable, values, command=None, width=18):
    bg = parent.cget("bg") if hasattr(parent, "cget") else BG
    frame = tk.Frame(parent, bg=bg)
    swatch = tk.Label(frame, text="", width=2, relief="solid", bd=1, bg=WHITE)
    swatch.pack(side="left", padx=(0, 4), ipady=7)
    om = tk.OptionMenu(frame, variable, *values)
    om.configure(bg=WHITE, fg=TEXT, activebackground="#DDEAF7", relief="solid", bd=1, font=body_font(10), width=width)
    try:
        menu = om["menu"]
        menu.configure(font=body_font(10))
        menu.delete(0, "end")
        for val in values:
            def _cmd(v=val):
                variable.set(v)
                _update()
                if command:
                    command()
            menu.add_command(label=(val or "(leer)"), command=_cmd)
            idx = menu.index("end")
            if val:
                c = _wz442_cat_color(app, val)
                try:
                    menu.entryconfig(idx, background=_wz442_light_color(c), foreground=TEXT, activebackground=c, activeforeground=_wz_cat_fg(c))
                except Exception:
                    pass
    except Exception:
        pass
    om.pack(side="left")
    def _update(*_):
        val = (variable.get() or "").strip()
        if val:
            c = _wz442_cat_color(app, val)
            swatch.configure(bg=c)
            try:
                om.configure(bg=_wz442_light_color(c), activebackground=c)
            except Exception:
                pass
        else:
            swatch.configure(bg=WHITE)
            try:
                om.configure(bg=WHITE)
            except Exception:
                pass
    try:
        variable.trace_add("write", _update)
    except Exception:
        pass
    _update()
    return frame

def _wz442_render_hits_pane(self, x, y, w, h):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text="Treffer", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
    entries = self.kb_filtered_entries()
    listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="flat", activestyle="none", exportselection=False)
    listbox.pack(fill="both", expand=True, padx=16, pady=(0, 14))
    id_map = []
    if entries:
        for idx, entry in enumerate(entries):
            cats = entry.get("categories", []) or []
            cats_txt = ", ".join(cats)
            line = f"{self.kb_display_date(entry.get('updated_at'))}  |  {entry.get('title','')}"
            if cats_txt:
                line += f"  [{cats_txt}]"
            listbox.insert("end", line)
            id_map.append(entry.get("id"))
            if cats:
                c = _wz442_cat_color(self, cats[0])
                try:
                    listbox.itemconfig(idx, bg=_wz442_light_color(c), fg=TEXT, selectbackground=c, selectforeground=_wz_cat_fg(c))
                except Exception:
                    pass
    else:
        listbox.insert("end", "Noch keine Treffer vorhanden.")
        listbox.insert("end", "Filter und Suche wirken auf die Wissensdatenbank.")
    def _select(event=None):
        sel = listbox.curselection()
        if not sel or not id_map or sel[0] >= len(id_map):
            return
        if not self.kb_confirm_unsaved_before_switch():
            return
        self.kb_selected_entry_id = id_map[sel[0]]
        self.knowledge_view = "detail"
        self.knowledge_start_overlay = False
        self.render_page()
    listbox.bind("<<ListboxSelect>>", _select)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

def _wz442_render_work_area(self):
    self.kb_ensure_state_vars()
    try:
        self.root.bind("<Escape>", self.kb_handle_escape, add="+")
    except Exception:
        pass
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    nav_y = 145
    bx = 28
    bx = self.draw_kb_button(bx, nav_y, "Start", self.kb_show_start_overlay, False, width=92)
    bx = self.draw_kb_button(bx, nav_y, "Übersicht", lambda: self.kb_switch_view_from_start("all"), False, width=112)
    if self.kb_can_create_or_edit():
        bx = self.draw_kb_button(bx, nav_y, "Neuer Eintrag", lambda: self.kb_switch_view_from_start("new"), False, width=128)
    bx = self.draw_kb_button(bx, nav_y, "To-Dos", lambda: self.kb_switch_view_from_start("todos"), False, width=92)
    bx = self.draw_kb_button(bx, nav_y, "Veraltete Einträge", lambda: self.kb_switch_view_from_start("outdated"), False, width=148)
    if self.kb_can_manage_categories():
        bx = self.draw_kb_button(bx, nav_y, "Kategorien verwalten", lambda: self.kb_switch_view_from_start("categories"), False, width=168)

    search_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(search_frame)
    tk.Label(search_frame, text="Suche", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w")
    search_entry = tk.Entry(search_frame, textvariable=self.kb_search_var, font=body_font(11), bg=WHITE, fg=TEXT, relief="solid", bd=1)
    search_entry.grid(row=1, column=0, sticky="ew", ipady=5)
    search_entry.bind("<Return>", self.kb_on_search_return)
    tk.Button(search_frame, text="Suchen", command=self.kb_apply_filters, bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10)).grid(row=1, column=1, padx=(8, 0), sticky="ns")
    search_frame.grid_columnconfigure(0, weight=1)
    self.canvas.create_window(ui_s(28), ui_s(192), window=search_frame, anchor="nw", width=ui_s(min(790, max(560, w-80))), height=ui_s(60))

    filter_frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(filter_frame)
    values = [""] + self.kb_get_categories()
    tk.Label(filter_frame, text="Kategorie-Filter", bg=BG, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky="w", columnspan=6)
    for idx, var in enumerate(self.kb_filter_vars):
        block = tk.Frame(filter_frame, bg=BG)
        block.grid(row=1, column=idx, padx=(0 if idx == 0 else 12, 0), sticky="w")
        tk.Label(block, text=f"Kategorie {idx+1}", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        opt = _wz442_option_menu(block, self, var, values, command=self.kb_apply_filters, width=14)
        opt.pack(anchor="w")
    self.canvas.create_window(ui_s(28), ui_s(252), window=filter_frame, anchor="nw", width=ui_s(min(1050, max(760, w-70))), height=ui_s(90))

    left_x, left_y = 20, 350
    left_w = max(340, min(460, int(w * 0.31)))
    pane_h = max(360, h-left_y-24)
    right_x = left_x + left_w + 24
    right_w = max(760, w-right_x-24)
    self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
    if self.knowledge_view == "new":
        self.render_kb_new_entry_area(right_x, left_y, right_w, pane_h)
    elif self.knowledge_view == "todos":
        self.render_kb_list_area(right_x, left_y, right_w, pane_h, title="To-Dos")
    elif self.knowledge_view == "outdated":
        self.render_kb_list_area(right_x, left_y, right_w, pane_h, title="Veraltete Einträge", status_filter="Veraltet")
    elif self.knowledge_view == "categories":
        self.render_kb_categories_area(right_x, left_y, right_w, pane_h)
    elif self.knowledge_view == "detail":
        self.render_kb_detail_area(right_x, left_y, right_w, pane_h)
    else:
        self.render_kb_list_area(right_x, left_y, right_w, pane_h, title="Übersicht")

def _wz442_photo(path, max_w=180, max_h=130):
    try:
        return _wz441_cached_photo(path, max_w, max_h)
    except Exception:
        pass
    try:
        if not PIL_AVAILABLE or not path or not os.path.exists(path):
            return None
        img = Image.open(path)
        img.thumbnail((int(max_w), int(max_h)))
        return ImageTk.PhotoImage(img)
    except Exception:
        return None

def _wz442_place_text_images(self):
    parent = getattr(self, "kb_text_overlay_parent", None)
    if not parent:
        return
    for child in list(parent.children.values()):
        if getattr(child, "_wz_float_image", False):
            child.destroy()
    self._kb_text_image_refs = []
    for idx, item in enumerate(getattr(self, "kb_inline_images", []) or []):
        photo = _wz442_photo(item.get("path") or item.get("source"), item.get("w", 180), item.get("h", 130))
        if not photo:
            continue
        self._kb_text_image_refs.append(photo)
        lbl = tk.Label(parent, image=photo, bg="#F8FAFC", bd=1, relief="solid", cursor="fleur")
        lbl._wz_float_image = True
        lbl.place(x=int(item.get("x", 20)), y=int(item.get("y", 20)))
        def press(event, i=idx):
            self._kb_float_idx = i
            self._kb_float_start = (event.x_root, event.y_root)
            self._kb_float_orig = (int(self.kb_inline_images[i].get("x", 20)), int(self.kb_inline_images[i].get("y", 20)))
        def drag(event, i=idx):
            sx, sy = getattr(self, "_kb_float_start", (event.x_root, event.y_root))
            ox, oy = getattr(self, "_kb_float_orig", (20, 20))
            self.kb_inline_images[i]["x"] = max(0, ox + event.x_root - sx)
            self.kb_inline_images[i]["y"] = max(0, oy + event.y_root - sy)
            self.kb_mark_unsaved()
            _wz442_place_text_images(self)
        def wheel(event, i=idx):
            factor = 1.08 if getattr(event, "delta", 0) > 0 else 0.92
            self.kb_inline_images[i]["w"] = max(40, min(900, int(float(self.kb_inline_images[i].get("w", 180)) * factor)))
            self.kb_inline_images[i]["h"] = max(30, min(700, int(float(self.kb_inline_images[i].get("h", 130)) * factor)))
            self.kb_mark_unsaved()
            _wz442_place_text_images(self)
            return "break"
        lbl.bind("<Button-1>", press)
        lbl.bind("<B1-Motion>", drag)
        lbl.bind("<MouseWheel>", wheel)

def _wz442_add_image(self, path=None):
    if not path:
        try:
            from tkinter import filedialog
            path = filedialog.askopenfilename(title="Bild in Textfeld einfügen", filetypes=[("Bilder", "*.png *.jpg *.jpeg *.bmp *.gif"), ("Alle Dateien", "*.*")])
        except Exception:
            path = ""
    if not path:
        return
    imgs = list(getattr(self, "kb_inline_images", []) or [])
    imgs.append({"source": path, "path": path, "name": os.path.basename(path), "x": 30, "y": 30, "w": 180, "h": 130})
    self.kb_inline_images = imgs
    self.kb_mark_unsaved()
    _wz442_place_text_images(self)

def _wz442_normalize_images(self, entry_id):
    out = []
    target = os.path.join(self.kb_attachment_dir(entry_id), "InlineImages")
    try:
        os.makedirs(target, exist_ok=True)
    except Exception:
        pass
    import shutil as _shutil
    for item in getattr(self, "kb_inline_images", []) or []:
        if not isinstance(item, dict):
            continue
        src = item.get("source") or item.get("path")
        name = item.get("name") or os.path.basename(str(src or "bild"))
        dst = item.get("path") or src
        try:
            if src and os.path.exists(src):
                safe = re.sub(r"[^A-Za-z0-9_.äöüÄÖÜß-]+", "_", os.path.basename(name))[:100]
                dst = os.path.join(target, safe)
                if os.path.abspath(src) != os.path.abspath(dst):
                    _shutil.copy2(src, dst)
        except Exception:
            pass
        out.append({"name": name, "path": dst, "x": int(item.get("x", 20)), "y": int(item.get("y", 20)), "w": int(item.get("w", 180)), "h": int(item.get("h", 130))})
    return out

def _wz442_save_entry_from_form(self):
    self.kb_ensure_state_vars()
    title = (self.kb_title_var.get() or "").strip()
    if not title:
        try:
            messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
        except Exception:
            pass
        return False
    categories = []
    for var in getattr(self, "kb_entry_category_vars", []) or []:
        value = (var.get() or "").strip()
        if value and value not in categories:
            categories.append(value)
    categories = categories[:4]
    user = (self.kb_user_var.get() or "").strip()
    status = (self.kb_status_var.get() or "Aktiv").strip()
    rhythm = (self.kb_rhythm_var.get() or "").strip() if any(c.lower() == "to-do" for c in categories) else ""
    text_value = self.kb_text_widget.get("1.0", "end-1c") if hasattr(self, "kb_text_widget") else getattr(self, "kb_text_initial", "")
    try:
        formatting = _wz439_capture_formatting(self.kb_text_widget)
    except Exception:
        formatting = getattr(self, "kb_text_formatting_initial", []) or []
    entries = self.kb_load_entries()
    now = self.kb_now() if hasattr(self, "kb_now") else datetime.now().isoformat(timespec="seconds")
    selected_id = getattr(self, "kb_edit_entry_id", None) or (self.kb_make_entry_id() if hasattr(self, "kb_make_entry_id") else hashlib.sha1((title + now).encode("utf-8")).hexdigest()[:16])
    inline_images = _wz442_normalize_images(self, selected_id)
    found = False
    for entry in entries:
        if entry.get("id") == selected_id:
            entry.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "text_formatting": formatting, "inline_images": inline_images, "updated_at": now})
            found = True
            break
    if not found:
        entries.append({"id": selected_id, "title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "text_formatting": formatting, "inline_images": inline_images, "created_at": now, "updated_at": now, "comments": [], "attachments": []})
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = selected_id
        self.kb_edit_entry_id = None
        self.knowledge_unsaved = False
        self.knowledge_view = "detail"
        self.render_page()
        return True
    return False

def _wz442_render_new_entry(self, x, y, w, h):
    if not hasattr(self, "kb_title_var"):
        self.kb_prepare_new_entry()
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text=("Eintrag bearbeiten" if getattr(self, "kb_edit_entry_id", None) else "Neuer Eintrag"), bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(10, 4))
    buttons = tk.Frame(frame, bg=WHITE)
    buttons.pack(side="bottom", fill="x", padx=18, pady=(8, 12))
    tk.Button(buttons, text="Bild in Text einfügen", command=lambda: _wz442_add_image(self), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Anhang hinzufügen", command=self.kb_add_pending_attachment, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Als Entwurf speichern", command=lambda: (self.kb_status_var.set("Entwurf"), self.kb_save_entry_from_form()), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=12, ipady=4)
    tk.Button(buttons, text="Speichern", command=self.kb_save_entry_from_form, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Abbrechen", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
    form = tk.Frame(frame, bg=WHITE)
    form.pack(side="top", fill="x", padx=18, pady=(0, 6))
    tk.Label(form, text="Titel des Eintrags", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=1, column=0, columnspan=4, sticky="ew", ipady=4, pady=(0, 6))
    values = [""] + self.kb_get_categories()
    for i, var in enumerate(self.kb_entry_category_vars):
        tk.Label(form, text=f"Kategorie {i+1}", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=i, sticky="w", padx=(0 if i == 0 else 8, 0))
        box = tk.Frame(form, bg=WHITE)
        box.grid(row=3, column=i, sticky="ew", padx=(0 if i == 0 else 8, 0), pady=(0, 6))
        cb = ttk.Combobox(box, textvariable=var, values=values, state="normal", font=body_font(11), width=15)
        cb.pack(side="left", fill="x", expand=True)
        sw = tk.Label(box, width=2, relief="solid", bd=1, bg=WHITE)
        sw.pack(side="left", padx=(4, 0), ipady=6)
        def _upd(*_, v=var, s=sw):
            val = (v.get() or "").strip()
            s.configure(bg=_wz442_cat_color(self, val) if val else WHITE)
        try:
            var.trace_add("write", _upd)
        except Exception:
            pass
        cb.bind("<<ComboboxSelected>>", lambda e: (self.kb_mark_unsaved(), self.render_page()))
        _upd()
    tk.Label(form, text="Assoziierter Benutzer", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=0, sticky="w")
    tk.Entry(form, textvariable=self.kb_user_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=5, column=0, sticky="ew", pady=(0, 4))
    tk.Label(form, text="Status", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=1, sticky="w", padx=(8, 0))
    ttk.Combobox(form, textvariable=self.kb_status_var, values=["Aktiv", "Entwurf", "Veraltet"], state="readonly", font=body_font(11)).grid(row=5, column=1, sticky="ew", padx=(8, 0), pady=(0, 4))
    if any((v.get() or "").strip().lower() == "to-do" for v in self.kb_entry_category_vars):
        tk.Label(form, text="To-Do-Rhythmus", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=2, sticky="w", padx=(8, 0))
        ttk.Combobox(form, textvariable=self.kb_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(11)).grid(row=5, column=2, sticky="ew", padx=(8, 0), pady=(0, 4))
    else:
        self.kb_rhythm_var.set("")
    for col in range(4):
        form.grid_columnconfigure(col, weight=1)
    work = tk.Frame(frame, bg=WHITE)
    work.pack(fill="both", expand=True, padx=18, pady=(0, 0))
    tk.Label(work, text="Freitext / Prozessdokumentation / Leitfaden (Bilder liegen frei über dem Text)", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    overlay = tk.Frame(work, bg=WHITE)
    overlay.pack(fill="both", expand=True)
    self.kb_text_overlay_parent = overlay
    self.kb_text_widget = tk.Text(overlay, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word", undo=True)
    yscroll = tk.Scrollbar(overlay, orient="vertical", command=self.kb_text_widget.yview)
    self.kb_text_widget.configure(yscrollcommand=yscroll.set)
    try:
        _wz439_toolbar(work, self.kb_text_widget, self.kb_mark_unsaved)
    except Exception:
        pass
    self.kb_text_widget.pack(side="left", fill="both", expand=True)
    yscroll.pack(side="right", fill="y")
    self.kb_text_widget.insert("1.0", getattr(self, "kb_text_initial", ""))
    try:
        _wz439_apply_formatting(self.kb_text_widget, getattr(self, "kb_text_formatting_initial", []) or [])
    except Exception:
        pass
    self.kb_text_widget.bind("<KeyRelease>", lambda e: self.kb_mark_unsaved())
    try:
        self.kb_text_widget.bind("<MouseWheel>", lambda e: _wz439_text_mousewheel(self.kb_text_widget, e))
    except Exception:
        pass
    _wz442_place_text_images(self)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

def _wz442_render_detail(self, x, y, w, h):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry:
        self.render_kb_list_area(x, y, w, h, title="Übersicht")
        return
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text=entry.get("title", ""), bg=WHITE, fg=BLUE, font=body_font(16, weight="bold")).pack(anchor="w", padx=18, pady=(14, 5))
    tk.Label(frame, text=f"Geändert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {entry.get('user','')}    Status: {entry.get('status','')}", bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor="w", padx=18)
    try:
        _wz439_badges(frame, self, entry.get("categories", []) or [])
    except Exception:
        tk.Label(frame, text="Kategorien: " + (", ".join(entry.get("categories", []) or []) or "Keine Kategorien"), bg=WHITE, fg=TEXT, font=body_font(10, weight="bold")).pack(anchor="w", padx=18, pady=(8, 4))
    text_frame = tk.Frame(frame, bg=WHITE)
    text_frame.pack(fill="both", expand=True, padx=18, pady=(6, 8))
    self.kb_text_overlay_parent = text_frame
    txt = tk.Text(text_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word", height=8)
    yscroll = tk.Scrollbar(text_frame, orient="vertical", command=txt.yview)
    txt.configure(yscrollcommand=yscroll.set)
    txt.pack(side="left", fill="both", expand=True)
    yscroll.pack(side="right", fill="y")
    txt.insert("1.0", entry.get("text", ""))
    try:
        _wz439_apply_formatting(txt, entry.get("text_formatting", []) or [])
    except Exception:
        pass
    txt.configure(state="disabled")
    try:
        txt.bind("<MouseWheel>", lambda ev: _wz439_text_mousewheel(txt, ev))
    except Exception:
        pass
    # Bilder in der Ansicht ebenfalls als Overlay über dem Text anzeigen, nicht als separates Bildfeld.
    self.kb_inline_images = [dict(x) for x in (entry.get("inline_images", []) or []) if isinstance(x, dict)]
    _wz442_place_text_images(self)
    lower = tk.Frame(frame, bg=WHITE)
    lower.pack(fill="x", padx=18, pady=(0, 8))
    left = tk.Frame(lower, bg=WHITE)
    left.pack(side="left", fill="both", expand=True)
    right = tk.Frame(lower, bg=WHITE)
    right.pack(side="right", fill="both", expand=True, padx=(12, 0))
    tk.Label(left, text="Anhänge", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    for a in entry.get("attachments", []) or []:
        tk.Label(left, text="• " + a.get("name", ""), bg=WHITE, fg=TEXT, font=body_font(9)).pack(anchor="w")
    tk.Label(right, text="Kommentare", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    for c in (entry.get("comments", []) or [])[-3:]:
        tk.Label(right, text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}", bg=WHITE, fg=TEXT, font=body_font(9), wraplength=380, justify="left").pack(anchor="w")
    comment = tk.Text(right, height=2, bg="#F8FAFC", fg=TEXT, font=body_font(9), relief="solid", bd=1, wrap="word")
    comment.pack(fill="x", pady=(4, 0))
    buttons = tk.Frame(frame, bg=WHITE)
    buttons.pack(fill="x", padx=18, pady=(0, 14))
    if self.kb_can_create_or_edit():
        tk.Button(buttons, text="Bearbeiten", command=self.kb_edit_selected_entry, bg=WHITE, fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Kommentar speichern", command=lambda: self.kb_add_comment_to_selected(comment), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Word-Export", command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Zur Übersicht", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

# Finale Restore-Zuweisungen: Nach v0.441, damit Performance + Features gleichzeitig aktiv sind.
FiBuMateApp.render_kb_hits_pane = _wz442_render_hits_pane
FiBuMateApp.render_knowledge_work_area = _wz442_render_work_area
FiBuMateApp.render_kb_new_entry_area = _wz442_render_new_entry
FiBuMateApp.render_kb_detail_area = _wz442_render_detail
FiBuMateApp.kb_save_entry_from_form = _wz442_save_entry_from_form



# ------------------------------------------------------------------
# Wissenszentrale - Bild-Rechtsklick-Löschen FINAL 2026-06-18
# Version 0.443
# Zweck:
# - Baut auf dem aktuellen v0.442-Stand auf und entfernt keine bestehenden Features.
# - Eingefügte Bild-Overlays im Textfeld können per Rechtsklick über Kontextmenü gelöscht werden.
# - Löschen entfernt das Bild aus dem Wissenseintrag und markiert den Eintrag als ungespeichert.
# - Die zugrunde liegende Bilddatei bleibt bewusst erhalten; gelöscht wird nur die Verknüpfung im Eintrag.
# ------------------------------------------------------------------

def _wz443_delete_inline_image(self, idx):
    try:
        # Löschen soll nur im Erstellen-/Bearbeitenmodus aktiv sein.
        # In der Detailansicht würde sonst ohne Speicherkontext eine Ansicht verändert.
        if getattr(self, "knowledge_view", "") != "new":
            try:
                messagebox.showinfo("Wissenszentrale", "Bitte den Eintrag zuerst über 'Bearbeiten' öffnen, um Bilder zu löschen.")
            except Exception:
                pass
            return
        images = list(getattr(self, "kb_inline_images", []) or [])
        if idx is None or idx < 0 or idx >= len(images):
            return
        images.pop(idx)
        self.kb_inline_images = images
        self.kb_mark_unsaved()
        _wz443_place_text_images(self)
    except Exception as exc:
        try:
            messagebox.showerror("Wissenszentrale", "Bild konnte nicht gelöscht werden:\n" + str(exc))
        except Exception:
            pass

def _wz443_show_image_context_menu(self, event, idx):
    try:
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label="Bild löschen", command=lambda i=idx: _wz443_delete_inline_image(self, i))
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
    except Exception:
        # Fallback: falls Kontextmenü in der lokalen Tk-Umgebung nicht öffnet, direkt löschen.
        _wz443_delete_inline_image(self, idx)
    return "break"

def _wz443_place_text_images(self):
    parent = getattr(self, "kb_text_overlay_parent", None)
    if not parent:
        return
    try:
        for child in list(parent.children.values()):
            if getattr(child, "_wz_float_image", False):
                child.destroy()
    except Exception:
        pass
    self._kb_text_image_refs = []
    for idx, item in enumerate(getattr(self, "kb_inline_images", []) or []):
        try:
            photo = _wz442_photo(item.get("path") or item.get("source"), item.get("w", 180), item.get("h", 130))
        except Exception:
            photo = None
        if not photo:
            continue
        self._kb_text_image_refs.append(photo)
        lbl = tk.Label(parent, image=photo, bg="#F8FAFC", bd=1, relief="solid", cursor="fleur")
        lbl._wz_float_image = True
        lbl.place(x=int(item.get("x", 20)), y=int(item.get("y", 20)))

        def press(event, i=idx):
            self._kb_float_idx = i
            self._kb_float_start = (event.x_root, event.y_root)
            self._kb_float_orig = (
                int(self.kb_inline_images[i].get("x", 20)),
                int(self.kb_inline_images[i].get("y", 20)),
            )

        def drag(event, i=idx):
            # Verschieben nur im Erstellen-/Bearbeitenmodus.
            if getattr(self, "knowledge_view", "") != "new":
                return "break"
            sx, sy = getattr(self, "_kb_float_start", (event.x_root, event.y_root))
            ox, oy = getattr(self, "_kb_float_orig", (20, 20))
            self.kb_inline_images[i]["x"] = max(0, ox + event.x_root - sx)
            self.kb_inline_images[i]["y"] = max(0, oy + event.y_root - sy)
            self.kb_mark_unsaved()
            _wz443_place_text_images(self)
            return "break"

        def wheel(event, i=idx):
            # Skalieren nur im Erstellen-/Bearbeitenmodus.
            if getattr(self, "knowledge_view", "") != "new":
                return "break"
            factor = 1.08 if getattr(event, "delta", 0) > 0 else 0.92
            self.kb_inline_images[i]["w"] = max(40, min(900, int(float(self.kb_inline_images[i].get("w", 180)) * factor)))
            self.kb_inline_images[i]["h"] = max(30, min(700, int(float(self.kb_inline_images[i].get("h", 130)) * factor)))
            self.kb_mark_unsaved()
            _wz443_place_text_images(self)
            return "break"

        lbl.bind("<Button-1>", press)
        lbl.bind("<B1-Motion>", drag)
        lbl.bind("<MouseWheel>", wheel)
        lbl.bind("<Button-3>", lambda event, i=idx: _wz443_show_image_context_menu(self, event, i))

# Finale Zuweisung: bestehende v0.442-Renderlogik nutzt den globalen Namen _wz442_place_text_images.
# Dadurch bleibt die komplette v0.442-Funktionalität erhalten und nur das Bild-Handling wird erweitert.
_wz442_place_text_images = _wz443_place_text_images



# ------------------------------------------------------------------
# FiBu Mate - Lokaler Login-Fallback FINAL 2026-06-18
# Version 0.444
# Zweck:
# - Baut auf dem aktuellen v0.443-Stand auf und entfernt keine bestehenden Features.
# - Behebt lokalen Build-Login, wenn keine zentrale Benutzerdatei erreichbar/gefüllt ist.
# - Nur der technische Superuser "wagnerm" darf lokal automatisch gebootstrappt werden.
# - Keine automatische Anlage beliebiger Benutzer; vorhandene Berechtigungslogik bleibt erhalten.
# - Speichern der Benutzerdaten fällt bei nicht beschreibbarem Zentralpfad auf eine lokale Datei neben der Anwendung zurück.
# ------------------------------------------------------------------

try:
    _fm444_original_get_user_data_path = FiBuMateApp.get_user_data_path
except Exception:
    _fm444_original_get_user_data_path = None
try:
    _fm444_original_load_user_data = FiBuMateApp.load_user_data
except Exception:
    _fm444_original_load_user_data = None
try:
    _fm444_original_save_user_data = FiBuMateApp.save_user_data
except Exception:
    _fm444_original_save_user_data = None
try:
    _fm444_original_login_user = FiBuMateApp.login_user
except Exception:
    _fm444_original_login_user = None

def _fm444_local_user_data_path():
    try:
        return os.path.join(SCRIPT_DIR, "fibu_mate_users_local.json")
    except Exception:
        return r"C:\python\fibu_mate_users_local.json"

def _fm444_default_user_data():
    return {"last_username_prefill": "", "users": {}, "settings": {"auto_close_mail_enabled": True}}

def _fm444_normalize_user_data(data):
    if not isinstance(data, dict):
        data = _fm444_default_user_data()
    data.setdefault("last_username_prefill", "")
    data.setdefault("users", {})
    if not isinstance(data.get("users"), dict):
        data["users"] = {}
    data.setdefault("settings", {})
    if not isinstance(data.get("settings"), dict):
        data["settings"] = {}
    data["settings"].setdefault("auto_close_mail_enabled", True)
    return data

def _fm444_read_json_file(path):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return _fm444_normalize_user_data(json.load(f))
    except Exception:
        pass
    return None

def _fm444_write_json_file(path, data):
    folder = os.path.dirname(path)
    if folder:
        os.makedirs(folder, exist_ok=True)
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(_fm444_normalize_user_data(data), f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)

def _fm444_user_data_candidates():
    candidates = []
    try:
        p = resolve_user_data_path()
        if p:
            candidates.append(p)
    except Exception:
        pass
    try:
        if LEGACY_USER_DATA_PATH:
            candidates.append(LEGACY_USER_DATA_PATH)
    except Exception:
        pass
    candidates.append(_fm444_local_user_data_path())
    out = []
    seen = set()
    for p in candidates:
        try:
            key = os.path.abspath(p).casefold()
        except Exception:
            key = str(p).casefold()
        if p and key not in seen:
            seen.add(key)
            out.append(p)
    return out

def _fm444_load_user_data(self):
    # 1) Bestehende zentrale/konfigurierte Datei mit Benutzern bevorzugen.
    first_readable = None
    for path in _fm444_user_data_candidates():
        data = _fm444_read_json_file(path)
        if data is None:
            continue
        if first_readable is None:
            first_readable = (path, data)
        if data.get("users"):
            self._fm444_active_user_data_path = path
            return data
    # 2) Wenn eine leere Datei existiert, diese verwenden, aber den aktiven Pfad merken.
    if first_readable is not None:
        self._fm444_active_user_data_path = first_readable[0]
        return first_readable[1]
    # 3) Keine Datei erreichbar: lokale Datei neben der Anwendung erzeugen.
    data = _fm444_default_user_data()
    local_path = _fm444_local_user_data_path()
    try:
        _fm444_write_json_file(local_path, data)
        self._fm444_active_user_data_path = local_path
    except Exception:
        self._fm444_active_user_data_path = local_path
    return data

def _fm444_get_user_data_path(self):
    active = getattr(self, "_fm444_active_user_data_path", "")
    if active:
        return active
    try:
        return resolve_user_data_path()
    except Exception:
        return _fm444_local_user_data_path()

def _fm444_save_user_data(self):
    data = _fm444_normalize_user_data(getattr(self, "user_data", None))
    preferred = getattr(self, "_fm444_active_user_data_path", "") or _fm444_get_user_data_path(self)
    local_path = _fm444_local_user_data_path()
    errors = []
    for path in [preferred, local_path]:
        if not path:
            continue
        try:
            _fm444_write_json_file(path, data)
            self._fm444_active_user_data_path = path
            try:
                self._user_data_mtime = os.path.getmtime(path) if os.path.exists(path) else 0
            except Exception:
                pass
            return
        except Exception as exc:
            errors.append(f"{path}: {exc}")
    try:
        messagebox.showerror("FiBu Mate", "Benutzerdaten konnten nicht gespeichert werden:\n\n" + "\n".join(errors))
    except Exception:
        pass

def _fm444_bootstrap_superuser_if_allowed(self, username, key):
    if key != SUPERUSER_KEY:
        return False
    users = self.user_data.setdefault("users", {})
    if key in users:
        return True
    # Nur der definierte Superuser darf automatisch erzeugt werden; keine Anlage beliebiger Benutzer.
    users[key] = {
        "display_name": username or SUPERUSER_KEY,
        "first_name": "",
        "full_name": username or SUPERUSER_KEY,
        "favorites": [],
        "email": "",
        "auth": {"password_hash": None, "enabled": False},
        "permission": ROLE_E4,
    }
    try:
        self._fm444_local_bootstrap_performed = True
    except Exception:
        pass
    return True

def _fm444_login_user(self, username):
    username = " ".join(str(username).strip().split())
    key = normalize_username(username)
    if not key:
        messagebox.showwarning("FiBu Mate", "Bitte einen Nutzernamen eingeben.")
        return
    self.user_data = _fm444_normalize_user_data(getattr(self, "user_data", None))
    users = self.user_data.setdefault("users", {})
    if key not in users:
        if not _fm444_bootstrap_superuser_if_allowed(self, username, key):
            messagebox.showwarning(
                "FiBu Mate",
                "Benutzer nicht gefunden.\nBitte wende dich an eine Administratorin / einen Administrator."
            )
            return
    # Ab hier bewusst identisch zur bestehenden Login-Logik, damit keine Berechtigungs-/Featurelogik verloren geht.
    users[key].setdefault("display_name", username)
    users[key].setdefault("first_name", "")
    users[key].setdefault("full_name", " ".join(x for x in [users[key].get("first_name", "").strip(), users[key].get("display_name", username).strip()] if x).strip() or username)
    users[key].setdefault("favorites", [])
    users[key].setdefault("email", "")
    users[key].setdefault("auth", {"password_hash": None, "enabled": False})
    users[key]["permission"] = ROLE_MIGRATION.get(users[key].get("permission", ROLE_E1), ROLE_E1)
    if key == SUPERUSER_KEY:
        users[key]["permission"] = ROLE_E4
    self.ensure_permissions_defaults()
    self.current_user_key = key
    self.current_user_display = users[key].get("display_name", username)
    self.favorites = set(fav for fav in users[key].get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
    users[key]["favorites"] = sorted(self.favorites)
    self.save_user_data()
    try:
        if getattr(self, "_fm444_local_bootstrap_performed", False):
            messagebox.showinfo(
                "FiBu Mate",
                "Lokaler Entwicklungs-Login wurde initialisiert.\n\n"
                "Benutzer: wagnerm\n"
                "Berechtigung: E4 - System-Administrator\n\n"
                "Es wurden keine beliebigen Benutzer automatisch angelegt."
            )
            self._fm444_local_bootstrap_performed = False
    except Exception:
        pass
    self.start_live_permissions_refresh()
    self.page_history = []
    self.breadcrumb = []
    self.show_page("main", "Hauptmenü", add_to_history=False)

# Finale Zuweisungen: Login-Fix nach allen bisherigen Patches aktivieren.
FiBuMateApp.load_user_data = _fm444_load_user_data
FiBuMateApp.get_user_data_path = _fm444_get_user_data_path
FiBuMateApp.save_user_data = _fm444_save_user_data
FiBuMateApp.login_user = _fm444_login_user



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale DPI/Word/Installer-Version FINAL 2026-06-19
# Version 0.445
# Zweck:
# - Baut auf dem aktuellen v0.444-Stand auf und entfernt keine bestehenden Features.
# - Wissenszentrale wird gegen abweichende Windows-Skalierung stabilisiert.
# - Schrift bleibt über body_font/readability gut lesbar; nur die DPI-Aufblähung wird neutralisiert.
# - Word-Export nutzt einen robusten Exportpfad ohne APPDATA_DIR-Abhängigkeit.
# - Versionsabfrage nutzt künftig app_version.json aus der installierten/Patcher-Version statt Build-Python-Stand.
# ------------------------------------------------------------------

_FM445_TK_SCALING_100 = 96.0 / 72.0

def _fm445_safe_read_json(path, default=None):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else (default if default is not None else {})
    except Exception:
        pass
    return default if default is not None else {}

def _fm445_app_version_candidates():
    candidates = []
    try:
        candidates.append(os.path.join(SCRIPT_DIR, "config", "app_version.json"))
        candidates.append(os.path.join(SCRIPT_DIR, "app_version.json"))
        candidates.append(os.path.join(BIN_DIR, "config", "app_version.json"))
        candidates.append(os.path.join(BIN_DIR, "app_version.json"))
        candidates.append(os.path.join(CENTRAL_RELEASE_DIR, "latest.json"))
    except Exception:
        pass
    seen = set(); out = []
    for p in candidates:
        key = os.path.abspath(str(p)).casefold() if p else ""
        if key and key not in seen:
            seen.add(key); out.append(p)
    return out

def _fm445_installed_version(self=None):
    for path in _fm445_app_version_candidates():
        data = _fm445_safe_read_json(path, {})
        version = str(data.get("installed_version") or data.get("app_version") or data.get("version") or data.get("latest_version") or "").strip()
        if version:
            return version[1:] if version.lower().startswith("v") else version
    try:
        build = int(getattr(self, "version_state", {}).get("build", DEFAULT_BUILD)) if self is not None else DEFAULT_BUILD
        return f"{VERSION_PREFIX}.{build}"
    except Exception:
        return "0.0.0"

def _fm445_current_app_version(self):
    return _fm445_installed_version(self)

def _fm445_current_app_version_display(self):
    version = _fm445_installed_version(self)
    return "v" + str(version).lstrip("v")

def _fm445_version_label_text(self):
    return _fm445_current_app_version_display(self)

def _fm445_apply_wz_dpi_immunity(self):
    try:
        current = float(self.root.tk.call("tk", "scaling"))
        if not hasattr(self, "_fm445_original_tk_scaling"):
            self._fm445_original_tk_scaling = current
        if current > (_FM445_TK_SCALING_100 + 0.01):
            self.root.tk.call("tk", "scaling", _FM445_TK_SCALING_100)
            self._fm445_wz_dpi_compensated = True
        else:
            self._fm445_wz_dpi_compensated = False
    except Exception:
        pass

def _fm445_export_dir(self):
    candidates = []
    for env_name in ("APPDATA", "LOCALAPPDATA", "USERPROFILE"):
        try:
            base = os.environ.get(env_name, "").strip()
            if base:
                candidates.append(os.path.join(base, "FiBu Mate", "Wissenszentrale_Export"))
        except Exception:
            pass
    try:
        candidates.append(os.path.join(SCRIPT_DIR, "Wissenszentrale_Export"))
    except Exception:
        pass
    candidates.append(os.path.join(os.getcwd(), "Wissenszentrale_Export"))
    for path in candidates:
        try:
            os.makedirs(path, exist_ok=True)
            return path
        except Exception:
            continue
    return os.getcwd()

def _fm445_export_word(self):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry:
        try: messagebox.showwarning("Wissenszentrale", "Kein Eintrag für den Word-Export ausgewählt.")
        except Exception: pass
        return
    try:
        from docx import Document
        import re as _re
        doc = Document()
        doc.add_heading(entry.get("title", "Wissenseintrag") or "Wissenseintrag", level=1)
        meta = [
            ("Geändert", self.kb_display_date(entry.get("updated_at"))),
            ("Benutzer", entry.get("user", "")),
            ("Status", entry.get("status", "")),
            ("Kategorien", ", ".join(entry.get("categories", []) or [])),
            ("To-Do-Rhythmus", entry.get("rhythm", "")),
        ]
        for label, val in meta:
            if val:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(val))
        doc.add_heading("Inhalt", level=2)
        text_value = entry.get("text", "") or ""
        if text_value:
            for line in str(text_value).splitlines() or [""]:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("-")
        images = entry.get("inline_images", []) or []
        if images:
            doc.add_heading("Bilder", level=2)
            for img in images:
                name = img.get("name") or os.path.basename(str(img.get("path") or img.get("source") or ""))
                path = img.get("path") or img.get("source") or ""
                doc.add_paragraph(f"- {name} ({path})")
        doc.add_heading("Anhänge", level=2)
        attachments = entry.get("attachments", []) or []
        if attachments:
            for a in attachments:
                doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        else:
            doc.add_paragraph("Keine Anhänge")
        doc.add_heading("Kommentare", level=2)
        comments = entry.get("comments", []) or []
        if comments:
            for c in comments:
                doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        else:
            doc.add_paragraph("Keine Kommentare")
        safe = _re.sub(r"[^A-Za-z0-9_äöüÄÖÜß.-]+", "_", entry.get("title", "Wissenseintrag"))[:80].strip("._") or "Wissenseintrag"
        out = os.path.join(self.kb_export_dir(), f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(out)
        messagebox.showinfo("Wissenszentrale", "Word-Export erstellt:\n" + out)
    except Exception as exc:
        try: messagebox.showerror("Wissenszentrale", "Word-Export fehlgeschlagen:\n" + str(exc))
        except Exception: pass

try:
    _fm445_old_render_knowledge_work_area = FiBuMateApp.render_knowledge_work_area
    def _fm445_render_knowledge_work_area(self):
        _fm445_apply_wz_dpi_immunity(self)
        return _fm445_old_render_knowledge_work_area(self)
    FiBuMateApp.render_knowledge_work_area = _fm445_render_knowledge_work_area
except Exception:
    pass

FiBuMateApp.current_app_version = _fm445_current_app_version
FiBuMateApp.current_app_version_display = _fm445_current_app_version_display
FiBuMateApp.version_label_text = _fm445_version_label_text
FiBuMateApp.kb_export_dir = _fm445_export_dir
FiBuMateApp.kb_export_selected_to_word = _fm445_export_word



# ------------------------------------------------------------------
# Wissenszentrale - Formatleiste + Farbformat + Clipboard-Bilder FINAL 2026-06-19
# Version 0.446
# Zweck:
# - Baut auf v0.445 auf und erhält alle bestehenden Wissenszentrale-Funktionen.
# - Stellt die sichtbare Formatleiste für Einträge final wieder her.
# - Unterstützt Schriftgröße, Fett, Kursiv, Unterstrichen und Schriftfarbe Standard/Rot/Blau.
# - Speichert und lädt Textformatierungen inkl. Farbe rückwärtskompatibel zu alten Einträgen.
# - Bilder können zusätzlich per Strg+C/Strg+V aus Zwischenablage oder aus kopierten Bilddateien eingefügt werden.
# ------------------------------------------------------------------

_WZ446_COLOR_MAP = {
    "standard": "#000000",
    "rot": "#E30613",
    "blau": "#004B93",
}
_WZ446_COLOR_LABELS = {"standard": "Standard", "rot": "Rot", "blau": "Blau"}

def _wz446_color_key(value):
    raw = str(value or "standard").strip().lower()
    if raw in ("schwarz", "black", "text", "normal", "standard", ""):
        return "standard"
    if raw in ("red", "rot", "#e30613", "#dc2626"):
        return "rot"
    if raw in ("blue", "blau", "#004b93", "#2563eb"):
        return "blau"
    return "standard"

def _wz446_font_size_from_widget(widget):
    try:
        f = tkfont.Font(font=widget.cget("font"))
        size = int(f.cget("size"))
        return abs(size) if size else 10
    except Exception:
        return 10

def _wz446_tag_name(size=10, bold=False, italic=False, underline=False, color="standard"):
    return f"wzfmt_s{int(size)}_b{int(bool(bold))}_i{int(bool(italic))}_u{int(bool(underline))}_c{_wz446_color_key(color)}"

def _wz446_parse_tag(tag):
    tag = str(tag or "")
    # Neue Tags inkl. Farbe
    m = re.fullmatch(r"wzfmt_s(\d+)_b([01])_i([01])_u([01])_c([A-Za-z0-9#]+)", tag)
    if m:
        return {
            "size": int(m.group(1)),
            "bold": m.group(2) == "1",
            "italic": m.group(3) == "1",
            "underline": m.group(4) == "1",
            "color": _wz446_color_key(m.group(5)),
        }
    # Alte Tags ohne Farbe bleiben lesbar.
    m = re.fullmatch(r"wzfmt_s(\d+)_b([01])_i([01])_u([01])", tag)
    if m:
        return {
            "size": int(m.group(1)),
            "bold": m.group(2) == "1",
            "italic": m.group(3) == "1",
            "underline": m.group(4) == "1",
            "color": "standard",
        }
    return None

def _wz446_font_tuple(fmt):
    styles = []
    if fmt.get("bold"):
        styles.append("bold")
    if fmt.get("italic"):
        styles.append("italic")
    if fmt.get("underline"):
        styles.append("underline")
    return tuple(["Segoe UI", int(fmt.get("size") or 10)] + styles)

def _wz446_ensure_tag(widget, size=10, bold=False, italic=False, underline=False, color="standard"):
    key = _wz446_color_key(color)
    tag = _wz446_tag_name(size, bold, italic, underline, key)
    try:
        widget.tag_configure(tag, font=_wz446_font_tuple({"size": size, "bold": bold, "italic": italic, "underline": underline}), foreground=_WZ446_COLOR_MAP.get(key, "#000000"))
    except Exception:
        pass
    return tag

def _wz446_capture_formatting(widget):
    out = []
    try:
        for tag in widget.tag_names():
            fmt = _wz446_parse_tag(tag)
            if not fmt:
                continue
            ranges = widget.tag_ranges(tag)
            for i in range(0, len(ranges), 2):
                out.append({"start": str(ranges[i]), "end": str(ranges[i + 1]), **fmt})
    except Exception:
        pass
    return out

def _wz446_apply_formatting(widget, formatting):
    try:
        for item in formatting or []:
            if not isinstance(item, dict):
                continue
            tag = _wz446_ensure_tag(
                widget,
                item.get("size", 10),
                item.get("bold"),
                item.get("italic"),
                item.get("underline"),
                item.get("color", "standard"),
            )
            widget.tag_add(tag, item.get("start", "1.0"), item.get("end", "1.0"))
    except Exception:
        pass

def _wz446_format_at(widget, index):
    fmt = {"size": _wz446_font_size_from_widget(widget), "bold": False, "italic": False, "underline": False, "color": "standard"}
    try:
        for tag in widget.tag_names(index):
            parsed = _wz446_parse_tag(tag)
            if parsed:
                fmt.update(parsed)
    except Exception:
        pass
    return fmt

def _wz446_remove_format_tags(widget, start, end):
    try:
        for tag in list(widget.tag_names()):
            if _wz446_parse_tag(tag):
                widget.tag_remove(tag, start, end)
    except Exception:
        pass

def _wz446_apply_selection(widget, size=None, toggle=None, color=None):
    try:
        start, end = widget.index("sel.first"), widget.index("sel.last")
    except Exception:
        try:
            widget.focus_set()
        except Exception:
            pass
        return
    fmt = _wz446_format_at(widget, start)
    if size is not None:
        try:
            fmt["size"] = int(size)
        except Exception:
            pass
    if toggle in ("bold", "italic", "underline"):
        fmt[toggle] = not bool(fmt.get(toggle))
    if color is not None:
        fmt["color"] = _wz446_color_key(color)
    _wz446_remove_format_tags(widget, start, end)
    widget.tag_add(_wz446_ensure_tag(widget, **fmt), start, end)
    try:
        widget.edit_modified(True)
    except Exception:
        pass

def _wz446_clipboard_dir():
    candidates = []
    try:
        candidates.append(os.path.join(SCRIPT_DIR, "Wissenszentrale_Clipboard"))
    except Exception:
        pass
    for env in ("TEMP", "TMP", "LOCALAPPDATA"):
        try:
            base = os.environ.get(env, "").strip()
            if base:
                candidates.append(os.path.join(base, "FiBuMate", "Wissenszentrale_Clipboard"))
        except Exception:
            pass
    candidates.append(os.path.join(os.getcwd(), "Wissenszentrale_Clipboard"))
    for path in candidates:
        try:
            os.makedirs(path, exist_ok=True)
            return path
        except Exception:
            continue
    return os.getcwd()

def _wz446_is_image_file(path):
    try:
        return str(path or "").lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp")) and os.path.exists(path)
    except Exception:
        return False

def _wz446_add_clipboard_image(self, event=None):
    """Fügt Bilder aus der Zwischenablage in den Eintrag ein.
    Unterstützt echte Bilddaten aus der Zwischenablage sowie im Explorer kopierte Bilddateien.
    Gibt 'break' zurück, wenn ein Bild eingefügt wurde; normales Text-Paste bleibt unverändert möglich.
    """
    paths = []
    try:
        from PIL import ImageGrab
        data = ImageGrab.grabclipboard()
        if data is None:
            return None
        # Windows/Explorer: Liste kopierter Dateien
        if isinstance(data, (list, tuple)):
            for item in data:
                p = str(item)
                if _wz446_is_image_file(p):
                    paths.append(p)
        else:
            # Direkte Bilddaten aus Screenshot/Office/Browser etc.
            try:
                from PIL import Image
                if hasattr(data, "save"):
                    name = "clipboard_image_" + datetime.now().strftime("%Y%m%d_%H%M%S_%f") + ".png"
                    out = os.path.join(_wz446_clipboard_dir(), name)
                    img = data.convert("RGBA") if hasattr(data, "convert") else data
                    img.save(out, "PNG")
                    paths.append(out)
            except Exception:
                pass
    except Exception:
        # Fallback: manche Tk-Clipboard-Inhalte sind Dateipfade als Text.
        try:
            raw = self.root.clipboard_get()
            for part in re.split(r"[\r\n]+", raw or ""):
                p = part.strip().strip('"')
                if _wz446_is_image_file(p):
                    paths.append(p)
        except Exception:
            pass
    if not paths:
        return None
    for path in paths:
        try:
            _wz442_add_image(self, path)
        except Exception:
            pass
    try:
        self.kb_mark_unsaved()
        _wz442_place_text_images(self)
    except Exception:
        pass
    return "break"

def _wz446_toolbar(parent, widget, mark_unsaved):
    bar = tk.Frame(parent, bg=WHITE)
    bar.pack(fill="x", pady=(0, 6))
    tk.Label(bar, text="Format:", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="left", padx=(0, 6))
    size_var = tk.StringVar(value="10")
    size_cb = ttk.Combobox(bar, textvariable=size_var, values=["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "28", "32"], width=4, state="readonly", font=body_font(9))
    size_cb.pack(side="left", padx=(0, 6))
    size_cb.bind("<<ComboboxSelected>>", lambda e: (_wz446_apply_selection(widget, size=int(size_var.get())), mark_unsaved()))
    for label, key in [("B", "bold"), ("I", "italic"), ("U", "underline")]:
        tk.Button(
            bar,
            text=label,
            command=lambda k=key: (_wz446_apply_selection(widget, toggle=k), mark_unsaved()),
            bg=WHITE,
            fg=TEXT,
            font=body_font(10, weight="bold" if key == "bold" else None, underline=(key == "underline")),
            relief="solid",
            bd=1,
            width=3,
        ).pack(side="left", padx=(0, 4))
    tk.Label(bar, text="Farbe:", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="left", padx=(10, 4))
    for key in ("standard", "rot", "blau"):
        tk.Button(
            bar,
            text=_WZ446_COLOR_LABELS[key],
            command=lambda c=key: (_wz446_apply_selection(widget, color=c), mark_unsaved()),
            bg=WHITE,
            fg=_WZ446_COLOR_MAP[key],
            font=body_font(9, weight="bold" if key != "standard" else None),
            relief="solid",
            bd=1,
        ).pack(side="left", padx=(0, 4), ipadx=5)
    app = getattr(mark_unsaved, "__self__", None)
    if app is not None:
        tk.Button(
            bar,
            text="Bild aus Zwischenablage",
            command=lambda: _wz446_add_clipboard_image(app),
            bg=WHITE,
            fg=TEXT,
            font=body_font(9),
            relief="solid",
            bd=1,
        ).pack(side="left", padx=(10, 4), ipadx=6)
        try:
            widget.bind("<Control-v>", lambda e, a=app: _wz446_add_clipboard_image(a, e) or None, add="+")
            widget.bind("<Control-V>", lambda e, a=app: _wz446_add_clipboard_image(a, e) or None, add="+")
            widget.bind("<<Paste>>", lambda e, a=app: _wz446_add_clipboard_image(a, e) or None, add="+")
        except Exception:
            pass
    tk.Label(bar, text="Text markieren und Format wählen.", bg=WHITE, fg=TEXT2, font=body_font(8)).pack(side="left", padx=(8, 0))
    try:
        widget.bind("<Control-b>", lambda e: (_wz446_apply_selection(widget, toggle="bold"), mark_unsaved(), "break")[-1], add="+")
        widget.bind("<Control-i>", lambda e: (_wz446_apply_selection(widget, toggle="italic"), mark_unsaved(), "break")[-1], add="+")
        widget.bind("<Control-u>", lambda e: (_wz446_apply_selection(widget, toggle="underline"), mark_unsaved(), "break")[-1], add="+")
    except Exception:
        pass
    return bar

# Globale Kompatibilitätsnamen überschreiben, weil vorhandene Renderfunktionen diese Namen direkt aufrufen.
_wz439_tag_name = _wz446_tag_name
_wz439_parse_tag = _wz446_parse_tag
_wz439_font_tuple = _wz446_font_tuple
_wz439_ensure_tag = _wz446_ensure_tag
_wz439_capture_formatting = _wz446_capture_formatting
_wz439_apply_formatting = _wz446_apply_formatting
_wz439_format_at = _wz446_format_at
_wz439_apply_selection = _wz446_apply_selection
_wz439_toolbar = _wz446_toolbar

# Falls alte Kompatibilitätsnamen aus v0.437/v0.439 noch referenziert werden.
_wz437_parse = _wz446_parse_tag
_wz437_ensure_tag = _wz446_ensure_tag
_wz437_capture = _wz446_capture_formatting
_wz437_apply = _wz446_apply_formatting
_wz437_apply_selection = _wz446_apply_selection



# ------------------------------------------------------------------
# Wissenszentrale - Layout + sichtbare Formatleiste FINAL 2026-06-19
# Version 0.448
# Zweck:
# - Textbearbeitungs-Optionen im Editor sichtbar: Schriftgröße, Fett, Kursiv, Unterstrichen, Standard/Rot/Blau.
# - Kategorie-Filter rechts neben der Suche.
# - Treffer- und Bearbeiten-/Detailfenster größer durch früheren Start der Arbeitsbereiche.
# ------------------------------------------------------------------

def _wz448_apply_editor_format(widget, app, size=None, toggle=None, color=None):
    try:
        _wz446_apply_selection(widget, size=size, toggle=toggle, color=color)
    except Exception:
        try:
            _wz439_apply_selection(widget, size=size, toggle=toggle)
        except Exception:
            pass
    try: app.kb_mark_unsaved()
    except Exception: pass
    return "break"

def _wz448_visible_format_toolbar(parent, widget, app):
    bar=tk.Frame(parent,bg=WHITE); bar.pack(fill="x",pady=(2,5))
    tk.Label(bar,text="Textformat:",bg=WHITE,fg=TEXT2,font=body_font(9,weight="bold")).pack(side="left",padx=(0,6))
    size_var=tk.StringVar(value="10")
    cb=ttk.Combobox(bar,textvariable=size_var,values=["8","9","10","11","12","14","16","18","20","22","24","28","32"],width=4,state="readonly",font=body_font(9)); cb.pack(side="left",padx=(0,6))
    cb.bind("<<ComboboxSelected>>",lambda e:_wz448_apply_editor_format(widget,app,size=int(size_var.get())))
    for label,key in [("Fett","bold"),("Kursiv","italic"),("Unterstr.","underline")]:
        tk.Button(bar,text=label,command=lambda k=key:_wz448_apply_editor_format(widget,app,toggle=k),bg=WHITE,fg=TEXT,font=body_font(9,weight="bold" if key=="bold" else None,underline=(key=="underline")),relief="solid",bd=1).pack(side="left",padx=(0,4),ipadx=5,ipady=1)
    tk.Label(bar,text="Farbe:",bg=WHITE,fg=TEXT2,font=body_font(9)).pack(side="left",padx=(10,4))
    for label,key,fg in [("Standard","standard","#000000"),("Rot","rot",RED),("Blau","blau",BLUE)]:
        tk.Button(bar,text=label,command=lambda c=key:_wz448_apply_editor_format(widget,app,color=c),bg=WHITE,fg=fg,font=body_font(9,weight="bold" if key!="standard" else None),relief="solid",bd=1).pack(side="left",padx=(0,4),ipadx=6,ipady=1)
    try:
        tk.Button(bar,text="Bild aus Zwischenablage",command=lambda:_wz446_add_clipboard_image(app),bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).pack(side="left",padx=(10,0),ipadx=6,ipady=1)
        widget.bind("<Control-v>",lambda e:_wz446_add_clipboard_image(app,e) or None,add="+"); widget.bind("<Control-V>",lambda e:_wz446_add_clipboard_image(app,e) or None,add="+"); widget.bind("<<Paste>>",lambda e:_wz446_add_clipboard_image(app,e) or None,add="+")
    except Exception: pass
    try:
        widget.bind("<Control-b>",lambda e:_wz448_apply_editor_format(widget,app,toggle="bold"),add="+"); widget.bind("<Control-i>",lambda e:_wz448_apply_editor_format(widget,app,toggle="italic"),add="+"); widget.bind("<Control-u>",lambda e:_wz448_apply_editor_format(widget,app,toggle="underline"),add="+")
    except Exception: pass
    return bar

def _wz448_render_work_area(self):
    self.kb_ensure_state_vars()
    try: self.root.bind("<Escape>",self.kb_handle_escape,add="+"); self.root.option_add("*TCombobox*Listbox.font",body_font(11))
    except Exception: pass
    w,h=self.canvas.winfo_width(),self.canvas.winfo_height()
    nav_y=135; bx=28
    bx=self.draw_kb_button(bx,nav_y,"Start",self.kb_show_start_overlay,False,width=92); bx=self.draw_kb_button(bx,nav_y,"Übersicht",lambda:self.kb_switch_view_from_start("all"),False,width=112)
    if self.kb_can_create_or_edit(): bx=self.draw_kb_button(bx,nav_y,"Neuer Eintrag",lambda:self.kb_switch_view_from_start("new"),False,width=128)
    bx=self.draw_kb_button(bx,nav_y,"To-Dos",lambda:self.kb_switch_view_from_start("todos"),False,width=92); bx=self.draw_kb_button(bx,nav_y,"Veraltete Einträge",lambda:self.kb_switch_view_from_start("outdated"),False,width=148)
    if self.kb_can_manage_categories(): self.draw_kb_button(bx,nav_y,"Kategorien verwalten",lambda:self.kb_switch_view_from_start("categories"),False,width=168)
    ff=tk.Frame(self.root,bg=BG); self.widget_items.append(ff); ff.grid_columnconfigure(0,weight=1); ff.grid_columnconfigure(1,weight=0)
    sb=tk.Frame(ff,bg=BG); sb.grid(row=0,column=0,sticky="nsew",padx=(0,14)); tk.Label(sb,text="Suche",bg=BG,fg=TEXT2,font=body_font(9)).pack(anchor="w")
    sr=tk.Frame(sb,bg=BG); sr.pack(fill="x"); ent=tk.Entry(sr,textvariable=self.kb_search_var,font=body_font(10),bg=WHITE,fg=TEXT,relief="solid",bd=1); ent.pack(side="left",fill="x",expand=True,ipady=4); ent.bind("<Return>",self.kb_on_search_return); tk.Button(sr,text="Suchen",command=self.kb_apply_filters,bg=WHITE,fg=TEXT,relief="solid",bd=1,font=body_font(9)).pack(side="left",padx=(8,0),ipady=2)
    kb=tk.Frame(ff,bg=BG); kb.grid(row=0,column=1,sticky="ne"); tk.Label(kb,text="Kategorie-Filter",bg=BG,fg=TEXT2,font=body_font(9)).grid(row=0,column=0,columnspan=5,sticky="w")
    vals=[""]+self.kb_get_categories()
    for i,var in enumerate(self.kb_filter_vars):
        box=tk.Frame(kb,bg=BG); box.grid(row=1,column=i,padx=(0 if i==0 else 9,0),sticky="nw"); tk.Label(box,text=f"Kategorie {i+1}",bg=BG,fg=TEXT2,font=body_font(8)).pack(anchor="w"); c=ttk.Combobox(box,textvariable=var,values=vals,state="readonly",font=body_font(9),width=13); c.pack(anchor="w"); c.bind("<<ComboboxSelected>>",self.kb_apply_filters)
    selected=[(v.get() or "").strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view=="todos" or "to-do" in selected:
        box=tk.Frame(kb,bg=BG); box.grid(row=1,column=4,padx=(9,0),sticky="nw"); tk.Label(box,text="Rhythmus",bg=BG,fg=TEXT2,font=body_font(8)).pack(anchor="w"); c=ttk.Combobox(box,textvariable=self.kb_todo_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(9),width=12); c.pack(anchor="w"); c.bind("<<ComboboxSelected>>",self.kb_apply_filters)
    else:
        try: self.kb_todo_rhythm_var.set("")
        except Exception: pass
    self.canvas.create_window(ui_s(28),ui_s(172),window=ff,anchor="nw",width=ui_s(max(980,w-56)),height=ui_s(62))
    left_x,left_y=20,255; left_w=max(360,min(500,int(w*0.26))); pane_h=max(430,h-left_y-28); right_x=left_x+left_w+24; right_w=max(780,w-right_x-24)
    self.render_kb_hits_pane(left_x,left_y,left_w,pane_h)
    if self.knowledge_view=="new": self.render_kb_new_entry_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view=="todos": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="To-Dos")
    elif self.knowledge_view=="outdated": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Veraltete Einträge",status_filter="Veraltet")
    elif self.knowledge_view=="categories": self.render_kb_categories_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view=="detail": self.render_kb_detail_area(right_x,left_y,right_w,pane_h)
    else: self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Übersicht")

def _wz448_render_new_entry(self,x,y,w,h):
    if not hasattr(self,"kb_title_var"): self.kb_prepare_new_entry()
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    tk.Label(frame,text=("Eintrag bearbeiten" if getattr(self,"kb_edit_entry_id",None) else "Neuer Eintrag"),bg=WHITE,fg=BLUE,font=body_font(15,weight="bold")).pack(anchor="w",padx=18,pady=(10,4))
    buttons=tk.Frame(frame,bg=WHITE); buttons.pack(side="bottom",fill="x",padx=18,pady=(8,10))
    tk.Button(buttons,text="Bild in Text einfügen",command=lambda:_wz442_add_image(self),bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=10,ipady=3); tk.Button(buttons,text="Anhang hinzufügen",command=self.kb_add_pending_attachment,bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=10,ipady=3); tk.Button(buttons,text="Als Entwurf speichern",command=lambda:(self.kb_status_var.set("Entwurf"),self.kb_save_entry_from_form()),bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=10,ipady=3); tk.Button(buttons,text="Speichern",command=self.kb_save_entry_from_form,bg="#CFEAD6",fg=TEXT,font=body_font(9,weight="bold"),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=14,ipady=3); tk.Button(buttons,text="Abbrechen",command=lambda:self.kb_switch_view_from_start("all"),bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).pack(side="left",ipadx=14,ipady=3)
    form=tk.Frame(frame,bg=WHITE); form.pack(side="top",fill="x",padx=18,pady=(0,5)); tk.Label(form,text="Titel des Eintrags",bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=0,column=0,sticky="w"); tk.Entry(form,textvariable=self.kb_title_var,bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).grid(row=1,column=0,columnspan=4,sticky="ew",ipady=3,pady=(0,5))
    values=[""]+self.kb_get_categories()
    for i,var in enumerate(self.kb_entry_category_vars):
        tk.Label(form,text=f"Kategorie {i+1}",bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=2,column=i,sticky="w",padx=(0 if i==0 else 8,0)); box=tk.Frame(form,bg=WHITE); box.grid(row=3,column=i,sticky="ew",padx=(0 if i==0 else 8,0),pady=(0,5)); cb=ttk.Combobox(box,textvariable=var,values=values,state="normal",font=body_font(9),width=15); cb.pack(side="left",fill="x",expand=True); sw=tk.Label(box,width=2,relief="solid",bd=1,bg=WHITE); sw.pack(side="left",padx=(4,0),ipady=5)
        def _upd(*_,v=var,s=sw):
            val=(v.get() or "").strip()
            try: s.configure(bg=_wz442_cat_color(self,val) if val else WHITE)
            except Exception: s.configure(bg=WHITE)
        try: var.trace_add("write",_upd)
        except Exception: pass
        cb.bind("<<ComboboxSelected>>",lambda e:(self.kb_mark_unsaved(),self.render_page())); _upd()
    tk.Label(form,text="Assoziierter Benutzer",bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=0,sticky="w"); tk.Entry(form,textvariable=self.kb_user_var,bg=WHITE,fg=TEXT,font=body_font(9),relief="solid",bd=1).grid(row=5,column=0,sticky="ew",pady=(0,3)); tk.Label(form,text="Status",bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=1,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_status_var,values=["Aktiv","Entwurf","Veraltet"],state="readonly",font=body_font(9)).grid(row=5,column=1,sticky="ew",padx=(8,0),pady=(0,3))
    if any((v.get() or "").strip().lower()=="to-do" for v in self.kb_entry_category_vars):
        tk.Label(form,text="To-Do-Rhythmus",bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=2,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(9)).grid(row=5,column=2,sticky="ew",padx=(8,0),pady=(0,3))
    else: self.kb_rhythm_var.set("")
    for col in range(4): form.grid_columnconfigure(col,weight=1)
    work=tk.Frame(frame,bg=WHITE); work.pack(side="top",fill="both",expand=True,padx=18,pady=(0,0)); tk.Label(work,text="Freitext / Prozessdokumentation / Leitfaden (Bilder liegen frei über dem Text)",bg=WHITE,fg=TEXT2,font=body_font(8)).pack(anchor="w")
    overlay=tk.Frame(work,bg=WHITE); self.kb_text_overlay_parent=overlay; self.kb_text_widget=tk.Text(overlay,bg="#F8FAFC",fg=TEXT,font=body_font(10),relief="solid",bd=1,wrap="word",undo=True)
    _wz448_visible_format_toolbar(work,self.kb_text_widget,self)
    overlay.pack(fill="both",expand=True); yscroll=tk.Scrollbar(overlay,orient="vertical",command=self.kb_text_widget.yview); self.kb_text_widget.configure(yscrollcommand=yscroll.set); self.kb_text_widget.pack(side="left",fill="both",expand=True); yscroll.pack(side="right",fill="y")
    self.kb_text_widget.insert("1.0",getattr(self,"kb_text_initial",""))
    try: _wz446_apply_formatting(self.kb_text_widget,getattr(self,"kb_text_formatting_initial",[]) or [])
    except Exception:
        try: _wz439_apply_formatting(self.kb_text_widget,getattr(self,"kb_text_formatting_initial",[]) or [])
        except Exception: pass
    self.kb_text_widget.bind("<KeyRelease>",lambda e:self.kb_mark_unsaved())
    try: self.kb_text_widget.bind("<MouseWheel>",lambda e:_wz439_text_mousewheel(self.kb_text_widget,e))
    except Exception: pass
    _wz442_place_text_images(self)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

FiBuMateApp.render_knowledge_work_area=_wz448_render_work_area
FiBuMateApp.render_kb_new_entry_area=_wz448_render_new_entry



# ------------------------------------------------------------------
# Wissenszentrale - Editor UX FINAL 2026-06-19
# Version 0.449
# Zweck:
# - Behebt TclError invalid command name durch entfernte Kategorie-Farb-Labels.
# - Vier Kategorieauswahlen im Editor mit farbig hinterlegter Auswahl und farbigem Dropdown.
# - Ungespeicherter Text wird vor internen Re-Renders erhalten; Wechsel fragt Speichern/Verwerfen/Abbrechen.
# - Standard-Freitextgröße 16; Formatänderung ohne Markierung wirkt für neu zu schreibenden Text.
# - Bilder mit Rahmen-/Ecken-Resize, Rechtsklick: Größe wiederherstellen / Bild löschen.
# - Tabellen einfügen mit Popup; Tabellen/Zeilen/Spalten grob per Maus vergrößerbar und Rechtsklick-Menü.
# - Kopfdaten kompakter und ein-/ausklappbar.
# ------------------------------------------------------------------

_WZ449_DEFAULT_TEXT_SIZE = 16

def _wz449_safe_cat_color(app, name):
    try:
        return app.kb_get_category_color(name) if name else WHITE
    except Exception:
        try:
            return _wz442_cat_color(app, name) if name else WHITE
        except Exception:
            return WHITE

def _wz449_cat_fg(color):
    try:
        return _wz_cat_fg(color)
    except Exception:
        return TEXT

def _wz449_capture_editor_state(self):
    """Sichert aktuellen Editorzustand vor internen Re-Renders, damit Text nie still verschwindet."""
    try:
        if hasattr(self, 'kb_text_widget') and self.kb_text_widget.winfo_exists():
            self.kb_text_initial = self.kb_text_widget.get('1.0', 'end-1c')
            try:
                self.kb_text_formatting_initial = _wz446_capture_formatting(self.kb_text_widget)
            except Exception:
                try: self.kb_text_formatting_initial = _wz439_capture_formatting(self.kb_text_widget)
                except Exception: pass
    except Exception:
        pass
    try:
        self.kb_inline_images = [dict(x) for x in (getattr(self, 'kb_inline_images', []) or []) if isinstance(x, dict)]
    except Exception:
        pass
    try:
        self.kb_inline_tables = [dict(x) for x in (getattr(self, 'kb_inline_tables', []) or []) if isinstance(x, dict)]
    except Exception:
        pass

def _wz449_confirm_unsaved_before_destructive(self):
    try:
        _wz449_capture_editor_state(self)
    except Exception:
        pass
    if not bool(getattr(self, 'knowledge_unsaved', False)):
        return True
    try:
        result = messagebox.askyesnocancel('Ungespeicherte Änderungen', 'Der aktuelle Eintrag enthält ungespeicherte Änderungen.\n\nJa = speichern\nNein = verwerfen\nAbbrechen = im Eintrag bleiben')
    except Exception:
        result = None
    if result is True:
        try:
            return bool(self.kb_save_entry_from_form())
        except Exception:
            return False
    if result is False:
        self.knowledge_unsaved = False
        return True
    return False

try:
    _wz449_old_switch_view = FiBuMateApp.kb_switch_view_from_start
    def _wz449_switch_view_from_start(self, view):
        if getattr(self, 'knowledge_view', '') == 'new' and not _wz449_confirm_unsaved_before_destructive(self):
            return
        return _wz449_old_switch_view(self, view)
    FiBuMateApp.kb_switch_view_from_start = _wz449_switch_view_from_start
except Exception:
    pass

def _wz449_prepare_entry(self, entry=None):
    self.kb_ensure_state_vars()
    entry = entry or {}
    self.kb_edit_entry_id = entry.get('id') if entry else None
    self.kb_title_var = tk.StringVar(value=entry.get('title', ''))
    cats = list(entry.get('categories', []) or [])[:4]
    self.kb_entry_category_vars = [tk.StringVar(value=(cats[i] if i < len(cats) else '')) for i in range(4)]
    self.kb_user_var = tk.StringVar(value=entry.get('user', getattr(self, 'current_user_display', '') or getattr(self, 'current_user_key', '') or ''))
    self.kb_status_var = tk.StringVar(value=entry.get('status', 'Aktiv'))
    self.kb_rhythm_var = tk.StringVar(value=entry.get('rhythm', ''))
    self.kb_text_initial = entry.get('text', '')
    self.kb_text_formatting_initial = entry.get('text_formatting', []) or []
    self.kb_inline_images = [dict(x) for x in (entry.get('inline_images', []) or []) if isinstance(x, dict)]
    self.kb_inline_tables = [dict(x) for x in (entry.get('inline_tables', []) or []) if isinstance(x, dict)]
    self.kb_pending_attachments = []
    self.kb_header_collapsed = bool(getattr(self, 'kb_header_collapsed', False))
    self.kb_current_text_format = {'size': _WZ449_DEFAULT_TEXT_SIZE, 'bold': False, 'italic': False, 'underline': False, 'color': 'standard'}
    self.knowledge_unsaved = False

FiBuMateApp.kb_prepare_new_entry = _wz449_prepare_entry

def _wz449_colored_category_selector(parent, app, var, values, width=19, allow_custom=True, mark=True):
    box = tk.Frame(parent, bg=WHITE)
    entry = tk.Entry(box, textvariable=var, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1, width=width)
    entry.pack(side='left', fill='x', expand=True, ipady=3)
    btn = tk.Menubutton(box, text='▼', bg=WHITE, fg=TEXT, relief='solid', bd=1, font=body_font(8), width=2)
    menu = tk.Menu(btn, tearoff=False)
    btn.configure(menu=menu)
    btn.pack(side='left', padx=(3,0), ipady=1)
    def apply_bg(*_):
        try:
            if not entry.winfo_exists(): return
            val = (var.get() or '').strip()
            color = _wz449_safe_cat_color(app, val)
            entry.configure(bg=color if val else WHITE, fg=_wz449_cat_fg(color) if val else TEXT)
        except Exception:
            pass
    def select(value):
        var.set(value)
        apply_bg()
        if mark:
            try: app.kb_mark_unsaved()
            except Exception: pass
    try:
        for value in values:
            color = _wz449_safe_cat_color(app, value)
            if not value:
                menu.add_command(label='(leer)', command=lambda v='': select(v))
            else:
                menu.add_command(label='  ' + str(value), command=lambda v=value: select(v), background=color, foreground=_wz449_cat_fg(color), activebackground=color, activeforeground=_wz449_cat_fg(color))
    except Exception:
        pass
    if allow_custom:
        entry.bind('<KeyRelease>', lambda e: (apply_bg(), app.kb_mark_unsaved()))
    else:
        entry.configure(state='readonly')
    apply_bg()
    return box

def _wz449_text_tag(widget, fmt):
    try:
        return _wz446_ensure_tag(widget, fmt.get('size', _WZ449_DEFAULT_TEXT_SIZE), fmt.get('bold'), fmt.get('italic'), fmt.get('underline'), fmt.get('color', 'standard'))
    except Exception:
        return ''

def _wz449_apply_editor_format(widget, app, size=None, toggle=None, color=None):
    fmt = dict(getattr(app, 'kb_current_text_format', {'size': _WZ449_DEFAULT_TEXT_SIZE, 'bold': False, 'italic': False, 'underline': False, 'color': 'standard'}))
    try:
        start, end = widget.index('sel.first'), widget.index('sel.last')
        try:
            at = _wz446_format_at(widget, start)
            if at: fmt.update(at)
        except Exception: pass
        has_selection = True
    except Exception:
        start = end = None
        has_selection = False
    if size is not None:
        try: fmt['size'] = int(size)
        except Exception: pass
    if toggle in ('bold','italic','underline'):
        fmt[toggle] = not bool(fmt.get(toggle))
    if color is not None:
        try: fmt['color'] = _wz446_color_key(color)
        except Exception: fmt['color'] = color
    app.kb_current_text_format = dict(fmt)
    if has_selection:
        try:
            _wz446_remove_format_tags(widget, start, end)
            widget.tag_add(_wz449_text_tag(widget, fmt), start, end)
        except Exception:
            pass
    try: app.kb_mark_unsaved()
    except Exception: pass
    return 'break'

def _wz449_apply_current_format_to_typed_char(widget, app, event=None):
    try:
        if event and getattr(event, 'keysym', '') in ('BackSpace','Delete','Left','Right','Up','Down','Home','End','Prior','Next','Shift_L','Shift_R','Control_L','Control_R','Alt_L','Alt_R'):
            return
        fmt = getattr(app, 'kb_current_text_format', None)
        if not fmt: return
        tag = _wz449_text_tag(widget, fmt)
        if tag:
            widget.tag_add(tag, 'insert-1c', 'insert')
    except Exception:
        pass

def _wz449_visible_format_toolbar(parent, widget, app):
    bar = tk.Frame(parent, bg=WHITE)
    bar.pack(fill='x', pady=(1,4))
    tk.Label(bar, text='Textformat:', bg=WHITE, fg=TEXT2, font=body_font(8, weight='bold')).pack(side='left', padx=(0,5))
    size_var = tk.StringVar(value=str(getattr(app, 'kb_current_text_format', {}).get('size', _WZ449_DEFAULT_TEXT_SIZE)))
    cb = ttk.Combobox(bar, textvariable=size_var, values=['8','9','10','11','12','14','16','18','20','22','24','28','32'], width=4, state='readonly', font=body_font(8))
    cb.pack(side='left', padx=(0,5))
    cb.bind('<<ComboboxSelected>>', lambda e: _wz449_apply_editor_format(widget, app, size=int(size_var.get())))
    for label,key in [('Fett','bold'),('Kursiv','italic'),('Unterstr.','underline')]:
        tk.Button(bar, text=label, command=lambda k=key:_wz449_apply_editor_format(widget, app, toggle=k), bg=WHITE, fg=TEXT, font=body_font(8, weight='bold' if key=='bold' else None, underline=(key=='underline')), relief='solid', bd=1).pack(side='left', padx=(0,3), ipadx=4, ipady=1)
    tk.Label(bar, text='Farbe:', bg=WHITE, fg=TEXT2, font=body_font(8)).pack(side='left', padx=(7,3))
    for label,key,fg in [('Standard','standard','#000000'),('Rot','rot',RED),('Blau','blau',BLUE)]:
        tk.Button(bar, text=label, command=lambda c=key:_wz449_apply_editor_format(widget, app, color=c), bg=WHITE, fg=fg, font=body_font(8, weight='bold' if key!='standard' else None), relief='solid', bd=1).pack(side='left', padx=(0,3), ipadx=4, ipady=1)
    tk.Button(bar, text='Bild aus Zwischenablage', command=lambda:_wz446_add_clipboard_image(app), bg=WHITE, fg=TEXT, font=body_font(8), relief='solid', bd=1).pack(side='left', padx=(7,3), ipadx=5, ipady=1)
    tk.Button(bar, text='Tabelle einfügen', command=lambda:_wz449_table_popup(app), bg=WHITE, fg=TEXT, font=body_font(8), relief='solid', bd=1).pack(side='left', padx=(0,3), ipadx=5, ipady=1)
    try:
        widget.bind('<KeyRelease>', lambda e: (_wz449_apply_current_format_to_typed_char(widget, app, e), app.kb_mark_unsaved()), add='+')
        widget.bind('<Control-b>', lambda e:_wz449_apply_editor_format(widget, app, toggle='bold'), add='+')
        widget.bind('<Control-i>', lambda e:_wz449_apply_editor_format(widget, app, toggle='italic'), add='+')
        widget.bind('<Control-u>', lambda e:_wz449_apply_editor_format(widget, app, toggle='underline'), add='+')
        widget.bind('<Control-v>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
        widget.bind('<Control-V>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
        widget.bind('<<Paste>>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
    except Exception:
        pass
    return bar

def _wz449_photo(path, max_w=180, max_h=130):
    try: return _wz441_cached_photo(path, max_w, max_h)
    except Exception: pass
    try:
        if not PIL_AVAILABLE or not path or not os.path.exists(path): return None
        img = Image.open(path); img.thumbnail((int(max_w), int(max_h)))
        return ImageTk.PhotoImage(img)
    except Exception:
        return None

def _wz449_add_image(self, path=None):
    if not path:
        try:
            from tkinter import filedialog
            path = filedialog.askopenfilename(title='Bild in Text einfügen', filetypes=[('Bilder','*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff *.webp'),('Alle Dateien','*.*')])
        except Exception: path = ''
    if not path: return
    w0,h0 = 180,130
    try:
        if PIL_AVAILABLE and os.path.exists(path):
            im = Image.open(path); ow,oh = im.size; scale = min(1.0, 220/max(1,ow), 160/max(1,oh)); w0,h0 = max(40,int(ow*scale)), max(30,int(oh*scale))
    except Exception: pass
    imgs = list(getattr(self,'kb_inline_images',[]) or [])
    imgs.append({'source':path,'path':path,'name':os.path.basename(path),'x':30+len(imgs)*18,'y':55+len(imgs)*12,'w':w0,'h':h0,'orig_w':w0,'orig_h':h0})
    self.kb_inline_images = imgs
    try: self.kb_mark_unsaved()
    except Exception: pass
    _wz449_place_objects(self)

_wz442_add_image = _wz449_add_image

def _wz449_image_context(self, idx, event):
    menu = tk.Menu(self.root, tearoff=False)
    def restore():
        try:
            it=self.kb_inline_images[idx]; it['w']=int(it.get('orig_w',180)); it['h']=int(it.get('orig_h',130)); self.kb_mark_unsaved(); _wz449_place_objects(self)
        except Exception: pass
    def delete():
        try:
            self.kb_inline_images.pop(idx); self.kb_mark_unsaved(); _wz449_place_objects(self)
        except Exception: pass
    menu.add_command(label='Größe wiederherstellen', command=restore)
    menu.add_command(label='Bild löschen', command=delete)
    try: menu.tk_popup(event.x_root, event.y_root)
    finally:
        try: menu.grab_release()
        except Exception: pass
    return 'break'

def _wz449_place_image(self, parent, idx, item):
    x,y,w,h = int(item.get('x',20)), int(item.get('y',20)), int(item.get('w',180)), int(item.get('h',130))
    frame = tk.Frame(parent, bg=BLUE, bd=1, relief='solid', cursor='fleur')
    frame._wz_float_image=True
    frame.place(x=x, y=y, width=max(35,w), height=max(30,h))
    photo = _wz449_photo(item.get('path') or item.get('source'), max(20,w-4), max(20,h-4))
    if photo:
        self._kb_text_image_refs.append(photo)
        lbl = tk.Label(frame, image=photo, bg='#F8FAFC')
    else:
        lbl = tk.Label(frame, text='Bild', bg='#FEE2E2', fg=RED, font=body_font(8))
    lbl.place(x=2,y=2,width=max(10,w-4),height=max(10,h-4))
    def start_move(e):
        self._wz449_drag=('move_img',idx,e.x_root,e.y_root,int(item.get('x',20)),int(item.get('y',20)),0,0)
    def do_drag(e):
        d=getattr(self,'_wz449_drag',None)
        if not d or d[0]!='move_img': return
        _,i,sx,sy,ox,oy,_,_=d; self.kb_inline_images[i]['x']=max(0,ox+e.x_root-sx); self.kb_inline_images[i]['y']=max(0,oy+e.y_root-sy); self.kb_mark_unsaved(); _wz449_place_objects(self)
    frame.bind('<Button-1>', start_move); lbl.bind('<Button-1>', start_move)
    frame.bind('<B1-Motion>', do_drag); lbl.bind('<B1-Motion>', do_drag)
    frame.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e)); lbl.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e))
    handles=[('nw',0,0),('n',w//2-4,0),('ne',w-8,0),('e',w-8,h//2-4),('se',w-8,h-8),('s',w//2-4,h-8),('sw',0,h-8),('w',0,h//2-4)]
    for name,hx,hy in handles:
        hd=tk.Frame(frame,bg=WHITE,highlightbackground=BLUE,highlightthickness=1,cursor='sizing')
        hd.place(x=max(0,hx),y=max(0,hy),width=8,height=8)
        def hs(e, nm=name):
            self._wz449_drag=('resize_img',idx,e.x_root,e.y_root,int(item.get('x',20)),int(item.get('y',20)),int(item.get('w',180)),int(item.get('h',130)),nm)
        def hr(e):
            d=getattr(self,'_wz449_drag',None)
            if not d or d[0]!='resize_img': return
            _,i,sx,sy,ox,oy,ow,oh,nm=d; dx=e.x_root-sx; dy=e.y_root-sy; nx,ny,nw,nh=ox,oy,ow,oh
            if 'e' in nm: nw=max(30,ow+dx)
            if 's' in nm: nh=max(25,oh+dy)
            if 'w' in nm: nx=max(0,ox+dx); nw=max(30,ow-dx)
            if 'n' in nm: ny=max(0,oy+dy); nh=max(25,oh-dy)
            self.kb_inline_images[i].update({'x':nx,'y':ny,'w':nw,'h':nh}); self.kb_mark_unsaved(); _wz449_place_objects(self)
        hd.bind('<Button-1>',hs); hd.bind('<B1-Motion>',hr); hd.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e))

def _wz449_default_table(rows, cols):
    rows=max(1,int(rows)); cols=max(1,int(cols))
    return {'x':40,'y':80,'col_widths':[100]*cols,'row_heights':[28]*rows,'cells':[['' for _ in range(cols)] for __ in range(rows)]}

def _wz449_table_popup(app):
    try:
        win=tk.Toplevel(app.root); win.title('Tabelle einfügen'); win.configure(bg=WHITE); win.resizable(False,False)
        rows=tk.IntVar(value=3); cols=tk.IntVar(value=3)
        tk.Label(win,text='Zeilenanzahl',bg=WHITE,fg=TEXT,font=body_font(10)).grid(row=0,column=0,padx=14,pady=(14,4),sticky='w')
        tk.Spinbox(win,from_=1,to=30,textvariable=rows,width=8,font=body_font(10)).grid(row=0,column=1,padx=14,pady=(14,4))
        tk.Label(win,text='Spaltenanzahl',bg=WHITE,fg=TEXT,font=body_font(10)).grid(row=1,column=0,padx=14,pady=4,sticky='w')
        tk.Spinbox(win,from_=1,to=20,textvariable=cols,width=8,font=body_font(10)).grid(row=1,column=1,padx=14,pady=4)
        def insert():
            tables=list(getattr(app,'kb_inline_tables',[]) or []); t=_wz449_default_table(rows.get(),cols.get()); t['x']=50+len(tables)*20; t['y']=90+len(tables)*20; tables.append(t); app.kb_inline_tables=tables; app.kb_mark_unsaved(); _wz449_place_objects(app); win.destroy()
        br=tk.Frame(win,bg=WHITE); br.grid(row=2,column=0,columnspan=2,pady=14)
        tk.Button(br,text='Einfügen',command=insert,bg='#CFEAD6',fg=TEXT,font=body_font(10,weight='bold'),relief='solid',bd=1).pack(side='left',padx=(0,8),ipadx=12,ipady=3)
        tk.Button(br,text='Verwerfen',command=win.destroy,bg=WHITE,fg=TEXT,font=body_font(10),relief='solid',bd=1).pack(side='left',ipadx=12,ipady=3)
        win.transient(app.root); win.grab_set(); win.focus_set()
    except Exception as exc:
        try: messagebox.showerror('Wissenszentrale','Tabelle konnte nicht eingefügt werden:\n'+str(exc))
        except Exception: pass

def _wz449_table_context(self, idx, event, row=None, col=None):
    menu=tk.Menu(self.root,tearoff=False)
    def redraw():
        self.kb_mark_unsaved(); _wz449_place_objects(self)
    def del_table():
        try: self.kb_inline_tables.pop(idx); redraw()
        except Exception: pass
    def add_row():
        t=self.kb_inline_tables[idx]; c=len(t.get('col_widths',[])) or 1; t.setdefault('cells',[]).append(['' for _ in range(c)]); t.setdefault('row_heights',[]).append(28); redraw()
    def add_col():
        t=self.kb_inline_tables[idx]; t.setdefault('col_widths',[]).append(100)
        for r in t.setdefault('cells',[]): r.append('')
        redraw()
    def del_row():
        if row is None: return
        t=self.kb_inline_tables[idx]
        if len(t.get('cells',[]))>1: t['cells'].pop(row); t.get('row_heights',[]).pop(row); redraw()
    def del_col():
        if col is None: return
        t=self.kb_inline_tables[idx]
        if len(t.get('col_widths',[]))>1:
            t['col_widths'].pop(col)
            for r in t.get('cells',[]):
                if col < len(r): r.pop(col)
            redraw()
    menu.add_command(label='Tabelle löschen', command=del_table)
    menu.add_separator(); menu.add_command(label='Spalten anfügen', command=add_col); menu.add_command(label='Zeilen anfügen', command=add_row)
    if row is not None or col is not None:
        menu.add_separator()
        if col is not None: menu.add_command(label='Spalte löschen', command=del_col)
        if row is not None: menu.add_command(label='Zeile löschen', command=del_row)
    try: menu.tk_popup(event.x_root,event.y_root)
    finally:
        try: menu.grab_release()
        except Exception: pass
    return 'break'

def _wz449_place_table(self,parent,idx,t):
    x,y=int(t.get('x',40)),int(t.get('y',80)); colw=[int(v) for v in t.get('col_widths',[100])]; rowh=[int(v) for v in t.get('row_heights',[28])]; cells=t.get('cells',[]) or []
    rows=len(rowh); cols=len(colw); totalw=sum(colw); totalh=sum(rowh)
    fr=tk.Frame(parent,bg=LINE,bd=1,relief='solid',cursor='fleur'); fr._wz_float_table=True; fr.place(x=x,y=y,width=max(60,totalw+2),height=max(30,totalh+2))
    def move_start(e): self._wz449_drag=('move_tbl',idx,e.x_root,e.y_root,x,y,0,0)
    def move_drag(e):
        d=getattr(self,'_wz449_drag',None)
        if not d or d[0]!='move_tbl': return
        _,i,sx,sy,ox,oy,_,_=d; self.kb_inline_tables[i]['x']=max(0,ox+e.x_root-sx); self.kb_inline_tables[i]['y']=max(0,oy+e.y_root-sy); self.kb_mark_unsaved(); _wz449_place_objects(self)
    fr.bind('<Button-1>',move_start); fr.bind('<B1-Motion>',move_drag); fr.bind('<Button-3>',lambda e,i=idx:_wz449_table_context(self,i,e))
    yy=1
    for r in range(rows):
        xx=1
        for c in range(cols):
            val=''
            try: val=cells[r][c]
            except Exception: pass
            cell=tk.Entry(fr,bg=WHITE,fg=TEXT,font=body_font(10),relief='solid',bd=1)
            cell.insert(0,val); cell.place(x=xx,y=yy,width=max(25,colw[c]),height=max(20,rowh[r]))
            def save_cell(e, rr=r, cc=c, ce=cell):
                try:
                    while len(self.kb_inline_tables[idx].setdefault('cells',[]))<=rr: self.kb_inline_tables[idx]['cells'].append(['' for _ in range(cols)])
                    while len(self.kb_inline_tables[idx]['cells'][rr])<=cc: self.kb_inline_tables[idx]['cells'][rr].append('')
                    self.kb_inline_tables[idx]['cells'][rr][cc]=ce.get(); self.kb_mark_unsaved()
                except Exception: pass
            cell.bind('<KeyRelease>',save_cell); cell.bind('<Button-3>',lambda e,rr=r,cc=c,i=idx:_wz449_table_context(self,i,e,rr,cc))
            xx+=colw[c]
        yy+=rowh[r]
    # Spalten-/Zeilen-Resize-Zonen
    xx=1
    for c,wc in enumerate(colw[:-1]):
        xx+=wc; hd=tk.Frame(fr,bg=BLUE,cursor='sb_h_double_arrow'); hd.place(x=xx-2,y=1,width=4,height=max(20,totalh))
        def cs(e, cc=c): self._wz449_drag=('col_resize',idx,cc,e.x_root,0,self.kb_inline_tables[idx]['col_widths'][cc],self.kb_inline_tables[idx]['col_widths'][cc+1],0)
        def cd(e):
            d=getattr(self,'_wz449_drag',None)
            if not d or d[0]!='col_resize': return
            _,i,cc,sx,_,w1,w2,_=d; dx=e.x_root-sx; self.kb_inline_tables[i]['col_widths'][cc]=max(25,w1+dx); self.kb_inline_tables[i]['col_widths'][cc+1]=max(25,w2-dx); self.kb_mark_unsaved(); _wz449_place_objects(self)
        hd.bind('<Button-1>',cs); hd.bind('<B1-Motion>',cd)
    yy=1
    for r,hr in enumerate(rowh[:-1]):
        yy+=hr; hd=tk.Frame(fr,bg=BLUE,cursor='sb_v_double_arrow'); hd.place(x=1,y=yy-2,width=max(40,totalw),height=4)
        def rs(e, rr=r): self._wz449_drag=('row_resize',idx,rr,e.y_root,0,self.kb_inline_tables[idx]['row_heights'][rr],self.kb_inline_tables[idx]['row_heights'][rr+1],0)
        def rd(e):
            d=getattr(self,'_wz449_drag',None)
            if not d or d[0]!='row_resize': return
            _,i,rr,sy,_,h1,h2,_=d; dy=e.y_root-sy; self.kb_inline_tables[i]['row_heights'][rr]=max(20,h1+dy); self.kb_inline_tables[i]['row_heights'][rr+1]=max(20,h2-dy); self.kb_mark_unsaved(); _wz449_place_objects(self)
        hd.bind('<Button-1>',rs); hd.bind('<B1-Motion>',rd)

def _wz449_place_objects(self):
    parent=getattr(self,'kb_text_overlay_parent',None)
    if not parent: return
    for child in list(parent.children.values()):
        if getattr(child,'_wz_float_image',False) or getattr(child,'_wz_float_table',False):
            try: child.destroy()
            except Exception: pass
    self._kb_text_image_refs=[]
    for idx,item in enumerate(getattr(self,'kb_inline_images',[]) or []):
        try: _wz449_place_image(self,parent,idx,item)
        except Exception: pass
    for idx,t in enumerate(getattr(self,'kb_inline_tables',[]) or []):
        try: _wz449_place_table(self,parent,idx,t)
        except Exception: pass

_wz442_place_text_images = _wz449_place_objects

def _wz449_normalize_images(self, entry_id):
    out=[]; target=os.path.join(self.kb_attachment_dir(entry_id),'InlineImages')
    try: os.makedirs(target,exist_ok=True)
    except Exception: pass
    import shutil as _shutil
    for item in getattr(self,'kb_inline_images',[]) or []:
        if not isinstance(item,dict): continue
        src=item.get('source') or item.get('path'); name=item.get('name') or os.path.basename(str(src or 'bild')); dst=item.get('path') or src
        try:
            if src and os.path.exists(src):
                safe=re.sub(r'[^A-Za-z0-9_.äöüÄÖÜß-]+','_',os.path.basename(name))[:100]; dst=os.path.join(target,safe)
                if os.path.abspath(src)!=os.path.abspath(dst): _shutil.copy2(src,dst)
        except Exception: pass
        out.append({'name':name,'path':dst,'x':int(item.get('x',20)),'y':int(item.get('y',20)),'w':int(item.get('w',180)),'h':int(item.get('h',130)),'orig_w':int(item.get('orig_w',item.get('w',180))),'orig_h':int(item.get('orig_h',item.get('h',130)))})
    return out

_wz442_normalize_images = _wz449_normalize_images

def _wz449_save_entry_from_form(self):
    self.kb_ensure_state_vars(); _wz449_capture_editor_state(self)
    title=(self.kb_title_var.get() or '').strip()
    if not title:
        try: messagebox.showwarning('Wissenszentrale','Bitte einen Titel erfassen.')
        except Exception: pass
        return False
    categories=[]
    for var in getattr(self,'kb_entry_category_vars',[]) or []:
        value=(var.get() or '').strip()
        if value and value not in categories: categories.append(value)
    categories=categories[:4]
    user=(self.kb_user_var.get() or '').strip(); status=(self.kb_status_var.get() or 'Aktiv').strip(); rhythm=(self.kb_rhythm_var.get() or '').strip() if any(c.lower()=='to-do' for c in categories) else ''
    text_value=getattr(self,'kb_text_initial',''); formatting=getattr(self,'kb_text_formatting_initial',[]) or []
    entries=self.kb_load_entries(); now=self.kb_now() if hasattr(self,'kb_now') else datetime.now().isoformat(timespec='seconds')
    selected_id=getattr(self,'kb_edit_entry_id',None) or (self.kb_make_entry_id() if hasattr(self,'kb_make_entry_id') else hashlib.sha1((title+now).encode('utf-8')).hexdigest()[:16])
    inline_images=_wz449_normalize_images(self,selected_id); inline_tables=[dict(t) for t in (getattr(self,'kb_inline_tables',[]) or []) if isinstance(t,dict)]
    found=False
    for entry in entries:
        if entry.get('id')==selected_id:
            entry.update({'title':title,'categories':categories,'user':user,'status':status,'rhythm':rhythm,'text':text_value,'text_formatting':formatting,'inline_images':inline_images,'inline_tables':inline_tables,'updated_at':now}); found=True; break
    if not found:
        entries.append({'id':selected_id,'title':title,'categories':categories,'user':user,'status':status,'rhythm':rhythm,'text':text_value,'text_formatting':formatting,'inline_images':inline_images,'inline_tables':inline_tables,'created_at':now,'updated_at':now,'comments':[],'attachments':[]})
    pending=list(getattr(self,'kb_pending_attachments',[]) or [])
    if pending:
        import shutil as _shutil
        copied=[]; target=self.kb_attachment_dir(selected_id)
        for src in pending:
            try:
                name=os.path.basename(src); dst=os.path.join(target,name); _shutil.copy2(src,dst); copied.append({'name':name,'path':dst,'added_at':now})
            except Exception: pass
        for entry in entries:
            if entry.get('id')==selected_id: entry.setdefault('attachments',[]).extend(copied); break
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id=selected_id; self.kb_edit_entry_id=None; self.kb_pending_attachments=[]; self.knowledge_unsaved=False; self.knowledge_view='detail'; self.render_page(); return True
    return False

FiBuMateApp.kb_save_entry_from_form = _wz449_save_entry_from_form

def _wz449_toggle_header(self):
    _wz449_capture_editor_state(self)
    self.kb_header_collapsed = not bool(getattr(self,'kb_header_collapsed',False))
    self.render_page()

def _wz449_render_new_entry(self,x,y,w,h):
    if not hasattr(self,'kb_title_var'): self.kb_prepare_new_entry()
    # Alte trace_add-Callbacks auf bereits gelöschte Labels neutralisieren: Kategorie-Variablen neu aufbauen.
    try:
        old=[v.get() for v in getattr(self,'kb_entry_category_vars',[])]; self.kb_entry_category_vars=[tk.StringVar(value=(old[i] if i<len(old) else '')) for i in range(4)]
    except Exception: pass
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    top=tk.Frame(frame,bg=WHITE); top.pack(fill='x',padx=14,pady=(8,2))
    tk.Label(top,text=('Eintrag bearbeiten' if getattr(self,'kb_edit_entry_id',None) else 'Neuer Eintrag'),bg=WHITE,fg=BLUE,font=body_font(13,weight='bold')).pack(side='left')
    tk.Button(top,text=('Kopfdaten anzeigen' if getattr(self,'kb_header_collapsed',False) else 'Kopfdaten ausblenden'),command=lambda:_wz449_toggle_header(self),bg=WHITE,fg=TEXT,font=body_font(8),relief='solid',bd=1).pack(side='right',ipadx=8,ipady=1)
    buttons=tk.Frame(frame,bg=WHITE); buttons.pack(side='bottom',fill='x',padx=14,pady=(6,8))
    tk.Button(buttons,text='Bild in Text einfügen',command=lambda:_wz449_add_image(self),bg=WHITE,fg=TEXT,font=body_font(8),relief='solid',bd=1).pack(side='left',padx=(0,6),ipadx=8,ipady=2)
    tk.Button(buttons,text='Anhang hinzufügen',command=self.kb_add_pending_attachment,bg=WHITE,fg=TEXT,font=body_font(8),relief='solid',bd=1).pack(side='left',padx=(0,6),ipadx=8,ipady=2)
    tk.Button(buttons,text='Als Entwurf speichern',command=lambda:(self.kb_status_var.set('Entwurf'),self.kb_save_entry_from_form()),bg=WHITE,fg=TEXT,font=body_font(8),relief='solid',bd=1).pack(side='left',padx=(0,6),ipadx=8,ipady=2)
    tk.Button(buttons,text='Speichern',command=self.kb_save_entry_from_form,bg='#CFEAD6',fg=TEXT,font=body_font(8,weight='bold'),relief='solid',bd=1).pack(side='left',padx=(0,6),ipadx=12,ipady=2)
    tk.Button(buttons,text='Abbrechen',command=lambda:self.kb_switch_view_from_start('all'),bg=WHITE,fg=TEXT,font=body_font(8),relief='solid',bd=1).pack(side='left',ipadx=12,ipady=2)
    if not getattr(self,'kb_header_collapsed',False):
        form=tk.Frame(frame,bg=WHITE); form.pack(side='top',fill='x',padx=14,pady=(0,3))
        tk.Label(form,text='Titel',bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=0,column=0,sticky='w')
        tk.Entry(form,textvariable=self.kb_title_var,bg=WHITE,fg=TEXT,font=body_font(9),relief='solid',bd=1).grid(row=1,column=0,columnspan=4,sticky='ew',ipady=2,pady=(0,3))
        values=['']+self.kb_get_categories()
        for i,var in enumerate(self.kb_entry_category_vars):
            tk.Label(form,text=f'Kategorie {i+1}',bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=2,column=i,sticky='w',padx=(0 if i==0 else 8,0))
            sel=_wz449_colored_category_selector(form,self,var,values,width=16,allow_custom=True,mark=True); sel.grid(row=3,column=i,sticky='ew',padx=(0 if i==0 else 8,0),pady=(0,3))
        tk.Label(form,text='Benutzer',bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=0,sticky='w')
        tk.Entry(form,textvariable=self.kb_user_var,bg=WHITE,fg=TEXT,font=body_font(9),relief='solid',bd=1).grid(row=5,column=0,sticky='ew',pady=(0,2))
        tk.Label(form,text='Status',bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=1,sticky='w',padx=(8,0))
        ttk.Combobox(form,textvariable=self.kb_status_var,values=['Aktiv','Entwurf','Veraltet'],state='readonly',font=body_font(9)).grid(row=5,column=1,sticky='ew',padx=(8,0),pady=(0,2))
        if any((v.get() or '').strip().lower()=='to-do' for v in self.kb_entry_category_vars):
            tk.Label(form,text='Rhythmus',bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=4,column=2,sticky='w',padx=(8,0))
            ttk.Combobox(form,textvariable=self.kb_rhythm_var,values=['','täglich','wöchentlich','monatlich','quartalsweise','jährlich','bei Bedarf'],state='readonly',font=body_font(9)).grid(row=5,column=2,sticky='ew',padx=(8,0),pady=(0,2))
        else: self.kb_rhythm_var.set('')
        for c in range(4): form.grid_columnconfigure(c,weight=1)
    work=tk.Frame(frame,bg=WHITE); work.pack(side='top',fill='both',expand=True,padx=14,pady=(0,0))
    tk.Label(work,text='Freitext / Prozessdokumentation / Leitfaden',bg=WHITE,fg=TEXT2,font=body_font(8)).pack(anchor='w')
    overlay=tk.Frame(work,bg=WHITE); self.kb_text_overlay_parent=overlay
    self.kb_text_widget=tk.Text(overlay,bg='#F8FAFC',fg=TEXT,font=body_font(_WZ449_DEFAULT_TEXT_SIZE),relief='solid',bd=1,wrap='word',undo=True)
    _wz449_visible_format_toolbar(work,self.kb_text_widget,self)
    overlay.pack(fill='both',expand=True); yscroll=tk.Scrollbar(overlay,orient='vertical',command=self.kb_text_widget.yview); self.kb_text_widget.configure(yscrollcommand=yscroll.set); self.kb_text_widget.pack(side='left',fill='both',expand=True); yscroll.pack(side='right',fill='y')
    self.kb_text_widget.insert('1.0',getattr(self,'kb_text_initial',''))
    try: _wz446_apply_formatting(self.kb_text_widget,getattr(self,'kb_text_formatting_initial',[]) or [])
    except Exception:
        try: _wz439_apply_formatting(self.kb_text_widget,getattr(self,'kb_text_formatting_initial',[]) or [])
        except Exception: pass
    self.kb_text_widget.bind('<KeyRelease>',lambda e: (_wz449_capture_editor_state(self), self.kb_mark_unsaved()), add='+')
    try: self.kb_text_widget.bind('<MouseWheel>',lambda e:_wz439_text_mousewheel(self.kb_text_widget,e))
    except Exception: pass
    _wz449_place_objects(self)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor='nw',width=ui_s(w),height=ui_s(h))

FiBuMateApp.render_kb_new_entry_area = _wz449_render_new_entry

# Detailansicht: Standard-Freitext ebenfalls größer und Tabellen/Bilder rudimentär anzeigen.
try:
    _wz449_old_render_detail = FiBuMateApp.render_kb_detail_area
    def _wz449_render_detail(self,x,y,w,h):
        return _wz449_old_render_detail(self,x,y,w,h)
    FiBuMateApp.render_kb_detail_area = _wz449_render_detail
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate / Wissenszentrale - Export, Sortierung, Versionsverlauf-Admin FINAL 2026-06-19
# Version 0.450
# Zweck:
# - Word-Export fragt Ablagepfad und Dateinamen per Dialog ab.
# - Übersicht: Klick auf Spaltenüberschrift sortiert auf-/absteigend.
# - Versionsverlauf: E4 darf Versionsverlauf-Einträge bearbeiten, ausblenden/einblenden und löschen.
# - Grundlage für bessere Flächennutzung und Lesbarkeit: konsequente größere Treeview-Zeilen und kompaktere Header.
# ------------------------------------------------------------------

def _fm450_is_e4(self):
    try:
        role = self.my_role()
        if str(role) == ROLE_E4 or self.role_rank(role) >= 4:
            return True
    except Exception:
        pass
    try:
        return int(self.kb_current_permission_level()) >= 4
    except Exception:
        pass
    try:
        key = str(getattr(self, 'current_user_key', '') or '').lower()
        disp = str(getattr(self, 'current_user_display', '') or '').lower()
        return 'wagnerm' in key or 'wagnerm' in disp
    except Exception:
        return False

def _fm450_version_key(entry):
    return (
        str(entry.get('version','')),
        str(entry.get('published_at','')),
        str(entry.get('date','')),
        str(entry.get('update_id','')),
        str(entry.get('title','')),
    )

def _fm450_version_entries_raw(self):
    try:
        return list(self.load_version_history().get('entries', []) or [])
    except Exception:
        return []

def _fm450_save_version_entries_admin(self, entries):
    data={'entries': entries}
    errors=[]
    targets=[]
    try:
        targets.append(self._central_version_history_path())
    except Exception:
        pass
    try:
        os.makedirs(USER_DIR, exist_ok=True)
        targets.append(os.path.join(USER_DIR, VERSION_HISTORY_FILE))
    except Exception:
        pass
    seen=set(); wrote=False
    for path in targets:
        if not path or path in seen: continue
        seen.add(path)
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            tmp=path+'.tmp'
            with open(tmp,'w',encoding='utf-8') as f:
                json.dump(data,f,ensure_ascii=False,indent=2)
            os.replace(tmp,path)
            wrote=True
        except Exception as exc:
            errors.append(f'{path}: {exc}')
    if not wrote:
        try: messagebox.showerror('Versionsverlauf', 'Versionsverlauf konnte nicht gespeichert werden:\n'+'\n'.join(errors))
        except Exception: pass
        return False
    return True

def _fm450_export_word(self):
    entry = self.kb_get_entry(getattr(self, 'kb_selected_entry_id', None))
    if not entry:
        try: messagebox.showwarning('Wissenszentrale', 'Kein Eintrag für den Word-Export ausgewählt.')
        except Exception: pass
        return
    try:
        from tkinter import filedialog
        from docx import Document
        import re as _re
        safe = _re.sub(r'[^A-Za-z0-9_äöüÄÖÜß.-]+', '_', entry.get('title', 'Wissenseintrag'))[:80].strip('._') or 'Wissenseintrag'
        initial_name = f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        try:
            initial_dir = self.kb_export_dir()
        except Exception:
            initial_dir = os.getcwd()
        out = filedialog.asksaveasfilename(
            title='Word-Export speichern unter',
            defaultextension='.docx',
            initialdir=initial_dir,
            initialfile=initial_name,
            filetypes=[('Word-Dokument', '*.docx'), ('Alle Dateien', '*.*')]
        )
        if not out:
            return
        doc = Document()
        doc.add_heading(entry.get('title', 'Wissenseintrag') or 'Wissenseintrag', level=1)
        meta = [
            ('Geändert', self.kb_display_date(entry.get('updated_at'))),
            ('Benutzer', entry.get('user', '')),
            ('Status', entry.get('status', '')),
            ('Kategorien', ', '.join(entry.get('categories', []) or [])),
            ('To-Do-Rhythmus', entry.get('rhythm', '')),
        ]
        for label, val in meta:
            if val:
                p = doc.add_paragraph(); p.add_run(f'{label}: ').bold = True; p.add_run(str(val))
        doc.add_heading('Inhalt', level=2)
        text_value = entry.get('text', '') or ''
        if text_value:
            for line in str(text_value).splitlines() or ['']:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph('-')
        images = entry.get('inline_images', []) or []
        if images:
            doc.add_heading('Bilder', level=2)
            for img in images:
                name = img.get('name') or os.path.basename(str(img.get('path') or img.get('source') or ''))
                path = img.get('path') or img.get('source') or ''
                doc.add_paragraph(f'- {name} ({path})')
        tables = entry.get('inline_tables', []) or []
        if tables:
            doc.add_heading('Tabellen', level=2)
            for t_index, table_data in enumerate(tables, 1):
                cells = table_data.get('cells', []) or []
                rows = max(1, len(cells))
                cols = max(1, max((len(r) for r in cells), default=1))
                doc.add_paragraph(f'Tabelle {t_index}')
                tbl = doc.add_table(rows=rows, cols=cols)
                tbl.style = 'Table Grid'
                for r in range(rows):
                    for c in range(cols):
                        try: tbl.cell(r,c).text = str(cells[r][c])
                        except Exception: tbl.cell(r,c).text = ''
        doc.add_heading('Anhänge', level=2)
        attachments = entry.get('attachments', []) or []
        if attachments:
            for a in attachments:
                doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        else:
            doc.add_paragraph('Keine Anhänge')
        doc.add_heading('Kommentare', level=2)
        comments = entry.get('comments', []) or []
        if comments:
            for c in comments:
                doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        else:
            doc.add_paragraph('Keine Kommentare')
        doc.save(out)
        messagebox.showinfo('Wissenszentrale', 'Word-Export erstellt:\n' + out)
    except Exception as exc:
        try: messagebox.showerror('Wissenszentrale', 'Word-Export fehlgeschlagen:\n' + str(exc))
        except Exception: pass

FiBuMateApp.kb_export_selected_to_word = _fm450_export_word

def _fm450_sort_entries(self, entries, column):
    state = getattr(self, 'kb_overview_sort_state', {'column': 'updated_at', 'descending': True}) or {}
    descending = bool(state.get('descending', True))
    if state.get('column') == column:
        descending = not descending
    else:
        descending = True if column in ('updated_at','date') else False
    self.kb_overview_sort_state = {'column': column, 'descending': descending}
    def key(entry):
        if column == 'date':
            return str(entry.get('updated_at') or entry.get('created_at') or '')
        if column == 'title':
            return str(entry.get('title','')).casefold()
        if column == 'categories':
            return ', '.join(entry.get('categories',[]) or []).casefold()
        if column == 'user':
            return str(entry.get('user','')).casefold()
        if column == 'status':
            return str(entry.get('status','')).casefold()
        return str(entry.get(column,'')).casefold()
    try:
        return sorted(entries, key=key, reverse=descending)
    except Exception:
        return entries

def _fm450_render_kb_list_area(self, x, y, w, h, title='Gesamtliste aller Einträge', status_filter=None):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    header = tk.Frame(frame, bg=WHITE)
    header.pack(fill='x', padx=16, pady=(12, 6))
    tk.Label(header, text=title, bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(side='left')
    sort_state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
    direction = 'absteigend' if sort_state.get('descending', True) else 'aufsteigend'
    tk.Label(header, text=f"Sortierung: {sort_state.get('column','date')} ({direction})", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')
    columns = ('date','title','categories','user','status')
    labels = {'date':'Geändert','title':'Titel','categories':'Kategorien','user':'Benutzer','status':'Status'}
    tree = ttk.Treeview(frame, columns=columns, show='headings', height=14)
    try:
        style = ttk.Style()
        style.configure('Treeview', rowheight=28, font=body_font(10))
        style.configure('Treeview.Heading', font=body_font(10, weight='bold'))
    except Exception:
        pass
    for col in columns:
        marker = ''
        if sort_state.get('column') == col:
            marker = ' ↓' if sort_state.get('descending', True) else ' ↑'
        tree.heading(col, text=labels[col]+marker, command=lambda c=col: (setattr(self, 'kb_overview_sort_state', {'column': c, 'descending': not (getattr(self, 'kb_overview_sort_state', {}) or {}).get('descending', True) if (getattr(self, 'kb_overview_sort_state', {}) or {}).get('column') == c else (True if c == 'date' else False)}), self.render_page()))
    tree.column('date', width=125, stretch=False)
    tree.column('title', width=max(260, int(w*0.36)), stretch=True)
    tree.column('categories', width=max(180, int(w*0.22)), stretch=True)
    tree.column('user', width=145, stretch=False)
    tree.column('status', width=100, stretch=False)
    yscroll = ttk.Scrollbar(frame, orient='vertical', command=tree.yview)
    tree.configure(yscrollcommand=yscroll.set)
    tree.pack(side='left', fill='both', expand=True, padx=(16,0), pady=(0,16))
    yscroll.pack(side='right', fill='y', padx=(0,16), pady=(0,16))
    entries = self.kb_filtered_entries()
    if status_filter:
        entries = [e for e in entries if str(e.get('status','')).lower() == status_filter.lower()]
    state = getattr(self, 'kb_overview_sort_state', None)
    if not state:
        self.kb_overview_sort_state = {'column':'date','descending':True}
    entries = sorted(entries, key=lambda e: (
        str(e.get('updated_at') or e.get('created_at') or '') if self.kb_overview_sort_state.get('column')=='date' else
        str(e.get('title','')).casefold() if self.kb_overview_sort_state.get('column')=='title' else
        ', '.join(e.get('categories',[]) or []).casefold() if self.kb_overview_sort_state.get('column')=='categories' else
        str(e.get('user','')).casefold() if self.kb_overview_sort_state.get('column')=='user' else
        str(e.get('status','')).casefold()
    ), reverse=bool(self.kb_overview_sort_state.get('descending', True)))
    id_map = {}
    for entry in entries:
        iid = entry.get('id') or self.kb_make_entry_id()
        base_iid = iid; n = 1
        while iid in id_map:
            n += 1; iid = f'{base_iid}_{n}'
        id_map[iid] = entry.get('id')
        tree.insert('', 'end', iid=iid, values=(self.kb_display_date(entry.get('updated_at')), entry.get('title',''), ', '.join(entry.get('categories',[]) or []), entry.get('user',''), entry.get('status','')))
    if not entries:
        tree.insert('', 'end', values=('', 'Noch keine Einträge vorhanden', '', '', ''))
    def _open(event=None):
        sel = tree.selection()
        if sel and sel[0] in id_map:
            self.kb_select_entry(id_map[sel[0]])
    tree.bind('<Double-Button-1>', _open)
    tree.bind('<Return>', _open)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))

FiBuMateApp.render_kb_list_area = _fm450_render_kb_list_area

def _fm450_edit_version_dialog(self, index, entries):
    if not _fm450_is_e4(self):
        return
    entry = dict(entries[index])
    try:
        win = tk.Toplevel(self.root)
        win.title('Versionsverlauf bearbeiten')
        win.configure(bg=WHITE)
        win.geometry('720x460')
        version_var = tk.StringVar(value=str(entry.get('version','')))
        title_var = tk.StringVar(value=str(entry.get('title','')))
        tk.Label(win, text='Version', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w', padx=16, pady=(14,2))
        tk.Entry(win, textvariable=version_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', padx=16, ipady=4)
        tk.Label(win, text='Titel / Kurzbeschreibung', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w', padx=16, pady=(10,2))
        tk.Entry(win, textvariable=title_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', padx=16, ipady=4)
        tk.Label(win, text='Eintrag / Stichpunkte', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w', padx=16, pady=(10,2))
        txt = tk.Text(win, bg='#F8FAFC', fg=TEXT, font=body_font(11), relief='solid', bd=1, wrap='word')
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            txt.insert('1.0', '\n'.join(str(b) for b in bullets))
        else:
            txt.insert('1.0', str(bullets or entry.get('text','') or ''))
        txt.pack(fill='both', expand=True, padx=16, pady=(0,10))
        br = tk.Frame(win, bg=WHITE); br.pack(fill='x', padx=16, pady=(0,14))
        def save():
            entry['version'] = version_var.get().strip()
            entry['title'] = title_var.get().strip()
            raw = txt.get('1.0','end-1c').strip()
            entry['bullets'] = [line.strip() for line in raw.splitlines() if line.strip()]
            entries[index] = entry
            if _fm450_save_version_entries_admin(self, entries):
                win.destroy(); self.render_page()
        tk.Button(br, text='Speichern', command=save, bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=14, ipady=4)
        tk.Button(br, text='Abbrechen', command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', ipadx=14, ipady=4)
        win.transient(self.root); win.grab_set(); win.focus_set()
    except Exception as exc:
        try: messagebox.showerror('Versionsverlauf', 'Bearbeiten nicht möglich:\n'+str(exc))
        except Exception: pass

def _fm450_render_versions_menu(self):
    entries = _fm450_version_entries_raw(self)
    is_e4 = _fm450_is_e4(self)
    if not is_e4:
        entries = [e for e in entries if not bool(e.get('hidden'))]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top)
    container = tk.Frame(self.root, bg=BG); self.widget_items.append(container)
    self.canvas.create_window(0, area_top, window=container, anchor='nw', width=w, height=view_h)
    scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0)
    scrollbar = tk.Scrollbar(container, orient='vertical', command=scroll_canvas.yview)
    content = tk.Frame(scroll_canvas, bg=BG)
    content_window = scroll_canvas.create_window((0,0), window=content, anchor='nw')
    def update_scrollregion(_event=None):
        scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width()))
        scroll_canvas.configure(scrollregion=scroll_canvas.bbox('all'))
        try: self._sync_scrollbar_visibility(scroll_canvas, scrollbar)
        except Exception: pass
    content.bind('<Configure>', update_scrollregion)
    scroll_canvas.bind('<Configure>', update_scrollregion)
    scroll_canvas.configure(yscrollcommand=scrollbar.set)
    scroll_canvas.pack(side='left', fill='both', expand=True)
    scrollbar.pack(side='right', fill='y')
    tk.Label(content, text='Versionsverlauf', bg=BG, fg=BLUE, font=body_font(20, weight='bold')).pack(anchor='w', padx=24, pady=(18,4))
    source_txt = f"Versionsverlauf: lokale Historie plus zentrale Release-Historie ({self._central_version_history_path()})"
    tk.Label(content, text=source_txt, bg=BG, fg=TEXT2, font=body_font(9), wraplength=max(520, w-70), justify='left').pack(anchor='w', padx=24, pady=(0,10))
    if is_e4:
        tk.Label(content, text='E4-Administration aktiv: Einträge können bearbeitet, ausgeblendet/eingeblendet oder gelöscht werden.', bg=BG, fg=RED, font=body_font(10, weight='bold'), wraplength=max(520,w-70), justify='left').pack(anchor='w', padx=24, pady=(0,10))
    if not entries:
        tk.Label(content, text='Noch kein Versionsverlauf vorhanden.', bg=BG, fg=TEXT2, font=body_font(12)).pack(anchor='w', padx=24, pady=24)
        return
    for idx, entry in enumerate(entries):
        card = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=1)
        card.pack(fill='x', padx=24, pady=8)
        top = tk.Frame(card, bg=WHITE); top.pack(fill='x', padx=14, pady=(10,4))
        version = str(entry.get('version','')).strip() or 'ohne Version'
        title = str(entry.get('title','')).strip()
        hidden = bool(entry.get('hidden'))
        head = version + (f' - {title}' if title else '') + ('  [AUSGEBLENDET]' if hidden else '')
        tk.Label(top, text=head, bg=WHITE, fg=(RED if hidden else BLUE), font=body_font(13, weight='bold')).pack(side='left', anchor='w')
        date_txt = entry.get('date') or entry.get('published_at') or ''
        if date_txt:
            tk.Label(top, text=str(date_txt), bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            for b in bullets:
                tk.Label(card, text='• '+str(b), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520,w-120), justify='left').pack(anchor='w', padx=18, pady=1)
        else:
            tk.Label(card, text=str(bullets), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520,w-120), justify='left').pack(anchor='w', padx=18, pady=1)
        if is_e4:
            br = tk.Frame(card, bg=WHITE); br.pack(fill='x', padx=14, pady=(8,10))
            def edit(i=idx): _fm450_edit_version_dialog(self, i, entries)
            def toggle(i=idx):
                entries[i]['hidden'] = not bool(entries[i].get('hidden'))
                if _fm450_save_version_entries_admin(self, entries): self.render_page()
            def delete(i=idx):
                try:
                    ok = messagebox.askyesno('Versionsverlauf', 'Diesen Versionsverlauf-Eintrag unwiderruflich löschen?')
                except Exception: ok = False
                if ok:
                    try: entries.pop(i)
                    except Exception: pass
                    if _fm450_save_version_entries_admin(self, entries): self.render_page()
            tk.Button(br, text='Bearbeiten', command=edit, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0,6), ipadx=10, ipady=2)
            tk.Button(br, text=('Einblenden' if hidden else 'Ausblenden'), command=toggle, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0,6), ipadx=10, ipady=2)
            tk.Button(br, text='Löschen', command=delete, bg='#FEE2E2', fg=RED, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='left', ipadx=10, ipady=2)
    try:
        scroll_canvas.bind_all('<MouseWheel>', lambda e: scroll_canvas.yview_scroll(int(-1*(e.delta/120)), 'units'))
    except Exception:
        pass

FiBuMateApp.render_versions_menu = _fm450_render_versions_menu



# ------------------------------------------------------------------
# Wissenszentrale - Exportoptionen, Bildtracking, Widgetdesign, Filterleiste FINAL 2026-06-19
# Version 0.451
# Zweck:
# - Word-Export mit Optionsdialog: Titel, Metadaten, Prozessdoku, Bilder, Tabellen, Anhänge, Kommentare, Intersport-Layout.
# - Bildverschiebung/-skalierung verbessert: root-koordinatenbasiertes Tracking, größere Anfasser, weniger Flackern.
# - Bildrahmen liegt exakt an den sichtbaren Bildkonturen; Bild wird auf gewählte Größe verzerrt/gestreckt, kein erzwungenes Seitenverhältnis.
# - Dezentes, aber gestalteteres Button-/Widgetdesign für Wissenszentrale-Aktionen.
# - Suchleiste und Kategorie-Dropdowns auf gleicher Y-Höhe; Kategorien in Dropdowns farbig hinterlegt.
# ------------------------------------------------------------------

_FM451_BTN_BG = '#F8FAFC'
_FM451_BTN_ACTIVE = '#E6EEF8'
_FM451_BTN_ACCENT = '#DDF3E4'
_FM451_BTN_BORDER = '#B7C7D8'
_FM451_SHADOW = '#D6E0EA'


def _fm451_soften(color, factor=0.88):
    try:
        color = _wz_cat_hex(color)
        r, g, b = int(color[1:3],16), int(color[3:5],16), int(color[5:7],16)
        r = int(r + (255-r)*factor); g = int(g + (255-g)*factor); b = int(b + (255-b)*factor)
        return f'#{r:02X}{g:02X}{b:02X}'
    except Exception:
        return '#F3F6FA'


def _fm451_button(parent, text, command=None, accent=False, danger=False, font_size=9, width=None):
    bg = '#FEE2E2' if danger else (_FM451_BTN_ACCENT if accent else _FM451_BTN_BG)
    active = '#FCA5A5' if danger else (_FM451_BTN_ACCENT if accent else _FM451_BTN_ACTIVE)
    fg = RED if danger else TEXT
    btn = tk.Button(
        parent, text=text, command=command, bg=bg, fg=fg,
        activebackground=active, activeforeground=fg,
        font=body_font(font_size, weight='bold' if accent or danger else None),
        relief='flat', bd=0, cursor='hand2', padx=10, pady=4,
        highlightthickness=1, highlightbackground=_FM451_BTN_BORDER, highlightcolor=BLUE
    )
    if width:
        try: btn.configure(width=width)
        except Exception: pass
    return btn


def _fm451_draw_kb_button(self, x, y, text, command=None, accent=False, width=None):
    btn = _fm451_button(self.root, text, command=command, accent=accent, font_size=9)
    btn_w = width or max(112, 24 + len(text) * 8)
    self.canvas.create_window(ui_s(x), ui_s(y), window=btn, anchor='nw', width=ui_s(btn_w), height=ui_s(34))
    self.widget_items.append(btn)
    return x + btn_w + 8

FiBuMateApp.draw_kb_button = _fm451_draw_kb_button


def _fm451_cat_color(app, name):
    try:
        return app.kb_get_category_color(name) if name else WHITE
    except Exception:
        try: return _wz449_safe_cat_color(app, name)
        except Exception: return WHITE


def _fm451_cat_fg(color):
    try: return _wz_cat_fg(color)
    except Exception: return TEXT


def _fm451_filter_dropdown(parent, app, var, values, width=17, command=None):
    box = tk.Frame(parent, bg=BG)
    value_lbl = tk.Label(box, text=(var.get() or '(alle)'), bg=WHITE, fg=TEXT, font=body_font(10), anchor='w', relief='flat', padx=8)
    value_lbl.pack(side='left', fill='x', expand=True, ipady=5)
    arrow = tk.Menubutton(box, text='▼', bg=_FM451_BTN_BG, fg=TEXT, activebackground=_FM451_BTN_ACTIVE,
                          relief='flat', bd=0, font=body_font(8), cursor='hand2', highlightthickness=1,
                          highlightbackground=_FM451_BTN_BORDER, width=2)
    menu = tk.Menu(arrow, tearoff=False)
    arrow.configure(menu=menu)
    arrow.pack(side='left', padx=(3,0), ipady=3)
    def repaint():
        try:
            val = (var.get() or '').strip()
            color = _fm451_cat_color(app, val)
            value_lbl.configure(text=(val if val else '(alle)'), bg=(_fm451_soften(color, 0.78) if val else WHITE), fg=(_fm451_cat_fg(color) if val else TEXT))
        except Exception:
            pass
    def select(value):
        var.set(value)
        repaint()
        if command:
            try: command()
            except TypeError: command(None)
            except Exception: pass
    menu.add_command(label='(alle)', command=lambda: select(''), background=WHITE, foreground=TEXT)
    for value in values:
        if not value: continue
        color = _fm451_cat_color(app, value)
        menu.add_command(label='  ' + str(value), command=lambda v=value: select(v),
                         background=_fm451_soften(color, 0.40), foreground=_fm451_cat_fg(color),
                         activebackground=color, activeforeground=_fm451_cat_fg(color))
    repaint()
    return box


def _fm451_render_knowledge_work_area(self):
    self.kb_ensure_state_vars()
    try:
        self.root.bind('<Escape>', self.kb_handle_escape, add='+')
    except Exception:
        pass
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    nav_y = 145
    bx = 28
    bx = self.draw_kb_button(bx, nav_y, 'Start', self.kb_show_start_overlay, False, width=92)
    bx = self.draw_kb_button(bx, nav_y, 'Übersicht', lambda: self.kb_switch_view_from_start('all'), False, width=112)
    if self.kb_can_create_or_edit():
        bx = self.draw_kb_button(bx, nav_y, 'Neuer Eintrag', lambda: self.kb_switch_view_from_start('new'), False, width=128)
    bx = self.draw_kb_button(bx, nav_y, 'To-Dos', lambda: self.kb_switch_view_from_start('todos'), False, width=92)
    bx = self.draw_kb_button(bx, nav_y, 'Veraltete Einträge', lambda: self.kb_switch_view_from_start('outdated'), False, width=148)
    if self.kb_can_manage_categories():
        bx = self.draw_kb_button(bx, nav_y, 'Kategorien verwalten', lambda: self.kb_switch_view_from_start('categories'), False, width=168)

    bar = tk.Frame(self.root, bg=BG)
    self.widget_items.append(bar)
    bar.grid_columnconfigure(0, weight=1)
    tk.Label(bar, text='Suche', bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky='w')
    for i in range(4):
        tk.Label(bar, text=f'Kategorie {i+1}', bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=i+2, sticky='w', padx=(12,0))
    search_row = tk.Frame(bar, bg=BG)
    search_row.grid(row=1, column=0, sticky='ew')
    search_row.grid_columnconfigure(0, weight=1)
    ent = tk.Entry(search_row, textvariable=self.kb_search_var, font=body_font(11), bg=WHITE, fg=TEXT, relief='flat', bd=0,
                   highlightthickness=1, highlightbackground=_FM451_BTN_BORDER, highlightcolor=BLUE)
    ent.grid(row=0, column=0, sticky='ew', ipady=7)
    ent.bind('<Return>', self.kb_on_search_return)
    _fm451_button(search_row, 'Suchen', command=self.kb_apply_filters, font_size=9).grid(row=0, column=1, padx=(8,0), sticky='ns')
    vals = [''] + self.kb_get_categories()
    for i, var in enumerate(self.kb_filter_vars):
        dd = _fm451_filter_dropdown(bar, self, var, vals, command=self.kb_apply_filters)
        dd.grid(row=1, column=i+2, sticky='ew', padx=(12,0))
    selected = [(v.get() or '').strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view == 'todos' or 'to-do' in selected:
        tk.Label(bar, text='Rhythmus', bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=6, sticky='w', padx=(12,0))
        rhythm = ttk.Combobox(bar, textvariable=self.kb_todo_rhythm_var, values=['','täglich','wöchentlich','monatlich','quartalsweise','jährlich','bei Bedarf'], state='readonly', font=body_font(10), width=15)
        rhythm.grid(row=1, column=6, sticky='ew', padx=(12,0), ipady=4)
        rhythm.bind('<<ComboboxSelected>>', self.kb_apply_filters)
    else:
        self.kb_todo_rhythm_var.set('')
    self.canvas.create_window(ui_s(28), ui_s(192), window=bar, anchor='nw', width=ui_s(min(max(1100, int(w-64)), w-56)), height=ui_s(68))

    left_x = 20
    left_y = 282
    left_w = max(340, min(460, int(w * 0.31)))
    pane_h = max(500, h - left_y - 24)
    right_x = left_x + left_w + 24
    right_w = max(760, w - right_x - 24)
    self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
    if self.knowledge_view == 'new': self.render_kb_new_entry_area(right_x, left_y, right_w, pane_h)
    elif self.knowledge_view == 'todos': self.render_kb_list_area(right_x, left_y, right_w, pane_h, title='To-Dos')
    elif self.knowledge_view == 'outdated': self.render_kb_list_area(right_x, left_y, right_w, pane_h, title='Veraltete Einträge', status_filter='Veraltet')
    elif self.knowledge_view == 'categories': self.render_kb_categories_area(right_x, left_y, right_w, pane_h)
    elif self.knowledge_view == 'detail': self.render_kb_detail_area(right_x, left_y, right_w, pane_h)
    else: self.render_kb_list_area(right_x, left_y, right_w, pane_h, title='Übersicht')

FiBuMateApp.render_knowledge_work_area = _fm451_render_knowledge_work_area


def _fm451_photo_exact(path, width, height):
    if not PIL_AVAILABLE or not path or not os.path.exists(path):
        return None
    try:
        width = max(24, int(width)); height = max(24, int(height))
        img = Image.open(path).convert('RGBA')
        img = img.resize((width, height))
        return ImageTk.PhotoImage(img)
    except Exception:
        return None


def _fm451_place_image(self, parent, idx, item):
    x, y = int(item.get('x', 20)), int(item.get('y', 20))
    w, h = max(30, int(item.get('w', 180))), max(25, int(item.get('h', 130)))
    item['w'], item['h'] = w, h
    frame = tk.Frame(parent, bg=WHITE, highlightthickness=1, highlightbackground=BLUE, highlightcolor=BLUE, cursor='fleur')
    frame._wz_float_image = True
    frame.place(x=x, y=y, width=w, height=h)
    photo = _fm451_photo_exact(item.get('path') or item.get('source'), w, h)
    if photo:
        self._kb_text_image_refs.append(photo)
        lbl = tk.Label(frame, image=photo, bg=WHITE, bd=0, highlightthickness=0)
    else:
        lbl = tk.Label(frame, text='Bild', bg='#FEE2E2', fg=RED, font=body_font(8), bd=0, highlightthickness=0)
    lbl.place(x=0, y=0, width=w, height=h)
    item['draw_w'], item['draw_h'] = w, h

    def begin_move(e):
        self._fm451_img_drag = {'mode':'move', 'idx':idx, 'sx':e.x_root, 'sy':e.y_root, 'x':int(item.get('x',x)), 'y':int(item.get('y',y)), 'frame':frame}
        return 'break'
    def move(e):
        d = getattr(self, '_fm451_img_drag', None)
        if not d or d.get('idx') != idx or d.get('mode') != 'move': return 'break'
        nx = max(0, d['x'] + e.x_root - d['sx']); ny = max(0, d['y'] + e.y_root - d['sy'])
        self.kb_inline_images[idx]['x'] = nx; self.kb_inline_images[idx]['y'] = ny
        try: d['frame'].place_configure(x=nx, y=ny)
        except Exception: pass
        return 'break'
    def end_drag(e):
        try: self.kb_mark_unsaved()
        except Exception: pass
        self._fm451_img_drag = None
        return 'break'
    frame.bind('<ButtonPress-1>', begin_move); lbl.bind('<ButtonPress-1>', begin_move)
    frame.bind('<B1-Motion>', move); lbl.bind('<B1-Motion>', move)
    frame.bind('<ButtonRelease-1>', end_drag); lbl.bind('<ButtonRelease-1>', end_drag)
    frame.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e)); lbl.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e))

    handle_size = 11
    handles=[('nw',0,0),('n',w//2-handle_size//2,0),('ne',w-handle_size,0),('e',w-handle_size,h//2-handle_size//2),('se',w-handle_size,h-handle_size),('s',w//2-handle_size//2,h-handle_size),('sw',0,h-handle_size),('w',0,h//2-handle_size//2)]
    for name,hx,hy in handles:
        hd = tk.Frame(frame, bg=WHITE, highlightbackground=BLUE, highlightthickness=1, cursor='sizing')
        hd.place(x=max(0,hx), y=max(0,hy), width=handle_size, height=handle_size)
        def start_resize(e, nm=name):
            self._fm451_img_drag={'mode':'resize','idx':idx,'handle':nm,'sx':e.x_root,'sy':e.y_root,'x':int(item.get('x',x)),'y':int(item.get('y',y)),'w':int(item.get('w',w)),'h':int(item.get('h',h)),'frame':frame}
            return 'break'
        def resize(e):
            d=getattr(self,'_fm451_img_drag',None)
            if not d or d.get('idx')!=idx or d.get('mode')!='resize': return 'break'
            dx=e.x_root-d['sx']; dy=e.y_root-d['sy']; nx,ny,nw,nh=d['x'],d['y'],d['w'],d['h']; nm=d['handle']
            if 'e' in nm: nw=max(30,d['w']+dx)
            if 's' in nm: nh=max(25,d['h']+dy)
            if 'w' in nm: nx=max(0,d['x']+dx); nw=max(30,d['w']-dx)
            if 'n' in nm: ny=max(0,d['y']+dy); nh=max(25,d['h']-dy)
            self.kb_inline_images[idx].update({'x':nx,'y':ny,'w':nw,'h':nh})
            try: d['frame'].place_configure(x=nx,y=ny,width=nw,height=nh)
            except Exception: pass
            return 'break'
        def finish_resize(e):
            try: self.kb_mark_unsaved()
            except Exception: pass
            self._fm451_img_drag=None
            _wz449_place_objects(self)
            return 'break'
        hd.bind('<ButtonPress-1>', start_resize)
        hd.bind('<B1-Motion>', resize)
        hd.bind('<ButtonRelease-1>', finish_resize)
        hd.bind('<Button-3>', lambda e,i=idx:_wz449_image_context(self,i,e))

_wz449_place_image = _fm451_place_image


def _fm451_widget_toolbar(parent, widget, app):
    bar = tk.Frame(parent, bg=WHITE)
    bar.pack(fill='x', pady=(1,4))
    tk.Label(bar, text='Textformat:', bg=WHITE, fg=TEXT2, font=body_font(8, weight='bold')).pack(side='left', padx=(0,5))
    size_var = tk.StringVar(value=str(getattr(app, 'kb_current_text_format', {}).get('size', _WZ449_DEFAULT_TEXT_SIZE if '_WZ449_DEFAULT_TEXT_SIZE' in globals() else 16)))
    cb = ttk.Combobox(bar, textvariable=size_var, values=['8','9','10','11','12','14','16','18','20','22','24','28','32'], width=4, state='readonly', font=body_font(8))
    cb.pack(side='left', padx=(0,5))
    cb.bind('<<ComboboxSelected>>', lambda e: _wz449_apply_editor_format(widget, app, size=int(size_var.get())))
    for label,key in [('Fett','bold'),('Kursiv','italic'),('Unterstr.','underline')]:
        _fm451_button(bar, label, command=lambda k=key:_wz449_apply_editor_format(widget, app, toggle=k), font_size=8).pack(side='left', padx=(0,3), ipadx=4, ipady=1)
    tk.Label(bar, text='Farbe:', bg=WHITE, fg=TEXT2, font=body_font(8)).pack(side='left', padx=(7,3))
    for label,key,fg in [('Standard','standard','#000000'),('Rot','rot',RED),('Blau','blau',BLUE)]:
        b = _fm451_button(bar, label, command=lambda c=key:_wz449_apply_editor_format(widget, app, color=c), font_size=8)
        try: b.configure(fg=fg)
        except Exception: pass
        b.pack(side='left', padx=(0,3), ipadx=4, ipady=1)
    _fm451_button(bar, 'Bild aus Zwischenablage', command=lambda:_wz446_add_clipboard_image(app), font_size=8).pack(side='left', padx=(7,3), ipadx=5, ipady=1)
    _fm451_button(bar, 'Tabelle einfügen', command=lambda:_wz449_table_popup(app), font_size=8).pack(side='left', padx=(0,3), ipadx=5, ipady=1)
    try:
        widget.bind('<KeyRelease>', lambda e: (_wz449_apply_current_format_to_typed_char(widget, app, e), app.kb_mark_unsaved()), add='+')
        widget.bind('<Control-b>', lambda e:_wz449_apply_editor_format(widget, app, toggle='bold'), add='+')
        widget.bind('<Control-i>', lambda e:_wz449_apply_editor_format(widget, app, toggle='italic'), add='+')
        widget.bind('<Control-u>', lambda e:_wz449_apply_editor_format(widget, app, toggle='underline'), add='+')
        widget.bind('<Control-v>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
        widget.bind('<Control-V>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
        widget.bind('<<Paste>>', lambda e:_wz446_add_clipboard_image(app,e) or None, add='+')
    except Exception: pass
    return bar

_wz449_visible_format_toolbar = _fm451_widget_toolbar


def _fm451_word_options_dialog(self, entry):
    result = {'ok': False}
    win = tk.Toplevel(self.root)
    win.title('Word-Export Optionen')
    win.configure(bg=WHITE)
    win.geometry('520x520')
    vars_ = {
        'title_page': tk.BooleanVar(value=True),
        'metadata': tk.BooleanVar(value=True),
        'process_doc': tk.BooleanVar(value=True),
        'images': tk.BooleanVar(value=True),
        'tables': tk.BooleanVar(value=True),
        'attachments': tk.BooleanVar(value=True),
        'comments': tk.BooleanVar(value=True),
        'intersport_layout': tk.BooleanVar(value=True),
    }
    tk.Label(win, text='Word-Export konfigurieren', bg=WHITE, fg=BLUE, font=body_font(15, weight='bold')).pack(anchor='w', padx=18, pady=(16,6))
    tk.Label(win, text='Bitte auswählen, welche Bestandteile in das Word-Dokument übernommen werden sollen.', bg=WHITE, fg=TEXT2, font=body_font(10), wraplength=460, justify='left').pack(anchor='w', padx=18, pady=(0,12))
    labels = [
        ('title_page', 'Titel / Deckbereich exportieren'),
        ('metadata', 'Kopfdaten / Metadaten exportieren'),
        ('process_doc', 'Prozessdokumentation / Freitext exportieren'),
        ('images', 'Bilderliste exportieren'),
        ('tables', 'Tabellen exportieren'),
        ('attachments', 'Anhänge auflisten'),
        ('comments', 'Kommentare exportieren'),
        ('intersport_layout', 'Intersport-Layout verwenden (Farben, Kopfbereich, Logo wenn verfügbar)'),
    ]
    for key, label in labels:
        tk.Checkbutton(win, text=label, variable=vars_[key], bg=WHITE, fg=TEXT, activebackground=WHITE, font=body_font(10), anchor='w').pack(fill='x', padx=22, pady=4)
    btns = tk.Frame(win, bg=WHITE); btns.pack(fill='x', padx=18, pady=(18,14))
    def ok():
        result.update({k: bool(v.get()) for k,v in vars_.items()}); result['ok']=True; win.destroy()
    _fm451_button(btns, 'Export starten', command=ok, accent=True, font_size=10).pack(side='left', padx=(0,8), ipadx=14, ipady=4)
    _fm451_button(btns, 'Verwerfen', command=win.destroy, font_size=10).pack(side='left', ipadx=14, ipady=4)
    win.transient(self.root); win.grab_set(); win.focus_set(); self.root.wait_window(win)
    return result


def _fm451_export_word(self):
    entry = self.kb_get_entry(getattr(self, 'kb_selected_entry_id', None))
    if not entry:
        try: messagebox.showwarning('Wissenszentrale', 'Kein Eintrag für den Word-Export ausgewählt.')
        except Exception: pass
        return
    try:
        opts = _fm451_word_options_dialog(self, entry)
        if not opts.get('ok'):
            return
        from tkinter import filedialog
        from docx import Document
        from docx.shared import RGBColor, Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re as _re
        safe = _re.sub(r'[^A-Za-z0-9_äöüÄÖÜß.-]+', '_', entry.get('title', 'Wissenseintrag'))[:80].strip('._') or 'Wissenseintrag'
        initial_name = f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        try: initial_dir = self.kb_export_dir()
        except Exception: initial_dir = os.getcwd()
        out = filedialog.asksaveasfilename(title='Word-Export speichern unter', defaultextension='.docx', initialdir=initial_dir, initialfile=initial_name, filetypes=[('Word-Dokument','*.docx'),('Alle Dateien','*.*')])
        if not out: return
        doc = Document()
        if opts.get('intersport_layout'):
            sec = doc.sections[0]
            sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65); sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
            header = sec.header.paragraphs[0]
            header.text = 'INTERSPORT | FiBu Mate Wissenszentrale'
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                header.runs[0].font.color.rgb = RGBColor(0,75,147); header.runs[0].font.bold = True
            except Exception: pass
        if opts.get('title_page'):
            title = doc.add_heading(entry.get('title','Wissenseintrag') or 'Wissenseintrag', level=0)
            if opts.get('intersport_layout'):
                try:
                    title.runs[0].font.color.rgb = RGBColor(0,75,147)
                except Exception: pass
            p = doc.add_paragraph('Export aus FiBu Mate - Wissenszentrale')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p.runs[0].font.size = Pt(11); p.runs[0].font.color.rgb = RGBColor(68,83,100)
            except Exception: pass
            if opts.get('intersport_layout'):
                doc.add_paragraph('────────────────────────────────────────')
        if opts.get('metadata'):
            doc.add_heading('Kopfdaten', level=1)
            meta=[('Geändert', self.kb_display_date(entry.get('updated_at'))),('Benutzer', entry.get('user','')),('Status', entry.get('status','')),('Kategorien', ', '.join(entry.get('categories',[]) or [])),('To-Do-Rhythmus', entry.get('rhythm',''))]
            for label,val in meta:
                if val:
                    p=doc.add_paragraph(); p.add_run(f'{label}: ').bold=True; p.add_run(str(val))
        if opts.get('process_doc'):
            doc.add_heading('Prozessdokumentation', level=1)
            text_value = entry.get('text','') or ''
            if text_value:
                for line in str(text_value).splitlines() or ['']:
                    doc.add_paragraph(line)
            else: doc.add_paragraph('-')
        if opts.get('images'):
            images=entry.get('inline_images',[]) or []
            if images:
                doc.add_heading('Bilder', level=1)
                for img in images:
                    name=img.get('name') or os.path.basename(str(img.get('path') or img.get('source') or ''))
                    path=img.get('path') or img.get('source') or ''
                    doc.add_paragraph(f'- {name} ({path})')
        if opts.get('tables'):
            tables=entry.get('inline_tables',[]) or []
            if tables:
                doc.add_heading('Tabellen', level=1)
                for t_index, table_data in enumerate(tables,1):
                    cells=table_data.get('cells',[]) or []
                    rows=max(1,len(cells)); cols=max(1,max((len(r) for r in cells), default=1))
                    doc.add_paragraph(f'Tabelle {t_index}')
                    tbl=doc.add_table(rows=rows, cols=cols); tbl.style='Table Grid'
                    for r in range(rows):
                        for c in range(cols):
                            try: tbl.cell(r,c).text=str(cells[r][c])
                            except Exception: tbl.cell(r,c).text=''
        if opts.get('attachments'):
            doc.add_heading('Anhänge', level=1)
            attachments=entry.get('attachments',[]) or []
            if attachments:
                for a in attachments: doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
            else: doc.add_paragraph('Keine Anhänge')
        if opts.get('comments'):
            doc.add_heading('Kommentare', level=1)
            comments=entry.get('comments',[]) or []
            if comments:
                for c in comments: doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
            else: doc.add_paragraph('Keine Kommentare')
        doc.save(out)
        messagebox.showinfo('Wissenszentrale', 'Word-Export erstellt:\n' + out)
    except Exception as exc:
        try: messagebox.showerror('Wissenszentrale', 'Word-Export fehlgeschlagen:\n' + str(exc))
        except Exception: pass

FiBuMateApp.kb_export_selected_to_word = _fm451_export_word



# ------------------------------------------------------------------
# Globales Button-/Widgetdesign Variante 4 FINAL SAFE 2026-06-19
# Version 0.455
# Zweck:
# - Variante 4 "Soft-3D / dezenter Verlauf" global für Buttons einführen.
# - Gilt für tk.Button, tk.Menubutton, ttk.Button sowie später geladene Module/Unterfenster.
# - Menükacheln bleiben ausdrücklich unverändert, weil sie Canvas-Tiles sind.
# - Keine Rekursion in configure(), kein ttk.Style() ohne vorhandenes Root-Fenster, kein implizites Tk-Root.
# ------------------------------------------------------------------

_FM455_INSTALLED = False
_FM455_ORIG_TK_BUTTON_INIT = None
_FM455_ORIG_TK_BUTTON_CONFIGURE = None
_FM455_ORIG_TK_MENUBUTTON_INIT = None
_FM455_ORIG_TK_MENUBUTTON_CONFIGURE = None
_FM455_ORIG_TTK_BUTTON_INIT = None
_FM455_CONFIG_GUARD = set()

_FM455_BTN_NEUTRAL_BG = '#F7FAFD'
_FM455_BTN_NEUTRAL_ACTIVE = '#E7F0FA'
_FM455_BTN_PRIMARY_BG = '#EAF2FB'
_FM455_BTN_PRIMARY_ACTIVE = '#DDEAF7'
_FM455_BTN_SUCCESS_BG = '#E1F5E7'
_FM455_BTN_SUCCESS_ACTIVE = '#CFEAD6'
_FM455_BTN_DANGER_BG = '#FFF1F1'
_FM455_BTN_DANGER_ACTIVE = '#FEE2E2'
_FM455_BTN_DISABLED_BG = '#D8E1EA'
_FM455_BTN_DISABLED_FG = '#718096'
_FM455_BTN_BORDER = '#7F96AD'


def _fm455_norm_color(value):
    try:
        value = str(value or '').strip()
        if value.startswith('#') and len(value) == 7:
            return value.upper()
    except Exception:
        pass
    return ''


def _fm455_button_text(widget):
    try:
        return str(widget.cget('text') or '').strip()
    except Exception:
        return ''


def _fm455_should_skip_button(widget):
    """Schützt Farbfeld-/Swatch-Buttons; Menükacheln sind Canvas und werden ohnehin nicht gepatcht."""
    try:
        text_value = _fm455_button_text(widget)
        bg = _fm455_norm_color(widget.cget('bg')) or _fm455_norm_color(widget.cget('background'))
        if not text_value and bg and bg not in ('#F7FAFD', '#FFFFFF', '#E8EEF5'):
            return True
        if bool(getattr(widget, '_fibumate_keep_native_button_style', False)):
            return True
    except Exception:
        pass
    return False


def _fm455_button_kind(widget):
    text_value = _fm455_button_text(widget).lower()
    try:
        bg = _fm455_norm_color(widget.cget('bg')) or _fm455_norm_color(widget.cget('background'))
    except Exception:
        bg = ''
    if any(x in text_value for x in ('löschen', 'entfernen', 'abbrechen', 'verwerfen', 'deaktivieren', 'zurücksetzen')) or bg in (_fm455_norm_color(RED), '#FEE2E2', '#FFF1F1'):
        return 'danger'
    if any(x in text_value for x in ('speichern', 'ok', 'übernehmen', 'export starten', 'erstellen', 'anlegen', 'bestätigen', 'fertig', 'ausführen')) or bg in ('#CFEAD6', '#DDF3E4', '#E1F5E7'):
        return 'success'
    if any(x in text_value for x in ('suchen', 'start', 'übersicht', 'word-export', 'bearbeiten', 'hilfe', '[i] hilfe', 'öffnen', 'anzeigen')) or bg == _fm455_norm_color(BLUE):
        return 'primary'
    return 'neutral'


def _fm455_int_option(widget, option, default):
    try:
        value = str(widget.cget(option) or '').strip()
        if value.isdigit():
            return int(value)
    except Exception:
        pass
    return default


def _fm455_raw_config(widget, options):
    """Wichtig: nutzt Original-configure, damit kein rekursiver configure-Patch entsteht."""
    try:
        if isinstance(widget, tk.Menubutton) and _FM455_ORIG_TK_MENUBUTTON_CONFIGURE:
            return _FM455_ORIG_TK_MENUBUTTON_CONFIGURE(widget, None, **options)
    except Exception:
        pass
    try:
        if _FM455_ORIG_TK_BUTTON_CONFIGURE:
            return _FM455_ORIG_TK_BUTTON_CONFIGURE(widget, None, **options)
    except Exception:
        pass
    return None


def _fm455_apply_soft3d_button(widget, *, force=False):
    wid = id(widget)
    if wid in _FM455_CONFIG_GUARD:
        return widget
    _FM455_CONFIG_GUARD.add(wid)
    try:
        if not force and _fm455_should_skip_button(widget):
            return widget
        kind = _fm455_button_kind(widget)
        if kind == 'danger':
            bg, active, fg = _FM455_BTN_DANGER_BG, _FM455_BTN_DANGER_ACTIVE, RED
        elif kind == 'success':
            bg, active, fg = _FM455_BTN_SUCCESS_BG, _FM455_BTN_SUCCESS_ACTIVE, TEXT
        elif kind == 'primary':
            bg, active, fg = _FM455_BTN_PRIMARY_BG, _FM455_BTN_PRIMARY_ACTIVE, BLUE
        else:
            bg, active, fg = _FM455_BTN_NEUTRAL_BG, _FM455_BTN_NEUTRAL_ACTIVE, TEXT
        try:
            if str(widget.cget('state')) == 'disabled':
                bg, active, fg = _FM455_BTN_DISABLED_BG, _FM455_BTN_DISABLED_BG, _FM455_BTN_DISABLED_FG
        except Exception:
            pass
        padx = max(8, _fm455_int_option(widget, 'padx', 8))
        pady = max(3, _fm455_int_option(widget, 'pady', 3))
        if len(_fm455_button_text(widget)) <= 2:
            padx, pady = 6, 2
        options = {
            'bg': bg,
            'fg': fg,
            'activebackground': active,
            'activeforeground': fg,
            'relief': 'raised',
            'bd': 2,
            'highlightthickness': 1,
            'highlightbackground': _FM455_BTN_BORDER,
            'highlightcolor': BLUE,
            'cursor': 'hand2',
            'overrelief': 'ridge',
            'padx': padx,
            'pady': pady,
        }
        for key, value in options.items():
            try:
                _fm455_raw_config(widget, {key: value})
            except Exception:
                pass
        return widget
    finally:
        try:
            _FM455_CONFIG_GUARD.discard(wid)
        except Exception:
            pass


def _fm455_button_init(self, *args, **kwargs):
    _FM455_ORIG_TK_BUTTON_INIT(self, *args, **kwargs)
    _fm455_apply_soft3d_button(self)


def _fm455_button_configure(self, cnf=None, **kw):
    result = _FM455_ORIG_TK_BUTTON_CONFIGURE(self, cnf, **kw)
    if cnf or kw:
        _fm455_apply_soft3d_button(self)
    return result


def _fm455_menubutton_init(self, *args, **kwargs):
    _FM455_ORIG_TK_MENUBUTTON_INIT(self, *args, **kwargs)
    _fm455_apply_soft3d_button(self, force=True)


def _fm455_menubutton_configure(self, cnf=None, **kw):
    result = _FM455_ORIG_TK_MENUBUTTON_CONFIGURE(self, cnf, **kw)
    if cnf or kw:
        _fm455_apply_soft3d_button(self, force=True)
    return result


def _fm455_ttk_button_init(self, *args, **kwargs):
    if 'style' not in kwargs:
        kwargs['style'] = 'FiBuSoft3D.TButton'
    _FM455_ORIG_TTK_BUTTON_INIT(self, *args, **kwargs)


def _fm455_install_ttk_style(root=None):
    """Kein implizites Tk-Root: Style nur installieren, wenn ein echtes Root existiert."""
    try:
        if root is None:
            root = getattr(tk, '_default_root', None)
        if root is None:
            return False
        try:
            if not bool(root.winfo_exists()):
                return False
        except Exception:
            return False
        style = ttk.Style(root)
        style.configure('FiBuSoft3D.TButton',
                        background=_FM455_BTN_NEUTRAL_BG,
                        foreground=TEXT,
                        borderwidth=2,
                        relief='raised',
                        padding=(10, 5),
                        font=body_font(10))
        style.map('FiBuSoft3D.TButton',
                  background=[('active', _FM455_BTN_NEUTRAL_ACTIVE), ('pressed', '#D6E2EF'), ('disabled', _FM455_BTN_DISABLED_BG)],
                  foreground=[('disabled', _FM455_BTN_DISABLED_FG)])
        return True
    except Exception:
        return False


def install_fm455_global_button_design(root=None):
    """Globales Buttondesign Variante 4; Menükacheln bleiben als Canvas-Tiles unverändert."""
    global _FM455_INSTALLED, _FM455_ORIG_TK_BUTTON_INIT, _FM455_ORIG_TK_BUTTON_CONFIGURE
    global _FM455_ORIG_TK_MENUBUTTON_INIT, _FM455_ORIG_TK_MENUBUTTON_CONFIGURE, _FM455_ORIG_TTK_BUTTON_INIT
    if _FM455_ORIG_TK_BUTTON_INIT is None:
        _FM455_ORIG_TK_BUTTON_INIT = tk.Button.__init__
    if _FM455_ORIG_TK_BUTTON_CONFIGURE is None:
        _FM455_ORIG_TK_BUTTON_CONFIGURE = tk.Button.configure
    if _FM455_ORIG_TK_MENUBUTTON_INIT is None:
        _FM455_ORIG_TK_MENUBUTTON_INIT = tk.Menubutton.__init__
    if _FM455_ORIG_TK_MENUBUTTON_CONFIGURE is None:
        _FM455_ORIG_TK_MENUBUTTON_CONFIGURE = tk.Menubutton.configure
    if _FM455_ORIG_TTK_BUTTON_INIT is None:
        _FM455_ORIG_TTK_BUTTON_INIT = ttk.Button.__init__
    try:
        tk.Button.__init__ = _fm455_button_init
        tk.Button.config = _fm455_button_configure
        tk.Button.configure = _fm455_button_configure
    except Exception:
        pass
    try:
        tk.Menubutton.__init__ = _fm455_menubutton_init
        tk.Menubutton.config = _fm455_menubutton_configure
        tk.Menubutton.configure = _fm455_menubutton_configure
    except Exception:
        pass
    try:
        ttk.Button.__init__ = _fm455_ttk_button_init
    except Exception:
        pass
    _fm455_install_ttk_style(root)
    _FM455_INSTALLED = True


# Klassenpatch sofort installieren: erzeugt kein Tk-Root und wirkt dadurch global auch für später geladene Module.
install_fm455_global_button_design()

try:
    _fm455_old_run = FiBuMateApp.run
    def _fm455_run(self):
        try:
            install_fm455_global_button_design(getattr(self, 'root', None))
        except Exception:
            pass
        return _fm455_old_run(self)
    FiBuMateApp.run = _fm455_run
except Exception:
    pass



# ------------------------------------------------------------------
# Wissenszentrale - aktive Formatbuttons FINAL 2026-06-18
# Version 0.456
# Zweck:
# - Textformat-Buttons bleiben optisch eingedrückt, solange die Option an der aktuellen Auswahl/Cursorposition aktiv ist.
# - Gilt für Fett, Kursiv, Unterstrichen sowie Textfarben Standard/Rot/Blau.
# - Nutzt die vorhandene v0.446-Formatlogik inkl. Farbtags; überschreibt nur die Toolbar-Darstellung und Aktualisierung.
# ------------------------------------------------------------------

_FM456_ACTIVE_BG = "#DDEAF7"
_FM456_ACTIVE_BG_RED = "#FEE2E2"
_FM456_ACTIVE_BG_BLUE = "#DDEAF7"
_FM456_INACTIVE_BG = WHITE


def _fm456_current_format(widget):
    try:
        try:
            index = widget.index("sel.first")
        except Exception:
            index = widget.index("insert")
        return _wz446_format_at(widget, index)
    except Exception:
        return {"size": 10, "bold": False, "italic": False, "underline": False, "color": "standard"}


def _fm456_button_style(button, active=False, color_key=None):
    try:
        if active:
            bg = _FM456_ACTIVE_BG_RED if color_key == "rot" else (_FM456_ACTIVE_BG_BLUE if color_key == "blau" else _FM456_ACTIVE_BG)
            button.configure(relief="sunken", bd=2, bg=bg, activebackground=bg)
        else:
            button.configure(relief="raised", bd=1, bg=_FM456_INACTIVE_BG, activebackground="#E7F0FA")
    except Exception:
        pass


def _fm456_update_format_buttons(widget, buttons):
    fmt = _fm456_current_format(widget)
    try:
        _fm456_button_style(buttons.get("bold"), bool(fmt.get("bold")))
        _fm456_button_style(buttons.get("italic"), bool(fmt.get("italic")))
        _fm456_button_style(buttons.get("underline"), bool(fmt.get("underline")))
        active_color = _wz446_color_key(fmt.get("color", "standard"))
        for key in ("standard", "rot", "blau"):
            _fm456_button_style(buttons.get("color_" + key), active_color == key, key)
    except Exception:
        pass


def _fm456_apply_and_update(widget, buttons, mark_unsaved, *, toggle=None, color=None, size=None):
    try:
        _wz446_apply_selection(widget, toggle=toggle, color=color, size=size)
        mark_unsaved()
    except Exception:
        pass
    try:
        widget.after_idle(lambda: _fm456_update_format_buttons(widget, buttons))
    except Exception:
        _fm456_update_format_buttons(widget, buttons)


def _fm456_toolbar(parent, widget, mark_unsaved):
    bar = tk.Frame(parent, bg=WHITE)
    bar.pack(fill="x", pady=(0, 6))
    buttons = {}
    tk.Label(bar, text="Format:", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="left", padx=(0, 6))
    size_var = tk.StringVar(value="10")
    size_cb = ttk.Combobox(bar, textvariable=size_var, values=["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "28", "32"], width=4, state="readonly", font=body_font(9))
    size_cb.pack(side="left", padx=(0, 6))
    size_cb.bind("<<ComboboxSelected>>", lambda e: _fm456_apply_and_update(widget, buttons, mark_unsaved, size=int(size_var.get())))
    for label, key in [("B", "bold"), ("I", "italic"), ("U", "underline")]:
        btn = tk.Button(bar, text=label, command=lambda k=key: _fm456_apply_and_update(widget, buttons, mark_unsaved, toggle=k), bg=WHITE, fg=TEXT, font=body_font(10, weight="bold" if key == "bold" else None, underline=(key == "underline")), relief="raised", bd=1, width=3)
        btn.pack(side="left", padx=(0, 4)); buttons[key] = btn
    tk.Label(bar, text="Farbe:", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="left", padx=(10, 4))
    for key in ("standard", "rot", "blau"):
        btn = tk.Button(bar, text=_WZ446_COLOR_LABELS[key], command=lambda c=key: _fm456_apply_and_update(widget, buttons, mark_unsaved, color=c), bg=WHITE, fg=_WZ446_COLOR_MAP[key], font=body_font(9, weight="bold" if key != "standard" else None), relief="raised", bd=1)
        btn.pack(side="left", padx=(0, 4), ipadx=5); buttons["color_" + key] = btn
    app = getattr(mark_unsaved, "__self__", None)
    if app is not None:
        try:
            tk.Button(bar, text="Bild aus Zwischenablage", command=lambda: _wz446_add_clipboard_image(app), bg=WHITE, fg=TEXT, font=body_font(9), relief="raised", bd=1).pack(side="left", padx=(10, 4), ipadx=5)
            tk.Button(bar, text="Bild aus Datei", command=lambda: _wz446_insert_image_file(app), bg=WHITE, fg=TEXT, font=body_font(9), relief="raised", bd=1).pack(side="left", padx=(0, 4), ipadx=5)
        except Exception:
            pass
    tk.Label(bar, text="Text markieren und Format wählen.", bg=WHITE, fg=TEXT2, font=body_font(8)).pack(side="left", padx=(8, 0))
    def schedule_update(event=None):
        try:
            widget.after_idle(lambda: _fm456_update_format_buttons(widget, buttons))
        except Exception:
            _fm456_update_format_buttons(widget, buttons)
        return None
    try:
        for seq in ("<KeyRelease>", "<ButtonRelease-1>", "<FocusIn>", "<<Selection>>"):
            widget.bind(seq, schedule_update, add="+")
        widget.bind("<Control-b>", lambda e: (_fm456_apply_and_update(widget, buttons, mark_unsaved, toggle="bold"), "break")[-1], add="+")
        widget.bind("<Control-i>", lambda e: (_fm456_apply_and_update(widget, buttons, mark_unsaved, toggle="italic"), "break")[-1], add="+")
        widget.bind("<Control-u>", lambda e: (_fm456_apply_and_update(widget, buttons, mark_unsaved, toggle="underline"), "break")[-1], add="+")
    except Exception:
        pass
    schedule_update()
    return bar

# Finale Kompatibilitätszuweisung: vorhandene Renderfunktionen rufen _wz439_toolbar oder _wz446_toolbar direkt auf.
_wz446_toolbar = _fm456_toolbar
_wz439_toolbar = _fm456_toolbar


# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Kategorie-Filter + Untermenue-Icon FINAL 2026-06-18
# Version 0.465
# Zweck:
# - Kategorie-Filter der Wissenszentrale darf bei gemischten/alten Kategorie-Datenformaten nicht mehr abstuerzen.
# - Filterausfuehrung aus Combobox-Events wird entkoppelt, damit Tk-Widgets nicht waehrend des Auswahl-Callbacks zerstoert werden.
# - Menuekacheln, die in ein Untermenue fuehren, nutzen das neue Entrance-/Login-Pfeil-Icon.
# - Hauptmenuekacheln bleiben bewusst unveraendert.
# ------------------------------------------------------------------

_FM465_SUBMENU_ICON_FILE = "arrow_entrance_in_internet_log_login_security_icon_127060.ico"
try:
    ICON_FILES["submenu_enter"] = _FM465_SUBMENU_ICON_FILE
except Exception:
    pass


def _fm465_text(value):
    """Robuste Textnormalisierung fuer Suche/Filter ohne TypeError bei Listen, Dicts oder None."""
    try:
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            parts = []
            for key in ("name", "title", "label", "value", "text", "user", "status", "rhythm"):
                if key in value:
                    parts.append(_fm465_text(value.get(key)))
            if parts:
                return " ".join(p for p in parts if p)
            return " ".join(_fm465_text(v) for v in value.values())
        if isinstance(value, (list, tuple, set)):
            return " ".join(_fm465_text(v) for v in value)
        return str(value)
    except Exception:
        try:
            return str(value)
        except Exception:
            return ""


def _fm465_norm_category_value(value):
    """Liefert einen sauberen Kategorienamen fuer String/Dict/sonstige Altformate."""
    try:
        if value is None:
            return ""
        if isinstance(value, dict):
            for key in ("name", "title", "label", "value", "category"):
                candidate = _fm465_norm_category_value(value.get(key))
                if candidate:
                    return candidate
            return ""
        value = str(value).strip()
        if not value:
            return ""
        return " ".join(value.split())
    except Exception:
        return ""


def _fm465_categories(raw):
    """Normalisiert Kategorien auf Liste[str] - tolerant gegen Liste, Text, Dict und Altimporte."""
    result = []
    try:
        if raw is None:
            values = []
        elif isinstance(raw, dict):
            values = [raw]
        elif isinstance(raw, (list, tuple, set)):
            values = list(raw)
        else:
            values = [raw]
        seen = set()
        for item in values:
            if isinstance(item, str) and re.search(r"[;,|]", item):
                sub_values = re.split(r"[;,|]", item)
            else:
                sub_values = [item]
            for sub_item in sub_values:
                name = _fm465_norm_category_value(sub_item)
                if not name:
                    continue
                marker = name.casefold()
                if marker not in seen:
                    seen.add(marker)
                    result.append(name)
    except Exception:
        pass
    return result


def _fm465_filter_values(self):
    values = []
    try:
        self.kb_ensure_state_vars()
        for var in getattr(self, "kb_filter_vars", []) or []:
            try:
                name = _fm465_norm_category_value(var.get())
            except Exception:
                name = ""
            if name:
                marker = name.casefold()
                if marker not in values:
                    values.append(marker)
    except Exception:
        pass
    return values


def _fm465_kb_filtered_entries(self):
    """Absturzsichere Filterlogik fuer Suche, Kategorie-Filter und To-Do-Rhythmus."""
    try:
        self.kb_ensure_state_vars()
    except Exception:
        pass
    try:
        entries = self.kb_load_entries()
    except Exception as exc:
        try:
            messagebox.showerror("Wissenszentrale", "Wissenseinträge konnten nicht geladen werden:\n" + str(exc))
        except Exception:
            pass
        entries = []
    if not isinstance(entries, list):
        entries = []

    try:
        search = _fm465_text(getattr(self, "kb_search_var", tk.StringVar(value="")).get()).strip().casefold()
    except Exception:
        search = ""
    filters = _fm465_filter_values(self)
    try:
        rhythm = _fm465_text(getattr(self, "kb_todo_rhythm_var", tk.StringVar(value="")).get()).strip().casefold()
    except Exception:
        rhythm = ""
    view = str(getattr(self, "knowledge_view", "all") or "all")

    result = []
    for entry in entries:
        try:
            if not isinstance(entry, dict):
                continue
            categories = _fm465_categories(entry.get("categories", []))
            cats_cf = [c.casefold() for c in categories]
            if view == "todos" and "to-do" not in cats_cf:
                continue
            if view == "todos" and rhythm and _fm465_text(entry.get("rhythm", "")).strip().casefold() != rhythm:
                continue
            if filters and not all(f in cats_cf for f in filters):
                continue
            if search:
                hay = " ".join([
                    _fm465_text(entry.get("title", "")),
                    _fm465_text(entry.get("text", "")),
                    _fm465_text(entry.get("user", "")),
                    _fm465_text(entry.get("status", "")),
                    _fm465_text(entry.get("rhythm", "")),
                    " ".join(categories),
                ]).casefold()
                if search not in hay:
                    continue
            result.append(entry)
        except Exception:
            continue

    def _sort_key(entry):
        try:
            return _fm465_text(entry.get("updated_at") or entry.get("created_at") or "")
        except Exception:
            return ""
    return sorted(result, key=_sort_key, reverse=True)


def _fm465_kb_apply_filters(self, event=None):
    """Filter anwenden, aber Combobox-Callbacks sicher nach Tk-idle verschieben."""
    def _run():
        try:
            self._fm465_filter_render_pending = False
        except Exception:
            pass
        try:
            self.kb_selected_entry_id = None
            current_view = getattr(self, "knowledge_view", "all")
            if current_view not in ("new", "categories", "outdated"):
                self.knowledge_view = "todos" if current_view == "todos" else "all"
            self.render_page()
        except Exception as exc:
            try:
                messagebox.showerror("Wissenszentrale", "Filter konnte nicht angewendet werden:\n" + str(exc))
            except Exception:
                pass
    try:
        if event is not None and hasattr(self, "root") and self.root:
            if getattr(self, "_fm465_filter_render_pending", False):
                return "break"
            self._fm465_filter_render_pending = True
            self.root.after_idle(_run)
            return "break"
    except Exception:
        pass
    _run()
    return "break"


def _fm465_kb_on_search_return(self, event=None):
    return _fm465_kb_apply_filters(self, event)


try:
    FiBuMateApp.kb_filtered_entries = _fm465_kb_filtered_entries
    FiBuMateApp.kb_apply_filters = _fm465_kb_apply_filters
    FiBuMateApp.kb_on_search_return = _fm465_kb_on_search_return
except Exception:
    pass


try:
    _fm465_old_module_icon_type = FiBuMateApp.module_icon_type
except Exception:
    _fm465_old_module_icon_type = None


def _fm465_module_icon_type(self, module_id):
    """Untermenue-Kacheln im Modulmenue erhalten das neue Eingangspfeil-Icon."""
    try:
        if str(module_id or "").startswith("page:"):
            return "submenu_enter"
    except Exception:
        pass
    if _fm465_old_module_icon_type:
        return _fm465_old_module_icon_type(self, module_id)
    return "modules"


try:
    FiBuMateApp.module_icon_type = _fm465_module_icon_type
except Exception:
    pass


def _fm465_render_center_menu(self, items, title=""):
    """Zentrale Untermenue-/Funktionsauswahl mit neuem Untermenue-Icon.

    Wird nicht fuer Hauptmenuekacheln verwendet; das Hauptmenue bleibt dadurch optisch unveraendert.
    """
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = int(max(260, min(360, w * 0.2)))
    tile_h = int(max(110, min(160, h * 0.14)))
    gap = int(max(18, h * 0.03))
    start_y = y_pct(h, 55)
    start_x = x_pct(w, 50)
    for i, (label, page) in enumerate(items or []):
        y = start_y + i * (tile_h + gap)
        cmd = (lambda p=page, l=label: self.show_page(p, l, True))
        tile = Tile(self.root, self, f"center_{page}", label, cmd, favorite_enabled=False, center_text=True, icon_type="submenu_enter")
        tile.resize_tile(tile_w, tile_h)
        self.widget_items.append(tile)
        self.focusable_tiles.append(tile)
        self.canvas.create_window(start_x, y, window=tile, anchor="center")


try:
    FiBuMateApp.render_center_menu = _fm465_render_center_menu
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Kategorie-Filter Auswahlabsturz FIX 2026-06-23
# Version 0.466
# Zweck:
# - Behebt den Absturz beim Auswaehlen einer Filterkategorie in der Wissenszentrale.
# - Ursache: Der aktive Dropdown-/Menue-Callback konnte sofort ein render_page() ausloesen,
#   wodurch Tk-Widgets waehrend der laufenden Menueauswahl zerstoert wurden.
# - Loesung: Filter-Renderings werden konsequent per after_idle entkoppelt und re-entrant geschuetzt.
# - Die farbige Kategorieauswahl bleibt erhalten; Suche, To-Dos und Veraltete Eintraege bleiben kompatibel.
# ------------------------------------------------------------------

_FM466_FILTER_PATCH_VERSION = "0.466"


def _fm466_safe_call_filter_command(app, command):
    """Fuehrt Filterkommandos erst nach Abschluss des aktuellen Tk-Menuecallbacks aus."""
    if not command:
        return
    def _run():
        try:
            command()
        except TypeError:
            try:
                command(None)
            except Exception as exc:
                try:
                    messagebox.showerror("Wissenszentrale", "Filter konnte nicht angewendet werden:\n" + str(exc))
                except Exception:
                    pass
        except Exception as exc:
            try:
                messagebox.showerror("Wissenszentrale", "Filter konnte nicht angewendet werden:\n" + str(exc))
            except Exception:
                pass
    try:
        root = getattr(app, "root", None)
        if root:
            root.after_idle(_run)
            return
    except Exception:
        pass
    _run()


def _fm466_filter_dropdown(parent, app, var, values, width=17, command=None):
    """Stabiles farbiges Kategorie-Dropdown ohne Sofort-Render im Menuecallback."""
    box = tk.Frame(parent, bg=BG)
    value_lbl = tk.Label(box, text=(var.get() or '(alle)'), bg=WHITE, fg=TEXT,
                         font=body_font(10), anchor='w', relief='flat', padx=8)
    value_lbl.pack(side='left', fill='x', expand=True, ipady=5)
    arrow = tk.Menubutton(
        box, text='▼', bg=_FM451_BTN_BG if '_FM451_BTN_BG' in globals() else WHITE,
        fg=TEXT, activebackground=_FM451_BTN_ACTIVE if '_FM451_BTN_ACTIVE' in globals() else '#DDEAF7',
        relief='flat', bd=0, font=body_font(8), cursor='hand2', highlightthickness=1,
        highlightbackground=_FM451_BTN_BORDER if '_FM451_BTN_BORDER' in globals() else LINE, width=2)
    menu = tk.Menu(arrow, tearoff=False)
    arrow.configure(menu=menu)
    arrow.pack(side='left', padx=(3, 0), ipady=3)

    def _cat_color(value):
        try:
            return _fm451_cat_color(app, value)
        except Exception:
            try:
                return app.kb_get_category_color(value)
            except Exception:
                return _wz_cat_default_color(value)

    def _soften(color, factor=0.78):
        try:
            return _fm451_soften(color, factor)
        except Exception:
            try:
                return _wz442_light_color(color, factor)
            except Exception:
                return WHITE

    def repaint():
        try:
            val = (var.get() or '').strip()
            if val:
                color = _cat_color(val)
                value_lbl.configure(text=val, bg=_soften(color, 0.78), fg=_wz_cat_fg(color))
            else:
                value_lbl.configure(text='(alle)', bg=WHITE, fg=TEXT)
        except Exception:
            pass

    def select(value):
        try:
            var.set(value)
            repaint()
        finally:
            _fm466_safe_call_filter_command(app, command)

    try:
        menu.add_command(label='(alle)', command=lambda: select(''), background=WHITE, foreground=TEXT)
    except Exception:
        menu.add_command(label='(alle)', command=lambda: select(''))
    seen = set()
    for value in values or []:
        if not value:
            continue
        marker = str(value).casefold()
        if marker in seen:
            continue
        seen.add(marker)
        color = _cat_color(value)
        try:
            menu.add_command(label='  ' + str(value), command=lambda v=value: select(v),
                             background=_soften(color, 0.40), foreground=_wz_cat_fg(color),
                             activebackground=color, activeforeground=_wz_cat_fg(color))
        except Exception:
            menu.add_command(label='  ' + str(value), command=lambda v=value: select(v))
    repaint()
    return box


def _fm466_kb_apply_filters(self, event=None):
    """Absturzsicheres Filter-Rendering: nie direkt aus Auswahl-/Menuecallbacks rendern."""
    def _run():
        try:
            self._fm466_filter_render_pending = False
            self._fm466_filter_render_running = True
            self.kb_selected_entry_id = None
            current_view = getattr(self, "knowledge_view", "all")
            if current_view not in ("new", "categories", "outdated"):
                self.knowledge_view = "todos" if current_view == "todos" else "all"
            self.render_page()
        except Exception as exc:
            try:
                messagebox.showerror("Wissenszentrale", "Filter konnte nicht angewendet werden:\n" + str(exc))
            except Exception:
                pass
        finally:
            try:
                self._fm466_filter_render_running = False
            except Exception:
                pass
    try:
        if getattr(self, "_fm466_filter_render_running", False):
            return "break"
        root = getattr(self, "root", None)
        if root:
            if getattr(self, "_fm466_filter_render_pending", False):
                return "break"
            self._fm466_filter_render_pending = True
            root.after_idle(_run)
            return "break"
    except Exception:
        pass
    _run()
    return "break"


def _fm466_kb_on_search_return(self, event=None):
    return _fm466_kb_apply_filters(self, event)


# Finale Zuweisungen direkt vor Programmstart: ueberschreibt v0.451/v0.465 Filter-Callbacks.
try:
    _fm451_filter_dropdown = _fm466_filter_dropdown
except Exception:
    pass
try:
    FiBuMateApp.kb_apply_filters = _fm466_kb_apply_filters
    FiBuMateApp.kb_on_search_return = _fm466_kb_on_search_return
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Aktualisiert + Aktuell-Kennzeichnung FINAL 2026-06-23
# Version 0.467
# Zweck:
# - Wissenszentrale: "Geändert" wird fachlich zu "Aktualisiert".
# - Trefferliste zeigt den Benutzer/Vollnamen zusätzlich an.
# - Detailansicht und Übersicht erhalten den Button "Eintrag als aktuell kennzeichnen".
# - Der Button setzt nach Ja/Nein-Abfrage updated_at auf das aktuelle Datum/die aktuelle Zeit.
# - Kategorie-Filter zeigen ausgewählte Kategorien mit voller Farbe statt entsättigter Farbe.
# - Untermenü-Kacheln außerhalb des Hauptmenüs nutzen das Entrance-/Login-Pfeil-Icon.
# ------------------------------------------------------------------

_FM467_PATCH_VERSION = "0.467"
_FM467_SUBMENU_ICON_FILE = "arrow_entrance_in_internet_log_login_security_icon_127060.ico"

try:
    ICON_FILES["submenu_enter"] = _FM467_SUBMENU_ICON_FILE
except Exception:
    pass


def _fm467_text(value):
    try:
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        if isinstance(value, (list, tuple, set)):
            return " ".join(_fm467_text(x) for x in value if x is not None)
        if isinstance(value, dict):
            return " ".join(_fm467_text(v) for v in value.values() if v is not None)
        return str(value)
    except Exception:
        return ""


def _fm467_user_fullname(self, raw_user):
    """Bestmoegliche Vollnamenanzeige fuer Benutzerfelder in der Wissenszentrale."""
    raw = _fm467_text(raw_user).strip()
    if not raw:
        return ""
    # Wenn bereits ein sprechender Anzeigename gespeichert ist, direkt verwenden.
    if " " in raw or "," in raw:
        return raw
    wanted = raw.casefold()
    candidates = []
    try:
        data = getattr(self, "user_data", None)
        if isinstance(data, dict):
            candidates.extend(data.values())
            candidates.append(data)
        elif isinstance(data, list):
            candidates.extend(data)
    except Exception:
        pass
    seen = set()
    stack = list(candidates)
    while stack:
        item = stack.pop(0)
        marker = id(item)
        if marker in seen:
            continue
        seen.add(marker)
        if isinstance(item, dict):
            keys = ["username", "user", "key", "login", "kuerzel", "kürzel", "id", "name"]
            vals = [_fm467_text(item.get(k)).strip() for k in keys]
            if any(v and v.casefold() == wanted for v in vals):
                for dk in ("display_name", "display", "fullname", "full_name", "vollname", "name"):
                    val = _fm467_text(item.get(dk)).strip()
                    if val:
                        return val
            for v in item.values():
                if isinstance(v, (dict, list, tuple)):
                    stack.append(v)
        elif isinstance(item, (list, tuple)):
            stack.extend(item)
    return raw


def _fm467_now(self):
    try:
        return self.kb_now()
    except Exception:
        try:
            return datetime.now().isoformat(timespec="seconds")
        except Exception:
            return str(datetime.now())


def _fm467_set_entry_current(self, entry_id=None):
    entry_id = entry_id or getattr(self, "kb_selected_entry_id", None)
    if not entry_id:
        try:
            messagebox.showwarning("Wissenszentrale", "Bitte zuerst einen Eintrag auswählen.")
        except Exception:
            pass
        return False
    try:
        ok = messagebox.askyesno("Wissenszentrale", "Den Inhalt dieses Eintrags als aktuell kennzeichnen?")
    except Exception:
        ok = False
    if not ok:
        return False
    entries = self.kb_load_entries()
    now = _fm467_now(self)
    changed = False
    for entry in entries:
        try:
            if entry.get("id") == entry_id:
                entry["updated_at"] = now
                changed = True
                break
        except Exception:
            continue
    if not changed:
        try:
            messagebox.showwarning("Wissenszentrale", "Der ausgewählte Eintrag wurde nicht gefunden.")
        except Exception:
            pass
        return False
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = entry_id
        try:
            self.knowledge_unsaved = False
        except Exception:
            pass
        try:
            self.render_page()
        except Exception:
            pass
        return True
    return False


def _fm467_cat_color(app, value):
    try:
        return app.kb_get_category_color(value)
    except Exception:
        try:
            return _wz442_cat_color(app, value)
        except Exception:
            return _wz_cat_default_color(value)


def _fm467_filter_dropdown(parent, app, var, values, width=17, command=None):
    """Kategorie-Filter mit voller Farbe fuer die aktuell ausgewaehlte Kategorie."""
    box = tk.Frame(parent, bg=BG)
    value_lbl = tk.Label(box, text=(var.get() or '(alle)'), bg=WHITE, fg=TEXT,
                         font=body_font(10), anchor='w', relief='flat', padx=8)
    value_lbl.pack(side='left', fill='x', expand=True, ipady=5)
    arrow = tk.Menubutton(
        box, text='▼', bg=_FM451_BTN_BG if '_FM451_BTN_BG' in globals() else WHITE,
        fg=TEXT, activebackground=_FM451_BTN_ACTIVE if '_FM451_BTN_ACTIVE' in globals() else '#DDEAF7',
        relief='flat', bd=0, font=body_font(8), cursor='hand2', highlightthickness=1,
        highlightbackground=_FM451_BTN_BORDER if '_FM451_BTN_BORDER' in globals() else LINE, width=2)
    menu = tk.Menu(arrow, tearoff=False)
    arrow.configure(menu=menu)
    arrow.pack(side='left', padx=(3, 0), ipady=3)

    def repaint():
        try:
            val = (var.get() or '').strip()
            if val:
                color = _fm467_cat_color(app, val)
                value_lbl.configure(text=val, bg=color, fg=_wz_cat_fg(color))
            else:
                value_lbl.configure(text='(alle)', bg=WHITE, fg=TEXT)
        except Exception:
            pass

    def _deferred_command():
        if not command:
            return
        try:
            if hasattr(app, 'root') and app.root:
                app.root.after_idle(command)
                return
        except Exception:
            pass
        try:
            command()
        except TypeError:
            command(None)

    def select(value):
        var.set(value)
        repaint()
        _deferred_command()

    try:
        menu.add_command(label='(alle)', command=lambda: select(''), background=WHITE, foreground=TEXT)
    except Exception:
        menu.add_command(label='(alle)', command=lambda: select(''))
    seen = set()
    for value in values or []:
        if not value:
            continue
        key = str(value).casefold()
        if key in seen:
            continue
        seen.add(key)
        color = _fm467_cat_color(app, value)
        try:
            menu.add_command(label='  ' + str(value), command=lambda v=value: select(v),
                             background=color, foreground=_wz_cat_fg(color),
                             activebackground=color, activeforeground=_wz_cat_fg(color))
        except Exception:
            menu.add_command(label='  ' + str(value), command=lambda v=value: select(v))
    repaint()
    return box


def _fm467_render_hits_pane(self, x, y, w, h):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text="Treffer", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
    entries = self.kb_filtered_entries()
    listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="flat", activestyle="none", exportselection=False)
    listbox.pack(fill="both", expand=True, padx=16, pady=(0, 14))
    id_map = []
    if entries:
        for idx, entry in enumerate(entries):
            cats = entry.get("categories", []) or []
            cats_txt = ", ".join(_fm467_text(c) for c in cats if _fm467_text(c).strip())
            user_txt = _fm467_user_fullname(self, entry.get("user", ""))
            line = f"{self.kb_display_date(entry.get('updated_at'))}  |  {entry.get('title','')}"
            if user_txt:
                line += f"  |  {user_txt}"
            if cats_txt:
                line += f"  [{cats_txt}]"
            listbox.insert("end", line)
            id_map.append(entry.get("id"))
            if cats:
                c = _fm467_cat_color(self, cats[0])
                try:
                    listbox.itemconfig(idx, bg=_wz442_light_color(c), fg=TEXT, selectbackground=c, selectforeground=_wz_cat_fg(c))
                except Exception:
                    pass
    else:
        listbox.insert("end", "Noch keine Treffer vorhanden.")
        listbox.insert("end", "Filter und Suche wirken auf die Wissensdatenbank.")
    def _select(event=None):
        sel = listbox.curselection()
        if not sel or not id_map or sel[0] >= len(id_map):
            return
        if not self.kb_confirm_unsaved_before_switch():
            return
        self.kb_selected_entry_id = id_map[sel[0]]
        self.knowledge_view = "detail"
        self.knowledge_start_overlay = False
        self.render_page()
    listbox.bind("<<ListboxSelect>>", _select)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))


def _fm467_render_kb_list_area(self, x, y, w, h, title='Gesamtliste aller Einträge', status_filter=None):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    header = tk.Frame(frame, bg=WHITE)
    header.pack(fill='x', padx=16, pady=(12, 6))
    tk.Label(header, text=title, bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(side='left')
    action_btn = tk.Button(header, text='Eintrag als aktuell kennzeichnen', command=lambda: _current_from_tree(),
                           bg='#CFEAD6', fg=TEXT, font=body_font(9, weight='bold'), relief='solid', bd=1)
    action_btn.pack(side='right', padx=(8, 0), ipadx=8, ipady=3)
    sort_state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
    direction = 'absteigend' if sort_state.get('descending', True) else 'aufsteigend'
    tk.Label(header, text=f"Sortierung: {sort_state.get('column','date')} ({direction})", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')
    columns = ('date','title','categories','user','status')
    labels = {'date':'Aktualisiert','title':'Titel','categories':'Kategorien','user':'Benutzer','status':'Status'}
    tree = ttk.Treeview(frame, columns=columns, show='headings', height=14)
    try:
        style = ttk.Style()
        style.configure('Treeview', rowheight=28, font=body_font(10))
        style.configure('Treeview.Heading', font=body_font(10, weight='bold'))
    except Exception:
        pass
    for col in columns:
        marker = ''
        if sort_state.get('column') == col:
            marker = ' ↓' if sort_state.get('descending', True) else ' ↑'
        tree.heading(col, text=labels[col]+marker, command=lambda c=col: (setattr(self, 'kb_overview_sort_state', {'column': c, 'descending': not (getattr(self, 'kb_overview_sort_state', {}) or {}).get('descending', True) if (getattr(self, 'kb_overview_sort_state', {}) or {}).get('column') == c else (True if c == 'date' else False)}), self.render_page()))
    tree.column('date', width=125, stretch=False)
    tree.column('title', width=max(250, int(w*0.34)), stretch=True)
    tree.column('categories', width=max(170, int(w*0.20)), stretch=True)
    tree.column('user', width=190, stretch=False)
    tree.column('status', width=100, stretch=False)
    yscroll = ttk.Scrollbar(frame, orient='vertical', command=tree.yview)
    tree.configure(yscrollcommand=yscroll.set)
    tree.pack(side='left', fill='both', expand=True, padx=(16,0), pady=(0,16))
    yscroll.pack(side='right', fill='y', padx=(0,16), pady=(0,16))
    entries = self.kb_filtered_entries()
    if status_filter:
        entries = [e for e in entries if str(e.get('status','')).lower() == status_filter.lower()]
    state = getattr(self, 'kb_overview_sort_state', None)
    if not state:
        self.kb_overview_sort_state = {'column':'date','descending':True}
    entries = sorted(entries, key=lambda e: (
        str(e.get('updated_at') or e.get('created_at') or '') if self.kb_overview_sort_state.get('column')=='date' else
        str(e.get('title','')).casefold() if self.kb_overview_sort_state.get('column')=='title' else
        ', '.join(e.get('categories',[]) or []).casefold() if self.kb_overview_sort_state.get('column')=='categories' else
        _fm467_user_fullname(self, e.get('user','')).casefold() if self.kb_overview_sort_state.get('column')=='user' else
        str(e.get('status','')).casefold()
    ), reverse=bool(self.kb_overview_sort_state.get('descending', True)))
    id_map = {}
    for entry in entries:
        iid = entry.get('id') or self.kb_make_entry_id()
        base_iid = iid; n = 1
        while iid in id_map:
            n += 1; iid = f'{base_iid}_{n}'
        id_map[iid] = entry.get('id')
        tree.insert('', 'end', iid=iid, values=(
            self.kb_display_date(entry.get('updated_at')),
            entry.get('title',''),
            ', '.join(entry.get('categories',[]) or []),
            _fm467_user_fullname(self, entry.get('user','')),
            entry.get('status','')
        ))
    if not entries:
        tree.insert('', 'end', values=('', 'Noch keine Einträge vorhanden', '', '', ''))
    def _open(event=None):
        sel = tree.selection()
        if sel and sel[0] in id_map:
            self.kb_select_entry(id_map[sel[0]])
    def _current_from_tree():
        sel = tree.selection()
        if not sel or sel[0] not in id_map:
            try:
                messagebox.showwarning('Wissenszentrale', 'Bitte zuerst einen Eintrag in der Übersicht auswählen.')
            except Exception:
                pass
            return
        _fm467_set_entry_current(self, id_map[sel[0]])
    tree.bind('<Double-Button-1>', _open)
    tree.bind('<Return>', _open)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


def _fm467_render_kb_detail_area(self, x, y, w, h):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry:
        self.render_kb_list_area(x, y, w, h, title="Übersicht")
        return
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    top = tk.Frame(frame, bg=WHITE)
    top.pack(fill='x', padx=18, pady=(14, 5))
    tk.Label(top, text=entry.get("title", ""), bg=WHITE, fg=BLUE, font=body_font(16, weight="bold")).pack(side='left', anchor='w')
    tk.Button(top, text='Eintrag als aktuell kennzeichnen', command=lambda: _fm467_set_entry_current(self, entry.get('id')),
              bg='#CFEAD6', fg=TEXT, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='right', padx=(8, 0), ipadx=8, ipady=3)
    tk.Label(frame, text=f"Aktualisiert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {_fm467_user_fullname(self, entry.get('user',''))}    Status: {entry.get('status','')}", bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor="w", padx=18)
    try:
        _wz439_badges(frame, self, entry.get("categories", []) or [])
    except Exception:
        tk.Label(frame, text="Kategorien: " + (", ".join(entry.get("categories", []) or []) or "Keine Kategorien"), bg=WHITE, fg=TEXT, font=body_font(10, weight="bold")).pack(anchor="w", padx=18, pady=(8,4))
    try:
        _wz439_render_image_view(frame, self, entry.get("inline_images", []) or [])
    except Exception:
        pass
    text_frame = tk.Frame(frame, bg=WHITE)
    text_frame.pack(fill="both", expand=True, padx=18, pady=(6, 8))
    txt = tk.Text(text_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word", height=8)
    yscroll = tk.Scrollbar(text_frame, orient="vertical", command=txt.yview)
    txt.configure(yscrollcommand=yscroll.set)
    txt.pack(side="left", fill="both", expand=True)
    yscroll.pack(side="right", fill="y")
    txt.insert("1.0", entry.get("text", ""))
    try:
        _wz439_apply_formatting(txt, entry.get("text_formatting", []) or [])
    except Exception:
        pass
    txt.configure(state="disabled")
    try:
        txt.bind("<MouseWheel>", lambda ev: _wz439_text_mousewheel(txt, ev))
    except Exception:
        pass
    lower = tk.Frame(frame, bg=WHITE)
    lower.pack(fill="x", padx=18, pady=(0, 8))
    left = tk.Frame(lower, bg=WHITE)
    left.pack(side="left", fill="both", expand=True)
    right = tk.Frame(lower, bg=WHITE)
    right.pack(side="right", fill="both", expand=True, padx=(12, 0))
    tk.Label(left, text="Anhänge", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    attachments = entry.get("attachments", []) or []
    if attachments:
        for a in attachments:
            tk.Label(left, text="• " + a.get("name", ""), bg=WHITE, fg=TEXT, font=body_font(9)).pack(anchor="w")
    else:
        tk.Label(left, text="Keine Anhänge", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    tk.Label(right, text="Kommentare", bg=WHITE, fg=BLUE, font=body_font(10, weight="bold")).pack(anchor="w")
    for c in (entry.get("comments", []) or [])[-3:]:
        tk.Label(right, text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}", bg=WHITE, fg=TEXT, font=body_font(9), wraplength=380, justify="left").pack(anchor="w")
    comment = tk.Text(right, height=2, bg="#F8FAFC", fg=TEXT, font=body_font(9), relief="solid", bd=1, wrap="word")
    comment.pack(fill="x", pady=(4, 0))
    buttons = tk.Frame(frame, bg=WHITE)
    buttons.pack(fill="x", padx=18, pady=(0, 14))
    if self.kb_can_create_or_edit():
        tk.Button(buttons, text="Bearbeiten", command=self.kb_edit_selected_entry, bg=WHITE, fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Kommentar speichern", command=lambda: self.kb_add_comment_to_selected(comment), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Word-Export", command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
    tk.Button(buttons, text="Zur Übersicht", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))


def _fm467_export_word(self):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry:
        try: messagebox.showwarning("Wissenszentrale", "Kein Eintrag für den Word-Export ausgewählt.")
        except Exception: pass
        return
    try:
        from docx import Document
        import re as _re
        doc = Document()
        doc.add_heading(entry.get("title", "Wissenseintrag") or "Wissenseintrag", level=1)
        meta = [
            ("Aktualisiert", self.kb_display_date(entry.get("updated_at"))),
            ("Benutzer", _fm467_user_fullname(self, entry.get("user", ""))),
            ("Status", entry.get("status", "")),
            ("Kategorien", ", ".join(entry.get("categories", []) or [])),
            ("To-Do-Rhythmus", entry.get("rhythm", "")),
        ]
        for label, val in meta:
            if val:
                p = doc.add_paragraph(); p.add_run(f"{label}: ").bold = True; p.add_run(str(val))
        doc.add_heading("Inhalt", level=2)
        text_value = entry.get("text", "") or ""
        if text_value:
            for line in str(text_value).splitlines() or [""]:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("-")
        doc.add_heading("Anhänge", level=2)
        attachments = entry.get("attachments", []) or []
        if attachments:
            for a in attachments:
                doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        else:
            doc.add_paragraph("Keine Anhänge")
        doc.add_heading("Kommentare", level=2)
        comments = entry.get("comments", []) or []
        if comments:
            for c in comments:
                doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        else:
            doc.add_paragraph("Keine Kommentare")
        safe = _re.sub(r"[^A-Za-z0-9_äöüÄÖÜß.-]+", "_", entry.get("title", "Wissenseintrag"))[:80].strip("._") or "Wissenseintrag"
        out = os.path.join(self.kb_export_dir(), f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(out)
        messagebox.showinfo("Wissenszentrale", "Word-Export erstellt:\n" + out)
    except Exception as exc:
        try: messagebox.showerror("Wissenszentrale", "Word-Export fehlgeschlagen:\n" + str(exc))
        except Exception: pass


# Icon-Render-Sicherheit: submenu_enter muss auch tatsaechlich gezeichnet werden koennen.
try:
    _fm467_old_draw_tile_icon_image = FiBuMateApp.draw_tile_icon_image
except Exception:
    _fm467_old_draw_tile_icon_image = None


def _fm467_draw_tile_icon_image(self, tile, icon_type, cx, cy):
    if icon_type == "submenu_enter":
        photo = self.get_icon_photo("submenu_enter", 48, 48)
        if photo:
            tile.create_image(cx, cy, image=photo)
            return True
        return False
    if _fm467_old_draw_tile_icon_image:
        return _fm467_old_draw_tile_icon_image(self, tile, icon_type, cx, cy)
    return False


def _fm467_module_icon_type(self, module_id):
    try:
        if str(module_id or "").startswith("page:"):
            return "submenu_enter"
    except Exception:
        pass
    try:
        return _fm465_module_icon_type(self, module_id)
    except Exception:
        return "modules"


# Finale Zuweisungen direkt vor Programmstart.
try:
    _fm451_filter_dropdown = _fm467_filter_dropdown
    _fm466_filter_dropdown = _fm467_filter_dropdown
except Exception:
    pass
try:
    FiBuMateApp.kb_mark_entry_current = _fm467_set_entry_current
    FiBuMateApp.render_kb_hits_pane = _fm467_render_hits_pane
    FiBuMateApp.render_kb_list_area = _fm467_render_kb_list_area
    FiBuMateApp.render_kb_detail_area = _fm467_render_kb_detail_area
    FiBuMateApp.kb_export_selected_to_word = _fm467_export_word
    FiBuMateApp.draw_tile_icon_image = _fm467_draw_tile_icon_image
    FiBuMateApp.module_icon_type = _fm467_module_icon_type
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Übersicht Kategorie-Chips FINAL 2026-06-23
# Version 0.468
# Zweck:
# - Übersichtstabelle der Wissenszentrale zeigt Kategorien als farbige, nicht-interaktive Chips.
# - Ersetzt die reine Kategorien-Textspalte in der Übersicht durch visuelle Farbbadges je Kategorie.
# - Erhält Sortierung, Auswahl, Doppelklick-Öffnen und Button "Eintrag als aktuell kennzeichnen".
# ------------------------------------------------------------------

_FM468_PATCH_VERSION = "0.468"


def _fm468_cat_color(app, value):
    try:
        return _fm467_cat_color(app, value)
    except Exception:
        try:
            return app.kb_get_category_color(value)
        except Exception:
            return _wz_cat_default_color(value)


def _fm468_text(value):
    try:
        return _fm467_text(value)
    except Exception:
        return "" if value is None else str(value)


def _fm468_user_fullname(app, value):
    try:
        return _fm467_user_fullname(app, value)
    except Exception:
        return _fm468_text(value).strip()


def _fm468_sort_entries_for_overview(self, entries):
    state = getattr(self, 'kb_overview_sort_state', None)
    if not state:
        self.kb_overview_sort_state = {'column': 'date', 'descending': True}
        state = self.kb_overview_sort_state
    col = state.get('column', 'date')
    def key(e):
        try:
            if col == 'date':
                return str(e.get('updated_at') or e.get('created_at') or '')
            if col == 'title':
                return str(e.get('title','')).casefold()
            if col == 'categories':
                return ', '.join(e.get('categories',[]) or []).casefold()
            if col == 'user':
                return _fm468_user_fullname(self, e.get('user','')).casefold()
            return str(e.get('status','')).casefold()
        except Exception:
            return ''
    return sorted(entries or [], key=key, reverse=bool(state.get('descending', True)))


def _fm468_render_category_chips(parent, app, categories, max_width=330):
    chip_frame = tk.Frame(parent, bg=WHITE)
    used = 0
    if not categories:
        tk.Label(chip_frame, text='—', bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='left')
        return chip_frame
    for cat in categories:
        cat = _fm468_text(cat).strip()
        if not cat:
            continue
        # einfache Breitenbegrenzung; weitere Kategorien werden als +n zusammengefasst
        est = 18 + len(cat) * 7
        remaining = len([c for c in categories if _fm468_text(c).strip()])
        if used and used + est > max_width:
            tk.Label(chip_frame, text='  +' + str(max(1, remaining)) + '  ', bg='#E5E7EB', fg=TEXT, font=body_font(8, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0, 5), ipady=1)
            break
        color = _fm468_cat_color(app, cat)
        fg = _wz_cat_fg(color)
        tk.Label(chip_frame, text='  ' + cat + '  ', bg=color, fg=fg, font=body_font(8, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0, 5), ipady=1)
        used += est
    return chip_frame


def _fm468_render_kb_list_area(self, x, y, w, h, title='Gesamtliste aller Einträge', status_filter=None):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)

    header = tk.Frame(frame, bg=WHITE)
    header.pack(fill='x', padx=16, pady=(12, 6))
    tk.Label(header, text=title, bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(side='left')
    tk.Button(header, text='Eintrag als aktuell kennzeichnen', command=lambda: _current_from_selected(),
              bg='#CFEAD6', fg=TEXT, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='right', padx=(8, 0), ipadx=8, ipady=3)
    sort_state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
    direction = 'absteigend' if sort_state.get('descending', True) else 'aufsteigend'
    tk.Label(header, text=f"Sortierung: {sort_state.get('column','date')} ({direction})", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')

    table = tk.Frame(frame, bg=WHITE)
    table.pack(fill='both', expand=True, padx=16, pady=(0, 16))
    table.grid_columnconfigure(0, weight=1)
    table.grid_rowconfigure(1, weight=1)

    # Spaltenlayout wird als proportionaler Grid-Aufbau gerendert; Kategorie-Spalte nutzt echte Tk-Chips.
    columns = [
        ('date', 'Aktualisiert', 0, 16),
        ('title', 'Titel', 1, 32),
        ('categories', 'Kategorien', 2, 30),
        ('user', 'Benutzer', 3, 16),
        ('status', 'Status', 4, 10),
    ]
    head = tk.Frame(table, bg='#F1F5F9')
    head.grid(row=0, column=0, sticky='ew')
    for _, _, col_index, weight in columns:
        head.grid_columnconfigure(col_index, weight=weight)
    for key, label, col_index, _weight in columns:
        marker = ''
        if sort_state.get('column') == key:
            marker = ' ↓' if sort_state.get('descending', True) else ' ↑'
        def sort_cmd(c=key):
            current = getattr(self, 'kb_overview_sort_state', {}) or {}
            descending = not current.get('descending', True) if current.get('column') == c else (True if c == 'date' else False)
            self.kb_overview_sort_state = {'column': c, 'descending': descending}
            self.render_page()
        tk.Button(head, text=label + marker, command=sort_cmd, bg='#F1F5F9', fg=TEXT, activebackground='#E2E8F0',
                  font=body_font(9, weight='bold'), relief='flat', anchor='w').grid(row=0, column=col_index, sticky='ew', padx=(4, 4), pady=3)

    canvas = tk.Canvas(table, bg=WHITE, highlightthickness=0)
    scroll = ttk.Scrollbar(table, orient='vertical', command=canvas.yview)
    canvas.configure(yscrollcommand=scroll.set)
    canvas.grid(row=1, column=0, sticky='nsew')
    scroll.grid(row=1, column=1, sticky='ns')
    body = tk.Frame(canvas, bg=WHITE)
    win = canvas.create_window((0,0), window=body, anchor='nw')
    def _on_body_config(event=None):
        try:
            canvas.configure(scrollregion=canvas.bbox('all'))
        except Exception:
            pass
    def _on_canvas_config(event=None):
        try:
            canvas.itemconfigure(win, width=canvas.winfo_width())
        except Exception:
            pass
    body.bind('<Configure>', _on_body_config)
    canvas.bind('<Configure>', _on_canvas_config)
    try:
        canvas.bind('<MouseWheel>', lambda e: (canvas.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
        body.bind('<MouseWheel>', lambda e: (canvas.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
    except Exception:
        pass

    entries = self.kb_filtered_entries()
    if status_filter:
        entries = [e for e in entries if str(e.get('status','')).lower() == status_filter.lower()]
    entries = _fm468_sort_entries_for_overview(self, entries)
    selected = {'id': getattr(self, 'kb_selected_entry_id', None)}
    row_frames = {}

    def _paint_selection():
        for eid, row in row_frames.items():
            try:
                row.configure(bg=('#EAF2FB' if eid == selected.get('id') else WHITE))
                for child in row.winfo_children():
                    if not getattr(child, '_fm468_chip_container', False):
                        try: child.configure(bg=row.cget('bg'))
                        except Exception: pass
            except Exception:
                pass

    def _select_entry(eid):
        selected['id'] = eid
        self.kb_selected_entry_id = eid
        _paint_selection()

    def _open_entry(eid):
        if not eid:
            return
        if not self.kb_confirm_unsaved_before_switch():
            return
        self.kb_selected_entry_id = eid
        self.knowledge_view = 'detail'
        self.knowledge_start_overlay = False
        self.render_page()

    if not entries:
        empty = tk.Label(body, text='Noch keine Einträge vorhanden', bg=WHITE, fg=TEXT2, font=body_font(10), anchor='w')
        empty.pack(fill='x', padx=8, pady=10)
    else:
        for row_no, entry in enumerate(entries):
            eid = entry.get('id') or self.kb_make_entry_id()
            row = tk.Frame(body, bg=WHITE if row_no % 2 == 0 else '#FAFAFA', highlightbackground='#E5E7EB', highlightthickness=1)
            row.pack(fill='x', expand=True)
            for _, _, col_index, weight in columns:
                row.grid_columnconfigure(col_index, weight=weight)
            row_frames[eid] = row
            vals = {
                'date': self.kb_display_date(entry.get('updated_at')),
                'title': entry.get('title',''),
                'user': _fm468_user_fullname(self, entry.get('user','')),
                'status': entry.get('status',''),
            }
            tk.Label(row, text=vals['date'], bg=row.cget('bg'), fg=TEXT, font=body_font(9), anchor='w').grid(row=0, column=0, sticky='ew', padx=(6,4), pady=6)
            tk.Label(row, text=vals['title'], bg=row.cget('bg'), fg=TEXT, font=body_font(9, weight='bold'), anchor='w').grid(row=0, column=1, sticky='ew', padx=4, pady=6)
            chips = _fm468_render_category_chips(row, self, entry.get('categories', []) or [], max_width=max(220, int(w*0.24)))
            chips._fm468_chip_container = True
            chips.configure(bg=row.cget('bg'))
            chips.grid(row=0, column=2, sticky='w', padx=4, pady=5)
            tk.Label(row, text=vals['user'], bg=row.cget('bg'), fg=TEXT, font=body_font(9), anchor='w').grid(row=0, column=3, sticky='ew', padx=4, pady=6)
            tk.Label(row, text=vals['status'], bg=row.cget('bg'), fg=TEXT, font=body_font(9), anchor='w').grid(row=0, column=4, sticky='ew', padx=4, pady=6)
            for child in row.winfo_children():
                try:
                    child.bind('<Button-1>', lambda e, id=eid: _select_entry(id))
                    child.bind('<Double-Button-1>', lambda e, id=eid: _open_entry(id))
                    child.bind('<MouseWheel>', lambda e: (canvas.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
                except Exception:
                    pass
            row.bind('<Button-1>', lambda e, id=eid: _select_entry(id))
            row.bind('<Double-Button-1>', lambda e, id=eid: _open_entry(id))
            row.bind('<MouseWheel>', lambda e: (canvas.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
    _paint_selection()

    def _current_from_selected():
        eid = selected.get('id') or getattr(self, 'kb_selected_entry_id', None)
        if not eid:
            try:
                messagebox.showwarning('Wissenszentrale', 'Bitte zuerst einen Eintrag in der Übersicht auswählen.')
            except Exception:
                pass
            return
        try:
            _fm467_set_entry_current(self, eid)
        except Exception:
            try:
                self.kb_mark_entry_current(eid)
            except Exception:
                pass

    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


# Finale Zuweisung direkt vor Programmstart: Übersicht nutzt ab v0.468 echte Kategorie-Chips.
try:
    FiBuMateApp.render_kb_list_area = _fm468_render_kb_list_area
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Übersicht stabil + ESC Start FINAL 2026-06-23
# Version 0.469
# Zweck:
# - Übersichtstabelle robuster: feste Spaltenkoordinaten, alle Inhalte linksbündig.
# - Kategorie-Chips werden gezeichnet, nicht als Layout-Widgets gesetzt; dadurch kein Verschieben.
# - Escape aus dem Startbereich der Wissenszentrale geht direkt ins Hauptmenü (Seite main), nicht nach main_menu.
# ------------------------------------------------------------------

_FM469_PATCH_VERSION = "0.469"


def _fm469_txt(v):
    try: return _fm468_text(v)
    except Exception: return "" if v is None else str(v)


def _fm469_user(app, v):
    try: return _fm468_user_fullname(app, v)
    except Exception:
        try: return _fm467_user_fullname(app, v)
        except Exception: return _fm469_txt(v).strip()


def _fm469_cat_color(app, v):
    try: return _fm468_cat_color(app, v)
    except Exception:
        try: return app.kb_get_category_color(v)
        except Exception: return _wz_cat_default_color(v)


def _fm469_ell(v, max_chars):
    s = _fm469_txt(v).replace('\n', ' ').strip()
    try: max_chars = int(max_chars)
    except Exception: max_chars = 30
    if max_chars < 5: return s[:max_chars]
    return s if len(s) <= max_chars else s[:max_chars-1] + '…'


def _fm469_layout(total_width):
    tw = max(760, int(total_width or 760))
    date_w = 124
    status_w = 96
    user_w = max(150, min(210, int(tw * 0.17)))
    cat_w = max(250, min(380, int(tw * 0.30)))
    title_w = max(230, tw - date_w - cat_w - user_w - status_w)
    widths = [date_w, title_w, cat_w, user_w, status_w]
    starts = [0]
    for ww in widths[:-1]: starts.append(starts[-1] + ww)
    return starts, widths


def _fm469_sort(self, entries):
    try: return _fm468_sort_entries_for_overview(self, entries)
    except Exception:
        state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
        col = state.get('column', 'date')
        def key(e):
            if col == 'date': return str(e.get('updated_at') or e.get('created_at') or '')
            if col == 'title': return str(e.get('title','')).casefold()
            if col == 'categories': return ', '.join(e.get('categories',[]) or []).casefold()
            if col == 'user': return _fm469_user(self, e.get('user','')).casefold()
            return str(e.get('status','')).casefold()
        return sorted(entries or [], key=key, reverse=bool(state.get('descending', True)))


def _fm469_render_kb_list_area(self, x, y, w, h, title='Gesamtliste aller Einträge', status_filter=None):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    top = tk.Frame(frame, bg=WHITE); top.pack(fill='x', padx=16, pady=(12, 6))
    tk.Label(top, text=title, bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(side='left')
    tk.Button(top, text='Eintrag als aktuell kennzeichnen', command=lambda: _mark_current(), bg='#CFEAD6', fg=TEXT, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='right', padx=(8,0), ipadx=8, ipady=3)
    state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
    tk.Label(top, text=f"Sortierung: {state.get('column','date')} ({'absteigend' if state.get('descending', True) else 'aufsteigend'})", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')

    table = tk.Frame(frame, bg=WHITE); table.pack(fill='both', expand=True, padx=16, pady=(0,16))
    table.grid_columnconfigure(0, weight=1); table.grid_rowconfigure(1, weight=1)
    head = tk.Canvas(table, height=31, bg='#F1F5F9', highlightthickness=0); head.grid(row=0, column=0, sticky='ew')
    body = tk.Canvas(table, bg=WHITE, highlightthickness=0); body.grid(row=1, column=0, sticky='nsew')
    sb = ttk.Scrollbar(table, orient='vertical', command=body.yview); sb.grid(row=1, column=1, sticky='ns')
    body.configure(yscrollcommand=sb.set)
    selected = {'id': getattr(self, 'kb_selected_entry_id', None)}
    row_rects = {}
    keys = ['date','title','categories','user','status']
    labels = {'date':'Aktualisiert','title':'Titel','categories':'Kategorien','user':'Benutzer','status':'Status'}

    entries = self.kb_filtered_entries()
    if status_filter: entries = [e for e in entries if str(e.get('status','')).lower() == status_filter.lower()]
    entries = _fm469_sort(self, entries)

    def _sort(c):
        cur = getattr(self, 'kb_overview_sort_state', {}) or {}
        desc = not cur.get('descending', True) if cur.get('column') == c else (True if c == 'date' else False)
        self.kb_overview_sort_state = {'column': c, 'descending': desc}
        self.render_page()

    def _draw_chip(cv, x0, y0, txt, color, right, tag):
        txt = _fm469_txt(txt).strip()
        if not txt: return x0
        ww = max(42, min(145, 18 + len(txt)*7))
        if x0 + ww > right: return x0
        fg = _wz_cat_fg(color)
        cv.create_rectangle(x0, y0, x0+ww, y0+20, fill=color, outline=color, tags=(tag,))
        cv.create_text(x0+9, y0+10, text=_fm469_ell(txt, max(4, int((ww-14)/7))), anchor='w', fill=fg, font=body_font(8, weight='bold'), tags=(tag,))
        return x0 + ww + 6

    def _paint_selection():
        for eid, (rect, row_no) in row_rects.items():
            fill = '#EAF2FB' if eid == selected.get('id') else (WHITE if row_no % 2 == 0 else '#FAFAFA')
            try: body.itemconfigure(rect, fill=fill)
            except Exception: pass

    def _select(eid):
        selected['id'] = eid; self.kb_selected_entry_id = eid; _paint_selection()

    def _open(eid):
        if not eid: return
        if not self.kb_confirm_unsaved_before_switch(): return
        self.kb_selected_entry_id = eid; self.knowledge_view = 'detail'; self.knowledge_start_overlay = False; self.render_page()

    def _redraw(width=None):
        width = max(760, int(width or body.winfo_width() or (w-44)))
        starts, widths = _fm469_layout(width-18)
        head.delete('all'); body.delete('all'); row_rects.clear()
        head.create_rectangle(0,0,width,31, fill='#F1F5F9', outline='')
        for i,k in enumerate(keys):
            x0 = starts[i]
            marker = ' ↓' if state.get('column') == k and state.get('descending', True) else (' ↑' if state.get('column') == k else '')
            tag = 'fm469_sort_' + k
            head.create_text(x0+8, 16, text=labels[k]+marker, anchor='w', fill=TEXT, font=body_font(9, weight='bold'), tags=(tag,))
            head.tag_bind(tag, '<Button-1>', lambda e, c=k: _sort(c))
            if i > 0: head.create_line(x0,0,x0,31, fill='#D7DEE8')
        row_h = 34
        if not entries:
            body.create_text(8, 18, text='Noch keine Einträge vorhanden', anchor='w', fill=TEXT2, font=body_font(10))
            body.configure(scrollregion=(0,0,width,40)); return
        for r,e in enumerate(entries):
            eid = e.get('id') or self.kb_make_entry_id()
            tag = 'fm469_row_' + str(r); y0 = r*row_h
            fill = '#EAF2FB' if eid == selected.get('id') else (WHITE if r % 2 == 0 else '#FAFAFA')
            rect = body.create_rectangle(0,y0,width,y0+row_h, fill=fill, outline='#E5E7EB', tags=(tag,))
            row_rects[eid] = (rect, r)
            for ix in starts[1:]: body.create_line(ix, y0, ix, y0+row_h, fill='#EEF2F7', tags=(tag,))
            vals = [self.kb_display_date(e.get('updated_at')), e.get('title',''), '', _fm469_user(self, e.get('user','')), e.get('status','')]
            chars = [14, max(18, int(widths[1]/7)), 0, max(12, int(widths[3]/7)), max(8, int(widths[4]/7))]
            for i,val in enumerate(vals):
                if i == 2: continue
                body.create_text(starts[i]+8, y0+row_h/2, text=_fm469_ell(val, chars[i]), anchor='w', fill=TEXT, font=body_font(9, weight='bold' if i==1 else None), tags=(tag,))
            cx, right = starts[2]+8, starts[2]+widths[2]-8
            cats = [c for c in (e.get('categories',[]) or []) if _fm469_txt(c).strip()]
            if not cats:
                body.create_text(cx, y0+row_h/2, text='—', anchor='w', fill=TEXT2, font=body_font(9), tags=(tag,))
            else:
                shown = 0
                for cat in cats:
                    nx = _draw_chip(body, cx, y0+7, cat, _fm469_cat_color(self, cat), right, tag)
                    if nx == cx:
                        rest = len(cats)-shown
                        if rest and cx+42 <= right: _draw_chip(body, cx, y0+7, '+'+str(rest), '#E5E7EB', right, tag)
                        break
                    cx = nx; shown += 1
            body.tag_bind(tag, '<Button-1>', lambda ev, id=eid: _select(id))
            body.tag_bind(tag, '<Double-Button-1>', lambda ev, id=eid: _open(id))
        body.configure(scrollregion=(0,0,width,max(row_h*len(entries),row_h)))

    def _mark_current():
        eid = selected.get('id') or getattr(self, 'kb_selected_entry_id', None)
        if not eid:
            try: messagebox.showwarning('Wissenszentrale','Bitte zuerst einen Eintrag in der Übersicht auswählen.')
            except Exception: pass
            return
        try: _fm467_set_entry_current(self, eid)
        except Exception:
            try: self.kb_mark_entry_current(eid)
            except Exception: pass

    head.bind('<Configure>', lambda e: _redraw(e.width))
    body.bind('<Configure>', lambda e: _redraw(e.width))
    try: body.bind('<MouseWheel>', lambda e: (body.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
    except Exception: pass
    _redraw(max(760, int(w)-44))
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


def _fm469_go_main_direct(self):
    try: self.knowledge_start_overlay = False
    except Exception: pass
    try:
        self.current_page = 'main'; self.current_title = 'Hauptmenü'; self.breadcrumb = [('main','Hauptmenü')]; self.page_history = []
        self.render_page()
    except Exception:
        try: self.show_page('main', 'Hauptmenü', False)
        except Exception: pass


def _fm469_kb_handle_escape(self, event=None):
    if getattr(self, 'current_page', '') == 'knowledge_base':
        try:
            if not bool(getattr(self, 'knowledge_start_overlay', False)):
                self.knowledge_start_overlay = True; self.render_page(); return 'break'
        except Exception: pass
        _fm469_go_main_direct(self); return 'break'
    return None

try:
    FiBuMateApp.render_kb_list_area = _fm469_render_kb_list_area
    FiBuMateApp.kb_handle_escape = _fm469_kb_handle_escape
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Kategorie löschen + Auto-Speichern + Footer FINAL 2026-06-23
# Version 0.470
# Zweck:
# - Kategorien können ab Berechtigungsstufe 3 gelöscht werden; Sonderkategorie To-Do bleibt geschützt.
# - Wissenseinträge im Bearbeitungsmodus werden automatisch gespeichert, sowohl lokal als auch live über die bestehende Speicherlogik.
# - Benutzer werden konsequent mit Vor- und Nachnamen aus der Benutzerverwaltung angezeigt.
# - Startkachel "Neuer Eintrag": Öffnen-Button erhält die gleiche neutrale Farbgebung wie die anderen Öffnen-Buttons.
# - Fußzeile wird nochmals ca. 15 % niedriger; Text bleibt in Originalgröße und wird nur bei Überlauf reduziert.
# ------------------------------------------------------------------

_FM470_PATCH_VERSION = "0.470"
_FM470_AUTOSAVE_DELAY_MS = 3000
_FM470_AUTOSAVE_INTERVAL_MS = 30000


def _fm470_norm(v):
    return " ".join(str(v or "").strip().split())


def _fm470_user_fullname(self, raw_user):
    raw = _fm470_norm(raw_user)
    if not raw:
        return ""
    wanted = raw.casefold()
    try:
        data = getattr(self, 'user_data', {}) or {}
    except Exception:
        data = {}
    users = data.get('users', {}) if isinstance(data, dict) else {}

    def full_from_item(key, item):
        if not isinstance(item, dict):
            return ""
        first = _fm470_norm(item.get('first_name') or item.get('vorname') or item.get('first') or item.get('given_name'))
        last = _fm470_norm(item.get('last_name') or item.get('nachname') or item.get('surname') or item.get('display_name') or item.get('display') or item.get('name') or key)
        full = _fm470_norm(item.get('full_name') or item.get('fullname') or item.get('vollname'))
        if full and (' ' in full or ',' in full):
            return full
        if first and last:
            return (first + ' ' + last).strip()
        return full or last or first or _fm470_norm(key)

    if isinstance(users, dict):
        for key, item in users.items():
            keys = {_fm470_norm(key).casefold()}
            if isinstance(item, dict):
                for k in ('username','user','user_key','login','kuerzel','kürzel','id','display_name','full_name','fullname','name','email'):
                    val = _fm470_norm(item.get(k))
                    if val:
                        keys.add(val.casefold())
            if wanted in keys:
                return full_from_item(key, item)
    # Falls der Eintrag bereits wie ein Vollname aussieht, beibehalten.
    if ' ' in raw or ',' in raw:
        return raw
    return raw


# Globale Anzeigehelfer der letzten Wissenszentrale-Patches auf Vor-/Nachname umbiegen.
try:
    _fm467_user_fullname = _fm470_user_fullname
except Exception:
    pass
try:
    _fm468_user_fullname = _fm470_user_fullname
except Exception:
    pass


def _fm470_permission_level(self):
    try:
        return int(self.kb_current_permission_level())
    except Exception:
        pass
    try:
        rank = ROLE_RANK.get(self.my_role(), 1)
        return int(rank)
    except Exception:
        return 1


def _fm470_can_delete_category(self):
    return _fm470_permission_level(self) >= 3


def _fm470_delete_category_from_entries(self, cat_name):
    cat_name = _wz_cat_norm(cat_name)
    entries = self.kb_load_entries()
    changed = False
    for entry in entries:
        old = list(entry.get('categories', []) or [])
        new = [c for c in old if _wz_cat_norm(c).casefold() != cat_name.casefold()]
        if new != old:
            entry['categories'] = new
            changed = True
    return self.kb_save_entries(entries) if changed else True


def _fm470_render_kb_categories_area(self, x, y, w, h):
    if not self.kb_can_manage_categories():
        self.render_kb_list_area(x, y, w, h, title='Übersicht')
        return
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text='Kategorien verwalten', bg=WHITE, fg=BLUE, font=body_font(15, weight='bold')).pack(anchor='w', padx=18, pady=(14, 6))
    hint = 'Kategorien können erstellt, umbenannt und farblich gepflegt werden. Löschen ist ab Berechtigungsstufe 3 möglich; To-Do bleibt als Sonderkategorie geschützt.'
    tk.Label(frame, text=hint, bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(540, int(w*0.86)), justify='left').pack(anchor='w', padx=18, pady=(0, 10))
    main = tk.Frame(frame, bg=WHITE); main.pack(fill='both', expand=True, padx=18, pady=(0, 14))
    main.grid_columnconfigure(0, weight=1); main.grid_columnconfigure(1, weight=1); main.grid_rowconfigure(0, weight=1)
    left = tk.Frame(main, bg=WHITE); left.grid(row=0, column=0, sticky='nsew', padx=(0,12))
    right = tk.Frame(main, bg=WHITE); right.grid(row=0, column=1, sticky='nsew', padx=(12,0))
    tk.Label(left, text='Vorhandene Kategorien', bg=WHITE, fg=TEXT2, font=body_font(10, weight='bold')).pack(anchor='w')
    lb = tk.Listbox(left, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1, exportselection=False)
    lb.pack(fill='both', expand=True, pady=(6,0))
    name_var = tk.StringVar(value=''); color_var = tk.StringVar(value='#2563EB'); info_var = tk.StringVar(value='Keine Kategorie ausgewählt')
    tk.Label(right, text='Kategorie erstellen / bearbeiten', bg=WHITE, fg=TEXT2, font=body_font(10, weight='bold')).pack(anchor='w')
    tk.Label(right, textvariable=info_var, bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor='w', pady=(3,8))
    form = tk.Frame(right, bg=WHITE); form.pack(fill='x')
    tk.Label(form, text='Name', bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky='w')
    name_entry = tk.Entry(form, textvariable=name_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1)
    name_entry.grid(row=1, column=0, sticky='ew', ipady=5, pady=(0,8))
    tk.Label(form, text='Farbe (#RRGGBB)', bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=0, sticky='w')
    color_entry = tk.Entry(form, textvariable=color_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1)
    color_entry.grid(row=3, column=0, sticky='ew', ipady=5, pady=(0,8)); form.grid_columnconfigure(0, weight=1)
    preview = tk.Label(right, text='Farbvorschau', bg='#2563EB', fg='#FFFFFF', font=body_font(10, weight='bold'), relief='solid', bd=1)
    preview.pack(fill='x', ipady=8, pady=(0,10))
    swatches = tk.Frame(right, bg=WHITE); swatches.pack(anchor='w', pady=(0,10))
    for c in _WZ_CAT_PALETTE:
        tk.Button(swatches, text='', command=lambda cc=c: (color_var.set(cc), apply_preview()), bg=c, activebackground=c, relief='solid', bd=1, width=3).pack(side='left', padx=(0,4), ipady=6)

    def apply_preview(*_):
        c = _wz_cat_hex(color_var.get(), _wz_cat_default_color(name_var.get() or 'Kategorie'))
        preview.configure(bg=c, fg=_wz_cat_fg(c), text='Farbvorschau  ' + c)

    def refresh(selected=None):
        selected = _wz_cat_norm(selected or getattr(self, 'kb_selected_category_name', ''))
        lb.delete(0, 'end'); sel_idx = None
        for idx, cat in enumerate(self.kb_get_categories()):
            color = self.kb_get_category_color(cat)
            lb.insert('end', cat)
            try:
                if cat == selected:
                    lb.itemconfig(idx, bg=color, fg=_wz_cat_fg(color), selectbackground=color, selectforeground=_wz_cat_fg(color)); sel_idx = idx
                else:
                    lb.itemconfig(idx, bg=_wz442_light_color(color), fg=TEXT, selectbackground=color, selectforeground=_wz_cat_fg(color))
            except Exception:
                pass
        if sel_idx is not None:
            lb.selection_set(sel_idx); lb.see(sel_idx)

    def select_cat(event=None):
        sel = lb.curselection()
        if not sel: return
        cat = lb.get(sel[0]); self.kb_selected_category_name = cat
        name_var.set(cat); color_var.set(self.kb_get_category_color(cat)); info_var.set('Ausgewählt: ' + cat); apply_preview(); refresh(cat)
    lb.bind('<<ListboxSelect>>', select_cat)

    def create_cat():
        name = _wz_cat_norm(name_var.get())
        if not name:
            messagebox.showwarning('Wissenszentrale', 'Bitte einen Kategorienamen eingeben.'); return
        if name.lower() in {c.lower() for c in self.kb_get_categories()}:
            messagebox.showwarning('Wissenszentrale', 'Diese Kategorie existiert bereits. Bitte Änderungen speichern verwenden.'); return
        meta = self.kb_load_category_meta(); meta[name] = {'name': name, 'color': _wz_cat_hex(color_var.get(), _wz_cat_default_color(name))}
        if self.kb_save_category_meta(meta):
            self.kb_selected_category_name = name; info_var.set('Neu erstellt: ' + name); refresh(name); self.render_page()

    def save_cat():
        old = _wz_cat_norm(getattr(self, 'kb_selected_category_name', '')); new = _wz_cat_norm(name_var.get())
        if not new:
            messagebox.showwarning('Wissenszentrale', 'Bitte einen Kategorienamen eingeben.'); return
        meta = self.kb_load_category_meta()
        if old and old != new:
            if old.lower() == 'to-do':
                messagebox.showwarning('Wissenszentrale', "Die Sonderkategorie 'To-Do' darf nicht umbenannt werden. Die Farbe kann gespeichert werden."); new = old; name_var.set(old)
            elif new.lower() in {c.lower() for c in self.kb_get_categories() if c.lower() != old.lower()}:
                messagebox.showwarning('Wissenszentrale', 'Der neue Kategoriename existiert bereits.'); return
            else:
                for k in list(meta.keys()):
                    if k.lower() == old.lower(): meta.pop(k, None)
                if not _wz_cat_rename_entries(self, old, new): return
        meta[new] = {'name': new, 'color': _wz_cat_hex(color_var.get(), _wz_cat_default_color(new))}
        if self.kb_save_category_meta(meta):
            self.kb_selected_category_name = new; info_var.set('Gespeichert: ' + new); refresh(new); self.render_page()

    def delete_cat():
        if not _fm470_can_delete_category(self):
            messagebox.showwarning('Wissenszentrale', 'Kategorien können erst ab Berechtigungsstufe 3 gelöscht werden.'); return
        cat = _wz_cat_norm(getattr(self, 'kb_selected_category_name', '') or name_var.get())
        if not cat:
            messagebox.showwarning('Wissenszentrale', 'Bitte zuerst eine Kategorie auswählen.'); return
        if cat.lower() == 'to-do':
            messagebox.showwarning('Wissenszentrale', "Die Sonderkategorie 'To-Do' darf nicht gelöscht werden."); return
        if not messagebox.askyesno('Wissenszentrale', f"Kategorie '{cat}' wirklich löschen?\n\nDie Kategorie wird aus allen Einträgen entfernt."):
            return
        meta = self.kb_load_category_meta()
        for k in list(meta.keys()):
            if k.lower() == cat.lower(): meta.pop(k, None)
        if not _fm470_delete_category_from_entries(self, cat): return
        if self.kb_save_category_meta(meta):
            self.kb_selected_category_name = ''; name_var.set(''); color_var.set('#2563EB'); info_var.set('Kategorie gelöscht: ' + cat); apply_preview(); refresh(); self.render_page()

    def blank_cat():
        self.kb_selected_category_name = ''; name_var.set(''); color_var.set('#2563EB'); info_var.set('Neue Kategorie erfassen'); apply_preview(); refresh(); name_entry.focus_set()

    color_var.trace_add('write', apply_preview)
    btns = tk.Frame(right, bg=WHITE); btns.pack(fill='x', pady=(4,0))
    tk.Button(btns, text='Kategorie erstellen', command=create_cat, bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=10, ipady=4)
    tk.Button(btns, text='Änderungen speichern', command=save_cat, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=10, ipady=4)
    del_bg = '#FEE2E2' if _fm470_can_delete_category(self) else '#E5E7EB'
    tk.Button(btns, text='Kategorie löschen', command=delete_cat, bg=del_bg, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=10, ipady=4)
    tk.Button(btns, text='Neu leeren', command=blank_cat, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', ipadx=10, ipady=4)
    apply_preview(); refresh(getattr(self, 'kb_selected_category_name', ''))
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


def _fm470_autosave_now(self):
    try:
        if getattr(self, 'current_page', '') != 'knowledge_base' or getattr(self, 'knowledge_view', '') != 'new':
            return False
        self.kb_ensure_state_vars()
        title = _fm470_norm(getattr(self, 'kb_title_var', tk.StringVar(value='')).get())
        try:
            text_value = self.kb_text_widget.get('1.0', 'end-1c')
            formatting = _wz439_capture_formatting(self.kb_text_widget) if '_wz439_capture_formatting' in globals() else []
        except Exception:
            text_value = getattr(self, 'kb_text_initial', '')
            formatting = getattr(self, 'kb_text_formatting_initial', []) or []
        categories = []
        for var in getattr(self, 'kb_entry_category_vars', []) or []:
            v = _fm470_norm(var.get())
            if v and v not in categories: categories.append(v)
        user = _fm470_norm(getattr(self, 'kb_user_var', tk.StringVar(value='')).get())
        status = _fm470_norm(getattr(self, 'kb_status_var', tk.StringVar(value='Aktiv')).get()) or 'Entwurf'
        rhythm = _fm470_norm(getattr(self, 'kb_rhythm_var', tk.StringVar(value='')).get()) if any(c.lower() == 'to-do' for c in categories) else ''
        if not title and not text_value.strip() and not categories:
            return False
        if not title:
            title = 'Automatisch gespeicherter Entwurf'
            try: self.kb_title_var.set(title)
            except Exception: pass
        entries = self.kb_load_entries(); now = self.kb_now() if hasattr(self, 'kb_now') else datetime.now().isoformat(timespec='seconds')
        selected_id = getattr(self, 'kb_edit_entry_id', None) or getattr(self, 'kb_autosave_entry_id', None)
        if not selected_id:
            selected_id = self.kb_make_entry_id() if hasattr(self, 'kb_make_entry_id') else hashlib.sha1((title+now).encode('utf-8')).hexdigest()[:16]
            self.kb_autosave_entry_id = selected_id
            self.kb_edit_entry_id = selected_id
        found = False
        for entry in entries:
            if entry.get('id') == selected_id:
                entry.update({'title': title, 'categories': categories[:4], 'user': user, 'status': status, 'rhythm': rhythm, 'text': text_value, 'text_formatting': formatting, 'updated_at': now, 'autosaved_at': now})
                found = True; break
        if not found:
            entries.append({'id': selected_id, 'title': title, 'categories': categories[:4], 'user': user, 'status': status, 'rhythm': rhythm, 'text': text_value, 'text_formatting': formatting, 'inline_images': getattr(self, 'kb_inline_images', []) or [], 'created_at': now, 'updated_at': now, 'autosaved_at': now, 'comments': [], 'attachments': []})
        ok = self.kb_save_entries(entries)
        if ok:
            self.kb_selected_entry_id = selected_id; self.knowledge_unsaved = False; self.kb_autosave_last_at = now
        return bool(ok)
    except Exception:
        return False


def _fm470_schedule_autosave(self, delay=None):
    try:
        if getattr(self, 'current_page', '') != 'knowledge_base' or getattr(self, 'knowledge_view', '') != 'new': return
        delay = _FM470_AUTOSAVE_DELAY_MS if delay is None else int(delay)
        token = object(); self.kb_autosave_token = token
        def run():
            if getattr(self, 'kb_autosave_token', None) is token:
                _fm470_autosave_now(self)
                _fm470_schedule_autosave(self, _FM470_AUTOSAVE_INTERVAL_MS)
        self.root.after(delay, run)
    except Exception:
        pass


def _fm470_install_autosave(self):
    try:
        if getattr(self, 'knowledge_view', '') != 'new': return
        def mark(event=None):
            try: self.kb_mark_unsaved()
            except Exception: pass
            _fm470_schedule_autosave(self, _FM470_AUTOSAVE_DELAY_MS)
        tw = getattr(self, 'kb_text_widget', None)
        if tw:
            tw.bind('<KeyRelease>', mark, add='+'); tw.bind('<<Paste>>', mark, add='+'); tw.bind('<<Cut>>', mark, add='+')
        for var in [getattr(self, 'kb_title_var', None), getattr(self, 'kb_user_var', None), getattr(self, 'kb_status_var', None), getattr(self, 'kb_rhythm_var', None)] + list(getattr(self, 'kb_entry_category_vars', []) or []):
            try: var.trace_add('write', lambda *_: _fm470_schedule_autosave(self, _FM470_AUTOSAVE_DELAY_MS))
            except Exception: pass
        _fm470_schedule_autosave(self, _FM470_AUTOSAVE_INTERVAL_MS)
    except Exception:
        pass

try:
    _fm470_old_render_new_entry_area = FiBuMateApp.render_kb_new_entry_area
except Exception:
    _fm470_old_render_new_entry_area = None

def _fm470_render_kb_new_entry_area(self, x, y, w, h):
    if _fm470_old_render_new_entry_area:
        result = _fm470_old_render_new_entry_area(self, x, y, w, h)
    else:
        result = None
    _fm470_install_autosave(self)
    return result

try:
    _fm470_old_render_start_overlay = FiBuMateApp.render_knowledge_start_overlay
except Exception:
    _fm470_old_render_start_overlay = None

def _fm470_render_knowledge_start_overlay(self):
    if not _fm470_old_render_start_overlay:
        return None
    original = self.draw_kb_button
    def neutral_open_button(x, y, text, command=None, accent=False, width=None):
        if text == 'Öffnen':
            accent = False
        return original(x, y, text, command, accent, width)
    self.draw_kb_button = neutral_open_button
    try:
        return _fm470_old_render_start_overlay(self)
    finally:
        self.draw_kb_button = original


def _fm470_footer_font(size=9):
    return ('Segoe UI', int(size))


def _fm470_fit_label(label, max_width, base_size=9, min_size=7):
    try:
        label.update_idletasks()
        size = base_size
        while size > min_size and label.winfo_reqwidth() > max_width:
            size -= 1
            label.configure(font=_fm470_footer_font(size))
            label.update_idletasks()
    except Exception:
        pass


def _fm470_create_footer(self):
    self.footer = tk.Frame(self.root, bg='black', height=ui_s(42)); self.footer._zoom_exclude = True
    self.footer.pack(side='bottom', fill='x'); self.footer.pack_propagate(False)
    self.user_label = tk.Label(self.footer, bg='black', fg='white', font=_fm470_footer_font(9))
    self.user_label.place(relx=0, rely=0.5, anchor='w', x=10)
    self.version_label = tk.Label(self.footer, text=self.version_label_text(), bg='black', fg='white', font=_fm470_footer_font(9))
    self.version_label.place(relx=0.5, rely=0.5, anchor='center')
    self.clock_label = tk.Label(self.footer, bg='black', fg='white', font=_fm470_footer_font(9))
    self.clock_label.place(relx=1, rely=0.5, anchor='e', x=-10)
    self.update_clock()


def _fm470_update_clock(self):
    try:
        self.clock_label.config(text=datetime.now().strftime('%H:%M:%S'), font=_fm470_footer_font(9))
        display = _fm470_user_fullname(self, getattr(self, 'current_user_key', '') or getattr(self, 'current_user_display', ''))
        self.user_label.config(text=(f'Benutzer {display}' if display else ''), font=_fm470_footer_font(9))
        self.version_label.config(text=self.version_label_text(), font=_fm470_footer_font(9))
        try:
            w = max(600, self.root.winfo_width())
            _fm470_fit_label(self.user_label, int(w*0.36), 9, 7)
            _fm470_fit_label(self.version_label, int(w*0.28), 9, 7)
            _fm470_fit_label(self.clock_label, int(w*0.16), 9, 7)
        except Exception:
            pass
        self.root.after(1000, self.update_clock)
    except Exception:
        pass

try:
    FiBuMateApp.render_kb_categories_area = _fm470_render_kb_categories_area
    FiBuMateApp.render_kb_new_entry_area = _fm470_render_kb_new_entry_area
    FiBuMateApp.render_knowledge_start_overlay = _fm470_render_knowledge_start_overlay
    FiBuMateApp.create_footer = _fm470_create_footer
    FiBuMateApp.update_clock = _fm470_update_clock
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale Eintrag löschen + Hauptmenü Kachelgrößen FINAL 2026-06-23
# Version 0.471
# Zweck:
# - Wissenseinträge können ab E3 gelöscht werden; darunter ist der Button sichtbar, aber deaktiviert/grau.
# - Löschen-Button steht links neben "Eintrag als aktuell kennzeichnen" und nutzt das etablierte Papierkorb-Symbol.
# - Hauptmenü: Zeile 1 +5 %, Zeile 2 -5 %; neue Kachel "Dateiausgabe" in Zeile 2 als Platzhalter.
# ------------------------------------------------------------------

_FM471_PATCH_VERSION = "0.471"
_FM471_DELETE_SYMBOL = "🗑"


def _fm471_permission_rank(self):
    try:
        return int(self.kb_current_permission_level())
    except Exception:
        pass
    try:
        return int(ROLE_RANK.get(self.my_role(), 1))
    except Exception:
        return 1


def _fm471_can_delete_entry(self):
    return _fm471_permission_rank(self) >= 3


def _fm471_entry_title(entry):
    try:
        return str((entry or {}).get('title') or 'Wissenseintrag')
    except Exception:
        return 'Wissenseintrag'


def _fm471_delete_entry(self, entry_id=None):
    entry_id = entry_id or getattr(self, 'kb_selected_entry_id', None)
    if not entry_id:
        try: messagebox.showwarning('Wissenszentrale', 'Bitte zuerst einen Eintrag auswählen.')
        except Exception: pass
        return False
    if not _fm471_can_delete_entry(self):
        return False
    entries = self.kb_load_entries()
    target = None
    remaining = []
    for e in entries:
        if e.get('id') == entry_id:
            target = e
        else:
            remaining.append(e)
    if not target:
        try: messagebox.showwarning('Wissenszentrale', 'Der ausgewählte Eintrag wurde nicht gefunden.')
        except Exception: pass
        return False
    try:
        if not messagebox.askyesno('Eintrag löschen', f"Eintrag wirklich löschen?\n\n{_fm471_entry_title(target)}"):
            return False
    except Exception:
        return False
    if self.kb_save_entries(remaining):
        try:
            if hasattr(self, 'kb_autosave_entry_id') and self.kb_autosave_entry_id == entry_id:
                self.kb_autosave_entry_id = None
        except Exception:
            pass
        self.kb_selected_entry_id = None
        self.kb_edit_entry_id = None
        self.knowledge_unsaved = False
        self.knowledge_view = 'all'
        self.knowledge_start_overlay = False
        self.render_page()
        return True
    return False


def _fm471_delete_button(parent, self, entry_id=None):
    can_delete = _fm471_can_delete_entry(self)
    return tk.Button(parent, text=f'{_FM471_DELETE_SYMBOL} Löschen', command=(lambda eid=entry_id: _fm471_delete_entry(self, eid)) if can_delete else (lambda: None),
                     bg=('#FEE2E2' if can_delete else '#E5E7EB'), fg=(TEXT if can_delete else '#7A7F87'),
                     activebackground=('#FECACA' if can_delete else '#E5E7EB'), font=body_font(10), relief='solid', bd=1,
                     state=('normal' if can_delete else 'disabled'), cursor=('hand2' if can_delete else 'arrow'))


# --- Übersicht inkl. Löschbutton links neben "Eintrag als aktuell kennzeichnen" ---
def _fm471_render_kb_list_area(self, x, y, w, h, title='Gesamtliste aller Einträge', status_filter=None):
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    top = tk.Frame(frame, bg=WHITE); top.pack(fill='x', padx=16, pady=(12, 6))
    tk.Label(top, text=title, bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(side='left')
    selected = {'id': getattr(self, 'kb_selected_entry_id', None)}
    row_rects = {}

    def _mark_current():
        eid = selected.get('id') or getattr(self, 'kb_selected_entry_id', None)
        if not eid:
            try: messagebox.showwarning('Wissenszentrale', 'Bitte zuerst einen Eintrag in der Übersicht auswählen.')
            except Exception: pass
            return
        try: _fm467_set_entry_current(self, eid)
        except Exception:
            try: self.kb_mark_entry_current(eid)
            except Exception: pass

    tk.Button(top, text='Eintrag als aktuell kennzeichnen', command=_mark_current, bg='#CFEAD6', fg=TEXT, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='right', padx=(8,0), ipadx=8, ipady=3)
    _fm471_delete_button(top, self, None).pack(side='right', padx=(8,0), ipadx=8, ipady=3)
    state = getattr(self, 'kb_overview_sort_state', {'column':'date','descending':True}) or {}
    tk.Label(top, text=f"Sortierung: {state.get('column','date')} ({'absteigend' if state.get('descending', True) else 'aufsteigend'})", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')

    table = tk.Frame(frame, bg=WHITE); table.pack(fill='both', expand=True, padx=16, pady=(0,16))
    table.grid_columnconfigure(0, weight=1); table.grid_rowconfigure(1, weight=1)
    head = tk.Canvas(table, height=31, bg='#F1F5F9', highlightthickness=0); head.grid(row=0, column=0, sticky='ew')
    body = tk.Canvas(table, bg=WHITE, highlightthickness=0); body.grid(row=1, column=0, sticky='nsew')
    sb = ttk.Scrollbar(table, orient='vertical', command=body.yview); sb.grid(row=1, column=1, sticky='ns')
    body.configure(yscrollcommand=sb.set)
    keys = ['date','title','categories','user','status']
    labels = {'date':'Aktualisiert','title':'Titel','categories':'Kategorien','user':'Benutzer','status':'Status'}
    entries = self.kb_filtered_entries()
    if status_filter: entries = [e for e in entries if str(e.get('status','')).lower() == status_filter.lower()]
    try: entries = _fm469_sort(self, entries)
    except Exception: entries = _fm469_sort_entries_for_overview(self, entries)

    def _sort(c):
        cur = getattr(self, 'kb_overview_sort_state', {}) or {}
        desc = not cur.get('descending', True) if cur.get('column') == c else (True if c == 'date' else False)
        self.kb_overview_sort_state = {'column': c, 'descending': desc}
        self.render_page()

    def _draw_chip(cv, x0, y0, txt, color, right, tag):
        txt = str(txt or '').strip()
        if not txt: return x0
        ww = max(42, min(145, 18 + len(txt)*7))
        if x0 + ww > right: return x0
        fg = _wz_cat_fg(color)
        ell = txt if len(txt) <= max(4, int((ww-14)/7)) else txt[:max(4, int((ww-14)/7))-1] + '…'
        cv.create_rectangle(x0, y0, x0+ww, y0+20, fill=color, outline=color, tags=(tag,))
        cv.create_text(x0+9, y0+10, text=ell, anchor='w', fill=fg, font=body_font(8, weight='bold'), tags=(tag,))
        return x0 + ww + 6

    def _layout(total_width):
        try: return _fm469_layout(total_width)
        except Exception: return _fm469_column_layout(total_width)

    def _ell(v, n):
        try: return _fm469_ell(v, n)
        except Exception:
            s=str(v or '').replace('\n',' ').strip(); return s if len(s)<=n else s[:max(1,n-1)]+'…'

    def _user(v):
        try: return _fm470_user_fullname(self, v)
        except Exception: return str(v or '')

    def _cat_color(v):
        try: return _fm469_cat_color(self, v)
        except Exception:
            try: return _fm470_cat_color(self, v)
            except Exception: return self.kb_get_category_color(v)

    def _paint_selection():
        for eid, (rect, row_no) in row_rects.items():
            fill = '#EAF2FB' if eid == selected.get('id') else (WHITE if row_no % 2 == 0 else '#FAFAFA')
            try: body.itemconfigure(rect, fill=fill)
            except Exception: pass

    def _select(eid):
        selected['id'] = eid; self.kb_selected_entry_id = eid; _paint_selection()

    def _open(eid):
        if not eid: return
        if not self.kb_confirm_unsaved_before_switch(): return
        self.kb_selected_entry_id = eid; self.knowledge_view = 'detail'; self.knowledge_start_overlay = False; self.render_page()

    def _redraw(width=None):
        width = max(760, int(width or body.winfo_width() or (w-44)))
        starts, widths = _layout(width-18)
        head.delete('all'); body.delete('all'); row_rects.clear()
        head.create_rectangle(0,0,width,31, fill='#F1F5F9', outline='')
        for i,k in enumerate(keys):
            x0 = starts[i]
            marker = ' ↓' if state.get('column') == k and state.get('descending', True) else (' ↑' if state.get('column') == k else '')
            tag = 'fm471_sort_' + k
            head.create_text(x0+8, 16, text=labels[k]+marker, anchor='w', fill=TEXT, font=body_font(9, weight='bold'), tags=(tag,))
            head.tag_bind(tag, '<Button-1>', lambda e, c=k: _sort(c))
            if i > 0: head.create_line(x0,0,x0,31, fill='#D7DEE8')
        row_h = 34
        if not entries:
            body.create_text(8, 18, text='Noch keine Einträge vorhanden', anchor='w', fill=TEXT2, font=body_font(10))
            body.configure(scrollregion=(0,0,width,40)); return
        for r,e in enumerate(entries):
            eid = e.get('id') or self.kb_make_entry_id()
            tag = 'fm471_row_' + str(r); y0 = r*row_h
            fill = '#EAF2FB' if eid == selected.get('id') else (WHITE if r % 2 == 0 else '#FAFAFA')
            rect = body.create_rectangle(0,y0,width,y0+row_h, fill=fill, outline='#E5E7EB', tags=(tag,))
            row_rects[eid] = (rect, r)
            for ix in starts[1:]: body.create_line(ix, y0, ix, y0+row_h, fill='#EEF2F7', tags=(tag,))
            vals = [self.kb_display_date(e.get('updated_at')), e.get('title',''), '', _user(e.get('user','')), e.get('status','')]
            chars = [14, max(18, int(widths[1]/7)), 0, max(12, int(widths[3]/7)), max(8, int(widths[4]/7))]
            for i,val in enumerate(vals):
                if i == 2: continue
                body.create_text(starts[i]+8, y0+row_h/2, text=_ell(val, chars[i]), anchor='w', fill=TEXT, font=body_font(9, weight='bold' if i==1 else None), tags=(tag,))
            cx, right = starts[2]+8, starts[2]+widths[2]-8
            cats = [c for c in (e.get('categories',[]) or []) if str(c or '').strip()]
            if not cats:
                body.create_text(cx, y0+row_h/2, text='—', anchor='w', fill=TEXT2, font=body_font(9), tags=(tag,))
            else:
                shown = 0
                for cat in cats:
                    nx = _draw_chip(body, cx, y0+7, cat, _cat_color(cat), right, tag)
                    if nx == cx:
                        rest = len(cats)-shown
                        if rest and cx+42 <= right: _draw_chip(body, cx, y0+7, '+'+str(rest), '#E5E7EB', right, tag)
                        break
                    cx = nx; shown += 1
            body.tag_bind(tag, '<Button-1>', lambda ev, id=eid: _select(id))
            body.tag_bind(tag, '<Double-Button-1>', lambda ev, id=eid: _open(id))
        body.configure(scrollregion=(0,0,width,max(row_h*len(entries),row_h)))

    head.bind('<Configure>', lambda e: _redraw(e.width))
    body.bind('<Configure>', lambda e: _redraw(e.width))
    try: body.bind('<MouseWheel>', lambda e: (body.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
    except Exception: pass
    _redraw(max(760, int(w)-44))
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


# --- Detailansicht inkl. Löschbutton ---
def _fm471_render_kb_detail_area(self, x, y, w, h):
    entry = self.kb_get_entry(getattr(self, 'kb_selected_entry_id', None))
    if not entry:
        self.render_kb_list_area(x, y, w, h, title='Übersicht')
        return
    frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
    self.widget_items.append(frame)
    tk.Label(frame, text=entry.get('title', ''), bg=WHITE, fg=BLUE, font=body_font(16, weight='bold')).pack(anchor='w', padx=18, pady=(14, 5))
    try: user_display = _fm470_user_fullname(self, entry.get('user',''))
    except Exception: user_display = entry.get('user','')
    tk.Label(frame, text=f"Aktualisiert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {user_display}    Status: {entry.get('status','')}", bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w', padx=18)
    try: _wz439_badges(frame, self, entry.get('categories', []) or [])
    except Exception:
        tk.Label(frame, text='Kategorien: ' + (', '.join(entry.get('categories', []) or []) or 'Keine Kategorien'), bg=WHITE, fg=TEXT, font=body_font(10, weight='bold')).pack(anchor='w', padx=18, pady=(8,4))
    try: _wz439_render_image_view(frame, self, entry.get('inline_images', []) or [])
    except Exception: pass
    text_frame = tk.Frame(frame, bg=WHITE); text_frame.pack(fill='both', expand=True, padx=18, pady=(6, 8))
    txt = tk.Text(text_frame, bg='#F8FAFC', fg=TEXT, font=body_font(10), relief='solid', bd=1, wrap='word', height=8)
    yscroll = tk.Scrollbar(text_frame, orient='vertical', command=txt.yview); txt.configure(yscrollcommand=yscroll.set)
    txt.pack(side='left', fill='both', expand=True); yscroll.pack(side='right', fill='y')
    txt.insert('1.0', entry.get('text', ''))
    try: _wz439_apply_formatting(txt, entry.get('text_formatting', []) or [])
    except Exception: pass
    txt.configure(state='disabled')
    try: txt.bind('<MouseWheel>', lambda ev: _wz439_text_mousewheel(txt, ev))
    except Exception: pass
    lower = tk.Frame(frame, bg=WHITE); lower.pack(fill='x', padx=18, pady=(0, 8))
    left = tk.Frame(lower, bg=WHITE); left.pack(side='left', fill='both', expand=True)
    right = tk.Frame(lower, bg=WHITE); right.pack(side='right', fill='both', expand=True, padx=(12,0))
    tk.Label(left, text='Anhänge', bg=WHITE, fg=BLUE, font=body_font(10, weight='bold')).pack(anchor='w')
    for a in entry.get('attachments', []) or []:
        tk.Label(left, text='• ' + a.get('name',''), bg=WHITE, fg=TEXT, font=body_font(9)).pack(anchor='w')
    tk.Label(right, text='Kommentare', bg=WHITE, fg=BLUE, font=body_font(10, weight='bold')).pack(anchor='w')
    for c in (entry.get('comments', []) or [])[-3:]:
        tk.Label(right, text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}", bg=WHITE, fg=TEXT, font=body_font(9), wraplength=380, justify='left').pack(anchor='w')
    comment = tk.Text(right, height=2, bg='#F8FAFC', fg=TEXT, font=body_font(9), relief='solid', bd=1, wrap='word'); comment.pack(fill='x', pady=(4,0))
    buttons = tk.Frame(frame, bg=WHITE); buttons.pack(fill='x', padx=18, pady=(0,14))
    if self.kb_can_create_or_edit():
        tk.Button(buttons, text='Bearbeiten', command=self.kb_edit_selected_entry, bg=WHITE, fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=16, ipady=4)
    tk.Button(buttons, text='Kommentar speichern', command=lambda: self.kb_add_comment_to_selected(comment), bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=16, ipady=4)
    tk.Button(buttons, text='Word-Export', command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=16, ipady=4)
    tk.Button(buttons, text='Zur Übersicht', command=lambda: self.kb_switch_view_from_start('all'), bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0,8), ipadx=16, ipady=4)
    _fm471_delete_button(buttons, self, entry.get('id')).pack(side='right', padx=(8,0), ipadx=12, ipady=4)
    tk.Button(buttons, text='Eintrag als aktuell kennzeichnen', command=lambda eid=entry.get('id'): _fm467_set_entry_current(self, eid), bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='right', ipadx=12, ipady=4)
    self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor='nw', width=ui_s(w), height=ui_s(h))


# --- Hauptmenü: Zeile 1 +5 %, Zeile 2 -5 %, neue Kachel Dateiausgabe ---
def _fm471_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w * 0.165))); th = max(125, min(170, int(h * 0.155)))
    gap_x = max(38, int(w * 0.04))
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count * tile_w + (count - 1) * gap
        start_x = (w - total_width) / 2 + tile_w / 2
        return [start_x + i * (tile_w + gap) for i in range(count)]
    def create_tile_group(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')
    top_y = y_pct(h, 64); tools_y = y_pct(h, 43)
    create_tile_group(top_tiles, top_y, 'main_menuzeile_1', max(250, int(tw*1.05)), max(130, int(th*1.05)))
    create_tile_group(middle_tiles, tools_y, 'main_menuzeile_2', max(200, int(tw*0.86*0.95)), max(100, int(th*0.86*0.95)))
    self.draw_continuous_relief_line(self.canvas, (top_y + tools_y) / 2, x_pct(w, 9.5), x_pct(w, 92))
    mini_tw = int(tw * 0.72); mini_th = int(th * 0.72)
    mini_center_y = y_pct(h, 16.5) + th / 2 - mini_th / 2
    mini_top = mini_center_y - mini_th / 2
    create_tile_group(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top - 13, x_pct(w, 9.5), x_pct(w, 92))
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_kb_list_area = _fm471_render_kb_list_area
    FiBuMateApp.render_kb_detail_area = _fm471_render_kb_detail_area
    FiBuMateApp.render_main_menu = _fm471_render_main_menu
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Modul Dateiausgabe Übersicht + Vorschau FINAL 2026-06-23
# Version 0.472
# Zweck:
# - Integriert das neue Hauptmenü-Modul "Dateiausgabe" direkt in die Hauptdatei.
# - Zeigt zuletzt mit FiBu Mate erstellte Dateien mit Name, Benutzer, Datum und Kurzpfad.
# - Unterstützt Sammelpositionen mit aufklappbaren Unterpositionen.
# - Bietet Datei öffnen, Speicherort öffnen, Outlook-Entwurf und Teams-Weitergabe.
# - Zweigeteilte Ansicht: links Übersicht, rechts Vorschau der gewählten Position.
# ------------------------------------------------------------------

_FM472_PATCH_VERSION = "0.472"
_FM472_OUTPUT_EXTENSIONS = {'.pdf', '.xlsx', '.xls', '.docx', '.doc', '.csv', '.txt', '.png', '.jpg', '.jpeg', '.bmp', '.gif'}


def _fm472_now_iso():
    try:
        return datetime.now().isoformat(timespec='seconds')
    except Exception:
        return ''


def _fm472_display_dt(ts):
    try:
        if isinstance(ts, (int, float)):
            return datetime.fromtimestamp(ts).strftime('%d.%m.%Y %H:%M')
        s = str(ts or '')
        if not s:
            return ''
        return datetime.fromisoformat(s.replace('Z', '+00:00')).strftime('%d.%m.%Y %H:%M')
    except Exception:
        return str(ts or '')[:16]


def _fm472_safe_user(self, value=None):
    try:
        if value:
            return _fm470_user_fullname(self, value)
    except Exception:
        pass
    try:
        return _fm470_user_fullname(self, getattr(self, 'current_user_key', '') or getattr(self, 'current_user_display', ''))
    except Exception:
        return str(value or getattr(self, 'current_user_display', '') or getattr(self, 'current_user_key', '') or '')


def _fm472_output_roots():
    roots = []
    candidates = []
    try: candidates.append(os.path.join(NETWORK_ROOT, 'Dateiausgabe'))
    except Exception: pass
    try: candidates.append(os.path.join(SCRIPT_DIR, 'Dateiausgabe'))
    except Exception: pass
    try: candidates.append(os.path.join(os.getcwd(), 'Dateiausgabe'))
    except Exception: pass
    candidates.append(r'C:\python\Dateiausgabe')
    candidates.append(r'C:\python\Dateiausgabe\FiBu Mate')
    seen = set()
    for p in candidates:
        try:
            ap = os.path.abspath(p)
            if ap not in seen:
                seen.add(ap); roots.append(ap)
        except Exception:
            pass
    return roots


def _fm472_primary_output_root():
    for p in _fm472_output_roots():
        try:
            os.makedirs(p, exist_ok=True)
            return p
        except Exception:
            pass
    return os.getcwd()


def _fm472_log_paths():
    paths = []
    for root in _fm472_output_roots():
        try:
            paths.append(os.path.join(root, 'fibu_mate_dateiausgabe_log.json'))
        except Exception:
            pass
    return paths


def _fm472_read_json(path):
    try:
        if path and os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _fm472_write_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def _fm472_short_path(path, max_len=58):
    s = str(path or '')
    if len(s) <= max_len:
        return s
    parts = s.replace('/', os.sep).split(os.sep)
    if len(parts) >= 3:
        compact = parts[0] + os.sep + '…' + os.sep + os.sep.join(parts[-2:])
        if len(compact) <= max_len:
            return compact
    return '…' + s[-max_len+1:]


def _fm472_file_key(path):
    try:
        return hashlib.sha1(os.path.abspath(path).casefold().encode('utf-8', errors='ignore')).hexdigest()[:18]
    except Exception:
        return hashlib.sha1(str(path).encode('utf-8', errors='ignore')).hexdigest()[:18]


def _fm472_parent_dir_under_roots(path):
    try:
        ap = os.path.abspath(path)
        for root in _fm472_output_roots():
            try:
                rp = os.path.relpath(ap, root)
                if not rp.startswith('..'):
                    parts = rp.split(os.sep)
                    if len(parts) > 1:
                        return os.path.join(root, parts[0])
            except Exception:
                pass
    except Exception:
        pass
    return os.path.dirname(path)


def _fm472_guess_module(path):
    s = str(path or '').lower()
    if 'debitor' in s and 'serien' in s:
        return 'Debitoren-Serienbrief'
    if 'aramark' in s:
        return 'Aramark Monatsabrechnungen'
    if 'nike' in s:
        return 'Nike-Tools'
    if 'afi' in s or 'upload' in s:
        return 'AFI-Upload'
    return 'FiBu Mate'


def _fm472_register_output_file(self, path, user=None, module=None, group_id=None, group_title=None, created_at=None, children=None):
    """Zentrale Registrierfunktion für künftige Exportmodule.
    Einzeldatei: file_output_register(path)
    Sammelexport: gleicher group_id/group_title für mehrere Dateien.
    """
    try:
        if not path:
            return False
        entry = {
            'id': _fm472_file_key(path),
            'type': 'file',
            'name': os.path.basename(str(path)),
            'path': os.path.abspath(str(path)),
            'user': user or _fm472_safe_user(self),
            'module': module or _fm472_guess_module(path),
            'group_id': group_id or '',
            'group_title': group_title or '',
            'created_at': created_at or _fm472_now_iso(),
        }
        data = {}
        target = _fm472_log_paths()[0]
        for p in _fm472_log_paths():
            d = _fm472_read_json(p)
            if isinstance(d, dict) and d.get('items'):
                data = d; break
        items = data.get('items', []) if isinstance(data, dict) else []
        items = [x for x in items if isinstance(x, dict) and x.get('id') != entry['id']]
        items.append(entry)
        data = {'items': items[-1000:]}
        _fm472_write_json(target, data)
        return True
    except Exception:
        return False


def _fm472_load_registered_items(self):
    by_id = {}
    for p in _fm472_log_paths():
        d = _fm472_read_json(p)
        for item in (d.get('items', []) if isinstance(d, dict) else []):
            if isinstance(item, dict) and item.get('path'):
                by_id[item.get('id') or _fm472_file_key(item.get('path'))] = dict(item)
    return list(by_id.values())


def _fm472_scan_output_files(self):
    items = []
    seen = set()
    # 1) Registrierte Dateien bevorzugt übernehmen.
    for item in _fm472_load_registered_items(self):
        p = item.get('path')
        if not p:
            continue
        try:
            ap = os.path.abspath(p); seen.add(ap.casefold())
            if os.path.exists(ap):
                st = os.stat(ap)
                item.setdefault('created_at', datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds'))
                item.setdefault('name', os.path.basename(ap)); item['path'] = ap
                item.setdefault('user', _fm472_safe_user(self, item.get('user')))
                item.setdefault('module', _fm472_guess_module(ap)); item.setdefault('type', 'file')
                items.append(item)
        except Exception:
            pass
    # 2) Dateiausgabe-Ordner scannen, damit auch bisherige Exporte ohne Log sichtbar werden.
    for root in _fm472_output_roots():
        if not os.path.isdir(root):
            continue
        try:
            for dirpath, dirnames, filenames in os.walk(root):
                # technische Logs/Backups ausblenden
                if any(part.lower() in {'__pycache__', 'logs'} for part in dirpath.split(os.sep)):
                    continue
                for fn in filenames:
                    ext = os.path.splitext(fn)[1].lower()
                    if ext not in _FM472_OUTPUT_EXTENSIONS:
                        continue
                    ap = os.path.abspath(os.path.join(dirpath, fn))
                    if ap.casefold() in seen:
                        continue
                    seen.add(ap.casefold())
                    try:
                        st = os.stat(ap)
                        items.append({
                            'id': _fm472_file_key(ap), 'type': 'file', 'name': fn, 'path': ap,
                            'user': _fm472_safe_user(self), 'module': _fm472_guess_module(ap),
                            'created_at': datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds'),
                            'group_id': '', 'group_title': ''
                        })
                    except Exception:
                        pass
        except Exception:
            pass
    return items


def _fm472_build_groups(self, items):
    # Registrierte group_id/group_title zuerst, sonst Unterordner als Sammelposition, wenn mehrere Dateien im Unterordner liegen.
    groups = {}
    singles = []
    dir_bucket = {}
    for item in items:
        gid = str(item.get('group_id') or '').strip()
        if gid:
            groups.setdefault(gid, {'id':'grp_'+gid, 'type':'group', 'name': item.get('group_title') or item.get('module') or 'Sammelexport', 'path':'', 'user': item.get('user',''), 'module': item.get('module',''), 'created_at': item.get('created_at',''), 'children': []})['children'].append(item)
        else:
            base_dir = _fm472_parent_dir_under_roots(item.get('path',''))
            dir_bucket.setdefault(base_dir, []).append(item)
    for base_dir, arr in dir_bucket.items():
        if len(arr) > 1 and base_dir and any(x.lower() in os.path.basename(base_dir).lower() for x in ('serien', 'sammel', 'export', 'dateiausgabe')):
            latest = max(arr, key=lambda x: str(x.get('created_at','')))
            gid = 'dir_' + _fm472_file_key(base_dir)
            groups[gid] = {'id':gid, 'type':'group', 'name': os.path.basename(base_dir) or 'Sammelexport', 'path': base_dir, 'user': latest.get('user',''), 'module': latest.get('module','Sammelexport'), 'created_at': latest.get('created_at',''), 'children': sorted(arr, key=lambda x: str(x.get('name','')).lower())}
        else:
            singles.extend(arr)
    out = list(groups.values()) + singles
    out.sort(key=lambda x: str(x.get('created_at','')), reverse=True)
    return out


def _fm472_filter_items(items, query):
    q = str(query or '').strip().casefold()
    if not q:
        return items
    def hit(item):
        vals = [item.get('name',''), item.get('user',''), item.get('module',''), item.get('created_at',''), _fm472_display_dt(item.get('created_at')), item.get('path','')]
        return q in ' '.join(str(v) for v in vals).casefold()
    result = []
    for item in items:
        if item.get('type') == 'group':
            children = [c for c in item.get('children', []) if hit(c)]
            if hit(item) or children:
                clone = dict(item); clone['children'] = children if children else item.get('children', [])
                result.append(clone)
        elif hit(item):
            result.append(item)
    return result


def _fm472_open_path(path):
    try:
        if not path or not os.path.exists(path):
            messagebox.showwarning('Dateiausgabe', 'Datei oder Ordner wurde nicht gefunden.'); return False
        if sys.platform.startswith('win'):
            os.startfile(path)
        elif sys.platform == 'darwin':
            import subprocess; subprocess.Popen(['open', path])
        else:
            import subprocess; subprocess.Popen(['xdg-open', path])
        return True
    except Exception as exc:
        try: messagebox.showerror('Dateiausgabe', 'Öffnen fehlgeschlagen:\n' + str(exc))
        except Exception: pass
        return False


def _fm472_open_file(item):
    return _fm472_open_path((item or {}).get('path',''))


def _fm472_open_folder(item):
    p = (item or {}).get('path','')
    if item and item.get('type') == 'group' and p and os.path.isdir(p):
        return _fm472_open_path(p)
    return _fm472_open_path(os.path.dirname(p))


def _fm472_send_outlook(item):
    try:
        path = (item or {}).get('path','')
        if item and item.get('type') == 'group':
            path = item.get('path') or os.path.dirname((item.get('children') or [{}])[0].get('path',''))
        if not path:
            return False
        try:
            import win32com.client
            outlook = win32com.client.Dispatch('Outlook.Application')
            mail = outlook.CreateItem(0)
            mail.Subject = 'FiBu Mate Dateiausgabe: ' + str((item or {}).get('name',''))
            mail.Body = 'Hallo,\n\nanbei / im Pfad findest du die Datei aus FiBu Mate:\n' + path + '\n\nViele Grüße'
            if os.path.isfile(path):
                mail.Attachments.Add(path)
            mail.Display()
            return True
        except Exception:
            import webbrowser
            from urllib.parse import quote
            webbrowser.open('mailto:?subject=' + quote('FiBu Mate Dateiausgabe: ' + str((item or {}).get('name',''))) + '&body=' + quote('Datei/Pfad aus FiBu Mate:\n' + path))
            return True
    except Exception as exc:
        try: messagebox.showerror('Dateiausgabe', 'Outlook-Versand konnte nicht vorbereitet werden:\n' + str(exc))
        except Exception: pass
        return False


def _fm472_send_teams(self, item):
    try:
        path = (item or {}).get('path','')
        if item and item.get('type') == 'group':
            path = item.get('path') or os.path.dirname((item.get('children') or [{}])[0].get('path',''))
        msg = 'FiBu Mate Dateiausgabe: ' + str((item or {}).get('name','')) + '\n' + path
        try:
            self.root.clipboard_clear(); self.root.clipboard_append(msg)
        except Exception:
            pass
        import webbrowser
        webbrowser.open('msteams:/')
        try: messagebox.showinfo('Dateiausgabe', 'Der Dateipfad wurde in die Zwischenablage kopiert. Teams wurde geöffnet; bitte in den gewünschten Chat einfügen.')
        except Exception: pass
        return True
    except Exception as exc:
        try: messagebox.showerror('Dateiausgabe', 'Teams-Weitergabe konnte nicht vorbereitet werden:\n' + str(exc))
        except Exception: pass
        return False


def _fm472_preview_pdf(canvas, path, width):
    pages = 1
    texts = []
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        pages = max(1, len(reader.pages))
        for i, page in enumerate(reader.pages[:8]):
            try: texts.append((page.extract_text() or '').strip()[:500])
            except Exception: texts.append('')
    except Exception:
        pass
    y = 12
    for i in range(pages):
        h = 190
        canvas.create_rectangle(14, y, width-28, y+h, fill='#FFFFFF', outline=LINE)
        canvas.create_text(28, y+20, text=f'PDF – Seite {i+1} von {pages}', anchor='w', fill=BLUE, font=body_font(11, weight='bold'))
        snippet = texts[i] if i < len(texts) and texts[i] else 'Vorschau als Seitenplatzhalter. Zum vollständigen Anzeigen bitte Datei öffnen.'
        canvas.create_text(28, y+48, text=snippet, anchor='nw', fill=TEXT, font=body_font(9), width=max(300, width-80))
        y += h + 14
    return y


def _fm472_preview_image(canvas, path, width):
    y = 12
    try:
        if PIL_AVAILABLE and os.path.exists(path):
            img = Image.open(path)
            img.thumbnail((max(220, width-80), 680))
            photo = ImageTk.PhotoImage(img)
            canvas._fm472_preview_refs = getattr(canvas, '_fm472_preview_refs', []) + [photo]
            canvas.create_image(28, y, image=photo, anchor='nw')
            y += photo.height() + 24
            return y
    except Exception:
        pass
    canvas.create_text(28, y, text='Bildvorschau konnte nicht geladen werden.', anchor='nw', fill=TEXT2, font=body_font(10))
    return y + 40


def _fm472_preview_text(canvas, path, width):
    y = 12
    txt = ''
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            txt = f.read(6000)
    except Exception as exc:
        txt = 'Textvorschau konnte nicht geladen werden:\n' + str(exc)
    canvas.create_rectangle(14, y, width-28, y+460, fill='#FFFFFF', outline=LINE)
    canvas.create_text(28, y+18, text=txt, anchor='nw', fill=TEXT, font=body_font(9), width=max(300, width-80))
    return y+480


def _fm472_preview_office(canvas, path, width):
    y = 12
    ext = os.path.splitext(path)[1].lower()
    title = {' .docx':'Word-Dokument', '.docx':'Word-Dokument', '.xlsx':'Excel-Arbeitsmappe', '.xls':'Excel-Arbeitsmappe', '.doc':'Word-Dokument'}.get(ext, 'Dokument')
    canvas.create_rectangle(14, y, width-28, y+260, fill='#FFFFFF', outline=LINE)
    canvas.create_text(28, y+20, text=title, anchor='w', fill=BLUE, font=body_font(12, weight='bold'))
    meta = [os.path.basename(path), _fm472_short_path(path, 92)]
    try:
        st = os.stat(path); meta.append('Geändert: ' + _fm472_display_dt(st.st_mtime)); meta.append('Größe: ' + str(round(st.st_size/1024, 1)) + ' KB')
    except Exception: pass
    # DOCX-Textauszug, falls möglich.
    snippet = ''
    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(path)
            snippet = '\n'.join(p.text for p in doc.paragraphs[:12] if p.text).strip()[:1200]
        except Exception:
            snippet = ''
    canvas.create_text(28, y+54, text='\n'.join(meta) + ('\n\n' + snippet if snippet else '\n\nFür die vollständige Vorschau bitte Datei öffnen.'), anchor='nw', fill=TEXT, font=body_font(9), width=max(300, width-80))
    return y+280


def _fm472_draw_preview(self, canvas, item):
    canvas.delete('all'); canvas._fm472_preview_refs = []
    width = max(460, canvas.winfo_width() or 620)
    if not item:
        canvas.create_text(18, 20, text='Bitte links eine Datei oder Sammelposition auswählen.', anchor='nw', fill=TEXT2, font=body_font(11))
        canvas.configure(scrollregion=(0,0,width,80)); return
    title = str(item.get('name','Datei'))
    path = item.get('path','')
    y = 12
    canvas.create_text(18, y, text=title, anchor='nw', fill=BLUE, font=body_font(14, weight='bold'), width=width-40)
    y += 34
    info = f"Benutzer: {item.get('user','')}    Erstellt/geändert: {_fm472_display_dt(item.get('created_at'))}\nPfad: {_fm472_short_path(path, 110)}"
    canvas.create_text(18, y, text=info, anchor='nw', fill=TEXT2, font=body_font(9), width=width-40)
    y += 54
    if item.get('type') == 'group':
        children = item.get('children', []) or []
        canvas.create_rectangle(14, y, width-28, y+max(140, 34+len(children)*26), fill='#FFFFFF', outline=LINE)
        canvas.create_text(28, y+18, text=f'Sammelposition mit {len(children)} Unterposition(en)', anchor='w', fill=BLUE, font=body_font(11, weight='bold'))
        yy = y+48
        for c in children[:40]:
            canvas.create_text(34, yy, text='• ' + str(c.get('name','')), anchor='w', fill=TEXT, font=body_font(9), width=width-80)
            yy += 24
        y = yy + 20
    elif path and os.path.exists(path):
        ext = os.path.splitext(path)[1].lower()
        if ext == '.pdf': y += _fm472_preview_pdf(canvas, path, width)
        elif ext in {'.png','.jpg','.jpeg','.bmp','.gif'}: y += _fm472_preview_image(canvas, path, width)
        elif ext in {'.txt','.csv'}: y += _fm472_preview_text(canvas, path, width)
        elif ext in {'.docx','.doc','.xlsx','.xls'}: y += _fm472_preview_office(canvas, path, width)
        else:
            canvas.create_text(18, y, text='Für diesen Dateityp ist nur eine Metadaten-Vorschau verfügbar.', anchor='nw', fill=TEXT, font=body_font(10)); y += 60
    else:
        canvas.create_text(18, y, text='Datei wurde nicht gefunden.', anchor='nw', fill=RED, font=body_font(10)); y += 60
    canvas.configure(scrollregion=(0,0,width,max(y+40, canvas.winfo_height())))


def _fm472_render_file_output(self):
    self.focusable_tiles.clear()
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    x0, y0 = 24, 135
    full_w, full_h = max(900, w-48), max(520, h-y0-64)
    outer = tk.Frame(self.root, bg=BG)
    self.widget_items.append(outer)
    self.canvas.create_window(ui_s(x0), ui_s(y0), window=outer, anchor='nw', width=ui_s(full_w), height=ui_s(full_h))
    search_var = tk.StringVar(value=getattr(self, 'file_output_search_text', ''))
    selected_item = {'item': getattr(self, 'file_output_selected_item', None)}
    all_items = self.file_output_build_items()

    top = tk.Frame(outer, bg=BG); top.pack(fill='x', pady=(0,8))
    tk.Label(top, text='Suche', bg=BG, fg=TEXT2, font=body_font(10)).pack(side='left', padx=(0,8))
    ent = tk.Entry(top, textvariable=search_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1)
    ent.pack(side='left', fill='x', expand=True, ipady=5)
    tk.Button(top, text='Aktualisieren', command=lambda: self.render_page(), bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(8,0), ipadx=10, ipady=4)

    content = tk.Frame(outer, bg=BG); content.pack(fill='both', expand=True)
    left = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=2); left.pack(side='left', fill='both', expand=True, padx=(0,12))
    right = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=2); right.pack(side='right', fill='both', expand=True)
    left.configure(width=max(520, int(full_w*0.48))); right.configure(width=max(520, int(full_w*0.50)))

    head = tk.Frame(left, bg=WHITE); head.pack(fill='x', padx=12, pady=(10,6))
    tk.Label(head, text='Zuletzt erstellte Dateien', bg=WHITE, fg=BLUE, font=body_font(13, weight='bold')).pack(side='left')
    action_bar = tk.Frame(left, bg=WHITE); action_bar.pack(fill='x', padx=12, pady=(0,8))
    for label, cmd in [
        ('Datei öffnen', lambda: _fm472_open_file(selected_item.get('item') or {})),
        ('Speicherort öffnen', lambda: _fm472_open_folder(selected_item.get('item') or {})),
        ('Per Outlook senden', lambda: _fm472_send_outlook(selected_item.get('item') or {})),
        ('Per Teams senden', lambda: _fm472_send_teams(self, selected_item.get('item') or {})),
    ]:
        tk.Button(action_bar, text=label, command=cmd, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0,6), ipadx=8, ipady=3)

    tree_frame = tk.Frame(left, bg=WHITE); tree_frame.pack(fill='both', expand=True, padx=12, pady=(0,12))
    columns = ('name','user','date','path')
    tree = ttk.Treeview(tree_frame, columns=columns, show='tree headings', selectmode='browse')
    tree.heading('#0', text='Typ')
    tree.heading('name', text='Dateiname / Sammelposition')
    tree.heading('user', text='Benutzer')
    tree.heading('date', text='Erstellt')
    tree.heading('path', text='Ablagepfad kurz')
    tree.column('#0', width=54, stretch=False)
    tree.column('name', width=210, stretch=True)
    tree.column('user', width=120, stretch=False)
    tree.column('date', width=122, stretch=False)
    tree.column('path', width=190, stretch=True)
    vsb = ttk.Scrollbar(tree_frame, orient='vertical', command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    tree.pack(side='left', fill='both', expand=True); vsb.pack(side='right', fill='y')

    prev_head = tk.Frame(right, bg=WHITE); prev_head.pack(fill='x', padx=12, pady=(10,6))
    tk.Label(prev_head, text='Vorschau', bg=WHITE, fg=BLUE, font=body_font(13, weight='bold')).pack(side='left')
    prev_canvas = tk.Canvas(right, bg='#F8FAFC', highlightthickness=0)
    prev_scroll = ttk.Scrollbar(right, orient='vertical', command=prev_canvas.yview)
    prev_canvas.configure(yscrollcommand=prev_scroll.set)
    prev_canvas.pack(side='left', fill='both', expand=True, padx=(12,0), pady=(0,12))
    prev_scroll.pack(side='right', fill='y', padx=(0,12), pady=(0,12))

    id_map = {}
    def fill_tree():
        tree.delete(*tree.get_children())
        q = search_var.get(); self.file_output_search_text = q
        shown = _fm472_filter_items(all_items, q)
        for item in shown:
            iid = item.get('id') or _fm472_file_key(item.get('path','') + item.get('name',''))
            id_map[iid] = item
            typ = '▸' if item.get('type') == 'group' else '•'
            tree.insert('', 'end', iid=iid, text=typ, values=(item.get('name',''), item.get('user',''), _fm472_display_dt(item.get('created_at')), _fm472_short_path(item.get('path',''))))
            if item.get('type') == 'group':
                for c in item.get('children', []) or []:
                    cid = c.get('id') or _fm472_file_key(c.get('path',''))
                    if cid in id_map: cid = iid + '_' + cid
                    id_map[cid] = c
                    tree.insert(iid, 'end', iid=cid, text='•', values=(c.get('name',''), c.get('user',''), _fm472_display_dt(c.get('created_at')), _fm472_short_path(c.get('path',''))))
    def on_select(event=None):
        sel = tree.selection()
        item = id_map.get(sel[0]) if sel else None
        selected_item['item'] = item
        self.file_output_selected_item = item
        _fm472_draw_preview(self, prev_canvas, item)
    def on_open(event=None):
        item = selected_item.get('item')
        if item and item.get('type') == 'group':
            try: tree.item(tree.selection()[0], open=not bool(tree.item(tree.selection()[0], 'open')))
            except Exception: pass
        elif item:
            _fm472_open_file(item)
    tree.bind('<<TreeviewSelect>>', on_select)
    tree.bind('<Double-Button-1>', on_open)
    search_var.trace_add('write', lambda *_: (fill_tree(), _fm472_draw_preview(self, prev_canvas, selected_item.get('item'))))
    try: prev_canvas.bind('<MouseWheel>', lambda e: (prev_canvas.yview_scroll(int(-1*(e.delta/120)), 'units'), 'break'))
    except Exception: pass
    fill_tree()
    if tree.get_children():
        first = tree.get_children()[0]
        tree.selection_set(first); tree.focus(first); on_select()
    else:
        _fm472_draw_preview(self, prev_canvas, None)


def _fm472_build_items(self):
    return _fm472_build_groups(self, _fm472_scan_output_files(self))


try:
    FiBuMateApp.file_output_register = _fm472_register_output_file
    FiBuMateApp.file_output_build_items = _fm472_build_items
    FiBuMateApp.render_file_output = _fm472_render_file_output
except Exception:
    pass

try:
    _fm472_old_render_page = FiBuMateApp.render_page
except Exception:
    _fm472_old_render_page = None

def _fm472_render_page(self):
    if getattr(self, 'current_page', '') == 'file_output':
        self.init_responsive_scaling()
        self.clear_content()
        self.draw_background()
        self.draw_header(getattr(self, 'current_title', 'Dateiausgabe') or 'Dateiausgabe')
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        self.render_file_output()
        return
    if _fm472_old_render_page:
        return _fm472_old_render_page(self)

try:
    FiBuMateApp.render_page = _fm472_render_page
except Exception:
    pass

try:
    MODULE_DESCRIPTIONS['file_output'] = 'Zentrale Übersicht zuletzt mit FiBu Mate erstellter Dateien inklusive Vorschau, Öffnen, Speicherort, Outlook- und Teams-Weitergabe.'
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Dateiausgabe Vorschau/Versand/Löschen/Sortierung FINAL 2026-06-23
# Version 0.473
# Zweck:
# - Entfernt den inneren Rahmen in der Vorschau.
# - Ergänzt echte Tabellen-Vorschau für Excel/CSV.
# - Versand versendet/übergibt Dateien, protokolliert Versand und zeigt Versandindikator mit Mouseover.
# - Buttons einheitlich grau und mit Outlook/Teams/Explorer-Icons.
# - Sammelpositionen nur noch bei gleicher group_id = gleicher Exportvorgang.
# - B4 darf Exporte löschen; Sortierung per Spaltenkopf, Standard: neueste zuerst.
# ------------------------------------------------------------------

_FM473_PATCH_VERSION = "0.473"
_FM473_SENT_ICON = r"C:\python\bin\Imgs\Icons\1485969917-3-right_78909.ico"
_FM473_OUTLOOK_ICON = r"C:\python\bin\Imgs\Icons\microsoft_office_outlook_logo_icon_145721.ico"
_FM473_TEAMS_ICON = r"C:\python\bin\Imgs\Icons\microsoft_office_teams_logo_icon_145726.ico"
_FM473_EXPLORER_ICON = r"C:\python\bin\Imgs\Icons\File_Explorer_23583.ico"
_FM473_BUTTON_GREY = "#E5E7EB"
_FM473_BUTTON_GREY_ACTIVE = "#D1D5DB"


def _fm473_can_delete_exports(self):
    try:
        return int(ROLE_RANK.get(self.my_role(), 1)) >= 4
    except Exception:
        try: return int(self.kb_current_permission_level()) >= 4
        except Exception: return False


def _fm473_icon_photo(self, path, size=18):
    try:
        cache = getattr(self, '_fm473_icon_cache', {})
        key = (path, int(size))
        if key in cache:
            return cache[key]
        if PIL_AVAILABLE and path and os.path.exists(path):
            img = Image.open(path)
            img.thumbnail((int(size), int(size)))
            ph = ImageTk.PhotoImage(img)
            cache[key] = ph
            self._fm473_icon_cache = cache
            return ph
    except Exception:
        pass
    return None


def _fm473_make_button(parent, text, command, icon_path=None, state='normal'):
    img = None
    try:
        app = parent
        while app and not hasattr(app, '_fm473_icon_cache') and hasattr(app, 'master'):
            app = app.master
    except Exception:
        app = None
    # Buttonbilder werden im Render selbst nachträglich gesetzt, wenn App-Kontext verfügbar ist.
    return tk.Button(parent, text=text, command=command, bg=_FM473_BUTTON_GREY, fg=TEXT,
                     activebackground=_FM473_BUTTON_GREY_ACTIVE, font=body_font(9), relief='solid', bd=1,
                     state=state, cursor='hand2' if state == 'normal' else 'arrow', compound='left', padx=8, pady=3)


def _fm473_log_target():
    paths = _fm472_log_paths()
    return paths[0] if paths else os.path.join(os.getcwd(), 'fibu_mate_dateiausgabe_log.json')


def _fm473_load_all_log_items():
    by_id = {}
    for p in _fm472_log_paths():
        d = _fm472_read_json(p)
        for item in (d.get('items', []) if isinstance(d, dict) else []):
            if isinstance(item, dict):
                iid = item.get('id') or _fm472_file_key(item.get('path',''))
                if iid: by_id[iid] = dict(item)
    return by_id


def _fm473_write_log_items(items_by_id):
    target = _fm473_log_target()
    try:
        _fm472_write_json(target, {'items': list(items_by_id.values())[-1000:]})
        return True
    except Exception:
        return False


def _fm473_mark_sent(self, item, channel):
    now = datetime.now().isoformat(timespec='seconds')
    ids = []
    if item.get('type') == 'group':
        for c in item.get('children', []) or []:
            ids.append(c.get('id') or _fm472_file_key(c.get('path','')))
            c['sent_at'] = now; c['sent_channel'] = channel
        item['sent_at'] = now; item['sent_channel'] = channel
    else:
        ids.append(item.get('id') or _fm472_file_key(item.get('path','')))
        item['sent_at'] = now; item['sent_channel'] = channel
    logs = _fm473_load_all_log_items()
    for iid in ids:
        current = logs.get(iid, {})
        current.setdefault('id', iid)
        if item.get('type') != 'group':
            current.update({k:item.get(k) for k in ('name','path','user','module','created_at','group_id','group_title') if item.get(k) is not None})
        current['sent_at'] = now; current['sent_channel'] = channel
        logs[iid] = current
    _fm473_write_log_items(logs)
    return now


def _fm473_files_for_item(item):
    files = []
    if not item:
        return files
    if item.get('type') == 'group':
        for c in item.get('children', []) or []:
            p = c.get('path','')
            if p and os.path.isfile(p): files.append(p)
    else:
        p = item.get('path','')
        if p and os.path.isfile(p): files.append(p)
    return files


def _fm473_send_outlook(self, item):
    files = _fm473_files_for_item(item)
    if not files:
        try: messagebox.showwarning('Dateiausgabe', 'Keine versendbare Datei gefunden.')
        except Exception: pass
        return False
    try:
        import win32com.client
        outlook = win32com.client.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)
        mail.Subject = 'FiBu Mate Dateiausgabe: ' + str((item or {}).get('name',''))
        mail.Body = 'Hallo,\n\nanbei die Datei(en) aus FiBu Mate.\n\nViele Grüße'
        for f in files:
            mail.Attachments.Add(f)
        mail.Display()
        _fm473_mark_sent(self, item, 'Outlook')
        return True
    except Exception as exc:
        try:
            messagebox.showerror('Dateiausgabe', 'Outlook-Versand mit Dateianhang konnte nicht vorbereitet werden:\n' + str(exc))
        except Exception: pass
        return False


def _fm473_copy_files_to_clipboard(paths):
    # Windows: CF_HDROP über pywin32, wenn verfügbar. Danach kann in Teams eingefügt werden.
    try:
        import win32clipboard, win32con, struct
        files = '\0'.join(paths) + '\0\0'
        data = files.encode('utf-16le')
        dropfiles = struct.pack('IiiII', 20, 0, 0, 0, 1)
        win32clipboard.OpenClipboard()
        try:
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardData(win32con.CF_HDROP, dropfiles + data)
        finally:
            win32clipboard.CloseClipboard()
        return True
    except Exception:
        return False


def _fm473_send_teams(self, item):
    files = _fm473_files_for_item(item)
    if not files:
        try: messagebox.showwarning('Dateiausgabe', 'Keine versendbare Datei gefunden.')
        except Exception: pass
        return False
    ok_clip = _fm473_copy_files_to_clipboard(files)
    try:
        import webbrowser
        webbrowser.open('msteams:/')
    except Exception:
        pass
    if ok_clip:
        try: messagebox.showinfo('Dateiausgabe', 'Die Datei(en) wurden in die Zwischenablage gelegt. Teams wurde geöffnet; bitte im gewünschten Chat einfügen.')
        except Exception: pass
        _fm473_mark_sent(self, item, 'Teams')
        return True
    # Fallback: Ordner öffnen und trotzdem Versandvorbereitung nicht als erfolgreich protokollieren.
    try:
        _fm472_open_folder(item)
        messagebox.showwarning('Dateiausgabe', 'Automatisches Anhängen in Teams ist auf diesem System nicht verfügbar. Der Speicherort wurde geöffnet; bitte Datei in Teams ziehen oder einfügen.')
    except Exception:
        pass
    return False


def _fm473_register_output_file(self, path, user=None, module=None, group_id=None, group_title=None, created_at=None, children=None):
    # group_id wird nur verwendet, wenn sie explizit vom Exportvorgang übergeben wird.
    try:
        if not path: return False
        entry = {'id': _fm472_file_key(path), 'type':'file', 'name':os.path.basename(str(path)),
                 'path':os.path.abspath(str(path)), 'user':user or _fm472_safe_user(self),
                 'module':module or _fm472_guess_module(path), 'group_id':group_id or '',
                 'group_title':group_title or '', 'created_at':created_at or _fm472_now_iso()}
        logs = _fm473_load_all_log_items()
        old = logs.get(entry['id'], {})
        old.update(entry)
        logs[entry['id']] = old
        return _fm473_write_log_items(logs)
    except Exception:
        return False


def _fm473_load_registered_items(self):
    return list(_fm473_load_all_log_items().values())


def _fm473_scan_output_files(self):
    items, seen = [], set()
    logs = _fm473_load_all_log_items()
    for item in logs.values():
        p = item.get('path')
        if p:
            try:
                ap = os.path.abspath(p); seen.add(ap.casefold())
                if os.path.exists(ap):
                    st = os.stat(ap)
                    item = dict(item); item.setdefault('created_at', datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds'))
                    item.setdefault('name', os.path.basename(ap)); item['path'] = ap
                    item.setdefault('user', _fm472_safe_user(self, item.get('user'))); item.setdefault('module', _fm472_guess_module(ap)); item.setdefault('type','file')
                    items.append(item)
            except Exception: pass
    for root in _fm472_output_roots():
        if not os.path.isdir(root): continue
        try:
            for dirpath, dirnames, filenames in os.walk(root):
                if any(part.lower() in {'__pycache__','logs'} for part in dirpath.split(os.sep)): continue
                for fn in filenames:
                    ext=os.path.splitext(fn)[1].lower()
                    if ext not in _FM472_OUTPUT_EXTENSIONS: continue
                    ap=os.path.abspath(os.path.join(dirpath, fn))
                    if ap.casefold() in seen: continue
                    seen.add(ap.casefold())
                    try:
                        st=os.stat(ap); iid=_fm472_file_key(ap)
                        li=logs.get(iid,{})
                        items.append({'id':iid,'type':'file','name':fn,'path':ap,'user':li.get('user') or _fm472_safe_user(self),
                                      'module':li.get('module') or _fm472_guess_module(ap),
                                      'created_at':li.get('created_at') or datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds'),
                                      'group_id':li.get('group_id',''),'group_title':li.get('group_title',''),
                                      'sent_at':li.get('sent_at',''),'sent_channel':li.get('sent_channel','')})
                    except Exception: pass
        except Exception: pass
    return items


def _fm473_build_groups(self, items):
    groups, singles = {}, []
    for item in items:
        gid = str(item.get('group_id') or '').strip()
        # Nur explizite group_id erzeugt Sammelposition = gleicher Exportvorgang.
        if gid:
            g = groups.setdefault(gid, {'id':'grp_'+gid,'type':'group','name':item.get('group_title') or item.get('module') or 'Sammelexport',
                                        'path':'','user':item.get('user',''),'module':item.get('module',''),
                                        'created_at':item.get('created_at',''),'children':[], 'sent_at':item.get('sent_at',''), 'sent_channel':item.get('sent_channel','')})
            g['children'].append(item)
            if str(item.get('created_at','')) > str(g.get('created_at','')): g['created_at'] = item.get('created_at','')
            if item.get('sent_at') and str(item.get('sent_at')) > str(g.get('sent_at','')): g['sent_at']=item.get('sent_at'); g['sent_channel']=item.get('sent_channel','')
        else:
            singles.append(item)
    out = list(groups.values()) + singles
    out.sort(key=lambda x: str(x.get('created_at','')), reverse=True)
    return out


def _fm473_build_items(self):
    return _fm473_build_groups(self, _fm473_scan_output_files(self))


def _fm473_preview_table(canvas, rows, width, start_y=96, title='Tabellenvorschau'):
    y = start_y
    canvas.create_text(18, y, text=title, anchor='nw', fill=BLUE, font=body_font(12, weight='bold'))
    y += 30
    if not rows:
        canvas.create_text(18, y, text='Keine Tabellendaten gefunden.', anchor='nw', fill=TEXT2, font=body_font(10)); return y+40
    max_cols = min(8, max(len(r) for r in rows))
    col_w = max(80, int((width-48)/max_cols))
    row_h = 26
    for r_idx, row in enumerate(rows[:26]):
        bg = '#EEF2F7' if r_idx == 0 else ('#FFFFFF' if r_idx % 2 else '#F8FAFC')
        y0 = y + r_idx*row_h
        canvas.create_rectangle(18, y0, 18+col_w*max_cols, y0+row_h, fill=bg, outline='#E5E7EB')
        for c in range(max_cols):
            x0 = 18 + c*col_w
            if c > 0: canvas.create_line(x0, y0, x0, y0+row_h, fill='#E5E7EB')
            val = '' if c >= len(row) or row[c] is None else str(row[c])
            if len(val) > 22: val = val[:21] + '…'
            canvas.create_text(x0+6, y0+row_h/2, text=val, anchor='w', fill=TEXT, font=body_font(8, weight='bold' if r_idx==0 else None), width=col_w-10)
    return y + min(26, len(rows))*row_h + 30


def _fm473_preview_excel(canvas, path, width):
    rows=[]
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        for i, row in enumerate(ws.iter_rows(max_rows=26, max_cols=8, values_only=True)):
            rows.append([v for v in row])
        try: wb.close()
        except Exception: pass
    except Exception as exc:
        rows = [['Excel-Vorschau konnte nicht geladen werden', str(exc)]]
    return _fm473_preview_table(canvas, rows, width, 96, 'Excel-Vorschau')


def _fm473_preview_csv(canvas, path, width):
    rows=[]
    try:
        import csv
        with open(path, 'r', encoding='utf-8-sig', errors='replace', newline='') as f:
            sample = f.read(4096); f.seek(0)
            try: dialect = csv.Sniffer().sniff(sample, delimiters=';,\t,')
            except Exception: dialect = csv.excel
            reader = csv.reader(f, dialect)
            for i, row in enumerate(reader):
                rows.append(row)
                if i >= 25: break
    except Exception as exc:
        rows = [['CSV-Vorschau konnte nicht geladen werden', str(exc)]]
    return _fm473_preview_table(canvas, rows, width, 96, 'CSV-Vorschau')


def _fm473_draw_preview(self, canvas, item):
    canvas.delete('all'); canvas._fm472_preview_refs=[]
    width=max(460, canvas.winfo_width() or 620)
    if not item:
        canvas.create_text(18, 20, text='Bitte links eine Datei oder Sammelposition auswählen.', anchor='nw', fill=TEXT2, font=body_font(11))
        canvas.configure(scrollregion=(0,0,width,80)); return
    title=str(item.get('name','Datei')); path=item.get('path','')
    y=12
    canvas.create_text(18,y,text=title,anchor='nw',fill=BLUE,font=body_font(14,weight='bold'),width=width-40); y+=34
    info=f"Benutzer: {item.get('user','')}    Erstellt/geändert: {_fm472_display_dt(item.get('created_at'))}"
    if item.get('sent_at'): info += f"    Versand: {_fm472_display_dt(item.get('sent_at'))} ({item.get('sent_channel','')})"
    info += f"\nPfad: {_fm472_short_path(path, 110)}"
    canvas.create_text(18,y,text=info,anchor='nw',fill=TEXT2,font=body_font(9),width=width-40); y+=54
    if item.get('type')=='group':
        children=item.get('children',[]) or []
        canvas.create_text(18,y,text=f'Sammelposition mit {len(children)} Unterposition(en)',anchor='nw',fill=BLUE,font=body_font(11,weight='bold')); y+=32
        for c in children[:60]:
            s='• '+str(c.get('name',''))
            if c.get('sent_at'): s += '   ✓ Versand: ' + _fm472_display_dt(c.get('sent_at'))
            canvas.create_text(28,y,text=s,anchor='nw',fill=TEXT,font=body_font(9),width=width-70); y+=24
    elif path and os.path.exists(path):
        ext=os.path.splitext(path)[1].lower()
        if ext=='.pdf': y += _fm472_preview_pdf(canvas,path,width)
        elif ext in {'.png','.jpg','.jpeg','.bmp','.gif'}: y += _fm472_preview_image(canvas,path,width)
        elif ext in {'.xlsx','.xls'}: y = _fm473_preview_excel(canvas,path,width)
        elif ext=='.csv': y = _fm473_preview_csv(canvas,path,width)
        elif ext=='.txt': y += _fm472_preview_text(canvas,path,width)
        elif ext in {'.docx','.doc'}: y += _fm472_preview_office(canvas,path,width)
        else:
            canvas.create_text(18,y,text='Für diesen Dateityp ist nur eine Metadaten-Vorschau verfügbar.',anchor='nw',fill=TEXT,font=body_font(10)); y+=60
    else:
        canvas.create_text(18,y,text='Datei wurde nicht gefunden.',anchor='nw',fill=RED,font=body_font(10)); y+=60
    canvas.configure(scrollregion=(0,0,width,max(y+40,canvas.winfo_height())))


def _fm473_delete_export(self, item):
    if not _fm473_can_delete_exports(self):
        return False
    files = _fm473_files_for_item(item)
    if not files:
        try: messagebox.showwarning('Dateiausgabe', 'Keine löschbare Datei gefunden.')
        except Exception: pass
        return False
    try:
        if not messagebox.askyesno('Export löschen', 'Export wirklich löschen?\n\n' + str(item.get('name',''))):
            return False
    except Exception:
        return False
    errors=[]
    for f in files:
        try: os.remove(f)
        except Exception as exc: errors.append(f'{os.path.basename(f)}: {exc}')
    logs = _fm473_load_all_log_items()
    for f in files:
        logs.pop(_fm472_file_key(f), None)
    _fm473_write_log_items(logs)
    if errors:
        try: messagebox.showwarning('Dateiausgabe', 'Einige Dateien konnten nicht gelöscht werden:\n' + '\n'.join(errors))
        except Exception: pass
    self.file_output_selected_item = None
    self.render_page()
    return not errors


def _fm473_sort_items(items, column, descending):
    def key(item):
        if column == 'name': return str(item.get('name','')).casefold()
        if column == 'user': return str(item.get('user','')).casefold()
        if column == 'path': return str(item.get('path','')).casefold()
        return str(item.get('created_at',''))
    out = sorted(items, key=key, reverse=bool(descending))
    for item in out:
        if item.get('type') == 'group':
            item['children'] = sorted(item.get('children',[]) or [], key=key, reverse=bool(descending))
    return out


def _fm473_render_file_output(self):
    self.focusable_tiles.clear()
    w,h=self.canvas.winfo_width(), self.canvas.winfo_height()
    x0,y0=24,135; full_w,full_h=max(900,w-48),max(520,h-y0-64)
    outer=tk.Frame(self.root,bg=BG); self.widget_items.append(outer)
    self.canvas.create_window(ui_s(x0),ui_s(y0),window=outer,anchor='nw',width=ui_s(full_w),height=ui_s(full_h))
    search_var=tk.StringVar(value=getattr(self,'file_output_search_text',''))
    selected_item={'item':getattr(self,'file_output_selected_item',None)}
    sort_state=getattr(self,'file_output_sort_state',{'column':'date','descending':True}) or {'column':'date','descending':True}
    all_items=self.file_output_build_items()

    top=tk.Frame(outer,bg=BG); top.pack(fill='x',pady=(0,8))
    tk.Label(top,text='Suche',bg=BG,fg=TEXT2,font=body_font(10)).pack(side='left',padx=(0,8))
    ent=tk.Entry(top,textvariable=search_var,bg=WHITE,fg=TEXT,font=body_font(11),relief='solid',bd=1); ent.pack(side='left',fill='x',expand=True,ipady=5)
    tk.Button(top,text='Aktualisieren',command=lambda:self.render_page(),bg=_FM473_BUTTON_GREY,fg=TEXT,activebackground=_FM473_BUTTON_GREY_ACTIVE,font=body_font(10),relief='solid',bd=1).pack(side='left',padx=(8,0),ipadx=10,ipady=4)
    content=tk.Frame(outer,bg=BG); content.pack(fill='both',expand=True)
    left=tk.Frame(content,bg=WHITE,highlightbackground=LINE,highlightthickness=2); left.pack(side='left',fill='both',expand=True,padx=(0,12))
    right=tk.Frame(content,bg=WHITE,highlightbackground=LINE,highlightthickness=2); right.pack(side='right',fill='both',expand=True)
    left.configure(width=max(470,int(full_w*0.42))); right.configure(width=max(580,int(full_w*0.56)))
    head=tk.Frame(left,bg=WHITE); head.pack(fill='x',padx=12,pady=(10,6))
    tk.Label(head,text='Zuletzt erstellte Dateien',bg=WHITE,fg=BLUE,font=body_font(13,weight='bold')).pack(side='left')
    action_bar=tk.Frame(left,bg=WHITE); action_bar.pack(fill='x',padx=12,pady=(0,8))
    def add_btn(text, cmd, icon=None, enabled=True):
        b=tk.Button(action_bar,text=text,command=cmd,bg=_FM473_BUTTON_GREY,fg=TEXT,activebackground=_FM473_BUTTON_GREY_ACTIVE,font=body_font(9),relief='solid',bd=1,state='normal' if enabled else 'disabled',compound='left',cursor='hand2' if enabled else 'arrow')
        ph=_fm473_icon_photo(self, icon, 18) if icon else None
        if ph: b.configure(image=ph); b._img=ph
        b.pack(side='left',padx=(0,6),ipadx=8,ipady=3); return b
    add_btn('Datei öffnen', lambda:_fm472_open_file(selected_item.get('item') or {}))
    add_btn('Speicherort öffnen', lambda:_fm472_open_folder(selected_item.get('item') or {}), _FM473_EXPLORER_ICON)
    add_btn('Outlook Versand', lambda:(_fm473_send_outlook(self, selected_item.get('item') or {}) and self.render_page()), _FM473_OUTLOOK_ICON)
    add_btn('Teams Versand', lambda:(_fm473_send_teams(self, selected_item.get('item') or {}) and self.render_page()), _FM473_TEAMS_ICON)
    add_btn('Export löschen', lambda:_fm473_delete_export(self, selected_item.get('item') or {}), None, _fm473_can_delete_exports(self))

    tree_frame=tk.Frame(left,bg=WHITE); tree_frame.pack(fill='both',expand=True,padx=12,pady=(0,12))
    columns=('name','user','date','path')
    tree=ttk.Treeview(tree_frame,columns=columns,show='tree headings',selectmode='browse')
    sent_img=_fm473_icon_photo(self,_FM473_SENT_ICON,16)
    self._fm473_sent_img=sent_img
    def heading_text(col,label):
        marker=' ↓' if sort_state.get('column')==col and sort_state.get('descending',True) else (' ↑' if sort_state.get('column')==col else '')
        return label+marker
    def set_sort(col):
        cur=getattr(self,'file_output_sort_state',{'column':'date','descending':True})
        desc=not cur.get('descending',True) if cur.get('column')==col else (True if col=='date' else False)
        self.file_output_sort_state={'column':col,'descending':desc}; self.render_page()
    tree.heading('#0',text='')
    tree.heading('name',text=heading_text('name','Dateiname / Sammelposition'),command=lambda:set_sort('name'))
    tree.heading('user',text=heading_text('user','Benutzer'),command=lambda:set_sort('user'))
    tree.heading('date',text=heading_text('date','Erstellt'),command=lambda:set_sort('date'))
    tree.heading('path',text=heading_text('path','Ablagepfad kurz'),command=lambda:set_sort('path'))
    tree.column('#0',width=24,minwidth=24,stretch=False,anchor='w'); tree.column('name',width=220,stretch=True); tree.column('user',width=120,stretch=False); tree.column('date',width=122,stretch=False); tree.column('path',width=190,stretch=True)
    vsb=ttk.Scrollbar(tree_frame,orient='vertical',command=tree.yview); tree.configure(yscrollcommand=vsb.set)
    tree.pack(side='left',fill='both',expand=True); vsb.pack(side='right',fill='y')
    prev_head=tk.Frame(right,bg=WHITE); prev_head.pack(fill='x',padx=12,pady=(10,6))
    tk.Label(prev_head,text='Vorschau',bg=WHITE,fg=BLUE,font=body_font(13,weight='bold')).pack(side='left')
    prev_canvas=tk.Canvas(right,bg='#F8FAFC',highlightthickness=0,bd=0)
    prev_scroll=ttk.Scrollbar(right,orient='vertical',command=prev_canvas.yview); prev_canvas.configure(yscrollcommand=prev_scroll.set)
    prev_canvas.pack(side='left',fill='both',expand=True,padx=(12,0),pady=(0,12)); prev_scroll.pack(side='right',fill='y',padx=(0,12),pady=(0,12))
    id_map={}; tooltip={'win':None}
    def fill_tree():
        tree.delete(*tree.get_children()); id_map.clear()
        q=search_var.get(); self.file_output_search_text=q
        shown=_fm472_filter_items(all_items,q)
        shown=_fm473_sort_items(shown, sort_state.get('column','date'), sort_state.get('descending',True))
        for item in shown:
            iid=item.get('id') or _fm472_file_key(item.get('path','')+item.get('name',''))
            id_map[iid]=item
            typ='▸' if item.get('type')=='group' else ''
            img=sent_img if item.get('sent_at') and sent_img else ''
            tree.insert('', 'end', iid=iid, text=typ, image=img, values=(item.get('name',''), item.get('user',''), _fm472_display_dt(item.get('created_at')), _fm472_short_path(item.get('path',''))))
            if item.get('type')=='group':
                for c in item.get('children',[]) or []:
                    cid=c.get('id') or _fm472_file_key(c.get('path',''))
                    if cid in id_map: cid=iid+'_'+cid
                    id_map[cid]=c
                    cimg=sent_img if c.get('sent_at') and sent_img else ''
                    tree.insert(iid,'end',iid=cid,text='',image=cimg,values=(c.get('name',''),c.get('user',''),_fm472_display_dt(c.get('created_at')),_fm472_short_path(c.get('path',''))))
    def on_select(event=None):
        sel=tree.selection(); item=id_map.get(sel[0]) if sel else None
        selected_item['item']=item; self.file_output_selected_item=item; _fm473_draw_preview(self,prev_canvas,item)
    def on_open(event=None):
        item=selected_item.get('item')
        if item and item.get('type')=='group':
            try: tree.item(tree.selection()[0],open=not bool(tree.item(tree.selection()[0],'open')))
            except Exception: pass
        elif item: _fm472_open_file(item)
    def hide_tip():
        if tooltip.get('win'):
            try: tooltip['win'].destroy()
            except Exception: pass
            tooltip['win']=None
    def on_motion(ev):
        row=tree.identify_row(ev.y); col=tree.identify_column(ev.x)
        item=id_map.get(row) if row else None
        if item and item.get('sent_at') and col=='#0':
            if tooltip.get('row')==row and tooltip.get('win'): return
            hide_tip(); tooltip['row']=row
            tw=tk.Toplevel(tree); tw.wm_overrideredirect(True); tw.wm_geometry(f'+{tree.winfo_rootx()+ev.x+18}+{tree.winfo_rooty()+ev.y+18}')
            tk.Label(tw,text='Versendet: '+_fm472_display_dt(item.get('sent_at'))+' '+str(item.get('sent_channel','')),bg='#FFF7CC',fg=TEXT,relief='solid',bd=1,font=body_font(9)).pack(ipadx=6,ipady=3)
            tooltip['win']=tw
        else:
            hide_tip(); tooltip['row']=None
    tree.bind('<<TreeviewSelect>>',on_select); tree.bind('<Double-Button-1>',on_open); tree.bind('<Motion>',on_motion); tree.bind('<Leave>',lambda e:hide_tip())
    search_var.trace_add('write',lambda *_:(fill_tree(),_fm473_draw_preview(self,prev_canvas,selected_item.get('item'))))
    try: prev_canvas.bind('<MouseWheel>',lambda e:(prev_canvas.yview_scroll(int(-1*(e.delta/120)),'units'),'break'))
    except Exception: pass
    fill_tree()
    if tree.get_children():
        first=tree.get_children()[0]; tree.selection_set(first); tree.focus(first); on_select()
    else:
        _fm473_draw_preview(self,prev_canvas,None)

try:
    FiBuMateApp.file_output_register = _fm473_register_output_file
    FiBuMateApp.file_output_build_items = _fm473_build_items
    FiBuMateApp.render_file_output = _fm473_render_file_output
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Dateiausgabe Layout/Exportpfade/Wissenszentrale/Userrechte FINAL 2026-06-23
# Version 0.474
# ------------------------------------------------------------------
_FM474_PATCH_VERSION = "0.474"
_FM474_STANDARD_EXPORT_DIR = os.path.join(NETWORK_ROOT, "Dateiausgabe")
_FM474_WZ_EXPORT_DIR = os.path.join(NETWORK_ROOT, "Wissenszentrale")
_FM474_BTN_GREY = "#D6DCE4"
_FM474_BTN_GREY_ACTIVE = "#C8D0DA"

try:
    _FM473_BUTTON_GREY = _FM474_BTN_GREY
    _FM473_BUTTON_GREY_ACTIVE = _FM474_BTN_GREY_ACTIVE
except Exception:
    pass

def _fm474_ensure_dir(path):
    try: os.makedirs(path, exist_ok=True)
    except Exception: pass
    return path

def _fm474_standard_export_dir():
    return _fm474_ensure_dir(_FM474_STANDARD_EXPORT_DIR)

def _fm474_wz_export_dir():
    return _fm474_ensure_dir(_FM474_WZ_EXPORT_DIR)

def _fm474_output_roots():
    roots=[]
    for p in (_FM474_STANDARD_EXPORT_DIR, _FM474_WZ_EXPORT_DIR, os.path.join(SCRIPT_DIR, 'Dateiausgabe'), r'C:\python\Dateiausgabe'):
        try:
            ap=os.path.abspath(p)
            if ap not in roots: roots.append(ap)
        except Exception: pass
    return roots

def _fm474_primary_output_root():
    return _fm474_standard_export_dir()

try:
    _fm472_output_roots = _fm474_output_roots
    _fm472_primary_output_root = _fm474_primary_output_root
except Exception:
    pass

def _fm474_install_export_dialog_defaults():
    try:
        import tkinter.filedialog as _fd
        if getattr(_fd, '_fm474_export_defaults_installed', False): return
        _fd._fm474_export_defaults_installed=True
        _fd._fm474_orig_asksaveasfilename=_fd.asksaveasfilename
        _fd._fm474_orig_asksaveasfile=_fd.asksaveasfile
        _fd._fm474_orig_askdirectory=_fd.askdirectory
        def _target(kwargs):
            title=str(kwargs.get('title','')).casefold()
            return _fm474_wz_export_dir() if ('wissens' in title or 'wz' in title) else _fm474_standard_export_dir()
        def asksaveasfilename(*args, **kwargs):
            kwargs.setdefault('initialdir', _target(kwargs)); return _fd._fm474_orig_asksaveasfilename(*args, **kwargs)
        def asksaveasfile(*args, **kwargs):
            kwargs.setdefault('initialdir', _target(kwargs)); return _fd._fm474_orig_asksaveasfile(*args, **kwargs)
        def askdirectory(*args, **kwargs):
            title=str(kwargs.get('title','')).casefold()
            if any(x in title for x in ('export','ausgabe','speicher','ziel')): kwargs.setdefault('initialdir', _target(kwargs))
            return _fd._fm474_orig_askdirectory(*args, **kwargs)
        _fd.asksaveasfilename=asksaveasfilename; _fd.asksaveasfile=asksaveasfile; _fd.askdirectory=askdirectory
    except Exception: pass
_fm474_install_export_dialog_defaults()

def _fm474_kb_export_dir(self):
    return _fm474_wz_export_dir()

def _fm474_export_word(self):
    entry=self.kb_get_entry(getattr(self,'kb_selected_entry_id',None))
    if not entry: return
    try:
        from docx import Document
        import re as _re
        out_dir=self.kb_export_dir(); doc=Document(); doc.add_heading(entry.get('title','Wissenseintrag'), level=1)
        for label,val in [('Geändert', self.kb_display_date(entry.get('updated_at'))),('Benutzer',entry.get('user','')),('Status',entry.get('status','')),('Kategorien',', '.join(entry.get('categories',[]) or [])),('To-Do-Rhythmus',entry.get('rhythm',''))]:
            if val: doc.add_paragraph(f'{label}: {val}')
        doc.add_heading('Inhalt', level=2); doc.add_paragraph(entry.get('text',''))
        doc.add_heading('Anhänge', level=2)
        for a in entry.get('attachments',[]) or []: doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        doc.add_heading('Kommentare', level=2)
        for c in entry.get('comments',[]) or []: doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        safe=_re.sub(r'[^A-Za-z0-9_äöüÄÖÜß-]+','_',entry.get('title','Wissenseintrag'))[:80].strip('_') or 'Wissenseintrag'
        out=os.path.join(out_dir, f'{safe}.docx'); doc.save(out)
        try: self.file_output_register(out, user=getattr(self,'current_user_display','') or getattr(self,'current_user_key',''), module='Wissenszentrale', created_at=datetime.now().isoformat(timespec='seconds'))
        except Exception: pass
        messagebox.showinfo('Wissenszentrale','Word-Export erstellt:\n'+out)
    except Exception as exc:
        messagebox.showerror('Wissenszentrale','Word-Export fehlgeschlagen:\n'+str(exc))

try:
    FiBuMateApp.kb_export_dir = _fm474_kb_export_dir
    FiBuMateApp.kb_export_selected_to_word = _fm474_export_word
except Exception: pass

def _fm474_preview_table(canvas, rows, width, start_y=96, title='Tabellenvorschau'):
    y=start_y
    canvas.create_text(18,y,text=title,anchor='nw',fill=BLUE,font=body_font(12,weight='bold')); y+=30
    if not rows:
        canvas.create_text(18,y,text='Keine Tabellendaten gefunden.',anchor='nw',fill=TEXT2,font=body_font(10)); return y+40
    max_cols=min(8, max(1, max(len(r) for r in rows)))
    col_w=max(96, int(max(360,width-54)/max_cols)); row_h=30
    def trim(v):
        s='' if v is None else str(v).replace('\n',' ').replace('\r',' ').strip()
        n=max(6,int((col_w-12)/7)); return s if len(s)<=n else s[:n-1]+'…'
    for r_idx,row in enumerate(rows[:28]):
        bg='#EEF2F7' if r_idx==0 else ('#FFFFFF' if r_idx%2 else '#F8FAFC')
        y0=y+r_idx*row_h
        for c in range(max_cols):
            x0=18+c*col_w
            canvas.create_rectangle(x0,y0,x0+col_w,y0+row_h,fill=bg,outline='#E5E7EB')
            canvas.create_text(x0+6,y0+row_h/2,text=trim(row[c] if c<len(row) else ''),anchor='w',fill=TEXT,font=body_font(8,weight='bold' if r_idx==0 else None))
    return y+min(28,len(rows))*row_h+30

try: _fm473_preview_table = _fm474_preview_table
except Exception: pass

# Dateiausgabe-Render auf Basis v0.473 gezielt ersetzen: rechte Vorschau breiter, V als Überschrift entfernt, Buttons grau.
try:
    _fm474_old_render_file_output = FiBuMateApp.render_file_output
except Exception:
    _fm474_old_render_file_output = None

def _fm474_render_file_output(self):
    # v0.473-Renderer bleibt fachlich Basis; Konstanten/Roots/Vorschautabelle sind oben bereits überschrieben.
    return _fm474_old_render_file_output(self) if _fm474_old_render_file_output else None

# Nachgelagerte Tk-Widget-Korrektur für Dateiausgabe: wird nach Render ausgeführt.
def _fm474_render_file_output_wrapper(self):
    if _fm474_old_render_file_output:
        result=_fm474_old_render_file_output(self)
    else:
        result=None
    try:
        # Alle normalen Dateiausgabe-Buttons auf Grau angleichen; Export löschen bleibt rot.
        def walk(w):
            for ch in w.winfo_children():
                try:
                    if isinstance(ch, tk.Button):
                        txt=str(ch.cget('text'))
                        if 'löschen' in txt.casefold():
                            ch.configure(bg=RED, fg=WHITE, activebackground=darken(RED,0.08), activeforeground=WHITE)
                        else:
                            ch.configure(bg=_FM474_BTN_GREY, fg=TEXT, activebackground=_FM474_BTN_GREY_ACTIVE, activeforeground=TEXT)
                except Exception: pass
                walk(ch)
        walk(self.root)
    except Exception: pass
    return result

try: FiBuMateApp.render_file_output = _fm474_render_file_output_wrapper
except Exception: pass

# Benutzerverwaltung vollständig überschrieben: Wagnerm darf sich selbst bearbeiten und andere bis E4 stufen.
def _fm474_render_users_menu(self):
    zfont=self.zoomed_content_font
    if not self.can_view_user_management(): self.render_menu_text('Keine Berechtigung für die Benutzerverwaltung.'); return
    users=self.user_data.setdefault('users',{})
    visible=[self.current_user_key] if self.my_role()==ROLE_E1 and self.current_user_key in users else sorted(users.keys())
    w,h=self.canvas.winfo_width(),self.canvas.winfo_height(); top=132; vh=int(max(top+260,h-92)-top)
    cont=tk.Frame(self.root,bg=BG); self.widget_items.append(cont); self.canvas.create_window(0,top,window=cont,anchor='nw',width=w,height=vh)
    body=tk.Frame(cont,bg=BG); body.pack(fill='both',expand=True,padx=24)
    footer=tk.Frame(body,bg=BG); footer.pack(side='bottom',fill='x',pady=(8,0))
    cv=tk.Canvas(body,bg=BG,highlightthickness=0,bd=0); sb=tk.Scrollbar(body,orient='vertical',command=cv.yview)
    frame=tk.Frame(cv,bg=BG); win=cv.create_window((0,0),window=frame,anchor='nw'); cv.configure(yscrollcommand=sb.set); cv.pack(side='left',fill='both',expand=True); sb.pack(side='right',fill='y')
    try: self.register_scroll_canvas(cv,sb)
    except Exception: pass
    def upd(e=None):
        cv.itemconfigure(win,width=max(1,cv.winfo_width())); cv.configure(scrollregion=cv.bbox('all'))
    frame.bind('<Configure>',upd); cv.bind('<Configure>',upd); cv.bind('<MouseWheel>',lambda e:self._route_mousewheel_to_canvas(e,cv))
    tk.Label(frame,text='Benutzerverwaltung',font=zfont(('Segoe UI',15,'bold')),bg=BG,fg=TEXT).grid(row=0,column=0,columnspan=8,sticky='w',pady=(0,14))
    for c,head in enumerate(['Benutzer','Anzeigename/Nachname','Vorname','E-Mail','Rolle','Passwort','','']):
        tk.Label(frame,text=head,bg=BG,fg=TEXT,font=('Segoe UI',10,'bold')).grid(row=1,column=c,sticky='w',padx=(0,14))
    is_super=self.current_user_key==SUPERUSER_KEY
    def can_edit(k):
        if is_super: return True
        if k==SUPERUSER_KEY: return False
        if k==self.current_user_key: return True
        return self.my_role() in (ROLE_E3,ROLE_E4) and self.role_rank(users.get(k,{}).get('permission',ROLE_E1)) < self.role_rank(self.my_role())
    def role_vals(k):
        if is_super and k!=SUPERUSER_KEY: return ROLE_ORDER
        if k==SUPERUSER_KEY: return [ROLE_E4]
        return [r for r in ROLE_ORDER if self.role_rank(r)<=self.role_rank(self.my_role()) and r!=ROLE_E4]
    def save_row(k,nv,fv,ev,rv):
        if k not in users or not can_edit(k): return
        data=users[k]; new_name=' '.join(nv.get().strip().split()) or data.get('display_name',k); new_key=normalize_username(new_name)
        if k==SUPERUSER_KEY: new_key=SUPERUSER_KEY
        if is_super and k!=SUPERUSER_KEY and new_key!=k:
            if not new_key or new_key in users: messagebox.showwarning('FiBu Mate','Benutzername ungültig oder bereits vorhanden.'); return
            data=users.pop(k); users[new_key]=data
        data['display_name']=new_name; data['first_name']=fv.get().strip(); data['email']=ev.get().strip()
        data['full_name']=' '.join(x for x in [data.get('first_name',''),data.get('display_name',new_key)] if x).strip() or data.get('display_name',new_key)
        role=rv.get() or data.get('permission',ROLE_E1)
        if role not in role_vals(new_key): messagebox.showwarning('FiBu Mate','Diese Rolle darf nicht vergeben werden.'); return
        data['permission']=ROLE_E4 if new_key==SUPERUSER_KEY else role
        self.ensure_permissions_defaults(); self.save_user_data(); messagebox.showinfo('FiBu Mate','Benutzerdaten wurden gespeichert.'); self.render_page()
    def del_user(k):
        if k in (SUPERUSER_KEY,self.current_user_key) or not can_edit(k): messagebox.showwarning('FiBu Mate','Dieser Benutzer kann nicht gelöscht werden.'); return
        if messagebox.askyesno('Benutzer löschen',f"Benutzer wirklich löschen?\n\n{users.get(k,{}).get('display_name',k)}"):
            users.pop(k,None); self.save_user_data(); self.render_page()
    row=2
    for k in visible:
        u=users[k]; edit=can_edit(k); nv=tk.StringVar(value=u.get('display_name',k)); fv=tk.StringVar(value=u.get('first_name','')); ev=tk.StringVar(value=u.get('email','')); rv=tk.StringVar(value=u.get('permission',ROLE_E1))
        tk.Label(frame,text=k,bg=BG,fg=TEXT2,font=('Segoe UI',10)).grid(row=row,column=0,sticky='w',padx=(0,14),pady=4)
        tk.Entry(frame,textvariable=nv,font=('Segoe UI',10),width=22,bg='#E8EEF5',fg=TEXT,relief='solid',bd=1,state='normal' if edit and (is_super or k not in (self.current_user_key,SUPERUSER_KEY)) else 'disabled').grid(row=row,column=1,sticky='w',padx=(0,14),pady=4)
        tk.Entry(frame,textvariable=fv,font=('Segoe UI',10),width=18,bg='#E8EEF5',fg=TEXT,relief='solid',bd=1,state='normal' if edit else 'disabled').grid(row=row,column=2,sticky='w',padx=(0,14),pady=4)
        tk.Entry(frame,textvariable=ev,font=('Segoe UI',10),width=32,bg='#E8EEF5',fg=TEXT,relief='solid',bd=1,state='normal' if edit else 'disabled').grid(row=row,column=3,sticky='w',padx=(0,14),pady=4)
        ttk.Combobox(frame,textvariable=rv,values=role_vals(k),state='readonly' if edit and (is_super or k!=SUPERUSER_KEY) else 'disabled',font=('Segoe UI',10),width=24).grid(row=row,column=4,sticky='w',padx=(0,14),pady=4)
        tk.Label(frame,text='aktiv' if u.get('auth',{}).get('enabled') else 'nicht aktiv',bg=BG,fg=TEXT2,font=('Segoe UI',10)).grid(row=row,column=5,sticky='w',padx=(0,14),pady=4)
        tk.Button(frame,text='Speichern',command=lambda k=k,nv=nv,fv=fv,ev=ev,rv=rv:save_row(k,nv,fv,ev,rv),bg=BLUE if edit else GREY_DISABLED,fg='white',bd=0,width=9,padx=10,pady=5,state='normal' if edit else 'disabled').grid(row=row,column=6,sticky='w',padx=(18,10),pady=4)
        ok=edit and k not in (self.current_user_key,SUPERUSER_KEY) and self.my_role() in (ROLE_E3,ROLE_E4)
        tk.Button(frame,text='Löschen',command=lambda k=k:del_user(k),bg=RED if ok else GREY_DISABLED,fg='white',bd=0,width=9,padx=10,pady=5,state='normal' if ok else 'disabled').grid(row=row,column=7,sticky='w',padx=(0,8),pady=4)
        row+=1
    upd()

try: FiBuMateApp.render_users_menu = _fm474_render_users_menu
except Exception: pass



# ------------------------------------------------------------------
# FiBu Mate - Wissenszentrale-Löschen + Dateiausgabe-Farb/Iconfix FINAL 2026-06-23
# Version 0.475
# Zweck:
# - Wissenszentrale: Löschen entfernt Einträge endgültig aus zentraler und lokaler JSON statt über Merge wiederherzustellen.
# - Dateiausgabe: erste Spalte linksbündig, keine überlaufenden Punkte/Icons.
# - Dateiausgabe: Aktionsbuttons final einheitlich grau, nur Export löschen bleibt rot.
# ------------------------------------------------------------------
_FM475_PATCH_VERSION = "0.475"
_FM475_BUTTON_GREY = "#D6DCE4"
_FM475_BUTTON_GREY_ACTIVE = "#C8D0DA"
_FM475_DELETE_RED = RED
_FM475_TREE_ICON_COL_W = 24

try:
    _FM473_BUTTON_GREY = _FM475_BUTTON_GREY
    _FM473_BUTTON_GREY_ACTIVE = _FM475_BUTTON_GREY_ACTIVE
    _FM474_BTN_GREY = _FM475_BUTTON_GREY
    _FM474_BTN_GREY_ACTIVE = _FM475_BUTTON_GREY_ACTIVE
except Exception:
    pass


def _fm475_wz_entry_paths():
    paths = []
    candidates = []
    try: candidates.append(_wz441_entries_g())
    except Exception: pass
    try: candidates.append(_wz441_entries_l())
    except Exception: pass
    try: candidates.append(os.path.join(NETWORK_ROOT, 'Fibu_Mate_Doc', 'Database', 'Wissenszentrale', 'knowledge_entries.json'))
    except Exception: pass
    candidates.append(r'C:\python\knowledge_entries.json')
    for p in candidates:
        try:
            if p and p not in paths:
                paths.append(p)
        except Exception:
            pass
    return paths


def _fm475_write_entries_replace(entries):
    """Schreibt den übergebenen Eintragsbestand bewusst ersetzend.
    Wichtig fürs Löschen: Die bisherige Merge-Speicherlogik darf gelöschte Einträge nicht aus einer zweiten Quelle zurückholen.
    """
    ok = False
    errors = []
    payload = {'entries': list(entries or [])}
    for path in _fm475_wz_entry_paths():
        try:
            folder = os.path.dirname(path)
            if folder:
                os.makedirs(folder, exist_ok=True)
            tmp = path + '.tmp'
            with open(tmp, 'w', encoding='utf-8') as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            os.replace(tmp, path)
            ok = True
        except Exception as exc:
            errors.append(f'{path}: {exc}')
    try:
        if '_WZ441_ENTRIES_CACHE' in globals():
            _WZ441_ENTRIES_CACHE['ts'] = 0.0
            _WZ441_ENTRIES_CACHE['sig'] = ''
            _WZ441_ENTRIES_CACHE['entries'] = list(entries or [])
    except Exception:
        pass
    if not ok:
        raise RuntimeError('\n'.join(errors) or 'Keine Wissenszentrale-Speicherdatei konnte geschrieben werden.')
    return True


def _fm475_delete_entry(self, entry_id=None):
    entry_id = entry_id or getattr(self, 'kb_selected_entry_id', None)
    if not entry_id:
        try: messagebox.showwarning('Wissenszentrale', 'Bitte zuerst einen Eintrag auswählen.')
        except Exception: pass
        return False
    try:
        if not _fm471_can_delete_entry(self):
            messagebox.showwarning('Wissenszentrale', 'Keine Berechtigung zum Löschen dieses Eintrags.')
            return False
    except Exception:
        pass
    entries = self.kb_load_entries()
    target = None
    remaining = []
    for e in entries:
        if e.get('id') == entry_id:
            target = e
        else:
            remaining.append(e)
    if not target:
        try: messagebox.showwarning('Wissenszentrale', 'Der ausgewählte Eintrag wurde nicht gefunden.')
        except Exception: pass
        return False
    try:
        title = str(target.get('title') or 'Wissenseintrag')
        if not messagebox.askyesno('Eintrag löschen', f'Eintrag wirklich löschen?\n\n{title}'):
            return False
    except Exception:
        return False
    try:
        _fm475_write_entries_replace(remaining)
        try:
            if hasattr(self, 'kb_autosave_entry_id') and self.kb_autosave_entry_id == entry_id:
                self.kb_autosave_entry_id = None
            self.kb_selected_entry_id = None
            self.kb_edit_entry_id = None
            self.knowledge_view = 'all'
            self.knowledge_unsaved = False
        except Exception:
            pass
        try: messagebox.showinfo('Wissenszentrale', 'Eintrag wurde gelöscht.')
        except Exception: pass
        self.render_page()
        return True
    except Exception as exc:
        try: messagebox.showerror('Wissenszentrale', 'Eintrag konnte nicht gelöscht werden:\n' + str(exc))
        except Exception: pass
        return False


def _fm475_delete_button(parent, self, entry_id=None):
    try:
        can_delete = _fm471_can_delete_entry(self)
    except Exception:
        can_delete = False
    return tk.Button(parent, text=f'{_FM471_DELETE_SYMBOL} Löschen', command=(lambda eid=entry_id: _fm475_delete_entry(self, eid)) if can_delete else (lambda: None),
                     bg=('#FEE2E2' if can_delete else '#E5E7EB'), fg=(TEXT if can_delete else '#7A7F87'),
                     activebackground=('#FECACA' if can_delete else '#E5E7EB'), font=body_font(9, weight='bold'), relief='solid', bd=1,
                     state=('normal' if can_delete else 'disabled'), cursor=('hand2' if can_delete else 'arrow'))

try:
    _fm471_delete_entry = _fm475_delete_entry
    _fm471_delete_button = _fm475_delete_button
    FiBuMateApp.kb_delete_entry = _fm475_delete_entry
except Exception:
    pass


def _fm475_normalize_dateiausgabe_buttons(root):
    """Setzt im aktuell gerenderten Dateiausgabe-Bereich alle Aktionsbuttons konsistent.
    Export löschen bleibt bewusst rot; alle anderen Buttons werden grau.
    """
    def walk(widget):
        for child in widget.winfo_children():
            try:
                if isinstance(child, tk.Button):
                    txt = str(child.cget('text') or '')
                    if 'löschen' in txt.casefold():
                        child.configure(bg=_FM475_DELETE_RED, fg=WHITE, activebackground=darken(_FM475_DELETE_RED, 0.08), activeforeground=WHITE,
                                        relief='solid', bd=1, highlightthickness=0)
                    else:
                        child.configure(bg=_FM475_BUTTON_GREY, fg=TEXT, activebackground=_FM475_BUTTON_GREY_ACTIVE, activeforeground=TEXT,
                                        relief='solid', bd=1, highlightthickness=0)
            except Exception:
                pass
            walk(child)
    try:
        walk(root)
    except Exception:
        pass


def _fm475_normalize_treeviews(root):
    def walk(widget):
        for child in widget.winfo_children():
            try:
                if isinstance(child, ttk.Treeview):
                    try: child.heading('#0', text='')
                    except Exception: pass
                    try: child.column('#0', width=_FM475_TREE_ICON_COL_W, minwidth=_FM475_TREE_ICON_COL_W, stretch=False, anchor='w')
                    except Exception: pass
            except Exception:
                pass
            walk(child)
    try:
        walk(root)
    except Exception:
        pass

try:
    _fm475_previous_render_file_output = FiBuMateApp.render_file_output
except Exception:
    _fm475_previous_render_file_output = None


def _fm475_render_file_output(self):
    result = None
    if _fm475_previous_render_file_output:
        result = _fm475_previous_render_file_output(self)
    try:
        _fm475_normalize_dateiausgabe_buttons(self.root)
        _fm475_normalize_treeviews(self.root)
        try:
            self.root.after(50, lambda: (_fm475_normalize_dateiausgabe_buttons(self.root), _fm475_normalize_treeviews(self.root)))
        except Exception:
            pass
    except Exception:
        pass
    return result

try:
    FiBuMateApp.render_file_output = _fm475_render_file_output
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Dateiausgabe echte Dokumentvorschau + Teams-Popup FINAL 2026-06-23
# Version 0.476
# ------------------------------------------------------------------
_FM476_PATCH_VERSION = "0.476"
_FM476_PREVIEW_CACHE_DIR = os.path.join(os.environ.get('LOCALAPPDATA') or os.path.join(os.path.expanduser('~'), 'AppData', 'Local'), 'FiBuMate', 'PreviewCache')


def _fm476_msg(canvas, y, width, title, lines, color=TEXT2):
    canvas.create_text(18, y, text=title, anchor='nw', fill=BLUE if color != RED else RED, font=body_font(12, weight='bold'), width=max(280, width-45)); y += 30
    for line in lines:
        canvas.create_text(18, y, text=str(line), anchor='nw', fill=color, font=body_font(10), width=max(280, width-45)); y += 24
    return max(70, y - 96 + 10)


def _fm476_cache_pdf_for_word(path):
    try: os.makedirs(_FM476_PREVIEW_CACHE_DIR, exist_ok=True)
    except Exception: pass
    try:
        st=os.stat(path); key=hashlib.sha1((os.path.abspath(path)+str(st.st_mtime_ns)+str(st.st_size)).encode('utf-8',errors='ignore')).hexdigest()
    except Exception:
        key=hashlib.sha1(os.path.abspath(str(path)).encode('utf-8',errors='ignore')).hexdigest()
    out=os.path.join(_FM476_PREVIEW_CACHE_DIR, key+'.pdf')
    if os.path.exists(out) and os.path.getsize(out)>0: return out
    # Primär Windows/Word COM
    if sys.platform.startswith('win'):
        word=None; doc=None
        try:
            import win32com.client
            word=win32com.client.DispatchEx('Word.Application'); word.Visible=False; word.DisplayAlerts=0
            doc=word.Documents.Open(os.path.abspath(path), ReadOnly=True, AddToRecentFiles=False)
            doc.ExportAsFixedFormat(OutputFileName=os.path.abspath(out), ExportFormat=17, OpenAfterExport=False)
            try: doc.Close(False)
            except Exception: pass
            try: word.Quit()
            except Exception: pass
            if os.path.exists(out) and os.path.getsize(out)>0: return out
        except Exception:
            try:
                if doc: doc.Close(False)
            except Exception: pass
            try:
                if word: word.Quit()
            except Exception: pass
    # Fallback LibreOffice
    try:
        import subprocess, shutil
        exe=shutil.which('soffice') or shutil.which('libreoffice')
        if exe:
            subprocess.run([exe,'--headless','--convert-to','pdf','--outdir',_FM476_PREVIEW_CACHE_DIR,os.path.abspath(path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=45)
            cand=os.path.join(_FM476_PREVIEW_CACHE_DIR, os.path.splitext(os.path.basename(path))[0]+'.pdf')
            if os.path.exists(cand) and os.path.getsize(cand)>0:
                try: os.replace(cand,out)
                except Exception: out=cand
                return out
    except Exception: pass
    return ''


def _fm476_render_pdf(canvas, path, width, title):
    if not PIL_AVAILABLE:
        return _fm476_msg(canvas,96,width,title,['Pillow ist nicht verfügbar; Vorschau kann nicht angezeigt werden.'],RED)
    try:
        import fitz
        doc=fitz.open(path)
        if len(doc)<1:
            try: doc.close()
            except Exception: pass
            return _fm476_msg(canvas,96,width,title,['Das Dokument enthält keine Seite.'],RED)
        page=doc.load_page(0); max_w=max(260,int(width-46)); scale=min(2.2,max(0.7,max_w/max(1.0,float(page.rect.width))))
        pix=page.get_pixmap(matrix=fitz.Matrix(scale,scale), alpha=False)
        try: doc.close()
        except Exception: pass
        img=Image.frombytes('RGB',[pix.width,pix.height],pix.samples)
        if img.width>max_w:
            r=max_w/float(img.width); img=img.resize((max(1,int(img.width*r)), max(1,int(img.height*r))))
        photo=ImageTk.PhotoImage(img); refs=getattr(canvas,'_fm472_preview_refs',[]) or []; refs.append(photo); canvas._fm472_preview_refs=refs
        y=96; canvas.create_text(18,y,text=title,anchor='nw',fill=BLUE,font=body_font(12,weight='bold')); y+=32
        # Kein zusätzlicher Kasten/Rahmen: nur die gerenderte Seite.
        canvas.create_image(18,y,image=photo,anchor='nw')
        return 32+img.height+28
    except Exception as exc:
        return _fm476_msg(canvas,96,width,title,['Echte Vorschau konnte nicht gerendert werden.','Renderer benötigt: PyMuPDF/fitz.',f'Technische Ursache: {exc}'],RED)


def _fm476_preview_pdf(canvas, path, width):
    return _fm476_render_pdf(canvas, path, width, 'PDF-Vorschau')


def _fm476_preview_office(canvas, path, width):
    ext=os.path.splitext(path)[1].lower()
    if ext in {'.doc','.docx'}:
        pdf=_fm476_cache_pdf_for_word(path)
        if pdf: return _fm476_render_pdf(canvas, pdf, width, 'Word-Vorschau')
        return _fm476_msg(canvas,96,width,'Word-Vorschau',['Echte Word-Vorschau konnte nicht erzeugt werden.','Benötigt wird Microsoft Word COM oder LibreOffice.','Die frühere Metadaten-/Kasten-Vorschau wurde entfernt.'],RED)
    return _fm476_msg(canvas,96,width,'Dokumentvorschau',['Für diesen Office-Dateityp ist noch kein Renderer hinterlegt.'],RED)

try:
    _fm472_preview_pdf = _fm476_preview_pdf
    _fm472_preview_office = _fm476_preview_office
except Exception: pass


def _fm476_teams_preview_popup(self, item):
    files=_fm473_files_for_item(item)
    if not files:
        try: messagebox.showwarning('Dateiausgabe','Keine versendbare Datei gefunden.')
        except Exception: pass
        return False
    win=tk.Toplevel(self.root); win.title('Teams Versand vorbereiten'); win.configure(bg=BG)
    try: win.transient(self.root); win.grab_set(); win.geometry('760x560')
    except Exception: pass
    outer=tk.Frame(win,bg=BG); outer.pack(fill='both',expand=True,padx=18,pady=16)
    tk.Label(outer,text='Teams Versand – Vorschau',bg=BG,fg=BLUE,font=body_font(15,weight='bold')).pack(anchor='w')
    tk.Label(outer,text='Keine Zwischenablage: Kontakt und Nachricht erfassen, danach Teams-Chat vorbereiten.',bg=BG,fg=TEXT2,font=body_font(10),wraplength=700,justify='left').pack(anchor='w',pady=(4,12))
    tk.Label(outer,text='Kontakt / Teams-E-Mail',bg=BG,fg=TEXT,font=body_font(10,weight='bold')).pack(anchor='w')
    contact_var=tk.StringVar(value=''); contact=tk.Entry(outer,textvariable=contact_var,bg=WHITE,fg=TEXT,font=body_font(11),relief='solid',bd=1); contact.pack(fill='x',ipady=5,pady=(2,10))
    tk.Label(outer,text='Nachricht',bg=BG,fg=TEXT,font=body_font(10,weight='bold')).pack(anchor='w')
    msg=tk.Text(outer,height=7,bg=WHITE,fg=TEXT,font=body_font(10),relief='solid',bd=1,wrap='word')
    msg.insert('1.0','Hallo,\n\nanbei die folgende Datei bzw. die folgenden Dateien:\n\n'+'\n'.join('- '+os.path.basename(f) for f in files)+'\n\nViele Grüße')
    msg.pack(fill='x',pady=(2,12))
    tk.Label(outer,text='Anhang / Dateien',bg=BG,fg=TEXT,font=body_font(10,weight='bold')).pack(anchor='w')
    lf=tk.Frame(outer,bg=WHITE,highlightbackground=LINE,highlightthickness=1); lf.pack(fill='both',expand=True,pady=(2,12))
    lb=tk.Listbox(lf,bg=WHITE,fg=TEXT,font=body_font(9),relief='flat'); sb=tk.Scrollbar(lf,orient='vertical',command=lb.yview); lb.configure(yscrollcommand=sb.set)
    lb.pack(side='left',fill='both',expand=True); sb.pack(side='right',fill='y')
    for f in files: lb.insert('end',f)
    bar=tk.Frame(outer,bg=BG); bar.pack(fill='x')
    def picked():
        try:
            s=lb.curselection(); return lb.get(s[0]) if s else files[0]
        except Exception: return files[0]
    def open_teams():
        users=','.join([x.strip() for x in re.split(r'[;,]',contact_var.get().strip()) if x.strip()])
        if not users:
            try: messagebox.showwarning('Teams Versand','Bitte Kontakt bzw. Teams-E-Mail-Adresse angeben.')
            except Exception: pass
            return
        message_text=msg.get('1.0','end-1c').strip()+'\n\nDatei(en):\n'+'\n'.join(os.path.basename(f) for f in files)
        url='https://teams.microsoft.com/l/chat/0/0?users='+quote(users)+'&message='+quote(message_text)
        try:
            webbrowser.open(url); _fm473_mark_sent(self,item,'Teams')
            try: win.destroy()
            except Exception: pass
            try: self.render_page()
            except Exception: pass
        except Exception as exc:
            try: messagebox.showerror('Teams Versand','Teams-Vorschau konnte nicht geöffnet werden:\n'+str(exc))
            except Exception: pass
    tk.Button(bar,text='Datei öffnen',command=lambda:_fm472_open_file({'path':picked()}),bg='#D6DCE4',fg=TEXT,activebackground='#C8D0DA',font=body_font(10),relief='solid',bd=1).pack(side='left',padx=(0,8),ipadx=12,ipady=5)
    tk.Button(bar,text='Speicherort öffnen',command=lambda:_fm472_open_folder({'path':picked()}),bg='#D6DCE4',fg=TEXT,activebackground='#C8D0DA',font=body_font(10),relief='solid',bd=1).pack(side='left',padx=(0,8),ipadx=12,ipady=5)
    tk.Button(bar,text='Teams öffnen',command=open_teams,bg='#D6DCE4',fg=TEXT,activebackground='#C8D0DA',font=body_font(10,weight='bold'),relief='solid',bd=1).pack(side='right',padx=(8,0),ipadx=14,ipady=5)
    tk.Button(bar,text='Abbrechen',command=win.destroy,bg='#D6DCE4',fg=TEXT,activebackground='#C8D0DA',font=body_font(10),relief='solid',bd=1).pack(side='right',padx=(8,0),ipadx=14,ipady=5)
    try: contact.focus_set()
    except Exception: pass
    return True

try:
    _fm473_send_teams = _fm476_teams_preview_popup
except Exception: pass



# ------------------------------------------------------------------
# FiBu Mate - Dateiausgabe Doppelklick sperren + Iconspalte links FINAL 2026-06-23
# Version 0.477
# Zweck:
# - Dateiausgabe: Dateien/Sammelpositionen werden nicht mehr per Doppelklick geöffnet.
# - Dateiausgabe: erste Spalte wird als schmale, linksbündige Status-/Iconspalte normalisiert.
# ------------------------------------------------------------------
_FM477_PATCH_VERSION = "0.477"
_FM477_ICON_COL_W = 22


def _fm477_finalize_file_output_treeviews(root):
    """Finaler UI-Fix für alle Treeviews im aktuell gerenderten Dateiausgabe-Bereich.
    Doppelklick wird auf Widget-Ebene komplett abgefangen; bestehende ältere on_open-Bindings werden überschrieben.
    Die erste Tree-Spalte wird sehr schmal und linksbündig geführt, damit Status-Icons/Punkte nicht in Folgespalten laufen.
    """
    try:
        style = ttk.Style(root)
        try:
            style.configure('Treeview', indent=0)
        except Exception:
            pass
    except Exception:
        pass

    def walk(widget):
        try:
            children = widget.winfo_children()
        except Exception:
            children = []
        for child in children:
            try:
                if isinstance(child, ttk.Treeview):
                    try:
                        child.unbind('<Double-Button-1>')
                        child.bind('<Double-Button-1>', lambda event: 'break')
                    except Exception:
                        pass
                    try:
                        child.heading('#0', text='', anchor='w')
                    except Exception:
                        pass
                    try:
                        child.column('#0', width=_FM477_ICON_COL_W, minwidth=_FM477_ICON_COL_W, stretch=False, anchor='w')
                    except Exception:
                        pass
                    try:
                        # Punkte bei normalen Dateien entfernen; Gruppen behalten nur den kleinen Aufklapphinweis.
                        for iid in child.get_children(''):
                            has_children = bool(child.get_children(iid))
                            child.item(iid, text=('▸' if has_children else ''))
                            for cid in child.get_children(iid):
                                child.item(cid, text='')
                    except Exception:
                        pass
            except Exception:
                pass
            walk(child)
    walk(root)

try:
    _fm477_previous_render_file_output = FiBuMateApp.render_file_output
except Exception:
    _fm477_previous_render_file_output = None


def _fm477_render_file_output(self):
    result = None
    if _fm477_previous_render_file_output:
        result = _fm477_previous_render_file_output(self)
    try:
        _fm477_finalize_file_output_treeviews(self.root)
        try:
            self.root.after(50, lambda: _fm477_finalize_file_output_treeviews(self.root))
            self.root.after(250, lambda: _fm477_finalize_file_output_treeviews(self.root))
        except Exception:
            pass
    except Exception:
        pass
    return result

try:
    FiBuMateApp.render_file_output = _fm477_render_file_output
except Exception:
    pass



# ------------------------------------------------------------------
# FiBu Mate - Dateiausgabe Log-Only Exporte + Screenshot-Iconspalte FINAL 2026-07-06
# Version 0.478
# Zweck:
# - Dateiausgabe listet nur noch registrierte FiBu-Mate-Exporte aus dem Dateiausgabe-Log.
# - Nachtraeglich in Exportordner kopierte/verschobene Dateien werden nicht mehr per Ordnerscan aufgenommen.
# - Erste Treeview-Spalte wird passend zum Screenshot breiter und linksbuendig gehalten.
# ------------------------------------------------------------------
_FM478_PATCH_VERSION = "0.478"
_FM478_ICON_COL_W = 44


def _fm478_registered_export_item(self, item):
    try:
        if not isinstance(item, dict):
            return None
        p = item.get('path') or ''
        if not p:
            return None
        ap = os.path.abspath(str(p))
        if not os.path.isfile(ap):
            return None
        st = os.stat(ap)
        out = dict(item)
        out['id'] = out.get('id') or _fm472_file_key(ap)
        out['type'] = 'file'
        out['name'] = out.get('name') or os.path.basename(ap)
        out['path'] = ap
        out['user'] = out.get('user') or _fm472_safe_user(self)
        out['module'] = out.get('module') or _fm472_guess_module(ap)
        out['created_at'] = out.get('created_at') or datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds')
        out['group_id'] = out.get('group_id', '') or ''
        out['group_title'] = out.get('group_title', '') or ''
        out['sent_at'] = out.get('sent_at', '') or ''
        out['sent_channel'] = out.get('sent_channel', '') or ''
        out['created_by_fibu_mate'] = True
        return out
    except Exception:
        return None


def _fm478_scan_output_files(self):
    """Nur registrierte FiBu-Mate-Exporte; kein Ordnerscan von Exportpfaden."""
    items = []
    seen = set()
    try:
        logs = _fm473_load_all_log_items()
    except Exception:
        logs = {}
    for raw in logs.values():
        item = _fm478_registered_export_item(self, raw)
        if not item:
            continue
        key = os.path.abspath(item.get('path', '')).casefold()
        if not key or key in seen:
            continue
        seen.add(key)
        items.append(item)
    items.sort(key=lambda x: str(x.get('created_at', '')), reverse=True)
    return items


def _fm478_build_items(self):
    return _fm473_build_groups(self, _fm478_scan_output_files(self))


def _fm478_register_output_file(self, path, user=None, module=None, group_id=None, group_title=None, created_at=None, children=None):
    ok = False
    try:
        ok = _fm473_register_output_file(self, path, user=user, module=module, group_id=group_id, group_title=group_title, created_at=created_at, children=children)
        if ok and path:
            logs = _fm473_load_all_log_items()
            iid = _fm472_file_key(path)
            entry = logs.get(iid, {})
            entry['id'] = iid
            entry['type'] = 'file'
            entry['name'] = entry.get('name') or os.path.basename(str(path))
            entry['path'] = os.path.abspath(str(path))
            entry['user'] = entry.get('user') or user or _fm472_safe_user(self)
            entry['module'] = entry.get('module') or module or _fm472_guess_module(path)
            entry['group_id'] = entry.get('group_id') or group_id or ''
            entry['group_title'] = entry.get('group_title') or group_title or ''
            entry['created_at'] = entry.get('created_at') or created_at or _fm472_now_iso()
            entry['created_by_fibu_mate'] = True
            logs[iid] = entry
            _fm473_write_log_items(logs)
        return bool(ok)
    except Exception:
        return bool(ok)


def _fm478_finalize_file_output_treeviews(root):
    def walk(widget):
        try:
            children = widget.winfo_children()
        except Exception:
            children = []
        for child in children:
            try:
                if isinstance(child, ttk.Treeview):
                    try:
                        child.unbind('<Double-Button-1>')
                        child.bind('<Double-Button-1>', lambda event: 'break')
                    except Exception:
                        pass
                    try:
                        child.heading('#0', text='', anchor='w')
                    except Exception:
                        pass
                    try:
                        child.column('#0', width=_FM478_ICON_COL_W, minwidth=_FM478_ICON_COL_W, stretch=False, anchor='w')
                    except Exception:
                        pass
                    try:
                        for iid in child.get_children(''):
                            has_children = bool(child.get_children(iid))
                            child.item(iid, text=('▸' if has_children else ''))
                            for cid in child.get_children(iid):
                                child.item(cid, text='')
                    except Exception:
                        pass
            except Exception:
                pass
            walk(child)
    walk(root)

try:
    _fm478_previous_render_file_output = FiBuMateApp.render_file_output
except Exception:
    _fm478_previous_render_file_output = None


def _fm478_render_file_output(self):
    result = None
    if _fm478_previous_render_file_output:
        result = _fm478_previous_render_file_output(self)
    try:
        _fm478_finalize_file_output_treeviews(self.root)
        try:
            self.root.after(25, lambda: _fm478_finalize_file_output_treeviews(self.root))
            self.root.after(125, lambda: _fm478_finalize_file_output_treeviews(self.root))
            self.root.after(350, lambda: _fm478_finalize_file_output_treeviews(self.root))
            self.root.after(750, lambda: _fm478_finalize_file_output_treeviews(self.root))
        except Exception:
            pass
    except Exception:
        pass
    return result

try:
    _fm473_scan_output_files = _fm478_scan_output_files
    FiBuMateApp.file_output_build_items = _fm478_build_items
    FiBuMateApp.file_output_register = _fm478_register_output_file
    FiBuMateApp.render_file_output = _fm478_render_file_output
except Exception:
    pass



# ------------------------------------------------------------------
# FIBU_MATE_HELP_SANDBOX_TILE_MAINMENU_UI_V0482
# Datum: 2026-07-08
# Zweck:
# - Hilfe-Button global sichtbar in Menues/Modulen.
# - Hilfe-Sandbox je aktuellem Menue/Modul/Tool mit Leitfaeden und Videos.
# - Kacheldesign moderner mit leichtem 3D-Effekt, Hover bleibt aktiv.
# - Globale Schriftfamilie moderner ohne Groessenaenderung.
# - Hauptmenue-Abschnitte subtil beschriften: Module / Tools.
# ------------------------------------------------------------------
FIBU_MATE_HELP_SANDBOX_TILE_MAINMENU_UI_VERSION = "0.482"
FM_MODERN_FONT_FAMILY = "Aptos"
FM_HELP_SANDBOX_FILE = os.path.join(CENTRAL_CONFIG_DIR, "fibu_mate_help_sandbox.json")

try:
    from tkinter import filedialog as _fm482_filedialog
except Exception:
    _fm482_filedialog = None

# Moderne globale Schriftfamilie, Groessen bleiben unveraendert.
try:
    FONT_TITLE = (FM_MODERN_FONT_FAMILY, FONT_TITLE[1], *FONT_TITLE[2:])
    FONT_MENU = (FM_MODERN_FONT_FAMILY, FONT_MENU[1], *FONT_MENU[2:])
    FONT_TILE = (FM_MODERN_FONT_FAMILY, FONT_TILE[1], *FONT_TILE[2:])
    FONT_TILE_SMALL = (FM_MODERN_FONT_FAMILY, FONT_TILE_SMALL[1], *FONT_TILE_SMALL[2:])
    FONT_SMALL = (FM_MODERN_FONT_FAMILY, FONT_SMALL[1], *FONT_SMALL[2:])
    BASE_FONT_TITLE = FONT_TITLE
    BASE_FONT_MENU = FONT_MENU
    BASE_FONT_TILE = FONT_TILE
    BASE_FONT_TILE_SMALL = FONT_TILE_SMALL
    BASE_FONT_SMALL = FONT_SMALL
except Exception:
    pass

_body_font_before_v0482 = body_font if 'body_font' in globals() else None

def body_font(size=10, weight=None, underline=False, scale=1.0):
    if _body_font_before_v0482:
        base = _body_font_before_v0482(size, weight=weight, underline=underline, scale=scale)
        try:
            return tuple([FM_MODERN_FONT_FAMILY, base[1]] + list(base[2:]))
        except Exception:
            return base
    style = []
    if weight:
        style.append(weight)
    if underline:
        style.append('underline')
    return tuple([FM_MODERN_FONT_FAMILY, size] + style)


def _fm482_load_help_sandbox():
    try:
        if os.path.exists(FM_HELP_SANDBOX_FILE):
            with open(FM_HELP_SANDBOX_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
    except Exception:
        pass
    return {}


def _fm482_save_help_sandbox(data):
    try:
        os.makedirs(os.path.dirname(FM_HELP_SANDBOX_FILE), exist_ok=True)
        with open(FM_HELP_SANDBOX_FILE, 'w', encoding='utf-8') as f:
            json.dump(data or {}, f, ensure_ascii=False, indent=2)
        return True
    except Exception as exc:
        try:
            messagebox.showerror('Hilfe', f'Hilfe-Sandbox konnte nicht gespeichert werden:\n\n{exc}')
        except Exception:
            pass
        return False


def _fm482_help_context(self):
    page = getattr(self, 'current_page', '') or 'main'
    title = getattr(self, 'current_title', '') or 'FiBu Mate'
    key = page
    if page.startswith('tool:'):
        key = page
    return key, title


def _fm482_help_entry_for(self, key, title):
    data = _fm482_load_help_sandbox()
    entry = data.get(key)
    if not isinstance(entry, dict):
        entry = {'title': title, 'guide_text': '', 'guides': [], 'videos': [], 'updated_at': ''}
    entry.setdefault('title', title)
    entry.setdefault('guide_text', '')
    entry.setdefault('guides', [])
    entry.setdefault('videos', [])
    entry.setdefault('updated_at', '')
    return data, entry


def _fm482_open_help_target(value):
    value = str(value or '').strip()
    if not value:
        return
    try:
        if re.match(r'^[a-zA-Z][a-zA-Z0-9+.-]*://', value):
            webbrowser.open(value)
        elif os.path.exists(value):
            os.startfile(value) if hasattr(os, 'startfile') else webbrowser.open(value)
        else:
            webbrowser.open(value)
    except Exception as exc:
        try:
            messagebox.showerror('Hilfe', f'Der Eintrag konnte nicht geöffnet werden:\n\n{value}\n\n{exc}')
        except Exception:
            pass


def _fm482_lines_from_text(widget):
    try:
        return [line.strip() for line in widget.get('1.0', 'end').splitlines() if line.strip()]
    except Exception:
        return []


def _fm482_insert_lines(widget, lines):
    try:
        widget.delete('1.0', 'end')
        widget.insert('1.0', '\n'.join(lines or []))
    except Exception:
        pass


def _fm482_show_help_popup(self):
    key, title = _fm482_help_context(self)
    data, entry = _fm482_help_entry_for(self, key, title)
    popup = tk.Toplevel(self.root)
    popup.title(f'FiBu Mate - Hilfe: {title}')
    popup.configure(bg=BG)
    popup.transient(self.root)
    popup.resizable(True, True)
    popup.minsize(880, 620)
    popup_w, popup_h = 1120, 760
    try:
        self.root.update_idletasks()
        screen_w = self.root.winfo_screenwidth(); screen_h = self.root.winfo_screenheight()
        x = max(0, int((screen_w - popup_w) / 2)); y = max(0, int((screen_h - popup_h) / 2))
        popup.geometry(f'{popup_w}x{popup_h}+{x}+{y}')
    except Exception:
        popup.geometry(f'{popup_w}x{popup_h}')
    try:
        self._make_modal(popup)
    except Exception:
        pass

    header = tk.Frame(popup, bg=HEADER, height=64); header.pack(side='top', fill='x'); header.pack_propagate(False)
    tk.Label(header, text='Hilfe-Sandbox', bg=HEADER, fg=TEXT, font=(FM_MODERN_FONT_FAMILY, 18, 'bold')).pack(side='left', padx=18)
    tk.Label(header, text=f'Aktueller Kontext: {title}', bg=HEADER, fg=TEXT2, font=body_font(10, 'bold')).pack(side='left', padx=(10, 0))
    tk.Label(header, text=f'ID: {key}', bg=HEADER, fg=TEXT2, font=body_font(8)).pack(side='right', padx=18)

    outer = tk.Frame(popup, bg=BG); outer.pack(side='top', fill='both', expand=True, padx=16, pady=14)
    outer.columnconfigure(0, weight=2); outer.columnconfigure(1, weight=1); outer.rowconfigure(0, weight=1)

    left = tk.Frame(outer, bg=WHITE, bd=1, relief='solid'); left.grid(row=0, column=0, sticky='nsew', padx=(0, 10))
    right = tk.Frame(outer, bg=WHITE, bd=1, relief='solid'); right.grid(row=0, column=1, sticky='nsew')
    left.rowconfigure(1, weight=1); left.columnconfigure(0, weight=1)
    right.rowconfigure(2, weight=1); right.rowconfigure(5, weight=1); right.columnconfigure(0, weight=1)

    tk.Label(left, text='Leitfaden / individuelle Hilfestellung', bg=WHITE, fg=BLUE, font=body_font(12, 'bold')).grid(row=0, column=0, sticky='w', padx=14, pady=(12, 6))
    guide_text = tk.Text(left, bg='#F8FAFC', fg=TEXT, wrap='word', font=body_font(10), relief='solid', bd=1)
    guide_text.grid(row=1, column=0, sticky='nsew', padx=14, pady=(0, 10))
    guide_text.insert('1.0', entry.get('guide_text', ''))
    hint = 'Sandbox: Text frei pflegen, Leitfäden als Dateien/Pfade und Erklärvideos als Links oder Dateien hinterlegen. Diese Inhalte gelten nur für den aktuellen Kontext.'
    tk.Label(left, text=hint, bg=WHITE, fg=TEXT2, font=body_font(9), wraplength=650, justify='left').grid(row=2, column=0, sticky='ew', padx=14, pady=(0, 12))

    tk.Label(right, text='Leitfäden / Dateien', bg=WHITE, fg=BLUE, font=body_font(11, 'bold')).grid(row=0, column=0, sticky='w', padx=12, pady=(12, 4))
    guides_text = tk.Text(right, height=8, bg='#F8FAFC', fg=TEXT, wrap='none', font=body_font(9), relief='solid', bd=1)
    guides_text.grid(row=2, column=0, sticky='nsew', padx=12, pady=(0, 8))
    _fm482_insert_lines(guides_text, entry.get('guides', []))
    guide_btns = tk.Frame(right, bg=WHITE); guide_btns.grid(row=1, column=0, sticky='ew', padx=12, pady=(0, 4))
    def add_guide():
        if not _fm482_filedialog: return
        path = _fm482_filedialog.askopenfilename(title='Leitfaden auswählen', filetypes=[('Dokumente/Videos', '*.pdf *.docx *.xlsx *.pptx *.mp4 *.mov *.avi *.mkv'), ('Alle Dateien', '*.*')])
        if path:
            guides_text.insert('end', ('\n' if guides_text.get('1.0','end').strip() else '') + path)
    def open_guides_first():
        lines = _fm482_lines_from_text(guides_text)
        if lines: _fm482_open_help_target(lines[0])
    tk.Button(guide_btns, text='Datei hinzufügen', command=add_guide, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6))
    tk.Button(guide_btns, text='Ersten öffnen', command=open_guides_first, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left')

    tk.Label(right, text='Erklärvideos / Links', bg=WHITE, fg=BLUE, font=body_font(11, 'bold')).grid(row=3, column=0, sticky='w', padx=12, pady=(8, 4))
    video_btns = tk.Frame(right, bg=WHITE); video_btns.grid(row=4, column=0, sticky='ew', padx=12, pady=(0, 4))
    videos_text = tk.Text(right, height=8, bg='#F8FAFC', fg=TEXT, wrap='none', font=body_font(9), relief='solid', bd=1)
    videos_text.grid(row=5, column=0, sticky='nsew', padx=12, pady=(0, 12))
    _fm482_insert_lines(videos_text, entry.get('videos', []))
    def add_video_file():
        if not _fm482_filedialog: return
        path = _fm482_filedialog.askopenfilename(title='Erklärvideo auswählen', filetypes=[('Videos', '*.mp4 *.mov *.avi *.mkv *.wmv'), ('Alle Dateien', '*.*')])
        if path:
            videos_text.insert('end', ('\n' if videos_text.get('1.0','end').strip() else '') + path)
    def open_video_first():
        lines = _fm482_lines_from_text(videos_text)
        if lines: _fm482_open_help_target(lines[0])
    tk.Button(video_btns, text='Video-Datei', command=add_video_file, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6))
    tk.Button(video_btns, text='Erstes öffnen', command=open_video_first, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left')

    footer = tk.Frame(popup, bg=BG); footer.pack(side='bottom', fill='x', padx=16, pady=(0, 16))
    def save_help():
        data[key] = {
            'title': title,
            'guide_text': guide_text.get('1.0', 'end').strip(),
            'guides': _fm482_lines_from_text(guides_text),
            'videos': _fm482_lines_from_text(videos_text),
            'updated_at': datetime.now().isoformat(timespec='seconds'),
        }
        if _fm482_save_help_sandbox(data):
            try: messagebox.showinfo('Hilfe', 'Hilfe-Sandbox gespeichert.')
            except Exception: pass
    tk.Button(footer, text='Speichern', command=save_help, bg='#CFEAD6', fg=TEXT, font=body_font(10, 'bold'), bd=1, relief='solid', padx=ui_s(22), pady=ui_s(8), cursor='hand2').pack(side='right', padx=(8, 0))
    tk.Button(footer, text='Schließen', command=popup.destroy, bg=BLUE, fg='white', font=body_font(10, 'bold'), bd=0, padx=ui_s(22), pady=ui_s(9), cursor='hand2').pack(side='right')
    tk.Label(footer, text=f'Speicherort: {FM_HELP_SANDBOX_FILE}', bg=BG, fg=TEXT2, font=body_font(8)).pack(side='left')


def _fm482_draw_help_control(self):
    # Hilfe soll in jedem Menue/Modul sichtbar sein. Nur dort, wo draw_controls nicht aufgerufen wird, erscheint sie naturgemaess nicht.
    btn = tk.Canvas(self.root, width=MINI_WIDGET_W, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor='hand2'); btn._zoom_exclude = True
    btn.create_rectangle(1, 1, MINI_WIDGET_W - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
    self.draw_mini_help_icon(btn, 25, MINI_WIDGET_H / 2)
    btn.create_text(MINI_WIDGET_W / 2 + ui_s(6), MINI_WIDGET_H / 2, text='Hilfe', fill=TEXT, font=(FM_MODERN_FONT_FAMILY, max(8, ui_s(9)), 'bold'), anchor='center')
    btn.bind('<Button-1>', lambda _event=None: self.show_help_popup())
    btn.bind('<Enter>', lambda _e: self.show_small_tooltip(btn, 'Hilfe-Sandbox'))
    btn.bind('<Leave>', lambda _e: self.hide_small_tooltip())
    self.widget_items.append(btn)
    self.canvas.create_window(self.canvas.winfo_width() - ui_s(22), ui_s(8), window=btn, anchor='ne')

try:
    FiBuMateApp.show_help_popup = _fm482_show_help_popup
    FiBuMateApp.draw_help_control = _fm482_draw_help_control
except Exception:
    pass


# --- Modernisiertes Kacheldesign mit leichtem 3D-Effekt ---

def _fm482_tile_draw(self):
    self.delete('all')
    pad = ui_s(8)
    off = ui_s(1) if self.pressed else 0
    x0, y0 = pad + off, pad + off
    x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
    base = self.current_color()
    # weicher 3D-Schatten und dezente Tiefe; Icon-/Textpositionen bleiben technisch identisch.
    if not self.pressed:
        self.create_rectangle(pad + 8, pad + 10, pad + 8 + self.tile_width, pad + 10 + self.tile_height, fill=blend(SHADOW, '#1F2937', 0.12), outline='')
        self.create_rectangle(pad + 4, pad + 5, pad + 4 + self.tile_width, pad + 5 + self.tile_height, fill=blend(SHADOW, WHITE, 0.20), outline='')
    self.create_rectangle(x0, y0, x1, y1, fill=base, outline=darken(base, 0.18), width=1)
    # Lichtkante oben/links + dunklere Kante unten/rechts fuer kleinen 3D-Effekt.
    self.create_line(x0 + 2, y0 + 2, x1 - 2, y0 + 2, fill=blend(base, WHITE, 0.34), width=2)
    self.create_line(x0 + 2, y0 + 2, x0 + 2, y1 - 2, fill=blend(base, WHITE, 0.20), width=1)
    self.create_line(x0 + 2, y1 - 2, x1 - 2, y1 - 2, fill=darken(base, 0.18), width=2)
    self.create_line(x1 - 2, y0 + 2, x1 - 2, y1 - 2, fill=darken(base, 0.12), width=1)
    if self.hovered and self.enabled if hasattr(self, 'enabled') else self.hovered:
        self.create_rectangle(x0 + 4, y0 + 4, x1 - 4, y1 - 4, outline=blend(WHITE, base, 0.25), width=1)
    if self.corner_fold:
        self._draw_corner_fold(x1, y0)
    title_y = y0 + ui_s(14)
    title_color = BLUE if self.lock_tile else 'white'
    icon_to_draw = 'lock' if self.lock_tile else self.icon_type
    text_width = max(ui_s(110), self.tile_width - ui_s(44))
    if self.center_text and not icon_to_draw:
        centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
        self.create_text((x0 + x1) / 2, (y0 + y1) / 2, text=self.title, anchor='center', fill=title_color, font=centered_font, width=text_width, justify='center')
    else:
        title_max_h = max(ui_s(28), int((y1 - y0) * 0.36))
        title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
        self.create_text((x0 + x1) / 2, title_y, text=self.title, anchor='n', fill=title_color, font=title_font, width=text_width, justify='center')
        if icon_to_draw:
            min_icon_y = title_y + title_h + ui_s(20)
            default_icon_y = y0 + int((y1 - y0) * 0.66)
            icon_y = max(default_icon_y, min_icon_y)
            icon_y = min(icon_y, y1 - ui_s(34))
            close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, 'current_close_period_label') else ''
            if close_period and self.tile_id in ('monthly_close', 'quarterly_close', 'yearly_close'):
                icon_x = x0 + 98
                self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                self.create_text(icon_x + 48, icon_y, text=close_period, anchor='w', fill=title_color, font=self.app.zoomed_content_font((FM_MODERN_FONT_FAMILY, 13, 'italic')), width=max(120, x1 - icon_x - 58), justify='left')
            else:
                self._draw_main_icon(icon_to_draw, (x0 + x1) / 2, icon_y)
    self.star_bounds = None
    if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
        sx, sy = x1 - 24, y0 + 24
        self.star_bounds = (sx - 18, sy - 18, sx + 18, sy + 18)
        self.create_text(sx, sy, text='★', fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=(FM_MODERN_FONT_FAMILY, 24, 'bold'))

try:
    Tile.draw = _fm482_tile_draw
except Exception:
    pass


# --- Hauptmenue mit Abschnittsnamen Module / Tools ---

def _fm482_draw_section_label(self, text, x, y):
    try:
        self.canvas.create_rectangle(x - ui_s(8), y - ui_s(12), x + ui_s(132), y + ui_s(12), fill=blend(BG, WHITE, 0.40), outline=blend(BG, BLUE, 0.10))
        self.canvas.create_text(x, y, text=text, fill=blend(TEXT2, BLUE, 0.35), font=body_font(9, 'bold'), anchor='w')
    except Exception:
        pass


def _fm482_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w * 0.165))); th = max(125, min(170, int(h * 0.155)))
    gap_x = max(38, int(w * 0.04))
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count * tile_w + (count - 1) * gap
        start_x = (w - total_width) / 2 + tile_w / 2
        return [start_x + i * (tile_w + gap) for i in range(count)]
    def create_tile_group(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')
    # Bestehende Positionen bleiben erhalten.
    top_tile_w = max(250, int(tw*1.05)); top_tile_h = max(130, int(th*1.05))
    tool_tile_w = max(200, int(tw*0.86*0.95)); tool_tile_h = max(100, int(th*0.86*0.95))
    top_y = y_pct(h, 64); tools_y = y_pct(h, 43)
    mini_tw = int(tw * 0.72); mini_th = int(th * 0.72)
    mini_center_y = y_pct(h, 16.5) + th / 2 - mini_th / 2
    mini_top = mini_center_y - mini_th / 2
    left_x = x_pct(w, 9.5) + ui_s(2)
    _fm482_draw_section_label(self, 'Module', left_x, top_y - top_tile_h / 2 - ui_s(18))
    create_tile_group(top_tiles, top_y, 'main_menuzeile_1', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, (top_y + tools_y) / 2, x_pct(w, 9.5), x_pct(w, 92))
    _fm482_draw_section_label(self, 'Tools', left_x, tools_y - tool_tile_h / 2 - ui_s(18))
    create_tile_group(middle_tiles, tools_y, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    create_tile_group(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top - 13, x_pct(w, 9.5), x_pct(w, 92))
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_main_menu = _fm482_render_main_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_HELP_ROLE_DELETEICON_TILELABEL_FIX_V0483
# Datum: 2026-07-08
# Zweck:
# - Kachel nur noch mit einem Schatten, kein weisser Hover-Rand.
# - Abschnittstitel im Hauptmenue frei auf dem Hintergrund links oben im Abschnitt.
# - Globales Papierkorb-Icon fuer Loeschen-Buttons.
# - Hilfe-Sandbox nur fuer E4; E3 und niedriger erhalten reine Anzeigeansicht.
# ------------------------------------------------------------------
FIBU_MATE_HELP_ROLE_DELETEICON_TILELABEL_FIX_VERSION = "0.483"
FM_DELETE_ICON_PATH = r"C:\python\bin\Imgs\Icons\biggarbagebin_121980.ico"


def _fm483_is_e4(self):
    try:
        return self.my_role() == ROLE_E4
    except Exception:
        return False


def _fm483_help_readonly_text(entry, title, key):
    parts = [f"Hilfe - {title}", ""]
    guide = (entry or {}).get('guide_text', '').strip()
    guides = (entry or {}).get('guides', []) or []
    videos = (entry or {}).get('videos', []) or []
    if guide:
        parts.append(guide)
    else:
        parts.append('Für dieses Menü/Modul wurde noch keine Hilfestellung hinterlegt.')
    if guides:
        parts.extend(['', 'Leitfäden / Dokumente:'])
        parts.extend([f"- {g}" for g in guides if str(g).strip()])
    if videos:
        parts.extend(['', 'Erklärvideos / Links:'])
        parts.extend([f"- {v}" for v in videos if str(v).strip()])
    parts.extend(['', f'Kontext: {key}'])
    return '\n'.join(parts)


def _fm483_show_help_viewer(self, key, title, entry):
    popup = tk.Toplevel(self.root)
    popup.title(f'FiBu Mate - Hilfe: {title}')
    popup.configure(bg=BG)
    popup.transient(self.root)
    popup.resizable(True, True)
    popup.minsize(780, 520)
    popup_w, popup_h = 1000, 700
    try:
        self.root.update_idletasks()
        screen_w = self.root.winfo_screenwidth(); screen_h = self.root.winfo_screenheight()
        x = max(0, int((screen_w - popup_w) / 2)); y = max(0, int((screen_h - popup_h) / 2))
        popup.geometry(f'{popup_w}x{popup_h}+{x}+{y}')
    except Exception:
        popup.geometry(f'{popup_w}x{popup_h}')
    try:
        self._make_modal(popup)
    except Exception:
        pass
    # Eine einzige Anzeigeansicht: der Inhalt fuellt das Fenster.
    txt = tk.Text(popup, bg=WHITE, fg=TEXT, wrap='word', font=body_font(11), relief='flat', bd=0, padx=18, pady=18)
    txt.pack(fill='both', expand=True)
    txt.insert('1.0', _fm483_help_readonly_text(entry, title, key))
    txt.configure(state='disabled')
    try:
        txt.bind('<Double-Button-1>', lambda _e: _fm482_open_help_target(txt.get('insert linestart', 'insert lineend').replace('- ', '').strip()))
    except Exception:
        pass
    footer = tk.Frame(popup, bg=BG)
    footer.pack(fill='x')
    tk.Label(footer, text='Tipp: Dokument-/Videopfade per Doppelklick auf die jeweilige Zeile öffnen.', bg=BG, fg=TEXT2, font=body_font(8)).pack(side='left', padx=12, pady=6)
    tk.Button(footer, text='Schließen', command=popup.destroy, bg=BLUE, fg='white', font=body_font(9, 'bold'), bd=0, padx=14, pady=5).pack(side='right', padx=12, pady=6)


def _fm483_show_help_popup(self):
    key, title = _fm482_help_context(self) if '_fm482_help_context' in globals() else (getattr(self, 'current_page', 'main'), getattr(self, 'current_title', 'FiBu Mate'))
    data, entry = _fm482_help_entry_for(self, key, title) if '_fm482_help_entry_for' in globals() else ({}, {'title': title, 'guide_text': '', 'guides': [], 'videos': []})
    if _fm483_is_e4(self):
        return _fm482_show_help_popup(self) if '_fm482_show_help_popup' in globals() else None
    return _fm483_show_help_viewer(self, key, title, entry)

try:
    FiBuMateApp.show_help_popup = _fm483_show_help_popup
except Exception:
    pass


# --- Ein-Schatten-Kacheldesign ohne weissen Hover-Rand ---

def _fm483_tile_draw(self):
    self.delete('all')
    pad = ui_s(8)
    off = ui_s(1) if self.pressed else 0
    x0, y0 = pad + off, pad + off
    x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
    base = self.current_color()
    # Genau ein Schatten.
    if not self.pressed:
        self.create_rectangle(pad + 7, pad + 8, pad + 7 + self.tile_width, pad + 8 + self.tile_height, fill=blend(SHADOW, '#1F2937', 0.16), outline='')
    self.create_rectangle(x0, y0, x1, y1, fill=base, outline=darken(base, 0.20), width=1)
    # Dezenter 3D-Effekt ueber Kanten, aber kein weisser Hover-Rahmen.
    self.create_line(x0 + 2, y0 + 2, x1 - 2, y0 + 2, fill=blend(base, WHITE, 0.26), width=1)
    self.create_line(x0 + 2, y0 + 2, x0 + 2, y1 - 2, fill=blend(base, WHITE, 0.14), width=1)
    self.create_line(x0 + 2, y1 - 2, x1 - 2, y1 - 2, fill=darken(base, 0.18), width=2)
    self.create_line(x1 - 2, y0 + 2, x1 - 2, y1 - 2, fill=darken(base, 0.12), width=1)
    if self.corner_fold:
        self._draw_corner_fold(x1, y0)
    title_y = y0 + ui_s(14)
    title_color = BLUE if self.lock_tile else 'white'
    icon_to_draw = 'lock' if self.lock_tile else self.icon_type
    text_width = max(ui_s(110), self.tile_width - ui_s(44))
    if self.center_text and not icon_to_draw:
        centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
        self.create_text((x0 + x1) / 2, (y0 + y1) / 2, text=self.title, anchor='center', fill=title_color, font=centered_font, width=text_width, justify='center')
    else:
        title_max_h = max(ui_s(28), int((y1 - y0) * 0.36))
        title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
        self.create_text((x0 + x1) / 2, title_y, text=self.title, anchor='n', fill=title_color, font=title_font, width=text_width, justify='center')
        if icon_to_draw:
            min_icon_y = title_y + title_h + ui_s(20)
            default_icon_y = y0 + int((y1 - y0) * 0.66)
            icon_y = max(default_icon_y, min_icon_y)
            icon_y = min(icon_y, y1 - ui_s(34))
            close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, 'current_close_period_label') else ''
            if close_period and self.tile_id in ('monthly_close', 'quarterly_close', 'yearly_close'):
                icon_x = x0 + 98
                self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                self.create_text(icon_x + 48, icon_y, text=close_period, anchor='w', fill=title_color, font=self.app.zoomed_content_font((FM_MODERN_FONT_FAMILY, 13, 'italic')), width=max(120, x1 - icon_x - 58), justify='left')
            else:
                self._draw_main_icon(icon_to_draw, (x0 + x1) / 2, icon_y)
    self.star_bounds = None
    if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
        sx, sy = x1 - 24, y0 + 24
        self.star_bounds = (sx - 18, sy - 18, sx + 18, sy + 18)
        self.create_text(sx, sy, text='★', fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=(FM_MODERN_FONT_FAMILY, 24, 'bold'))

try:
    Tile.draw = _fm483_tile_draw
except Exception:
    pass


# --- Abschnittstitel frei auf dem Hauptmenue-Hintergrund an linker Abschnittskante ---

def _fm483_draw_section_label(self, text, x, y):
    try:
        self.canvas.create_text(x, y, text=text, fill=blend(TEXT2, BLUE, 0.30), font=body_font(9, 'bold'), anchor='w')
    except Exception:
        pass


def _fm483_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w * 0.165))); th = max(125, min(170, int(h * 0.155)))
    gap_x = max(38, int(w * 0.04))
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count * tile_w + (count - 1) * gap
        start_x = (w - total_width) / 2 + tile_w / 2
        return [start_x + i * (tile_w + gap) for i in range(count)]
    def create_tile_group(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')
    top_tile_w = max(250, int(tw*1.05)); top_tile_h = max(130, int(th*1.05))
    tool_tile_w = max(200, int(tw*0.86*0.95)); tool_tile_h = max(100, int(th*0.86*0.95))
    top_y = y_pct(h, 64); tools_y = y_pct(h, 43)
    mini_tw = int(tw * 0.72); mini_th = int(th * 0.72)
    mini_center_y = y_pct(h, 16.5) + th / 2 - mini_th / 2
    mini_top = mini_center_y - mini_th / 2
    # Abschnittstitel am linken Rand des Hintergrund-Abschnitts, ohne Textbox.
    label_x = ui_s(10)
    _fm483_draw_section_label(self, 'Module', label_x, ui_s(154))
    create_tile_group(top_tiles, top_y, 'main_menuzeile_1', top_tile_w, top_tile_h)
    first_sep_y = (top_y + tools_y) / 2
    self.draw_continuous_relief_line(self.canvas, first_sep_y, x_pct(w, 9.5), x_pct(w, 92))
    _fm483_draw_section_label(self, 'Tools', label_x, first_sep_y + ui_s(28))
    create_tile_group(middle_tiles, tools_y, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    create_tile_group(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top - 13, x_pct(w, 9.5), x_pct(w, 92))
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_main_menu = _fm483_render_main_menu
except Exception:
    pass


# --- Globales Loeschen-Icon ---
_FM483_ORIG_TK_BUTTON_INIT = tk.Button.__init__
_FM483_ORIG_TK_BUTTON_CONFIGURE = tk.Button.configure
_FM483_DELETE_ICON_CACHE = {}


def _fm483_delete_button_should_icon(widget):
    try:
        txt = str(widget.cget('text') or '')
        return 'Löschen' in txt or 'Loeschen' in txt or 'löschen' in txt.lower()
    except Exception:
        return False


def _fm483_get_delete_photo(widget):
    if not PIL_AVAILABLE:
        return None
    try:
        path = FM_DELETE_ICON_PATH
        if not os.path.exists(path):
            path = os.path.join(ICON_DIR, 'biggarbagebin_121980.ico')
        key = (str(path), ui_s(18))
        img = Image.open(path).resize((ui_s(18), ui_s(18)))
        photo = ImageTk.PhotoImage(img)
        _FM483_DELETE_ICON_CACHE[id(widget)] = photo
        return photo
    except Exception:
        return None


def _fm483_apply_delete_icon(widget):
    try:
        if not _fm483_delete_button_should_icon(widget):
            return
        photo = _fm483_get_delete_photo(widget)
        if photo:
            _FM483_ORIG_TK_BUTTON_CONFIGURE(widget, None, image=photo, compound='left')
            widget._fm483_delete_icon = photo
    except Exception:
        pass


def _fm483_button_init(self, master=None, cnf={}, **kw):
    _FM483_ORIG_TK_BUTTON_INIT(self, master, cnf, **kw)
    _fm483_apply_delete_icon(self)


def _fm483_button_configure(self, cnf=None, **kw):
    result = _FM483_ORIG_TK_BUTTON_CONFIGURE(self, cnf, **kw)
    try:
        if ('text' in kw) or _fm483_delete_button_should_icon(self):
            _fm483_apply_delete_icon(self)
    except Exception:
        pass
    return result

try:
    tk.Button.__init__ = _fm483_button_init
    tk.Button.configure = _fm483_button_configure
    tk.Button.config = _fm483_button_configure
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_HEADER_LINGO_MAINMENU_REWORK_V0484
# Datum: 2026-07-08
# Zweck:
# - Hauptmenue-Kacheln linksbuendig zur Trennlinienkante ausrichten, Groesse/Abstand beibehalten.
# - Header reduzieren, Breadcrumb/Favoriten unterhalb der Buttons positionieren und farblich saettigen.
# - Favoriten-Chips gleich gross und ueber E4-Menue "Lingo" abkuerzbar.
# - Hilfe rollenabhaengig: E4 Sandbox, <=E3 reine Anzeige.
# - Loeschen-Icon global auf Papierkorb anwenden.
# - Kachel: genau ein Schatten, kein weisser Hover-Rand.
# ------------------------------------------------------------------
FIBU_MATE_HEADER_LINGO_MAINMENU_REWORK_VERSION = "0.484"
FM_HEADER_H = ui_s(96)
FM_BAR_Y = ui_s(62)
FM_DELETE_ICON_PATH = r"C:\python\bin\Imgs\Icons\biggarbagebin_121980.ico"

try:
    from tkinter import filedialog as _fm484_filedialog
except Exception:
    _fm484_filedialog = None

# ---------- Header / Breadcrumb / Favoriten ----------
def _fm484_draw_background(self):
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    self.canvas.create_rectangle(0, 0, w, h, fill=BG, outline='')
    for i in range(10):
        y = 145 + i * max(52, h * 0.060)
        color = blend(BG, BLUE if i % 2 == 0 else RED, 0.040 + (i % 4) * 0.008)
        self.canvas.create_line(-80, y, w + 90, y + h * 0.075, fill=color, width=1)
    self.canvas.create_oval(w * 0.68, h * 0.18, w * 1.05, h * 0.78, outline=blend(BG, BLUE, 0.11), width=2)
    self.canvas.create_oval(w * 0.73, h * 0.24, w * 1.01, h * 0.70, outline=blend(BG, WHITE, 0.46), width=1)
    self.canvas.create_oval(-w * 0.14, h * 0.42, w * 0.25, h * 1.02, outline=blend(BG, RED, 0.075), width=2)
    self.canvas.create_polygon(w * 0.02, h * 0.90, w * 0.32, h * 0.72, w * 0.60, h * 1.04, fill=blend(BG, WHITE, 0.20), outline='')
    self.canvas.create_polygon(w * 0.62, h * 0.20, w * 0.96, h * 0.36, w * 1.06, h * 0.12, fill=blend(BG, BLUE, 0.040), outline='')
    self.canvas.create_polygon(w * 0.08, h * 0.18, w * 0.22, h * 0.26, w * 0.04, h * 0.34, fill=blend(BG, WHITE, 0.15), outline='')
    self.canvas.create_rectangle(0, 0, w, FM_HEADER_H, fill=HEADER, outline='')
    self.canvas.create_rectangle(0, FM_HEADER_H, w, FM_HEADER_H + ui_s(3), fill=LINE, outline='')


def _fm484_draw_header(self, title):
    w = self.canvas.winfo_width()
    self.canvas.create_rectangle(0, 0, w, FM_HEADER_H, fill=HEADER, outline='')
    self.canvas.create_line(0, FM_HEADER_H, w, FM_HEADER_H, fill=LINE, width=1)
    # Titel bleibt horizontal zentriert, wird aber im reduzierten Header sichtbar nach oben gezogen.
    self.canvas.create_text(w / 2, ui_s(83), text=title, font=FONT_MENU, fill=TEXT2, anchor='center')


def _fm484_draw_gradient_line(self, x1, x2, y):
    width = max(1, int(x2 - x1)); steps = max(20, min(180, width // 5))
    for i in range(steps):
        t = i / max(1, steps - 1)
        # Saettiger als bisher, damit Leisten gegen den Hintergrund klarer stehen.
        color = blend(RED, BLUE, t)
        self.canvas.create_line(x1 + width * i / steps, y, x1 + width * (i + 1) / steps, y, fill=color, width=2)


def _fm484_bar_width(self):
    w = self.canvas.winfo_width()
    return max(ui_s(440), min(ui_s(690), int(w * 0.34)))


def _fm484_draw_path_bar(self):
    y_mid = FM_BAR_Y
    x1 = ui_s(10)
    x2 = min(x1 + _fm484_bar_width(self), self.canvas.winfo_width() * 0.48)
    self.draw_gradient_line(x1, x2, y_mid - ui_s(10))
    self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
    x = x1 + ui_s(12)
    for idx, (page, title) in enumerate(self.breadcrumb):
        current = idx == len(self.breadcrumb) - 1
        tid = self.canvas.create_text(x, y_mid, text=title, font=body_font(9, 'bold') if current else body_font(9), fill=TEXT if current else BLUE, anchor='w')
        bbox = self.canvas.bbox(tid); tw = bbox[2] - bbox[0] if bbox else ui_s(70)
        if not current:
            self.canvas.tag_bind(tid, '<Button-1>', lambda e, i=idx: self.jump_to_breadcrumb(i))
            self.canvas.tag_bind(tid, '<Enter>', lambda e: self.canvas.config(cursor='hand2'))
            self.canvas.tag_bind(tid, '<Leave>', lambda e: self.canvas.config(cursor='arrow'))
        x += tw + ui_s(12)
        if idx < len(self.breadcrumb) - 1:
            self.canvas.create_polygon(x, y_mid - ui_s(5), x, y_mid + ui_s(5), x + ui_s(7), y_mid, fill=RED, outline=RED)
            x += ui_s(17)


def _fm484_lingo_map(self):
    try:
        settings = self.user_data.setdefault('settings', {})
        m = settings.setdefault('lingo_shortcuts', {})
        return m if isinstance(m, dict) else {}
    except Exception:
        return {}


def _fm484_lingo_label(self, tool_id):
    m = _fm484_lingo_map(self)
    if tool_id in m and str(m[tool_id]).strip():
        return str(m[tool_id]).strip()[:8]
    label = TOOL_REGISTRY.get(tool_id, {}).get('favorite_label', tool_id)
    # Default kuerzen, damit alle Favoriten-Chips gleich gross bleiben.
    return str(label).strip()[:8]


def _fm484_draw_favorites_bar(self):
    w = self.canvas.winfo_width(); y_mid = FM_BAR_Y
    bar_w = _fm484_bar_width(self)
    x2 = w - ui_s(10)
    x1 = max(w * 0.52, x2 - bar_w)
    self.draw_gradient_line(x1, x2, y_mid - ui_s(10))
    self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
    self.canvas.create_text(x1 + ui_s(8), y_mid, text='★', font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', max(10, ui_s(13)), 'bold'), fill=GOLD, anchor='w')
    self.canvas.create_text(x1 + ui_s(30), y_mid, text='Favoriten', font=body_font(9, 'bold'), fill=TEXT, anchor='w')
    chip_w = ui_s(58); chip_gap = ui_s(4); chip_h = ui_s(20)
    x = x1 + ui_s(108)
    max_x = x2 - ui_s(4)
    for fav in sorted(f for f in self.favorites if f in TOOL_REGISTRY and f not in HIDDEN_TOOL_IDS):
        if x + chip_w > max_x:
            break
        label = _fm484_lingo_label(self, fav)
        chip_color = self.current_tile_color() or BLUE
        chip = tk.Label(self.root, text=label, bg=chip_color, fg='white', font=body_font(8, 'bold'), width=7, anchor='center', padx=0, pady=0, cursor='hand2')
        chip._zoom_exclude = True
        chip.bind('<Button-1>', lambda e, fid=fav: self.execute_favorite(fid))
        self.widget_items.append(chip)
        self.canvas.create_window(x, y_mid + ui_s(8), window=chip, anchor='w', width=chip_w, height=chip_h)
        x += chip_w + chip_gap

try:
    FiBuMateApp.draw_background = _fm484_draw_background
    FiBuMateApp.draw_header = _fm484_draw_header
    FiBuMateApp.draw_gradient_line = _fm484_draw_gradient_line
    FiBuMateApp.draw_path_bar = _fm484_draw_path_bar
    FiBuMateApp.draw_favorites_bar = _fm484_draw_favorites_bar
except Exception:
    pass

# ---------- Dynamischere, gleich grosse Trennlinien ----------
def _fm484_draw_continuous_relief_line(self, canvas, y, x1, x2):
    # gleiche Groesse/Wirkungshoehe, aber feiner Verlauf statt harter Strich
    steps = 72
    for i in range(steps):
        t = i / max(1, steps - 1)
        fade = 0.42 + 0.38 * (1 - abs(t - 0.5) * 2)
        sx = x1 + (x2 - x1) * i / steps
        ex = x1 + (x2 - x1) * (i + 1) / steps
        canvas.create_line(sx, y, ex, y, fill=blend(BG, '#1F2933', 0.22 * fade), width=2)
        canvas.create_line(sx, y + 2, ex, y + 2, fill=blend(BG, WHITE, 0.60 * fade), width=1)

try:
    FiBuMateApp.draw_continuous_relief_line = _fm484_draw_continuous_relief_line
except Exception:
    pass

# ---------- Hilfe: E4 Sandbox, <=E3 Anzeige ----------
def _fm484_is_e4(self):
    try:
        return self.my_role() == ROLE_E4
    except Exception:
        return False


def _fm484_help_readonly_text(entry, title, key):
    parts = [f'Hilfe - {title}', '']
    guide = (entry or {}).get('guide_text', '').strip()
    guides = (entry or {}).get('guides', []) or []
    videos = (entry or {}).get('videos', []) or []
    parts.append(guide if guide else 'Für dieses Menü/Modul wurde noch keine Hilfestellung hinterlegt.')
    if guides:
        parts.extend(['', 'Leitfäden / Dokumente:'])
        parts.extend([f'- {g}' for g in guides if str(g).strip()])
    if videos:
        parts.extend(['', 'Erklärvideos / Links:'])
        parts.extend([f'- {v}' for v in videos if str(v).strip()])
    parts.extend(['', f'Kontext: {key}'])
    return '\n'.join(parts)


def _fm484_show_help_viewer(self, key, title, entry):
    popup = tk.Toplevel(self.root)
    popup.title(f'FiBu Mate - Hilfe: {title}')
    popup.configure(bg=BG); popup.transient(self.root); popup.resizable(True, True); popup.minsize(780, 520)
    try:
        self.root.update_idletasks(); sw = self.root.winfo_screenwidth(); sh = self.root.winfo_screenheight(); pw, ph = 1000, 700
        popup.geometry(f'{pw}x{ph}+{max(0,int((sw-pw)/2))}+{max(0,int((sh-ph)/2))}')
    except Exception:
        popup.geometry('1000x700')
    try: self._make_modal(popup)
    except Exception: pass
    txt = tk.Text(popup, bg=WHITE, fg=TEXT, wrap='word', font=body_font(11), relief='flat', bd=0, padx=18, pady=18)
    txt.pack(fill='both', expand=True)
    txt.insert('1.0', _fm484_help_readonly_text(entry, title, key))
    txt.configure(state='disabled')
    try:
        txt.bind('<Double-Button-1>', lambda _e: _fm482_open_help_target(txt.get('insert linestart', 'insert lineend').replace('- ', '').strip()) if '_fm482_open_help_target' in globals() else None)
    except Exception:
        pass


def _fm484_show_help_popup(self):
    key, title = _fm482_help_context(self) if '_fm482_help_context' in globals() else (getattr(self, 'current_page', 'main'), getattr(self, 'current_title', 'FiBu Mate'))
    data, entry = _fm482_help_entry_for(self, key, title) if '_fm482_help_entry_for' in globals() else ({}, {'title': title, 'guide_text': '', 'guides': [], 'videos': []})
    if _fm484_is_e4(self):
        return _fm482_show_help_popup(self) if '_fm482_show_help_popup' in globals() else None
    return _fm484_show_help_viewer(self, key, title, entry)

try:
    FiBuMateApp.show_help_popup = _fm484_show_help_popup
except Exception:
    pass

# ---------- Kachel: ein Schatten, kein weisser Hover-Rand ----------
def _fm484_tile_draw(self):
    self.delete('all')
    pad = ui_s(8); off = ui_s(1) if self.pressed else 0
    x0, y0 = pad + off, pad + off
    x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
    base = self.current_color()
    if not self.pressed:
        self.create_rectangle(pad + 7, pad + 8, pad + 7 + self.tile_width, pad + 8 + self.tile_height, fill=blend(SHADOW, '#1F2937', 0.16), outline='')
    self.create_rectangle(x0, y0, x1, y1, fill=base, outline=darken(base, 0.20), width=1)
    self.create_line(x0 + 2, y0 + 2, x1 - 2, y0 + 2, fill=blend(base, WHITE, 0.26), width=1)
    self.create_line(x0 + 2, y0 + 2, x0 + 2, y1 - 2, fill=blend(base, WHITE, 0.14), width=1)
    self.create_line(x0 + 2, y1 - 2, x1 - 2, y1 - 2, fill=darken(base, 0.18), width=2)
    self.create_line(x1 - 2, y0 + 2, x1 - 2, y1 - 2, fill=darken(base, 0.12), width=1)
    if self.corner_fold:
        self._draw_corner_fold(x1, y0)
    title_y = y0 + ui_s(14)
    title_color = BLUE if self.lock_tile else 'white'
    icon_to_draw = 'lock' if self.lock_tile else self.icon_type
    text_width = max(ui_s(110), self.tile_width - ui_s(44))
    if self.center_text and not icon_to_draw:
        centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
        self.create_text((x0+x1)/2, (y0+y1)/2, text=self.title, anchor='center', fill=title_color, font=centered_font, width=text_width, justify='center')
    else:
        title_max_h = max(ui_s(28), int((y1-y0)*0.36))
        title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
        self.create_text((x0+x1)/2, title_y, text=self.title, anchor='n', fill=title_color, font=title_font, width=text_width, justify='center')
        if icon_to_draw:
            min_icon_y = title_y + title_h + ui_s(20)
            default_icon_y = y0 + int((y1-y0)*0.66)
            icon_y = max(default_icon_y, min_icon_y); icon_y = min(icon_y, y1-ui_s(34))
            close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, 'current_close_period_label') else ''
            if close_period and self.tile_id in ('monthly_close','quarterly_close','yearly_close'):
                icon_x = x0 + 98
                self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                self.create_text(icon_x+48, icon_y, text=close_period, anchor='w', fill=title_color, font=self.app.zoomed_content_font((FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', 13, 'italic')), width=max(120, x1-icon_x-58), justify='left')
            else:
                self._draw_main_icon(icon_to_draw, (x0+x1)/2, icon_y)
    self.star_bounds = None
    if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
        sx, sy = x1 - 24, y0 + 24
        self.star_bounds = (sx-18, sy-18, sx+18, sy+18)
        self.create_text(sx, sy, text='★', fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', 24, 'bold'))

try:
    Tile.draw = _fm484_tile_draw
except Exception:
    pass

# ---------- Hauptmenue linksbuendig zu Trennlinien ----------
def _fm484_draw_section_label(self, text, x, y):
    try:
        # frei auf Hintergrund, keine Textbox
        self.canvas.create_text(x, y, text=text, fill=blend(TEXT2, BLUE, 0.30), font=body_font(9, 'bold'), anchor='w')
    except Exception:
        pass


def _fm484_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165))); th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04))
    line_x1 = x_pct(w, 9.5); line_x2 = x_pct(w, 92)
    def create_tile_row(items, left_edge, y, id_prefix, tile_w, tile_h):
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            cx = left_edge + tile_w/2 + i*(tile_w+gap_x)
            self.canvas.create_window(cx, y, window=tile, anchor='center')
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count*tile_w + (count-1)*gap
        start_x = (w-total_width)/2 + tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]
    def create_centered_row(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i,item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')
    top_tile_w=max(250,int(tw*1.05)); top_tile_h=max(130,int(th*1.05))
    tool_tile_w=max(200,int(tw*0.86*0.95)); tool_tile_h=max(100,int(th*0.86*0.95))
    # Gruppen bleiben gleich gross/abstaendig, werden aber linksbuendig zur Trennlinie und leicht nach oben gesetzt.
    top_y = y_pct(h,64) - ui_s(18)
    tools_y = y_pct(h,43) - ui_s(14)
    mini_tw=int(tw*0.72); mini_th=int(th*0.72)
    mini_center_y = y_pct(h,16.5) + th/2 - mini_th/2
    mini_top = mini_center_y - mini_th/2
    label_x = ui_s(8)
    # Module-Titel in Y etwa auf der unteren Kante des Header-/Abschnittbeginns.
    _fm484_draw_section_label(self,'Module', label_x, FM_HEADER_H + ui_s(42))
    create_tile_row(top_tiles, line_x1, top_y, 'main_menuzeile_1', top_tile_w, top_tile_h)
    first_sep_y = (top_y + tools_y)/2
    self.draw_continuous_relief_line(self.canvas, first_sep_y, line_x1, line_x2)
    _fm484_draw_section_label(self,'Tools', label_x, first_sep_y + ui_s(28))
    create_tile_row(middle_tiles, line_x1, tools_y, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    create_centered_row(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top-ui_s(13), line_x1, line_x2)
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_main_menu = _fm484_render_main_menu
except Exception:
    pass

# ---------- Lingo in Einstellungen ----------
def _fm484_lingo_items(self):
    items = []
    # Untermenues / App-Seiten
    for key, title in [
        ('main','Hauptmenü'),('closing_calendar','Abschlusskalender'),('knowledge_base','Wissenszentrale'),('data_prep','Tools - Hauptbuch'),('debitoren_tools','Tools - Debitoren'),('file_output','Dateiausgabe'),('settings','Einstellungen'),('information','Informationen'),('in_dev','In Entwicklung')
    ]:
        items.append((key,title))
    for key, meta in sorted(TOOL_REGISTRY.items()):
        if key in HIDDEN_TOOL_IDS:
            continue
        items.append((key, meta.get('title') or meta.get('favorite_label') or key))
    return items


def _fm484_show_lingo_popup(self):
    if self.my_role() != ROLE_E4:
        messagebox.showwarning('Lingo', 'Lingo kann nur von E4 gepflegt werden.'); return
    popup = tk.Toplevel(self.root); popup.title('Lingo - Favoriten-Abkürzungen'); popup.configure(bg=BG); popup.transient(self.root); popup.resizable(True, True); popup.minsize(760,560)
    try: self._make_modal(popup)
    except Exception: pass
    settings = self.user_data.setdefault('settings', {})
    shortcuts = settings.setdefault('lingo_shortcuts', {})
    tk.Label(popup, text='Lingo - Abkürzungen für Favoriten-Chips', bg=HEADER, fg=TEXT, font=body_font(14,'bold'), anchor='w').pack(fill='x', ipady=10, padx=0)
    box = tk.Frame(popup, bg=BG); box.pack(fill='both', expand=True, padx=14, pady=12)
    canvas = tk.Canvas(box, bg=BG, highlightthickness=0); canvas.pack(side='left', fill='both', expand=True)
    scroll = ttk.Scrollbar(box, orient='vertical', command=canvas.yview); scroll.pack(side='right', fill='y')
    canvas.configure(yscrollcommand=scroll.set)
    inner = tk.Frame(canvas, bg=BG); canvas.create_window(0,0,window=inner,anchor='nw')
    vars = {}
    tk.Label(inner, text='Kontext / Tool', bg=BG, fg=TEXT2, font=body_font(9,'bold')).grid(row=0,column=0,sticky='w',padx=(0,12),pady=(0,6))
    tk.Label(inner, text='Abkürzung', bg=BG, fg=TEXT2, font=body_font(9,'bold')).grid(row=0,column=1,sticky='w',pady=(0,6))
    for r,(key,title) in enumerate(_fm484_lingo_items(self),1):
        tk.Label(inner, text=title, bg=BG, fg=TEXT, font=body_font(9), anchor='w', width=34).grid(row=r,column=0,sticky='w',padx=(0,12),pady=2)
        var = tk.StringVar(value=str(shortcuts.get(key,'')))
        vars[key]=var
        tk.Entry(inner, textvariable=var, bg=WHITE, fg=TEXT, font=body_font(9), width=12, relief='solid', bd=1).grid(row=r,column=1,sticky='w',pady=2)
    def refresh_scroll(_e=None): canvas.configure(scrollregion=canvas.bbox('all'))
    inner.bind('<Configure>', refresh_scroll)
    footer=tk.Frame(popup,bg=BG); footer.pack(fill='x',padx=14,pady=(0,12))
    def save():
        shortcuts.clear()
        for k,v in vars.items():
            val=v.get().strip()
            if val: shortcuts[k]=val[:8]
        self.save_user_data()
        messagebox.showinfo('Lingo','Lingo-Abkürzungen gespeichert.')
        popup.destroy(); self.render_page()
    tk.Button(footer,text='Speichern',command=save,bg='#CFEAD6',fg=TEXT,font=body_font(10,'bold'),relief='solid',bd=1,padx=18,pady=6).pack(side='right',padx=(8,0))
    tk.Button(footer,text='Schließen',command=popup.destroy,bg=WHITE,fg=TEXT,font=body_font(10),relief='solid',bd=1,padx=18,pady=6).pack(side='right')

try:
    _fm484_prev_render_settings_menu = FiBuMateApp.render_settings_menu
except Exception:
    _fm484_prev_render_settings_menu = None

def _fm484_render_settings_menu(self):
    modules = [
        ('Kachelfarbe', 'tile_colors'),
        ('Benutzerverwaltung', 'users'),
        ('Berechtigungen', 'permissions'),
        ('Versionsverlauf', 'versions'),
    ]
    if self.my_role() == ROLE_E4:
        modules.append(('Lingo', '__lingo_popup__'))
    def open_mid(mid, ttl):
        if mid == '__lingo_popup__':
            return self.show_lingo_popup()
        return self.show_page(mid, ttl, True)
    w,h=self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w=max(250,min(330,int(w*0.18))); tile_h=max(105,min(140,int(h*0.13)))
    gap=max(28,int(w*0.025)); total=len(modules)*tile_w+(len(modules)-1)*gap
    start=(w-total)/2+tile_w/2; y=y_pct(h,56)
    for i,(title,mid) in enumerate(modules):
        tile=Tile(self.root,self,f'settings_{mid}',title,lambda m=mid,t=title: open_mid(m,t),favorite_enabled=False,icon_type='gear')
        tile.resize_tile(tile_w,tile_h)
        self.widget_items.append(tile); self.focusable_tiles.append(tile)
        self.canvas.create_window(start+i*(tile_w+gap),y,window=tile,anchor='center')
    self.draw_bottom_logo()

try:
    FiBuMateApp.show_lingo_popup = _fm484_show_lingo_popup
    FiBuMateApp.render_settings_menu = _fm484_render_settings_menu
except Exception:
    pass

# ---------- Globales Loeschen-Icon, auch sinngemaesse Texte ----------
_FM484_ORIG_TK_BUTTON_INIT = tk.Button.__init__
_FM484_ORIG_TK_BUTTON_CONFIGURE = tk.Button.configure
_FM484_DELETE_WORDS = ('löschen','loeschen','entfernen','remove','delete','verwerfen')


def _fm484_delete_button_should_icon(widget):
    try:
        txt = str(widget.cget('text') or '').lower()
        return any(w in txt for w in _FM484_DELETE_WORDS)
    except Exception:
        return False


def _fm484_get_delete_photo(widget):
    if not PIL_AVAILABLE:
        return None
    try:
        path = FM_DELETE_ICON_PATH
        if not os.path.exists(path):
            path = os.path.join(ICON_DIR, 'biggarbagebin_121980.ico')
        if not os.path.exists(path):
            return None
        img = Image.open(path).resize((ui_s(18), ui_s(18)))
        photo = ImageTk.PhotoImage(img)
        widget._fm484_delete_icon = photo
        return photo
    except Exception:
        return None


def _fm484_apply_delete_icon(widget):
    try:
        if not _fm484_delete_button_should_icon(widget):
            return
        photo = _fm484_get_delete_photo(widget)
        if photo:
            _FM484_ORIG_TK_BUTTON_CONFIGURE(widget, None, image=photo, compound='left')
    except Exception:
        pass


def _fm484_button_init(self, master=None, cnf={}, **kw):
    _FM484_ORIG_TK_BUTTON_INIT(self, master, cnf, **kw)
    _fm484_apply_delete_icon(self)


def _fm484_button_configure(self, cnf=None, **kw):
    result = _FM484_ORIG_TK_BUTTON_CONFIGURE(self, cnf, **kw)
    _fm484_apply_delete_icon(self)
    return result

try:
    tk.Button.__init__ = _fm484_button_init
    tk.Button.configure = _fm484_button_configure
    tk.Button.config = _fm484_button_configure
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_AERO_MAINMENU_FAVORITES_HEADER_V0485
# Datum: 2026-07-08
# Zweck:
# - Hauptmenue-Kacheln knapp unterhalb der Trennlinien, weiterhin linksbuendig zur Trennlinie.
# - Favoriten-Chips korrekt innerhalb der Favoritenleiste platzieren.
# - Favoriten-Stern groesser mit dezentem Umriss.
# - Button "Aenderung vorschlagen" gleich breit wie "Hilfe", ohne Textueberlauf.
# - Kacheldesign moderner im Windows-7-Aero-Stil ohne echte Transparenz.
# ------------------------------------------------------------------
FIBU_MATE_AERO_MAINMENU_FAVORITES_HEADER_VERSION = "0.485"


def _fm485_aero_color_steps(base, count=18):
    steps = []
    for i in range(count):
        t = i / max(1, count - 1)
        if t < 0.45:
            c = blend(base, WHITE, 0.18 - 0.10 * (t / 0.45))
        else:
            c = blend(base, '#0B2542', 0.05 + 0.08 * ((t - 0.45) / 0.55))
        steps.append(c)
    return steps


def _fm485_tile_draw(self):
    self.delete('all')
    pad = ui_s(8)
    off = ui_s(1) if self.pressed else 0
    x0, y0 = pad + off, pad + off
    x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
    base = self.current_color()
    # Unveraendert genau ein Schatten.
    if not self.pressed:
        self.create_rectangle(pad + 7, pad + 8, pad + 7 + self.tile_width, pad + 8 + self.tile_height, fill=blend(SHADOW, '#1F2937', 0.16), outline='')
    # Aero-artige Flaeche ohne Transparenz: feine Farbbaender + Gloss-Streifen.
    self.create_rectangle(x0, y0, x1, y1, fill=base, outline=darken(base, 0.22), width=1)
    band_count = 22
    band_h = max(1, int((y1 - y0) / band_count))
    for i, color in enumerate(_fm485_aero_color_steps(base, band_count)):
        yy0 = y0 + i * band_h
        yy1 = y1 if i == band_count - 1 else min(y1, yy0 + band_h + 1)
        self.create_rectangle(x0 + 1, yy0, x1 - 1, yy1, fill=color, outline=color)
    # Glas-/Aero-Highlight: rein deckende helle Oberflaeche, keine echte Transparenz.
    gloss_h = max(ui_s(24), int((y1 - y0) * 0.32))
    self.create_rectangle(x0 + 3, y0 + 3, x1 - 3, y0 + gloss_h, fill=blend(base, WHITE, 0.16), outline='')
    self.create_line(x0 + 3, y0 + 3, x1 - 3, y0 + 3, fill=blend(base, WHITE, 0.46), width=1)
    self.create_line(x0 + 3, y0 + gloss_h, x1 - 3, y0 + gloss_h, fill=blend(base, WHITE, 0.08), width=1)
    # dezente 3D-Kanten, kein weisser Hover-Rand.
    self.create_line(x0 + 2, y0 + 2, x1 - 2, y0 + 2, fill=blend(base, WHITE, 0.30), width=1)
    self.create_line(x0 + 2, y0 + 2, x0 + 2, y1 - 2, fill=blend(base, WHITE, 0.14), width=1)
    self.create_line(x0 + 2, y1 - 2, x1 - 2, y1 - 2, fill=darken(base, 0.20), width=2)
    self.create_line(x1 - 2, y0 + 2, x1 - 2, y1 - 2, fill=darken(base, 0.14), width=1)
    if self.corner_fold:
        self._draw_corner_fold(x1, y0)
    title_y = y0 + ui_s(14)
    title_color = BLUE if self.lock_tile else 'white'
    icon_to_draw = 'lock' if self.lock_tile else self.icon_type
    text_width = max(ui_s(110), self.tile_width - ui_s(44))
    if self.center_text and not icon_to_draw:
        centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
        self.create_text((x0 + x1) / 2, (y0 + y1) / 2, text=self.title, anchor='center', fill=title_color, font=centered_font, width=text_width, justify='center')
    else:
        title_max_h = max(ui_s(28), int((y1 - y0) * 0.36))
        title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
        self.create_text((x0 + x1) / 2, title_y, text=self.title, anchor='n', fill=title_color, font=title_font, width=text_width, justify='center')
        if icon_to_draw:
            min_icon_y = title_y + title_h + ui_s(20)
            default_icon_y = y0 + int((y1 - y0) * 0.66)
            icon_y = max(default_icon_y, min_icon_y)
            icon_y = min(icon_y, y1 - ui_s(34))
            close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, 'current_close_period_label') else ''
            if close_period and self.tile_id in ('monthly_close', 'quarterly_close', 'yearly_close'):
                icon_x = x0 + 98
                self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                self.create_text(icon_x + 48, icon_y, text=close_period, anchor='w', fill=title_color, font=self.app.zoomed_content_font((FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', 13, 'italic')), width=max(120, x1 - icon_x - 58), justify='left')
            else:
                self._draw_main_icon(icon_to_draw, (x0 + x1) / 2, icon_y)
    self.star_bounds = None
    if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
        sx, sy = x1 - 24, y0 + 24
        self.star_bounds = (sx - 18, sy - 18, sx + 18, sy + 18)
        self.create_text(sx, sy, text='★', fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', 24, 'bold'))

try:
    Tile.draw = _fm485_tile_draw
except Exception:
    pass


def _fm485_draw_favorites_bar(self):
    w = self.canvas.winfo_width(); y_mid = FM_BAR_Y if 'FM_BAR_Y' in globals() else ui_s(62)
    bar_w = _fm484_bar_width(self) if '_fm484_bar_width' in globals() else max(ui_s(440), min(ui_s(690), int(w * 0.34)))
    x2 = w - ui_s(10)
    x1 = max(w * 0.52, x2 - bar_w)
    self.draw_gradient_line(x1, x2, y_mid - ui_s(10))
    self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
    star_x = x1 + ui_s(8)
    star_size = max(13, ui_s(16))
    # Stern mit dezentem Umriss durch mehrfache leicht versetzte Zeichnung.
    for dx, dy in [(-1,0),(1,0),(0,-1),(0,1)]:
        self.canvas.create_text(star_x + dx, y_mid + dy, text='★', font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', star_size, 'bold'), fill=blend(TEXT2, BLUE, 0.25), anchor='w')
    self.canvas.create_text(star_x, y_mid, text='★', font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', star_size, 'bold'), fill=GOLD, anchor='w')
    self.canvas.create_text(x1 + ui_s(34), y_mid, text='Favoriten', font=body_font(9, 'bold'), fill=TEXT, anchor='w')
    chip_w = ui_s(58); chip_gap = ui_s(4); chip_h = ui_s(18)
    x = x1 + ui_s(108)
    max_x = x2 - ui_s(4)
    for fav in sorted(f for f in self.favorites if f in TOOL_REGISTRY and f not in HIDDEN_TOOL_IDS):
        if x + chip_w > max_x:
            break
        label = _fm484_lingo_label(self, fav) if '_fm484_lingo_label' in globals() else (TOOL_REGISTRY.get(fav, {}).get('favorite_label', fav)[:8])
        chip_color = darken(self.current_tile_color() or BLUE, 0.08)
        chip = tk.Label(self.root, text=label, bg=chip_color, fg='white', font=body_font(8, 'bold'), width=7, anchor='center', padx=0, pady=0, cursor='hand2')
        chip._zoom_exclude = True
        chip.bind('<Button-1>', lambda e, fid=fav: self.execute_favorite(fid))
        self.widget_items.append(chip)
        # Innerhalb der Leiste: zentriert zwischen oberer und unterer Linie.
        self.canvas.create_window(x, y_mid, window=chip, anchor='w', width=chip_w, height=chip_h)
        x += chip_w + chip_gap

try:
    FiBuMateApp.draw_favorites_bar = _fm485_draw_favorites_bar
except Exception:
    pass


def _fm485_draw_suggestion_control(self):
    if self.current_page == 'launch':
        return
    suggestion_w = MINI_WIDGET_W
    btn = tk.Canvas(self.root, width=suggestion_w, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor='hand2')
    btn._zoom_exclude = True
    btn.create_rectangle(1, 1, suggestion_w - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
    icon_x = ui_s(15)
    self.draw_mini_bulb_icon(btn, icon_x, MINI_WIDGET_H / 2)
    # gleiche Breite wie Hilfe; kleinerer Text und feste Breite verhindern Ueberlauf.
    text_x = ui_s(31)
    btn.create_text(text_x, MINI_WIDGET_H / 2, text='Änderung vorschlagen', fill=TEXT, font=(FM_MODERN_FONT_FAMILY if 'FM_MODERN_FONT_FAMILY' in globals() else 'Aptos', max(7, ui_s(8)), 'bold'), anchor='w', width=suggestion_w - text_x - ui_s(4))
    btn.bind('<Button-1>', lambda _e: self.open_suggestion_mail())
    btn.bind('<Enter>', lambda _e: self.show_small_tooltip(btn, 'Änderung vorschlagen'))
    btn.bind('<Leave>', lambda _e: self.hide_small_tooltip())
    self.widget_items.append(btn)
    # rechter Rand links vom Hilfe-Mini-Widget, beide Buttons gleich breit.
    x = self.canvas.winfo_width() - ui_s(22) - MINI_WIDGET_W - MINI_WIDGET_GAP
    self.canvas.create_window(x, ui_s(8), window=btn, anchor='ne')

try:
    FiBuMateApp.draw_suggestion_control = _fm485_draw_suggestion_control
except Exception:
    pass


def _fm485_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165))); th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04))
    line_x1 = x_pct(w, 9.5); line_x2 = x_pct(w, 92)
    def create_tile_row(items, visible_left_edge, y, id_prefix, tile_w, tile_h):
        # Tile-Canvas hat 8 px Innenpad; sichtbare Kachelkante wird exakt auf visible_left_edge gelegt.
        canvas_left = visible_left_edge - ui_s(8)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            cx = canvas_left + (tile_w + ui_s(16)) / 2 + i*(tile_w+gap_x)
            self.canvas.create_window(cx, y, window=tile, anchor='center')
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count*tile_w + (count-1)*gap
        start_x = (w-total_width)/2 + tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]
    def create_centered_row(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i,item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')
    top_tile_w=max(250,int(tw*1.05)); top_tile_h=max(130,int(th*1.05))
    tool_tile_w=max(200,int(tw*0.86*0.95)); tool_tile_h=max(100,int(th*0.86*0.95))
    # Nur knapp unterhalb der Trennlinien.
    top_edge = (FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)) + ui_s(10)
    top_y = top_edge + top_tile_h / 2
    first_sep_y = top_edge + top_tile_h + ui_s(24)
    tools_edge = first_sep_y + ui_s(10)
    tools_y = tools_edge + tool_tile_h / 2
    mini_tw=int(tw*0.72); mini_th=int(th*0.72)
    mini_center_y = y_pct(h,16.5) + th/2 - mini_th/2
    mini_top = mini_center_y - mini_th/2
    label_x = ui_s(8)
    _fm484_draw_section_label(self,'Module', label_x, (FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)) + ui_s(42)) if '_fm484_draw_section_label' in globals() else self.canvas.create_text(label_x, ui_s(154), text='Module', fill=TEXT2, font=body_font(9,'bold'), anchor='w')
    create_tile_row(top_tiles, line_x1, top_y, 'main_menuzeile_1', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, first_sep_y, line_x1, line_x2)
    _fm484_draw_section_label(self,'Tools', label_x, first_sep_y + ui_s(28)) if '_fm484_draw_section_label' in globals() else self.canvas.create_text(label_x, first_sep_y + ui_s(28), text='Tools', fill=TEXT2, font=body_font(9,'bold'), anchor='w')
    create_tile_row(middle_tiles, line_x1, tools_y, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    create_centered_row(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top-ui_s(13), line_x1, line_x2)
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_main_menu = _fm485_render_main_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_MAINMENU_TILES_BELOW_SEPARATORS_V0486
# Datum: 2026-07-08
# Zweck:
# - Hauptmenue-Kacheln duerfen Trennlinien nicht schneiden.
# - Kachel-Canvas inklusive Schatten liegt direkt unterhalb der jeweiligen Trennlinie.
# - Linksbündigkeit zur Trennlinienkante, Kachelgröße und Kachelabstand bleiben erhalten.
# ------------------------------------------------------------------
FIBU_MATE_MAINMENU_TILES_BELOW_SEPARATORS_VERSION = "0.486"


def _fm486_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165)))
    th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04))
    line_x1 = x_pct(w, 9.5)
    line_x2 = x_pct(w, 92)
    header_line_y = FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)

    def create_tile_row(items, visible_left_edge, canvas_top, id_prefix, tile_w, tile_h):
        # Tile-Canvas ist tile_w/tile_h + 16 gross. Die Kachel plus Schatten beginnt erst unterhalb der Linie.
        canvas_left = visible_left_edge - ui_s(8)
        canvas_h = tile_h + ui_s(16)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            cx = canvas_left + (tile_w + ui_s(16)) / 2 + i*(tile_w+gap_x)
            cy = canvas_top + canvas_h / 2
            self.canvas.create_window(cx, cy, window=tile, anchor='center')

    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count*tile_w + (count-1)*gap
        start_x = (w-total_width)/2 + tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]

    def create_centered_row(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i,item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w,tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')

    top_tile_w=max(250,int(tw*1.05))
    top_tile_h=max(130,int(th*1.05))
    tool_tile_w=max(200,int(tw*0.86*0.95))
    tool_tile_h=max(100,int(th*0.86*0.95))
    mini_tw=int(tw*0.72)
    mini_th=int(th*0.72)

    # Vollstaendige Kachel-Canvas inkl. Schatten knapp unterhalb der Trennlinie platzieren.
    gap_below_line = ui_s(5)
    top_canvas_top = header_line_y + gap_below_line
    top_canvas_bottom = top_canvas_top + top_tile_h + ui_s(16)
    first_sep_y = top_canvas_bottom + ui_s(18)
    tools_canvas_top = first_sep_y + gap_below_line
    tools_canvas_bottom = tools_canvas_top + tool_tile_h + ui_s(16)

    # Unterer Abschnitt bleibt weiterhin mittig.
    mini_center_y = y_pct(h,16.5) + th/2 - mini_th/2
    mini_top = mini_center_y - mini_th/2

    label_x = ui_s(8)
    if '_fm484_draw_section_label' in globals():
        _fm484_draw_section_label(self, 'Module', label_x, header_line_y + ui_s(42))
    else:
        self.canvas.create_text(label_x, header_line_y + ui_s(42), text='Module', fill=TEXT2, font=body_font(9,'bold'), anchor='w')

    create_tile_row(top_tiles, line_x1, top_canvas_top, 'main_menuzeile_1', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, first_sep_y, line_x1, line_x2)

    if '_fm484_draw_section_label' in globals():
        _fm484_draw_section_label(self, 'Tools', label_x, first_sep_y + ui_s(28))
    else:
        self.canvas.create_text(label_x, first_sep_y + ui_s(28), text='Tools', fill=TEXT2, font=body_font(9,'bold'), anchor='w')

    create_tile_row(middle_tiles, line_x1, tools_canvas_top, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    create_centered_row(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_continuous_relief_line(self.canvas, mini_top-ui_s(13), line_x1, line_x2)
    self.draw_bottom_logo()

try:
    FiBuMateApp.render_main_menu = _fm486_render_main_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_DATEIAUSGABE_WZ_START_REWORK_V0487
# Datum: 2026-07-09
# Zweck:
# - Dateiausgabe: Versandpfeil per Rechtsklick persistent toggeln; Outlook setzt Versandpfeil automatisch.
# - Dateiausgabe: registrierte FiBu-Mate-Exporte und Dateien im Dateiausgabe-Ausgabepfad anzeigen.
# - Dateiausgabe: Benutzerfilter per Spaltentitel, Versandstatusfilter, Teams-Button entfernt.
# - Dateiausgabe: Linksklick auf selektierte Zeile hebt Auswahl auf; Buttons ohne Auswahl deaktiviert; Mehrfachauswahl per Strg.
# - Dateiausgabe: Vorschau als echte Bild-/Thumbnail-Renderings; PDF alle Seiten, Bilder direkt, Excel/CSV als Tabellenbild, Word-Fallback Text.
# - Wissenszentrale: Startkarten ohne pixelige Kreise; Karten dezent in jeweiliger Farbe eingefärbt.
# ------------------------------------------------------------------
FIBU_MATE_DATEIAUSGABE_WZ_START_REWORK_VERSION = "0.487"
_FM487_SENT_FILTERS = ("Alle", "Versandt", "Nicht versandt")


def _fm487_file_exts():
    return set(_FM472_OUTPUT_EXTENSIONS) if '_FM472_OUTPUT_EXTENSIONS' in globals() else {'.pdf','.xlsx','.xls','.docx','.doc','.csv','.txt','.png','.jpg','.jpeg','.bmp','.gif'}


def _fm487_file_id(path):
    return _fm472_file_key(path) if '_fm472_file_key' in globals() else hashlib.sha1(os.path.abspath(str(path)).casefold().encode('utf-8', errors='ignore')).hexdigest()[:18]


def _fm487_unknown_user(value):
    v = str(value or '').strip()
    return v if v else 'unbekannt'


def _fm487_load_logs():
    try:
        return _fm473_load_all_log_items()
    except Exception:
        return {}


def _fm487_write_logs(logs):
    try:
        return _fm473_write_log_items(logs)
    except Exception:
        try:
            target = _fm473_log_target() if '_fm473_log_target' in globals() else (_fm472_log_paths()[0] if '_fm472_log_paths' in globals() and _fm472_log_paths() else os.path.join(os.getcwd(), 'fibu_mate_dateiausgabe_log.json'))
            _fm472_write_json(target, {'items': list((logs or {}).values())})
            return True
        except Exception:
            return False


def _fm487_output_roots():
    try:
        return _fm472_output_roots()
    except Exception:
        return [os.path.join(NETWORK_ROOT, 'Dateiausgabe')]


def _fm487_scan_output_files(self):
    """Alle Log-Dateien plus alle Dateien in hinterlegten Dateiausgabe-Wurzeln.
    Manuell ausserhalb gespeicherte FiBu-Mate-Dateien erscheinen ueber das Dateiausgabe-Log.
    """
    logs = _fm487_load_logs()
    items, seen = [], set()
    # 1) registrierte FiBu-Mate-Dateien, auch ausserhalb des Ausgabepfads
    for raw in logs.values():
        if not isinstance(raw, dict):
            continue
        p = raw.get('path') or ''
        if not p:
            continue
        try:
            ap = os.path.abspath(str(p))
            if not os.path.isfile(ap):
                continue
            st = os.stat(ap)
            iid = raw.get('id') or _fm487_file_id(ap)
            item = dict(raw)
            item.update({'id': iid, 'type': 'file', 'name': item.get('name') or os.path.basename(ap), 'path': ap})
            item['user'] = _fm487_unknown_user(item.get('user'))
            item['module'] = item.get('module') or (_fm472_guess_module(ap) if '_fm472_guess_module' in globals() else 'FiBu Mate')
            item['created_at'] = item.get('created_at') or datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds')
            item['created_by_fibu_mate'] = True
            items.append(item); seen.add(ap.casefold())
        except Exception:
            pass
    # 2) Dateien im hinterlegten Dateiausgabe-Pfad, Nutzer unbekannt wenn nicht registriert
    for root in _fm487_output_roots():
        if not os.path.isdir(root):
            continue
        try:
            for dirpath, dirnames, filenames in os.walk(root):
                if any(part.lower() in {'__pycache__','logs','previewcache'} for part in dirpath.split(os.sep)):
                    continue
                for fn in filenames:
                    if os.path.splitext(fn)[1].lower() not in _fm487_file_exts():
                        continue
                    ap = os.path.abspath(os.path.join(dirpath, fn))
                    if ap.casefold() in seen:
                        continue
                    try:
                        st = os.stat(ap); iid = _fm487_file_id(ap); raw = logs.get(iid, {})
                        items.append({
                            'id': iid, 'type': 'file', 'name': fn, 'path': ap,
                            'user': _fm487_unknown_user(raw.get('user') or ''),
                            'module': raw.get('module') or (_fm472_guess_module(ap) if '_fm472_guess_module' in globals() else 'FiBu Mate'),
                            'created_at': raw.get('created_at') or datetime.fromtimestamp(st.st_mtime).isoformat(timespec='seconds'),
                            'group_id': raw.get('group_id',''), 'group_title': raw.get('group_title',''),
                            'sent_at': raw.get('sent_at',''), 'sent_channel': raw.get('sent_channel',''),
                            'created_by_fibu_mate': True,
                        })
                        seen.add(ap.casefold())
                    except Exception:
                        pass
        except Exception:
            pass
    return items


def _fm487_build_groups(self, items):
    groups, singles = {}, []
    for item in items:
        gid = str(item.get('group_id') or '').strip()
        if gid:
            g = groups.setdefault(gid, {
                'id': 'grp_' + gid, 'type': 'group',
                'name': item.get('group_title') or item.get('module') or 'Sammelexport',
                'path': '', 'user': item.get('user',''), 'module': item.get('module',''),
                'created_at': item.get('created_at',''), 'children': [],
                'sent_at': item.get('sent_at',''), 'sent_channel': item.get('sent_channel','')
            })
            g['children'].append(item)
            if str(item.get('created_at','')) > str(g.get('created_at','')): g['created_at'] = item.get('created_at','')
            if item.get('sent_at') and str(item.get('sent_at')) > str(g.get('sent_at','')):
                g['sent_at'] = item.get('sent_at'); g['sent_channel'] = item.get('sent_channel','')
        else:
            singles.append(item)
    out = list(groups.values()) + singles
    out.sort(key=lambda x: str(x.get('created_at','')), reverse=True)
    for g in out:
        if g.get('type') == 'group':
            g['children'] = sorted(g.get('children', []) or [], key=lambda c: str(c.get('name','')).casefold())
    return out


def _fm487_build_items(self):
    return _fm487_build_groups(self, _fm487_scan_output_files(self))


def _fm487_register_output_file(self, path, user=None, module=None, group_id=None, group_title=None, created_at=None, children=None):
    try:
        if not path: return False
        ap = os.path.abspath(str(path)); iid = _fm487_file_id(ap)
        logs = _fm487_load_logs(); old = logs.get(iid, {})
        old.update({
            'id': iid, 'type': 'file', 'name': os.path.basename(ap), 'path': ap,
            'user': _fm487_unknown_user(user or old.get('user') or getattr(self, 'current_user_display', '') or getattr(self, 'current_user_key', '')),
            'module': module or old.get('module') or (_fm472_guess_module(ap) if '_fm472_guess_module' in globals() else 'FiBu Mate'),
            'group_id': group_id or old.get('group_id',''), 'group_title': group_title or old.get('group_title',''),
            'created_at': created_at or old.get('created_at') or _fm472_now_iso(),
            'created_by_fibu_mate': True,
        })
        logs[iid] = old
        return _fm487_write_logs(logs)
    except Exception:
        return False


def _fm487_files_for_item(item):
    try:
        return _fm473_files_for_item(item)
    except Exception:
        if not item: return []
        if item.get('type') == 'group':
            return [c.get('path') for c in item.get('children', []) or [] if c.get('path') and os.path.isfile(c.get('path'))]
        p = item.get('path')
        return [p] if p and os.path.isfile(p) else []


def _fm487_each_file_item(item):
    if not item:
        return []
    if item.get('type') == 'group':
        return [c for c in item.get('children', []) or [] if c.get('path')]
    return [item]


def _fm487_mark_sent_manual(self, item, sent=True, channel='Manuell'):
    logs = _fm487_load_logs()
    now = datetime.now().isoformat(timespec='seconds')
    for sub in _fm487_each_file_item(item):
        p = sub.get('path',''); iid = sub.get('id') or (_fm487_file_id(p) if p else '')
        if not iid: continue
        cur = logs.get(iid, {})
        cur.update({k: sub.get(k) for k in ('id','type','name','path','user','module','created_at','group_id','group_title') if sub.get(k) is not None})
        cur.setdefault('id', iid); cur.setdefault('type', 'file')
        cur.setdefault('user', _fm487_unknown_user(cur.get('user')))
        if sent:
            cur['sent_at'] = now; cur['sent_channel'] = channel
            sub['sent_at'] = now; sub['sent_channel'] = channel
        else:
            cur.pop('sent_at', None); cur.pop('sent_channel', None)
            sub['sent_at'] = ''; sub['sent_channel'] = ''
        logs[iid] = cur
    if item.get('type') == 'group':
        if sent: item['sent_at'] = now; item['sent_channel'] = channel
        else: item['sent_at'] = ''; item['sent_channel'] = ''
    else:
        if sent: item['sent_at'] = now; item['sent_channel'] = channel
        else: item['sent_at'] = ''; item['sent_channel'] = ''
    return _fm487_write_logs(logs)


def _fm487_toggle_sent(self, item):
    active = bool(item and item.get('sent_at'))
    return _fm487_mark_sent_manual(self, item, sent=not active, channel='Manuell')


def _fm487_send_outlook(self, items):
    if not isinstance(items, (list, tuple)):
        items = [items] if items else []
    files = []
    for item in items:
        files.extend(_fm487_files_for_item(item))
    files = [f for f in dict.fromkeys(files) if f and os.path.isfile(f)]
    if not files:
        try: messagebox.showwarning('Dateiausgabe', 'Keine versendbare Datei gefunden.')
        except Exception: pass
        return False
    try:
        import win32com.client
        outlook = win32com.client.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)
        mail.Subject = 'FiBu Mate Dateiausgabe'
        mail.Body = 'Hallo,\n\nanbei die Datei(en) aus FiBu Mate.\n\nViele Grüße'
        for f in files: mail.Attachments.Add(f)
        mail.Display()
        for item in items: _fm487_mark_sent_manual(self, item, sent=True, channel='Outlook')
        return True
    except Exception as exc:
        try: messagebox.showerror('Dateiausgabe', 'Outlook-Versand mit Dateianhang konnte nicht vorbereitet werden:\n' + str(exc))
        except Exception: pass
        return False


def _fm487_sort_items(items, column, descending):
    def key(item):
        if column == 'name': return str(item.get('name','')).casefold()
        if column == 'user': return str(item.get('user','')).casefold()
        if column == 'path': return str(item.get('path','')).casefold()
        return str(item.get('created_at',''))
    out = sorted(items, key=key, reverse=bool(descending))
    for item in out:
        if item.get('type') == 'group':
            item['children'] = sorted(item.get('children', []) or [], key=key, reverse=bool(descending))
    return out


def _fm487_filter_items(items, query='', users=None, sent_filter='Alle'):
    q = str(query or '').strip().casefold()
    users = set(users or [])
    all_users = not users
    sent_filter = sent_filter or 'Alle'
    def hit(item):
        vals = [item.get('name',''), item.get('user',''), item.get('module',''), item.get('created_at',''), _fm472_display_dt(item.get('created_at')) if '_fm472_display_dt' in globals() else '', item.get('path','')]
        if q and q not in ' '.join(str(v) for v in vals).casefold(): return False
        if not all_users and _fm487_unknown_user(item.get('user')) not in users: return False
        if sent_filter == 'Versandt' and not item.get('sent_at'): return False
        if sent_filter == 'Nicht versandt' and item.get('sent_at'): return False
        return True
    result = []
    for item in items:
        if item.get('type') == 'group':
            children = [c for c in item.get('children', []) or [] if hit(c)]
            if hit(item) or children:
                clone = dict(item); clone['children'] = children if children else item.get('children', [])
                result.append(clone)
        elif hit(item):
            result.append(item)
    return result


def _fm487_make_table_image(rows, title='Tabellenvorschau'):
    if not PIL_AVAILABLE:
        return None
    try:
        from PIL import ImageDraw, ImageFont
        max_cols = min(8, max([len(r) for r in rows] or [1]))
        row_count = min(32, len(rows))
        col_w, row_h = 132, 28
        w = max(420, max_cols * col_w + 34)
        h = max(120, row_count * row_h + 58)
        img = Image.new('RGB', (w, h), '#F8FAFC')
        draw = ImageDraw.Draw(img)
        try: font = ImageFont.truetype('DejaVuSans.ttf', 11); font_b = ImageFont.truetype('DejaVuSans-Bold.ttf', 11)
        except Exception: font = font_b = None
        draw.text((16, 12), title, fill='#004B93', font=font_b)
        y0 = 42
        for r_idx, row in enumerate(rows[:row_count]):
            y = y0 + r_idx * row_h
            bg = '#E8EEF5' if r_idx == 0 else ('#FFFFFF' if r_idx % 2 else '#F2F6FA')
            draw.rectangle((16, y, 16 + max_cols*col_w, y + row_h), fill=bg, outline='#D7DEE8')
            for c in range(max_cols):
                x = 16 + c * col_w
                if c > 0: draw.line((x, y, x, y+row_h), fill='#D7DEE8')
                val = '' if c >= len(row) or row[c] is None else str(row[c])
                if len(val) > 24: val = val[:23] + '…'
                draw.text((x+6, y+8), val, fill='#182431', font=font_b if r_idx == 0 else font)
        return img
    except Exception:
        return None


def _fm487_preview_images_for_path(path):
    images = []
    ext = os.path.splitext(path)[1].lower()
    if not path or not os.path.exists(path):
        raise RuntimeError('Datei wurde nicht gefunden.')
    if ext == '.pdf':
        try:
            import fitz
            doc = fitz.open(path)
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(1.15, 1.15), alpha=False)
                img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
                images.append(img)
            doc.close()
        except Exception as exc:
            raise RuntimeError('PDF-Vorschau konnte nicht gerendert werden: ' + str(exc))
    elif ext in {'.png','.jpg','.jpeg','.bmp','.gif'}:
        try:
            images.append(Image.open(path).convert('RGB'))
        except Exception as exc:
            raise RuntimeError('Bildvorschau konnte nicht geladen werden: ' + str(exc))
    elif ext in {'.xlsx','.xls'}:
        rows=[]
        try:
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            ws = wb.active
            for i,row in enumerate(ws.iter_rows(max_rows=32, max_cols=8, values_only=True)):
                rows.append([v for v in row])
            try: wb.close()
            except Exception: pass
            img = _fm487_make_table_image(rows, 'Excel - erstes Arbeitsblatt')
            if img: images.append(img)
        except Exception as exc:
            raise RuntimeError('Excel-Vorschau konnte nicht gerendert werden: ' + str(exc))
    elif ext == '.csv':
        rows=[]
        try:
            import csv
            with open(path, 'r', encoding='utf-8-sig', errors='replace', newline='') as f:
                sample = f.read(4096); f.seek(0)
                try: dialect = csv.Sniffer().sniff(sample, delimiters=';,\t,')
                except Exception: dialect = csv.excel
                for i,row in enumerate(csv.reader(f, dialect)):
                    rows.append(row)
                    if i >= 31: break
            img = _fm487_make_table_image(rows, 'CSV - Tabellenvorschau')
            if img: images.append(img)
        except Exception as exc:
            raise RuntimeError('CSV-Vorschau konnte nicht gerendert werden: ' + str(exc))
    return images


def _fm487_draw_preview(self, canvas, items):
    canvas.delete('all'); canvas._fm487_preview_refs=[]
    width = max(460, canvas.winfo_width() or 720)
    y = 14
    if not items:
        canvas.configure(scrollregion=(0,0,width,80)); return
    if not isinstance(items, (list, tuple)): items = [items]
    files = []
    word_fallbacks = []
    errors = []
    for item in items:
        for sub in _fm487_each_file_item(item):
            p = sub.get('path','')
            if not p: continue
            ext = os.path.splitext(p)[1].lower()
            if ext in {'.doc','.docx','.txt'}:
                # Word-Fallback bzw. Text nur fuer Word/Text-Ausnahme.
                word_fallbacks.append((sub, p))
            else:
                files.append((sub, p))
    for sub,p in files:
        try:
            imgs = _fm487_preview_images_for_path(p)
            for idx,img in enumerate(imgs or []):
                max_w = max(360, width - 44)
                max_h = 900
                img = img.copy(); img.thumbnail((max_w, max_h))
                ph = ImageTk.PhotoImage(img)
                canvas._fm487_preview_refs.append(ph)
                canvas.create_image(18, y, image=ph, anchor='nw')
                canvas.create_rectangle(18, y, 18+ph.width(), y+ph.height(), outline=LINE)
                y += ph.height() + 18
        except Exception as exc:
            errors.append(str(exc))
    for sub,p in word_fallbacks:
        try:
            txt = ''
            ext = os.path.splitext(p)[1].lower()
            if ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(p)
                    txt = '\n'.join(par.text for par in doc.paragraphs if par.text).strip()[:6000]
                except Exception as exc:
                    txt = 'Word-Vorschau konnte nicht gelesen werden: ' + str(exc)
            else:
                with open(p, 'r', encoding='utf-8', errors='replace') as f: txt = f.read(6000)
            canvas.create_rectangle(18, y, width-22, y+360, fill=WHITE, outline=LINE)
            canvas.create_text(34, y+18, text=txt or 'Keine Textvorschau verfügbar.', anchor='nw', fill=TEXT, font=body_font(9), width=width-80)
            y += 380
        except Exception as exc:
            errors.append(str(exc))
    if errors:
        canvas.create_rectangle(18, y, width-22, y+80, fill='#FFF7CC', outline=LINE)
        canvas.create_text(34, y+16, text='\n'.join(errors[:4]), anchor='nw', fill=RED, font=body_font(9), width=width-80)
        y += 100
    canvas.configure(scrollregion=(0,0,width,max(y+30, canvas.winfo_height())))


def _fm487_user_filter_popup(self, tree, all_items, refresh_cb):
    popup = tk.Toplevel(self.root)
    popup.title('Benutzerfilter')
    popup.configure(bg=BG)
    popup.transient(self.root)
    popup.resizable(True, True)
    try: popup.geometry('360x460')
    except Exception: pass
    qvar = tk.StringVar(value='')
    selected = set(getattr(self, 'file_output_user_filter_selected', set()) or [])
    users = sorted({_fm487_unknown_user(i.get('user')) for i in _fm487_scan_output_files(self)})
    vars = {}
    top = tk.Frame(popup, bg=BG); top.pack(fill='x', padx=12, pady=10)
    tk.Label(top, text='Suchen', bg=BG, fg=TEXT2, font=body_font(9)).pack(side='left')
    ent = tk.Entry(top, textvariable=qvar, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1); ent.pack(side='left', fill='x', expand=True, padx=(8,0))
    box = tk.Frame(popup, bg=WHITE, relief='solid', bd=1); box.pack(fill='both', expand=True, padx=12, pady=(0,10))
    canvas = tk.Canvas(box, bg=WHITE, highlightthickness=0); canvas.pack(side='left', fill='both', expand=True)
    scroll = ttk.Scrollbar(box, orient='vertical', command=canvas.yview); scroll.pack(side='right', fill='y')
    canvas.configure(yscrollcommand=scroll.set)
    inner = tk.Frame(canvas, bg=WHITE); canvas.create_window(0,0,window=inner,anchor='nw')
    def rebuild(*_):
        for w in inner.winfo_children(): w.destroy()
        vars.clear(); q = qvar.get().strip().casefold()
        for u in users:
            if q and q not in u.casefold(): continue
            var = tk.BooleanVar(value=(not selected or u in selected))
            vars[u] = var
            tk.Checkbutton(inner, text=u, variable=var, bg=WHITE, fg=TEXT, font=body_font(9), anchor='w').pack(fill='x', anchor='w')
        try: canvas.configure(scrollregion=canvas.bbox('all'))
        except Exception: pass
    def select_all(flag):
        for v in vars.values(): v.set(bool(flag))
    qvar.trace_add('write', rebuild); rebuild()
    btns = tk.Frame(popup, bg=BG); btns.pack(fill='x', padx=12, pady=(0,12))
    tk.Button(btns, text='Alle auswählen', command=lambda: select_all(True), bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0,6))
    tk.Button(btns, text='Alle abwählen', command=lambda: select_all(False), bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left')
    def apply():
        chosen = {u for u,v in vars.items() if v.get()}
        self.file_output_user_filter_selected = chosen
        refresh_cb(); popup.destroy()
    tk.Button(btns, text='Übernehmen', command=apply, bg='#CFEAD6', fg=TEXT, font=body_font(9,'bold'), relief='solid', bd=1).pack(side='right')


def _fm487_render_file_output(self):
    self.focusable_tiles.clear()
    w,h = self.canvas.winfo_width(), self.canvas.winfo_height()
    x0,y0 = 24,135; full_w,full_h = max(900,w-48), max(520,h-y0-64)
    outer = tk.Frame(self.root,bg=BG); self.widget_items.append(outer)
    self.canvas.create_window(ui_s(x0),ui_s(y0),window=outer,anchor='nw',width=ui_s(full_w),height=ui_s(full_h))
    search_var = tk.StringVar(value=getattr(self,'file_output_search_text',''))
    sent_filter_var = tk.StringVar(value=getattr(self,'file_output_sent_filter','Alle') or 'Alle')
    sort_state = getattr(self,'file_output_sort_state',{'column':'date','descending':True}) or {'column':'date','descending':True}
    selected_items = {'items': list(getattr(self,'file_output_selected_items', []) or [])}
    all_items = self.file_output_build_items()

    top = tk.Frame(outer,bg=BG); top.pack(fill='x',pady=(0,8))
    tk.Label(top,text='Suche',bg=BG,fg=TEXT2,font=body_font(10)).pack(side='left',padx=(0,8))
    ent=tk.Entry(top,textvariable=search_var,bg=WHITE,fg=TEXT,font=body_font(11),relief='solid',bd=1); ent.pack(side='left',fill='x',expand=True,ipady=5)
    tk.Label(top,text='Versand',bg=BG,fg=TEXT2,font=body_font(10)).pack(side='left',padx=(10,6))
    sent_cb=ttk.Combobox(top,textvariable=sent_filter_var,values=list(_FM487_SENT_FILTERS),state='readonly',width=14,font=body_font(9)); sent_cb.pack(side='left')
    tk.Button(top,text='Aktualisieren',command=lambda:self.render_page(),bg=_FM473_BUTTON_GREY if '_FM473_BUTTON_GREY' in globals() else '#D6DCE4',fg=TEXT,font=body_font(10),relief='solid',bd=1).pack(side='left',padx=(8,0),ipadx=10,ipady=4)
    content=tk.Frame(outer,bg=BG); content.pack(fill='both',expand=True)
    left=tk.Frame(content,bg=WHITE,highlightbackground=LINE,highlightthickness=2); left.pack(side='left',fill='both',expand=True,padx=(0,12))
    right=tk.Frame(content,bg=WHITE,highlightbackground=LINE,highlightthickness=2); right.pack(side='right',fill='both',expand=True)
    left.configure(width=max(470,int(full_w*0.42))); right.configure(width=max(580,int(full_w*0.56)))
    head=tk.Frame(left,bg=WHITE); head.pack(fill='x',padx=12,pady=(10,6))
    tk.Label(head,text='Zuletzt erstellte Dateien',bg=WHITE,fg=BLUE,font=body_font(13,weight='bold')).pack(side='left')
    action_bar=tk.Frame(left,bg=WHITE); action_bar.pack(fill='x',padx=12,pady=(0,8))
    buttons=[]
    def selected(): return list(selected_items.get('items') or [])
    def one_selected():
        arr=selected(); return arr[0] if arr else {}
    def set_buttons_state():
        state='normal' if selected() else 'disabled'
        for b in buttons:
            try: b.configure(state=state, cursor='hand2' if state=='normal' else 'arrow')
            except Exception: pass
    def add_btn(text, cmd, icon=None):
        b=tk.Button(action_bar,text=text,command=cmd,bg=_FM473_BUTTON_GREY if '_FM473_BUTTON_GREY' in globals() else '#D6DCE4',fg=TEXT,activebackground=_FM473_BUTTON_GREY_ACTIVE if '_FM473_BUTTON_GREY_ACTIVE' in globals() else '#C8D0DA',font=body_font(9),relief='solid',bd=1,state='disabled',compound='left',cursor='arrow')
        try:
            ph=_fm473_icon_photo(self, icon, 18) if icon else None
            if ph: b.configure(image=ph); b._img=ph
        except Exception: pass
        b.pack(side='left',padx=(0,6),ipadx=8,ipady=3); buttons.append(b); return b
    add_btn('Datei öffnen', lambda: _fm472_open_file(one_selected()) if '_fm472_open_file' in globals() else None)
    add_btn('Speicherort öffnen', lambda: _fm472_open_folder(one_selected()) if '_fm472_open_folder' in globals() else None, _FM473_EXPLORER_ICON if '_FM473_EXPLORER_ICON' in globals() else None)
    add_btn('Outlook Versand', lambda: (_fm487_send_outlook(self, selected()) and self.render_page()), _FM473_OUTLOOK_ICON if '_FM473_OUTLOOK_ICON' in globals() else None)
    add_btn('Export löschen', lambda: _fm473_delete_export(self, one_selected()) if '_fm473_delete_export' in globals() else None)

    tree_frame=tk.Frame(left,bg=WHITE); tree_frame.pack(fill='both',expand=True,padx=12,pady=(0,12))
    columns=('name','user','date','path')
    tree=ttk.Treeview(tree_frame,columns=columns,show='tree headings',selectmode='extended')
    try: sent_img=_fm473_icon_photo(self,_FM473_SENT_ICON,16)
    except Exception: sent_img=None
    self._fm487_sent_img = sent_img
    def heading_text(col,label):
        marker=' ↓' if sort_state.get('column')==col and sort_state.get('descending',True) else (' ↑' if sort_state.get('column')==col else '')
        return label+marker
    def set_sort(col):
        cur=getattr(self,'file_output_sort_state',{'column':'date','descending':True})
        desc=not cur.get('descending',True) if cur.get('column')==col else (True if col=='date' else False)
        self.file_output_sort_state={'column':col,'descending':desc}; self.render_page()
    tree.heading('#0',text='')
    tree.heading('name',text=heading_text('name','Dateiname / Sammelposition'),command=lambda:set_sort('name'))
    tree.heading('user',text=heading_text('user','Benutzer'),command=lambda:_fm487_user_filter_popup(self, tree, all_items, fill_tree))
    tree.heading('date',text=heading_text('date','Erstellt'),command=lambda:set_sort('date'))
    tree.heading('path',text=heading_text('path','Ablagepfad kurz'),command=lambda:set_sort('path'))
    tree.column('#0',width=44,minwidth=44,stretch=False,anchor='center'); tree.column('name',width=220,stretch=True); tree.column('user',width=120,stretch=False); tree.column('date',width=122,stretch=False); tree.column('path',width=190,stretch=True)
    vsb=ttk.Scrollbar(tree_frame,orient='vertical',command=tree.yview); tree.configure(yscrollcommand=vsb.set)
    tree.pack(side='left',fill='both',expand=True); vsb.pack(side='right',fill='y')
    prev_head=tk.Frame(right,bg=WHITE); prev_head.pack(fill='x',padx=12,pady=(10,6))
    tk.Label(prev_head,text='Vorschau',bg=WHITE,fg=BLUE,font=body_font(13,weight='bold')).pack(side='left')
    prev_canvas=tk.Canvas(right,bg='#F8FAFC',highlightthickness=0,bd=0)
    prev_scroll=ttk.Scrollbar(right,orient='vertical',command=prev_canvas.yview); prev_canvas.configure(yscrollcommand=prev_scroll.set)
    prev_canvas.pack(side='left',fill='both',expand=True,padx=(12,0),pady=(0,12)); prev_scroll.pack(side='right',fill='y',padx=(0,12),pady=(0,12))
    id_map={}; suppress_select={'flag':False}
    def fill_tree():
        id_map.clear(); tree.delete(*tree.get_children())
        q=search_var.get(); self.file_output_search_text=q
        self.file_output_sent_filter=sent_filter_var.get() or 'Alle'
        user_filter=getattr(self,'file_output_user_filter_selected',set()) or set()
        shown=_fm487_filter_items(all_items,q,user_filter,self.file_output_sent_filter)
        shown=_fm487_sort_items(shown,sort_state.get('column','date'),sort_state.get('descending',True))
        for item in shown:
            iid=item.get('id') or _fm487_file_id(item.get('path','')+item.get('name',''))
            id_map[iid]=item
            img=sent_img if item.get('sent_at') and sent_img else ''
            tree.insert('', 'end', iid=iid, text='', image=img, values=(item.get('name',''), item.get('user',''), _fm472_display_dt(item.get('created_at')) if '_fm472_display_dt' in globals() else item.get('created_at',''), _fm472_short_path(item.get('path','')) if '_fm472_short_path' in globals() else item.get('path','')))
            if item.get('type')=='group':
                for c in item.get('children',[]) or []:
                    cid=c.get('id') or _fm487_file_id(c.get('path',''))
                    if cid in id_map: cid=iid+'_'+cid
                    id_map[cid]=c
                    cimg=sent_img if c.get('sent_at') and sent_img else ''
                    tree.insert(iid,'end',iid=cid,text='',image=cimg,values=(c.get('name',''),c.get('user',''),_fm472_display_dt(c.get('created_at')) if '_fm472_display_dt' in globals() else c.get('created_at',''),_fm472_short_path(c.get('path','')) if '_fm472_short_path' in globals() else c.get('path','')))
        tree.selection_remove(tree.selection())
        selected_items['items']=[]; self.file_output_selected_items=[]; self.file_output_selected_item=None
        _fm487_draw_preview(self,prev_canvas,[]); set_buttons_state()
    def current_selection_items():
        return [id_map[i] for i in tree.selection() if i in id_map]
    def apply_selection():
        if suppress_select.get('flag'): return
        arr=current_selection_items()
        selected_items['items']=arr; self.file_output_selected_items=arr; self.file_output_selected_item=arr[0] if arr else None
        _fm487_draw_preview(self,prev_canvas,arr); set_buttons_state()
    def on_click(ev):
        row=tree.identify_row(ev.y)
        col=tree.identify_column(ev.x)
        if not row:
            suppress_select['flag']=True; tree.selection_remove(tree.selection()); suppress_select['flag']=False; apply_selection(); return 'break'
        if col == '#0':
            return None
        sel=set(tree.selection())
        ctrl=bool(getattr(ev,'state',0) & 0x0004)
        if row in sel and not ctrl and len(sel)==1:
            suppress_select['flag']=True; tree.selection_remove(row); suppress_select['flag']=False; apply_selection(); return 'break'
        return None
    def on_right_click(ev):
        row=tree.identify_row(ev.y); col=tree.identify_column(ev.x)
        if row and col == '#0' and row in id_map:
            _fm487_toggle_sent(self, id_map[row]); fill_tree(); return 'break'
        return None
    tree.bind('<Button-1>', on_click)
    tree.bind('<Button-3>', on_right_click)
    tree.bind('<<TreeviewSelect>>', lambda e: apply_selection())
    search_var.trace_add('write', lambda *_: fill_tree())
    sent_filter_var.trace_add('write', lambda *_: fill_tree())
    try: prev_canvas.bind('<MouseWheel>',lambda e:(prev_canvas.yview_scroll(int(-1*(e.delta/120)),'units'),'break'))
    except Exception: pass
    fill_tree()

try:
    FiBuMateApp.file_output_register = _fm487_register_output_file
    FiBuMateApp.file_output_build_items = _fm487_build_items
    FiBuMateApp.render_file_output = _fm487_render_file_output
except Exception:
    pass


# Wissenszentrale-Start: Kreise entfernen, Karten dezent einfaerben
def _fm487_kb_card_color(color):
    try:
        return blend(WHITE, color, 0.10)
    except Exception:
        return WHITE


def _fm487_render_knowledge_start_overlay(self):
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    y0 = ui_s(98); y1 = h - ui_s(34)
    self.canvas.create_rectangle(0, y0, w, y1, fill=blend(BG, WHITE, 0.86), outline=LINE, width=2)
    self.canvas.create_rectangle(0, y0, w, ui_s(158), fill=blend(HEADER, WHITE, 0.34), outline=LINE, width=1)
    self.canvas.create_text(ui_s(45), ui_s(123), text='Wissenszentrale – Start', fill=BLUE, font=body_font(22, weight='bold'), anchor='w')
    self.canvas.create_text(ui_s(45), ui_s(148), text='Startseite maximiert – Arbeitsbereich vollständig abgedeckt', fill=TEXT2, font=body_font(9), anchor='w')
    self.draw_kb_button(w - ui_s(180), 112, 'Schließen', self.kb_close_start_overlay, False, width=130)
    self.canvas.create_text(ui_s(45), ui_s(195), text='Was möchtest du tun?', fill=TEXT, font=body_font(15, weight='bold'), anchor='w')
    var = tk.BooleanVar(value=bool(getattr(self, 'knowledge_show_start', True)))
    def _toggle(): self.save_knowledge_start_preference(var.get())
    chk = tk.Checkbutton(self.root, text='Startseite anzeigen', variable=var, command=_toggle, bg=blend(BG, WHITE, 0.86), fg=TEXT, font=body_font(11), activebackground=blend(BG, WHITE, 0.86))
    self.canvas.create_window(ui_s(45), h - ui_s(88), window=chk, anchor='nw', width=ui_s(260), height=ui_s(34)); self.widget_items.append(chk)
    cards=[('Wissen suchen','Gesamtliste und Suche öffnen','all','#2563EB'),('Neuer Eintrag','Dokumentation oder Aufgabe erfassen','new','#059669'),('To-Dos anzeigen','wiederkehrende Aufgaben anzeigen','todos','#DC2626'),('Kategorien','Kategorien anzeigen','categories','#7C3AED'),('Veraltete Einträge','Prüfung / Archiv','outdated','#EC4899')]
    card_w=min(430,max(300,int((max(900,w-ui_s(90))-90)/3))); card_h=150
    start_x,start_y=45,235; gap_x,gap_y=45,38
    for idx,(title,desc,view,color) in enumerate(cards):
        col,row=idx%3,idx//3; x=start_x+col*(card_w+gap_x); y=start_y+row*(card_h+gap_y)
        fill=_fm487_kb_card_color(color); hover_fill=blend(fill,color,0.08)
        self.canvas.create_rectangle(ui_s(x+6),ui_s(y+8),ui_s(x+card_w+6),ui_s(y+card_h+8),fill=SHADOW,outline='')
        rect=self.canvas.create_rectangle(ui_s(x),ui_s(y),ui_s(x+card_w),ui_s(y+card_h),fill=fill,outline=blend(LINE,color,0.18))
        self.canvas.create_text(ui_s(x+26),ui_s(y+35),text=title,fill=BLUE,font=body_font(15,weight='bold'),anchor='w')
        self.canvas.create_text(ui_s(x+26),ui_s(y+75),text=desc,fill=TEXT2,font=body_font(11),anchor='w')
        self.draw_kb_button(x+26,y+105,'Öffnen',lambda v=view:self.kb_switch_view_from_start(v),view=='new')
        def _enter(_e, item=rect, hf=hover_fill):
            try: self.canvas.itemconfigure(item, fill=hf)
            except Exception: pass
        def _leave(_e, item=rect, f=fill):
            try: self.canvas.itemconfigure(item, fill=f)
            except Exception: pass
        self.canvas.tag_bind(rect,'<Enter>',_enter); self.canvas.tag_bind(rect,'<Leave>',_leave)
    self.canvas.create_text(ui_s(45),h-ui_s(45),text='Die Startseite liegt über allen Arbeitsansicht-Widgets und endet erst oberhalb der Fußleiste.',fill=TEXT2,font=body_font(10),anchor='w')

try:
    FiBuMateApp.render_knowledge_start_overlay = _fm487_render_knowledge_start_overlay
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_PASSWORD_AUTH_FINAL_V0488
# Datum: 2026-07-13
# Zweck:
# - Passwortfunktion aktivieren.
# - Erstanmeldung ohne Passwort: Benutzer legt nach Eingabe des Benutzernamens direkt ein Passwort fest.
# - Login mit gesetztem Passwort nur nach Passwortpruefung.
# - Passwort-Reset ab E3: alter Hash wird geloescht, Benutzer muss beim naechsten Login neu vergeben.
# - E-Mail-Benachrichtigung automatisch via Outlook, Fallback: vorbereitete Outlook-Mail anzeigen.
# - Passwoerter werden nie angezeigt und nur als PBKDF2-Hash gespeichert.
# ------------------------------------------------------------------
FIBU_MATE_PASSWORD_AUTH_VERSION = "0.488"


def _fm488_role_rank_value(role):
    try:
        return int(ROLE_RANK.get(role, 1))
    except Exception:
        return 1


def _fm488_user_role_from_data(user):
    try:
        return ROLE_MIGRATION.get((user or {}).get("permission", ROLE_E1), (user or {}).get("permission", ROLE_E1) or ROLE_E1)
    except Exception:
        return ROLE_E1


def _fm488_user_fullname_from_record(key, user):
    try:
        full = str((user or {}).get("full_name", "")).strip()
        if full:
            return full
        first = str((user or {}).get("first_name", "")).strip()
        disp = str((user or {}).get("display_name", key)).strip()
        return " ".join(x for x in [first, disp] if x).strip() or str(key)
    except Exception:
        return str(key)


def _fm488_normalize_auth(user):
    if not isinstance(user, dict):
        return {"password_hash": None, "enabled": False, "password_must_set": True}
    auth = user.setdefault("auth", {})
    if not isinstance(auth, dict):
        auth = {}
        user["auth"] = auth
    auth.setdefault("password_hash", None)
    auth.setdefault("enabled", bool(auth.get("password_hash")))
    auth.setdefault("password_must_set", not bool(auth.get("password_hash")))
    return auth


def _fm488_hash_password(password):
    import os as _os, base64 as _base64, hashlib as _hashlib
    salt = _os.urandom(16)
    iterations = 260000
    dk = _hashlib.pbkdf2_hmac("sha256", str(password).encode("utf-8"), salt, iterations)
    return "pbkdf2_sha256${}${}${}".format(iterations, _base64.b64encode(salt).decode("ascii"), _base64.b64encode(dk).decode("ascii"))


def _fm488_verify_password(password, stored):
    import base64 as _base64, hashlib as _hashlib, hmac as _hmac
    try:
        parts = str(stored or "").split("$")
        if len(parts) != 4 or parts[0] != "pbkdf2_sha256":
            return False
        iterations = int(parts[1])
        salt = _base64.b64decode(parts[2].encode("ascii"))
        expected = _base64.b64decode(parts[3].encode("ascii"))
        actual = _hashlib.pbkdf2_hmac("sha256", str(password).encode("utf-8"), salt, iterations)
        return _hmac.compare_digest(actual, expected)
    except Exception:
        return False


def _fm488_password_valid(password):
    p = str(password or "")
    if len(p) < 8:
        return False, "Das Passwort muss mindestens 8 Zeichen enthalten."
    if not re.search(r"[A-Za-z]", p):
        return False, "Das Passwort muss mindestens einen Buchstaben enthalten."
    if not re.search(r"\d", p):
        return False, "Das Passwort muss mindestens eine Zahl enthalten."
    return True, ""


def _fm488_center_popup(self, win, width=430, height=235):
    try:
        win.update_idletasks()
        x = self.root.winfo_rootx() + max(0, int((self.root.winfo_width() - width) / 2))
        y = self.root.winfo_rooty() + max(0, int((self.root.winfo_height() - height) / 2))
        win.geometry(f"{width}x{height}+{x}+{y}")
    except Exception:
        pass


def _fm488_password_set_dialog(self, title, intro, require_old=False, old_hash=None):
    result = {"password": None}
    win = tk.Toplevel(self.root)
    win.title(title)
    win.configure(bg=BG)
    win.resizable(False, False)
    try:
        win.transient(self.root)
        win.grab_set()
    except Exception:
        pass
    frame = tk.Frame(win, bg=BG)
    frame.pack(fill="both", expand=True, padx=18, pady=16)
    tk.Label(frame, text=title, bg=BG, fg=BLUE, font=body_font(13, "bold")).pack(anchor="w", pady=(0, 6))
    tk.Label(frame, text=intro, bg=BG, fg=TEXT, font=body_font(10), wraplength=390, justify="left").pack(anchor="w", pady=(0, 10))
    old_var = tk.StringVar(value="")
    new_var = tk.StringVar(value="")
    rep_var = tk.StringVar(value="")
    if require_old:
        tk.Label(frame, text="Altes Passwort", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        tk.Entry(frame, textvariable=old_var, show="*", bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10)).pack(fill="x", ipady=4, pady=(0, 6))
    tk.Label(frame, text="Neues Passwort", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    e1 = tk.Entry(frame, textvariable=new_var, show="*", bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10))
    e1.pack(fill="x", ipady=4, pady=(0, 6))
    tk.Label(frame, text="Neues Passwort wiederholen", bg=BG, fg=TEXT2, font=body_font(9)).pack(anchor="w")
    e2 = tk.Entry(frame, textvariable=rep_var, show="*", bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(10))
    e2.pack(fill="x", ipady=4, pady=(0, 10))
    info_var = tk.StringVar(value="Mindestens 8 Zeichen, 1 Buchstabe und 1 Zahl.")
    tk.Label(frame, textvariable=info_var, bg=BG, fg=TEXT2, font=body_font(9), wraplength=390, justify="left").pack(anchor="w", pady=(0, 10))
    row = tk.Frame(frame, bg=BG)
    row.pack(fill="x")
    def save():
        if require_old and not _fm488_verify_password(old_var.get(), old_hash):
            info_var.set("Das alte Passwort ist nicht korrekt.")
            return
        if new_var.get() != rep_var.get():
            info_var.set("Die beiden neuen Passwoerter stimmen nicht ueberein.")
            return
        ok, msg = _fm488_password_valid(new_var.get())
        if not ok:
            info_var.set(msg)
            return
        result["password"] = new_var.get()
        try:
            win.grab_release()
        except Exception:
            pass
        win.destroy()
    tk.Button(row, text="Speichern", command=save, bg="#CFEAD6", fg=TEXT, font=body_font(10, "bold"), relief="solid", bd=1, padx=18, pady=6).pack(side="right", padx=(8, 0))
    tk.Button(row, text="Abbrechen", command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1, padx=18, pady=6).pack(side="right")
    e1.focus_set()
    try:
        win.bind("<Return>", lambda e: save())
        _fm488_center_popup(self, win, 455, 330 if require_old else 285)
        self.root.wait_window(win)
    except Exception:
        pass
    return result.get("password")


def _fm488_prompt_password_dialog(self, display_name):
    result = {"password": None}
    win = tk.Toplevel(self.root)
    win.title("FiBu Mate - Passwort")
    win.configure(bg=BG)
    win.resizable(False, False)
    try:
        win.transient(self.root)
        win.grab_set()
    except Exception:
        pass
    frame = tk.Frame(win, bg=BG)
    frame.pack(fill="both", expand=True, padx=18, pady=16)
    tk.Label(frame, text="Passwort erforderlich", bg=BG, fg=BLUE, font=body_font(13, "bold")).pack(anchor="w", pady=(0, 8))
    tk.Label(frame, text="Benutzer: " + str(display_name), bg=BG, fg=TEXT, font=body_font(10)).pack(anchor="w", pady=(0, 10))
    var = tk.StringVar(value="")
    ent = tk.Entry(frame, textvariable=var, show="*", bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(11))
    ent.pack(fill="x", ipady=5, pady=(0, 12))
    row = tk.Frame(frame, bg=BG)
    row.pack(fill="x")
    def ok():
        result["password"] = var.get()
        try:
            win.grab_release()
        except Exception:
            pass
        win.destroy()
    tk.Button(row, text="Anmelden", command=ok, bg=BLUE, fg=WHITE, font=body_font(10, "bold"), relief="solid", bd=0, padx=18, pady=6).pack(side="right", padx=(8, 0))
    tk.Button(row, text="Abbrechen", command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1, padx=18, pady=6).pack(side="right")
    def _pw_focus():
        try:
            win.lift(); win.focus_force(); ent.focus_force(); ent.icursor("end")
        except Exception:
            try: ent.focus_set()
            except Exception: pass
    try:
        win.bind("<Map>", lambda e: win.after(20, _pw_focus), add="+")
        win.after_idle(_pw_focus); win.after(60, _pw_focus); win.after(150, _pw_focus); win.after(350, _pw_focus)
    except Exception:
        try: ent.focus_set()
        except Exception: pass
    try:
        win.bind("<Return>", lambda e: ok())
        _fm488_center_popup(self, win, 430, 190)
        try: win.wait_visibility()
        except Exception: pass
        _pw_focus()
        self.root.wait_window(win)
    except Exception:
        pass
    return result.get("password")


def _fm488_finalize_login(self, key, username):
    users = self.user_data.setdefault("users", {})
    u = users[key]
    u.setdefault("display_name", username)
    u.setdefault("first_name", "")
    u.setdefault("full_name", " ".join(x for x in [u.get("first_name", "").strip(), u.get("display_name", username).strip()] if x).strip() or username)
    u.setdefault("favorites", [])
    u.setdefault("email", "")
    u["permission"] = ROLE_MIGRATION.get(u.get("permission", ROLE_E1), ROLE_E1)
    if key == SUPERUSER_KEY:
        u["permission"] = ROLE_E4
    self.ensure_permissions_defaults()
    self.current_user_key = key
    self.current_user_display = u.get("display_name", username)
    self.favorites = set(fav for fav in u.get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
    u["favorites"] = sorted(self.favorites)
    self.save_user_data()
    try:
        self.start_live_permissions_refresh()
    except Exception:
        pass
    self.page_history = []
    self.breadcrumb = []
    self.show_page("main", "Hauptmenue", add_to_history=False)


def _fm488_login_user(self, username):
    username = " ".join(str(username).strip().split())
    key = normalize_username(username)
    if not key:
        messagebox.showwarning("FiBu Mate", "Bitte einen Nutzernamen eingeben.")
        return
    try:
        self.user_data = _fm444_normalize_user_data(getattr(self, "user_data", None))
    except Exception:
        self.user_data = getattr(self, "user_data", {}) or {"users": {}}
    users = self.user_data.setdefault("users", {})
    if key not in users:
        try:
            if key == SUPERUSER_KEY and _fm444_bootstrap_superuser_if_allowed(self, username, key):
                pass
            else:
                messagebox.showwarning("FiBu Mate", "Benutzer nicht gefunden.\nBitte wende dich an eine Administratorin / einen Administrator.")
                return
        except Exception:
            messagebox.showwarning("FiBu Mate", "Benutzer nicht gefunden.\nBitte wende dich an eine Administratorin / einen Administrator.")
            return
    u = users[key]
    u.setdefault("display_name", username)
    u.setdefault("first_name", "")
    u.setdefault("full_name", " ".join(x for x in [u.get("first_name", "").strip(), u.get("display_name", username).strip()] if x).strip() or username)
    u.setdefault("email", "")
    auth = _fm488_normalize_auth(u)
    if not auth.get("password_hash") or bool(auth.get("password_must_set")):
        new_pw = _fm488_password_set_dialog(self, "Passwort festlegen", "Fuer diesen Benutzer ist noch kein Passwort gesetzt oder das Passwort wurde zurueckgesetzt. Bitte jetzt ein neues Passwort vergeben.", require_old=False)
        if not new_pw:
            return
        auth["password_hash"] = _fm488_hash_password(new_pw)
        auth["enabled"] = True
        auth["password_must_set"] = False
        auth["password_set_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        auth.pop("password_reset_at", None)
        auth.pop("password_reset_by", None)
        self.save_user_data()
        try:
            messagebox.showinfo("FiBu Mate", "Passwort wurde gespeichert.")
        except Exception:
            pass
    else:
        entered = _fm488_prompt_password_dialog(self, _fm488_user_fullname_from_record(key, u))
        if entered is None:
            return
        if not _fm488_verify_password(entered, auth.get("password_hash")):
            messagebox.showwarning("FiBu Mate", "Passwort ist nicht korrekt.")
            return
    _fm488_finalize_login(self, key, username)


def _fm488_can_reset_password(self, target_key):
    users = self.user_data.setdefault("users", {})
    if target_key not in users:
        return False
    actor_role = self.my_role() if hasattr(self, "my_role") else users.get(getattr(self, "current_user_key", ""), {}).get("permission", ROLE_E1)
    actor_rank = _fm488_role_rank_value(actor_role)
    target_rank = _fm488_role_rank_value(_fm488_user_role_from_data(users.get(target_key, {})))
    return actor_rank >= 3 and target_rank <= actor_rank


def _fm488_send_password_reset_mail(self, target_user):
    email = str((target_user or {}).get("email", "")).strip()
    if not email:
        return "no_email"
    display = _fm488_user_fullname_from_record("", target_user)
    subject = "FiBu Mate - Passwort wurde zurueckgesetzt"
    body = ("Hallo " + display + ",\n\n"
            "dein FiBu-Mate-Passwort wurde zurueckgesetzt.\n"
            "Bitte melde dich in FiBu Mate mit deinem Benutzernamen an. "
            "Beim naechsten Anmelden wirst du automatisch aufgefordert, ein neues Passwort festzulegen.\n\n"
            "Viele Gruesse\nFiBu Mate")
    try:
        import win32com.client as win32
        outlook = win32.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail.To = email
        mail.Subject = subject
        mail.Body = body
        try:
            mail.Send()
            return "sent"
        except Exception:
            mail.Display()
            return "displayed"
    except Exception:
        try:
            import urllib.parse as _up, webbrowser as _wb
            url = "mailto:" + _up.quote(email) + "?subject=" + _up.quote(subject) + "&body=" + _up.quote(body)
            _wb.open(url)
            return "displayed"
        except Exception:
            return "failed"


def _fm488_reset_password_for_user(self, target_key):
    users = self.user_data.setdefault("users", {})
    if target_key not in users:
        return
    if not _fm488_can_reset_password(self, target_key):
        messagebox.showwarning("Keine Berechtigung", "Du darfst dieses Passwort nicht zuruecksetzen.")
        return
    target = users[target_key]
    display = _fm488_user_fullname_from_record(target_key, target)
    if not messagebox.askyesno("Passwort zuruecksetzen", "Passwort wirklich zuruecksetzen fuer:\n\n" + display + "?"):
        return
    auth = _fm488_normalize_auth(target)
    auth["password_hash"] = None
    auth["enabled"] = False
    auth["password_must_set"] = True
    auth["password_reset_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    auth["password_reset_by"] = getattr(self, "current_user_key", "") or ""
    self.save_user_data()
    mail_status = _fm488_send_password_reset_mail(self, target)
    if mail_status == "sent":
        messagebox.showinfo("FiBu Mate", "Passwort wurde zurueckgesetzt und die E-Mail wurde automatisch versendet.")
    elif mail_status == "displayed":
        messagebox.showinfo("FiBu Mate", "Passwort wurde zurueckgesetzt. Die E-Mail wurde vorbereitet/geoeffnet.")
    elif mail_status == "no_email":
        messagebox.showwarning("FiBu Mate", "Passwort wurde zurueckgesetzt. Fuer diesen Benutzer ist keine E-Mail-Adresse gepflegt.")
    else:
        messagebox.showwarning("FiBu Mate", "Passwort wurde zurueckgesetzt. Die E-Mail-Benachrichtigung konnte nicht automatisch erstellt werden.")
    try:
        self.render_page()
    except Exception:
        pass


def _fm488_change_own_password(self):
    key = getattr(self, "current_user_key", "")
    users = self.user_data.setdefault("users", {})
    if key not in users:
        return
    auth = _fm488_normalize_auth(users[key])
    require_old = bool(auth.get("password_hash")) and not bool(auth.get("password_must_set"))
    new_pw = _fm488_password_set_dialog(self, "Eigenes Passwort aendern", "Bitte neues Passwort vergeben. Passwoerter werden nicht angezeigt und nicht im Klartext gespeichert.", require_old=require_old, old_hash=auth.get("password_hash"))
    if not new_pw:
        return
    auth["password_hash"] = _fm488_hash_password(new_pw)
    auth["enabled"] = True
    auth["password_must_set"] = False
    auth["password_set_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    self.save_user_data()
    messagebox.showinfo("FiBu Mate", "Passwort wurde geaendert.")


try:
    _fm488_prev_render_users_menu = FiBuMateApp.render_users_menu
except Exception:
    _fm488_prev_render_users_menu = None


def _fm488_render_users_menu(self):
    if not self.can_view_user_management():
        self.render_menu_text("Keine Berechtigung fuer die Benutzerverwaltung.")
        return
    users = self.user_data.setdefault("users", {})
    visible = [self.current_user_key] if self.my_role() == ROLE_E1 and self.current_user_key in users else sorted(users.keys())
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    top = 132
    vh = int(max(top + 260, h - 92) - top)
    cont = tk.Frame(self.root, bg=BG)
    self.widget_items.append(cont)
    self.canvas.create_window(0, top, window=cont, anchor="nw", width=w, height=vh)
    body = tk.Frame(cont, bg=BG)
    body.pack(fill="both", expand=True, padx=24)
    cv = tk.Canvas(body, bg=BG, highlightthickness=0, bd=0)
    sb = tk.Scrollbar(body, orient="vertical", command=cv.yview)
    frame = tk.Frame(cv, bg=BG)
    win = cv.create_window((0, 0), window=frame, anchor="nw")
    cv.configure(yscrollcommand=sb.set)
    cv.pack(side="left", fill="both", expand=True)
    sb.pack(side="right", fill="y")
    try:
        self.register_scroll_canvas(cv, sb)
    except Exception:
        pass
    def upd(e=None):
        cv.itemconfigure(win, width=max(1, cv.winfo_width()))
        cv.configure(scrollregion=cv.bbox("all"))
    frame.bind("<Configure>", upd)
    cv.bind("<Configure>", upd)
    tk.Label(frame, text="Benutzerverwaltung", font=self.zoomed_content_font(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=9, sticky="w", pady=(0, 14))
    heads = ["Benutzer", "Anzeigename/Nachname", "Vorname", "E-Mail", "Rolle", "Passwort", "", "", ""]
    for c, head in enumerate(heads):
        tk.Label(frame, text=head, bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=c, sticky="w", padx=(0, 14))
    is_super = self.current_user_key == SUPERUSER_KEY
    def can_edit(k):
        if is_super:
            return True
        if k == SUPERUSER_KEY:
            return False
        if k == self.current_user_key:
            return True
        return self.my_role() in (ROLE_E3, ROLE_E4) and self.role_rank(users.get(k, {}).get("permission", ROLE_E1)) <= self.role_rank(self.my_role())
    def role_vals(k):
        if is_super and k != SUPERUSER_KEY:
            return ROLE_ORDER
        if k == SUPERUSER_KEY:
            return [ROLE_E4]
        return [r for r in ROLE_ORDER if self.role_rank(r) <= self.role_rank(self.my_role()) and (self.my_role() == ROLE_E4 or r != ROLE_E4)]
    def save_row(k, nv, fv, ev, rv):
        if k not in users or not can_edit(k):
            return
        data = users[k]
        new_name = " ".join(nv.get().strip().split()) or data.get("display_name", k)
        new_key = normalize_username(new_name)
        if k == SUPERUSER_KEY:
            new_key = SUPERUSER_KEY
        if is_super and k != SUPERUSER_KEY and new_key != k:
            if not new_key or new_key in users:
                messagebox.showwarning("FiBu Mate", "Benutzername ungueltig oder bereits vorhanden.")
                return
            data = users.pop(k)
            users[new_key] = data
        data["display_name"] = new_name
        data["first_name"] = fv.get().strip()
        data["email"] = ev.get().strip()
        data["full_name"] = " ".join(x for x in [data.get("first_name", ""), data.get("display_name", new_key)] if x).strip() or data.get("display_name", new_key)
        role = rv.get() or data.get("permission", ROLE_E1)
        if role not in role_vals(new_key):
            messagebox.showwarning("FiBu Mate", "Diese Rolle darf nicht vergeben werden.")
            return
        data["permission"] = ROLE_E4 if new_key == SUPERUSER_KEY else role
        _fm488_normalize_auth(data)
        self.ensure_permissions_defaults()
        self.save_user_data()
        messagebox.showinfo("FiBu Mate", "Benutzerdaten wurden gespeichert.")
        self.render_page()
    def del_user(k):
        if k in (SUPERUSER_KEY, self.current_user_key) or not can_edit(k):
            messagebox.showwarning("FiBu Mate", "Dieser Benutzer kann nicht geloescht werden.")
            return
        if messagebox.askyesno("Benutzer loeschen", "Benutzer wirklich loeschen?\n\n" + users.get(k, {}).get("display_name", k)):
            users.pop(k, None)
            self.save_user_data()
            self.render_page()
    row = 2
    for k in visible:
        u = users[k]
        auth = _fm488_normalize_auth(u)
        edit = can_edit(k)
        nv = tk.StringVar(value=u.get("display_name", k))
        fv = tk.StringVar(value=u.get("first_name", ""))
        ev = tk.StringVar(value=u.get("email", ""))
        rv = tk.StringVar(value=u.get("permission", ROLE_E1))
        pwd_status = "gesetzt" if auth.get("password_hash") and not auth.get("password_must_set") else "muss gesetzt werden"
        tk.Label(frame, text=k, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=0, sticky="w", padx=(0, 14), pady=4)
        tk.Entry(frame, textvariable=nv, font=("Segoe UI", 10), width=22, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1, state="normal" if edit and (is_super or k not in (self.current_user_key, SUPERUSER_KEY)) else "disabled").grid(row=row, column=1, sticky="w", padx=(0, 14), pady=4)
        tk.Entry(frame, textvariable=fv, font=("Segoe UI", 10), width=18, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1, state="normal" if edit else "disabled").grid(row=row, column=2, sticky="w", padx=(0, 14), pady=4)
        tk.Entry(frame, textvariable=ev, font=("Segoe UI", 10), width=30, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1, state="normal" if edit else "disabled").grid(row=row, column=3, sticky="w", padx=(0, 14), pady=4)
        ttk.Combobox(frame, textvariable=rv, values=role_vals(k), state="readonly" if edit and (is_super or k != SUPERUSER_KEY) else "disabled", font=("Segoe UI", 10), width=24).grid(row=row, column=4, sticky="w", padx=(0, 14), pady=4)
        tk.Label(frame, text=pwd_status, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=5, sticky="w", padx=(0, 14), pady=4)
        tk.Button(frame, text="Speichern", command=lambda k=k, nv=nv, fv=fv, ev=ev, rv=rv: save_row(k, nv, fv, ev, rv), bg=BLUE if edit else GREY_DISABLED, fg="white", bd=0, width=10, padx=10, pady=5, state="normal" if edit else "disabled").grid(row=row, column=6, sticky="w", padx=(12, 8), pady=4)
        can_reset = _fm488_can_reset_password(self, k)
        tk.Button(frame, text="PW Reset", command=lambda k=k: _fm488_reset_password_for_user(self, k), bg=WHITE if can_reset else GREY_DISABLED, fg=TEXT, bd=1, width=10, padx=8, pady=5, state="normal" if can_reset else "disabled").grid(row=row, column=7, sticky="w", padx=(0, 8), pady=4)
        ok = edit and k not in (self.current_user_key, SUPERUSER_KEY) and self.my_role() in (ROLE_E3, ROLE_E4)
        tk.Button(frame, text="Loeschen", command=lambda k=k: del_user(k), bg=RED if ok else GREY_DISABLED, fg="white", bd=0, width=10, padx=8, pady=5, state="normal" if ok else "disabled").grid(row=row, column=8, sticky="w", padx=(0, 8), pady=4)
        row += 1
    upd()


try:
    _fm488_prev_render_settings_menu = FiBuMateApp.render_settings_menu
except Exception:
    _fm488_prev_render_settings_menu = None


def _fm488_render_settings_menu(self):
    modules = [("Passwort aendern", "__change_password__"), ("Kachelfarbe", "tile_colors"), ("Benutzerverwaltung", "users"), ("Berechtigungen", "permissions"), ("Versionsverlauf", "versions")]
    if self.my_role() == ROLE_E4:
        modules.append(("Lingo", "__lingo_popup__"))
    def open_mid(mid, ttl):
        if mid == "__change_password__":
            return _fm488_change_own_password(self)
        if mid == "__lingo_popup__":
            return self.show_lingo_popup()
        return self.show_page(mid, ttl, True)
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = max(250, min(330, int(w * 0.18)))
    tile_h = max(105, min(140, int(h * 0.13)))
    gap = max(24, int(w * 0.022))
    total = len(modules) * tile_w + (len(modules) - 1) * gap
    start = (w - total) / 2 + tile_w / 2
    y = y_pct(h, 56)
    for i, (title, mid) in enumerate(modules):
        tile = Tile(self.root, self, f"settings_{mid}", title, lambda m=mid, t=title: open_mid(m, t), favorite_enabled=False, icon_type="gear")
        tile.resize_tile(tile_w, tile_h)
        self.widget_items.append(tile)
        self.focusable_tiles.append(tile)
        self.canvas.create_window(start + i * (tile_w + gap), y, window=tile, anchor="center")
    self.draw_bottom_logo()


FiBuMateApp.login_user = _fm488_login_user
FiBuMateApp.render_users_menu = _fm488_render_users_menu
FiBuMateApp.render_settings_menu = _fm488_render_settings_menu
FiBuMateApp.show_change_password_popup = _fm488_change_own_password




# ------------------------------------------------------------------
# FIBU_MATE_SETTINGS_WRAP_AND_VERSION_E4_SAVE_V0489
# Datum: 2026-07-13
# Zweck:
# - Einstellungen: Kacheln umbrechen, sobald eine Zeile mehr als 4 Kacheln haette.
# - Versionsverlauf: Bearbeiten/Ausblenden/Loeschen ausschliesslich fuer E4.
# - Versionsverlauf-Bearbeitung: Speichern-Button explizit vorhanden und schreibt die Aenderungen.
# - Hinweis Passwortfunktion/G-Laufwerk bleibt technisch unveraendert: zentrale Benutzerdatei nach resolve_user_data_path().
# ------------------------------------------------------------------
FIBU_MATE_SETTINGS_WRAP_AND_VERSION_E4_SAVE_VERSION = "0.489"


def _fm489_is_e4_only(self):
    try:
        return str(self.my_role()) == ROLE_E4 or self.role_rank(self.my_role()) >= 4
    except Exception:
        pass
    try:
        users = getattr(self, 'user_data', {}).get('users', {})
        role = users.get(getattr(self, 'current_user_key', ''), {}).get('permission', ROLE_E1)
        return role == ROLE_E4 or int(ROLE_RANK.get(role, 1)) >= 4
    except Exception:
        return False


def _fm489_render_settings_menu(self):
    modules = [
        ("Passwort aendern", "__change_password__"),
        ("Kachelfarbe", "tile_colors"),
        ("Benutzerverwaltung", "users"),
        ("Berechtigungen", "permissions"),
        ("Versionsverlauf", "versions"),
    ]
    if self.my_role() == ROLE_E4:
        modules.append(("Lingo", "__lingo_popup__"))

    def open_mid(mid, ttl):
        if mid == "__change_password__":
            try:
                return _fm488_change_own_password(self)
            except Exception:
                return getattr(self, 'show_change_password_popup', lambda: None)()
        if mid == "__lingo_popup__":
            return self.show_lingo_popup()
        return self.show_page(mid, ttl, True)

    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = max(250, min(330, int(w * 0.18)))
    tile_h = max(105, min(140, int(h * 0.13)))
    gap_x = max(24, int(w * 0.022))
    gap_y = max(22, int(h * 0.025))
    max_cols = 4
    rows = [modules[i:i + max_cols] for i in range(0, len(modules), max_cols)]
    total_h = len(rows) * tile_h + max(0, len(rows) - 1) * gap_y
    start_y = max(185, int((h - total_h) * 0.50))
    for r, row_items in enumerate(rows):
        cols = len(row_items)
        total_w = cols * tile_w + max(0, cols - 1) * gap_x
        start_x = max(24, (w - total_w) / 2)
        y = start_y + r * (tile_h + gap_y)
        for c, (title, mid) in enumerate(row_items):
            tile = Tile(self.root, self, f"settings_{mid}", title, lambda m=mid, t=title: open_mid(m, t), favorite_enabled=False, icon_type="gear")
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(start_x + c * (tile_w + gap_x), y, window=tile, anchor="nw")
    self.draw_bottom_logo()


def _fm489_edit_version_dialog(self, index, entries):
    if not _fm489_is_e4_only(self):
        try:
            messagebox.showwarning('Versionsverlauf', 'Nur E4 darf Versionsverlauf-Eintraege bearbeiten.')
        except Exception:
            pass
        return
    try:
        entry = dict(entries[index])
    except Exception:
        return
    try:
        win = tk.Toplevel(self.root)
        win.title('Versionsverlauf bearbeiten')
        win.configure(bg=WHITE)
        win.geometry('760x510')
        win.resizable(True, True)
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        outer = tk.Frame(win, bg=WHITE)
        outer.pack(fill='both', expand=True, padx=16, pady=14)
        version_var = tk.StringVar(value=str(entry.get('version', '')))
        title_var = tk.StringVar(value=str(entry.get('title', '')))
        date_var = tk.StringVar(value=str(entry.get('date') or entry.get('published_at') or ''))
        tk.Label(outer, text='Version', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=version_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Titel / Kurzbeschreibung', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=title_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Datum / Veroeffentlichung', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=date_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Eintrag / Stichpunkte', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        txt_frame = tk.Frame(outer, bg=WHITE)
        txt_frame.pack(fill='both', expand=True, pady=(0, 10))
        txt = tk.Text(txt_frame, bg='#F8FAFC', fg=TEXT, font=body_font(11), relief='solid', bd=1, wrap='word')
        ys = tk.Scrollbar(txt_frame, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            txt.insert('1.0', '\n'.join(str(b) for b in bullets))
        else:
            txt.insert('1.0', str(bullets or entry.get('text', '') or ''))
        txt.pack(side='left', fill='both', expand=True)
        ys.pack(side='right', fill='y')
        br = tk.Frame(outer, bg=WHITE)
        br.pack(fill='x')
        def save():
            entry['version'] = version_var.get().strip()
            entry['title'] = title_var.get().strip()
            if date_var.get().strip():
                entry['date'] = date_var.get().strip()
            raw = txt.get('1.0', 'end-1c').strip()
            entry['bullets'] = [line.strip() for line in raw.splitlines() if line.strip()]
            entries[index] = entry
            if _fm450_save_version_entries_admin(self, entries):
                try:
                    messagebox.showinfo('Versionsverlauf', 'Aenderungen wurden gespeichert.')
                except Exception:
                    pass
                win.destroy()
                self.render_page()
        tk.Button(br, text='Speichern', command=save, bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0, 8), ipadx=18, ipady=5)
        tk.Button(br, text='Abbrechen', command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', ipadx=18, ipady=5)
        txt.focus_set()
    except Exception as exc:
        try:
            messagebox.showerror('Versionsverlauf', 'Bearbeiten nicht moeglich:\n' + str(exc))
        except Exception:
            pass


def _fm489_render_versions_menu(self):
    entries = _fm450_version_entries_raw(self)
    is_e4 = _fm489_is_e4_only(self)
    if not is_e4:
        entries = [e for e in entries if not bool(e.get('hidden'))]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    area_top = 132
    area_bottom = max(area_top + 260, h - 92)
    view_h = int(area_bottom - area_top)
    container = tk.Frame(self.root, bg=BG)
    self.widget_items.append(container)
    self.canvas.create_window(0, area_top, window=container, anchor='nw', width=w, height=view_h)
    scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0)
    scrollbar = tk.Scrollbar(container, orient='vertical', command=scroll_canvas.yview)
    content = tk.Frame(scroll_canvas, bg=BG)
    content_window = scroll_canvas.create_window((0, 0), window=content, anchor='nw')
    def update_scrollregion(_event=None):
        scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width()))
        scroll_canvas.configure(scrollregion=scroll_canvas.bbox('all'))
        try:
            self._sync_scrollbar_visibility(scroll_canvas, scrollbar)
        except Exception:
            pass
    content.bind('<Configure>', update_scrollregion)
    scroll_canvas.bind('<Configure>', update_scrollregion)
    scroll_canvas.configure(yscrollcommand=scrollbar.set)
    scroll_canvas.pack(side='left', fill='both', expand=True)
    scrollbar.pack(side='right', fill='y')
    try:
        self.register_scroll_canvas(scroll_canvas, scrollbar)
    except Exception:
        pass
    tk.Label(content, text='Versionsverlauf', bg=BG, fg=BLUE, font=body_font(20, weight='bold')).pack(anchor='w', padx=24, pady=(18, 4))
    source_txt = f"Versionsverlauf: lokale Historie plus zentrale Release-Historie ({self._central_version_history_path()})"
    tk.Label(content, text=source_txt, bg=BG, fg=TEXT2, font=body_font(9), wraplength=max(520, w - 70), justify='left').pack(anchor='w', padx=24, pady=(0, 10))
    if is_e4:
        tk.Label(content, text='E4-Administration aktiv: Eintraege koennen bearbeitet, ausgeblendet/eingeblendet oder geloescht werden.', bg=BG, fg=RED, font=body_font(10, weight='bold'), wraplength=max(520, w - 70), justify='left').pack(anchor='w', padx=24, pady=(0, 10))
    else:
        tk.Label(content, text='Nur E4 darf Versionsverlauf-Eintraege bearbeiten.', bg=BG, fg=TEXT2, font=body_font(10), wraplength=max(520, w - 70), justify='left').pack(anchor='w', padx=24, pady=(0, 10))
    if not entries:
        tk.Label(content, text='Noch kein Versionsverlauf vorhanden.', bg=BG, fg=TEXT2, font=body_font(12)).pack(anchor='w', padx=24, pady=24)
        return
    for idx, entry in enumerate(entries):
        card = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=1)
        card.pack(fill='x', padx=24, pady=8)
        top = tk.Frame(card, bg=WHITE)
        top.pack(fill='x', padx=14, pady=(10, 4))
        version = str(entry.get('version', '')).strip() or 'ohne Version'
        title = str(entry.get('title', '')).strip()
        hidden = bool(entry.get('hidden'))
        head = version + (f' - {title}' if title else '') + ('  [AUSGEBLENDET]' if hidden else '')
        tk.Label(top, text=head, bg=WHITE, fg=(RED if hidden else BLUE), font=body_font(13, weight='bold')).pack(side='left', anchor='w')
        date_txt = entry.get('date') or entry.get('published_at') or ''
        if date_txt:
            tk.Label(top, text=str(date_txt), bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            for b in bullets:
                tk.Label(card, text='• ' + str(b), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520, w - 120), justify='left').pack(anchor='w', padx=18, pady=1)
        else:
            tk.Label(card, text=str(bullets), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520, w - 120), justify='left').pack(anchor='w', padx=18, pady=1)
        if is_e4:
            br = tk.Frame(card, bg=WHITE)
            br.pack(fill='x', padx=14, pady=(8, 10))
            def edit(i=idx):
                _fm489_edit_version_dialog(self, i, entries)
            def toggle(i=idx):
                entries[i]['hidden'] = not bool(entries[i].get('hidden'))
                if _fm450_save_version_entries_admin(self, entries):
                    self.render_page()
            def delete(i=idx):
                try:
                    ok = messagebox.askyesno('Versionsverlauf', 'Diesen Versionsverlauf-Eintrag unwiderruflich loeschen?')
                except Exception:
                    ok = False
                if ok:
                    try:
                        entries.pop(i)
                    except Exception:
                        pass
                    if _fm450_save_version_entries_admin(self, entries):
                        self.render_page()
            tk.Button(br, text='Bearbeiten', command=edit, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6), ipadx=10, ipady=2)
            tk.Button(br, text=('Einblenden' if hidden else 'Ausblenden'), command=toggle, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6), ipadx=10, ipady=2)
            tk.Button(br, text='Loeschen', command=delete, bg='#FEE2E2', fg=RED, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='left', ipadx=10, ipady=2)

FiBuMateApp.render_settings_menu = _fm489_render_settings_menu
FiBuMateApp.render_versions_menu = _fm489_render_versions_menu




# ------------------------------------------------------------------
# FIBU_MATE_VERSIONS_INFO_ONLY_APPLY_BUTTON_V0490
# Datum: 2026-07-13
# Zweck:
# - Versionsverlauf wieder nur unter Informationen anzeigen, nicht unter Einstellungen.
# - Einstellungen-Kacheln behalten den Umbruch nach maximal 4 Kacheln je Zeile.
# - Versionsverlauf-Bearbeitung fuer E4 mit explizitem Button "Uebernehmen".
# ------------------------------------------------------------------
FIBU_MATE_VERSIONS_INFO_ONLY_APPLY_BUTTON_VERSION = "0.490"


def _fm490_is_e4_only(self):
    try:
        return str(self.my_role()) == ROLE_E4 or self.role_rank(self.my_role()) >= 4
    except Exception:
        pass
    try:
        users = getattr(self, 'user_data', {}).get('users', {})
        role = users.get(getattr(self, 'current_user_key', ''), {}).get('permission', ROLE_E1)
        return role == ROLE_E4 or int(ROLE_RANK.get(role, 1)) >= 4
    except Exception:
        return False


def _fm490_render_settings_menu(self):
    # Versionsverlauf ist bewusst NICHT enthalten; er bleibt ausschliesslich unter Informationen.
    modules = [
        ("Passwort aendern", "__change_password__"),
        ("Kachelfarbe", "tile_colors"),
        ("Benutzerverwaltung", "users"),
        ("Berechtigungen", "permissions"),
    ]
    if self.my_role() == ROLE_E4:
        modules.append(("Lingo", "__lingo_popup__"))

    def open_mid(mid, ttl):
        if mid == "__change_password__":
            try:
                return _fm488_change_own_password(self)
            except Exception:
                return getattr(self, 'show_change_password_popup', lambda: None)()
        if mid == "__lingo_popup__":
            return self.show_lingo_popup()
        return self.show_page(mid, ttl, True)

    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = max(250, min(330, int(w * 0.18)))
    tile_h = max(105, min(140, int(h * 0.13)))
    gap_x = max(24, int(w * 0.022))
    gap_y = max(22, int(h * 0.025))
    max_cols = 4
    rows = [modules[i:i + max_cols] for i in range(0, len(modules), max_cols)]
    total_h = len(rows) * tile_h + max(0, len(rows) - 1) * gap_y
    start_y = max(185, int((h - total_h) * 0.50))
    for r, row_items in enumerate(rows):
        cols = len(row_items)
        total_w = cols * tile_w + max(0, cols - 1) * gap_x
        start_x = max(24, (w - total_w) / 2)
        y = start_y + r * (tile_h + gap_y)
        for c, (title, mid) in enumerate(row_items):
            tile = Tile(self.root, self, f"settings_{mid}", title, lambda m=mid, t=title: open_mid(m, t), favorite_enabled=False, icon_type="gear")
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(start_x + c * (tile_w + gap_x), y, window=tile, anchor="nw")
    self.draw_bottom_logo()


def _fm490_edit_version_dialog(self, index, entries):
    if not _fm490_is_e4_only(self):
        try:
            messagebox.showwarning('Versionsverlauf', 'Nur E4 darf Versionsverlauf-Eintraege bearbeiten.')
        except Exception:
            pass
        return
    try:
        entry = dict(entries[index])
    except Exception:
        return
    try:
        win = tk.Toplevel(self.root)
        win.title('Versionsverlauf bearbeiten')
        win.configure(bg=WHITE)
        win.geometry('760x520')
        win.resizable(True, True)
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass
        outer = tk.Frame(win, bg=WHITE)
        outer.pack(fill='both', expand=True, padx=16, pady=14)
        tk.Label(outer, text='Versionsverlauf-Eintrag bearbeiten', bg=WHITE, fg=BLUE, font=body_font(14, weight='bold')).pack(anchor='w', pady=(0, 10))
        version_var = tk.StringVar(value=str(entry.get('version', '')))
        title_var = tk.StringVar(value=str(entry.get('title', '')))
        date_var = tk.StringVar(value=str(entry.get('date') or entry.get('published_at') or ''))
        tk.Label(outer, text='Version', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=version_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Titel / Kurzbeschreibung', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=title_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Datum / Veroeffentlichung', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        tk.Entry(outer, textvariable=date_var, bg=WHITE, fg=TEXT, font=body_font(11), relief='solid', bd=1).pack(fill='x', ipady=4, pady=(0, 8))
        tk.Label(outer, text='Eintrag / Stichpunkte', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        txt_frame = tk.Frame(outer, bg=WHITE)
        txt_frame.pack(fill='both', expand=True, pady=(0, 10))
        txt = tk.Text(txt_frame, bg='#F8FAFC', fg=TEXT, font=body_font(11), relief='solid', bd=1, wrap='word')
        ys = tk.Scrollbar(txt_frame, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            txt.insert('1.0', '\n'.join(str(b) for b in bullets))
        else:
            txt.insert('1.0', str(bullets or entry.get('text', '') or ''))
        txt.pack(side='left', fill='both', expand=True)
        ys.pack(side='right', fill='y')
        status_var = tk.StringVar(value='')
        tk.Label(outer, textvariable=status_var, bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor='w', pady=(0, 8))
        br = tk.Frame(outer, bg=WHITE)
        br.pack(fill='x')

        def apply_changes(close_after=False):
            entry['version'] = version_var.get().strip()
            entry['title'] = title_var.get().strip()
            if date_var.get().strip():
                entry['date'] = date_var.get().strip()
            raw = txt.get('1.0', 'end-1c').strip()
            entry['bullets'] = [line.strip() for line in raw.splitlines() if line.strip()]
            entries[index] = entry
            if _fm450_save_version_entries_admin(self, entries):
                status_var.set('Aenderungen wurden uebernommen.')
                if close_after:
                    win.destroy()
                    self.render_page()

        # Explizit gewuenschter Button: Uebernehmen
        tk.Button(br, text='Uebernehmen', command=lambda: apply_changes(False), bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='left', padx=(0, 8), ipadx=18, ipady=5)
        tk.Button(br, text='Speichern und schliessen', command=lambda: apply_changes(True), bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', padx=(0, 8), ipadx=18, ipady=5)
        tk.Button(br, text='Abbrechen', command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='left', ipadx=18, ipady=5)
        txt.focus_set()
    except Exception as exc:
        try:
            messagebox.showerror('Versionsverlauf', 'Bearbeiten nicht moeglich:\n' + str(exc))
        except Exception:
            pass


def _fm490_render_versions_menu(self):
    entries = _fm450_version_entries_raw(self)
    is_e4 = _fm490_is_e4_only(self)
    if not is_e4:
        entries = [e for e in entries if not bool(e.get('hidden'))]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    area_top = 132
    area_bottom = max(area_top + 260, h - 92)
    view_h = int(area_bottom - area_top)
    container = tk.Frame(self.root, bg=BG)
    self.widget_items.append(container)
    self.canvas.create_window(0, area_top, window=container, anchor='nw', width=w, height=view_h)
    scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0)
    scrollbar = tk.Scrollbar(container, orient='vertical', command=scroll_canvas.yview)
    content = tk.Frame(scroll_canvas, bg=BG)
    content_window = scroll_canvas.create_window((0, 0), window=content, anchor='nw')
    def update_scrollregion(_event=None):
        scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width()))
        scroll_canvas.configure(scrollregion=scroll_canvas.bbox('all'))
        try:
            self._sync_scrollbar_visibility(scroll_canvas, scrollbar)
        except Exception:
            pass
    content.bind('<Configure>', update_scrollregion)
    scroll_canvas.bind('<Configure>', update_scrollregion)
    scroll_canvas.configure(yscrollcommand=scrollbar.set)
    scroll_canvas.pack(side='left', fill='both', expand=True)
    scrollbar.pack(side='right', fill='y')
    try:
        self.register_scroll_canvas(scroll_canvas, scrollbar)
    except Exception:
        pass
    tk.Label(content, text='Versionsverlauf', bg=BG, fg=BLUE, font=body_font(20, weight='bold')).pack(anchor='w', padx=24, pady=(18, 4))
    source_txt = f"Versionsverlauf: lokale Historie plus zentrale Release-Historie ({self._central_version_history_path()})"
    tk.Label(content, text=source_txt, bg=BG, fg=TEXT2, font=body_font(9), wraplength=max(520, w - 70), justify='left').pack(anchor='w', padx=24, pady=(0, 10))
    if is_e4:
        tk.Label(content, text='E4-Administration aktiv: Bearbeiten, Ausblenden/Einblenden und Loeschen sind nur hier unter Informationen verfuegbar.', bg=BG, fg=RED, font=body_font(10, weight='bold'), wraplength=max(520, w - 70), justify='left').pack(anchor='w', padx=24, pady=(0, 10))
    if not entries:
        tk.Label(content, text='Noch kein Versionsverlauf vorhanden.', bg=BG, fg=TEXT2, font=body_font(12)).pack(anchor='w', padx=24, pady=24)
        return
    for idx, entry in enumerate(entries):
        card = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=1)
        card.pack(fill='x', padx=24, pady=8)
        top = tk.Frame(card, bg=WHITE)
        top.pack(fill='x', padx=14, pady=(10, 4))
        version = str(entry.get('version', '')).strip() or 'ohne Version'
        title = str(entry.get('title', '')).strip()
        hidden = bool(entry.get('hidden'))
        head = version + (f' - {title}' if title else '') + ('  [AUSGEBLENDET]' if hidden else '')
        tk.Label(top, text=head, bg=WHITE, fg=(RED if hidden else BLUE), font=body_font(13, weight='bold')).pack(side='left', anchor='w')
        date_txt = entry.get('date') or entry.get('published_at') or ''
        if date_txt:
            tk.Label(top, text=str(date_txt), bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side='right')
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            for b in bullets:
                tk.Label(card, text='• ' + str(b), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520, w - 120), justify='left').pack(anchor='w', padx=18, pady=1)
        else:
            tk.Label(card, text=str(bullets), bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(520, w - 120), justify='left').pack(anchor='w', padx=18, pady=1)
        if is_e4:
            br = tk.Frame(card, bg=WHITE)
            br.pack(fill='x', padx=14, pady=(8, 10))
            def edit(i=idx):
                _fm490_edit_version_dialog(self, i, entries)
            def toggle(i=idx):
                entries[i]['hidden'] = not bool(entries[i].get('hidden'))
                if _fm450_save_version_entries_admin(self, entries):
                    self.render_page()
            def delete(i=idx):
                try:
                    ok = messagebox.askyesno('Versionsverlauf', 'Diesen Versionsverlauf-Eintrag unwiderruflich loeschen?')
                except Exception:
                    ok = False
                if ok:
                    try:
                        entries.pop(i)
                    except Exception:
                        pass
                    if _fm450_save_version_entries_admin(self, entries):
                        self.render_page()
            tk.Button(br, text='Bearbeiten', command=edit, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6), ipadx=10, ipady=2)
            tk.Button(br, text=('Einblenden' if hidden else 'Ausblenden'), command=toggle, bg=WHITE, fg=TEXT, font=body_font(9), relief='solid', bd=1).pack(side='left', padx=(0, 6), ipadx=10, ipady=2)
            tk.Button(br, text='Loeschen', command=delete, bg='#FEE2E2', fg=RED, font=body_font(9, weight='bold'), relief='solid', bd=1).pack(side='left', ipadx=10, ipady=2)

FiBuMateApp.render_settings_menu = _fm490_render_settings_menu
FiBuMateApp.render_versions_menu = _fm490_render_versions_menu




# ------------------------------------------------------------------
# FIBU_MATE_VERSION_DIALOG_HEADER_APPLY_V0491
# Datum: 2026-07-13
# Zweck:
# - Versionsverlauf bleibt nur unter Informationen, nicht unter Einstellungen.
# - Bearbeiten-Popup im Versionsverlauf: Uebernehmen/Speichern direkt in der Kopfzeile.
# - Der Button ist sichtbar, auch wenn das Textfeld gross ist.
# ------------------------------------------------------------------
FIBU_MATE_VERSION_DIALOG_HEADER_APPLY_VERSION = "0.491"


def _fm491_is_e4_only(self):
    try:
        return str(self.my_role()) == ROLE_E4 or self.role_rank(self.my_role()) >= 4
    except Exception:
        pass
    try:
        users = getattr(self, 'user_data', {}).get('users', {})
        role = users.get(getattr(self, 'current_user_key', ''), {}).get('permission', ROLE_E1)
        return role == ROLE_E4 or int(ROLE_RANK.get(role, 1)) >= 4
    except Exception:
        return False


def _fm491_render_settings_menu(self):
    # Versionsverlauf bleibt bewusst ausschliesslich unter Informationen.
    modules = [
        ("Passwort aendern", "__change_password__"),
        ("Kachelfarbe", "tile_colors"),
        ("Benutzerverwaltung", "users"),
        ("Berechtigungen", "permissions"),
    ]
    if self.my_role() == ROLE_E4:
        modules.append(("Lingo", "__lingo_popup__"))

    def open_mid(mid, ttl):
        if mid == "__change_password__":
            try:
                return _fm488_change_own_password(self)
            except Exception:
                return getattr(self, 'show_change_password_popup', lambda: None)()
        if mid == "__lingo_popup__":
            return self.show_lingo_popup()
        return self.show_page(mid, ttl, True)

    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = max(250, min(330, int(w * 0.18)))
    tile_h = max(105, min(140, int(h * 0.13)))
    gap_x = max(24, int(w * 0.022))
    gap_y = max(22, int(h * 0.025))
    max_cols = 4
    rows = [modules[i:i + max_cols] for i in range(0, len(modules), max_cols)]
    total_h = len(rows) * tile_h + max(0, len(rows) - 1) * gap_y
    start_y = max(185, int((h - total_h) * 0.50))
    for r, row_items in enumerate(rows):
        cols = len(row_items)
        total_w = cols * tile_w + max(0, cols - 1) * gap_x
        start_x = max(24, (w - total_w) / 2)
        y = start_y + r * (tile_h + gap_y)
        for c, (title, mid) in enumerate(row_items):
            tile = Tile(self.root, self, f"settings_{mid}", title, lambda m=mid, t=title: open_mid(m, t), favorite_enabled=False, icon_type="gear")
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(start_x + c * (tile_w + gap_x), y, window=tile, anchor="nw")
    self.draw_bottom_logo()


def _fm491_edit_version_dialog(self, index, entries):
    if not _fm491_is_e4_only(self):
        try:
            messagebox.showwarning('Versionsverlauf', 'Nur E4 darf Versionsverlauf-Eintraege bearbeiten.')
        except Exception:
            pass
        return
    try:
        entry = dict(entries[index])
    except Exception:
        return
    try:
        win = tk.Toplevel(self.root)
        win.title('Versionsverlauf bearbeiten')
        win.configure(bg=WHITE)
        win.geometry('780x560')
        win.minsize(720, 480)
        win.resizable(True, True)
        try:
            win.transient(self.root)
            win.grab_set()
        except Exception:
            pass

        outer = tk.Frame(win, bg=WHITE)
        outer.pack(fill='both', expand=True, padx=14, pady=12)

        header = tk.Frame(outer, bg=WHITE)
        header.pack(fill='x', pady=(0, 10))
        tk.Label(header, text='Versionsverlauf-Eintrag bearbeiten', bg=WHITE, fg=BLUE, font=body_font(13, weight='bold')).pack(side='left', anchor='w')
        status_var = tk.StringVar(value='')

        version_var = tk.StringVar(value=str(entry.get('version', '')))
        title_var = tk.StringVar(value=str(entry.get('title', '')))
        date_var = tk.StringVar(value=str(entry.get('date') or entry.get('published_at') or ''))

        form = tk.Frame(outer, bg=WHITE)
        form.pack(fill='x', pady=(0, 8))
        tk.Label(form, text='Version', bg=WHITE, fg=TEXT2, font=body_font(10)).grid(row=0, column=0, sticky='w')
        tk.Entry(form, textvariable=version_var, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).grid(row=1, column=0, sticky='ew', ipady=4, pady=(0, 8))
        tk.Label(form, text='Titel / Kurzbeschreibung', bg=WHITE, fg=TEXT2, font=body_font(10)).grid(row=2, column=0, sticky='w')
        tk.Entry(form, textvariable=title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).grid(row=3, column=0, sticky='ew', ipady=4, pady=(0, 8))
        tk.Label(form, text='Datum / Veroeffentlichung', bg=WHITE, fg=TEXT2, font=body_font(10)).grid(row=4, column=0, sticky='w')
        tk.Entry(form, textvariable=date_var, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).grid(row=5, column=0, sticky='ew', ipady=4)
        form.grid_columnconfigure(0, weight=1)

        tk.Label(outer, text='Eintrag / Stichpunkte', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w')
        txt_frame = tk.Frame(outer, bg=WHITE)
        txt_frame.pack(fill='both', expand=True, pady=(2, 8))
        txt = tk.Text(txt_frame, bg='#F8FAFC', fg=TEXT, font=body_font(10), relief='solid', bd=1, wrap='word')
        ys = tk.Scrollbar(txt_frame, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        bullets = entry.get('bullets', [])
        if isinstance(bullets, list):
            txt.insert('1.0', '\n'.join(str(b) for b in bullets))
        else:
            txt.insert('1.0', str(bullets or entry.get('text', '') or ''))
        txt.pack(side='left', fill='both', expand=True)
        ys.pack(side='right', fill='y')

        status = tk.Label(outer, textvariable=status_var, bg=WHITE, fg=TEXT2, font=body_font(9))
        status.pack(anchor='w', pady=(0, 6))

        def apply_changes(close_after=False):
            entry['version'] = version_var.get().strip()
            entry['title'] = title_var.get().strip()
            if date_var.get().strip():
                entry['date'] = date_var.get().strip()
            raw = txt.get('1.0', 'end-1c').strip()
            entry['bullets'] = [line.strip() for line in raw.splitlines() if line.strip()]
            entries[index] = entry
            if _fm450_save_version_entries_admin(self, entries):
                status_var.set('Aenderungen wurden uebernommen.')
                if close_after:
                    win.destroy()
                    self.render_page()

        # Kopfzeilen-Button: sichtbar direkt oben rechts im Popup.
        tk.Button(header, text='Uebernehmen', command=lambda: apply_changes(False), bg='#CFEAD6', fg=TEXT, font=body_font(10, weight='bold'), relief='solid', bd=1).pack(side='right', padx=(8, 0), ipadx=16, ipady=4)
        tk.Button(header, text='Speichern', command=lambda: apply_changes(True), bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='right', padx=(8, 0), ipadx=16, ipady=4)
        tk.Button(header, text='Abbrechen', command=win.destroy, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1).pack(side='right', ipadx=16, ipady=4)
        txt.focus_set()
    except Exception as exc:
        try:
            messagebox.showerror('Versionsverlauf', 'Bearbeiten nicht moeglich:\n' + str(exc))
        except Exception:
            pass

# Wichtig: _fm490_render_versions_menu ruft den globalen Namen _fm490_edit_version_dialog auf.
# Deshalb wird dieser globale Name final auf die neue Kopfzeilen-Variante umgebogen.
_fm490_edit_version_dialog = _fm491_edit_version_dialog
try:
    _fm489_edit_version_dialog = _fm491_edit_version_dialog
except Exception:
    pass

FiBuMateApp.render_settings_menu = _fm491_render_settings_menu
try:
    FiBuMateApp.render_versions_menu = _fm490_render_versions_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_FOOTER_INTERSPORT_MARK_V0492
# Datum: 2026-07-13
# Zweck:
# - Unteren INTERSPORT-Schriftzug dezent durch INTERSPORT_banner_mark_CMYK.png ersetzen.
# ------------------------------------------------------------------
FIBU_MATE_FOOTER_INTERSPORT_MARK_VERSION = "0.492"

try:
    _fm492_prev_draw_bottom_logo = FiBuMateApp.draw_bottom_logo
except Exception:
    _fm492_prev_draw_bottom_logo = None


def _fm492_footer_logo_candidates():
    candidates = []
    try:
        candidates.append(r"C:\python\bin\Imgs\INTERSPORT_banner_mark_CMYK.png")
    except Exception:
        pass
    try:
        candidates.append(os.path.join(IMG_DIR, "INTERSPORT_banner_mark_CMYK.png"))
    except Exception:
        pass
    try:
        candidates.append(os.path.join(BIN_DIR, "Imgs", "INTERSPORT_banner_mark_CMYK.png"))
    except Exception:
        pass
    out = []
    seen = set()
    for p in candidates:
        if p and p not in seen:
            seen.add(p); out.append(p)
    return out


def _fm492_draw_bottom_logo(self):
    try:
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if PIL_AVAILABLE:
            for path in _fm492_footer_logo_candidates():
                if os.path.exists(path):
                    img = Image.open(path)
                    # Dezent: klar lesbar, aber kompakt am unteren Rand.
                    max_w = max(180, min(360, int(w * 0.22)))
                    max_h = 42
                    ph = ImageTk.PhotoImage(resize_keep_ratio(img, max_w, max_h))
                    self.image_refs.append(ph)
                    self.canvas.create_image(w / 2, h - 24, image=ph)
                    return
        # Fallback beibehalten, falls die Datei lokal nicht vorhanden ist.
        if _fm492_prev_draw_bottom_logo:
            return _fm492_prev_draw_bottom_logo(self)
        self.canvas.create_text(w / 2, h - 24, text="INTERSPORT", font=("Segoe UI", 18, "bold"), fill=BLUE)
    except Exception:
        try:
            if _fm492_prev_draw_bottom_logo:
                return _fm492_prev_draw_bottom_logo(self)
        except Exception:
            pass

try:
    FiBuMateApp.draw_bottom_logo = _fm492_draw_bottom_logo
    FiBuMateApp.draw_intersport_logo_above_footer = lambda self, show_mini_logo=True: self.draw_bottom_logo()
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_KONTENAUSWERTUNG_V0501
# Datum: 2026-07-13
# Zweck:
# - Neues Tool "Kontenauswertung" unter Tools -> Tools - Treasury.
# - Import von SFIRM-Excel-Exporten mit Entgeltabschluss-Aufschlüsselung.
# - SQLite-Datenbank für Monatsabschlüsse, Gebührenpositionen, Stückkosten und Entwicklung.
# - Auswertung mit Monatsübersicht, Detailpositionen, Kosten-/Mengen-/Stückkostenentwicklung, Plausibilitätsprüfung und Excel-Export.
# - Hauptmenü: Tools - Treasury ergänzt; Dateiausgabe in unterste Modulleiste rechts neben Einstellungen verschoben.
# ------------------------------------------------------------------
FIBU_MATE_KONTENAUSWERTUNG_VERSION = "0.501"


def _fm501_konto_db_path():
    for base in (
        os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Database"),
        os.path.join(SCRIPT_DIR, "bin", "Fibu_Mate_Doc", "Database"),
        os.path.join(SCRIPT_DIR, "Database"),
    ):
        try:
            os.makedirs(base, exist_ok=True)
            return os.path.join(base, "kontenauswertung.sqlite3")
        except Exception:
            continue
    return os.path.join(os.getcwd(), "kontenauswertung.sqlite3")


def _fm501_sqlite():
    import sqlite3
    return sqlite3


def _fm501_db_connect():
    sqlite3 = _fm501_sqlite()
    con = sqlite3.connect(_fm501_konto_db_path())
    con.row_factory = sqlite3.Row
    return con


def _fm501_db_init():
    con = _fm501_db_connect()
    try:
        cur = con.cursor()
        cur.execute('''CREATE TABLE IF NOT EXISTS imports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            file_mtime TEXT,
            imported_at TEXT NOT NULL,
            imported_by TEXT,
            row_count INTEGER DEFAULT 0,
            closing_count INTEGER DEFAULT 0,
            detail_count INTEGER DEFAULT 0,
            file_hash TEXT,
            UNIQUE(filename, file_hash)
        )''')
        cur.execute('''CREATE TABLE IF NOT EXISTS closings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            import_id INTEGER,
            source_file TEXT,
            source_row INTEGER,
            buchungsdatum TEXT,
            wertstellungsdatum TEXT,
            abschlussdatum TEXT,
            monat_key TEXT,
            zeitraum_von TEXT,
            zeitraum_bis TEXT,
            rechnungsnummer TEXT,
            betrag_brutto REAL,
            betrag_netto REAL,
            ust REAL,
            waehrung TEXT,
            gvc TEXT,
            btc TEXT,
            raw_text TEXT,
            created_at TEXT NOT NULL,
            UNIQUE(rechnungsnummer, abschlussdatum)
        )''')
        cur.execute('''CREATE TABLE IF NOT EXISTS fee_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            closing_id INTEGER NOT NULL,
            abschlussdatum TEXT,
            monat_key TEXT,
            geschaeftsvorfall TEXT,
            positionstyp TEXT,
            anzahl INTEGER,
            stueckkosten REAL,
            gesamtbetrag REAL,
            berechneter_betrag REAL,
            abweichung REAL,
            plausibel INTEGER,
            sort_order INTEGER,
            raw_line TEXT,
            UNIQUE(closing_id, geschaeftsvorfall, positionstyp, sort_order)
        )''')
        con.commit()
    finally:
        con.close()


def _fm501_de_float(value):
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).strip()
    if not s:
        return None
    neg = s.endswith('-') or s.startswith('-')
    s = s.strip('-').replace('EUR', '').replace('€', '').strip()
    s = s.replace('.', '').replace(',', '.')
    try:
        val = float(s)
        return -val if neg and False else val
    except Exception:
        return None


def _fm501_date_iso(value):
    if not value:
        return ""
    try:
        if hasattr(value, 'date'):
            return value.date().isoformat()
    except Exception:
        pass
    text = str(value).strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except Exception:
            pass
    return text


def _fm501_date_de_to_iso(text):
    if not text:
        return ""
    try:
        return datetime.strptime(str(text).strip(), "%d.%m.%Y").date().isoformat()
    except Exception:
        return str(text).strip()


def _fm501_month_key_from_iso(iso_date):
    try:
        d = datetime.strptime(str(iso_date), "%Y-%m-%d")
        return d.strftime("%m/%y")
    except Exception:
        return str(iso_date or "")[:7]


def _fm501_hash_file(path):
    try:
        h = hashlib.sha256()
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b''):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return ""


def _fm501_parse_entgelt_text(text):
    import re as _re
    text = str(text or "")
    result = {
        'abschlussdatum': '', 'zeitraum_von': '', 'zeitraum_bis': '', 'betrag_netto': None,
        'ust': None, 'betrag_brutto_text': None, 'rechnungsnummer': '', 'items': []
    }
    m = _re.search(r'Abrechnung\s+(\d{2}\.\d{2}\.\d{4})', text)
    if m:
        result['abschlussdatum'] = _fm501_date_de_to_iso(m.group(1))
    m = _re.search(r'Entgelte\s+vom\s+(\d{2}\.\d{2}\.\d{4})\s+bis\s+(\d{2}\.\d{2}\.\d{4})\s+([\d\.]+,\d{2})-', text)
    if m:
        result['zeitraum_von'] = _fm501_date_de_to_iso(m.group(1))
        result['zeitraum_bis'] = _fm501_date_de_to_iso(m.group(2))
        result['betrag_netto'] = _fm501_de_float(m.group(3))
    m = _re.search(r'USt:\s*([\d\.]+,\d{2})-', text)
    if m:
        result['ust'] = _fm501_de_float(m.group(1))
    m = _re.search(r'Abrechnung\s+\d{2}\.\d{2}\.\d{4}\s+([\d\.]+,\d{2})-', text.split('Rechnungsnummer')[0])
    if m:
        result['betrag_brutto_text'] = _fm501_de_float(m.group(1))
    m = _re.search(r'Rechnungsnummer:\s*(\S+)', text)
    if m:
        result['rechnungsnummer'] = m.group(1).strip()

    gf_pattern = _re.compile(r'(?P<label>[^\n]+)\n\s*Geschäftsvorfälle\s+(?P<count>[\d\.]+)\s*x\s*(?P<unit>\d+,\d{2})\s+(?P<total>\d+,\d{2})-', _re.M)
    occupied = []
    order = 0
    for g in gf_pattern.finditer(text):
        order += 1
        label = g.group('label').strip()
        count = int(g.group('count').replace('.', ''))
        unit = _fm501_de_float(g.group('unit'))
        total = _fm501_de_float(g.group('total'))
        calc = round((count or 0) * (unit or 0), 2)
        diff = round((total or 0) - calc, 2)
        result['items'].append({
            'geschaeftsvorfall': label, 'positionstyp': 'Geschäftsvorfälle', 'anzahl': count,
            'stueckkosten': unit, 'gesamtbetrag': total, 'berechneter_betrag': calc,
            'abweichung': diff, 'plausibel': 1 if abs(diff) < 0.011 else 0,
            'sort_order': order, 'raw_line': g.group(0).strip()
        })
        occupied.append((g.start(), g.end()))

    # Weitere entgeltrelevante Zeilen wie "Kontoauszug signiert 22 x 1,10 24,20-".
    other_pattern = _re.compile(r'(?P<label>[^\n\d][^\n]*?)\s+(?P<count>[\d\.]+)\s*x\s*(?P<unit>\d+,\d{2})\s+(?P<total>\d+,\d{2})-', _re.M)
    for g in other_pattern.finditer(text):
        if any(g.start() >= a and g.start() < b for a, b in occupied):
            continue
        label = g.group('label').strip()
        if not label or 'Geschäftsvorfälle' in label:
            continue
        order += 1
        count = int(g.group('count').replace('.', ''))
        unit = _fm501_de_float(g.group('unit'))
        total = _fm501_de_float(g.group('total'))
        calc = round((count or 0) * (unit or 0), 2)
        diff = round((total or 0) - calc, 2)
        result['items'].append({
            'geschaeftsvorfall': label, 'positionstyp': 'Sonstige Entgeltposition', 'anzahl': count,
            'stueckkosten': unit, 'gesamtbetrag': total, 'berechneter_betrag': calc,
            'abweichung': diff, 'plausibel': 1 if abs(diff) < 0.011 else 0,
            'sort_order': order, 'raw_line': g.group(0).strip()
        })

    m = _re.search(r'Mindestbetrag\s+([\d\.]+,\d{2})', text)
    if m:
        order += 1
        amount = _fm501_de_float(m.group(1))
        result['items'].append({
            'geschaeftsvorfall': 'Mindestbetrag', 'positionstyp': 'Mindestbetrag', 'anzahl': None,
            'stueckkosten': None, 'gesamtbetrag': amount, 'berechneter_betrag': amount,
            'abweichung': 0.0, 'plausibel': 1, 'sort_order': order, 'raw_line': m.group(0).strip()
        })
    return result


def _fm501_import_excel(path, app=None):
    from openpyxl import load_workbook
    _fm501_db_init()
    wb = load_workbook(path, data_only=True, read_only=False)
    ws = wb.active
    headers = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    h = {str(v).strip(): i + 1 for i, v in enumerate(headers) if v is not None}
    required = ['Buchung', 'Wertstellung', 'Buchungstext', 'Verwendungszweck', 'Betrag']
    missing = [x for x in required if x not in h]
    if missing:
        raise ValueError('Fehlende Spalten: ' + ', '.join(missing))
    file_hash = _fm501_hash_file(path)
    filename = os.path.basename(path)
    user = getattr(app, 'current_user_display', '') or getattr(app, 'current_user_key', '') or ''
    now = datetime.now().isoformat(timespec='seconds')
    con = _fm501_db_connect()
    closing_count = 0; detail_count = 0; skipped = 0
    try:
        cur = con.cursor()
        cur.execute('INSERT OR IGNORE INTO imports(filename,file_mtime,imported_at,imported_by,row_count,closing_count,detail_count,file_hash) VALUES(?,?,?,?,?,?,?,?)',
                    (filename, str(os.path.getmtime(path)) if os.path.exists(path) else '', now, user, ws.max_row - 1, 0, 0, file_hash))
        con.commit()
        cur.execute('SELECT id FROM imports WHERE filename=? AND file_hash=?', (filename, file_hash))
        import_id = cur.fetchone()[0]
        for r in range(2, ws.max_row + 1):
            bt = ws.cell(r, h['Buchungstext']).value
            if not (bt and 'ENTGELTABSCHLUSS' in str(bt).upper()):
                continue
            vz = str(ws.cell(r, h.get('Verwendungszweck')).value or '')
            if 'Verwendungszweck 2' in h and ws.cell(r, h['Verwendungszweck 2']).value:
                # VZ 2 ist bei SFIRM oft nur Restfragment, wird beim Rawtext aber mitgesichert.
                raw_text = vz + '\n' + str(ws.cell(r, h['Verwendungszweck 2']).value or '')
            else:
                raw_text = vz
            parsed = _fm501_parse_entgelt_text(vz)
            buchungsdatum = _fm501_date_iso(ws.cell(r, h['Buchung']).value)
            wertstellung = _fm501_date_iso(ws.cell(r, h['Wertstellung']).value)
            abschluss = parsed.get('abschlussdatum') or wertstellung or buchungsdatum
            monat_key = _fm501_month_key_from_iso(abschluss)
            betrag = ws.cell(r, h['Betrag']).value
            brutto = abs(float(betrag)) if isinstance(betrag, (int, float)) else _fm501_de_float(betrag)
            rechnung = parsed.get('rechnungsnummer') or f'{filename}:row:{r}'
            cur.execute('''INSERT OR IGNORE INTO closings(import_id,source_file,source_row,buchungsdatum,wertstellungsdatum,abschlussdatum,monat_key,zeitraum_von,zeitraum_bis,rechnungsnummer,betrag_brutto,betrag_netto,ust,waehrung,gvc,btc,raw_text,created_at)
                           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                        (import_id, filename, r, buchungsdatum, wertstellung, abschluss, monat_key, parsed.get('zeitraum_von',''), parsed.get('zeitraum_bis',''), rechnung,
                         brutto, parsed.get('betrag_netto'), parsed.get('ust'), ws.cell(r, h.get('Währung', h['Betrag'])).value if 'Währung' in h else 'EUR',
                         ws.cell(r, h['GVC']).value if 'GVC' in h else '', ws.cell(r, h['BTC']).value if 'BTC' in h else '', raw_text, now))
            cur.execute('SELECT id FROM closings WHERE rechnungsnummer=? AND abschlussdatum=?', (rechnung, abschluss))
            row = cur.fetchone()
            if not row:
                skipped += 1; continue
            closing_id = row[0]
            closing_count += 1
            # Bei Reimport alte Positionen ersetzen, Kopf bleibt eindeutig.
            cur.execute('DELETE FROM fee_items WHERE closing_id=?', (closing_id,))
            for item in parsed.get('items', []):
                cur.execute('''INSERT OR IGNORE INTO fee_items(closing_id,abschlussdatum,monat_key,geschaeftsvorfall,positionstyp,anzahl,stueckkosten,gesamtbetrag,berechneter_betrag,abweichung,plausibel,sort_order,raw_line)
                               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                            (closing_id, abschluss, monat_key, item.get('geschaeftsvorfall'), item.get('positionstyp'), item.get('anzahl'), item.get('stueckkosten'), item.get('gesamtbetrag'), item.get('berechneter_betrag'), item.get('abweichung'), item.get('plausibel'), item.get('sort_order'), item.get('raw_line')))
                detail_count += 1
        cur.execute('UPDATE imports SET closing_count=?, detail_count=? WHERE id=?', (closing_count, detail_count, import_id))
        con.commit()
    finally:
        con.close()
    return {'filename': filename, 'closings': closing_count, 'details': detail_count, 'skipped': skipped, 'db': _fm501_konto_db_path()}


def _fm501_fetch_closings():
    _fm501_db_init()
    con = _fm501_db_connect()
    try:
        return [dict(r) for r in con.execute('SELECT * FROM closings ORDER BY abschlussdatum DESC, id DESC').fetchall()]
    finally:
        con.close()


def _fm501_fetch_items(closing_id=None):
    _fm501_db_init()
    con = _fm501_db_connect()
    try:
        if closing_id:
            return [dict(r) for r in con.execute('SELECT * FROM fee_items WHERE closing_id=? ORDER BY sort_order, id', (closing_id,)).fetchall()]
        return [dict(r) for r in con.execute('SELECT * FROM fee_items ORDER BY abschlussdatum, geschaeftsvorfall').fetchall()]
    finally:
        con.close()


def _fm501_summary_stats():
    _fm501_db_init()
    con = _fm501_db_connect()
    try:
        row = con.execute('SELECT COUNT(*) c, SUM(betrag_brutto) brutto, SUM(betrag_netto) netto, SUM(ust) ust FROM closings').fetchone()
        detail = con.execute('SELECT COUNT(*) c, SUM(gesamtbetrag) gesamt FROM fee_items').fetchone()
        months = con.execute('SELECT monat_key, betrag_brutto FROM closings ORDER BY abschlussdatum').fetchall()
        top = con.execute('SELECT geschaeftsvorfall, SUM(gesamtbetrag) gesamt, SUM(COALESCE(anzahl,0)) anzahl, AVG(stueckkosten) avg_unit FROM fee_items GROUP BY geschaeftsvorfall ORDER BY gesamt DESC LIMIT 8').fetchall()
        return {'closings': row['c'] or 0, 'brutto': row['brutto'] or 0, 'netto': row['netto'] or 0, 'ust': row['ust'] or 0, 'details': detail['c'] or 0, 'detail_sum': detail['gesamt'] or 0, 'months': [dict(x) for x in months], 'top': [dict(x) for x in top]}
    finally:
        con.close()


def _fm501_fmt_eur(v):
    try:
        return f"{float(v):,.2f} €".replace(',', 'X').replace('.', ',').replace('X', '.')
    except Exception:
        return "-"


def _fm501_fmt_num(v):
    try:
        return f"{float(v):,.0f}".replace(',', '.')
    except Exception:
        return "-"


def _fm501_export_analysis_excel(app=None):
    try:
        from tkinter import filedialog
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from openpyxl.chart import BarChart, LineChart, Reference
        path = filedialog.asksaveasfilename(title='Kontenauswertung exportieren', defaultextension='.xlsx', initialfile='Kontenauswertung_Export_' + datetime.now().strftime('%Y_%m_%d') + '.xlsx', filetypes=[('Excel-Dateien','*.xlsx')])
        if not path:
            return None
        closings = _fm501_fetch_closings(); items = _fm501_fetch_items()
        wb = Workbook(); ws = wb.active; ws.title = 'Entgeltabschluesse'
        header_fill = PatternFill('solid', fgColor='D9EAF7'); sub_fill = PatternFill('solid', fgColor='EAF7EF')
        thin = Side(style='thin', color='A6B7C8')
        border = Border(bottom=thin)
        headers = ['Monat','Abschlussdatum','Buchungsdatum','Wertstellung','Rechnungsnummer','Brutto','Netto','USt','Quelle','Zeile']
        ws.append(headers)
        for c in ws[1]: c.font = Font(bold=True); c.fill = header_fill; c.border = border
        for c in closings:
            ws.append([c.get('monat_key'), c.get('abschlussdatum'), c.get('buchungsdatum'), c.get('wertstellungsdatum'), c.get('rechnungsnummer'), c.get('betrag_brutto'), c.get('betrag_netto'), c.get('ust'), c.get('source_file'), c.get('source_row')])
        for row in ws.iter_rows(min_row=2, min_col=6, max_col=8):
            for cell in row: cell.number_format = '#,##0.00 €'
        for col, width in {'A':12,'B':14,'C':14,'D':14,'E':34,'F':14,'G':14,'H':14,'I':34,'J':10}.items(): ws.column_dimensions[col].width = width
        ws2 = wb.create_sheet('Einzelpositionen')
        headers2 = ['Monat','Abschlussdatum','Geschäftsvorfall','Positionstyp','Anzahl','Stückkosten','Gesamtbetrag','Berechnet','Abweichung','Plausibel']
        ws2.append(headers2)
        for c in ws2[1]: c.font = Font(bold=True); c.fill = sub_fill; c.border = border
        for i in items:
            ws2.append([i.get('monat_key'), i.get('abschlussdatum'), i.get('geschaeftsvorfall'), i.get('positionstyp'), i.get('anzahl'), i.get('stueckkosten'), i.get('gesamtbetrag'), i.get('berechneter_betrag'), i.get('abweichung'), 'Ja' if i.get('plausibel') else 'Nein'])
        for row in ws2.iter_rows(min_row=2, min_col=6, max_col=9):
            for cell in row: cell.number_format = '#,##0.00 €'
        for col, width in {'A':12,'B':14,'C':32,'D':24,'E':12,'F':14,'G':14,'H':14,'I':14,'J':12}.items(): ws2.column_dimensions[col].width = width
        ws3 = wb.create_sheet('Pivot Geschäftsvorfälle')
        ws3.append(['Geschäftsvorfall','Gesamtkosten','Anzahl gesamt','Ø Stückkosten'])
        for c in ws3[1]: c.font = Font(bold=True); c.fill = header_fill; c.border = border
        stats = _fm501_summary_stats()
        for t in stats.get('top', []):
            ws3.append([t.get('geschaeftsvorfall'), t.get('gesamt'), t.get('anzahl'), t.get('avg_unit')])
        for row in ws3.iter_rows(min_row=2, min_col=2, max_col=4):
            for cell in row: cell.number_format = '#,##0.00 €' if cell.column in (2,4) else '#,##0'
        ws3.column_dimensions['A'].width = 32; ws3.column_dimensions['B'].width = 16; ws3.column_dimensions['C'].width = 14; ws3.column_dimensions['D'].width = 16
        if ws3.max_row >= 2:
            chart = BarChart(); chart.title = 'Top Gebührenpositionen'; chart.y_axis.title = 'EUR'; chart.x_axis.title = 'Geschäftsvorfall'
            data = Reference(ws3, min_col=2, min_row=1, max_row=ws3.max_row)
            cats = Reference(ws3, min_col=1, min_row=2, max_row=ws3.max_row)
            chart.add_data(data, titles_from_data=True); chart.set_categories(cats); chart.height = 8; chart.width = 18
            ws3.add_chart(chart, 'F2')
        wb.save(path)
        try: messagebox.showinfo('Kontenauswertung', 'Excel-Export erstellt:\n' + path)
        except Exception: pass
        return path
    except Exception as exc:
        try: messagebox.showerror('Kontenauswertung', 'Excel-Export fehlgeschlagen:\n' + str(exc))
        except Exception: pass
        return None


def _fm501_draw_bar_chart(canvas, x, y, w, h, data, title, value_key, label_key):
    canvas.create_rectangle(x, y, x+w, y+h, fill=WHITE, outline=LINE)
    canvas.create_text(x+14, y+14, text=title, fill=BLUE, font=body_font(12, 'bold'), anchor='w')
    if not data:
        canvas.create_text(x+w/2, y+h/2, text='Noch keine Daten vorhanden.', fill=TEXT2, font=body_font(10))
        return
    max_v = max(float(d.get(value_key) or 0) for d in data) or 1
    left = x + 16; bottom = y + h - 34; top = y + 42; bar_area_h = max(20, bottom-top)
    n = len(data); gap = 8; bw = max(16, int((w - 40 - (n-1)*gap) / max(1,n)))
    for i, d in enumerate(data):
        val = float(d.get(value_key) or 0)
        bh = int(bar_area_h * val / max_v)
        bx = left + i * (bw + gap)
        canvas.create_rectangle(bx, bottom-bh, bx+bw, bottom, fill=blend(BLUE, WHITE, 0.25), outline=BLUE)
        label = str(d.get(label_key) or '')[:8]
        canvas.create_text(bx+bw/2, bottom+13, text=label, fill=TEXT2, font=body_font(8))
        canvas.create_text(bx+bw/2, bottom-bh-8, text=_fm501_fmt_eur(val), fill=TEXT, font=body_font(8))


def _fm501_draw_top_list(parent, top):
    for child in parent.winfo_children(): child.destroy()
    tk.Label(parent, text='Top-Kostentreiber', bg=WHITE, fg=BLUE, font=body_font(12,'bold')).pack(anchor='w', padx=10, pady=(8,4))
    if not top:
        tk.Label(parent, text='Noch keine Detailpositionen importiert.', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor='w', padx=10, pady=8)
        return
    for row in top:
        txt = f"{row.get('geschaeftsvorfall','')}  |  {_fm501_fmt_eur(row.get('gesamt'))}  |  Anzahl {_fm501_fmt_num(row.get('anzahl'))}  |  Ø {_fm501_fmt_eur(row.get('avg_unit'))}"
        tk.Label(parent, text=txt, bg=WHITE, fg=TEXT, font=body_font(9), anchor='w', justify='left').pack(fill='x', padx=10, pady=1)


def _fm501_render_treasury_tools_menu(self):
    items = [
        ('Kontenauswertung', 'account_analysis'),
    ]
    self.render_center_menu(items, title='Tools - Treasury')
    self.draw_bottom_logo()


def _fm501_render_account_analysis_tool(self):
    _fm501_db_init()
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    top_y = 132
    frame = tk.Frame(self.root, bg=BG)
    self.widget_items.append(frame)
    self.canvas.create_window(0, top_y, window=frame, anchor='nw', width=w, height=max(400, h-top_y-55))
    header = tk.Frame(frame, bg=WHITE, highlightbackground=LINE, highlightthickness=1)
    header.pack(fill='x', padx=20, pady=(12,8))
    tk.Label(header, text='Kontenauswertung', bg=WHITE, fg=BLUE, font=body_font(18,'bold')).pack(side='left', padx=14, pady=12)
    tk.Label(header, text='SFIRM-Excel-Import für Entgeltabschlüsse, Gebührenentwicklung, Mengen und Stückkosten', bg=WHITE, fg=TEXT2, font=body_font(10)).pack(side='left', padx=(10,0))

    status_var = tk.StringVar(value='Datenbank: ' + _fm501_konto_db_path())
    selected_closing_id = {'id': None}
    content = tk.Frame(frame, bg=BG)
    content.pack(fill='both', expand=True, padx=20, pady=(0,10))
    controls = tk.Frame(content, bg=BG); controls.pack(fill='x', pady=(0,8))

    def refresh_all():
        stats = _fm501_summary_stats()
        kpi_var.set(f"Entgeltabschlüsse: {stats['closings']}   |   Detailpositionen: {stats['details']}   |   Brutto gesamt: {_fm501_fmt_eur(stats['brutto'])}   |   Netto gesamt: {_fm501_fmt_eur(stats['netto'])}   |   USt: {_fm501_fmt_eur(stats['ust'])}")
        closing_tree.delete(*closing_tree.get_children())
        for c in _fm501_fetch_closings():
            closing_tree.insert('', 'end', iid=str(c['id']), values=(c.get('monat_key'), c.get('abschlussdatum'), c.get('buchungsdatum'), _fm501_fmt_eur(c.get('betrag_brutto')), _fm501_fmt_eur(c.get('betrag_netto')), _fm501_fmt_eur(c.get('ust')), c.get('rechnungsnummer'), c.get('source_file')))
        item_tree.delete(*item_tree.get_children())
        cid = selected_closing_id.get('id')
        for i in _fm501_fetch_items(cid):
            item_tree.insert('', 'end', values=(i.get('monat_key'), i.get('geschaeftsvorfall'), i.get('positionstyp'), _fm501_fmt_num(i.get('anzahl')) if i.get('anzahl') is not None else '', _fm501_fmt_eur(i.get('stueckkosten')) if i.get('stueckkosten') is not None else '', _fm501_fmt_eur(i.get('gesamtbetrag')), _fm501_fmt_eur(i.get('abweichung')), 'Ja' if i.get('plausibel') else 'Nein'))
        charts.delete('all')
        _fm501_draw_bar_chart(charts, 8, 8, max(420, int(w*0.48)), 205, stats.get('months', [])[-12:], 'Brutto-Entgelte je Abschlussmonat', 'betrag_brutto', 'monat_key')
        _fm501_draw_bar_chart(charts, max(440, int(w*0.50)), 8, max(420, int(w*0.45)), 205, stats.get('top', []), 'Top Gebührenpositionen', 'gesamt', 'geschaeftsvorfall')
        _fm501_draw_top_list(top_frame, stats.get('top', []))

    def import_file():
        try:
            from tkinter import filedialog
            path = filedialog.askopenfilename(title='SFIRM-Excel-Export auswählen', filetypes=[('Excel-Dateien','*.xlsx *.xlsm *.xltx *.xltm'), ('Alle Dateien','*.*')])
            if not path: return
            res = _fm501_import_excel(path, self)
            status_var.set(f"Import abgeschlossen: {res['closings']} Entgeltabschlüsse, {res['details']} Detailpositionen. DB: {res['db']}")
            try: messagebox.showinfo('Kontenauswertung', f"Import abgeschlossen.\n\nEntgeltabschlüsse: {res['closings']}\nDetailpositionen: {res['details']}\nDatenbank:\n{res['db']}")
            except Exception: pass
            refresh_all()
        except Exception as exc:
            try: messagebox.showerror('Kontenauswertung', 'Import fehlgeschlagen:\n' + str(exc))
            except Exception: pass

    tk.Button(controls, text='Excel importieren', command=import_file, bg='#CFEAD6', fg=TEXT, font=body_font(10,'bold'), relief='solid', bd=1, padx=14, pady=6).pack(side='left', padx=(0,8))
    tk.Button(controls, text='Aktualisieren', command=refresh_all, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1, padx=14, pady=6).pack(side='left', padx=(0,8))
    tk.Button(controls, text='Excel-Auswertung exportieren', command=lambda: _fm501_export_analysis_excel(self), bg=WHITE, fg=BLUE, font=body_font(10), relief='solid', bd=1, padx=14, pady=6).pack(side='left', padx=(0,8))
    tk.Label(controls, textvariable=status_var, bg=BG, fg=TEXT2, font=body_font(8)).pack(side='right')

    kpi_var = tk.StringVar(value='')
    tk.Label(content, textvariable=kpi_var, bg=WHITE, fg=TEXT, font=body_font(10,'bold'), anchor='w', padx=12, pady=8, relief='solid', bd=1).pack(fill='x', pady=(0,8))

    charts = tk.Canvas(content, bg=BG, height=225, highlightthickness=0)
    charts.pack(fill='x', pady=(0,8))

    split = tk.Frame(content, bg=BG); split.pack(fill='both', expand=True)
    left = tk.Frame(split, bg=WHITE, highlightbackground=LINE, highlightthickness=1); left.pack(side='left', fill='both', expand=True, padx=(0,8))
    right = tk.Frame(split, bg=WHITE, highlightbackground=LINE, highlightthickness=1); right.pack(side='left', fill='both', expand=True, padx=(8,0))
    tk.Label(left, text='Entgeltabschlüsse', bg=WHITE, fg=BLUE, font=body_font(12,'bold')).pack(anchor='w', padx=10, pady=(8,4))
    cols = ('monat','abschluss','buchung','brutto','netto','ust','rechnung','quelle')
    closing_tree = ttk.Treeview(left, columns=cols, show='headings', height=10)
    for col, label, width_col in [('monat','Monat',70),('abschluss','Abschluss',90),('buchung','Buchung',90),('brutto','Brutto',90),('netto','Netto',90),('ust','USt',80),('rechnung','Rechnung',220),('quelle','Quelle',200)]:
        closing_tree.heading(col, text=label); closing_tree.column(col, width=width_col, stretch=(col in ('rechnung','quelle')))
    ys1 = ttk.Scrollbar(left, orient='vertical', command=closing_tree.yview); closing_tree.configure(yscrollcommand=ys1.set)
    closing_tree.pack(side='left', fill='both', expand=True, padx=(10,0), pady=(0,10)); ys1.pack(side='right', fill='y', pady=(0,10), padx=(0,10))

    def on_select(_event=None):
        sel = closing_tree.selection()
        selected_closing_id['id'] = int(sel[0]) if sel else None
        refresh_all()
    closing_tree.bind('<<TreeviewSelect>>', on_select)

    tk.Label(right, text='Einzelpositionen / Stückkosten', bg=WHITE, fg=BLUE, font=body_font(12,'bold')).pack(anchor='w', padx=10, pady=(8,4))
    cols2 = ('monat','gv','typ','anzahl','stk','gesamt','abw','ok')
    item_tree = ttk.Treeview(right, columns=cols2, show='headings', height=10)
    for col, label, width_col in [('monat','Monat',60),('gv','Geschäftsvorfall',190),('typ','Typ',140),('anzahl','Anzahl',70),('stk','Stückkosten',90),('gesamt','Gesamt',90),('abw','Abw.',70),('ok','OK',50)]:
        item_tree.heading(col, text=label); item_tree.column(col, width=width_col, stretch=(col in ('gv','typ')))
    ys2 = ttk.Scrollbar(right, orient='vertical', command=item_tree.yview); item_tree.configure(yscrollcommand=ys2.set)
    item_tree.pack(side='left', fill='both', expand=True, padx=(10,0), pady=(0,10)); ys2.pack(side='right', fill='y', pady=(0,10), padx=(0,10))

    top_frame = tk.Frame(content, bg=WHITE, highlightbackground=LINE, highlightthickness=1)
    top_frame.pack(fill='x', pady=(8,0))
    refresh_all()


def _fm501_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Tools - Treasury','cmd':lambda: self.show_page('treasury_tools','Tools - Treasury',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165))); th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04)); line_x1 = x_pct(w, 9.5); line_x2 = x_pct(w, 92)
    def create_tile_row(items, left_edge, y, id_prefix, tile_w, tile_h):
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(left_edge + i*(tile_w+gap_x), y, window=tile, anchor='nw')
    top_tile_w = max(250, int(tw*1.05)); top_tile_h = max(130, int(th*1.05))
    tool_tile_w = max(200, int(tw*0.82)); tool_tile_h = max(100, int(th*0.82))
    mini_tw = int(tw*0.68); mini_th = int(th*0.68)
    y_modules = y_pct(h, 55); y_tools = y_pct(h, 35); mini_y = y_pct(h, 15)
    self.draw_continuous_relief_line(self.canvas, y_modules-22, line_x1, line_x2)
    self.draw_section_label('Module', line_x1, y_modules-38)
    create_tile_row(top_tiles, line_x1, y_modules, 'main_module', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, y_tools-20, line_x1, line_x2)
    self.draw_section_label('Tools', line_x1, y_tools-36)
    create_tile_row(middle_tiles, line_x1, y_tools, 'main_tools', tool_tile_w, tool_tile_h)
    self.draw_continuous_relief_line(self.canvas, mini_y-18, line_x1, line_x2)
    create_tile_row(mini_tiles, line_x1, mini_y, 'main_mini', mini_tw, mini_th)
    self.draw_bottom_logo()

try:
    _fm501_prev_render_page = FiBuMateApp.render_page
except Exception:
    _fm501_prev_render_page = None


def _fm501_render_page(self):
    if self.current_page in ('treasury_tools', 'account_analysis'):
        self.clear_widgets()
        self.draw_background()
        self.draw_header(self.current_title)
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        if self.current_page == 'treasury_tools':
            self.render_treasury_tools_menu()
        elif self.current_page == 'account_analysis':
            self.render_account_analysis_tool()
        if self.focusable_tiles:
            self.focus_index = 0
            self.focusable_tiles[0].focus_set()
        return
    if _fm501_prev_render_page:
        return _fm501_prev_render_page(self)

try:
    FiBuMateApp.render_treasury_tools_menu = _fm501_render_treasury_tools_menu
    FiBuMateApp.render_account_analysis_tool = _fm501_render_account_analysis_tool
    FiBuMateApp.render_main_menu = _fm501_render_main_menu
    FiBuMateApp.render_page = _fm501_render_page
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_MAINMENU_KONTENAUSWERTUNG_FIX_V0502
# Datum: 2026-07-13
# Zweck:
# - Hauptmenue nach v0.501 wieder sichtbar machen.
# - Ursache: v0.501 rief self.draw_section_label auf, obwohl diese Methode nicht final an FiBuMateApp gebunden war.
# - Schreibweise "Hauptmenü" beim Login wieder korrigieren.
# - Tools - Treasury bleibt in der Tools-Zeile, Dateiausgabe bleibt rechts neben Einstellungen in der unteren Modulleiste.
# ------------------------------------------------------------------
FIBU_MATE_MAINMENU_KONTENAUSWERTUNG_FIX_VERSION = "0.502"


def _fm502_draw_section_label_safe(self, text, x, y):
    try:
        if '_fm484_draw_section_label' in globals():
            return _fm484_draw_section_label(self, text, x, y)
    except Exception:
        pass
    try:
        self.canvas.create_text(x, y, text=text, fill=blend(TEXT2, BLUE, 0.30), font=body_font(9, 'bold'), anchor='w')
    except Exception:
        pass


def _fm502_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Tools - Treasury','cmd':lambda: self.show_page('treasury_tools','Tools - Treasury',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165)))
    th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04))
    line_x1 = x_pct(w, 9.5)
    line_x2 = x_pct(w, 92)
    header_line_y = FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)

    def create_tile_row(items, visible_left_edge, canvas_top, id_prefix, tile_w, tile_h):
        canvas_left = visible_left_edge - ui_s(8)
        canvas_h = tile_h + ui_s(16)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            cx = canvas_left + (tile_w + ui_s(16)) / 2 + i*(tile_w+gap_x)
            cy = canvas_top + canvas_h / 2
            self.canvas.create_window(cx, cy, window=tile, anchor='center')

    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count*tile_w + (count-1)*gap
        start_x = (w-total_width)/2 + tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]

    def create_centered_row(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i,item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w,tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')

    top_tile_w=max(250,int(tw*1.05))
    top_tile_h=max(130,int(th*1.05))
    tool_tile_w=max(200,int(tw*0.86*0.95))
    tool_tile_h=max(100,int(th*0.86*0.95))
    mini_tw=int(tw*0.72)
    mini_th=int(th*0.72)

    gap_below_line = ui_s(5)
    top_canvas_top = header_line_y + gap_below_line
    top_canvas_bottom = top_canvas_top + top_tile_h + ui_s(16)
    first_sep_y = top_canvas_bottom + ui_s(18)
    tools_canvas_top = first_sep_y + gap_below_line
    mini_center_y = y_pct(h,16.5) + th/2 - mini_th/2
    mini_top = mini_center_y - mini_th/2
    label_x = ui_s(8)

    _fm502_draw_section_label_safe(self, 'Module', label_x, header_line_y + ui_s(42))
    create_tile_row(top_tiles, line_x1, top_canvas_top, 'main_menuzeile_1', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, first_sep_y, line_x1, line_x2)

    _fm502_draw_section_label_safe(self, 'Tools', label_x, first_sep_y + ui_s(28))
    create_tile_row(middle_tiles, line_x1, tools_canvas_top, 'main_menuzeile_2', tool_tile_w, tool_tile_h)

    self.draw_continuous_relief_line(self.canvas, mini_top-ui_s(13), line_x1, line_x2)
    create_centered_row(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_bottom_logo()


# Login-Finalisierung korrigieren: Hauptmenü mit Umlaut anzeigen.
try:
    _fm502_prev_finalize_login = _fm488_finalize_login
except Exception:
    _fm502_prev_finalize_login = None


def _fm502_finalize_login(self, key, username):
    if _fm502_prev_finalize_login:
        _fm502_prev_finalize_login(self, key, username)
        try:
            if getattr(self, 'current_page', '') == 'main':
                self.current_title = 'Hauptmenü'
                self.breadcrumb = [('main', 'Hauptmenü')]
        except Exception:
            pass
        return
    try:
        self.show_page('main', 'Hauptmenü', add_to_history=False)
    except Exception:
        pass

try:
    FiBuMateApp.draw_section_label = _fm502_draw_section_label_safe
    FiBuMateApp.render_main_menu = _fm502_render_main_menu
    _fm488_finalize_login = _fm502_finalize_login
    FiBuMateApp.finalize_login = _fm502_finalize_login
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_TREASURY_OPEN_FIX_V0503
# Datum: 2026-07-13
# Zweck:
# - Klick auf "Tools - Treasury" robust machen.
# - Direkter Öffner fuer Treasury-Menue statt ausschliesslich verschachteltem lambda/show_page-Pfad.
# - show_page und render_page fangen treasury_tools/account_analysis final ab.
# - Kontenauswertung bleibt unter Tools -> Tools - Treasury erreichbar.
# ------------------------------------------------------------------
FIBU_MATE_TREASURY_OPEN_FIX_VERSION = "0.503"


def _fm503_render_special_page(self, page_name, title):
    try:
        self.clear_widgets()
        self.draw_background()
        self.draw_header(title or self.current_title or '')
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        if page_name == 'treasury_tools':
            self.render_treasury_tools_menu()
        elif page_name == 'account_analysis':
            self.render_account_analysis_tool()
        if self.focusable_tiles:
            self.focus_index = 0
            try: self.focusable_tiles[0].focus_set()
            except Exception: pass
        return True
    except Exception as exc:
        try:
            messagebox.showerror('FiBu Mate', 'Seite konnte nicht geöffnet werden:\n' + str(exc))
        except Exception:
            pass
        return True


def _fm503_open_treasury_tools_menu(self):
    try:
        if not self.confirm_unsaved_changes():
            return
    except Exception:
        pass
    try:
        if getattr(self, 'current_page', None):
            self.page_history.append((self.current_page, self.current_title))
    except Exception:
        pass
    self.current_page = 'treasury_tools'
    self.current_title = 'Tools - Treasury'
    try:
        self.update_breadcrumb('treasury_tools', 'Tools - Treasury')
    except Exception:
        self.breadcrumb = [('main', 'Hauptmenü'), ('treasury_tools', 'Tools - Treasury')]
    _fm503_render_special_page(self, 'treasury_tools', 'Tools - Treasury')


def _fm503_open_account_analysis_tool(self):
    try:
        if not self.confirm_unsaved_changes():
            return
    except Exception:
        pass
    try:
        if getattr(self, 'current_page', None):
            self.page_history.append((self.current_page, self.current_title))
    except Exception:
        pass
    self.current_page = 'account_analysis'
    self.current_title = 'Kontenauswertung'
    try:
        self.update_breadcrumb('account_analysis', 'Kontenauswertung')
    except Exception:
        self.breadcrumb = [('main', 'Hauptmenü'), ('treasury_tools', 'Tools - Treasury'), ('account_analysis', 'Kontenauswertung')]
    _fm503_render_special_page(self, 'account_analysis', 'Kontenauswertung')

try:
    _fm503_prev_show_page = FiBuMateApp.show_page
except Exception:
    _fm503_prev_show_page = None


def _fm503_show_page(self, page_name, title='', add_to_history=True):
    if page_name == 'treasury_tools':
        return _fm503_open_treasury_tools_menu(self)
    if page_name == 'account_analysis':
        return _fm503_open_account_analysis_tool(self)
    if _fm503_prev_show_page:
        return _fm503_prev_show_page(self, page_name, title, add_to_history)

try:
    _fm503_prev_render_page = FiBuMateApp.render_page
except Exception:
    _fm503_prev_render_page = None


def _fm503_render_page(self):
    page = getattr(self, 'current_page', '')
    if page in ('treasury_tools', 'account_analysis'):
        return _fm503_render_special_page(self, page, getattr(self, 'current_title', '') or ('Tools - Treasury' if page == 'treasury_tools' else 'Kontenauswertung'))
    if _fm503_prev_render_page:
        return _fm503_prev_render_page(self)


def _fm503_render_treasury_tools_menu(self):
    # Eigene Kachel statt render_center_menu, damit der Klick sicher ueber den direkten Öffner laeuft.
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile_w = int(max(280, min(380, w * 0.21)))
    tile_h = int(max(120, min(170, h * 0.15)))
    x = x_pct(w, 50)
    y = y_pct(h, 55)
    tile = Tile(self.root, self, 'treasury_kontenauswertung', 'Kontenauswertung', command=lambda: _fm503_open_account_analysis_tool(self), favorite_enabled=False, center_text=True, icon_type='modules')
    tile.resize_tile(tile_w, tile_h)
    self.widget_items.append(tile)
    self.focusable_tiles.append(tile)
    self.canvas.create_window(x, y, window=tile, anchor='center')
    self.draw_bottom_logo()


def _fm503_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Tools - Treasury','cmd':lambda: _fm503_open_treasury_tools_menu(self),'fixed':None,'lock':False,'icon':'modules','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165)))
    th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04))
    line_x1 = x_pct(w, 9.5)
    line_x2 = x_pct(w, 92)
    header_line_y = FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)

    def create_tile_row(items, visible_left_edge, canvas_top, id_prefix, tile_w, tile_h):
        canvas_left = visible_left_edge - ui_s(8)
        canvas_h = tile_h + ui_s(16)
        for i, item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w, tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            cx = canvas_left + (tile_w + ui_s(16)) / 2 + i*(tile_w+gap_x)
            cy = canvas_top + canvas_h / 2
            self.canvas.create_window(cx, cy, window=tile, anchor='center')

    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width = count*tile_w + (count-1)*gap
        start_x = (w-total_width)/2 + tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]

    def create_centered_row(items, y, id_prefix, tile_w, tile_h):
        xs = centered_x_positions(len(items), tile_w)
        for i,item in enumerate(items):
            tile = Tile(self.root, self, f'{id_prefix}_{i}', item['title'], item['cmd'], favorite_enabled=False, fixed_color=item['fixed'], lock_tile=item['lock'], icon_type=item['icon'], corner_fold=item['fold'])
            tile.resize_tile(tile_w,tile_h)
            self.widget_items.append(tile)
            self.focusable_tiles.append(tile)
            self.canvas.create_window(xs[i], y, window=tile, anchor='center')

    top_tile_w=max(250,int(tw*1.05)); top_tile_h=max(130,int(th*1.05))
    tool_tile_w=max(200,int(tw*0.86*0.95)); tool_tile_h=max(100,int(th*0.86*0.95))
    mini_tw=int(tw*0.72); mini_th=int(th*0.72)
    gap_below_line = ui_s(5)
    top_canvas_top = header_line_y + gap_below_line
    top_canvas_bottom = top_canvas_top + top_tile_h + ui_s(16)
    first_sep_y = top_canvas_bottom + ui_s(18)
    tools_canvas_top = first_sep_y + gap_below_line
    mini_center_y = y_pct(h,16.5) + th/2 - mini_th/2
    mini_top = mini_center_y - mini_th/2
    label_x = ui_s(8)
    try:
        _fm502_draw_section_label_safe(self, 'Module', label_x, header_line_y + ui_s(42))
    except Exception:
        self.canvas.create_text(label_x, header_line_y + ui_s(42), text='Module', fill=TEXT2, font=body_font(9,'bold'), anchor='w')
    create_tile_row(top_tiles, line_x1, top_canvas_top, 'main_menuzeile_1', top_tile_w, top_tile_h)
    self.draw_continuous_relief_line(self.canvas, first_sep_y, line_x1, line_x2)
    try:
        _fm502_draw_section_label_safe(self, 'Tools', label_x, first_sep_y + ui_s(28))
    except Exception:
        self.canvas.create_text(label_x, first_sep_y + ui_s(28), text='Tools', fill=TEXT2, font=body_font(9,'bold'), anchor='w')
    create_tile_row(middle_tiles, line_x1, tools_canvas_top, 'main_menuzeile_2', tool_tile_w, tool_tile_h)
    self.draw_continuous_relief_line(self.canvas, mini_top-ui_s(13), line_x1, line_x2)
    create_centered_row(mini_tiles, mini_center_y, 'main_mini_menuezeile', mini_tw, mini_th)
    self.draw_bottom_logo()

try:
    FiBuMateApp.open_treasury_tools_menu = _fm503_open_treasury_tools_menu
    FiBuMateApp.open_account_analysis_tool = _fm503_open_account_analysis_tool
    FiBuMateApp.show_page = _fm503_show_page
    FiBuMateApp.render_page = _fm503_render_page
    FiBuMateApp.render_treasury_tools_menu = _fm503_render_treasury_tools_menu
    FiBuMateApp.render_main_menu = _fm503_render_main_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_KONTENAUSWERTUNG_STABIL_V0504
# Datum: 2026-07-13
# Zweck: clear_widgets-Fehler aus v0.503 korrigieren und Treasury/Kontenauswertung stabil oeffnen.
# ------------------------------------------------------------------
FIBU_MATE_KONTENAUSWERTUNG_STABIL_VERSION = "0.504"


def _fm504_clear_widgets_compat(self):
    try:
        return self.clear_content()
    except Exception:
        try:
            for wdg in list(getattr(self, 'widget_items', []) or []):
                try: wdg.destroy()
                except Exception: pass
            self.widget_items = []
            self.focusable_tiles = []
        except Exception:
            pass


def _fm504_render_special_page(self, page_name, title):
    try:
        self.clear_content()
        self.draw_background()
        self.draw_header(title or getattr(self, 'current_title', '') or '')
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        if page_name == 'treasury_tools':
            self.render_treasury_tools_menu()
        elif page_name == 'account_analysis':
            self.render_account_analysis_tool()
        else:
            return False
        if getattr(self, 'focusable_tiles', None):
            self.focus_index = 0
            try: self.focusable_tiles[0].focus_set()
            except Exception: pass
        return True
    except Exception as exc:
        try: messagebox.showerror('FiBu Mate', 'Seite konnte nicht geöffnet werden:\n' + str(exc))
        except Exception: pass
        return True


def _fm504_open_page_direct(self, page_name, title):
    try:
        if page_name != getattr(self, 'current_page', '') and not self.confirm_unsaved_changes():
            return
    except Exception:
        pass
    try:
        if getattr(self, 'current_page', None) and getattr(self, 'current_page', None) != page_name:
            self.page_history.append((self.current_page, self.current_title))
    except Exception:
        pass
    self.current_page = page_name
    self.current_title = title
    try: self.update_breadcrumb(page_name, title)
    except Exception:
        self.breadcrumb = [('main', 'Hauptmenü'), (page_name, title)]
    return _fm504_render_special_page(self, page_name, title)


def _fm504_open_treasury_tools_menu(self):
    return _fm504_open_page_direct(self, 'treasury_tools', 'Tools - Treasury')


def _fm504_open_account_analysis_tool(self):
    return _fm504_open_page_direct(self, 'account_analysis', 'Kontenauswertung')

try:
    _fm504_prev_show_page = FiBuMateApp.show_page
except Exception:
    _fm504_prev_show_page = None


def _fm504_show_page(self, page_name, title='', add_to_history=True):
    if page_name == 'treasury_tools':
        return _fm504_open_treasury_tools_menu(self)
    if page_name == 'account_analysis':
        return _fm504_open_account_analysis_tool(self)
    if _fm504_prev_show_page:
        return _fm504_prev_show_page(self, page_name, title, add_to_history)

try:
    _fm504_prev_render_page = FiBuMateApp.render_page
except Exception:
    _fm504_prev_render_page = None


def _fm504_render_page(self):
    page = getattr(self, 'current_page', '')
    if page in ('treasury_tools', 'account_analysis'):
        return _fm504_render_special_page(self, page, getattr(self, 'current_title', '') or ('Tools - Treasury' if page == 'treasury_tools' else 'Kontenauswertung'))
    if _fm504_prev_render_page:
        return _fm504_prev_render_page(self)


def _fm504_render_treasury_tools_menu(self):
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tile = Tile(self.root, self, 'treasury_kontenauswertung', 'Kontenauswertung', command=lambda: _fm504_open_account_analysis_tool(self), favorite_enabled=False, center_text=True, icon_type='modules')
    tile.resize_tile(int(max(280, min(380, w * 0.21))), int(max(120, min(170, h * 0.15))))
    self.widget_items.append(tile)
    self.focusable_tiles.append(tile)
    self.canvas.create_window(x_pct(w, 50), y_pct(h, 55), window=tile, anchor='center')
    self.draw_bottom_logo()


def _fm504_render_main_menu(self):
    top_tiles = [
        {'title':'Abschlusskalender','cmd':lambda: self.show_page('closing_calendar','Abschlusskalender',True),'fixed':None,'lock':False,'icon':'calendar','fold':False},
        {'title':'Wissenszentrale','cmd':lambda: self.open_knowledge_base(),'fixed':None,'lock':False,'icon':'knowledge','fold':False},
    ]
    middle_tiles = [
        {'title':'Tools - Hauptbuch','cmd':lambda: self.show_page('data_prep','Tools - Hauptbuch',True),'fixed':None,'lock':False,'icon':'pdf_xls','fold':False},
        {'title':'Tools - Debitoren','cmd':lambda: self.show_page('debitoren_tools','Tools - Debitoren',True),'fixed':None,'lock':False,'icon':'modules','fold':False},
        {'title':'Tools - Treasury','cmd':lambda: _fm504_open_treasury_tools_menu(self),'fixed':None,'lock':False,'icon':'modules','fold':False},
    ]
    mini_tiles = [
        {'title':'In Entwicklung','cmd':self.try_open_in_dev,'fixed':GREY_TILE,'lock':True,'icon':'lock','fold':False},
        {'title':'Informationen','cmd':lambda: self.show_page('information','Informationen',True),'fixed':None,'lock':False,'icon':'info','fold':True},
        {'title':'Einstellungen','cmd':lambda: self.show_page('settings','Einstellungen',True),'fixed':None,'lock':False,'icon':'gear','fold':True},
        {'title':'Dateiausgabe','cmd':lambda: self.show_page('file_output','Dateiausgabe',True),'fixed':None,'lock':False,'icon':'doc_file','fold':False},
    ]
    w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
    tw = max(240, min(330, int(w*0.165))); th = max(125, min(170, int(h*0.155)))
    gap_x = max(38, int(w*0.04)); line_x1 = x_pct(w, 9.5); line_x2 = x_pct(w, 92)
    header_line_y = FM_HEADER_H if 'FM_HEADER_H' in globals() else ui_s(96)
    def create_tile_row(items, visible_left_edge, canvas_top, id_prefix, tile_w, tile_h):
        canvas_left=visible_left_edge-ui_s(8); canvas_h=tile_h+ui_s(16)
        for i,item in enumerate(items):
            tile=Tile(self.root,self,f'{id_prefix}_{i}',item['title'],item['cmd'],favorite_enabled=False,fixed_color=item['fixed'],lock_tile=item['lock'],icon_type=item['icon'],corner_fold=item['fold'])
            tile.resize_tile(tile_w,tile_h); self.widget_items.append(tile); self.focusable_tiles.append(tile)
            self.canvas.create_window(canvas_left+(tile_w+ui_s(16))/2+i*(tile_w+gap_x), canvas_top+canvas_h/2, window=tile, anchor='center')
    def centered_x_positions(count, tile_w, gap=gap_x):
        total_width=count*tile_w+(count-1)*gap; start_x=(w-total_width)/2+tile_w/2
        return [start_x+i*(tile_w+gap) for i in range(count)]
    def create_centered_row(items,y,id_prefix,tile_w,tile_h):
        xs=centered_x_positions(len(items),tile_w)
        for i,item in enumerate(items):
            tile=Tile(self.root,self,f'{id_prefix}_{i}',item['title'],item['cmd'],favorite_enabled=False,fixed_color=item['fixed'],lock_tile=item['lock'],icon_type=item['icon'],corner_fold=item['fold'])
            tile.resize_tile(tile_w,tile_h); self.widget_items.append(tile); self.focusable_tiles.append(tile); self.canvas.create_window(xs[i],y,window=tile,anchor='center')
    top_tile_w=max(250,int(tw*1.05)); top_tile_h=max(130,int(th*1.05)); tool_tile_w=max(200,int(tw*0.86*0.95)); tool_tile_h=max(100,int(th*0.86*0.95)); mini_tw=int(tw*0.72); mini_th=int(th*0.72)
    top_canvas_top=header_line_y+ui_s(5); top_canvas_bottom=top_canvas_top+top_tile_h+ui_s(16); first_sep_y=top_canvas_bottom+ui_s(18); tools_canvas_top=first_sep_y+ui_s(5); mini_center_y=y_pct(h,16.5)+th/2-mini_th/2; mini_top=mini_center_y-mini_th/2; label_x=ui_s(8)
    try: _fm502_draw_section_label_safe(self,'Module',label_x,header_line_y+ui_s(42))
    except Exception: self.canvas.create_text(label_x,header_line_y+ui_s(42),text='Module',fill=TEXT2,font=body_font(9,'bold'),anchor='w')
    create_tile_row(top_tiles,line_x1,top_canvas_top,'main_menuzeile_1',top_tile_w,top_tile_h)
    self.draw_continuous_relief_line(self.canvas,first_sep_y,line_x1,line_x2)
    try: _fm502_draw_section_label_safe(self,'Tools',label_x,first_sep_y+ui_s(28))
    except Exception: self.canvas.create_text(label_x,first_sep_y+ui_s(28),text='Tools',fill=TEXT2,font=body_font(9,'bold'),anchor='w')
    create_tile_row(middle_tiles,line_x1,tools_canvas_top,'main_menuzeile_2',tool_tile_w,tool_tile_h)
    self.draw_continuous_relief_line(self.canvas,mini_top-ui_s(13),line_x1,line_x2)
    create_centered_row(mini_tiles,mini_center_y,'main_mini_menuezeile',mini_tw,mini_th)
    self.draw_bottom_logo()

try:
    FiBuMateApp.clear_widgets = _fm504_clear_widgets_compat
    FiBuMateApp.open_treasury_tools_menu = _fm504_open_treasury_tools_menu
    FiBuMateApp.open_account_analysis_tool = _fm504_open_account_analysis_tool
    FiBuMateApp.show_page = _fm504_show_page
    FiBuMateApp.render_page = _fm504_render_page
    FiBuMateApp.render_treasury_tools_menu = _fm504_render_treasury_tools_menu
    FiBuMateApp.render_main_menu = _fm504_render_main_menu
except Exception:
    pass




# ------------------------------------------------------------------
# FIBU_MATE_KONTENAUSWERTUNG_AUSBAU_V0505
# Datum: 2026-07-13
# Zweck:
# - Kopfbereich in Kontenauswertung entfernen, Arbeitsbereich nach oben ziehen.
# - Button "Excel-Daten aufnehmen" statt "Excel importieren".
# - Importe nach Bank/Konto erfassen, Doppelimporte erkennen und importierte Datensätze löschen können.
# - Entgeltabschlüsse, Einzelpositionen/Stückkosten und Kostentreiber filterbar + sortierbar.
# - Top-Kostentreiber als scrollbar formatierte Tabelle.
# - Schaubilder personalisierbar/filterbar und erst nach Button-Klick aktualisieren.
# - Bank-/Konto-Auswertungen inkl. Ø Stückkosten und Stückkostenentwicklung je Geschäftsvorfall pro Bank.
# ------------------------------------------------------------------
FIBU_MATE_KONTENAUSWERTUNG_AUSBAU_VERSION = "0.505"


def _fm505_db_columns(con, table):
    try:
        return {str(r['name']) for r in con.execute(f'PRAGMA table_info({table})').fetchall()}
    except Exception:
        return set()


def _fm505_add_col(con, table, name, spec):
    cols = _fm505_db_columns(con, table)
    if name not in cols:
        con.execute(f'ALTER TABLE {table} ADD COLUMN {name} {spec}')


def _fm505_db_init():
    _fm501_db_init()
    con = _fm501_db_connect()
    try:
        _fm505_add_col(con, 'imports', 'bank', 'TEXT')
        _fm505_add_col(con, 'imports', 'konto', 'TEXT')
        _fm505_add_col(con, 'imports', 'deleted', 'INTEGER DEFAULT 0')
        _fm505_add_col(con, 'imports', 'deleted_at', 'TEXT')
        _fm505_add_col(con, 'imports', 'deleted_by', 'TEXT')
        _fm505_add_col(con, 'closings', 'bank', 'TEXT')
        _fm505_add_col(con, 'closings', 'konto', 'TEXT')
        _fm505_add_col(con, 'fee_items', 'bank', 'TEXT')
        _fm505_add_col(con, 'fee_items', 'konto', 'TEXT')
        con.commit()
    finally:
        con.close()


def _fm505_distinct_values(table, col):
    _fm505_db_init()
    con = _fm501_db_connect()
    try:
        rows = con.execute(f"SELECT DISTINCT COALESCE({col},'') v FROM {table} WHERE COALESCE({col},'')<>'' ORDER BY v").fetchall()
        return [r['v'] for r in rows]
    finally:
        con.close()


def _fm505_guess_konto_from_filename(path):
    try:
        base = os.path.basename(str(path or ''))
        m = re.search(r'(\d{4,12})(?=\D*(?:\.xlsx|\.xlsm|\.xls|$))', base)
        if m:
            return m.group(1)
    except Exception:
        pass
    return ''


def _fm505_ask_bank_konto(app, path):
    default_konto = _fm505_guess_konto_from_filename(path)
    result = {'ok': False, 'bank': '', 'konto': default_konto}
    try:
        win = tk.Toplevel(app.root)
        win.title('Bank und Konto zum Import')
        win.configure(bg=WHITE)
        win.transient(app.root)
        win.grab_set()
        tk.Label(win, text='Excel-Daten aufnehmen', bg=WHITE, fg=BLUE, font=body_font(14, 'bold')).grid(row=0, column=0, columnspan=2, sticky='w', padx=18, pady=(14,6))
        tk.Label(win, text='Bitte Bank und Konto für diesen SFIRM-Export angeben. Das Konto wird aus dem Dateinamen vorgeschlagen und kann überschrieben werden.', bg=WHITE, fg=TEXT2, font=body_font(9), wraplength=480, justify='left').grid(row=1, column=0, columnspan=2, sticky='w', padx=18, pady=(0,12))
        bank_var = tk.StringVar(value='')
        konto_var = tk.StringVar(value=default_konto)
        bank_values = _fm505_distinct_values('imports', 'bank')
        tk.Label(win, text='Bank', bg=WHITE, fg=TEXT, font=body_font(10, 'bold')).grid(row=2, column=0, sticky='w', padx=18, pady=(0,3))
        cb = ttk.Combobox(win, textvariable=bank_var, values=bank_values, state='normal', font=body_font(10), width=34)
        cb.grid(row=3, column=0, sticky='ew', padx=18, pady=(0,10))
        tk.Label(win, text='Konto', bg=WHITE, fg=TEXT, font=body_font(10, 'bold')).grid(row=2, column=1, sticky='w', padx=(8,18), pady=(0,3))
        ent = tk.Entry(win, textvariable=konto_var, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1)
        ent.grid(row=3, column=1, sticky='ew', padx=(8,18), pady=(0,10), ipady=4)
        win.grid_columnconfigure(0, weight=1); win.grid_columnconfigure(1, weight=1)
        def ok():
            if not bank_var.get().strip():
                try: messagebox.showwarning('Kontenauswertung', 'Bitte eine Bank eingeben oder auswählen.')
                except Exception: pass
                return
            if not konto_var.get().strip():
                try: messagebox.showwarning('Kontenauswertung', 'Bitte ein Konto eingeben.')
                except Exception: pass
                return
            result.update({'ok': True, 'bank': bank_var.get().strip(), 'konto': konto_var.get().strip()})
            win.destroy()
        def cancel():
            win.destroy()
        btns = tk.Frame(win, bg=WHITE); btns.grid(row=4, column=0, columnspan=2, sticky='e', padx=18, pady=(2,14))
        tk.Button(btns, text='Übernehmen', command=ok, bg='#CFEAD6', fg=TEXT, font=body_font(10,'bold'), relief='solid', bd=1, padx=14, pady=5).pack(side='left', padx=(0,8))
        tk.Button(btns, text='Abbrechen', command=cancel, bg=WHITE, fg=TEXT, font=body_font(10), relief='solid', bd=1, padx=14, pady=5).pack(side='left')
        try: win.wait_window()
        except Exception: pass
    except Exception:
        # Nicht-UI/Testumgebung: Pflichtwerte leerlassen; Aufrufer darf explizit übergeben.
        pass
    return result if result.get('ok') else None



def _fm505b_normalize_header(value):
    import unicodedata as _ud
    text=_ud.normalize('NFKC',str(value or ''))
    text=text.replace('\ufeff','').replace('\u200b','').replace('\xa0',' ')
    text=re.sub(r'\s+',' ',text).strip().casefold()
    return re.sub(r'[^a-z0-9äöüß]+','',text)


def _fm505b_find_header(workbook):
    aliases={
        'Buchung': {'buchung','buchungsdatum','buchungstag'},
        'Wertstellung': {'wertstellung','wertstellungsdatum','valuta','valutadatum'},
        'Buchungstext': {'buchungstext','buchungsart','umsatzart'},
        'Verwendungszweck': {'verwendungszweck','verwendungszweck1','vz','zahlungsreferenz'},
        'Verwendungszweck 2': {'verwendungszweck2','verwendungszweckii','vz2'},
        'Betrag': {'betrag','umsatz','umsatzbetrag','betragwaehrung'},
        'Währung': {'währung','waehrung','currency'},
        'GVC': {'gvc'}, 'BTC': {'btc'},
    }
    alias_norm={k:{_fm505b_normalize_header(x) for x in vals} for k,vals in aliases.items()}
    required={'Buchung','Wertstellung','Buchungstext','Verwendungszweck','Betrag'}
    best=None
    for ws in workbook.worksheets:
        for row_idx in range(1,min(max(1,int(ws.max_row or 1)),50)+1):
            mapping={}; raw=[]
            for col_idx in range(1,min(max(1,int(ws.max_column or 1)),100)+1):
                value=ws.cell(row=row_idx,column=col_idx).value
                if value is None or str(value).strip()=='': continue
                raw.append(str(value).strip())
                norm=_fm505b_normalize_header(value)
                for canonical,variants in alias_norm.items():
                    if norm in variants and canonical not in mapping:
                        mapping[canonical]=col_idx; break
            score=len(required.intersection(mapping))
            if best is None or score>best[0]: best=(score,ws,row_idx,mapping,raw)
            if required.issubset(mapping): return ws,row_idx,mapping
    score,ws,row_idx,mapping,raw=best or (0,workbook.active,1,{},[])
    missing=sorted(required.difference(mapping))
    raise ValueError('Die SFIRM-Kopfzeile konnte nicht vollständig erkannt werden. Fehlende Spalten: '+', '.join(missing)+f'\n\nBestes Ergebnis: Tabellenblatt "{ws.title}", Zeile {row_idx}. Gefundene Spalten: '+(', '.join(raw[:20]) if raw else 'keine'))

def _fm505_import_excel(path, app=None, bank=None, konto=None, replace_existing=False):
    from openpyxl import load_workbook
    _fm505_db_init()
    if not bank or not konto:
        if app is not None:
            sel = _fm505_ask_bank_konto(app, path)
            if not sel:
                return {'filename': os.path.basename(path), 'closings': 0, 'details': 0, 'skipped': 0, 'duplicate': False, 'cancelled': True, 'db': _fm501_konto_db_path()}
            bank, konto = sel['bank'], sel['konto']
        else:
            bank = bank or 'Unbekannte Bank'
            konto = konto or _fm505_guess_konto_from_filename(path) or 'Unbekanntes Konto'
    bank = str(bank).strip()
    konto = str(konto).strip()
    wb = load_workbook(path, data_only=True, read_only=False)
    ws, header_row, h = _fm505b_find_header(wb)
    file_hash = _fm501_hash_file(path)
    filename = os.path.basename(path)
    user = getattr(app, 'current_user_display', '') or getattr(app, 'current_user_key', '') or ''
    now = datetime.now().isoformat(timespec='seconds')
    con = _fm501_db_connect()
    closing_count = 0; detail_count = 0; skipped = 0
    duplicate = False
    try:
        cur = con.cursor()
        # Doppelimport global über Dateiinhalt erkennen: Die Datenbank hat UNIQUE(filename,file_hash).
        # Wenn der Nutzer das Ersetzen bestätigt, werden die bestehenden Detaildaten gelöscht
        # und derselbe Import-Datensatz neu aufgebaut. Dadurch bleibt filename+file_hash eindeutig.
        existing = cur.execute("SELECT id, COALESCE(bank,'') bank, COALESCE(konto,'') konto FROM imports WHERE filename=? AND file_hash=?", (filename, file_hash)).fetchone()
        if existing and not replace_existing:
            return {'filename': filename, 'closings': 0, 'details': 0, 'skipped': 0, 'duplicate': True, 'cancelled': False, 'db': _fm501_konto_db_path(), 'bank': bank, 'konto': konto, 'existing_bank': existing['bank'], 'existing_konto': existing['konto'], 'existing_import_id': int(existing['id'])}
        if existing and replace_existing:
            import_id = int(existing['id'])
            old_closing_ids = cur.execute('SELECT id FROM closings WHERE import_id=?', (import_id,)).fetchall()
            for old in old_closing_ids:
                old_id = old['id'] if hasattr(old, 'keys') else old[0]
                cur.execute('DELETE FROM fee_items WHERE closing_id=?', (old_id,))
            cur.execute('DELETE FROM closings WHERE import_id=?', (import_id,))
            cur.execute('UPDATE imports SET file_mtime=?, imported_at=?, imported_by=?, row_count=?, closing_count=0, detail_count=0, bank=?, konto=?, deleted=0, deleted_at=NULL, deleted_by=NULL WHERE id=?',
                        (str(os.path.getmtime(path)) if os.path.exists(path) else '', now, user, max(0, ws.max_row - header_row), bank, konto, import_id))
        else:
            cur.execute('INSERT INTO imports(filename,file_mtime,imported_at,imported_by,row_count,closing_count,detail_count,file_hash,bank,konto,deleted) VALUES(?,?,?,?,?,?,?,?,?,?,0)',
                        (filename, str(os.path.getmtime(path)) if os.path.exists(path) else '', now, user, max(0, ws.max_row - header_row), 0, 0, file_hash, bank, konto))
            import_id = cur.lastrowid
        for r in range(header_row + 1, ws.max_row + 1):
            bt = ws.cell(r, h['Buchungstext']).value
            if not (bt and 'ENTGELTABSCHLUSS' in str(bt).upper()):
                continue
            vz = str(ws.cell(r, h.get('Verwendungszweck')).value or '')
            raw_text = vz
            if 'Verwendungszweck 2' in h and ws.cell(r, h['Verwendungszweck 2']).value:
                raw_text = vz + '\n' + str(ws.cell(r, h['Verwendungszweck 2']).value or '')
            parsed = _fm501_parse_entgelt_text(vz)
            buchungsdatum = _fm501_date_iso(ws.cell(r, h['Buchung']).value)
            wertstellung = _fm501_date_iso(ws.cell(r, h['Wertstellung']).value)
            abschluss = parsed.get('abschlussdatum') or wertstellung or buchungsdatum
            monat_key = _fm501_month_key_from_iso(abschluss)
            betrag = ws.cell(r, h['Betrag']).value
            brutto = abs(float(betrag)) if isinstance(betrag, (int, float)) else _fm501_de_float(betrag)
            base_rechnung = parsed.get('rechnungsnummer') or f'{filename}:row:{r}'
            rechnung = f'{bank}|{konto}|{base_rechnung}'
            cur.execute('''INSERT OR IGNORE INTO closings(import_id,source_file,source_row,buchungsdatum,wertstellungsdatum,abschlussdatum,monat_key,zeitraum_von,zeitraum_bis,rechnungsnummer,betrag_brutto,betrag_netto,ust,waehrung,gvc,btc,raw_text,created_at,bank,konto)
                           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                        (import_id, filename, r, buchungsdatum, wertstellung, abschluss, monat_key, parsed.get('zeitraum_von',''), parsed.get('zeitraum_bis',''), rechnung,
                         brutto, parsed.get('betrag_netto'), parsed.get('ust'), ws.cell(r, h.get('Währung', h['Betrag'])).value if 'Währung' in h else 'EUR',
                         ws.cell(r, h['GVC']).value if 'GVC' in h else '', ws.cell(r, h['BTC']).value if 'BTC' in h else '', raw_text, now, bank, konto))
            row = cur.execute('SELECT id FROM closings WHERE rechnungsnummer=? AND abschlussdatum=?', (rechnung, abschluss)).fetchone()
            if not row:
                skipped += 1; continue
            closing_id = row[0]
            closing_count += 1
            cur.execute('DELETE FROM fee_items WHERE closing_id=?', (closing_id,))
            for item in parsed.get('items', []):
                cur.execute('''INSERT OR IGNORE INTO fee_items(closing_id,abschlussdatum,monat_key,geschaeftsvorfall,positionstyp,anzahl,stueckkosten,gesamtbetrag,berechneter_betrag,abweichung,plausibel,sort_order,raw_line,bank,konto)
                               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
                            (closing_id, abschluss, monat_key, item.get('geschaeftsvorfall'), item.get('positionstyp'), item.get('anzahl'), item.get('stueckkosten'), item.get('gesamtbetrag'), item.get('berechneter_betrag'), item.get('abweichung'), item.get('plausibel'), item.get('sort_order'), item.get('raw_line'), bank, konto))
                detail_count += 1
        cur.execute('UPDATE imports SET closing_count=?, detail_count=? WHERE id=?', (closing_count, detail_count, import_id))
        con.commit()
    finally:
        con.close()
    return {'filename': filename, 'closings': closing_count, 'details': detail_count, 'skipped': skipped, 'duplicate': duplicate, 'cancelled': False, 'db': _fm501_konto_db_path(), 'bank': bank, 'konto': konto}


def _fm505_filter_match(row, filters):
    if not filters:
        return True
    def get(k): return str(row.get(k, '') if row.get(k, '') is not None else '')
    if filters.get('bank') and get('bank') != filters['bank']:
        return False
    if filters.get('konto') and get('konto') != filters['konto']:
        return False
    if filters.get('monat') and filters['monat'].lower() not in get('monat_key').lower():
        return False
    if filters.get('von') and get('abschlussdatum') and get('abschlussdatum') < filters['von']:
        return False
    if filters.get('bis') and get('abschlussdatum') and get('abschlussdatum') > filters['bis']:
        return False
    if filters.get('gv') and filters['gv'].lower() not in get('geschaeftsvorfall').lower():
        return False
    if filters.get('typ') and get('positionstyp') != filters['typ']:
        return False
    if filters.get('plausibel') in ('Ja','Nein'):
        want = 1 if filters['plausibel']=='Ja' else 0
        if int(row.get('plausibel') or 0) != want:
            return False
    if filters.get('mindest') in ('Ja','Nein'):
        is_min = 'mindestbetrag' in get('positionstyp').lower() or 'mindestbetrag' in get('geschaeftsvorfall').lower()
        if (filters['mindest']=='Ja') != is_min:
            return False
    q = str(filters.get('q') or '').strip().lower()
    if q:
        blob = ' '.join(str(v or '') for v in row.values()).lower()
        if q not in blob:
            return False
    return True


def _fm505_fetch_imports(filters=None):
    _fm505_db_init()
    con = _fm501_db_connect()
    try:
        rows=[dict(r) for r in con.execute("SELECT * FROM imports WHERE COALESCE(deleted,0)=0 ORDER BY imported_at DESC, id DESC").fetchall()]
        if filters:
            rows=[r for r in rows if _fm505_filter_match(r, filters)]
        return rows
    finally:
        con.close()


def _fm505_fetch_closings(filters=None):
    _fm505_db_init()
    con = _fm501_db_connect()
    try:
        rows=[dict(r) for r in con.execute('SELECT * FROM closings WHERE import_id NOT IN (SELECT id FROM imports WHERE COALESCE(deleted,0)=1) ORDER BY abschlussdatum DESC, id DESC').fetchall()]
        if filters:
            rows=[r for r in rows if _fm505_filter_match(r, filters)]
        return rows
    finally:
        con.close()


def _fm505_fetch_items(filters=None, closing_id=None):
    _fm505_db_init()
    con = _fm501_db_connect()
    try:
        if closing_id:
            rows=[dict(r) for r in con.execute('SELECT * FROM fee_items WHERE closing_id=? ORDER BY sort_order, id', (closing_id,)).fetchall()]
        else:
            rows=[dict(r) for r in con.execute('SELECT * FROM fee_items WHERE closing_id NOT IN (SELECT id FROM closings WHERE import_id IN (SELECT id FROM imports WHERE COALESCE(deleted,0)=1)) ORDER BY abschlussdatum DESC, geschaeftsvorfall').fetchall()]
        if filters:
            rows=[r for r in rows if _fm505_filter_match(r, filters)]
        return rows
    finally:
        con.close()


def _fm505_summary_stats(filters=None):
    closings = _fm505_fetch_closings(filters)
    items = _fm505_fetch_items(filters)
    brutto=sum(float(r.get('betrag_brutto') or 0) for r in closings)
    netto=sum(float(r.get('betrag_netto') or 0) for r in closings)
    ust=sum(float(r.get('ust') or 0) for r in closings)
    months=[]
    by_month={}
    for r in closings:
        mk=r.get('monat_key') or r.get('abschlussdatum') or ''
        by_month[mk]=by_month.get(mk,0)+float(r.get('betrag_brutto') or 0)
    for mk,val in sorted(by_month.items()):
        months.append({'monat_key':mk,'betrag_brutto':val})
    top_map={}
    for i in items:
        key=i.get('geschaeftsvorfall') or ''
        rec=top_map.setdefault(key, {'geschaeftsvorfall':key,'gesamt':0.0,'anzahl':0.0,'weighted_unit_num':0.0,'unit_values':[]})
        ges=float(i.get('gesamtbetrag') or 0); anz=float(i.get('anzahl') or 0); stk=i.get('stueckkosten')
        rec['gesamt']+=ges; rec['anzahl']+=anz
        if anz and stk is not None:
            rec['weighted_unit_num']+=float(stk)*anz
            rec['unit_values'].append(float(stk))
    top=[]
    for rec in top_map.values():
        rec['avg_unit']=(rec['weighted_unit_num']/rec['anzahl']) if rec.get('anzahl') else (sum(rec['unit_values'])/len(rec['unit_values']) if rec['unit_values'] else 0)
        top.append(rec)
    top.sort(key=lambda x:x.get('gesamt') or 0, reverse=True)
    return {'closings':len(closings),'details':len(items),'brutto':brutto,'netto':netto,'ust':ust,'months':months,'top':top[:30]}


def _fm505_delete_import(import_id, app=None):
    _fm505_db_init()
    user = getattr(app, 'current_user_display', '') or getattr(app, 'current_user_key', '') or ''
    now = datetime.now().isoformat(timespec='seconds')
    con = _fm501_db_connect()
    try:
        con.execute('UPDATE imports SET deleted=1, deleted_at=?, deleted_by=? WHERE id=?', (now,user,import_id))
        con.commit()
        return True
    finally:
        con.close()


def _fm505_tree_sort(tree, col, descending=False, numeric=False):
    items=[]
    for iid in tree.get_children(''):
        val=tree.set(iid, col)
        if numeric:
            try:
                key=float(str(val).replace('€','').replace('.','').replace(',','.').strip())
            except Exception:
                key=0.0
        else:
            key=str(val).lower()
        items.append((key,iid))
    items.sort(reverse=descending)
    for idx, (_, iid) in enumerate(items):
        tree.move(iid,'',idx)
    tree.heading(col, command=lambda: _fm505_tree_sort(tree, col, not descending, numeric))


def _fm505_draw_line_chart(canvas, x, y, w, h, data, title, value_key, label_key):
    canvas.create_rectangle(x,y,x+w,y+h,fill=WHITE,outline=LINE)
    canvas.create_text(x+14,y+14,text=title,fill=BLUE,font=body_font(12,'bold'),anchor='w')
    if not data:
        canvas.create_text(x+w/2,y+h/2,text='Keine Daten für die gewählten Filter.',fill=TEXT2,font=body_font(10))
        return
    vals=[float(d.get(value_key) or 0) for d in data]
    max_v=max(vals) or 1
    min_v=min(vals) if vals else 0
    if max_v==min_v: min_v=0
    left=x+35; right=x+w-20; top=y+42; bottom=y+h-34
    canvas.create_line(left,bottom,right,bottom,fill=LINE)
    canvas.create_line(left,top,left,bottom,fill=LINE)
    n=len(data)
    points=[]
    for i,d in enumerate(data):
        xx=left+(right-left)*(i/(max(1,n-1)))
        yy=bottom-(bottom-top)*((float(d.get(value_key) or 0)-min_v)/(max_v-min_v if max_v!=min_v else 1))
        points.append((xx,yy,d))
    for i in range(len(points)-1):
        canvas.create_line(points[i][0],points[i][1],points[i+1][0],points[i+1][1],fill=BLUE,width=2)
    for xx,yy,d in points:
        canvas.create_oval(xx-3,yy-3,xx+3,yy+3,fill=RED,outline=RED)
        canvas.create_text(xx,bottom+12,text=str(d.get(label_key) or '')[:8],fill=TEXT2,font=body_font(8))


def _fm505_chart_data(filters, metric='brutto', chart_type='months'):
    items=_fm505_fetch_items(filters)
    closings=_fm505_fetch_closings(filters)
    if chart_type=='top':
        return _fm505_summary_stats(filters).get('top',[])[:10], 'gesamt', 'geschaeftsvorfall'
    if chart_type=='unit_dev':
        gv=str(filters.get('gv') or '').strip()
        data_map={}
        for i in items:
            if gv and gv.lower() not in str(i.get('geschaeftsvorfall') or '').lower():
                continue
            mk=i.get('monat_key') or ''
            anz=float(i.get('anzahl') or 0); ges=float(i.get('gesamtbetrag') or 0)
            rec=data_map.setdefault(mk, {'monat_key':mk,'gesamt':0.0,'anzahl':0.0})
            rec['gesamt']+=ges; rec['anzahl']+=anz
        out=[]
        for mk,rec in sorted(data_map.items()):
            out.append({'monat_key':mk,'avg_unit':(rec['gesamt']/rec['anzahl']) if rec['anzahl'] else 0})
        return out, 'avg_unit', 'monat_key'
    by={}
    for c in closings:
        mk=c.get('monat_key') or ''
        if metric=='netto': val=float(c.get('betrag_netto') or 0)
        elif metric=='ust': val=float(c.get('ust') or 0)
        else: val=float(c.get('betrag_brutto') or 0)
        by[mk]=by.get(mk,0)+val
    return [{'monat_key':mk,'value':val} for mk,val in sorted(by.items())], 'value', 'monat_key'


def _fm505_render_account_analysis_tool(self):
    _fm505_db_init()
    w,h=self.canvas.winfo_width(),self.canvas.winfo_height()
    top_y=112
    outer=tk.Frame(self.root,bg=BG)
    self.widget_items.append(outer)
    self.canvas.create_window(0,top_y,window=outer,anchor='nw',width=w,height=max(460,h-top_y-55))
    main=tk.Frame(outer,bg=BG)
    main.pack(fill='both',expand=True,padx=16,pady=(6,8))
    status_var=tk.StringVar(value='Datenbank: '+_fm501_konto_db_path())
    kpi_var=tk.StringVar(value='')
    selected={'closing_id':None,'import_id':None}
    sort_state={}
    filters={'bank':'','konto':'','monat':'','von':'','bis':'','gv':'','typ':'','mindest':'','plausibel':'','q':''}

    controls=tk.Frame(main,bg=BG); controls.pack(fill='x',pady=(0,6))
    def make_combo(parent,var,values,width=14):
        cb=ttk.Combobox(parent,textvariable=var,values=values,state='normal',font=body_font(9),width=width)
        return cb
    def current_filters():
        for k,v in filter_vars.items(): filters[k]=v.get().strip()
        return dict(filters)
    def current_chart_filters():
        f=dict(current_filters())
        f.update({
            'bank': chart_bank.get().strip(), 'konto': chart_konto.get().strip(), 'von': chart_von.get().strip(), 'bis': chart_bis.get().strip(), 'gv': chart_gv.get().strip()
        })
        return f

    def refresh_filter_values():
        banks=['']+_fm505_distinct_values('imports','bank')
        konten=['']+_fm505_distinct_values('imports','konto')
        for cb in [bank_cb, chart_bank_cb]: cb.configure(values=banks)
        for cb in [konto_cb, chart_konto_cb]: cb.configure(values=konten)
        gvs=['']+sorted({r.get('geschaeftsvorfall') or '' for r in _fm505_fetch_items({}) if r.get('geschaeftsvorfall')})
        types=['']+sorted({r.get('positionstyp') or '' for r in _fm505_fetch_items({}) if r.get('positionstyp')})
        gv_cb.configure(values=gvs); chart_gv_cb.configure(values=gvs); typ_cb.configure(values=types)

    def fill_tree(tree, rows, cols, formatter):
        tree.delete(*tree.get_children())
        for row in rows:
            iid=str(row.get('_iid') or row.get('id') or '')
            try: tree.insert('', 'end', iid=iid if iid else None, values=formatter(row))
            except Exception: tree.insert('', 'end', values=formatter(row))

    def refresh_tables():
        f=current_filters()
        imports=_fm505_fetch_imports(f)
        closings=_fm505_fetch_closings(f)
        items=_fm505_fetch_items(f, selected.get('closing_id')) if selected.get('closing_id') else _fm505_fetch_items(f)
        stats=_fm505_summary_stats(f)
        kpi_var.set(f"Entgeltabschlüsse: {stats['closings']}   |   Einzelpositionen: {stats['details']}   |   Brutto: {_fm501_fmt_eur(stats['brutto'])}   |   Netto: {_fm501_fmt_eur(stats['netto'])}   |   USt: {_fm501_fmt_eur(stats['ust'])}")
        fill_tree(import_tree, imports, None, lambda r:(r.get('bank') or '', r.get('konto') or '', r.get('filename') or '', r.get('imported_at') or '', r.get('closing_count') or 0, r.get('detail_count') or 0))
        fill_tree(closing_tree, closings, None, lambda r:(r.get('bank') or '', r.get('konto') or '', r.get('monat_key') or '', r.get('abschlussdatum') or '', r.get('buchungsdatum') or '', _fm501_fmt_eur(r.get('betrag_brutto')), _fm501_fmt_eur(r.get('betrag_netto')), _fm501_fmt_eur(r.get('ust')), str(r.get('rechnungsnummer') or '').split('|')[-1], r.get('source_file') or ''))
        fill_tree(item_tree, items, None, lambda r:(r.get('bank') or '', r.get('konto') or '', r.get('monat_key') or '', r.get('geschaeftsvorfall') or '', r.get('positionstyp') or '', _fm501_fmt_num(r.get('anzahl')) if r.get('anzahl') is not None else '', _fm501_fmt_eur(r.get('stueckkosten')) if r.get('stueckkosten') is not None else '', _fm501_fmt_eur(r.get('gesamtbetrag')), _fm501_fmt_eur(r.get('abweichung')), 'Ja' if r.get('plausibel') else 'Nein'))
        top=_fm505_summary_stats(f).get('top',[])
        fill_tree(top_tree, top, None, lambda r:(r.get('geschaeftsvorfall') or '', _fm501_fmt_eur(r.get('gesamt')), _fm501_fmt_num(r.get('anzahl')), _fm501_fmt_eur(r.get('avg_unit'))))
        status_var.set(f"Datenbank: {_fm501_konto_db_path()}   |   Sichtbare Importe: {len(imports)}")

    def draw_charts():
        # Tk meldet vor dem ersten Layoutdurchlauf Breite 1. Erst Layout abschließen,
        # dann zwei gleich breite Diagrammflächen mit festen Außen-/Innenabständen berechnen.
        try: charts.update_idletasks()
        except Exception: pass
        charts.delete('all')
        f=current_chart_filters()
        metric=chart_metric.get()
        measured=int(charts.winfo_width() or 0)
        available=max(900, measured, int(w)-32)
        outer=12; gap=18
        each=max(420, int((available-(2*outer)-gap)/2))
        chart_h=244
        d1,key1,label1=_fm505_chart_data(f, metric, 'months')
        d2,key2,label2=_fm505_chart_data(f, metric, 'unit_dev' if chart_kind.get()=='Stückkostenentwicklung' else 'top')
        _fm501_draw_bar_chart(charts, outer, 6, each, chart_h, d1, 'Entwicklung nach Monat', key1, label1)
        second_x=outer+each+gap
        if chart_kind.get()=='Stückkostenentwicklung':
            _fm505_draw_line_chart(charts, second_x, 6, each, chart_h, d2, 'Ø Stückkosten je Geschäftsvorfall pro Bank', key2, label2)
        else:
            _fm501_draw_bar_chart(charts, second_x, 6, each, chart_h, d2, 'Top Gebührenpositionen', key2, label2)
        charts.configure(scrollregion=(0,0,second_x+each+outer,chart_h+12))

    def refresh_all():
        refresh_filter_values(); refresh_tables()

    def import_file():
        try:
            from tkinter import filedialog
            path=filedialog.askopenfilename(title='SFIRM-Excel-Export auswählen', filetypes=[('Excel-Dateien','*.xlsx *.xlsm *.xltx *.xltm'),('Alle Dateien','*.*')])
            if not path: return
            res=_fm505_import_excel(path,self)
            if res.get('cancelled'): return
            if res.get('duplicate'):
                replace=False
                try:
                    replace=messagebox.askyesno(
                        'Kontenauswertung',
                        f"Diese Excel-Datei wurde bereits verarbeitet.\n\nBisher: {res.get('existing_bank') or '-'} / {res.get('existing_konto') or '-'}\nNeu: {res.get('bank') or '-'} / {res.get('konto') or '-'}\n\nBestehende Daten durch den aktuellen Import ersetzen?"
                    )
                except Exception:
                    replace=False
                if replace:
                    res=_fm505_import_excel(path,self,bank=res.get('bank'),konto=res.get('konto'),replace_existing=True)
                    try: messagebox.showinfo('Kontenauswertung', f"Daten ersetzt.\n\nBank: {res.get('bank')}\nKonto: {res.get('konto')}\nEntgeltabschlüsse: {res['closings']}\nDetailpositionen: {res['details']}")
                    except Exception: pass
                else:
                    try: messagebox.showinfo('Kontenauswertung','Bestehende Daten wurden nicht verändert.')
                    except Exception: pass
            else:
                try: messagebox.showinfo('Kontenauswertung', f"Daten aufgenommen.\n\nBank: {res.get('bank')}\nKonto: {res.get('konto')}\nEntgeltabschlüsse: {res['closings']}\nDetailpositionen: {res['details']}")
                except Exception: pass
            refresh_all()
        except Exception as exc:
            try: messagebox.showerror('Kontenauswertung','Aufnahme fehlgeschlagen:\n'+str(exc))
            except Exception: pass

    def delete_selected_import():
        sel=import_tree.selection()
        if not sel:
            try: messagebox.showwarning('Kontenauswertung','Bitte zuerst einen Importdatensatz auswählen.')
            except Exception: pass
            return
        iid=int(sel[0])
        ok=True
        try: ok=messagebox.askyesno('Kontenauswertung','Ausgewählten Importdatensatz inklusive Entgeltabschlüssen und Einzelpositionen aus der aktiven Auswertung entfernen?')
        except Exception: ok=True
        if ok:
            _fm505_delete_import(iid,self); selected['import_id']=None; selected['closing_id']=None; refresh_all(); draw_charts()

    def on_closing_select(_=None):
        sel=closing_tree.selection(); selected['closing_id']=int(sel[0]) if sel else None; refresh_tables()
    def on_import_select(_=None):
        sel=import_tree.selection(); selected['import_id']=int(sel[0]) if sel else None

    tk.Button(controls,text='Excel-Daten aufnehmen',command=import_file,bg='#CFEAD6',fg=TEXT,font=body_font(10,'bold'),relief='solid',bd=1,padx=12,pady=5).pack(side='left',padx=(0,8))
    tk.Button(controls,text='Aktualisieren',command=refresh_all,bg=WHITE,fg=TEXT,font=body_font(10),relief='solid',bd=1,padx=12,pady=5).pack(side='left',padx=(0,8))
    tk.Button(controls,text='Import löschen',command=delete_selected_import,bg='#FDDDDD',fg=TEXT,font=body_font(10),relief='solid',bd=1,padx=12,pady=5).pack(side='left',padx=(0,8))
    tk.Button(controls,text='Excel-Auswertung exportieren',command=lambda:_fm501_export_analysis_excel(self),bg=WHITE,fg=BLUE,font=body_font(10),relief='solid',bd=1,padx=12,pady=5).pack(side='left',padx=(0,8))
    tk.Label(controls,textvariable=status_var,bg=BG,fg=TEXT2,font=body_font(8)).pack(side='right')

    filter_box=tk.Frame(main,bg=WHITE,highlightbackground=LINE,highlightthickness=1); filter_box.pack(fill='x',pady=(0,6))
    filter_vars={k:tk.StringVar(value='') for k in ['bank','konto','monat','von','bis','gv','typ','mindest','plausibel','q']}
    tk.Label(filter_box,text='Tabellenfilter',bg=WHITE,fg=BLUE,font=body_font(10,'bold')).grid(row=0,column=0,sticky='w',padx=8,pady=(5,1))
    bank_cb=make_combo(filter_box,filter_vars['bank'],[''],12); konto_cb=make_combo(filter_box,filter_vars['konto'],[''],10); gv_cb=make_combo(filter_box,filter_vars['gv'],[''],18); typ_cb=make_combo(filter_box,filter_vars['typ'],[''],18)
    widgets=[('Bank',bank_cb),('Konto',konto_cb),('Monat',tk.Entry(filter_box,textvariable=filter_vars['monat'],font=body_font(9),width=7)),('Von',tk.Entry(filter_box,textvariable=filter_vars['von'],font=body_font(9),width=10)),('Bis',tk.Entry(filter_box,textvariable=filter_vars['bis'],font=body_font(9),width=10)),('Geschäftsvorfall',gv_cb),('Typ',typ_cb),('Mindestbetrag',ttk.Combobox(filter_box,textvariable=filter_vars['mindest'],values=['','Ja','Nein'],state='readonly',width=8,font=body_font(9))),('Plausibel',ttk.Combobox(filter_box,textvariable=filter_vars['plausibel'],values=['','Ja','Nein'],state='readonly',width=8,font=body_font(9))),('Suche',tk.Entry(filter_box,textvariable=filter_vars['q'],font=body_font(9),width=18))]
    for idx,(lab,wdg) in enumerate(widgets):
        tk.Label(filter_box,text=lab,bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=1,column=idx,sticky='w',padx=(8,2),pady=(0,1))
        wdg.grid(row=2,column=idx,sticky='ew',padx=(8,2),pady=(0,6))
    tk.Button(filter_box,text='Tabellen filtern',command=refresh_tables,bg=WHITE,fg=TEXT,font=body_font(9),relief='solid',bd=1).grid(row=2,column=len(widgets),padx=8,pady=(0,6),sticky='e')
    for c in range(len(widgets)): filter_box.grid_columnconfigure(c,weight=1)

    tk.Label(main,textvariable=kpi_var,bg=WHITE,fg=TEXT,font=body_font(10,'bold'),anchor='w',padx=10,pady=6,relief='solid',bd=1).pack(fill='x',pady=(0,6))

    chart_filter=tk.Frame(main,bg=WHITE,highlightbackground=LINE,highlightthickness=1); chart_filter.pack(fill='x',pady=(0,6))
    tk.Label(chart_filter,text='Schaubilder filtern',bg=WHITE,fg=BLUE,font=body_font(10,'bold')).grid(row=0,column=0,sticky='w',padx=8,pady=(5,1))
    chart_bank=tk.StringVar(value=''); chart_konto=tk.StringVar(value=''); chart_von=tk.StringVar(value=''); chart_bis=tk.StringVar(value=''); chart_gv=tk.StringVar(value=''); chart_metric=tk.StringVar(value='brutto'); chart_kind=tk.StringVar(value='Top-Kostentreiber')
    chart_bank_cb=make_combo(chart_filter,chart_bank,[''],12); chart_konto_cb=make_combo(chart_filter,chart_konto,[''],10); chart_gv_cb=make_combo(chart_filter,chart_gv,[''],18)
    cw=[('Bank',chart_bank_cb),('Konto',chart_konto_cb),('Von',tk.Entry(chart_filter,textvariable=chart_von,font=body_font(9),width=10)),('Bis',tk.Entry(chart_filter,textvariable=chart_bis,font=body_font(9),width=10)),('Kennzahl',ttk.Combobox(chart_filter,textvariable=chart_metric,values=['brutto','netto','ust'],state='readonly',font=body_font(9),width=10)),('Geschäftsvorfall',chart_gv_cb),('Diagramm 2',ttk.Combobox(chart_filter,textvariable=chart_kind,values=['Top-Kostentreiber','Stückkostenentwicklung'],state='readonly',font=body_font(9),width=18))]
    for idx,(lab,wdg) in enumerate(cw):
        tk.Label(chart_filter,text=lab,bg=WHITE,fg=TEXT2,font=body_font(8)).grid(row=1,column=idx,sticky='w',padx=(8,2),pady=(0,1))
        wdg.grid(row=2,column=idx,sticky='ew',padx=(8,2),pady=(0,6))
    tk.Button(chart_filter,text='Schaubilder aktualisieren',command=draw_charts,bg=WHITE,fg=BLUE,font=body_font(9,'bold'),relief='solid',bd=1).grid(row=2,column=len(cw),padx=8,pady=(0,6),sticky='e')
    charts=tk.Canvas(main,bg=BG,height=262,highlightthickness=0); charts.pack(fill='x',pady=(0,8))

    split=tk.Frame(main,bg=BG); split.pack(fill='both',expand=True)
    left=tk.Frame(split,bg=WHITE,highlightbackground=LINE,highlightthickness=1); left.pack(side='left',fill='both',expand=True,padx=(0,6))
    right=tk.Frame(split,bg=WHITE,highlightbackground=LINE,highlightthickness=1); right.pack(side='left',fill='both',expand=True,padx=(6,0))
    bottom=tk.Frame(main,bg=WHITE,highlightbackground=LINE,highlightthickness=1); bottom.pack(fill='both',expand=True,pady=(6,0))

    def build_tree(parent,title,cols,defs,height=7):
        tk.Label(parent,text=title,bg=WHITE,fg=BLUE,font=body_font(11,'bold')).pack(anchor='w',padx=8,pady=(6,2))
        wrap=tk.Frame(parent,bg=WHITE); wrap.pack(fill='both',expand=True,padx=8,pady=(0,8))
        tree=ttk.Treeview(wrap,columns=cols,show='headings',height=height)
        ysb=ttk.Scrollbar(wrap,orient='vertical',command=tree.yview); xsb=ttk.Scrollbar(wrap,orient='horizontal',command=tree.xview)
        tree.configure(yscrollcommand=ysb.set,xscrollcommand=xsb.set)
        for col,label,width_col,num in defs:
            tree.heading(col,text=label,command=lambda c=col,n=num:_fm505_tree_sort(tree,c,False,n))
            tree.column(col,width=width_col,stretch=True)
        tree.grid(row=0,column=0,sticky='nsew'); ysb.grid(row=0,column=1,sticky='ns'); xsb.grid(row=1,column=0,sticky='ew')
        wrap.grid_rowconfigure(0,weight=1); wrap.grid_columnconfigure(0,weight=1)
        return tree

    import_tree=build_tree(left,'Importierte Datensätze',('bank','konto','datei','zeit','abschl','pos'),[('bank','Bank',90,False),('konto','Konto',80,False),('datei','Datei',180,False),('zeit','Importiert',140,False),('abschl','Abschl.',60,True),('pos','Pos.',60,True)],height=4)
    closing_tree=build_tree(left,'Entgeltabschlüsse',('bank','konto','monat','abschluss','buchung','brutto','netto','ust','rechnung','quelle'),[('bank','Bank',80,False),('konto','Konto',80,False),('monat','Monat',60,False),('abschluss','Abschluss',90,False),('buchung','Buchung',90,False),('brutto','Brutto',90,True),('netto','Netto',90,True),('ust','USt',80,True),('rechnung','Rechnung',150,False),('quelle','Quelle',160,False)],height=7)
    item_tree=build_tree(right,'Einzelpositionen / Stückkosten',('bank','konto','monat','gv','typ','anzahl','stk','gesamt','abw','ok'),[('bank','Bank',75,False),('konto','Konto',75,False),('monat','Monat',60,False),('gv','Geschäftsvorfall',180,False),('typ','Typ',140,False),('anzahl','Anzahl',70,True),('stk','Stückkosten',90,True),('gesamt','Gesamt',90,True),('abw','Abw.',70,True),('ok','OK',50,False)],height=14)
    top_tree=build_tree(bottom,'Top-Kostentreiber',('gv','gesamt','anzahl','avg'),[('gv','Geschäftsvorfall',260,False),('gesamt','Gesamtkosten',120,True),('anzahl','Anzahl gesamt',110,True),('avg','Ø Stückkosten gewichtet',150,True)],height=6)
    closing_tree.bind('<<TreeviewSelect>>',on_closing_select); import_tree.bind('<<TreeviewSelect>>',on_import_select)
    refresh_all()
    try: self.root.after_idle(draw_charts)
    except Exception: draw_charts()

# Finale v0.505-Zuweisungen.
try:
    FiBuMateApp.render_account_analysis_tool = _fm505_render_account_analysis_tool
except Exception:
    pass


# FIBU_MATE_V0505_FIX_HEADER_ROW_UND_SCHAUBILD_LAYOUT
# Datum: 2026-07-13
# Behebt NameError header_row und negative Diagrammbreiten vor dem ersten Tk-Layout.


# FIBU_MATE_V0505_FIX_UNIQUE_IMPORT_UND_PASSWORTFOKUS_FINAL
# Datum: 2026-07-13
# Keine Layoutänderung: nur Doppelimport-Fehler und Passwort-Fokus korrigiert.

# FIBU_MATE_V0505_HOTFIX_UNIQUE_IMPORT_DELETED_ROWS
# Datum: 2026-07-13
# Doppelimport wird nun auch dann vor INSERT erkannt, wenn ein gleiches filename+file_hash bereits als gelöscht markiert ist.

# FIBU_MATE_V0505_REPLACE_EXISTING_IMPORT_FINAL
# Datum: 2026-07-13
# Doppelimport: bestehende Importdaten können nach Abfrage durch aktuellen Import ersetzt werden.
if __name__ == "__main__":
    FiBuMateApp().run()