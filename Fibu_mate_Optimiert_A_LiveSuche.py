# FiBuMate_PATCH_MARKER: 20260609_MENUZEILEN_DEBITOREN_PROTOCOL
# FiBuMate_PATCH_MARKER: 20260609_v0436_DIREKT_ABSCHLUSSKALENDER_EIN_MODUL
# FiBuMate_PATCH_MARKER: 20260609_v0436_DREI_MODULE_STICHTAGSPFLEGE_OHNE_IDS
# FiBuMate_PATCH_MARKER: 20260609_150049 (V0.436_TEAM_RELEASE_SCALING_MENU_CLEANUP_READABILITY_SAFE_TEXT)
import os
import sys
import json
import importlib


def _load_tool_module_from_file(module_path: str):
    """Lädt ein Tool-Modul immer direkt aus C:\python\bin\tools\*.py.
    Damit werden Import-Konflikte mit anderen Python-Umgebungen ausgeschlossen."""
    try:
        import importlib.util
        parts = (module_path or '').split('.')
        name = parts[-1] if parts else module_path
        file_path = os.path.join(BIN_DIR, 'tools', f"{name}.py")
        if not os.path.exists(file_path):
            # Fallback: ggf. liegt die Datei direkt im bin-Ordner
            file_path2 = os.path.join(BIN_DIR, f"{name}.py")
            if os.path.exists(file_path2):
                file_path = file_path2
        spec_name = f"fibumate_local_{name}"
        spec = importlib.util.spec_from_file_location(spec_name, file_path)
        if spec and spec.loader:
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            return mod
    except Exception:
        pass
    return None

import webbrowser
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.font as tkfont
from tkinter import messagebox
from datetime import datetime
from urllib.parse import quote

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

APP_NAME = "FiBu Mate"
VERSION_PREFIX = "0.4"
DEFAULT_BUILD = 0
VERSION_STATE_FILE = "version_state.json"
VERSION_HISTORY_FILE = "version_history.json"
# v0.436: Manuelle Zoomprofile entfernt; Darstellung skaliert automatisch anhand Fenster-/Monitorgröße.
ZOOM_PROFILE_FILE = None

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BIN_DIR = os.path.join(SCRIPT_DIR, "bin")
IMG_DIR = os.path.join(BIN_DIR, "Imgs")
ICON_DIR = os.path.join(IMG_DIR, "Icons")
USER_DIR = os.path.join(BIN_DIR, "User")
LEGACY_USER_DATA_PATH = os.path.join(USER_DIR, "fibu_mate_users.json")
NETWORK_ROOT = r"G:\BUC\FM Anwendung"
CENTRAL_CONFIG_DIR = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Config")
CENTRAL_USER_DATA_PATH = os.path.join(CENTRAL_CONFIG_DIR, "fibu_mate_users.json")
CENTRAL_RELEASE_DIR = os.path.join(NETWORK_ROOT, "Fibu_Mate_Doc", "Releases")
LATEST_UPDATE_FILE = "latest.json"
UPDATE_EVENT_FILE = "update_event.json"
UPDATE_CHECK_INTERVAL_MS = 60 * 1000
HIDDEN_TOOL_IDS = {"enbw_strom_tanken_upload"}  # nicht löschen: nur aus Oberfläche/Favoriten ausblenden


def _safe_json_load(path):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        return {}
    return {}


def resolve_user_data_path():
    env_path = os.environ.get("FIBUMATE_USER_DATA_PATH", "").strip()
    if env_path:
        return env_path
    for cfg in (os.path.join(SCRIPT_DIR, "config", "local_config.json"), os.path.join(CENTRAL_CONFIG_DIR, "local_config.json")):
        data = _safe_json_load(cfg)
        p = str(data.get("user_data_path", "")).strip()
        if p:
            return p
    return CENTRAL_USER_DATA_PATH


USER_DATA_PATH = resolve_user_data_path()

if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

BANNER_GROSS_PATH = os.path.join(IMG_DIR, "FMBanner_Gross.png")
BANNER_KLEIN_PATH = os.path.join(IMG_DIR, "FMBanner_Klein.png")
HELP_IMAGE_PATH = os.path.join(IMG_DIR, "FM_Help_Menu.png")
ICON_FILES = {
    "help": "metrostatus_question_support_6987.ico",
    "idea": "2799205-creative-idea-light_99776.ico",
    "doc_add": "addfileinterfacesymbolofpapersheetwithtextlinesandplussign_79821.ico",
    "doc_file": "fileinterfacesymboloftextpapersheet_79740.ico",
    "attach": "-attach-file_90371.ico",
    "calendar": "calendar-icon_34471.ico",
    "gear": "1486504840-cog-cogwheel-gear-repr-options-setting_81360.ico",
    "info": "1486504328-bullet-list-menu-lines-points-items-options_81334.ico",
    "lock": "lock_lock_15063.ico",
    "unlock": "lock_unlock_15064.ico",
    "xls": "ext_xls_filetype_icon_176238.ico",
    "pdf": "ext_pdf_filetype_icon_176234.ico",
    "compliance": "1486503790-bank-building-government-house-real-estate-panteon_81294.ico",
    "tax_reporting": "1486504352-checklist-clipboard-inventory-list-report-tasks-todo_81326.ico",
    "audit": "1486506277-like-thumbs-up-hands-gesture-finger-vote_81482.ico",
    "documentation": "1486485527-account-albums-screens-tabs_81163.ico",
    "search": "1486503763-bigger-enlarge-search-magnifier-magnify-zoom_81256.ico",
    "filter": "1486504837-descending-filter-filtering-tool-funnel-sort_81363.ico",
    "edit": "1486504369-change-edit-options-pencil-settings-tools-write_81307.ico",
}
INTERSPORT_LOGO_CANDIDATES = ["IS_Banner_lang.png", "Intersport_Logo.png", "INTERSPORT_Logo.png", "intersport_logo.png", "INTERSPORT.png"]

BLUE = "#004B93"
RED = "#E30613"
GOLD = "#FFD700"
STAR_GREY = "#7A7F87"
BG = "#E8EEF5"
HEADER = "#D3DEE9"
LINE = "#91A3B5"
TEXT = "#182431"
TEXT2 = "#445364"
SHADOW = "#A8B5C3"
WHITE = "#FFFFFF"
GREY_DISABLED = "#B9C3CF"
GREY_TILE = "#D6DCE4"

FONT_TITLE = ("Segoe UI", 26, "bold")
FONT_MENU = ("Segoe UI", 22, "bold")
FONT_TILE = ("Segoe UI", 18, "bold")
FONT_TILE_SMALL = ("Segoe UI", 15, "bold")
FONT_SMALL = ("Segoe UI", 10)
BASE_FONT_TITLE = FONT_TITLE
BASE_FONT_MENU = FONT_MENU
BASE_FONT_TILE = FONT_TILE
BASE_FONT_TILE_SMALL = FONT_TILE_SMALL
BASE_FONT_SMALL = FONT_SMALL
UI_SCALE = 1.0
UI_TEXT_SCALE = 1.00  # v0.436: kein manueller Textzoom; automatische, begrenzte UI-Skalierung.
UI_BODY_TEXT_SCALE = 1.00  # v0.433 Korrektur Paket 1g: 1920x1080 ist Body-Referenz; dynamische Lesbarkeit über body_font().
GLOBAL_TEXT_ZOOM_MIN = 0.70
GLOBAL_TEXT_ZOOM_MAX = 1.80
GLOBAL_TEXT_ZOOM_STEP = 0.025  # v0.433 Korrektur Paket 1h: feine Körnung pro Mausrad-Raster.
GLOBAL_TEXT_ZOOM = 1.00  # Fallback; v0.434 nutzt bereichsbezogene Zoomprofile.

def ui_s(value):
    try:
        return max(1, int(round(float(value) * UI_SCALE)))
    except Exception:
        return value

def scaled_font(font_tuple):
    try:
        family, size, *rest = font_tuple
        return tuple([family, max(7, int(round(size * UI_SCALE * UI_TEXT_SCALE)))] + rest)
    except Exception:
        return font_tuple

def ui_icon_size(base=36):
    """v0.433 Korrektur: zentrale Icon-Skalierung für kleinere Monitore."""
    try:
        return max(14, int(round(float(base) * UI_SCALE)))
    except Exception:
        return base

def body_scale_value():
    """v0.433 Korrektur Paket 1g: Dynamische Body-Skalierung.
    Referenz: 1920x1080 => Faktor 1.00. Kleinere Displays schrumpfen Body-/Tabellentexte nicht weiter,
    größere Displays dürfen moderat wachsen. Kacheln/Überschriften bleiben über scaled_font() getrennt.
    """
    try:
        return max(1.0, min(1.18, float(UI_SCALE))) * UI_BODY_TEXT_SCALE
    except Exception:
        return 1.0


def body_font(size=10, weight=None, underline=False, scale=1.0):
    """Zentrale Schrift für Pfadleiste, Fußleiste, Modulbeschreibungen, Modul-/Tabellentexte und Bedienbuttons.

    Lesbarkeitsregel V3: Sehr kleine Body-Schriften werden sichtbar auf mindestens 12 pt angehoben.
    Modulbeschreibungen und Canvas-Texte werden zusätzlich zentral abgefangen.
    """
    try:
        scaled = int(round(float(size) * UI_TEXT_SCALE * body_scale_value() * float(scale)))
        scaled = max(12, scaled)
    except Exception:
        scaled = max(12, int(size) if str(size).isdigit() else 12)
    style = []
    if weight:
        style.append(weight)
    if underline:
        style.append('underline')
    return tuple(['Segoe UI', scaled] + style)


# Zentrale Lesbarkeitsstufe V3: etwas größerer Fließtext, Überschriften bleiben unverändert.
READABILITY_MIN_LABEL_FONT = 13
READABILITY_MIN_BUTTON_FONT = 12
READABILITY_MIN_ENTRY_FONT = 12
READABILITY_MIN_TEXT_FONT = 12
READABILITY_MIN_TTK_FONT = 12
READABILITY_MIN_CANVAS_FONT = 13
READABILITY_TREEVIEW_ROWHEIGHT = 28
_CENTRAL_READABILITY_INSTALLED = False
_TK_ORIGINAL_INITS = {}
_TK_ORIGINAL_CONFIGS = {}
_TK_ORIGINAL_CANVAS_CREATE_TEXT = None


def _readability_font_tuple(font_value, minimum=11):
    """Hebt nur sehr kleine Tk-Fonts konservativ an; große/formatierte Fonts bleiben erhalten."""
    if font_value is None or font_value == "":
        return ("Segoe UI", minimum)
    try:
        if isinstance(font_value, tuple):
            if len(font_value) >= 2 and isinstance(font_value[1], int):
                if font_value[1] < minimum:
                    return tuple([font_value[0], minimum] + list(font_value[2:]))
                return font_value
        if isinstance(font_value, list):
            if len(font_value) >= 2 and isinstance(font_value[1], int):
                result = list(font_value)
                if result[1] < minimum:
                    result[1] = minimum
                return tuple(result)
        # Tk erlaubt Font-Strings; diese werden bewusst nicht geparst, um keine Systemfonts zu zerstören.
    except Exception:
        pass
    return font_value


def _readability_should_skip_canvas_text(text_value):
    """Schützt sehr kleine Symbol-/Icon-Texte vor unbeabsichtigter Vergrößerung."""
    try:
        value = str(text_value or "").strip()
        return value in {"", "★", "+", "−", "-", "HELP"}
    except Exception:
        return False


def _patch_tk_widget_font(widget_cls, minimum=11):
    if widget_cls in _TK_ORIGINAL_INITS:
        return
    original_init = widget_cls.__init__
    original_configure = getattr(widget_cls, "configure", None)
    original_config = getattr(widget_cls, "config", None)
    _TK_ORIGINAL_INITS[widget_cls] = original_init
    _TK_ORIGINAL_CONFIGS[widget_cls] = (original_configure, original_config)

    def patched_init(self, master=None, cnf=None, **kw):
        cnf = {} if cnf is None else dict(cnf)
        if "font" in kw:
            kw["font"] = _readability_font_tuple(kw.get("font"), minimum)
        elif "font" in cnf:
            cnf["font"] = _readability_font_tuple(cnf.get("font"), minimum)
        else:
            kw["font"] = ("Segoe UI", minimum)
        return original_init(self, master, cnf, **kw)

    def patched_configure(self, cnf=None, **kw):
        cnf = {} if cnf is None else dict(cnf)
        if "font" in kw:
            kw["font"] = _readability_font_tuple(kw.get("font"), minimum)
        if "font" in cnf:
            cnf["font"] = _readability_font_tuple(cnf.get("font"), minimum)
        return original_configure(self, cnf, **kw)

    widget_cls.__init__ = patched_init
    if original_configure is not None:
        widget_cls.configure = patched_configure
        widget_cls.config = patched_configure


def _patch_canvas_create_text():
    global _TK_ORIGINAL_CANVAS_CREATE_TEXT
    if _TK_ORIGINAL_CANVAS_CREATE_TEXT is not None:
        return
    _TK_ORIGINAL_CANVAS_CREATE_TEXT = tk.Canvas.create_text

    def patched_create_text(self, *args, **kw):
        if not _readability_should_skip_canvas_text(kw.get("text", "")):
            if "font" in kw:
                kw["font"] = _readability_font_tuple(kw.get("font"), READABILITY_MIN_CANVAS_FONT)
            else:
                kw["font"] = ("Segoe UI", READABILITY_MIN_CANVAS_FONT)
        return _TK_ORIGINAL_CANVAS_CREATE_TEXT(self, *args, **kw)

    tk.Canvas.create_text = patched_create_text


def install_central_readability(root=None):
    """Zentrale Lesbarkeitsverbesserung V3 für FiBu Mate und nachgeladene Tool-Module.

    V3 erhöht die bereits sichtbare V2-Wirkung moderat und wirkt weiter auf Canvas-Texte sowie spätere .configure(font=...)-Aufrufe.
    Dadurch werden Menübeschreibungen und viele modulinterne Texte noch etwas größer,
    ohne einzelne Moduldateien ändern zu müssen.
    """
    global _CENTRAL_READABILITY_INSTALLED
    if _CENTRAL_READABILITY_INSTALLED:
        return
    _CENTRAL_READABILITY_INSTALLED = True
    try:
        _patch_tk_widget_font(tk.Label, READABILITY_MIN_LABEL_FONT)
        _patch_tk_widget_font(tk.Button, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Entry, READABILITY_MIN_ENTRY_FONT)
        _patch_tk_widget_font(tk.Checkbutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Radiobutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Menubutton, READABILITY_MIN_BUTTON_FONT)
        _patch_tk_widget_font(tk.Listbox, READABILITY_MIN_TEXT_FONT)
        _patch_tk_widget_font(tk.Text, READABILITY_MIN_TEXT_FONT)
        _patch_canvas_create_text()
    except Exception:
        pass
    try:
        style = ttk.Style(root) if root is not None else ttk.Style()
        style.configure("TLabel", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TButton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TCheckbutton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TRadiobutton", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("TNotebook.Tab", font=("Segoe UI", READABILITY_MIN_TTK_FONT))
        style.configure("Treeview", font=("Segoe UI", READABILITY_MIN_TTK_FONT), rowheight=max(READABILITY_TREEVIEW_ROWHEIGHT, ui_s(READABILITY_TREEVIEW_ROWHEIGHT)))
        style.configure("Treeview.Heading", font=("Segoe UI", READABILITY_MIN_TTK_FONT, "bold"))
    except Exception:
        pass

MINI_WIDGET_W = 174
MINI_WIDGET_H = 30
MINI_WIDGET_GAP = 8

COLOR_PALETTE = [("Blau", BLUE), ("Grün", "#059669"), ("Rot", RED), ("Gelb", "#F59E0B"), ("Lila", "#7C3AED"), ("Pink", "#EC4899"), ("Dunkelgrau", "#334155"), ("Orange", "#F97316"), ("Türkis", "#06B6D4")]

ROLE_E1 = "E1 - Standard"
ROLE_E2 = "E2 - Erweitert"
ROLE_E3 = "E3 - Administrator"
ROLE_E4 = "E4 - System-Administrator"
ROLE_STANDARD = ROLE_E1
ROLE_ADMIN = ROLE_E3
OLD_ROLE_E4 = "Wagnerm"
ROLE_WAGNERM = ROLE_E4
SUPERUSER_KEY = "wagnerm"
ROLE_ORDER = [ROLE_E1, ROLE_E2, ROLE_E3, ROLE_E4]
ROLE_RANK = {
    ROLE_E1: 1, ROLE_E2: 2, ROLE_E3: 3, ROLE_E4: 4,
    "Standard": 1, "Ebene 1": 1, "E1": 1,
    "Ebene 2": 2, "E2": 2,
    "Administrator": 3, "Ebene 3": 3, "E3": 3,
    "System-Administrator": 4, OLD_ROLE_E4: 4, "Ebene 4": 4, "E4": 4,
}
ROLE_MIGRATION = {"Standard": ROLE_E1, "Administrator": ROLE_E3, "System-Administrator": ROLE_E4, OLD_ROLE_E4: ROLE_E4, "Ebene 1": ROLE_E1, "Ebene 2": ROLE_E2, "Ebene 3": ROLE_E3, "Ebene 4": ROLE_E4, "E1": ROLE_E1, "E2": ROLE_E2, "E3": ROLE_E3, "E4": ROLE_E4}

TOOL_REGISTRY = {
    "nike_pdf_to_excel": {"title": "Nike - PDF zu Excel", "module": "bin.tools.nike_pdf_to_excel", "favorite_label": "Nike PDF"},
    "nike_op_liste_pdf_check": {"title": "Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen", "module": "bin.tools.nike_op_liste_pdf_check", "favorite_label": "Nike OP PDF"},
    "invoice_pdf_collector": {"title": "Nike - Rechnungs-PDFs in Sammelordner", "module": "bin.tools.invoice_pdf_collector", "favorite_label": "Nike RE sammeln"},
    "enbw_strom_tanken_upload": {"title": "EnBW - Strom-Tanken Upload-Erstellung", "module": "bin.tools.enbw_strom_tanken_upload", "favorite_label": "EnBW Strom"},
    "supplier_invoice_afi_upload": {"title": "Lieferanten-Rechnung zu AFI-Upload", "module": "bin.tools.supplier_invoice_afi_upload", "favorite_label": "Lieferanten AFI"},
    "aramark_monatsabrechnungen_pdf_to_excel": {"title": "Aramark Monatsabrechnungen - PDF zu Excel", "module": "bin.tools.aramark_monatsabrechnungen_pdf_to_excel", "favorite_label": "Aramark Monat"},
    "debitoren_serienbrief": {"title": "Debitoren-Serienbrief", "module": "bin.tools.debitoren_serienbrief", "favorite_label": "Debitoren SB"},
    "monthly_close": {"title": "Monatsabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Monatsabschluss"},
    "quarterly_close": {"title": "Quartalsabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Quartalsabschluss"},
    "yearly_close": {"title": "Jahresabschluss", "module": "bin.tools.abschlusskalender", "favorite_label": "Jahresabschluss"},
    "deadline_maintenance": {"title": "Stichtagspflege", "module": "bin.tools.deadline_maintenance", "favorite_label": "Stichtage"},
    "x001_sap_test": {"title": "X001 SAP - Test", "module": "bin.tools.x001_sap_test", "favorite_label": "X001"},
    "tax_reporting": {"title": "Steuermeldungs-Cockpit", "module": "bin.tools.compliance_tax_reporting", "favorite_label": "Steuermeldungen"},
    "audit_cockpit": {"title": "Audit-Cockpit", "module": "bin.tools.compliance_audit_cockpit", "favorite_label": "Audit"},
    "documentation_center": {"title": "Dokumentationszentrale", "module": "bin.tools.compliance_documentation_center", "favorite_label": "Dokumente"},
}

MODULE_DESCRIPTIONS = {
    "nike_pdf_to_excel": "Mit diesem Modul lassen sich große Mengen an Nike PDF-Rechnungen in Excel ein Excel-Format ausgeben. Die auszugebenden Daten lassen sich filtern und sind individuell anpassbar.",
    "nike_op_liste_pdf_check": "Prüft die Vollständigkeit von Nike PDF-Rechnungen gegen eine OP-Liste: Rechnungsnummern aus PDF-Dateinamen werden mit Spalte B der Excel-Datei abgeglichen.",
    "monthly_close": "Interaktives Monatsabschluss-Cockpit mit Teamfortschritt, Aufgabenstatus, Fristwarnungen und Anlagen je Aufgabe.",
    "quarterly_close": "Interaktives Quartalsabschluss-Cockpit auf Basis der Monatsabschluss-Struktur mit Quartalsperioden.",
    "tax_reporting": "Überwacht steuerliche Meldungen und Meldefristen je Zeitraum inklusive Status, Zuständigkeit, Nachweisen, Historie und Abschlussbericht-Integration.",
    "audit_cockpit": "Zentrale Übersicht über kritische Systemereignisse wie Wiederöffnungen, Berechtigungsänderungen, Benutzeränderungen und nachträgliche Änderungen nach Abschluss.",
    "documentation_center": "Zentrale Suche und Prüfung aller Nachweise, Anlagen, Dokumentationspfade und Berichte aus FiBu Mate inklusive fehlender oder ungültiger Dokumentationen.",
    "yearly_close": "Interaktives Jahresabschluss-Cockpit für Geschäftsjahre vom 01.10. bis 30.09.",
    "deadline_maintenance": "Pflege der Abschluss-Stichtage (Dekadenabschluss, 18-Uhr, 08-Uhr, Monatsabschluss) inkl. Feiertage BW und automatischer Übernahme in Monats-/Quartals-/Jahresabschluss.",
    "x001_sap_test": "SAP-Scripting-Test; Scripting in SAP deaktiviert.",
    "invoice_pdf_collector": "In Excel gefilterte Rechnungsnummern aus PDF-Verzeichnis wählen und in neuen Sammelordner kopieren.",
    "enbw_strom_tanken_upload": "Erstellt aus EnBW E-Tankkosten-Abrechnungen eine SAP-AFI-uploadfähige CSV anhand der bestehenden Upload-Vorlage; Zuordnung nach Kennzeichen, Steuerlogik, Grundgebühren und Hinweis-Popup bei Abweichungen.",
    "page:compliance_audit": "Entwicklungsbereich für Compliance- und Audit-Funktionen: Steuermeldungen, Audit-Cockpit und Dokumentationszentrale bleiben gebündelt, werden aber nicht im produktiven Hauptmenü angezeigt.",
    "supplier_invoice_afi_upload": "Generisches CSV-Modul für Lieferantenrechnungen: erkennt relevante Spalten logisch, gleicht gegen eine wählbare AFI-/Kontierungsvorlage ab und exportiert eine uploadfähige AFI-CSV mit Spalten A-J.",
    "aramark_monatsabrechnungen_pdf_to_excel": "Erstellt aus einer oder mehreren Aramark-Monatsabrechnungs-PDFs eine gemeinsame Excel-Datei nach der festen Aramark-Vorlage; pro PDF entsteht eine frei benennbare eigene Umsatz-Spalte inklusive Plausibilitätsprüfung und Exportvorschau.",
}
DESCRIPTION_FONT = ("Segoe UI", 11)
DESCRIPTION_COLOR = TEXT2
DESCRIPTION_X_OFFSET = 14
DESCRIPTION_Y_OFFSET = 18


def maximize_window(window: tk.Tk):
    try:
        window.state("zoomed")
    except Exception:
        try:
            window.attributes("-zoomed", True)
        except Exception:
            pass


def x_pct(width: int, percent: float) -> float:
    return width * percent / 100


def y_pct(height: int, percent: float) -> float:
    return height * (1 - percent / 100)


def normalize_username(username: str) -> str:
    return " ".join(str(username).strip().split()).casefold()


def _rgb(hex_color: str):
    v = hex_color.lstrip("#")
    return tuple(int(v[i:i + 2], 16) for i in (0, 2, 4))


def _hex(rgb):
    return "#{:02x}{:02x}{:02x}".format(*rgb)


def blend(a: str, b: str, t: float) -> str:
    ar, ag, ab = _rgb(a)
    br, bg, bb = _rgb(b)
    return _hex((int(ar + (br - ar) * t), int(ag + (bg - ag) * t), int(ab + (bb - ab) * t)))


def darken(color: str, amount: float = 0.18) -> str:
    return blend(color, "#000000", amount)


def now_date_str() -> str:
    return datetime.now().strftime("%y%m%d")


def load_image(path: str):
    if not PIL_AVAILABLE:
        return None
    try:
        if path and os.path.exists(path):
            return Image.open(path)
    except Exception:
        return None
    return None


def resize_keep_ratio(image, max_w, max_h):
    ow, oh = image.size
    scale = min(1, max_w / ow, max_h / oh)
    return image.resize((max(1, int(ow * scale)), max(1, int(oh * scale))))


def find_intersport_logo():
    for name in INTERSPORT_LOGO_CANDIDATES:
        path = os.path.join(IMG_DIR, name)
        if os.path.exists(path):
            return path
    return None


class ArrowIndicator(tk.Canvas):
    def __init__(self, parent, direction, command, size=46):
        super().__init__(parent, width=size, height=size, bg=BG, highlightthickness=0, bd=0)
        self.direction = direction
        self.command = command
        self.size = size
        self.enabled = True
        self.hover = False
        self.bind("<Enter>", self._enter)
        self.bind("<Leave>", self._leave)
        self.bind("<Button-1>", self._click)
        self._draw()

    def tip_offset_from_center(self) -> float:
        s = self.size
        tip_y = (s - 11) if self.direction == "down" else 11
        return tip_y - (s / 2)
        self.bump_version_once(
            "2026-05-30_bu31_compliance_audit_deadlines_niketools_unsaved_fix",
            [
                "BU31: Stichtagspflege vollständig nach Compliance & Audit verschoben und aus Abschlusskalender entfernt.",
                "BU31: Tools - Hauptbuch erhält ein funktionsfähiges Untermenü Nike-Tools mit den drei Nike-Modulen im Standard-Modul-Menü-Layout.",
                "BU31: Modul Rechnungen aus Ordner sammeln in Nike - Rechnungs-PDFs in Sammelordner umbenannt und Modulbeschreibung aktualisiert.",
                "BU31: Dialog für ungespeicherte Änderungen nur bei echtem Dirty-State; nach Speichern + Übernehmen wird der Status bereinigt.",
                "BU31: Audit-Cockpit zeigt Details als öffnen mit ausführlichem Popup inklusive Zeitstempel.",
                "BU31: Default-Abschlussstichtag auf ersten Werktag nach Periodenende mit BW-Feiertagen umgestellt und Stichtage in Abschluss-/Steuermeldungsdaten synchronisiert.",
                "BU31: Benutzerregel angepasst: Nur Wagnerm darf Wagnerm umbenennen; E3 und niedriger dürfen sich nicht selbst umbenennen.",
            ],
        )
        self.bump_version_once(
            "2026-05-30_bu32_icons_deadline_sync_version_429",
            [
                "BU32: Nike-Tools-Kacheln zeigen wieder PDF- bzw. PDF/Excel-Icons im Untermenü Nike-Tools.",
                "BU32: Stichtagspflege synchronisiert gepflegte Abschluss-Stichtage verbindlich in Monats-, Quartals- und Jahresabschluss.",
                "BU32: Steuermeldungs-Cockpit übernimmt bei Kalender-Sync die gepflegten Stichtage als Fälligkeit je Zeitraum.",
                "BU32: Versionierung auf v0.429 fortgeschrieben und Versionsverlauf ergänzt.",
            ],
        )


    def set_enabled(self, enabled: bool):
        self.enabled = bool(enabled)
        self.configure(cursor="hand2" if self.enabled else "arrow")
        self._draw()

    def _enter(self, *_):
        self.hover = True
        self._draw()

    def _leave(self, *_):
        self.hover = False
        self._draw()

    def _click(self, *_):
        if self.enabled and self.command:
            self.command()

    def _draw(self):
        self.delete("all")
        s = self.size
        if self.hover and self.enabled:
            self.create_oval(4, 4, s - 4, s - 4, fill=blend(BG, WHITE, 0.55), outline="")
        c = BLUE if self.enabled else GREY_DISABLED
        if self.direction == "up":
            pts = (s / 2, 11, 13, s - 14, s - 13, s - 14)
            self.create_polygon(*pts, fill=c, outline=c)
            self.create_rectangle(s / 2 - 4, s - 16, s / 2 + 4, s - 11, fill=c, outline=c)
        else:
            pts = (13, 14, s - 13, 14, s / 2, s - 11)
            self.create_polygon(*pts, fill=c, outline=c)
            self.create_rectangle(s / 2 - 4, 11, s / 2 + 4, 16, fill=c, outline=c)


class Tile(tk.Canvas):
    def __init__(self, parent, app, tile_id, title, command=None, favorite_enabled=False, fixed_color=None, lock_tile=False, center_text=False, icon_type=None, corner_fold=False):
        super().__init__(parent, highlightthickness=0, bd=0, bg=BG, cursor="arrow", takefocus=True)
        self.app = app
        self.tile_id = tile_id
        self.title = title
        self.command = command
        self.favorite_enabled = favorite_enabled
        self.fixed_color = fixed_color
        self.lock_tile = lock_tile
        self.center_text = center_text
        self.icon_type = icon_type
        self.corner_fold = corner_fold
        self.tile_width = ui_s(300)
        self.tile_height = ui_s(150)
        self.hovered = False
        self.pressed = False
        self.star_bounds = None
        self.star_click_started = False
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<ButtonPress-1>", self.on_press)
        self.bind("<ButtonRelease-1>", self.on_release)
        self.bind("<FocusIn>", lambda e: self.app.set_focused_tile(self))
        self.bind("<Key-Return>", self.on_keyboard_activate)
        self.bind("<space>", self.on_keyboard_activate)
        self.draw()

    def base_color(self):
        return self.fixed_color or (self.app.current_tile_color() or BLUE)

    def current_color(self):
        c = self.base_color()
        if self.pressed:
            return darken(c, 0.28)
        if self.hovered:
            return darken(c, 0.16)
        return c

    def resize_tile(self, width, height):
        self.tile_width = int(width)
        self.tile_height = int(height)
        self.configure(width=self.tile_width + 16, height=self.tile_height + 16)
        self.draw()

    def title_font(self):
        return FONT_TILE_SMALL if len(self.title) > 28 else FONT_TILE

    def _font_with_size(self, font_spec, size):
        try:
            actual = tkfont.Font(root=self, font=font_spec).actual()
            family = actual.get('family') or 'Segoe UI'
            styles = []
            if (actual.get('weight') or '').lower() == 'bold':
                styles.append('bold')
            if (actual.get('slant') or '').lower() == 'italic':
                styles.append('italic')
            if bool(actual.get('underline')):
                styles.append('underline')
            return tuple([family, max(8, int(size))] + styles)
        except Exception:
            try:
                family = font_spec[0]
                rest = list(font_spec[2:]) if len(font_spec) > 2 else []
                return tuple([family, max(8, int(size))] + rest)
            except Exception:
                return ('Segoe UI', max(8, int(size)), 'bold')

    def _measure_wrapped_text_height(self, text, font_spec, max_width):
        try:
            f = tkfont.Font(root=self, font=font_spec)
            words = str(text or '').split() or ['']
            lines = []
            current = ''
            for word in words:
                candidate = word if not current else current + ' ' + word
                if f.measure(candidate) <= max_width:
                    current = candidate
                    continue
                if current:
                    lines.append(current)
                segment = ''
                for ch in word:
                    cand = segment + ch
                    if not segment or f.measure(cand) <= max_width:
                        segment = cand
                    else:
                        lines.append(segment)
                        segment = ch
                current = segment
            if current:
                lines.append(current)
            line_h = max(1, int(f.metrics('linespace')))
            return max(1, len(lines)) * line_h, max(1, len(lines))
        except Exception:
            return 0, 1

    def fitted_title_font(self, max_width, max_height, centered=False):
        """Lesbarkeitsschutz v0.436: bevorzugt größere Schrift, bleibt aber im Sollbereich."""
        base_font = self.app.zoomed_content_font(FONT_TILE_SMALL if len(self.title) > 24 else self.title_font())
        try:
            actual = tkfont.Font(root=self, font=base_font).actual()
            base_size = abs(int(actual.get('size') or 12))
        except Exception:
            base_size = 12
        preferred_min = 13 if not centered else 12
        start_size = max(base_size, ui_s(preferred_min))
        hard_max = max(start_size, 18 if not centered else 17)
        min_size = 9 if not centered else 8
        max_lines = 3 if not centered else 4
        chosen = self._font_with_size(base_font, min_size)
        chosen_h = 0
        for size in range(hard_max, min_size - 1, -1):
            candidate = self._font_with_size(base_font, size)
            text_h, text_lines = self._measure_wrapped_text_height(self.title, candidate, max_width)
            if text_h <= max_height and text_lines <= max_lines:
                return candidate, text_h
            chosen = candidate
            chosen_h = text_h
        return chosen, chosen_h

    def _draw_corner_fold(self, x1, y0):
        size = min(ui_s(30), max(18, int(self.tile_height * 0.20)))
        self.create_polygon(x1 - size, y0, x1, y0, x1, y0 + size, fill="#D6DCE4", outline="#C2CAD5")
        self.create_line(x1 - size, y0, x1, y0 + size, fill="#EEF2F6", width=1)

    def _draw_worksheet_icon(self, cx, cy):
        color = BLUE if self.lock_tile else "white"
        lw = 2
        # Nur das vordere Blatt: A4-Verhältnis, ohne hinteres Rahmen-/Blatt-Element
        page_w = 34
        page_h = 48
        x0 = cx - page_w / 2
        y0 = cy - page_h / 2
        x1 = cx + page_w / 2
        y1 = cy + page_h / 2
        fold = 10
        self.create_line(x0, y0, x1 - fold, y0, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1 - fold, y0, x1, y0 + fold, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1, y0 + fold, x1, y1, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1, y1, x0, y1, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x0, y1, x0, y0, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x1 - fold, y0, x1 - fold, y0 + fold, fill=color, width=lw)
        self.create_line(x1 - fold, y0 + fold, x1, y0 + fold, fill=color, width=lw)
        # kurze Textlinien oben wie in der Vorlage
        self.create_line(x0 + 6, y0 + 9, x0 + 18, y0 + 9, fill=color, width=lw + 1, capstyle="round")
        self.create_line(x0 + 6, y0 + 15, x0 + 24, y0 + 15, fill=color, width=lw + 1, capstyle="round")
        # Tabellenbereich im unteren Blattbereich
        tx0, ty0, tx1, ty1 = x0 + 5, y0 + 23, x1 - 5, y1 - 5
        self.create_rectangle(tx0, ty0, tx1, ty1, outline=color, width=lw)
        for x in (tx0 + (tx1 - tx0) / 4, tx0 + (tx1 - tx0) / 2, tx0 + 3 * (tx1 - tx0) / 4):
            self.create_line(x, ty0, x, ty1, fill=color, width=lw)
        for y in (ty0 + (ty1 - ty0) / 4, ty0 + (ty1 - ty0) / 2, ty0 + 3 * (ty1 - ty0) / 4):
            self.create_line(tx0, y, tx1, y, fill=color, width=lw)

    def _draw_calendar_icon(self, cx, cy):
        color = "white"
        lw = 3
        # rounded outline built with arcs + lines
        self.create_arc(cx - 25, cy - 19, cx - 7, cy - 1, start=90, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx + 7, cy - 19, cx + 25, cy - 1, start=0, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx - 25, cy + 8, cx - 7, cy + 26, start=180, extent=90, outline=color, width=lw, style="arc")
        self.create_arc(cx + 7, cy + 8, cx + 25, cy + 26, start=270, extent=90, outline=color, width=lw, style="arc")
        self.create_line(cx - 16, cy - 19, cx + 16, cy - 19, fill=color, width=lw)
        self.create_line(cx - 25, cy - 10, cx - 25, cy + 17, fill=color, width=lw)
        self.create_line(cx + 25, cy - 10, cx + 25, cy + 17, fill=color, width=lw)
        self.create_line(cx - 16, cy + 26, cx + 16, cy + 26, fill=color, width=lw)
        self.create_line(cx - 25, cy - 5, cx + 25, cy - 5, fill=color, width=lw)
        # rings
        self.create_line(cx - 13, cy - 25, cx - 13, cy - 13, fill=color, width=5, capstyle="round")
        self.create_line(cx + 13, cy - 25, cx + 13, cy - 13, fill=color, width=5, capstyle="round")
        # day boxes
        for x, y in [(cx - 12, cy + 4), (cx, cy + 4), (cx + 12, cy + 4), (cx - 12, cy + 16), (cx, cy + 16)]:
            self.create_rectangle(x - 4, y - 4, x + 4, y + 4, outline=color, width=2)
        # check mark
        self.create_line(cx + 10, cy + 16, cx + 15, cy + 21, fill=color, width=lw, capstyle="round")
        self.create_line(cx + 15, cy + 21, cx + 24, cy + 10, fill=color, width=lw, capstyle="round")

    def _draw_module_menu_icon(self, cx, cy):
        self._draw_worksheet_icon(cx, cy)

    def _draw_gear_icon(self, cx, cy):
        color = "white"
        for dx, dy in [(0, -16), (0, 16), (-16, 0), (16, 0), (-11, -11), (11, -11), (-11, 11), (11, 11)]:
            self.create_rectangle(cx + dx - 3, cy + dy - 3, cx + dx + 3, cy + dy + 3, fill=color, outline=color)
        self.create_oval(cx - 13, cy - 13, cx + 13, cy + 13, outline=color, width=3)
        self.create_oval(cx - 5, cy - 5, cx + 5, cy + 5, outline=color, width=2)

    def _draw_info_icon(self, cx, cy):
        color = "white"
        self.create_oval(cx - 18, cy - 18, cx + 18, cy + 18, outline=color, width=4)
        self.create_oval(cx - 3, cy - 12, cx + 3, cy - 6, fill=color, outline=color)
        self.create_rectangle(cx - 4, cy - 1, cx + 4, cy + 13, fill=color, outline=color)
        self.create_rectangle(cx - 8, cy + 9, cx + 8, cy + 15, fill=color, outline=color)

    def _draw_lock_icon(self, cx, cy):
        color = BLUE if self.lock_tile else "white"
        lw = 4
        self.create_arc(cx - 19, cy - 27, cx + 19, cy + 11, start=0, extent=180, style="arc", outline=color, width=lw)
        self.create_line(cx - 19, cy - 8, cx - 19, cy - 1, fill=color, width=lw)
        self.create_line(cx + 19, cy - 8, cx + 19, cy - 1, fill=color, width=lw)
        self.create_rectangle(cx - 23, cy - 4, cx + 23, cy + 26, outline=color, width=lw)
        self.create_oval(cx - 5, cy + 6, cx + 5, cy + 16, outline=color, width=lw)
        self.create_line(cx, cy + 15, cx, cy + 23, fill=color, width=lw)

    def _draw_main_icon(self, icon_type, cx, cy):
        if hasattr(self.app, "draw_tile_icon_image") and self.app.draw_tile_icon_image(self, icon_type, cx, cy):
            return
        if icon_type == "modules":
            self._draw_module_menu_icon(cx, cy)
        elif icon_type == "worksheet":
            self._draw_worksheet_icon(cx, cy)
        elif icon_type == "calendar":
            self._draw_calendar_icon(cx, cy)
        elif icon_type == "gear":
            self._draw_gear_icon(cx, cy)
        elif icon_type == "info":
            self._draw_info_icon(cx, cy)
        elif icon_type == "lock":
            self._draw_lock_icon(cx, cy)

    def draw(self):
        self.delete("all")
        pad = ui_s(8)
        off = ui_s(1) if self.pressed else 0
        x0, y0 = pad + off, pad + off
        x1, y1 = x0 + self.tile_width - off * 2, y0 + self.tile_height - off * 2
        if not self.pressed:
            self.create_rectangle(pad + 5, pad + 6, pad + 5 + self.tile_width, pad + 6 + self.tile_height, fill=SHADOW, outline=SHADOW)
        self.create_rectangle(x0, y0, x1, y1, fill=self.current_color(), outline=self.current_color())
        if self.corner_fold:
            self._draw_corner_fold(x1, y0)
        title_y = y0 + ui_s(14)
        title_color = BLUE if self.lock_tile else "white"
        icon_to_draw = "lock" if self.lock_tile else self.icon_type
        text_width = max(ui_s(110), self.tile_width - ui_s(44))
        if self.center_text and not icon_to_draw:
            centered_font, _ = self.fitted_title_font(text_width, max(ui_s(40), self.tile_height - ui_s(28)), centered=True)
            self.create_text((x0 + x1) / 2, (y0 + y1) / 2, text=self.title, anchor="center", fill=title_color, font=centered_font, width=text_width, justify="center")
        else:
            title_max_h = max(ui_s(28), int((y1 - y0) * 0.36))
            title_font, title_h = self.fitted_title_font(text_width, title_max_h, centered=False)
            self.create_text((x0 + x1) / 2, title_y, text=self.title, anchor="n", fill=title_color, font=title_font, width=text_width, justify="center")
            if icon_to_draw:
                min_icon_y = title_y + title_h + ui_s(20)
                default_icon_y = y0 + int((y1 - y0) * 0.66)
                icon_y = max(default_icon_y, min_icon_y)
                icon_y = min(icon_y, y1 - ui_s(34))
                close_period = self.app.current_close_period_label(self.tile_id) if hasattr(self.app, "current_close_period_label") else ""
                if close_period and self.tile_id in ("monthly_close", "quarterly_close", "yearly_close"):
                    icon_x = x0 + 98
                    self._draw_main_icon(icon_to_draw, icon_x, icon_y)
                    self.create_text(icon_x + 48, icon_y, text=close_period, anchor="w", fill=title_color, font=self.app.zoomed_content_font(("Segoe UI", 13, "italic")), width=max(120, x1 - icon_x - 58), justify="left")
                else:
                    self._draw_main_icon(icon_to_draw, (x0 + x1) / 2, icon_y)
        self.star_bounds = None
        if self.favorite_enabled and self.tile_id in TOOL_REGISTRY and (self.hovered or self.tile_id in self.app.favorites):
            sx, sy = x1 - 24, y0 + 24
            self.star_bounds = (sx - 18, sy - 18, sx + 18, sy + 18)
            self.create_text(sx, sy, text="★", fill=GOLD if self.tile_id in self.app.favorites else STAR_GREY, font=("Segoe UI", 24, "bold"))

    def point_in_star(self, x, y):
        return bool(self.star_bounds and self.star_bounds[0] <= x <= self.star_bounds[2] and self.star_bounds[1] <= y <= self.star_bounds[3])

    def on_enter(self, *_):
        self.hovered = True
        self.configure(cursor="hand2")
        self.draw()

    def on_leave(self, *_):
        self.hovered = False
        self.pressed = False
        self.star_click_started = False
        self.configure(cursor="arrow")
        self.draw()

    def on_press(self, event=None):
        self.focus_set()
        self.app.set_focused_tile(self)
        self.star_click_started = bool(event and self.point_in_star(event.x, event.y))
        self.pressed = not self.star_click_started
        self.draw()

    def on_release(self, event=None):
        if self.star_click_started:
            self.star_click_started = False
            if event and self.point_in_star(event.x, event.y):
                self.app.toggle_favorite(self.tile_id)
            return
        was = self.pressed
        self.pressed = False
        self.draw()
        if was and self.command:
            self.after(80, self.command)

    def on_keyboard_activate(self, *_):
        if self.command:
            self.command()


class FiBuMateApp:
    def zoomed_content_font(self, font_tuple):
        """v0.436: Kompatibilitätswrapper ohne manuellen Bereichs-/Textzoom."""
        return font_tuple

    def _calculate_ui_scale(self):
        try:
            sw = max(1, self.root.winfo_screenwidth())
            sh = max(1, self.root.winfo_screenheight())
            ww = max(1, self.root.winfo_width() or sw)
            wh = max(1, self.root.winfo_height() or sh)
            scale = min(ww / 1920.0, wh / 1080.0)
            if ww <= 1 or wh <= 1:
                scale = min(sw / 1920.0, sh / 1080.0)
            return max(0.72, min(1.18, scale))
        except Exception:
            return 1.0

    def init_responsive_scaling(self):
        """v0.436: zentrale automatische Skalierung ohne Benutzer-Zoomprofile."""
        global UI_SCALE, FONT_TITLE, FONT_MENU, FONT_TILE, FONT_TILE_SMALL, FONT_SMALL, MINI_WIDGET_W, MINI_WIDGET_H, MINI_WIDGET_GAP
        try:
            UI_SCALE = self._calculate_ui_scale()
            self.ui_scale = UI_SCALE
            FONT_TITLE = scaled_font(BASE_FONT_TITLE)
            FONT_MENU = scaled_font(BASE_FONT_MENU)
            FONT_TILE = scaled_font(BASE_FONT_TILE)
            FONT_TILE_SMALL = scaled_font(BASE_FONT_TILE_SMALL)
            FONT_SMALL = scaled_font(BASE_FONT_SMALL)
            MINI_WIDGET_W = ui_s(150)
            MINI_WIDGET_H = ui_s(26)
            MINI_WIDGET_GAP = ui_s(7)
            try:
                self.root.tk.call('tk', 'scaling', max(0.80, min(1.12, UI_SCALE)))
            except Exception:
                pass
        except Exception:
            self.ui_scale = 1.0

    def __init__(self):
        self.root = tk.Tk()
        install_central_readability(self.root)
        self.init_responsive_scaling()
        self.root.title(APP_NAME)
        self.root.protocol("WM_DELETE_WINDOW", self.confirm_exit)
        self._install_modal_toplevel_patch()
        maximize_window(self.root)
        self.page_history = []
        self.breadcrumb = []
        self.current_page = "launch"
        self.current_title = ""
        self.knowledge_view = "all"
        self.knowledge_show_start = self.load_knowledge_start_preference()
        self.knowledge_start_overlay = False
        self.knowledge_unsaved = False
        self.knowledge_overlay_offset_x = 0
        self.knowledge_overlay_offset_y = 0
        self.knowledge_overlay_drag = None
        self.widget_items = []
        self.focusable_tiles = []
        self.focus_index = -1
        self._suppress_next_global_return = False
        self._closing_in_progress = False
        self.global_text_zoom = 1.0
        self.zoom_profiles = {}
        self.current_scope_zoom = 1.0
        self.favorites = set()
        self.current_user_key = None
        self.current_user_display = ""
        self.user_data = self.load_user_data()
        self._user_data_mtime = self._get_user_data_mtime()
        self._live_permissions_started = False
        self._live_permissions_popup_open = False
        self._update_watcher_started = False
        self._update_popup_open = False
        self._notified_update_keys = set()
        self._last_update_event_seen = ""
        self.ensure_permissions_defaults()
        self.version_state = self.load_version_state()
        self.bump_version_once("2026-06-12_wissenszentrale_next_blocks_complete", ["Wissenszentrale Next Blocks: Word-Export, Anhänge, Kommentare, B2-Kategorienverwaltung, ESC-Logik, neutrale Navigation und To-Do-Rhythmus-Logik ergänzt."])
        self.bump_version_once("2026-06-12_wissenszentrale_functional_complete", ["Wissenszentrale: Layout kollisionsfrei neu strukturiert und Hauptfenster nach unten verschoben.", "Einträge können jetzt erstellt, gespeichert, bearbeitet, gesucht und über Kategorien gefiltert werden.", "Trefferliste, Gesamtliste, Detailansicht, To-Do-Filter und Kategorienanzeige technisch angebunden.", "Datenhaltung als lokale JSON-Datei unter APPDATA vorbereitet."])
        self.bump_version_once("2026-06-12_wissenszentrale_block5_overlay_underlay_panes", ["Wissenszentrale Block 5: Startseiten-Overlay maximiert zwischen FiBu-Mate-Kopf- und Fußleiste.", "Unterliegende Arbeitsansicht-Widgets werden während der Startseite ausgeblendet, damit keine Buttons/Dropdowns durchdrücken.", "Trefferfenster und Hauptfenster gemäß Markierung deutlich größer und höher ausgerichtet.", "Neuer-Eintrag-Ansicht nutzt den vergrößerten Hauptbereich besser aus."])
        self.bump_version_once("2026-06-12_wissenszentrale_block4_list_filters_overlay_fixed", ["Wissenszentrale Block 4: Startseiten-Overlay auf konstante Größe umgestellt; Drag verschiebt nur noch die Position.", "Treffer- und Hauptbereich deutlich vergrößert/dominanter dargestellt.", "Standardansicht im Hauptbereich auf Gesamtliste aller Einträge umgestellt, chronologisch neueste oben.", "Kategorie-Filterleiste mit vier Dropdowns ergänzt; To-Do-/Rhythmusfilter erscheint nur noch in der To-Do-Ansicht."])
        self.bump_version_once("2026-06-12_wissenszentrale_block3_drag_layout_cleanup", ["Wissenszentrale Block 3: Startseiten-Overlay mit sauberer Klick-/Drag-Erkennung korrigiert.", "Arbeitsansicht-Buttons nach unten versetzt, damit die Überschrift nicht mehr überlagert wird.", "Start-Overlay nach oben vergrößert und Drag-Ende strikt an das Loslassen der linken Maustaste gebunden.", "Beispieleinträge und Beispielkategorien aus der Wissenszentrale entfernt; sichtbar bleiben nur To-Do-bezogene Filter/Rhythmuswerte.", "Tools-Menüzeile im Hauptmenü kleiner skaliert als die Hauptmodul-Zeile darüber."])
        self.bump_version_once("2026-06-12_wissenszentrale_start_overlay", ["Wissenszentrale als eigener Hauptmenüpunkt nach Abschlusskalender vorbereitet.", "Arbeitsansicht nach Vorschlag A und vorgeschaltete Startseite nach Vorschlag E als Overlay umgesetzt.", "Startseiten-Checkbox speichert, ob die Startseite beim Modulstart angezeigt wird; Home-Button 'Start' öffnet die Startseite jederzeit über der aktiven Arbeitsansicht.", "Beim Wechsel der Arbeitsansicht über das Start-Overlay wird bei ungespeicherten Änderungen Speichern/Verwerfen/Abbrechen abgefragt."])
        self.bump_version_once("2026-06-12_central_readability_uplift_v3", ["Zentrale Lesbarkeitsverbesserung V3: Fließtexte, Beschreibungen, Buttons und Tabellen-/ttk-Texte moderat weiter vergrößert.", "Überschriften bleiben unverändert; nur Mindestgrößen für kleine Texte werden angehoben.", "Canvas-Texte mindestens 13 pt, Labels mindestens 13 pt, Buttons/Eingaben/ttk mindestens 12 pt, Treeview-Zeilenhöhe mindestens 28 px."])
        self.bump_version_once("2026-06-12_central_readability_uplift_v2", ["Zentrale Lesbarkeitsverbesserung V2: Mindestschriftgrößen sichtbar angehoben.", "Canvas-Texte werden jetzt ebenfalls zentral erfasst, damit Menü- und Modulbeschreibungen tatsächlich größer dargestellt werden.", "Spätere .configure(font=...)-Aufrufe werden abgefangen; dadurch wirkt die Verbesserung auch stärker in nachgeladenen Tool-Modulen."])
        self.bump_version_once("2026-06-12_central_readability_uplift", ["Zentrale Lesbarkeitsverbesserung ergänzt: sehr kleine Label-, Button-, Eingabe-, Text- und ttk-Schriften werden konservativ auf mindestens 10 pt angehoben.", "Die Verbesserung geht von Fibu_mate.py aus und wirkt auch auf nachgeladene Tool-Module, ohne diese einzeln anzupassen.", "Layouts bleiben geschützt: Es werden keine Tabellenstrukturen, Modulgrößen oder Buttonbreiten pauschal verändert."])
        self.bump_version_once(
            "2026-05-11_close_cutoff_subtasks_team_members_due_rules",
            [
                "Abschlusskalender angepasst: Kreditoren heißt nun Zentralregulierung, Controlling heißt nun Treasury.",
                "Monats- und Quartalsabschluss erweitert: Team-Mitglieder können im Bearbeitungsmodus je Übersichtskachel gepflegt und unter der nächsten Frist angezeigt werden.",
                "Monats- und Quartalsabschluss erweitert: Aufgaben können Unteraufgaben erhalten; Unteraufgaben sind separat abhakbar und steuern den Erledigt-Status der Hauptaufgabe.",
                "Fälligkeitsarten im Monats- und Quartalsabschluss auf Abschluss-Stichtag, x. Werktag des Folgemonats, x. Tag des Kalendermonats und Konkretes Datum umgestellt.",
                "Abschluss-Stichtag kann in der Abschlussübersicht gepflegt werden und wird automatisch auf Aufgaben mit Fälligkeitsart Abschluss-Stichtag übertragen.",
            ],
        )
        self.bump_version_once(
            "2026-05-11_nike_opliste_export_thread_performance_fix",
            [
                "Nike OP-Liste Modul stabilisiert: Export läuft im Hintergrund, damit FiBu Mate während großer .xlsx-Auswertungen nicht einfriert.",
                "Nike OP-Liste Modul optimiert: OP-Liste wird performanter über Zeilen-Iteratoren eingelesen.",
                "Nike OP-Liste Modul verbessert: Export-Button wird während der Verarbeitung gesperrt und Fortschritt/Fehler werden sauber zurückgemeldet.",
            ],
        )
        self.bump_version_once(
            "2026-05-11_nike_obsolete_removed_opliste_pdf_check_added",
            [
                "Nike-Module bereinigt: „Nike - Differenzen: PDF vs ZMIR6“ wurde entfernt, da das Modul obsolet ist.",
                "Nike-Module bereinigt: „Nike - Differenzbericht Sales und Lager“ wurde entfernt, da das Modul obsolet ist.",
                "Neues Modul ergänzt: „Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen“.",
                "Neues Nike-Modul vergleicht Rechnungsnummern aus OP-Liste Excel Spalte B mit Rechnungsnummern aus PDF-Dateinamen und exportiert zwei Auswertungsblätter.",
            ],
        )
        self.bump_version_once(
            "2026-05-12_monthly_close_recurring_workday_catalog_due_visibility",
            [
                "Monatsabschluss erweitert: Aufgaben können im Bearbeitungsmodus als wiederkehrend markiert werden.",
                "Monatsabschluss erweitert: zentraler Aufgabenkatalog für wiederkehrende Aufgaben eingeführt.",
                "Monatsabschluss erweitert: wiederkehrende Aufgaben werden automatisch in Folgemonate ab dem aktuellen Monat übernommen.",
                "Monatsabschluss erweitert: Fälligkeit kann als konkretes Datum, x. Werktag des Monats oder x. Werktag des Folgemonats gepflegt werden.",
                "Monatsabschluss erweitert: Werktagsberechnung berücksichtigt Montag bis Freitag sowie Feiertage Baden-Württemberg und verschiebt auf den nächsten Werktag.",
                "Monatsabschluss verbessert: berechnetes Fälligkeitsdatum wird im Aufgaben-Dialog direkt als Vorschau angezeigt.",
                "Monatsabschluss verbessert: bei Werktags-Fälligkeit wird das Eingabefeld für ein konkretes Fälligkeitsdatum ausgeblendet.",
                "Monatsabschluss verbessert: Tabelle zeigt bei regelbasierter Fälligkeit zusätzlich den Regelhinweis an.",
                "Monatsabschluss angepasst: ZM/Z4/Z5a bleiben vorerst manuell über Aufgabe, Fristart, Priorität und Fälligkeitsregel pflegbar.",
            ],
        )

        self.bump_version_once(
            "2026-05-13_user_permissions_close_sync_transfer_deadline_defaults",
            [
                "Benutzerverwaltung erweitert: Standardbenutzer sehen nur eigene Benutzerdaten, dürfen die eigene E-Mail-Adresse pflegen und die Benutzerliste ist scrollbar mit sichtbarer Neuanlage für System-Administrator.",
                "Berechtigungen angepasst: Rolle Wagnerm heißt nun System-Administrator, Standardbenutzer können das Berechtigungsmenü nicht mehr öffnen und das Berechtigungsmenü ist scrollbar.",
                "Monats- und Quartalsabschluss erweitert: Aufgaben können inklusive Unteraufgaben in den jeweils anderen Abschluss übernommen werden.",
                "Monats- und Quartalsabschluss angepasst: Standardbenutzer dürfen Aufgaben nur als erledigt markieren, wenn sie selbst als zuständig eingetragen sind.",
                "Teammitglieder werden auf folgende Perioden übertragen und zwischen Monats- und Quartalsabschluss synchronisiert.",
                "Fristart keine entfernt; Standard-Fristart ist intern.",
            ],
        )
        self.bump_version_once(
            "2026-05-13_close_team_detail_scroll_collapsible_subtasks_recurring_transfer_fix",
            [
                "Monats- und Quartalsabschluss erweitert: Unteraufgaben sind in der Team-Detailsansicht standardmäßig zugeklappt und können je Aufgabe über „Unteraufgaben ausklappen >“ ein- und ausgeblendet werden.",
                "Monats- und Quartalsabschluss korrigiert: Beim Übernehmen einer Aufgabe in den jeweils anderen Abschluss wird der Status Wiederkehrend mit übernommen.",
                "Monats- und Quartalsabschluss erweitert: Aufgabenübersichten in den Team-Detailsichten sind scrollbar.",
            ],
        )

        self.bump_version_once(
            "2026-05-15_ui_adjustments_icon_background_refinement",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )

        self.bump_version_once(
            "2026-05-15_ui_adjustments_mini_widgets_help_line_worksheet",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )

        self.bump_version_once(
            "2026-05-15_close_calendar_documentation_attachments_detail_table",
            [
                "Abschlusskalender erweitert: In den Aufgaben-Detailansichten von Monats- und Quartalsabschluss wurde eine neue Spalte Dokumentation ergänzt; je Aufgabe und Unteraufgabe kann ein Dokumentationspfad als Leitfaden bzw. Aufgabenbeschreibung hinterlegt, geöffnet und geändert werden.",
                "Abschlusskalender erweitert: Das Anlagen-Popup unterstützt mehrere Anlagenpfade je Aufgabe und Unteraufgabe, manuelle Pfadeingabe, Auswahl per Dateidialog und Bemerkungen/Informationen zur Bearbeitung für alle Rollen.",
                "Abschlusskalender angepasst: Spaltenbreiten, Ausrichtungen und Fälligkeitsdarstellung in den Detailtabellen wurden optimiert; Fällig wird nun als Datum plus Fälligkeitsart angezeigt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_ico_icon_replacements",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_menu_tile_alignment_doc_attach_buttons",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_documentation_remove_option",
            [
                "Monats- und Quartalsabschluss erweitert: Im Dokumentations-Popup kann ein hinterlegter Dokumentationspfad über ein Papierkorb-Icon entfernt werden.",
                "Monats- und Quartalsabschluss erweitert: Das Entfernen einer Dokumentation ist durch die Sicherheitsabfrage „Dokumentation entfernen?“ mit Ja/Nein geschützt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_ui_adjustments_task_table_column_separators",
            [
                "Anpassung der Benutzeroberfläche",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_edit_icon_delegate_owner",
            [
                "Monats- und Quartalsabschluss angepasst: Der Bearbeitungsmodus-Button verwendet nun das Stift-Icon als direkte Schaltfläche ohne blaue Kachel.",
                "Monats- und Quartalsabschluss erweitert: Administratoren und System-Administratoren können die Zuständigkeit einer Aufgabe im geöffneten Zeitraum per Delegieren-Button ändern.",
                "Monats- und Quartalsabschluss erweitert: Wird eine Aufgabe mit Unteraufgaben delegiert, werden alle Unteraufgaben im betreffenden Zeitraum mitdelegiert.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_delete_scope_cleanup_sort_fix",
            [
                "Monats- und Quartalsabschluss korrigiert: Beim Löschen wird die ausgewählte Aufgabe eindeutig über die konkrete Aufgabeninstanz bzw. eine stabile Aufgabenkennung identifiziert, damit keine falschen Aufgaben gelöscht werden.",
                "Monats- und Quartalsabschluss erweitert: Beim Löschen kann gewählt werden, ob die Aufgabe nur im aktuellen Zeitraum oder auch in allen folgenden Zeiträumen gelöscht wird.",
                "Monats- und Quartalsabschluss erweitert: Im Bearbeitungsmodus gibt es den Button „Alle Folgezeiträume bereinigen“, der folgende Perioden an den aktuellen Aufgabenbestand anpasst.",
                "Monats- und Quartalsabschluss angepasst: Aufgaben und Unteraufgaben werden alphabetisch nach Titel sortiert angezeigt.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_complete_ids_linking_pdf_no_reportlab",
            [
                "Abschlusskalender erweitert: Jahresabschluss und Aufgaben-Historie wurden eingebunden.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Aufgaben-IDs sind im Bearbeitungsdialog sichtbar und für Administrator/System-Administrator editierbar; bestehende fachliche Aufgaben erhalten initial QM001 bis QM016.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Aufgaben werden über identische Aufgaben-ID verknüpft; bei ID-Änderung wird die alte ID mit Deaktivierungsdatum archiviert.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Delegierungen können einmalig oder permanent auf Folgezeiträume übertragen werden.",
                "PDF-Berichte korrigiert: Export funktioniert ohne externe reportlab-Abhängigkeit über einen integrierten einfachen PDF-Generator.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_task_linking_mail_delegation",
            [
                "Aufgaben-Historie erweitert: Dezenter Reiter „Aufgaben verknüpfen“ ergänzt, inklusive Vorschlägen für gleichnamige Aufgaben in Monats-, Quartals- und Jahresabschluss.",
                "Abschlusskalender erweitert: Aufgaben-Verknüpfung setzt gemeinsame IDs automatisch nach Priorität M > Q > J und archiviert alte Aufgaben-IDs mit Deaktivierungsdatum.",
                "Monats-, Quartals- und Jahresabschluss erweitert: Bei Delegierung wird eine E-Mail an die neu zuständige Person vorbereitet; bei permanenter Delegierung mit Hinweis „bis auf Weiteres“.",
            ],
        )
        self.bump_version_once(
            "2026-05-15_close_period_lock_e_roles_reporting",
            [
                "Berechtigungen auf E1 bis E4 migriert: E1 Standard, E2 Erweitert, E3 Administrator, E4 System-Administrator; unbekannte Logins legen keine Benutzer mehr automatisch an.",
                "Benutzerverwaltung erweitert: Vorname und vollständiger Name werden für Ansprache und E-Mail-Kommunikation vorbereitet.",
                "Abschlusskalender erweitert: Zeiträume können ab E3 nach Ablauf des Abschluss-Stichtags geschlossen und wieder geöffnet werden; Wiederöffnung erfordert Begründung und automatische E-Mail-Benachrichtigung an E3/E4.",
                "Abschlussberichte erweitert: Berichtsdialog mit Signatur-/Freigabeoption, ausführlichere Inhalte, Änderungsprotokoll und ReportLab-Unterstützung mit einfachem Fallback.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_ui_login_usermgmt_closetiles_automail_fix",
            [
                "Benutzeranmeldung korrigiert: Unbekannte Benutzernamen werden endgültig nicht mehr automatisch angelegt.",
                "Benutzerverwaltung angepasst: Speichern- und Löschen-Buttons sind getrennt, gleich groß und ohne Aktionsspalten-Überschrift dargestellt.",
                "Mini-Widgets angepasst: Änderung-vorschlagen-Icon wurde auf die neue Icon-Datei umgestellt.",
                "Abschlusskalender angepasst: Monats-, Quartals- und Jahresabschluss-Kacheln zeigen den aktuellen Zeitraum direkt neben dem Kalendericon.",
                "Abschlusskalender erweitert: E4 kann Auto-Mail kompakt ein- und ausschalten; Abschluss- und Wiederöffnungs-Mails beachten diese Einstellung.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_v04_compliance_audit_tax_doc_modules",
            [
                "Major-Version auf v0.4 angehoben.",
                "Neues Hauptmenü Compliance & Audit ergänzt und Hauptmenü-Anordnung angepasst.",
                "Neues Modul Steuermeldungs-Cockpit mit Meldearten, Status, Nachweisen, Freigabe, Historie, PDF-Bericht und Audit-Anbindung ergänzt.",
                "Neues Modul Audit-Cockpit mit zentralem Audit-Log, Risikostufen, Filterung, PDF-Bericht und manuellem Archivierungsbutton für E3/E4 ergänzt.",
                "Neue Dokumentationszentrale mit Suche, Pfadprüfung, manuellen Dokumenten, PDF-/Excel-Export und Fehlende-Dokumentation-Bericht ergänzt.",
            ],
        )
        self.bump_version_once(
            "2026-05-27_scroll_modal_dropdown_shell_fix",
            [
                "UI-Interaktion zentral verbessert: Scrollbereiche reagieren nur noch, wenn tatsächlich mehr Inhalt vorhanden ist als angezeigt werden kann.",
                "Popup- und Dialogfenster werden modal behandelt; ESC schließt bevorzugt das aktive Dialogfenster statt den Hintergrund zu bedienen.",
                "Mousewheel-Routing angepasst: Dropdowns und aktive Dialoge erhalten Vorrang, damit nicht versehentlich Menüs im Hintergrund scrollen.",
            ],
        )
        self.banner_big = load_image(BANNER_GROSS_PATH)
        self.banner_small = load_image(BANNER_KLEIN_PATH)
        self.help_image = load_image(HELP_IMAGE_PATH)
        logo_path = find_intersport_logo()
        self.intersport_logo = load_image(logo_path) if logo_path else None
        self.image_refs = []
        self.ensure_version_429_once()
        self.ensure_version_430_once()
        self.ensure_version_431_once()
        self.ensure_version_432_once()
        self.ensure_version_433_once()
        self.ensure_version_434_once()
        self.ensure_version_435_once()
        self.normalize_version_after_zoom_patch()

        self.bump_version_once(
            "2026-06-05_afi_uploads_supplier_invoice_module",
            [
                "Tools - Hauptbuch erweitert: Neues Untermenü AFI-Uploads hinzugefügt.",
                "EnBW - Strom-Tanken Upload-Erstellung in das Untermenü AFI-Uploads verschoben.",
                "Neues generisches Modul Lieferanten-Rechnung zu AFI-Upload für robuste CSV-Rechnungsanalyse und AFI-CSV-Export ergänzt.",
            ],
        )

        self.bump_version_once(
            "2026-06-05_enbw_strom_tanken_upload",
            [
                "Tools - Hauptbuch erweitert: Neues Modul EnBW - Strom-Tanken Upload-Erstellung hinzugefügt.",
                "EnBW-Modul erstellt SAP-AFI-uploadfähige CSV mit unveränderter Spaltenlogik A-J, Zuordnung nach normalisiertem Kennzeichen und Hinweis-Popup bei Namensabweichungen bzw. fehlender Vorlage.",
                "EnBW-Modul berücksichtigt gerundete Netto-/MwSt-Beträge, Steuerkennzeichen VD/V2/V0 und 19%-Grundgebühren aus der EnBW-Rechnung.",
            ],
        )

        self.bump_version_once(
            "2026-06-09_v0436_team_release_scaling_menu_cleanup",
            [
                "Version 0.436: Compliance & Audit als eigene Kachel in den Bereich In Entwicklung verschoben.",
                "Abschlusskalender angepasst: Stichtagspflege zwischen Jahresabschluss und Aufgaben-Historie einsortiert und bleibt federführend für Stichtage/Zuständigkeiten.",
                "Globale Zoomfunktion inklusive sichtbarer Zoomleiste, Zoomprofile und Strg+Mausrad-Zoom deaktiviert; FiBu Mate skaliert automatisch mit Fenster- und Monitorgröße auf Referenz 1920x1080.",
                "Kopfzeile nach Finance-Mate-Vorbild optisch bereinigt: kleinere Überschrift, getrennte Breadcrumb-/Favoriten-/Mini-Widget-Zonen und optimierte Positionen für Zurück, Änderung vorschlagen und Hilfe.",
                "Fußleiste um 15% reduziert; Inhalte skalieren mit der automatischen UI-Skalierung.",
                "Lesbarkeitsschutz ergänzt: zu kleine Texte werden innerhalb ihres Sollbereichs bevorzugt größer dargestellt, ohne Strukturen zu berühren oder unschön umzubrechen.",
            ],
        )
        self.create_footer()
        self.canvas = tk.Canvas(self.root, highlightthickness=0, bg=BG, cursor="arrow")
        self.canvas.pack(side="top", fill="both", expand=True)
        self.canvas.bind("<Configure>", self.on_resize)
        self.active_scroll_canvas = None
        # v0.436: keine globalen Zoom-Bindings mehr; automatische Skalierung übernimmt die Darstellung.
        for key, handler in [("<Escape>", self.handle_escape), ("<Return>", self.handle_enter), ("<Tab>", self.handle_tab), ("<Shift-Tab>", self.handle_shift_tab), ("<ISO_Left_Tab>", self.handle_shift_tab)]:
            self.root.bind_all(key, handler)
        self.show_page("launch", add_to_history=False)

    def install_zoom_mouse_bindings(self):
        """v0.434 Paket 1C: robuste Strg+Mausrad-Bindings für Windows/Linux und Canvas/Widgets."""
        try:
            sequences = (
                "<MouseWheel>",
                "<Control-MouseWheel>",
                "<Control-Button-4>", "<Control-Button-5>",
                "<Control-ButtonPress-4>", "<Control-ButtonPress-5>",
                "<Control-0>",
            )
            for seq in sequences:
                handler = self.reset_global_text_zoom if seq == "<Control-0>" else self.on_global_mousewheel
                try:
                    self.root.bind_all(seq, handler, add="+")
                except Exception:
                    pass
                try:
                    self.root.bind(seq, handler, add="+")
                except Exception:
                    pass
                try:
                    if hasattr(self, "canvas") and self.canvas is not None:
                        self.canvas.bind(seq, handler, add="+")
                except Exception:
                    pass
        except Exception:
            pass

    # === FiBu Mate UI interaction helpers: START ===
    def _install_modal_toplevel_patch(self):
        try:
            if getattr(tk, "_fibu_mate_modal_patch_installed", False):
                return
            original_toplevel = tk.Toplevel
            app = self
            class FiBuMateModalToplevel(original_toplevel):
                def __init__(self, *args, **kwargs):
                    super().__init__(*args, **kwargs)
                    try:
                        self.after_idle(lambda w=self: app._auto_modalize_toplevel(w))
                    except Exception:
                        pass
            tk._fibu_mate_original_toplevel = original_toplevel
            tk.Toplevel = FiBuMateModalToplevel
            tk._fibu_mate_modal_patch_installed = True
        except Exception:
            pass

    def _auto_modalize_toplevel(self, win):
        try:
            if not win or not win.winfo_exists() or win is self.root:
                return
            if getattr(win, "_fibu_mate_no_auto_modal", False):
                return
            try:
                if bool(win.overrideredirect()):
                    return
            except Exception:
                pass
            self._make_modal(win)
            try:
                self.root.after_idle(self.apply_global_text_zoom)
            except Exception:
                pass
        except Exception:
            pass

    def _make_modal(self, win):
        try:
            if not win or not win.winfo_exists():
                return win
            if getattr(win, "_fibu_mate_modalized", False):
                return win
            win._fibu_mate_modalized = True
            try:
                win.transient(self.root)
            except Exception:
                pass
            try:
                win.grab_set()
            except Exception:
                pass
            try:
                win.focus_force()
            except Exception:
                pass
            try:
                win.bind("<Escape>", lambda event, w=win: (self._close_modal_window(w), "break")[-1], add="+")
            except Exception:
                pass
            try:
                win.protocol("WM_DELETE_WINDOW", lambda w=win: self._close_modal_window(w))
            except Exception:
                pass
            return win
        except Exception:
            return win

    def _close_modal_window(self, win=None):
        try:
            target = win
            if target is None:
                try:
                    target = self.root.grab_current()
                except Exception:
                    target = None
            if target and target is not self.root and target.winfo_exists():
                try:
                    target.grab_release()
                except Exception:
                    pass
                target.destroy()
                return True
        except Exception:
            pass
        return False

    def _has_modal_dialog(self):
        try:
            current = self.root.grab_current()
            return bool(current and current is not self.root and current.winfo_exists())
        except Exception:
            return False

    def _focus_is_dropdown_or_menu(self):
        try:
            fw = self.root.focus_get()
            if fw is None:
                return False
            cls = fw.winfo_class().lower()
            if cls in ("combobox", "tcombobox", "menu", "menubutton", "optionmenu"):
                return True
            name = str(fw).lower()
            return "popdown" in name or "combobox" in name
        except Exception:
            return False

    def _scroll_canvas_has_overflow(self, canvas):
        try:
            if canvas is None or not canvas.winfo_exists():
                return False
            canvas.update_idletasks()
            bbox = canvas.bbox("all")
            if not bbox:
                return False
            return max(0, bbox[3] - bbox[1]) > max(1, canvas.winfo_height()) + 2
        except Exception:
            return False

    def _sync_scrollbar_visibility(self, canvas, scrollbar=None):
        try:
            needs_scroll = self._scroll_canvas_has_overflow(canvas)
            if scrollbar is not None and scrollbar.winfo_exists():
                manager = scrollbar.winfo_manager()
                if needs_scroll:
                    if not scrollbar.winfo_ismapped():
                        if manager == "grid":
                            scrollbar.grid()
                        else:
                            scrollbar.pack(side="right", fill="y")
                else:
                    if scrollbar.winfo_ismapped():
                        if manager == "grid":
                            scrollbar.grid_remove()
                        else:
                            scrollbar.pack_forget()
                    try:
                        canvas.yview_moveto(0)
                    except Exception:
                        pass
            return needs_scroll
        except Exception:
            return False

    def register_scroll_canvas(self, canvas, scrollbar=None):
        try:
            if canvas is None:
                return
            canvas._fibu_mate_scrollbar = scrollbar
            def _activate(_event=None, c=canvas):
                if not self._has_modal_dialog() and self._scroll_canvas_has_overflow(c):
                    self.active_scroll_canvas = c
            def _deactivate(_event=None, c=canvas):
                if self.active_scroll_canvas is c:
                    self.active_scroll_canvas = None
            canvas.bind("<Enter>", _activate, add="+")
            canvas.bind("<Leave>", _deactivate, add="+")
            canvas.bind("<MouseWheel>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c), add="+")
            canvas.bind("<Button-4>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c, linux_delta=-1), add="+")
            canvas.bind("<Button-5>", lambda e, c=canvas: self._route_mousewheel_to_canvas(e, c, linux_delta=1), add="+")
            self._sync_scrollbar_visibility(canvas, scrollbar)
        except Exception:
            pass

    def _route_mousewheel_to_canvas(self, event, canvas=None, linux_delta=None):
        try:
            if self._has_modal_dialog():
                return "break"
            if self._focus_is_dropdown_or_menu():
                return None
            target = canvas or self.active_scroll_canvas
            if target is None or not target.winfo_exists():
                return None
            scrollbar = getattr(target, "_fibu_mate_scrollbar", None)
            if not self._sync_scrollbar_visibility(target, scrollbar):
                return "break"
            if linux_delta is not None:
                units = linux_delta
            else:
                delta = getattr(event, "delta", 0)
                units = int(-delta / 120) if delta else 0
                if units == 0 and delta:
                    units = -1 if delta > 0 else 1
            if units:
                target.yview_scroll(units, "units")
                try:
                    if hasattr(self, "_update_scroll_indicators"):
                        self._update_scroll_indicators()
                except Exception:
                    pass
            return "break"
        except Exception:
            return "break"
    # === FiBu Mate UI interaction helpers: END ===
    def get_icon_photo(self, icon_key, max_w=32, max_h=32):
        if not PIL_AVAILABLE:
            return None
        if not hasattr(self, "icon_cache"):
            self.icon_cache = {}
        cache_key = (icon_key, int(max_w), int(max_h))
        if cache_key in self.icon_cache:
            return self.icon_cache[cache_key]
        file_name = ICON_FILES.get(icon_key, icon_key)
        path = os.path.join(ICON_DIR, file_name)
        img = load_image(path)
        if not img:
            return None
        try:
            img = img.convert("RGBA")
            img = resize_keep_ratio(img, max_w, max_h)
            photo = ImageTk.PhotoImage(img)
            self.icon_cache[cache_key] = photo
            return photo
        except Exception:
            return None

    def draw_canvas_icon(self, canvas, icon_key, x, y, max_w=22, max_h=22):
        photo = self.get_icon_photo(icon_key, max_w, max_h)
        if not photo:
            return False
        canvas.create_image(x, y, image=photo)
        return True

    def draw_tile_icon_image(self, tile, icon_type, cx, cy):
        mapping = {"calendar": "calendar", "gear": "gear", "info": "info", "lock": "lock", "worksheet": "xls", "modules": "xls", "compliance": "compliance", "tax_reporting": "tax_reporting", "audit": "audit", "documentation": "documentation"}
        if icon_type == "pdf_xls":
            p1 = self.get_icon_photo("pdf", 50, 50)
            p2 = self.get_icon_photo("xls", 50, 50)
            if p1 and p2:
                tile.create_image(cx - 31, cy, image=p1)
                tile.create_image(cx + 31, cy, image=p2)
                return True
            return False
        icon_key = mapping.get(icon_type)
        if not icon_key:
            return False
        photo = self.get_icon_photo(icon_key, 48, 48)
        if not photo:
            return False
        tile.create_image(cx, cy, image=photo)
        return True

    def version_label_text(self) -> str:
        build = int(self.version_state.get("build", DEFAULT_BUILD))
        return f"v{VERSION_PREFIX}{build}.{now_date_str()}"


    def ensure_version_429_once(self):
        """BU33a: Version v0.429 und Versionsverlauf robust sicherstellen.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.429 dem Build 29,
        nicht 429. Diese Methode korrigiert auch bereits falsch gespeicherte
        Builds wie 429, die als v0.4429 angezeigt werden.
        """
        update_id = "2026-05-30_bu33a_force_version_429"
        bullets = [
            "BU33a: Versionierung robust korrigiert; sichtbare Version wird auf v0.429 gesetzt.",
            "BU33a: Falsch gespeicherter Build 429 wird auf Build 29 korrigiert, damit nicht v0.4429 angezeigt wird.",
            "BU33a: Versionsverlauf-Eintrag ergänzt, auch wenn vorherige Update-Marker bereits als angewendet gespeichert waren.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 29 or current_build > 100:
                self.version_state["build"] = 29
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.429",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_430_once(self):
        """BU33b: Version v0.430 und Versionsverlauf robust sicherstellen.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.430 dem Build 30.
        Diese Methode wird nach allen regulären bump_version_once-Blöcken und vor
        create_footer() ausgeführt, damit das Versionslabel direkt korrekt ist.
        """
        update_id = "2026-05-30_bu33b_closing_deadline_sync_version_430"
        bullets = [
            "BU33b: Abschlusskalender-Synchronisation versioniert.",
            "BU33b: Monatsabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Quartalsabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Jahresabschluss übernimmt gepflegte Stichtage aus der Stichtagspflege und aktualisiert Aufgaben mit Abschluss-Stichtag.",
            "BU33b: Bestehende Periodendateien aktualisieren closing_cutoff_date und die Fälligkeiten betroffener Aufgaben.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 30 or current_build > 100:
                self.version_state["build"] = 30
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.430",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_431_once(self):
        """v0.431: Abschluss-Stichtag in Kalendern nur noch aus Stichtagspflege.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.431 dem Build 31.
        """
        update_id = "2026-05-30_v0_431_calendar_cutoff_readonly_deadline_source"
        bullets = [
            "v0.431: Abschluss-Stichtag in Monats-, Quartals- und Jahresabschluss wird in den Zeitraumsübersichten nur noch aus der Stichtagspflege angezeigt.",
            "v0.431: Die manuelle Änderung des Abschluss-Stichtags in den Kalender-Zeitraumsübersichten wurde entfernt.",
            "v0.431: Aufgaben mit Fälligkeitsart Abschluss-Stichtag verwenden den zentral gepflegten Stichtag aus der Stichtagspflege als Basis.",
            "v0.431: Bestehende Periodendaten werden beim Öffnen der Kalender erneut gegen die Stichtagspflege normalisiert.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 31 or current_build > 100:
                self.version_state["build"] = 31
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.431",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def ensure_version_432_once(self):
        """v0.432: Zeitraumlogik und Dokumentationszentrale vorbereitet.
        VERSION_PREFIX ist "0.4"; daher entspricht v0.432 dem Build 32.
        """
        update_id = "2026-05-30_v0_432_period_logic_documentation_center"
        bullets = [
            "v0.432: Abschlusszeiträume starten ab 05/2026; Jahresabschluss beginnt mit GJ 2025/2026.",
            "v0.432: Monats- und Quartalszeiträume werden in der Zukunft nur bis Ende des freigegebenen Geschäftsjahres angelegt.",
            "v0.432: Folge-Geschäftsjahr wird automatisch ab dem Abschluss-Stichtag August des aktuellen Geschäftsjahres freigegeben.",
            "v0.432: Stichtagspflege verwendet das aktuelle Geschäftsjahr als Standard und arbeitet mit Geschäftsjahresperioden.",
            "v0.432: Anzeigeformate TT.MM.JJJJ, MM/JJ und QQ/JJ sowie Auswahl nur angelegter Zeiträume ergänzt.",
            "v0.432: Vorhandene ältere Zeiträume werden nicht gelöscht, aber nicht mehr angezeigt oder automatisch angelegt.",
            "v0.432: Dokumentationszentrale zeigt nur Gegenwart/Vergangenheit, erlaubt Positionsauswahl, positionsbezogenes Anhängen und einen zentralen Exportdialog.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 32 or current_build > 100:
                self.version_state["build"] = 32
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {"version": "v0.432", "date": now_date_str(), "update_id": update_id, "bullets": bullets})
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

        try:
            self.start_update_watcher()
        except Exception:
            pass

    def load_version_state(self):
        os.makedirs(USER_DIR, exist_ok=True)
        path = os.path.join(USER_DIR, VERSION_STATE_FILE)
        default = {"build": DEFAULT_BUILD}
        try:
            if not os.path.exists(path):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(default, f, ensure_ascii=False, indent=2)
                return default
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            data.setdefault("build", default["build"])
            return data
        except Exception:
            return default

    def save_version_state(self):
        try:
            with open(os.path.join(USER_DIR, VERSION_STATE_FILE), "w", encoding="utf-8") as f:
                json.dump(self.version_state, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def load_version_history(self):
        os.makedirs(USER_DIR, exist_ok=True)
        path = os.path.join(USER_DIR, VERSION_HISTORY_FILE)
        local_history = {"entries": []}
        try:
            if not os.path.exists(path):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(local_history, f, ensure_ascii=False, indent=2)
            else:
                with open(path, "r", encoding="utf-8") as f:
                    local_history = json.load(f)
        except Exception:
            local_history = {"entries": []}
        try:
            central_path = os.path.join(CENTRAL_RELEASE_DIR, VERSION_HISTORY_FILE)
            if os.path.exists(central_path):
                with open(central_path, "r", encoding="utf-8") as f:
                    central_history = json.load(f)
                merged = []
                seen = set()
                for entry in (central_history.get("entries") or central_history.get("versions") or []):
                    key = (str(entry.get("version", "")), str(entry.get("published_at", "")), str(entry.get("update_id", "")))
                    if key not in seen:
                        merged.append(entry); seen.add(key)
                for entry in local_history.get("entries", []):
                    key = (str(entry.get("version", "")), str(entry.get("published_at", "")), str(entry.get("update_id", "")))
                    if key not in seen:
                        merged.append(entry); seen.add(key)
                return {"entries": merged}
        except Exception:
            pass
        return local_history

    def save_version_history(self, history):
        try:
            with open(os.path.join(USER_DIR, VERSION_HISTORY_FILE), "w", encoding="utf-8") as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def bump_version(self, bullets):
        if not isinstance(bullets, list):
            bullets = [str(bullets)]
        self.version_state["build"] = int(self.version_state.get("build", DEFAULT_BUILD)) + 1
        self.save_version_state()
        history = self.load_version_history()
        history.setdefault("entries", [])
        history["entries"].insert(0, {"version": f"v{VERSION_PREFIX}{self.version_state['build']}", "date": now_date_str(), "bullets": bullets})
        self.save_version_history(history)
        try:
            self.version_label.config(text=self.version_label_text())
        except Exception:
            pass

    def bump_version_once(self, update_id, bullets):
        """
        Schreibt einen Versionsverlauf-Eintrag nur einmal.
        Damit wird die Version nicht bei jedem Programmstart erneut erhöht.
        """
        try:
            applied = self.version_state.setdefault("applied_updates", [])
            if update_id in applied:
                return
            self.bump_version(bullets)
            applied.append(update_id)
            self.version_state["applied_updates"] = applied
            self.save_version_state()
        except Exception:
            pass

    def _safe_json_load_file(self, path, default=None):
        try:
            if path and os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass
        return {} if default is None else default

    def _normalize_version_for_compare(self, version):
        raw = str(version or "").strip()
        if raw.lower().startswith("v"):
            raw = raw[1:]
        return raw or "0.0.0"

    def _version_tuple_for_compare(self, version):
        raw = self._normalize_version_for_compare(version)
        parts = []
        for token in raw.replace("-", ".").split("."):
            digits = "".join(ch for ch in token if ch.isdigit())
            try:
                parts.append(int(digits or "0"))
            except Exception:
                parts.append(0)
        return tuple(parts or [0])

    def current_app_version(self):
        try:
            build = int(self.version_state.get("build", DEFAULT_BUILD))
        except Exception:
            build = DEFAULT_BUILD
        return f"{VERSION_PREFIX}.{build}"

    def current_app_version_display(self):
        try:
            build = int(self.version_state.get("build", DEFAULT_BUILD))
            return f"v{VERSION_PREFIX}{build}"
        except Exception:
            return f"v{self.current_app_version()}"

    def _latest_update_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, LATEST_UPDATE_FILE)

    def _update_event_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, UPDATE_EVENT_FILE)

    def _central_version_history_path(self):
        return os.path.join(CENTRAL_RELEASE_DIR, VERSION_HISTORY_FILE)

    def load_latest_update_info(self):
        latest = self._safe_json_load_file(self._latest_update_path(), {})
        event = self._safe_json_load_file(self._update_event_path(), {})
        if event and event.get("version") and not latest.get("version"):
            latest.update(event)
        return latest

    def _patch_text_from_update_info(self, info):
        patch_text = info.get("patch_text") or info.get("bullets") or []
        if isinstance(patch_text, str):
            lines = [line.strip(" -") for line in patch_text.splitlines() if line.strip()]
        elif isinstance(patch_text, list):
            lines = [str(line).strip(" -") for line in patch_text if str(line).strip()]
        else:
            lines = []
        return lines

    def _format_update_message(self, info, start_check=False):
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        update_type = str(info.get("update_type") or "Patch").strip()
        mandatory = bool(info.get("mandatory"))
        published_at = str(info.get("published_at") or "").strip()
        patch_file = str(info.get("patch_file") or "").strip()
        lines = self._patch_text_from_update_info(info)
        bullet_text = "\n".join(f"• {line}" for line in lines[:12]) if lines else "• Kein Patchtext hinterlegt."
        title_line = "Beim Start wurde eine veraltete FiBu-Mate-Version erkannt." if start_check else "Während der Nutzung wurde ein neuer FiBu-Mate-Patch veröffentlicht."
        mandatory_text = "Ja" if mandatory else "Nein"
        return (
            f"{title_line}\n\n"
            f"Installierte Version: {self.current_app_version_display()}\n"
            f"Verfügbare Version: v{latest_version.lstrip('v')}\n"
            f"Patch-Art: {update_type}\n"
            f"Pflichtupdate: {mandatory_text}\n"
            f"Veröffentlicht: {published_at or '-'}\n"
            f"Patchdatei: {patch_file or '-'}\n\n"
            f"Änderungen:\n{bullet_text}\n\n"
            f"Der vollständige Versionsverlauf wird aus folgender Datei gelesen:\n"
            f"{self._central_version_history_path()}"
        )

    def is_newer_update_available(self, info):
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        if not latest_version:
            return False
        try:
            return self._version_tuple_for_compare(latest_version) > self._version_tuple_for_compare(self.current_app_version())
        except Exception:
            return False

    def show_update_notification(self, info, start_check=False):
        if self._update_popup_open:
            return
        latest_version = str(info.get("version") or info.get("latest_version") or "").strip()
        published_at = str(info.get("published_at") or "").strip()
        key = f"{latest_version}|{published_at}|{'start' if start_check else 'live'}"
        if key in self._notified_update_keys and not start_check:
            return
        self._notified_update_keys.add(key)
        self._update_popup_open = True
        try:
            title = "FiBu Mate - veraltete Version" if start_check else "FiBu Mate - neuer Patch verfügbar"
            messagebox.showinfo(title, self._format_update_message(info, start_check=start_check))
        finally:
            self._update_popup_open = False

    def check_update_on_start(self):
        try:
            info = self.load_latest_update_info()
            if self.is_newer_update_available(info):
                self.show_update_notification(info, start_check=True)
        except Exception:
            pass

    def check_update_live(self):
        try:
            info = self.load_latest_update_info()
            event = self._safe_json_load_file(self._update_event_path(), {})
            event_key = "|".join([str(event.get("version", "")), str(event.get("published_at", "")), str(event.get("patch_file", ""))])
            if event_key and event_key != self._last_update_event_seen:
                self._last_update_event_seen = event_key
                if self.is_newer_update_available(info):
                    self.show_update_notification(info, start_check=False)
            elif self.is_newer_update_available(info):
                self.show_update_notification(info, start_check=False)
        except Exception:
            pass
        try:
            self.root.after(UPDATE_CHECK_INTERVAL_MS, self.check_update_live)
        except Exception:
            pass

    def start_update_watcher(self):
        if getattr(self, "_update_watcher_started", False):
            return
        self._update_watcher_started = True
        try:
            self.root.after(2500, self.check_update_on_start)
            self.root.after(UPDATE_CHECK_INTERVAL_MS, self.check_update_live)
        except Exception:
            pass

    def get_user_data_path(self):
        global USER_DATA_PATH
        USER_DATA_PATH = resolve_user_data_path()
        return USER_DATA_PATH

    def _get_user_data_mtime(self):
        try:
            path = self.get_user_data_path()
            return os.path.getmtime(path) if os.path.exists(path) else 0
        except Exception:
            return 0

    def load_user_data(self):
        default = {"last_username_prefill": "", "users": {}, "settings": {"auto_close_mail_enabled": True}}
        try:
            path = self.get_user_data_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            if not os.path.exists(path):
                if os.path.exists(LEGACY_USER_DATA_PATH) and os.path.abspath(LEGACY_USER_DATA_PATH) != os.path.abspath(path):
                    with open(LEGACY_USER_DATA_PATH, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                else:
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(default, f, ensure_ascii=False, indent=2)
                    return default
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            data.setdefault("last_username_prefill", "")
            data.setdefault("users", {})
            data.setdefault("settings", {})
            data["settings"].setdefault("auto_close_mail_enabled", True)
            return data
        except Exception:
            return default

    def save_user_data(self):
        try:
            path = self.get_user_data_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.user_data, f, ensure_ascii=False, indent=2)
            self._user_data_mtime = self._get_user_data_mtime()
        except Exception as e:
            messagebox.showerror("FiBu Mate", "Benutzerdaten konnten nicht gespeichert werden:\n\n" + str(e))

    def ensure_permissions_defaults(self):
        users = self.user_data.setdefault("users", {})
        for key, u in users.items():
            u.setdefault("display_name", key)
            u.setdefault("favorites", [])
            u.setdefault("email", "")
            u.setdefault("auth", {"password_hash": None, "enabled": False})
            u["favorites"] = [fav for fav in u.get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS]
            u.setdefault("first_name", "")
            u.setdefault("full_name", " ".join(x for x in [u.get("first_name", "").strip(), u.get("display_name", key).strip()] if x).strip() or u.get("display_name", key))
            current_permission = ROLE_MIGRATION.get(u.get("permission"), u.get("permission", ROLE_E1))
            if key == SUPERUSER_KEY:
                u["permission"] = ROLE_E4
            else:
                if current_permission == ROLE_E4:
                    u["permission"] = ROLE_E1
                else:
                    u["permission"] = current_permission if current_permission in ROLE_ORDER else ROLE_E1

    def my_role(self):
        if not self.current_user_key:
            return ROLE_E1
        if self.current_user_key == SUPERUSER_KEY:
            return ROLE_E4
        return ROLE_MIGRATION.get(self.user_data.get("users", {}).get(self.current_user_key, {}).get("permission", ROLE_E1), ROLE_E1)

    def role_rank(self, role):
        return ROLE_RANK.get(ROLE_MIGRATION.get(role, role), ROLE_RANK[ROLE_E1])

    def can_view_user_management(self):
        return True

    def can_create_users(self):
        return self.my_role() in (ROLE_E3, ROLE_E4)

    def can_manage_permissions(self):
        return self.my_role() in (ROLE_E3, ROLE_E4)

    def auto_close_mail_enabled(self):
        return bool(self.user_data.setdefault("settings", {}).setdefault("auto_close_mail_enabled", True))

    def toggle_auto_close_mail(self):
        settings = self.user_data.setdefault("settings", {})
        settings["auto_close_mail_enabled"] = not bool(settings.get("auto_close_mail_enabled", True))
        self.save_user_data()
        self.render_page()

    def log_audit_event(self, event_type="Info", module="FiBu Mate", title="", details="", risk="Info", period="", related_id="", public=True):
        try:
            try:
                from bin.tools import compliance_common as cc
            except Exception:
                import compliance_common as cc
            return cc.log_audit(self, event_type, module, title, details, risk, period, related_id, public)
        except Exception:
            return None

    def current_close_period_label(self, module_id):
        now = datetime.now()
        if module_id == "monthly_close":
            names = ["Januar", "Februar", "März", "April", "Mai", "Juni", "Juli", "August", "September", "Oktober", "November", "Dezember"]
            return f"{names[now.month - 1]} {now.year}"
        if module_id == "quarterly_close":
            q = (now.month - 1) // 3 + 1
            return f"Q{q} {now.year}"
        if module_id == "yearly_close":
            start_year = now.year if now.month >= 10 else now.year - 1
            return f"GJ {start_year}/{start_year + 1}"
        return ""

    def max_assignable_role_rank(self):
        return self.role_rank(self.my_role())

    def current_tile_color(self):
        if self.current_user_key:
            return self.user_data.get("users", {}).get(self.current_user_key, {}).get("tile_color")
        return None

    def create_footer(self):
        self.footer = tk.Frame(self.root, bg="black", height=ui_s(49)); self.footer._zoom_exclude = True
        self.footer.pack(side="bottom", fill="x")
        self.footer.pack_propagate(False)
        self.user_label = tk.Label(self.footer, bg="black", fg="white", font=body_font(9))
        self.user_label.place(relx=0, rely=0.5, anchor="w", x=12)
        self.version_label = tk.Label(self.footer, text=self.version_label_text(), bg="black", fg="white", font=body_font(9))
        self.version_label.place(relx=0.5, rely=0.5, anchor="center")
        self.clock_label = tk.Label(self.footer, bg="black", fg="white", font=body_font(9))
        self.clock_label.place(relx=1, rely=0.5, anchor="e", x=-12)
        self.update_clock()

    def update_clock(self):
        self.clock_label.config(text=datetime.now().strftime("%H:%M:%S"))
        self.user_label.config(text=f"Benutzer {self.current_user_display}" if self.current_user_display else "")
        self.root.after(1000, self.update_clock)


    def register_unsaved_changes_provider(self, has_unsaved_callback=None, save_callback=None, discard_callback=None):
        self._unsaved_provider = has_unsaved_callback
        self._unsaved_save_callback = save_callback
        self._unsaved_discard_callback = discard_callback

    def clear_unsaved_changes_provider(self):
        self._unsaved_provider = None
        self._unsaved_save_callback = None
        self._unsaved_discard_callback = None

    def _has_unsaved_changes(self):
        try:
            return bool(self._unsaved_provider and self._unsaved_provider())
        except Exception:
            return False

    def confirm_unsaved_changes(self):
        if not self._has_unsaved_changes():
            return True
        dlg = tk.Toplevel(self.root)
        dlg.title("Ungespeicherte Änderungen")
        dlg.configure(bg=BG)
        dlg.transient(self.root)
        dlg.grab_set()
        result = {"value": None}
        tk.Label(dlg, text="Es liegen ungespeicherte Änderungen vor.", bg=BG, fg=TEXT, font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=18, pady=(16, 6))
        tk.Label(dlg, text="Möchten Sie die Änderungen speichern und übernehmen oder verwerfen?", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).pack(anchor="w", padx=18, pady=(0, 14))
        btns = tk.Frame(dlg, bg=BG)
        btns.pack(fill="x", padx=18, pady=(0, 16))
        def choose(v):
            result["value"] = v
            dlg.destroy()
        tk.Button(btns, text="Speichern + Übernehmen", command=lambda: choose("save"), bg=BLUE, fg="white", bd=0, padx=12, pady=7).pack(side="left", padx=(0, 8))
        discard_btn = tk.Button(btns, text="Änderungen Verwerfen", command=lambda: choose("discard"), bg=WHITE, fg=RED, bd=1, padx=12, pady=7)
        try:
            if PIL_AVAILABLE:
                icon_path = os.path.join(ICON_DIR, "biggarbagebin_121980.ico")
                if os.path.exists(icon_path):
                    img = Image.open(icon_path).resize((18, 18))
                    photo = ImageTk.PhotoImage(img)
                    discard_btn.configure(image=photo, compound="left")
                    discard_btn.image = photo
        except Exception:
            pass
        discard_btn.pack(side="left", padx=(0, 8))
        tk.Button(btns, text="Abbrechen", command=lambda: choose("cancel"), bg=WHITE, fg=TEXT, bd=1, padx=12, pady=7).pack(side="right")
        self.root.wait_window(dlg)
        if result["value"] == "save":
            try:
                if self._unsaved_save_callback:
                    self._unsaved_save_callback()
                self.clear_unsaved_changes_provider()
                return True
            except Exception as exc:
                messagebox.showerror("FiBu Mate", "Speichern fehlgeschlagen: " + str(exc))
                return False
        if result["value"] == "discard":
            try:
                if self._unsaved_discard_callback:
                    self._unsaved_discard_callback()
            finally:
                self.clear_unsaved_changes_provider()
            return True
        return False

    def show_page(self, page_name, title="", add_to_history=True):
        if page_name != self.current_page and not self.confirm_unsaved_changes():
            return
        if add_to_history and self.current_page:
            self.page_history.append((self.current_page, self.current_title))
        self.current_page = page_name
        self.current_title = title
        self.update_breadcrumb(page_name, title)
        self.render_page()

    def update_breadcrumb(self, page_name, title):
        if page_name == "launch":
            self.breadcrumb = []
        elif page_name == "main":
            self.breadcrumb = [("main", "Hauptmenü")]
        else:
            if not self.breadcrumb:
                self.breadcrumb = [("main", "Hauptmenü")]
            existing = next((i for i, (p, _) in enumerate(self.breadcrumb) if p == page_name), None)
            self.breadcrumb = self.breadcrumb[: existing + 1] if existing is not None else self.breadcrumb + [(page_name, title)]

    def jump_to_breadcrumb(self, index):
        if not self.confirm_unsaved_changes():
            return
        if 0 <= index < len(self.breadcrumb):
            self.current_page, self.current_title = self.breadcrumb[index]
            self.breadcrumb = self.breadcrumb[: index + 1]
            self.page_history = self.breadcrumb[:-1].copy()
            self.render_page()

    def go_back(self):
        if not self.confirm_unsaved_changes():
            return
        if self.page_history:
            page, title = self.page_history.pop()
            self.current_page, self.current_title = page, title
            self.update_breadcrumb(page, title)
            self.render_page()

    def display_zoom_key(self):
        try:
            return f"{self.root.winfo_screenwidth()}x{self.root.winfo_screenheight()}"
        except Exception:
            return "unknown"

    def zoom_user_key(self):
        return self.current_user_key or "__anonymous__"

    def current_zoom_scope_key(self):
        try:
            if str(self.current_page or "").startswith("tool:"):
                return self.current_page
            return self.current_page or "main"
        except Exception:
            return "main"

    def load_zoom_profiles(self):
        path = os.path.join(USER_DIR, ZOOM_PROFILE_FILE)
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else {}
        except Exception:
            pass
        return {}

    def save_zoom_profiles(self):
        try:
            os.makedirs(USER_DIR, exist_ok=True)
            with open(os.path.join(USER_DIR, ZOOM_PROFILE_FILE), "w", encoding="utf-8") as f:
                json.dump(self.zoom_profiles, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def get_scope_zoom(self, scope=None):
        try:
            scope = scope or self.current_zoom_scope_key()
            user = self.zoom_user_key(); display = self.display_zoom_key()
            return float(self.zoom_profiles.get(user, {}).get(display, {}).get(scope, 1.0))
        except Exception:
            return 1.0

    def set_scope_zoom(self, zoom, scope=None):
        try:
            scope = scope or self.current_zoom_scope_key()
            user = self.zoom_user_key(); display = self.display_zoom_key()
            self.zoom_profiles.setdefault(user, {}).setdefault(display, {})[scope] = round(float(zoom), 3)
            self.current_scope_zoom = round(float(zoom), 3)
            self.save_zoom_profiles()
        except Exception:
            pass

    def prepare_scope_zoom(self):
        global GLOBAL_TEXT_ZOOM
        try:
            self.current_scope_zoom = self.get_scope_zoom()
            GLOBAL_TEXT_ZOOM = 1.0
        except Exception:
            self.current_scope_zoom = 1.0

    def is_header_footer_or_fixed_widget(self, widget):
        try:
            w = widget
            while w is not None:
                if w is getattr(self, "footer", None):
                    return True
                if getattr(w, "_zoom_exclude", False):
                    return True
                w = w.master
        except Exception:
            pass
        try:
            # Kopf-/Mini-Widgets/Breadcrumb nur im Hauptfenster über y-Koordinate ausschließen.
            if widget.winfo_toplevel() is self.root:
                y = widget.winfo_rooty() - self.canvas.winfo_rooty()
                if y < 126:
                    return True
        except Exception:
            pass
        return False

    def _event_allows_content_zoom(self, event):
        try:
            widget = getattr(event, "widget", None)
            y_root = getattr(event, "y_root", None)
            # Hauptcanvas: nicht pauschal als Header behandeln; nur die echte Kopfzone y < 126 sperren.
            if widget is getattr(self, "canvas", None):
                if y_root is not None:
                    try:
                        y = y_root - self.canvas.winfo_rooty()
                        return y >= 126
                    except Exception:
                        return True
                return True
            if widget is not None and self.is_header_footer_or_fixed_widget(widget):
                return False
            # Nur Ereignisse im Hauptfenster im Headerbereich blocken; Toplevel/Popup-Inhalte bleiben zoombar.
            if widget is not None and y_root is not None:
                try:
                    if widget.winfo_toplevel() is self.root:
                        y = y_root - self.canvas.winfo_rooty()
                        if y < 126:
                            return False
                except Exception:
                    pass
        except Exception:
            pass
        return True

    def _is_ctrl_mousewheel(self, event):
        try:
            # Windows/Tk: ControlMask ist 0x0004. Bei Control-spezifischen Bindings kann state je nach Umgebung variieren;
            # num 4/5 bzw. delta != 0 mit Control-Binding wird deshalb zusätzlich akzeptiert.
            if getattr(event, "state", 0) & 0x0004:
                return True
        except Exception:
            pass
        try:
            if getattr(event, "num", None) in (4, 5):
                return True
            if int(getattr(event, "delta", 0) or 0) != 0 and str(getattr(event, "type", "")) in ("MouseWheel", "38"):
                return bool(getattr(event, "state", 0) & 0x0004)
        except Exception:
            pass
        return False

    def _mousewheel_direction(self, event):
        try:
            if getattr(event, "num", None) == 4:
                return 1
            if getattr(event, "num", None) == 5:
                return -1
        except Exception:
            pass
        try:
            delta = int(getattr(event, "delta", 0))
            return 1 if delta > 0 else (-1 if delta < 0 else 0)
        except Exception:
            return 0

    def on_global_mousewheel(self, event):
        if self._is_ctrl_mousewheel(event):
            if not self.is_zoom_control_visible():
                return "break"
            if not self._event_allows_content_zoom(event):
                return "break"
            direction = self._mousewheel_direction(event)
            if direction:
                self.adjust_global_text_zoom(direction)
            return "break"
        return self._route_mousewheel_to_canvas(event)

    def adjust_global_text_zoom(self, direction):
        try:
            direction = 1 if direction > 0 else -1
            current = self.get_scope_zoom()
            new_zoom = current + direction * GLOBAL_TEXT_ZOOM_STEP
            new_zoom = max(GLOBAL_TEXT_ZOOM_MIN, min(GLOBAL_TEXT_ZOOM_MAX, round(new_zoom, 3)))
            if abs(new_zoom - current) < 0.0001:
                return
            self.set_scope_zoom(new_zoom)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass

    def refresh_zoomed_content(self):
        """Erzwingt die sichtbare Anwendung des Bereichszooms durch Neu-Rendern des aktuellen Inhalts."""
        try:
            if getattr(self, "current_page", "") == "launch":
                return
            self.root.after_idle(self.render_page)
        except Exception:
            try:
                self.render_page()
            except Exception:
                pass

    def reset_global_text_zoom(self, event=None):
        try:
            self.set_scope_zoom(1.0)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass
        return "break"

    def _font_actual_tuple(self, widget, font_spec):
        try:
            f = tkfont.Font(root=widget, font=font_spec)
            actual = f.actual()
            family = actual.get('family') or 'Segoe UI'
            size = abs(int(actual.get('size') or 10))
            weight = actual.get('weight') or 'normal'
            slant = actual.get('slant') or 'roman'
            underline = bool(actual.get('underline'))
            overstrike = bool(actual.get('overstrike'))
            return (family, size, weight, slant, underline, overstrike)
        except Exception:
            return ('Segoe UI', 10, 'normal', 'roman', False, False)

    def _scaled_font_from_base(self, base):
        try:
            family, size, weight, slant, underline, overstrike = base
            scaled_size = max(6, int(round(float(size) * float(self.current_scope_zoom))))
            args = [family, scaled_size]
            if weight and weight != 'normal':
                args.append(weight)
            if slant and slant != 'roman':
                args.append(slant)
            if underline:
                args.append('underline')
            if overstrike:
                args.append('overstrike')
            return tuple(args)
        except Exception:
            return ('Segoe UI', 10)

    def _apply_zoom_to_widget_tree(self, widget):
        try:
            # Canvas zuerst behandeln: Der Hauptcanvas startet bei y=0, enthält aber auch den zoombaren Inhaltsbereich.
            # Header-Texte werden in _apply_zoom_to_canvas weiterhin anhand der Item-BBox ausgeschlossen.
            if isinstance(widget, tk.Canvas):
                if not getattr(widget, "_zoom_exclude", False):
                    self._apply_zoom_to_canvas(widget)
                # Font-Konfiguration des Canvas selbst ist nicht relevant; Kinder trotzdem prüfen.
                for child in widget.winfo_children():
                    self._apply_zoom_to_widget_tree(child)
                return
            if self.is_header_footer_or_fixed_widget(widget):
                return
            try:
                font_spec = widget.cget('font')
                key = str(widget)
                if key not in self._zoom_base_fonts:
                    self._zoom_base_fonts[key] = self._font_actual_tuple(widget, font_spec)
                widget.configure(font=self._scaled_font_from_base(self._zoom_base_fonts[key]))
            except Exception:
                pass
            for child in widget.winfo_children():
                self._apply_zoom_to_widget_tree(child)
        except Exception:
            pass

    def _apply_zoom_to_canvas(self, canvas):
        try:
            ckey = str(canvas)
            base_map = self._zoom_base_canvas_fonts.setdefault(ckey, {})
            for item in canvas.find_all():
                try:
                    if canvas.type(item) != 'text':
                        continue
                    try:
                        bbox = canvas.bbox(item)
                        if bbox and bbox[1] < 126:
                            continue
                    except Exception:
                        pass
                    if item not in base_map:
                        base_map[item] = self._font_actual_tuple(canvas, canvas.itemcget(item, 'font'))
                    canvas.itemconfigure(item, font=self._scaled_font_from_base(base_map[item]))
                except Exception:
                    pass
        except Exception:
            pass

    def apply_global_text_zoom(self):
        try:
            self.prepare_scope_zoom()
            self._apply_zoom_to_widget_tree(self.root)
            for win in list(self.root.winfo_children()):
                try:
                    if isinstance(win, tk.Toplevel) and win.winfo_exists():
                        self._apply_zoom_to_widget_tree(win)
                except Exception:
                    pass
        except Exception:
            pass
    def clear_content(self):
        self.clear_unsaved_changes_provider()
        self.active_scroll_canvas = None
        if hasattr(self, "module_escape_handler"):
            delattr(self, "module_escape_handler")
        if hasattr(self, "_update_scroll_indicators"):
            delattr(self, "_update_scroll_indicators")
        self.canvas.delete("all")
        for w in self.widget_items:
            try:
                w.destroy()
            except Exception:
                pass
        self.widget_items.clear()
        self.focusable_tiles.clear()
        self.focus_index = -1
        self.image_refs.clear()

    def render_page(self):
        self.init_responsive_scaling()
        self.clear_content()
        self.draw_background()
        if self.current_page == "launch":
            self.render_launch()
            return
        self.draw_header(self.current_title)
        self.draw_controls()
        self.draw_path_bar()
        self.draw_favorites_bar()
        if self.current_page == "main": self.render_main_menu()
        elif self.current_page == "data_prep": self.render_data_prep_menu()
        elif self.current_page == "debitoren_tools": self.render_debitoren_tools_menu()
        elif self.current_page == "nike_tools": self.render_nike_tools_menu()
        elif self.current_page == "afi_uploads": self.render_afi_uploads_menu()
        elif self.current_page == "closing_calendar": self.render_closing_calendar_menu()
        elif self.current_page == "knowledge_base": self.render_knowledge_base()
        elif self.current_page == "compliance_audit": self.render_compliance_audit_menu()
        elif self.current_page == "in_dev": self.render_in_dev_menu()
        elif self.current_page == "settings": self.render_settings_menu()
        elif self.current_page == "tile_colors": self.render_tile_colors_menu()
        elif self.current_page == "users": self.render_users_menu()
        elif self.current_page == "permissions": self.render_permissions_menu()
        elif self.current_page == "information": self.render_information_menu()
        elif self.current_page == "versions": self.render_versions_menu()
        elif self.current_page.startswith("tool:"): self.render_external_tool(self.current_page.replace("tool:", "", 1))
        else: self.render_menu_text("Menü in Arbeit.")
        if self.focusable_tiles:
            self.focus_index = 0
            self.focusable_tiles[0].focus_set()

    def on_resize(self, *_):
        self.root.after(30, self.render_page)

    def draw_background(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        self.canvas.create_rectangle(0, 0, w, h, fill=BG, outline="")
        # Dezentes dynamisches Hintergrunddesign: weiche Linien und Flächen in FiBu-Mate-/INTERSPORT-Farben
        for i in range(10):
            y = 145 + i * max(52, h * 0.060)
            color = blend(BG, BLUE if i % 2 == 0 else RED, 0.040 + (i % 4) * 0.008)
            self.canvas.create_line(-80, y, w + 90, y + h * 0.075, fill=color, width=1)
        self.canvas.create_oval(w * 0.68, h * 0.18, w * 1.05, h * 0.78, outline=blend(BG, BLUE, 0.11), width=2)
        self.canvas.create_oval(w * 0.73, h * 0.24, w * 1.01, h * 0.70, outline=blend(BG, WHITE, 0.46), width=1)
        self.canvas.create_oval(-w * 0.14, h * 0.42, w * 0.25, h * 1.02, outline=blend(BG, RED, 0.075), width=2)
        self.canvas.create_polygon(w * 0.02, h * 0.90, w * 0.32, h * 0.72, w * 0.60, h * 1.04, fill=blend(BG, WHITE, 0.20), outline="")
        self.canvas.create_polygon(w * 0.62, h * 0.20, w * 0.96, h * 0.36, w * 1.06, h * 0.12, fill=blend(BG, BLUE, 0.040), outline="")
        self.canvas.create_polygon(w * 0.08, h * 0.18, w * 0.22, h * 0.26, w * 0.04, h * 0.34, fill=blend(BG, WHITE, 0.15), outline="")
        self.canvas.create_rectangle(0, 0, w, 109, fill=HEADER, outline="")
        self.canvas.create_rectangle(0, 109, w, 112, fill=LINE, outline="")

    def draw_header(self, title):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        header_h = ui_s(92)
        self.canvas.create_rectangle(0, 0, w, header_h, fill=HEADER, outline="")
        self.canvas.create_line(0, header_h, w, header_h, fill=LINE, width=1)
        self.canvas.create_text(w / 2, ui_s(70), text=title, font=FONT_MENU, fill=TEXT2, anchor="center")

    def draw_gradient_line(self, x1, x2, y):
        width = max(1, int(x2 - x1)); steps = max(20, min(180, width // 5))
        for i in range(steps):
            t = i / max(1, steps - 1)
            color = blend(HEADER, RED, t / 0.14) if t < 0.14 else blend(BLUE, HEADER, (t - 0.86) / 0.14) if t > 0.86 else blend(RED, BLUE, (t - 0.14) / 0.72)
            self.canvas.create_line(x1 + width * i / steps, y, x1 + width * (i + 1) / steps, y, fill=color, width=2)

    def draw_path_bar(self):
        w = self.canvas.winfo_width(); x1 = ui_s(12); x2 = min(w * 0.42, ui_s(760)); y_mid = ui_s(38)
        self.draw_gradient_line(x1, x2, y_mid - ui_s(10)); self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
        x = x1 + ui_s(96)
        for idx, (page, title) in enumerate(self.breadcrumb):
            current = idx == len(self.breadcrumb) - 1
            tid = self.canvas.create_text(x, y_mid, text=title, font=body_font(9, "bold") if current else body_font(9), fill=TEXT if current else BLUE, anchor="w")
            bbox = self.canvas.bbox(tid); tw = bbox[2] - bbox[0] if bbox else ui_s(70)
            if not current:
                self.canvas.tag_bind(tid, "<Button-1>", lambda e, i=idx: self.jump_to_breadcrumb(i))
                self.canvas.tag_bind(tid, "<Enter>", lambda e: self.canvas.config(cursor="hand2"))
                self.canvas.tag_bind(tid, "<Leave>", lambda e: self.canvas.config(cursor="arrow"))
            x += tw + ui_s(14)
            if idx < len(self.breadcrumb) - 1:
                self.canvas.create_polygon(x, y_mid - ui_s(6), x, y_mid + ui_s(6), x + ui_s(7), y_mid, fill=RED, outline=RED); x += ui_s(20)

    def draw_favorites_bar(self):
        w = self.canvas.winfo_width(); x1, x2 = max(w * 0.58, w - ui_s(720)), w - ui_s(360)
        if x2 <= x1 + ui_s(120): return
        y_mid = ui_s(38)
        self.draw_gradient_line(x1, x2, y_mid - ui_s(10)); self.draw_gradient_line(x1, x2, y_mid + ui_s(10))
        self.canvas.create_text(x1 + ui_s(16), y_mid, text="★", font=("Segoe UI", max(10, ui_s(13)), "bold"), fill=GOLD, anchor="w")
        self.canvas.create_text(x1 + ui_s(40), y_mid, text="Favoriten", font=body_font(9, "bold"), fill=TEXT, anchor="w")
        x = x1 + ui_s(118); chip_color = self.current_tile_color() or BLUE; max_x = x2 - ui_s(8)
        for fav in sorted(f for f in self.favorites if f in TOOL_REGISTRY and f not in HIDDEN_TOOL_IDS):
            if x + ui_s(96) > max_x: break
            label = TOOL_REGISTRY.get(fav, {}).get("favorite_label", fav)
            chip = tk.Label(self.root, text=label, bg=chip_color, fg="white", font=body_font(8), padx=ui_s(6), pady=ui_s(1), cursor="hand2"); chip._zoom_exclude = True
            chip.bind("<Button-1>", lambda e, fid=fav: self.execute_favorite(fid))
            self.widget_items.append(chip); self.canvas.create_window(x, y_mid + ui_s(8), window=chip, anchor="w"); x += ui_s(108)

    def draw_controls(self):
        self.draw_logout_control() if self.current_page == "main" else self.draw_back_control()
        self.draw_suggestion_control()
        self.draw_help_control()
        self.draw_close_control()

    def draw_back_control(self):
        frame = tk.Frame(self.root, bg=HEADER, cursor="hand2")
        arrow = tk.Label(frame, text="←", fg=RED, bg=HEADER, font=body_font(12, "bold"), cursor="hand2")
        txt = tk.Label(frame, text="zurück", fg=BLUE, bg=HEADER, font=body_font(9, underline=True), cursor="hand2")
        arrow.pack(side="left"); txt.pack(side="left", padx=ui_s(4))
        arrow.bind("<Button-1>", lambda e: self.go_back()); txt.bind("<Button-1>", lambda e: self.go_back())
        self.widget_items.append(frame); self.canvas.create_window(ui_s(14), ui_s(8), window=frame, anchor="nw")

    def draw_logout_control(self):
        btn = tk.Button(self.root, text="Abmelden", command=self.logout, bg=self.current_tile_color() or BLUE, fg="white", font=body_font(10, "bold", scale=0.75), bd=0, cursor="hand2", padx=ui_s(9), pady=ui_s(4))
        self.widget_items.append(btn); self.canvas.create_window(ui_s(14), ui_s(8), window=btn, anchor="nw")

    def draw_mini_bulb_icon(self, canvas, x, y):
        if self.draw_canvas_icon(canvas, "idea", x, y, 18, 18):
            return
        canvas.create_oval(x - 6, y - 8, x + 6, y + 4, fill="#FFE45C", outline="#FACC15", width=1)
        canvas.create_rectangle(x - 4, y + 4, x + 4, y + 9, outline=BLUE, width=1)

    def draw_mini_help_icon(self, canvas, x, y):
        if self.draw_canvas_icon(canvas, "help", x, y, 19, 19):
            return
        canvas.create_rectangle(x - 17, y - 8, x + 17, y + 8, outline=BLUE, width=1)
        canvas.create_text(x, y, text="HELP", fill=BLUE, font=("Segoe UI", 7, "bold"))

    def is_zoom_control_visible(self):
        """v0.434 technische Korrektur: Zoomleiste nur in Modulen und Einstellungen-Untermenüs anzeigen."""
        try:
            page = str(getattr(self, "current_page", "") or "")
            if page.startswith("tool:"):
                return True
            return page in {"tile_colors", "users", "permissions", "information", "versions"}
        except Exception:
            return False

    def draw_zoom_control(self):
        """v0.434 Paket 1E: sichtbare Zoom-Leiste als Alternative zu Strg+Mausrad."""
        if not self.is_zoom_control_visible():
            return
        try:
            w = ui_s(190)
            h = MINI_WIDGET_H
            btn = tk.Canvas(self.root, width=w, height=h, bg=HEADER, highlightthickness=0, bd=0, cursor="arrow"); btn._zoom_exclude = True
            btn.create_rectangle(1, 1, w - 1, h - 1, fill=HEADER, outline=BLUE, width=1)
            minus = btn.create_text(ui_s(18), h / 2, text="−", fill=BLUE, font=("Segoe UI", 13, "bold"), anchor="center")
            label = btn.create_text(w / 2, h / 2, text=self.zoom_percent_text(), fill=TEXT, font=("Segoe UI", 9, "bold"), anchor="center")
            plus = btn.create_text(w - ui_s(18), h / 2, text="+", fill=BLUE, font=("Segoe UI", 13, "bold"), anchor="center")
            btn.create_line(ui_s(36), 4, ui_s(36), h - 4, fill=LINE)
            btn.create_line(w - ui_s(36), 4, w - ui_s(36), h - 4, fill=LINE)
            btn._zoom_label_item = label
            btn._zoom_minus_item = minus
            btn._zoom_plus_item = plus
            btn.tag_bind(minus, "<Button-1>", lambda _e: self.zoom_bar_step(-1))
            btn.tag_bind(plus, "<Button-1>", lambda _e: self.zoom_bar_step(1))
            btn.tag_bind(label, "<Button-1>", lambda _e: self.zoom_bar_reset())
            btn.bind("<Button-1>", lambda e: self.zoom_bar_click(e, btn, w))
            btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Zoom: − / Prozent zurücksetzen / +"))
            btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
            self._zoom_control = btn
            self.widget_items.append(btn)
            self.canvas.create_window(ui_s(168), 7, window=btn, anchor="nw")
        except Exception:
            pass

    def zoom_percent_text(self):
        try:
            return f"{int(round(float(self.get_scope_zoom()) * 100))}%"
        except Exception:
            return "100%"

    def update_zoom_control_label(self):
        try:
            btn = getattr(self, "_zoom_control", None)
            if btn and btn.winfo_exists():
                btn.itemconfigure(getattr(btn, "_zoom_label_item", None), text=self.zoom_percent_text())
        except Exception:
            pass

    def zoom_bar_step(self, direction):
        try:
            self.adjust_global_text_zoom(direction)
            self.update_zoom_control_label()
        except Exception:
            pass
        return "break"

    def zoom_bar_reset(self):
        try:
            self.set_scope_zoom(1.0)
            self.apply_global_text_zoom()
            self.update_zoom_control_label()
            self.refresh_zoomed_content()
        except Exception:
            pass
        return "break"

    def zoom_bar_click(self, event, canvas, width):
        try:
            x = getattr(event, "x", 0)
            if x <= ui_s(38):
                return self.zoom_bar_step(-1)
            if x >= width - ui_s(38):
                return self.zoom_bar_step(1)
            return self.zoom_bar_reset()
        except Exception:
            return "break"

    def draw_suggestion_control(self):
        if self.current_page == "launch":
            return
        # v0.434 Paket 1D: Mini-Widget bleibt zoom-excluded; Text wird bewusst mit fixer Mini-Schrift gezeichnet.
        # Dadurch kann „Änderung vorschlagen“ nicht mehr durch UI_TEXT_SCALE aus dem Button laufen.
        suggestion_w = max(MINI_WIDGET_W, ui_s(202))
        btn = tk.Canvas(self.root, width=suggestion_w, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor="hand2"); btn._zoom_exclude = True
        btn.create_rectangle(1, 1, suggestion_w - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
        icon_x = ui_s(16)
        text_x = ui_s(36)
        self.draw_mini_bulb_icon(btn, icon_x, MINI_WIDGET_H / 2)
        btn.create_text(text_x, MINI_WIDGET_H / 2, text="Änderung vorschlagen", fill=TEXT, font=("Segoe UI", max(8, ui_s(9)), "bold"), anchor="w")
        btn.bind("<Button-1>", lambda _e: self.open_suggestion_mail())
        btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Änderung vorschlagen"))
        btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
        self.widget_items.append(btn)
        # rechter Rand: links vom Hilfe-Mini-Widget; anchor=ne erwartet die rechte Kante, daher nicht nochmals um die eigene Breite verschieben.
        x = self.canvas.winfo_width() - ui_s(176) - MINI_WIDGET_GAP
        self.canvas.create_window(x, ui_s(8), window=btn, anchor="ne")

    def show_small_tooltip(self, widget, text):
        self.hide_small_tooltip()
        try:
            self._small_tooltip = tk.Toplevel(widget)
            self._small_tooltip.wm_overrideredirect(True)
            self._small_tooltip.geometry(f"+{widget.winfo_rootx() - 90}+{widget.winfo_rooty() + 34}")
            tk.Label(self._small_tooltip, text=text, bg="#111827", fg="white", font=("Segoe UI", 9), padx=7, pady=4).pack()
        except Exception:
            self._small_tooltip = None

    def hide_small_tooltip(self):
        tip = getattr(self, "_small_tooltip", None)
        if tip:
            try:
                tip.destroy()
            except Exception:
                pass
        self._small_tooltip = None

    def open_suggestion_mail(self):
        messagebox.showinfo(
            "Änderung vorschlagen",
            "Vielen Dank für deinen Input! Bitte beschreibe in der folgenden Vorlage deinen Vorschlag so ausführlich wie möglich - Füge gerne Screenshots mit an",
        )
        user_name = self.current_user_display or self.current_user_key or ""
        subject = "Änderungsvorschlag Fibu Mate"
        body = (
            f"Änderungsvorschlag von {user_name},\n\n"
            "Folgenden Änderungsvorschlag würde ich gerne mitteilen / Folgende Anpassung wünsche ich mir:\n\n"
            "[Text des Vorschlagenden]"
        )
        mailto = (
            "mailto:matthias.wagner@intersport.de"
            "?cc=" + quote("matze.wagner1@yahoo.de")
            + "&subject=" + quote(subject)
            + "&body=" + quote(body)
        )
        try:
            webbrowser.open(mailto)
        except Exception as exc:
            messagebox.showerror("Änderung vorschlagen", f"Outlook konnte nicht geöffnet werden:\n\n{exc}")

    def draw_help_control(self):
        if self.current_page == "launch":
            return
        btn = tk.Canvas(self.root, width=MINI_WIDGET_W, height=MINI_WIDGET_H, bg=HEADER, highlightthickness=0, bd=0, cursor="hand2"); btn._zoom_exclude = True
        btn.create_rectangle(1, 1, MINI_WIDGET_W - 1, MINI_WIDGET_H - 1, fill=HEADER, outline=BLUE, width=1)
        self.draw_mini_help_icon(btn, 25, MINI_WIDGET_H / 2)
        btn.create_text(MINI_WIDGET_W / 2 + ui_s(6), MINI_WIDGET_H / 2, text="Hilfe", fill=TEXT, font=("Segoe UI", max(8, ui_s(9)), "bold"), anchor="center")
        def open_help(_event=None):
            self.show_help_popup()
        btn.bind("<Button-1>", open_help)
        btn.bind("<Enter>", lambda _e: self.show_small_tooltip(btn, "Hilfe"))
        btn.bind("<Leave>", lambda _e: self.hide_small_tooltip())
        self.widget_items.append(btn)
        self.canvas.create_window(self.canvas.winfo_width() - ui_s(22), ui_s(8), window=btn, anchor="ne")

    def draw_close_control(self):
        """v0.433 Korrektur Paket 1c/1e: Separater X-Button in den Mini-Widgets deaktiviert.
        Das Schließen erfolgt ausschließlich über das native Fenster-X bzw. die bestehende App-Logik.
        """
        return

    def show_help_popup(self):
        popup = tk.Toplevel(self.root)
        popup.title("FiBu Mate - Hilfe")
        popup.configure(bg=BG)
        popup.transient(self.root)
        popup.resizable(True, True)
        popup.minsize(780, 520)
        popup_w, popup_h = 1000, 700
        try:
            self.root.update_idletasks()
            screen_w = self.root.winfo_screenwidth()
            screen_h = self.root.winfo_screenheight()
            x = max(0, int((screen_w - popup_w) / 2))
            y = max(0, int((screen_h - popup_h) / 2))
            popup.geometry(f"{popup_w}x{popup_h}+{x}+{y}")
        except Exception:
            popup.geometry(f"{popup_w}x{popup_h}")
        self._make_modal(popup)
        header = tk.Frame(popup, bg=HEADER, height=64); header.pack(side="top", fill="x"); header.pack_propagate(False)
        tk.Label(header, text="Hilfe", bg=HEADER, fg=TEXT, font=("Segoe UI", 18, "bold")).pack(side="left", padx=18)
        content = tk.Frame(popup, bg=BG); content.pack(side="top", fill="both", expand=True, padx=16, pady=16)
        if self.help_image and PIL_AVAILABLE:
            try:
                popup.update_idletasks(); img = resize_keep_ratio(self.help_image, max(600, popup.winfo_width() - 70), max(400, popup.winfo_height() - 150)); ph = ImageTk.PhotoImage(img)
                label = tk.Label(content, image=ph, bg=BG); label.image = ph; label.pack(expand=True)
            except Exception as error:
                tk.Label(content, text=f"Das Hilfe-Bild konnte nicht geladen werden:\n\n{error}", bg=BG, fg=TEXT, font=body_font(11), justify="left").pack(anchor="nw")
        else:
            tk.Label(content, text=("Das Hilfe-Menü ist vorbereitet, aber das Hilfe-Bild konnte nicht geladen werden.\n\n" f"Erwarteter Pfad:\n{HELP_IMAGE_PATH}"), bg=BG, fg=TEXT, font=body_font(11), justify="left").pack(anchor="nw")
        footer = tk.Frame(popup, bg=BG); footer.pack(side="bottom", fill="x", padx=16, pady=(0, 16))
        tk.Button(footer, text="Schließen", command=popup.destroy, bg=BLUE, fg="white", font=body_font(10, "bold"), bd=0, padx=ui_s(24), pady=ui_s(10), cursor="hand2").pack(side="right")

    def draw_bottom_logo(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if self.intersport_logo and PIL_AVAILABLE:
            ph = ImageTk.PhotoImage(resize_keep_ratio(self.intersport_logo, 420, 55)); self.image_refs.append(ph); self.canvas.create_image(w / 2, h - 24, image=ph)
        else:
            self.canvas.create_text(w / 2, h - 24, text="INTERSPORT", font=("Segoe UI", 22, "bold"), fill=BLUE)

    def draw_intersport_logo_above_footer(self, show_mini_logo=True): return self.draw_bottom_logo()

    def toggle_favorite(self, tool_id):
        if tool_id not in TOOL_REGISTRY: return
        self.favorites.remove(tool_id) if tool_id in self.favorites else self.favorites.add(tool_id)
        if self.current_user_key:
            self.user_data["users"][self.current_user_key]["favorites"] = sorted(self.favorites); self.save_user_data()
        self.render_page()

    def execute_favorite(self, tool_id):
        if tool_id in TOOL_REGISTRY: self.open_tool(tool_id)

    def render_launch(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if self.banner_big and PIL_AVAILABLE:
            ph = ImageTk.PhotoImage(resize_keep_ratio(self.banner_big, 900, 230)); self.image_refs.append(ph); self.canvas.create_image(w / 2, h * 0.235, image=ph)
        else:
            self.canvas.create_text(w / 2, h * 0.235, text="FiBu Mate", font=("Segoe UI", 46, "bold"), fill=TEXT)
        panel = tk.Frame(self.root, bg=BG); username_var = tk.StringVar(value="")
        tk.Label(panel, text="Benutzername", bg=BG, fg=TEXT, font=("Segoe UI", 11, "bold")).pack(anchor="w", pady=(0, 6))
        row = tk.Frame(panel, bg=BG); row.pack(fill="x", pady=(0, 18))
        entry = tk.Entry(row, textvariable=username_var, font=body_font(13), width=37, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1); entry.pack(side="left", ipady=ui_s(7))
        btn = tk.Button(row, text="Anmelden", command=lambda: self.login_user(username_var.get()), bg=BLUE, fg="white", font=body_font(11, "bold", scale=0.75), bd=0, cursor="hand2"); btn.pack(side="left", padx=(12, 0), ipady=6)
        self.widget_items.append(panel); self.canvas.create_window(x_pct(w, 50), y_pct(h, 39), window=panel, anchor="center")
        entry.focus_set(); entry.bind("<Return>", lambda e: self.login_user_from_entry(username_var.get()))
        self.draw_bottom_logo(); self.draw_close_control()



    # ------------------------------------------------------------------
    # Wissenszentrale - Startseite als Overlay + Arbeitsansicht A
    # ------------------------------------------------------------------
    def knowledge_pref_path(self):
        try:
            return os.path.join(APPDATA_DIR, "knowledge_start_pref.json")
        except Exception:
            return "knowledge_start_pref.json"

    def kb_entries_path(self):
        try:
            os.makedirs(APPDATA_DIR, exist_ok=True)
            return os.path.join(APPDATA_DIR, "knowledge_entries.json")
        except Exception:
            return "knowledge_entries.json"

    def load_knowledge_start_preference(self):
        try:
            path = self.knowledge_pref_path()
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return bool(data.get("show_start", True))
        except Exception:
            pass
        return True

    def save_knowledge_start_preference(self, value):
        self.knowledge_show_start = bool(value)
        try:
            path = self.knowledge_pref_path()
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"show_start": self.knowledge_show_start}, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def kb_load_entries(self):
        try:
            path = self.kb_entries_path()
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                entries = data.get("entries", []) if isinstance(data, dict) else []
                return entries if isinstance(entries, list) else []
        except Exception:
            pass
        return []

    def kb_save_entries(self, entries):
        try:
            path = self.kb_entries_path()
            dir_name = os.path.dirname(path)
            if dir_name:
                os.makedirs(dir_name, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"entries": entries}, f, ensure_ascii=False, indent=2)
            return True
        except Exception as exc:
            try:
                messagebox.showerror("Wissenszentrale", f"Einträge konnten nicht gespeichert werden:\n{exc}")
            except Exception:
                pass
            return False

    def kb_default_categories(self):
        return ["To-Do"]

    def kb_get_categories(self):
        cats = set(self.kb_default_categories())
        for entry in self.kb_load_entries():
            for cat in entry.get("categories", []) or []:
                cat = str(cat).strip()
                if cat:
                    cats.add(cat)
        return sorted(cats, key=lambda x: x.lower())

    def kb_now(self):
        try:
            return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            return ""

    def kb_display_date(self, value):
        try:
            dt = datetime.strptime(str(value), "%Y-%m-%d %H:%M:%S")
            return dt.strftime("%d.%m.%Y %H:%M")
        except Exception:
            return str(value or "")

    def kb_make_entry_id(self):
        try:
            return datetime.now().strftime("KB%Y%m%d%H%M%S%f")
        except Exception:
            return "KB_ENTRY"

    def open_knowledge_base(self):
        self.knowledge_view = "all"
        self.knowledge_start_overlay = bool(getattr(self, "knowledge_show_start", True))
        self.show_page("knowledge_base", "Wissenszentrale", True)

    def draw_kb_button(self, x, y, text, command=None, accent=False, width=None):
        btn = tk.Button(self.root, text=text, command=command, font=body_font(10, weight="bold" if accent else None),
                        bg=("#CFEAD6" if accent else WHITE), fg=TEXT, activebackground="#DDEAF7",
                        relief="solid", bd=1, cursor="hand2", padx=8, pady=3)
        btn_w = width or max(110, 18 + len(text) * 8)
        self.canvas.create_window(ui_s(x), ui_s(y), window=btn, anchor="nw", width=ui_s(btn_w), height=ui_s(34))
        self.widget_items.append(btn)
        return x + btn_w + 8

    def draw_kb_badge(self, x, y, text, color):
        w = max(54, 20 + len(str(text)) * 8)
        self.canvas.create_rectangle(ui_s(x), ui_s(y), ui_s(x+w), ui_s(y+24), fill=color, outline=color, width=1)
        self.canvas.create_text(ui_s(x+10), ui_s(y+12), text=text, font=body_font(9, weight="bold"), fill=WHITE, anchor="w")
        return x + w + 8

    def kb_has_unsaved_changes(self):
        return bool(getattr(self, "knowledge_unsaved", False))

    def kb_mark_unsaved(self):
        self.knowledge_unsaved = True

    def kb_save_current_work(self):
        # Speichern wird in der Wissenszentrale durch kb_save_entry_from_form übernommen.
        self.knowledge_unsaved = False
        return True

    def kb_confirm_unsaved_before_switch(self):
        if not self.kb_has_unsaved_changes():
            return True
        try:
            result = messagebox.askyesnocancel(
                "Ungespeicherte Änderungen",
                "In der aktuellen Arbeitsansicht liegen ungespeicherte Änderungen vor.\n\n"
                "Ja = speichern und wechseln\nNein = verwerfen und wechseln\nAbbrechen = in der aktuellen Ansicht bleiben"
            )
        except Exception:
            result = False
        if result is True:
            return self.kb_save_current_work()
        if result is False:
            self.knowledge_unsaved = False
            return True
        return False

    def kb_switch_view_from_start(self, view):
        if not self.kb_confirm_unsaved_before_switch():
            return
        self.knowledge_view = view
        if view == "new":
            self.kb_prepare_new_entry()
        self.knowledge_start_overlay = False
        self.render_page()

    def kb_show_start_overlay(self):
        self.knowledge_start_overlay = True
        self.render_page()

    def kb_close_start_overlay(self):
        self.knowledge_start_overlay = False
        self.render_page()

    def render_knowledge_base(self):
        underlay_start = len(self.widget_items)
        self.render_knowledge_work_area()
        if bool(getattr(self, "knowledge_start_overlay", False)):
            for widget in list(self.widget_items[underlay_start:]):
                try:
                    widget.destroy()
                except Exception:
                    pass
            del self.widget_items[underlay_start:]
            self.render_knowledge_start_overlay()

    def kb_ensure_state_vars(self):
        if not hasattr(self, "kb_search_var"):
            self.kb_search_var = tk.StringVar(value="")
        if not hasattr(self, "kb_filter_vars") or len(getattr(self, "kb_filter_vars", [])) != 4:
            self.kb_filter_vars = [tk.StringVar(value="") for _ in range(4)]
        if not hasattr(self, "kb_todo_rhythm_var"):
            self.kb_todo_rhythm_var = tk.StringVar(value="")
        if not hasattr(self, "kb_selected_entry_id"):
            self.kb_selected_entry_id = None
        if not hasattr(self, "kb_edit_entry_id"):
            self.kb_edit_entry_id = None

    def kb_prepare_new_entry(self, entry=None):
        self.kb_ensure_state_vars()
        self.kb_edit_entry_id = entry.get("id") if entry else None
        self.kb_title_var = tk.StringVar(value=(entry or {}).get("title", ""))
        cats = list((entry or {}).get("categories", []) or [])[:4]
        self.kb_entry_category_vars = [tk.StringVar(value=(cats[i] if i < len(cats) else "")) for i in range(4)]
        self.kb_user_var = tk.StringVar(value=(entry or {}).get("user", self.current_user_display or self.current_user_key or ""))
        self.kb_status_var = tk.StringVar(value=(entry or {}).get("status", "Aktiv"))
        self.kb_rhythm_var = tk.StringVar(value=(entry or {}).get("rhythm", ""))
        self.kb_text_initial = (entry or {}).get("text", "")
        self.knowledge_unsaved = bool(entry)

    def kb_filtered_entries(self):
        self.kb_ensure_state_vars()
        entries = self.kb_load_entries()
        search = (self.kb_search_var.get() or "").strip().lower()
        filters = [(v.get() or "").strip().lower() for v in self.kb_filter_vars]
        filters = [f for f in filters if f]
        rhythm = (self.kb_todo_rhythm_var.get() or "").strip().lower()
        result = []
        for entry in entries:
            cats = [str(c).strip().lower() for c in (entry.get("categories", []) or [])]
            if getattr(self, "knowledge_view", "all") == "todos" and "to-do" not in cats:
                continue
            if getattr(self, "knowledge_view", "all") == "todos" and rhythm:
                if str(entry.get("rhythm", "")).strip().lower() != rhythm:
                    continue
            if filters and not all(f in cats for f in filters):
                continue
            if search:
                hay = " ".join([
                    str(entry.get("title", "")), str(entry.get("text", "")),
                    str(entry.get("user", "")), " ".join(entry.get("categories", []) or [])
                ]).lower()
                if search not in hay:
                    continue
            result.append(entry)
        return sorted(result, key=lambda e: e.get("updated_at") or e.get("created_at") or "", reverse=True)

    def kb_apply_filters(self, event=None):
        self.kb_selected_entry_id = None
        if getattr(self, "knowledge_view", "all") not in ("new", "categories", "outdated"):
            self.knowledge_view = "all" if getattr(self, "knowledge_view", "all") != "todos" else "todos"
        self.render_page()

    def kb_on_search_return(self, event=None):
        self.kb_apply_filters()

    def kb_select_entry(self, entry_id):
        self.kb_selected_entry_id = entry_id
        self.knowledge_view = "detail"
        self.knowledge_unsaved = False
        self.render_page()

    def kb_get_entry(self, entry_id):
        for entry in self.kb_load_entries():
            if entry.get("id") == entry_id:
                return entry
        return None

    def kb_edit_selected_entry(self):
        entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
        if not entry:
            return
        self.kb_prepare_new_entry(entry)
        self.knowledge_view = "new"
        self.render_page()

    def kb_save_entry_from_form(self):
        self.kb_ensure_state_vars()
        title = (getattr(self, "kb_title_var", tk.StringVar(value="")).get() or "").strip()
        if not title:
            try:
                messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
            except Exception:
                pass
            return
        categories = []
        for var in getattr(self, "kb_entry_category_vars", []):
            value = (var.get() or "").strip()
            if value and value not in categories:
                categories.append(value)
        categories = categories[:4]
        user = (getattr(self, "kb_user_var", tk.StringVar(value="")).get() or "").strip()
        status = (getattr(self, "kb_status_var", tk.StringVar(value="Aktiv")).get() or "Aktiv").strip()
        rhythm = (getattr(self, "kb_rhythm_var", tk.StringVar(value="")).get() or "").strip()
        text_value = ""
        try:
            text_value = self.kb_text_widget.get("1.0", "end-1c")
        except Exception:
            text_value = getattr(self, "kb_text_initial", "")
        entries = self.kb_load_entries()
        now = self.kb_now()
        edit_id = getattr(self, "kb_edit_entry_id", None)
        if edit_id:
            for entry in entries:
                if entry.get("id") == edit_id:
                    entry.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "updated_at": now})
                    break
            selected_id = edit_id
        else:
            selected_id = self.kb_make_entry_id()
            entries.append({
                "id": selected_id, "title": title, "categories": categories, "user": user,
                "status": status, "rhythm": rhythm, "text": text_value,
                "created_at": now, "updated_at": now, "comments": [], "attachments": []
            })
        if self.kb_save_entries(entries):
            self.kb_selected_entry_id = selected_id
            self.kb_edit_entry_id = None
            self.knowledge_unsaved = False
            self.knowledge_view = "detail"
            self.render_page()

    def kb_export_selected_to_word(self):
        # Platzhalter: Der echte Word-Export wird später mit python-docx an die finale Datenstruktur angebunden.
        try:
            messagebox.showinfo("Wissenszentrale", "Word-Export wird in einem Folgeblock an die finale Datenstruktur angebunden.")
        except Exception:
            pass

    def render_knowledge_work_area(self):
        self.kb_ensure_state_vars()
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        # Alle Elemente sind bewusst in vertikal getrennten Zonen aufgebaut:
        # Navigation 145, Suche 192, Filter 245, Hauptfenster ab 365.
        nav_y = 145
        x = 28
        x = self.draw_kb_button(x, nav_y, "Start", self.kb_show_start_overlay, False, width=92)
        x = self.draw_kb_button(x, nav_y, "Neuer Eintrag", lambda: self.kb_switch_view_from_start("new"), True, width=128)
        x = self.draw_kb_button(x, nav_y, "To-Dos", lambda: self.kb_switch_view_from_start("todos"), False, width=92)
        x = self.draw_kb_button(x, nav_y, "Veraltete Einträge", lambda: self.kb_switch_view_from_start("outdated"), False, width=148)
        x = self.draw_kb_button(x, nav_y, "Kategorien verwalten", lambda: self.kb_switch_view_from_start("categories"), False, width=168)

        search_frame = tk.Frame(self.root, bg=BG)
        self.widget_items.append(search_frame)
        tk.Label(search_frame, text="Suche", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
        entry = tk.Entry(search_frame, textvariable=self.kb_search_var, font=body_font(10), bg=WHITE, fg=TEXT, relief="solid", bd=1)
        entry.grid(row=1, column=0, sticky="ew", ipady=4)
        entry.bind("<Return>", self.kb_on_search_return)
        btn = tk.Button(search_frame, text="Suchen", command=self.kb_apply_filters, bg=WHITE, fg=TEXT, relief="solid", bd=1, font=body_font(9), cursor="hand2")
        btn.grid(row=1, column=1, padx=(8, 0), sticky="ns")
        search_frame.grid_columnconfigure(0, weight=1)
        self.canvas.create_window(ui_s(28), ui_s(192), window=search_frame, anchor="nw", width=ui_s(min(760, max(520, w-80))), height=ui_s(56))

        filter_frame = tk.Frame(self.root, bg=BG)
        self.widget_items.append(filter_frame)
        if self.knowledge_view == "todos":
            tk.Label(filter_frame, text="To-Do-Filter / Rhythmus", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w", columnspan=2)
            rhythms = ["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"]
            cb = ttk.Combobox(filter_frame, textvariable=self.kb_todo_rhythm_var, values=rhythms, state="readonly", font=body_font(9), width=18)
            cb.grid(row=1, column=0, sticky="w")
            cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
            tk.Label(filter_frame, text="Leer = alle Rhythmen", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=1, column=1, padx=(12,0), sticky="w")
        else:
            tk.Label(filter_frame, text="Kategorie-Filter", bg=BG, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w", columnspan=5)
            values = [""] + self.kb_get_categories()
            for idx, var in enumerate(self.kb_filter_vars):
                block = tk.Frame(filter_frame, bg=BG)
                block.grid(row=1, column=idx, padx=(0 if idx == 0 else 10, 0), sticky="w")
                tk.Label(block, text=f"Kategorie {idx+1}", bg=BG, fg=TEXT2, font=body_font(8)).pack(anchor="w")
                cb = ttk.Combobox(block, textvariable=var, values=values, state="readonly", font=body_font(9), width=18)
                cb.pack(anchor="w")
                cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
            tk.Label(filter_frame, text="Leere Filter = keine Filterung", bg=BG, fg=TEXT2, font=body_font(8)).grid(row=1, column=4, padx=(12,0), sticky="s")
        self.canvas.create_window(ui_s(28), ui_s(252), window=filter_frame, anchor="nw", width=ui_s(min(860, max(620, w-70))), height=ui_s(82))

        # Hauptfenster bewusst tiefer gesetzt, damit Filter, Treffer und Detail niemals kollidieren.
        left_x = 20
        left_y = 365
        left_w = max(340, min(460, int(w * 0.31)))
        pane_h = max(390, h - left_y - 50)
        right_x = left_x + left_w + 24
        right_y = left_y
        right_w = max(760, w - right_x - 24)
        self.render_kb_hits_pane(left_x, left_y, left_w, pane_h)
        if self.knowledge_view == "new":
            self.render_kb_new_entry_area(right_x, right_y, right_w, pane_h)
        elif self.knowledge_view == "todos":
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="To-Dos")
        elif self.knowledge_view == "outdated":
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Veraltete Einträge", status_filter="Veraltet")
        elif self.knowledge_view == "categories":
            self.render_kb_categories_area(right_x, right_y, right_w, pane_h)
        elif self.knowledge_view == "detail":
            self.render_kb_detail_area(right_x, right_y, right_w, pane_h)
        else:
            self.render_kb_list_area(right_x, right_y, right_w, pane_h, title="Gesamtliste aller Einträge")

    def render_kb_hits_pane(self, x, y, w, h):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text="Treffer", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
        entries = self.kb_filtered_entries()
        listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="flat", activestyle="none", exportselection=False)
        listbox.pack(fill="both", expand=True, padx=16, pady=(0, 14))
        id_map = []
        if entries:
            for entry in entries:
                cats = ", ".join(entry.get("categories", []) or [])
                line = f"{self.kb_display_date(entry.get('updated_at'))}  |  {entry.get('title','')}"
                if cats:
                    line += f"  [{cats}]"
                listbox.insert("end", line)
                id_map.append(entry.get("id"))
        else:
            listbox.insert("end", "Noch keine Treffer vorhanden.")
            listbox.insert("end", "Filter und Suche wirken auf die Wissensdatenbank.")
        def _select(event=None):
            sel = listbox.curselection()
            if sel and sel[0] < len(id_map):
                self.kb_select_entry(id_map[sel[0]])
        listbox.bind("<Double-Button-1>", _select)
        listbox.bind("<Return>", _select)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_list_area(self, x, y, w, h, title="Gesamtliste aller Einträge", status_filter=None):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        header = tk.Frame(frame, bg=WHITE)
        header.pack(fill="x", padx=18, pady=(16, 8))
        tk.Label(header, text=title, bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(side="left")
        tk.Label(header, text="neueste Einträge oben", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(side="right")
        columns = ("date", "title", "categories", "user", "status")
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=12)
        tree.heading("date", text="Geändert")
        tree.heading("title", text="Titel")
        tree.heading("categories", text="Kategorien")
        tree.heading("user", text="Benutzer")
        tree.heading("status", text="Status")
        tree.column("date", width=120, stretch=False)
        tree.column("title", width=max(220, int(w*0.33)), stretch=True)
        tree.column("categories", width=max(170, int(w*0.22)), stretch=True)
        tree.column("user", width=130, stretch=False)
        tree.column("status", width=90, stretch=False)
        yscroll = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=yscroll.set)
        tree.pack(side="left", fill="both", expand=True, padx=(18,0), pady=(0,18))
        yscroll.pack(side="right", fill="y", padx=(0,18), pady=(0,18))
        entries = self.kb_filtered_entries()
        if status_filter:
            entries = [e for e in entries if str(e.get("status", "")).lower() == status_filter.lower()]
        id_map = {}
        for entry in entries:
            iid = entry.get("id") or self.kb_make_entry_id()
            id_map[iid] = entry.get("id")
            tree.insert("", "end", iid=iid, values=(self.kb_display_date(entry.get("updated_at")), entry.get("title", ""), ", ".join(entry.get("categories", []) or []), entry.get("user", ""), entry.get("status", "")))
        if not entries:
            tree.insert("", "end", values=("", "Noch keine Einträge vorhanden", "", "", ""))
        def _open(event=None):
            sel = tree.selection()
            if sel and sel[0] in id_map:
                self.kb_select_entry(id_map[sel[0]])
        tree.bind("<Double-Button-1>", _open)
        tree.bind("<Return>", _open)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_detail_area(self, x, y, w, h):
        entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
        if not entry:
            self.render_kb_list_area(x, y, w, h, title="Gesamtliste aller Einträge")
            return
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text=entry.get("title", ""), bg=WHITE, fg=BLUE, font=body_font(16, weight="bold")).pack(anchor="w", padx=18, pady=(16, 6))
        meta = f"Geändert: {self.kb_display_date(entry.get('updated_at'))}    Benutzer: {entry.get('user','')}    Status: {entry.get('status','')}"
        tk.Label(frame, text=meta, bg=WHITE, fg=TEXT2, font=body_font(10)).pack(anchor="w", padx=18)
        cats = ", ".join(entry.get("categories", []) or []) or "Keine Kategorien"
        tk.Label(frame, text=f"Kategorien: {cats}", bg=WHITE, fg=TEXT, font=body_font(10, weight="bold")).pack(anchor="w", padx=18, pady=(8, 4))
        txt = tk.Text(frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word")
        txt.pack(fill="both", expand=True, padx=18, pady=(6, 12))
        txt.insert("1.0", entry.get("text", ""))
        txt.configure(state="disabled")
        btnrow = tk.Frame(frame, bg=WHITE)
        btnrow.pack(fill="x", padx=18, pady=(0, 16))
        tk.Button(btnrow, text="Bearbeiten", command=self.kb_edit_selected_entry, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Word-Export", command=self.kb_export_selected_to_word, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Zur Liste", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_new_entry_area(self, x, y, w, h):
        if not hasattr(self, "kb_title_var"):
            self.kb_prepare_new_entry()
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text=("Eintrag bearbeiten" if getattr(self, "kb_edit_entry_id", None) else "Neuer Eintrag"), bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 6))
        form = tk.Frame(frame, bg=WHITE)
        form.pack(fill="x", padx=18, pady=(4, 8))
        tk.Label(form, text="Titel des Eintrags", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=0, column=0, sticky="w")
        title_entry = tk.Entry(form, textvariable=self.kb_title_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1)
        title_entry.grid(row=1, column=0, columnspan=4, sticky="ew", ipady=5, pady=(0, 8))
        categories = [""] + self.kb_get_categories()
        for idx, var in enumerate(self.kb_entry_category_vars):
            tk.Label(form, text=f"Kategorie {idx+1}", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=2, column=idx, sticky="w", padx=(0 if idx==0 else 8,0))
            cb = ttk.Combobox(form, textvariable=var, values=categories, state="normal", font=body_font(9), width=18)
            cb.grid(row=3, column=idx, sticky="ew", padx=(0 if idx==0 else 8,0), pady=(0, 8))
        tk.Label(form, text="Assoziierter Benutzer", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=0, sticky="w")
        tk.Entry(form, textvariable=self.kb_user_var, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).grid(row=5, column=0, sticky="ew", pady=(0, 8))
        tk.Label(form, text="Status", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=1, sticky="w", padx=(8,0))
        ttk.Combobox(form, textvariable=self.kb_status_var, values=["Aktiv", "Entwurf", "Veraltet"], state="readonly", font=body_font(9)).grid(row=5, column=1, sticky="ew", padx=(8,0), pady=(0, 8))
        tk.Label(form, text="To-Do-Rhythmus", bg=WHITE, fg=TEXT2, font=body_font(9)).grid(row=4, column=2, sticky="w", padx=(8,0))
        ttk.Combobox(form, textvariable=self.kb_rhythm_var, values=["", "täglich", "wöchentlich", "monatlich", "quartalsweise", "jährlich", "bei Bedarf"], state="readonly", font=body_font(9)).grid(row=5, column=2, sticky="ew", padx=(8,0), pady=(0, 8))
        for col in range(4):
            form.grid_columnconfigure(col, weight=1)
        body_frame = tk.Frame(frame, bg=WHITE)
        body_frame.pack(fill="both", expand=True, padx=18, pady=(0, 10))
        tk.Label(body_frame, text="Freitext / Prozessdokumentation / Leitfaden", bg=WHITE, fg=TEXT2, font=body_font(9)).pack(anchor="w")
        self.kb_text_widget = tk.Text(body_frame, bg="#F8FAFC", fg=TEXT, font=body_font(10), relief="solid", bd=1, wrap="word")
        self.kb_text_widget.pack(fill="both", expand=True)
        self.kb_text_widget.insert("1.0", getattr(self, "kb_text_initial", ""))
        self.kb_text_widget.bind("<KeyRelease>", lambda e: self.kb_mark_unsaved())
        btnrow = tk.Frame(frame, bg=WHITE)
        btnrow.pack(fill="x", padx=18, pady=(0, 16))
        tk.Button(btnrow, text="Als Entwurf speichern", command=lambda: (self.kb_status_var.set("Entwurf"), self.kb_save_entry_from_form()), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Speichern", command=self.kb_save_entry_from_form, bg="#CFEAD6", fg=TEXT, font=body_font(10, weight="bold"), relief="solid", bd=1).pack(side="left", padx=(0, 8), ipadx=16, ipady=4)
        tk.Button(btnrow, text="Abbrechen", command=lambda: self.kb_switch_view_from_start("all"), bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1).pack(side="left", ipadx=16, ipady=4)
        title_entry.focus_set()
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def render_kb_categories_area(self, x, y, w, h):
        frame = tk.Frame(self.root, bg=WHITE, highlightbackground=LINE, highlightthickness=2)
        self.widget_items.append(frame)
        tk.Label(frame, text="Kategorien verwalten", bg=WHITE, fg=BLUE, font=body_font(15, weight="bold")).pack(anchor="w", padx=18, pady=(16, 8))
        tk.Label(frame, text="Kategorien entstehen aktuell automatisch aus den Einträgen. Neue Kategorien können direkt beim Erfassen in den Kategorie-Feldern eingetragen werden.", bg=WHITE, fg=TEXT, font=body_font(10), wraplength=max(500, int(w*0.85)), justify="left").pack(anchor="w", padx=18, pady=(0, 12))
        listbox = tk.Listbox(frame, bg=WHITE, fg=TEXT, font=body_font(10), relief="solid", bd=1)
        listbox.pack(fill="both", expand=True, padx=18, pady=(0, 18))
        for cat in self.kb_get_categories():
            listbox.insert("end", cat)
        self.canvas.create_window(ui_s(x), ui_s(y), window=frame, anchor="nw", width=ui_s(w), height=ui_s(h))

    def kb_overlay_start_drag(self, event):
        self.knowledge_overlay_drag = None

    def kb_overlay_drag(self, event):
        return

    def kb_overlay_end_drag(self, event=None):
        self.knowledge_overlay_drag = None

    def render_knowledge_start_overlay(self):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        y0 = ui_s(98)
        y1 = h - ui_s(34)
        self.canvas.create_rectangle(0, y0, w, y1, fill=blend(BG, WHITE, 0.86), outline=LINE, width=2)
        self.canvas.create_rectangle(0, y0, w, ui_s(158), fill=blend(HEADER, WHITE, 0.34), outline=LINE, width=1)
        self.canvas.create_text(ui_s(45), ui_s(123), text="Wissenszentrale – Start", fill=BLUE, font=body_font(22, weight="bold"), anchor="w")
        self.canvas.create_text(ui_s(45), ui_s(148), text="Startseite maximiert – Arbeitsbereich vollständig abgedeckt", fill=TEXT2, font=body_font(9), anchor="w")
        self.draw_kb_button(w - ui_s(180), 112, "Schließen", self.kb_close_start_overlay, False, width=130)
        self.canvas.create_text(ui_s(45), ui_s(195), text="Was möchtest du tun?", fill=TEXT, font=body_font(15, weight="bold"), anchor="w")
        var = tk.BooleanVar(value=bool(getattr(self, "knowledge_show_start", True)))
        def _toggle():
            self.save_knowledge_start_preference(var.get())
        chk = tk.Checkbutton(self.root, text="Startseite anzeigen", variable=var, command=_toggle, bg=blend(BG, WHITE, 0.86), fg=TEXT, font=body_font(11), activebackground=blend(BG, WHITE, 0.86))
        self.canvas.create_window(ui_s(45), h - ui_s(88), window=chk, anchor="nw", width=ui_s(260), height=ui_s(34))
        self.widget_items.append(chk)
        cards = [("Wissen suchen", "Gesamtliste und Suche öffnen", "all", "#2563EB"), ("Neuer Eintrag", "Dokumentation oder Aufgabe erfassen", "new", "#059669"), ("To-Dos anzeigen", "wiederkehrende Aufgaben anzeigen", "todos", "#DC2626"), ("Kategorien", "Kategorien anzeigen", "categories", "#7C3AED"), ("Veraltete Einträge", "Prüfung / Archiv", "outdated", "#EC4899")]
        card_w = min(430, max(300, int((max(900, w-ui_s(90)) - 90) / 3)))
        card_h = 150
        start_x, start_y = 45, 235
        gap_x, gap_y = 45, 38
        for idx, (title, desc, view, color) in enumerate(cards):
            col, row = idx % 3, idx // 3
            x = start_x + col * (card_w + gap_x)
            y = start_y + row * (card_h + gap_y)
            self.canvas.create_rectangle(ui_s(x+6), ui_s(y+8), ui_s(x+card_w+6), ui_s(y+card_h+8), fill=SHADOW, outline="")
            self.canvas.create_rectangle(ui_s(x), ui_s(y), ui_s(x+card_w), ui_s(y+card_h), fill=WHITE, outline=LINE)
            self.canvas.create_oval(ui_s(x+24), ui_s(y+26), ui_s(x+74), ui_s(y+76), fill=color, outline=color)
            self.canvas.create_text(ui_s(x+95), ui_s(y+35), text=title, fill=BLUE, font=body_font(15, weight="bold"), anchor="w")
            self.canvas.create_text(ui_s(x+95), ui_s(y+75), text=desc, fill=TEXT2, font=body_font(11), anchor="w")
            self.draw_kb_button(x+95, y+105, "Öffnen", lambda v=view: self.kb_switch_view_from_start(v), view == "new")
        self.canvas.create_text(ui_s(45), h - ui_s(45), text="Die Startseite liegt über allen Arbeitsansicht-Widgets und endet erst oberhalb der Fußleiste.", fill=TEXT2, font=body_font(10), anchor="w")

    def render_main_menu(self):
        # Menüzeile 1: Abschlusskalender
        top_tiles = [
            {"title": "Abschlusskalender", "cmd": lambda: self.show_page("closing_calendar", "Abschlusskalender", True), "fixed": None, "lock": False, "icon": "calendar", "fold": False},
            {"title": "Wissenszentrale", "cmd": lambda: self.open_knowledge_base(), "fixed": None, "lock": False, "icon": "info", "fold": False},
        ]
        # Menüzeile 2: Tools-Menüs
        middle_tiles = [
            {"title": "Tools - Hauptbuch", "cmd": lambda: self.show_page("data_prep", "Tools - Hauptbuch", True), "fixed": None, "lock": False, "icon": "pdf_xls", "fold": False},
            {"title": "Tools - Debitoren", "cmd": lambda: self.show_page("debitoren_tools", "Tools - Debitoren", True), "fixed": None, "lock": False, "icon": "modules", "fold": False},
        ]
        # Mini-Menüzeile: kleinere Kacheln; Eselsohr rechts oben außer bei "In Entwicklung".
        mini_tiles = [
            {"title": "In Entwicklung", "cmd": self.try_open_in_dev, "fixed": GREY_TILE, "lock": True, "icon": "lock", "fold": False},
            {"title": "Informationen", "cmd": lambda: self.show_page("information", "Informationen", True), "fixed": None, "lock": False, "icon": "info", "fold": True},
            {"title": "Einstellungen", "cmd": lambda: self.show_page("settings", "Einstellungen", True), "fixed": None, "lock": False, "icon": "gear", "fold": True},
        ]
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        tw = max(240, min(330, int(w * 0.165)))
        th = max(125, min(170, int(h * 0.155)))
        gap_x = max(38, int(w * 0.04))
        def centered_x_positions(count, tile_w=tw, gap=gap_x):
            if count <= 0:
                return []
            total_width = count * tile_w + (count - 1) * gap
            start_x = (w - total_width) / 2 + tile_w / 2
            return [start_x + i * (tile_w + gap) for i in range(count)]
        def create_tile_group(items, y, id_prefix, tile_w=tw, tile_h=th):
            xs = centered_x_positions(len(items), tile_w)
            for i, item in enumerate(items):
                tile = Tile(self.root, self, f"{id_prefix}_{i}", item["title"], item["cmd"], favorite_enabled=False, fixed_color=item["fixed"], lock_tile=item["lock"], icon_type=item["icon"], corner_fold=item["fold"])
                tile.resize_tile(tile_w, tile_h)
                self.widget_items.append(tile)
                self.focusable_tiles.append(tile)
                self.canvas.create_window(xs[i], y, window=tile, anchor="center")
        top_y = y_pct(h, 64)
        tools_y = y_pct(h, 43)
        create_tile_group(top_tiles, top_y, "main_menuzeile_1")
        create_tile_group(middle_tiles, tools_y, "main_menuzeile_2", tile_w=max(210, int(tw * 0.86)), tile_h=max(105, int(th * 0.86)))
        # Trennlinie zwischen Tools-Menüs und Abschlusskalender – identisch zur Mini-Menüzeilen-Trennlinie.
        self.draw_continuous_relief_line(self.canvas, (top_y + tools_y) / 2, x_pct(w, 9.5), x_pct(w, 92))
        mini_tw = int(tw * 0.72)
        mini_th = int(th * 0.72)
        mini_center_y = y_pct(h, 16.5) + th / 2 - mini_th / 2
        mini_top = mini_center_y - mini_th / 2
        create_tile_group(mini_tiles, mini_center_y, "main_mini_menuezeile", tile_w=mini_tw, tile_h=mini_th)
        self.draw_continuous_relief_line(self.canvas, mini_top - 13, x_pct(w, 9.5), x_pct(w, 92))
        self.draw_bottom_logo()

    def try_open_in_dev(self):
        if self.my_role() != ROLE_E4: return
        self.show_page("in_dev", "In Entwicklung", True)

    def render_data_prep_menu(self):
        modules = [
            ("Nike-Tools", "page:nike_tools"),
            ("AFI-Uploads", "page:afi_uploads"),
            ("Aramark Monatsabrechnungen - PDF zu Excel", "aramark_monatsabrechnungen_pdf_to_excel"),
        ]
        self.render_module_menu(modules, show_descriptions=False)
        self.draw_bottom_logo()


    def render_debitoren_tools_menu(self):
        modules = [("Debitoren-Serienbrief", "debitoren_serienbrief")]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_afi_uploads_menu(self):
        modules = [
            # EnBW ist fachlich in "Lieferanten-Rechnung zu AFI-Upload" integriert.
            # Der alte Toolcode bleibt im TOOL_REGISTRY erhalten, wird aber nicht mehr angezeigt.
            ("Lieferanten-Rechnung zu AFI-Upload", "supplier_invoice_afi_upload"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_nike_tools_menu(self):
        modules = [
            ("Nike - PDF zu Excel", "nike_pdf_to_excel"),
            ("Nike - OP-Liste: Vollständigkeit PDF-Rechnungen prüfen", "nike_op_liste_pdf_check"),
            ("Nike - Rechnungs-PDFs in Sammelordner", "invoice_pdf_collector"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_closing_calendar_menu(self):
        modules = [("Monatsabschluss", "monthly_close"), ("Quartalsabschluss", "quarterly_close"), ("Jahresabschluss", "yearly_close"), ("Stichtagspflege", "deadline_maintenance")]
        self.render_module_menu(modules, show_descriptions=True)
        if self.my_role() == ROLE_E4:
            text = "Auto-Mail: Ein" if self.auto_close_mail_enabled() else "Auto-Mail: Aus"
            btn = tk.Button(self.root, text=text, command=self.toggle_auto_close_mail, bg=BLUE if self.auto_close_mail_enabled() else GREY_DISABLED, fg="white", bd=0, padx=10, pady=3, cursor="hand2", font=("Segoe UI", 9, "bold"))
            self.widget_items.append(btn)
            self.canvas.create_window(self.canvas.winfo_width() / 2, 136, window=btn, anchor="n")
        self.draw_bottom_logo()

    def render_compliance_audit_menu(self):
        modules = [
            ("Steuermeldungs-Cockpit", "tax_reporting"),
            ("Audit-Cockpit", "audit_cockpit"),
            ("Dokumentationszentrale", "documentation_center"),
        ]
        self.render_module_menu(modules, show_descriptions=True)
        self.draw_bottom_logo()

    def render_in_dev_menu(self):
        modules = [("Compliance & Audit", "page:compliance_audit"), ("X001 SAP - Test", "x001_sap_test")]
        self.render_module_menu(modules, show_descriptions=True); self.draw_bottom_logo()


    def render_settings_menu(self):
        items = [("Farbschema", "tile_colors")]
        if self.can_view_user_management():
            items.append(("Benutzerverwaltung", "users"))
        if self.can_manage_permissions():
            items.append(("Berechtigungen", "permissions"))
        self.render_center_menu(items, title="Einstellungen"); self.draw_bottom_logo()

    def render_information_menu(self):
        self.render_center_menu([("Versionsverlauf", "versions")], title="Informationen"); self.draw_bottom_logo()

    def render_center_menu(self, items, title=""):
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); tile_w = int(max(260, min(360, w * 0.2))); tile_h = int(max(110, min(160, h * 0.14))); gap = int(max(18, h * 0.03)); start_y = y_pct(h, 55); start_x = x_pct(w, 50)
        for i, (label, page) in enumerate(items):
            y = start_y + i * (tile_h + gap); cmd = (lambda p=page, l=label: self.show_page(p, l, True))
            tile = Tile(self.root, self, f"center_{page}", label, cmd, favorite_enabled=False, center_text=True); tile.resize_tile(tile_w, tile_h); self.widget_items.append(tile); self.focusable_tiles.append(tile); self.canvas.create_window(start_x, y, window=tile, anchor="center")

    def render_tile_colors_menu(self):
        zfont = self.zoomed_content_font
        frame = tk.Frame(self.root, bg=BG); self.widget_items.append(frame)
        tk.Label(frame, text="Standardfarben", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=5, pady=(0, 14), sticky="w")
        selected_color = self.current_tile_color() or BLUE
        def set_color(col):
            if self.current_user_key:
                self.user_data["users"][self.current_user_key]["tile_color"] = col; self.save_user_data()
            self.render_page()
        for idx, (name, color) in enumerate(COLOR_PALETTE):
            r = 1 + idx // 5; c = idx % 5; sw = tk.Canvas(frame, width=140, height=78, bg=BG, highlightthickness=0, bd=0, cursor="hand2")
            if color == selected_color:
                sw.create_rectangle(4, 4, 136, 58, outline=WHITE, width=4)
            sw.create_rectangle(8, 8, 132, 54, fill=color, outline=LINE, width=2); sw.create_text(70, 66, text=name, fill=TEXT, font=zfont(("Segoe UI", 9, "bold"))); sw.bind("<Button-1>", lambda e, col=color: set_color(col)); sw.grid(row=r, column=c, padx=10, pady=10)
        def reset():
            if self.current_user_key:
                self.user_data["users"][self.current_user_key].pop("tile_color", None); self.save_user_data()
            self.render_page()
        tk.Button(frame, text="Standard wiederherstellen", command=reset, bg=selected_color, fg="white", bd=0, padx=16, pady=10, cursor="hand2").grid(row=r + 1, column=0, columnspan=5, pady=(18, 6))
        self.canvas.create_window(self.canvas.winfo_width() / 2, y_pct(self.canvas.winfo_height(), 48), window=frame, anchor="center")


    def render_users_menu(self):
        zfont = self.zoomed_content_font
        if not self.can_view_user_management():
            self.render_menu_text("Keine Berechtigung für die Benutzerverwaltung."); return
        users = self.user_data.setdefault("users", {})
        visible_keys = [self.current_user_key] if self.my_role() == ROLE_E1 and self.current_user_key in users else sorted(users.keys())
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)
        arrow_col = tk.Frame(container, bg=BG, width=54); arrow_col.pack(side="left", fill="y", padx=(18, 0)); arrow_col.pack_propagate(False)
        body = tk.Frame(container, bg=BG); body.pack(side="left", fill="both", expand=True)
        footer = tk.Frame(body, bg=BG); footer.pack(side="bottom", fill="x", pady=(8, 0))
        scroll_canvas = tk.Canvas(body, bg=BG, highlightthickness=0, bd=0); scrollbar = tk.Scrollbar(body, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG); content_window = scroll_canvas.create_window((0, 0), window=content, anchor="nw")
        scroll_canvas.configure(yscrollcommand=scrollbar.set); scroll_canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y"); self.register_scroll_canvas(scroll_canvas, scrollbar)
        def update_arrows():
            try:
                first, last = scroll_canvas.yview(); up_arrow.set_enabled(first > 0.001); down_arrow.set_enabled(last < 0.999)
            except Exception: pass
        def scroll_units(n): scroll_canvas.yview_scroll(n, "units"); update_arrows()
        up_arrow = ArrowIndicator(arrow_col, "up", lambda: scroll_units(-5), size=42); down_arrow = ArrowIndicator(arrow_col, "down", lambda: scroll_units(5), size=42)
        up_arrow.pack(side="top", pady=(78, 10)); down_arrow.pack(side="top", pady=(0, 10))
        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width())); scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all")); self._sync_scrollbar_visibility(scroll_canvas, scrollbar); update_arrows()
        self._update_scroll_indicators = update_arrows; content.bind("<Configure>", update_scrollregion); scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<MouseWheel>", lambda e: self._route_mousewheel_to_canvas(e, scroll_canvas))
        frame = content
        tk.Label(frame, text="Benutzerverwaltung", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=8, sticky="w", pady=(0, 14))
        for col, header in enumerate(["Benutzer", "Anzeigename/Nachname", "Vorname", "E-Mail", "Rolle", "Passwort"]):
            tk.Label(frame, text=header, bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=col, sticky="w", padx=(0, 14))
        def can_edit_user(target_key):
            if target_key == self.current_user_key: return True
            if target_key == SUPERUSER_KEY: return False
            target_role = users.get(target_key, {}).get("permission", ROLE_E1)
            return self.my_role() in (ROLE_E3, ROLE_E4) or self.role_rank(target_role) < self.role_rank(self.my_role())
        def save_user_row(old_key, name_var, email_var, first_name_var=None):
            if old_key not in users: return
            data = users[old_key]; data["email"] = email_var.get().strip(); data["first_name"] = first_name_var.get().strip() if first_name_var else data.get("first_name", ""); data["full_name"] = " ".join(x for x in [data.get("first_name", "").strip(), data.get("display_name", old_key).strip()] if x).strip() or data.get("display_name", old_key)
            may_rename = ((old_key == SUPERUSER_KEY and self.current_user_key == SUPERUSER_KEY) or (self.my_role() in (ROLE_E3, ROLE_E4) and old_key not in (self.current_user_key, SUPERUSER_KEY) and can_edit_user(old_key)))
            if may_rename:
                new_name = " ".join(name_var.get().strip().split()); new_key = normalize_username(new_name)
                if not new_key: messagebox.showwarning("FiBu Mate", "Bitte einen Benutzernamen eingeben."); return
                if old_key != SUPERUSER_KEY and new_key == SUPERUSER_KEY: messagebox.showwarning("FiBu Mate", "Der Benutzer Wagnerm kann nicht überschrieben werden."); return
                if new_key != old_key and new_key in users: messagebox.showwarning("FiBu Mate", "Dieser Benutzername existiert bereits."); return
                data = users.pop(old_key); data["display_name"] = new_name; data["full_name"] = " ".join(x for x in [data.get("first_name", "").strip(), new_name] if x).strip() or new_name; users[new_key] = data
            self.save_user_data(); messagebox.showinfo("FiBu Mate", "Benutzerdaten wurden gespeichert."); self.render_page()
        def delete_user(user_key):
            if user_key in (SUPERUSER_KEY, self.current_user_key): messagebox.showwarning("FiBu Mate", "Dieser Benutzer kann nicht gelöscht werden."); return
            if not can_edit_user(user_key): messagebox.showwarning("FiBu Mate", "Keine Berechtigung für diesen Benutzer."); return
            if not messagebox.askyesno("Benutzer löschen", f"Benutzer wirklich löschen?\n\n{users.get(user_key, {}).get('display_name', user_key)}"): return
            users.pop(user_key, None); self.save_user_data(); messagebox.showinfo("FiBu Mate", "Benutzer wurde gelöscht."); self.render_page()
        row = 2
        for key in visible_keys:
            user = users[key]; user.setdefault("email", "")
            tk.Label(frame, text=key, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=0, sticky="w", padx=(0, 14), pady=4)
            name_var = tk.StringVar(value=user.get("display_name", key)); first_name_var = tk.StringVar(value=user.get("first_name", "")); email_var = tk.StringVar(value=user.get("email", ""))
            name_state = "normal" if self.my_role() in (ROLE_E3, ROLE_E4) and key not in (self.current_user_key, SUPERUSER_KEY) else "disabled"
            tk.Entry(frame, textvariable=name_var, font=("Segoe UI", 10), width=22, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1, state=name_state).grid(row=row, column=1, sticky="w", padx=(0, 14), pady=4)
            tk.Entry(frame, textvariable=first_name_var, font=("Segoe UI", 10), width=18, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=row, column=2, sticky="w", padx=(0, 14), pady=4)
            tk.Entry(frame, textvariable=email_var, font=("Segoe UI", 10), width=32, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=row, column=3, sticky="w", padx=(0, 14), pady=4)
            tk.Label(frame, text=user.get("permission", ROLE_E1), bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=4, sticky="w", padx=(0, 14), pady=4)
            tk.Label(frame, text="aktiv" if user.get("auth", {}).get("enabled") else "nicht aktiv", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=5, sticky="w", padx=(0, 14), pady=4)
            editable = can_edit_user(key)
            tk.Button(frame, text="Speichern", command=lambda k=key, n=name_var, e=email_var, f=first_name_var: save_user_row(k, n, e, f), bg=BLUE if editable else GREY_DISABLED, fg="white", bd=0, width=9, padx=10, pady=5, cursor="hand2" if editable else "arrow", state="normal" if editable else "disabled").grid(row=row, column=6, sticky="w", padx=(18, 10), pady=4)
            del_ok = editable and key not in (self.current_user_key, SUPERUSER_KEY) and self.my_role() in (ROLE_E3, ROLE_E4)
            tk.Button(frame, text="Löschen", command=lambda k=key: delete_user(k), bg=RED if del_ok else GREY_DISABLED, fg="white", bd=0, width=9, padx=10, pady=5, cursor="hand2" if del_ok else "arrow", state="normal" if del_ok else "disabled").grid(row=row, column=7, sticky="w", padx=(0, 8), pady=4)
            row += 1
        if self.can_create_users():
            tk.Label(footer, text="Neuen Benutzer anlegen", font=("Segoe UI", 12, "bold"), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=6, sticky="w", pady=(0, 8))
            username_var = tk.StringVar(); first_name_new_var = tk.StringVar(); email_new_var = tk.StringVar()
            tk.Label(footer, text="Benutzername", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=0, sticky="w", padx=(0, 8))
            tk.Entry(footer, textvariable=username_var, font=("Segoe UI", 10), width=28, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=1, sticky="w", padx=(0, 14))
            tk.Label(footer, text="Vorname", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=2, sticky="e", padx=(0, 4))
            tk.Entry(footer, textvariable=first_name_new_var, font=("Segoe UI", 10), width=18, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=3, sticky="w", padx=(0, 14))
            tk.Label(footer, text="E-Mail", bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=1, column=4, sticky="e", padx=(0, 4))
            tk.Entry(footer, textvariable=email_new_var, font=("Segoe UI", 10), width=32, bg="#E8EEF5", fg=TEXT, relief="solid", bd=1).grid(row=1, column=5, sticky="w", padx=(0, 14))
            def create_user():
                raw_name = " ".join(username_var.get().strip().split()); key = normalize_username(raw_name)
                if not key: messagebox.showwarning("FiBu Mate", "Bitte einen Benutzernamen eingeben."); return
                if key in users: messagebox.showwarning("FiBu Mate", "Dieser Benutzer existiert bereits."); return
                users[key] = {"display_name": raw_name, "first_name": first_name_new_var.get().strip(), "full_name": " ".join(x for x in [first_name_new_var.get().strip(), raw_name] if x).strip() or raw_name, "email": email_new_var.get().strip(), "favorites": [], "auth": {"password_hash": None, "enabled": False}, "permission": ROLE_E4 if key == SUPERUSER_KEY else ROLE_E1}
                self.ensure_permissions_defaults(); self.save_user_data(); messagebox.showinfo("FiBu Mate", f"Benutzer wurde angelegt:\n\n{raw_name}"); self.render_page()
            tk.Button(footer, text="Benutzer anlegen", command=create_user, bg=BLUE, fg="white", bd=0, padx=14, pady=8, cursor="hand2").grid(row=1, column=6, sticky="w")
        update_scrollregion()


    def render_permissions_menu(self):
        zfont = self.zoomed_content_font
        if not self.can_manage_permissions():
            self.render_menu_text("Keine Berechtigung für das Menü Berechtigungen."); return
        users = self.user_data.setdefault("users", {})
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)
        arrow_col = tk.Frame(container, bg=BG, width=54); arrow_col.pack(side="left", fill="y", padx=(18, 0)); arrow_col.pack_propagate(False)
        body = tk.Frame(container, bg=BG); body.pack(side="left", fill="both", expand=True)
        scroll_canvas = tk.Canvas(body, bg=BG, highlightthickness=0, bd=0); scrollbar = tk.Scrollbar(body, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG); content_window = scroll_canvas.create_window((0,0), window=content, anchor="nw")
        scroll_canvas.configure(yscrollcommand=scrollbar.set); scroll_canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y"); self.register_scroll_canvas(scroll_canvas, scrollbar)
        def update_arrows():
            try:
                first, last = scroll_canvas.yview(); up_arrow.set_enabled(first > 0.001); down_arrow.set_enabled(last < 0.999)
            except Exception: pass
        def scroll_units(n): scroll_canvas.yview_scroll(n, "units"); update_arrows()
        up_arrow = ArrowIndicator(arrow_col, "up", lambda: scroll_units(-5), size=42); down_arrow = ArrowIndicator(arrow_col, "down", lambda: scroll_units(5), size=42)
        up_arrow.pack(side="top", pady=(78, 10)); down_arrow.pack(side="top", pady=(0, 10))
        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width())); scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all")); self._sync_scrollbar_visibility(scroll_canvas, scrollbar); update_arrows()
        self._update_scroll_indicators = update_arrows; content.bind("<Configure>", update_scrollregion); scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<MouseWheel>", lambda e: self._route_mousewheel_to_canvas(e, scroll_canvas))
        frame = content
        tk.Label(frame, text="Berechtigungen", font=zfont(("Segoe UI", 15, "bold")), bg=BG, fg=TEXT).grid(row=0, column=0, columnspan=5, sticky="w", pady=(0, 10))
        def show_role_info():
            popup = tk.Toplevel(self.root); popup.title("Berechtigungen der Rollen anzeigen"); popup.configure(bg=BG); popup.geometry("820x620"); popup.transient(self.root)
            text = tk.Text(popup, bg="white", fg=TEXT, wrap="word", font=("Segoe UI", 10)); text.pack(fill="both", expand=True, padx=14, pady=14)
            content = "E1 - Standard\n- eigene Benutzerdaten ansehen/pflegen\n- Abschlusskalender lesen\n- eigene zuständige Aufgaben erledigen\n- Berichte, Historie und Änderungsprotokolle ansehen\n\nE2 - Erweitert\n- aktuell gleiche Berechtigungen wie E1\n- Platzhalter für spätere Erweiterungen\n\nE3 - Administrator\n- Benutzer anlegen, bearbeiten und löschen\n- Berechtigungen maximal bis E3 vergeben\n- Aufgaben administrieren und Aufgaben-IDs bearbeiten/verknüpfen\n- Zeiträume nach Stichtag schließen/öffnen\n- Berichte und Änderungsprotokolle ansehen\n\nE4 - System-Administrator\n- Berechtigungen analog E3\n- System-Administrator"
            text.insert("1.0", content); text.config(state="disabled")
        tk.Button(frame, text="Berechtigungen der Rollen anzeigen", command=show_role_info, bg=HEADER, fg=TEXT, bd=1, padx=10, pady=5, cursor="hand2").grid(row=0, column=5, sticky="e", pady=(0,10))
        tk.Label(frame, text="Berechtigungen können von berechtigten Rollen maximal bis zur eigenen Rolle vergeben werden.", bg=BG, fg=TEXT2, font=("Segoe UI", 10), justify="left").grid(row=1, column=0, columnspan=5, sticky="w", pady=(0, 16))
        for col, header in enumerate(["Benutzer", "Anzeigename", "Aktuelle Rolle", "Ändern"]): tk.Label(frame, text=header, bg=BG, fg=TEXT, font=("Segoe UI", 10, "bold")).grid(row=2, column=col, sticky="w", padx=(0, 24))
        row = 3
        for key in sorted(users.keys()):
            user = users[key]; current_role = user.get("permission", ROLE_E1)
            tk.Label(frame, text=key, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=0, sticky="w", padx=(0, 24), pady=5)
            tk.Label(frame, text=user.get("display_name", key), bg=BG, fg=TEXT2, font=("Segoe UI", 10), wraplength=260, justify="left").grid(row=row, column=1, sticky="w", padx=(0, 24), pady=5)
            tk.Label(frame, text=current_role, bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=2, sticky="w", padx=(0, 24), pady=5)
            if self.my_role() == ROLE_E3 and (key == SUPERUSER_KEY or current_role == ROLE_E4):
                tk.Label(frame, text="gesperrt", bg=BG, fg=TEXT2, font=("Segoe UI", 10)).grid(row=row, column=3, sticky="w", pady=5); row += 1; continue
            available_roles = [role for role in ROLE_ORDER if self.role_rank(role) <= self.max_assignable_role_rank()]
            if key != SUPERUSER_KEY and ROLE_E4 in available_roles: available_roles.remove(ROLE_E4)
            if key == SUPERUSER_KEY: available_roles = [ROLE_E4]
            role_var = tk.StringVar(value=current_role if current_role in available_roles else available_roles[0])
            dropdown = tk.OptionMenu(frame, role_var, *available_roles); dropdown.config(bg="#E8EEF5", fg=TEXT, bd=1, highlightthickness=0, cursor="hand2"); dropdown.grid(row=row, column=3, sticky="w", pady=5)
            def save_role(user_key=key, var=role_var):
                new_role = var.get()
                if user_key == SUPERUSER_KEY: new_role = ROLE_E4
                if user_key != SUPERUSER_KEY and new_role == ROLE_E4: messagebox.showwarning("FiBu Mate", "Die Rolle System-Administrator darf nur der Benutzer Wagnerm tragen."); return
                if self.role_rank(new_role) > self.max_assignable_role_rank(): messagebox.showwarning("FiBu Mate", "Du kannst keine Rolle vergeben, die höher als deine eigene Rolle ist."); return
                users[user_key]["permission"] = new_role; self.ensure_permissions_defaults(); self.save_user_data(); messagebox.showinfo("FiBu Mate", "Berechtigung wurde gespeichert."); self.render_page()
            tk.Button(frame, text="Speichern", command=save_role, bg=BLUE, fg="white", bd=0, padx=10, pady=5, cursor="hand2").grid(row=row, column=4, sticky="w", padx=(8, 0), pady=5)
            row += 1
        update_scrollregion()

    def render_versions_menu(self):
        history = self.load_version_history().get("entries", [])
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height()
        area_top = 132
        area_bottom = max(area_top + 260, h - 92)
        view_h = int(area_bottom - area_top)

        container = tk.Frame(self.root, bg=BG)
        self.widget_items.append(container)
        self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h)

        scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=scroll_canvas.yview)
        content = tk.Frame(scroll_canvas, bg=BG)
        content_window = scroll_canvas.create_window((0, 0), window=content, anchor="nw")

        def update_scrollregion(_event=None):
            scroll_canvas.itemconfigure(content_window, width=max(1, scroll_canvas.winfo_width()))
            scroll_canvas.configure(scrollregion=scroll_canvas.bbox("all"))
            self._sync_scrollbar_visibility(scroll_canvas, scrollbar)

        content.bind("<Configure>", update_scrollregion)
        scroll_canvas.bind("<Configure>", update_scrollregion)
        scroll_canvas.configure(yscrollcommand=scrollbar.set)
        scroll_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.register_scroll_canvas(scroll_canvas, scrollbar)

        def _on_versions_mousewheel(event):
            return self._route_mousewheel_to_canvas(event, scroll_canvas)

        scroll_canvas.bind("<MouseWheel>", _on_versions_mousewheel)
        content.bind("<MouseWheel>", _on_versions_mousewheel)

        history_source_text = f"Versionsverlauf: lokale Historie plus zentrale Release-Historie ({self._central_version_history_path()})"
        tk.Label(content, text=history_source_text, bg=BG, fg=TEXT2, font=("Segoe UI", 11, "bold"), wraplength=max(700, w - 80), justify="left").pack(anchor="w", padx=24, pady=(12, 0))
        if not history:
            tk.Label(content, text="Noch kein Versionsverlauf vorhanden.", bg=BG, fg=TEXT2, font=("Segoe UI", 12)).pack(anchor="w", padx=24, pady=24)
            update_scrollregion()
            return

        for entry in history:
            entry_frame = tk.Frame(content, bg=BG)
            entry_frame.pack(fill="x", padx=24, pady=(16, 0), anchor="n")
            ver = entry.get("version", "")
            date = entry.get("date", "")
            bullets = entry.get("bullets", [])
            tk.Label(entry_frame, text=f"{ver}.{date}", bg=BG, fg=TEXT, font=("Segoe UI", 12, "bold"), anchor="w").pack(fill="x", anchor="w")
            for bullet in bullets:
                tk.Label(entry_frame, text=f"• {bullet}", bg=BG, fg=TEXT2, font=body_font(11), justify="left", anchor="w", wraplength=max(300, w - 110)).pack(fill="x", anchor="w", padx=(16, 0), pady=(2, 0))
            separator = tk.Canvas(content, bg=BG, height=18, highlightthickness=0, bd=0)
            separator.pack(fill="x", padx=10, pady=(8, 0))
            try:
                self.draw_relief_line(separator, 8, 0, max(100, w - 70))
            except Exception:
                separator.create_line(0, 8, max(100, w - 70), 8, fill=LINE, width=1)
        update_scrollregion()

    def render_menu_text(self, text):
        lab = tk.Label(self.root, text=text, font=("Segoe UI", 16), bg=BG, fg=TEXT); self.widget_items.append(lab); self.canvas.create_window(x_pct(self.canvas.winfo_width(), 50), y_pct(self.canvas.winfo_height(), 60), window=lab, anchor="center"); self.draw_bottom_logo()

    def draw_continuous_relief_line(self, canvas, y, x1, x2):
        canvas.create_line(x1, y, x2, y, fill=blend(BG, "#1F2933", 0.30), width=2)
        canvas.create_line(x1, y + 2, x2, y + 2, fill=blend(BG, WHITE, 0.65), width=1)

    def draw_relief_line(self, canvas, y, x1, x2):
        for i in range(120):
            t = i / 119; fade = 1 - abs(t - 0.5) * 2; sx = x1 + (x2 - x1) * i / 120; ex = x1 + (x2 - x1) * (i + 1) / 120
            canvas.create_line(sx, y, ex, y, fill=blend(BG, "#1F2933", 0.30 * fade), width=2); canvas.create_line(sx, y + 2, ex, y + 2, fill=blend(BG, WHITE, 0.65 * fade), width=1)

    def draw_module_description(self, canvas, module_id, x1, y_top, width):
        """v0.436 Lesbarkeitsschutz: Beschreibung so groß wie möglich, aber im Sollbereich."""
        txt = MODULE_DESCRIPTIONS.get(module_id, "")
        if not txt:
            return
        max_w = max(120, int(width) - 2 * DESCRIPTION_X_OFFSET)
        max_h = max(ui_s(52), int(getattr(self, "_module_desc_tile_h", ui_s(130))) - 2 * DESCRIPTION_Y_OFFSET)
        chosen = body_font(10)
        try:
            for size in range(max(11, ui_s(12)), 8, -1):
                candidate = ("Segoe UI", size)
                f = tkfont.Font(root=self.root, font=candidate)
                words = str(txt).split() or [""]
                line_count = 1; cur = ""
                for word in words:
                    cand = word if not cur else cur + " " + word
                    if f.measure(cand) <= max_w:
                        cur = cand
                    else:
                        line_count += 1; cur = word
                if line_count * max(1, int(f.metrics("linespace"))) <= max_h:
                    chosen = candidate
                    break
        except Exception:
            chosen = body_font(10)
        canvas.create_text(x1 + DESCRIPTION_X_OFFSET, y_top + DESCRIPTION_Y_OFFSET, text=txt, anchor="nw", fill=DESCRIPTION_COLOR, font=chosen, width=max_w, justify="left")

    def module_icon_type(self, module_id):
        if module_id == "tax_reporting":
            return "tax_reporting"
        if module_id == "page:compliance_audit":
            return "compliance"
        if module_id == "audit_cockpit":
            return "audit"
        if module_id == "documentation_center":
            return "documentation"
        if self.current_page in ("data_prep", "nike_tools", "afi_uploads", "debitoren_tools"):
            if module_id == "enbw_strom_tanken_upload":
                return "xls"
            if module_id == "supplier_invoice_afi_upload":
                return "xls"
            if module_id in ("nike_pdf_to_excel", "nike_op_liste_pdf_check"):
                return "pdf_xls"
            if module_id == "invoice_pdf_collector":
                return "pdf_xls"
            if module_id == "debitoren_serienbrief":
                return "modules"
            if str(module_id).startswith("page:"):
                return "pdf_xls"
            return "pdf_xls"
        if self.current_page == "compliance_audit":
            return "compliance"
        if self.current_page == "closing_calendar":
            return "calendar"
        if module_id in ("monthly_close", "quarterly_close", "yearly_close", "deadline_maintenance"):
            if module_id == "deadline_maintenance" and self.role_rank() < 3:
                messagebox.showwarning("Keine Berechtigung", "Dieses Modul ist erst ab E3 verfügbar.")
                return
            return "calendar"
        if str(module_id).startswith("nike_") or module_id == "invoice_pdf_collector":
            return "pdf_xls"
        return "modules"

    def render_module_menu(self, modules, show_descriptions=True):
        modules = [(title, module_id) for title, module_id in modules if module_id not in HIDDEN_TOOL_IDS]
        w, h = self.canvas.winfo_width(), self.canvas.winfo_height(); tile_w = max(290, min(390, int(w * 0.22))); tile_h = max(120, min(160, int(h * 0.15))); self._module_desc_tile_h = tile_h; gap = max(18, int(h * 0.025)); area_top = 132; area_bottom = max(area_top + 260, h - 92); view_h = int(area_bottom - area_top); first_center_x = x_pct(w, 25); first_center_y = y_pct(h, 70); left_x = max(0, first_center_x - tile_w / 2 - 8); top = max(0, first_center_y - tile_h / 2 - area_top - 8)
        container = tk.Frame(self.root, bg=BG); self.widget_items.append(container); self.canvas.create_window(0, area_top, window=container, anchor="nw", width=w, height=view_h); canvas_w = w - left_x - 10; scroll_canvas = tk.Canvas(container, bg=BG, highlightthickness=0, bd=0); scroll_canvas.place(x=left_x, y=0, width=canvas_w, height=view_h); content_h = top + len(modules) * (tile_h + gap) + 20; scroll_canvas.configure(scrollregion=(0, 0, canvas_w, content_h)); self.register_scroll_canvas(scroll_canvas); desc_x1 = tile_w + 90; desc_x2 = canvas_w - 20
        for idx, (title, module_id) in enumerate(modules):
            y = top + idx * (tile_h + gap); cmd = (lambda mid=module_id, ttl=title: self.show_page(str(mid).replace("page:", "", 1), ttl, True)) if str(module_id).startswith("page:") else ((lambda mid=module_id: self.open_tool(mid)) if module_id in TOOL_REGISTRY else self.show_placeholder)
            tile = Tile(scroll_canvas, self, module_id, title, cmd, favorite_enabled=module_id in TOOL_REGISTRY, icon_type=self.module_icon_type(module_id)); tile.resize_tile(tile_w, tile_h); self.focusable_tiles.append(tile); scroll_canvas.create_window(0, y, window=tile, anchor="nw")
            if show_descriptions: self.draw_module_description(scroll_canvas, module_id, desc_x1, y, desc_x2 - desc_x1)
            if idx < len(modules) - 1: self.draw_relief_line(scroll_canvas, y + tile_h + gap / 2, desc_x1, desc_x2)

    def render_external_tool(self, tool_id):
        if tool_id in HIDDEN_TOOL_IDS:
            messagebox.showinfo("FiBu Mate", "Dieses Tool ist in 'Lieferanten-Rechnung zu AFI-Upload' integriert und wird nicht mehr separat gestartet.")
            return
        try:
            module_path = TOOL_REGISTRY[tool_id]["module"]
            module = _load_tool_module_from_file(module_path) or importlib.import_module(module_path)
            try:
                importlib.reload(module)
            except Exception:
                pass
            if not hasattr(module, "render"): raise RuntimeError("Modul hat keine render(app)-Funktion")
            # v0.436: Tool-Kontext setzen, damit das zusammengefasste Abschlusskalender-Modul
            # korrekt zwischen Monats-, Quartals- und Jahresabschluss routet.
            self.current_tool_id = tool_id
            self.current_tool_title = TOOL_REGISTRY.get(tool_id, {}).get("title", tool_id)
            module.render(self)
        except Exception as e:
            messagebox.showerror("FiBu Mate", f"Fehler beim Laden des Moduls:\n\n{TOOL_REGISTRY.get(tool_id, {}).get('title', tool_id)}\n\n{e}"); self.draw_bottom_logo()

    def open_tool(self, tool_id):
        if str(tool_id).startswith("page:"):
            page = str(tool_id).split(":", 1)[1]
            titles = {"nike_tools": "Nike-Tools", "afi_uploads": "AFI-Uploads", "debitoren_tools": "Tools - Debitoren"}
            self.show_page(page, titles.get(page, page), True)
            return
        if tool_id in TOOL_REGISTRY:
            self.show_page(f"tool:{tool_id}", TOOL_REGISTRY[tool_id]["title"], True)

    def show_placeholder(self): messagebox.showinfo("FiBu Mate", "Hinter diesem Widget entsteht gerade ein Modul.")

    def login_user_from_entry(self, username):
        self._suppress_next_global_return = True
        self.login_user(username)
        return "break"

    def login_user(self, username):
        username = " ".join(str(username).strip().split())
        key = normalize_username(username)
        if not key:
            messagebox.showwarning("FiBu Mate", "Bitte einen Nutzernamen eingeben.")
            return
        users = self.user_data.setdefault("users", {})
        if key not in users:
            messagebox.showwarning("FiBu Mate", "Benutzer nicht gefunden.\nBitte wende dich an eine Administratorin / einen Administrator.")
            return
        users[key].setdefault("display_name", username)
        users[key].setdefault("first_name", "")
        users[key].setdefault("full_name", " ".join(x for x in [users[key].get("first_name", "").strip(), users[key].get("display_name", username).strip()] if x).strip() or username)
        users[key].setdefault("favorites", [])
        users[key].setdefault("email", "")
        users[key].setdefault("auth", {"password_hash": None, "enabled": False})
        users[key]["permission"] = ROLE_MIGRATION.get(users[key].get("permission", ROLE_E1), ROLE_E1)
        if key == SUPERUSER_KEY:
            users[key]["permission"] = ROLE_E4
        self.ensure_permissions_defaults()
        self.current_user_key = key
        self.current_user_display = users[key].get("display_name", username)
        self.favorites = set(fav for fav in users[key].get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
        users[key]["favorites"] = sorted(self.favorites)
        self.save_user_data()
        self.start_live_permissions_refresh()
        self.page_history = []
        self.breadcrumb = []
        self.show_page("main", "Hauptmenü", add_to_history=False)

    def start_live_permissions_refresh(self):
        if self._live_permissions_started:
            return
        self._live_permissions_started = True
        self.root.after(3000, self.check_live_user_data_updates)

    def _show_single_live_popup(self, title, text):
        if self._live_permissions_popup_open:
            return
        self._live_permissions_popup_open = True
        try:
            messagebox.showinfo(title, text)
        finally:
            self._live_permissions_popup_open = False

    def check_live_user_data_updates(self):
        try:
            if not self.current_user_key:
                self.root.after(3000, self.check_live_user_data_updates)
                return
            mtime = self._get_user_data_mtime()
            if mtime and mtime != getattr(self, "_user_data_mtime", 0):
                old_role = self.my_role()
                old_display = self.current_user_display
                old_favorites = set(self.favorites)
                self.user_data = self.load_user_data()
                self.ensure_permissions_defaults()
                users = self.user_data.get("users", {})
                key = self.current_user_key
                if key not in users:
                    self.current_user_key = None
                    self.current_user_display = ""
                    self.favorites = set()
                    self.page_history = []
                    self.show_page("launch", add_to_history=False)
                    self._show_single_live_popup("FiBu Mate", "Ihr Benutzer wurde entfernt oder ist nicht mehr verfügbar. Sie wurden abgemeldet.")
                    self._user_data_mtime = mtime
                    return
                user = users[key]
                self.current_user_display = user.get("display_name", old_display)
                self.favorites = set(fav for fav in user.get("favorites", []) if fav in TOOL_REGISTRY and fav not in HIDDEN_TOOL_IDS)
                new_role = self.my_role()
                self._user_data_mtime = mtime
                if new_role != old_role or self.current_user_display != old_display or self.favorites != old_favorites:
                    try:
                        self.update_header()
                    except Exception:
                        pass
                    try:
                        self.render_page()
                    except Exception:
                        pass
                    self._show_single_live_popup("FiBu Mate", "Ihre Benutzer-/Berechtigungsdaten wurden live aktualisiert.")
        except Exception:
            pass
        try:
            self.root.after(3000, self.check_live_user_data_updates)
        except Exception:
            pass

    def logout(self):
        self.current_user_key = None; self.current_user_display = ""; self.favorites = set(); self.page_history = []; self.show_page("launch", add_to_history=False)

    def set_focused_tile(self, tile):
        if tile in self.focusable_tiles: self.focus_index = self.focusable_tiles.index(tile)

    def handle_escape(self, *_):
        if self._close_modal_window():
            return "break"
        handler = getattr(self, "module_escape_handler", None)
        if handler:
            try:
                if handler():
                    return "break"
            except Exception:
                pass
        if self.current_page != "launch":
            self.go_back()
        return "break"
    def handle_enter(self, *_):
        if getattr(self, "_suppress_next_global_return", False):
            self._suppress_next_global_return = False
            return "break"
        focus_widget = self.root.focus_get()
        if isinstance(focus_widget, tk.Entry):
            return "break"
        if 0 <= self.focus_index < len(self.focusable_tiles):
            try:
                self.focusable_tiles[self.focus_index].on_keyboard_activate()
            except Exception:
                pass
        return "break"

    def handle_tab(self, *_):
        if self.focusable_tiles: self.focus_index = (self.focus_index + 1) % len(self.focusable_tiles); self.focusable_tiles[self.focus_index].focus_set()
        return "break"

    def handle_shift_tab(self, *_):
        if self.focusable_tiles: self.focus_index = (self.focus_index - 1) % len(self.focusable_tiles); self.focusable_tiles[self.focus_index].focus_set()
        return "break"

    def confirm_exit(self):
        if getattr(self, "_closing_in_progress", False):
            return
        self._closing_in_progress = True
        try:
            if not self.confirm_unsaved_changes():
                self._closing_in_progress = False
                return
            if not messagebox.askyesno(APP_NAME, "FiBu Mate wirklich schließen?"):
                self._closing_in_progress = False
                return
            try:
                for w in list(self.root.winfo_children()):
                    if isinstance(w, tk.Toplevel):
                        try:
                            w.destroy()
                        except Exception:
                            pass
            except Exception:
                pass
            try:
                self.root.quit()
            except Exception:
                pass
            try:
                self.root.destroy()
            except Exception:
                pass
        except Exception:
            try:
                self.root.destroy()
            except Exception:
                pass


    def ensure_version_433_once(self):
        self.bump_version_once(
            "2026-05-30_v0433_responsive_scaling_documentation_center_table",
            [
                "v0.433: Responsive Skalierung für kleinere Monitore ergänzt; zentrale Schrift-, Kachel- und Widgetgrößen werden an die Bildschirmgröße angepasst.",
                "v0.433: Dokumentationszentrale-Tabelle umgebaut: Aufgabenzuordnung als erste Spalte, Dokument statt Dokumentname, Status und Pfad ausgeblendet.",
                "v0.433: Dokumentnamen in der Dokumentationszentrale sind anklickbar und öffnen die hinterlegte Datei; Datei- und Anhang-Icons ergänzt.",
                "v0.433: Dokumentationszentrale erhält horizontale Scroll-Unterstützung für Tabellenüberlauf und optimierte Spaltenbreiten.",
                "v0.433 Korrektur Paket 1b: Kachel-/Widget-Skalierung nachjustiert, Schließen-Logik gegen doppelten Dialog bzw. weißes Restfenster abgesichert und Mini-Widget-Schließen neben Hilfe neutralisiert.",
            ],
        )

    def run(self):
        try:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_force()
            self.root.update_idletasks()
            self.root.update()
        except Exception:
            pass
        self.root.mainloop()


    def ensure_version_435_once(self):
        """Version 0.435: Audit-Cockpit-Korrektur und Zuständigkeitspflege mit Aufgabenverknüpfung."""
        update_id = "2026-05-29_v0435_audit_cockpit_zustaendigkeit_verknuepfung"
        bullets = [
            "v0.435: Audit-Cockpit-Fehler korrigiert; Zeitraumformatierung akzeptiert den Zeitstempel-Kontext.",
            "v0.435: Zuständigkeitspflege erweitert um Aufgabenverknüpfung mit gemeinsamer Aufgaben-ID.",
            "v0.435: Mehrfachpräfixe M/Q/J umgesetzt, z. B. MQ oder QJ; Aufgaben erscheinen in allen passenden Übersichtsreitern.",
            "v0.435: Zuständigkeit und Benutzer-Key werden gemeinsam als Feld Zuständigkeit gepflegt.",
            "v0.435: Drop-downs in allen editierbaren Zuständigkeitspflege-Zellen, Checkboxen in der Aufgaben-ID-Spalte, Buttons Aufgaben verknüpfen und Löschen ergänzt.",
        ]
        try:
            self.version_state.setdefault("applied_updates", [])
            changed = False
            current_build = int(self.version_state.get("build", DEFAULT_BUILD))
            if current_build < 35 or current_build > 100:
                self.version_state["build"] = 35
                changed = True
            if update_id not in self.version_state["applied_updates"]:
                history = self.load_version_history()
                history.setdefault("entries", [])
                if not any(e.get("update_id") == update_id for e in history.get("entries", [])):
                    history["entries"].insert(0, {
                        "version": "v0.435",
                        "date": now_date_str(),
                        "update_id": update_id,
                        "bullets": bullets,
                    })
                    self.save_version_history(history)
                self.version_state["applied_updates"].append(update_id)
                changed = True
            if changed:
                self.save_version_state()
        except Exception:
            pass

    def normalize_version_after_zoom_patch(self):
        """Korrektur: versehentliche Versionssprünge aus Zoom-Testpaketen zurücknehmen, v0.435 aber beibehalten."""
        try:
            changed = False
            try:
                build = int(self.version_state.get("build", DEFAULT_BUILD))
            except Exception:
                build = DEFAULT_BUILD
            if build > 35:
                self.version_state["build"] = 35
                changed = True
            applied = self.version_state.setdefault("applied_updates", [])
            for upd in ("2026-05-29_v0434_paket1c_fibu_mate_zoom_miniwidget", "2026-05-29_v0434_paket1d_fibu_mate_zoom_miniwidget_nachkorrektur", "2026-05-29_v0434_paket1e_zoomleiste"):
                if upd not in applied:
                    applied.append(upd)
                    changed = True
            self.version_state["applied_updates"] = applied
            if changed:
                self.save_version_state()
            history = self.load_version_history()
            entries = history.setdefault("entries", [])
            cleaned = []
            for entry in entries:
                version = str(entry.get("version", ""))
                update_id = str(entry.get("update_id", ""))
                bullets = "\n".join(str(b) for b in entry.get("bullets", [])).lower()
                accidental = version in ("v0.436", "v0.437") or update_id in ("2026-05-29_v0434_paket1c_fibu_mate_zoom_miniwidget", "2026-05-29_v0434_paket1d_fibu_mate_zoom_miniwidget_nachkorrektur", "2026-05-29_v0434_paket1e_zoomleiste") or "paket 1c" in bullets or "paket 1d" in bullets or "paket 1e" in bullets or "zoom-leiste" in bullets or "zoomleiste" in bullets
                if not accidental:
                    cleaned.append(entry)
            if len(cleaned) != len(entries):
                history["entries"] = cleaned
                self.save_version_history(history)
            try:
                self.version_label.config(text=self.version_label_text())
            except Exception:
                pass
        except Exception:
            pass

    def ensure_version_434_once(self):
        """v0.434: Bereichsbezogener Zoom und Lesbarkeitspaket 1A."""
        try:
            self.bump_version_once(
                "2026-05-29_v0434_scope_zoom_readability_paket1a",
                [
                    "v0.434 Paket 1A: Zoom-Architektur von global auf bereichsbezogen umgestellt.",
                    "v0.434 Paket 1A: Zoomprofile werden pro Benutzer, Bildschirmgröße und Bereich/Modul gespeichert.",
                    "v0.434 Paket 1A: Kopf- und Fußleiste sind vom Strg+Mausrad-Zoom ausgenommen; Content, Module, Popups und Canvas-Inhalte können bereichsbezogen zoomen.",
                    "v0.434 Paket 1A: Kachel-Hover erhält den aktiven Menü-Zoom und setzt skalierte Kacheltexte nicht mehr zurück.",
                    "v0.434 Paket 1A: Steuermeldungs- und Audit-Cockpit erhalten größere Standard-Tabellen-/Dialogschriften.",
                ],
            )
        except Exception:
            pass

    def ensure_version_434_1c_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return

    def ensure_version_434_1d_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return

    def ensure_version_434_1e_once(self):
        """Technische Korrektur innerhalb v0.434; kein Versions-Bump."""
        return


# ------------------------------------------------------------------
# Wissenszentrale - Next Blocks Final Patch 2026-06-12
# Ergänzt per Monkey-Patch: Word-Export, Anhänge, Kommentare,
# B2-Kategorienverwaltung, B1-Ausblendung, ESC-Logik, neutrale Navigation,
# Übersicht als Standard, vergrößerte Dropdowns, To-Do-Rhythmus bedingt.
# ------------------------------------------------------------------
def _kbp_permission_level(self):
    key = str(getattr(self, "current_user_key", "") or "").lower()
    disp = str(getattr(self, "current_user_display", "") or "").lower()
    if "wagnerm" in key or "wagnerm" in disp or "admin" in key or "admin" in disp:
        return 4
    raw = ""
    try:
        data = getattr(self, "user_data", {}) or {}
        raw = str(data).lower()
    except Exception:
        raw = ""
    check = " ".join([key, disp, raw])
    if "e4" in check or "stufe 4" in check or "level 4" in check or "berechtigung 4" in check:
        return 4
    if "e3" in check or "stufe 3" in check or "level 3" in check or "berechtigung 3" in check:
        return 3
    if "b2" in check or "stufe 2" in check or "level 2" in check or "berechtigung 2" in check:
        return 2
    return 1

def _kbp_can_edit(self):
    return self.kb_current_permission_level() >= 2

def _kbp_can_manage_categories(self):
    return self.kb_current_permission_level() >= 2

def _kbp_draw_button(self, x, y, text, command=None, accent=False, width=None):
    btn = tk.Button(self.root, text=text, command=command, font=body_font(10, weight="bold" if accent else None),
                    bg=WHITE, fg=TEXT, activebackground="#DDEAF7", relief="solid", bd=1, cursor="hand2", padx=8, pady=3)
    btn_w = width or max(110, 18 + len(text) * 8)
    self.canvas.create_window(ui_s(x), ui_s(y), window=btn, anchor="nw", width=ui_s(btn_w), height=ui_s(34))
    self.widget_items.append(btn)
    return x + btn_w + 8

def _kbp_handle_escape(self, event=None):
    if getattr(self, "current_page", "") == "knowledge_base":
        if not bool(getattr(self, "knowledge_start_overlay", False)):
            self.knowledge_start_overlay = True
            self.render_page()
            return "break"
        self.knowledge_start_overlay = False
        self.show_page("main_menu", "Hauptmenü", False)
        return "break"
    return None

def _kbp_export_dir(self):
    path = os.path.join(APPDATA_DIR, "Wissenszentrale_Export")
    try: os.makedirs(path, exist_ok=True)
    except Exception: pass
    return path

def _kbp_attachment_dir(self, entry_id):
    path = os.path.join(APPDATA_DIR, "Wissenszentrale_Anhaenge", str(entry_id))
    try: os.makedirs(path, exist_ok=True)
    except Exception: pass
    return path

def _kbp_add_pending_attachment(self):
    try:
        from tkinter import filedialog
        files = filedialog.askopenfilenames(title="Anhänge auswählen")
    except Exception:
        files = []
    if not files:
        return
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    for f in files:
        if f and f not in pending:
            pending.append(f)
    self.kb_pending_attachments = pending
    self.kb_mark_unsaved()
    self.render_page()

def _kbp_add_comment_to_selected(self, text_widget):
    entry_id = getattr(self, "kb_selected_entry_id", None)
    try: comment = text_widget.get("1.0", "end-1c").strip()
    except Exception: comment = ""
    if not entry_id or not comment:
        return
    entries = self.kb_load_entries()
    for e in entries:
        if e.get("id") == entry_id:
            e.setdefault("comments", []).append({"user": self.current_user_display or self.current_user_key or "", "text": comment, "created_at": self.kb_now()})
            e["updated_at"] = self.kb_now()
            break
    if self.kb_save_entries(entries):
        self.render_page()

def _kbp_save_entry_from_form(self):
    if not self.kb_can_create_or_edit():
        return False
    self.kb_ensure_state_vars()
    title = (getattr(self, "kb_title_var", tk.StringVar(value="")).get() or "").strip()
    if not title:
        try: messagebox.showwarning("Wissenszentrale", "Bitte einen Titel erfassen.")
        except Exception: pass
        return False
    categories = []
    for var in getattr(self, "kb_entry_category_vars", []):
        value = (var.get() or "").strip()
        if value and value not in categories:
            categories.append(value)
    categories = categories[:4]
    user = (getattr(self, "kb_user_var", tk.StringVar(value="")).get() or "").strip()
    status = (getattr(self, "kb_status_var", tk.StringVar(value="Aktiv")).get() or "Aktiv").strip()
    rhythm = (getattr(self, "kb_rhythm_var", tk.StringVar(value="")).get() or "").strip() if any(c.lower()=="to-do" for c in categories) else ""
    try: text_value = self.kb_text_widget.get("1.0", "end-1c")
    except Exception: text_value = getattr(self, "kb_text_initial", "")
    entries = self.kb_load_entries(); now = self.kb_now(); edit_id = getattr(self, "kb_edit_entry_id", None)
    if edit_id:
        selected_id = edit_id
        for e in entries:
            if e.get("id") == edit_id:
                e.update({"title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "updated_at": now})
                break
    else:
        selected_id = self.kb_make_entry_id()
        entries.append({"id": selected_id, "title": title, "categories": categories, "user": user, "status": status, "rhythm": rhythm, "text": text_value, "created_at": now, "updated_at": now, "comments": [], "attachments": []})
    pending = list(getattr(self, "kb_pending_attachments", []) or [])
    if pending:
        import shutil as _shutil
        copied=[]; target=self.kb_attachment_dir(selected_id)
        for src in pending:
            try:
                name=os.path.basename(src); dst=os.path.join(target,name); _shutil.copy2(src,dst); copied.append({"name": name, "path": dst, "added_at": now})
            except Exception: pass
        for e in entries:
            if e.get("id") == selected_id:
                e.setdefault("attachments", []).extend(copied); break
    if self.kb_save_entries(entries):
        self.kb_selected_entry_id = selected_id; self.kb_edit_entry_id = None; self.kb_pending_attachments=[]; self.knowledge_unsaved=False; self.knowledge_view="detail"; self.render_page(); return True
    return False

def _kbp_confirm_unsaved_before_switch(self):
    if not self.kb_has_unsaved_changes():
        return True
    try:
        result = messagebox.askyesnocancel("Ungespeicherte Änderungen", "Änderungen speichern?\n\nJa = speichern und wechseln\nNein = verwerfen und wechseln\nAbbrechen = bleiben")
    except Exception:
        result = False
    if result is True:
        return self.kb_save_entry_from_form() if getattr(self, "knowledge_view", "") == "new" else self.kb_save_current_work()
    if result is False:
        self.knowledge_unsaved = False; return True
    return False

def _kbp_switch_view_from_start(self, view):
    if not self.kb_confirm_unsaved_before_switch(): return
    if view == "new" and not self.kb_can_create_or_edit(): view = "all"
    if view == "categories" and not self.kb_can_manage_categories(): view = "all"
    self.knowledge_view = view
    if view == "new": self.kb_prepare_new_entry()
    self.knowledge_start_overlay = False
    self.render_page()

def _kbp_export_word(self):
    entry = self.kb_get_entry(getattr(self, "kb_selected_entry_id", None))
    if not entry: return
    try:
        from docx import Document
        import re as _re
        doc = Document(); doc.add_heading(entry.get("title", "Wissenseintrag"), level=1)
        for label, val in [("Geändert", self.kb_display_date(entry.get("updated_at"))), ("Benutzer", entry.get("user","")), ("Status", entry.get("status","")), ("Kategorien", ", ".join(entry.get("categories", []) or [])), ("To-Do-Rhythmus", entry.get("rhythm", ""))]:
            if val: doc.add_paragraph(f"{label}: {val}")
        doc.add_heading("Inhalt", level=2); doc.add_paragraph(entry.get("text", ""))
        doc.add_heading("Anhänge", level=2)
        for a in entry.get("attachments", []) or []: doc.add_paragraph(f"- {a.get('name','')} ({a.get('path','')})")
        doc.add_heading("Kommentare", level=2)
        for c in entry.get("comments", []) or []: doc.add_paragraph(f"{self.kb_display_date(c.get('created_at'))} - {c.get('user','')}: {c.get('text','')}")
        safe = _re.sub(r"[^A-Za-z0-9_äöüÄÖÜß-]+", "_", entry.get("title", "Wissenseintrag"))[:80]
        out = os.path.join(self.kb_export_dir(), f"{safe}.docx")
        doc.save(out); messagebox.showinfo("Wissenszentrale", "Word-Export erstellt:\n" + out)
    except Exception as exc:
        messagebox.showerror("Wissenszentrale", "Word-Export fehlgeschlagen:\n" + str(exc))

def _kbp_render_work(self):
    self.kb_ensure_state_vars()
    try: self.root.option_add("*TCombobox*Listbox.font", body_font(12)); self.root.bind("<Escape>", self.kb_handle_escape, add="+")
    except Exception: pass
    w,h=self.canvas.winfo_width(), self.canvas.winfo_height()
    if self.knowledge_view == "new" and not self.kb_can_create_or_edit(): self.knowledge_view="all"
    if self.knowledge_view == "categories" and not self.kb_can_manage_categories(): self.knowledge_view="all"
    x=28; y=145
    x=self.draw_kb_button(x,y,"Start",self.kb_show_start_overlay,False,92)
    x=self.draw_kb_button(x,y,"Übersicht",lambda:self.kb_switch_view_from_start("all"),False,112)
    if self.kb_can_create_or_edit(): x=self.draw_kb_button(x,y,"Neuer Eintrag",lambda:self.kb_switch_view_from_start("new"),False,128)
    x=self.draw_kb_button(x,y,"To-Dos",lambda:self.kb_switch_view_from_start("todos"),False,92)
    x=self.draw_kb_button(x,y,"Veraltete Einträge",lambda:self.kb_switch_view_from_start("outdated"),False,148)
    if self.kb_can_manage_categories(): x=self.draw_kb_button(x,y,"Kategorien verwalten",lambda:self.kb_switch_view_from_start("categories"),False,168)
    sf=tk.Frame(self.root,bg=BG); self.widget_items.append(sf)
    tk.Label(sf,text="Suche",bg=BG,fg=TEXT2,font=body_font(10)).grid(row=0,column=0,sticky="w")
    ent=tk.Entry(sf,textvariable=self.kb_search_var,font=body_font(11),bg=WHITE,fg=TEXT,relief="solid",bd=1); ent.grid(row=1,column=0,sticky="ew",ipady=5); ent.bind("<Return>", self.kb_on_search_return); ent.bind("<KeyRelease>", lambda e: self.kb_apply_filters())
    tk.Button(sf,text="Suchen",command=self.kb_apply_filters,bg=WHITE,fg=TEXT,relief="solid",bd=1,font=body_font(10)).grid(row=1,column=1,padx=(8,0),sticky="ns")
    sf.grid_columnconfigure(0,weight=1); self.canvas.create_window(ui_s(28),ui_s(192),window=sf,anchor="nw",width=ui_s(min(790,max(560,w-80))),height=ui_s(60))
    ff=tk.Frame(self.root,bg=BG); self.widget_items.append(ff); vals=[""]+self.kb_get_categories(); tk.Label(ff,text="Kategorie-Filter",bg=BG,fg=TEXT2,font=body_font(10)).grid(row=0,column=0,sticky="w",columnspan=6)
    for i,var in enumerate(self.kb_filter_vars):
        b=tk.Frame(ff,bg=BG); b.grid(row=1,column=i,padx=(0 if i==0 else 12,0),sticky="w"); tk.Label(b,text=f"Kategorie {i+1}",bg=BG,fg=TEXT2,font=body_font(9)).pack(anchor="w"); cb=ttk.Combobox(b,textvariable=var,values=vals,state="readonly",font=body_font(11),width=18); cb.pack(anchor="w"); cb.bind("<<ComboboxSelected>>",self.kb_apply_filters)
    selected=[(v.get() or "").strip().lower() for v in self.kb_filter_vars]
    if self.knowledge_view == "todos" or "to-do" in selected:
        b=tk.Frame(ff,bg=BG); b.grid(row=1,column=4,padx=(12,0),sticky="w"); tk.Label(b,text="To-Do-Rhythmus",bg=BG,fg=TEXT2,font=body_font(9)).pack(anchor="w"); cb=ttk.Combobox(b,textvariable=self.kb_todo_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(11),width=16); cb.pack(anchor="w"); cb.bind("<<ComboboxSelected>>", self.kb_apply_filters)
    else:
        self.kb_todo_rhythm_var.set(""); tk.Label(ff,text="Leere Filter = keine Filterung",bg=BG,fg=TEXT2,font=body_font(9)).grid(row=1,column=4,padx=(14,0),sticky="s")
    self.canvas.create_window(ui_s(28),ui_s(252),window=ff,anchor="nw",width=ui_s(min(1050,max(760,w-70))),height=ui_s(90))
    left_x=20; left_y=365; left_w=max(340,min(460,int(w*0.31))); pane_h=max(420,h-left_y-24); right_x=left_x+left_w+24; right_w=max(760,w-right_x-24)
    self.render_kb_hits_pane(left_x,left_y,left_w,pane_h)
    if self.knowledge_view == "new": self.render_kb_new_entry_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view == "todos": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="To-Dos")
    elif self.knowledge_view == "outdated": self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Veraltete Einträge",status_filter="Veraltet")
    elif self.knowledge_view == "categories": self.render_kb_categories_area(right_x,left_y,right_w,pane_h)
    elif self.knowledge_view == "detail": self.render_kb_detail_area(right_x,left_y,right_w,pane_h)
    else: self.render_kb_list_area(right_x,left_y,right_w,pane_h,title="Übersicht")

def _kbp_render_new(self,x,y,w,h):
    if not self.kb_can_create_or_edit(): self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    if not hasattr(self,"kb_title_var"): self.kb_prepare_new_entry()
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    tk.Label(frame,text=("Eintrag bearbeiten" if getattr(self,"kb_edit_entry_id",None) else "Neuer Eintrag"),bg=WHITE,fg=BLUE,font=body_font(15,weight="bold")).pack(anchor="w",padx=18,pady=(16,6))
    form=tk.Frame(frame,bg=WHITE); form.pack(fill="x",padx=18,pady=(4,8)); tk.Label(form,text="Titel des Eintrags",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=0,column=0,sticky="w")
    tk.Entry(form,textvariable=self.kb_title_var,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).grid(row=1,column=0,columnspan=4,sticky="ew",ipady=5,pady=(0,8))
    cats=[""]+self.kb_get_categories()
    for i,var in enumerate(self.kb_entry_category_vars):
        tk.Label(form,text=f"Kategorie {i+1}",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=2,column=i,sticky="w",padx=(0 if i==0 else 8,0)); cb=ttk.Combobox(form,textvariable=var,values=cats,state="normal",font=body_font(11),width=18); cb.grid(row=3,column=i,sticky="ew",padx=(0 if i==0 else 8,0),pady=(0,8)); cb.bind("<<ComboboxSelected>>",lambda e:self.render_page())
    tk.Label(form,text="Assoziierter Benutzer",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=0,sticky="w"); tk.Entry(form,textvariable=self.kb_user_var,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).grid(row=5,column=0,sticky="ew",pady=(0,8))
    tk.Label(form,text="Status",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=1,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_status_var,values=["Aktiv","Entwurf","Veraltet"],state="readonly",font=body_font(11)).grid(row=5,column=1,sticky="ew",padx=(8,0),pady=(0,8))
    if any((v.get() or "").strip().lower()=="to-do" for v in self.kb_entry_category_vars):
        tk.Label(form,text="To-Do-Rhythmus",bg=WHITE,fg=TEXT2,font=body_font(9)).grid(row=4,column=2,sticky="w",padx=(8,0)); ttk.Combobox(form,textvariable=self.kb_rhythm_var,values=["","täglich","wöchentlich","monatlich","quartalsweise","jährlich","bei Bedarf"],state="readonly",font=body_font(11)).grid(row=5,column=2,sticky="ew",padx=(8,0),pady=(0,8))
    else: self.kb_rhythm_var.set("")
    for c in range(4): form.grid_columnconfigure(c,weight=1)
    bf=tk.Frame(frame,bg=WHITE); bf.pack(fill="both",expand=True,padx=18,pady=(0,10)); tk.Label(bf,text="Freitext / Prozessdokumentation / Leitfaden",bg=WHITE,fg=TEXT2,font=body_font(9)).pack(anchor="w"); self.kb_text_widget=tk.Text(bf,bg="#F8FAFC",fg=TEXT,font=body_font(10),relief="solid",bd=1,wrap="word"); self.kb_text_widget.pack(fill="both",expand=True); self.kb_text_widget.insert("1.0",getattr(self,"kb_text_initial","")); self.kb_text_widget.bind("<KeyRelease>",lambda e:self.kb_mark_unsaved())
    pending=getattr(self,"kb_pending_attachments",[]) or []
    if pending: tk.Label(frame,text="Vorgemerkte Anhänge: "+", ".join([os.path.basename(x) for x in pending]),bg=WHITE,fg=TEXT2,font=body_font(9)).pack(anchor="w",padx=18)
    br=tk.Frame(frame,bg=WHITE); br.pack(fill="x",padx=18,pady=(0,16)); tk.Button(br,text="Anhang hinzufügen",command=self.kb_add_pending_attachment,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=12,ipady=4); tk.Button(br,text="Als Entwurf speichern",command=lambda:(self.kb_status_var.set("Entwurf"),self.kb_save_entry_from_form()),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=12,ipady=4); tk.Button(br,text="Speichern",command=self.kb_save_entry_from_form,bg=WHITE,fg=TEXT,font=body_font(10,weight="bold"),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Abbrechen",command=lambda:self.kb_switch_view_from_start("all"),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",ipadx=16,ipady=4)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

def _kbp_render_detail(self,x,y,w,h):
    e=self.kb_get_entry(getattr(self,"kb_selected_entry_id",None))
    if not e: self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame)
    tk.Label(frame,text=e.get("title",""),bg=WHITE,fg=BLUE,font=body_font(16,weight="bold")).pack(anchor="w",padx=18,pady=(16,6)); tk.Label(frame,text=f"Geändert: {self.kb_display_date(e.get('updated_at'))}    Benutzer: {e.get('user','')}    Status: {e.get('status','')}",bg=WHITE,fg=TEXT2,font=body_font(10)).pack(anchor="w",padx=18)
    tk.Label(frame,text="Kategorien: "+(", ".join(e.get("categories",[]) or []) or "Keine Kategorien"),bg=WHITE,fg=TEXT,font=body_font(10,weight="bold")).pack(anchor="w",padx=18,pady=(8,4))
    txt=tk.Text(frame,bg="#F8FAFC",fg=TEXT,font=body_font(10),relief="solid",bd=1,wrap="word",height=8); txt.pack(fill="both",expand=True,padx=18,pady=(6,8)); txt.insert("1.0",e.get("text","")); txt.configure(state="disabled")
    lower=tk.Frame(frame,bg=WHITE); lower.pack(fill="x",padx=18,pady=(0,8)); left=tk.Frame(lower,bg=WHITE); left.pack(side="left",fill="both",expand=True); right=tk.Frame(lower,bg=WHITE); right.pack(side="right",fill="both",expand=True,padx=(12,0))
    tk.Label(left,text="Anhänge",bg=WHITE,fg=BLUE,font=body_font(10,weight="bold")).pack(anchor="w")
    for a in e.get("attachments",[]) or []: tk.Label(left,text="• "+a.get("name",""),bg=WHITE,fg=TEXT,font=body_font(9)).pack(anchor="w")
    tk.Label(right,text="Kommentare",bg=WHITE,fg=BLUE,font=body_font(10,weight="bold")).pack(anchor="w")
    for c in (e.get("comments",[]) or [])[-3:]: tk.Label(right,text=f"{self.kb_display_date(c.get('created_at'))}: {c.get('text','')[:80]}",bg=WHITE,fg=TEXT,font=body_font(9),wraplength=380,justify="left").pack(anchor="w")
    comment=tk.Text(right,height=2,bg="#F8FAFC",fg=TEXT,font=body_font(9),relief="solid",bd=1,wrap="word"); comment.pack(fill="x",pady=(4,0))
    br=tk.Frame(frame,bg=WHITE); br.pack(fill="x",padx=18,pady=(0,14))
    if self.kb_can_create_or_edit(): tk.Button(br,text="Bearbeiten",command=self.kb_edit_selected_entry,bg=WHITE,fg=TEXT,font=body_font(10,weight="bold"),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4)
    tk.Button(br,text="Kommentar speichern",command=lambda:self.kb_add_comment_to_selected(comment),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Word-Export",command=self.kb_export_selected_to_word,bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",padx=(0,8),ipadx=16,ipady=4); tk.Button(br,text="Zur Übersicht",command=lambda:self.kb_switch_view_from_start("all"),bg=WHITE,fg=TEXT,font=body_font(10),relief="solid",bd=1).pack(side="left",ipadx=16,ipady=4)
    self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

def _kbp_render_categories(self,x,y,w,h):
    if not self.kb_can_manage_categories(): self.render_kb_list_area(x,y,w,h,title="Übersicht"); return
    frame=tk.Frame(self.root,bg=WHITE,highlightbackground=LINE,highlightthickness=2); self.widget_items.append(frame); tk.Label(frame,text="Kategorien verwalten",bg=WHITE,fg=BLUE,font=body_font(15,weight="bold")).pack(anchor="w",padx=18,pady=(16,8)); tk.Label(frame,text="B2, E3, E4 und Admin dürfen Kategorien lesen und verwalten. Neue Kategorien können direkt beim Erfassen in den Kategorie-Feldern ergänzt werden.",bg=WHITE,fg=TEXT,font=body_font(10),wraplength=max(500,int(w*0.85)),justify="left").pack(anchor="w",padx=18,pady=(0,12)); lb=tk.Listbox(frame,bg=WHITE,fg=TEXT,font=body_font(12),relief="solid",bd=1); lb.pack(fill="both",expand=True,padx=18,pady=(0,18)); [lb.insert("end",c) for c in self.kb_get_categories()]; self.canvas.create_window(ui_s(x),ui_s(y),window=frame,anchor="nw",width=ui_s(w),height=ui_s(h))

FiBuMateApp.kb_current_permission_level = _kbp_permission_level
FiBuMateApp.kb_can_create_or_edit = _kbp_can_edit
FiBuMateApp.kb_can_manage_categories = _kbp_can_manage_categories
FiBuMateApp.draw_kb_button = _kbp_draw_button
FiBuMateApp.kb_handle_escape = _kbp_handle_escape
FiBuMateApp.kb_export_dir = _kbp_export_dir
FiBuMateApp.kb_attachment_dir = _kbp_attachment_dir
FiBuMateApp.kb_add_pending_attachment = _kbp_add_pending_attachment
FiBuMateApp.kb_add_comment_to_selected = _kbp_add_comment_to_selected
FiBuMateApp.kb_save_entry_from_form = _kbp_save_entry_from_form
FiBuMateApp.kb_confirm_unsaved_before_switch = _kbp_confirm_unsaved_before_switch
FiBuMateApp.kb_switch_view_from_start = _kbp_switch_view_from_start
FiBuMateApp.kb_export_selected_to_word = _kbp_export_word
FiBuMateApp.render_knowledge_work_area = _kbp_render_work
FiBuMateApp.render_kb_new_entry_area = _kbp_render_new
FiBuMateApp.render_kb_detail_area = _kbp_render_detail
FiBuMateApp.render_kb_categories_area = _kbp_render_categories

if __name__ == "__main__":
    FiBuMateApp().run()